from __future__ import annotations

import os
import sys
from collections.abc import Iterator
from pathlib import Path

import pytest

PROJECT_ROOT = Path(__file__).resolve().parents[4]
LOCAL_SRC_PATHS = (
    PROJECT_ROOT,
    PROJECT_ROOT / "packages" / "core" / "src",
    PROJECT_ROOT / "packages" / "application" / "src",
    PROJECT_ROOT / "packages" / "runtime" / "src",
    PROJECT_ROOT / "packages" / "bundle" / "src",
    PROJECT_ROOT / "packages" / "apps" / "gui" / "src",
)
for path in reversed(LOCAL_SRC_PATHS):
    if str(path) not in sys.path:
        sys.path.insert(0, str(path))

from tests.support.gui import shutdown_main_window

PROJECT_CONFIGS = PROJECT_ROOT / "configs"


def write_minimal_base_config_tree(base_dir: Path) -> None:
    """Create an empty TOML file for every spec in the registry under *base_dir*.

    GUI tests that only read/write a few specific files use this so the
    three-layer loader (which requires every base file to exist) can
    construct without FileNotFoundError.
    """
    from docwen_runtime.config.registry import CONFIG_FILES

    for spec in CONFIG_FILES:
        path = base_dir / spec.rel_path
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text("\n", encoding="utf-8")


def _create_sample_docx(output_path: Path) -> Path:
    from docx import Document

    doc = Document()
    doc.add_heading("E2E Test Document", level=1)
    doc.add_paragraph("This is a test document for GUI end-to-end smoke testing.")
    doc.add_heading("Section One", level=2)
    doc.add_paragraph("A paragraph with normal text for conversion verification.")
    doc.save(str(output_path))
    return output_path


@pytest.fixture(scope="session")
def qapp():
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

    from PySide6.QtWidgets import QApplication

    app = QApplication.instance()
    if app is None:
        app = QApplication(["pytest"])
    yield app


@pytest.fixture(autouse=True)
def cleanup_qt_top_level_widgets() -> Iterator[None]:
    """Delete stray top-level widgets left by GUI tests.

    Many widget tests instantiate tabs or dialogs without a parent. Closing
    them hides the widget but keeps it in QApplication.topLevelWidgets(), which
    can make global theme sync progressively slower across a long test run.
    """
    yield
    try:
        from PySide6.QtCore import QCoreApplication, QEvent
        from PySide6.QtWidgets import QApplication
    except Exception:
        return
    app = QApplication.instance()
    if not isinstance(app, QApplication):
        return
    for widget in list(app.topLevelWidgets()):
        try:
            widget.close()
            widget.deleteLater()
        except RuntimeError:
            continue
    QCoreApplication.sendPostedEvents(None, QEvent.Type.DeferredDelete)
    app.processEvents()


@pytest.fixture
def main_window(qapp) -> Iterator[object]:
    from docwen_gui.main_window import MainWindow
    from docwen_gui.view_models.main_window_vm import MainWindowViewModel

    vm = MainWindowViewModel(controller=None)
    window = MainWindow(view_model=vm)
    window.setup_ui()
    yield window
    window.close()


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    return _create_sample_docx(tmp_path / "e2e_test.docx")


@pytest.fixture
def main_window_with_controller(qapp, tmp_path: Path) -> Iterator[object]:
    from PySide6.QtWidgets import QApplication

    from docwen_application.controller import ApplicationController
    from docwen_bundle.config_port import ConfigPortAdapter
    from docwen_bundle.runtime_factory import create_runtime_port
    from docwen_gui.app import create_main_window
    from docwen_gui.qt_bridge.task_event_bridge import TaskEventBridge

    bridge = TaskEventBridge()

    def _event_callback(event) -> None:
        payload = {"task_id": event.task_id, **dict(event.payload)}
        bridge.enqueue(event.event_type, payload)

    runtime_port = create_runtime_port(event_callback=_event_callback)
    controller = ApplicationController(
        runtime_port=runtime_port,
        config_port=ConfigPortAdapter(base_dir=PROJECT_CONFIGS, user_dir=tmp_path / "configs"),
    )
    controller.start()
    window = create_main_window(controller=controller, task_event_bridge=bridge)

    app = QApplication.instance()
    if app is not None:
        app.processEvents()

    yield window
    shutdown_main_window(window)
