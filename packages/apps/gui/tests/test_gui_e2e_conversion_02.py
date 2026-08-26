"""Focused tests split from test_gui_e2e_conversion.py."""

from __future__ import annotations

from ._gui_e2e_conversion_support import (
    _E2E_CONVERSION_TIMEOUT_MS,
    Path,
    _wait_for,
    pytest,
    shutdown_main_window,
    t,
    threading,
    time,
)


@pytest.mark.gui
class TestGuiCancellationExecution:
    def test_cancel_button_interrupts_application_preconversion(
        self,
        main_window_with_controller,
        tmp_path: Path,
        monkeypatch,
    ) -> None:
        from PySide6.QtWidgets import QApplication

        from docwen_application.preconversion.pre_converter import PreConversionFailure
        from docwen_core.models import (
            AdmissionDecision,
            DetectionConfidence,
            DetectionMethod,
            FileInspection,
            FormatRelation,
            StructureStatus,
        )
        from docwen_gui.main_window import _normalize_path

        window = main_window_with_controller
        source = tmp_path / "cancel-preconversion.doc"
        source.write_bytes(b"legacy document fixture")
        normalized = _normalize_path(str(source))
        source_stat = source.stat()
        admitted_doc = FileInspection(
            file_path=str(source.resolve()),
            size_bytes=source_stat.st_size,
            mtime_ns=source_stat.st_mtime_ns,
            extension=".doc",
            declared_format="doc",
            declared_category="document",
            detected_format="doc",
            detected_category="document",
            workflow_category="document",
            detection_method=DetectionMethod.CONTAINER,
            confidence=DetectionConfidence.CERTAIN,
            structure_status=StructureStatus.VALID,
            relation=FormatRelation.EXACT_MATCH,
            decision=AdmissionDecision.ALLOW,
            declared_supported=True,
            detected_supported=True,
        )
        bridge_entered = threading.Event()
        bridge_saw_cancel = threading.Event()
        staging_roots: list[Path] = []

        def fake_pre_convert(
            _input_path: str,
            _source_format: str,
            *,
            staging_dir: str,
            cancel: object,
            **_kwargs: object,
        ) -> PreConversionFailure:
            staging_roots.append(Path(staging_dir).parent)
            bridge_entered.set()
            deadline = time.monotonic() + 5.0
            while time.monotonic() < deadline:
                if bool(getattr(cancel, "is_cancelled", False)):
                    bridge_saw_cancel.set()
                    return PreConversionFailure(message="cancelled", cancelled=True)
                time.sleep(0.01)
            return PreConversionFailure(message="test bridge timed out")

        # This test owns the cancellation boundary, not OLE parsing.  Feed the
        # same concrete, content-admitted identity to GUI and Application so
        # production ``resolve_chain('doc', 'md')`` selects preconversion.
        window.view_model._file_inspector = lambda _path: admitted_doc
        monkeypatch.setattr("docwen_core.detection.inspect_file", lambda _path: admitted_doc)
        monkeypatch.setattr("docwen_core.detection._validation.inspect_file", lambda _path: admitted_doc)
        monkeypatch.setattr(
            "docwen_application.preconversion.pre_converter.pre_convert",
            fake_pre_convert,
        )

        window.view_model.add_files([str(source)])
        app = QApplication.instance()
        if app is not None:
            app.processEvents()
        window._action_area_vm.request_conversion("md")

        assert _wait_for(bridge_entered.is_set, timeout_ms=5000, interval_ms=20)
        assert window._action_area.cancel_button is not None
        window._action_area.cancel_button.click()

        def _cancelled() -> bool:
            entry = window._batch_list_vm.get_file_entry(normalized)
            return entry is not None and entry.status == "cancelled"

        assert _wait_for(_cancelled, timeout_ms=10000, interval_ms=20)
        assert bridge_saw_cancel.is_set()
        assert _wait_for(lambda: not window._active_threads, timeout_ms=5000, interval_ms=20)
        entry = window._batch_list_vm.get_file_entry(normalized)
        assert entry is not None
        assert entry.status == "cancelled"
        assert not entry.output_path
        assert window._info_area_vm._task_summary.state == "cancelled"
        assert window._info_area_vm._task_summary.failed_count == 0
        assert window._info_area_vm._task_summary.cancelled_count == 1
        assert staging_roots and all(not root.exists() for root in staging_roots)

    def test_cancel_button_cancels_running_runtime_task(self, qapp, tmp_path: Path) -> None:
        from PySide6.QtWidgets import QApplication

        from docwen_application.controller import ApplicationController
        from docwen_core.models.artifact import ARTIFACT_KIND_PRIMARY, ArtifactManifest
        from docwen_core.models.manifest import PluginManifest, RouteSpec
        from docwen_core.models.result import ConversionResult
        from docwen_gui.app import create_main_window
        from docwen_gui.main_window import _normalize_path
        from docwen_gui.qt_bridge.task_event_bridge import TaskEventBridge
        from docwen_runtime.adapters import RuntimePortAdapter
        from docwen_runtime.capabilities import build_runtime_capability_projection
        from docwen_runtime.engine.route_resolver import RouteResolver
        from docwen_runtime.engine.task_manager import TaskManager
        from docwen_runtime.output.finalizer import OutputFinalizer
        from docwen_runtime.plugin_registry.registry import PluginRegistry
        from docwen_runtime.workspace.manager import WorkspaceManager

        class SlowCancellablePlugin:
            @property
            def manifest(self) -> PluginManifest:
                return PluginManifest(
                    plugin_id="slow_cancellable",
                    name="Slow Cancellable",
                    version="0.1.0",
                    routes=[
                        RouteSpec(
                            source_format="markdown",
                            target_format="docx",
                            label="Markdown -> DOCX",
                        )
                    ],
                )

            def can_handle(self, source_format: str, target_format: str, action_name: str = "") -> bool:
                return source_format == "markdown" and target_format == "docx" and not action_name

            def convert(self, context) -> ConversionResult:
                for index in range(200):
                    context.cancellation.check()
                    if index % 10 == 0:
                        context.progress.report_progress(float(index) / 2.0, "Working")
                    time.sleep(0.01)

                output_path = context.workspace.create_artifact_path("slow", ".docx")
                Path(output_path).write_text("finished", encoding="utf-8")
                artifact = ArtifactManifest(
                    artifact_id="slow-docx",
                    kind=ARTIFACT_KIND_PRIMARY,
                    staging_path=output_path,
                    suggested_name=f"{Path(context.workspace.input_path).stem}.docx",
                    is_primary=True,
                )
                return ConversionResult(task_id=context.request.request_id, success=True, artifacts=[artifact])

        bridge = TaskEventBridge()

        def _event_callback(event) -> None:
            bridge.enqueue(event.event_type, {"task_id": event.task_id, **dict(event.payload)})

        registry = PluginRegistry()
        registry.register(SlowCancellablePlugin())
        workspace_manager = WorkspaceManager(root_dir=str(tmp_path / "workspace"))
        task_manager = TaskManager(
            registry,
            RouteResolver(registry),
            workspace_manager,
            OutputFinalizer(),
        )
        controller = ApplicationController(
            runtime_port=RuntimePortAdapter(
                task_manager,
                event_callback=_event_callback,
                capability_provider=lambda: build_runtime_capability_projection(
                    registry.list_manifests(),
                    platform_id="windows",
                    egress_guard_status={},
                ),
            )
        )
        controller.start()
        window = create_main_window(controller=controller, task_event_bridge=bridge)

        try:
            source = tmp_path / "cancel-me.md"
            source.write_text("# Cancel me", encoding="utf-8")
            normalized = _normalize_path(str(source))
            window.view_model.add_files([str(source)])

            app = QApplication.instance()
            if app is not None:
                app.processEvents()

            window._action_area_vm.request_conversion("docx")

            def _running() -> bool:
                entry = window._batch_list_vm.get_file_entry(normalized)
                return bool(window._active_threads) and entry is not None and entry.status == "processing"

            assert _wait_for(_running, timeout_ms=_E2E_CONVERSION_TIMEOUT_MS, interval_ms=20)
            assert _wait_for(
                lambda: window.view_model.status_message.startswith(t("main_window.task_progress_prefix")),
                timeout_ms=_E2E_CONVERSION_TIMEOUT_MS,
                interval_ms=20,
            )
            assert window._info_area_vm.status_source == "transient"
            assert window._info_area_vm.status_summary_text.startswith(t("main_window.task_progress_prefix"))
            assert window._info_area_vm.activity_enabled is True
            assert window._action_area_vm.cancel_visible is True
            assert window._action_area.cancel_button is not None
            window._action_area.cancel_button.click()

            def _cancelled() -> bool:
                entry = window._batch_list_vm.get_file_entry(normalized)
                return entry is not None and entry.status == "cancelled"

            assert _wait_for(_cancelled, timeout_ms=_E2E_CONVERSION_TIMEOUT_MS, interval_ms=20)
            entry = window._batch_list_vm.get_file_entry(normalized)
            assert entry is not None
            assert entry.status == "cancelled"
            assert not entry.output_path
            assert entry.error_message
            assert window._action_area_vm.cancel_visible is False
            assert window._info_area_vm.has_task_summary
            assert window._info_area_vm._task_summary.state == "cancelled"
            assert window._info_area_vm._task_summary.failed_count == 0
            assert window._info_area_vm._task_summary.cancelled_count == 1
            assert window._info_area_vm.guide_visible is True
            assert window._info_area_vm.guide_actions == [{"action_key": "add_more_files", "target_path": ""}]
        finally:
            shutdown_main_window(window)


@pytest.mark.gui
class TestGuiBatchExecution:
    def test_batch_mode_conversion_uses_execute_batch_and_reports_partial_summary(self, qapp, tmp_path: Path) -> None:
        from PySide6.QtWidgets import QApplication

        from docwen_application.controller import ApplicationController
        from docwen_core.models.artifact import ARTIFACT_KIND_PRIMARY, ArtifactManifest
        from docwen_core.models.manifest import PluginManifest, RouteSpec
        from docwen_core.models.result import ConversionErrorInfo, ConversionResult
        from docwen_gui.app import create_main_window
        from docwen_gui.main_window import _normalize_path
        from docwen_runtime.adapters import RuntimePortAdapter
        from docwen_runtime.capabilities import build_runtime_capability_projection
        from docwen_runtime.engine.route_resolver import RouteResolver
        from docwen_runtime.engine.task_manager import TaskManager
        from docwen_runtime.output.finalizer import OutputFinalizer
        from docwen_runtime.plugin_registry.registry import PluginRegistry
        from docwen_runtime.workspace.manager import WorkspaceManager

        class PartiallyFailingMarkdownPlugin:
            @property
            def manifest(self) -> PluginManifest:
                return PluginManifest(
                    plugin_id="batch_partial",
                    name="Batch Partial",
                    version="0.1.0",
                    routes=[
                        RouteSpec(
                            source_format="markdown",
                            target_format="docx",
                            label="Markdown -> DOCX",
                        )
                    ],
                )

            def can_handle(self, source_format: str, target_format: str, action_name: str = "") -> bool:
                return source_format == "markdown" and target_format == "docx" and not action_name

            def convert(self, context) -> ConversionResult:
                input_path = Path(context.workspace.input_path)
                if "bad" in input_path.name:
                    return ConversionResult(
                        task_id=context.request.request_id,
                        success=False,
                        error=ConversionErrorInfo(
                            error_type="conversion_failed",
                            message="batch boom",
                        ),
                    )

                output_path = context.workspace.create_artifact_path(input_path.stem, ".docx")
                Path(output_path).write_text(f"converted {input_path.name}", encoding="utf-8")
                artifact = ArtifactManifest(
                    artifact_id=f"{input_path.stem}-docx",
                    kind=ARTIFACT_KIND_PRIMARY,
                    staging_path=output_path,
                    suggested_name=f"{input_path.stem}.docx",
                    is_primary=True,
                )
                return ConversionResult(task_id=context.request.request_id, success=True, artifacts=[artifact])

        registry = PluginRegistry()
        registry.register(PartiallyFailingMarkdownPlugin())
        workspace_manager = WorkspaceManager(root_dir=str(tmp_path / "workspace"))
        task_manager = TaskManager(
            registry,
            RouteResolver(registry),
            workspace_manager,
            OutputFinalizer(),
        )
        controller = ApplicationController(
            runtime_port=RuntimePortAdapter(
                task_manager,
                capability_provider=lambda: build_runtime_capability_projection(
                    registry.list_manifests(),
                    platform_id="windows",
                    egress_guard_status={},
                ),
            )
        )
        controller.start()
        window = create_main_window(controller=controller)

        try:
            ok_source = tmp_path / "ok.md"
            bad_source = tmp_path / "bad.md"
            ok_source.write_text("# OK", encoding="utf-8")
            bad_source.write_text("# Bad", encoding="utf-8")
            ok_norm = _normalize_path(str(ok_source))
            bad_norm = _normalize_path(str(bad_source))

            window.view_model.mode = "batch"
            window.view_model.add_files([str(ok_source), str(bad_source)])
            app = QApplication.instance()
            if app is not None:
                app.processEvents()

            window._action_area_vm.request_conversion("docx")

            def _batch_finished() -> bool:
                ok_entry = window._batch_list_vm.get_file_entry(ok_norm)
                bad_entry = window._batch_list_vm.get_file_entry(bad_norm)
                return (
                    ok_entry is not None
                    and bad_entry is not None
                    and ok_entry.status == "completed"
                    and bad_entry.status == "failed"
                )

            assert _wait_for(_batch_finished, timeout_ms=_E2E_CONVERSION_TIMEOUT_MS, interval_ms=20)

            ok_entry = window._batch_list_vm.get_file_entry(ok_norm)
            bad_entry = window._batch_list_vm.get_file_entry(bad_norm)
            assert ok_entry is not None
            assert bad_entry is not None
            assert ok_entry.output_path
            assert Path(ok_entry.output_path).exists()
            assert bad_entry.error_message == "batch boom"
            assert window._action_area_vm.cancel_visible is False
            assert window._info_area_vm.has_task_summary
            summary = window._info_area_vm._task_summary
            assert summary.state == "partial"
            assert summary.completed_count == 2
            assert summary.total_count == 2
            assert summary.failed_count == 1
            assert summary.navigate_path == bad_norm
            assert [action["action_key"] for action in window._info_area_vm.guide_actions] == [
                "open_output_dir",
                "view_failed_details",
                "retry_failed",
                "add_more_files",
            ]
        finally:
            shutdown_main_window(window)


@pytest.mark.gui
class TestMarkdownTemplateWorkflow:
    def test_selected_docx_template_is_used_in_gui_markdown_to_docx_run(
        self,
        main_window_with_controller,
        monkeypatch: pytest.MonkeyPatch,
        tmp_path: Path,
    ) -> None:
        from docx import Document
        from PySide6.QtWidgets import QApplication

        from docwen_gui.main_window import _normalize_path
        from docwen_runtime.templates import TemplateRegistry

        templates_dir = tmp_path / "templates"
        templates_dir.mkdir()
        template_path = templates_dir / "Corporate Report.docx"
        template_doc = Document()
        template_doc.add_paragraph("TEMPLATE SENTINEL BEFORE")
        template_doc.add_paragraph("{{body}}")
        template_doc.add_paragraph("TEMPLATE SENTINEL AFTER")
        template_doc.save(str(template_path))

        registry = TemplateRegistry(templates_dir)
        monkeypatch.setattr(TemplateRegistry, "default", staticmethod(lambda: registry))

        source = tmp_path / "brief.md"
        source.write_text("# Smoke Title\n\nBody paragraph from GUI template workflow.", encoding="utf-8")

        window = main_window_with_controller
        normalized = _normalize_path(str(source))
        window._load_templates_into_main_selector()
        docx_selector = window._template_selector.get_selector("docx")
        assert docx_selector is not None
        assert docx_selector.has_template("Corporate Report")
        docx_selector.select_template("Corporate Report", selection_source="user")

        window.view_model.add_files([str(source)])
        app = QApplication.instance()
        if app is not None:
            app.processEvents()

        assert window._action_area_vm.visible
        window._action_area_vm.request_conversion("docx")

        if app is not None:
            app.processEvents()

        def _conversion_finished() -> bool:
            entry = window._batch_list_vm.get_file_entry(normalized)
            return entry is not None and entry.status in ("completed", "failed")

        assert _wait_for(_conversion_finished, timeout_ms=_E2E_CONVERSION_TIMEOUT_MS)
        entry = window._batch_list_vm.get_file_entry(normalized)
        assert entry is not None
        assert entry.status == "completed"
        assert entry.output_path

        output_path = Path(entry.output_path)
        assert output_path.exists()
        assert output_path.name == "brief.docx"
        output_doc = Document(str(output_path))
        paragraph_text = [paragraph.text for paragraph in output_doc.paragraphs if paragraph.text.strip()]

        assert "TEMPLATE SENTINEL BEFORE" in paragraph_text
        assert "TEMPLATE SENTINEL AFTER" in paragraph_text
        assert "Smoke Title" in paragraph_text
        assert "Body paragraph from GUI template workflow." in paragraph_text
        assert window._info_area_vm.has_task_summary
        assert window._info_area_vm._task_summary.state == "success"
