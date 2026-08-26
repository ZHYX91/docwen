"""DOCX adapter for the provider-neutral semantic document model.

The public model remains free of python-docx types.  This adapter imports
python-docx lazily and writes only standard WordprocessingML structures.
"""

from __future__ import annotations

import re
from copy import deepcopy
from dataclasses import dataclass
from itertools import pairwise
from typing import Any
from urllib.parse import urlsplit

from docwen_core.docx_bookmarks import (
    BOOKMARK_ID_MAX,
    DocxBookmarkInventory,
    bookmark_id_key,
    build_docx_bookmark_inventory,
    is_legal_bookmark_name,
    prove_bookmark_name,
)
from docwen_core.docx_parsing.document_semantics import (
    BIBLIOGRAPHY_BOOKMARK_NAME,
    BIBLIOGRAPHY_ENTRY_BOOKMARK_PREFIX,
    CITATION_BOOKMARK_PREFIX,
    decode_bibliography_entry_bookmark,
    decode_citation_bookmark,
    decode_target_bookmark,
    encode_bibliography_entry_bookmark,
    encode_caption_pairing_bookmark,
    encode_citation_bookmark,
    encode_object_bookmark,
    encode_object_pairing_bookmark,
    encode_shorthand_bookmark,
    encode_target_bookmark,
    extract_neutral_semantic_caption,
    extract_semantic_table_metadata,
    inspect_caption_pairing_markers,
    inspect_caption_target_markers,
    inspect_object_pairing_markers,
    inspect_object_target_markers,
)
from docwen_core.docx_parsing.table_extraction import DocxTableGeometryError, build_docx_table_semantic_grid
from docwen_core.models.semantic_document import (
    SemanticBibliographyEntry,
    SemanticBibliographyFragment,
    SemanticBibliographyRun,
    SemanticCaption,
    SemanticCitationCluster,
    SemanticCitationItem,
    SemanticDiagnostic,
    SemanticDocument,
    SemanticDocumentValidationError,
    SemanticImportResult,
    SemanticParagraph,
    SemanticReference,
    SemanticTable,
    SemanticTableCell,
    SemanticText,
    derive_table_header_shape,
    is_portable_semantic_id,
    validate_semantic_document,
)


class DocxSemanticRenderer:
    """Render neutral semantic objects to standard DOCX structures."""

    def __init__(self, document: Any, *, bookmark_id_start: int = 1000) -> None:
        if type(bookmark_id_start) is not int:
            raise TypeError("bookmark_id_start must be an integer")
        if not 0 <= bookmark_id_start <= BOOKMARK_ID_MAX:
            raise ValueError(f"bookmark_id_start must be between 0 and {BOOKMARK_ID_MAX}")
        self._document = document
        self._bookmark_id = bookmark_id_start
        self._pairing_token = 1

    def render_blocks(self, blocks: tuple[SemanticParagraph | SemanticTable, ...]) -> list[Any]:
        """Append validated neutral blocks and return their python-docx objects."""

        semantic_document = SemanticDocument(blocks=blocks)
        diagnostics = validate_semantic_document(semantic_document)
        if diagnostics:
            raise SemanticDocumentValidationError(diagnostics)

        rendered: list[Any] = []
        for block in blocks:
            if isinstance(block, SemanticParagraph):
                rendered.append(self.render_paragraph(block))
            else:
                rendered.extend(self.render_table(block))
        return rendered

    def render_paragraph(self, paragraph_data: SemanticParagraph) -> Any:
        paragraph = self._document.add_paragraph()
        for inline in paragraph_data.inlines:
            if isinstance(inline, SemanticText):
                paragraph.add_run(inline.value)
            elif isinstance(inline, SemanticReference):
                self.render_reference(paragraph, inline)
            else:
                self.render_citation(paragraph, inline)
        return paragraph

    def render_caption(
        self,
        caption: SemanticCaption,
        *,
        source_form: str = "imported",
    ) -> Any:
        """Render a caption with native SEQ and a visible cached result."""

        bookmark_names: list[str] = []
        shorthand_name = None
        if source_form == "shorthand" and caption.target_id is not None:
            shorthand_name = encode_shorthand_bookmark(caption.target_id)
            bookmark_names.append(shorthand_name)
        target_name = None
        if caption.target_id is not None:
            target_name = encode_target_bookmark(caption.target_id)
            bookmark_names.append(target_name)
        reservations = self._reserve_semantic_bookmarks(tuple(bookmark_names))
        return self._render_caption_with_reservations(
            caption,
            shorthand_name=shorthand_name,
            target_name=target_name,
            pairing_name=None,
            reservations=reservations,
        )

    def render_caption_for_table(
        self,
        table: Any,
        caption: SemanticCaption,
        *,
        source_form: str = "imported",
    ) -> Any:
        """Atomically render and bind a caption to an existing physical table."""

        table_parent, object_anchor = self._validate_table_binding_target(table)
        if caption.kind != "table":
            diagnostic = SemanticDiagnostic(
                "error",
                "semantic.caption.object_kind_mismatch",
                "A physical table caption must use kind 'table'.",
                "caption",
            )
            raise SemanticDocumentValidationError((diagnostic,))

        shorthand_name = None
        target_name = None
        pairing_name = None
        object_name: str
        if caption.target_id is None:
            pairing_token = self._next_pairing_token()
            pairing_name = encode_caption_pairing_bookmark(pairing_token)
            object_name = encode_object_pairing_bookmark(pairing_token)
        else:
            target_name = encode_target_bookmark(caption.target_id)
            object_name = encode_object_bookmark(caption.target_id)
            if source_form == "shorthand":
                shorthand_name = encode_shorthand_bookmark(caption.target_id)

        bookmark_names = tuple(
            name for name in (shorthand_name, target_name, pairing_name, object_name) if name is not None
        )
        reservations = self._reserve_semantic_bookmarks(bookmark_names)
        paragraph = self._render_caption_with_reservations(
            caption,
            shorthand_name=shorthand_name,
            target_name=target_name,
            pairing_name=pairing_name,
            reservations=reservations,
        )
        caption_parent = paragraph._element.getparent()
        if caption_parent is None:
            raise ValueError("rendered caption is detached")
        caption_parent.remove(paragraph._element)
        table_parent.insert(table_parent.index(table._element), paragraph._element)
        append_zero_width_bookmark(
            object_anchor,
            object_name,
            reservations[object_name],
        )
        return paragraph

    def _render_caption_with_reservations(
        self,
        caption: SemanticCaption,
        *,
        shorthand_name: str | None,
        target_name: str | None,
        pairing_name: str | None,
        reservations: dict[str, str],
    ) -> Any:
        paragraph = self._document.add_paragraph()
        paragraph.add_run(f"{caption.label} ")
        if pairing_name is not None:
            append_zero_width_bookmark(paragraph, pairing_name, reservations[pairing_name])
        if shorthand_name is not None:
            append_zero_width_bookmark(paragraph, shorthand_name, reservations[shorthand_name])
        append_complex_field(
            paragraph,
            instruction=f" SEQ {caption.kind.title()} \\* ARABIC ",
            cached_result=caption.cached_number,
            bookmark_name=target_name,
            bookmark_id=(reservations[target_name] if target_name is not None else None),
        )
        paragraph.add_run(f": {caption.content}")
        return paragraph

    def render_reference(self, paragraph: Any, reference: SemanticReference) -> None:
        """Render a target-bound REF field with an explicit cached result."""

        self.append_complex_field(
            paragraph,
            instruction=f" REF {encode_target_bookmark(reference.target_id)} \\h ",
            cached_result=reference.cached_result,
        )

    def render_citation(self, paragraph: Any, citation: SemanticCitationCluster) -> None:
        """Render one locked, clean Word citation field inside a cluster marker."""

        self._validate_paragraph_target(paragraph, purpose="citation")
        diagnostics = validate_semantic_document(SemanticDocument(blocks=(SemanticParagraph((citation,)),)))
        if diagnostics:
            raise SemanticDocumentValidationError(diagnostics)
        bookmark_name = encode_citation_bookmark(citation.cluster_id)
        reservations = self._reserve_semantic_bookmarks((bookmark_name,))
        append_complex_field(
            paragraph,
            instruction=_citation_instruction(citation),
            cached_result=citation.cached_result,
            bookmark_name=bookmark_name,
            bookmark_id=reservations[bookmark_name],
            locked=True,
            dirty=False,
        )

    def render_table(self, table_data: SemanticTable) -> list[Any]:
        """Render an explicit anchor-only table and its bound caption."""

        diagnostics = validate_semantic_document(SemanticDocument(blocks=(table_data,)))
        if diagnostics:
            raise SemanticDocumentValidationError(diagnostics)

        table = self._document.add_table(rows=table_data.row_count, cols=table_data.column_count)
        for cell_data in table_data.cells:
            anchor = table.cell(cell_data.row, cell_data.column)
            anchor.text = cell_data.text
        for cell_data in table_data.cells:
            if cell_data.row_span == 1 and cell_data.column_span == 1:
                continue
            table.cell(cell_data.row, cell_data.column).merge(
                table.cell(
                    cell_data.row + cell_data.row_span - 1,
                    cell_data.column + cell_data.column_span - 1,
                )
            )

        header_rows, header_columns = derive_table_header_shape(table_data)
        apply_semantic_table_roles(
            table,
            header_rows=header_rows,
            header_columns=header_columns,
            repeat_header=table_data.repeat_header,
        )
        rendered: list[Any] = []
        if table_data.caption is not None:
            caption = self.render_caption_for_table(table, table_data.caption)
            rendered.append(caption)
        rendered.append(table)
        return rendered

    def bind_object_target(self, table: Any, target_id: str) -> None:
        """Bind a physical table to its caption target without adjacency rules."""

        _table_parent, object_anchor = self._validate_table_binding_target(table)
        bookmark_name = encode_object_bookmark(target_id)
        reservations = self._reserve_semantic_bookmarks((bookmark_name,))
        append_zero_width_bookmark(
            object_anchor,
            bookmark_name,
            reservations[bookmark_name],
        )

    def bind_object_pairing(self, table: Any, pairing_token: str) -> None:
        """Bind an unaddressable table to its caption with an internal marker."""

        _table_parent, object_anchor = self._validate_table_binding_target(table)
        bookmark_name = encode_object_pairing_bookmark(pairing_token)
        reservations = self._reserve_semantic_bookmarks((bookmark_name,))
        append_zero_width_bookmark(
            object_anchor,
            bookmark_name,
            reservations[bookmark_name],
        )

    def render_bibliography_fragment(
        self,
        fragment: SemanticBibliographyFragment,
        *,
        placeholder_anchor: Any,
        fallback_style_id: str | None = None,
        hyperlink_style_id: str | None = None,
    ) -> tuple[Any, ...]:
        """Replace one explicit anchor with marked, already-formatted entries."""

        self._validate_paragraph_target(placeholder_anchor, purpose="bibliography anchor")
        parent = placeholder_anchor._element.getparent()
        if parent is None:
            raise ValueError("bibliography placeholder anchor is detached")
        diagnostics = validate_semantic_document(SemanticDocument(blocks=(), bibliography=fragment))
        if diagnostics:
            raise SemanticDocumentValidationError(diagnostics)
        if fallback_style_id is not None:
            _require_style_type(self._document, fallback_style_id, "paragraph")
        if hyperlink_style_id is not None:
            _require_style_type(self._document, hyperlink_style_id, "character")
        if (
            any(run.href is not None for entry in fragment.entries for run in entry.runs)
            and hyperlink_style_id != "Hyperlink"
        ):
            raise _bibliography_diagnostic(
                "semantic.docx.bibliography.style_invalid",
                "Linked bibliography runs require the canonical Hyperlink character style.",
            )

        anchor_element = placeholder_anchor._element
        anchor_ppr = anchor_element.pPr
        anchor_styles = [] if anchor_ppr is None else list(anchor_ppr.findall(_qn("w:pStyle")))
        if len(anchor_styles) > 1:
            raise _bibliography_diagnostic(
                "semantic.docx.bibliography.style_invalid",
                "Bibliography anchor must not contain duplicate paragraph-style references.",
            )
        if anchor_styles:
            anchor_style_id = anchor_styles[0].get(_qn("w:val"))
            if anchor_style_id is None or not _has_style_type(self._document, anchor_style_id, "paragraph"):
                raise _bibliography_diagnostic(
                    "semantic.docx.bibliography.style_invalid",
                    "Bibliography anchor paragraph style must resolve to exactly one paragraph style.",
                )
        section_properties = [] if anchor_ppr is None else list(anchor_ppr.findall(_qn("w:sectPr")))
        if len(section_properties) > 1:
            raise _bibliography_section_error("Bibliography anchor has multiple section property elements.")
        if not fragment.entries:
            if section_properties:
                previous = anchor_element.getprevious()
                if previous is None or previous.tag != _qn("w:p"):
                    raise _bibliography_section_error(
                        "An empty bibliography cannot safely transfer section properties without a preceding paragraph."
                    )
                previous_ppr = previous.pPr
                if previous_ppr is not None and previous_ppr.find(_qn("w:sectPr")) is not None:
                    raise _bibliography_section_error(
                        "An empty bibliography cannot overwrite preceding section properties."
                    )
                _transfer_section_properties_and_remove_anchor(
                    parent,
                    anchor_element,
                    previous,
                    previous_ppr,
                    section_properties[0],
                )
                return ()
            parent.remove(anchor_element)
            return ()

        entry_names = tuple(encode_bibliography_entry_bookmark(entry.item_id) for entry in fragment.entries)
        bookmark_id_before = self._bookmark_id
        reservations = self._reserve_semantic_bookmarks((BIBLIOGRAPHY_BOOKMARK_NAME, *entry_names))
        existing_relationship_ids = set(self._document.part.rels)
        paragraphs: list[Any] = []
        inserted_elements: list[Any] = []
        try:
            for offset, entry in enumerate(fragment.entries):
                paragraph = _detached_bibliography_paragraph(
                    self._document,
                    anchor_ppr,
                    keep_section=offset == len(fragment.entries) - 1,
                    fallback_style_id=fallback_style_id,
                )
                for run_data in entry.runs:
                    _append_bibliography_run(
                        paragraph,
                        run_data,
                        hyperlink_style_id=hyperlink_style_id,
                    )
                paragraphs.append(paragraph)
            boundary_id = reservations[BIBLIOGRAPHY_BOOKMARK_NAME]
            for offset, (paragraph, entry_name) in enumerate(zip(paragraphs, entry_names, strict=True)):
                prepend_bookmark_start(paragraph, entry_name, reservations[entry_name])
                if offset == 0:
                    prepend_bookmark_start(paragraph, BIBLIOGRAPHY_BOOKMARK_NAME, boundary_id)
                append_bookmark_end(paragraph, reservations[entry_name])
                if offset == len(paragraphs) - 1:
                    append_bookmark_end(paragraph, boundary_id)
            anchor_index = parent.index(anchor_element)
            for offset, paragraph in enumerate(paragraphs):
                parent.insert(anchor_index + offset, paragraph._element)
                inserted_elements.append(paragraph._element)
            parent.remove(anchor_element)
        except Exception:
            for element in inserted_elements:
                element_parent = element.getparent()
                if element_parent is not None:
                    element_parent.remove(element)
            self._bookmark_id = bookmark_id_before
            for relationship_id in set(self._document.part.rels) - existing_relationship_ids:
                del self._document.part.rels[relationship_id]
            raise
        return tuple(paragraphs)

    def append_zero_width_bookmark(self, paragraph: Any, bookmark_name: str) -> None:
        reservations = self._reserve_semantic_bookmarks((bookmark_name,))
        append_zero_width_bookmark(paragraph, bookmark_name, reservations[bookmark_name])

    def append_complex_field(
        self,
        paragraph: Any,
        *,
        instruction: str,
        cached_result: str,
        bookmark_name: str | None = None,
        locked: bool = False,
        dirty: bool = True,
    ) -> None:
        reservations = self._reserve_semantic_bookmarks((bookmark_name,)) if bookmark_name is not None else {}
        append_complex_field(
            paragraph,
            instruction=instruction,
            cached_result=cached_result,
            bookmark_name=bookmark_name,
            bookmark_id=(reservations[bookmark_name] if bookmark_name is not None else None),
            locked=locked,
            dirty=dirty,
        )

    def _reserve_semantic_bookmarks(self, bookmark_names: tuple[str, ...]) -> dict[str, str]:
        if not bookmark_names:
            return {}
        name_keys = tuple(name.casefold() for name in bookmark_names)
        if len(name_keys) != len(set(name_keys)):
            raise ValueError("semantic bookmark names must be unique within one reservation")
        invalid_names = [name for name in bookmark_names if not is_legal_bookmark_name(name)]
        if invalid_names:
            raise ValueError("semantic bookmark name is outside the portable Word subset")

        inventory = build_docx_bookmark_inventory(self._document)
        conflicts = [name for name in bookmark_names if name.casefold() in inventory.used_name_keys]
        if conflicts:
            diagnostic = SemanticDiagnostic(
                "error",
                "semantic.docx.bookmark.name_conflict",
                f"Semantic bookmark name {conflicts[0]} already exists in the DOCX package.",
                "document",
            )
            raise SemanticDocumentValidationError((diagnostic,))

        used_ids = set(inventory.used_id_keys)
        reservations: dict[str, str] = {}
        candidate = self._bookmark_id
        for name in bookmark_names:
            while candidate <= BOOKMARK_ID_MAX and bookmark_id_key(str(candidate)) in used_ids:
                candidate += 1
            if candidate > BOOKMARK_ID_MAX:
                raise OverflowError("no portable DOCX bookmark IDs remain")
            bookmark_id = str(candidate)
            reservations[name] = bookmark_id
            used_ids.add(bookmark_id_key(bookmark_id))
            candidate += 1
        self._bookmark_id = candidate
        return reservations

    def _validate_table_binding_target(self, table: Any) -> tuple[Any, Any]:
        if table.part is not self._document.part:
            raise ValueError("table must belong to the renderer document part")
        table_parent = table._element.getparent()
        if table_parent is None:
            raise ValueError("a caption cannot be bound to a detached table")
        if not table.rows or not table.rows[0].cells:
            raise ValueError("a caption can be bound only to a non-empty table")
        first_cell_paragraphs = table.rows[0].cells[0].paragraphs
        if not first_cell_paragraphs:
            raise ValueError("a caption can be bound only to a table with an anchor paragraph")
        return table_parent, first_cell_paragraphs[0]

    def _validate_paragraph_target(self, paragraph: Any, *, purpose: str) -> None:
        if paragraph.part is not self._document.part:
            raise ValueError(f"{purpose} paragraph must belong to the renderer document part")
        parent = paragraph._element.getparent()
        if parent is None:
            raise ValueError(f"{purpose} paragraph must be attached")
        if parent is not self._document.element.body:
            raise ValueError(f"{purpose} paragraph must be a direct main document body paragraph")

    def _next_pairing_token(self) -> str:
        used_name_keys = build_docx_bookmark_inventory(self._document).used_name_keys
        while True:
            token = str(self._pairing_token)
            self._pairing_token += 1
            caption_name = encode_caption_pairing_bookmark(token)
            object_name = encode_object_pairing_bookmark(token)
            if caption_name.casefold() not in used_name_keys and object_name.casefold() not in used_name_keys:
                return token


def _qn(tag: str) -> str:
    from docx.oxml.ns import qn

    return qn(tag)


def _bibliography_diagnostic(code: str, message: str) -> SemanticDocumentValidationError:
    return SemanticDocumentValidationError(
        (
            SemanticDiagnostic(
                "error",
                code,
                message,
                "document/bibliography",
            ),
        )
    )


def _bibliography_section_error(message: str) -> SemanticDocumentValidationError:
    return _bibliography_diagnostic("semantic.docx.bibliography.section_unsafe", message)


def _require_style_type(document: Any, style_id: str, expected_type: str) -> None:
    matching = [
        style for style in document.styles.element.findall(_qn("w:style")) if style.get(_qn("w:styleId")) == style_id
    ]
    if len(matching) != 1 or matching[0].get(_qn("w:type")) != expected_type:
        raise _bibliography_diagnostic(
            "semantic.docx.bibliography.style_invalid",
            f"Bibliography style {style_id!r} must resolve to exactly one {expected_type} style.",
        )


def _has_style_type(document: Any, style_id: str, expected_type: str) -> bool:
    matching = [
        style for style in document.styles.element.findall(_qn("w:style")) if style.get(_qn("w:styleId")) == style_id
    ]
    return len(matching) == 1 and matching[0].get(_qn("w:type")) == expected_type


def _transfer_section_properties_and_remove_anchor(
    parent: Any,
    anchor_element: Any,
    previous: Any,
    previous_ppr: Any | None,
    section_properties: Any,
) -> None:
    from docx.oxml import OxmlElement

    original_ppr = None if previous_ppr is None else deepcopy(previous_ppr)
    replacement_ppr = deepcopy(previous_ppr) if previous_ppr is not None else OxmlElement("w:pPr")
    section_copy = deepcopy(section_properties)
    change_tracking = replacement_ppr.find(_qn("w:pPrChange"))
    if change_tracking is None:
        replacement_ppr.append(section_copy)
    else:
        replacement_ppr.insert(replacement_ppr.index(change_tracking), section_copy)
    try:
        if previous_ppr is None:
            previous.insert(0, replacement_ppr)
        else:
            previous.replace(previous_ppr, replacement_ppr)
        parent.remove(anchor_element)
    except Exception:
        current_ppr = previous.pPr
        if original_ppr is None:
            if current_ppr is not None:
                previous.remove(current_ppr)
        elif current_ppr is not None:
            previous.replace(current_ppr, original_ppr)
        raise


def _detached_bibliography_paragraph(
    document: Any,
    anchor_ppr: Any | None,
    *,
    keep_section: bool,
    fallback_style_id: str | None,
) -> Any:
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    paragraph_element: Any = OxmlElement("w:p")
    paragraph_ppr = deepcopy(anchor_ppr) if anchor_ppr is not None else None
    if paragraph_ppr is not None and not keep_section:
        for section_properties in paragraph_ppr.findall(_qn("w:sectPr")):
            paragraph_ppr.remove(section_properties)
    if paragraph_ppr is None and fallback_style_id is not None:
        paragraph_ppr = OxmlElement("w:pPr")
    if paragraph_ppr is not None and fallback_style_id is not None:
        style = paragraph_ppr.find(_qn("w:pStyle"))
        if style is None:
            style = OxmlElement("w:pStyle")
            style.set(_qn("w:val"), fallback_style_id)
            paragraph_ppr.insert(0, style)
    if paragraph_ppr is not None:
        paragraph_element.append(paragraph_ppr)
    return Paragraph(paragraph_element, document)


def _append_bibliography_run(
    paragraph: Any,
    run_data: SemanticBibliographyRun,
    *,
    hyperlink_style_id: str | None,
) -> None:
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement
    from docx.text.run import Run

    run_element: Any = OxmlElement("w:r")
    run = Run(run_element, paragraph)
    if run_data.bold:
        run.bold = True
    if run_data.italic:
        run.italic = True
    _append_exact_run_text(run_element, run_data.text)
    if run_data.href is None:
        paragraph._p.append(run_element)
        return

    relationship_id = paragraph.part.relate_to(run_data.href, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(_qn("r:id"), relationship_id)
    if hyperlink_style_id is not None:
        run_properties = run_element.get_or_add_rPr()
        style = OxmlElement("w:rStyle")
        style.set(_qn("w:val"), hyperlink_style_id)
        run_properties.insert(0, style)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def _append_exact_run_text(run_element: Any, text: str) -> None:
    from docx.oxml import OxmlElement

    plain: list[str] = []

    def flush_plain() -> None:
        if not plain:
            return
        value = "".join(plain)
        text_element = OxmlElement("w:t")
        if value[:1].isspace() or value[-1:].isspace():
            text_element.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        text_element.text = value
        run_element.append(text_element)
        plain.clear()

    for character in text:
        if character not in {"\t", "\n", "\r"}:
            plain.append(character)
            continue
        flush_plain()
        token = {"\t": "w:tab", "\n": "w:br", "\r": "w:cr"}[character]
        run_element.append(OxmlElement(token))
    flush_plain()


def _citation_instruction(citation: SemanticCitationCluster) -> str:
    parts = ["CITATION", citation.items[0].item_id]
    for item in citation.items[1:]:
        parts.extend((r"\m", item.item_id))
    return f" {' '.join(parts)} "


@dataclass(frozen=True, slots=True)
class _CaptionBindingCandidate:
    index: int
    caption: SemanticCaption
    valid: bool
    invalid_reported: bool = False


@dataclass(frozen=True, slots=True)
class _ObjectBindingCandidate:
    index: int
    valid: bool
    invalid_reported: bool = False


def _proven_target_table_ids(
    body_elements: list[Any],
    *,
    bookmark_inventory: DocxBookmarkInventory,
) -> frozenset[str]:
    """Return targets with exactly one valid physical table binding."""

    from docx.oxml.ns import qn

    candidates: dict[str, list[bool]] = {}
    for element in body_elements:
        if element.tag != qn("w:tbl"):
            continue
        target_markers = inspect_object_target_markers(
            element,
            bookmark_inventory=bookmark_inventory,
        )
        pairing = inspect_object_pairing_markers(
            element,
            bookmark_inventory=bookmark_inventory,
        )
        if not target_markers.present or len(target_markers.target_ids) != 1:
            continue
        target_id = target_markers.target_ids[0]
        candidates.setdefault(target_id, []).append(not target_markers.malformed and not pairing.present)
    return frozenset(target_id for target_id, bindings in candidates.items() if len(bindings) == 1 and bindings[0])


def _proven_pairing_table_tokens(
    body_elements: list[Any],
    *,
    bookmark_inventory: DocxBookmarkInventory,
) -> frozenset[str]:
    """Return pairing tokens with exactly one valid physical table binding."""

    from docx.oxml.ns import qn

    candidates: dict[str, list[bool]] = {}
    for element in body_elements:
        if element.tag != qn("w:tbl"):
            continue
        target_markers = inspect_object_target_markers(
            element,
            bookmark_inventory=bookmark_inventory,
        )
        pairing = inspect_object_pairing_markers(
            element,
            bookmark_inventory=bookmark_inventory,
        )
        if not pairing.present or len(pairing.tokens) != 1:
            continue
        token = pairing.tokens[0]
        candidates.setdefault(token, []).append(not pairing.malformed and not target_markers.present)
    return frozenset(token for token, bindings in candidates.items() if len(bindings) == 1 and bindings[0])


def _proven_caption_table_kind(
    paragraph_element: Any,
    *,
    bookmark_inventory: DocxBookmarkInventory,
    proven_target_table_ids: frozenset[str],
    proven_pairing_table_tokens: frozenset[str],
) -> str | None:
    """Infer ``table`` only from matching, conflict-free caption/object proof."""

    target_markers = inspect_caption_target_markers(
        paragraph_element,
        bookmark_inventory=bookmark_inventory,
    )
    pairing = inspect_caption_pairing_markers(
        paragraph_element,
        bookmark_inventory=bookmark_inventory,
    )
    target_bound = (
        not target_markers.malformed
        and len(target_markers.target_ids) == 1
        and not pairing.present
        and target_markers.target_ids[0] in proven_target_table_ids
    )
    pair_bound = (
        not pairing.malformed
        and len(pairing.tokens) == 1
        and not target_markers.present
        and pairing.tokens[0] in proven_pairing_table_tokens
    )
    return "table" if target_bound or pair_bound else None


class DocxSemanticImporter:
    """Import only physically proven DOCX semantics into the neutral model."""

    def import_document(self, document: Any) -> SemanticImportResult:
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph

        body_elements = list(document.element.body)
        paragraphs = {
            index: Paragraph(element, document)
            for index, element in enumerate(body_elements)
            if element.tag == qn("w:p")
        }
        diagnostics: list[SemanticDiagnostic] = []
        try:
            bookmark_inventory = build_docx_bookmark_inventory(document)
        except ValueError:
            diagnostic = SemanticDiagnostic(
                "error",
                "semantic.docx.bookmark.inventory_invalid",
                "A reachable WordprocessingML XML part could not be safely inventoried.",
                "document",
            )
            return _import_without_proven_semantics(
                body_elements,
                paragraphs,
                diagnostic=diagnostic,
            )
        bibliography, bibliography_indices = _extract_bibliography(
            body_elements,
            paragraphs,
            diagnostics,
            bookmark_inventory=bookmark_inventory,
        )
        proven_target_table_ids = _proven_target_table_ids(
            body_elements,
            bookmark_inventory=bookmark_inventory,
        )
        proven_pairing_table_tokens = _proven_pairing_table_tokens(
            body_elements,
            bookmark_inventory=bookmark_inventory,
        )

        target_captions: dict[str, list[_CaptionBindingCandidate]] = {}
        paired_captions: dict[str, list[_CaptionBindingCandidate]] = {}
        for index, paragraph in paragraphs.items():
            if index in bibliography_indices:
                continue
            extracted = extract_neutral_semantic_caption(
                paragraph._p,
                bookmark_inventory=bookmark_inventory,
                proven_object_kind=_proven_caption_table_kind(
                    paragraph._p,
                    bookmark_inventory=bookmark_inventory,
                    proven_target_table_ids=proven_target_table_ids,
                    proven_pairing_table_tokens=proven_pairing_table_tokens,
                ),
            )
            if extracted is None:
                continue
            caption = SemanticCaption(
                kind=extracted.kind,  # type: ignore[arg-type]
                target_id=extracted.target_id,
                cached_number=extracted.cached_number,
                label=extracted.label,
                content=extracted.content,
            )
            conflict = extracted.target_marker_present and extracted.pairing_marker_present
            invalid_reported = False
            if conflict:
                diagnostics.append(
                    SemanticDiagnostic(
                        "error",
                        "semantic.docx.binding.caption_marker_conflict",
                        "A target-bound caption may not also contain an internal pairing marker.",
                        f"document/body[{index}]",
                    )
                )
                invalid_reported = True
            if extracted.shorthand_marker_malformed:
                diagnostics.append(
                    SemanticDiagnostic(
                        "error",
                        "semantic.docx.binding.shorthand_marker_invalid",
                        "Caption shorthand provenance must be one balanced bookmark matching its target.",
                        f"document/body[{index}]",
                    )
                )
                invalid_reported = True

            if extracted.target_marker_present:
                if len(extracted.target_ids) != 1:
                    diagnostics.append(
                        SemanticDiagnostic(
                            "error",
                            "semantic.docx.binding.caption_target_marker_invalid",
                            "A target-bound caption must contain exactly one proven target bookmark.",
                            f"document/body[{index}]",
                        )
                    )
                else:
                    target_id = extracted.target_ids[0]
                    target_caption = SemanticCaption(
                        kind=caption.kind,
                        target_id=target_id,
                        cached_number=caption.cached_number,
                        label=caption.label,
                        content=caption.content,
                    )
                    target_captions.setdefault(target_id, []).append(
                        _CaptionBindingCandidate(
                            index=index,
                            caption=target_caption,
                            valid=(
                                not extracted.target_marker_malformed
                                and not conflict
                                and not extracted.shorthand_marker_malformed
                            ),
                            invalid_reported=invalid_reported,
                        )
                    )

            if extracted.pairing_marker_present:
                if len(extracted.pairing_tokens) != 1:
                    diagnostics.append(
                        SemanticDiagnostic(
                            "error",
                            "semantic.docx.binding.caption_marker_invalid",
                            "A targetless caption must contain exactly one proven internal pairing marker.",
                            f"document/body[{index}]",
                        )
                    )
                else:
                    token = extracted.pairing_tokens[0]
                    paired_captions.setdefault(token, []).append(
                        _CaptionBindingCandidate(
                            index=index,
                            caption=SemanticCaption(
                                kind=caption.kind,
                                target_id=None,
                                cached_number=caption.cached_number,
                                label=caption.label,
                                content=caption.content,
                            ),
                            valid=(
                                not extracted.pairing_marker_malformed
                                and not conflict
                                and not extracted.shorthand_marker_malformed
                            ),
                            invalid_reported=invalid_reported,
                        )
                    )

        table_elements: dict[int, Any] = {}
        target_tables: dict[str, list[_ObjectBindingCandidate]] = {}
        paired_tables: dict[str, list[_ObjectBindingCandidate]] = {}
        for index, element in enumerate(body_elements):
            if element.tag != qn("w:tbl"):
                continue
            table_elements[index] = element
            target_markers = inspect_object_target_markers(
                element,
                bookmark_inventory=bookmark_inventory,
            )
            pairing = inspect_object_pairing_markers(
                element,
                bookmark_inventory=bookmark_inventory,
            )
            conflict = target_markers.present and pairing.present
            invalid_reported = False
            if conflict:
                diagnostics.append(
                    SemanticDiagnostic(
                        "error",
                        "semantic.docx.binding.object_marker_conflict",
                        "A target-bound table may not also contain an internal pairing marker.",
                        f"document/body[{index}]",
                    )
                )
                invalid_reported = True
            if target_markers.present:
                if len(target_markers.target_ids) != 1:
                    diagnostics.append(
                        SemanticDiagnostic(
                            "error",
                            "semantic.docx.binding.object_target_marker_invalid",
                            "A target-bound table must contain exactly one proven object bookmark.",
                            f"document/body[{index}]",
                        )
                    )
                else:
                    target_id = target_markers.target_ids[0]
                    target_tables.setdefault(target_id, []).append(
                        _ObjectBindingCandidate(
                            index=index,
                            valid=not target_markers.malformed and not conflict,
                            invalid_reported=invalid_reported,
                        )
                    )
            if pairing.present:
                if len(pairing.tokens) != 1:
                    diagnostics.append(
                        SemanticDiagnostic(
                            "error",
                            "semantic.docx.binding.object_marker_invalid",
                            "A targetless table must contain exactly one proven internal pairing marker.",
                            f"document/body[{index}]",
                        )
                    )
                else:
                    token = pairing.tokens[0]
                    paired_tables.setdefault(token, []).append(
                        _ObjectBindingCandidate(
                            index=index,
                            valid=not pairing.malformed and not conflict,
                            invalid_reported=invalid_reported,
                        )
                    )

        caption_indices: set[int] = set()
        caption_by_table_index: dict[int, SemanticCaption] = {}
        _resolve_target_bindings(
            target_captions,
            target_tables,
            caption_indices=caption_indices,
            caption_by_table_index=caption_by_table_index,
            diagnostics=diagnostics,
        )
        _resolve_pair_bindings(
            paired_captions,
            paired_tables,
            caption_indices=caption_indices,
            caption_by_table_index=caption_by_table_index,
            diagnostics=diagnostics,
        )
        imported_paragraph_indices = set(paragraphs) - bibliography_indices - caption_indices
        _audit_package_citation_markers(
            body_elements,
            diagnostics,
            bookmark_inventory=bookmark_inventory,
            imported_paragraph_indices=imported_paragraph_indices,
        )

        table_by_index: dict[int, SemanticTable] = {}
        for index, element in table_elements.items():
            try:
                table = _import_table(element, caption=caption_by_table_index.get(index))
            except DocxTableGeometryError as exc:
                diagnostics.append(_table_geometry_diagnostic(exc, index=index))
                continue
            table_by_index[index] = table

        blocks: list[SemanticParagraph | SemanticTable] = []
        for index, _element in enumerate(body_elements):
            if index in bibliography_indices or index in caption_indices:
                continue
            if index in table_by_index:
                blocks.append(table_by_index[index])
                continue
            paragraph = paragraphs.get(index)
            if paragraph is None:
                continue
            blocks.append(
                _import_paragraph(
                    paragraph._p,
                    bookmark_inventory=bookmark_inventory,
                    diagnostics=diagnostics,
                    location=f"document/body[{index}]",
                )
            )

        semantic_document = SemanticDocument(blocks=tuple(blocks), bibliography=bibliography)
        diagnostics.extend(validate_semantic_document(semantic_document))
        return SemanticImportResult(document=semantic_document, diagnostics=tuple(diagnostics))


def _resolve_target_bindings(
    captions: dict[str, list[_CaptionBindingCandidate]],
    objects: dict[str, list[_ObjectBindingCandidate]],
    *,
    caption_indices: set[int],
    caption_by_table_index: dict[int, SemanticCaption],
    diagnostics: list[SemanticDiagnostic],
) -> None:
    for target_id in sorted(set(captions) | set(objects)):
        caption_candidates = captions.get(target_id, [])
        object_candidates = objects.get(target_id, [])
        if len(caption_candidates) > 1:
            diagnostics.append(
                SemanticDiagnostic(
                    "error",
                    "semantic.docx.caption.target_duplicate",
                    f"DOCX contains duplicate caption target {target_id}.",
                    f"document/body[{caption_candidates[1].index}]",
                )
            )
        if len(object_candidates) > 1:
            diagnostics.append(
                SemanticDiagnostic(
                    "error",
                    "semantic.docx.binding.object_target_duplicate",
                    f"DOCX contains duplicate object target {target_id}.",
                    f"document/body[{object_candidates[1].index}]",
                )
            )
        if len(caption_candidates) == 1:
            candidate = caption_candidates[0]
            if not candidate.valid and not candidate.invalid_reported:
                diagnostics.append(
                    SemanticDiagnostic(
                        "error",
                        "semantic.docx.binding.caption_target_marker_invalid",
                        "Caption target bookmark is not globally unique and balanced.",
                        f"document/body[{candidate.index}]",
                    )
                )
        if len(object_candidates) == 1:
            candidate = object_candidates[0]
            if not candidate.valid and not candidate.invalid_reported:
                diagnostics.append(
                    SemanticDiagnostic(
                        "error",
                        "semantic.docx.binding.object_target_marker_invalid",
                        "Object target bookmark is not globally unique and balanced.",
                        f"document/body[{candidate.index}]",
                    )
                )
        if len(caption_candidates) != 1 or len(object_candidates) != 1:
            if len(caption_candidates) == 1 and caption_candidates[0].valid and not object_candidates:
                diagnostics.append(
                    SemanticDiagnostic(
                        "error",
                        "semantic.docx.caption.object_missing",
                        f"Caption target {target_id} is not bound to a DOCX object.",
                        f"document/body[{caption_candidates[0].index}]",
                    )
                )
            if len(object_candidates) == 1 and object_candidates[0].valid and not caption_candidates:
                diagnostics.append(
                    SemanticDiagnostic(
                        "error",
                        "semantic.docx.table.caption_missing",
                        f"Table target {target_id} has no proven caption.",
                        f"document/body[{object_candidates[0].index}]",
                    )
                )
            continue
        caption_candidate = caption_candidates[0]
        object_candidate = object_candidates[0]
        if not caption_candidate.valid or not object_candidate.valid:
            continue
        caption_indices.add(caption_candidate.index)
        caption_by_table_index[object_candidate.index] = caption_candidate.caption


def _resolve_pair_bindings(
    captions: dict[str, list[_CaptionBindingCandidate]],
    objects: dict[str, list[_ObjectBindingCandidate]],
    *,
    caption_indices: set[int],
    caption_by_table_index: dict[int, SemanticCaption],
    diagnostics: list[SemanticDiagnostic],
) -> None:
    for token in sorted(set(captions) | set(objects)):
        caption_candidates = captions.get(token, [])
        object_candidates = objects.get(token, [])
        if len(caption_candidates) > 1:
            diagnostics.append(
                SemanticDiagnostic(
                    "error",
                    "semantic.docx.binding.caption_duplicate",
                    f"DOCX contains duplicate caption pairing token {token}.",
                    f"document/body[{caption_candidates[1].index}]",
                )
            )
        if len(object_candidates) > 1:
            diagnostics.append(
                SemanticDiagnostic(
                    "error",
                    "semantic.docx.binding.object_duplicate",
                    f"DOCX contains duplicate object pairing token {token}.",
                    f"document/body[{object_candidates[1].index}]",
                )
            )
        if len(caption_candidates) == 1:
            candidate = caption_candidates[0]
            if not candidate.valid and not candidate.invalid_reported:
                diagnostics.append(
                    SemanticDiagnostic(
                        "error",
                        "semantic.docx.binding.caption_marker_invalid",
                        "Caption pairing bookmark is not globally unique and balanced.",
                        f"document/body[{candidate.index}]",
                    )
                )
        if len(object_candidates) == 1:
            candidate = object_candidates[0]
            if not candidate.valid and not candidate.invalid_reported:
                diagnostics.append(
                    SemanticDiagnostic(
                        "error",
                        "semantic.docx.binding.object_marker_invalid",
                        "Object pairing bookmark is not globally unique and balanced.",
                        f"document/body[{candidate.index}]",
                    )
                )
        if len(caption_candidates) != 1 or len(object_candidates) != 1:
            if len(caption_candidates) == 1 and caption_candidates[0].valid and not object_candidates:
                diagnostics.append(
                    SemanticDiagnostic(
                        "error",
                        "semantic.docx.binding.object_missing",
                        f"Caption pairing token {token} has no DOCX object marker.",
                        f"document/body[{caption_candidates[0].index}]",
                    )
                )
            if len(object_candidates) == 1 and object_candidates[0].valid and not caption_candidates:
                diagnostics.append(
                    SemanticDiagnostic(
                        "error",
                        "semantic.docx.binding.caption_missing",
                        f"Object pairing token {token} has no proven caption marker.",
                        f"document/body[{object_candidates[0].index}]",
                    )
                )
            continue
        caption_candidate = caption_candidates[0]
        object_candidate = object_candidates[0]
        if not caption_candidate.valid or not object_candidate.valid:
            continue
        caption_indices.add(caption_candidate.index)
        caption_by_table_index[object_candidate.index] = caption_candidate.caption


def _import_without_proven_semantics(
    body_elements: list[Any],
    paragraphs: dict[int, Any],
    *,
    diagnostic: SemanticDiagnostic,
) -> SemanticImportResult:
    from docx.oxml.ns import qn

    diagnostics = [diagnostic]
    blocks: list[SemanticParagraph | SemanticTable] = []
    for index, element in enumerate(body_elements):
        if element.tag == qn("w:tbl"):
            try:
                blocks.append(_import_table(element, caption=None))
            except DocxTableGeometryError as exc:
                diagnostics.append(_table_geometry_diagnostic(exc, index=index))
        elif index in paragraphs:
            blocks.append(
                _import_paragraph(
                    paragraphs[index]._p,
                    bookmark_inventory=None,
                    diagnostics=None,
                    location=f"document/body[{index}]",
                )
            )
    return SemanticImportResult(
        document=SemanticDocument(blocks=tuple(blocks)),
        diagnostics=tuple(diagnostics),
    )


def append_zero_width_bookmark(paragraph: Any, bookmark_name: str, bookmark_id: str) -> None:
    """Append one standard zero-width bookmark."""

    append_bookmark_start(paragraph, bookmark_name, bookmark_id)
    append_bookmark_end(paragraph, bookmark_id)


def append_bookmark_start(paragraph: Any, bookmark_name: str, bookmark_id: str) -> None:
    paragraph._p.append(_bookmark_start_element(bookmark_name, bookmark_id))


def prepend_bookmark_start(paragraph: Any, bookmark_name: str, bookmark_id: str) -> None:
    """Insert a bookmark start before all paragraph content but after ``w:pPr``."""

    insert_at = 1 if paragraph._p.pPr is not None else 0
    paragraph._p.insert(insert_at, _bookmark_start_element(bookmark_name, bookmark_id))


def _bookmark_start_element(bookmark_name: str, bookmark_id: str) -> Any:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), bookmark_id)
    bookmark_start.set(qn("w:name"), bookmark_name)
    return bookmark_start


def append_bookmark_end(paragraph: Any, bookmark_id: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), bookmark_id)
    paragraph._p.append(bookmark_end)


def append_complex_field(
    paragraph: Any,
    *,
    instruction: str,
    cached_result: str,
    bookmark_name: str | None = None,
    bookmark_id: str | None = None,
    locked: bool = False,
    dirty: bool = True,
) -> None:
    """Append a Word complex field with a visible cached result."""

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if bookmark_name is not None:
        if bookmark_id is None:
            raise ValueError("bookmark_id is required when bookmark_name is provided")
        append_bookmark_start(paragraph, bookmark_name, bookmark_id)

    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    if locked:
        begin.set(qn("w:fldLock"), "true")
    if dirty:
        begin.set(qn("w:dirty"), "true")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instruction_text.text = instruction
    instruction_run._r.append(instruction_text)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)
    paragraph.add_run(cached_result)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)
    if bookmark_name is not None and bookmark_id is not None:
        append_bookmark_end(paragraph, bookmark_id)


def apply_semantic_table_roles(
    table: Any,
    *,
    header_rows: int,
    header_columns: int,
    repeat_header: str,
) -> None:
    """Encode explicit header roles and repeat policy with standard OOXML."""

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table_properties = table._tbl.tblPr
    if table_properties is None:
        table_properties = OxmlElement("w:tblPr")
        table._tbl.insert(0, table_properties)
    for existing in table_properties.findall(qn("w:tblLook")):
        table_properties.remove(existing)
    table_look = OxmlElement("w:tblLook")
    table_look.set(qn("w:firstRow"), "1" if header_rows else "0")
    table_look.set(qn("w:lastRow"), "0")
    table_look.set(qn("w:firstColumn"), "1" if header_columns else "0")
    table_look.set(qn("w:lastColumn"), "0")
    table_look.set(qn("w:noHBand"), "0")
    table_look.set(qn("w:noVBand"), "1")
    table_properties.append(table_look)

    for row_index, row in enumerate(table.rows):
        row_properties = row._tr.get_or_add_trPr()
        for existing in row_properties.findall(qn("w:cnfStyle")):
            row_properties.remove(existing)
        for existing in row_properties.findall(qn("w:tblHeader")):
            row_properties.remove(existing)
        if row_index < header_rows:
            conditional = OxmlElement("w:cnfStyle")
            conditional.set(qn("w:firstRow"), "1")
            row_properties.append(conditional)
            if repeat_header in {"always", "never"}:
                repeat = OxmlElement("w:tblHeader")
                repeat.set(qn("w:val"), "1" if repeat_header == "always" else "0")
                row_properties.append(repeat)

        seen_cells: set[int] = set()
        for column_index, cell in enumerate(row.cells):
            cell_key = id(cell._tc)
            if cell_key in seen_cells:
                continue
            seen_cells.add(cell_key)
            cell_properties = cell._tc.get_or_add_tcPr()
            for existing in cell_properties.findall(qn("w:cnfStyle")):
                cell_properties.remove(existing)
            if column_index < header_columns:
                conditional = OxmlElement("w:cnfStyle")
                conditional.set(qn("w:firstColumn"), "1")
                cell_properties.append(conditional)


def _audit_package_citation_markers(
    body_elements: list[Any],
    diagnostics: list[SemanticDiagnostic],
    *,
    bookmark_inventory: DocxBookmarkInventory,
    imported_paragraph_indices: set[int],
) -> None:
    citation_starts = tuple(
        start for start in bookmark_inventory.starts if _is_citation_bookmark_name(start.name or "")
    )
    main_body_name_keys = {
        (start.name or "").casefold()
        for start in citation_starts
        if _containing_direct_body_paragraph_index(start.element, body_elements) in imported_paragraph_indices
    }
    reported_name_keys: set[str] = set()
    for start in citation_starts:
        body_index = _containing_direct_body_paragraph_index(start.element, body_elements)
        if body_index in imported_paragraph_indices:
            continue
        name_key = (start.name or "").casefold()
        if name_key in main_body_name_keys or name_key in reported_name_keys:
            continue
        reported_name_keys.add(name_key)
        if body_index is None:
            diagnostics.append(
                SemanticDiagnostic(
                    "error",
                    "semantic.docx.citation.marker_outside_body",
                    "Citation cluster bookmarks are supported only inside direct main-body paragraphs.",
                    start.part_name,
                )
            )
        else:
            diagnostics.append(
                SemanticDiagnostic(
                    "error",
                    "semantic.docx.citation.marker_unconsumed",
                    "Citation cluster bookmark occurs in a body paragraph owned by another semantic object.",
                    f"document/body[{body_index}]",
                )
            )


def _is_citation_bookmark_name(bookmark_name: str) -> bool:
    return bookmark_name.casefold().startswith(CITATION_BOOKMARK_PREFIX.casefold())


def _containing_direct_body_paragraph_index(marker: Any, body_elements: list[Any]) -> int | None:
    from docx.oxml.ns import qn

    current = marker.getparent()
    while current is not None:
        if current in body_elements:
            return body_elements.index(current) if current.tag == qn("w:p") else None
        current = current.getparent()
    return None


def _extract_bibliography(
    body_elements: list[Any],
    paragraphs: dict[int, Any],
    diagnostics: list[SemanticDiagnostic],
    *,
    bookmark_inventory: DocxBookmarkInventory,
) -> tuple[SemanticBibliographyFragment | None, set[int]]:
    total_starts = bookmark_inventory.starts_named(BIBLIOGRAPHY_BOOKMARK_NAME)
    entry_starts = tuple(
        item
        for item in bookmark_inventory.starts
        if (item.name or "").casefold().startswith(BIBLIOGRAPHY_ENTRY_BOOKMARK_PREFIX.casefold())
    )
    if not total_starts:
        if entry_starts:
            diagnostics.append(
                SemanticDiagnostic(
                    "error",
                    "semantic.docx.bibliography.entry_orphan",
                    "Bibliography entry markers require one proven total boundary.",
                    "document/body",
                )
            )
        return None, set()

    total_proof = prove_bookmark_name(bookmark_inventory, BIBLIOGRAPHY_BOOKMARK_NAME)
    boundary_location = "document/body"
    if total_proof.start is not None:
        boundary_index = _direct_body_paragraph_index(total_proof.start.element, body_elements)
        if boundary_index is not None:
            boundary_location = f"document/body[{boundary_index}]"
    if (
        not total_proof.valid
        or total_proof.start is None
        or total_proof.end is None
        or total_proof.start.name != BIBLIOGRAPHY_BOOKMARK_NAME
    ):
        diagnostics.append(
            SemanticDiagnostic(
                "error",
                "semantic.docx.bibliography.boundary_unbalanced",
                "DOCX bibliography must contain one globally unique balanced boundary.",
                boundary_location,
            )
        )
        return None, set()

    start_index = _direct_body_paragraph_index(total_proof.start.element, body_elements)
    end_index = _direct_body_paragraph_index(total_proof.end.element, body_elements)
    if start_index is None or end_index is None or end_index < start_index:
        diagnostics.append(
            SemanticDiagnostic(
                "error",
                "semantic.docx.bibliography.boundary_invalid",
                "DOCX bibliography boundary must be paragraph-local in the main document body.",
                "document/body",
            )
        )
        return None, set()

    indices = set(range(start_index, end_index + 1))
    if any(index not in paragraphs for index in indices):
        diagnostics.append(
            SemanticDiagnostic(
                "error",
                "semantic.docx.bibliography.boundary_invalid",
                "DOCX bibliography boundary may contain only entry paragraphs.",
                f"document/body[{start_index}]",
            )
        )
        return None, set()

    start_paragraph = paragraphs[start_index]._p
    end_paragraph = paragraphs[end_index]._p
    start_children = list(start_paragraph)
    end_children = list(end_paragraph)
    if (
        total_proof.start.element.getparent() is not start_paragraph
        or total_proof.end.element.getparent() is not end_paragraph
        or total_proof.start.element not in start_children
        or total_proof.end.element not in end_children
        or _first_content_index(start_paragraph) != start_children.index(total_proof.start.element)
        or end_children.index(total_proof.end.element) != len(end_children) - 1
    ):
        diagnostics.append(
            SemanticDiagnostic(
                "error",
                "semantic.docx.bibliography.boundary_invalid",
                "Bibliography total boundary must exactly enclose all entry paragraphs.",
                f"document/body[{start_index}]",
            )
        )
        return None, set()

    entries_by_index: dict[int, SemanticBibliographyEntry] = {}
    failed = False
    decoded_item_ids: set[str] = set()
    for entry_start in entry_starts:
        bookmark_name = entry_start.name or ""
        item_id = decode_bibliography_entry_bookmark(bookmark_name)
        proof = prove_bookmark_name(bookmark_inventory, bookmark_name)
        if item_id is None or not proof.valid or proof.start is None or proof.end is None:
            diagnostics.append(
                SemanticDiagnostic(
                    "error",
                    "semantic.docx.bibliography.entry_marker_invalid",
                    "Bibliography entry bookmark must be canonical, globally unique, and balanced.",
                    "document/body",
                )
            )
            failed = True
            continue
        entry_index = _direct_body_paragraph_index(proof.start.element, body_elements)
        end_entry_index = _direct_body_paragraph_index(proof.end.element, body_elements)
        if (
            entry_index is None
            or end_entry_index != entry_index
            or entry_index not in indices
            or proof.start.element.getparent() is not paragraphs[entry_index]._p
            or proof.end.element.getparent() is not paragraphs[entry_index]._p
        ):
            diagnostics.append(
                SemanticDiagnostic(
                    "error",
                    "semantic.docx.bibliography.entry_outside_boundary",
                    "Bibliography entry bookmark must stay inside one total-boundary paragraph.",
                    "document/body",
                )
            )
            failed = True
            continue
        paragraph_element = paragraphs[entry_index]._p
        paragraph_children = list(paragraph_element)
        entry_start_index = paragraph_children.index(proof.start.element)
        entry_end_index = paragraph_children.index(proof.end.element)
        entry_elements = paragraph_children[entry_start_index + 1 : entry_end_index]
        before_entry = paragraph_children[_first_content_index(paragraph_element) : entry_start_index]
        after_entry = paragraph_children[entry_end_index + 1 :]
        expected_before = [total_proof.start.element] if entry_index == start_index else []
        expected_after = [total_proof.end.element] if entry_index == end_index else []
        runs = _import_bibliography_runs(paragraphs[entry_index], entry_elements)
        if (
            entry_end_index <= entry_start_index + 1
            or runs is None
            or not any(run.text.strip() for run in runs)
            or not _same_elements(before_entry, expected_before)
            or not _same_elements(after_entry, expected_after)
        ):
            diagnostics.append(
                SemanticDiagnostic(
                    "error",
                    "semantic.docx.bibliography.entry_marker_invalid",
                    "Bibliography entry bookmark must exactly enclose one non-empty formatted paragraph.",
                    f"document/body[{entry_index}]",
                )
            )
            failed = True
            continue
        if item_id in decoded_item_ids or entry_index in entries_by_index:
            diagnostics.append(
                SemanticDiagnostic(
                    "error",
                    "semantic.docx.bibliography.entry_duplicate",
                    "Bibliography item markers and entry paragraphs must be one-to-one.",
                    f"document/body[{entry_index}]",
                )
            )
            failed = True
            continue
        decoded_item_ids.add(item_id)
        assert runs is not None
        entries_by_index[entry_index] = SemanticBibliographyEntry(item_id=item_id, runs=runs)

    if set(entries_by_index) != indices:
        diagnostics.append(
            SemanticDiagnostic(
                "error",
                "semantic.docx.bibliography.entries_incomplete",
                "Every paragraph inside the bibliography boundary must have exactly one entry marker.",
                f"document/body[{start_index}]",
            )
        )
        failed = True
    if not failed:
        first_entry = bookmark_inventory.starts_named(
            encode_bibliography_entry_bookmark(entries_by_index[start_index].item_id)
        )[0]
        last_entry_name = encode_bibliography_entry_bookmark(entries_by_index[end_index].item_id)
        last_entry_proof = prove_bookmark_name(bookmark_inventory, last_entry_name)
        if (
            start_children.index(total_proof.start.element) >= start_children.index(first_entry.element)
            or not last_entry_proof.valid
            or last_entry_proof.end is None
            or end_children.index(last_entry_proof.end.element) >= end_children.index(total_proof.end.element)
        ):
            diagnostics.append(
                SemanticDiagnostic(
                    "error",
                    "semantic.docx.bibliography.boundary_invalid",
                    "Bibliography total boundary must be outside every entry marker.",
                    f"document/body[{start_index}]",
                )
            )
            failed = True
    if failed:
        return None, set()
    entries = tuple(entries_by_index[index] for index in range(start_index, end_index + 1))
    return SemanticBibliographyFragment(entries=entries), indices


def _import_bibliography_runs(
    paragraph: Any,
    entry_elements: list[Any],
) -> tuple[SemanticBibliographyRun, ...] | None:
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    imported: list[SemanticBibliographyRun] = []
    for element in entry_elements:
        if element.tag == _qn("w:r"):
            if element.attrib:
                return None
            run = _import_bibliography_run(element, href=None, hyperlink=False)
            if run is None:
                return None
            imported.append(run)
            continue
        if element.tag != _qn("w:hyperlink"):
            return None
        relationship_id = element.get(_qn("r:id"))
        if (
            relationship_id is None
            or set(element.attrib) != {_qn("r:id")}
            or element.get(_qn("w:anchor")) is not None
            or any(child.tag != _qn("w:r") for child in element)
            or not list(element)
        ):
            return None
        try:
            relationship = paragraph.part.rels[relationship_id]
        except KeyError:
            return None
        if relationship.reltype != RT.HYPERLINK or not relationship.is_external:
            return None
        href = relationship.target_ref
        if not _is_absolute_http_url(href):
            return None
        if not _has_style_type(paragraph.part.document, "Hyperlink", "character"):
            return None
        for child in element:
            if child.attrib:
                return None
            run = _import_bibliography_run(child, href=href, hyperlink=True)
            if run is None:
                return None
            imported.append(run)
    return tuple(imported) if imported else None


def _import_bibliography_run(
    run_element: Any,
    *,
    href: str | None,
    hyperlink: bool,
) -> SemanticBibliographyRun | None:
    children = list(run_element)
    position = 0
    bold = False
    italic = False
    if children and children[0].tag == _qn("w:rPr"):
        formatting = _import_bibliography_run_properties(children[0], hyperlink=hyperlink)
        if formatting is None:
            return None
        bold, italic = formatting
        position = 1
    text_parts: list[str] = []
    for child in children[position:]:
        if child.tag == _qn("w:t"):
            if set(child.attrib) - {"{http://www.w3.org/XML/1998/namespace}space"} or list(child):
                return None
            text_parts.append(child.text or "")
        elif child.tag == _qn("w:tab"):
            if child.attrib or list(child):
                return None
            text_parts.append("\t")
        elif child.tag == _qn("w:br"):
            if child.attrib or list(child):
                return None
            text_parts.append("\n")
        elif child.tag == _qn("w:cr"):
            if child.attrib or list(child):
                return None
            text_parts.append("\r")
        else:
            return None
    text = "".join(text_parts)
    if not text:
        return None
    return SemanticBibliographyRun(text=text, bold=bold, italic=italic, href=href)


def _import_bibliography_run_properties(
    run_properties: Any,
    *,
    hyperlink: bool,
) -> tuple[bool, bool] | None:
    bold: bool | None = None
    italic: bool | None = None
    style_seen = False
    seen_tags: set[str] = set()
    if run_properties.attrib:
        return None
    for child in run_properties:
        if child.tag in seen_tags or list(child):
            return None
        seen_tags.add(child.tag)
        if child.tag == _qn("w:rStyle"):
            if (
                not hyperlink
                or style_seen
                or set(child.attrib) != {_qn("w:val")}
                or child.get(_qn("w:val")) != "Hyperlink"
            ):
                return None
            style_seen = True
            continue
        if child.tag == _qn("w:b"):
            if set(child.attrib) - {_qn("w:val")}:
                return None
            value = _word_toggle_value(child)
            if value is not True or bold is not None:
                return None
            bold = value
            continue
        if child.tag == _qn("w:i"):
            if set(child.attrib) - {_qn("w:val")}:
                return None
            value = _word_toggle_value(child)
            if value is not True or italic is not None:
                return None
            italic = value
            continue
        return None
    if hyperlink and not style_seen:
        return None
    return bool(bold), bool(italic)


def _word_toggle_value(element: Any) -> bool | None:
    value = element.get(_qn("w:val"))
    if value is None or value.casefold() in {"1", "true", "on"}:
        return True
    if value.casefold() in {"0", "false", "off"}:
        return False
    return None


def _is_absolute_http_url(value: object) -> bool:
    if not isinstance(value, str) or not value:
        return False
    try:
        parsed = urlsplit(value)
    except ValueError:
        return False
    return parsed.scheme.casefold() in {"http", "https"} and parsed.hostname is not None


def _direct_body_paragraph_index(marker: Any, body_elements: list[Any]) -> int | None:
    from docx.oxml.ns import qn

    paragraph = marker.getparent()
    if paragraph is None or paragraph.tag != qn("w:p"):
        return None
    parent = paragraph.getparent()
    if parent is None or paragraph not in body_elements:
        return None
    return body_elements.index(paragraph)


def _first_content_index(paragraph_element: Any) -> int:
    from docx.oxml.ns import qn

    children = list(paragraph_element)
    return 1 if children and children[0].tag == qn("w:pPr") else 0


def _same_elements(actual: list[Any], expected: list[Any]) -> bool:
    return len(actual) == len(expected) and all(
        actual_element is expected_element for actual_element, expected_element in zip(actual, expected, strict=True)
    )


def _import_table(element: Any, *, caption: SemanticCaption | None) -> SemanticTable:
    from docx.oxml.ns import qn

    grid = build_docx_table_semantic_grid(
        element,
        cell_text_resolver=lambda cell, _row, _column: "".join(text.text or "" for text in cell.iter(qn("w:t"))),
    )
    metadata = extract_semantic_table_metadata(element, default_first_row=False)
    cells: list[SemanticTableCell] = []
    for row in grid:
        for cell in row:
            if cell.is_covered:
                continue
            role = _role_for_position(
                cell.row,
                cell.col,
                header_rows=metadata.header_rows,
                header_columns=metadata.header_columns,
            )
            cells.append(
                SemanticTableCell(
                    row=cell.row,
                    column=cell.col,
                    text=cell.anchor_text,
                    role=role,  # type: ignore[arg-type]
                    row_span=cell.rowspan,
                    column_span=cell.colspan,
                )
            )
    return SemanticTable(
        row_count=len(grid),
        column_count=len(grid[0]),
        cells=tuple(cells),
        repeat_header=metadata.repeat_header,  # type: ignore[arg-type]
        caption=caption,
    )


def _table_geometry_diagnostic(error: DocxTableGeometryError, *, index: int) -> SemanticDiagnostic:
    return SemanticDiagnostic(
        "error",
        f"semantic.docx.table.{error.code}",
        str(error),
        f"document/body[{index}]",
    )


def _import_paragraph(
    paragraph_element: Any,
    *,
    bookmark_inventory: DocxBookmarkInventory | None,
    diagnostics: list[SemanticDiagnostic] | None,
    location: str,
) -> SemanticParagraph:
    inlines: list[SemanticText | SemanticReference | SemanticCitationCluster] = []
    children = list(paragraph_element)
    if bookmark_inventory is None:
        _import_generic_elements(
            children,
            inlines,
            bookmark_inventory=bookmark_inventory,
        )
        return SemanticParagraph(inlines=tuple(inlines))

    citation_ranges, marker_failed = _citation_ranges(
        paragraph_element,
        bookmark_inventory=bookmark_inventory,
        diagnostics=diagnostics,
        location=location,
    )
    if marker_failed:
        _append_text(inlines, _visible_text(children))
        return SemanticParagraph(inlines=tuple(inlines))

    covered_instruction_indices = {
        index
        for citation_range in citation_ranges
        for index in range(citation_range.start_index + 1, citation_range.end_index)
    }
    orphan_instruction_indices = _citation_instruction_child_indices(paragraph_element) - covered_instruction_indices
    if orphan_instruction_indices and diagnostics is not None:
        diagnostics.append(
            SemanticDiagnostic(
                "error",
                "semantic.docx.citation.marker_missing",
                "A DOCX CITATION field is not enclosed by one proven cluster bookmark.",
                location,
            )
        )
    if orphan_instruction_indices:
        _append_text(inlines, _visible_text(children))
        return SemanticParagraph(inlines=tuple(inlines))

    cursor = 0
    for citation_range in citation_ranges:
        _import_generic_elements(
            children[cursor : citation_range.start_index],
            inlines,
            bookmark_inventory=bookmark_inventory,
        )
        field_elements = children[citation_range.start_index + 1 : citation_range.end_index]
        evidence = _parse_exact_complex_field(field_elements)
        citation = _citation_from_evidence(
            citation_range.cluster_id,
            evidence,
            diagnostics=diagnostics,
            location=location,
        )
        if citation is None:
            _append_text(inlines, _visible_text(field_elements))
        else:
            inlines.append(citation)
        cursor = citation_range.end_index + 1
    _import_generic_elements(
        children[cursor:],
        inlines,
        bookmark_inventory=bookmark_inventory,
    )
    return SemanticParagraph(inlines=tuple(inlines))


@dataclass(frozen=True, slots=True)
class _CitationRange:
    cluster_id: str
    start_index: int
    end_index: int


@dataclass(frozen=True, slots=True)
class _ComplexFieldEvidence:
    structure_valid: bool
    instruction: str = ""
    cached_result: str = ""
    locked: bool = False
    dirty: bool = True


def _citation_ranges(
    paragraph_element: Any,
    *,
    bookmark_inventory: DocxBookmarkInventory,
    diagnostics: list[SemanticDiagnostic] | None,
    location: str,
) -> tuple[list[_CitationRange], bool]:
    from docx.oxml.ns import qn

    children = list(paragraph_element)
    child_index = {id(element): index for index, element in enumerate(children)}
    marker_starts = [
        element
        for element in paragraph_element.iter(qn("w:bookmarkStart"))
        if _is_citation_bookmark_name(element.get(qn("w:name")) or "")
    ]
    ranges: list[_CitationRange] = []
    failed = False
    for marker_start in marker_starts:
        bookmark_name = marker_start.get(qn("w:name")) or ""
        cluster_id = decode_citation_bookmark(bookmark_name)
        proof = prove_bookmark_name(
            bookmark_inventory,
            bookmark_name,
            scope_element=paragraph_element,
        )
        if (
            cluster_id is None
            or not proof.valid
            or proof.start is None
            or proof.end is None
            or proof.start.element.getparent() is not paragraph_element
            or proof.end.element.getparent() is not paragraph_element
            or id(proof.start.element) not in child_index
            or id(proof.end.element) not in child_index
        ):
            failed = True
            if diagnostics is not None:
                diagnostics.append(
                    SemanticDiagnostic(
                        "error",
                        "semantic.docx.citation.marker_invalid",
                        "Citation cluster bookmark must be canonical, globally unique, balanced, and paragraph-local.",
                        location,
                    )
                )
            continue
        start_index = child_index[id(proof.start.element)]
        end_index = child_index[id(proof.end.element)]
        if end_index <= start_index + 1:
            failed = True
            if diagnostics is not None:
                diagnostics.append(
                    SemanticDiagnostic(
                        "error",
                        "semantic.docx.citation.marker_invalid",
                        "Citation cluster bookmark must contain exactly one physical field range.",
                        location,
                    )
                )
            continue
        ranges.append(
            _CitationRange(
                cluster_id=cluster_id,
                start_index=start_index,
                end_index=end_index,
            )
        )

    ranges.sort(key=lambda item: item.start_index)
    for previous, current in pairwise(ranges):
        if current.start_index <= previous.end_index:
            failed = True
            if diagnostics is not None:
                diagnostics.append(
                    SemanticDiagnostic(
                        "error",
                        "semantic.docx.citation.marker_invalid",
                        "Citation cluster bookmark ranges must not overlap or nest.",
                        location,
                    )
                )
            break
    return ranges, failed


def _citation_instruction_child_indices(paragraph_element: Any) -> set[int]:
    from docx.oxml.ns import qn

    indices: set[int] = set()
    field_indices: set[int] = set()
    instruction_parts: list[str] = []
    in_field = False
    for index, child in enumerate(paragraph_element):
        for element in child.iter():
            if element.tag == qn("w:fldSimple"):
                instruction = element.get(qn("w:instr")) or ""
                if re.match(r"\s*CITATION\b", instruction, re.IGNORECASE):
                    indices.add(index)
                continue
            if element.tag == qn("w:fldChar"):
                field_type = element.get(qn("w:fldCharType")) or ""
                if field_type == "begin":
                    if in_field and re.match(
                        r"\s*CITATION\b",
                        "".join(instruction_parts),
                        re.IGNORECASE,
                    ):
                        indices.update(field_indices)
                    in_field = True
                    field_indices = {index}
                    instruction_parts = []
                elif in_field:
                    field_indices.add(index)
                    if field_type == "end":
                        if re.match(
                            r"\s*CITATION\b",
                            "".join(instruction_parts),
                            re.IGNORECASE,
                        ):
                            indices.update(field_indices)
                        in_field = False
                continue
            if element.tag == qn("w:instrText") and in_field:
                field_indices.add(index)
                instruction_parts.append(element.text or "")
            elif element.tag == qn("w:instrText") and re.match(
                r"\s*CITATION\b",
                element.text or "",
                re.IGNORECASE,
            ):
                indices.add(index)
    if in_field and re.match(r"\s*CITATION\b", "".join(instruction_parts), re.IGNORECASE):
        indices.update(field_indices)
    return indices


def _parse_exact_complex_field(elements: list[Any]) -> _ComplexFieldEvidence:
    from docx.oxml.ns import qn

    state = "before"
    instruction_parts: list[str] = []
    cached_parts: list[str] = []
    begin = None
    dirty = False
    for element in elements:
        if element.tag != qn("w:r"):
            return _ComplexFieldEvidence(False)
        semantic_child_seen = False
        for child in element:
            if child.tag == qn("w:rPr"):
                continue
            semantic_child_seen = True
            if child.tag == qn("w:fldChar"):
                field_type = child.get(qn("w:fldCharType")) or ""
                dirty_value = child.get(qn("w:dirty"))
                if dirty_value not in {None, "0", "false", "off"}:
                    dirty = True
                if field_type == "begin" and state == "before":
                    begin = child
                    state = "instruction"
                elif field_type == "separate" and state == "instruction":
                    state = "result"
                elif field_type == "end" and state == "result":
                    state = "done"
                else:
                    return _ComplexFieldEvidence(False)
            elif child.tag == qn("w:instrText") and state == "instruction":
                instruction_parts.append(child.text or "")
            elif child.tag == qn("w:t") and state == "result":
                cached_parts.append(child.text or "")
            elif child.tag == qn("w:tab") and state == "result":
                cached_parts.append("\t")
            elif child.tag in {qn("w:br"), qn("w:cr")} and state == "result":
                cached_parts.append("\n")
            else:
                return _ComplexFieldEvidence(False)
        if not semantic_child_seen:
            return _ComplexFieldEvidence(False)
    if state != "done" or begin is None or not instruction_parts:
        return _ComplexFieldEvidence(False)
    locked_value = begin.get(qn("w:fldLock"))
    return _ComplexFieldEvidence(
        structure_valid=True,
        instruction="".join(instruction_parts),
        cached_result="".join(cached_parts),
        locked=locked_value in {"1", "true", "on"},
        dirty=dirty,
    )


def _citation_from_evidence(
    cluster_id: str,
    evidence: _ComplexFieldEvidence,
    *,
    diagnostics: list[SemanticDiagnostic] | None,
    location: str,
) -> SemanticCitationCluster | None:
    code = ""
    message = ""
    items: tuple[SemanticCitationItem, ...] | None = None
    if not evidence.structure_valid:
        code = "semantic.docx.citation.field_invalid"
        message = "Citation bookmark must contain exactly one complete complex field."
    elif not evidence.locked:
        code = "semantic.docx.citation.field_unlocked"
        message = "Citation complex fields must be locked."
    elif evidence.dirty:
        code = "semantic.docx.citation.field_dirty"
        message = "Citation complex fields must not be marked dirty."
    elif (items := _parse_citation_instruction(evidence.instruction)) is None:
        code = "semantic.docx.citation.instruction_invalid"
        message = "Citation field instruction must contain only ordered portable item IDs."
    elif not evidence.cached_result.strip():
        code = "semantic.docx.citation.cached_result_empty"
        message = "Citation field cached result must not be empty."
    if code:
        if diagnostics is not None:
            diagnostics.append(SemanticDiagnostic("error", code, message, location))
        return None
    assert items is not None
    return SemanticCitationCluster(
        cluster_id=cluster_id,
        items=items,
        cached_result=evidence.cached_result,
    )


def _parse_citation_instruction(instruction: str) -> tuple[SemanticCitationItem, ...] | None:
    tokens = instruction.split()
    if len(tokens) < 2 or tokens[0].upper() != "CITATION":
        return None
    item_ids = [tokens[1]]
    position = 2
    while position < len(tokens):
        if position + 1 >= len(tokens) or tokens[position].lower() != r"\m":
            return None
        item_ids.append(tokens[position + 1])
        position += 2
    if any(not is_portable_semantic_id(item_id) for item_id in item_ids):
        return None
    return tuple(SemanticCitationItem(item_id) for item_id in item_ids)


def _visible_text(elements: list[Any]) -> str:
    from docx.oxml.ns import qn

    output: list[str] = []
    for element in elements:
        for item in element.iter():
            if item.tag == qn("w:t"):
                output.append(item.text or "")
            elif item.tag == qn("w:tab"):
                output.append("\t")
            elif item.tag in {qn("w:br"), qn("w:cr")}:
                output.append("\n")
    return "".join(output)


def _import_generic_elements(
    roots: list[Any],
    inlines: list[SemanticText | SemanticReference | SemanticCitationCluster],
    *,
    bookmark_inventory: DocxBookmarkInventory | None,
) -> None:
    from docx.oxml.ns import qn

    instruction_parts: list[str] = []
    result_parts: list[str] = []
    in_field = False
    in_result = False
    for root in roots:
        for element in root.iter():
            if element.tag == qn("w:fldSimple"):
                instruction = element.get(qn("w:instr")) or ""
                result = _visible_text([element])
                _append_imported_field(
                    inlines,
                    instruction,
                    result,
                    bookmark_inventory=bookmark_inventory,
                )
                continue
            if element.tag == qn("w:fldChar"):
                field_type = element.get(qn("w:fldCharType")) or ""
                if field_type == "begin":
                    in_field = True
                    in_result = False
                    instruction_parts = []
                    result_parts = []
                elif in_field and field_type == "separate":
                    in_result = True
                elif in_field and field_type == "end":
                    _append_imported_field(
                        inlines,
                        "".join(instruction_parts),
                        "".join(result_parts),
                        bookmark_inventory=bookmark_inventory,
                    )
                    in_field = False
                    in_result = False
                continue
            if element.tag == qn("w:instrText") and in_field:
                instruction_parts.append(element.text or "")
                continue
            visible_value = None
            if element.tag == qn("w:t"):
                visible_value = element.text or ""
            elif element.tag == qn("w:tab"):
                visible_value = "\t"
            elif element.tag in {qn("w:br"), qn("w:cr")}:
                visible_value = "\n"
            if visible_value is None:
                continue
            if in_field:
                if in_result:
                    result_parts.append(visible_value)
            elif not _has_ancestor_tag(element, qn("w:fldSimple")):
                _append_text(inlines, visible_value)


def _append_imported_field(
    inlines: list[SemanticText | SemanticReference | SemanticCitationCluster],
    instruction: str,
    cached_result: str,
    *,
    bookmark_inventory: DocxBookmarkInventory | None,
) -> None:
    match = re.match(r"\s*REF\s+(\S+)", instruction, re.IGNORECASE)
    bookmark_name = match.group(1) if match is not None else ""
    target_id = decode_target_bookmark(bookmark_name)
    if (
        target_id is not None
        and bookmark_inventory is not None
        and prove_bookmark_name(bookmark_inventory, bookmark_name).valid
    ):
        inlines.append(SemanticReference(target_id=target_id, cached_result=cached_result))
    else:
        _append_text(inlines, cached_result)


def _append_text(
    inlines: list[SemanticText | SemanticReference | SemanticCitationCluster],
    value: str,
) -> None:
    if not value:
        return
    if inlines and isinstance(inlines[-1], SemanticText):
        previous = inlines[-1]
        inlines[-1] = SemanticText(previous.value + value)
    else:
        inlines.append(SemanticText(value))


def _has_ancestor_tag(element: Any, tag: str) -> bool:
    parent = element.getparent()
    while parent is not None:
        if parent.tag == tag:
            return True
        parent = parent.getparent()
    return False


def _role_for_position(row: int, column: int, *, header_rows: int, header_columns: int) -> str:
    if row < header_rows and column < header_columns:
        return "corner_header"
    if row < header_rows:
        return "column_header"
    if column < header_columns:
        return "row_header"
    return "data"


__all__ = [
    "DocxSemanticImporter",
    "DocxSemanticRenderer",
    "append_bookmark_end",
    "append_bookmark_start",
    "append_complex_field",
    "append_zero_width_bookmark",
    "apply_semantic_table_roles",
]
