"""Focused tests split from test_docx_citations.py."""

from __future__ import annotations

from ._docx_citations_support import (
    BIBLIOGRAPHY_BOOKMARK_NAME,
    RT,
    WD_STYLE_TYPE,
    Document,
    DocxSemanticImporter,
    DocxSemanticRenderer,
    SemanticBibliographyEntry,
    SemanticBibliographyFragment,
    SemanticBibliographyRun,
    SemanticDocumentValidationError,
    SemanticParagraph,
    SemanticText,
    _bookmark_range,
    _remove_element,
    _render_bibliography,
    append_zero_width_bookmark,
    deepcopy,
    docx_semantics_module,
    encode_bibliography_entry_bookmark,
    pytest,
    qn,
)

pytestmark = pytest.mark.contract


def test_empty_bibliography_transfers_section_properties_to_preceding_paragraph() -> None:
    document = Document()
    previous = document.add_paragraph("References")
    anchor = document.add_paragraph("exclusive bibliography anchor")
    body = document.element.body  # type: ignore[attr-defined]
    section_properties = deepcopy(body.sectPr)
    assert section_properties is not None
    anchor._p.get_or_add_pPr().append(section_properties)

    rendered = DocxSemanticRenderer(document).render_bibliography_fragment(
        SemanticBibliographyFragment(entries=()),
        placeholder_anchor=anchor,
    )

    assert rendered == ()
    assert [paragraph.text for paragraph in document.paragraphs] == ["References"]
    assert previous._p.pPr is not None
    assert previous._p.pPr.find(qn("w:sectPr")) is not None


def test_empty_bibliography_inserts_section_properties_before_tracked_ppr_change() -> None:
    from docx.oxml import OxmlElement

    document = Document()
    previous = document.add_paragraph("References")
    previous_properties = previous._p.get_or_add_pPr()
    previous_properties.append(OxmlElement("w:pPrChange"))
    anchor = document.add_paragraph("exclusive bibliography anchor")
    body = document.element.body  # type: ignore[attr-defined]
    section_properties = deepcopy(body.sectPr)
    assert section_properties is not None
    anchor._p.get_or_add_pPr().append(section_properties)

    DocxSemanticRenderer(document).render_bibliography_fragment(
        SemanticBibliographyFragment(entries=()),
        placeholder_anchor=anchor,
    )

    final_properties = previous._p.pPr
    assert final_properties is not None
    final_section = final_properties.find(qn("w:sectPr"))
    final_change = final_properties.find(qn("w:pPrChange"))
    assert final_section is not None
    assert final_change is not None
    children = list(final_properties)
    assert children.index(final_section) < children.index(final_change)


@pytest.mark.parametrize("unsafe_predecessor", ["none", "table", "section"])
def test_empty_bibliography_rejects_unsafe_section_transfer_without_write(unsafe_predecessor: str) -> None:
    document = Document()
    if unsafe_predecessor == "table":
        document.add_table(rows=1, cols=1)
    elif unsafe_predecessor == "section":
        previous = document.add_paragraph("References")
        body = document.element.body  # type: ignore[attr-defined]
        existing = deepcopy(body.sectPr)
        assert existing is not None
        previous._p.get_or_add_pPr().append(existing)
    anchor = document.add_paragraph("exclusive bibliography anchor")
    body = document.element.body  # type: ignore[attr-defined]
    section_properties = deepcopy(body.sectPr)
    assert section_properties is not None
    anchor._p.get_or_add_pPr().append(section_properties)
    if unsafe_predecessor == "none":
        body = document.element.body  # type: ignore[attr-defined]
        body.remove(anchor._p)
        body.insert(0, anchor._p)
    before = document.element.xml

    with pytest.raises(SemanticDocumentValidationError) as error:
        DocxSemanticRenderer(document).render_bibliography_fragment(
            SemanticBibliographyFragment(entries=()),
            placeholder_anchor=anchor,
        )

    assert [item.code for item in error.value.diagnostics] == ["semantic.docx.bibliography.section_unsafe"]
    assert document.element.xml == before


def test_bibliography_render_failure_rolls_back_relationships_and_bookmark_cursor(monkeypatch) -> None:
    document = Document()
    document.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
    anchor = document.add_paragraph("exclusive bibliography anchor")
    retained_id = document.part.relate_to(
        "https://example.org/retained",
        RT.HYPERLINK,
        is_external=True,
    )
    fragment = SemanticBibliographyFragment(
        entries=(
            SemanticBibliographyEntry(
                "first",
                (
                    SemanticBibliographyRun("Retained", href="https://example.org/retained"),
                    SemanticBibliographyRun("New", href="https://example.org/new"),
                ),
            ),
        )
    )
    renderer = DocxSemanticRenderer(document, bookmark_id_start=4100)
    before_xml = document.element.xml
    before_relationships = {
        relationship_id: (relationship.reltype, relationship.target_ref, relationship.is_external)
        for relationship_id, relationship in document.part.rels.items()
    }
    original_append = docx_semantics_module.append_bookmark_end

    def fail_after_relationships(*_args, **_kwargs) -> None:
        raise RuntimeError("injected marker failure")

    monkeypatch.setattr(docx_semantics_module, "append_bookmark_end", fail_after_relationships)
    with pytest.raises(RuntimeError, match="injected marker failure"):
        renderer.render_bibliography_fragment(
            fragment,
            placeholder_anchor=anchor,
            hyperlink_style_id="Hyperlink",
        )

    assert document.element.xml == before_xml
    assert {
        relationship_id: (relationship.reltype, relationship.target_ref, relationship.is_external)
        for relationship_id, relationship in document.part.rels.items()
    } == before_relationships
    assert retained_id in document.part.rels

    monkeypatch.setattr(docx_semantics_module, "append_bookmark_end", original_append)
    renderer.render_bibliography_fragment(
        fragment,
        placeholder_anchor=anchor,
        hyperlink_style_id="Hyperlink",
    )
    bookmark_ids = [
        item.get(qn("w:id"))
        for item in document.element.iter(qn("w:bookmarkStart"))
        if (item.get(qn("w:name")) or "").startswith("_DW")
    ]
    assert bookmark_ids == ["4100", "4101"]


@pytest.mark.parametrize("damage", ["missing_relationship", "unsupported_run_property"])
def test_bibliography_rich_ooxml_damage_fails_closed_and_preserves_visible_text(damage: str) -> None:
    document, _fragment = _render_bibliography()
    hyperlink = next(document.element.iter(qn("w:hyperlink")))
    if damage == "missing_relationship":
        hyperlink.set(qn("r:id"), "rIdMissing")
    else:
        from docx.oxml import OxmlElement

        run_properties = next(hyperlink.iter(qn("w:rPr")))
        run_properties.append(OxmlElement("w:u"))

    imported = DocxSemanticImporter().import_document(document)

    assert imported.document.bibliography is None
    assert "semantic.docx.bibliography.entry_marker_invalid" in {item.code for item in imported.diagnostics}
    assert "Neutral documents" in "".join(
        inline.value
        for block in imported.document.blocks
        if isinstance(block, SemanticParagraph)
        for inline in block.inlines
        if isinstance(inline, SemanticText)
    )


def test_bibliography_renderer_rejects_an_anchor_from_another_document() -> None:
    document = Document()
    foreign_document = Document()
    foreign_anchor = foreign_document.add_paragraph("foreign anchor")
    fragment = SemanticBibliographyFragment(
        entries=(
            SemanticBibliographyEntry(
                "smith2025",
                (SemanticBibliographyRun("Formatted entry."),),
            ),
        )
    )
    document_before = document.element.xml
    anchor_before = foreign_anchor._p.xml

    with pytest.raises(ValueError, match="must belong to the renderer document part"):
        DocxSemanticRenderer(document).render_bibliography_fragment(
            fragment,
            placeholder_anchor=foreign_anchor,
        )

    assert document.element.xml == document_before
    assert foreign_anchor._p.xml == anchor_before


@pytest.mark.parametrize("position", ["before", "inside"])
def test_extra_bookmark_outside_or_inside_entry_breaks_exact_containment(position: str) -> None:
    document, _ = _render_bibliography()
    entry_name = encode_bibliography_entry_bookmark("smith2025")
    entry_start, entry_end = _bookmark_range(document, entry_name)
    paragraph = document.paragraphs[1]
    append_zero_width_bookmark(paragraph, "_unrelated", "9000")
    unrelated_start, unrelated_end = _bookmark_range(document, "_unrelated")
    anchor = entry_start if position == "before" else entry_end
    anchor.addprevious(unrelated_start)
    anchor.addprevious(unrelated_end)

    imported = DocxSemanticImporter().import_document(document)

    assert imported.document.bibliography is None
    assert "semantic.docx.bibliography.entry_marker_invalid" in {item.code for item in imported.diagnostics}
    assert "Smith, A. (2025). Neutral documents." in {
        block.inlines[0].value
        for block in imported.document.blocks
        if isinstance(block, SemanticParagraph)
        and len(block.inlines) == 1
        and isinstance(block.inlines[0], SemanticText)
    }


def test_bibliography_entry_outside_total_boundary_fails_closed() -> None:
    document, _ = _render_bibliography()
    total_start, total_end = _bookmark_range(document, BIBLIOGRAPHY_BOOKMARK_NAME)
    first_entry = document.paragraphs[1]._p
    first_entry.append(total_end)
    assert total_start.getparent() is first_entry

    imported = DocxSemanticImporter().import_document(document)

    assert imported.document.bibliography is None
    assert "semantic.docx.bibliography.entry_outside_boundary" in {item.code for item in imported.diagnostics}


def test_duplicate_bibliography_entry_marker_fails_closed() -> None:
    document, _ = _render_bibliography()
    duplicate = document.add_paragraph("unrelated")
    append_zero_width_bookmark(
        duplicate,
        encode_bibliography_entry_bookmark("smith2025"),
        "9000",
    )

    imported = DocxSemanticImporter().import_document(document)

    assert imported.document.bibliography is None
    assert "semantic.docx.bibliography.entry_marker_invalid" in {item.code for item in imported.diagnostics}


def test_orphan_bibliography_entry_markers_fail_closed_and_preserve_text() -> None:
    document, _ = _render_bibliography()
    total_start, total_end = _bookmark_range(document, BIBLIOGRAPHY_BOOKMARK_NAME)
    _remove_element(total_start)
    _remove_element(total_end)

    imported = DocxSemanticImporter().import_document(document)

    assert imported.document.bibliography is None
    assert [item.code for item in imported.diagnostics] == ["semantic.docx.bibliography.entry_orphan"]
    assert "Smith, A. (2025). Neutral documents." in {
        block.inlines[0].value
        for block in imported.document.blocks
        if isinstance(block, SemanticParagraph)
        and len(block.inlines) == 1
        and isinstance(block.inlines[0], SemanticText)
    }


def test_duplicate_bibliography_total_marker_fails_closed() -> None:
    document, _ = _render_bibliography()
    duplicate = document.add_paragraph("unrelated")
    append_zero_width_bookmark(duplicate, BIBLIOGRAPHY_BOOKMARK_NAME, "9000")

    imported = DocxSemanticImporter().import_document(document)

    assert imported.document.bibliography is None
    assert [item.code for item in imported.diagnostics] == ["semantic.docx.bibliography.boundary_unbalanced"]


def test_bibliography_entry_with_a_field_fails_closed() -> None:
    document, _ = _render_bibliography()
    _entry_start, entry_end = _bookmark_range(
        document,
        encode_bibliography_entry_bookmark("smith2025"),
    )
    scratch = document.add_paragraph()
    DocxSemanticRenderer(document).append_complex_field(
        scratch,
        instruction=" DATE ",
        cached_result="2026-08-12",
    )
    for element in list(scratch._p):
        entry_end.addprevious(element)
    _remove_element(scratch._p)

    imported = DocxSemanticImporter().import_document(document)

    assert imported.document.bibliography is None
    assert "semantic.docx.bibliography.entry_marker_invalid" in {item.code for item in imported.diagnostics}
