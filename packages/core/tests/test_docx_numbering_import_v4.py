"""Proof-only DOCX numbering import and mutation gates."""

from __future__ import annotations

from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import lxml.etree as etree
import pytest
from docx import Document
from docx.oxml.ns import qn

from docwen_core.docx_numbering_import import (
    AMBIGUOUS_VISIBLE_PREFIX_DIAGNOSTIC,
    DocxNumberingImportError,
    HeadingNumberingProofIndex,
    import_heading_without_source_mutation,
)
from docwen_core.docx_numbering_ooxml import (
    apply_heading_numbering,
    create_heading_numbering_projection,
    existing_numbering_ids,
    write_heading_numbering_projection,
)
from docwen_core.models.resolved_numbering import (
    HeadingCounterSegment,
    HeadingDefinition,
    HeadingInstance,
    HeadingLevelDefinition,
    HeadingListMaterialization,
    HeadingStart,
    NumberingTarget,
    ResolvedNumberingPlan,
)

pytestmark = pytest.mark.unit


def _plan() -> tuple[ResolvedNumberingPlan, NumberingTarget]:
    target = NumberingTarget(
        source_start=0,
        source_end=12,
        kind="heading",
        enabled=True,
        target_id=None,
        derived_number="3",
        materialization=HeadingListMaterialization("main", "scope", 1),
    )
    plan = ResolvedNumberingPlan(
        heading_definitions=(
            HeadingDefinition(
                "main",
                (
                    HeadingLevelDefinition(
                        level=1,
                        start=1,
                        number_format="arabic_half",
                        display=(HeadingCounterSegment(1, "arabic_half"),),
                        suffix="space",
                        restart_after_level=None,
                    ),
                ),
            ),
        ),
        heading_instances=(HeadingInstance("scope", "main", (HeadingStart(1, 3),)),),
        targets=(target,),
    )
    return plan, target


def _numbered_document(path: Path) -> None:
    document = Document()
    plan, target = _plan()
    projection = create_heading_numbering_projection(
        plan,
        heading_style_ids={1: "Heading1"},
        existing_abstract_ids=existing_numbering_ids(document)[0],
        existing_num_ids=existing_numbering_ids(document)[1],
    )
    paragraph = document.add_heading("2.3 标题", level=1)
    apply_heading_numbering(
        paragraph,
        target,
        projection,
        heading_style_ids={1: "Heading1"},
    )
    document.save(str(path))
    write_heading_numbering_projection(path, projection)


def _replace_member(path: Path, member: str, data: bytes) -> None:
    replacement = path.with_suffix(".replacement.docx")
    with ZipFile(path) as source, ZipFile(replacement, "w", ZIP_DEFLATED) as output:
        for info in source.infolist():
            output.writestr(info, data if info.filename == member else source.read(info.filename))
    replacement.replace(path)


def test_proven_word_list_is_separate_from_complete_authored_heading(tmp_path: Path) -> None:
    source = tmp_path / "proven.docx"
    _numbered_document(source)
    reopened = Document(str(source))
    paragraph = reopened.paragraphs[0]

    imported = import_heading_without_source_mutation(
        paragraph,
        HeadingNumberingProofIndex.load(source),
        suspected_visible_prefix=True,
    )

    assert imported.authored_text == "2.3 标题"
    assert imported.diagnostics == ()
    assert imported.numbering is not None
    assert imported.numbering.level == 1
    assert imported.numbering.number_format == "decimal"
    assert imported.numbering.level_text == "%1"
    assert imported.numbering.start == 3
    assert imported.numbering.style_id == "Heading1"


def test_unproven_visible_prefix_remains_authored_and_gets_stable_diagnostic(
    tmp_path: Path,
) -> None:
    source = tmp_path / "plain.docx"
    document = Document()
    document.add_heading("2.3 标题", level=1)
    document.save(str(source))
    paragraph = Document(str(source)).paragraphs[0]

    imported = import_heading_without_source_mutation(
        paragraph,
        HeadingNumberingProofIndex.load(source),
        suspected_visible_prefix=True,
    )

    assert imported.authored_text == "2.3 标题"
    assert imported.numbering is None
    assert [item.code for item in imported.diagnostics] == [AMBIGUOUS_VISIBLE_PREFIX_DIAGNOSTIC]


def test_missing_numbering_instance_fails_instead_of_parsing_visible_text(tmp_path: Path) -> None:
    source = tmp_path / "missing-instance.docx"
    _numbered_document(source)
    document_xml: bytes
    with ZipFile(source) as package:
        document = etree.fromstring(package.read("word/document.xml"))
        num_id = document.find(f".//{qn('w:numPr')}/{qn('w:numId')}")
        assert num_id is not None
        num_id.set(qn("w:val"), "2147483647")
        document_xml = etree.tostring(document, encoding="UTF-8", xml_declaration=True)
    _replace_member(source, "word/document.xml", document_xml)
    paragraph = Document(str(source)).paragraphs[0]

    with pytest.raises(DocxNumberingImportError, match="missing numbering instance"):
        import_heading_without_source_mutation(
            paragraph,
            HeadingNumberingProofIndex.load(source),
            suspected_visible_prefix=True,
        )
    assert paragraph.text == "2.3 标题"


def test_duplicate_numbering_definition_fails_package_proof(tmp_path: Path) -> None:
    source = tmp_path / "duplicate-definition.docx"
    _numbered_document(source)
    with ZipFile(source) as package:
        numbering = etree.fromstring(package.read("word/numbering.xml"))
        first = numbering.find(qn("w:abstractNum"))
        assert first is not None
        numbering.append(etree.fromstring(etree.tostring(first)))
        data = etree.tostring(numbering, encoding="UTF-8", xml_declaration=True)
    _replace_member(source, "word/numbering.xml", data)

    with pytest.raises(DocxNumberingImportError, match="repeats an abstractNumId"):
        HeadingNumberingProofIndex.load(source)


def test_inline_level_override_never_inherits_abstract_level_semantics(tmp_path: Path) -> None:
    source = tmp_path / "inline-level-override.docx"
    _numbered_document(source)
    with ZipFile(source) as package:
        numbering = etree.fromstring(package.read("word/numbering.xml"))
    override = numbering.find(f"{qn('w:num')}/{qn('w:lvlOverride')}")
    assert override is not None
    inline_level = etree.SubElement(override, qn("w:lvl"))
    inline_level.set(qn("w:ilvl"), "0")
    etree.SubElement(inline_level, qn("w:numFmt")).set(qn("w:val"), "upperRoman")
    _replace_member(
        source,
        "word/numbering.xml",
        etree.tostring(numbering, encoding="UTF-8", xml_declaration=True),
    )
    paragraph = Document(str(source)).paragraphs[0]

    with pytest.raises(DocxNumberingImportError, match="exact ancestry"):
        HeadingNumberingProofIndex.load(source)
    assert paragraph.text == "2.3 标题"


def test_mc_alternate_content_cannot_hide_an_effective_level_override(tmp_path: Path) -> None:
    source = tmp_path / "alternate-content-override.docx"
    _numbered_document(source)
    with ZipFile(source) as package:
        numbering = etree.fromstring(package.read("word/numbering.xml"))
    override = numbering.find(f"{qn('w:num')}/{qn('w:lvlOverride')}")
    assert override is not None
    instance = override.getparent()
    assert instance is not None
    instance.remove(override)
    alternate = etree.SubElement(
        instance,
        "{http://schemas.openxmlformats.org/markup-compatibility/2006}AlternateContent",
    )
    choice = etree.SubElement(
        alternate,
        "{http://schemas.openxmlformats.org/markup-compatibility/2006}Choice",
    )
    choice.set("Requires", "w14")
    choice.append(override)
    _replace_member(
        source,
        "word/numbering.xml",
        etree.tostring(numbering, encoding="UTF-8", xml_declaration=True),
    )

    with pytest.raises(DocxNumberingImportError, match="markup-compatibility topology"):
        HeadingNumberingProofIndex.load(source)


def test_numid_zero_is_explicitly_unnumbered_and_never_changes_text(tmp_path: Path) -> None:
    source = tmp_path / "disabled.docx"
    document = Document()
    paragraph = document.add_heading("第二章 概述", level=1)
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = etree.SubElement(p_pr, qn("w:numPr"))
    etree.SubElement(num_pr, qn("w:ilvl")).set(qn("w:val"), "0")
    etree.SubElement(num_pr, qn("w:numId")).set(qn("w:val"), "0")
    document.save(str(source))
    reopened = Document(str(source))

    imported = import_heading_without_source_mutation(
        reopened.paragraphs[0],
        HeadingNumberingProofIndex.load(source),
        suspected_visible_prefix=True,
    )

    assert imported.authored_text == "第二章 概述"
    assert imported.numbering is None
    assert imported.diagnostics[0].code == AMBIGUOUS_VISIBLE_PREFIX_DIAGNOSTIC
