"""Document-consistency mutations for the resolved-numbering DOCX session."""

from __future__ import annotations

import hashlib
from copy import deepcopy
from dataclasses import replace
from pathlib import Path
from typing import Any, cast
from zipfile import ZipFile

import lxml.etree as etree
import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

from docwen_core._docx_semantics_v3_model import (
    CaptionStyleBindingV3,
    CaptionStyleKeyV3,
)
from docwen_core.docx_resolved_numbering import (
    ResolvedNumberingDocxError,
    ResolvedNumberingDocxSession,
)
from docwen_core.models.resolved_numbering import (
    CaptionMaterialization,
    HeadingCounterSegment,
    HeadingDefinition,
    HeadingInstance,
    HeadingLevelDefinition,
    HeadingListMaterialization,
    NumberingExportPlanEnvelope,
    NumberingTarget,
    ResolvedCitation,
    ResolvedCitationItem,
    ResolvedDocument,
    ResolvedDocumentEnvelope,
    ResolvedDocumentTarget,
    ResolvedNumberingPlan,
    ResolvedNumberingPort,
    ResolvedReference,
)

pytestmark = pytest.mark.unit


def _sha(value: str) -> str:
    return hashlib.sha256(value.encode()).hexdigest()


def _caption_bindings(document: Any) -> tuple[CaptionStyleBindingV3, ...]:
    bindings: list[CaptionStyleBindingV3] = []
    for key, style_id, name in (
        ("figure_caption", "DWFigureCaption", "Figure Caption V4"),
        ("table_caption", "DWTableCaption", "Table Caption V4"),
        ("equation_caption", "DWEquationCaption", "Equation Caption V4"),
        ("code_block_caption", "DWCodeCaption", "Code Caption V4"),
    ):
        style = document.styles.add_style(style_id, WD_STYLE_TYPE.PARAGRAPH)
        style.name = name
        bindings.append(CaptionStyleBindingV3(cast(CaptionStyleKeyV3, key), style.style_id, style.name))
    return tuple(bindings)


def _caption_materialization(sequence: int) -> CaptionMaterialization:
    return CaptionMaterialization(
        type="simple_seq",
        counter="Table",
        number_format="arabic_half",
        sequence_action="continue",
        start_value=None,
        chapter_heading_level=None,
        chapter_heading_style=None,
        chapter_separator=None,
        restart_heading_level=None,
        restart_heading_style=None,
        chapter_cached_number=None,
        sequence_cached_number=str(sequence),
        localized_label="Table",
        label_separator=" ",
    )


def _two_table_port(state: str) -> ResolvedNumberingPort:
    source = "Table: One\n|a|\n|-|\n|1|\n\nTable: Two\n|a|\n|-|\n|2|\n"
    separator = source.index("\n\n")
    ranges = ((0, separator), (separator + 2, len(source) - 1))
    enabled = state != "disabled_idless"
    ids = ("table-a", "table-b") if state == "id_bearing" else (None, None)
    document_targets = tuple(
        ResolvedDocumentTarget(start, end, _sha(source[start:end]), "table", target_id, None, title)
        for (start, end), target_id, title in zip(ranges, ids, ("One", "Two"), strict=True)
    )
    plan_targets = tuple(
        NumberingTarget(
            start,
            end,
            "table",
            enabled,
            target_id,
            str(index) if enabled else None,
            _caption_materialization(index) if enabled else None,
        )
        for index, ((start, end), target_id) in enumerate(zip(ranges, ids, strict=True), start=1)
    )
    source_sha = _sha(source)
    plan_sha = "b" * 64
    return ResolvedNumberingPort(
        ResolvedDocumentEnvelope(
            "source",
            source_sha,
            plan_sha,
            ResolvedDocument(source, document_targets, (), (), (), ()),
        ),
        NumberingExportPlanEnvelope(
            "source",
            source_sha,
            plan_sha,
            ResolvedNumberingPlan((), (), plan_targets),
        ),
    )


def _render_two_tables(tmp_path: Path, state: str) -> tuple[ResolvedNumberingDocxSession, Path]:
    document = Document()
    port = _two_table_port(state)
    session = ResolvedNumberingDocxSession(
        document,
        port,
        heading_style_ids={},
        heading_style_names={},
        caption_style_bindings=_caption_bindings(document),
    )
    for target in port.document.targets:
        caption = document.add_paragraph(style="Table Caption V4")
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = target.authored_text
        session.bind_caption(
            caption,
            (table._element,),
            source_start=target.source_start,
            source_end=target.source_end,
            kind="table",
        )
    output = tmp_path / f"two-tables-{state}.docx"
    session.write_package(output)
    return session, output


@pytest.mark.parametrize("state", ["disabled_idless", "enabled_idless", "id_bearing"])
def test_caption_object_slots_cannot_exchange_same_kind_ooxml(tmp_path: Path, state: str) -> None:
    session, output = _render_two_tables(tmp_path, state)
    with ZipFile(output) as package:
        infos = package.infolist()
        data = {item.filename: package.read(item.filename) for item in infos}
    root = etree.fromstring(data["word/document.xml"])
    tables = list(root.iter(qn("w:tbl")))
    assert len(tables) == 2
    first_parent = tables[0].getparent()
    second_parent = tables[1].getparent()
    assert first_parent is not None and second_parent is not None
    first_index = first_parent.index(tables[0])
    second_index = second_parent.index(tables[1])
    first_parent.remove(tables[0])
    second_parent.remove(tables[1])
    first_parent.insert(first_index, deepcopy(tables[1]))
    second_parent.insert(second_index, deepcopy(tables[0]))
    data["word/document.xml"] = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    replacement = output.with_suffix(".mutated.docx")
    with ZipFile(replacement, "w") as package:
        for info in infos:
            package.writestr(info, data[info.filename])
    replacement.replace(output)

    with pytest.raises(ResolvedNumberingDocxError, match="object OOXML"):
        session.prove_package(output)


def _rich_heading_port() -> ResolvedNumberingPort:
    token = "@[[#^head-a|标题]]"
    citation_token = "@cite"
    title = f"literal {token} then {token} and {citation_token}"
    source = f"# {title} ^head-a\n"
    target_end = len(source) - 1
    reference_start = source.index(token, source.index(token) + 1)
    citation_start = source.index(citation_token)
    target = ResolvedDocumentTarget(0, target_end, _sha(source[:target_end]), "heading", "head-a", 1, title)
    reference = ResolvedReference(
        reference_start,
        reference_start + len(token),
        _sha(token),
        token,
        0,
        target_end,
        "heading",
        "head-a",
        "1",
        "标题",
    )
    citation = ResolvedCitation(
        citation_start,
        citation_start + len(citation_token),
        _sha(citation_token),
        citation_token,
        "narrative",
        "cluster-a",
        (ResolvedCitationItem("cite", "reference-record:98", _sha("record"), "Citation"),),
        "Citation",
    )
    definition = HeadingDefinition(
        "main",
        (HeadingLevelDefinition(1, 1, "arabic_half", (HeadingCounterSegment(1, "arabic_half"),), "space", None),),
    )
    plan_target = NumberingTarget(
        0,
        target_end,
        "heading",
        True,
        "head-a",
        "1",
        HeadingListMaterialization("main", "document", 1),
    )
    source_sha = _sha(source)
    plan_sha = "c" * 64
    return ResolvedNumberingPort(
        ResolvedDocumentEnvelope(
            "source",
            source_sha,
            plan_sha,
            ResolvedDocument(source, (target,), (reference,), (), (citation,), ()),
        ),
        NumberingExportPlanEnvelope(
            "source",
            source_sha,
            plan_sha,
            ResolvedNumberingPlan((definition,), (HeadingInstance("document", "main", ()),), (plan_target,)),
        ),
    )


def _heading_session(document: Any, port: ResolvedNumberingPort) -> ResolvedNumberingDocxSession:
    return ResolvedNumberingDocxSession(
        document,
        port,
        heading_style_ids={level: f"Heading{level}" for level in range(1, 10)},
        heading_style_names={f"heading_{level}": f"Heading {level}" for level in range(1, 10)},
        caption_style_bindings=_caption_bindings(document),
    )


def test_heading_nested_carriers_preserve_exact_source_position_and_merge_suffix(tmp_path: Path) -> None:
    document = Document()
    port = _rich_heading_port()
    session = _heading_session(document, port)
    heading = document.add_heading(level=1)
    reference = port.document.references[0]
    citation = port.document.citations[0]
    heading.add_run(f"literal {reference.authored_token} then ")
    session.render_reference(heading, source_start=reference.source_start, source_end=reference.source_end)
    heading.add_run(" and ")
    session.render_citation(heading, source_start=citation.source_start, source_end=citation.source_end)
    heading.add_run(" merged suffix")
    target = port.document.targets[0]
    session.bind_heading(heading, source_start=target.source_start, source_end=target.source_end)
    output = tmp_path / "rich-heading.docx"
    session.write_package(output)
    session.prove_package(output)


def test_heading_rejects_same_token_carrier_at_wrong_source_occurrence() -> None:
    document = Document()
    port = _rich_heading_port()
    session = _heading_session(document, port)
    heading = document.add_heading(level=1)
    reference = port.document.references[0]
    citation = port.document.citations[0]
    heading.add_run("literal ")
    session.render_reference(heading, source_start=reference.source_start, source_end=reference.source_end)
    heading.add_run(f" then {reference.authored_token} and ")
    session.render_citation(heading, source_start=citation.source_start, source_end=citation.source_end)
    target = port.document.targets[0]

    with pytest.raises(ResolvedNumberingDocxError, match="source-relative projection"):
        session.bind_heading(heading, source_start=target.source_start, source_end=target.source_end)


@pytest.mark.parametrize("carrier", ["reference", "citation"])
def test_heading_rejects_nested_carrier_bound_to_another_paragraph(carrier: str) -> None:
    document = Document()
    port = _rich_heading_port()
    session = _heading_session(document, port)
    heading = document.add_heading(level=1)
    decoy = document.add_paragraph()
    reference = port.document.references[0]
    citation = port.document.citations[0]
    heading.add_run(f"literal {reference.authored_token} then ")
    session.render_reference(
        decoy if carrier == "reference" else heading,
        source_start=reference.source_start,
        source_end=reference.source_end,
    )
    heading.add_run(" and ")
    session.render_citation(
        decoy if carrier == "citation" else heading,
        source_start=citation.source_start,
        source_end=citation.source_end,
    )
    target = port.document.targets[0]

    with pytest.raises(ResolvedNumberingDocxError, match="belongs to another paragraph"):
        session.bind_heading(heading, source_start=target.source_start, source_end=target.source_end)


def test_session_revalidates_direct_typed_heading_level_before_projection() -> None:
    document = Document()
    port = _rich_heading_port()
    forged_target = replace(port.document.targets[0], heading_level=7)
    forged_document = replace(port.document, targets=(forged_target,))
    forged_port = replace(
        port,
        document_envelope=replace(port.document_envelope, document=forged_document),
    )

    with pytest.raises(ResolvedNumberingDocxError, match="runtime revalidation"):
        _heading_session(document, forged_port)


def test_session_keeps_unused_definition_levels_seven_through_nine() -> None:
    document = Document()
    port = _rich_heading_port()
    definition = port.plan.heading_definitions[0]
    extra_levels = tuple(
        HeadingLevelDefinition(
            level,
            1,
            "arabic_half",
            (HeadingCounterSegment(level, "arabic_half"),),
            "space",
            None,
        )
        for level in range(2, 10)
    )
    expanded_definition = replace(definition, levels=(*definition.levels, *extra_levels))
    expanded_plan = replace(port.plan, heading_definitions=(expanded_definition,))
    expanded_port = replace(port, plan_envelope=replace(port.plan_envelope, plan=expanded_plan))

    session = _heading_session(document, expanded_port)

    assert len(list(session.projection.abstract_nums[0].iter(qn("w:lvl")))) == 9
