"""Closed caption-style map and exact caption-kind authentication tests."""

from __future__ import annotations

import hashlib
import re
import uuid
from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile

import lxml.etree as etree
import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docwen_core._docx_semantics_v3_styles import caption_style_binding_map_xml
from docwen_core.docx_semantics_v3 import (
    CAPTION_STYLE_BINDING_MAP_NAMESPACE,
    CaptionStyleBindingV3,
    DocxSemanticsV3Error,
    DocxSemanticsV3Recovery,
    DocxSemanticsV3Session,
    append_complex_field,
)

pytestmark = pytest.mark.contract

_DS_NAMESPACE = "http://schemas.openxmlformats.org/officeDocument/2006/customXml"
_REL_NAMESPACE = "http://schemas.openxmlformats.org/package/2006/relationships"
_CT_NAMESPACE = "http://schemas.openxmlformats.org/package/2006/content-types"
_CUSTOM_XML_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml"


def test_addressable_and_idless_captions_use_persisted_collision_binding(tmp_path: Path) -> None:
    document = Document()
    bindings, styles = _add_caption_styles(document, table_collision=True)
    session = DocxSemanticsV3Session(
        document,
        source_sha256=hashlib.sha256(b"two tables").hexdigest(),
        caption_style_bindings=bindings,
    )
    first_caption, first_table = _add_table_caption(document, styles["table_caption"], "Addressable")
    session.bind_caption(
        first_caption,
        (first_table._tbl,),
        {"kind": "table", "id": "results", "number": "1", "title": "Addressable"},
    )
    second_caption, second_table = _add_table_caption(document, styles["table_caption"], "ID-less", number="2")
    session.bind_caption(
        second_caption,
        (second_table._tbl,),
        {"kind": "table", "id": None, "number": "2", "title": "ID-less"},
    )
    output = _write(session, document, tmp_path / "collision-roundtrip.docx")

    recovery = DocxSemanticsV3Recovery.load(output, Document(str(output)))
    assert recovery.caption_style_bindings == bindings
    assert recovery.caption_signatures == (
        ("table", "results", "Addressable", "1"),
        ("table", None, "ID-less", "2"),
    )
    with ZipFile(output) as package:
        item_name, _props_name = _owned_item_and_props(package)
        xml = package.read(item_name)
    assert xml == caption_style_binding_map_xml(bindings)
    assert b'resolved_style_id="DocWenTableCaptionDocWen1"' in xml


def test_unmapped_prefix_style_is_never_caption_authority(tmp_path: Path) -> None:
    document = Document()
    forged = document.styles.add_style("Forged Caption", WD_STYLE_TYPE.PARAGRAPH)
    forged_element = forged._element
    assert forged_element is not None
    forged_element.set(qn("w:styleId"), "DocWenTableCaptionForged")
    _add_table_caption(document, forged, "Forged")
    session = DocxSemanticsV3Session(document, source_sha256=hashlib.sha256(b"forged").hexdigest())
    session.finalize_document()
    output = tmp_path / "forged.docx"
    document.save(str(output))
    recovery = DocxSemanticsV3Recovery.load(output, Document(str(output)))
    assert recovery.recovered_captions == ()
    assert recovery.caption_style_bindings == ()


def test_bind_caption_rejects_cross_kind_pstyle() -> None:
    document = Document()
    bindings, styles = _add_caption_styles(document)
    session = DocxSemanticsV3Session(
        document,
        source_sha256=hashlib.sha256(b"cross-kind").hexdigest(),
        caption_style_bindings=bindings,
    )
    caption, table = _add_table_caption(document, styles["figure_caption"], "Wrong")
    with pytest.raises(DocxSemanticsV3Error, match="pStyle does not match"):
        session.bind_caption(
            caption,
            (table._tbl,),
            {"kind": "table", "id": None, "number": "1", "title": "Wrong"},
        )


def test_recovery_rejects_missing_caption_style_map(tmp_path: Path) -> None:
    output = _valid_table_package(tmp_path)
    _remove_caption_style_map(output)
    with pytest.raises(DocxSemanticsV3Error, match="lacks its authenticated caption-style"):
        DocxSemanticsV3Recovery.load(output, Document(str(output)))


@pytest.mark.parametrize(
    ("damage", "message"),
    (
        ("missing", "does not resolve one exact"),
        ("type", "does not resolve one exact"),
        ("name", "visible name is not exact"),
        ("alias", "must not use aliases"),
        ("duplicate_id", "does not resolve one exact"),
    ),
)
def test_styles_registry_damage_fails_closed(tmp_path: Path, damage: str, message: str) -> None:
    output = _valid_table_package(tmp_path, name=f"style-{damage}.docx")

    def mutate(root) -> None:
        style = next(item for item in root.findall(qn("w:style")) if item.get(qn("w:styleId")) == "DocWenTableCaption")
        if damage == "missing":
            root.remove(style)
        elif damage == "type":
            style.set(qn("w:type"), "character")
        elif damage == "name":
            style.find(qn("w:name")).set(qn("w:val"), "Not Table Caption")
        elif damage == "alias":
            alias = OxmlElement("w:aliases")
            alias.set(qn("w:val"), "Alias")
            style.append(alias)
        else:
            duplicate = deepcopy(style)
            root.append(duplicate)

    _mutate_xml_member(output, "word/styles.xml", mutate)
    with pytest.raises(DocxSemanticsV3Error, match=message):
        DocxSemanticsV3Recovery.load(output, Document(str(output)))


@pytest.mark.parametrize("damage", ("missing", "duplicate", "cross_kind"))
def test_caption_paragraph_pstyle_tampering_fails_closed(tmp_path: Path, damage: str) -> None:
    output = _valid_table_package(tmp_path)

    def mutate(root) -> None:
        style = next(item for item in root.iter(qn("w:pStyle")) if item.get(qn("w:val")) == "DocWenTableCaption")
        if damage == "missing":
            style.getparent().remove(style)
        elif damage == "duplicate":
            style.addnext(deepcopy(style))
        else:
            style.set(qn("w:val"), "DocWenFigureCaption")

    _mutate_xml_member(output, "word/document.xml", mutate)
    message = "one exact direct pStyle" if damage != "cross_kind" else "pStyle does not match"
    with pytest.raises(DocxSemanticsV3Error, match=message):
        DocxSemanticsV3Recovery.load(output, Document(str(output)))


def test_caption_style_map_uuid_binds_exact_bytes(tmp_path: Path) -> None:
    output = _valid_table_package(tmp_path)
    with ZipFile(output) as package:
        item_name, _props_name = _owned_item_and_props(package)
        mutated = package.read(item_name).replace(b"Table Caption", b"Table Captiox", 1)
    _replace_zip_members(output, {item_name: mutated})
    with pytest.raises(DocxSemanticsV3Error, match="UUID does not match"):
        DocxSemanticsV3Recovery.load(output, Document(str(output)))


def test_caption_style_map_is_closed_after_uuid_rebind(tmp_path: Path) -> None:
    output = _valid_table_package(tmp_path)
    with ZipFile(output) as package:
        item_name, props_name = _owned_item_and_props(package)
        item = package.read(item_name)
        props = package.read(props_name)
    mutated = item.replace(b'<binding semantic_key="table_caption"', b'<binding extra="x" semantic_key="table_caption"')
    _replace_zip_members(
        output,
        {item_name: mutated, props_name: _rebind_props_uuid(props, mutated)},
    )
    with pytest.raises(DocxSemanticsV3Error, match="record is not closed"):
        DocxSemanticsV3Recovery.load(output, Document(str(output)))


def test_caption_style_map_reordered_records_fail_after_uuid_rebind(tmp_path: Path) -> None:
    output = _valid_table_package(tmp_path)
    with ZipFile(output) as package:
        item_name, props_name = _owned_item_and_props(package)
        original = package.read(item_name)
        props = package.read(props_name)
    pattern = re.compile(
        rb'(<binding semantic_key="figure_caption"[^>]*/>)(<binding semantic_key="table_caption"[^>]*/>)'
    )
    mutated, count = pattern.subn(rb"\2\1", original, count=1)
    assert count == 1
    _replace_zip_members(
        output,
        {item_name: mutated, props_name: _rebind_props_uuid(props, mutated)},
    )
    with pytest.raises(DocxSemanticsV3Error, match="canonically ordered"):
        DocxSemanticsV3Recovery.load(output, Document(str(output)))


def test_caption_style_map_rejects_duplicate_resolved_id() -> None:
    document = Document()
    bindings, _styles = _add_caption_styles(document)
    duplicate = tuple(
        CaptionStyleBindingV3(
            item.semantic_key,
            "DocWenFigureCaption" if item.semantic_key == "table_caption" else item.resolved_style_id,
            item.visible_name,
        )
        for item in bindings
    )
    with pytest.raises(DocxSemanticsV3Error, match="style IDs are not unique"):
        caption_style_binding_map_xml(duplicate)


def _valid_table_package(tmp_path: Path, *, name: str = "valid-table.docx") -> Path:
    document = Document()
    bindings, styles = _add_caption_styles(document)
    session = DocxSemanticsV3Session(
        document,
        source_sha256=hashlib.sha256(name.encode()).hexdigest(),
        caption_style_bindings=bindings,
    )
    caption, table = _add_table_caption(document, styles["table_caption"], "Results")
    session.bind_caption(
        caption,
        (table._tbl,),
        {"kind": "table", "id": "results", "number": "1", "title": "Results"},
    )
    return _write(session, document, tmp_path / name)


def _add_caption_styles(document, *, table_collision: bool = False):
    definitions = (
        ("figure_caption", "DocWenFigureCaption", "Figure Caption"),
        ("table_caption", "DocWenTableCaption", "Table Caption"),
        ("equation_caption", "DocWenEquationCaption", "Equation Caption"),
        ("code_block_caption", "DocWenCodeBlockCaption", "Code Block Caption"),
    )
    if table_collision:
        conflict = document.styles.add_style("Conflicting Table Identity", WD_STYLE_TYPE.CHARACTER)
        conflict._element.set(qn("w:styleId"), "DocWenTableCaption")
    styles = {}
    bindings = []
    for key, requested_id, visible_name in definitions:
        style = document.styles.add_style(visible_name, WD_STYLE_TYPE.PARAGRAPH)
        resolved = "DocWenTableCaptionDocWen1" if table_collision and key == "table_caption" else requested_id
        style._element.set(qn("w:styleId"), resolved)
        styles[key] = style
        bindings.append(CaptionStyleBindingV3(key, resolved, visible_name))  # type: ignore[arg-type]
    return tuple(bindings), styles


def _add_table_caption(document, style, title: str, *, number: str = "1"):
    caption = document.add_paragraph(style=style)
    caption.add_run("Table ")
    append_complex_field(caption, instruction=" SEQ Table \\* ARABIC ", cached_result=number)
    caption.add_run(f": {title}")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = title
    return caption, table


def _write(session, document, output: Path) -> Path:
    session.finalize_document()
    document.save(str(output))
    session.write_package(output)
    session.prove_package(output)
    return output


def _owned_item_and_props(package: ZipFile) -> tuple[str, str]:
    item = next(
        name
        for name in package.namelist()
        if re.fullmatch(r"customXml/item[1-9][0-9]*\.xml", name)
        and CAPTION_STYLE_BINDING_MAP_NAMESPACE.encode() in package.read(name)
    )
    number = re.fullmatch(r"customXml/item([1-9][0-9]*)\.xml", item).group(1)  # type: ignore[union-attr]
    return item, f"customXml/itemProps{number}.xml"


def _rebind_props_uuid(props: bytes, item: bytes) -> bytes:
    value = (
        "{"
        + str(
            uuid.uuid5(
                uuid.NAMESPACE_URL,
                f"{CAPTION_STYLE_BINDING_MAP_NAMESPACE}\0{hashlib.sha256(item).hexdigest()}",
            )
        ).upper()
        + "}"
    )
    return re.sub(rb"\{[0-9A-F-]{36}\}", value.encode(), props, count=1)


def _remove_caption_style_map(path: Path) -> None:
    with ZipFile(path) as package:
        item_name, props_name = _owned_item_and_props(package)
        number = re.fullmatch(r"customXml/item([1-9][0-9]*)\.xml", item_name).group(1)  # type: ignore[union-attr]
        rels_name = f"customXml/_rels/item{number}.xml.rels"
        replacements = {name: package.read(name) for name in package.namelist()}
    for name in (item_name, props_name, rels_name):
        replacements.pop(name)
    document_rels = etree.fromstring(replacements["word/_rels/document.xml.rels"])
    for relation in list(document_rels):
        if relation.get("Type") == _CUSTOM_XML_REL and relation.get("Target") == f"../{item_name}":
            document_rels.remove(relation)
    replacements["word/_rels/document.xml.rels"] = etree.tostring(
        document_rels, encoding="UTF-8", xml_declaration=True, standalone=True
    )
    types = etree.fromstring(replacements["[Content_Types].xml"])
    for item in list(types):
        if item.get("PartName") in {f"/{item_name}", f"/{props_name}"}:
            types.remove(item)
    replacements["[Content_Types].xml"] = etree.tostring(types, encoding="UTF-8", xml_declaration=True, standalone=True)
    _rewrite_zip(path, replacements)


def _mutate_xml_member(path: Path, name: str, mutate) -> None:
    root = etree.fromstring(_read(path, name))
    mutate(root)
    _replace_zip_members(
        path,
        {name: etree.tostring(root, encoding="UTF-8", xml_declaration=True, standalone=True)},
    )


def _read(path: Path, name: str) -> bytes:
    with ZipFile(path) as package:
        return package.read(name)


def _replace_zip_members(path: Path, replacements: dict[str, bytes]) -> None:
    with ZipFile(path) as package:
        all_members = {name: package.read(name) for name in package.namelist()}
    all_members.update(replacements)
    _rewrite_zip(path, all_members)


def _rewrite_zip(path: Path, members: dict[str, bytes]) -> None:
    temporary = path.with_suffix(".tmp")
    with ZipFile(temporary, "w") as target:
        for name, data in members.items():
            target.writestr(name, data)
    temporary.replace(path)
