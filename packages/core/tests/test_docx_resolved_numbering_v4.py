"""Request-level v4 numbering DOCX session gates."""

from __future__ import annotations

import hashlib
from collections.abc import Callable
from pathlib import Path
from typing import cast
from zipfile import ZipFile, ZipInfo

import lxml.etree as etree
import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

from docwen_core._docx_semantics_v3_model import CaptionStyleBindingV3, CaptionStyleKeyV3
from docwen_core.docx_numbering_occurrence import NUMBERING_OCCURRENCE_MAP_NAMESPACE
from docwen_core.docx_resolved_numbering import (
    ResolvedNumberingDocxError,
    ResolvedNumberingDocxSession,
)
from docwen_core.models.resolved_numbering import (
    HeadingCounterSegment,
    HeadingDefinition,
    HeadingInstance,
    HeadingLevelDefinition,
    HeadingListMaterialization,
    NumberingExportPlanEnvelope,
    NumberingTarget,
    ResolvedDocument,
    ResolvedDocumentEnvelope,
    ResolvedDocumentTarget,
    ResolvedNumberingPlan,
    ResolvedNumberingPort,
    ResolvedReference,
)

pytestmark = pytest.mark.unit


def _sha(value: str) -> str:
    return hashlib.sha256(value.encode()).hexdigest()


def _caption_bindings(document) -> tuple[CaptionStyleBindingV3, ...]:
    output: list[CaptionStyleBindingV3] = []
    for semantic_key, style_id, name in (
        ("figure_caption", "DWFigureCaption", "Figure Caption V4"),
        ("table_caption", "DWTableCaption", "Table Caption V4"),
        ("equation_caption", "DWEquationCaption", "Equation Caption V4"),
        ("code_block_caption", "DWCodeCaption", "Code Caption V4"),
    ):
        style = document.styles.add_style(style_id, WD_STYLE_TYPE.PARAGRAPH)
        style.name = name
        output.append(CaptionStyleBindingV3(cast(CaptionStyleKeyV3, semantic_key), style.style_id, style.name))
    return tuple(output)


def _port() -> ResolvedNumberingPort:
    source = "# 2.3 标题\n\nTable: 数据\n|a|\n|-|\n|b|\n\n@[[#^head-a|标题]]\n"
    heading_start = 0
    heading_end = source.index("\n")
    table_start = source.index("Table:")
    table_end = source.index("\n\n@[[")
    token = "@[[#^head-a|标题]]"
    reference_start = source.index(token)
    reference_end = reference_start + len(token)
    heading_target = ResolvedDocumentTarget(
        source_start=heading_start,
        source_end=heading_end,
        source_slice_sha256=_sha(source[heading_start:heading_end]),
        kind="heading",
        target_id="head-a",
        heading_level=1,
        authored_text="2.3 标题",
    )
    table_target = ResolvedDocumentTarget(
        source_start=table_start,
        source_end=table_end,
        source_slice_sha256=_sha(source[table_start:table_end]),
        kind="table",
        target_id=None,
        heading_level=None,
        authored_text="数据",
    )
    reference = ResolvedReference(
        source_start=reference_start,
        source_end=reference_end,
        source_slice_sha256=_sha(token),
        authored_token=token,
        target_source_start=heading_start,
        target_source_end=heading_end,
        target_kind="heading",
        target_id="head-a",
        cached_number="1",
        alias="标题",
    )
    definition = HeadingDefinition(
        definition_id="main",
        levels=(
            HeadingLevelDefinition(
                level=1,
                start=1,
                number_format="arabic_half",
                display=(HeadingCounterSegment(1, "arabic_half"),),
                suffix="space",
                restart_after_level=None,
            ),
        ),
    )
    plan = ResolvedNumberingPlan(
        heading_definitions=(definition,),
        heading_instances=(HeadingInstance("document", "main", ()),),
        targets=(
            NumberingTarget(
                source_start=heading_start,
                source_end=heading_end,
                kind="heading",
                enabled=True,
                target_id="head-a",
                derived_number="1",
                materialization=HeadingListMaterialization("main", "document", 1),
            ),
            NumberingTarget(
                source_start=table_start,
                source_end=table_end,
                kind="table",
                enabled=False,
                target_id=None,
                derived_number=None,
                materialization=None,
            ),
        ),
    )
    plan_sha = "a" * 64
    return ResolvedNumberingPort(
        ResolvedDocumentEnvelope(
            input_id="source",
            source_sha256=_sha(source),
            plan_sha256=plan_sha,
            document=ResolvedDocument(
                source,
                (heading_target, table_target),
                (reference,),
                (),
                (),
                (),
            ),
        ),
        NumberingExportPlanEnvelope(
            input_id="source",
            source_sha256=_sha(source),
            plan_sha256=plan_sha,
            plan=plan,
        ),
    )


def test_session_writes_heading_ref_and_disabled_idless_occurrence(tmp_path: Path) -> None:
    document = Document()
    bindings = _caption_bindings(document)
    session = ResolvedNumberingDocxSession(
        document,
        _port(),
        heading_style_ids={1: "Heading1"},
        heading_style_names={"heading_1": "Heading 1"},
        caption_style_bindings=bindings,
    )
    heading = document.add_heading("2.3 标题", level=1)
    reference = document.add_paragraph()
    caption = document.add_paragraph(style="Table Caption V4")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "b"
    targets = session.port.document.targets
    session.bind_heading(
        heading,
        source_start=targets[0].source_start,
        source_end=targets[0].source_end,
    )
    session.bind_caption(
        caption,
        (table._element,),
        source_start=targets[1].source_start,
        source_end=targets[1].source_end,
        kind="table",
    )
    resolved_reference = session.port.document.references[0]
    session.render_reference(
        reference,
        source_start=resolved_reference.source_start,
        source_end=resolved_reference.source_end,
    )
    session.finalize_document()
    output = tmp_path / "resolved.docx"
    document.save(str(output))
    session.write_package(output)
    session.prove_package(output)

    assert heading.text == "2.3 标题"
    with ZipFile(output) as package:
        numbering = etree.fromstring(package.read("word/numbering.xml"))
        document_xml = etree.fromstring(package.read("word/document.xml"))
        assert any(
            NUMBERING_OCCURRENCE_MAP_NAMESPACE.encode() in package.read(name)
            for name in package.namelist()
            if name.startswith("customXml/item") and name.endswith(".xml")
        )
    assert numbering.find(qn("w:abstractNum")) is not None
    heading_paragraph = next(
        paragraph
        for paragraph in document_xml.iter(qn("w:p"))
        if "".join(item.text or "" for item in paragraph.iter(qn("w:t"))) == "2.3 标题"
    )
    assert heading_paragraph.find(f"{qn('w:pPr')}/{qn('w:numPr')}") is not None
    assert not list(heading_paragraph.iter(qn("w:instrText")))
    instructions = [item.text for item in document_xml.iter(qn("w:instrText"))]
    assert any(value and " REF " in value and r"\n \h" in value for value in instructions)
    assert "1" in [item.text for item in document_xml.iter(qn("w:t"))]
    assert " 标题" in [item.text for item in document_xml.iter(qn("w:t"))]


def test_session_rejects_incomplete_bindings_before_package_mutation(tmp_path: Path) -> None:
    document = Document()
    session = ResolvedNumberingDocxSession(
        document,
        _port(),
        heading_style_ids={1: "Heading1"},
        heading_style_names={"heading_1": "Heading 1"},
        caption_style_bindings=_caption_bindings(document),
    )
    target = session.port.document.targets[0]
    session.bind_heading(
        document.add_heading("2.3 标题", level=1),
        source_start=target.source_start,
        source_end=target.source_end,
    )
    output = tmp_path / "must-not-exist.docx"

    with pytest.raises(ResolvedNumberingDocxError, match="not every resolved target"):
        session.write_package(output)
    assert not output.exists()


def test_session_never_uses_authored_heading_prefix_as_number() -> None:
    document = Document()
    session = ResolvedNumberingDocxSession(
        document,
        _port(),
        heading_style_ids={1: "Heading1"},
        heading_style_names={"heading_1": "Heading 1"},
        caption_style_bindings=_caption_bindings(document),
    )
    target = session.port.document.targets[0]
    heading = document.add_heading("2.3 标题", level=1)
    session.bind_heading(
        heading,
        source_start=target.source_start,
        source_end=target.source_end,
    )

    assert heading.text == "2.3 标题"
    assert all((item.text or "") != "1" for item in heading._p.iter(qn("w:t")))


def test_heading_merge_snapshot_preserves_authored_prefix_and_full_rendered_text(tmp_path: Path) -> None:
    document = Document()
    session = ResolvedNumberingDocxSession(
        document,
        _port(),
        heading_style_ids={1: "Heading1"},
        heading_style_names={"heading_1": "Heading 1"},
        caption_style_bindings=_caption_bindings(document),
    )
    heading = document.add_heading("2.3 标题：合并正文", level=1)
    targets = session.port.document.targets
    session.bind_heading(
        heading,
        source_start=targets[0].source_start,
        source_end=targets[0].source_end,
    )
    caption = document.add_paragraph(style="Table Caption V4")
    table = document.add_table(rows=1, cols=1)
    session.bind_caption(
        caption,
        (table._element,),
        source_start=targets[1].source_start,
        source_end=targets[1].source_end,
        kind="table",
    )
    reference = document.add_paragraph()
    resolved_reference = session.port.document.references[0]
    session.render_reference(
        reference,
        source_start=resolved_reference.source_start,
        source_end=resolved_reference.source_end,
    )
    output = tmp_path / "merged-heading.docx"
    session.write_package(output)
    session.prove_package(output)

    _replace_zip_member(
        output,
        "word/document.xml",
        "合并正文".encode(),
        "篡改正文".encode(),
    )
    with pytest.raises(ValueError, match="rendered text changed"):
        session.prove_package(output)


def test_heading_bind_rejects_rendered_text_without_authored_prefix() -> None:
    document = Document()
    session = ResolvedNumberingDocxSession(
        document,
        _port(),
        heading_style_ids={1: "Heading1"},
        heading_style_names={"heading_1": "Heading 1"},
        caption_style_bindings=_caption_bindings(document),
    )
    target = session.port.document.targets[0]

    with pytest.raises(ResolvedNumberingDocxError, match="authored title prefix"):
        session.bind_heading(
            document.add_heading("不同标题", level=1),
            source_start=target.source_start,
            source_end=target.source_end,
        )


@pytest.mark.parametrize(
    ("member", "old", "new"),
    [
        ("word/numbering.xml", None, None),
        ("word/document.xml", None, None),
        ("word/document.xml", b">" + "数据".encode() + b"<", b">" + "篡改".encode() + b"<"),
        ("word/document.xml", b">1<", b">9<"),
        ("word/document.xml", b"docwen-target-v1:", b"docwen-target-v2:"),
    ],
)
def test_session_reopen_proof_rejects_numbering_caption_ref_and_target_tamper(
    tmp_path: Path,
    member: str,
    old: bytes | None,
    new: bytes | None,
) -> None:
    session, output = _render_resolved_package(tmp_path)
    if member == "word/numbering.xml":
        _mutate_owned_number_format(output, session)
    elif old is None:
        owned_num_id = str(session.projection.num_id("document")).encode()
        _replace_zip_member(
            output,
            member,
            b'w:numId w:val="' + owned_num_id + b'"',
            b'w:numId w:val="99"',
        )
    else:
        assert new is not None
        _replace_zip_member(output, member, old, new)

    with pytest.raises(ValueError):
        session.prove_package(output)


def _render_resolved_package(tmp_path: Path) -> tuple[ResolvedNumberingDocxSession, Path]:
    document = Document()
    session = ResolvedNumberingDocxSession(
        document,
        _port(),
        heading_style_ids={1: "Heading1"},
        heading_style_names={"heading_1": "Heading 1"},
        caption_style_bindings=_caption_bindings(document),
    )
    heading = document.add_heading("2.3 标题", level=1)
    reference = document.add_paragraph()
    caption = document.add_paragraph(style="Table Caption V4")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "b"
    targets = session.port.document.targets
    session.bind_heading(heading, source_start=targets[0].source_start, source_end=targets[0].source_end)
    session.bind_caption(
        caption,
        (table._element,),
        source_start=targets[1].source_start,
        source_end=targets[1].source_end,
        kind="table",
    )
    resolved_reference = session.port.document.references[0]
    session.render_reference(
        reference,
        source_start=resolved_reference.source_start,
        source_end=resolved_reference.source_end,
    )
    output = tmp_path / "resolved-proof.docx"
    document.save(str(output))
    session.write_package(output)
    return session, output


def _replace_zip_member(path: Path, member: str, old: bytes, new: bytes) -> None:
    with ZipFile(path) as package:
        infos = package.infolist()
        data = {info.filename: package.read(info.filename) for info in infos}
    assert old in data[member]
    data[member] = data[member].replace(old, new, 1)
    temporary = path.with_suffix(".rewrite.docx")
    with ZipFile(temporary, "w") as output:
        for info in infos:
            output.writestr(info, data[info.filename])
        known = {info.filename for info in infos}
        for name in sorted(set(data) - known):
            output.writestr(ZipInfo(name), data[name])
    temporary.replace(path)


def _mutate_owned_number_format(path: Path, session: ResolvedNumberingDocxSession) -> None:
    abstract_id = str(session.projection.abstract_id("main"))

    def mutate(payload: bytes) -> bytes:
        root = etree.fromstring(payload)
        candidates = [
            item for item in root.findall(qn("w:abstractNum")) if item.get(qn("w:abstractNumId")) == abstract_id
        ]
        assert len(candidates) == 1
        number_format = candidates[0].find("./" + qn("w:lvl") + "/" + qn("w:numFmt"))
        assert number_format is not None
        assert number_format.get(qn("w:val")) == "decimal"
        number_format.set(qn("w:val"), "upperRoman")
        return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)

    _mutate_zip_member(path, "word/numbering.xml", mutate)


def _mutate_zip_member(path: Path, member: str, mutate: Callable[[bytes], bytes]) -> None:
    with ZipFile(path) as package:
        infos = package.infolist()
        data = {info.filename: package.read(info.filename) for info in infos}
    data[member] = mutate(data[member])
    temporary = path.with_suffix(".rewrite.docx")
    with ZipFile(temporary, "w") as output:
        for info in infos:
            output.writestr(info, data[info.filename])
    temporary.replace(path)
