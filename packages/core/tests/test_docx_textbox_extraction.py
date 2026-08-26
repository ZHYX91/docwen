from __future__ import annotations

import pytest
from docx import Document
from docx.document import Document as DocumentType
from lxml import etree

from docwen_core.docx_parsing.textbox_extraction import extract_textbox_paragraphs
from docwen_core.docx_parsing.xml_ns import NS_W, NS_WPS

pytestmark = pytest.mark.unit


def _append_drawingml_textbox(doc: DocumentType, text: str) -> None:
    outer_para = doc.add_paragraph()._p
    run = etree.SubElement(outer_para, f"{{{NS_W}}}r")
    drawing = etree.SubElement(run, f"{{{NS_W}}}drawing")
    textbox = etree.SubElement(drawing, f"{{{NS_WPS}}}txbx")
    content = etree.SubElement(textbox, f"{{{NS_W}}}txbxContent")
    inner_para = etree.SubElement(content, f"{{{NS_W}}}p")
    inner_run = etree.SubElement(inner_para, f"{{{NS_W}}}r")
    for index, line in enumerate(text.split("\n")):
        if index:
            etree.SubElement(inner_run, f"{{{NS_W}}}br")
        text_element = etree.SubElement(inner_run, f"{{{NS_W}}}t")
        text_element.text = line


def test_distinct_drawingml_textboxes_are_not_lost_to_wrapper_id_reuse() -> None:
    doc = Document()
    expected = [f"Textbox {index}" for index in range(32)]
    for text in [*expected, expected[5]]:
        _append_drawingml_textbox(doc, text)

    extracted = extract_textbox_paragraphs(doc)

    assert [item.text for item in extracted] == [*expected, expected[5]]
    assert extracted[5].anchor_index != extracted[-1].anchor_index


def test_drawingml_textbox_uses_accepted_revision_and_visible_run_projection() -> None:
    doc = Document()
    _append_drawingml_textbox(doc, "placeholder")
    outer_paragraph = doc.paragraphs[-1]._p
    textbox_paragraph = outer_paragraph.find(f".//{{{NS_WPS}}}txbx//{{{NS_W}}}p")
    assert textbox_paragraph is not None
    for child in list(textbox_paragraph):
        textbox_paragraph.remove(child)

    inserted = etree.SubElement(textbox_paragraph, f"{{{NS_W}}}ins")
    inserted_run = etree.SubElement(inserted, f"{{{NS_W}}}r")
    inserted_text = etree.SubElement(inserted_run, f"{{{NS_W}}}t")
    inserted_text.text = "accepted"
    deleted = etree.SubElement(textbox_paragraph, f"{{{NS_W}}}del")
    deleted_run = etree.SubElement(deleted, f"{{{NS_W}}}r")
    deleted_text = etree.SubElement(deleted_run, f"{{{NS_W}}}delText")
    deleted_text.text = "deleted"

    extracted = extract_textbox_paragraphs(doc)

    assert [item.text for item in extracted] == ["accepted"]


def test_drawingml_textbox_preserves_manual_line_breaks() -> None:
    doc = Document()
    _append_drawingml_textbox(doc, "Describe the outcome:\n- specific")

    extracted = extract_textbox_paragraphs(doc)

    assert [item.text for item in extracted] == ["Describe the outcome:\n- specific"]


def test_body_textbox_reports_its_top_level_anchor_index() -> None:
    doc = Document()
    doc.add_paragraph("Before")
    _append_drawingml_textbox(doc, "Anchored textbox")
    doc.add_paragraph("After")

    extracted = extract_textbox_paragraphs(doc)

    assert [(item.text, item.anchor_index, item.source) for item in extracted] == [("Anchored textbox", 1, "textbox")]
