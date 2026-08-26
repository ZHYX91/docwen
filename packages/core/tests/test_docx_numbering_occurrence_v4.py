"""Executable disabled ID-less caption occurrence authority gates."""

from __future__ import annotations

import hashlib
from pathlib import Path
from typing import Any, cast
from zipfile import ZipFile

import lxml.etree as etree
import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docwen_core._docx_semantics_v3_model import (
    CaptionStyleBindingV3,
    CaptionStyleKeyV3,
    DocxSemanticsV3Error,
)
from docwen_core._docx_semantics_v3_package import (
    inject_custom_xml_parts,
    read_owned_map_parts,
    verify_custom_xml_support,
)
from docwen_core.docx_numbering_occurrence import (
    NUMBERING_OCCURRENCE_MAP_NAMESPACE,
    derive_numbering_occurrence,
    numbering_occurrence_map_xml,
    parse_numbering_occurrence_map,
    prove_numbering_occurrence_sdt,
    wrap_numbering_occurrence,
)
from docwen_core.docx_resolved_numbering import (
    ResolvedNumberingDocxError,
    ResolvedNumberingDocxSession,
)
from docwen_core.models.resolved_numbering import (
    CaptionMaterialization,
    HeadingCounterSegment,
    HeadingDefinition,
    HeadingInstance,
    HeadingLevelDefinition,
    HeadingListMaterialization,
    NumberingExportPlanEnvelope,
    NumberingTarget,
    ResolvedCitation,
    ResolvedCitationItem,
    ResolvedDocument,
    ResolvedDocumentEnvelope,
    ResolvedDocumentTarget,
    ResolvedNumberingPlan,
    ResolvedNumberingPort,
    ResolvedReference,
)

pytestmark = pytest.mark.unit

SOURCE_SHA = hashlib.sha256(b"Table: authored\n|a|\n|-|\n|b|\n").hexdigest()
PLAN_SHA = hashlib.sha256(b"plan").hexdigest()


def _sha(value: str) -> str:
    return hashlib.sha256(value.encode()).hexdigest()


def _caption_bindings(document: Any) -> tuple[CaptionStyleBindingV3, ...]:
    output: list[CaptionStyleBindingV3] = []
    for key, style_id, name in (
        ("figure_caption", "DWFigureCaption", "Figure Caption V4"),
        ("table_caption", "DWTableCaption", "Table Caption V4"),
        ("equation_caption", "DWEquationCaption", "Equation Caption V4"),
        ("code_block_caption", "DWCodeCaption", "Code Caption V4"),
    ):
        style = document.styles.add_style(style_id, WD_STYLE_TYPE.PARAGRAPH)
        style.name = name
        output.append(CaptionStyleBindingV3(cast(CaptionStyleKeyV3, key), style.style_id, style.name))
    return tuple(output)


def _rich_port(*, enabled: bool = False) -> ResolvedNumberingPort:
    token = "@[[#^head-a|标题]]"
    citation_token = "@cite"
    caption_text = f"literal {token} then {token} and {citation_token}"
    source = f"# Heading ^head-a\n\nTable: {caption_text}\n|a|\n|-|\n|b|\n"
    heading_end = source.index("\n")
    table_start = source.index("Table:")
    reference_start = source.index(token, source.index(token) + 1)
    citation_start = source.index(citation_token)
    heading = ResolvedDocumentTarget(0, heading_end, _sha(source[:heading_end]), "heading", "head-a", 1, "Heading")
    table = ResolvedDocumentTarget(
        table_start,
        len(source) - 1,
        _sha(source[table_start:-1]),
        "table",
        "table-a" if enabled else None,
        None,
        caption_text,
    )
    reference = ResolvedReference(
        reference_start,
        reference_start + len(token),
        _sha(token),
        token,
        0,
        heading_end,
        "heading",
        "head-a",
        "1",
        "标题",
    )
    citation = ResolvedCitation(
        citation_start,
        citation_start + len(citation_token),
        _sha(citation_token),
        citation_token,
        "narrative",
        "cluster-a",
        (ResolvedCitationItem("cite", "reference-record:98", _sha("record"), "Citation"),),
        "Citation",
    )
    definition = HeadingDefinition(
        "main",
        (HeadingLevelDefinition(1, 1, "arabic_half", (HeadingCounterSegment(1, "arabic_half"),), "space", None),),
    )
    caption_materialization = (
        CaptionMaterialization(
            "simple_seq",
            "Table",
            "arabic_half",
            "continue",
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            "1",
            "Table",
            " ",
        )
        if enabled
        else None
    )
    plan = ResolvedNumberingPlan(
        (definition,),
        (HeadingInstance("document", "main", ()),),
        (
            NumberingTarget(
                0, heading_end, "heading", True, "head-a", "1", HeadingListMaterialization("main", "document", 1)
            ),
            NumberingTarget(
                table_start,
                len(source) - 1,
                "table",
                enabled,
                "table-a" if enabled else None,
                "1" if enabled else None,
                caption_materialization,
            ),
        ),
    )
    source_sha = _sha(source)
    plan_sha = "a" * 64
    return ResolvedNumberingPort(
        ResolvedDocumentEnvelope(
            "source",
            source_sha,
            plan_sha,
            ResolvedDocument(source, (heading, table), (reference,), (), (citation,), ()),
        ),
        NumberingExportPlanEnvelope("source", source_sha, plan_sha, plan),
    )


def _two_table_port() -> ResolvedNumberingPort:
    source = "Table: One\n|a|\n|-|\n|1|\n\nTable: Two\n|a|\n|-|\n|2|\n"
    separator = source.index("\n\n")
    ranges = ((0, separator), (separator + 2, len(source) - 1))
    targets = tuple(
        ResolvedDocumentTarget(start, end, _sha(source[start:end]), "table", None, None, name)
        for (start, end), name in zip(ranges, ("One", "Two"), strict=True)
    )
    plan_targets = tuple(NumberingTarget(start, end, "table", False, None, None, None) for start, end in ranges)
    source_sha = _sha(source)
    plan_sha = "b" * 64
    return ResolvedNumberingPort(
        ResolvedDocumentEnvelope(
            "source",
            source_sha,
            plan_sha,
            ResolvedDocument(source, targets, (), (), (), ()),
        ),
        NumberingExportPlanEnvelope(
            "source",
            source_sha,
            plan_sha,
            ResolvedNumberingPlan((), (), plan_targets),
        ),
    )


def _identity(*, start: int = 0, end: int = 30):
    return derive_numbering_occurrence(
        source_sha256=SOURCE_SHA,
        source_start=start,
        source_end=end,
        kind="table",
        plan_sha256=PLAN_SHA,
    )


def _wrapped_table_document():
    document = Document()
    caption = document.add_paragraph("authored")
    caption.style = "Caption"
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "cell"
    identity = _identity()
    wrap_numbering_occurrence(caption._p, (table._element,), identity)
    sdt = next(item for item in cast(Any, document.element).body if item.tag == qn("w:sdt"))
    return document, sdt, identity


def test_occurrence_digest_and_map_are_exact_and_round_trip() -> None:
    first = _identity()
    second = derive_numbering_occurrence(
        source_sha256=SOURCE_SHA,
        source_start=31,
        source_end=60,
        kind="figure",
        plan_sha256=PLAN_SHA,
    )

    data = numbering_occurrence_map_xml([first, second])
    assert data.startswith(b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n')
    assert data.endswith(b"\n") and data.count(b"\n") == 2
    root = etree.fromstring(data)
    assert parse_numbering_occurrence_map(root) == [first, second]
    occurrence = next(iter(root))
    assert tuple(occurrence.attrib) == (
        "tag",
        "source_sha256",
        "source_start",
        "source_end",
        "kind",
        "enabled",
        "target_id",
        "derived_number",
        "plan_sha256",
        "sha256",
    )
    assert occurrence.get("enabled") == "false"
    assert occurrence.get("target_id") == occurrence.get("derived_number") == ""


def test_occurrence_sdt_has_exact_caption_object_order_and_no_semantic_target() -> None:
    _document, sdt, identity = _wrapped_table_document()

    caption, logical_object = prove_numbering_occurrence_sdt(
        sdt,
        identity,
        caption_style_id="Caption",
    )

    assert caption.tag == qn("w:p")
    assert logical_object.tag == qn("w:tbl")
    assert not list(sdt.iter(qn("w:bookmarkStart")))
    assert not list(sdt.iter(qn("w:instrText")))
    assert all(not (item.get(qn("w:val")) or "").startswith("docwen-target-v1:") for item in sdt.iter(qn("w:tag")))


def test_reversed_table_pair_fails_closed() -> None:
    _document, sdt, identity = _wrapped_table_document()
    content = sdt.find(qn("w:sdtContent"))
    assert content is not None
    first, second = tuple(content)
    content.remove(first)
    content.remove(second)
    content.extend((second, first))

    with pytest.raises(DocxSemanticsV3Error, match="caption slot"):
        prove_numbering_occurrence_sdt(sdt, identity, caption_style_id="Caption")


@pytest.mark.parametrize("mutation", ["bookmark", "seq", "style"])
def test_occurrence_physical_mutations_fail_closed(mutation: str) -> None:
    _document, sdt, identity = _wrapped_table_document()
    content = sdt.find(qn("w:sdtContent"))
    assert content is not None
    caption = content[0]
    if mutation == "bookmark":
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), "1")
        start.set(qn("w:name"), "DW_T_0123456789abcdef0123456789abcdef012")
        caption.append(start)
    elif mutation == "seq":
        run = OxmlElement("w:r")
        instruction = OxmlElement("w:instrText")
        instruction.text = " SEQ Table \\* ARABIC "
        run.append(instruction)
        caption.append(run)
    else:
        caption.find(f"{qn('w:pPr')}/{qn('w:pStyle')}").set(qn("w:val"), "Wrong")

    with pytest.raises(DocxSemanticsV3Error):
        prove_numbering_occurrence_sdt(sdt, identity, caption_style_id="Caption")


def _render_rich_package(
    tmp_path: Path,
    *,
    enabled: bool = False,
) -> tuple[ResolvedNumberingDocxSession, Path]:
    document = Document()
    port = _rich_port(enabled=enabled)
    session = ResolvedNumberingDocxSession(
        document,
        port,
        heading_style_ids={1: "Heading1"},
        heading_style_names={"heading_1": "Heading 1"},
        caption_style_bindings=_caption_bindings(document),
    )
    heading = document.add_heading("Heading", level=1)
    caption = document.add_paragraph(style="Table Caption V4")
    table = document.add_table(rows=1, cols=1)
    session.bind_heading(
        heading,
        source_start=port.document.targets[0].source_start,
        source_end=port.document.targets[0].source_end,
    )
    reference = port.document.references[0]
    caption.add_run(f"literal {reference.authored_token} then ")
    session.render_reference(caption, source_start=reference.source_start, source_end=reference.source_end)
    caption.add_run(" and ")
    citation = port.document.citations[0]
    session.render_citation(caption, source_start=citation.source_start, source_end=citation.source_end)
    target = port.document.targets[1]
    session.bind_rendered_caption(
        caption,
        (table._element,),
        source_start=target.source_start,
        source_end=target.source_end,
        kind="table",
    )
    output = tmp_path / ("rich-enabled.docx" if enabled else "rich-disabled.docx")
    session.write_package(output)
    return session, output


def test_rich_caption_preserves_literal_duplicate_and_source_order(tmp_path: Path) -> None:
    session, output = _render_rich_package(tmp_path)
    session.prove_package(output)
    with ZipFile(output) as package:
        root = etree.fromstring(package.read("word/document.xml"))
    literal = session.port.document.references[0].authored_token
    assert sum(literal in (item.text or "") for item in root.iter(qn("w:t"))) == 1
    tags = [
        item.get(qn("w:val"), "")
        for item in root.iter(qn("w:tag"))
        if (item.get(qn("w:val"), "")).startswith(("docwen-ref-occurrence-v1:", "docwen-citation-occurrence-v1:"))
    ]
    assert len(tags) == 2


def test_rich_caption_number_bookmark_encloses_only_seq_fields(tmp_path: Path) -> None:
    session, output = _render_rich_package(tmp_path, enabled=True)
    session.prove_package(output)
    with ZipFile(output) as package:
        root = etree.fromstring(package.read("word/document.xml"))
    caption = next(
        paragraph
        for paragraph in root.iter(qn("w:p"))
        if (style := paragraph.find(f"{qn('w:pPr')}/{qn('w:pStyle')}")) is not None
        and style.get(qn("w:val")) == "DWTableCaption"
    )
    children = [item for item in caption if item.tag != qn("w:pPr")]
    start = next(index for index, item in enumerate(children) if item.tag == qn("w:bookmarkStart"))
    end = next(index for index, item in enumerate(children) if item.tag == qn("w:bookmarkEnd"))
    assert "SEQ Table" in "".join(item.text or "" for child in children[start:end] for item in child.iter())
    assert all(item.tag != qn("w:sdt") for item in children[start : end + 1])
    assert any(item.tag == qn("w:sdt") for item in children[end + 1 :])


def test_rich_caption_rejects_carrier_at_wrong_identical_token_occurrence() -> None:
    document = Document()
    port = _rich_port()
    session = ResolvedNumberingDocxSession(
        document,
        port,
        heading_style_ids={1: "Heading1"},
        heading_style_names={"heading_1": "Heading 1"},
        caption_style_bindings=_caption_bindings(document),
    )
    caption = document.add_paragraph(style="Table Caption V4")
    table = document.add_table(rows=1, cols=1)
    reference = port.document.references[0]
    caption.add_run("literal ")
    session.render_reference(caption, source_start=reference.source_start, source_end=reference.source_end)
    caption.add_run(f" then {reference.authored_token} and ")
    citation = port.document.citations[0]
    session.render_citation(caption, source_start=citation.source_start, source_end=citation.source_end)
    target = port.document.targets[1]

    with pytest.raises(ResolvedNumberingDocxError, match="source-relative projection"):
        session.bind_rendered_caption(
            caption,
            (table._element,),
            source_start=target.source_start,
            source_end=target.source_end,
            kind="table",
        )


def test_plain_bind_rejects_nested_authority_before_claiming_target() -> None:
    document = Document()
    port = _rich_port()
    session = ResolvedNumberingDocxSession(
        document,
        port,
        heading_style_ids={1: "Heading1"},
        heading_style_names={"heading_1": "Heading 1"},
        caption_style_bindings=_caption_bindings(document),
    )
    caption = document.add_paragraph(style="Table Caption V4")
    table = document.add_table(rows=1, cols=1)
    target = port.document.targets[1]
    with pytest.raises(ResolvedNumberingDocxError, match="bind_rendered_caption"):
        session.bind_caption(
            caption,
            (table._element,),
            source_start=target.source_start,
            source_end=target.source_end,
            kind="table",
        )
    with pytest.raises(ResolvedNumberingDocxError, match="unbound"):
        session.bind_rendered_caption(
            caption,
            (table._element,),
            source_start=target.source_start,
            source_end=target.source_end,
            kind="table",
        )


@pytest.mark.parametrize("mutation", ["wrong_paragraph", "forged_plain"])
def test_rich_caption_rejects_wrong_paragraph_or_forged_plain_projection(mutation: str) -> None:
    document = Document()
    port = _rich_port()
    session = ResolvedNumberingDocxSession(
        document,
        port,
        heading_style_ids={1: "Heading1"},
        heading_style_names={"heading_1": "Heading 1"},
        caption_style_bindings=_caption_bindings(document),
    )
    caption = document.add_paragraph(style="Table Caption V4")
    table = document.add_table(rows=1, cols=1)
    reference_paragraph = document.add_paragraph() if mutation == "wrong_paragraph" else caption
    reference = port.document.references[0]
    caption.add_run("forged plain " if mutation == "forged_plain" else f"literal {reference.authored_token} then ")
    session.render_reference(
        reference_paragraph,
        source_start=reference.source_start,
        source_end=reference.source_end,
    )
    caption.add_run(" and ")
    citation = port.document.citations[0]
    session.render_citation(caption, source_start=citation.source_start, source_end=citation.source_end)
    target = port.document.targets[1]
    expected = "another paragraph" if mutation == "wrong_paragraph" else "source-relative projection"
    with pytest.raises(ResolvedNumberingDocxError, match=expected):
        session.bind_rendered_caption(
            caption,
            (table._element,),
            source_start=target.source_start,
            source_end=target.source_end,
            kind="table",
        )


def _render_two_occurrences(tmp_path: Path) -> tuple[ResolvedNumberingDocxSession, Path]:
    document = Document()
    port = _two_table_port()
    session = ResolvedNumberingDocxSession(
        document,
        port,
        heading_style_ids={1: "Heading1"},
        heading_style_names={"heading_1": "Heading 1"},
        caption_style_bindings=_caption_bindings(document),
    )
    for target in port.document.targets:
        caption = document.add_paragraph(style="Table Caption V4")
        table = document.add_table(rows=1, cols=1)
        session.bind_caption(
            caption,
            (table._element,),
            source_start=target.source_start,
            source_end=target.source_end,
            kind="table",
        )
    output = tmp_path / "two-occurrences.docx"
    session.write_package(output)
    return session, output


@pytest.mark.parametrize("mutation", ["swap", "empty_decoy"])
def test_occurrence_reopen_proof_binds_exact_source_order_and_caption(
    tmp_path: Path,
    mutation: str,
) -> None:
    session, output = _render_two_occurrences(tmp_path)
    with ZipFile(output) as package:
        infos = package.infolist()
        data = {item.filename: package.read(item.filename) for item in infos}
    root = etree.fromstring(data["word/document.xml"])
    body = root.find(qn("w:body"))
    assert body is not None

    def occurrence_tag(element: Any) -> str:
        tag_element = element.find(f"{qn('w:sdtPr')}/{qn('w:tag')}")
        assert tag_element is not None
        value = tag_element.get(qn("w:val"))
        assert value is not None
        return value

    occurrences = [
        item
        for item in body
        if item.tag == qn("w:sdt") and occurrence_tag(item).startswith("docwen-numbering-occurrence-v1:")
    ]
    assert len(occurrences) == 2
    if mutation == "swap":
        first_tag = occurrences[0].find(f"{qn('w:sdtPr')}/{qn('w:tag')}")
        second_tag = occurrences[1].find(f"{qn('w:sdtPr')}/{qn('w:tag')}")
        assert first_tag is not None and second_tag is not None
        first = occurrence_tag(occurrences[0])
        second = occurrence_tag(occurrences[1])
        first_tag.set(qn("w:val"), second)
        second_tag.set(qn("w:val"), first)
    else:
        decoy = OxmlElement("w:sdt")
        properties = OxmlElement("w:sdtPr")
        tag = OxmlElement("w:tag")
        original_tag = occurrences[0].find(f"{qn('w:sdtPr')}/{qn('w:tag')}")
        assert original_tag is not None
        tag.set(qn("w:val"), occurrence_tag(occurrences[0]))
        properties.append(tag)
        decoy.extend((properties, OxmlElement("w:sdtContent")))
        body.insert(len(body) - 1, decoy)
    data["word/document.xml"] = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    replacement = output.with_suffix(".mutated.docx")
    with ZipFile(replacement, "w") as package:
        for info in infos:
            package.writestr(info, data[info.filename])
    replacement.replace(output)

    with pytest.raises(ResolvedNumberingDocxError, match="physical order/cardinality"):
        session.prove_package(output)


def test_occurrence_custom_xml_uses_independent_canonical_trio(tmp_path: Path) -> None:
    document, _sdt, identity = _wrapped_table_document()
    output = tmp_path / "occurrence.docx"
    document.save(str(output))
    map_bytes = numbering_occurrence_map_xml([identity])

    inject_custom_xml_parts(output, [(NUMBERING_OCCURRENCE_MAP_NAMESPACE, map_bytes)])

    with ZipFile(output) as package:
        owned = read_owned_map_parts(package)
        assert NUMBERING_OCCURRENCE_MAP_NAMESPACE in owned
        item_number, root = owned[NUMBERING_OCCURRENCE_MAP_NAMESPACE]
        verify_custom_xml_support(package, item_number, NUMBERING_OCCURRENCE_MAP_NAMESPACE)
        assert parse_numbering_occurrence_map(root) == [identity]


def test_occurrence_map_mutations_fail_closed() -> None:
    identity = _identity()
    root = etree.fromstring(numbering_occurrence_map_xml([identity]))
    root[0].set("derived_number", "1")
    with pytest.raises(DocxSemanticsV3Error, match="canonical"):
        parse_numbering_occurrence_map(root)


def test_occurrence_ranges_must_not_overlap() -> None:
    first = _identity(start=0, end=30)
    second = derive_numbering_occurrence(
        source_sha256=SOURCE_SHA,
        source_start=29,
        source_end=40,
        kind="figure",
        plan_sha256=PLAN_SHA,
    )
    with pytest.raises(DocxSemanticsV3Error, match="overlap"):
        numbering_occurrence_map_xml([first, second])
