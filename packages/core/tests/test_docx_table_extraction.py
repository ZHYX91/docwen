"""DOCX table projection through the shared semantic grid."""

from __future__ import annotations

import pytest
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docwen_core.docx_parsing.table_extraction import (
    DocxTableGeometryError,
    build_docx_table_semantic_grid,
    markdown_table_lines,
    render_docx_table_rows,
)

pytestmark = pytest.mark.contract


def _plain_text(cell, _row: int, _column: int) -> str:
    return "".join(node.text or "" for node in cell.iter(qn("w:t"))).replace("|", "\\|")


def test_docx_table_markdown_preserves_rows_columns_and_escaped_pipes() -> None:
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "列|甲"
    table.cell(0, 1).text = "列乙"
    table.cell(1, 0).text = "值甲"
    table.cell(1, 1).text = "值乙"

    rendered = render_docx_table_rows(
        table._tbl,
        cell_text_resolver=_plain_text,
    )
    markdown = "\n".join(markdown_table_lines(rendered))

    assert markdown == "| 列\\|甲 | 列乙 |\n| --- | --- |\n| 值甲 | 值乙 |"


def test_shared_table_geometry_marks_a_horizontal_merge_once() -> None:
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "合并表头"
    table.cell(1, 0).text = "甲"
    table.cell(1, 1).text = "乙"

    rendered = render_docx_table_rows(
        table._tbl,
        cell_text_resolver=_plain_text,
        strategy="marker",
    )

    assert markdown_table_lines(rendered) == [
        "| 合并表头 | < |",
        "| --- | --- |",
        "| 甲 | 乙 |",
    ]


def test_structural_markdown_projection_places_delimiter_after_all_headers() -> None:
    rendered = [
        ["Region", "Sales", "<"],
        ["Quarter", "Q1", "Q2"],
        ["North", "10", "12"],
    ]

    assert markdown_table_lines(rendered, header_rows=2, header_columns=1) == [
        "| Region | Sales | < |",
        "| Quarter | Q1 | Q2 |",
        "| --- || --- | --- |",
        "| North | 10 | 12 |",
    ]


def test_structural_row_rendering_escapes_literal_merge_markers_only() -> None:
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "literal"
    table.cell(0, 1).text = "<"
    table.cell(1, 0).merge(table.cell(1, 1)).text = "merged"

    rendered = render_docx_table_rows(
        table._tbl,
        cell_text_resolver=_plain_text,
        strategy="marker",
        escape_literal_merge_markers=True,
    )

    assert rendered == [["literal", "\\<"], ["merged", "<"]]


def test_fill_row_rendering_escapes_repeated_literal_merge_markers() -> None:
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "^"

    rendered = render_docx_table_rows(
        table._tbl,
        cell_text_resolver=_plain_text,
        strategy="fill",
        escape_literal_merge_markers=True,
    )

    assert rendered == [["\\^", "\\^"]]


def test_shared_table_geometry_preserves_a_rectangular_horizontal_vertical_merge() -> None:
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(1, 1)).text = "combined"

    grid = build_docx_table_semantic_grid(table._tbl, cell_text_resolver=_plain_text)

    assert (grid[0][0].rowspan, grid[0][0].colspan, grid[0][0].anchor_text) == (2, 2, "combined")
    assert all(grid[row][column].anchor_text == "combined" for row in range(2) for column in range(2))


@pytest.mark.parametrize("wrapper", ["customXml", "sdt"])
@pytest.mark.parametrize("wrapped_kind", ["row", "cell"])
def test_shared_table_geometry_walks_only_legal_ordered_wrappers(wrapper: str, wrapped_kind: str) -> None:
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "wrapped"
    if wrapped_kind == "row":
        _wrap_ooxml_child(table._tbl, table.rows[0]._tr, wrapper=wrapper)
    else:
        _wrap_ooxml_child(table.rows[0]._tr, table.cell(0, 0)._tc, wrapper=wrapper)

    grid = build_docx_table_semantic_grid(table._tbl, cell_text_resolver=_plain_text)

    assert len(grid) == len(grid[0]) == 1
    assert grid[0][0].anchor_text == "wrapped"


def test_shared_table_geometry_does_not_collect_cells_from_a_nested_table() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=1)
    row = table.rows[0]._tr
    row.remove(table.cell(0, 0)._tc)
    wrapper = OxmlElement("w:customXml")
    nested_table = OxmlElement("w:tbl")
    nested_row = OxmlElement("w:tr")
    nested_row.append(OxmlElement("w:tc"))
    nested_table.append(nested_row)
    wrapper.append(nested_table)
    row.append(wrapper)

    with pytest.raises(DocxTableGeometryError) as rejected:
        build_docx_table_semantic_grid(table._tbl, cell_text_resolver=_plain_text)

    assert rejected.value.code == "cells_missing"


def test_shared_table_geometry_honors_leading_and_trailing_grid_offsets() -> None:
    document = Document()
    table = document.add_table(rows=3, cols=2)
    right_cell = table.cell(1, 1)
    left_cell = table.cell(2, 0)
    trailing_cell = table.cell(2, 1)._tc
    _remove_cell(table.rows[1]._tr, table.cell(1, 0)._tc)
    right_cell.text = "right"
    _set_row_grid_offset(table.rows[1]._tr, "gridBefore", "1")
    _remove_cell(table.rows[2]._tr, trailing_cell)
    left_cell.text = "left"
    _set_row_grid_offset(table.rows[2]._tr, "gridAfter", "1")

    grid = build_docx_table_semantic_grid(table._tbl, cell_text_resolver=_plain_text)

    assert [cell.anchor_text for cell in grid[1]] == ["", "right"]
    assert [cell.anchor_text for cell in grid[2]] == ["left", ""]


def test_shared_table_geometry_matches_vertical_merge_after_grid_before() -> None:
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 1).text = "vertical"
    _set_vertical_merge(table.cell(0, 1)._tc, "restart")
    continuation = table.cell(1, 1)._tc
    _remove_cell(table.rows[1]._tr, table.cell(1, 0)._tc)
    _set_row_grid_offset(table.rows[1]._tr, "gridBefore", "1")
    _set_vertical_merge(continuation, "continue")

    grid = build_docx_table_semantic_grid(table._tbl, cell_text_resolver=_plain_text)

    assert (grid[0][1].rowspan, grid[0][1].anchor_text) == (2, "vertical")
    assert (grid[1][1].is_covered, grid[1][1].anchor_text) == (True, "vertical")


def test_shared_table_geometry_infers_an_omitted_table_grid() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "inferred"
    table._tbl.remove(table._tbl.tblGrid)

    grid = build_docx_table_semantic_grid(table._tbl, cell_text_resolver=_plain_text)

    assert len(grid) == len(grid[0]) == 1
    assert grid[0][0].anchor_text == "inferred"


def test_shared_table_geometry_augments_a_short_table_grid_for_grid_span() -> None:
    document = Document()
    table = _table_with_grid_span(document, "2")
    table.cell(0, 0).text = "wide"

    grid = build_docx_table_semantic_grid(table._tbl, cell_text_resolver=_plain_text)

    assert len(grid[0]) == 2
    assert (grid[0][0].colspan, grid[0][0].anchor_text) == (2, "wide")


def test_shared_table_geometry_applies_offset_against_an_augmented_grid() -> None:
    document = Document()
    table = document.add_table(rows=2, cols=1)
    table.cell(1, 0).text = "offset"
    grid_span = OxmlElement("w:gridSpan")
    grid_span.set(qn("w:val"), "3")
    table.cell(0, 0)._tc.get_or_add_tcPr().append(grid_span)
    _set_row_grid_offset(table.rows[1]._tr, "gridBefore", "2")

    grid = build_docx_table_semantic_grid(table._tbl, cell_text_resolver=_plain_text)

    assert len(grid[0]) == 3
    assert [cell.anchor_text for cell in grid[1]] == ["", "", "offset"]


def test_shared_table_geometry_ignores_conflicting_grid_offsets() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "kept"
    _set_row_grid_offset(table.rows[0]._tr, "gridBefore", "2")
    _set_row_grid_offset(table.rows[0]._tr, "gridAfter", "1")

    grid = build_docx_table_semantic_grid(table._tbl, cell_text_resolver=_plain_text)

    assert len(grid[0]) == 1
    assert grid[0][0].anchor_text == "kept"


@pytest.mark.parametrize(
    ("prepare", "code"),
    [
        (lambda document: document.add_table(rows=0, cols=0), "rows_missing"),
        (lambda document: _table_row_without_cells(document), "cells_missing"),
        (lambda document: _table_with_grid_span(document, "not-an-integer"), "grid_span_invalid"),
        (lambda document: _table_with_grid_span(document, "0"), "grid_span_invalid"),
        (lambda document: _table_with_grid_span(document, "-1"), "grid_span_invalid"),
        (lambda document: _table_with_grid_span(document, None), "grid_span_invalid"),
        (lambda document: _table_with_vertical_merge(document, "orphan"), "vmerge_invalid"),
        (lambda document: _table_with_vertical_merge(document, "mismatched"), "vmerge_invalid"),
        (lambda document: _table_with_vertical_merge(document, "invalid_enum"), "vmerge_invalid"),
        (lambda document: _table_with_vertical_merge(document, "uppercase"), "vmerge_invalid"),
        (lambda document: _table_with_grid_offset(document, "gridBefore", None), "grid_offset_invalid"),
        (lambda document: _table_with_grid_offset(document, "gridAfter", "-1"), "grid_offset_invalid"),
        (lambda document: _table_with_grid_offset(document, "gridBefore", "not-an-integer"), "grid_offset_invalid"),
    ],
)
def test_shared_table_geometry_rejects_malformed_ooxml(prepare, code: str) -> None:
    document = Document()
    table = prepare(document)

    with pytest.raises(DocxTableGeometryError) as rejected:
        build_docx_table_semantic_grid(table._tbl, cell_text_resolver=_plain_text)

    assert rejected.value.code == code


def _table_row_without_cells(document: DocumentType):
    table = document.add_table(rows=1, cols=1)
    row = table.rows[0]._tr
    row.remove(table.cell(0, 0)._tc)
    return table


def _table_with_grid_span(document: DocumentType, value: str | None):
    table = document.add_table(rows=1, cols=1)
    grid_span = OxmlElement("w:gridSpan")
    if value is not None:
        grid_span.set(qn("w:val"), value)
    table.cell(0, 0)._tc.get_or_add_tcPr().append(grid_span)
    return table


def _wrap_ooxml_child(parent, child, *, wrapper: str) -> None:
    index = parent.index(child)
    parent.remove(child)
    container = OxmlElement(f"w:{wrapper}")
    if wrapper == "sdt":
        content = OxmlElement("w:sdtContent")
        content.append(child)
        container.append(content)
    else:
        container.append(child)
    parent.insert(index, container)


def _remove_cell(row, cell) -> None:
    row.remove(cell)


def _set_vertical_merge(cell, value: str | None) -> None:
    vertical_merge = OxmlElement("w:vMerge")
    if value is not None:
        vertical_merge.set(qn("w:val"), value)
    cell.get_or_add_tcPr().append(vertical_merge)


def _set_row_grid_offset(row, name: str, value: str | None) -> None:
    offset = OxmlElement(f"w:{name}")
    if value is not None:
        offset.set(qn("w:val"), value)
    row.get_or_add_trPr().append(offset)


def _table_with_vertical_merge(document: DocumentType, malformation: str):
    if malformation == "mismatched":
        table = document.add_table(rows=2, cols=2)
        first_row = table.rows[0]._tr
        anchor = table.cell(0, 0)._tc
        _remove_cell(first_row, table.cell(0, 1)._tc)
        grid_span = OxmlElement("w:gridSpan")
        grid_span.set(qn("w:val"), "2")
        anchor.get_or_add_tcPr().append(grid_span)
        _set_vertical_merge(anchor, "restart")
        _set_vertical_merge(table.cell(1, 0)._tc, "continue")
        return table
    table = document.add_table(rows=1, cols=1)
    _set_vertical_merge(
        table.cell(0, 0)._tc,
        {"invalid_enum": "bogus", "uppercase": "RESTART"}.get(malformation),
    )
    return table


def _table_with_grid_offset(document: DocumentType, name: str, value: str | None):
    table = document.add_table(rows=1, cols=1)
    _set_row_grid_offset(table.rows[0]._tr, name, value)
    return table
