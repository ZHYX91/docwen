"""Post-review hostile package and authenticated-group proofs for semantics v3."""

from __future__ import annotations

import hashlib
import re
import uuid
from pathlib import Path
from zipfile import ZipFile

import lxml.etree as etree
import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docwen_core.docx_semantics_v3 import (
    ANCHOR_TAG_PREFIX,
    SOFT_REFERENCE_TAG_PREFIX,
    CaptionStyleBindingV3,
    DocxSemanticsV3Error,
    DocxSemanticsV3Recovery,
    DocxSemanticsV3Session,
    append_complex_field,
)

pytestmark = pytest.mark.contract


def test_soft_reference_physical_order_and_direct_paragraph_owner_are_proven(
    tmp_path: Path,
) -> None:
    session, output = _write_two_soft_references(tmp_path)

    def swap(root) -> None:
        sdts = [item for item in root.iter(qn("w:sdt")) if (_sdt_tag(item) or "").startswith(SOFT_REFERENCE_TAG_PREFIX)]
        assert len(sdts) == 2 and sdts[0].getparent() is sdts[1].getparent()
        parent = sdts[0].getparent()
        first_index = parent.index(sdts[0])
        second_index = parent.index(sdts[1])
        parent.remove(sdts[1])
        parent.insert(first_index, sdts[1])
        parent.remove(sdts[0])
        parent.insert(second_index, sdts[0])

    _mutate_xml_member(output, "word/document.xml", swap)
    with pytest.raises(DocxSemanticsV3Error, match="physical order"):
        session.prove_package(output)

    _session, nested = _write_two_soft_references(tmp_path, name="nested-soft.docx")

    def nest(root) -> None:
        sdt = next(
            item for item in root.iter(qn("w:sdt")) if (_sdt_tag(item) or "").startswith(SOFT_REFERENCE_TAG_PREFIX)
        )
        run = OxmlElement("w:r")
        parent = sdt.getparent()
        parent.replace(sdt, run)
        run.append(sdt)

    _mutate_xml_member(nested, "word/document.xml", nest)
    with pytest.raises(DocxSemanticsV3Error, match="direct paragraph child"):
        DocxSemanticsV3Recovery.load(nested, Document(str(nested)))


def test_soft_reference_map_rejects_overlapping_ranges(tmp_path: Path) -> None:
    session, output = _write_two_soft_references(tmp_path)
    item_name, props_name = _owned_item_and_props(output, "document-soft-reference-map")
    with ZipFile(output) as package:
        item = package.read(item_name)
        props = package.read(props_name)
    root = etree.fromstring(item)
    records = list(root)
    record = records[1]
    authored_token = record.get("authored_token") or ""
    source_sha256 = record.get("source_sha256") or ""
    record.set("source_start", "8")
    record.set("source_end", str(8 + len(authored_token)))
    source_end = record.get("source_end") or ""
    preimage = "docwen-soft-ref-map-v1\0" + source_sha256 + "\08\0" + source_end + "\0" + authored_token
    record.set("tag", f"docwen-soft-ref-v1:{hashlib.sha256(preimage.encode()).hexdigest()[:32]}")
    mutated = b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n' + etree.tostring(root) + b"\n"
    _replace_zip_members(output, {item_name: mutated, props_name: _rebind_props_uuid(props, mutated)})
    with pytest.raises(DocxSemanticsV3Error, match="overlap"):
        session.prove_package(output)


def test_unmapped_owned_anchor_and_multi_paragraph_paragraph_anchor_fail_closed(
    tmp_path: Path,
) -> None:
    session, output = _write_paragraph_anchor(tmp_path)

    def add_unmapped(root) -> None:
        body = root.find(f".//{qn('w:body')}")
        fake = OxmlElement("w:sdt")
        props = OxmlElement("w:sdtPr")
        tag = OxmlElement("w:tag")
        tag.set(qn("w:val"), f"{ANCHOR_TAG_PREFIX}{'0' * 32}")
        props.append(tag)
        content = OxmlElement("w:sdtContent")
        content.append(OxmlElement("w:p"))
        fake.extend((props, content))
        body.insert(len(body) - 1, fake)

    _mutate_xml_member(output, "word/document.xml", add_unmapped)
    with pytest.raises(DocxSemanticsV3Error, match="no authenticated map record"):
        session.prove_package(output)

    second_session, second = _write_paragraph_anchor(tmp_path, name="two-paragraphs.docx")

    def add_paragraph(root) -> None:
        anchor = next(item for item in root.iter(qn("w:sdt")) if (_sdt_tag(item) or "").startswith(ANCHOR_TAG_PREFIX))
        anchor.find(qn("w:sdtContent")).append(OxmlElement("w:p"))

    _mutate_xml_member(second, "word/document.xml", add_paragraph)
    with pytest.raises(DocxSemanticsV3Error, match="cardinality"):
        second_session.prove_package(second)


def test_heading_target_requires_heading_style(tmp_path: Path) -> None:
    source = "# Intro ^intro\n"
    document = Document()
    session = DocxSemanticsV3Session(document, source_sha256=hashlib.sha256(source.encode()).hexdigest())
    heading = document.add_heading("Intro", level=1)
    session.bind_heading(heading, {"kind": "heading", "id": "intro"})
    session.finalize_document()
    output = tmp_path / "heading-style.docx"
    document.save(str(output))
    session.write_package(output)

    def mutate(root) -> None:
        style = root.find(f".//{qn('w:pPr')}/{qn('w:pStyle')}")
        assert style is not None
        style.set(qn("w:val"), "Normal")

    _mutate_xml_member(output, "word/document.xml", mutate)
    with pytest.raises(DocxSemanticsV3Error, match=r"Heading1\.\.Heading9"):
        session.prove_package(output)


def test_all_custom_xml_item_ids_and_owned_props_bytes_are_proven(tmp_path: Path) -> None:
    session, output = _write_paragraph_anchor(tmp_path)
    with ZipFile(output) as package:
        _item_name, props_name = _owned_item_and_props(
            output,
            "document-target-map",
        )
        props = package.read(props_name)
        unrelated = package.read("customXml/item1.xml")
    item_id = re.search(rb'ds:itemID="(\{[0-9A-F-]{36}\})"', props).group(1)  # type: ignore[union-attr]
    injected = unrelated.replace(
        b"<b:Sources ",
        b'<b:Sources ds:itemID="'
        + item_id
        + b'" xmlns:ds="http://schemas.openxmlformats.org/officeDocument/2006/customXml" ',
        1,
    )
    _replace_zip_members(output, {"customXml/item1.xml": injected})
    with pytest.raises(DocxSemanticsV3Error, match="UUID is duplicated"):
        session.prove_package(output)

    second_session, second = _write_paragraph_anchor(tmp_path, name="props-whitespace.docx")
    _item_name, second_props_name = _owned_item_and_props(second, "document-target-map")
    with ZipFile(second) as package:
        second_props = package.read(second_props_name)
    mutated_props = second_props.replace(b"><ds:schemaRefs>", b"> <ds:schemaRefs>", 1)
    _replace_zip_members(second, {second_props_name: mutated_props})
    with pytest.raises(DocxSemanticsV3Error, match="properties"):
        second_session.prove_package(second)


def test_duplicate_targetmode_internal_relationship_and_owned_tails_fail_closed(
    tmp_path: Path,
) -> None:
    session, output = _write_paragraph_anchor(tmp_path)
    item_name, _props_name = _owned_item_and_props(output, "document-target-map")

    def duplicate(root) -> None:
        owned = next(item for item in root if item.get("Target") == f"../{item_name}")
        duplicate = etree.fromstring(etree.tostring(owned))
        duplicate.set("Id", "rId999")
        duplicate.set("TargetMode", "Internal")
        root.append(duplicate)

    _mutate_xml_member(output, "word/_rels/document.xml.rels", duplicate)
    with pytest.raises(DocxSemanticsV3Error, match="exactly one"):
        session.prove_package(output)

    for part_name, marker in (
        ("word/_rels/document.xml.rels", "relationship"),
        ("[Content_Types].xml", "content types"),
    ):
        tail_session, tail_output = _write_paragraph_anchor(
            tmp_path,
            name=f"tail-{len(marker)}.docx",
        )
        tail_item, _ = _owned_item_and_props(tail_output, "document-target-map")

        def add_tail(root, *, owned_part=part_name, item_name=tail_item) -> None:
            if owned_part.startswith("word/"):
                owned = next(item for item in root if item.get("Target") == f"../{item_name}")
            else:
                owned = next(item for item in root if item.get("PartName") == f"/{item_name}")
            owned.tail = "JUNK"

        _mutate_xml_member(tail_output, part_name, add_tail)
        with pytest.raises(DocxSemanticsV3Error, match=marker):
            tail_session.prove_package(tail_output)


def test_ordinary_anchor_group_reports_exact_group_and_member_index(tmp_path: Path) -> None:
    document = Document()
    session = DocxSemanticsV3Session(document, source_sha256=hashlib.sha256(b"- A\n- B\n\n^list\n").hexdigest())
    first = document.add_paragraph("A", style="List Bullet")
    second = document.add_paragraph("B", style="List Bullet")
    session.bind_ordinary_anchor((first._p, second._p), {"block_kind": "list", "id": "list"})
    session.finalize_document()
    output = tmp_path / "list-group.docx"
    document.save(str(output))
    session.write_package(output)
    recovery = session.prove_package(output)
    loaded = Document(str(output))
    reopened = DocxSemanticsV3Recovery.load(output, loaded)
    logical = reopened.logical_body_elements(loaded)
    groups = [reopened.ordinary_anchor_group(item) for item in logical if item.tag == qn("w:p")]
    groups = [item for item in groups if item is not None]
    assert [item.index for item in groups] == [0, 1]
    assert all(item.anchor.source_id == "list" for item in groups)
    assert groups[0].elements == groups[1].elements
    assert recovery.anchor_identities == reopened.anchor_identities


def test_caption_target_can_own_an_inner_ordinary_anchor_group(tmp_path: Path) -> None:
    source = "Table: Results ^results\n\n| A |\n|---|\n\n^raw-table\n"
    document = Document()
    caption_styles = _add_caption_styles(document)
    session = DocxSemanticsV3Session(
        document,
        source_sha256=hashlib.sha256(source.encode()).hexdigest(),
        caption_style_bindings=caption_styles,
    )
    caption = document.add_paragraph(style="DocWenTableCaption")
    caption.add_run("Table ")
    append_complex_field(caption, instruction=" SEQ Table \\* ARABIC ", cached_result="1")
    caption.add_run(": Results")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "A"
    session.bind_ordinary_anchor(
        (table._tbl,),
        {"block_kind": "table", "id": "raw-table"},
    )
    session.bind_caption(
        caption,
        (table._tbl,),
        {
            "kind": "table",
            "id": "results",
            "number": "1",
            "title": "Results",
        },
    )
    session.finalize_document()
    output = tmp_path / "caption-inner-anchor.docx"
    document.save(str(output))
    session.write_package(output)
    session.prove_package(output)

    loaded = Document(str(output))
    recovery = DocxSemanticsV3Recovery.load(output, loaded)
    logical = recovery.logical_body_elements(loaded)
    table_element = next(item for item in logical if item.tag == qn("w:tbl"))
    group = recovery.ordinary_anchor_group(table_element)
    assert group is not None
    assert group.anchor.source_id == "raw-table"
    assert group.anchor.block_kind == "table"
    assert group.elements == (table_element,)
    assert group.index == 0
    assert recovery.caption_for_object(table_element).source_id == "results"  # type: ignore[union-attr]


def _add_caption_styles(document) -> tuple[CaptionStyleBindingV3, ...]:
    values = (
        ("figure_caption", "DocWenFigureCaption"),
        ("table_caption", "DocWenTableCaption"),
        ("equation_caption", "DocWenEquationCaption"),
        ("code_block_caption", "DocWenCodeBlockCaption"),
    )
    for _key, style_id in values:
        document.styles.add_style(style_id, WD_STYLE_TYPE.PARAGRAPH)
    return tuple(
        CaptionStyleBindingV3(key, style_id, style_id)  # type: ignore[arg-type]
        for key, style_id in values
    )


def _write_two_soft_references(
    tmp_path: Path,
    *,
    name: str = "two-soft.docx",
) -> tuple[DocxSemanticsV3Session, Path]:
    source = "# One\n\n@[[#One]] @[[#One|Alias]]\n"
    document = Document()
    session = DocxSemanticsV3Session(document, source_sha256=hashlib.sha256(source.encode()).hexdigest())
    document.add_heading("One", level=1)
    paragraph = document.add_paragraph()
    session.render_reference(paragraph, _soft("@[[#One]]", 7, "1"))
    paragraph.add_run(" ")
    session.render_reference(paragraph, _soft("@[[#One|Alias]]", 17, "1", alias="Alias"))
    session.finalize_document()
    output = tmp_path / name
    document.save(str(output))
    session.write_package(output)
    session.prove_package(output)
    return session, output


def _soft(raw: str, start: int, number: str, *, alias: str | None = None) -> dict[str, object]:
    return {
        "selector_kind": "heading_path",
        "resolution_status": "resolved",
        "resolved_kind": "heading",
        "heading_path": ["One"],
        "alias": alias,
        "cached_number": number,
        "raw": raw,
        "range": {"start": start, "end": start + len(raw)},
    }


def _write_paragraph_anchor(
    tmp_path: Path,
    *,
    name: str = "paragraph-anchor.docx",
) -> tuple[DocxSemanticsV3Session, Path]:
    document = Document()
    session = DocxSemanticsV3Session(document, source_sha256=hashlib.sha256(b"P ^raw\n").hexdigest())
    paragraph = document.add_paragraph("P")
    session.bind_paragraph_anchor(paragraph, {"block_kind": "paragraph", "id": "raw"})
    session.finalize_document()
    output = tmp_path / name
    document.save(str(output))
    session.write_package(output)
    session.prove_package(output)
    return session, output


def _sdt_tag(sdt) -> str | None:
    tag = sdt.find(f"./{qn('w:sdtPr')}/{qn('w:tag')}")
    return None if tag is None else tag.get(qn("w:val"))


def _owned_item_and_props(path: Path, marker: str) -> tuple[str, str]:
    with ZipFile(path) as package:
        item = next(name for name in package.namelist() if marker.encode() in package.read(name))
    number = re.fullmatch(r"customXml/item([1-9][0-9]*)\.xml", item).group(1)  # type: ignore[union-attr]
    return item, f"customXml/itemProps{number}.xml"


def _rebind_props_uuid(props: bytes, item: bytes) -> bytes:
    namespace = "https://docwen.dev/schema/document-soft-reference-map/v1"
    value = "{" + str(uuid.uuid5(uuid.NAMESPACE_URL, f"{namespace}\0{hashlib.sha256(item).hexdigest()}")).upper() + "}"
    return re.sub(rb"\{[0-9A-F-]{36}\}", value.encode(), props, count=1)


def _replace_zip_members(path: Path, replacements: dict[str, bytes]) -> None:
    temporary = path.with_suffix(".tmp")
    with ZipFile(path) as source, ZipFile(temporary, "w") as target:
        for info in source.infolist():
            target.writestr(info, replacements.get(info.filename, source.read(info.filename)))
    temporary.replace(path)


def _mutate_xml_member(path: Path, name: str, mutate) -> None:
    def transform(value: bytes) -> bytes:
        root = etree.fromstring(value)
        mutate(root)
        return etree.tostring(root, encoding="UTF-8", xml_declaration=True, standalone=True)

    with ZipFile(path) as package:
        original = package.read(name)
    _replace_zip_members(path, {name: transform(original)})
