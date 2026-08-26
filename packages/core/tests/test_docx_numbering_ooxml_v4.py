"""Headless OOXML gates for the resolved-numbering physical layer."""

from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace
from zipfile import ZIP_DEFLATED, ZipFile

import lxml.etree as etree
import pytest
from docx import Document
from docx.oxml.ns import qn

from docwen_core.docx_numbering_ooxml import (
    WML_NS,
    ResolvedNumberingOoxmlError,
    apply_heading_numbering,
    create_heading_numbering_projection,
    existing_numbering_ids,
    inline_reference_sdt,
    materialize_caption_number,
    write_heading_numbering_projection,
)
from docwen_core.models.resolved_numbering import (
    CaptionMaterialization,
    HeadingCounterSegment,
    HeadingDefinition,
    HeadingInstance,
    HeadingLevelDefinition,
    HeadingListMaterialization,
    HeadingLiteralSegment,
    HeadingStart,
    NumberingTarget,
    ResolvedNumberingPlan,
)

pytestmark = pytest.mark.unit


def _plan(*, target: NumberingTarget | None = None) -> ResolvedNumberingPlan:
    definition = HeadingDefinition(
        definition_id="main",
        levels=(
            HeadingLevelDefinition(
                level=1,
                start=2,
                number_format="arabic_half",
                display=(HeadingCounterSegment(1, "arabic_half"), HeadingLiteralSegment(".")),
                suffix="space",
                restart_after_level=None,
            ),
            HeadingLevelDefinition(
                level=2,
                start=1,
                number_format="letter_lower",
                display=(
                    HeadingCounterSegment(1, "arabic_half"),
                    HeadingLiteralSegment("-"),
                    HeadingCounterSegment(2, "letter_lower"),
                ),
                suffix="nothing",
                restart_after_level=1,
            ),
        ),
    )
    return ResolvedNumberingPlan(
        heading_definitions=(definition,),
        heading_instances=(
            HeadingInstance(
                instance_id="scope-a",
                definition_id="main",
                starts=(HeadingStart(level=2, value=3),),
            ),
        ),
        targets=() if target is None else (target,),
    )


def _heading_target(*, enabled: bool = True) -> NumberingTarget:
    return NumberingTarget(
        source_start=0,
        source_end=12,
        kind="heading",
        enabled=enabled,
        target_id="heading-a",
        derived_number="2-c" if enabled else None,
        materialization=(HeadingListMaterialization("main", "scope-a", 2) if enabled else None),
    )


def _caption_target(
    materialization: CaptionMaterialization | None,
    *,
    derived: str | None,
    enabled: bool = True,
    target_id: str | None = "figure-a",
) -> NumberingTarget:
    return NumberingTarget(
        source_start=20,
        source_end=50,
        kind="figure",
        enabled=enabled,
        target_id=target_id,
        derived_number=derived,
        materialization=materialization,
    )


def _caption_materialization(
    type_: str,
    action: str,
    *,
    start: int | None = None,
    chapter_level: int | None = None,
    restart_level: int | None = None,
    chapter_cache: str | None = None,
    sequence_cache: str = "1",
) -> CaptionMaterialization:
    return CaptionMaterialization(
        type=type_,  # type: ignore[arg-type]
        counter="Figure",
        number_format="arabic_half",
        sequence_action=action,  # type: ignore[arg-type]
        start_value=start,
        chapter_heading_level=chapter_level,
        chapter_heading_style=f"heading_{chapter_level}" if chapter_level is not None else None,
        chapter_separator="-" if chapter_level is not None else None,
        restart_heading_level=restart_level,
        restart_heading_style=f"heading_{restart_level}" if restart_level is not None else None,
        chapter_cached_number=chapter_cache,
        sequence_cached_number=sequence_cache,
        localized_label="图",
        label_separator=" ",
    )


def _instructions(element) -> list[str]:
    return [item.text or "" for item in element.iter(qn("w:instrText"))]


def _replace_member(path: Path, member: str, data: bytes) -> None:
    replacement = path.with_suffix(".replacement.docx")
    with ZipFile(path) as source, ZipFile(replacement, "w", ZIP_DEFLATED) as output:
        for info in source.infolist():
            output.writestr(info, data if info.filename == member else source.read(info.filename))
    replacement.replace(path)


def test_heading_projection_is_exact_and_never_rewrites_authored_title(tmp_path: Path) -> None:
    document = Document()
    target = _heading_target()
    projection = create_heading_numbering_projection(
        _plan(target=target),
        heading_style_ids={1: "Heading1", 2: "Heading2"},
        existing_abstract_ids=existing_numbering_ids(document)[0],
        existing_num_ids=existing_numbering_ids(document)[1],
    )
    paragraph = document.add_heading("2.3 标题", level=2)
    apply_heading_numbering(
        paragraph,
        target,
        projection,
        heading_style_ids={1: "Heading1", 2: "Heading2"},
    )

    assert paragraph.text == "2.3 标题"
    assert not list(paragraph._p.iter(qn("w:fldChar")))
    num_pr = paragraph._p.find(f"{{{WML_NS}}}pPr/{{{WML_NS}}}numPr")
    assert num_pr is not None
    ilvl = num_pr.find(f"{{{WML_NS}}}ilvl")
    num_id = num_pr.find(f"{{{WML_NS}}}numId")
    assert ilvl is not None and ilvl.get(f"{{{WML_NS}}}val") == "1"
    assert num_id is not None and num_id.get(f"{{{WML_NS}}}val") == str(projection.num_id("scope-a"))

    output = tmp_path / "numbered.docx"
    document.save(str(output))
    write_heading_numbering_projection(output, projection)
    with ZipFile(output) as package:
        numbering = etree.fromstring(package.read("word/numbering.xml"))
    abstract = numbering.find(f'{{{WML_NS}}}abstractNum[@{{{WML_NS}}}abstractNumId="{projection.abstract_id("main")}"]')
    assert abstract is not None
    multi_level = abstract.find(f"{{{WML_NS}}}multiLevelType")
    assert multi_level is not None and multi_level.get(f"{{{WML_NS}}}val") == "multilevel"
    levels = abstract.findall(f"{{{WML_NS}}}lvl")
    assert len(levels) == 2
    first_text = levels[0].find(f"{{{WML_NS}}}lvlText")
    first_restart = levels[0].find(f"{{{WML_NS}}}lvlRestart")
    second_text = levels[1].find(f"{{{WML_NS}}}lvlText")
    second_restart = levels[1].find(f"{{{WML_NS}}}lvlRestart")
    assert first_text is not None and first_text.get(f"{{{WML_NS}}}val") == "%1."
    assert first_restart is not None and first_restart.get(f"{{{WML_NS}}}val") == "0"
    assert second_text is not None and second_text.get(f"{{{WML_NS}}}val") == "%1-%2"
    assert second_restart is not None and second_restart.get(f"{{{WML_NS}}}val") == "1"
    instance = numbering.find(f'{{{WML_NS}}}num[@{{{WML_NS}}}numId="{projection.num_id("scope-a")}"]')
    assert instance is not None
    override = instance.find(f"{{{WML_NS}}}lvlOverride")
    assert override is not None and override.get(f"{{{WML_NS}}}ilvl") == "1"
    start_override = override.find(f"{{{WML_NS}}}startOverride")
    assert start_override is not None and start_override.get(f"{{{WML_NS}}}val") == "3"


@pytest.mark.parametrize(
    ("materialization", "derived", "instruction"),
    [
        (
            _caption_materialization("simple_seq", "continue", sequence_cache="7"),
            "7",
            r" SEQ Figure \* ARABIC ",
        ),
        (
            _caption_materialization("simple_seq", "reset_to_start", start=9, sequence_cache="9"),
            "9",
            r" SEQ Figure \r 9 \* ARABIC ",
        ),
        (
            _caption_materialization(
                "simple_seq",
                "restart_by_heading_level",
                start=1,
                restart_level=3,
                sequence_cache="2",
            ),
            "2",
            r" SEQ Figure \s 3 \* ARABIC ",
        ),
    ],
)
def test_simple_caption_fields_have_exact_cached_result_and_bookmark(
    materialization: CaptionMaterialization,
    derived: str,
    instruction: str,
) -> None:
    document = Document()
    paragraph = document.add_paragraph("must be replaced")
    target = _caption_target(materialization, derived=derived)

    visible = materialize_caption_number(
        paragraph,
        target,
        authored_content="曲线",
        heading_style_names={"heading_3": "Heading 3"},
        bookmark_name="DW_T_0123456789abcdef0123456789abcdef012",
        bookmark_id="1",
    )

    assert visible == f"图 {derived} 曲线"
    assert _instructions(paragraph._p) == [instruction]
    assert [item.text for item in paragraph._p.iter(qn("w:t"))] == ["图 ", derived, " 曲线"]
    starts = list(paragraph._p.iter(qn("w:bookmarkStart")))
    ends = list(paragraph._p.iter(qn("w:bookmarkEnd")))
    assert [item.get(qn("w:id")) for item in starts] == ["1"]
    assert [item.get(qn("w:id")) for item in ends] == ["1"]


@pytest.mark.parametrize(
    ("materialization", "expected_seq"),
    [
        (
            _caption_materialization(
                "chapter_seq",
                "continue",
                chapter_level=1,
                chapter_cache="2",
                sequence_cache="4",
            ),
            r" SEQ Figure \* ARABIC ",
        ),
        (
            _caption_materialization(
                "chapter_seq",
                "reset_to_start",
                start=9,
                chapter_level=1,
                chapter_cache="2",
                sequence_cache="4",
            ),
            r" SEQ Figure \r 9 \* ARABIC ",
        ),
        (
            _caption_materialization(
                "chapter_seq",
                "restart_by_heading_level",
                start=1,
                chapter_level=1,
                restart_level=3,
                chapter_cache="2",
                sequence_cache="4",
            ),
            r" SEQ Figure \s 3 \* ARABIC ",
        ),
    ],
)
def test_chapter_caption_bookmark_owns_styleref_separator_and_seq(
    materialization: CaptionMaterialization,
    expected_seq: str,
) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    visible = materialize_caption_number(
        paragraph,
        _caption_target(materialization, derived="2-4"),
        authored_content="曲线",
        heading_style_names={"heading_1": "Heading 1", "heading_3": "Heading 3"},
        bookmark_name="DW_T_0123456789abcdef0123456789abcdef012",
        bookmark_id="3",
    )

    assert visible == "图 2-4 曲线"
    assert _instructions(paragraph._p) == [
        ' STYLEREF "Heading 1" \\n ',
        expected_seq,
    ]
    start = paragraph._p.index(next(paragraph._p.iter(qn("w:bookmarkStart"))))
    end = paragraph._p.index(next(paragraph._p.iter(qn("w:bookmarkEnd"))))
    field_positions = [paragraph._p.index(run) for run in paragraph._p if list(run.iter(qn("w:fldChar")))]
    assert start < min(field_positions) < max(field_positions) < end
    separator_run = next(run for run in paragraph._p if "".join(t.text or "" for t in run.iter(qn("w:t"))) == "-")
    assert start < paragraph._p.index(separator_run) < end


def test_chapter_field_caches_are_explicit_and_never_split_on_separator() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    materialization = _caption_materialization(
        "chapter_seq",
        "continue",
        chapter_level=1,
        chapter_cache="1-2",
        sequence_cache="3",
    )

    materialize_caption_number(
        paragraph,
        _caption_target(materialization, derived="1-2-3"),
        authored_content="多级章号",
        heading_style_names={"heading_1": "Heading 1"},
    )

    assert [item.text for item in paragraph._p.iter(qn("w:t"))] == [
        "图 ",
        "1-2",
        "-",
        "3",
        " 多级章号",
    ]


def test_caption_cache_contradiction_fails_before_partial_result_is_accepted() -> None:
    document = Document()
    paragraph = document.add_paragraph("original")
    materialization = _caption_materialization(
        "chapter_seq",
        "continue",
        chapter_level=1,
        chapter_cache="1-2",
        sequence_cache="3",
    )

    with pytest.raises(ResolvedNumberingOoxmlError, match="contradict"):
        materialize_caption_number(
            paragraph,
            _caption_target(materialization, derived="1-2-4"),
            authored_content="无效",
            heading_style_names={"heading_1": "Heading 1"},
        )
    assert paragraph.text == "original"


def test_disabled_id_bearing_caption_has_zero_width_bookmark_and_no_field() -> None:
    document = Document()
    paragraph = document.add_paragraph("old")
    visible = materialize_caption_number(
        paragraph,
        _caption_target(None, derived=None, enabled=False),
        authored_content="曲线",
        heading_style_names={},
        bookmark_name="DW_T_0123456789abcdef0123456789abcdef012",
        bookmark_id="5",
    )

    assert visible == "曲线"
    children = list(paragraph._p)
    start_index = next(index for index, item in enumerate(children) if item.tag == qn("w:bookmarkStart"))
    end_index = next(index for index, item in enumerate(children) if item.tag == qn("w:bookmarkEnd"))
    text_index = next(index for index, item in enumerate(children) if item.tag == qn("w:r"))
    assert start_index + 1 == end_index < text_index
    assert not _instructions(paragraph._p)
    assert paragraph.text == "曲线"


def test_reference_uses_ref_cache_and_keeps_alias_outside_field() -> None:
    sdt = inline_reference_sdt(
        "docwen-ref-occurrence-v1:0123456789abcdef0123456789abcdef",
        bookmark_name="DW_T_0123456789abcdef0123456789abcdef012",
        cached_number="3.2",
        heading_number_only=True,
        alias="当前标题",
    )

    assert _instructions(sdt) == [r" REF DW_T_0123456789abcdef0123456789abcdef012 \n \h "]
    assert [item.text for item in sdt.iter(qn("w:t"))] == ["3.2", " 当前标题"]
    content = sdt.find(qn("w:sdtContent"))
    assert content is not None
    end_index = next(
        index
        for index, run in enumerate(content)
        if any(item.get(qn("w:fldCharType")) == "end" for item in run.iter(qn("w:fldChar")))
    )
    alias_index = next(index for index, run in enumerate(content) if "当前标题" in "".join(run.itertext()))
    assert alias_index > end_index


def test_post_save_numbering_id_collision_fails_closed(tmp_path: Path) -> None:
    document = Document()
    projection = create_heading_numbering_projection(
        _plan(),
        heading_style_ids={1: "Heading1", 2: "Heading2"},
        existing_abstract_ids=existing_numbering_ids(document)[0],
        existing_num_ids=existing_numbering_ids(document)[1],
    )
    collision = etree.Element(qn("w:num"))
    collision.set(qn("w:numId"), str(projection.num_id("scope-a")))
    ref = etree.SubElement(collision, qn("w:abstractNumId"))
    ref.set(qn("w:val"), "0")
    document.part.numbering_part.element.append(collision)
    output = tmp_path / "collision.docx"
    document.save(str(output))

    with pytest.raises(ResolvedNumberingOoxmlError, match="collide"):
        write_heading_numbering_projection(output, projection)


def test_allocator_reserves_dangling_story_style_and_abstract_references() -> None:
    document = Document()
    base_abstract, base_num = existing_numbering_ids(document)
    reserved_nums = tuple(max(base_num, default=0) + offset for offset in (1, 2, 3))
    reserved_abstract = max(base_abstract, default=-1) + 1

    body_paragraph = document.add_paragraph("dangling body")
    body_num_pr = etree.SubElement(body_paragraph._p.get_or_add_pPr(), qn("w:numPr"))
    etree.SubElement(body_num_pr, qn("w:numId")).set(qn("w:val"), str(reserved_nums[0]))
    style = document.styles["Normal"]
    style_p_pr = style.element.find(qn("w:pPr"))
    if style_p_pr is None:
        style_p_pr = etree.SubElement(style.element, qn("w:pPr"))
    style_num_pr = etree.SubElement(style_p_pr, qn("w:numPr"))
    etree.SubElement(style_num_pr, qn("w:numId")).set(qn("w:val"), str(reserved_nums[1]))
    header_paragraph = document.sections[0].header.paragraphs[0]
    header_num_pr = etree.SubElement(header_paragraph._p.get_or_add_pPr(), qn("w:numPr"))
    etree.SubElement(header_num_pr, qn("w:numId")).set(qn("w:val"), str(reserved_nums[2]))
    dangling_num = etree.SubElement(document.part.numbering_part.element, qn("w:num"))
    dangling_num.set(qn("w:numId"), str(reserved_nums[2] + 1))
    etree.SubElement(dangling_num, qn("w:abstractNumId")).set(qn("w:val"), str(reserved_abstract))

    abstract_ids, num_ids = existing_numbering_ids(document)
    projection = create_heading_numbering_projection(
        _plan(),
        heading_style_ids={1: "Heading1", 2: "Heading2"},
        existing_abstract_ids=abstract_ids,
        existing_num_ids=num_ids,
    )

    assert set(reserved_nums) <= num_ids
    assert reserved_nums[2] + 1 in num_ids
    assert reserved_abstract in abstract_ids
    assert projection.num_id("scope-a") not in {*reserved_nums, reserved_nums[2] + 1}
    assert projection.abstract_id("main") != reserved_abstract


@pytest.mark.parametrize("mutation", ["relationship", "content_type"])
def test_existing_numbering_part_requires_exact_opc_support(
    tmp_path: Path,
    mutation: str,
) -> None:
    document = Document()
    projection = create_heading_numbering_projection(
        _plan(),
        heading_style_ids={1: "Heading1", 2: "Heading2"},
        existing_abstract_ids=existing_numbering_ids(document)[0],
        existing_num_ids=existing_numbering_ids(document)[1],
    )
    output = tmp_path / f"missing-{mutation}.docx"
    document.save(str(output))
    with ZipFile(output) as package:
        if mutation == "relationship":
            member = "word/_rels/document.xml.rels"
            root = etree.fromstring(package.read(member))
            victim = next(item for item in root if item.get("Target") == "numbering.xml")
        else:
            member = "[Content_Types].xml"
            root = etree.fromstring(package.read(member))
            victim = next(item for item in root if item.get("PartName") == "/word/numbering.xml")
        root.remove(victim)
    _replace_member(output, member, etree.tostring(root, encoding="UTF-8", xml_declaration=True))
    before = output.read_bytes()

    with pytest.raises(ResolvedNumberingOoxmlError, match="lacks its exact"):
        write_heading_numbering_projection(output, projection)
    assert output.read_bytes() == before


def test_disabled_heading_removes_numpr_without_touching_authored_text() -> None:
    document = Document()
    paragraph = document.add_heading("第二章 概述", level=1)
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = etree.SubElement(p_pr, qn("w:numPr"))
    etree.SubElement(num_pr, qn("w:numId")).set(qn("w:val"), "4")
    empty_projection = SimpleNamespace()

    apply_heading_numbering(
        paragraph,
        _heading_target(enabled=False),
        empty_projection,  # type: ignore[arg-type]
        heading_style_ids={1: "Heading1"},
    )

    assert paragraph._p.find(f"{{{WML_NS}}}pPr/{{{WML_NS}}}numPr") is None
    assert paragraph.text == "第二章 概述"
