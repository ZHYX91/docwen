"""Focused OOXML/package tests for the Markdown semantics v3 runtime slice."""

from __future__ import annotations

import hashlib
import re
import uuid
from pathlib import Path
from zipfile import ZipFile

import lxml.etree as etree
import pytest
from docx import Document
from docx.oxml.ns import qn

from docwen_core._docx_semantics_v3_package import reference_occurrence_map_xml
from docwen_core.docx_bookmarks import build_docx_bookmark_inventory
from docwen_core.docx_semantics_v3 import (
    ANCHOR_TAG_PREFIX,
    REFERENCE_OCCURRENCE_MAP_NAMESPACE,
    REFERENCE_OCCURRENCE_TAG_PREFIX,
    SOFT_REFERENCE_MAP_NAMESPACE,
    TARGET_MAP_NAMESPACE,
    DocxSemanticsV3Error,
    DocxSemanticsV3Recovery,
    DocxSemanticsV3Session,
    derive_anchor_identity_v3,
    derive_reference_occurrence_identity_v3,
    derive_soft_reference_identity_v3,
    derive_target_identity_v3,
)

pytestmark = pytest.mark.contract


def test_frozen_v3_identity_preimages_are_exact() -> None:
    target = derive_target_identity_v3("heading", "intro")
    target_digest = hashlib.sha256(b"docwen-target-map-v1\0heading\0intro").hexdigest()
    assert target.sha256 == target_digest
    assert target.bookmark_name == f"DW_T_{target_digest[:35]}"
    assert target.tag == f"docwen-target-v1:{target_digest[:32]}"

    anchor = derive_anchor_identity_v3("paragraph", "raw")
    anchor_digest = hashlib.sha256(b"docwen-anchor-map-v1\0anchor\0raw").hexdigest()
    assert anchor.sha256 == anchor_digest
    assert anchor.tag == f"{ANCHOR_TAG_PREFIX}{anchor_digest[:32]}"

    source_sha256 = hashlib.sha256(b"source").hexdigest()
    soft = derive_soft_reference_identity_v3(
        source_sha256=source_sha256,
        source_start=12,
        source_end=29,
        authored_token="@[[#Other|Alias]]",
        cached_number="2",
    )
    soft_preimage = "docwen-soft-ref-map-v1\0" + source_sha256 + "\0" + "12" + "\0" + "29" + "\0@[[#Other|Alias]]"
    soft_digest = hashlib.sha256(soft_preimage.encode()).hexdigest()
    assert soft.tag == f"docwen-soft-ref-v1:{soft_digest[:32]}"


def test_heading_anchor_references_and_citation_round_trip_headlessly(tmp_path: Path) -> None:
    source = """# Intro ^intro

# Other

Plain block ^raw

See [[Page#^raw]], ![[Page#^raw]], @[[#^intro]], @[[#Other|Alias]], and @fig-legacy.
"""
    source_sha256 = hashlib.sha256(source.encode()).hexdigest()
    document = Document()
    session = DocxSemanticsV3Session(document, source_sha256=source_sha256)

    intro = document.add_heading("Intro", level=1)
    session.bind_heading(intro, {"kind": "heading", "id": "intro"})
    document.add_heading("Other", level=1)
    ordinary = document.add_paragraph("Plain block")
    session.bind_paragraph_anchor(ordinary, {"block_kind": "paragraph", "id": "raw"})
    paragraph = document.add_paragraph("See [[Page#^raw]], ![[Page#^raw]], ")
    session.render_reference(
        paragraph,
        {
            "selector_kind": "stable_id",
            "resolution_status": "resolved",
            "resolved_kind": "heading",
            "target_id": "intro",
            "cached_number": "1",
            "raw": "@[[#^intro]]",
            "range": {"start": 70, "end": 82},
        },
    )
    paragraph.add_run(", ")
    session.render_reference(
        paragraph,
        {
            "selector_kind": "heading_path",
            "resolution_status": "resolved",
            "resolved_kind": "heading",
            "heading_path": ["Other"],
            "alias": "Alias",
            "cached_number": "2",
            "raw": "@[[#Other|Alias]]",
            "range": {"start": 84, "end": 101},
        },
    )
    paragraph.add_run(", and @fig-legacy.")

    session.finalize_document()
    output = tmp_path / "v3-roundtrip.docx"
    document.save(str(output))
    session.write_package(output)
    proven = session.prove_package(output)

    target = derive_target_identity_v3("heading", "intro")
    anchor = derive_anchor_identity_v3("paragraph", "raw")
    with ZipFile(output) as package:
        names = set(package.namelist())
        document_xml = package.read("word/document.xml").decode()
        semantic_name = _custom_xml_part_for_namespace(package, TARGET_MAP_NAMESPACE)
        soft_name = _custom_xml_part_for_namespace(package, SOFT_REFERENCE_MAP_NAMESPACE)
        semantic_xml = package.read(semantic_name).decode()
        soft_xml = package.read(soft_name).decode()
    assert semantic_name in names and soft_name in names
    assert document_xml.count(target.bookmark_name) == 2  # target and REF instruction
    assert " REF " + target.bookmark_name + r" \n \h " in document_xml
    assert anchor.tag in document_xml
    assert "@fig-legacy" in document_xml
    assert "[[Page#^raw]]" in document_xml
    assert f'source_id="raw" tag="{anchor.tag}"' in semantic_xml
    assert 'bookmark_name="' + target.bookmark_name + '"' in semantic_xml
    assert 'authored_token="@[[#Other|Alias]]"' in soft_xml

    loaded = Document(str(output))
    recovery = DocxSemanticsV3Recovery.load(output, loaded)
    body = recovery.logical_body_elements(loaded)
    paragraphs = [item for item in body if item.tag == qn("w:p")]
    assert len(paragraphs) == 4
    assert recovery.source_anchor(paragraphs[0]).source_id == "intro"  # type: ignore[union-attr]
    assert recovery.source_anchor(paragraphs[2]).source_id == "raw"  # type: ignore[union-attr]
    assert recovery.source_anchor(paragraphs[2]).owner_kind == "ordinary_anchor"  # type: ignore[union-attr]
    assert recovery.render_paragraph_text(paragraphs[3]) == (
        "See [[Page#^raw]], ![[Page#^raw]], @[[#^intro]], @[[#Other|Alias]], and @fig-legacy."
    )
    assert proven.target_identities == recovery.target_identities
    assert proven.anchor_identities == recovery.anchor_identities
    assert proven.soft_reference_identities == recovery.soft_reference_identities
    assert proven.stable_reference_target_ids == ("intro",)

    inventory = build_docx_bookmark_inventory(loaded)
    assert inventory.starts_named(target.bookmark_name)
    assert not inventory.starts_named(anchor.tag)
    assert all((item.name or "") != "raw" for item in inventory.starts)
    assert "SEQ " not in document_xml


def test_soft_reference_tampering_fails_closed(tmp_path: Path) -> None:
    source_sha256 = hashlib.sha256(b"# Other\n\n@[[#Other|Alias]]\n").hexdigest()
    document = Document()
    session = DocxSemanticsV3Session(document, source_sha256=source_sha256)
    document.add_heading("Other", level=1)
    paragraph = document.add_paragraph()
    session.render_reference(
        paragraph,
        {
            "selector_kind": "heading_path",
            "resolution_status": "resolved",
            "resolved_kind": "heading",
            "heading_path": ["Other"],
            "alias": "Alias",
            "cached_number": "1",
            "raw": "@[[#Other|Alias]]",
            "range": {"start": 9, "end": 26},
        },
    )
    session.finalize_document()
    output = tmp_path / "tampered.docx"
    document.save(str(output))
    session.write_package(output)

    _replace_zip_member(
        output,
        "word/document.xml",
        lambda value: value.replace(b">1 Alias<", b">9 Alias<", 1),
    )

    with pytest.raises(DocxSemanticsV3Error, match="cached text"):
        DocxSemanticsV3Recovery.load(output, Document(str(output)))


def test_ordinary_anchor_never_becomes_word_target(tmp_path: Path) -> None:
    document = Document()
    session = DocxSemanticsV3Session(document, source_sha256=hashlib.sha256(b"P ^raw\n").hexdigest())
    paragraph = document.add_paragraph("P")
    session.bind_paragraph_anchor(paragraph, {"block_kind": "paragraph", "id": "raw"})
    session.finalize_document()
    output = tmp_path / "ordinary.docx"
    document.save(str(output))
    session.write_package(output)

    with ZipFile(output) as package:
        xml = package.read("word/document.xml").decode()
    assert ANCHOR_TAG_PREFIX in xml
    assert "bookmarkStart" not in xml
    assert "SEQ " not in xml
    assert " REF " not in xml


def test_package_proof_rejects_noncanonical_owned_ref(tmp_path: Path) -> None:
    document = Document()
    session = DocxSemanticsV3Session(
        document,
        source_sha256=hashlib.sha256(b"# Intro ^intro\n\n@[[#^intro]]\n").hexdigest(),
    )
    heading = document.add_heading("Intro", level=1)
    session.bind_heading(heading, {"kind": "heading", "id": "intro"})
    reference = document.add_paragraph()
    session.render_reference(
        reference,
        {
            "selector_kind": "stable_id",
            "resolution_status": "resolved",
            "resolved_kind": "heading",
            "target_id": "intro",
            "cached_number": "1",
            "raw": "@[[#^intro]]",
            "range": {"start": 16, "end": 28},
        },
    )
    session.finalize_document()
    output = tmp_path / "bad-ref.docx"
    document.save(str(output))
    session.write_package(output)
    _replace_zip_member(
        output,
        "word/document.xml",
        lambda value: value.replace(b" \\n \\h ", b" \\h ", 1),
    )

    with pytest.raises(DocxSemanticsV3Error, match="instruction does not match"):
        session.prove_package(output)


def test_reference_occurrence_physical_order_is_source_bound(tmp_path: Path) -> None:
    session, output, source_sha256 = _write_two_occurrence_package(tmp_path)
    first = derive_reference_occurrence_identity_v3(
        source_sha256=source_sha256,
        source_start=16,
        source_end=28,
        authored_token="@[[#^intro]]",
        resolved_bookmark_name=derive_target_identity_v3("heading", "intro").bookmark_name,
        cached_number="1",
    )
    second = derive_reference_occurrence_identity_v3(
        source_sha256=source_sha256,
        source_start=29,
        source_end=41,
        authored_token="@[[#^intro]]",
        resolved_bookmark_name=derive_target_identity_v3("heading", "intro").bookmark_name,
        cached_number="1",
    )
    _replace_zip_member(
        output,
        "word/document.xml",
        lambda value: (
            value.replace(first.tag.encode(), b"SWAP_TAG", 1)
            .replace(second.tag.encode(), first.tag.encode(), 1)
            .replace(b"SWAP_TAG", second.tag.encode(), 1)
        ),
    )

    with pytest.raises(DocxSemanticsV3Error, match="physical order"):
        session.prove_package(output)


@pytest.mark.parametrize(
    ("old", "new", "message"),
    [
        (
            b'<w:r><w:fldChar w:fldCharType="separate"/></w:r>',
            b"",
            "canonical REF runs",
        ),
        (
            b"</w:instrText></w:r>",
            b'</w:instrText><w:fldChar w:fldCharType="begin"/></w:r>',
            "run payload",
        ),
    ],
)
def test_reference_occurrence_field_skeleton_is_closed(
    tmp_path: Path,
    old: bytes,
    new: bytes,
    message: str,
) -> None:
    session, output, _source_sha256 = _write_two_occurrence_package(tmp_path)

    def mutate(value: bytes) -> bytes:
        assert old in value
        return value.replace(old, new, 1)

    _replace_zip_member(output, "word/document.xml", mutate)
    with pytest.raises(DocxSemanticsV3Error, match=message):
        session.prove_package(output)


def test_reference_occurrence_map_rejects_nonempty_record_with_recomputed_uuid(tmp_path: Path) -> None:
    session, output, _source_sha256 = _write_two_occurrence_package(tmp_path)
    with ZipFile(output) as package:
        item_name = _custom_xml_part_for_namespace(package, REFERENCE_OCCURRENCE_MAP_NAMESPACE)
        item_number = re.fullmatch(r"customXml/item([1-9][0-9]*)\.xml", item_name).group(1)  # type: ignore[union-attr]
        item_bytes = package.read(item_name)
        props_name = f"customXml/itemProps{item_number}.xml"
        props_bytes = package.read(props_name)
    mutated_item = item_bytes.replace(b"/>", b">junk</referenceOccurrence>", 1)
    expected_uuid = (
        "{"
        + str(
            uuid.uuid5(
                uuid.NAMESPACE_URL,
                f"{REFERENCE_OCCURRENCE_MAP_NAMESPACE}\0{hashlib.sha256(mutated_item).hexdigest()}",
            )
        ).upper()
        + "}"
    ).encode()
    mutated_props = re.sub(rb"\{[0-9A-F-]{36}\}", expected_uuid, props_bytes, count=1)
    _replace_zip_members(output, {item_name: mutated_item, props_name: mutated_props})

    with pytest.raises(DocxSemanticsV3Error, match="record is not closed"):
        session.prove_package(output)


def test_writer_rejects_preclaimed_future_owned_item_uuid(tmp_path: Path) -> None:
    source = "# Intro ^intro\n\n@[[#^intro]]\n"
    source_sha256 = hashlib.sha256(source.encode()).hexdigest()
    document = Document()
    session = DocxSemanticsV3Session(document, source_sha256=source_sha256)
    heading = document.add_heading("Intro", level=1)
    session.bind_heading(heading, {"kind": "heading", "id": "intro"})
    reference = document.add_paragraph()
    session.render_reference(reference, _stable_reference(range_start=16, alias=None))
    session.finalize_document()
    output = tmp_path / "preclaimed.docx"
    document.save(str(output))
    target = derive_target_identity_v3("heading", "intro")
    record = derive_reference_occurrence_identity_v3(
        source_sha256=source_sha256,
        source_start=16,
        source_end=28,
        authored_token="@[[#^intro]]",
        resolved_bookmark_name=target.bookmark_name,
        cached_number="1",
    )
    item_bytes = reference_occurrence_map_xml([record])
    future_uuid = (
        "{"
        + str(
            uuid.uuid5(
                uuid.NAMESPACE_URL,
                f"{REFERENCE_OCCURRENCE_MAP_NAMESPACE}\0{hashlib.sha256(item_bytes).hexdigest()}",
            )
        ).upper()
        + "}"
    )
    props = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<ds:datastoreItem xmlns:ds="http://schemas.openxmlformats.org/officeDocument/2006/customXml" '
        f'ds:itemID="{future_uuid}"><ds:schemaRefs><ds:schemaRef ds:uri="urn:unrelated"/>'
        "</ds:schemaRefs></ds:datastoreItem>\n"
    ).encode()
    with ZipFile(output, "a") as package:
        package.writestr("customXml/itemProps99.xml", props)

    with pytest.raises(DocxSemanticsV3Error, match="UUID collides"):
        session.write_package(output)


def test_writer_rejects_owned_schema_signal_with_wrong_root(tmp_path: Path) -> None:
    document = Document()
    session = DocxSemanticsV3Session(document, source_sha256=hashlib.sha256(b"source").hexdigest())
    paragraph = document.add_paragraph("P")
    session.bind_paragraph_anchor(paragraph, {"block_kind": "paragraph", "id": "raw"})
    session.finalize_document()
    output = tmp_path / "wrong-root.docx"
    document.save(str(output))
    props = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<ds:datastoreItem xmlns:ds="http://schemas.openxmlformats.org/officeDocument/2006/customXml" '
        'ds:itemID="{00000000-0000-0000-0000-000000000099}"><ds:schemaRefs>'
        f'<ds:schemaRef ds:uri="{REFERENCE_OCCURRENCE_MAP_NAMESPACE}"/>'
        "</ds:schemaRefs></ds:datastoreItem>\n"
    ).encode()
    with ZipFile(output, "a") as package:
        package.writestr("customXml/item99.xml", b"<unrelated/>")
        package.writestr("customXml/itemProps99.xml", props)

    with pytest.raises(DocxSemanticsV3Error, match="non-owned root"):
        session.write_package(output)


def test_reader_rejects_duplicate_members_and_unknown_owned_relationship_attributes(tmp_path: Path) -> None:
    session, output, _source_sha256 = _write_two_occurrence_package(tmp_path)
    with ZipFile(output, "a") as package, pytest.warns(UserWarning):
        package.writestr("word/document.xml", package.read("word/document.xml"))
    with pytest.raises(DocxSemanticsV3Error, match="duplicate ZIP members"):
        session.prove_package(output)

    _session, clean, _source_sha256 = _write_two_occurrence_package(tmp_path, name="unknown-rel.docx")
    with ZipFile(clean) as package:
        item_name = _custom_xml_part_for_namespace(package, REFERENCE_OCCURRENCE_MAP_NAMESPACE)
        item_number = re.fullmatch(r"customXml/item([1-9][0-9]*)\.xml", item_name).group(1)  # type: ignore[union-attr]
        rels_name = f"customXml/_rels/item{item_number}.xml.rels"
    _replace_zip_member(
        clean,
        rels_name,
        lambda value: value.replace(b' Target="', b' Unknown="x" Target="', 1),
    )
    with pytest.raises(DocxSemanticsV3Error, match="relationship is not canonical"):
        DocxSemanticsV3Recovery.load(clean, Document(str(clean)))


@pytest.mark.parametrize("mutation", ["child", "text"])
def test_item_relationship_leaf_topology_is_closed(tmp_path: Path, mutation: str) -> None:
    session, output, _source_sha256 = _write_two_occurrence_package(tmp_path)
    item_number = _owned_item_number(output, REFERENCE_OCCURRENCE_MAP_NAMESPACE)
    rels_name = f"customXml/_rels/item{item_number}.xml.rels"

    def mutate(value: bytes) -> bytes:
        suffix = (
            b"><Unexpected/></Relationship></Relationships>"
            if mutation == "child"
            else b">junk</Relationship></Relationships>"
        )
        assert b"/></Relationships>" in value
        return value.replace(b"/></Relationships>", suffix, 1)

    _replace_zip_member(output, rels_name, mutate)
    with pytest.raises(DocxSemanticsV3Error, match="item relationship is not canonical"):
        session.prove_package(output)


@pytest.mark.parametrize("mutation", ["wrong_tag", "child", "text", "duplicate_id"])
def test_document_relationship_topology_is_closed(tmp_path: Path, mutation: str) -> None:
    session, output, _source_sha256 = _write_two_occurrence_package(tmp_path)
    item_number = _owned_item_number(output, REFERENCE_OCCURRENCE_MAP_NAMESPACE)

    def mutate(root) -> None:
        namespace = "http://schemas.openxmlformats.org/package/2006/relationships"
        owned = next(item for item in root if item.get("Target") == f"../customXml/item{item_number}.xml")
        if mutation == "wrong_tag":
            owned.tag = f"{{{namespace}}}NotRelationship"
        elif mutation == "child":
            etree.SubElement(owned, f"{{{namespace}}}Unexpected")
        elif mutation == "text":
            owned.text = "junk"
        else:
            other = next(item for item in root if item is not owned)
            other.set("Id", owned.get("Id"))

    _mutate_xml_member(output, "word/_rels/document.xml.rels", mutate)
    with pytest.raises(DocxSemanticsV3Error, match="relationship"):
        session.prove_package(output)


@pytest.mark.parametrize("part_kind", ["item", "props"])
@pytest.mark.parametrize("mutation", ["wrong_tag", "child", "text"])
def test_owned_content_type_override_topology_is_closed(
    tmp_path: Path,
    part_kind: str,
    mutation: str,
) -> None:
    session, output, _source_sha256 = _write_two_occurrence_package(tmp_path)
    item_number = _owned_item_number(output, REFERENCE_OCCURRENCE_MAP_NAMESPACE)
    part_name = f"/customXml/item{item_number}.xml" if part_kind == "item" else f"/customXml/itemProps{item_number}.xml"

    def mutate(root) -> None:
        namespace = "http://schemas.openxmlformats.org/package/2006/content-types"
        owned = next(item for item in root if item.get("PartName") == part_name)
        if mutation == "wrong_tag":
            owned.tag = f"{{{namespace}}}NotOverride"
        elif mutation == "child":
            etree.SubElement(owned, f"{{{namespace}}}Unexpected")
        else:
            owned.text = "junk"

    _mutate_xml_member(output, "[Content_Types].xml", mutate)
    with pytest.raises(DocxSemanticsV3Error, match="content types"):
        session.prove_package(output)


@pytest.mark.parametrize(
    "mutation",
    ["sdt_attribute", "sdt_text", "content_attribute", "content_text", "fldchar_text"],
)
def test_reference_occurrence_envelope_is_closed(tmp_path: Path, mutation: str) -> None:
    session, output, _source_sha256 = _write_two_occurrence_package(tmp_path)

    def mutate_document(root) -> None:
        occurrence = _first_occurrence_sdt(root)
        content = occurrence.find(qn("w:sdtContent"))
        assert content is not None
        if mutation == "sdt_attribute":
            occurrence.set(qn("w:rsidR"), "00000000")
        elif mutation == "sdt_text":
            occurrence.text = "junk"
        elif mutation == "content_attribute":
            content.set(qn("w:rsidR"), "00000000")
        elif mutation == "content_text":
            content.text = "junk"
        else:
            field = content.find(f".//{qn('w:fldChar')}")
            assert field is not None
            field.text = "junk"

    _mutate_xml_member(output, "word/document.xml", mutate_document)
    with pytest.raises(DocxSemanticsV3Error, match="canonical"):
        session.prove_package(output)


def test_reference_occurrence_alias_tail_is_closed(tmp_path: Path) -> None:
    session, output, _source_sha256 = _write_two_occurrence_package(
        tmp_path,
        second_alias="Again",
    )

    def mutate_document(root) -> None:
        occurrence = next(
            sdt
            for sdt in root.iter(qn("w:sdt"))
            if len(list(sdt.find(qn("w:sdtContent")))) == 6  # type: ignore[arg-type]
        )
        text = list(occurrence.find(qn("w:sdtContent")))[-1].find(qn("w:t"))  # type: ignore[arg-type]
        assert text is not None
        text.tail = "junk"

    _mutate_xml_member(output, "word/document.xml", mutate_document)
    with pytest.raises(DocxSemanticsV3Error, match="Alias"):
        session.prove_package(output)


def _write_two_occurrence_package(
    tmp_path: Path,
    *,
    name: str = "two-occurrences.docx",
    second_alias: str | None = None,
) -> tuple[DocxSemanticsV3Session, Path, str]:
    second_raw = "@[[#^intro]]" if second_alias is None else f"@[[#^intro|{second_alias}]]"
    source = f"# Intro ^intro\n\n@[[#^intro]] {second_raw}\n"
    source_sha256 = hashlib.sha256(source.encode()).hexdigest()
    document = Document()
    session = DocxSemanticsV3Session(document, source_sha256=source_sha256)
    heading = document.add_heading("Intro", level=1)
    session.bind_heading(heading, {"kind": "heading", "id": "intro"})
    paragraph = document.add_paragraph()
    session.render_reference(paragraph, _stable_reference(range_start=16, alias=None))
    paragraph.add_run(" ")
    session.render_reference(paragraph, _stable_reference(range_start=29, alias=second_alias))
    session.finalize_document()
    output = tmp_path / name
    document.save(str(output))
    session.write_package(output)
    session.prove_package(output)
    return session, output, source_sha256


def _stable_reference(*, range_start: int, alias: str | None) -> dict[str, object]:
    raw = "@[[#^intro]]" if alias is None else f"@[[#^intro|{alias}]]"
    return {
        "selector_kind": "stable_id",
        "resolution_status": "resolved",
        "resolved_kind": "heading",
        "target_id": "intro",
        "alias": alias,
        "cached_number": "1",
        "raw": raw,
        "range": {"start": range_start, "end": range_start + len(raw)},
    }


def _custom_xml_part_for_namespace(package: ZipFile, namespace: str) -> str:
    for name in package.namelist():
        if not name.startswith("customXml/item") or not name.endswith(".xml"):
            continue
        if namespace.encode() in package.read(name):
            return name
    raise AssertionError(f"missing custom XML map for {namespace}")


def _replace_zip_member(path: Path, name: str, transform) -> None:
    temporary = path.with_suffix(".tmp")
    with ZipFile(path) as source, ZipFile(temporary, "w") as target:
        for info in source.infolist():
            data = source.read(info.filename)
            target.writestr(info, transform(data) if info.filename == name else data)
    temporary.replace(path)


def _replace_zip_members(path: Path, replacements: dict[str, bytes]) -> None:
    temporary = path.with_suffix(".tmp")
    with ZipFile(path) as source, ZipFile(temporary, "w") as target:
        for info in source.infolist():
            target.writestr(info, replacements.get(info.filename, source.read(info.filename)))
    temporary.replace(path)


def _owned_item_number(path: Path, namespace: str) -> str:
    with ZipFile(path) as package:
        name = _custom_xml_part_for_namespace(package, namespace)
    match = re.fullmatch(r"customXml/item([1-9][0-9]*)\.xml", name)
    assert match is not None
    return match.group(1)


def _mutate_xml_member(path: Path, name: str, mutate) -> None:

    def transform(value: bytes) -> bytes:
        root = etree.fromstring(value)
        mutate(root)
        return etree.tostring(root, encoding="UTF-8", xml_declaration=True, standalone=True)

    _replace_zip_member(path, name, transform)


def _first_occurrence_sdt(root):
    for sdt in root.iter(qn("w:sdt")):
        tag = sdt.find(f"./{qn('w:sdtPr')}/{qn('w:tag')}")
        if tag is not None and (tag.get(qn("w:val")) or "").startswith(REFERENCE_OCCURRENCE_TAG_PREFIX):
            return sdt
    raise AssertionError("missing reference-occurrence SDT")
