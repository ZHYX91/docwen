"""Cross-owner source-carrier gates for the resolved-numbering DOCX session."""

from __future__ import annotations

import hashlib
from pathlib import Path
from typing import Any, cast
from zipfile import ZipFile

import lxml.etree as etree
import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docwen_core._docx_semantics_v3_fenced import (
    FENCED_SOURCE_TAG_PREFIX,
    derive_fenced_source_identity_v3,
)
from docwen_core._docx_semantics_v3_model import (
    ANCHOR_TAG_PREFIX,
    TARGET_TAG_PREFIX,
    CaptionStyleBindingV3,
    CaptionStyleKeyV3,
)
from docwen_core.docx_resolved_numbering import (
    ResolvedNumberingDocxError,
    ResolvedNumberingDocxSession,
)
from docwen_core.models.resolved_numbering import (
    CaptionMaterialization,
    HeadingCounterSegment,
    HeadingDefinition,
    HeadingInstance,
    HeadingLevelDefinition,
    HeadingListMaterialization,
    NumberingExportPlanEnvelope,
    NumberingTarget,
    ResolvedDocument,
    ResolvedDocumentEnvelope,
    ResolvedDocumentTarget,
    ResolvedNumberingPlan,
    ResolvedNumberingPort,
)

pytestmark = pytest.mark.unit


def _sha(value: str) -> str:
    return hashlib.sha256(value.encode()).hexdigest()


def _caption_bindings(document: Any) -> tuple[CaptionStyleBindingV3, ...]:
    output: list[CaptionStyleBindingV3] = []
    for key, style_id, name in (
        ("figure_caption", "DWFigureCaption", "Figure Caption V4"),
        ("table_caption", "DWTableCaption", "Table Caption V4"),
        ("equation_caption", "DWEquationCaption", "Equation Caption V4"),
        ("code_block_caption", "DWCodeCaption", "Code Caption V4"),
    ):
        style = document.styles.add_style(style_id, WD_STYLE_TYPE.PARAGRAPH)
        style.name = name
        output.append(CaptionStyleBindingV3(cast(CaptionStyleKeyV3, key), style.style_id, style.name))
    return tuple(output)


def _port(
    source: str,
    targets: tuple[ResolvedDocumentTarget, ...] = (),
    plan_targets: tuple[NumberingTarget, ...] = (),
    *,
    definitions: tuple[HeadingDefinition, ...] = (),
    instances: tuple[HeadingInstance, ...] = (),
) -> ResolvedNumberingPort:
    source_sha = _sha(source)
    plan_sha = _sha("resolved-carrier-plan")
    return ResolvedNumberingPort(
        ResolvedDocumentEnvelope(
            "source",
            source_sha,
            plan_sha,
            ResolvedDocument(source, targets, (), (), (), ()),
        ),
        NumberingExportPlanEnvelope(
            "source",
            source_sha,
            plan_sha,
            ResolvedNumberingPlan(definitions, instances, plan_targets),
        ),
    )


def _session(document: Any, port: ResolvedNumberingPort) -> ResolvedNumberingDocxSession:
    return ResolvedNumberingDocxSession(
        document,
        port,
        heading_style_ids={1: "Heading1"},
        heading_style_names={"heading_1": "Heading 1"},
        caption_style_bindings=_caption_bindings(document),
    )


def _anchor(
    source: str,
    source_id: str,
    block_kind: str,
    block_start: int,
    block_end: int,
    *,
    placement: str,
    marker_start: int | None = None,
    container_path: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    start = source.index(f"^{source_id}", block_start) if marker_start is None else marker_start
    return {
        "id": source_id,
        "block_kind": block_kind,
        "placement": placement,
        "range": {"start": start, "end": start + len(source_id) + 1},
        "block_range": {"start": block_start, "end": block_end},
        "container_path": [] if container_path is None else container_path,
    }


@pytest.mark.parametrize(
    ("source", "outer_kind"),
    (
        ("> # Head ^head\n\n^outer\n", "block_quote"),
        ("- # Head ^head\n\n^outer\n", "list"),
    ),
)
def test_equal_physical_outer_container_owns_addressable_heading(
    tmp_path: Path,
    source: str,
    outer_kind: str,
) -> None:
    target_end = source.index("\n\n^outer")
    target = ResolvedDocumentTarget(0, target_end, _sha(source[:target_end]), "heading", "head", 1, "Head")
    definition = HeadingDefinition(
        "main",
        (HeadingLevelDefinition(1, 1, "arabic_half", (HeadingCounterSegment(1, "arabic_half"),), "space", None),),
    )
    plan_target = NumberingTarget(
        0,
        target_end,
        "heading",
        True,
        "head",
        "1",
        HeadingListMaterialization("main", "document", 1),
    )
    document = Document()
    session = _session(
        document,
        _port(
            source,
            (target,),
            (plan_target,),
            definitions=(definition,),
            instances=(HeadingInstance("document", "main", ()),),
        ),
    )
    heading = document.add_heading("Head", level=1)
    session.bind_heading(heading, source_start=0, source_end=target_end)
    session.bind_ordinary_anchor(
        (heading._p,),
        _anchor(source, "outer", outer_kind, 0, target_end, placement="post_block"),
    )
    output = tmp_path / f"{outer_kind}-heading.docx"
    session.write_package(output)

    tags = _block_tags(output)
    assert tags[0].startswith(ANCHOR_TAG_PREFIX)
    assert tags[1].startswith(TARGET_TAG_PREFIX)
    session.prove_package(output)


def test_outer_quote_target_and_raw_image_form_one_authenticated_chain(tmp_path: Path) -> None:
    source = "> Figure: Cap ^fig\n>\n> ![x](a.png) ^raw\n\n^outer\n"
    target_end = source.index("\n\n^outer")
    target = ResolvedDocumentTarget(0, target_end, _sha(source[:target_end]), "figure", "fig", None, "Cap")
    materialization = CaptionMaterialization(
        "simple_seq",
        "Figure",
        "arabic_half",
        "continue",
        None,
        None,
        None,
        None,
        None,
        None,
        None,
        "1",
        "Figure",
        " ",
    )
    plan_target = NumberingTarget(0, target_end, "figure", True, "fig", "1", materialization)
    document = Document()
    session = _session(document, _port(source, (target,), (plan_target,)))
    image = document.add_paragraph()
    image.add_run()._r.append(OxmlElement("w:drawing"))
    caption = document.add_paragraph(style="Figure Caption V4")
    raw_start = source.index("> ![x]")
    raw = _anchor(
        source,
        "raw",
        "image",
        raw_start,
        target_end,
        placement="inline",
        container_path=[{"block_kind": "block_quote", "block_range": {"start": 0, "end": target_end}}],
    )
    outer = _anchor(source, "outer", "block_quote", 0, target_end, placement="post_block")
    session.bind_ordinary_anchor((image._p,), raw, direct_parent_source_id="outer")
    session.bind_caption(caption, (image._p,), source_start=0, source_end=target_end, kind="figure")
    session.bind_ordinary_anchor((image._p, caption._p), outer)
    output = tmp_path / "quote-figure-raw.docx"
    session.write_package(output)

    tags = _block_tags(output)
    assert [
        tag.startswith(ANCHOR_TAG_PREFIX) if index != 1 else tag.startswith(TARGET_TAG_PREFIX)
        for index, tag in enumerate(tags)
    ] == [True, True, True]
    session.prove_package(output)


def test_strict_outer_list_range_owns_heading_and_authored_tail(tmp_path: Path) -> None:
    source = "- # Head ^head\n  authored tail\n\n^outer\n"
    target_end = source.index("\n")
    outer_end = source.index("\n\n^outer")
    target = ResolvedDocumentTarget(0, target_end, _sha(source[:target_end]), "heading", "head", 1, "Head")
    definition = HeadingDefinition(
        "main",
        (HeadingLevelDefinition(1, 1, "arabic_half", (HeadingCounterSegment(1, "arabic_half"),), "space", None),),
    )
    plan_target = NumberingTarget(
        0,
        target_end,
        "heading",
        True,
        "head",
        "1",
        HeadingListMaterialization("main", "document", 1),
    )
    document = Document()
    session = _session(
        document,
        _port(
            source,
            (target,),
            (plan_target,),
            definitions=(definition,),
            instances=(HeadingInstance("document", "main", ()),),
        ),
    )
    heading = document.add_heading("Head", level=1)
    tail = document.add_paragraph("authored tail")
    session.bind_heading(heading, source_start=0, source_end=target_end)
    session.bind_ordinary_anchor(
        (heading._p, tail._p),
        _anchor(source, "outer", "list", 0, outer_end, placement="post_block"),
    )
    output = tmp_path / "strict-list-heading.docx"
    session.write_package(output)

    assert _block_tags(output)[0].startswith(ANCHOR_TAG_PREFIX)
    session.prove_package(output)


def test_outer_quote_owns_disabled_idless_occurrence(tmp_path: Path) -> None:
    source = "> Table: Cap\n>\n> |a|\n> |- |\n> |b|\n\n^outer\n"
    target_end = source.index("\n\n^outer")
    target = ResolvedDocumentTarget(0, target_end, _sha(source[:target_end]), "table", None, None, "Cap")
    plan_target = NumberingTarget(0, target_end, "table", False, None, None, None)
    document = Document()
    session = _session(document, _port(source, (target,), (plan_target,)))
    caption = document.add_paragraph(style="Table Caption V4")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "b"
    session.bind_caption(caption, (table._tbl,), source_start=0, source_end=target_end, kind="table")
    session.bind_ordinary_anchor(
        (caption._p, table._tbl),
        _anchor(source, "outer", "block_quote", 0, target_end, placement="post_block"),
    )
    output = tmp_path / "outer-disabled-table.docx"
    session.write_package(output)

    tags = _block_tags(output)
    assert tags[0].startswith(ANCHOR_TAG_PREFIX)
    assert tags[1].startswith("docwen-numbering-occurrence-v1:")
    session.prove_package(output)


def test_code_target_owns_raw_anchor_and_exact_fenced_payload(tmp_path: Path) -> None:
    source = "Code: Graph ^code\n\n```mermaid\ngraph TD\n```\n\n^raw-code\n"
    fence_start = source.index("```")
    fence_end = source.index("\n\n^raw-code") + 1
    target_end = len(source) - 1
    target = ResolvedDocumentTarget(0, target_end, _sha(source[:target_end]), "code_block", "code", None, "Graph")
    materialization = CaptionMaterialization(
        "simple_seq",
        "Code",
        "arabic_half",
        "continue",
        None,
        None,
        None,
        None,
        None,
        None,
        None,
        "1",
        "Code",
        " ",
    )
    plan_target = NumberingTarget(0, target_end, "code_block", True, "code", "1", materialization)
    identity = derive_fenced_source_identity_v3(
        source_sha256=_sha(source),
        source_start=fence_start,
        source_end=fence_end,
        block_sha256=_sha(source[fence_start:fence_end]),
        body_sha256=_sha("graph TD\n"),
        fence_character="`",
        opening_length=3,
        opening_prefix="",
        info="mermaid",
        opening_eol="\n",
        body_prefixes=("",),
        closing_state="present",
        closing_length=3,
        closing_prefix="",
        closing_suffix="",
        closing_eol="\n",
    )
    document = Document()
    session = _session(document, _port(source, (target,), (plan_target,)))
    caption = document.add_paragraph(style="Code Caption V4")
    code = document.add_paragraph("renderer placeholder")
    session.bind_fenced_source(code, identity, logical_body="graph TD\n")
    session.bind_ordinary_anchor(
        (code._p,),
        _anchor(source, "raw-code", "fenced_block", fence_start, fence_end, placement="post_block"),
    )
    session.bind_caption(caption, (code._p,), source_start=0, source_end=target_end, kind="code_block")
    output = tmp_path / "target-anchor-fence.docx"
    session.write_package(output)

    tags = _block_tags(output)
    assert tags[0].startswith(TARGET_TAG_PREFIX)
    assert tags[1].startswith(ANCHOR_TAG_PREFIX)
    with ZipFile(output) as package:
        assert FENCED_SOURCE_TAG_PREFIX.encode() in package.read("word/document.xml")
    session.prove_package(output)


def test_target_and_anchor_share_one_source_id_namespace() -> None:
    source = "> # Head ^same\n\n^same\n"
    target_end = source.index("\n\n^same")
    target = ResolvedDocumentTarget(0, target_end, _sha(source[:target_end]), "heading", "same", 1, "Head")
    definition = HeadingDefinition(
        "main",
        (HeadingLevelDefinition(1, 1, "arabic_half", (HeadingCounterSegment(1, "arabic_half"),), "space", None),),
    )
    plan_target = NumberingTarget(
        0,
        target_end,
        "heading",
        True,
        "same",
        "1",
        HeadingListMaterialization("main", "document", 1),
    )
    document = Document()
    session = _session(
        document,
        _port(
            source,
            (target,),
            (plan_target,),
            definitions=(definition,),
            instances=(HeadingInstance("document", "main", ()),),
        ),
    )
    heading = document.add_heading("Head", level=1)
    session.bind_heading(heading, source_start=0, source_end=target_end)

    with pytest.raises(ResolvedNumberingDocxError, match="source ID is duplicated"):
        session.bind_ordinary_anchor(
            (heading._p,),
            _anchor(source, "same", "block_quote", 0, target_end, placement="post_block", marker_start=target_end + 2),
        )


def test_partial_cross_anchor_overlap_fails_before_any_wrapper() -> None:
    source = "A ^left\n\nB\n\nC ^right\n"
    document = Document()
    session = _session(document, _port(source))
    first = document.add_paragraph("same")
    middle = document.add_paragraph("same")
    last = document.add_paragraph("same")
    left_end = source.index("\n\nB")
    right_start = source.index("C ")
    session.bind_ordinary_anchor(
        (first._p, middle._p),
        _anchor(source, "left", "list", 0, left_end, placement="inline"),
    )
    session.bind_ordinary_anchor(
        (middle._p, last._p),
        _anchor(source, "right", "block_quote", right_start, len(source), placement="inline"),
    )

    with pytest.raises(ResolvedNumberingDocxError, match="ownership is invalid"):
        session.finalize_document()
    assert not _live_block_tags(document)


def test_same_text_disjoint_anchor_tag_swap_is_rejected_after_reopen(tmp_path: Path) -> None:
    source = "same ^first\n\nsame ^second\n"
    document = Document()
    session = _session(document, _port(source))
    first = document.add_paragraph("same")
    second = document.add_paragraph("same")
    split = source.index("\n\n")
    session.bind_ordinary_anchor(
        (first._p,),
        _anchor(source, "first", "paragraph", 0, split, placement="inline"),
    )
    session.bind_ordinary_anchor(
        (second._p,),
        _anchor(source, "second", "paragraph", split + 2, len(source), placement="inline"),
    )
    output = tmp_path / "same-text.docx"
    session.write_package(output)
    _swap_anchor_tags(output)

    with pytest.raises(ResolvedNumberingDocxError, match="reopen proof"):
        session.prove_package(output)


def _block_tags(path: Path) -> list[str]:
    with ZipFile(path) as package:
        root = etree.fromstring(package.read("word/document.xml"))
    return [
        tag
        for item in root.iter(qn("w:sdt"))
        if (tag := _sdt_tag(item)) is not None
        and tag.startswith((ANCHOR_TAG_PREFIX, TARGET_TAG_PREFIX, "docwen-numbering-occurrence-v1:"))
    ]


def _live_block_tags(document: Any) -> list[str]:
    return [
        tag
        for item in document.element.body.iter(qn("w:sdt"))
        if (tag := _sdt_tag(item)) is not None
        and tag.startswith((ANCHOR_TAG_PREFIX, TARGET_TAG_PREFIX, "docwen-numbering-occurrence-v1:"))
    ]


def _sdt_tag(sdt: Any) -> str | None:
    tag = sdt.find(f"./{qn('w:sdtPr')}/{qn('w:tag')}")
    return None if tag is None else tag.get(qn("w:val"))


def _swap_anchor_tags(path: Path) -> None:
    with ZipFile(path) as package:
        infos = package.infolist()
        data = {item.filename: package.read(item.filename) for item in infos}
    root = etree.fromstring(data["word/document.xml"])
    tags = [item for item in root.iter(qn("w:tag")) if (item.get(qn("w:val")) or "").startswith(ANCHOR_TAG_PREFIX)]
    assert len(tags) == 2
    first = tags[0].get(qn("w:val"))
    second = tags[1].get(qn("w:val"))
    assert first is not None and second is not None
    tags[0].set(qn("w:val"), second)
    tags[1].set(qn("w:val"), first)
    data["word/document.xml"] = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    with ZipFile(path, "w") as package:
        for info in infos:
            package.writestr(info, data[info.filename])
