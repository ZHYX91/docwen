"""Focused DOCX writer/importer tests for neutral semantics."""

from __future__ import annotations

from copy import deepcopy

import pytest
from docx import Document
from docx.document import Document as DocumentType
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docwen_core.docx_bookmarks import BOOKMARK_ID_MAX, build_docx_bookmark_inventory
from docwen_core.docx_parsing.document_semantics import (
    BIBLIOGRAPHY_BOOKMARK_NAME,
    encode_object_bookmark,
    encode_object_pairing_bookmark,
    encode_shorthand_bookmark,
    encode_target_bookmark,
    extract_neutral_semantic_caption,
)
from docwen_core.docx_semantics import (
    DocxSemanticImporter,
    DocxSemanticRenderer,
    append_bookmark_end,
    append_zero_width_bookmark,
)
from docwen_core.models.semantic_document import (
    SemanticBibliographyEntry,
    SemanticBibliographyFragment,
    SemanticBibliographyRun,
    SemanticCaption,
    SemanticDocument,
    SemanticDocumentValidationError,
    SemanticParagraph,
    SemanticReference,
    SemanticTable,
    SemanticTableCell,
    SemanticText,
)

pytestmark = pytest.mark.contract


def _targetless_table() -> SemanticTable:
    return SemanticTable(
        row_count=1,
        column_count=1,
        cells=(SemanticTableCell(0, 0, "value"),),
        caption=SemanticCaption(
            kind="table",
            target_id=None,
            cached_number="3",
            label="Table",
            content="Unreferenced table",
        ),
    )


def _targeted_table(target_id: str = "tbl-sales") -> SemanticTable:
    return SemanticTable(
        row_count=1,
        column_count=1,
        cells=(SemanticTableCell(0, 0, "value"),),
        caption=SemanticCaption(
            kind="table",
            target_id=target_id,
            cached_number="7",
            label="Table",
            content="Targeted table",
        ),
    )


def _render_targetless_table() -> tuple[DocumentType, SemanticTable]:
    document = Document()
    table = _targetless_table()
    DocxSemanticRenderer(document).render_table(table)
    return document, table


def _render_targeted_table() -> tuple[DocumentType, SemanticTable]:
    document = Document()
    table = _targeted_table()
    DocxSemanticRenderer(document).render_table(table)
    return document, table


def _remove_marker_end(element, *, prefix: str) -> None:
    marker_start = next(
        item for item in element.iter(qn("w:bookmarkStart")) if (item.get(qn("w:name")) or "").startswith(prefix)
    )
    marker_id = marker_start.get(qn("w:id"))
    marker_end = next(item for item in element.iter(qn("w:bookmarkEnd")) if item.get(qn("w:id")) == marker_id)
    _remove_element(marker_end)


def _remove_element(element) -> None:
    parent = element.getparent()
    assert parent is not None
    parent.remove(element)


def _insert_simple_seq_before_complex_field(paragraph, *, identifier: str, cached_result: str) -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), f" SEQ {identifier} \\* ARABIC ")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = cached_result
    run.append(text)
    field.append(run)

    for index, child in enumerate(paragraph._p):
        if any(item.get(qn("w:fldCharType")) == "begin" for item in child.iter(qn("w:fldChar"))):
            paragraph._p.insert(index, field)
            return
    raise AssertionError("paragraph has no complex field")


def _wrap_table_child(parent, child, *, wrapper: str) -> None:
    index = parent.index(child)
    parent.remove(child)
    container = OxmlElement(f"w:{wrapper}")
    if wrapper == "sdt":
        content = OxmlElement("w:sdtContent")
        content.append(child)
        container.append(content)
    else:
        container.append(child)
    parent.insert(index, container)


def _append_vertical_merge(cell, value: str | None) -> None:
    vertical_merge = OxmlElement("w:vMerge")
    if value is not None:
        vertical_merge.set(qn("w:val"), value)
    cell.get_or_add_tcPr().append(vertical_merge)


def _prepare_malformed_table(document: DocumentType, malformation: str):
    if malformation == "rows_missing":
        return document.add_table(rows=0, cols=0)
    if malformation == "vmerge_mismatched":
        table = document.add_table(rows=2, cols=2)
        anchor = table.cell(0, 0)._tc
        removed = table.cell(0, 1)._tc
        table.rows[0]._tr.remove(removed)
        grid_span = OxmlElement("w:gridSpan")
        grid_span.set(qn("w:val"), "2")
        anchor.get_or_add_tcPr().append(grid_span)
        _append_vertical_merge(anchor, "restart")
        _append_vertical_merge(table.cell(1, 0)._tc, "continue")
        return table
    table = document.add_table(rows=1, cols=1)
    if malformation == "cells_missing":
        table.rows[0]._tr.remove(table.cell(0, 0)._tc)
    elif malformation.startswith("grid_span_"):
        grid_span = OxmlElement("w:gridSpan")
        value = {
            "grid_span_non_integer": "not-an-integer",
            "grid_span_zero": "0",
            "grid_span_negative": "-1",
            "grid_span_missing": None,
        }[malformation]
        if value is not None:
            grid_span.set(qn("w:val"), value)
        table.cell(0, 0)._tc.get_or_add_tcPr().append(grid_span)
    elif malformation in {"vmerge_orphan", "vmerge_invalid_enum", "vmerge_uppercase"}:
        _append_vertical_merge(
            table.cell(0, 0)._tc,
            {"vmerge_invalid_enum": "bogus", "vmerge_uppercase": "RESTART"}.get(malformation),
        )
    else:
        offset = OxmlElement("w:gridAfter" if malformation == "grid_after_negative" else "w:gridBefore")
        value = {
            "grid_before_missing": None,
            "grid_after_negative": "-1",
            "grid_before_non_integer": "not-an-integer",
        }[malformation]
        if value is not None:
            offset.set(qn("w:val"), value)
        table.rows[0]._tr.get_or_add_trPr().append(offset)
    return table


def _force_bookmark_inventory_fallback(document: DocumentType) -> None:
    malformed_part = Part(
        PackURI("/word/comments.xml"),
        CT.WML_COMMENTS,
        b"not-xml",
        document.part.package,
    )
    document.part.relate_to(malformed_part, RT.COMMENTS)


__all__ = (
    "BIBLIOGRAPHY_BOOKMARK_NAME",
    "BOOKMARK_ID_MAX",
    "CT",
    "RT",
    "Document",
    "DocxSemanticImporter",
    "DocxSemanticRenderer",
    "PackURI",
    "Part",
    "SemanticBibliographyEntry",
    "SemanticBibliographyFragment",
    "SemanticBibliographyRun",
    "SemanticDocument",
    "SemanticDocumentValidationError",
    "SemanticParagraph",
    "SemanticReference",
    "SemanticTable",
    "SemanticTableCell",
    "SemanticText",
    "_force_bookmark_inventory_fallback",
    "_insert_simple_seq_before_complex_field",
    "_prepare_malformed_table",
    "_remove_element",
    "_remove_marker_end",
    "_render_targeted_table",
    "_render_targetless_table",
    "_targeted_table",
    "_targetless_table",
    "_wrap_table_child",
    "append_bookmark_end",
    "append_zero_width_bookmark",
    "build_docx_bookmark_inventory",
    "deepcopy",
    "encode_object_bookmark",
    "encode_object_pairing_bookmark",
    "encode_shorthand_bookmark",
    "encode_target_bookmark",
    "extract_neutral_semantic_caption",
    "pytest",
    "pytestmark",
    "qn",
)
