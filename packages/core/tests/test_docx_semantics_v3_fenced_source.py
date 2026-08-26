"""Closed source/package proofs for the v4 fenced-source occurrence carrier."""

from __future__ import annotations

import hashlib
from pathlib import Path
from typing import Any, Literal, TypedDict, cast
from zipfile import ZipFile

import lxml.etree as etree
import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

from docwen_core._docx_semantics_v3_fenced_map import (
    fenced_source_map_xml,
    parse_fenced_source_map,
)
from docwen_core.docx_semantics_v3 import (
    FENCED_SOURCE_MAP_NAMESPACE,
    FENCED_SOURCE_TAG_PREFIX,
    CaptionStyleBindingV3,
    DocxSemanticsV3Error,
    DocxSemanticsV3Recovery,
    DocxSemanticsV3Session,
    FencedSourceIdentityV3,
    append_complex_field,
    derive_fenced_source_identity_v3,
    fenced_source_identity_from_mapping_v3,
    fenced_source_mapping_v3,
)

pytestmark = pytest.mark.contract


class _Framing(TypedDict):
    fence_character: str
    opening_length: int
    opening_prefix: str
    info: str
    opening_eol: str
    body_prefixes: tuple[str, ...]
    closing_state: Literal["present", "omitted_eof"]
    closing_length: int
    closing_prefix: str
    closing_suffix: str
    closing_eol: str


@pytest.mark.parametrize(
    ("source", "logical_body", "framing"),
    (
        (
            "```  rust linenums\nfn main() {}\n```  \n",
            "fn main() {}\n",
            {
                "fence_character": "`",
                "opening_length": 3,
                "opening_prefix": "",
                "info": "  rust linenums",
                "opening_eol": "\n",
                "body_prefixes": ("",),
                "closing_state": "present",
                "closing_length": 3,
                "closing_prefix": "",
                "closing_suffix": "  ",
                "closing_eol": "\n",
            },
        ),
        (
            "> ~~~~mermaid\r\n> graph TD\r\n> ~~~~~\t\r\n",
            "graph TD\r\n",
            {
                "fence_character": "~",
                "opening_length": 4,
                "opening_prefix": "> ",
                "info": "mermaid",
                "opening_eol": "\r\n",
                "body_prefixes": ("> ",),
                "closing_state": "present",
                "closing_length": 5,
                "closing_prefix": "> ",
                "closing_suffix": "\t",
                "closing_eol": "\r\n",
            },
        ),
        (
            "- ```query\n  tag:#project\n  ```",
            "tag:#project\n",
            {
                "fence_character": "`",
                "opening_length": 3,
                "opening_prefix": "- ",
                "info": "query",
                "opening_eol": "\n",
                "body_prefixes": ("  ",),
                "closing_state": "present",
                "closing_length": 3,
                "closing_prefix": "  ",
                "closing_suffix": "",
                "closing_eol": "",
            },
        ),
        (
            "```view\ncolumns: 2",
            "columns: 2",
            {
                "fence_character": "`",
                "opening_length": 3,
                "opening_prefix": "",
                "info": "view",
                "opening_eol": "\n",
                "body_prefixes": ("",),
                "closing_state": "omitted_eof",
                "closing_length": 0,
                "closing_prefix": "",
                "closing_suffix": "",
                "closing_eol": "",
            },
        ),
    ),
)
def test_direct_blockquote_list_and_omitted_eof_round_trip_exactly(
    tmp_path: Path,
    source: str,
    logical_body: str,
    framing: _Framing,
) -> None:
    identity = _identity(source, logical_body, **framing)
    document = Document()
    paragraph = _body_paragraph(document, logical_body)
    session = DocxSemanticsV3Session(document, source_sha256=identity.source_sha256)
    session.bind_fenced_source(
        paragraph,
        fenced_source_mapping_v3(identity),
        logical_body=logical_body,
    )
    output = _write(session, document, tmp_path / f"{identity.tag[-8:]}.docx")

    loaded = Document(str(output))
    recovery = DocxSemanticsV3Recovery.load(output, loaded)
    recovered_paragraph = next(cast(Any, loaded.element).body.iter(qn("w:p")))
    assert recovery.render_fenced_source(recovered_paragraph) == source
    assert recovery.fenced_source_identities == (identity,)
    with ZipFile(output) as package:
        document_xml = package.read("word/document.xml")
        map_xml = package.read(_map_part(package))
    assert identity.tag.encode() in document_xml
    assert FENCED_SOURCE_MAP_NAMESPACE.encode() in map_xml
    assert b"source_id=" not in map_xml
    assert b"target" not in map_xml.lower()
    assert b"bookmark" not in map_xml.lower()
    assert b"SEQ " not in document_xml and b" REF " not in document_xml


def test_carrier_nests_inside_ordinary_anchor_without_acquiring_target_fields(tmp_path: Path) -> None:
    source = "```rust\nfn main() {}\n```\n\n^raw-code\n"
    block = source[: source.index("\n\n")]
    identity = _identity(
        block,
        "fn main() {}\n",
        fence_character="`",
        opening_length=3,
        opening_prefix="",
        info="rust",
        opening_eol="\n",
        body_prefixes=("",),
        closing_state="present",
        closing_length=3,
        closing_prefix="",
        closing_suffix="",
        closing_eol="",
        source_document=source,
    )
    document = Document()
    paragraph = _body_paragraph(document, "fn main() {}")
    session = DocxSemanticsV3Session(document, source_sha256=identity.source_sha256)
    session.bind_fenced_source(paragraph, identity, logical_body="fn main() {}\n")
    session.bind_ordinary_anchor((paragraph._p,), {"block_kind": "fenced_block", "id": "raw-code"})
    output = _write(session, document, tmp_path / "ordinary-anchor.docx")

    with ZipFile(output) as package:
        root = etree.fromstring(package.read("word/document.xml"))
    fence = next(item for item in root.iter(qn("w:sdt")) if (_sdt_tag(item) or "").startswith(FENCED_SOURCE_TAG_PREFIX))
    owners = list(fence.iterancestors(qn("w:sdt")))
    assert len(owners) == 1 and (_sdt_tag(owners[0]) or "").startswith("docwen-anchor-v1:")
    assert not list(root.iter(qn("w:bookmarkStart")))
    assert not list(root.iter(qn("w:instrText")))


def test_session_canonicalizes_tabs_and_empty_lines_from_authenticated_body(tmp_path: Path) -> None:
    source = "```text\n\talpha\n\n omega \n```\n"
    logical_body = "\talpha\n\n omega \n"
    identity = _identity(
        source,
        logical_body,
        fence_character="`",
        opening_length=3,
        opening_prefix="",
        info="text",
        opening_eol="\n",
        body_prefixes=("", "", ""),
        closing_state="present",
        closing_length=3,
        closing_prefix="",
        closing_suffix="",
        closing_eol="\n",
    )
    document = Document()
    paragraph = document.add_paragraph("renderer placeholder")
    session = DocxSemanticsV3Session(document, source_sha256=identity.source_sha256)
    session.bind_fenced_source(paragraph, identity, logical_body=logical_body)
    output = _write(session, document, tmp_path / "tabs-empty-lines.docx")

    loaded = Document(str(output))
    recovery = DocxSemanticsV3Recovery.load(output, loaded)
    recovered_paragraph = next(cast(Any, loaded.element).body.iter(qn("w:p")))
    assert recovery.render_fenced_source(recovered_paragraph) == source
    with ZipFile(output) as package:
        root = etree.fromstring(package.read("word/document.xml"))
    carrier = next(
        item for item in root.iter(qn("w:sdt")) if (_sdt_tag(item) or "").startswith(FENCED_SOURCE_TAG_PREFIX)
    )
    assert len(list(carrier.iter(qn("w:tab")))) == 1
    assert len(list(carrier.iter(qn("w:br")))) == 3


def test_internal_crlf_and_lf_are_distinct_authenticated_body_breaks(tmp_path: Path) -> None:
    source = "```text\nalpha\r\nbeta\ngamma\n```\n"
    logical_body = "alpha\r\nbeta\ngamma\n"
    identity = _identity(
        source,
        logical_body,
        fence_character="`",
        opening_length=3,
        opening_prefix="",
        info="text",
        opening_eol="\n",
        body_prefixes=("", "", ""),
        closing_state="present",
        closing_length=3,
        closing_prefix="",
        closing_suffix="",
        closing_eol="\n",
    )
    document = Document()
    paragraph = document.add_paragraph("placeholder")
    session = DocxSemanticsV3Session(document, source_sha256=identity.source_sha256)
    session.bind_fenced_source(paragraph, identity, logical_body=logical_body)
    output = _write(session, document, tmp_path / "mixed-internal-eol.docx")

    loaded = Document(str(output))
    recovery = DocxSemanticsV3Recovery.load(output, loaded)
    recovered_paragraph = next(cast(Any, loaded.element).body.iter(qn("w:p")))
    assert recovery.render_fenced_source(recovered_paragraph) == source
    with ZipFile(output) as package:
        root = etree.fromstring(package.read("word/document.xml"))
    assert len(list(root.iter(qn("w:cr")))) == 1
    assert len(list(root.iter(qn("w:br")))) == 2


def test_captioned_code_uses_same_independent_inner_carrier(tmp_path: Path) -> None:
    source = "Code: Demo ^code-demo\n\n```python\nprint(1)\n```\n"
    start = source.index("```")
    block = source[start:]
    identity = _identity(
        block,
        "print(1)\n",
        fence_character="`",
        opening_length=3,
        opening_prefix="",
        info="python",
        opening_eol="\n",
        body_prefixes=("",),
        closing_state="present",
        closing_length=3,
        closing_prefix="",
        closing_suffix="",
        closing_eol="\n",
        source_document=source,
        source_start=start,
    )
    document = Document()
    bindings, code_style = _caption_styles(document)
    session = DocxSemanticsV3Session(
        document,
        source_sha256=identity.source_sha256,
        caption_style_bindings=bindings,
    )
    caption = document.add_paragraph(style=code_style)
    caption.add_run("Code ")
    append_complex_field(caption, instruction=" SEQ Code \\* ARABIC ", cached_result="1")
    caption.add_run(": Demo")
    paragraph = _body_paragraph(document, "print(1)")
    session.bind_fenced_source(paragraph, identity, logical_body="print(1)\n")
    session.bind_caption(
        caption,
        (paragraph._p,),
        {"kind": "code_block", "id": "code-demo", "number": "1", "title": "Demo"},
    )
    output = _write(session, document, tmp_path / "captioned-code.docx")

    loaded = Document(str(output))
    recovery = DocxSemanticsV3Recovery.load(output, loaded)
    code_paragraph = next(
        item
        for item in cast(Any, loaded.element).body.iter(qn("w:p"))
        if recovery.render_fenced_source(item) is not None
    )
    assert recovery.render_fenced_source(code_paragraph) == block
    assert recovery.caption_signatures == (("code_block", "code-demo", "Demo", "1"),)


def test_map_is_closed_canonical_and_rejects_scalar_tampering() -> None:
    source = "``` rust\nbody\n```"
    identity = _identity(
        source,
        "body\n",
        fence_character="`",
        opening_length=3,
        opening_prefix="",
        info=" rust",
        opening_eol="\n",
        body_prefixes=("",),
        closing_state="present",
        closing_length=3,
        closing_prefix="",
        closing_suffix="",
        closing_eol="",
    )
    data = fenced_source_map_xml([identity])
    root = etree.fromstring(data)
    assert parse_fenced_source_map(root) == [identity]
    assert fenced_source_map_xml(parse_fenced_source_map(root)) == data

    mapping = fenced_source_mapping_v3(identity)
    mapping["info_b64"] = "!!!!"
    with pytest.raises(DocxSemanticsV3Error, match="base64"):
        fenced_source_identity_from_mapping_v3(mapping)

    mapping = fenced_source_mapping_v3(identity)
    mapping["opening_length"] = "03"
    with pytest.raises(DocxSemanticsV3Error, match="canonical decimal"):
        fenced_source_identity_from_mapping_v3(mapping)

    mapping = fenced_source_mapping_v3(identity)
    mapping["source_end"] = "9" * 5000
    with pytest.raises(DocxSemanticsV3Error, match="canonical decimal"):
        fenced_source_identity_from_mapping_v3(mapping)

    mapping = fenced_source_mapping_v3(identity)
    mapping["opening_prefix_b64"] = "dGV4dCA="
    with pytest.raises(DocxSemanticsV3Error, match="non-container"):
        fenced_source_identity_from_mapping_v3(mapping)

    mapping = fenced_source_mapping_v3(identity)
    mapping["closing_state"] = "omitted_eof"
    mapping["closing_length"] = 0
    mapping["closing_suffix_b64"] = "IA=="
    with pytest.raises(DocxSemanticsV3Error, match="must not synthesize"):
        fenced_source_identity_from_mapping_v3(mapping)

    mapping = fenced_source_mapping_v3(identity)
    mapping["source_id"] = "forbidden"
    with pytest.raises(DocxSemanticsV3Error, match="fields are not closed"):
        fenced_source_identity_from_mapping_v3(mapping)

    root = etree.fromstring(data)
    root[0].set("unknown", "x")
    with pytest.raises(DocxSemanticsV3Error, match="record is not closed"):
        parse_fenced_source_map(root)

    with pytest.raises(DocxSemanticsV3Error, match="at least one"):
        fenced_source_map_xml([])


def test_visible_body_and_complete_payload_tampering_fail_closed(tmp_path: Path) -> None:
    session, output = _simple_package(tmp_path)

    def change_body(root: etree._Element) -> None:
        sdt = next(
            item for item in root.iter(qn("w:sdt")) if (_sdt_tag(item) or "").startswith(FENCED_SOURCE_TAG_PREFIX)
        )
        text_element = sdt.find(f".//{qn('w:t')}")
        assert text_element is not None
        text_element.text = "tampered"

    _mutate_document_xml(output, change_body)
    with pytest.raises(DocxSemanticsV3Error, match="visible body hash"):
        session.prove_package(output)

    second_session, second = _simple_package(tmp_path, name="outside-run.docx")

    def add_outside_payload(root: etree._Element) -> None:
        sdt = next(
            item for item in root.iter(qn("w:sdt")) if (_sdt_tag(item) or "").startswith(FENCED_SOURCE_TAG_PREFIX)
        )
        paragraph = sdt.getparent()
        assert paragraph is not None
        run = etree.Element(qn("w:r"))
        text = etree.SubElement(run, qn("w:t"))
        text.text = "outside"
        paragraph.append(run)

    _mutate_document_xml(second, add_outside_payload)
    with pytest.raises(DocxSemanticsV3Error, match="complete paragraph payload"):
        second_session.prove_package(second)


def test_unmapped_carrier_and_overlapping_source_ranges_fail_closed(tmp_path: Path) -> None:
    source = "```rust\nbody\n```\n"
    identity = _identity(
        source,
        "body\n",
        fence_character="`",
        opening_length=3,
        opening_prefix="",
        info="rust",
        opening_eol="\n",
        body_prefixes=("",),
        closing_state="present",
        closing_length=3,
        closing_prefix="",
        closing_suffix="",
        closing_eol="\n",
    )
    document = Document()
    paragraph = _body_paragraph(document, "body")
    session = DocxSemanticsV3Session(document, source_sha256=identity.source_sha256)
    session.bind_fenced_source(paragraph, identity, logical_body="body\n")
    session.finalize_document()
    output = tmp_path / "unmapped.docx"
    document.save(str(output))
    with pytest.raises(DocxSemanticsV3Error, match="unmapped"):
        DocxSemanticsV3Recovery.load(output, Document(str(output)))

    source = "```\n@[[#Other]]\n```\n"
    identity = _identity(
        source,
        "@[[#Other]]\n",
        fence_character="`",
        opening_length=3,
        opening_prefix="",
        info="",
        opening_eol="\n",
        body_prefixes=("",),
        closing_state="present",
        closing_length=3,
        closing_prefix="",
        closing_suffix="",
        closing_eol="\n",
    )
    document = Document()
    paragraph = _body_paragraph(document, "@[[#Other]]")
    overlap = document.add_paragraph()
    overlap_session = DocxSemanticsV3Session(document, source_sha256=identity.source_sha256)
    overlap_session.bind_fenced_source(paragraph, identity, logical_body="@[[#Other]]\n")
    overlap_session.render_reference(
        overlap,
        {
            "selector_kind": "heading_path",
            "resolution_status": "resolved",
            "resolved_kind": "heading",
            "heading_path": ["Other"],
            "cached_number": "1",
            "raw": "@[[#Other]]",
            "range": {"start": 4, "end": 15},
        },
    )
    overlap_session.finalize_document()
    overlap_path = tmp_path / "overlap.docx"
    document.save(str(overlap_path))
    overlap_session.write_package(overlap_path)
    with pytest.raises(DocxSemanticsV3Error, match="ranges overlap"):
        overlap_session.prove_package(overlap_path)


def _simple_package(tmp_path: Path, *, name: str = "simple.docx") -> tuple[DocxSemanticsV3Session, Path]:
    source = "```rust\nbody\n```\n"
    identity = _identity(
        source,
        "body\n",
        fence_character="`",
        opening_length=3,
        opening_prefix="",
        info="rust",
        opening_eol="\n",
        body_prefixes=("",),
        closing_state="present",
        closing_length=3,
        closing_prefix="",
        closing_suffix="",
        closing_eol="\n",
    )
    document = Document()
    paragraph = _body_paragraph(document, "body")
    session = DocxSemanticsV3Session(document, source_sha256=identity.source_sha256)
    session.bind_fenced_source(paragraph, identity, logical_body="body\n")
    return session, _write(session, document, tmp_path / name)


def _identity(
    block: str,
    logical_body: str,
    *,
    fence_character: str,
    opening_length: int,
    opening_prefix: str,
    info: str,
    opening_eol: str,
    body_prefixes: tuple[str, ...],
    closing_state: Literal["present", "omitted_eof"],
    closing_length: int,
    closing_prefix: str,
    closing_suffix: str,
    closing_eol: str,
    source_document: str | None = None,
    source_start: int = 0,
) -> FencedSourceIdentityV3:
    source = block if source_document is None else source_document
    return derive_fenced_source_identity_v3(
        source_sha256=hashlib.sha256(source.encode()).hexdigest(),
        source_start=source_start,
        source_end=source_start + len(block),
        block_sha256=hashlib.sha256(block.encode()).hexdigest(),
        body_sha256=hashlib.sha256(logical_body.encode()).hexdigest(),
        fence_character=fence_character,
        opening_length=opening_length,
        opening_prefix=opening_prefix,
        info=info,
        opening_eol=opening_eol,
        body_prefixes=body_prefixes,
        closing_state=closing_state,
        closing_length=closing_length,
        closing_prefix=closing_prefix,
        closing_suffix=closing_suffix,
        closing_eol=closing_eol,
    )


def _body_paragraph(document: Any, logical_body: str):
    paragraph = document.add_paragraph()
    lines = logical_body.split("\n")
    for index, line in enumerate(lines):
        if index:
            run = paragraph.add_run()
            run.add_break()
        paragraph.add_run(line)
    return paragraph


def _caption_styles(document: Any):
    definitions = (
        ("figure_caption", "DocWenFigureCaption", "Figure Caption"),
        ("table_caption", "DocWenTableCaption", "Table Caption"),
        ("equation_caption", "DocWenEquationCaption", "Equation Caption"),
        ("code_block_caption", "DocWenCodeBlockCaption", "Code Block Caption"),
    )
    bindings: list[CaptionStyleBindingV3] = []
    code_style: Any = None
    for key, style_id, name in definitions:
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style._element.set(qn("w:styleId"), style_id)
        bindings.append(CaptionStyleBindingV3(key, style_id, name))  # type: ignore[arg-type]
        if key == "code_block_caption":
            code_style = style
    assert code_style is not None
    return tuple(bindings), code_style


def _write(session: DocxSemanticsV3Session, document: Any, output: Path) -> Path:
    session.finalize_document()
    document.save(str(output))
    session.write_package(output)
    session.prove_package(output)
    return output


def _map_part(package: ZipFile) -> str:
    return next(
        name
        for name in package.namelist()
        if name.startswith("customXml/item")
        and name.endswith(".xml")
        and FENCED_SOURCE_MAP_NAMESPACE.encode() in package.read(name)
    )


def _sdt_tag(sdt: etree._Element) -> str | None:
    tag = sdt.find(f"./{qn('w:sdtPr')}/{qn('w:tag')}")
    return None if tag is None else tag.get(qn("w:val"))


def _mutate_document_xml(path: Path, mutate) -> None:
    with ZipFile(path) as package:
        infos = package.infolist()
        payloads = {item.filename: package.read(item.filename) for item in infos}
    root = etree.fromstring(payloads["word/document.xml"])
    mutate(root)
    payloads["word/document.xml"] = etree.tostring(
        root,
        encoding="UTF-8",
        xml_declaration=True,
        standalone=True,
    )
    temporary = path.with_suffix(".mutated.docx")
    with ZipFile(temporary, "w") as output:
        for info in infos:
            output.writestr(info, payloads[info.filename])
    temporary.replace(path)
