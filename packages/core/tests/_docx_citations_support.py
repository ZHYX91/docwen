"""Focused tests for provider-neutral DOCX citation and bibliography semantics."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from docx.document import Document as DocumentType
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml import etree

import docwen_core.docx_semantics as docx_semantics_module
from docwen_core.docx_parsing.document_semantics import (
    BIBLIOGRAPHY_BOOKMARK_NAME,
    encode_bibliography_entry_bookmark,
    encode_citation_bookmark,
)
from docwen_core.docx_semantics import (
    DocxSemanticImporter,
    DocxSemanticRenderer,
    append_bookmark_start,
    append_zero_width_bookmark,
)
from docwen_core.models.semantic_document import (
    SemanticBibliographyEntry,
    SemanticBibliographyFragment,
    SemanticBibliographyRun,
    SemanticCaption,
    SemanticCitationCluster,
    SemanticCitationItem,
    SemanticDocument,
    SemanticDocumentValidationError,
    SemanticParagraph,
    SemanticTable,
    SemanticTableCell,
    SemanticText,
)

pytestmark = pytest.mark.contract


def _citation(
    cluster_id: str = "cluster-one",
    *item_ids: str,
    cached_result: str = "[1]",
) -> SemanticCitationCluster:
    return SemanticCitationCluster(
        cluster_id=cluster_id,
        items=tuple(SemanticCitationItem(item_id) for item_id in (item_ids or ("smith2025",))),
        cached_result=cached_result,
    )


def _render_citation(
    citation: SemanticCitationCluster | None = None,
) -> tuple[DocumentType, SemanticCitationCluster]:
    document = Document()
    rendered_citation = citation or _citation()
    DocxSemanticRenderer(document).render_blocks((SemanticParagraph((rendered_citation,)),))
    return document, rendered_citation


def _citation_field_char(document: DocumentType, field_type: str):
    return next(item for item in document.element.iter(qn("w:fldChar")) if item.get(qn("w:fldCharType")) == field_type)


def _assert_citation_fell_back_to_text(document: DocumentType, expected: str = "[1]") -> None:
    imported = DocxSemanticImporter().import_document(document)
    assert imported.document == SemanticDocument(blocks=(SemanticParagraph((SemanticText(expected),)),))
    assert all(
        not isinstance(inline, SemanticCitationCluster)
        for block in imported.document.blocks
        if isinstance(block, SemanticParagraph)
        for inline in block.inlines
    )


def _render_bibliography() -> tuple[DocumentType, SemanticBibliographyFragment]:
    document = Document()
    document.styles.add_style("Bibliography", WD_STYLE_TYPE.PARAGRAPH)
    document.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
    document.add_paragraph("References")
    anchor = document.add_paragraph("exclusive bibliography anchor")
    fragment = SemanticBibliographyFragment(
        entries=(
            SemanticBibliographyEntry(
                "smith2025",
                (
                    SemanticBibliographyRun("Smith, A. (2025). ", bold=True),
                    SemanticBibliographyRun(
                        "Neutral documents",
                        italic=True,
                        href="https://example.org/neutral-documents",
                    ),
                    SemanticBibliographyRun("."),
                ),
            ),
            SemanticBibliographyEntry(
                "wang2024",
                (SemanticBibliographyRun("Wang, B. (2024). Structured tables."),),
            ),
        )
    )
    DocxSemanticRenderer(document).render_bibliography_fragment(
        fragment,
        placeholder_anchor=anchor,
        fallback_style_id="Bibliography",
        hyperlink_style_id="Hyperlink",
    )
    return document, fragment


def _bookmark_range(document: DocumentType, bookmark_name: str):
    start = next(
        item for item in document.element.iter(qn("w:bookmarkStart")) if item.get(qn("w:name")) == bookmark_name
    )
    bookmark_id = start.get(qn("w:id"))
    end = next(item for item in document.element.iter(qn("w:bookmarkEnd")) if item.get(qn("w:id")) == bookmark_id)
    return start, end


def _remove_element(element) -> None:
    parent = element.getparent()
    assert parent is not None
    parent.remove(element)


__all__ = (
    "BIBLIOGRAPHY_BOOKMARK_NAME",
    "CT",
    "RT",
    "WD_STYLE_TYPE",
    "Any",
    "Document",
    "DocxSemanticImporter",
    "DocxSemanticRenderer",
    "Inches",
    "PackURI",
    "Part",
    "Path",
    "Pt",
    "SemanticBibliographyEntry",
    "SemanticBibliographyFragment",
    "SemanticBibliographyRun",
    "SemanticCaption",
    "SemanticCitationCluster",
    "SemanticDocument",
    "SemanticDocumentValidationError",
    "SemanticParagraph",
    "SemanticTable",
    "SemanticTableCell",
    "SemanticText",
    "_assert_citation_fell_back_to_text",
    "_bookmark_range",
    "_citation",
    "_citation_field_char",
    "_remove_element",
    "_render_bibliography",
    "_render_citation",
    "append_bookmark_start",
    "append_zero_width_bookmark",
    "deepcopy",
    "docx_semantics_module",
    "encode_bibliography_entry_bookmark",
    "encode_citation_bookmark",
    "etree",
    "pytest",
    "pytestmark",
    "qn",
)
