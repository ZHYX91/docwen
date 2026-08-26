"""Independent map and physical proofs for nested ordinary anchors."""

from __future__ import annotations

import hashlib
import re
import uuid
from collections.abc import Callable
from pathlib import Path
from typing import Any
from zipfile import ZipFile

import lxml.etree as etree
import pytest
from docx import Document
from docx.oxml.ns import qn

from docwen_core._docx_semantics_v3_package import inject_custom_xml_parts, semantic_map_xml
from docwen_core._docx_semantics_v3_topology import (
    anchor_topology_map_xml,
    parse_anchor_topology_map,
)
from docwen_core.docx_semantics_v3 import (
    ANCHOR_TAG_PREFIX,
    ANCHOR_TOPOLOGY_MAP_NAMESPACE,
    TARGET_MAP_NAMESPACE,
    TARGET_TAG_PREFIX,
    DocxSemanticsV3Error,
    DocxSemanticsV3Recovery,
    DocxSemanticsV3Session,
    derive_anchor_identity_v3,
    derive_anchor_topology_edge_v3,
)

pytestmark = pytest.mark.contract

_CUSTOM_XML_PROPS_NS = "http://schemas.openxmlformats.org/officeDocument/2006/customXml"


def test_topology_edge_exact_hash_and_canonical_bytes() -> None:
    child = f"{ANCHOR_TAG_PREFIX}{'1' * 32}"
    parent = f"{ANCHOR_TAG_PREFIX}{'2' * 32}"
    edge = derive_anchor_topology_edge_v3(child, parent)

    assert edge.sha256 == "28a486c7939e34bd8d6654ec694c0a7fdbf3f1af2aceb37d76db22d6b01124de"
    assert (
        anchor_topology_map_xml([edge])
        == (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
            f'<documentAnchorTopologyMap xmlns="{ANCHOR_TOPOLOGY_MAP_NAMESPACE}" version="1">'
            f'<edge child_tag="{child}" parent_tag="{parent}" sha256="{edge.sha256}"/>'
            "</documentAnchorTopologyMap>\n"
        ).encode()
    )


def test_equal_range_emits_separate_map_and_outer_parent_wrapper(tmp_path: Path) -> None:
    session, path = _write_equal_range(tmp_path / "equal.docx", outer_kind="block_quote", inner_kind="code_block")
    outer = derive_anchor_identity_v3("block_quote", "outer")
    inner = derive_anchor_identity_v3("code_block", "inner")
    edge = derive_anchor_topology_edge_v3(inner.tag, outer.tag)

    with ZipFile(path) as package:
        topology_number, topology_bytes = _owned_item(package, ANCHOR_TOPOLOGY_MAP_NAMESPACE)
        _target_number, target_bytes = _owned_item(package, TARGET_MAP_NAMESPACE)
        names = set(package.namelist())
        assert {
            f"customXml/itemProps{topology_number}.xml",
            f"customXml/_rels/item{topology_number}.xml.rels",
        }.issubset(names)
        props = etree.fromstring(package.read(f"customXml/itemProps{topology_number}.xml"))
        expected_uuid = (
            "{"
            + str(
                uuid.uuid5(
                    uuid.NAMESPACE_URL,
                    f"{ANCHOR_TOPOLOGY_MAP_NAMESPACE}\0{hashlib.sha256(topology_bytes).hexdigest()}",
                )
            ).upper()
            + "}"
        )
        assert props.get(f"{{{_CUSTOM_XML_PROPS_NS}}}itemID") == expected_uuid
    assert topology_bytes == anchor_topology_map_xml([edge])
    assert target_bytes == semantic_map_xml([], [outer, inner])

    with ZipFile(path) as package:
        document_root = etree.fromstring(package.read("word/document.xml"))
    wrappers = _anchor_wrappers(document_root)
    assert [_sdt_tag(item) for item in wrappers] == [outer.tag, inner.tag]
    assert _sdt_tag(next(wrappers[0].iter(qn("w:sdt")))) == outer.tag
    inner_content = wrappers[1].getparent()
    assert inner_content is not None
    assert inner_content.getparent() is wrappers[0]
    for wrapper in wrappers:
        properties = wrapper.find(qn("w:sdtPr"))
        assert properties is not None
        assert [item.tag for item in properties] == [qn("w:tag")]
    assert not list(document_root.iter(qn("w:bookmarkStart")))
    assert not list(document_root.iter(qn("w:instrText")))
    assert session.prove_package(path).anchor_topology_edges == (edge,)


def test_heading_target_transparently_recovers_one_structured_ordinary_anchor(tmp_path: Path) -> None:
    session, path = _write_heading_with_outer_anchor(tmp_path / "heading-anchor.docx")

    loaded = Document(str(path))
    recovery = DocxSemanticsV3Recovery.load(path, loaded)
    heading = recovery.logical_body_elements(loaded)[0]
    source_anchor = recovery.source_anchor(heading)
    assert source_anchor is not None and source_anchor.source_id == "section"
    assert [group.anchor.source_id for group in recovery.ordinary_anchor_groups(heading)] == ["outer-quote"]
    assert recovery.anchor_topology_edges == ()
    with ZipFile(path) as package:
        root = etree.fromstring(package.read("word/document.xml"))
        assert not any(
            ANCHOR_TOPOLOGY_MAP_NAMESPACE.encode() in package.read(name)
            for name in package.namelist()
            if re.fullmatch(r"customXml/item[1-9][0-9]*\.xml", name)
        )
    target = next(item for item in root.iter(qn("w:sdt")) if (_sdt_tag(item) or "").startswith(TARGET_TAG_PREFIX))
    inner = next(item for item in target.iter(qn("w:sdt")) if (_sdt_tag(item) or "").startswith(ANCHOR_TAG_PREFIX))
    content = inner.getparent()
    assert content is not None and content.getparent() is target
    assert session.prove_package(path).anchor_topology_edges == ()


def test_heading_target_rejects_tampered_multi_block_inner_anchor(tmp_path: Path) -> None:
    session, path = _write_heading_with_outer_anchor(tmp_path / "heading-anchor-tamper.docx")
    _mutate_package(path, _append_paragraph_to_heading_anchor)
    with pytest.raises(DocxSemanticsV3Error, match="Heading target SDT must resolve to exactly one paragraph"):
        session.prove_package(path)


@pytest.mark.parametrize(
    ("outer_kind", "inner_kind"),
    (("block_quote", "list"), ("list", "block_quote")),
)
def test_explicit_source_parent_not_kind_rank_controls_equal_direction(
    tmp_path: Path,
    outer_kind: str,
    inner_kind: str,
) -> None:
    _session, path = _write_equal_range(
        tmp_path / f"{outer_kind}-{inner_kind}.docx",
        outer_kind=outer_kind,
        inner_kind=inner_kind,
    )
    loaded = Document(str(path))
    recovery = DocxSemanticsV3Recovery.load(path, loaded)
    paragraph = recovery.logical_body_elements(loaded)[0]
    assert [group.anchor.block_kind for group in recovery.ordinary_anchor_groups(paragraph)] == [
        inner_kind,
        outer_kind,
    ]


def test_strict_nested_range_requires_direct_source_parent_before_mutation() -> None:
    document = Document()
    first = document.add_paragraph("A")
    second = document.add_paragraph("B")
    session = _session(document)
    session.bind_ordinary_anchor((first._p, second._p), {"block_kind": "list", "id": "outer"})
    session.bind_ordinary_anchor((first._p,), {"block_kind": "paragraph", "id": "inner"})

    with pytest.raises(DocxSemanticsV3Error, match="lacks its authenticated source parent"):
        session.finalize_document()
    body = document.element.find(qn("w:body"))
    assert body is not None
    assert not _anchor_wrappers(body)


@pytest.mark.parametrize(
    ("parents", "message"),
    (
        (("missing", None), "not an authenticated ordinary anchor"),
        (("right", "left"), "cycle"),
    ),
)
def test_writer_rejects_unknown_parent_or_cycle(
    parents: tuple[str | None, str | None],
    message: str,
) -> None:
    document = Document()
    paragraph = document.add_paragraph("same")
    session = _session(document)
    session.bind_ordinary_anchor(
        (paragraph._p,),
        {"block_kind": "block_quote", "id": "left"},
        direct_parent_source_id=parents[0],
    )
    session.bind_ordinary_anchor(
        (paragraph._p,),
        {"block_kind": "list", "id": "right"},
        direct_parent_source_id=parents[1],
    )
    with pytest.raises(DocxSemanticsV3Error, match=message):
        session.finalize_document()


@pytest.mark.parametrize(
    ("mutation", "message"),
    (
        ("strip", "presence differs"),
        ("swap", "edges differ"),
        ("hash", "hash does not recompute"),
        ("uuid", "UUID does not match"),
        ("relationship", "exactly one custom XML relationship"),
    ),
)
def test_recovery_rejects_missing_swap_hash_or_uuid(
    tmp_path: Path,
    mutation: str,
    message: str,
) -> None:
    session, path = _write_equal_range(tmp_path / "tamper.docx", outer_kind="block_quote", inner_kind="code_block")
    mutations: dict[str, Callable[[dict[str, bytes]], None]] = {
        "strip": _strip_topology_item,
        "swap": _swap_equal_wrapper_tags,
        "hash": _tamper_topology_hash,
        "uuid": _tamper_topology_uuid,
        "relationship": _append_wrong_type_owned_relationship,
    }
    _mutate_package(path, mutations[mutation])
    with pytest.raises(DocxSemanticsV3Error, match=message):
        session.prove_package(path)


def test_parser_rejects_empty_second_parent_cycle_and_unknown_endpoint() -> None:
    left = f"{ANCHOR_TAG_PREFIX}{'1' * 32}"
    right = f"{ANCHOR_TAG_PREFIX}{'2' * 32}"
    third = f"{ANCHOR_TAG_PREFIX}{'3' * 32}"
    with pytest.raises(DocxSemanticsV3Error, match="at least one edge"):
        parse_anchor_topology_map(_topology_root([]))
    with pytest.raises(DocxSemanticsV3Error, match="more than one parent"):
        parse_anchor_topology_map(
            _topology_root([derive_anchor_topology_edge_v3(left, right), derive_anchor_topology_edge_v3(left, third)])
        )
    with pytest.raises(DocxSemanticsV3Error, match="cycle"):
        parse_anchor_topology_map(
            _topology_root([derive_anchor_topology_edge_v3(left, right), derive_anchor_topology_edge_v3(right, left)])
        )


def test_recovery_rejects_extra_map_for_flat_ordinary_anchors(tmp_path: Path) -> None:
    document = Document()
    first = document.add_paragraph("A")
    second = document.add_paragraph("B")
    session = _session(document)
    session.bind_paragraph_anchor(first, {"block_kind": "paragraph", "id": "first"})
    session.bind_paragraph_anchor(second, {"block_kind": "paragraph", "id": "second"})
    session.finalize_document()
    path = tmp_path / "extra.docx"
    document.save(str(path))
    identities = [derive_anchor_identity_v3("paragraph", value) for value in ("first", "second")]
    extra = derive_anchor_topology_edge_v3(identities[0].tag, identities[1].tag)
    inject_custom_xml_parts(
        path,
        [
            (TARGET_MAP_NAMESPACE, semantic_map_xml([], identities)),
            (ANCHOR_TOPOLOGY_MAP_NAMESPACE, anchor_topology_map_xml([extra])),
        ],
    )
    with pytest.raises(DocxSemanticsV3Error, match="presence differs from physical nesting"):
        DocxSemanticsV3Recovery.load(path, Document(str(path)))


def test_recovery_rejects_unknown_endpoint_before_physical_guess(tmp_path: Path) -> None:
    document = Document()
    first = document.add_paragraph("A")
    second = document.add_paragraph("B")
    session = _session(document)
    session.bind_paragraph_anchor(first, {"block_kind": "paragraph", "id": "first"})
    session.bind_paragraph_anchor(second, {"block_kind": "paragraph", "id": "second"})
    session.finalize_document()
    path = tmp_path / "unknown.docx"
    document.save(str(path))
    identities = [derive_anchor_identity_v3("paragraph", value) for value in ("first", "second")]
    unknown = f"{ANCHOR_TAG_PREFIX}{'f' * 32}"
    edge = derive_anchor_topology_edge_v3(unknown, identities[0].tag)
    inject_custom_xml_parts(
        path,
        [
            (TARGET_MAP_NAMESPACE, semantic_map_xml([], identities)),
            (ANCHOR_TOPOLOGY_MAP_NAMESPACE, anchor_topology_map_xml([edge])),
        ],
    )
    with pytest.raises(DocxSemanticsV3Error, match="unknown ordinary-anchor endpoint"):
        DocxSemanticsV3Recovery.load(path, Document(str(path)))


def test_recovery_requires_nearest_ordinary_descendant_edge(tmp_path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph("same")
    session = _session(document)
    session.bind_ordinary_anchor((paragraph._p,), {"block_kind": "block_quote", "id": "outer"})
    session.bind_ordinary_anchor(
        (paragraph._p,),
        {"block_kind": "list", "id": "middle"},
        direct_parent_source_id="outer",
    )
    session.bind_ordinary_anchor(
        (paragraph._p,),
        {"block_kind": "paragraph", "id": "inner"},
        direct_parent_source_id="middle",
    )
    session.finalize_document()
    path = tmp_path / "nearest.docx"
    document.save(str(path))
    identities = [
        derive_anchor_identity_v3(kind, source_id)
        for kind, source_id in (("block_quote", "outer"), ("list", "middle"), ("paragraph", "inner"))
    ]
    outer, middle, inner = identities
    wrong_edges = [
        derive_anchor_topology_edge_v3(inner.tag, outer.tag),
        derive_anchor_topology_edge_v3(middle.tag, outer.tag),
    ]
    inject_custom_xml_parts(
        path,
        [
            (TARGET_MAP_NAMESPACE, semantic_map_xml([], identities)),
            (ANCHOR_TOPOLOGY_MAP_NAMESPACE, anchor_topology_map_xml(wrong_edges)),
        ],
    )
    with pytest.raises(DocxSemanticsV3Error, match="nearest physical nesting"):
        DocxSemanticsV3Recovery.load(path, Document(str(path)))


def _write_equal_range(path: Path, *, outer_kind: str, inner_kind: str) -> tuple[DocxSemanticsV3Session, Path]:
    document = Document()
    paragraph = document.add_paragraph("graph TD")
    session = _session(document)
    session.bind_ordinary_anchor((paragraph._p,), {"block_kind": outer_kind, "id": "outer"})
    session.bind_ordinary_anchor(
        (paragraph._p,),
        {"block_kind": inner_kind, "id": "inner"},
        direct_parent_source_id="outer",
    )
    session.finalize_document()
    document.save(str(path))
    session.write_package(path)
    session.prove_package(path)
    return session, path


def _write_heading_with_outer_anchor(path: Path) -> tuple[DocxSemanticsV3Session, Path]:
    document = Document()
    heading = document.add_heading("Section", level=1)
    session = _session(document)
    session.bind_heading(heading, {"kind": "heading", "id": "section"})
    session.bind_ordinary_anchor(
        (heading._p,),
        {"block_kind": "block_quote", "id": "outer-quote"},
    )
    session.finalize_document()
    document.save(str(path))
    session.write_package(path)
    session.prove_package(path)
    return session, path


def _session(document: Any) -> DocxSemanticsV3Session:
    return DocxSemanticsV3Session(document, source_sha256=hashlib.sha256(b"source").hexdigest())


def _owned_item(package: ZipFile, namespace: str) -> tuple[int, bytes]:
    matches = [
        (int(match.group(1)), package.read(name))
        for name in package.namelist()
        if (match := re.fullmatch(r"customXml/item([1-9][0-9]*)\.xml", name)) is not None
        and namespace.encode() in package.read(name)
    ]
    assert len(matches) == 1
    return matches[0]


def _anchor_wrappers(root: etree._Element) -> list[etree._Element]:
    return [item for item in root.iter(qn("w:sdt")) if (_sdt_tag(item) or "").startswith(ANCHOR_TAG_PREFIX)]


def _sdt_tag(sdt: etree._Element) -> str | None:
    item = sdt.find(f"./{qn('w:sdtPr')}/{qn('w:tag')}")
    return None if item is None else item.get(qn("w:val"))


def _topology_root(edges: list[Any]) -> etree._Element:
    return (
        etree.fromstring(anchor_topology_map_xml(edges))
        if edges
        else etree.fromstring(f'<documentAnchorTopologyMap xmlns="{ANCHOR_TOPOLOGY_MAP_NAMESPACE}" version="1"/>')
    )


def _strip_topology_item(payloads: dict[str, bytes]) -> None:
    item_name = _topology_item_name(payloads)
    number = re.search(r"([1-9][0-9]*)", item_name)
    assert number is not None
    item_number = number.group(1)
    for name in (
        item_name,
        f"customXml/itemProps{item_number}.xml",
        f"customXml/_rels/item{item_number}.xml.rels",
    ):
        del payloads[name]
    rels = etree.fromstring(payloads["word/_rels/document.xml.rels"])
    relation = next(item for item in rels if item.get("Target") == f"../customXml/item{item_number}.xml")
    rels.remove(relation)
    payloads["word/_rels/document.xml.rels"] = etree.tostring(
        rels, encoding="UTF-8", xml_declaration=True, standalone=True
    )
    content_types = etree.fromstring(payloads["[Content_Types].xml"])
    for item in list(content_types):
        if item.get("PartName") in {
            f"/customXml/item{item_number}.xml",
            f"/customXml/itemProps{item_number}.xml",
        }:
            content_types.remove(item)
    payloads["[Content_Types].xml"] = etree.tostring(
        content_types, encoding="UTF-8", xml_declaration=True, standalone=True
    )


def _swap_equal_wrapper_tags(payloads: dict[str, bytes]) -> None:
    root = etree.fromstring(payloads["word/document.xml"])
    outer, inner = _anchor_wrappers(root)
    outer_tag = outer.find(f"./{qn('w:sdtPr')}/{qn('w:tag')}")
    inner_tag = inner.find(f"./{qn('w:sdtPr')}/{qn('w:tag')}")
    assert outer_tag is not None and inner_tag is not None
    outer_value = outer_tag.get(qn("w:val"))
    inner_value = inner_tag.get(qn("w:val"))
    assert outer_value is not None and inner_value is not None
    outer_tag.set(qn("w:val"), inner_value)
    inner_tag.set(qn("w:val"), outer_value)
    payloads["word/document.xml"] = etree.tostring(root, encoding="UTF-8", xml_declaration=True, standalone=True)


def _tamper_topology_hash(payloads: dict[str, bytes]) -> None:
    name = _topology_item_name(payloads)
    payloads[name] = re.sub(rb'sha256="[0-9a-f]{64}"', b'sha256="' + (b"0" * 64) + b'"', payloads[name])


def _tamper_topology_uuid(payloads: dict[str, bytes]) -> None:
    item_name = _topology_item_name(payloads)
    number = re.search(r"([1-9][0-9]*)", item_name)
    assert number is not None
    props_name = f"customXml/itemProps{number.group(1)}.xml"
    payloads[props_name] = re.sub(
        rb'ds:itemID="\{[0-9A-F-]{36}\}"',
        b'ds:itemID="{00000000-0000-0000-0000-000000000000}"',
        payloads[props_name],
    )


def _append_wrong_type_owned_relationship(payloads: dict[str, bytes]) -> None:
    item_name = _topology_item_name(payloads)
    number = re.search(r"([1-9][0-9]*)", item_name)
    assert number is not None
    name = "word/_rels/document.xml.rels"
    root = etree.fromstring(payloads[name])
    relationship = etree.Element(f"{{{etree.QName(root).namespace}}}Relationship")
    relationship.set("Id", "rId999999")
    relationship.set("Type", "urn:docwen:test:wrong-type")
    relationship.set("Target", f"../customXml/item{number.group(1)}.xml")
    root.append(relationship)
    payloads[name] = etree.tostring(root, encoding="UTF-8", xml_declaration=True, standalone=True)


def _append_paragraph_to_heading_anchor(payloads: dict[str, bytes]) -> None:
    name = "word/document.xml"
    root = etree.fromstring(payloads[name])
    target = next(item for item in root.iter(qn("w:sdt")) if (_sdt_tag(item) or "").startswith(TARGET_TAG_PREFIX))
    inner = next(item for item in target.iter(qn("w:sdt")) if (_sdt_tag(item) or "").startswith(ANCHOR_TAG_PREFIX))
    content = inner.find(qn("w:sdtContent"))
    assert content is not None
    content.append(etree.Element(qn("w:p")))
    payloads[name] = etree.tostring(root, encoding="UTF-8", xml_declaration=True, standalone=True)


def _topology_item_name(payloads: dict[str, bytes]) -> str:
    names = [
        name
        for name, data in payloads.items()
        if re.fullmatch(r"customXml/item[1-9][0-9]*\.xml", name) and ANCHOR_TOPOLOGY_MAP_NAMESPACE.encode() in data
    ]
    assert len(names) == 1
    return names[0]


def _mutate_package(path: Path, mutate: Callable[[dict[str, bytes]], None]) -> None:
    with ZipFile(path) as package:
        infos = package.infolist()
        payloads = {item.filename: package.read(item.filename) for item in infos}
    mutate(payloads)
    temporary = path.with_suffix(".mutated.docx")
    with ZipFile(temporary, "w") as output:
        for info in infos:
            if info.filename in payloads:
                output.writestr(info, payloads[info.filename])
    temporary.replace(path)
