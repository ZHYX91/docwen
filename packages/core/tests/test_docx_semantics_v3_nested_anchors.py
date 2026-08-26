"""Exact containment proofs for nested ordinary-anchor block SDTs."""

from __future__ import annotations

import copy
import hashlib
from collections.abc import Callable
from pathlib import Path
from typing import Any
from zipfile import ZipFile

import lxml.etree as etree
import pytest
from docx import Document
from docx.oxml.ns import qn

from docwen_core.docx_semantics_v3 import (
    ANCHOR_TAG_PREFIX,
    DocxSemanticsV3Error,
    DocxSemanticsV3Recovery,
    DocxSemanticsV3Session,
)

pytestmark = pytest.mark.contract


def test_multi_paragraph_item_and_whole_list_recover_inner_to_outer_once(tmp_path: Path) -> None:
    first_session, first_path = _write_nested_list(tmp_path / "outer-first.docx", outer_first=True)
    _second_session, second_path = _write_nested_list(tmp_path / "inner-first.docx", outer_first=False)

    loaded = Document(str(first_path))
    recovery = DocxSemanticsV3Recovery.load(first_path, loaded)
    paragraphs = [item for item in recovery.logical_body_elements(loaded) if item.tag == qn("w:p")]
    assert len(paragraphs) == 3
    chains = [tuple(group.anchor.source_id for group in recovery.ordinary_anchor_groups(item)) for item in paragraphs]
    assert chains == [
        ("first-item", "whole-list"),
        ("first-item", "whole-list"),
        ("whole-list",),
    ]
    assert recovery.ordinary_anchor_group(paragraphs[0]).anchor.source_id == "first-item"  # type: ignore[union-attr]
    assert recovery.source_anchor(paragraphs[0]).source_id == "first-item"  # type: ignore[union-attr]
    assert recovery.source_anchor(paragraphs[2]).source_id == "whole-list"  # type: ignore[union-attr]

    emitted: list[str] = []
    for paragraph in paragraphs:
        for group in recovery.ordinary_anchor_groups(paragraph):
            if group.index == len(group.elements) - 1:
                emitted.append(group.anchor.source_id)
    assert emitted == ["first-item", "whole-list"]

    with ZipFile(first_path) as package:
        first_root = etree.fromstring(package.read("word/document.xml"))
    with ZipFile(second_path) as package:
        second_root = etree.fromstring(package.read("word/document.xml"))
    assert _anchor_signature(first_root) == _anchor_signature(second_root)
    assert not list(first_root.iter(qn("w:bookmarkStart")))
    assert not list(first_root.iter(qn("w:bookmarkEnd")))
    assert not list(first_root.iter(qn("w:instrText")))
    assert not list(first_root.iter(qn("w:fldChar")))
    assert first_session.prove_package(first_path).anchor_identities == recovery.anchor_identities


def test_nested_quote_and_list_preserve_both_exact_groups(tmp_path: Path) -> None:
    document = Document()
    first = document.add_paragraph("Quoted item A", style="List Bullet")
    second = document.add_paragraph("Quoted item B", style="List Bullet")
    tail = document.add_paragraph("Quoted tail")
    session = DocxSemanticsV3Session(document, source_sha256=hashlib.sha256(b"> - A\n> - B\n> tail").hexdigest())
    session.bind_ordinary_anchor(
        (first._p, second._p),
        {"block_kind": "list", "id": "quoted-list"},
        direct_parent_source_id="whole-quote",
    )
    session.bind_ordinary_anchor(
        (first._p, second._p, tail._p),
        {"block_kind": "block_quote", "id": "whole-quote"},
    )
    output = _write(session, document, tmp_path / "quote-list.docx")

    loaded = Document(str(output))
    recovery = DocxSemanticsV3Recovery.load(output, loaded)
    paragraphs = [item for item in recovery.logical_body_elements(loaded) if item.tag == qn("w:p")]
    assert [
        tuple((group.anchor.block_kind, group.anchor.source_id) for group in recovery.ordinary_anchor_groups(item))
        for item in paragraphs
    ] == [
        (("list", "quoted-list"), ("block_quote", "whole-quote")),
        (("list", "quoted-list"), ("block_quote", "whole-quote")),
        (("block_quote", "whole-quote"),),
    ]


@pytest.mark.parametrize(
    ("left", "right", "message"),
    (
        ((0, 1), (1, 2), "partially overlap"),
        ((0, 1), (0, 1), "duplicate ownership"),
    ),
)
def test_writer_rejects_partial_or_duplicate_ownership_before_mutation(
    left: tuple[int, int],
    right: tuple[int, int],
    message: str,
) -> None:
    document = Document()
    paragraphs = [document.add_paragraph(value) for value in ("A", "B", "C")]
    session = DocxSemanticsV3Session(document, source_sha256=hashlib.sha256(b"A\nB\nC").hexdigest())
    session.bind_ordinary_anchor(
        tuple(paragraphs[index]._p for index in range(left[0], left[1] + 1)),
        {"block_kind": "list", "id": "left"},
    )
    session.bind_ordinary_anchor(
        tuple(paragraphs[index]._p for index in range(right[0], right[1] + 1)),
        {"block_kind": "block_quote", "id": "right"},
    )
    with pytest.raises(DocxSemanticsV3Error, match=message):
        session.finalize_document()
    body = document.element.find(qn("w:body"))
    assert body is not None
    assert not [item for item in body.iter(qn("w:sdt")) if _tag(item)]


@pytest.mark.parametrize(
    ("mutation", "message"),
    (
        ("swap", "wrong block kind or cardinality"),
        ("tamper", "invalid nesting owner"),
        ("unmapped", "no authenticated map record"),
        ("duplicate", "duplicate ordinary-anchor"),
    ),
)
def test_recovery_rejects_swapped_tampered_unmapped_or_duplicate_nested_carriers(
    tmp_path: Path,
    mutation: str,
    message: str,
) -> None:
    session, output = _write_nested_table_list(tmp_path / f"{mutation}.docx")
    mutations: dict[str, Callable[[etree._Element], None]] = {
        "swap": _swap_anchor_tags,
        "tamper": _move_inner_anchor_under_paragraph,
        "unmapped": _append_unmapped_anchor,
        "duplicate": _append_duplicate_anchor,
    }
    _mutate_document_xml(output, mutations[mutation])
    with pytest.raises(DocxSemanticsV3Error, match=message):
        session.prove_package(output)


def _write_nested_list(path: Path, *, outer_first: bool) -> tuple[DocxSemanticsV3Session, Path]:
    document = Document()
    first = document.add_paragraph("A", style="List Bullet")
    continuation = document.add_paragraph("A continuation")
    second = document.add_paragraph("B", style="List Bullet")
    session = DocxSemanticsV3Session(document, source_sha256=hashlib.sha256(b"- A\n  continuation\n- B").hexdigest())
    bindings = (
        ((first._p, continuation._p, second._p), {"block_kind": "list", "id": "whole-list"}, None),
        (
            (first._p, continuation._p),
            {"block_kind": "list_item", "id": "first-item"},
            "whole-list",
        ),
    )
    if not outer_first:
        bindings = tuple(reversed(bindings))
    for elements, anchor, direct_parent in bindings:
        session.bind_ordinary_anchor(elements, anchor, direct_parent_source_id=direct_parent)
    return session, _write(session, document, path)


def _write_nested_table_list(path: Path) -> tuple[DocxSemanticsV3Session, Path]:
    document = Document()
    paragraph = document.add_paragraph("List item", style="List Bullet")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "nested"
    session = DocxSemanticsV3Session(document, source_sha256=hashlib.sha256(b"- item\n  | nested |").hexdigest())
    session.bind_ordinary_anchor((paragraph._p, table._tbl), {"block_kind": "list", "id": "whole-list"})
    session.bind_ordinary_anchor(
        (table._tbl,),
        {"block_kind": "table", "id": "nested-table"},
        direct_parent_source_id="whole-list",
    )
    return session, _write(session, document, path)


def _write(session: DocxSemanticsV3Session, document: Any, path: Path) -> Path:
    session.finalize_document()
    document.save(str(path))
    session.write_package(path)
    session.prove_package(path)
    return path


def _owned_anchor_sdts(root: etree._Element) -> list[etree._Element]:
    return [item for item in root.iter(qn("w:sdt")) if (_tag(item) or "").startswith(ANCHOR_TAG_PREFIX)]


def _tag(sdt: etree._Element) -> str | None:
    item = sdt.find(f"./{qn('w:sdtPr')}/{qn('w:tag')}")
    return None if item is None else item.get(qn("w:val"))


def _anchor_signature(root: etree._Element) -> tuple[Any, ...]:
    def visit(item: etree._Element) -> tuple[Any, ...]:
        if item.tag == qn("w:sdt") and (_tag(item) or "").startswith(ANCHOR_TAG_PREFIX):
            content = item.find(qn("w:sdtContent"))
            assert content is not None
            return (_tag(item), *(visit(child) for child in content))
        if item.tag == qn("w:p"):
            return ("p", "".join(child.text or "" for child in item.iter(qn("w:t"))))
        if item.tag == qn("w:tbl"):
            return ("tbl",)
        return (etree.QName(item).localname,)

    body = next(root.iter(qn("w:body")))
    return tuple(visit(item) for item in body)


def _swap_anchor_tags(root: etree._Element) -> None:
    outer, inner = _owned_anchor_sdts(root)
    outer_tag = outer.find(f"./{qn('w:sdtPr')}/{qn('w:tag')}")
    inner_tag = inner.find(f"./{qn('w:sdtPr')}/{qn('w:tag')}")
    assert outer_tag is not None and inner_tag is not None
    outer_value = outer_tag.get(qn("w:val"))
    inner_value = inner_tag.get(qn("w:val"))
    assert outer_value is not None and inner_value is not None
    outer_tag.set(qn("w:val"), inner_value)
    inner_tag.set(qn("w:val"), outer_value)


def _move_inner_anchor_under_paragraph(root: etree._Element) -> None:
    outer, inner = _owned_anchor_sdts(root)
    paragraph = next(outer.iter(qn("w:p")))
    parent = inner.getparent()
    assert parent is not None
    parent.remove(inner)
    paragraph.append(inner)


def _append_unmapped_anchor(root: etree._Element) -> None:
    body = next(root.iter(qn("w:body")))
    clone = copy.deepcopy(_owned_anchor_sdts(root)[-1])
    tag = clone.find(f"./{qn('w:sdtPr')}/{qn('w:tag')}")
    assert tag is not None
    tag.set(qn("w:val"), f"{ANCHOR_TAG_PREFIX}{'0' * 32}")
    body.insert(len(body) - 1, clone)


def _append_duplicate_anchor(root: etree._Element) -> None:
    body = next(root.iter(qn("w:body")))
    body.insert(len(body) - 1, copy.deepcopy(_owned_anchor_sdts(root)[-1]))


def _mutate_document_xml(path: Path, mutate: Callable[[etree._Element], None]) -> None:
    with ZipFile(path) as package:
        infos = package.infolist()
        payloads = {item.filename: package.read(item.filename) for item in infos}
    root = etree.fromstring(payloads["word/document.xml"])
    mutate(root)
    payloads["word/document.xml"] = etree.tostring(root, encoding="UTF-8", xml_declaration=True, standalone=True)
    temporary = path.with_suffix(".mutated.docx")
    with ZipFile(temporary, "w") as output:
        for info in infos:
            output.writestr(info, payloads[info.filename])
    temporary.replace(path)
