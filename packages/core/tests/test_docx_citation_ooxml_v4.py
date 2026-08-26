"""Closed v4 resolved-citation physical authority gates."""

from __future__ import annotations

import hashlib
from dataclasses import replace
from pathlib import Path
from typing import Any, cast
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import lxml.etree as etree
import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docwen_core._docx_semantics_v3_model import CaptionStyleBindingV3, CaptionStyleKeyV3
from docwen_core._docx_semantics_v3_package import read_owned_map_parts
from docwen_core.docx_citation_ooxml import (
    CITATION_ITEM_MAP_NAMESPACE,
    CITATION_OCCURRENCE_MAP_NAMESPACE,
    CitationItemMap,
    CitationOccurrenceMap,
    ResolvedCitationOoxmlError,
    build_resolved_citation_projection,
    citation_instruction,
    citation_item_map_xml,
    citation_occurrence_map_xml,
    derive_citation_item,
    derive_citation_item_reference,
    derive_citation_occurrence,
    parse_citation_item_map,
    parse_citation_occurrence_map,
    preflight_citation_document,
    read_proven_resolved_citations,
    validate_citation_authorities,
    validate_citation_item_map,
    validate_citation_occurrence_map,
)
from docwen_core.docx_resolved_numbering import ResolvedNumberingDocxSession
from docwen_core.models.resolved_numbering import (
    NumberingExportPlanEnvelope,
    ResolvedCitation,
    ResolvedCitationItem,
    ResolvedDocument,
    ResolvedDocumentEnvelope,
    ResolvedNumberingPlan,
    ResolvedNumberingPort,
)

pytestmark = pytest.mark.unit


def _sha(value: str) -> str:
    return hashlib.sha256(value.encode()).hexdigest()


def _citations(source: str) -> tuple[ResolvedCitation, ...]:
    narrative = "@fig-legacy"
    cluster = "[@smith; @wang]"

    def occurrence(
        token: str,
        form: str,
        cluster_id: str,
        items: tuple[ResolvedCitationItem, ...],
        cache: str,
    ) -> ResolvedCitation:
        start = source.index(token)
        return ResolvedCitation(
            source_start=start,
            source_end=start + len(token),
            source_slice_sha256=_sha(token),
            authored_token=token,
            form=form,  # type: ignore[arg-type]
            cluster_id=cluster_id,
            items=items,
            cached_result=cache,
        )

    legacy = ResolvedCitationItem(
        citation_key="fig-legacy",
        record_id="reference-record:98",
        record_sha256=_sha("record-v98"),
        presentation="Legacy Figure Record",
    )
    smith = ResolvedCitationItem("smith", "record:smith", _sha("smith-v1"), "Smith (2024)")
    wang = ResolvedCitationItem("wang", "record:wang", _sha("wang-v2"), "Wang (2025)")
    return (
        occurrence(narrative, "narrative", "cluster-a", (legacy,), "Legacy Figure Record"),
        occurrence(cluster, "parenthetical", "cluster-b", (smith, wang), "(Smith, 2024; Wang, 2025)"),
    )


def _port(source: str) -> ResolvedNumberingPort:
    citations = _citations(source)
    source_sha = _sha(source)
    plan_sha = "a" * 64
    document = ResolvedDocument(source, (), (), (), citations, ())
    plan = ResolvedNumberingPlan((), (), ())
    return ResolvedNumberingPort(
        ResolvedDocumentEnvelope("source", source_sha, plan_sha, document),
        NumberingExportPlanEnvelope("source", source_sha, plan_sha, plan),
    )


def _caption_bindings(document: Any) -> tuple[CaptionStyleBindingV3, ...]:
    output: list[CaptionStyleBindingV3] = []
    for semantic_key, style_id, name in (
        ("figure_caption", "DWFigureCaption", "Figure Caption V4"),
        ("table_caption", "DWTableCaption", "Table Caption V4"),
        ("equation_caption", "DWEquationCaption", "Equation Caption V4"),
        ("code_block_caption", "DWCodeCaption", "Code Caption V4"),
    ):
        style = document.styles.add_style(style_id, WD_STYLE_TYPE.PARAGRAPH)
        style.name = name
        output.append(CaptionStyleBindingV3(cast(CaptionStyleKeyV3, semantic_key), style.style_id, style.name))
    return tuple(output)


def _projection():
    source = "@fig-legacy and [@smith; @wang]\n"
    projection = build_resolved_citation_projection(_sha(source), _citations(source))
    assert projection is not None
    return source, projection


def test_exact_item_and_occurrence_preimages_keep_provider_identity() -> None:
    source, projection = _projection()
    legacy = next(item for item in projection.item_map.items if item.record_id == "reference-record:98")
    presentation_sha = _sha("Legacy Figure Record")
    expected_item_digest = _sha(
        "\0".join(
            (
                "docwen-citation-item-map-v1",
                _sha(source),
                "reference-record:98",
                _sha("record-v98"),
                presentation_sha,
            )
        )
    )

    assert legacy.presentation_sha256 == presentation_sha
    assert legacy.sha256 == expected_item_digest
    assert legacy.word_tag == f"DWCIT_{expected_item_digest[:32]}"
    assert len(legacy.word_tag) == 38
    first = projection.occurrence_map.occurrences[0]
    ref_digest = _sha(
        "\0".join(
            (
                "docwen-citation-item-ref-v1",
                "fig-legacy",
                legacy.word_tag,
                legacy.sha256,
            )
        )
    )
    expected_occurrence_digest = _sha(
        "\0".join(
            (
                "docwen-citation-occurrence-map-v1",
                _sha(source),
                str(first.source_start),
                str(first.source_end),
                _sha("@fig-legacy"),
                "narrative",
                "cluster-a",
                _sha("Legacy Figure Record"),
                ref_digest,
            )
        )
    )
    assert first.item_refs[0].sha256 == ref_digest
    assert first.sha256 == expected_occurrence_digest
    assert first.tag == f"docwen-citation-occurrence-v1:{expected_occurrence_digest[:32]}"
    assert first.bookmark_name == f"_DWC_{expected_occurrence_digest[:35]}"
    assert len(first.bookmark_name) == 40


def test_both_maps_have_exact_canonical_roundtrip_and_nonempty_shape() -> None:
    _source, projection = _projection()
    item_bytes = citation_item_map_xml(projection.item_map)
    occurrence_bytes = citation_occurrence_map_xml(projection.occurrence_map)

    assert item_bytes.count(b"\n") == 2
    assert occurrence_bytes.count(b"\n") == 2
    assert b'record_id="reference-record:98"' in item_bytes
    assert parse_citation_item_map(etree.fromstring(item_bytes)) == projection.item_map
    assert parse_citation_occurrence_map(etree.fromstring(occurrence_bytes)) == projection.occurrence_map
    with pytest.raises(ResolvedCitationOoxmlError, match="must not be empty"):
        citation_item_map_xml(CitationItemMap(projection.item_map.source_sha256, ()))
    with pytest.raises(ResolvedCitationOoxmlError, match="must not be empty"):
        citation_occurrence_map_xml(CitationOccurrenceMap(projection.occurrence_map.source_sha256, ()))


def test_exact_tuple_deduplicates_but_record_versions_and_presentations_do_not() -> None:
    source = "@one @two @three\n"
    shared = ResolvedCitationItem("one", "record:one", _sha("v1"), "One")
    same = replace(shared, citation_key="two")
    changed = replace(shared, citation_key="three", record_sha256=_sha("v2"), presentation="One revised")
    citations = tuple(
        ResolvedCitation(
            source_start=source.index(token),
            source_end=source.index(token) + len(token),
            source_slice_sha256=_sha(token),
            authored_token=token,
            form="narrative",
            cluster_id=f"cluster-{index}",
            items=(item,),
            cached_result=item.presentation,
        )
        for index, (token, item) in enumerate((("@one", shared), ("@two", same), ("@three", changed)), 1)
    )
    projection = build_resolved_citation_projection(_sha(source), citations)
    assert projection is not None

    assert len(projection.item_map.items) == 2
    assert len({item.word_tag for item in projection.item_map.items}) == 2


@pytest.mark.parametrize("collision", ["full", "truncated"])
def test_item_digest_and_word_tag_collisions_fail_closed(collision: str) -> None:
    _source, projection = _projection()
    first, second = projection.item_map.items[:2]
    if collision == "full":
        second = replace(second, sha256=first.sha256, word_tag=first.word_tag)
        match = "full digest collides"
    else:
        second = replace(second, word_tag=first.word_tag)
        match = "truncated Word tag collides"
    colliding = CitationItemMap(
        projection.item_map.source_sha256, tuple(sorted((first, second), key=lambda x: x.word_tag))
    )

    with pytest.raises(ResolvedCitationOoxmlError, match=match):
        validate_citation_item_map(colliding)


def test_dangling_and_unused_item_records_fail_closed() -> None:
    _source, projection = _projection()
    unused_map = CitationItemMap(
        projection.item_map.source_sha256,
        (
            *projection.item_map.items,
            derive_citation_item(
                source_sha256=projection.item_map.source_sha256,
                record_id="record:unused",
                record_sha256=_sha("unused"),
                presentation="Unused",
            ),
        ),
    )
    unused_map = CitationItemMap(
        unused_map.source_sha256, tuple(sorted(unused_map.items, key=lambda item: item.word_tag))
    )
    with pytest.raises(ResolvedCitationOoxmlError, match="unused record"):
        validate_citation_authorities(unused_map, projection.occurrence_map)

    fake = derive_citation_item(
        source_sha256=projection.item_map.source_sha256,
        record_id="record:missing",
        record_sha256=_sha("missing"),
        presentation="Missing",
    )
    fake_ref = derive_citation_item_reference(
        citation_key="fig-legacy", word_tag=fake.word_tag, item_sha256=fake.sha256
    )
    old = projection.occurrence_map.occurrences[0]
    dangling = derive_citation_occurrence(
        source_sha256=old.source_sha256,
        source_start=old.source_start,
        source_end=old.source_end,
        source_slice_sha256=old.source_slice_sha256,
        authored_token=old.authored_token,
        form=old.form,
        cluster_id=old.cluster_id,
        cached_result=old.cached_result,
        item_refs=(fake_ref,),
    )
    occurrence_map = CitationOccurrenceMap(
        projection.occurrence_map.source_sha256,
        (dangling, *projection.occurrence_map.occurrences[1:]),
    )
    with pytest.raises(ResolvedCitationOoxmlError, match="dangling or cross-linked"):
        validate_citation_authorities(projection.item_map, occurrence_map)


@pytest.mark.parametrize("collision", ["full", "sdt", "bookmark"])
def test_occurrence_digest_and_physical_name_collisions_fail_closed(collision: str) -> None:
    _source, projection = _projection()
    first, second = projection.occurrence_map.occurrences
    if collision == "full":
        second = replace(second, sha256=first.sha256, tag=first.tag, bookmark_name=first.bookmark_name)
        match = "full digest collides"
    elif collision == "sdt":
        second = replace(second, tag=first.tag)
        match = "truncated SDT tag collides"
    else:
        second = replace(second, bookmark_name=first.bookmark_name)
        match = "truncated bookmark collides"
    colliding = CitationOccurrenceMap(projection.occurrence_map.source_sha256, (first, second))

    with pytest.raises(ResolvedCitationOoxmlError, match=match):
        validate_citation_occurrence_map(colliding)


def test_map_record_source_cache_and_item_order_mutations_fail_closed() -> None:
    _source, projection = _projection()
    item_root = etree.fromstring(citation_item_map_xml(projection.item_map))
    item_root[0].set("record_id", "reference-record:99")
    with pytest.raises(ResolvedCitationOoxmlError):
        parse_citation_item_map(item_root)

    occurrence_root = etree.fromstring(citation_occurrence_map_xml(projection.occurrence_map))
    occurrence_root[0].set("source_start", "00")
    with pytest.raises(ResolvedCitationOoxmlError, match="canonical decimal"):
        parse_citation_occurrence_map(occurrence_root)

    occurrence_root = etree.fromstring(citation_occurrence_map_xml(projection.occurrence_map))
    second = occurrence_root[1]
    second[:] = list(reversed(list(second)))
    with pytest.raises(ResolvedCitationOoxmlError):
        parse_citation_occurrence_map(occurrence_root)


def test_session_writes_two_independent_maps_and_proof_only_roundtrip(tmp_path: Path) -> None:
    source = "@fig-legacy and [@smith; @wang]\n"
    port = _port(source)
    document = Document()
    session = ResolvedNumberingDocxSession(
        document,
        port,
        heading_style_ids={},
        heading_style_names={},
        caption_style_bindings=_caption_bindings(document),
    )
    first = document.add_paragraph()
    second = document.add_paragraph()
    session.render_citation(
        first,
        source_start=port.document.citations[0].source_start,
        source_end=port.document.citations[0].source_end,
    )
    session.render_citation(
        second,
        source_start=port.document.citations[1].source_start,
        source_end=port.document.citations[1].source_end,
    )
    output = tmp_path / "citations.docx"
    document.save(str(output))
    session.write_package(output)

    assert read_proven_resolved_citations(output) == port.document.citations
    with ZipFile(output) as package:
        owned = read_owned_map_parts(package)
        assert CITATION_ITEM_MAP_NAMESPACE in owned
        assert CITATION_OCCURRENCE_MAP_NAMESPACE in owned
        assert owned[CITATION_ITEM_MAP_NAMESPACE][0] != owned[CITATION_OCCURRENCE_MAP_NAMESPACE][0]
        document_xml = etree.fromstring(package.read("word/document.xml"))
    instructions = [item.text for item in document_xml.iter(qn("w:instrText"))]
    assert session.citation_projection is not None
    assert instructions == [
        citation_instruction(session.citation_projection.occurrence_map.occurrences[0].item_refs),
        citation_instruction(session.citation_projection.occurrence_map.occurrences[1].item_refs),
    ]
    begins = [item for item in document_xml.iter(qn("w:fldChar")) if item.get(qn("w:fldCharType")) == "begin"]
    assert all(item.get(qn("w:fldLock")) == "true" and item.get(qn("w:dirty")) is None for item in begins)


@pytest.mark.parametrize(
    ("old", "new"),
    [
        (b'w:fldLock="true"', b'w:dirty="true"'),
        (b" CITATION ", b" CITATIONX "),
        (b"Legacy Figure Record", b"Tampered Cache Value"),
        (b"docwen-citation-occurrence-v1:", b"docwen-citation-occurrence-v2:"),
    ],
)
def test_physical_field_sdt_and_cache_tamper_fail_closed(tmp_path: Path, old: bytes, new: bytes) -> None:
    output = _render_package(tmp_path)
    _replace_zip_member(output, "word/document.xml", old, new)

    with pytest.raises(ResolvedCitationOoxmlError):
        read_proven_resolved_citations(output)


def test_missing_map_and_opc_trio_fail_closed(tmp_path: Path) -> None:
    output = _render_package(tmp_path)
    with ZipFile(output) as package:
        owned = read_owned_map_parts(package)
        item_number = owned[CITATION_ITEM_MAP_NAMESPACE][0]
    _drop_zip_member(output, f"customXml/itemProps{item_number}.xml")

    with pytest.raises(ResolvedCitationOoxmlError):
        read_proven_resolved_citations(output)


def test_duplicate_map_namespace_and_uuid_tamper_fail_closed(tmp_path: Path) -> None:
    output = _render_package(tmp_path)
    with ZipFile(output) as package:
        owned = read_owned_map_parts(package)
        item_number = owned[CITATION_ITEM_MAP_NAMESPACE][0]
        item_bytes = package.read(f"customXml/item{item_number}.xml")
    _add_zip_member(output, "customXml/item99.xml", item_bytes)
    with pytest.raises(ResolvedCitationOoxmlError):
        read_proven_resolved_citations(output)

    output = _render_package(tmp_path)
    with ZipFile(output) as package:
        owned = read_owned_map_parts(package)
        item_number = owned[CITATION_ITEM_MAP_NAMESPACE][0]
    _replace_zip_member(output, f"customXml/itemProps{item_number}.xml", b'ds:itemID="{', b'ds:itemID="{F')
    with pytest.raises(ResolvedCitationOoxmlError):
        read_proven_resolved_citations(output)


def test_visible_dwcit_text_without_maps_is_not_an_owned_signal(tmp_path: Path) -> None:
    document = Document()
    document.add_paragraph("Authored text DWCIT_00000000000000000000000000000000 is ordinary.")
    output = tmp_path / "ordinary.docx"
    document.save(str(output))

    assert read_proven_resolved_citations(output) == ()


def test_existing_nonowned_word_tag_and_bookmark_collisions_preflight() -> None:
    _source, projection = _projection()
    document = Document()
    paragraph = document.add_paragraph()
    run = OxmlElement("w:r")
    instruction = OxmlElement("w:instrText")
    instruction.text = f" CITATION {projection.item_map.items[0].word_tag} "
    run.append(instruction)
    paragraph._p.append(run)

    with pytest.raises(ResolvedCitationOoxmlError, match="non-owned field"):
        preflight_citation_document(document, projection)

    document = Document()
    paragraph = document.add_paragraph()
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), "1")
    start.set(qn("w:name"), projection.occurrence_map.occurrences[0].bookmark_name.swapcase())
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), "1")
    paragraph._p.extend((start, end))
    with pytest.raises(ResolvedCitationOoxmlError, match="existing package bookmark"):
        preflight_citation_document(document, projection)


def _render_package(tmp_path: Path) -> Path:
    source = "@fig-legacy and [@smith; @wang]\n"
    port = _port(source)
    document = Document()
    session = ResolvedNumberingDocxSession(
        document,
        port,
        heading_style_ids={},
        heading_style_names={},
        caption_style_bindings=_caption_bindings(document),
    )
    for citation in port.document.citations:
        paragraph = document.add_paragraph()
        session.render_citation(
            paragraph,
            source_start=citation.source_start,
            source_end=citation.source_end,
        )
    output = tmp_path / "citations.docx"
    document.save(str(output))
    session.write_package(output)
    return output


def _replace_zip_member(path: Path, member: str, old: bytes, new: bytes) -> None:
    with ZipFile(path) as package:
        infos = package.infolist()
        data = {info.filename: package.read(info.filename) for info in infos}
    assert data[member].count(old) >= 1
    data[member] = data[member].replace(old, new, 1)
    _rewrite_zip(path, infos, data)


def _drop_zip_member(path: Path, member: str) -> None:
    with ZipFile(path) as package:
        infos = [info for info in package.infolist() if info.filename != member]
        data = {info.filename: package.read(info.filename) for info in infos}
    _rewrite_zip(path, infos, data)


def _add_zip_member(path: Path, member: str, payload: bytes) -> None:
    with ZipFile(path) as package:
        infos = package.infolist()
        data = {info.filename: package.read(info.filename) for info in infos}
    assert member not in data
    data[member] = payload
    _rewrite_zip(path, infos, data)


def _rewrite_zip(path: Path, infos: list[ZipInfo], data: dict[str, bytes]) -> None:
    temporary = path.with_suffix(".rewrite.docx")
    with ZipFile(temporary, "w") as output:
        for info in infos:
            output.writestr(info, data[info.filename])
        known = {info.filename for info in infos}
        for name in sorted(set(data) - known):
            info = ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = ZIP_DEFLATED
            output.writestr(info, data[name])
    temporary.replace(path)
