"""Exact-neutral DOCX recovery map unit and integration gates."""

from __future__ import annotations

import hashlib
from pathlib import Path
from typing import Any, cast
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from lxml import etree

from docwen_core._docx_recovery_map import (
    RESOLVED_V4_RECOVERY_MAP_NAMESPACE,
    ResolvedV4RecoveryInput,
    ResolvedV4RecoveryMap,
    build_recovery_map,
    compute_physical_projection,
    inject_recovery_map,
    parse_recovery_map,
    read_recovery_map,
    recovery_map_xml,
)
from docwen_core._docx_semantics_v3_model import CaptionStyleBindingV3, CaptionStyleKeyV3
from docwen_core._docx_semantics_v3_package import DocxSemanticsV3Error
from docwen_core.docx_resolved_numbering import ResolvedNumberingDocxSession
from docwen_core.docx_resolved_numbering_recovery import ResolvedNumberingV4Recovery
from docwen_core.models.resolved_numbering import (
    CaptionMaterialization,
    NumberingExportPlanEnvelope,
    NumberingTarget,
    ResolvedDocument,
    ResolvedDocumentEnvelope,
    ResolvedDocumentTarget,
    ResolvedNumberingPlan,
    ResolvedNumberingPort,
)

pytestmark = pytest.mark.unit


def _sha(value: str) -> str:
    return hashlib.sha256(value.encode()).hexdigest()


def _port() -> ResolvedNumberingPort:
    source = "Table: One\n|a|\n|-|\n|1|\n"
    target_end = len(source) - 1
    target = ResolvedDocumentTarget(0, target_end, _sha(source[:target_end]), "table", "table-a", None, "One")
    plan_target = NumberingTarget(
        0,
        target_end,
        "table",
        True,
        "table-a",
        "1",
        CaptionMaterialization(
            "simple_seq",
            "Table",
            "arabic_half",
            "continue",
            None,
            None,
            None,
            None,
            None,
            None,
            None,
            "1",
            "Table",
            " ",
        ),
    )
    source_sha = _sha(source)
    plan_sha = "b" * 64
    return ResolvedNumberingPort(
        ResolvedDocumentEnvelope(
            "source",
            source_sha,
            plan_sha,
            ResolvedDocument(source, (target,), (), (), (), ()),
        ),
        NumberingExportPlanEnvelope(
            "source",
            source_sha,
            plan_sha,
            ResolvedNumberingPlan((), (), (plan_target,)),
        ),
    )


def _input() -> ResolvedV4RecoveryInput:
    return ResolvedV4RecoveryInput(
        neutral_raw=b'{"document": {"authored_markdown": "Table: One\\n|a|\\n|-|\\n|1|\\n"}}',
        plan_raw=b'{"plan": {"targets": []}}',
        authored_source=b"Table: One\n|a|\n|-|\n|1|\n",
        neutral_name="neutral-document.json",
        plan_name="numbering-export-plan.json",
        authored_name="authored-source.md",
        bibliography_owner="_DWB_BIBLIOGRAPHY",
        bibliography_placeholder="{{ bibliography }}",
        bibliography_media_type="application/vnd.docwen.semantic-bibliography+json",
    )


def _recovery_map() -> ResolvedV4RecoveryMap:
    return build_recovery_map(
        _port(),
        input_bytes=_input(),
        physical_sha256="c" * 64,
    )


def test_recovery_map_serialization_round_trips_canonical_bytes() -> None:
    value = _recovery_map()
    data = recovery_map_xml(value)
    assert data.startswith(b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n')
    root = etree.fromstring(data)
    parsed = parse_recovery_map(root)
    assert parsed == value
    assert parsed.source_sha256 == _sha("Table: One\n|a|\n|-|\n|1|\n")
    assert {item.role for item in parsed.pointers} == {"neutral_raw", "plan_raw", "authored_source"}


def test_recovery_map_rejects_wrong_pointer_roles_and_digests() -> None:
    value = _recovery_map()
    root = etree.fromstring(recovery_map_xml(value))
    pointer = root.find(
        f"{{{RESOLVED_V4_RECOVERY_MAP_NAMESPACE}}}pointers/{'{' + RESOLVED_V4_RECOVERY_MAP_NAMESPACE + '}'}pointer"
    )
    assert pointer is not None
    pointer.set("role", "evil_raw")
    with pytest.raises(DocxSemanticsV3Error, match="pointer role"):
        parse_recovery_map(root)
    pointer.set("role", "neutral_raw")
    pointer.set("sha256", "z" * 64)
    with pytest.raises(DocxSemanticsV3Error, match="digest"):
        parse_recovery_map(root)


def test_recovery_map_rejects_content_digest_tamper() -> None:
    value = _recovery_map()
    root = etree.fromstring(recovery_map_xml(value))
    root.set("recovery_sha256", "f" * 64)
    with pytest.raises(DocxSemanticsV3Error, match="content digest"):
        parse_recovery_map(root)


def test_recovery_map_rejects_missing_bibliography_or_projection() -> None:
    value = _recovery_map()
    root = etree.fromstring(recovery_map_xml(value))
    bibliography = root.find(f"{{{RESOLVED_V4_RECOVERY_MAP_NAMESPACE}}}bibliography")
    assert bibliography is not None
    root.remove(bibliography)
    with pytest.raises(DocxSemanticsV3Error, match="projection or bibliography"):
        parse_recovery_map(root)


def test_recovery_map_rejects_foreign_namespace_child() -> None:
    value = _recovery_map()
    root = etree.fromstring(recovery_map_xml(value))
    etree.SubElement(root, "{urn:evil}extra")
    with pytest.raises(DocxSemanticsV3Error, match="foreign child"):
        parse_recovery_map(root)


def test_projection_digest_is_byte_bound_and_excludes_only_recovery_trio(tmp_path: Path) -> None:
    package = tmp_path / "projection.docx"
    document = Document()
    document.add_paragraph("before")
    document.save(str(package))
    before = compute_physical_projection(package, exclude_item_numbers=set())

    value = _recovery_map()
    inject_recovery_map(package, value)
    with ZipFile(package) as opened:
        recovery = read_recovery_map(opened)
    assert recovery is not None
    item_number, parsed = recovery
    assert parsed.physical_sha256 == "c" * 64
    after = compute_physical_projection(package, exclude_item_numbers={item_number})
    assert after == before

    changed = tmp_path / "changed.docx"
    changed.write_bytes(package.read_bytes())
    with ZipFile(changed) as opened:
        infos = opened.infolist()
        members = {info.filename: opened.read(info.filename) for info in infos}
    members["word/document.xml"] = members["word/document.xml"].replace(b"<w:t>before</w:t>", b"<w:t>after</w:t>")
    with ZipFile(changed, "w", ZIP_DEFLATED) as output:
        for info in infos:
            output.writestr(info, members[info.filename])
    changed_digest = compute_physical_projection(changed, exclude_item_numbers={item_number})
    assert changed_digest != after


def test_session_writes_recovery_map_and_recovery_proves_it(tmp_path: Path) -> None:
    document = Document()
    port = _port()
    session = ResolvedNumberingDocxSession(
        document,
        port,
        heading_style_ids={},
        heading_style_names={},
        caption_style_bindings=_caption_bindings(document),
        recovery_input=_input(),
    )
    target = port.document.targets[0]
    caption = document.add_paragraph(style="Table Caption V4")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "One"
    session.bind_caption(
        caption, (table._element,), source_start=target.source_start, source_end=target.source_end, kind="table"
    )
    output = tmp_path / "recovery-map.docx"

    def add_request_owned_part(path: Path) -> None:
        with ZipFile(path, "a", ZIP_DEFLATED) as package:
            package.writestr("docProps/request-owned.txt", b"bound before recovery")

    session.write_package(output, pre_recovery_package_transform=add_request_owned_part)
    assert session.recovery_map is not None
    with ZipFile(output) as package:
        assert package.read("docProps/request-owned.txt") == b"bound before recovery"

    reopened = Document(str(output))
    recovery = ResolvedNumberingV4Recovery.load_if_present(output, reopened)
    assert recovery is not None
    assert recovery.source_recovery_available
    recovery.prove_exact_recovery_raw(
        neutral_raw=_input().neutral_raw,
        plan_raw=_input().plan_raw,
        authored_source=_input().authored_source,
    )


def test_session_rejects_recovery_when_port_has_no_recovery_input(tmp_path: Path) -> None:
    document = Document()
    port = _port()
    session = ResolvedNumberingDocxSession(
        document,
        port,
        heading_style_ids={},
        heading_style_names={},
        caption_style_bindings=_caption_bindings(document),
    )
    target = port.document.targets[0]
    caption = document.add_paragraph(style="Table Caption V4")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "One"
    session.bind_caption(
        caption, (table._element,), source_start=target.source_start, source_end=target.source_end, kind="table"
    )
    output = tmp_path / "no-recovery.docx"
    session.write_package(output)
    assert session.recovery_map is None
    reopened = Document(str(output))
    recovery = ResolvedNumberingV4Recovery.load_if_present(output, reopened)
    assert recovery is not None
    assert not recovery.source_recovery_available


def test_recovery_disables_exact_source_for_a_semantically_proven_package_edit(tmp_path: Path) -> None:
    document = Document()
    port = _port()
    session = ResolvedNumberingDocxSession(
        document,
        port,
        heading_style_ids={},
        heading_style_names={},
        caption_style_bindings=_caption_bindings(document),
        recovery_input=_input(),
    )
    target = port.document.targets[0]
    caption = document.add_paragraph(style="Table Caption V4")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "One"
    session.bind_caption(
        caption, (table._element,), source_start=target.source_start, source_end=target.source_end, kind="table"
    )
    output = tmp_path / "tampered.docx"
    session.write_package(output)

    with ZipFile(output) as opened:
        infos = opened.infolist()
        members = {info.filename: opened.read(info.filename) for info in infos}
    members["word/document.xml"] = members["word/document.xml"].replace(b"<w:t>One</w:t>", b"<w:t>Two</w:t>")
    with ZipFile(output, "w", ZIP_DEFLATED) as rewritten:
        for info in infos:
            rewritten.writestr(info, members[info.filename])

    reopened = Document(str(output))
    recovery = ResolvedNumberingV4Recovery.load_if_present(output, reopened)
    assert recovery is not None
    assert not recovery.source_recovery_available
    assert recovery.caption_signatures == (("table", "table-a", "One", "1"),)


def _caption_bindings(document: Any) -> tuple[CaptionStyleBindingV3, ...]:
    output: list[CaptionStyleBindingV3] = []
    for semantic_key, style_id, name in (
        ("table_caption", "DWTableCaption", "Table Caption V4"),
        ("figure_caption", "DWFigureCaption", "Figure Caption V4"),
        ("equation_caption", "DWEquationCaption", "Equation Caption V4"),
        ("code_block_caption", "DWCodeCaption", "Code Caption V4"),
    ):
        style = document.styles.add_style(style_id, WD_STYLE_TYPE.PARAGRAPH)
        style.name = name
        output.append(CaptionStyleBindingV3(cast(CaptionStyleKeyV3, semantic_key), style.style_id, style.name))
    return tuple(output)
