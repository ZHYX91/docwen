"""Golden tests for gongwen extraction (GOLDEN-008)."""

from __future__ import annotations

import json
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document

pytestmark = pytest.mark.contract

_PROJECT_ROOT = Path(__file__).resolve().parents[5]
_GONGWEN_OLD_SYSTEM_FIXTURE = _PROJECT_ROOT / "tests" / "fixtures" / "golden" / "old_system_gongwen_semantics.json"


class _ExactSchemeRegistry:
    def __init__(self, *, enabled: bool = True, levels: dict[str, str] | None = None) -> None:
        self._scheme = SimpleNamespace(
            enabled=enabled,
            levels={"level_1": "{1.arabic_half} "} if levels is None else levels,
        )

    def get_scheme(self, scheme_id: str) -> object:
        if scheme_id != "exact":
            raise LookupError(scheme_id)
        return self._scheme


def _load_gongwen_old_system_fixture() -> dict:
    return json.loads(_GONGWEN_OLD_SYSTEM_FIXTURE.read_text(encoding="utf-8"))


# Fixture: create a minimal gongwen DOCX for testing
@pytest.fixture
def gongwen_docx_path(tmp_path):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()

    # Title — 小标宋, 22pt, centered
    title = doc.add_paragraph("关于进一步规范公文处理工作的通知")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0] if title.runs else title.add_run("关于进一步规范公文处理工作的通知")
    run.font.name = "小标宋"
    run.font.size = Pt(22)

    # Doc number
    dn = doc.add_paragraph("国办发〔2024〕5号")
    if dn.runs:
        dn.runs[0].font.name = "仿宋"
        dn.runs[0].font.size = Pt(15)

    # Recipient
    rec = doc.add_paragraph("各省、自治区、直辖市人民政府办公厅：")
    if rec.runs:
        rec.runs[0].font.name = "仿宋"
        rec.runs[0].font.size = Pt(15)

    # Body
    body1 = doc.add_paragraph("为进一步规范公文处理工作，现就有关事项通知如下：")
    body2 = doc.add_paragraph("一、严格公文格式标准。")
    body3 = doc.add_paragraph("二、规范公文流转程序。")
    for p_inner in [body1, body2, body3]:
        if p_inner.runs:
            p_inner.runs[0].font.name = "仿宋"
            p_inner.runs[0].font.size = Pt(15)

    # Attachment
    att = doc.add_paragraph("附件：1. 公文格式标准")
    att2 = doc.add_paragraph("          2. 公文流转程序")
    for p_inner in [att, att2]:
        if p_inner.runs:
            p_inner.runs[0].font.name = "仿宋"
            p_inner.runs[0].font.size = Pt(15)

    # Copy to + printing
    cc = doc.add_paragraph("抄送：省委组织部、省人民政府办公厅")
    pr = doc.add_paragraph("国务院办公厅　　　　2024年1月15日印发")
    for p_inner in [cc, pr]:
        if p_inner.runs:
            p_inner.runs[0].font.name = "仿宋"
            p_inner.runs[0].font.size = Pt(15)

    path = tmp_path / "gongwen.docx"
    doc.save(str(path))
    return path


# Fixture: create a gongwen DOCX with explicit attachment documents.
@pytest.fixture
def gongwen_docx_with_attachments_path(tmp_path):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()

    # Title
    title = doc.add_paragraph("关于印发项目管理办法的通知")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0] if title.runs else title.add_run("关于印发项目管理办法的通知")
    run.font.name = "小标宋"
    run.font.size = Pt(22)

    # Doc number
    dn = doc.add_paragraph("国办发〔2024〕10号")
    if dn.runs:
        dn.runs[0].font.name = "仿宋"
        dn.runs[0].font.size = Pt(15)

    # Recipient
    rec = doc.add_paragraph("各省人民政府办公厅：")
    if rec.runs:
        rec.runs[0].font.name = "仿宋"
        rec.runs[0].font.size = Pt(15)

    # Body
    body = doc.add_paragraph("为规范项目管理工作，现制定本办法。")
    if body.runs:
        body.runs[0].font.name = "仿宋"
        body.runs[0].font.size = Pt(15)

    # Attachment header
    att = doc.add_paragraph("附件：1. 项目管理办法")
    att2 = doc.add_paragraph("          2. 实施细则")
    att3 = doc.add_paragraph("          3. 考核标准")
    for p_inner in [att, att2, att3]:
        if p_inner.runs:
            p_inner.runs[0].font.name = "仿宋"
            p_inner.runs[0].font.size = Pt(15)

    # Attachment content
    att_content = doc.add_paragraph("附件内容：项目管理办法全文......")
    if att_content.runs:
        att_content.runs[0].font.name = "仿宋"
        att_content.runs[0].font.size = Pt(15)

    path = tmp_path / "gongwen_att.docx"
    doc.save(str(path))
    return path


@pytest.mark.parametrize(
    ("scheme", "registry", "error_type", "diagnostic_code"),
    [
        ("", _ExactSchemeRegistry(), "invalid_input", "NUMBERING-SCHEME-REQUIRED"),
        ("exact", None, "capability_unavailable", "NUMBERING-REGISTRY-UNAVAILABLE"),
        ("missing", _ExactSchemeRegistry(), "resource_not_found", "NUMBERING-SCHEME-NOT-FOUND"),
        (
            "exact",
            _ExactSchemeRegistry(enabled=False),
            "capability_unavailable",
            "NUMBERING-SCHEME-DISABLED",
        ),
        ("exact", _ExactSchemeRegistry(levels={}), "invalid_input", "NUMBERING-SCHEME-NO-LEVELS"),
    ],
)
def test_gongwen_rejects_unusable_exact_numbering_scheme(
    tmp_path: Path,
    gongwen_docx_path: Path,
    scheme: str,
    registry: object,
    error_type: str,
    diagnostic_code: str,
) -> None:
    from tests.support.config import FakeConfigView
    from tests.support.execution import FakeExecutionContext
    from tests.support.logging import FakePluginLogger
    from tests.support.progress import FakeProgressSink
    from tests.support.workspace import FakeWorkspaceHandle

    from docwen_core.cancellation import CancellationToken
    from docwen_core.models.file_ref import FileRef
    from docwen_core.models.request import ConversionRequest, OutputPolicy
    from docwen_plugin_optimizer_gongwen.plugin import GongwenOptimizerPlugin

    staging = tmp_path / "staging"
    staging.mkdir()
    context = FakeExecutionContext(
        request=ConversionRequest(
            request_id="gongwen-numbering-failure",
            input_refs=[FileRef(path=str(gongwen_docx_path), format="docx", category="document")],
            target_format="md",
            action_name="gongwen",
            options={"add_numbering": True, "numbering_scheme": scheme},
            output_policy=OutputPolicy(),
        ),
        workspace=FakeWorkspaceHandle(str(gongwen_docx_path), str(staging)),
        config=FakeConfigView(),
        progress=FakeProgressSink(),
        cancellation=CancellationToken().view(),
        logger=FakePluginLogger(),
        numbering_registry=registry,
    )

    result = GongwenOptimizerPlugin().convert(context)

    assert not result.success
    assert result.error is not None
    assert result.error.error_type == error_type
    assert result.error.diagnostic_code == diagnostic_code
    assert result.artifacts == []


class TestGongwenGolden:
    def test_gongwen_matches_old_system_semantic_fixture(self, gongwen_docx_path):
        """Current gongwen mode should preserve old-system core metadata semantics."""
        from docwen_plugin_optimizer_gongwen.pipeline import convert_docx_to_md_gongwen

        fixture = _load_gongwen_old_system_fixture()
        assert fixture["golden_id"] == "GOLDEN-008"

        doc = Document(str(gongwen_docx_path))
        result = convert_docx_to_md_gongwen(doc, str(gongwen_docx_path), {})

        assert result["success"]
        assert len(result["yaml_info"]) == fixture["expected_semantics"]["yaml_field_count"]
        assert result["yaml_info"]["标题"] == fixture["expected_semantics"]["title"]
        assert result["yaml_info"]["发文字号"] == fixture["expected_semantics"]["doc_number"]
        assert result["yaml_info"]["主送机关"] == fixture["expected_semantics"]["recipient"]
        assert result["yaml_info"]["附件说明"] == fixture["expected_semantics"]["attachment_description"]
        assert result["yaml_info"]["发文机关署名"] == ""
        assert result["yaml_info"]["印发机关"] == fixture["expected_semantics"]["printing_authority"]
        assert result["yaml_info"]["印发日期"] == fixture["expected_semantics"]["printing_date"]
        assert fixture["expected_semantics"]["printing_line"] not in result["yaml_info"]["附件说明"]

        for token in fixture["expected_semantics"]["markdown_contains"]:
            assert token in result["markdown"]
        assert result["attachment_documents"] == []

        for project_name in ("docwen-ref-tk", "docwen-ref-pyside6", "docwen-current"):
            project = fixture["projects"][project_name]
            assert project["success"] is True
            assert project["yaml_field_count"] == fixture["expected_semantics"]["yaml_field_count"]
            assert project["contains_core_markdown"] is True

    def test_gongwen_old_system_fixture_finalizes_through_runtime(self, gongwen_docx_path, tmp_path):
        """Gongwen primary and attachment Markdown should land in the output dir."""
        from docwen_core.models.file_ref import FileRef
        from docwen_core.models.request import ConversionRequest, OutputPolicy
        from docwen_plugin_optimizer_gongwen.plugin import GongwenOptimizerPlugin
        from docwen_runtime.engine.route_resolver import RouteResolver
        from docwen_runtime.engine.task_manager import TaskManager
        from docwen_runtime.output.finalizer import OutputFinalizer
        from docwen_runtime.plugin_registry.registry import PluginRegistry
        from docwen_runtime.workspace.manager import WorkspaceManager

        fixture = _load_gongwen_old_system_fixture()
        expected = fixture["expected_semantics"]
        output_dir = tmp_path / "out"
        output_dir.mkdir()
        workspace_root = tmp_path / "workspace"
        registry = PluginRegistry()
        registry.register(GongwenOptimizerPlugin())
        task_mgr = TaskManager(
            registry,
            RouteResolver(registry),
            WorkspaceManager(root_dir=str(workspace_root)),
            OutputFinalizer(),
        )
        request = ConversionRequest(
            request_id="gongwen-finalizer-old-system-fixture",
            input_refs=[
                FileRef(
                    path=str(gongwen_docx_path),
                    format="docx",
                    category="document",
                    size_bytes=gongwen_docx_path.stat().st_size,
                )
            ],
            target_format="md",
            action_name="gongwen",
            options={},
            output_policy=OutputPolicy(output_dir=str(output_dir)),
        )

        result = task_mgr.execute_single(request)

        assert result.success, f"unexpected error: {result.error}"
        assert len(result.artifacts) == 2
        primary = next(artifact for artifact in result.artifacts if artifact.is_primary)
        manifest = next(artifact for artifact in result.artifacts if artifact.kind == "manifest")
        primary_path = Path(primary.staging_path)
        assert primary_path.parent.parent == output_dir
        assert primary_path.stem == primary_path.parent.name
        assert Path(manifest.staging_path) == primary_path.parent / "docwen-node.json"
        assert primary.media_type == "text/markdown"
        assert primary.metadata["paragraph_count"] == len(fixture["input_docx"]["paragraphs"])
        assert primary.metadata["gongwen_fields"] >= 4
        assert primary.metadata["gongwen_needs_review"] is True
        assert primary.metadata["gongwen_missing_required"] == ["issue_date", "issuing_authority_signature"]
        assert primary.metadata["gongwen_review_reasons"] == []
        assert any(d.code == "GONGWEN-OK" for d in result.diagnostics)
        review_diagnostic = next(d for d in result.diagnostics if d.code == "GONGWEN-NEEDS-REVIEW")
        assert review_diagnostic.level == "warning"
        assert review_diagnostic.message == "缺少必需字段：成文日期、发文机关署名"

        primary_markdown = primary_path.read_text(encoding="utf-8")
        for token in expected["markdown_contains"]:
            assert token in primary_markdown
        assert expected["printing_line"] not in primary_markdown
        assert str(workspace_root) not in primary_markdown

    def test_gongwen_image_artifact_finalizes_with_relative_markdown_link(
        self,
        gongwen_docx_path,
        tmp_path,
        monkeypatch,
    ):
        """A retained Gongwen image must survive staging and finalization."""
        from docwen_core.models.file_ref import FileRef
        from docwen_core.models.request import ConversionRequest, OutputPolicy
        from docwen_plugin_optimizer_gongwen.plugin import GongwenOptimizerPlugin
        from docwen_runtime.engine.route_resolver import RouteResolver
        from docwen_runtime.engine.task_manager import TaskManager
        from docwen_runtime.output.finalizer import OutputFinalizer
        from docwen_runtime.plugin_registry.registry import PluginRegistry
        from docwen_runtime.workspace.manager import WorkspaceManager

        image_bytes = (
            b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR"
            b"\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde"
            b"\x00\x00\x00\x0cIDATx\x9cc```\x00\x00\x00\x04\x00\x01\xf6\x178U"
            b"\x00\x00\x00\x00IEND\xaeB`\x82"
        )
        captured_output_dir: list[Path] = []

        def fake_convert(_doc, _input_path, options, **_kwargs):
            from docwen_plugin_optimizer_gongwen.models import AttachmentDocument

            staging_dir = Path(options["output_dir"])
            captured_output_dir.append(staging_dir)
            image_path = staging_dir / "gongwen-seal.png"
            image_path.write_bytes(image_bytes)
            return {
                "success": True,
                "yaml_info": {f"field_{index}": "" for index in range(18)},
                "markdown": "---\n---\n\n![](gongwen-seal.png)\n",
                "attachment_documents": [
                    AttachmentDocument(
                        ordinal=1,
                        title="附件汇总",
                        markdown="# Attachment\n",
                        paragraph_indices=(),
                    )
                ],
                "stats": {"paragraphs": 1},
                "image_paths": [str(image_path)],
                "metadata": {
                    "confidence": {"overall": "high"},
                    "recognition_review_signals": {
                        "missing_required": [],
                        "needs_review_reasons": [],
                        "recognition_summary": {"needs_review": False},
                    },
                },
            }

        monkeypatch.setattr(
            "docwen_plugin_optimizer_gongwen.plugin.convert_docx_to_md_gongwen",
            fake_convert,
        )

        output_dir = tmp_path / "out"
        output_dir.mkdir()
        workspace_root = tmp_path / "workspace"
        registry = PluginRegistry()
        registry.register(GongwenOptimizerPlugin())
        task_mgr = TaskManager(
            registry,
            RouteResolver(registry),
            WorkspaceManager(root_dir=str(workspace_root)),
            OutputFinalizer(),
        )
        request = ConversionRequest(
            request_id="gongwen-image-finalizer",
            input_refs=[
                FileRef(
                    path=str(gongwen_docx_path),
                    format="docx",
                    category="document",
                    size_bytes=gongwen_docx_path.stat().st_size,
                )
            ],
            target_format="md",
            action_name="gongwen",
            options={"to_md_keep_images": True, "image_mode": "file"},
            output_policy=OutputPolicy(output_dir=str(output_dir)),
        )

        result = task_mgr.execute_single(request)

        assert result.success, f"unexpected error: {result.error}"
        assert len(result.artifacts) == 4
        diagnostic_codes = {diagnostic.code for diagnostic in result.diagnostics}
        assert {"GONGWEN-OK", "FINALIZER_DONE"} <= diagnostic_codes
        assert "GONGWEN-NEEDS-REVIEW" not in diagnostic_codes
        assert len(captured_output_dir) == 1
        assert captured_output_dir[0].is_relative_to(workspace_root)
        assert captured_output_dir[0] != output_dir

        primary = next(artifact for artifact in result.artifacts if artifact.is_primary)
        image = next(artifact for artifact in result.artifacts if artifact.kind == "image")
        attachment = next(
            artifact for artifact in result.artifacts if artifact.metadata.get("source_kind") == "gongwen_attachment"
        )
        assert Path(primary.staging_path).parent.parent == output_dir
        assert Path(image.staging_path).parent == Path(primary.staging_path).parent
        assert Path(image.staging_path).read_bytes() == image_bytes
        assert image.suggested_name == "gongwen-seal.png"
        assert image.media_type == "image/png"
        assert image.metadata["source_format"] == "docx"
        assert image.metadata["source_kind"] == "gongwen_image"
        assert image.metadata["document_node_role"] == "resource"
        assert image.metadata["logical_path"] == image.logical_path
        attachment_path = Path(attachment.staging_path)
        assert attachment_path.parent.parent == Path(primary.staging_path).parent
        assert attachment_path.stem == attachment_path.parent.name

        primary_markdown = Path(primary.staging_path).read_text(encoding="utf-8")
        assert "![](gongwen-seal.png)" in primary_markdown
        assert attachment_path.parent.name in primary_markdown
        assert str(workspace_root) not in primary_markdown

    def test_extracts_title(self, gongwen_docx_path):
        """Verify the pipeline extracts the document title."""
        from docwen_plugin_optimizer_gongwen.pipeline import convert_docx_to_md_gongwen

        doc = Document(str(gongwen_docx_path))
        result = convert_docx_to_md_gongwen(doc, str(gongwen_docx_path), {})
        assert result["success"]
        assert "关于进一步规范公文处理工作的通知" in result["yaml_info"]["标题"]

    def test_extracts_doc_number(self, gongwen_docx_path):
        """Verify the pipeline extracts the document number."""
        from docwen_plugin_optimizer_gongwen.pipeline import convert_docx_to_md_gongwen

        doc = Document(str(gongwen_docx_path))
        result = convert_docx_to_md_gongwen(doc, str(gongwen_docx_path), {})
        assert "国办发〔2024〕5号" in result["yaml_info"]["发文字号"]

    def test_extracts_recipient(self, gongwen_docx_path):
        """Verify the pipeline extracts the recipient."""
        from docwen_plugin_optimizer_gongwen.pipeline import convert_docx_to_md_gongwen

        doc = Document(str(gongwen_docx_path))
        result = convert_docx_to_md_gongwen(doc, str(gongwen_docx_path), {})
        assert "人民政府办公厅" in result["yaml_info"]["主送机关"]

    def test_output_markdown_has_yaml_frontmatter(self, gongwen_docx_path):
        """Verify the output markdown includes YAML frontmatter."""
        from docwen_plugin_optimizer_gongwen.pipeline import convert_docx_to_md_gongwen

        doc = Document(str(gongwen_docx_path))
        result = convert_docx_to_md_gongwen(doc, str(gongwen_docx_path), {})
        assert "---" in result["markdown"]
        assert "关于进一步规范公文处理工作的通知" in result["markdown"]

    def test_roundtrip(self, gongwen_docx_path):
        """DOCX→MD gongwen mode preserves core fields in YAML frontmatter."""
        from docwen_plugin_optimizer_gongwen.pipeline import convert_docx_to_md_gongwen

        doc = Document(str(gongwen_docx_path))
        result = convert_docx_to_md_gongwen(doc, str(gongwen_docx_path), {})
        # After extraction, yaml_info should have meaningful content
        non_empty = {k: v for k, v in result["yaml_info"].items() if v and not (isinstance(v, list) and len(v) == 0)}
        assert len(non_empty) >= 3  # title, doc_number, recipient

    def test_pipeline_produces_valid_dict_structure(self, gongwen_docx_path):
        """Verify the pipeline result has all expected keys."""
        from docwen_plugin_optimizer_gongwen.pipeline import convert_docx_to_md_gongwen

        doc = Document(str(gongwen_docx_path))
        result = convert_docx_to_md_gongwen(doc, str(gongwen_docx_path), {})
        assert "success" in result
        assert "yaml_info" in result
        assert "markdown" in result
        assert "stats" in result
        assert "metadata" in result
        assert isinstance(result["yaml_info"], dict)
        assert len(result["yaml_info"]) == 18

    def test_generates_typed_attachment_documents(self, gongwen_docx_with_attachments_path):
        """Pipeline should generate typed attachment documents when detected."""
        from docwen_plugin_optimizer_gongwen.pipeline import convert_docx_to_md_gongwen

        doc = Document(str(gongwen_docx_with_attachments_path))
        result = convert_docx_to_md_gongwen(doc, str(gongwen_docx_with_attachments_path), {})
        assert result["success"]
        # Attachments present in the document should create typed documents.
        if result["yaml_info"].get("附件说明"):
            assert result["attachment_documents"]
            attachment_markdown = result["attachment_documents"][0].markdown
            assert "附件" in attachment_markdown
            assert "来源文件" in attachment_markdown
            assert "附件内容：项目管理办法全文......" not in result["yaml_info"]["附件说明"]
            assert "附件内容：项目管理办法全文......" in attachment_markdown


def test_gongwen_all_ocr_outcomes_warn_and_continue_later_images(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Every OCR outcome warns while base Gongwen Markdown and later OCR survive."""
    import docwen_core.text.ocr as ocr
    import docwen_plugin_optimizer_gongwen.extraction.paragraph_reader as paragraph_reader
    from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
    from docwen_plugin_optimizer_gongwen.pipeline import convert_docx_to_md_gongwen

    features = [
        ParagraphFeature(
            index=0,
            text="gongwen base content",
            extracted_images=[str(tmp_path / f"gongwen-ocr-{index}.png") for index in range(7)],
        )
    ]
    for image_path in features[0].extracted_images:
        Path(image_path).write_bytes(b"\x89PNG\r\n\x1a\ncontent-format-fixture")
    monkeypatch.setattr(paragraph_reader, "read_paragraphs", lambda *_args, **_kwargs: features)

    outcomes = iter(
        [
            RuntimeError("private unexpected adapter failure"),
            ocr.OcrOutcome(ocr.OcrStatus.UNAVAILABLE, message="private unavailable detail"),
            ocr.OcrOutcome(ocr.OcrStatus.MODEL_MISSING, message="private model path"),
            ocr.OcrOutcome(ocr.OcrStatus.INITIALIZATION_FAILED, message="private init detail"),
            ocr.OcrOutcome(ocr.OcrStatus.RECOGNITION_FAILED, message="private recognition detail"),
            ocr.OcrOutcome(ocr.OcrStatus.NO_TEXT),
            ocr.OcrOutcome(ocr.OcrStatus.SUCCESS, text="later gongwen OCR"),
        ]
    )
    calls: list[str] = []

    def _run_ocr_outcome(path: str, **_kwargs: object) -> ocr.OcrOutcome:
        calls.append(str(path))
        outcome = next(outcomes)
        if isinstance(outcome, Exception):
            raise outcome
        return outcome

    monkeypatch.setattr(ocr, "run_ocr_outcome", _run_ocr_outcome)

    class _RecordingProgress:
        def __init__(self) -> None:
            self.diagnostics: list[tuple[str, str, str, str]] = []

        def report_progress(self, percent: float, message: str = "") -> None:
            return

        def report_diagnostic(
            self,
            level: str,
            message: str,
            code: str = "",
            location: str = "",
        ) -> None:
            self.diagnostics.append((level, message, code, location))

        def report_artifact_ready(self, artifact_id: str, suggested_name: str) -> None:
            return

    progress = _RecordingProgress()
    input_path = tmp_path / "gongwen-ocr-best-effort.docx"
    doc = Document()
    doc.add_paragraph("input placeholder")
    doc.save(str(input_path))

    result = convert_docx_to_md_gongwen(
        doc,
        str(input_path),
        {"to_md_enable_ocr": True},
        progress=progress,
    )

    assert result["success"] is True
    assert "gongwen base content" in result["markdown"]
    assert "later gongwen OCR" in result["markdown"]
    assert len(calls) == 7
    assert len(progress.diagnostics) == 7
    assert [diagnostic[0] for diagnostic in progress.diagnostics] == ["warning"] * 7
    assert [diagnostic[2] for diagnostic in progress.diagnostics] == ["OCR-BEST-EFFORT"] * 7
    assert [
        message.split("status=", 1)[1].split(";", 1)[0] for _level, message, _code, _location in progress.diagnostics
    ] == [
        "recognition_failed",
        "unavailable",
        "model_missing",
        "initialization_failed",
        "recognition_failed",
        "no_text",
        "success",
    ]
    assert all(location for _level, _message, _code, location in progress.diagnostics)
    assert all("private" not in message for _level, message, _code, _location in progress.diagnostics)
