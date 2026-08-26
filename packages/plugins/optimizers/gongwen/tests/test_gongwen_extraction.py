"""Tests for gongwen paragraph feature extraction."""

from __future__ import annotations

import pytest

pytestmark = pytest.mark.contract


def test_word_numbering_failure_keeps_paragraph_and_reports_fallback(tmp_path) -> None:
    """Malformed numbering metadata falls back without becoming silent loss."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from tests.support.progress import FakeProgressSink

    from docwen_plugin_optimizer_gongwen.extraction.paragraph_reader import read_paragraphs

    source = tmp_path / "broken-numbering.docx"
    document = Document()
    paragraph = document.add_paragraph("普通正文内容")
    num_pr = OxmlElement("w:numPr")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "7")
    num_pr.append(num_id)
    paragraph._p.get_or_add_pPr().append(num_pr)
    document.save(source)

    class BrokenNumberingIndex:
        def lookup(self, _num_id: int, _level: int):
            raise ValueError("invalid numbering definition")

    progress = FakeProgressSink()
    features = read_paragraphs(
        Document(source),
        numbering_index=BrokenNumberingIndex(),
        diagnostic_sink=progress,
    )

    assert [feature.text for feature in features] == ["普通正文内容"]
    assert any(
        level == "warning" and code == "GONGWEN-HEADING-NUMBERING-FALLBACK" and location == "paragraph 0"
        for level, _message, code, location in progress.diagnostics
    )


class TestParagraphReader:
    def test_extracts_paragraphs_from_docx(self, tmp_path):
        from docx import Document

        from docwen_plugin_optimizer_gongwen.extraction.paragraph_reader import (
            read_paragraphs,
        )

        doc = Document()
        doc.add_paragraph("Test paragraph one")
        doc.add_paragraph("Test paragraph two")

        # Save and reload to ensure proper structure
        path = tmp_path / "test.docx"
        doc.save(str(path))

        doc2 = Document(str(path))
        features = read_paragraphs(doc2)
        assert len(features) >= 2
        for f in features:
            assert f.index >= 0
            assert f.text is not None
            assert hasattr(f, "font_name")

    def test_skips_empty_body_paragraphs(self, tmp_path):
        from docx import Document

        from docwen_plugin_optimizer_gongwen.extraction.paragraph_reader import (
            read_paragraphs,
        )

        doc = Document()
        doc.add_paragraph("  ")  # whitespace-only
        doc.add_paragraph("Real text")

        path = tmp_path / "test2.docx"
        doc.save(str(path))

        doc2 = Document(str(path))
        features = read_paragraphs(doc2)
        # Should only have "Real text" (whitespace skipped)
        texts = [f.text for f in features]
        assert "Real text" in texts
        assert [feature.index for feature in features] == list(range(len(features)))

    def test_preserves_empty_paragraphs_that_own_rich_content_or_breaks(self, tmp_path):
        import base64
        from io import BytesIO

        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_BREAK
        from docx.oxml import OxmlElement

        from docwen_plugin_optimizer_gongwen.extraction.paragraph_reader import read_paragraphs

        doc = Document()
        doc.add_paragraph("   ")

        image_paragraph = doc.add_paragraph()
        pixel = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
        image_paragraph.add_run().add_picture(BytesIO(pixel))

        formula_paragraph = doc.add_paragraph()
        formula_paragraph._p.append(OxmlElement("m:oMath"))

        page_break_paragraph = doc.add_paragraph()
        page_break_paragraph.add_run().add_break(WD_BREAK.PAGE)

        doc.add_section(WD_SECTION.NEW_PAGE)
        doc.add_paragraph("正文")
        path = tmp_path / "rich-empty-paragraphs.docx"
        doc.save(path)

        image_dir = tmp_path / "images"
        image_dir.mkdir()
        features = read_paragraphs(Document(path), output_dir=str(image_dir))

        assert any(feature.has_image and feature.text == "" for feature in features)
        assert any(feature.has_formula and feature.text == "" for feature in features)
        assert any(feature.has_page_break and feature.text == "" for feature in features)
        assert any(feature.has_section_break and feature.text == "" for feature in features)
        assert [feature.index for feature in features] == list(range(len(features)))

    def test_embedded_image_requires_a_caller_owned_output_directory(self, tmp_path):
        import base64
        from io import BytesIO

        import pytest
        from docx import Document

        from docwen_plugin_optimizer_gongwen.extraction.paragraph_reader import read_paragraphs

        pixel = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
        document = Document()
        document.add_paragraph().add_run().add_picture(BytesIO(pixel))
        source = tmp_path / "unowned-image-output.docx"
        document.save(source)

        with pytest.raises(ValueError, match=r"^gongwen_image_output_dir_required$"):
            read_paragraphs(Document(source))

        assert not list(tmp_path.glob("gongwen_images_*"))

    def test_request_cleanup_rules_are_the_only_rules_used(self, tmp_path):
        from docx import Document

        from docwen_core.text.heading_numbering import (
            compile_clean_rules_from_data,
        )
        from docwen_plugin_optimizer_gongwen.extraction.paragraph_reader import read_paragraphs

        request_rules = compile_clean_rules_from_data(
            [{"id": "request", "enabled": True, "pattern": r"^REQ:\s*", "level": 2}]
        )
        doc = Document()
        doc.add_paragraph("REQ: Request title")
        doc.add_paragraph("GLOBAL: Global title")
        path = tmp_path / "gongwen-request-cleanup.docx"
        doc.save(str(path))

        features = read_paragraphs(Document(str(path)), cleanup_rules=request_rules)

        assert (features[0].text, features[0].heading_level, features[0].heading_numbering_text) == (
            "Request title",
            2,
            "REQ: ",
        )
        assert features[1].text == "GLOBAL: Global title"

    def test_mixed_heading_prefers_direct_run_format_boundary(self, tmp_path):
        from docx import Document
        from docx.shared import Pt

        from docwen_plugin_optimizer_gongwen.extraction.paragraph_reader import (
            read_paragraphs,
        )

        title = "言之有谋，强化顶层设计："
        body = "坚持全局眼光，注重前瞻谋划。"
        doc = Document()
        paragraph = doc.add_paragraph(style="Heading 2")
        paragraph.add_run(title)
        body_run = paragraph.add_run(body)
        body_run.font.name = "Calibri"
        body_run.font.size = Pt(10.5)
        body_run.bold = False
        path = tmp_path / "mixed-run-format.docx"
        doc.save(path)

        feature = read_paragraphs(Document(path))[0]

        assert feature.text == f"{title}{body}"
        assert feature.heading_level == 2
        assert feature.heading_body_boundary == len(title)
        assert feature.heading_body_boundary_source == "run_format"

    def test_mixed_heading_detects_body_when_its_first_run_is_emphasis_or_inline_code(self, tmp_path):
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE

        from docwen_plugin_optimizer_gongwen.extraction.paragraph_reader import read_paragraphs
        from docwen_plugin_optimizer_gongwen.models import GongwenMetadata
        from docwen_plugin_optimizer_gongwen.rendering.markdown_renderer import render

        cases = (("Emphasis", "强调正文"), ("Inline Code", "代码正文"))
        for style_name, body in cases:
            doc = Document()
            doc.styles.add_style("Inline Code", WD_STYLE_TYPE.CHARACTER)
            paragraph = doc.add_paragraph(style="Heading 2")
            title = "言之有谋，强化顶层设计："
            paragraph.add_run("言之有谋，")
            internal_run = paragraph.add_run("强化顶层")
            internal_run.style = "Inline Code" if style_name == "Emphasis" else "Emphasis"
            paragraph.add_run("设计：")
            body_run = paragraph.add_run(body)
            body_run.style = style_name
            path = tmp_path / f"mixed-{style_name.replace(' ', '-').lower()}.docx"
            doc.save(path)

            feature = read_paragraphs(Document(path))[0]
            markdown = render(
                GongwenMetadata.default(),
                [feature.text],
                feature_map={0: feature},
                skip_indices=[],
            )

            assert feature.heading_body_boundary == len(title)
            assert feature.heading_body_boundary_source == "run_format"
            assert f"## {title}\n{body}" in markdown
            assert f"## {title}\n\n{body}" not in markdown

    def test_mixed_heading_ignores_transient_emphasis_after_an_earlier_sentence(self, tmp_path):
        from docx import Document
        from docx.shared import Pt

        from docwen_plugin_optimizer_gongwen.extraction.paragraph_reader import read_paragraphs

        title = "第一要求。重点和安排："
        body = "坚持落实各项任务。"
        doc = Document()
        paragraph = doc.add_paragraph(style="Heading 2")
        paragraph.add_run("第一要求。")
        emphasized = paragraph.add_run("重点")
        emphasized.style = "Emphasis"
        paragraph.add_run("和安排：")
        body_run = paragraph.add_run(body)
        body_run.font.name = "Calibri"
        body_run.font.size = Pt(10.5)
        body_run.bold = False
        path = tmp_path / "mixed-transient-emphasis.docx"
        doc.save(path)

        feature = read_paragraphs(Document(path))[0]

        assert feature.text == f"{title}{body}"
        assert feature.heading_level == 2
        assert feature.heading_body_boundary == len(title)
        assert feature.heading_body_boundary_source == "run_format"

    def test_heading_internal_inline_style_without_body_does_not_create_a_false_boundary(self, tmp_path):
        from docx import Document

        from docwen_plugin_optimizer_gongwen.extraction.paragraph_reader import read_paragraphs

        doc = Document()
        paragraph = doc.add_paragraph(style="Heading 2")
        paragraph.add_run("工作")
        emphasized = paragraph.add_run("要求")
        emphasized.style = "Emphasis"
        path = tmp_path / "heading-internal-emphasis.docx"
        doc.save(path)

        feature = read_paragraphs(Document(path))[0]

        assert feature.text == "工作要求"
        assert feature.heading_level == 2
        assert feature.heading_body_boundary is None

    def test_punctuation_fallback_uses_earliest_delimiter_by_text_position(self, tmp_path):
        from docx import Document

        from docwen_plugin_optimizer_gongwen.extraction.paragraph_reader import read_paragraphs

        text = "第一要求。第二要求：正文内容"
        doc = Document()
        doc.add_paragraph(text, style="Heading 2")
        path = tmp_path / "mixed-earliest-delimiter.docx"
        doc.save(path)

        feature = read_paragraphs(Document(path))[0]

        assert feature.heading_body_boundary == len("第一要求。")
        assert feature.heading_body_boundary_source == "punctuation_fallback"

    def test_single_run_mixed_heading_uses_punctuation_fallback(self, tmp_path):
        from docx import Document

        from docwen_core.text.heading_numbering import compile_clean_rules_from_data
        from docwen_plugin_optimizer_gongwen.extraction.paragraph_reader import (
            read_paragraphs,
        )

        title = "言之有谋，强化顶层设计："
        body = "坚持全局眼光，注重前瞻谋划。"
        rules = compile_clean_rules_from_data(
            [
                {
                    "id": "paren_chinese",
                    "enabled": True,
                    "pattern": r"^(（[一二三四五六七八九十]+）)",
                    "level": 2,
                }
            ]
        )
        doc = Document()
        doc.add_paragraph(f"（一）{title}{body}")
        path = tmp_path / "mixed-legacy.docx"
        doc.save(path)

        feature = read_paragraphs(Document(path), cleanup_rules=rules)[0]

        assert feature.text == f"{title}{body}"
        assert feature.heading_level == 2
        assert feature.heading_numbering_text == "（一）"
        assert feature.heading_body_boundary == len(title)
        assert feature.heading_body_boundary_source == "punctuation_fallback"

    def test_ordinary_colon_paragraph_is_not_split(self, tmp_path):
        from docx import Document

        from docwen_plugin_optimizer_gongwen.extraction.paragraph_reader import (
            read_paragraphs,
        )

        text = "说明：这是普通正文，不是标题。"
        doc = Document()
        doc.add_paragraph(text)
        path = tmp_path / "ordinary-colon.docx"
        doc.save(path)

        feature = read_paragraphs(Document(path))[0]

        assert feature.text == text
        assert feature.heading_level == 0
        assert feature.heading_body_boundary is None
        assert feature.heading_body_boundary_source == ""


def test_runtime_gongwen_route_uses_request_cleanup_for_headings_and_attachment_yaml(
    tmp_path,
):
    from pathlib import Path

    from docx import Document

    from docwen_core.models.file_ref import FileRef
    from docwen_core.models.request import ConversionRequest, OutputPolicy
    from docwen_plugin_optimizer_gongwen.plugin import GongwenOptimizerPlugin
    from docwen_runtime.config.loader import ConfigLoader
    from docwen_runtime.engine.route_resolver import RouteResolver
    from docwen_runtime.engine.task_manager import TaskManager
    from docwen_runtime.output.finalizer import OutputFinalizer
    from docwen_runtime.plugin_registry.registry import PluginRegistry
    from docwen_runtime.workspace.manager import WorkspaceManager

    project_configs = Path(__file__).resolve().parents[5] / "configs"
    snapshot = ConfigLoader(
        base_dir=project_configs,
        user_dir=tmp_path / "request-config",
    ).config.as_dict()
    snapshot["numbering"]["cleanup"] = {
        "settings": {"order": ["request"]},
        "rules": [
            {
                "id": "request",
                "enabled": True,
                "pattern": r"^REQ:\s*",
                "level": 1,
            }
        ],
    }
    source = tmp_path / "request-policy.docx"
    doc = Document()
    doc.add_paragraph("关于测试的通知")
    doc.add_paragraph("REQ: Request heading")
    doc.add_paragraph("GLOBAL: Global heading")
    doc.add_paragraph("附件：REQ: 项目清单")
    doc.save(source)
    output_dir = tmp_path / "out"
    output_dir.mkdir()
    plugins = PluginRegistry()
    plugins.register(GongwenOptimizerPlugin())
    manager = TaskManager(
        plugins,
        RouteResolver(plugins),
        WorkspaceManager(root_dir=str(tmp_path / "workspace")),
        OutputFinalizer(),
    )

    result = manager.execute_single(
        ConversionRequest(
            request_id="gongwen-request-policy",
            input_refs=[
                FileRef(
                    path=str(source),
                    format="docx",
                    category="document",
                    size_bytes=source.stat().st_size,
                )
            ],
            target_format="md",
            action_name="gongwen",
            config_snapshot=snapshot,
            output_policy=OutputPolicy(output_dir=str(output_dir)),
        )
    )

    assert result.success, result.error
    primary = next(artifact for artifact in result.artifacts if artifact.is_primary)
    markdown = Path(primary.staging_path).read_text(encoding="utf-8")
    assert "# Request heading" in markdown
    assert "GLOBAL: Global heading" in markdown
    assert "项目清单" in markdown
    assert "REQ:" not in markdown


class TestFormatFeatures:
    def test_extracts_font_name_from_first_run(self):
        from docx import Document
        from docx.shared import Pt

        from docwen_plugin_optimizer_gongwen.extraction.format_features import (
            extract_font_info,
        )

        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("test")
        run.font.name = "仿宋_GB2312"
        run.font.size = Pt(15)
        name, size = extract_font_info(p)
        assert "仿宋" in name
        assert size == 15.0

    def test_empty_paragraph_returns_defaults(self):
        from docx import Document

        from docwen_plugin_optimizer_gongwen.extraction.format_features import (
            extract_font_info,
        )

        doc = Document()
        p = doc.add_paragraph("no runs here")
        # Clear runs
        p.clear()
        name, size = extract_font_info(p)
        assert name == ""
        assert size is None


class TestTableAndTextboxExtraction:
    """Test table cell context and textbox detection (Task 6)."""

    def test_table_cell_context_detected(self):
        from docx import Document

        from docwen_plugin_optimizer_gongwen.extraction.special_content import detect_table_context

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        # Access paragraphs inside table cells
        cell_p1 = table.rows[0].cells[0].paragraphs[0]
        cell_p1.add_run("表头内容")
        cell_p2 = table.rows[1].cells[0].paragraphs[0]
        cell_p2.add_run("表格内容")

        result1 = detect_table_context(cell_p1)
        result2 = detect_table_context(cell_p2)
        assert result1 in ("body", "header")
        assert result2 in ("body", "header")

    def test_repeat_header_row_and_textbox_parent_chain_are_detected(self):
        from docx import Document
        from docx.oxml import OxmlElement

        from docwen_plugin_optimizer_gongwen.extraction.special_content import (
            detect_table_context,
            is_in_textbox,
        )

        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:tblHeader"))

        assert detect_table_context(table.cell(0, 0).paragraphs[0]) == "header"

        para = doc.add_paragraph("文本框内容")
        textbox_content = OxmlElement("w:txbxContent")
        textbox_content.append(para._p)

        assert is_in_textbox(para, doc) is True

    def test_textbox_paragraph_is_not_in_normal_paragraphs(self, tmp_path):
        from docx import Document

        from docwen_plugin_optimizer_gongwen.extraction.paragraph_reader import (
            read_paragraphs,
        )

        # A simple document without textboxes should have no textbox paragraphs
        doc = Document()
        doc.add_paragraph("普通段落")
        path = tmp_path / "plain.docx"
        doc.save(str(path))

        doc2 = Document(str(path))
        features = read_paragraphs(doc2)
        # No textbox in a plain document
        textbox_paras = [pf for pf in features if pf.is_in_textbox]
        assert len(textbox_paras) == 0

    def test_table_cell_context_empty_for_normal_paragraphs(self, tmp_path):
        from docx import Document

        from docwen_plugin_optimizer_gongwen.extraction.paragraph_reader import (
            read_paragraphs,
        )

        doc = Document()
        doc.add_paragraph("普通段落，不在表格中")
        path = tmp_path / "normal.docx"
        doc.save(str(path))

        doc2 = Document(str(path))
        features = read_paragraphs(doc2)
        for pf in features:
            assert pf.table_cell_context == ""

    def test_empty_paragraphs_tables_and_textboxes_share_one_dense_coordinate_system(self, monkeypatch):
        from docx import Document
        from docx.enum.text import WD_BREAK

        from docwen_core.docx_parsing.textbox_extraction import ExtractedParagraph
        from docwen_core.text.heading_numbering import compile_clean_rules_from_data
        from docwen_plugin_optimizer_gongwen.extraction import paragraph_reader

        doc = Document()
        doc.add_paragraph(" ")
        before = doc.add_paragraph("正文锚点")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "表格甲"
        table.cell(0, 1).text = "表格乙"
        break_paragraph = doc.add_paragraph()
        break_paragraph.add_run().add_break(WD_BREAK.PAGE)
        heading = doc.add_paragraph("一、后续标题")

        body = doc._element.body
        before_anchor = body.index(before._p)
        heading_anchor = body.index(heading._p)
        monkeypatch.setattr(
            paragraph_reader,
            "extract_textbox_paragraphs",
            lambda _doc: [
                ExtractedParagraph("文本框甲", before_anchor, "textbox", "document"),
                ExtractedParagraph("文本框乙", heading_anchor, "textbox", "document"),
            ],
        )

        cleanup_rules = compile_clean_rules_from_data(
            [{"id": "chinese", "enabled": True, "pattern": r"^[一二三四五六七八九十]+、", "level": 1}]
        )
        features = paragraph_reader.read_paragraphs(doc, cleanup_rules=cleanup_rules)

        assert [feature.text for feature in features] == [
            "正文锚点",
            "文本框甲",
            "表格甲",
            "表格乙",
            "",
            "后续标题",
            "文本框乙",
        ]
        assert [feature.index for feature in features] == list(range(7))
        assert [feature.source_index for feature in features] == [1, 1, 2, 2, 3, 4, 4]
        assert features[4].has_page_break is True
        assert features[5].heading_level == 1


class TestUtils:
    def test_contains_chinese(self):
        from docwen_plugin_optimizer_gongwen.utils import contains_chinese

        assert contains_chinese("中文测试") is True
        assert contains_chinese("English") is False
        assert contains_chinese("混合Mixed") is True

    def test_convert_date_format(self):
        from docwen_plugin_optimizer_gongwen.utils import convert_date_format

        assert convert_date_format("2024年1月15日") == "2024年1月15日"
        assert convert_date_format("2024-01-15") == "2024年1月15日"
        assert convert_date_format("2024/12/1") == "2024年12月1日"

    def test_remove_brackets(self):
        from docwen_plugin_optimizer_gongwen.utils import remove_brackets

        assert remove_brackets("（内容）") == "内容"
        assert remove_brackets("【内容】") == "内容"
        assert remove_brackets("no brackets") == "no brackets"

    def test_process_attachment_item(self):
        from docwen_core.text.heading_numbering import compile_clean_rules_from_data
        from docwen_plugin_optimizer_gongwen.utils import process_attachment_item

        rules = compile_clean_rules_from_data(
            [{"id": "arabic", "enabled": True, "pattern": r"^\d+[.．、]", "level": 3}]
        )
        assert process_attachment_item("1. 公文格式标准", cleanup_rules=rules) == "公文格式标准"
        assert process_attachment_item("2、公文流转程序", cleanup_rules=rules) == "公文流转程序"

    def test_process_attachment_item_accepts_request_cleanup_rules(self):
        from docwen_core.text.heading_numbering import compile_clean_rules_from_data
        from docwen_plugin_optimizer_gongwen.utils import process_attachment_item

        rules = compile_clean_rules_from_data([{"id": "request", "enabled": True, "pattern": r"^REQ:\s*", "level": 1}])

        assert process_attachment_item("REQ: 附件清单", cleanup_rules=rules) == "附件清单"
