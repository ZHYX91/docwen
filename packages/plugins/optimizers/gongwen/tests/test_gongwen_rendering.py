"""Tests for gongwen rendering layer."""

from __future__ import annotations

import pytest

pytestmark = pytest.mark.contract


class TestMarkdownRenderer:
    def test_renders_yaml_frontmatter_with_title(self):
        from docwen_plugin_optimizer_gongwen.models import GongwenMetadata
        from docwen_plugin_optimizer_gongwen.rendering.markdown_renderer import render

        m = GongwenMetadata.default()
        m.title = "关于xxx的通知"
        result = render(m, body_lines=["第一段正文。"], skip_indices=[], feature_map={})
        assert "---" in result
        assert "关于xxx的通知" in result
        assert "第一段正文。" in result

    def test_nonempty_metadata_renders_complete_18_field_schema(self):
        import yaml

        from docwen_plugin_optimizer_gongwen.models import GongwenMetadata
        from docwen_plugin_optimizer_gongwen.rendering.markdown_renderer import render

        metadata = GongwenMetadata.default()
        metadata.title = "关于xxx的通知"

        result = render(metadata, body_lines=["正文。"], skip_indices=[], feature_map={})
        frontmatter = yaml.safe_load(result.split("---", maxsplit=2)[1])

        assert frontmatter == metadata.to_dict()
        assert len(frontmatter) == 18

    def test_empty_metadata_omits_frontmatter(self):
        from docwen_plugin_optimizer_gongwen.models import GongwenMetadata
        from docwen_plugin_optimizer_gongwen.rendering.markdown_renderer import render

        m = GongwenMetadata.default()
        result = render(m, body_lines=["正文。"], skip_indices=[], feature_map={})
        assert "---" not in result

    def test_skips_structural_paragraphs(self):
        from docwen_plugin_optimizer_gongwen.models import GongwenMetadata
        from docwen_plugin_optimizer_gongwen.rendering.markdown_renderer import render

        m = GongwenMetadata.default()
        result = render(
            m,
            body_lines=["Line0", "Line1", "Line2"],
            skip_indices=[0, 2],
            feature_map={},
        )
        assert "Line0" not in result
        assert "Line1" in result
        assert "Line2" not in result

    # ── Rich content rendering tests (Task 2) ──

    def test_body_paragraph_with_inline_formula_renders_latex(self):
        from docwen_plugin_optimizer_gongwen.models import GongwenMetadata, ParagraphFeature
        from docwen_plugin_optimizer_gongwen.rendering.markdown_renderer import render

        pf = ParagraphFeature(index=0, text="能量公式", has_formula=True, formula_type="inline", formula_latex="E=mc^2")
        result = render(GongwenMetadata.default(), ["能量公式"], feature_map={0: pf}, skip_indices=[])
        assert "$E=mc^2$" in result

    def test_body_paragraph_with_block_formula(self):
        from docwen_plugin_optimizer_gongwen.models import GongwenMetadata, ParagraphFeature
        from docwen_plugin_optimizer_gongwen.rendering.markdown_renderer import render

        pf = ParagraphFeature(
            index=0, text="", has_formula=True, formula_type="block", formula_latex="\\sum_{i=1}^n x_i"
        )
        result = render(GongwenMetadata.default(), [""], feature_map={0: pf}, skip_indices=[])
        assert "$$" in result
        assert "\\sum_{i=1}^n x_i" in result

    def test_body_paragraph_with_image_renders_markdown_image(self):
        from docwen_plugin_optimizer_gongwen.models import GongwenMetadata, ParagraphFeature
        from docwen_plugin_optimizer_gongwen.rendering.markdown_renderer import render

        pf = ParagraphFeature(index=0, text="图片说明", extracted_images=["img_001.png"])
        result = render(GongwenMetadata.default(), ["图片说明"], feature_map={0: pf}, skip_indices=[])
        assert "![](img_001.png)" in result

    def test_skipped_structural_paragraph_preserves_relative_image_reference(self):
        from docwen_plugin_optimizer_gongwen.models import GongwenMetadata, ParagraphFeature
        from docwen_plugin_optimizer_gongwen.rendering.markdown_renderer import render

        pf = ParagraphFeature(
            index=0,
            text="issuing authority signature",
            extracted_images=[r"C:\staging\gongwen-seal.png"],
        )

        result = render(
            GongwenMetadata.default(),
            [pf.text],
            feature_map={0: pf},
            skip_indices=[0],
        )

        assert "issuing authority signature" not in result
        assert "![](gongwen-seal.png)" in result
        assert "C:\\staging" not in result

    def test_attachment_line_preserves_original_prefix_when_numbering_is_kept(self):
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.pipeline import _attachment_line_text

        feature = ParagraphFeature(
            index=0,
            text="指导思想",
            raw_text="一、指导思想",
            heading_level=1,
            heading_numbering_text="一、",
        )

        assert _attachment_line_text(feature, remove_numbering=False) == "一、指导思想"
        assert _attachment_line_text(feature, remove_numbering=True) == "指导思想"

    def test_heading_detected_and_rendered(self):
        from docwen_plugin_optimizer_gongwen.models import GongwenMetadata, ParagraphFeature
        from docwen_plugin_optimizer_gongwen.rendering.markdown_renderer import render

        pf = ParagraphFeature(index=0, text="重点工作", heading_level=2)
        result = render(GongwenMetadata.default(), ["重点工作"], feature_map={0: pf}, skip_indices=[])
        assert "## 重点工作" in result

    def test_mixed_heading_body_splits_once_without_shifting_following_feature(self):
        from docwen_plugin_optimizer_gongwen.models import GongwenMetadata, ParagraphFeature
        from docwen_plugin_optimizer_gongwen.rendering.markdown_renderer import render

        mixed = ParagraphFeature(
            index=0,
            text="言之有谋，强化顶层设计：坚持全局眼光。注重前瞻谋划。",
            heading_level=2,
            heading_body_boundary=len("言之有谋，强化顶层设计："),
            heading_body_boundary_source="run_format",
        )
        following = ParagraphFeature(
            index=1,
            text="下一页正文。",
            has_page_break=True,
        )

        result = render(
            GongwenMetadata.default(),
            [mixed.text, following.text],
            feature_map={0: mixed, 1: following},
            skip_indices=[],
        )

        assert "## 言之有谋，强化顶层设计：\n坚持全局眼光。注重前瞻谋划。" in result
        assert result.count("## ") == 1
        assert "---\n\n下一页正文。" in result

    def test_page_break_renders_separator(self):
        from docwen_plugin_optimizer_gongwen.models import GongwenMetadata, ParagraphFeature
        from docwen_plugin_optimizer_gongwen.rendering.markdown_renderer import render

        pf = ParagraphFeature(index=0, text="下一页内容", has_page_break=True)
        result = render(GongwenMetadata.default(), ["下一页内容"], feature_map={0: pf}, skip_indices=[])
        assert "---" in result

    def test_section_break_renders_separator(self):
        from docwen_plugin_optimizer_gongwen.models import GongwenMetadata, ParagraphFeature
        from docwen_plugin_optimizer_gongwen.rendering.markdown_renderer import render

        pf = ParagraphFeature(index=0, text="分节内容", has_section_break=True)
        result = render(GongwenMetadata.default(), ["分节内容"], feature_map={0: pf}, skip_indices=[])
        assert "---" in result

    # ── Numbering remove/add tests (Plan 1 阶段 C) ──

    def test_remove_numbering_true_strips_prefix(self):
        from docwen_plugin_optimizer_gongwen.models import GongwenMetadata, ParagraphFeature
        from docwen_plugin_optimizer_gongwen.rendering.markdown_renderer import render

        pf = ParagraphFeature(
            index=0,
            text="加强组织领导",
            heading_level=1,
            heading_numbering_text="一、",
        )
        result = render(
            GongwenMetadata.default(),
            ["加强组织领导"],
            feature_map={0: pf},
            skip_indices=[],
            remove_numbering=True,
        )
        assert "# 加强组织领导" in result
        assert "# 一、加强组织领导" not in result

    def test_remove_numbering_false_preserves_prefix(self):
        from docwen_plugin_optimizer_gongwen.models import GongwenMetadata, ParagraphFeature
        from docwen_plugin_optimizer_gongwen.rendering.markdown_renderer import render

        pf = ParagraphFeature(
            index=0,
            text="加强组织领导",
            heading_level=1,
            heading_numbering_text="一、",
        )
        result = render(
            GongwenMetadata.default(),
            ["加强组织领导"],
            feature_map={0: pf},
            skip_indices=[],
            remove_numbering=False,
        )
        assert "# 一、加强组织领导" in result

    def test_add_numbering_with_gongwen_standard(self):
        from docwen_core.text.heading_numbering import HeadingFormatter
        from docwen_plugin_optimizer_gongwen.models import GongwenMetadata, ParagraphFeature
        from docwen_plugin_optimizer_gongwen.rendering.markdown_renderer import render

        scheme_config = {
            "level_1": {"format": "{1.chinese_lower}、"},
            "level_2": {"format": "（{2.chinese_lower}）"},
        }
        formatter = HeadingFormatter(scheme_config)

        pf = ParagraphFeature(
            index=0,
            text="加强组织领导",
            heading_level=1,
            heading_numbering_text="一、",
        )
        result = render(
            GongwenMetadata.default(),
            ["加强组织领导"],
            feature_map={0: pf},
            skip_indices=[],
            remove_numbering=True,
            heading_formatter=formatter,
        )
        # Old "一、" removed, new "一、" added by gongwen_standard scheme
        assert "# 一、加强组织领导" in result

    def test_add_numbering_with_hierarchical_standard(self):
        from docwen_core.text.heading_numbering import HeadingFormatter
        from docwen_plugin_optimizer_gongwen.models import GongwenMetadata, ParagraphFeature
        from docwen_plugin_optimizer_gongwen.rendering.markdown_renderer import render

        scheme_config = {
            "level_1": {"format": "{1.arabic_half} "},
            "level_2": {"format": "{1.arabic_half}.{2.arabic_half} "},
        }
        formatter = HeadingFormatter(scheme_config)

        pf = ParagraphFeature(
            index=0,
            text="加强组织领导",
            heading_level=1,
            heading_numbering_text="一、",
        )
        result = render(
            GongwenMetadata.default(),
            ["加强组织领导"],
            feature_map={0: pf},
            skip_indices=[],
            remove_numbering=True,
            heading_formatter=formatter,
        )
        # Hierarchical: old Chinese numbering removed, new "1 " added
        assert "# 1 加强组织领导" in result

    def test_multiple_headings_with_formatter_counter(self):
        from docwen_core.text.heading_numbering import HeadingFormatter
        from docwen_plugin_optimizer_gongwen.models import GongwenMetadata, ParagraphFeature
        from docwen_plugin_optimizer_gongwen.rendering.markdown_renderer import render

        scheme_config = {
            "level_1": {"format": "{1.chinese_lower}、"},
            "level_2": {"format": "（{2.chinese_lower}）"},
        }
        formatter = HeadingFormatter(scheme_config)

        pf0 = ParagraphFeature(index=0, text="Introduction", heading_level=1)
        pf1 = ParagraphFeature(index=1, text="Background", heading_level=2)
        pf2 = ParagraphFeature(index=2, text="Methods", heading_level=1)

        body_lines = ["Introduction", "Background", "Methods"]
        feature_map = {0: pf0, 1: pf1, 2: pf2}

        result = render(
            GongwenMetadata.default(),
            body_lines,
            feature_map=feature_map,
            skip_indices=[],
            remove_numbering=True,
            heading_formatter=formatter,
        )
        assert "# 一、Introduction" in result
        assert "## （一）Background" in result
        assert "# 二、Methods" in result

    def test_remove_and_add_with_existing_numbering(self):
        from docwen_core.text.heading_numbering import HeadingFormatter
        from docwen_plugin_optimizer_gongwen.models import GongwenMetadata, ParagraphFeature
        from docwen_plugin_optimizer_gongwen.rendering.markdown_renderer import render

        scheme_config = {
            "level_1": {"format": "{1.chinese_lower}、"},
        }
        formatter = HeadingFormatter(scheme_config)

        pf = ParagraphFeature(
            index=0,
            text="Existing Title",
            heading_level=1,
            heading_numbering_text="三、",
        )
        result = render(
            GongwenMetadata.default(),
            ["Existing Title"],
            feature_map={0: pf},
            skip_indices=[],
            remove_numbering=True,
            heading_formatter=formatter,
        )
        # Old "三、" removed, new "一、" added by formatter starting from 1
        assert "# 一、Existing Title" in result
        assert "# 三、Existing Title" not in result


class TestAttachmentRenderer:
    def test_renders_attachment_header_and_body(self):
        from docwen_plugin_optimizer_gongwen.models import GongwenMetadata
        from docwen_plugin_optimizer_gongwen.rendering.attachment_renderer import (
            render_attachment,
        )

        m = GongwenMetadata.default()
        m.attachment = ["关于xxx的通知", "关于yyy的办法"]
        result = render_attachment(m, ["附件正文第一段。", "附件正文第二段。"])
        assert "## 附件" in result
        assert "1. 关于xxx的通知" in result
        assert "2. 关于yyy的办法" in result
        assert "附件正文第一段。" in result

    def test_empty_attachment_lines_returns_empty(self):
        from docwen_plugin_optimizer_gongwen.models import GongwenMetadata
        from docwen_plugin_optimizer_gongwen.rendering.attachment_renderer import (
            render_attachment,
        )

        m = GongwenMetadata.default()
        result = render_attachment(m, [])
        assert result == ""


class TestDocxFields:
    # ── DOCX XML manipulation tests (Task 3) ──

    def test_attachment_placeholder_replaced_in_docx(self, tmp_path):
        from docx import Document

        doc = Document()
        doc.add_paragraph("普通段落")
        doc.add_paragraph("{{附件说明}}")
        path = tmp_path / "test.docx"
        doc.save(str(path))

        from docwen_plugin_optimizer_gongwen.rendering.docx_fields import (
            replace_attachment_placeholder,
        )

        yaml_data = {"附件说明": ["附件：1. 预算报表", "     2. 项目清单"]}
        replace_attachment_placeholder(doc, yaml_data)

        texts = [p.text for p in doc.paragraphs]
        assert "{{附件说明}}" not in "".join(texts)
        assert "预算报表" in "".join(texts)

    def test_attachment_placeholder_removed_when_empty(self, tmp_path):
        from docx import Document

        doc = Document()
        doc.add_paragraph("普通段落")
        doc.add_paragraph("{{附件说明}}")
        path = tmp_path / "test.docx"
        doc.save(str(path))

        from docwen_plugin_optimizer_gongwen.rendering.docx_fields import (
            replace_attachment_placeholder,
        )

        yaml_data = {"附件说明": []}
        replace_attachment_placeholder(doc, yaml_data)

        texts = [p.text for p in doc.paragraphs]
        assert "{{附件说明}}" not in "".join(texts)
        # Only the first paragraph should remain
        assert len(doc.paragraphs) == 1

    def test_empty_field_deletes_row(self, tmp_path):
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "{{抄送机关}}"
        table.rows[1].cells[0].text = "其他内容"
        path = tmp_path / "test.docx"
        doc.save(str(path))

        from docwen_plugin_optimizer_gongwen.rendering.docx_fields import (
            apply_empty_field_rules,
        )

        apply_empty_field_rules(doc, {"抄送机关": ""})

        # Row containing {{抄送机关}} should be removed
        remaining_texts = []
        for row in table.rows:
            for cell in row.cells:
                remaining_texts.append(cell.text)
        assert "{{抄送机关}}" not in remaining_texts

    def test_empty_field_deletes_paragraph(self, tmp_path):
        from docx import Document

        doc = Document()
        doc.add_paragraph("{{密级和保密期限}}")
        doc.add_paragraph("{{发文字号}}")
        doc.add_paragraph("其他内容")
        path = tmp_path / "test.docx"
        doc.save(str(path))

        from docwen_plugin_optimizer_gongwen.rendering.docx_fields import (
            apply_empty_field_rules,
        )

        yaml_data = {"密级和保密期限": "", "发文字号": ""}
        apply_empty_field_rules(doc, yaml_data)

        remaining = [p.text for p in doc.paragraphs]
        assert "{{密级和保密期限}}" not in remaining
        assert "{{发文字号}}" not in remaining
        assert "其他内容" in remaining

    def test_no_placeholder_found_no_error(self, tmp_path):
        """No placeholder in document should not raise error."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("没有占位符")
        path = tmp_path / "test.docx"
        doc.save(str(path))

        from docwen_plugin_optimizer_gongwen.rendering.docx_fields import (
            replace_attachment_placeholder,
        )

        replace_attachment_placeholder(doc, {"附件说明": ["test"]})
        assert len(doc.paragraphs) == 1


class TestValidation:
    def test_validates_required_fields(self):
        from docwen_plugin_optimizer_gongwen.models import (
            RecognitionResult,
        )
        from docwen_plugin_optimizer_gongwen.validation import validate_result

        result = RecognitionResult(
            candidates={},
            yaml_info={},
            skip_indices=[],
            review_signals=[],
            missing_required=[],
            validation_finding_count=0,
        )

        result = validate_result(result)
        assert result.missing_required == ["issue_date", "issuing_authority_signature", "title"]

    def test_validates_with_all_required(self):
        from docwen_plugin_optimizer_gongwen.models import (
            RecognitionCandidate,
            RecognitionResult,
        )
        from docwen_plugin_optimizer_gongwen.validation import validate_result

        result = RecognitionResult(
            candidates={
                0: RecognitionCandidate(element_type="title", score=120, para_index=0),
                5: RecognitionCandidate(
                    element_type="issuing_authority_signature",
                    score=100,
                    para_index=5,
                ),
                6: RecognitionCandidate(element_type="issue_date", score=110, para_index=6),
            },
            yaml_info={},
            skip_indices=[],
            review_signals=[],
            missing_required=[],
            validation_finding_count=0,
        )

        result = validate_result(result)
        assert len(result.missing_required) == 0

    def test_detects_structural_ordering_issue(self):
        from docwen_plugin_optimizer_gongwen.models import (
            RecognitionCandidate,
            RecognitionResult,
        )
        from docwen_plugin_optimizer_gongwen.validation import validate_result

        # issue_date before issuing_authority_signature = wrong order
        result = RecognitionResult(
            candidates={
                0: RecognitionCandidate(element_type="issue_date", score=110, para_index=0),
                1: RecognitionCandidate(
                    element_type="issuing_authority_signature",
                    score=100,
                    para_index=1,
                ),
                2: RecognitionCandidate(element_type="title", score=120, para_index=2),
            },
            yaml_info={},
            skip_indices=[],
            review_signals=[],
            missing_required=[],
            validation_finding_count=0,
        )

        result = validate_result(result)
        # Should have at least 1 structural finding (sig should be before date, but date is at 0, sig at 1)
        assert result.validation_finding_count >= 1
        assert any("structural" in s for s in result.review_signals)

    def test_get_confidence_summary(self):
        from docwen_plugin_optimizer_gongwen.models import (
            RecognitionCandidate,
            RecognitionResult,
        )
        from docwen_plugin_optimizer_gongwen.validation import get_confidence_summary

        result = RecognitionResult(
            candidates={
                0: RecognitionCandidate(
                    element_type="title",
                    score=120,
                    para_index=0,
                    confidence="high",
                ),
                1: RecognitionCandidate(
                    element_type="issue_date",
                    score=100,
                    para_index=1,
                    confidence="medium",
                ),
            },
            yaml_info={},
            skip_indices=[],
            review_signals=[],
            missing_required=[],
            validation_finding_count=0,
        )

        summary = get_confidence_summary(result)
        assert summary["overall"] in ("high", "medium", "low", "none")
        assert "title" in summary["fields"]
        assert summary["fields"]["title"] == "high"

    def test_get_confidence_summary_empty(self):
        from docwen_plugin_optimizer_gongwen.models import RecognitionResult
        from docwen_plugin_optimizer_gongwen.validation import get_confidence_summary

        result = RecognitionResult(
            candidates={},
            yaml_info={},
            skip_indices=[],
            review_signals=[],
            missing_required=[],
            validation_finding_count=0,
        )

        summary = get_confidence_summary(result)
        assert summary["overall"] == "none"
        assert summary["fields"] == {}


class TestRuntimeConfig:
    def test_default_config(self):
        from docwen_plugin_optimizer_gongwen.runtime_config import (
            DEFAULT_CONFIG,
            GongwenContentRuntimeConfig,
        )

        assert isinstance(DEFAULT_CONFIG, GongwenContentRuntimeConfig)
        assert DEFAULT_CONFIG.horizontal_rule_enabled is False
        assert DEFAULT_CONFIG.page_break_marker == "---"

    def test_custom_config(self):
        from docwen_plugin_optimizer_gongwen.runtime_config import (
            configure_gongwen_content_runtime,
        )

        cfg = configure_gongwen_content_runtime(
            horizontal_rule_enabled=True,
            page_break_marker="***",
        )
        assert cfg.horizontal_rule_enabled is True
        assert cfg.page_break_marker == "***"
