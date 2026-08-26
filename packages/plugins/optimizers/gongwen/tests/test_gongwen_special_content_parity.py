from __future__ import annotations

import pytest
from docx import Document

pytestmark = pytest.mark.contract

from docwen_plugin_optimizer_gongwen.extraction.paragraph_reader import read_paragraphs
from docwen_plugin_optimizer_gongwen.models import GongwenMetadata
from docwen_plugin_optimizer_gongwen.rendering.markdown_renderer import render


def test_gongwen_heading_render_uses_cleaned_text():
    from docwen_core.text.heading_numbering import compile_clean_rules_from_data

    doc = Document()
    doc.add_paragraph("一、总体要求")

    rules = compile_clean_rules_from_data(
        [{"id": "chinese", "enabled": True, "pattern": r"^[一二三四五六七八九十]+、", "level": 1}]
    )
    features = read_paragraphs(doc, cleanup_rules=rules)
    assert features[0].raw_text == "一、总体要求"
    assert features[0].text == "总体要求"
    assert features[0].heading_numbering_text == "一、"

    md = render(
        GongwenMetadata.default(),
        [f.text for f in features],
        feature_map={0: features[0]},
        remove_numbering=False,
    )
    assert "# 一、总体要求" in md
    assert "# 总体要求" not in md


def test_gongwen_includes_injected_textbox_paragraph(monkeypatch):
    doc = Document()
    doc.add_paragraph("正文锚点")

    from docwen_core.docx_parsing.textbox_extraction import ExtractedParagraph
    from docwen_plugin_optimizer_gongwen.extraction import paragraph_reader

    monkeypatch.setattr(
        paragraph_reader,
        "extract_textbox_paragraphs",
        lambda document: [ExtractedParagraph("文本框内容", 0, "textbox", "document")],
    )

    features = read_paragraphs(doc)

    assert [feature.text for feature in features] == ["正文锚点", "文本框内容"]
    assert features[1].is_in_textbox is True
    assert features[1].source == "textbox"


def test_gongwen_includes_table_cell_text_for_recognition():
    doc = Document()
    doc.add_paragraph("正文")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "版记"
    table.cell(0, 1).text = "公开方式：主动公开"
    table.cell(1, 0).text = "抄送机关"
    table.cell(1, 1).text = "省政府办公厅"

    features = read_paragraphs(doc)
    texts = [feature.text for feature in features]

    assert "公开方式：主动公开" in texts
    assert "省政府办公厅" in texts
    table_features = [feature for feature in features if feature.source == "table"]
    assert table_features
    assert all(feature.table_cell_context in {"header", "body"} for feature in table_features)


def test_gongwen_merged_table_cells_are_not_duplicated():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "合并表头"
    table.cell(1, 0).text = "甲"
    table.cell(1, 1).text = "乙"

    features = read_paragraphs(doc)
    table_features = [feature for feature in features if feature.source == "table"]

    assert [feature.text for feature in table_features] == ["合并表头", "甲", "乙"]
    assert sum(feature.is_table_anchor for feature in table_features) == 1


def test_plain_projection_of_rich_table_emits_structured_review_warning():
    from docwen_plugin_optimizer_gongwen.pipeline import convert_docx_to_md_gongwen

    doc = Document()
    title = doc.add_paragraph("关于测试的通知")
    title.style = "Title"
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "表格内容"
    table.cell(0, 0).paragraphs[0].runs[0].bold = True

    result = convert_docx_to_md_gongwen(doc, "table.docx", {})
    warnings = result["metadata"]["recognition_review_signals"]["gongwen_warnings"]

    warning = next(item for item in warnings if item["code"] == "GW002")
    assert warning["scope"] == "table_semantics"
    assert warning["details"]["table_cell_feature_count"] == 1
    assert warning["details"]["fidelity_risks"] == ["rich_text"]


def test_nested_only_table_still_emits_a_precise_review_warning():
    from docwen_plugin_optimizer_gongwen.pipeline import convert_docx_to_md_gongwen

    doc = Document()
    title = doc.add_paragraph("关于测试的通知")
    title.style = "Title"
    outer = doc.add_table(rows=1, cols=1)
    nested = outer.cell(0, 0).add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "嵌套内容"

    result = convert_docx_to_md_gongwen(doc, "nested-table.docx", {})
    warnings = result["metadata"]["recognition_review_signals"]["gongwen_warnings"]

    warning = next(item for item in warnings if item["code"] == "GW002")
    assert "nested_table" in warning["details"]["fidelity_risks"]


def test_non_structural_table_renders_once_with_rows_and_columns_preserved():
    from docwen_plugin_optimizer_gongwen.pipeline import convert_docx_to_md_gongwen

    doc = Document()
    doc.add_paragraph("普通正文")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "列甲"
    table.cell(0, 1).text = "列乙"
    table.cell(1, 0).text = "值甲"
    table.cell(1, 1).text = "值乙"

    result = convert_docx_to_md_gongwen(doc, "table.docx", {})

    assert result["markdown"].count("| 列甲 | 列乙 |") == 1
    assert "| --- | --- |" in result["markdown"]
    assert result["markdown"].count("| 值甲 | 值乙 |") == 1
    warnings = result["metadata"]["recognition_review_signals"]["gongwen_warnings"]
    assert not any(item["code"] == "GW002" for item in warnings)


def test_mixed_metadata_and_body_table_is_not_dropped_from_markdown():
    from docwen_plugin_optimizer_gongwen.pipeline import convert_docx_to_md_gongwen

    doc = Document()
    title = doc.add_paragraph("关于测试的通知")
    title.style = "Title"
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "公开方式：主动公开"
    table.cell(0, 1).text = "版记"
    table.cell(1, 0).text = "工作事项"
    table.cell(1, 1).text = "正文内容"

    result = convert_docx_to_md_gongwen(doc, "mixed-table.docx", {})

    assert result["yaml_info"]["公开方式"] == "主动公开"
    assert "| 工作事项 | 正文内容 |" in result["markdown"]


def test_labelled_edition_table_stays_out_of_body_when_its_label_scores_as_body():
    from docwen_plugin_optimizer_gongwen.pipeline import convert_docx_to_md_gongwen

    doc = Document()
    title = doc.add_paragraph("关于测试的通知")
    title.style = "Title"
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "抄送：省委组织部。"
    table.cell(1, 0).text = "测试印发机关"
    table.cell(1, 1).text = "2025年7月5日印发"

    result = convert_docx_to_md_gongwen(doc, "edition-table.docx", {})

    assert result["yaml_info"]["印发机关"] == "测试印发机关"
    assert result["yaml_info"]["印发日期"] == "2025年7月5日"
    assert "| 抄送：省委组织部。" not in result["markdown"]


def test_rich_structural_table_reports_fidelity_loss_even_when_removed_from_body():
    from docwen_plugin_optimizer_gongwen.pipeline import convert_docx_to_md_gongwen

    doc = Document()
    title = doc.add_paragraph("关于测试的通知")
    title.style = "Title"
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "公开方式：主动公开"
    table.cell(0, 0).paragraphs[0].runs[0].bold = True

    result = convert_docx_to_md_gongwen(doc, "rich-structural-table.docx", {})
    warnings = result["metadata"]["recognition_review_signals"]["gongwen_warnings"]

    warning = next(item for item in warnings if item["code"] == "GW002")
    assert warning["details"]["rendered_table_count"] == 0
    assert warning["details"]["structural_table_count"] == 1
    assert warning["details"]["fidelity_risks"] == ["rich_text"]
