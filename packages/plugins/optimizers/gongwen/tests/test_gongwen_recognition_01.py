"""Focused tests split from test_gongwen_recognition.py."""

from __future__ import annotations

import pytest

pytestmark = pytest.mark.contract


def test_scoring_rule_failure_is_reported_once_per_rule(monkeypatch) -> None:
    """A fallible recognition heuristic must degrade visibly, not disappear."""
    from tests.support.progress import FakeProgressSink

    from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
    from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

    feature = ParagraphFeature(index=7, text="关于推进工作的通知")
    progress = FakeProgressSink()
    scorer = ElementScorer(diagnostic_sink=progress)

    def _broken_rule(_feature) -> bool:
        raise ValueError("malformed style probe")

    monkeypatch.setattr(scorer, "matches_title_pattern", _broken_rule)
    scorer.reset_context([feature])
    scorer.score_round(feature, round_group="round1")
    scorer.score_round(feature, round_group="round1")

    matching = [diagnostic for diagnostic in progress.diagnostics if diagnostic[2] == "GONGWEN-SCORING-RULE-SKIPPED"]
    assert len(matching) == 1
    assert matching[0][0] == "warning"
    assert matching[0][3] == "paragraph 7"
    assert scorer.rule_failures == (
        {
            "paragraph_index": 7,
            "element_type": "title",
            "condition": "matches_title_pattern",
            "exception_type": "ValueError",
        },
    )


class TestElementScorer:
    """Test the ElementScorer class and individual checker methods."""

    def test_round_groups_preserve_old_system_tie_break_order(self):
        """Equal scores must use the old deterministic element priority."""
        from docwen_plugin_optimizer_gongwen.constants import (
            ROUND1_ELEMENTS,
            ROUND2_ELEMENTS,
            ROUND3_ELEMENTS,
        )

        assert ROUND1_ELEMENTS == (
            "combined_id",
            "copy_id",
            "security",
            "urgency",
            "doc_number",
            "combined_doc_number_signer",
            "signer",
            "title",
            "recipient",
            "attachment_header",
            "issuing_authority_signature",
            "issue_date",
            "notes",
            "disclosure",
            "copy_to",
            "printing_date",
        )
        assert ROUND2_ELEMENTS == ("issuing_authority_mark", "printing_authority")
        assert ROUND3_ELEMENTS == (
            "title_following",
            "subtitle",
            "subtitle_following",
            "body",
            "attachment_following",
            "signer_following",
            "combined_doc_number_signer_following",
            "attachment_content",
        )

    def test_scores_title_paragraph_high(self):
        """A paragraph with 小标宋 font and 22pt size should score high for title."""
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        pf = ParagraphFeature(
            index=0,
            text="关于进一步加强安全生产工作的通知",
            font_name="小标宋",
            font_size_pt=22.0,
        )
        scorer = ElementScorer()
        scorer.reset_context([pf])
        candidates = scorer.score_round(pf, round_group="round1")
        title_candidates = [c for c in candidates if c.element_type == "title"]
        assert len(title_candidates) > 0
        assert title_candidates[0].score >= 80

    def test_effective_cjk_run_format_reaches_title_scoring(self, tmp_path):
        """A decorative leading run must not suppress official title facts."""
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt

        from docwen_plugin_optimizer_gongwen.extraction.paragraph_reader import (
            read_paragraphs,
        )
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        doc = Document()
        para = doc.add_paragraph()
        leading = para.add_run(" ")
        leading.font.name = "Calibri"
        leading.font.size = Pt(8)
        title = para.add_run("关于进一步加强安全生产工作的通知")
        r_pr = title._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        r_fonts.set(qn("w:eastAsia"), "方正小标宋简体")
        title.font.size = Pt(22)

        feature = read_paragraphs(doc, output_dir=str(tmp_path))[0]
        assert feature.font_name == "方正小标宋简体"
        assert feature.font_size_pt == 22.0

        scorer = ElementScorer()
        scorer.reset_context([feature])
        candidates = scorer.score_round(feature, round_group="round1")
        title_candidate = next(candidate for candidate in candidates if candidate.element_type == "title")
        assert title_candidate.score == 160

    def test_scores_doc_number_paragraph(self):
        """A paragraph matching doc number pattern should score high."""
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        pf = ParagraphFeature(index=0, text="国办发〔2024〕5号")
        scorer = ElementScorer()
        scorer.reset_context([pf])
        candidates = scorer.score_round(pf, round_group="round1")
        doc_candidates = [c for c in candidates if c.element_type == "doc_number"]
        assert len(doc_candidates) > 0
        assert doc_candidates[0].score >= 80

    def test_scores_security_paragraph(self):
        """A paragraph starting with security keyword should score for security."""
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        pf = ParagraphFeature(index=0, text="绝密")
        scorer = ElementScorer()
        scorer.reset_context([pf])
        candidates = scorer.score_round(pf, round_group="round1")
        sec = [c for c in candidates if c.element_type == "security"]
        assert len(sec) > 0

    def test_scores_urgency_paragraph(self):
        """A paragraph with urgency keyword should score for urgency."""
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        pf = ParagraphFeature(index=0, text="特急")
        scorer = ElementScorer()
        scorer.reset_context([pf])
        candidates = scorer.score_round(pf, round_group="round1")
        urg = [c for c in candidates if c.element_type == "urgency"]
        assert len(urg) > 0

    def test_scores_recipient_paragraph(self):
        """A paragraph with colon ending and 各 should score for recipient."""
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        pf = ParagraphFeature(index=3, text="各省、自治区、直辖市人民政府：")
        scorer = ElementScorer()
        # Need to simulate that a title was found before recipient
        scorer.reset_context([pf])
        scorer.update_context("title", 0)
        candidates = scorer.score_round(pf, round_group="round1")
        rec = [c for c in candidates if c.element_type == "recipient"]
        assert len(rec) > 0
        assert rec[0].score >= 80

    def test_scores_attachment_header(self):
        """A paragraph starting with 附件： should score for attachment."""
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        pf = ParagraphFeature(index=10, text="附件：1. 项目清单")
        scorer = ElementScorer()
        scorer.reset_context([pf])
        scorer.update_context("title", 0)
        candidates = scorer.score_round(pf, round_group="round1")
        att = [c for c in candidates if c.element_type == "attachment_header"]
        assert len(att) > 0

    def test_scores_issue_date_paragraph(self):
        """A standalone date should score for issue_date."""
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        pf = ParagraphFeature(index=20, text="2024年1月15日")
        scorer = ElementScorer()
        scorer.reset_context([pf])
        scorer.update_context("issuing_authority_signature", 19)
        candidates = scorer.score_round(pf, round_group="round1")
        date = [c for c in candidates if c.element_type == "issue_date"]
        assert len(date) > 0

    def test_empty_paragraph_scores_nothing(self):
        """An empty paragraph should not match any element type."""
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        scorer = ElementScorer()
        scorer.reset_context([ParagraphFeature(index=0, text="")])
        candidates = scorer.score_round(
            ParagraphFeature(index=0, text=""),
            round_group="round1",
        )
        assert len(candidates) == 0

    def test_scores_subtitle_paragraph(self):
        """A paragraph starting with —— should score for subtitle in round3."""
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        pf = ParagraphFeature(index=2, text="——以高质量发展为主题")
        scorer = ElementScorer()
        scorer.reset_context([pf])
        scorer.update_context("title", 1)
        candidates = scorer.score_round(pf, round_group="round3")
        sub = [c for c in candidates if c.element_type == "subtitle"]
        assert len(sub) > 0

    def test_scores_body_paragraph(self):
        """A Chinese-text paragraph after title should score for body."""
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        pf = ParagraphFeature(index=5, text="为贯彻落实国务院关于安全生产工作的决策部署，现就有关事项通知如下：")
        scorer = ElementScorer()
        scorer.reset_context([pf])
        scorer.update_context("title", 0)
        scorer.update_context("recipient", 4)
        candidates = scorer.score_round(pf, round_group="round3")
        body = [c for c in candidates if c.element_type == "body"]
        assert len(body) > 0

    def test_scores_signer_paragraph(self):
        """A paragraph with 签发人： should score for signer."""
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        pf = ParagraphFeature(index=2, text="签发人：张三")
        scorer = ElementScorer()
        scorer.reset_context([pf])
        scorer.update_context("issuing_authority_mark", 0)
        candidates = scorer.score_round(pf, round_group="round1")
        sig = [c for c in candidates if c.element_type == "signer"]
        assert len(sig) > 0

    def test_scores_combined_doc_number_signer(self):
        """A paragraph with doc number + 签发人 should score for combined type."""
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        pf = ParagraphFeature(index=2, text="国办发〔2024〕5号  签发人：张三")
        scorer = ElementScorer()
        scorer.reset_context([pf])
        scorer.update_context("issuing_authority_mark", 0)
        candidates = scorer.score_round(pf, round_group="round1")
        combined = [c for c in candidates if c.element_type == "combined_doc_number_signer"]
        assert len(combined) > 0

    def test_scores_issuing_authority_signature(self):
        """A paragraph ending with authority suffix should score for signature."""
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        pf = ParagraphFeature(index=15, text="国务院办公厅")
        scorer = ElementScorer()
        scorer.reset_context([pf])
        scorer.update_context("attachment_header", 10)
        candidates = scorer.score_round(pf, round_group="round1")
        sig = [c for c in candidates if c.element_type == "issuing_authority_signature"]
        assert len(sig) > 0

    def test_scores_printing_date(self):
        """A paragraph matching printing date format should score."""
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        pf = ParagraphFeature(index=25, text="2024年1月15日印发")
        scorer = ElementScorer()
        scorer.reset_context([pf])
        candidates = scorer.score_round(pf, round_group="round1")
        pd_ = [c for c in candidates if c.element_type == "printing_date"]
        assert len(pd_) > 0

    def test_scores_combined_printing_authority_and_date_line(self):
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        pf = ParagraphFeature(index=0, text="国务院办公厅　2024年1月15日印发")
        scorer = ElementScorer()
        scorer.reset_context([pf])

        candidates = scorer.score_round(pf, round_group="round1")

        assert any(candidate.element_type == "printing_date" for candidate in candidates)

    def test_scores_printing_line_with_word_spacing_inside_date(self):
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        feature = ParagraphFeature(index=0, text="国务院办公厅　2024 年 1 月 15 日 印发")

        candidates = ElementScorer().score_round(feature, round_group="round1")

        assert any(candidate.element_type == "printing_date" for candidate in candidates)

    def test_scores_disclosure(self):
        """A paragraph with disclosure label should score for disclosure."""
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        pf = ParagraphFeature(index=22, text="公开方式：主动公开")
        scorer = ElementScorer()
        scorer.reset_context([pf])
        scorer.update_context("issue_date", 21)
        candidates = scorer.score_round(pf, round_group="round1")
        disc = [c for c in candidates if c.element_type == "disclosure"]
        assert len(disc) > 0

    def test_scores_copy_to(self):
        """A paragraph with 抄送 label should score for copy_to."""
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        pf = ParagraphFeature(index=23, text="抄送：各省人民政府")
        scorer = ElementScorer()
        scorer.reset_context([pf])
        scorer.update_context("issue_date", 21)
        candidates = scorer.score_round(pf, round_group="round1")
        ct = [c for c in candidates if c.element_type == "copy_to"]
        assert len(ct) > 0

    def test_scores_combined_id(self):
        """A paragraph with 份号 + 发文字号 should score for combined_id."""
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        pf = ParagraphFeature(index=0, text="5国办发〔2024〕5号")
        scorer = ElementScorer()
        scorer.reset_context([pf])
        candidates = scorer.score_round(pf, round_group="round1")
        cid = [c for c in candidates if c.element_type == "combined_id"]
        assert len(cid) > 0

    def test_table_copy_to_organization_is_not_misclassified_as_signature(self):
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        pf = ParagraphFeature(
            index=8,
            text="省人民政府办公厅",
            source="table",
            table_cell_context="body",
        )
        scorer = ElementScorer()
        scorer.reset_context([pf])
        scorer.update_context("issue_date", 7)

        candidates = scorer.score_round(pf, round_group="round1")

        assert all(candidate.element_type != "issuing_authority_signature" for candidate in candidates)

    def test_printing_authority_directly_precedes_printing_date(self):
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        features = [
            ParagraphFeature(index=0, text="测试印发机关", source="table", table_cell_context="body"),
            ParagraphFeature(index=1, text="2025年7月5日印发", source="table", table_cell_context="body"),
        ]
        scorer = ElementScorer()
        scorer.reset_context(features)
        scorer.update_context("printing_date", 1)

        candidates = scorer.score_round(features[0], round_group="round2")

        assert any(candidate.element_type == "printing_authority" for candidate in candidates)

    def test_confidence_levels(self):
        """Test that confidence levels are correctly assigned."""
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        scorer = ElementScorer()
        # High confidence: title with strong signals
        pf = ParagraphFeature(
            index=0,
            text="关于进一步加强安全生产工作的通知",
            font_name="小标宋",
            font_size_pt=22.0,
        )
        scorer.reset_context([pf])
        candidates = scorer.score_round(pf, round_group="round1")
        title = [c for c in candidates if c.element_type == "title"]
        assert len(title) > 0
        assert title[0].confidence == "high"

    def test_does_not_score_title_for_plain_text(self):
        """Plain text without title signals should not score for title."""
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        pf = ParagraphFeature(
            index=10,
            text="普通文本段落",
            font_name="仿宋",
            font_size_pt=15.0,
        )
        scorer = ElementScorer()
        scorer.reset_context([pf])
        candidates = scorer.score_round(pf, round_group="round1")
        title = [c for c in candidates if c.element_type == "title"]
        assert len(title) == 0

    def test_scores_issuing_authority_mark(self):
        """issuing_authority_mark should score in round2 for early authority suffix."""
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        pf = ParagraphFeature(
            index=0,
            text="国务院办公厅",
            font_name="小标宋",
            font_size_pt=22.0,
        )
        scorer = ElementScorer()
        scorer.reset_context([pf])
        candidates = scorer.score_round(pf, round_group="round2")
        iam = [c for c in candidates if c.element_type == "issuing_authority_mark"]
        assert len(iam) > 0

    def test_scores_notes(self):
        """A bracketed paragraph should score for notes (附注)."""
        from docwen_plugin_optimizer_gongwen.models import ParagraphFeature
        from docwen_plugin_optimizer_gongwen.recognition.scorer import ElementScorer

        pf = ParagraphFeature(index=22, text="（此件公开发布）")
        scorer = ElementScorer()
        scorer.reset_context([pf])
        scorer.update_context("issue_date", 20)
        candidates = scorer.score_round(pf, round_group="round1")
        notes = [c for c in candidates if c.element_type == "notes"]
        assert len(notes) > 0
