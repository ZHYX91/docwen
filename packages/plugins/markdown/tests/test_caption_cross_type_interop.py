"""Cross-type caption ownership and native-table round-trip contracts."""

from __future__ import annotations

import base64
import hashlib
import json
from dataclasses import asdict
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from tests.support.config import FakeConfigView
from tests.support.execution import FakeExecutionContext
from tests.support.logging import FakePluginLogger
from tests.support.progress import FakeProgressSink
from tests.support.workspace import FakeWorkspaceHandle

from docwen_core.cancellation import CancellationToken
from docwen_core.docx_resolved_numbering_recovery import ResolvedNumberingV4Recovery
from docwen_core.docx_semantics_v3 import (
    CaptionStyleBindingV3,
    DocxSemanticsV3Recovery,
    DocxSemanticsV3Session,
)
from docwen_core.models.file_ref import FileRef
from docwen_core.models.request import ConversionRequest, OutputPolicy
from docwen_core.models.resolved_numbering import (
    NUMBERING_EXPORT_PLAN_MEDIA_TYPE,
    RESOLVED_DOCUMENT_MEDIA_TYPE,
    CaptionMaterialization,
    NumberingTarget,
    ResolvedDocument,
    ResolvedDocumentTarget,
    ResolvedEmbeddedResource,
    ResolvedNumberingPlan,
    ResolvedReference,
    ResolvedResourceOccurrence,
    canonicalize_numbering_plan,
)
from docwen_plugin_document.to_markdown.converter import DocxToMarkdownConverter
from docwen_plugin_markdown.document_semantics_v3 import analyze_markdown_semantics_v3
from docwen_plugin_markdown.mistune_extensions import parse_markdown_text
from docwen_plugin_markdown.renderer import MdToDocxRenderer
from docwen_plugin_markdown.runtime_semantics_v3 import (
    apply_runtime_semantics_v3,
    prepare_runtime_semantics_v3,
)
from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter
from docwen_plugin_markdown.to_docx.managed_styles import complete_managed_styles
from docwen_runtime.config.document_styles import build_document_style_catalog

from .conftest import PROJECT_ROOT

pytestmark = pytest.mark.contract

_CAPTIONS = {
    "figure": "Figure: Composite",
    "table": "Table: Composite",
    "equation": "Equation: Composite",
    "code_block": "Code: Composite",
}
_OBJECTS = {
    "figure": "![image](image.png)",
    "table": "| A |\n|---|\n| 1 |",
    "equation": "$$x=1$$",
    "code_block": "```text\nx\n```",
}


@pytest.mark.parametrize("caption_kind", tuple(_CAPTIONS))
@pytest.mark.parametrize("object_kind", tuple(_OBJECTS))
def test_all_caption_kinds_bind_all_captionable_object_kinds(
    caption_kind: str,
    object_kind: str,
) -> None:
    source = f"{_CAPTIONS[caption_kind]}\n\n{_OBJECTS[object_kind]}\n"

    analysis = analyze_markdown_semantics_v3(source, input_id="matrix.md")

    assert not analysis.has_errors
    [target] = analysis.projection["targets"]
    assert target["kind"] == caption_kind
    object_range = target["object_range"]
    assert source[object_range["start"] : object_range["end"]].strip() == _OBJECTS[object_kind]

    plan = prepare_runtime_semantics_v3(source, input_id="matrix.md")
    [owner] = apply_runtime_semantics_v3(parse_markdown_text(plan.shielded_source), plan)
    assert owner["_docwen_v3_caption_target"]["kind"] == caption_kind
    assert {
        "figure": "paragraph",
        "table": "table",
        "equation": "block_math",
        "code_block": "block_code",
    }[object_kind] == owner["type"]


@pytest.mark.parametrize("caption_first", [True, False])
@pytest.mark.parametrize("blank_lines", [0, 1])
@pytest.mark.parametrize("caption_kind", tuple(_CAPTIONS))
@pytest.mark.parametrize("object_kind", tuple(_OBJECTS))
def test_caption_binding_accepts_both_source_orders_with_zero_or_one_blank_line(
    caption_first: bool,
    blank_lines: int,
    caption_kind: str,
    object_kind: str,
) -> None:
    caption = _CAPTIONS[caption_kind]
    object_source = _OBJECTS[object_kind]
    separator = "\n" * (blank_lines + 1)
    source = f"{caption}{separator}{object_source}\n" if caption_first else f"{object_source}{separator}{caption}\n"

    analysis = analyze_markdown_semantics_v3(source, input_id="spacing.md")

    assert not analysis.has_errors
    assert [(item["kind"], item["title"]) for item in analysis.projection["targets"]] == [(caption_kind, "Composite")]
    plan = prepare_runtime_semantics_v3(source, input_id="spacing.md")
    owners = [
        node
        for node in apply_runtime_semantics_v3(parse_markdown_text(plan.shielded_source), plan)
        if "_docwen_v3_caption_target" in node
    ]
    assert len(owners) == 1
    assert owners[0]["_docwen_v3_caption_target"]["kind"] == caption_kind
    assert (
        owners[0]["type"]
        == {
            "figure": "paragraph",
            "table": "table",
            "equation": "block_math",
            "code_block": "block_code",
        }[object_kind]
    )


@pytest.mark.parametrize("caption_first", [True, False])
def test_two_blank_lines_break_caption_ownership(caption_first: bool) -> None:
    caption = "Figure: Composite"
    table = _OBJECTS["table"]
    source = f"{caption}\n\n\n{table}\n" if caption_first else f"{table}\n\n\n{caption}\n"

    analysis = analyze_markdown_semantics_v3(source, input_id="spacing.md")

    assert analysis.has_errors
    assert analysis.projection["targets"] == []
    assert [item["code"] for item in analysis.diagnostics] == ["docwen.markdown.caption.object_mismatch"]


@pytest.mark.parametrize(
    "source",
    [
        "![left](left.png)\n\nFigure: Ambiguous\n\n| right |\n|---|\n| 1 |\n",
        "Figure: First\n\n| shared |\n|---|\n| 1 |\n\nTable: Second\n",
    ],
)
def test_ambiguous_caption_object_graph_fails_closed(source: str) -> None:
    analysis = analyze_markdown_semantics_v3(source, input_id="ambiguous.md")

    assert analysis.has_errors
    assert analysis.projection["targets"] == []
    assert all(item["code"] == "docwen.markdown.caption.object_mismatch" for item in analysis.diagnostics)


def test_chain_does_not_use_global_matching_to_resolve_a_locally_ambiguous_caption() -> None:
    source = "Figure: First\n\n| first |\n|---|\n| 1 |\n\nTable: Second\n\n![second](second.png)\n"

    analysis = analyze_markdown_semantics_v3(source, input_id="ambiguous-chain.md")

    assert analysis.has_errors
    assert [(item["kind"], item["title"]) for item in analysis.projection["targets"]] == [("figure", "First")]
    assert [item["code"] for item in analysis.diagnostics] == ["docwen.markdown.caption.object_mismatch"]
    plan = prepare_runtime_semantics_v3(source, input_id="ambiguous-chain.md")
    assert plan.analysis.has_errors
    with pytest.raises(ValueError, match="invalid v3 analysis"):
        apply_runtime_semantics_v3(parse_markdown_text(plan.shielded_source), plan)


def test_next_line_caption_id_and_post_block_object_id_keep_distinct_owners() -> None:
    source = "| A |\n|---|\n| 1 |\n^object-id\nFigure: Composite\n^caption-id\n"

    analysis = analyze_markdown_semantics_v3(source, input_id="ids.md")
    assert not analysis.has_errors
    assert analysis.projection["targets"][0]["id"] == "caption-id"
    assert analysis.projection["anchors"][0]["id"] == "object-id"

    plan = prepare_runtime_semantics_v3(source, input_id="ids.md")
    [owner] = apply_runtime_semantics_v3(parse_markdown_text(plan.shielded_source), plan)
    assert owner["type"] == "table"
    assert owner["_docwen_v3_caption_target"]["id"] == "caption-id"
    assert owner["_docwen_v3_ordinary_anchor"]["id"] == "object-id"


@pytest.mark.parametrize("declaration", ["Code: ^snippet", "Code:\n^snippet"])
def test_empty_code_caption_is_valid_with_explicit_id(declaration: str) -> None:
    source = f"{declaration}\n```text\nx\n```\n"

    analysis = analyze_markdown_semantics_v3(source, input_id="code.md")

    assert not analysis.has_errors
    [target] = analysis.projection["targets"]
    assert (target["kind"], target["title"], target["id"]) == ("code_block", "", "snippet")


def test_figure_captioned_multi_image_table_round_trips_as_native_table(tmp_path: Path) -> None:
    png = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAIAAAACCAIAAAD91JpzAAAAEklEQVR4nGNUSFjAwMDAxAAGAA0qASTlOPBgAAAAAElFTkSuQmCC"
    )
    (tmp_path / "a.png").write_bytes(png)
    (tmp_path / "b.png").write_bytes(png)
    source_path = tmp_path / "composite.md"
    source = "Figure: Composite ^composite\n\n| A | B |\n|---|---|\n| ![a](a.png) | ![b](b.png) |\n"
    source_path.write_text(source, encoding="utf-8")
    output = _render_v3_docx(source_path, source, tmp_path / "composite.docx")

    reopened = Document(str(output))
    recovery = DocxSemanticsV3Recovery.load(output, reopened)
    body = recovery.logical_body_elements(reopened)
    table = next(item for item in body if item.tag == qn("w:tbl"))
    caption = recovery.caption_for_object(table)
    assert caption is not None
    assert (caption.kind, caption.source_id, caption.title, caption.cached_number) == (
        "figure",
        "composite",
        "Composite",
        "1",
    )
    assert "SEQ Figure" in caption.caption_element.xml
    assert len(list(table.iter(qn("w:drawing")))) == 2
    assert [item.tag for item in body[:2]] == [qn("w:tbl"), qn("w:p")]

    markdown = _convert_docx_to_markdown(output, tmp_path)
    declaration = markdown.index("Figure: Composite ^composite")
    table_start = markdown.index("| A | B |")
    assert declaration < table_start
    assert markdown.count("![[docx-image_") == 2
    assert "| --- | --- |" in markdown


def test_exact_two_figure_captioned_multi_image_table_round_trips_with_short_target_range(
    tmp_path: Path,
) -> None:
    neutral, plan = _write_exact_two_composite_pair(tmp_path)
    refs = (
        FileRef(
            path=str(neutral),
            format="markdown",
            category="document",
            size_bytes=neutral.stat().st_size,
            input_kind="document",
            input_role="neutral_document",
            logical_path="document.json",
            media_type=RESOLVED_DOCUMENT_MEDIA_TYPE,
        ),
        FileRef(
            path=str(plan),
            format="json",
            category="resource",
            size_bytes=plan.stat().st_size,
            input_kind="resource",
            input_role="numbering_export_plan",
            logical_path="numbering-plan.json",
            media_type=NUMBERING_EXPORT_PLAN_MEDIA_TYPE,
        ),
    )
    context = _md_to_docx_context(tmp_path / "exact-two", refs)
    result = MdToDocxConverter().convert(context)
    assert result.success, result.error
    output = Path(result.artifacts[0].staging_path)

    reopened = Document(str(output))
    recovery = ResolvedNumberingV4Recovery.load_if_present(output, reopened)
    assert recovery is not None
    body = recovery.logical_body_elements(reopened)
    table = next(item for item in body if item.tag == qn("w:tbl"))
    caption = recovery.caption_for_object(table)
    assert caption is not None
    assert (caption.kind, caption.source_id, caption.title, caption.cached_number) == (
        "figure",
        "composite",
        "Composite",
        "1",
    )
    assert "SEQ Figure" in caption.caption_element.xml
    instructions = [item.text or "" for item in reopened.element.iter(qn("w:instrText"))]
    assert any(" REF " in item for item in instructions)
    assert not any("SEQ Table" in item for item in instructions)
    assert len(list(table.iter(qn("w:drawing")))) == 2
    assert [item.tag for item in body[:2]] == [qn("w:tbl"), qn("w:p")]

    markdown = _convert_docx_to_markdown(output, tmp_path)
    authored = json.loads(neutral.read_text(encoding="utf-8"))["document"]["authored_markdown"]
    assert markdown == authored
    assert markdown.index("Figure: Composite\n^composite") < markdown.index("| A | B |")
    assert markdown.count("![[") == 2
    assert "|---|---|" in markdown
    assert "@[[#^composite]]" in markdown


def _write_exact_two_composite_pair(tmp_path: Path) -> tuple[Path, Path]:
    png = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAIAAAACCAIAAAD91JpzAAAAEklEQVR4nGNUSFjAwMDAxAAGAA0qASTlOPBgAAAAAElFTkSuQmCC"
    )
    source = (
        "Figure: Composite\n^composite\n| A | B |\n|---|---|\n| ![[a.png]] | ![[b.png]] |\n\nSee @[[#^composite]].\n"
    )
    declaration = "Figure: Composite"
    target = ResolvedDocumentTarget(
        source_start=0,
        source_end=len(declaration),
        source_slice_sha256=hashlib.sha256(declaration.encode()).hexdigest(),
        kind="figure",
        target_id="composite",
        heading_level=None,
        authored_text="Composite",
    )
    resources: list[ResolvedEmbeddedResource] = []
    occurrences: list[ResolvedResourceOccurrence] = []
    for index, name in enumerate(("a.png", "b.png"), start=1):
        token = f"![[{name}]]"
        start = source.index(token)
        resource_id = f"image-{index}"
        resources.append(
            ResolvedEmbeddedResource(
                resource_id=resource_id,
                role="linked_resource",
                media_type="image/png",
                size_bytes=len(png),
                sha256=hashlib.sha256(png).hexdigest(),
                content=png,
            )
        )
        occurrences.append(
            ResolvedResourceOccurrence(
                source_start=start,
                source_end=start + len(token),
                source_slice_sha256=hashlib.sha256(token.encode()).hexdigest(),
                authored_token=token,
                authored_locator=name,
                resource_id=resource_id,
            )
        )
    materialization = CaptionMaterialization(
        type="simple_seq",
        counter="Figure",
        number_format="arabic_half",
        sequence_action="reset_to_start",
        start_value=1,
        chapter_heading_level=None,
        chapter_heading_style=None,
        chapter_separator=None,
        restart_heading_level=None,
        restart_heading_style=None,
        chapter_cached_number=None,
        sequence_cached_number="1",
        localized_label="Figure",
        label_separator=" ",
    )
    numbering_plan = ResolvedNumberingPlan(
        heading_definitions=(),
        heading_instances=(),
        targets=(
            NumberingTarget(
                source_start=target.source_start,
                source_end=target.source_end,
                kind="figure",
                enabled=True,
                target_id="composite",
                derived_number="1",
                materialization=materialization,
            ),
        ),
    )
    plan_body = json.loads(json.dumps(asdict(numbering_plan)))
    plan_sha256 = hashlib.sha256(canonicalize_numbering_plan(plan_body)).hexdigest()
    reference_token = "@[[#^composite]]"
    reference_start = source.index(reference_token)
    reference = ResolvedReference(
        source_start=reference_start,
        source_end=reference_start + len(reference_token),
        source_slice_sha256=hashlib.sha256(reference_token.encode()).hexdigest(),
        authored_token=reference_token,
        target_source_start=target.source_start,
        target_source_end=target.source_end,
        target_kind="figure",
        target_id="composite",
        cached_number="1",
        alias=None,
    )
    document = ResolvedDocument(source, (target,), (reference,), tuple(occurrences), (), tuple(resources))
    document_body = asdict(document)
    for resource in document_body["resources"]:
        content = resource.pop("content")
        resource["content_base64"] = base64.b64encode(content).decode("ascii")
    source_sha256 = hashlib.sha256(source.encode()).hexdigest()
    neutral_payload = {
        "$schema": "urn:docwen:schema:resolved-document:v1",
        "schema": "docwen.resolved_document.v1",
        "input_id": "caption-cross-type",
        "source_sha256": source_sha256,
        "plan_sha256": plan_sha256,
        "document": document_body,
    }
    plan_payload = {
        "$schema": "urn:docwen:schema:numbering-export-plan:v1",
        "schema": "docwen.numbering_export_plan.v1",
        "input_id": "caption-cross-type",
        "source_sha256": source_sha256,
        "plan_sha256": plan_sha256,
        "plan": plan_body,
    }
    neutral = tmp_path / "composite-neutral.json"
    plan = tmp_path / "composite-plan.json"
    neutral.write_text(json.dumps(neutral_payload, separators=(",", ":")), encoding="utf-8")
    plan.write_text(json.dumps(plan_payload, separators=(",", ":")), encoding="utf-8")
    return neutral, plan


def _md_to_docx_context(tmp_path: Path, refs: tuple[FileRef, FileRef]) -> FakeExecutionContext:
    staging = tmp_path / "staging"
    staging.mkdir(parents=True)
    styles = build_document_style_catalog(
        {"gui": {"language": {"locale": "en_US"}}},
        locales_dir=PROJECT_ROOT / "i18n" / "locales",
    )
    return FakeExecutionContext(
        request=ConversionRequest(
            request_id="caption-cross-type-exact-two",
            input_refs=list(refs),
            target_format="docx",
            options={"locale": "en_US", "heading_merge_mode": "never"},
            output_policy=OutputPolicy(),
        ),
        workspace=FakeWorkspaceHandle(refs[0].path, str(staging), refs),
        config=FakeConfigView(),
        progress=FakeProgressSink(),
        cancellation=CancellationToken().view(),
        logger=FakePluginLogger(),
        numbering_registry=None,
        document_style_catalog=styles,
    )


def _render_v3_docx(source_path: Path, source: str, output: Path) -> Path:
    plan = prepare_runtime_semantics_v3(source, input_id=source_path.name)
    ast = apply_runtime_semantics_v3(parse_markdown_text(plan.shielded_source), plan)
    catalog = build_document_style_catalog(
        {"gui": {"language": {"locale": "en_US"}}},
        locales_dir=PROJECT_ROOT / "i18n" / "locales",
    )
    document, managed = complete_managed_styles(Document(), catalog)
    semantic_keys = ("figure_caption", "table_caption", "equation_caption", "code_block_caption")
    bindings = tuple(
        CaptionStyleBindingV3(key, managed.style_id(key), managed.get(key).name or "") for key in semantic_keys
    )
    session = DocxSemanticsV3Session(
        document,
        source_sha256=hashlib.sha256(source.encode()).hexdigest(),
        caption_style_bindings=bindings,
    )
    MdToDocxRenderer(
        document,
        managed_styles=managed,
        semantic_v3_session=session,
        source_file_path=str(source_path),
    ).render(ast)
    session.finalize_document()
    document.save(str(output))
    session.write_package(output)
    session.prove_package(output)
    return output


def _convert_docx_to_markdown(source: Path, tmp_path: Path) -> str:
    staging = tmp_path / "markdown-output"
    staging.mkdir()
    reference = FileRef(path=str(source), format="docx", category="document")
    context = FakeExecutionContext(
        request=ConversionRequest(
            request_id="caption-cross-type",
            input_refs=[reference],
            target_format="md",
            options={},
            output_policy=OutputPolicy(),
        ),
        workspace=FakeWorkspaceHandle(str(source), str(staging)),
        config=FakeConfigView(),
        progress=FakeProgressSink(),
        cancellation=CancellationToken().view(),
        logger=FakePluginLogger(),
        numbering_registry=None,
    )
    result = DocxToMarkdownConverter().convert(context)
    assert result.success, result.error
    return Path(result.artifacts[0].staging_path).read_text(encoding="utf-8")
