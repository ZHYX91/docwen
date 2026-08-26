"""Permanent contracts for grouped empty-placeholder cleanup."""

from __future__ import annotations

from typing import Any

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docwen_plugin_markdown.field_processors.gongwen import (
    process_attachment_placeholder,
    process_yaml,
)
from docwen_plugin_markdown.template_filler import (
    apply_placeholder_rules,
    apply_special_placeholder_handlers,
    fill_template,
    fill_yaml_placeholders,
)

pytestmark = pytest.mark.contract
from docwen_plugin_markdown.template_utils import scan_placeholders


def _apply_rules(document, yaml_data: dict[str, Any], rules: dict[str, list[list[str]]]) -> None:
    placeholder_map = scan_placeholders(document)
    fill_yaml_placeholders(document, yaml_data, placeholder_map)
    apply_placeholder_rules(document, yaml_data, placeholder_map, [rules])


def test_empty_paragraph_group_requires_the_complete_group_in_one_paragraph() -> None:
    document = Document()
    document.add_paragraph("remove prefix {{ first }} / {{second}} suffix")
    document.add_paragraph("keep first={{first}} suffix")
    document.add_paragraph("keep second={{second}} suffix")

    _apply_rules(
        document,
        {"first": "", "second": None},
        {"delete_paragraph_if_empty": [["first", "second"]]},
    )

    assert [paragraph.text for paragraph in document.paragraphs] == [
        "keep first= suffix",
        "keep second= suffix",
    ]


def test_nonempty_paragraph_group_preserves_container_and_replaces_each_value() -> None:
    document = Document()
    document.add_paragraph("prefix {{first}} / {{second}} suffix")

    _apply_rules(
        document,
        {"first": "kept", "second": ""},
        {"delete_paragraph_if_empty": [["first", "second"]]},
    )

    assert [paragraph.text for paragraph in document.paragraphs] == ["prefix kept /  suffix"]


def test_empty_row_group_requires_the_complete_group_in_one_row() -> None:
    document = Document()
    table = document.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "remove {{first}}"
    table.rows[0].cells[1].text = "and {{ second }}"
    table.rows[1].cells[0].text = "keep {{first}}"
    table.rows[1].cells[1].text = "first-only"
    table.rows[2].cells[0].text = "second-only"
    table.rows[2].cells[1].text = "keep {{second}}"

    _apply_rules(
        document,
        {"first": "", "second": ""},
        {"delete_row_if_empty": [["first", "second"]]},
    )

    assert [[cell.text for cell in row.cells] for row in table.rows] == [
        ["keep ", "first-only"],
        ["second-only", "keep "],
    ]


def test_nonempty_row_group_preserves_row_and_replaces_each_value() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "{{first}}"
    table.rows[0].cells[1].text = "{{second}}"

    _apply_rules(
        document,
        {"first": "kept", "second": ""},
        {"delete_row_if_empty": [["first", "second"]]},
    )

    assert [[cell.text for cell in row.cells] for row in table.rows] == [["kept", ""]]


def test_empty_cell_group_requires_the_complete_group_in_one_cell() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "clear label {{first}} / {{second}}"
    table.rows[0].cells[1].text = "keep label {{first}}"

    _apply_rules(
        document,
        {"first": "", "second": ""},
        {"delete_cell_if_empty": [["first", "second"]]},
    )

    assert len(table.rows) == 1
    assert len(table.rows[0].cells) == 2
    assert [cell.text for cell in table.rows[0].cells] == ["", "keep label "]


def test_nonempty_cell_group_preserves_cell_and_replaces_each_value() -> None:
    document = Document()
    cell = document.add_table(rows=1, cols=1).cell(0, 0)
    cell.text = "prefix {{first}} / {{second}} suffix"

    _apply_rules(
        document,
        {"first": "kept", "second": ""},
        {"delete_cell_if_empty": [["first", "second"]]},
    )

    assert cell.text == "prefix kept /  suffix"


def test_unconfigured_empty_placeholder_preserves_surrounding_paragraph() -> None:
    document = Document()
    document.add_paragraph("prefix {{missing}} suffix")

    _apply_rules(document, {}, {})

    assert [paragraph.text for paragraph in document.paragraphs] == ["prefix  suffix"]


def test_rules_only_consume_the_pre_render_template_snapshot() -> None:
    document = Document()
    body_placeholder = document.add_paragraph("{{body}}")
    document.add_paragraph("template {{first}} / {{second}}")
    placeholder_map = scan_placeholders(document)
    rendered = document.add_paragraph("literal {{first}} / {{second}}")

    fill_template(
        document,
        {"first": "", "second": ""},
        [rendered],
        body_placeholder,
        placeholder_map,
        placeholder_rules=[{"delete_paragraph_if_empty": [["first", "second"]]}],
    )

    assert [paragraph.text for paragraph in document.paragraphs] == ["literal {{first}} / {{second}}"]


def test_special_handler_does_not_consume_a_rendered_literal() -> None:
    document = Document()
    body_placeholder = document.add_paragraph("{{body}}")
    placeholder_map = scan_placeholders(document)
    rendered = document.add_paragraph("literal {{附件说明}}")

    fill_template(
        document,
        {"附件说明": ["附件：材料"]},
        [rendered],
        body_placeholder,
        placeholder_map,
        special_placeholder_handlers={"附件说明": process_attachment_placeholder},
    )

    assert [paragraph.text for paragraph in document.paragraphs] == ["literal {{附件说明}}"]


def test_deleting_the_last_matching_row_removes_the_empty_table() -> None:
    document = Document()
    document.add_table(rows=1, cols=1).cell(0, 0).text = "{{first}}"

    _apply_rules(document, {"first": ""}, {"delete_row_if_empty": [["first"]]})

    assert document.tables == []


def test_placeholder_fragments_in_different_cells_are_not_stitched() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "{{fir"
    table.cell(0, 1).text = "st}}"

    _apply_rules(document, {"first": ""}, {"delete_row_if_empty": [["first"]]})

    assert len(document.tables) == 1
    assert [cell.text for cell in table.rows[0].cells] == ["{{fir", "st}}"]


def test_nested_table_row_rules_use_the_nearest_original_row() -> None:
    document = Document()
    outer = document.add_table(rows=1, cols=1)
    nested = outer.cell(0, 0).add_table(rows=1, cols=2)
    nested.cell(0, 0).text = "{{first}}"
    nested.cell(0, 1).text = "{{second}}"

    _apply_rules(
        document,
        {"first": "", "second": ""},
        {"delete_row_if_empty": [["first", "second"]]},
    )

    assert len(document.tables) == 1
    assert outer.cell(0, 0).tables == []


def test_merged_cell_is_deduplicated_and_cleared_once() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=2)
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.text = "label {{first}} / {{second}}"

    _apply_rules(
        document,
        {"first": "", "second": ""},
        {"delete_cell_if_empty": [["first", "second"]]},
    )

    assert len(document.tables) == 1
    assert table.cell(0, 0)._tc is table.cell(0, 1)._tc
    assert table.cell(0, 0).text == ""


def test_row_rules_have_priority_over_cell_rules_across_processors() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "{{first}} / {{second}}"
    placeholder_map = scan_placeholders(document)

    apply_placeholder_rules(
        document,
        {"first": "", "second": ""},
        placeholder_map,
        [
            {"delete_cell_if_empty": [["first", "second"]]},
            {"delete_row_if_empty": [["first", "second"]]},
        ],
    )

    assert document.tables == []


def test_recursively_empty_mappings_and_sequences_trigger_group_cleanup() -> None:
    document = Document()
    document.add_paragraph("{{first}} / {{second}}")

    _apply_rules(
        document,
        {"first": {"nested": ""}, "second": [None, {"nested": "  "}]},
        {"delete_paragraph_if_empty": [["first", "second"]]},
    )

    assert document.paragraphs == []


def test_attachment_handler_processes_all_original_whitespace_variants() -> None:
    document = Document()
    document.add_paragraph("{{附件说明}}")
    document.add_paragraph("{{ 附件说明 }}")
    placeholder_map = scan_placeholders(document)

    handled = apply_special_placeholder_handlers(
        document,
        {"附件说明": ["附件：材料"]},
        {"附件说明": process_attachment_placeholder},
        placeholder_map=placeholder_map,
    )

    assert handled == {"附件说明"}
    assert [paragraph.text for paragraph in document.paragraphs] == ["附件：材料", "附件：材料"]


def test_special_body_handler_does_not_suppress_the_same_placeholder_in_a_table() -> None:
    document = Document()
    body_placeholder = document.add_paragraph("{{body}}")
    document.add_paragraph("{{附件说明}}")
    table_cell = document.add_table(rows=1, cols=1).cell(0, 0)
    table_cell.text = "table {{附件说明}}"
    placeholder_map = scan_placeholders(document)

    fill_template(
        document,
        {"附件说明": ["附件：材料"]},
        [],
        body_placeholder,
        placeholder_map,
        special_placeholder_handlers={"附件说明": process_attachment_placeholder},
    )

    assert "附件：材料" in [paragraph.text for paragraph in document.paragraphs]
    assert table_cell.text == "table 附件：材料"


def test_attachment_processor_treats_recursively_empty_values_as_empty() -> None:
    yaml_data: dict[str, Any] = {"附件说明": ["", None, {"nested": "  "}]}

    process_yaml(yaml_data)

    assert yaml_data["附件说明"] == []


def test_delete_table_rule_requires_the_complete_group_in_one_table() -> None:
    document = Document()
    remove_table = document.add_table(rows=2, cols=1)
    remove_table.cell(0, 0).text = "{{first}}"
    remove_table.cell(1, 0).text = "{{second}}"
    keep_first = document.add_table(rows=1, cols=1)
    keep_first.cell(0, 0).text = "keep {{first}}"
    keep_second = document.add_table(rows=1, cols=1)
    keep_second.cell(0, 0).text = "keep {{second}}"

    _apply_rules(
        document,
        {"first": "", "second": ""},
        {"delete_table_if_empty": [["first", "second"]]},
    )

    assert len(document.tables) == 2
    assert document.tables[0]._tbl is keep_first._tbl
    assert document.tables[1]._tbl is keep_second._tbl
    assert keep_first.cell(0, 0).text == "keep "
    assert keep_second.cell(0, 0).text == "keep "


def test_placeholder_inside_hyperlink_is_replaced_without_flattening_ooxml() -> None:
    document = Document()
    paragraph = document.add_paragraph("prefix ")
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), "rId-test")
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    bold = OxmlElement("w:b")
    run_properties.append(bold)
    text = OxmlElement("w:t")
    text.text = "{{missing}}"
    run.append(run_properties)
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    paragraph.add_run(" suffix")

    apply_placeholder_rules(document, {}, {"missing": [paragraph]})

    assert paragraph._p.find(qn("w:hyperlink")) is hyperlink
    assert hyperlink.find(qn("w:r")) is run
    assert run.find(qn("w:rPr")) is run_properties
    assert "".join(item.text or "" for item in paragraph._p.iter(qn("w:t"))) == "prefix  suffix"
