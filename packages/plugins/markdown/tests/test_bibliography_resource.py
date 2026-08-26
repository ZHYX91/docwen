"""Typed bibliography resource through the real Markdown-to-DOCX converter."""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from docwen_core.models.file_ref import FileRef
from docwen_core.semantic_bibliography import SEMANTIC_BIBLIOGRAPHY_MEDIA_TYPE
from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter
from docwen_runtime.toml_io import read_toml_file

from .conftest import PROJECT_ROOT, make_context

pytestmark = pytest.mark.contract


def _bibliography_payload(*, entries: list[object] | None = None) -> bytes:
    return json.dumps(
        {
            "schema": "docwen.semantic_bibliography.v1",
            "entries": entries
            if entries is not None
            else [
                {
                    "item_id": "smith2025",
                    "runs": [
                        {"text": "Smith, A. ", "bold": True},
                        {
                            "text": "Neutral documents",
                            "italic": True,
                            "href": "https://example.org/neutral-documents",
                        },
                    ],
                }
            ],
        }
    ).encode()


def _attach_bibliography(context, workspace, source: Path, resource: Path) -> None:
    source_ref = FileRef(
        path=str(source),
        format="markdown",
        category="markdown",
        input_kind="document",
        input_role="source",
        logical_path="source.md",
        media_type="text/markdown",
    )
    bibliography_ref = FileRef(
        path=str(resource),
        format="",
        category="other",
        input_kind="resource",
        input_role="bibliography",
        logical_path="bibliography.json",
        media_type=SEMANTIC_BIBLIOGRAPHY_MEDIA_TYPE,
    )
    context.request.input_refs = [source_ref, bibliography_ref]
    workspace._input_refs = (source_ref, bibliography_ref)


def test_nonempty_bibliography_synthesizes_after_unique_body_marker_and_renders_rich_runs(tmp_path: Path) -> None:
    source = tmp_path / "source.md"
    source.write_text("# Source\n\nBody.\n", encoding="utf-8")
    resource = tmp_path / "bibliography.json"
    resource.write_bytes(_bibliography_payload())
    context, workspace = make_context(str(source), options={})
    _attach_bibliography(context, workspace, source, resource)

    result = MdToDocxConverter().convert(context)

    assert result.success
    output = Document(result.artifacts[0].staging_path)
    texts = [paragraph.text for paragraph in output.paragraphs]
    assert "{{ bibliography }}" not in texts
    assert texts.index("Smith, A. Neutral documents") > texts.index("Body.")
    bibliography_paragraph = next(paragraph for paragraph in output.paragraphs if "Neutral documents" in paragraph.text)
    assert bibliography_paragraph.style is not None
    assert bibliography_paragraph.style.style_id == "Bibliography"
    hyperlink = next(bibliography_paragraph._p.iter(qn("w:hyperlink")))
    assert [style.get(qn("w:val")) for style in hyperlink.iter(qn("w:rStyle"))] == ["Hyperlink"]
    relationship_id = hyperlink.get(qn("r:id"))
    assert relationship_id is not None
    assert output.part.rels[relationship_id].target_ref == "https://example.org/neutral-documents"


@pytest.mark.parametrize(
    "locale",
    ("zh_CN", "en_US", "de_DE", "es_ES", "fr_FR", "ja_JP", "ko_KR", "pt_BR", "ru_RU", "vi_VN", "zh_TW"),
)
def test_every_shipped_template_synthesizes_bibliography_after_localized_body_marker(
    tmp_path: Path,
    locale: str,
) -> None:
    locale_table = read_toml_file(PROJECT_ROOT / "i18n" / "locales" / f"{locale}.toml")
    template_name = locale_table["meta"]["template_name"]
    template = PROJECT_ROOT / "templates" / f"{template_name}.docx"
    template_hash = template.read_bytes()
    source = tmp_path / f"source-{locale}.md"
    source.write_text("Body.\n", encoding="utf-8")
    resource = tmp_path / f"bibliography-{locale}.json"
    resource.write_bytes(_bibliography_payload())
    context, workspace = make_context(
        str(source),
        options={"template_name": str(template), "locale": locale},
    )
    _attach_bibliography(context, workspace, source, resource)

    result = MdToDocxConverter().convert(context)

    assert result.success, (locale, result.error, result.diagnostics)
    output = Document(result.artifacts[0].staging_path)
    texts = [paragraph.text for paragraph in output.paragraphs]
    assert texts.index("Smith, A. Neutral documents") > texts.index("Body.")
    assert template.read_bytes() == template_hash


def test_explicit_bibliography_anchor_is_reserved_from_yaml_and_preserves_heading(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    document = Document()
    document.add_paragraph("{{ body }}")
    document.add_paragraph("References")
    anchor = document.add_paragraph("{{ bibliography }}")
    anchor.paragraph_format.keep_with_next = True
    document.save(str(template))
    source = tmp_path / "source.md"
    source.write_text("---\nbibliography: yaml-must-not-own-this\n---\nBody.\n", encoding="utf-8")
    resource = tmp_path / "bibliography.json"
    resource.write_bytes(_bibliography_payload())
    context, workspace = make_context(str(source), options={"template_name": str(template)})
    _attach_bibliography(context, workspace, source, resource)

    result = MdToDocxConverter().convert(context)

    assert result.success
    output = Document(result.artifacts[0].staging_path)
    assert "yaml-must-not-own-this" not in "\n".join(paragraph.text for paragraph in output.paragraphs)
    assert output.paragraphs[1].text == "References"
    bibliography_paragraph = next(paragraph for paragraph in output.paragraphs if "Neutral documents" in paragraph.text)
    assert bibliography_paragraph.paragraph_format.keep_with_next


@pytest.mark.parametrize("shape", ["duplicate", "inline", "missing_body"])
def test_bibliography_placeholder_failures_are_stable_and_register_no_artifact(tmp_path: Path, shape: str) -> None:
    template = tmp_path / "template.docx"
    document = Document()
    if shape != "missing_body":
        document.add_paragraph("{{ body }}")
    if shape == "duplicate":
        document.add_paragraph("{{ bibliography }}")
        document.add_paragraph("{{ bibliography }}")
    elif shape == "inline":
        document.add_paragraph("prefix {{ bibliography }} suffix")
    document.save(str(template))
    source = tmp_path / "source.md"
    source.write_text("Body.\n", encoding="utf-8")
    resource = tmp_path / "bibliography.json"
    resource.write_bytes(_bibliography_payload())
    context, workspace = make_context(str(source), options={"template_name": str(template)})
    _attach_bibliography(context, workspace, source, resource)

    result = MdToDocxConverter().convert(context)

    assert not result.success
    assert result.error is not None
    assert result.error.diagnostic_code == "MD2DOCX-BIBLIOGRAPHY-PLACEHOLDER-INVALID"
    assert result.artifacts == []
    assert workspace.registered_artifacts == []


def test_invalid_bibliography_resource_fails_before_template_mutation(tmp_path: Path) -> None:
    source = tmp_path / "source.md"
    source.write_text("Body.\n", encoding="utf-8")
    resource = tmp_path / "bibliography.json"
    resource.write_bytes(b'{"schema":"wrong","entries":[]}')
    context, workspace = make_context(str(source))
    _attach_bibliography(context, workspace, source, resource)

    result = MdToDocxConverter().convert(context)

    assert not result.success
    assert result.error is not None
    assert result.error.diagnostic_code == "MD2DOCX-BIBLIOGRAPHY-RESOURCE-INVALID"
    assert workspace.registered_artifacts == []


def test_uppercase_yaml_placeholder_remains_yaml_owned_without_typed_resource(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    document = Document()
    document.add_paragraph("{{ body }}")
    document.add_paragraph("{{ Bibliography }}")
    document.save(str(template))
    source = tmp_path / "source.md"
    source.write_text("---\nBibliography: YAML bibliography\n---\nBody.\n", encoding="utf-8")
    context, _workspace = make_context(str(source), options={"template_name": str(template)})

    result = MdToDocxConverter().convert(context)

    assert result.success
    output = Document(result.artifacts[0].staging_path)
    assert [paragraph.text for paragraph in output.paragraphs if paragraph.text] == ["Body.", "YAML bibliography"]
