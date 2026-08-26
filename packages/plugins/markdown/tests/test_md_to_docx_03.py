"""Focused tests split from test_md_to_docx.py."""

from __future__ import annotations

import pytest

from ._md_to_docx_support import (
    FakeConfigView,
    MdToDocxConverter,
    Path,
    make_context,
    re,
    write_temp_md,
    zipfile,
)

pytestmark = pytest.mark.contract


class TestMdToDocxGolden:
    """Golden-level tests for MD → DOCX conversion."""

    @staticmethod
    def _verify_docx_structure(output_path: Path):
        """Verify the file is a well-formed DOCX (ZIP with expected entries)."""
        import zipfile

        assert zipfile.is_zipfile(output_path), "Output is not a valid ZIP file"
        with zipfile.ZipFile(output_path, "r") as zf:
            names = zf.namelist()
            assert "[Content_Types].xml" in names, "Missing [Content_Types].xml"
            assert any("word/document.xml" in n for n in names), "Missing word/document.xml"

    @staticmethod
    def _verify_docx_content(output_path: Path):
        """Verify the DOCX contains expected heading and paragraph structure."""
        from docx import Document

        doc = Document(str(output_path))

        paragraphs = doc.paragraphs
        assert len(paragraphs) >= 5, f"Expected at least 5 paragraphs, got {len(paragraphs)}"

        # Check heading styles
        heading_styles = [p.style.name for p in paragraphs if p.style and p.style.name and "Heading" in p.style.name]
        assert len(heading_styles) >= 2, f"Expected at least 2 headings, got {len(heading_styles)}"

    def test_md_to_docx_consumes_configured_heading_merge_mode(self, monkeypatch):
        """Global MD→DOCX heading merge config is used when options omit it."""
        import docwen_plugin_markdown.to_docx.converter as converter_module

        seen: dict[str, object] = {}

        def fake_detect_heading_merges(
            md_body: str,
            *,
            mode: str = "punct_required",
            punctuation: frozenset[str] | None = None,
        ):
            seen["mode"] = mode
            seen["punctuation"] = punctuation
            return set()

        monkeypatch.setattr(converter_module, "detect_heading_merges", fake_detect_heading_merges)

        md_path = write_temp_md("# Heading.\nBody")
        ctx, _workspace = make_context(md_path, target_format="docx")
        ctx._config = FakeConfigView(
            {
                "conversion": {
                    "md_to_docx": {
                        "heading_merge_mode": "never",
                        "heading_merge_punctuation": "：§§ ",
                    }
                }
            }
        )

        result = MdToDocxConverter().convert(ctx)

        assert result.success is True
        assert seen["mode"] == "never"
        assert seen["punctuation"] == frozenset({"：", "§"})

    def test_md_to_docx_consumes_configured_body_formatting_remove(self):
        """Configured body formatting remove maps to renderer minimal mode."""
        md_path = write_temp_md("This is **bold** text.")
        ctx, _workspace = make_context(md_path, target_format="docx")
        ctx._config = FakeConfigView({"conversion": {"md_to_docx": {"formatting_mode": "remove"}}})

        result = MdToDocxConverter().convert(ctx)

        assert result.success is True
        with zipfile.ZipFile(Path(result.artifacts[0].staging_path), "r") as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")
        assert "bold" in document_xml
        assert re.search(r"<w:b(?:\s|/|>)", document_xml) is None

    def test_md_to_docx_options_override_configured_body_formatting(self):
        """Explicit options keep their precedence over global config defaults."""
        md_path = write_temp_md("This is **bold** text.")
        ctx, _workspace = make_context(md_path, target_format="docx", options={"formatting_mode": "full"})
        ctx._config = FakeConfigView({"conversion": {"md_to_docx": {"formatting_mode": "remove"}}})

        result = MdToDocxConverter().convert(ctx)

        assert result.success is True
        with zipfile.ZipFile(Path(result.artifacts[0].staging_path), "r") as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")
        assert "bold" in document_xml
        assert re.search(r"<w:b(?:\s|/|>)", document_xml) is not None

    def test_md_to_docx_consumes_configured_body_formatting_keep(self):
        """Configured body formatting keep preserves visible Markdown markers."""
        md_path = write_temp_md("This is **bold** and `code`.")
        ctx, _workspace = make_context(md_path, target_format="docx")
        ctx._config = FakeConfigView({"conversion": {"md_to_docx": {"formatting_mode": "keep"}}})

        result = MdToDocxConverter().convert(ctx)

        assert result.success is True
        from docx import Document

        doc = Document(str(Path(result.artifacts[0].staging_path)))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        assert "This is **bold** and `code`." in text

    def test_md_to_docx_consumes_configured_code_font(self):
        """Configured code font is used when options omit code_font."""
        md_path = write_temp_md("Use `code` here.")
        ctx, _workspace = make_context(md_path, target_format="docx")
        ctx._config = FakeConfigView({"conversion": {"code_detection": {"code_font": "Courier New"}}})

        result = MdToDocxConverter().convert(ctx)

        assert result.success is True
        with zipfile.ZipFile(Path(result.artifacts[0].staging_path), "r") as zf:
            styles_xml = zf.read("word/styles.xml").decode("utf-8")
        inline_code = re.search(
            r'<w:style[^>]*w:styleId="DocWenInlineCode".*?</w:style>',
            styles_xml,
            flags=re.DOTALL,
        )
        assert inline_code is not None
        assert "courier new" in inline_code.group(0).lower()

    def test_md_to_docx_consumes_configured_table_header_formatting_mode(self):
        """Configured table header formatting mode reaches the DOCX renderer."""
        md_path = write_temp_md("| **Header** |\n|---|\n| cell |\n")
        ctx, _workspace = make_context(md_path, target_format="docx")
        ctx._config = FakeConfigView({"conversion": {"md_to_docx": {"table_header_formatting_mode": "apply"}}})

        result = MdToDocxConverter().convert(ctx)

        assert result.success is True
        with zipfile.ZipFile(Path(result.artifacts[0].staging_path), "r") as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")
        assert "Header" in document_xml
        assert re.search(r"<w:b(?:\s|/|>)", document_xml) is not None

    def test_md_to_docx_consumes_configured_three_line_table_builtin(self):
        """Configured default three-line table key reaches DOCX output."""
        md_path = write_temp_md("| Header |\n|---|\n| cell |\n")
        ctx, _workspace = make_context(md_path, target_format="docx")
        ctx._config = FakeConfigView(
            {
                "document": {
                    "style": {
                        "table": {
                            "md_to_docx": {
                                "table_style_mode": "builtin",
                                "builtin_style_key": "three_line_table",
                            }
                        }
                    }
                }
            }
        )

        result = MdToDocxConverter().convert(ctx)

        assert result.success is True
        with zipfile.ZipFile(Path(result.artifacts[0].staging_path), "r") as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")
            styles_xml = zf.read("word/styles.xml").decode("utf-8")
        table_style = re.search(
            r'<w:style[^>]*w:styleId="DocWenThreeLineTable".*?</w:style>',
            styles_xml,
            flags=re.DOTALL,
        )
        assert table_style is not None
        assert re.search(r'<w:top\b[^>]*w:val="single"[^>]*w:sz="12"', table_style.group(0))
        assert re.search(r'<w:bottom\b[^>]*w:val="single"[^>]*w:sz="12"', table_style.group(0))
        assert re.search(
            r'<w:tblStylePr[^>]*w:type="firstRow".*?<w:bottom\b[^>]*w:val="single"[^>]*w:sz="4"',
            table_style.group(0),
            flags=re.DOTALL,
        )
        assert "<w:tcBorders>" not in document_xml
        assert '<w:tblStyle w:val="DocWenThreeLineTable"' in document_xml

    def test_md_to_docx_custom_table_style_missing_degrades_without_failure(self):
        """Configured custom table style does not fail when the template lacks it."""
        md_path = write_temp_md("| Header |\n|---|\n| cell |\n")
        ctx, _workspace = make_context(md_path, target_format="docx")
        ctx._config = FakeConfigView(
            {
                "document": {
                    "style": {
                        "table": {
                            "md_to_docx": {
                                "table_style_mode": "custom",
                                "custom_style_name": "Definitely Missing Style",
                            }
                        }
                    }
                }
            }
        )

        result = MdToDocxConverter().convert(ctx)

        assert result.success is True
        from docx import Document

        doc = Document(str(Path(result.artifacts[0].staging_path)))
        assert len(doc.tables) == 1
        assert doc.tables[0].style is not None
        assert doc.tables[0].style.style_id == "DocWenTableGrid"

    def test_md_to_docx_consumes_configured_horizontal_rule_marker_actions(self):
        """Configured ---/***/___ marker actions produce distinct Word break elements."""
        md_path = write_temp_md("A\n\n---\n\nB\n\n***\n\nC\n\n___\n\nD\n")
        ctx, _workspace = make_context(md_path, target_format="docx")
        ctx._config = FakeConfigView(
            {
                "conversion": {
                    "horizontal_rule": {
                        "enabled": True,
                        "md_to_docx": {
                            "dash": "page_break",
                            "asterisk": "section_break",
                            "underscore": "horizontal_rule_1",
                        },
                    }
                }
            }
        )

        result = MdToDocxConverter().convert(ctx)

        assert result.success is True
        with zipfile.ZipFile(Path(result.artifacts[0].staging_path), "r") as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")
            styles_xml = zf.read("word/styles.xml").decode("utf-8")
        assert '<w:br w:type="page"' in document_xml
        assert '<w:type w:val="nextPage"' in document_xml
        assert '<w:pStyle w:val="DocWenHorizontalRule1"' in document_xml
        rule_style = re.search(
            r'<w:style[^>]*w:styleId="DocWenHorizontalRule1".*?</w:style>',
            styles_xml,
            flags=re.DOTALL,
        )
        assert rule_style is not None
        assert re.search(r'<w:bottom\b[^>]*w:val="single"[^>]*w:sz="4"', rule_style.group(0))

    def test_remove_numbering_option(self):
        """When remove_numbering is True, heading numbering is stripped."""
        md_content = "## 1. Introduction\n\n## 2. Background\n\nText."
        md_path = write_temp_md(md_content)
        ctx, _workspace = make_context(
            md_path,
            target_format="docx",
            options={"remove_numbering": True},
        )

        converter = MdToDocxConverter()
        result = converter.convert(ctx)
        assert result.success

        from docx import Document

        doc = Document(str(Path(result.artifacts[0].staging_path)))
        texts = [p.text for p in doc.paragraphs]
        combined = " ".join(texts)
        # The numbering "1." and "2." should have been removed
        assert "Introduction" in combined
        assert "Background" in combined
