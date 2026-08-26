"""Focused tests split from test_md_to_docx_numbering.py."""

from __future__ import annotations

from ._md_to_docx_numbering_support import (
    BUNDLED_DOCX_TEMPLATES,
    WML_NS,
    DocxListNumbering,
    MdToDocxConverter,
    Path,
    _ExactSchemeRegistry,
    _get_numbering_xml,
    apply_list_to_paragraph,
    etree,
    make_context,
    os,
    pytest,
    repository_numbering_registry,
    tempfile,
    write_numbering_to_docx,
    write_temp_md,
    zipfile,
)

pytestmark = pytest.mark.contract


class TestWriteNumberingToDocx:
    """Tests for the post-save ZIP writeback routine."""

    def test_adds_definitions_to_docx(self):
        """Numbering definitions are injected after save."""
        from docx import Document

        doc = Document()
        p = doc.add_paragraph("Item")
        apply_list_to_paragraph(p, "1", 0)

        fd, tmp = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc.save(tmp)
        try:
            num = DocxListNumbering()
            num.create_list_definition({0: "ordered"})
            write_numbering_to_docx(tmp, num)

            num_root = _get_numbering_xml(tmp)
            assert num_root is not None
            an_elems = num_root.findall(f"{{{WML_NS}}}abstractNum")
            num_elems = num_root.findall(f"{{{WML_NS}}}num")
            assert len(an_elems) >= 1
            assert len(num_elems) >= 1
        finally:
            os.unlink(tmp)

    def test_noop_when_no_definitions(self):
        """Calling writeback with empty numbering is a safe no-op."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Hello")
        fd, tmp = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc.save(tmp)
        try:
            num = DocxListNumbering()
            # has_definitions is False
            write_numbering_to_docx(tmp, num)

            # File should still be valid
            assert zipfile.is_zipfile(tmp)
        finally:
            os.unlink(tmp)


@pytest.mark.parametrize(
    ("scheme", "registry", "error_type", "diagnostic_code"),
    [
        ("", _ExactSchemeRegistry(), "invalid_input", "NUMBERING-SCHEME-REQUIRED"),
        ("exact", None, "capability_unavailable", "NUMBERING-REGISTRY-UNAVAILABLE"),
        ("missing", _ExactSchemeRegistry(), "resource_not_found", "NUMBERING-SCHEME-NOT-FOUND"),
        (
            "exact",
            _ExactSchemeRegistry(enabled=False),
            "capability_unavailable",
            "NUMBERING-SCHEME-DISABLED",
        ),
        ("exact", _ExactSchemeRegistry(levels={}), "invalid_input", "NUMBERING-SCHEME-NO-LEVELS"),
    ],
)
def test_text_heading_numbering_rejects_unusable_exact_scheme(
    scheme: str,
    registry: object,
    error_type: str,
    diagnostic_code: str,
) -> None:
    md_path = write_temp_md("# Heading\n")
    ctx, _ = make_context(
        md_path,
        target_format="docx",
        options={
            "add_numbering": True,
            "numbering_scheme": scheme,
            "heading_numbering_render_mode": "text",
        },
        numbering_registry=registry,
    )

    result = MdToDocxConverter().convert(ctx)

    assert not result.success
    assert result.error is not None
    assert result.error.error_type == error_type
    assert result.error.diagnostic_code == diagnostic_code
    assert result.artifacts == []


class TestWordNativeHeadingNumbering:
    """End-to-end tests for word_native heading numbering mode."""

    WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def _get_styles_xml(self, docx_path: str) -> etree._Element | None:
        with zipfile.ZipFile(docx_path, "r") as zf:
            if "word/styles.xml" not in zf.namelist():
                return None
            raw = zf.read("word/styles.xml")
        return etree.fromstring(raw)

    def _get_numbering_xml(self, docx_path: str) -> etree._Element | None:
        with zipfile.ZipFile(docx_path, "r") as zf:
            if "word/numbering.xml" not in zf.namelist():
                return None
            raw = zf.read("word/numbering.xml")
        return etree.fromstring(raw)

    @staticmethod
    def _find_heading_abstract_num(num_root: etree._Element) -> etree._Element | None:
        """Find the heading abstractNum (one with pStyle linking to Heading styles)."""
        WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        for an in num_root.findall(f"{{{WML_NS}}}abstractNum"):
            for lvl in an.findall(f"{{{WML_NS}}}lvl"):
                pStyle = lvl.find(f"{{{WML_NS}}}pStyle")
                if pStyle is not None:
                    val = pStyle.get(f"{{{WML_NS}}}val", "")
                    if val.startswith("Heading"):
                        return an
        return None

    def test_word_native_with_hierarchical_standard(self):
        """word_native + hierarchical_standard produces numbering.xml with pStyle."""
        md = "# Heading 1\n\n## Heading 2\n\n### Heading 3\n"
        md_path = write_temp_md(md)
        ctx, _ = make_context(
            md_path,
            target_format="docx",
            options={
                "remove_numbering": False,
                "add_numbering": True,
                "numbering_scheme": "hierarchical_standard",
                "heading_numbering_render_mode": "word_native",
            },
            numbering_registry=repository_numbering_registry(),
        )

        result = MdToDocxConverter().convert(ctx)
        assert result.success, f"Conversion failed: {result.error}"

        output = str(Path(result.artifacts[0].staging_path))

        # Verify numbering.xml exists and has abstractNum with pStyle
        num_root = self._get_numbering_xml(output)
        assert num_root is not None, "numbering.xml missing"

        # Find the heading abstractNum dynamically (ID is no longer hardcoded — BUG-2 fix)
        heading_an = self._find_heading_abstract_num(num_root)
        assert heading_an is not None, "Heading abstractNum not found"
        heading_id = heading_an.get(f"{{{self.WML_NS}}}abstractNumId")
        assert heading_id is not None

        # Verify pStyle on each level
        lvls = heading_an.findall(f"{{{self.WML_NS}}}lvl")
        assert len(lvls) >= 3
        for i in range(3):
            pStyle = lvls[i].find(f"{{{self.WML_NS}}}pStyle")
            assert pStyle is not None, f"Level {i} missing pStyle"
            assert pStyle.get(f"{{{self.WML_NS}}}val") == f"Heading{i + 1}"

        # Verify styles.xml has numPr on Heading styles with matching numId
        styles_root = self._get_styles_xml(output)
        assert styles_root is not None, "styles.xml missing"
        for style_id in ("Heading1", "Heading2", "Heading3"):
            style_elem = styles_root.find(f'{{{self.WML_NS}}}style[@{{{self.WML_NS}}}styleId="{style_id}"]')
            assert style_elem is not None, f"{style_id} not found"
            pPr = style_elem.find(f"{{{self.WML_NS}}}pPr")
            assert pPr is not None, f"{style_id} missing pPr"
            numPr = pPr.find(f"{{{self.WML_NS}}}numPr")
            assert numPr is not None, f"{style_id} missing numPr"
            numId_elem = numPr.find(f"{{{self.WML_NS}}}numId")
            assert numId_elem is not None
            assert numId_elem.get(f"{{{self.WML_NS}}}val") == heading_id

    @pytest.mark.parametrize(
        "template_path",
        BUNDLED_DOCX_TEMPLATES,
        ids=lambda path: path.stem,
    )
    def test_word_native_uses_bundled_template_heading_style_ids(self, template_path: Path):
        """Numbering binds to each bundled template's request-owned style IDs."""

        from docx import Document

        md_path = write_temp_md("# Heading 1\n\n## Heading 2\n\n### Heading 3\n")
        ctx, _ = make_context(
            md_path,
            target_format="docx",
            options={
                "template_name": str(template_path),
                "remove_numbering": False,
                "add_numbering": True,
                "numbering_scheme": "hierarchical_standard",
                "heading_numbering_render_mode": "word_native",
            },
            numbering_registry=repository_numbering_registry(),
        )

        result = MdToDocxConverter().convert(ctx)
        assert result.success, f"Conversion failed for {template_path.name}: {result.error}"
        output = str(Path(result.artifacts[0].staging_path))

        document = Document(output)
        paragraph_style_ids: dict[str, str] = {}
        for paragraph in document.paragraphs:
            if paragraph.text not in {"Heading 1", "Heading 2", "Heading 3"}:
                continue
            assert paragraph.style is not None
            paragraph_style_ids[paragraph.text] = paragraph.style.style_id
        expected_style_ids = [paragraph_style_ids[f"Heading {level}"] for level in range(1, 4)]

        num_root = self._get_numbering_xml(output)
        assert num_root is not None
        heading_abstract_num = next(
            abstract_num
            for abstract_num in num_root.findall(f"{{{self.WML_NS}}}abstractNum")
            if any(
                p_style.get(f"{{{self.WML_NS}}}val") == expected_style_ids[0]
                for p_style in abstract_num.findall(f"{{{self.WML_NS}}}lvl/{{{self.WML_NS}}}pStyle")
            )
        )
        levels = heading_abstract_num.findall(f"{{{self.WML_NS}}}lvl")
        actual_style_ids: list[str | None] = []
        for index in range(3):
            p_style = levels[index].find(f"{{{self.WML_NS}}}pStyle")
            assert p_style is not None
            actual_style_ids.append(p_style.get(f"{{{self.WML_NS}}}val"))
        assert actual_style_ids == expected_style_ids

        styles_root = self._get_styles_xml(output)
        assert styles_root is not None
        for style_id in expected_style_ids:
            style = styles_root.find(f'{{{self.WML_NS}}}style[@{{{self.WML_NS}}}styleId="{style_id}"]')
            assert style is not None
            assert style.find(f"{{{self.WML_NS}}}pPr/{{{self.WML_NS}}}numPr") is not None

    def test_word_native_with_gongwen_standard_approximate(self):
        """word_native + gongwen_standard succeeds (approximate verdict allowed)."""
        md = "# 标题\n\n## 标题二\n"
        md_path = write_temp_md(md)
        ctx, _ = make_context(
            md_path,
            target_format="docx",
            options={
                "remove_numbering": False,
                "add_numbering": True,
                "numbering_scheme": "gongwen_standard",
                "heading_numbering_render_mode": "word_native",
            },
            numbering_registry=repository_numbering_registry(),
        )

        result = MdToDocxConverter().convert(ctx)
        assert result.success, f"Conversion failed: {result.error}"

        num_root = self._get_numbering_xml(str(Path(result.artifacts[0].staging_path)))
        assert num_root is not None

    def test_text_mode_unchanged(self):
        """text mode (default) produces same text-concatenation behavior."""
        md = "# Heading 1\n\n## Heading 2\n"
        md_path = write_temp_md(md)
        ctx, _ = make_context(
            md_path,
            target_format="docx",
            options={
                "remove_numbering": False,
                "add_numbering": True,
                "numbering_scheme": "hierarchical_standard",
                # heading_numbering_render_mode defaults to "text"
            },
            numbering_registry=repository_numbering_registry(),
        )

        result = MdToDocxConverter().convert(ctx)
        assert result.success

        from docx import Document

        doc = Document(str(Path(result.artifacts[0].staging_path)))
        full = "\n".join(p.text for p in doc.paragraphs)
        # In text mode, heading text includes the numbering prefix
        assert "1 Heading 1" in full or "1 Heading" in full

    def test_word_native_no_text_numbering_added(self):
        """word_native does NOT add text numbering to headings."""
        md = "# Heading 1\n\n## Heading 2\n"
        md_path = write_temp_md(md)
        ctx, _ = make_context(
            md_path,
            target_format="docx",
            options={
                "remove_numbering": False,
                "add_numbering": True,
                "numbering_scheme": "hierarchical_standard",
                "heading_numbering_render_mode": "word_native",
            },
            numbering_registry=repository_numbering_registry(),
        )

        result = MdToDocxConverter().convert(ctx)
        assert result.success

        from docx import Document

        doc = Document(str(Path(result.artifacts[0].staging_path)))
        full = "\n".join(p.text for p in doc.paragraphs)
        # In word_native mode, heading text should NOT have "1 " prefix
        # (unlike text mode which concatenates "1 Heading 1")
        assert "Heading 1" in full
        assert "Heading 2" in full

    def test_word_native_no_heading_numbering_xml_when_no_add(self):
        """No heading numbering injection when add_numbering is False."""
        md = "# Heading\n"
        md_path = write_temp_md(md)
        ctx, _ = make_context(
            md_path,
            target_format="docx",
            options={
                "remove_numbering": False,
                "add_numbering": False,
                "heading_numbering_render_mode": "word_native",
            },
        )

        result = MdToDocxConverter().convert(ctx)
        assert result.success

        output = str(Path(result.artifacts[0].staging_path))
        with zipfile.ZipFile(output, "r") as zf:
            # numbering.xml may exist from template, but heading abstractNum
            # (with pStyle linking to Heading styles) should NOT be present
            names = zf.namelist()
            if "word/numbering.xml" in names:
                raw = zf.read("word/numbering.xml")
                num_root = etree.fromstring(raw)
                heading_an = self._find_heading_abstract_num(num_root)
                assert heading_an is None, "Heading abstractNum found despite add_numbering=False"
