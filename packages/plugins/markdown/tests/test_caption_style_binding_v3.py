"""Managed-style collision integration for the v3 DOCX caption map."""

from __future__ import annotations

import hashlib
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

from docwen_core.docx_semantics_v3 import (
    CAPTION_STYLE_BINDING_MAP_NAMESPACE,
    CaptionStyleBindingV3,
    DocxSemanticsV3Recovery,
    DocxSemanticsV3Session,
)
from docwen_core.docx_styles import SHIPPED_STYLE_LOCALES
from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter
from docwen_plugin_markdown.to_docx.managed_styles import complete_managed_styles
from docwen_runtime.config.document_styles import build_document_style_catalog

from .conftest import PROJECT_ROOT, make_context

pytestmark = pytest.mark.contract


@pytest.mark.parametrize("locale", SHIPPED_STYLE_LOCALES)
def test_all_shipped_locales_have_a_valid_complete_caption_binding_set(locale: str) -> None:
    catalog = build_document_style_catalog(
        {"gui": {"language": {"locale": locale}}},
        locales_dir=PROJECT_ROOT / "i18n" / "locales",
    )
    document, managed = complete_managed_styles(Document(), catalog)
    semantic_keys = (
        "figure_caption",
        "table_caption",
        "equation_caption",
        "code_block_caption",
    )
    bindings = tuple(
        CaptionStyleBindingV3(
            semantic_key=key,  # type: ignore[arg-type]
            resolved_style_id=managed.style_id(key),
            visible_name=managed.get(key).name or "",
        )
        for key in semantic_keys
    )
    session = DocxSemanticsV3Session(
        document,
        source_sha256=hashlib.sha256(locale.encode()).hexdigest(),
        caption_style_bindings=bindings,
    )
    assert session.has_projection is False


def test_real_converter_persists_and_recovers_request_local_caption_style(tmp_path: Path) -> None:
    catalog = build_document_style_catalog(
        {"gui": {"language": {"locale": "en_US"}}},
        locales_dir=PROJECT_ROOT / "i18n" / "locales",
    )
    template = tmp_path / "caption-style-conflict.docx"
    document = Document(str(PROJECT_ROOT / "scripts" / "maintenance" / "空白模板.docx"))
    conflict = document.styles.add_style(catalog.name_for("table_caption"), WD_STYLE_TYPE.CHARACTER)
    assert conflict._element is not None
    conflict._element.set(qn("w:styleId"), "DocWenTableCaption")
    document.save(str(template))
    template_sha256 = hashlib.sha256(template.read_bytes()).hexdigest()
    source = tmp_path / "source.md"
    source.write_text(
        "Table: Results ^results\n\n| A |\n|---|\n| 1 |\n",
        encoding="utf-8",
    )
    context, _workspace = make_context(
        str(source),
        target_format="docx",
        options={"template_name": str(template)},
        document_style_catalog=catalog,
    )

    result = MdToDocxConverter().convert(context)

    assert result.success is True, (result.error, result.diagnostics)
    assert len(result.artifacts) == 1
    output = Path(result.artifacts[0].staging_path)
    assert hashlib.sha256(template.read_bytes()).hexdigest() == template_sha256
    collision = next(item for item in result.diagnostics if item.code == "MD2DOCX-STYLE-COLLISION-PRESERVED")
    assert collision.location == ("style:table_caption;requested:DocWenTableCaption;resolved:DocWenTableCaptionDocWen1")
    with ZipFile(output) as package:
        caption_map = next(
            package.read(name)
            for name in package.namelist()
            if name.startswith("customXml/item")
            and name.endswith(".xml")
            and CAPTION_STYLE_BINDING_MAP_NAMESPACE.encode() in package.read(name)
        )
    assert b'semantic_key="table_caption"' in caption_map
    assert b'resolved_style_id="DocWenTableCaptionDocWen1"' in caption_map
    recovery = DocxSemanticsV3Recovery.load(output, Document(str(output)))
    assert recovery.caption_signatures == (("table", "results", "Results", "1"),)
    assert {item.semantic_key: item.resolved_style_id for item in recovery.caption_style_bindings}[
        "table_caption"
    ] == "DocWenTableCaptionDocWen1"
