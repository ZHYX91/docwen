"""OOXML contract tests for the complete managed DOCX style registry."""

from __future__ import annotations

import hashlib
import unicodedata
from copy import deepcopy
from functools import cache
from io import BytesIO
from pathlib import Path
from typing import Any, cast
from zipfile import ZipFile

import lxml.etree as etree
import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docwen_core.docx_styles import (
    BUILTIN_DOCUMENT_STYLES,
    CUSTOM_DOCUMENT_STYLES,
    MANAGED_DOCUMENT_STYLES,
    SHIPPED_STYLE_LOCALES,
    DocumentStyleCatalog,
)
from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter
from docwen_plugin_markdown.to_docx.managed_styles import (
    ManagedStyleCompletionError,
    _recognition_names,
    complete_managed_styles,
    validate_managed_style_package,
)
from docwen_runtime.config.document_styles import build_document_style_catalog
from docwen_runtime.toml_io import read_toml_file

from .conftest import PROJECT_ROOT, make_context

pytestmark = pytest.mark.contract

_WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_REF_TAGS = ("pStyle", "rStyle", "tblStyle", "basedOn", "next", "link", "numStyleLink", "styleLink")
_REF_QNAMES = tuple(qn(f"w:{name}") for name in _REF_TAGS)
_KIND_TO_TYPE = {
    "paragraph": WD_STYLE_TYPE.PARAGRAPH,
    "character": WD_STYLE_TYPE.CHARACTER,
    "table": WD_STYLE_TYPE.TABLE,
}


@cache
def _catalog(locale: str) -> DocumentStyleCatalog:
    return build_document_style_catalog(
        {"gui": {"language": {"locale": locale}}},
        locales_dir=PROJECT_ROOT / "i18n" / "locales",
    )


def _template_path(locale: str) -> Path:
    table = read_toml_file(PROJECT_ROOT / "i18n" / "locales" / f"{locale}.toml")
    template_name = table["meta"]["template_name"]
    return PROJECT_ROOT / "templates" / f"{template_name}.docx"


def _style_elements(document) -> dict[str, etree._Element]:
    return {style.style_id: style._element for style in document.styles}


def _raw_name(element: etree._Element) -> str:
    name = element.find(qn("w:name"))
    return name.get(qn("w:val"), "") if name is not None else ""


def _normalized_name(value: str) -> str:
    return unicodedata.normalize("NFC", value).casefold()


def _legacy_map(document, catalog: DocumentStyleCatalog) -> dict[str, str]:
    aliases: dict[str, tuple[str, str, bool]] = {
        _normalized_name("Normal"): ("Normal", "paragraph", True),
        _normalized_name("Default Paragraph Font"): ("DefaultParagraphFont", "character", True),
        _normalized_name("Normal Table"): ("TableNormal", "table", True),
        _normalized_name("Title"): ("Title", "paragraph", True),
    }
    for definition in MANAGED_DOCUMENT_STYLES:
        names = _recognition_names(definition, catalog)
        assert names[0] is not None
        for name in names:
            assert name is not None
            aliases[_normalized_name(name)] = (definition.style_id, definition.kind, definition.is_builtin)

    result: dict[str, str] = {}
    for style in document.styles:
        matched = aliases.get(_normalized_name(_raw_name(style._element)))
        if matched is None:
            continue
        stable_id, expected_kind, preserve_host_id = matched
        assert style._element.get(qn("w:type")) == expected_kind
        result[style.style_id] = style.style_id if preserve_host_id else stable_id
    return result


def _format_fingerprint(element: etree._Element, migrations: dict[str, str]) -> bytes:
    copy = deepcopy(element)
    for attribute in (qn("w:styleId"), qn("w:customStyle")):
        if attribute in copy.attrib:
            del copy.attrib[attribute]
    for child_name in ("name", "aliases"):
        for child in copy.findall(qn(f"w:{child_name}")):
            copy.remove(child)
    for tag in _REF_QNAMES:
        for reference in copy.iter(tag):
            value = reference.get(qn("w:val"), "")
            if value in migrations:
                reference.set(qn("w:val"), migrations[value])
    return etree.tostring(copy, method="c14n", exclusive=True)


def _package_snapshot(document) -> tuple[tuple[str, str], ...]:
    snapshot: list[tuple[str, str]] = []
    for part in document.part.package.parts:
        try:
            payload = etree.tostring(etree.fromstring(part.blob), method="c14n", exclusive=True)
        except etree.XMLSyntaxError:
            payload = part.blob
        snapshot.append((str(part.partname), hashlib.sha256(payload).hexdigest()))
    return tuple(sorted(snapshot))


def _document_bytes(document) -> bytes:
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _rewrite_xml_member(blob: bytes, member_name: str, transform) -> bytes:
    source = BytesIO(blob)
    target = BytesIO()
    with ZipFile(source, "r") as archive_in, ZipFile(target, "w", allowZip64=True) as archive_out:
        for item in archive_in.infolist():
            payload = archive_in.read(item.filename)
            if item.filename == member_name:
                root = etree.fromstring(payload)
                transform(root)
                payload = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            archive_out.writestr(item, payload)
    return target.getvalue()


def _append_duplicate_member(blob: bytes, member_name: str) -> bytes:
    source = BytesIO(blob)
    target = BytesIO()
    with ZipFile(source, "r") as archive_in:
        members = [(item, archive_in.read(item.filename)) for item in archive_in.infolist()]
        duplicate_payload = archive_in.read(member_name)
    with ZipFile(target, "w", allowZip64=True) as archive_out:
        for item, payload in members:
            archive_out.writestr(item, payload)
        archive_out.writestr(member_name, duplicate_payload)
    return target.getvalue()


def _attach_xml_part(document, *, name: str, content_type: str, relationship_type: str, payload: bytes) -> None:
    part = Part(PackURI(f"/{name}"), content_type, payload, document.part.package)
    document.part.relate_to(part, relationship_type)


def _audit_package(path: Path) -> None:
    with ZipFile(path) as archive:
        styles = etree.fromstring(archive.read("word/styles.xml"))
        types = {
            element.get(qn("w:styleId"), ""): element.get(qn("w:type"), "") for element in styles.findall(qn("w:style"))
        }
        for name in archive.namelist():
            if not name.startswith("word/") or not name.endswith(".xml") or name.startswith("word/glossary/"):
                continue
            root = etree.fromstring(archive.read(name))
            for tag in _REF_QNAMES:
                for element in root.iter(tag):
                    style_id = element.get(qn("w:val"), "")
                    assert style_id in types, (name, style_id)
                    expected = {
                        qn("w:pStyle"): "paragraph",
                        qn("w:rStyle"): "character",
                        qn("w:tblStyle"): "table",
                    }.get(tag)
                    if expected is not None:
                        assert types[style_id] == expected, (name, style_id, types[style_id])


def test_blank_mother_template_completes_all_43_and_reopens(tmp_path: Path) -> None:
    source = PROJECT_ROOT / "scripts" / "maintenance" / "空白模板.docx"
    source_hash = hashlib.sha256(source.read_bytes()).hexdigest()

    document, bindings = complete_managed_styles(Document(str(source)), _catalog("zh_CN"))
    output = tmp_path / "completed-blank.docx"
    document.save(str(output))
    reopened = Document(str(output))
    styles = _style_elements(reopened)

    assert len(bindings.styles) == 43
    assert len(reopened.styles) == 47
    assert {definition.style_id for definition in MANAGED_DOCUMENT_STYLES}.issubset(styles)
    for definition in MANAGED_DOCUMENT_STYLES:
        element = styles[definition.style_id]
        assert element.get(qn("w:type")) == definition.kind
        assert element.find(qn("w:aliases")) is None
    assert hashlib.sha256(source.read_bytes()).hexdigest() == source_hash
    _audit_package(output)


def test_v4_registry_is_exact_and_does_not_reserve_listing_or_excluded_word_styles(tmp_path: Path) -> None:
    assert len(BUILTIN_DOCUMENT_STYLES) == 16
    assert len(CUSTOM_DOCUMENT_STYLES) == 27
    assert len(MANAGED_DOCUMENT_STYLES) == 43
    keys = {definition.semantic_key for definition in MANAGED_DOCUMENT_STYLES}
    style_ids = {definition.style_id for definition in MANAGED_DOCUMENT_STYLES}
    assert "code_block_caption" in keys
    assert "DocWenCodeBlockCaption" in style_ids
    assert {"listing_caption", "list_caption"}.isdisjoint(keys)
    assert {"DocWenListingCaption", "DocWenListCaption"}.isdisjoint(style_ids)

    completed, _bindings = complete_managed_styles(
        Document(str(PROJECT_ROOT / "scripts" / "maintenance" / "空白模板.docx")),
        _catalog("en_US"),
    )
    output = tmp_path / "exact-v4-registry.docx"
    completed.save(str(output))
    reopened = Document(str(output))
    actual_ids = {style.style_id for style in reopened.styles}
    actual_names = {_normalized_name(_raw_name(cast(Any, style)._element)) for style in reopened.styles}
    assert len(actual_ids) == 47
    assert {
        "FollowedHyperlink",
        "InTextCitation",
        "CrossReference",
        "BibliographyHeading",
        "DocWenListingCaption",
        "DocWenListCaption",
    }.isdisjoint(actual_ids)
    assert {
        _normalized_name(value)
        for value in (
            "Followed Hyperlink",
            "In-text Citation",
            "Cross Reference",
            "Bibliography Heading",
            "Listing Caption",
            "List Caption",
        )
    }.isdisjoint(actual_names)


@pytest.mark.parametrize("locale", SHIPPED_STYLE_LOCALES)
def test_shipped_numeric_templates_reuse_builtin_ids_and_migrate_custom_ids_without_format_loss(
    tmp_path: Path,
    locale: str,
) -> None:
    source = _template_path(locale)
    source_hash = hashlib.sha256(source.read_bytes()).hexdigest()
    original = Document(str(source))
    migration = _legacy_map(original, _catalog(locale))
    managed_ids = {definition.style_id for definition in MANAGED_DOCUMENT_STYLES}
    existing_styles = {
        stable_id: _format_fingerprint(_style_elements(original)[legacy_id], migration)
        for legacy_id, stable_id in migration.items()
    }
    original_names = {_normalized_name(_raw_name(cast(Any, style)._element)) for style in original.styles}
    existing_custom_ids = {
        definition.style_id
        for definition in CUSTOM_DOCUMENT_STYLES
        if any(_normalized_name(name) in original_names for name in _recognition_names(definition, _catalog(locale)))
    }

    completed, bindings = complete_managed_styles(original, _catalog(locale))
    output = tmp_path / f"{locale}.docx"
    completed.save(str(output))
    reopened = Document(str(output))
    current = _style_elements(reopened)

    assert len(existing_styles) == 39
    assert set(existing_styles) & managed_ids == existing_custom_ids
    assert len(reopened.styles) == 47
    for stable_id, fingerprint in existing_styles.items():
        assert _format_fingerprint(current[stable_id], {}) == fingerprint
    original_by_name = {
        _normalized_name(_raw_name(cast(Any, style)._element)): style.style_id for style in original.styles
    }
    for definition in BUILTIN_DOCUMENT_STYLES:
        original_id = original_by_name.get(_normalized_name(definition.canonical_name or ""))
        if original_id is None:
            assert bindings.style_id(definition.semantic_key) == definition.style_id
            continue
        assert bindings.style_id(definition.semantic_key) == original_id
        element = current[original_id]
        assert _raw_name(element) == definition.canonical_name
        assert element.find(qn("w:aliases")) is None
        if original_id != definition.style_id:
            assert definition.style_id not in current
    assert hashlib.sha256(source.read_bytes()).hexdigest() == source_hash
    _audit_package(output)


def test_cross_locale_recognition_changes_only_visible_identity(tmp_path: Path) -> None:
    source = _template_path("de_DE")
    original = Document(str(source))
    catalog = _catalog("zh_CN")
    migration = _legacy_map(original, catalog)
    original_body = _format_fingerprint(
        _style_elements(original)[next(old for old, new in migration.items() if new == "DocWenBodyParagraph")],
        migration,
    )

    completed, _bindings = complete_managed_styles(original, catalog)
    output = tmp_path / "de-template-zh-output.docx"
    completed.save(str(output))
    reopened = Document(str(output))
    body = _style_elements(reopened)["DocWenBodyParagraph"]

    assert _raw_name(body) == catalog.name_for("body_paragraph")
    assert _format_fingerprint(body, {}) == original_body


def test_generic_footnote_and_story_references_are_migrated(tmp_path: Path) -> None:
    source = _template_path("en_US")
    original = Document(str(source))
    catalog = _catalog("en_US")
    migration = _legacy_map(original, catalog)
    legacy_body = next(old for old, new in migration.items() if new == "DocWenBodyParagraph")
    legacy_inline = next(old for old, new in migration.items() if new == "DocWenInlineCode")

    paragraph = original.add_paragraph("body")
    paragraph._p.get_or_add_pPr().get_or_add_pStyle().set(qn("w:val"), legacy_body)
    run = paragraph.add_run("code")
    run_properties = run._r.get_or_add_rPr()
    run_style = OxmlElement("w:rStyle")
    run_style.set(qn("w:val"), legacy_inline)
    run_properties.append(run_style)
    header = original.sections[0].header.paragraphs[0]
    header._p.get_or_add_pPr().get_or_add_pStyle().set(qn("w:val"), legacy_body)
    footer = original.sections[0].footer.paragraphs[0]
    footer._p.get_or_add_pPr().get_or_add_pStyle().set(qn("w:val"), legacy_body)
    footnotes_xml = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        b'<w:footnote w:id="1"><w:p><w:pPr><w:pStyle w:val="'
        + legacy_body.encode()
        + b'"/></w:pPr><w:r><w:rPr><w:rStyle w:val="'
        + legacy_inline.encode()
        + b'"/></w:rPr><w:t>x</w:t></w:r></w:p></w:footnote></w:footnotes>'
    )
    _attach_xml_part(
        original,
        name="word/footnotes.xml",
        content_type=CT.WML_FOOTNOTES,
        relationship_type=RT.FOOTNOTES,
        payload=footnotes_xml,
    )
    for part_name, root_name, item_name, content_type, relationship_type in (
        ("word/endnotes.xml", "endnotes", "endnote", CT.WML_ENDNOTES, RT.ENDNOTES),
        ("word/comments.xml", "comments", "comment", CT.WML_COMMENTS, RT.COMMENTS),
    ):
        payload = (
            b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            + f'<w:{root_name} xmlns:w="{_WML_NS}">'.encode()
            + f'<w:{item_name} w:id="1"><w:p><w:pPr><w:pStyle w:val="{legacy_body}"/></w:pPr>'.encode()
            + f'<w:r><w:rPr><w:rStyle w:val="{legacy_inline}"/></w:rPr><w:t>x</w:t></w:r></w:p></w:{item_name}>'.encode()
            + f"</w:{root_name}>".encode()
        )
        _attach_xml_part(
            original,
            name=part_name,
            content_type=content_type,
            relationship_type=relationship_type,
            payload=payload,
        )
    numbering_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<w:numbering xmlns:w="{_WML_NS}">'
        f'<w:abstractNum w:abstractNumId="424242"><w:lvl w:ilvl="0">'
        f'<w:pStyle w:val="{legacy_body}"/></w:lvl></w:abstractNum></w:numbering>'
    ).encode()
    _attach_xml_part(
        original,
        name="word/numbering.xml",
        content_type=CT.WML_NUMBERING,
        relationship_type=RT.NUMBERING,
        payload=numbering_xml,
    )
    untouched_custom_xml = f'<root xmlns:w="{_WML_NS}"><w:pStyle w:val="{legacy_body}"/></root>'.encode()
    _attach_xml_part(
        original,
        name="customXml/item99.xml",
        content_type="application/xml",
        relationship_type=RT.CUSTOM_XML,
        payload=untouched_custom_xml,
    )
    untouched_glossary = (
        f'<w:glossaryDocument xmlns:w="{_WML_NS}"><w:pStyle w:val="{legacy_body}"/></w:glossaryDocument>'.encode()
    )
    _attach_xml_part(
        original,
        name="word/glossary/document.xml",
        content_type=CT.WML_DOCUMENT_GLOSSARY,
        relationship_type=RT.GLOSSARY_DOCUMENT,
        payload=untouched_glossary,
    )

    completed, _bindings = complete_managed_styles(original, catalog)
    output = tmp_path / "all-parts.docx"
    completed.save(str(output))
    with ZipFile(output) as archive:
        for name in (
            "word/document.xml",
            "word/header1.xml",
            "word/footer1.xml",
            "word/footnotes.xml",
            "word/endnotes.xml",
            "word/comments.xml",
            "word/numbering.xml",
            "word/styles.xml",
        ):
            root = etree.fromstring(archive.read(name))
            reference_values = {element.get(qn("w:val"), "") for tag in _REF_QNAMES for element in root.iter(tag)}
            assert legacy_body not in reference_values
            assert legacy_inline not in reference_values
        assert b"DocWenBodyParagraph" in archive.read("word/footnotes.xml")
        assert b"DocWenInlineCode" in archive.read("word/footnotes.xml")
        assert archive.read("customXml/item99.xml") == untouched_custom_xml
        assert archive.read("word/glossary/document.xml") == untouched_glossary
    _audit_package(output)


def test_default_document_parallel_style_registry_stays_in_sync(tmp_path: Path) -> None:
    completed, _bindings = complete_managed_styles(Document(), _catalog("en_US"))
    output = tmp_path / "default.docx"
    completed.save(str(output))
    with ZipFile(output) as archive:
        primary = etree.fromstring(archive.read("word/styles.xml"))
        effects = etree.fromstring(archive.read("word/stylesWithEffects.xml"))
    primary_ids = {element.get(qn("w:styleId")) for element in primary.findall(qn("w:style"))}
    effects_ids = {element.get(qn("w:styleId")) for element in effects.findall(qn("w:style"))}
    managed_ids = {definition.style_id for definition in MANAGED_DOCUMENT_STYLES}
    assert managed_ids <= primary_ids
    assert managed_ids <= effects_ids


def test_parallel_style_registry_uses_its_own_legacy_ids_and_rewrites_graph(tmp_path: Path) -> None:
    legacy_effects_id = "LegacyEffectsHeadingOne"

    def mutate_effects(root: etree._Element) -> None:
        heading = next(element for element in root.findall(qn("w:style")) if element.get(qn("w:styleId")) == "Heading1")
        heading.set(qn("w:styleId"), legacy_effects_id)
        derived = OxmlElement("w:style")
        derived.set(qn("w:type"), "paragraph")
        derived.set(qn("w:styleId"), "EffectsDerived")
        name = OxmlElement("w:name")
        name.set(qn("w:val"), "Effects Derived")
        based_on = OxmlElement("w:basedOn")
        based_on.set(qn("w:val"), legacy_effects_id)
        derived.extend((name, based_on))
        root.append(derived)

    source_blob = _rewrite_xml_member(_document_bytes(Document()), "word/stylesWithEffects.xml", mutate_effects)
    source_effects = etree.fromstring(ZipFile(BytesIO(source_blob)).read("word/stylesWithEffects.xml"))
    source_heading = next(
        element
        for element in source_effects.findall(qn("w:style"))
        if element.get(qn("w:styleId")) == legacy_effects_id
    )
    source_fingerprint = _format_fingerprint(source_heading, {legacy_effects_id: "Heading1"})

    completed, _bindings = complete_managed_styles(Document(BytesIO(source_blob)), _catalog("en_US"))
    output = tmp_path / "independent-effects.docx"
    completed.save(str(output))
    with ZipFile(output) as archive:
        effects = etree.fromstring(archive.read("word/stylesWithEffects.xml"))
    effects_by_id = {element.get(qn("w:styleId"), ""): element for element in effects.findall(qn("w:style"))}

    assert legacy_effects_id not in effects_by_id
    assert _format_fingerprint(effects_by_id["Heading1"], {}) == source_fingerprint
    derived_based_on = effects_by_id["EffectsDerived"].find(qn("w:basedOn"))
    assert derived_based_on is not None
    assert derived_based_on.get(qn("w:val")) == "Heading1"


def test_parallel_builtin_style_conflict_fails_closed_without_caller_mutation() -> None:
    def break_effects(root: etree._Element) -> None:
        heading = next(element for element in root.findall(qn("w:style")) if element.get(qn("w:styleId")) == "Heading1")
        heading.set(qn("w:type"), "character")

    source_blob = _rewrite_xml_member(_document_bytes(Document()), "word/stylesWithEffects.xml", break_effects)
    document = Document(BytesIO(source_blob))
    before = _package_snapshot(document)

    with pytest.raises(ManagedStyleCompletionError) as raised:
        complete_managed_styles(document, _catalog("en_US"))

    assert _package_snapshot(document) == before
    assert raised.value.diagnostic_code == "MD2DOCX-STYLE-CONFLICT"
    assert raised.value.error_type == "invalid_input"


@pytest.mark.parametrize("conflict_kind", ("wrong_type", "split_identity"))
def test_custom_style_conflicts_preserve_user_style_and_allocate_request_local_id(conflict_kind: str) -> None:
    document = Document(str(PROJECT_ROOT / "scripts" / "maintenance" / "空白模板.docx"))
    if conflict_kind == "wrong_type":
        style = document.styles.add_style("Code Block Caption", WD_STYLE_TYPE.CHARACTER)
        style_element = style._element
        assert style_element is not None
        style_element.set(qn("w:styleId"), "DocWenCodeBlockCaption")
    else:
        by_id = document.styles.add_style("unrelated stable identity", WD_STYLE_TYPE.PARAGRAPH)
        by_id_element = by_id._element
        assert by_id_element is not None
        by_id_element.set(qn("w:styleId"), "DocWenCodeBlockCaption")
        document.styles.add_style(_catalog("en_US").name_for("code_block_caption"), WD_STYLE_TYPE.PARAGRAPH)
    preserved_style = _style_elements(document)["DocWenCodeBlockCaption"]
    preserved_style.append(OxmlElement("w:locked"))
    preserved_fingerprint = _format_fingerprint(preserved_style, {})
    before = _package_snapshot(document)

    completed, bindings = complete_managed_styles(document, _catalog("en_US"))
    current = _style_elements(completed)

    assert _package_snapshot(document) == before
    assert _format_fingerprint(current["DocWenCodeBlockCaption"], {}) == preserved_fingerprint
    assert current["DocWenCodeBlockCaption"].find(qn("w:aliases")) is None
    assert bindings.style_id("code_block_caption") == "DocWenCodeBlockCaptionDocWen1"
    assert current["DocWenCodeBlockCaptionDocWen1"].get(qn("w:type")) == "paragraph"
    assert [(item.code, item.semantic_key) for item in bindings.conflicts] == [
        ("MD2DOCX-STYLE-COLLISION-PRESERVED", "code_block_caption")
    ]


def test_converter_reports_preserved_style_conflict_and_creates_valid_artifact(tmp_path: Path) -> None:
    template = tmp_path / "conflicting-template.docx"
    document = Document(str(PROJECT_ROOT / "scripts" / "maintenance" / "空白模板.docx"))
    conflict = document.styles.add_style("Code Block Caption", WD_STYLE_TYPE.CHARACTER)
    conflict_element = conflict._element
    assert conflict_element is not None
    conflict_element.set(qn("w:styleId"), "DocWenCodeBlockCaption")
    document.save(str(template))
    template_hash = hashlib.sha256(template.read_bytes()).hexdigest()
    source = tmp_path / "source.md"
    source.write_text("# Source\n\nBody.\n", encoding="utf-8")
    context, _workspace = make_context(
        str(source),
        target_format="docx",
        options={"template_name": str(template)},
        document_style_catalog=_catalog("en_US"),
    )

    result = MdToDocxConverter().convert(context)

    assert result.success is True
    assert len(result.artifacts) == 1
    assert result.error is None
    collision = next(item for item in result.diagnostics if item.code == "MD2DOCX-STYLE-COLLISION-PRESERVED")
    assert collision.level == "warning"
    assert collision.location == (
        "style:code_block_caption;requested:DocWenCodeBlockCaption;resolved:DocWenCodeBlockCaptionDocWen1"
    )
    reopened = Document(result.artifacts[0].staging_path)
    by_id = _style_elements(reopened)
    assert by_id["DocWenCodeBlockCaption"].get(qn("w:type")) == "character"
    assert by_id["DocWenCodeBlockCaptionDocWen1"].get(qn("w:type")) == "paragraph"
    assert hashlib.sha256(template.read_bytes()).hexdigest() == template_hash


def test_converter_rejects_duplicate_template_zip_members_before_object_model_load(tmp_path: Path) -> None:
    template = tmp_path / "duplicate-member-template.docx"
    with pytest.warns(UserWarning, match="Duplicate name"):
        template.write_bytes(
            _append_duplicate_member(
                (PROJECT_ROOT / "scripts" / "maintenance" / "空白模板.docx").read_bytes(),
                "word/document.xml",
            )
        )
    source = tmp_path / "source.md"
    source.write_text("# Source\n\nBody.\n", encoding="utf-8")
    context, _workspace = make_context(
        str(source),
        target_format="docx",
        options={"template_name": str(template)},
        document_style_catalog=_catalog("en_US"),
    )

    result = MdToDocxConverter().convert(context)

    assert result.success is False
    assert result.artifacts == []
    assert result.error is not None
    assert result.error.error_type == "invalid_input"
    assert result.error.diagnostic_code == "MD2DOCX-TEMPLATE-PACKAGE-INVALID"
    assert [diagnostic.code for diagnostic in result.diagnostics] == ["MD2DOCX-TEMPLATE-PACKAGE-INVALID"]


@pytest.mark.parametrize("damage", ("duplicate", "missing_ref", "wrong_type_ref", "wrong_name", "aliases"))
def test_final_package_validator_rejects_post_render_identity_damage(damage: str) -> None:
    catalog = _catalog("en_US")
    completed, _bindings = complete_managed_styles(Document(), catalog)
    completed.add_paragraph("validation target")
    blob = _document_bytes(completed)

    if damage == "duplicate":
        with pytest.warns(UserWarning, match="Duplicate name"):
            damaged = _append_duplicate_member(blob, "word/document.xml")
    elif damage in {"missing_ref", "wrong_type_ref"}:

        def break_document(root: etree._Element) -> None:
            paragraph = root.find(f".//{qn('w:p')}")
            assert paragraph is not None
            properties = paragraph.find(qn("w:pPr"))
            if properties is None:
                properties = OxmlElement("w:pPr")
                paragraph.insert(0, properties)
            style = OxmlElement("w:pStyle")
            style.set(qn("w:val"), "MissingStyle" if damage == "missing_ref" else "DocWenInlineCode")
            properties.insert(0, style)

        damaged = _rewrite_xml_member(blob, "word/document.xml", break_document)
    else:

        def break_styles(root: etree._Element) -> None:
            body = next(
                item for item in root.findall(qn("w:style")) if item.get(qn("w:styleId")) == "DocWenBodyParagraph"
            )
            if damage == "wrong_name":
                name = body.find(qn("w:name"))
                assert name is not None
                name.set(qn("w:val"), "Wrong body identity")
            else:
                aliases = OxmlElement("w:aliases")
                aliases.set(qn("w:val"), "Legacy alias")
                body.append(aliases)

        damaged = _rewrite_xml_member(blob, "word/styles.xml", break_styles)

    with pytest.raises(ManagedStyleCompletionError) as raised:
        validate_managed_style_package(damaged, catalog)

    assert raised.value.diagnostic_code == "MD2DOCX-STYLE-COMPLETION-ERROR"
    assert raised.value.error_type == "conversion_failed"


def test_converter_final_style_audit_failure_leaves_no_artifact_or_staging_file(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "source.md"
    source.write_text("Body.\n", encoding="utf-8")
    context, workspace = make_context(
        str(source),
        target_format="docx",
        document_style_catalog=_catalog("en_US"),
    )

    def reject_final_package(
        _blob: bytes,
        _catalog: DocumentStyleCatalog,
        _bindings,
    ) -> None:
        raise ManagedStyleCompletionError(
            "MD2DOCX-STYLE-COMPLETION-ERROR",
            "Injected final audit failure.",
            error_type="conversion_failed",
        )

    monkeypatch.setattr(
        "docwen_plugin_markdown.to_docx.converter.validate_managed_style_package",
        reject_final_package,
    )

    result = MdToDocxConverter().convert(context)

    assert result.success is False
    assert result.artifacts == []
    assert workspace.registered_artifacts == []
    assert list(Path(workspace.staging_dir).glob("*.docx")) == []
    assert result.error is not None
    assert result.error.diagnostic_code == "MD2DOCX-STYLE-COMPLETION-ERROR"
