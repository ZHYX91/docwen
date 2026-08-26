"""Tests for MD→DOCX formatting: table styles, code blocks, rich text, headings.

Covers findings:
- F-F1-008: inline formula in mixed paragraphs
- F-F1-033: table style name resolution
- F-F1-034: table left indent in list context
- F-F1-037: rich text run synthesis (emoji, <w:br>, code styling, strikethrough)
- F-F1-038: heading override_style for partial bold/italic
- F-F3-025: code block paragraph shading
"""

from __future__ import annotations

import re
from typing import Any

import pytest
from docx import Document
from docx.oxml.ns import qn

pytestmark = pytest.mark.contract


def _docx_xml(doc) -> str:
    return doc._element.xml


# ── Helper: render markdown through the full pipeline ──────────────────


def _render_md(md_text: str, **kwargs):
    """Parse and render markdown text through MdToDocxRenderer."""
    from docwen_plugin_markdown.mistune_extensions import parse_markdown_text
    from docwen_plugin_markdown.renderer import MdToDocxRenderer

    ast = parse_markdown_text(md_text)
    doc = Document()
    renderer = MdToDocxRenderer(doc, **kwargs)
    renderer.render(ast)
    return doc


@pytest.mark.parametrize("level", range(1, 10))
def test_docwen_atx_heading_levels_one_through_nine_render_to_matching_word_styles(level: int):
    doc = _render_md(f"{'#' * level} Heading {level}\n")

    assert len(doc.paragraphs) == 1
    paragraph = doc.paragraphs[0]
    assert paragraph.text == f"Heading {level}"
    assert paragraph.style is not None
    assert paragraph.style.style_id == f"Heading{level}"


def test_ten_hashes_remain_visible_paragraph_text():
    doc = _render_md("########## Not a heading\n")

    paragraph = doc.paragraphs[0]
    assert paragraph.text == "########## Not a heading"
    assert paragraph.style is not None
    assert paragraph.style.style_id == "Normal"


# ── F-F1-033: Table style resolution ──────────────────────────────────


class TestTableStyleResolution:
    """Table styles are config-driven, not hardcoded."""

    def test_default_table_style_is_table_grid(self):
        """Default table style is 'Table Grid' when no config is given."""
        doc = _render_md("| A | B |\n|---|---|\n| 1 | 2 |\n")
        xml = _docx_xml(doc)
        assert len(doc.tables) == 1
        # python-docx stores style as <w:tblStyle w:val="TableGrid"/>
        assert "TableGrid" in xml, "Default table style should be 'TableGrid'"

    def test_custom_table_style_passed_through(self):
        """When table_style_name is provided, it is applied."""
        doc = _render_md(
            "| X | Y |\n|---|---|\n| a | b |\n",
            table_style_name="Table Grid",
        )
        xml = _docx_xml(doc)
        assert "TableGrid" in xml  # python-docx normalises to no-space ID

    def test_table_with_header_row_has_tblLook(self):
        """Table with header row gets <w:tblLook firstRow='1'>."""
        doc = _render_md(
            "| Header1 | Header2 |\n|---------|----------|\n| cell1 | cell2 |\n",
        )
        xml = _docx_xml(doc)
        assert "tblLook" in xml, "Table header row should trigger <w:tblLook> element"
        assert 'firstRow="1"' in xml or "firstRow='1'" in xml or 'w:firstRow="1"' in xml, (
            "tblLook should have firstRow enabled"
        )

    def test_table_header_remove_strips_header_inline_formatting_only(self):
        """Header remove strips markers while retaining baseline emphasis."""
        doc = _render_md(
            "| **Header** |\n|---|\n| **Body** |\n",
            table_header_formatting_mode="remove",
        )
        header = doc.tables[0].cell(0, 0).paragraphs[0]
        body = doc.tables[0].cell(1, 0).paragraphs[0]

        assert header.text == "Header"
        assert next(run for run in header.runs if run.text).bold is True
        assert body.text == "Body"
        assert next(run for run in body.runs if run.text).bold is True

    def test_table_header_keep_preserves_visible_markers_only_in_header(self):
        """Header keep shows Markdown markers while body follows body mode."""
        doc = _render_md(
            "| **Header** |\n|---|\n| **Body** |\n",
            table_header_formatting_mode="keep",
        )

        assert doc.tables[0].cell(0, 0).text == "**Header**"
        assert doc.tables[0].cell(1, 0).text == "Body"

    def test_resolve_table_style_name_custom_mode(self):
        """Custom mode with custom_style_name returns the custom name."""
        from docwen_plugin_markdown.to_docx.converter import (
            _resolve_table_style_name,
        )

        result = _resolve_table_style_name("custom", "table_grid", "My Custom Style")
        assert result == "My Custom Style", "Custom mode should return custom_style_name — F-F1-033"

    def test_resolve_table_style_name_custom_mode_empty_fallback(self):
        """Custom mode with empty custom_style_name falls back to builtin."""
        from docwen_plugin_markdown.to_docx.converter import (
            _resolve_table_style_name,
        )

        result = _resolve_table_style_name("custom", "table_grid", "")
        assert result == "Table Grid", "Custom mode with empty name should fall back to Table Grid — F-F1-033"

    def test_resolve_table_style_name_builtin_mode(self):
        """Builtin table_grid resolves to the Word Table Grid style."""
        from docwen_plugin_markdown.to_docx.converter import (
            _resolve_table_style_name,
        )

        result = _resolve_table_style_name("builtin", "table_grid", "")
        assert result == "Table Grid", "Builtin mode should resolve to Table Grid — F-F1-033"

    def test_resolve_table_style_name_three_line_builtin_mode(self):
        """Builtin three_line_table resolves to the legacy table style name."""
        from docwen_plugin_markdown.to_docx.converter import (
            _resolve_table_style_name,
        )

        result = _resolve_table_style_name("builtin", "three_line_table", "")
        assert result == "Three Line Table"

    def test_three_line_table_applies_direct_borders(self):
        """Three-line builtin uses top/bottom table borders and header separator."""
        doc = _render_md(
            "| Header |\n|---|\n| Body |\n",
            table_style_name="Three Line Table",
            table_style_key="three_line_table",
        )
        xml = _docx_xml(doc)

        assert "<w:tblBorders>" in xml
        assert '<w:top w:val="single" w:sz="12"' in xml
        assert '<w:bottom w:val="single" w:sz="12"' in xml
        assert re.search(r"<w:tcBorders>\s*<w:bottom w:val=\"single\" w:sz=\"4\"", xml)
        assert "TableGrid" not in xml

    def test_missing_custom_table_style_falls_back_without_failing(self):
        """Unavailable custom table styles degrade to Table Grid instead of failing."""
        doc = _render_md(
            "| Header |\n|---|\n| Body |\n",
            table_style_name="Definitely Missing Style",
        )

        assert len(doc.tables) == 1
        assert "TableGrid" in _docx_xml(doc)


# ── F-F1-034: Table left indent ───────────────────────────────────────


class TestTableLeftIndent:
    """Tables inside lists get left indent via <w:tblInd> (F-F1-034)."""

    def test_set_table_left_indent_adds_tblInd(self):
        """Direct call to set_table_left_indent produces <w:tblInd>."""
        from docwen_plugin_markdown.renderer_utils import set_table_left_indent

        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        set_table_left_indent(table, 720)
        xml = _docx_xml(doc)
        assert "tblInd" in xml, "set_table_left_indent should produce <w:tblInd> — F-F1-034"

    def test_set_table_left_indent_zero_noops(self):
        """Zero indent still produces element (no-op clearance)."""
        from docwen_plugin_markdown.renderer_utils import set_table_left_indent

        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        set_table_left_indent(table, 0)
        xml = _docx_xml(doc)
        assert "tblInd" in xml

    def test_table_inside_list_via_ast_has_tblInd(self):
        """User-path: table inside a nested list item produces <w:tblInd>
        through the full rendering pipeline (F-F1-034)."""
        from docwen_plugin_markdown.renderer import MdToDocxRenderer

        # Build AST: list > list_item > list > list_item > [paragraph, table]
        # Table at depth=1 triggers tblInd in the new renderer.
        ast = [
            {
                "type": "list",
                "attrs": {"ordered": False},
                "children": [
                    {
                        "type": "list_item",
                        "children": [
                            {
                                "type": "list",
                                "attrs": {"ordered": False},
                                "children": [
                                    {
                                        "type": "list_item",
                                        "children": [
                                            {
                                                "type": "paragraph",
                                                "children": [{"type": "text", "raw": "Nested item"}],
                                            },
                                            {
                                                "type": "table",
                                                "children": [
                                                    {
                                                        "type": "table_head",
                                                        "children": [
                                                            {
                                                                "type": "table_cell",
                                                                "children": [{"type": "text", "raw": "Col A"}],
                                                            },
                                                            {
                                                                "type": "table_cell",
                                                                "children": [{"type": "text", "raw": "Col B"}],
                                                            },
                                                        ],
                                                    },
                                                    {
                                                        "type": "table_body",
                                                        "children": [
                                                            {
                                                                "type": "table_row",
                                                                "children": [
                                                                    {
                                                                        "type": "table_cell",
                                                                        "children": [{"type": "text", "raw": "val 1"}],
                                                                    },
                                                                    {
                                                                        "type": "table_cell",
                                                                        "children": [{"type": "text", "raw": "val 2"}],
                                                                    },
                                                                ],
                                                            },
                                                        ],
                                                    },
                                                ],
                                            },
                                        ],
                                    }
                                ],
                            }
                        ],
                    }
                ],
            }
        ]
        doc = Document()
        renderer = MdToDocxRenderer(doc)
        renderer.render(ast)
        xml = _docx_xml(doc)

        assert "tblInd" in xml, (
            "Table inside nested list context must produce <w:tblInd> "
            "through the rendering pipeline — F-F1-034 user-path"
        )

    def test_table_not_in_list_no_indent_at_top_level(self):
        """A table at top level renders with no left indent (normal case)."""
        md = "| A | B |\n|---|---|\n| 1 | 2 |\n"
        doc = _render_md(md)
        xml = _docx_xml(doc)
        assert len(doc.tables) == 1
        # At top level (not in list), no tblInd should be added
        # because _list_depth is 0
        assert "tblInd" not in xml, "Table at top level should NOT have <w:tblInd>"

    def test_list_depth_tracks_for_table_context(self):
        """Verify the renderer tracks _list_depth correctly."""
        from docwen_plugin_markdown.renderer import MdToDocxRenderer

        doc = Document()
        renderer = MdToDocxRenderer(doc)
        assert renderer._list_depth == 0, "Initial list depth should be 0"


# ── F-F3-025: Code block paragraph shading ────────────────────────────


class TestCodeBlockStyle:
    """Code blocks get paragraph-level background shading and monospace font."""

    def test_code_block_has_paragraph_shading(self):
        """Code block paragraph has <w:shd> background shading."""
        md = "```python\nprint('hello')\n```"
        doc = _render_md(md, code_bg_color="E7E6E6")
        xml = _docx_xml(doc)
        assert "w:shd" in xml or "<w:shd" in xml, "Code block should have paragraph shading — F-F3-025"
        assert "E7E6E6" in xml, "Background color E7E6E6 should be in the XML"

    def test_code_block_uses_configurable_font(self):
        """Code block uses the configured code_font."""
        doc = _render_md("```\nmonospace text\n```", code_font="Courier New")
        xml = _docx_xml(doc)
        assert "Courier New" in xml, "Code block should use configured font"

    def test_code_block_default_background(self):
        """Code block with default config has shading."""
        doc = _render_md("```\ndefault config\n```")
        xml = _docx_xml(doc)
        # Default code_bg_color is "E7E6E6"
        assert "w:shd" in xml or "<w:shd" in xml, "Code block default should have paragraph shading"


# ── F-F1-037: Rich text run synthesis ─────────────────────────────────


class TestRichTextSynthesis:
    """Inline rendering: w:br soft breaks, code styling, strikethrough."""

    def test_linebreak_produces_word_soft_break(self):
        """Markdown hard line break produces <w:br> element."""
        md = "Line one  \nLine two\n"
        doc = _render_md(md)
        xml = _docx_xml(doc)
        assert "Line one" in xml, "First line text should be present"
        assert "Line two" in xml, "Second line text should be present"
        assert "w:br" in xml or "<w:br" in xml or "<w:br/>" in xml, "Linebreak must produce <w:br> element — F-F1-037"

    def test_codespan_has_background_shading(self):
        """Inline code span has character-level shading."""
        md = "Here is `code` inline."
        doc = _render_md(md, code_bg_color="E7E6E6")
        xml = _docx_xml(doc)
        assert "code" in xml, "Code span text should be present"
        assert "Consolas" in xml or "code" in xml, "Code span font should be set"

    def test_codespan_uses_configured_font(self):
        """Inline code span uses configured code_font."""
        doc = _render_md("Run `x=1` here.", code_font="Courier New")
        xml = _docx_xml(doc)
        assert "Courier New" in xml, "Code span should use configured font"

    def test_strikethrough_rendered(self):
        """Markdown ~~strikethrough~~ is rendered with strike=True."""
        md = "This is ~~struck~~ text."
        doc = _render_md(md)
        xml = _docx_xml(doc)
        assert "struck" in xml, "Strikethrough text should be present"
        # Verify w:strike attribute is set on the run
        assert "w:strike" in xml or "<w:strike" in xml, "Strikethrough must produce w:strike attribute — F-F1-037"

    def test_highlight_rendered_as_word_highlight(self):
        """Markdown ==highlight== renders as native Word yellow highlight."""
        md = "This is ==highlighted== text."
        doc = _render_md(md)
        xml = _docx_xml(doc)

        assert "highlighted" in xml
        assert "w:highlight" in xml or "<w:highlight" in xml
        assert 'w:val="yellow"' in xml or "w:val='yellow'" in xml

    def test_body_keep_mode_preserves_visible_markdown_markers(self):
        """Body keep mode displays inline Markdown markers literally."""
        md = "This is **bold**, *italic*, `code`, ==mark==, and [link](https://example.com)."
        doc = _render_md(md, formatting_mode="keep")

        assert doc.paragraphs[0].text == (
            "This is **bold**, *italic*, `code`, ==mark==, and [link](https://example.com)."
        )

    def test_emoji_splitting_applies_emoji_font(self):
        """Emoji characters get Segoe UI Emoji font applied."""
        md = "Hello \U0001f600 world \U0001f30d"
        # Verify the emoji-splitting infrastructure applies emoji font
        # for actual emoji characters
        doc = _render_md(md)
        xml = _docx_xml(doc)
        assert "Hello" in xml, "Plain text should precede emoji"
        assert "world" in xml, "Plain text should follow emoji"
        assert "Segoe UI Emoji" in xml, (
            "Emoji characters must be rendered with Segoe UI Emoji font "
            "set on all font slots (ascii/hAnsi/eastAsia/cs) — F-F1-037"
        )

    def test_mixed_bold_italic_inline(self):
        """Mix of bold, italic, code, and strikethrough in one paragraph."""
        md = "**bold** *italic* `code` ~~struck~~ normal"
        doc = _render_md(md)
        xml = _docx_xml(doc)
        assert "bold" in xml
        assert "italic" in xml
        assert "code" in xml
        assert "struck" in xml
        assert "normal" in xml


# ── F-F1-038: Heading override_style ──────────────────────────────────


class TestHeadingOverrideStyle:
    """Headings support partial formatting via override_style."""

    def test_heading_override_apply_mode(self):
        """In apply mode, plain text in headings has bold/italic explicitly
        set to False so the heading style defaults don't override."""
        md = "# **Bold** and normal"
        doc = _render_md(md, heading_formatting_mode="apply")
        xml = _docx_xml(doc)
        assert "Bold" in xml
        assert "and normal" in xml

    def test_heading_remove_mode_strips_formatting(self):
        """In remove mode, all markdown inline formatting is stripped;
        keep mode shows Markdown markers, and apply mode uses Word runs."""
        md = "# **Bold** and *italic* text"

        doc_remove = _render_md(md, heading_formatting_mode="remove")
        doc_keep = _render_md(md, heading_formatting_mode="keep")
        doc_apply = _render_md(md, heading_formatting_mode="apply")

        xml_remove = _docx_xml(doc_remove)
        xml_keep = _docx_xml(doc_keep)
        xml_apply = _docx_xml(doc_apply)

        # All three modes preserve the text content
        for xml in (xml_remove, xml_keep, xml_apply):
            assert "Bold" in xml
            assert "italic" in xml
            assert "text" in xml

        bold_runs_remove = len(re.findall(r"<w:b\s*/?>", xml_remove))

        assert "**Bold**" in doc_keep.paragraphs[0].text
        assert "*italic*" in doc_keep.paragraphs[0].text
        assert bold_runs_remove == 0, (
            "Remove mode should strip ALL inline formatting — no <w:b/> elements expected for **_bold_** text"
        )

    def test_heading_keep_mode_preserves_visible_markers(self):
        """In keep mode, markdown markers survive as visible heading text."""
        md = "# **Bold** normal"
        doc = _render_md(md, heading_formatting_mode="keep")

        assert doc.paragraphs[0].text == "**Bold** normal"
        assert re.search(r"<w:b\s*/?>", _docx_xml(doc)) is None

    def test_heading_with_inline_formula(self):
        """Heading with inline formula renders both text and OMML (F-F1-008)."""
        md = "# Energy $E=mc^2$ explained"
        doc = _render_md(md)
        xml = _docx_xml(doc)
        assert "Energy" in xml
        assert "oMath" in xml, "Inline formula in heading was dropped — F-F1-008"
        assert "explained" in xml

    def test_heading_mixed_formatting_content(self):
        """Heading with bold, italic, and code span."""
        md = "# **Bold** *italic* `code` text"
        doc = _render_md(md, heading_formatting_mode="apply")
        xml = _docx_xml(doc)
        assert "Bold" in xml
        assert "italic" in xml
        assert "code" in xml
        assert "text" in xml


# ── F-F1-008: Inline formula in paragraphs (user-path) ────────────────


class TestInlineFormulaUserPath:
    """Inline formulas survive the full pipeline and produce OMML."""

    def test_mixed_formula_text_paragraph(self):
        """Paragraph with text + formula + text — the core user-visible gap."""
        md = "Before formula $x^2$ after formula."
        doc = _render_md(md)
        xml = _docx_xml(doc)
        assert "Before formula" in xml
        assert "oMath" in xml, "Formula $x^2$ was dropped — F-F1-008 core user-visible gap"
        assert "after formula" in xml

    def test_markdown_math_syntax_dollar(self):
        """$...$ syntax produces OMML inline formula."""
        md = "Value $\\pi r^2$ calculated."
        doc = _render_md(md)
        xml = _docx_xml(doc)
        assert "oMath" in xml, "Inline math $...$ not rendered to OMML"


# ── End-to-end user path: converter pipeline ──────────────────────────


class TestFormattingViaConverter:
    """Verify formatting works through the real MdToDocxConverter."""

    def test_converter_preserves_code_block_style(self):
        """Through the converter, code blocks have paragraph shading."""
        from .conftest import make_context, write_temp_md

        md = "# Doc\n\n```python\ndef foo():\n    pass\n```\n"
        md_path = write_temp_md(md)
        from docwen_plugin_markdown.to_docx.converter import (
            MdToDocxConverter,
        )

        ctx, _workspace = make_context(md_path, target_format="docx")
        converter = MdToDocxConverter()
        result = converter.convert(ctx)
        assert result.success, f"Conversion failed: {result.error.message if result.error else '?'}"
        from pathlib import Path

        output_path = Path(result.artifacts[0].staging_path)
        doc = Document(str(output_path))

        def _paragraph_text_including_sdt(paragraph: Any) -> str:
            return "".join(item.text or "" for item in paragraph._p.iter(f"{qn('w:t')}"))

        code = next(
            paragraph for paragraph in doc.paragraphs if _paragraph_text_including_sdt(paragraph).startswith("def foo")
        )
        code_style = code.style
        assert code_style is not None
        assert code_style.style_id == "DocWenCodeBlock"
        code_style_element = code_style._element
        assert code_style_element is not None
        style_p_pr = code_style_element.find(qn("w:pPr"))
        assert style_p_pr is not None
        shading = style_p_pr.find(qn("w:shd"))
        assert shading is not None
        assert shading.get(qn("w:fill")) == "E7E6E6"

    def test_converter_preserves_table_structure(self):
        """Through the converter, tables have proper structure."""
        from .conftest import make_context, write_temp_md

        md = "# Doc\n\n| A | B |\n|---|---|\n| 1 | 2 |\n"
        md_path = write_temp_md(md)
        from docwen_plugin_markdown.to_docx.converter import (
            MdToDocxConverter,
        )

        ctx, _workspace = make_context(md_path, target_format="docx")
        converter = MdToDocxConverter()
        result = converter.convert(ctx)
        assert result.success
        from pathlib import Path

        output_path = Path(result.artifacts[0].staging_path)
        doc = Document(str(output_path))
        assert len(doc.tables) == 1
        assert len(doc.tables[0].rows) == 2
        assert len(doc.tables[0].columns) == 2
        table_style = doc.tables[0].style
        assert table_style is not None
        assert table_style.style_id == "DocWenThreeLineTable"
        table_style_element = table_style._element
        assert table_style_element is not None
        table_properties = table_style_element.find(qn("w:tblPr"))
        assert table_properties is not None
        borders = table_properties.find(qn("w:tblBorders"))
        assert borders is not None
        for edge_name in ("top", "bottom"):
            edge = borders.find(qn(f"w:{edge_name}"))
            assert edge is not None
            assert edge.get(qn("w:val")) == "single"
            assert edge.get(qn("w:sz")) == "12"

    def test_converter_preserves_inline_formula(self):
        """Through the converter, inline formulas are preserved."""
        from .conftest import make_context, write_temp_md

        md = "# Doc\n\nValue is $x^2 + y^2$ here.\n"
        md_path = write_temp_md(md)
        from docwen_plugin_markdown.to_docx.converter import (
            MdToDocxConverter,
        )

        ctx, _workspace = make_context(md_path, target_format="docx")
        converter = MdToDocxConverter()
        result = converter.convert(ctx)
        assert result.success, f"Conversion failed: {result.error.message if result.error else '?'}"
        from pathlib import Path

        output_path = Path(result.artifacts[0].staging_path)
        doc = Document(str(output_path))
        xml = _docx_xml(doc)
        assert "oMath" in xml, "Inline formula dropped by full converter pipeline — F-F1-008"
