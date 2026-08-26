"""Focused tests split from test_md_to_docx_notes.py."""

from __future__ import annotations

import pytest

from ._md_to_docx_notes_support import (
    MD_MIXED_NOTES,
    WML_NS,
    NoteContext,
    Path,
    _list_zip_entries,
    _mk_children,
    _q,
    _read_zip_entry,
    etree,
)

pytestmark = pytest.mark.contract


class TestWritebackMixedNotes:
    """F-F1-028: DOCX with both footnotes and endnotes."""

    def test_writeback_numbers_each_note_domain_by_first_reference(self):
        """Definition order does not control emitted Word note IDs."""
        from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter

        from .conftest import make_context, write_temp_md

        markdown = """First footnote[^footnote:b], repeated[^footnote:b], then[^a].
First endnote[^endnote:z], then[^endnote:y].

[^a]: footnote A
[^footnote:b]: footnote B
[^endnote:y]: endnote Y
[^endnote:z]: endnote Z
"""
        md_path = write_temp_md(markdown)
        source_path = Path(md_path)
        original = source_path.read_bytes()
        ctx, _workspace = make_context(md_path, target_format="docx")

        result = MdToDocxConverter().convert(ctx)

        assert result.success, f"Conversion failed: {result.error}"
        output_path = str(Path(result.artifacts[0].staging_path))
        doc_root = etree.fromstring(_read_zip_entry(output_path, "word/document.xml"))
        footnote_refs = [ref.get(_q("id")) for ref in doc_root.findall(f".//{{{WML_NS}}}footnoteReference")]
        endnote_refs = [ref.get(_q("id")) for ref in doc_root.findall(f".//{{{WML_NS}}}endnoteReference")]
        assert footnote_refs == ["1", "1", "2"]
        assert endnote_refs == ["1", "2"]

        footnote_root = etree.fromstring(_read_zip_entry(output_path, "word/footnotes.xml"))
        endnote_root = etree.fromstring(_read_zip_entry(output_path, "word/endnotes.xml"))

        def note_text(root: etree._Element, note_type: str, word_id: str) -> str:
            note = root.find(f"{{{WML_NS}}}{note_type}[@{{{WML_NS}}}id='{word_id}']")
            assert note is not None
            return "".join(item.text or "" for item in note.iter(f"{{{WML_NS}}}t"))

        assert "footnote B" in note_text(footnote_root, "footnote", "1")
        assert "footnote A" in note_text(footnote_root, "footnote", "2")
        assert "endnote Z" in note_text(endnote_root, "endnote", "1")
        assert "endnote Y" in note_text(endnote_root, "endnote", "2")
        assert source_path.read_bytes() == original

    def test_writeback_mixed_both_parts_present(self):
        """Markdown with footnotes + endnotes produces both parts."""
        from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter

        from .conftest import make_context, write_temp_md

        md_path = write_temp_md(MD_MIXED_NOTES)
        ctx, _workspace = make_context(md_path, target_format="docx")

        converter = MdToDocxConverter()
        result = converter.convert(ctx)
        output_path = str(Path(result.artifacts[0].staging_path))

        entries = _list_zip_entries(output_path)
        assert "word/footnotes.xml" in entries
        assert "word/endnotes.xml" in entries

    def test_writeback_mixed_openable(self):
        """Mixed notes DOCX is openable."""
        from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter

        from .conftest import make_context, write_temp_md

        md_path = write_temp_md(MD_MIXED_NOTES)
        ctx, _workspace = make_context(md_path, target_format="docx")

        converter = MdToDocxConverter()
        result = converter.convert(ctx)
        output_path = str(Path(result.artifacts[0].staging_path))

        from docx import Document

        doc = Document(output_path)
        assert doc is not None

        # Both footnote and endnote refs in body
        doc_xml = _read_zip_entry(output_path, "word/document.xml")
        root = etree.fromstring(doc_xml)
        fn_refs = root.findall(f".//{{{WML_NS}}}footnoteReference")
        en_refs = root.findall(f".//{{{WML_NS}}}endnoteReference")
        assert len(fn_refs) >= 1, f"Expected >=1 footnoteReference, got {len(fn_refs)}"
        assert len(en_refs) >= 1, f"Expected >=1 endnoteReference, got {len(en_refs)}"

    def test_writeback_mixed_relationships_valid(self):
        """Document relationships include both footnotes and endnotes entries."""
        from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter

        from .conftest import make_context, write_temp_md

        md_path = write_temp_md(MD_MIXED_NOTES)
        ctx, _workspace = make_context(md_path, target_format="docx")

        converter = MdToDocxConverter()
        result = converter.convert(ctx)
        output_path = str(Path(result.artifacts[0].staging_path))

        rels_xml = _read_zip_entry(output_path, "word/_rels/document.xml.rels")
        rels_root = etree.fromstring(rels_xml)

        targets = set()
        types = set()
        for rel in rels_root:
            tgt = rel.get("Target", "")
            typ = rel.get("Type", "")
            if tgt:
                targets.add(tgt)
            if typ:
                types.add(typ)

        assert "footnotes.xml" in targets, f"footnotes.xml not in rels targets: {targets}"
        assert "endnotes.xml" in targets, f"endnotes.xml not in rels targets: {targets}"

    def test_writeback_mixed_content_types_valid(self):
        """Content_Types includes both footnotes and endnotes overrides."""
        from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter

        from .conftest import make_context, write_temp_md

        md_path = write_temp_md(MD_MIXED_NOTES)
        ctx, _workspace = make_context(md_path, target_format="docx")

        converter = MdToDocxConverter()
        result = converter.convert(ctx)
        output_path = str(Path(result.artifacts[0].staging_path))

        ct_xml = _read_zip_entry(output_path, "[Content_Types].xml")
        ct_root = etree.fromstring(ct_xml)

        part_names = set()
        for override in ct_root:
            pn = override.get("PartName", "")
            if pn:
                part_names.add(pn)

        assert "/word/footnotes.xml" in part_names, f"/word/footnotes.xml not in Content_Types: {part_names}"
        assert "/word/endnotes.xml" in part_names, f"/word/endnotes.xml not in Content_Types: {part_names}"


class TestWritebackDirect:
    """F-F1-028: write_notes_to_docx unit-level tests."""

    def test_writeback_direct_function_call(self, tmp_path: Path):
        """Calling write_notes_to_docx directly on a saved DOCX works."""
        from docx import Document as DocxDocument

        from docwen_plugin_markdown.to_docx.notes import write_notes_to_docx

        # Build a minimal DOCX
        doc = DocxDocument()
        paragraph = doc.add_paragraph("Hello")

        # Create a NoteContext with footnote elements
        note_ctx = NoteContext()
        note_ctx._footnote_children = {"f1": _mk_children("Test footnote content")}
        reference = note_ctx.create_footnote_ref_run("f1")
        assert reference is not None
        paragraph._p.append(reference)

        docx_path = str(tmp_path / "test.docx")
        doc.save(docx_path)

        write_notes_to_docx(docx_path, note_ctx)

        # Verify part exists
        entries = _list_zip_entries(docx_path)
        assert "word/footnotes.xml" in entries

        # Verify openable
        doc2 = DocxDocument(docx_path)
        assert doc2 is not None

    def test_writeback_direct_no_elements_noop(self, tmp_path: Path):
        """Calling write_notes_to_docx with an empty NoteContext is a no-op."""
        from docx import Document as DocxDocument

        from docwen_plugin_markdown.to_docx.notes import write_notes_to_docx

        doc = DocxDocument()
        doc.add_paragraph("Hi")

        docx_path = str(tmp_path / "test.docx")
        doc.save(docx_path)

        entries_before = _list_zip_entries(docx_path)

        note_ctx = NoteContext()
        write_notes_to_docx(docx_path, note_ctx)

        entries_after = _list_zip_entries(docx_path)
        assert entries_before == entries_after, "No-op should not modify ZIP"

    def test_writeback_direct_idempotent(self, tmp_path: Path):
        """Calling write_notes_to_docx twice on the same file appends
        correctly (idempotency for already-existing parts)."""
        from docx import Document as DocxDocument

        from docwen_plugin_markdown.to_docx.notes import (
            prepare_note_context_for_document,
            write_notes_to_docx,
        )

        doc = DocxDocument()
        paragraph = doc.add_paragraph("Idempotent test")

        note_ctx = NoteContext()
        note_ctx._footnote_children = {"n1": _mk_children("Note one")}
        reference = note_ctx.create_footnote_ref_run("n1")
        assert reference is not None
        paragraph._p.append(reference)

        docx_path = str(tmp_path / "test.docx")
        doc.save(docx_path)

        write_notes_to_docx(docx_path, note_ctx)

        # Second call should not corrupt
        note_ctx2 = NoteContext()
        note_ctx2._footnote_children = {"n2": _mk_children("Note two")}
        reopened = DocxDocument(docx_path)
        prepare_note_context_for_document(reopened, note_ctx2)
        reference2 = note_ctx2.create_footnote_ref_run("n2")
        assert reference2 is not None
        reopened.paragraphs[0]._p.append(reference2)
        reopened.save(docx_path)

        write_notes_to_docx(docx_path, note_ctx2)

        # Verify it's still openable
        doc2 = DocxDocument(docx_path)
        assert doc2 is not None

        fn_xml = _read_zip_entry(docx_path, "word/footnotes.xml")
        root = etree.fromstring(fn_xml)
        footnotes = root.findall(f"{{{WML_NS}}}footnote")
        # separator + continuation + n1 + n2 = 4
        assert len(footnotes) == 4, f"Expected 4 footnotes after 2nd write, got {len(footnotes)}"
