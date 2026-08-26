"""Focused tests split from test_md_to_docx_numbering.py."""

from __future__ import annotations

import pytest

from ._md_to_docx_numbering_support import (
    WML_NS,
    DocxListNumbering,
    MdToDocxConverter,
    Path,
    _append_template_abstract_num,
    _get_numbering_xml,
    _list_paras_with_numpr,
    _our_abstract_num_id,
    apply_list_to_paragraph,
    copy,
    etree,
    hashlib,
    make_context,
    repository_numbering_registry,
    write_numbering_to_docx,
    write_temp_md,
)

pytestmark = pytest.mark.contract


def test_md_to_docx_text_and_word_native_use_request_cleanup_rules() -> None:
    from docx import Document

    from docwen_core.text.heading_numbering import (
        compile_clean_rules_from_data,
    )

    request_rules = compile_clean_rules_from_data(
        [{"id": "request", "enabled": True, "pattern": r"^REQ:\s*", "level": 1}]
    )

    for render_mode, add_numbering in (("text", False), ("word_native", True)):
        md_path = write_temp_md("# REQ: Request title\n\n# GLOBAL: Global title\n")
        ctx, _ = make_context(
            md_path,
            target_format="docx",
            options={
                "remove_numbering": render_mode == "text",
                "add_numbering": add_numbering,
                "numbering_scheme": "hierarchical_standard",
                "heading_numbering_render_mode": render_mode,
            },
            numbering_registry=repository_numbering_registry(),
            heading_cleanup_rules=request_rules,
        )

        result = MdToDocxConverter().convert(ctx)

        assert result.success, result.error
        output = Document(str(Path(result.artifacts[0].staging_path)))
        paragraph_text = "\n".join(paragraph.text for paragraph in output.paragraphs)
        assert "REQ:" not in paragraph_text
        assert "GLOBAL: Global title" in paragraph_text


class TestDocxListNumbering:
    """Unit tests for the numbering context collector."""

    def test_creates_list_definition(self):
        """A definition allocates a positive numId and stores OOXML elements."""
        num = DocxListNumbering()
        num_id = num.create_list_definition({0: "ordered", 1: "unordered"})
        assert int(num_id) > 0
        assert num.has_definitions

        an_elems, num_elems = num._get_elements()
        assert len(an_elems) == 1
        assert len(num_elems) == 1

        # abstractNum must have 9 levels
        lvls = an_elems[0].findall(f"{{{WML_NS}}}lvl")
        assert len(lvls) == 9

        # Level 0 → ordered (decimal), level 1 → unordered (bullet)
        lvl0_fmt = lvls[0].find(f"{{{WML_NS}}}numFmt")
        assert lvl0_fmt is not None, "Level 0 missing numFmt"
        assert lvl0_fmt.get(f"{{{WML_NS}}}val") == "decimal"
        lvl1_fmt = lvls[1].find(f"{{{WML_NS}}}numFmt")
        assert lvl1_fmt is not None, "Level 1 missing numFmt"
        assert lvl1_fmt.get(f"{{{WML_NS}}}val") == "bullet"

    def test_multiple_lists_get_unique_ids(self):
        """Each group receives a distinct numId."""
        num = DocxListNumbering()
        ids = [
            num.create_list_definition({0: "ordered"}),
            num.create_list_definition({0: "unordered"}),
            num.create_list_definition({0: "ordered"}),
        ]
        assert len(set(ids)) == 3

    def test_has_definitions_false_initially(self):
        """Fresh instance reports no definitions."""
        assert DocxListNumbering().has_definitions is False

    def test_multi_level_hybrid_multilevel(self):
        """Nested list definitions use hybridMultilevel."""
        num = DocxListNumbering()
        num.create_list_definition({0: "ordered", 1: "unordered", 2: "ordered"})
        an_elems, _ = num._get_elements()
        mlt = an_elems[0].find(f"{{{WML_NS}}}multiLevelType")
        assert mlt is not None
        assert mlt.get(f"{{{WML_NS}}}val") == "hybridMultilevel"

    def test_template_formats_are_stitched_newest_first_with_preset_fallback(self):
        from docx import Document

        document = Document()
        _append_template_abstract_num(
            document,
            20002,
            [
                {
                    "level": "0",
                    "tentative": "1",
                    "start": "9",
                    "numFmt": "upperLetter",
                    "lvlText": "%1-",
                }
            ],
        )
        _append_template_abstract_num(
            document,
            20001,
            [
                {
                    "level": "0",
                    "start": "3",
                    "numFmt": "lowerRoman",
                    "lvlText": "%1)",
                    "lvlJc": "right",
                    "left": "999",
                    "hanging": "222",
                }
            ],
        )
        _append_template_abstract_num(
            document,
            20000,
            [
                {
                    "level": "1",
                    "start": "1",
                    "numFmt": "bullet",
                    "lvlText": "◆",
                    "lvlJc": "center",
                    "left": "1333",
                    "hanging": "333",
                    "ascii": "Wingdings",
                    "hAnsi": "Symbol",
                }
            ],
        )

        numbering = DocxListNumbering(document)
        numbering.create_list_definition({0: "ordered", 1: "unordered", 2: "ordered"})
        abstract_num = numbering._get_elements()[0][0]
        levels = abstract_num.findall(f"{{{WML_NS}}}lvl")

        def value(level: etree._Element, tag: str) -> str | None:
            element = level.find(f"{{{WML_NS}}}{tag}")
            return element.get(f"{{{WML_NS}}}val") if element is not None else None

        assert value(levels[0], "start") == "3"
        assert value(levels[0], "numFmt") == "lowerRoman"
        assert value(levels[0], "lvlText") == "%1)"
        assert value(levels[0], "lvlJc") == "right"
        ordered_indent = levels[0].find(f"{{{WML_NS}}}pPr/{{{WML_NS}}}ind")
        assert ordered_indent is not None
        assert ordered_indent.get(f"{{{WML_NS}}}left") == "999"
        assert ordered_indent.get(f"{{{WML_NS}}}hanging") == "222"

        assert value(levels[1], "numFmt") == "bullet"
        assert value(levels[1], "lvlText") == "◆"
        bullet_fonts = levels[1].find(f"{{{WML_NS}}}rPr/{{{WML_NS}}}rFonts")
        assert bullet_fonts is not None
        assert bullet_fonts.get(f"{{{WML_NS}}}ascii") == "Wingdings"
        assert bullet_fonts.get(f"{{{WML_NS}}}hAnsi") == "Symbol"

        assert value(levels[2], "numFmt") == "decimal"
        assert value(levels[2], "lvlText") == "%3."

    def test_template_ids_are_avoided_and_schema_order_is_preserved(self, tmp_path: Path) -> None:
        from docx import Document

        document = Document()
        numbering_root = document.part.numbering_part.element
        existing_abstract = copy.deepcopy(numbering_root.findall(f"{{{WML_NS}}}abstractNum")[0])
        existing_abstract.set(f"{{{WML_NS}}}abstractNumId", "10000")
        first_num = numbering_root.find(f"{{{WML_NS}}}num")
        assert first_num is not None
        first_num.addprevious(existing_abstract)
        existing_num = copy.deepcopy(numbering_root.findall(f"{{{WML_NS}}}num")[0])
        existing_num.set(f"{{{WML_NS}}}numId", "10000")
        existing_ref = existing_num.find(f"{{{WML_NS}}}abstractNumId")
        assert existing_ref is not None
        existing_ref.set(f"{{{WML_NS}}}val", "10000")
        numbering_root.append(existing_num)

        numbering = DocxListNumbering(document)
        generated_num_id = numbering.create_list_definition({0: "ordered"})
        assert int(generated_num_id) > 10000
        apply_list_to_paragraph(document.add_paragraph("Generated item"), generated_num_id, 0)
        output = tmp_path / "numbering-id-order.docx"
        document.save(str(output))
        write_numbering_to_docx(str(output), numbering)

        result_root = _get_numbering_xml(str(output))
        assert result_root is not None
        abstract_ids = [
            element.get(f"{{{WML_NS}}}abstractNumId") for element in result_root.findall(f"{{{WML_NS}}}abstractNum")
        ]
        num_ids = [element.get(f"{{{WML_NS}}}numId") for element in result_root.findall(f"{{{WML_NS}}}num")]
        assert len(abstract_ids) == len(set(abstract_ids))
        assert len(num_ids) == len(set(num_ids))
        seen_num = False
        for child in result_root:
            local_name = etree.QName(child).localname
            if local_name == "num":
                seen_num = True
            elif local_name == "abstractNum":
                assert seen_num is False

    def test_list_nsid_is_stable_for_the_same_abstract_id(self) -> None:
        numbering = DocxListNumbering()
        numbering.create_list_definition({0: "ordered"})
        abstract_elements, _ = numbering._get_elements()
        nsid = abstract_elements[0].find(f"{{{WML_NS}}}nsid")
        assert nsid is not None
        expected = hashlib.sha256(b"docwen-list:10000").hexdigest()[:8].upper()
        assert nsid.get(f"{{{WML_NS}}}val") == expected


class TestApplyListToParagraph:
    """Tests for the per-paragraph numPr helper."""

    def test_adds_num_pr(self):
        """numPr is inserted into the paragraph's pPr with correct values."""
        from docx import Document

        doc = Document()
        p = doc.add_paragraph("Test")
        apply_list_to_paragraph(p, "42", 2)

        numPr = p._p.find(f"{{{WML_NS}}}pPr/{{{WML_NS}}}numPr")
        assert numPr is not None

        ilvl = numPr.find(f"{{{WML_NS}}}ilvl")
        assert ilvl is not None, "Missing w:ilvl in numPr"
        assert ilvl.get(f"{{{WML_NS}}}val") == "2"

        numId = numPr.find(f"{{{WML_NS}}}numId")
        assert numId is not None, "Missing w:numId in numPr"
        assert numId.get(f"{{{WML_NS}}}val") == "42"

    def test_replaces_existing_num_pr(self):
        """Calling apply_list_to_paragraph replaces any prior numPr."""
        from docx import Document

        doc = Document()
        p = doc.add_paragraph("Test")
        apply_list_to_paragraph(p, "10", 0)
        apply_list_to_paragraph(p, "20", 1)

        numPr = p._p.find(f"{{{WML_NS}}}pPr/{{{WML_NS}}}numPr")
        assert numPr is not None
        numId_elem = numPr.find(f"{{{WML_NS}}}numId")
        assert numId_elem is not None, "Missing numId"
        assert numId_elem.get(f"{{{WML_NS}}}val") == "20"
        ilvl_elem = numPr.find(f"{{{WML_NS}}}ilvl")
        assert ilvl_elem is not None, "Missing ilvl"
        assert ilvl_elem.get(f"{{{WML_NS}}}val") == "1"


class TestMdToDocxOrderedList:
    """End-to-end tests for ordered Markdown list → DOCX."""

    def test_ordered_list_has_num_pr(self):
        """Generated paragraphs carry ``w:numPr`` for ordered lists."""
        md = "1. First\n2. Second\n3. Third\n"
        md_path = write_temp_md(md)
        ctx, _ = make_context(md_path, target_format="docx")

        result = MdToDocxConverter().convert(ctx)
        assert result.success

        output = Path(result.artifacts[0].staging_path)
        list_paras = _list_paras_with_numpr(str(output))
        assert len(list_paras) >= 3, f"Expected >= 3 list paragraphs, got {len(list_paras)}"

        # All should share the same numId
        num_ids = set()
        for p in list_paras:
            numPr = p.find(f"{{{WML_NS}}}pPr/{{{WML_NS}}}numPr")
            assert numPr is not None, "numPr not found in list paragraph"
            numId = numPr.find(f"{{{WML_NS}}}numId")
            assert numId is not None, "numId not found in numPr"
            num_ids.add(numId.get(f"{{{WML_NS}}}val"))
        assert len(num_ids) == 1, "All items in one list group should share the same numId"

    def test_ordered_list_numbering_xml(self):
        """numbering.xml contains abstractNum + num for ordered lists."""
        md = "1. Alpha\n2. Beta\n"
        md_path = write_temp_md(md)
        ctx, _ = make_context(md_path, target_format="docx")

        result = MdToDocxConverter().convert(ctx)
        assert result.success

        num_root = _get_numbering_xml(str(Path(result.artifacts[0].staging_path)))
        assert num_root is not None, "numbering.xml missing"

        abstract_nums = num_root.findall(f"{{{WML_NS}}}abstractNum")
        nums = num_root.findall(f"{{{WML_NS}}}num")
        assert len(abstract_nums) >= 1
        assert len(nums) >= 1

        # Level 0 of our definition should be decimal for an ordered list
        our_an_id = _our_abstract_num_id(num_root)
        lvl0 = num_root.find(f'{{{WML_NS}}}abstractNum[@{{{WML_NS}}}abstractNumId="{our_an_id}"]/{{{WML_NS}}}lvl')
        assert lvl0 is not None
        numFmt = lvl0.find(f"{{{WML_NS}}}numFmt")
        assert numFmt is not None, "Level 0 missing numFmt in numbering.xml"
        assert numFmt.get(f"{{{WML_NS}}}val") == "decimal"

    def test_ordered_list_preserves_text(self):
        """List item text is preserved in the DOCX output."""
        md = "1. First item\n2. Second item\n"
        md_path = write_temp_md(md)
        ctx, _ = make_context(md_path, target_format="docx")

        result = MdToDocxConverter().convert(ctx)
        assert result.success

        from docx import Document

        doc = Document(str(Path(result.artifacts[0].staging_path)))
        full = "\n".join(p.text for p in doc.paragraphs)
        assert "First item" in full
        assert "Second item" in full


class TestMdToDocxUnorderedList:
    """End-to-end tests for unordered Markdown list → DOCX."""

    def test_unordered_list_bullet_numbering(self):
        """Unordered list items use bullet numFmt."""
        md = "- Item A\n- Item B\n- Item C\n"
        md_path = write_temp_md(md)
        ctx, _ = make_context(md_path, target_format="docx")

        result = MdToDocxConverter().convert(ctx)
        assert result.success

        num_root = _get_numbering_xml(str(Path(result.artifacts[0].staging_path)))
        assert num_root is not None

        # Locate the abstractNum that our num element references
        our_an_id = _our_abstract_num_id(num_root)
        lvl0 = num_root.find(f'{{{WML_NS}}}abstractNum[@{{{WML_NS}}}abstractNumId="{our_an_id}"]/{{{WML_NS}}}lvl')
        assert lvl0 is not None, f"No lvl found for abstractNumId={our_an_id}"
        numFmt = lvl0.find(f"{{{WML_NS}}}numFmt")
        assert numFmt is not None, "Level 0 missing numFmt (unordered)"
        assert numFmt.get(f"{{{WML_NS}}}val") == "bullet"

    def test_unordered_list_has_num_pr(self):
        """Unordered paragraphs also carry w:numPr."""
        md = "- Uno\n- Dos\n"
        md_path = write_temp_md(md)
        ctx, _ = make_context(md_path, target_format="docx")

        result = MdToDocxConverter().convert(ctx)
        assert result.success

        list_paras = _list_paras_with_numpr(str(Path(result.artifacts[0].staging_path)))
        assert len(list_paras) >= 2


class TestMdToDocxNestedList:
    """End-to-end tests for nested list levels."""

    def test_nested_levels_have_increasing_ilvl(self):
        """Deeper nesting produces higher ``w:ilvl`` values."""
        md = "1. Top\n   1. Nested A\n       - Deep bullet\n   2. Nested B\n2. Top 2\n"
        md_path = write_temp_md(md)
        ctx, _ = make_context(md_path, target_format="docx")

        result = MdToDocxConverter().convert(ctx)
        assert result.success

        output = str(Path(result.artifacts[0].staging_path))
        list_paras = _list_paras_with_numpr(output)

        ilvls: set[int] = set()
        for p in list_paras:
            ilvl_el = p.find(f"{{{WML_NS}}}pPr/{{{WML_NS}}}numPr/{{{WML_NS}}}ilvl")
            if ilvl_el is not None:
                ilvls.add(int(ilvl_el.get(f"{{{WML_NS}}}val", "0")))

        # We expect at least levels 0, 1, 2
        for expected in (0, 1, 2):
            assert expected in ilvls, f"Expected ilvl {expected} in nested list, got {sorted(ilvls)}"

    def test_nested_list_text_preserved(self):
        """Text from deeply nested items survives conversion."""
        md = "- L0\n  - L1\n    - L2\n"
        md_path = write_temp_md(md)
        ctx, _ = make_context(md_path, target_format="docx")

        result = MdToDocxConverter().convert(ctx)
        assert result.success

        from docx import Document

        doc = Document(str(Path(result.artifacts[0].staging_path)))
        full = "\n".join(p.text for p in doc.paragraphs)
        for item in ("L0", "L1", "L2"):
            assert item in full, f"'{item}' not found in output"

    def test_nested_ordered_in_unordered(self):
        """Ordered sub-list inside an unordered list works correctly."""
        md = "- Top\n  1. First numbered\n  2. Second numbered\n- Another\n"
        md_path = write_temp_md(md)
        ctx, _ = make_context(md_path, target_format="docx")

        result = MdToDocxConverter().convert(ctx)
        assert result.success

        num_root = _get_numbering_xml(str(Path(result.artifacts[0].staging_path)))
        assert num_root is not None

        # Locate our definition
        our_an_id = _our_abstract_num_id(num_root)
        # Level 0 → bullet, Level 1 → decimal
        lvl0 = num_root.find(
            f'{{{WML_NS}}}abstractNum[@{{{WML_NS}}}abstractNumId="{our_an_id}"]/{{{WML_NS}}}lvl[@{{{WML_NS}}}ilvl="0"]'
        )
        lvl1 = num_root.find(
            f'{{{WML_NS}}}abstractNum[@{{{WML_NS}}}abstractNumId="{our_an_id}"]/{{{WML_NS}}}lvl[@{{{WML_NS}}}ilvl="1"]'
        )
        assert lvl0 is not None, "Level 0 lvl not found"
        assert lvl1 is not None, "Level 1 lvl not found"
        lvl0_fmt = lvl0.find(f"{{{WML_NS}}}numFmt")
        assert lvl0_fmt is not None, "Level 0 missing numFmt"
        assert lvl0_fmt.get(f"{{{WML_NS}}}val") == "bullet"
        lvl1_fmt = lvl1.find(f"{{{WML_NS}}}numFmt")
        assert lvl1_fmt is not None, "Level 1 missing numFmt"
        assert lvl1_fmt.get(f"{{{WML_NS}}}val") == "decimal"

    def test_same_depth_mixed_list_conflict_deterministically_prefers_ordered(self) -> None:
        """F-F1-018: ordered numbering wins when one list tree mixes types at a depth."""
        md = "- Parent A\n  - bullet child\n- Parent B\n  1. ordered child\n"
        md_path = write_temp_md(md)
        ctx, _ = make_context(md_path, target_format="docx")

        result = MdToDocxConverter().convert(ctx)

        assert result.success
        num_root = _get_numbering_xml(str(Path(result.artifacts[0].staging_path)))
        assert num_root is not None
        abstract_num_id = _our_abstract_num_id(num_root)
        level_one = num_root.find(
            f'{{{WML_NS}}}abstractNum[@{{{WML_NS}}}abstractNumId="{abstract_num_id}"]/'
            f'{{{WML_NS}}}lvl[@{{{WML_NS}}}ilvl="1"]'
        )
        assert level_one is not None
        number_format = level_one.find(f"{{{WML_NS}}}numFmt")
        assert number_format is not None
        assert number_format.get(f"{{{WML_NS}}}val") == "decimal"


class TestMdToDocxMultipleLists:
    """Separated lists get independent numbering definitions."""

    def test_separated_lists_different_num_ids(self):
        """Two lists separated by a paragraph receive different numIds."""
        md = "1. First\n2. Second\n\nSome paragraph.\n\n- Bullet A\n- Bullet B\n"
        md_path = write_temp_md(md)
        ctx, _ = make_context(md_path, target_format="docx")

        result = MdToDocxConverter().convert(ctx)
        assert result.success

        list_paras = _list_paras_with_numpr(str(Path(result.artifacts[0].staging_path)))
        num_ids = set()
        for p in list_paras:
            numId_elem = p.find(f"{{{WML_NS}}}pPr/{{{WML_NS}}}numPr/{{{WML_NS}}}numId")
            assert numId_elem is not None, "numId not found in list paragraph"
            num_ids.add(numId_elem.get(f"{{{WML_NS}}}val"))

        # Two list groups → two numIds
        assert len(num_ids) == 2, f"Expected 2 unique numIds for 2 list groups, got {len(num_ids)}"

    def test_no_numbering_when_no_lists(self):
        """A doc without lists should still convert (no numbering regression)."""
        md = "# Title\n\nJust a paragraph.\n\nAnother.\n"
        md_path = write_temp_md(md)
        ctx, _ = make_context(md_path, target_format="docx")

        result = MdToDocxConverter().convert(ctx)
        assert result.success

        from docx import Document

        doc = Document(str(Path(result.artifacts[0].staging_path)))
        texts = [p.text for p in doc.paragraphs]
        assert any("Title" in t for t in texts)
        assert any("paragraph" in t for t in texts)
