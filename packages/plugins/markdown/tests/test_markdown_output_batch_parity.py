"""Broader Markdown output parity guards derived from the VIS-099 batch."""

from __future__ import annotations

import csv
import json
import re
import zipfile
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from openpyxl import load_workbook
from tests.support.config import FakeConfigView

pytestmark = pytest.mark.contract

from docwen_plugin_markdown.template_utils import extract_body_font, find_body_placeholder
from docwen_plugin_markdown.to_docx.converter import (
    MdToDocxConverter,
    _resolve_quote_style_levels,
    _resolve_template_style_keys,
)
from docwen_plugin_markdown.to_spreadsheet.converter import (
    MdToCsvConverter,
    MdToXlsxConverter,
)
from docwen_plugin_markdown.yaml_processor import (
    BODY_PLACEHOLDER_ALIASES,
    TITLE_PLACEHOLDER_ALIASES,
)

from .conftest import PROJECT_ROOT, make_context

FIXTURE_PATH = PROJECT_ROOT / "tests" / "fixtures" / "golden" / "old_system_markdown_output_batch_semantics.json"
FIXTURE: dict[str, Any] = json.loads(FIXTURE_PATH.read_text(encoding="utf-8"))
DOCX_TEMPLATES = tuple(FIXTURE["template_contract"]["docx_templates"])


def _visible_text(document: DocumentObject) -> list[str]:
    return [paragraph.text for paragraph in document.paragraphs if paragraph.text]


def _paragraph_text_including_sdt(paragraph: Any) -> str:
    """Return paragraph text including inline SDT-carrier content.

    Fenced-source carriers wrap visible payload runs in one inline SDT, which
    python-docx ``Paragraph.text`` does not traverse.  Recovery and other
    authenticated readers use the same full-w:t projection.
    """

    return "".join(item.text or "" for item in paragraph._p.iter(f"{qn('w:t')}"))


def test_bundled_template_placeholder_contract_covers_all_distributed_docx_templates() -> None:
    placeholder_pattern = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")
    for template_name in DOCX_TEMPLATES:
        document = Document(PROJECT_ROOT / "templates" / template_name)
        placeholders = [
            match.group(1).strip()
            for paragraph in document.paragraphs
            for match in placeholder_pattern.finditer(paragraph.text)
        ]
        assert len(placeholders) == 2, template_name
        assert placeholders[0] in TITLE_PLACEHOLDER_ALIASES, template_name
        assert placeholders[1] in BODY_PLACEHOLDER_ALIASES, template_name


@pytest.mark.parametrize("template_name", DOCX_TEMPLATES)
def test_bundled_multilingual_docx_templates_restore_alias_title_and_body_position(
    tmp_path: Path,
    template_name: str,
) -> None:
    source = tmp_path / "multilingual-fallback.md"
    source.write_text(
        "---\naliases:\n  - Fallback Title\n---\n\n# Body Heading\n\nBody text.\n",
        encoding="utf-8",
    )
    context, _workspace = make_context(
        str(source),
        target_format="docx",
        options={"template_name": str(PROJECT_ROOT / "templates" / template_name)},
    )

    result = MdToDocxConverter().convert(context)

    assert result.success is True
    document = Document(result.artifacts[0].staging_path)
    visible = _visible_text(document)
    assert visible[:3] == ["Fallback Title", "Body Heading", "Body text."], template_name
    assert all("{{" not in text and "}}" not in text for text in visible), template_name
    assert document.core_properties.title == "Fallback Title", template_name


def test_heading_merge_restores_one_mixed_style_word_paragraph(tmp_path: Path) -> None:
    source = tmp_path / "heading-merge.md"
    source.write_text(
        "## 言之有谋，强化顶层设计：\n坚持全局眼光，加强前瞻性思考、全局性谋划。\n\n## Later\nTail.\n",
        encoding="utf-8",
    )
    context, _workspace = make_context(str(source), target_format="docx")

    result = MdToDocxConverter().convert(context)

    assert result.success is True
    document = Document(result.artifacts[0].staging_path)
    visible_paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text]
    assert [paragraph.text for paragraph in visible_paragraphs] == [
        "言之有谋，强化顶层设计：坚持全局眼光，加强前瞻性思考、全局性谋划。",
        "Later",
        "Tail.",
    ]
    mixed_paragraph = visible_paragraphs[0]
    assert mixed_paragraph.style is not None
    assert mixed_paragraph.style.name == "Heading 2"
    assert len(mixed_paragraph.runs) >= 2
    assert mixed_paragraph.runs[-1].text == "坚持全局眼光，加强前瞻性思考、全局性谋划。"
    assert mixed_paragraph.runs[-1].font.name == "Calibri"
    assert mixed_paragraph.runs[-1].font.size is not None
    assert mixed_paragraph.runs[-1].font.size.pt == pytest.approx(10.5)
    mixed_body_fonts = mixed_paragraph.runs[-1]._r.get_or_add_rPr().get_or_add_rFonts()
    assert mixed_body_fonts.get(qn("w:eastAsia")) == "宋体"
    assert mixed_paragraph.runs[-1].bold is False


def test_halfwidth_heading_merge_inserts_one_readable_space(tmp_path: Path) -> None:
    source = tmp_path / "halfwidth-heading-merge.md"
    source.write_text("# Planning:\nBody text.\n", encoding="utf-8")
    context, _workspace = make_context(str(source), target_format="docx")

    result = MdToDocxConverter().convert(context)

    assert result.success is True
    document = Document(result.artifacts[0].staging_path)
    assert [paragraph.text for paragraph in document.paragraphs if paragraph.text] == ["Planning: Body text."]


def test_docx_core_metadata_preserves_declared_dcterms_qname_prefix(tmp_path: Path) -> None:
    source = tmp_path / "metadata-prefix.md"
    source.write_text(
        "---\naliases:\n  - Metadata Title\nsubtitle: Metadata Subject\n---\n# Body\n\nText.\n",
        encoding="utf-8",
    )
    template = PROJECT_ROOT / FIXTURE["templates"]["docx"]["path"]
    context, _workspace = make_context(
        str(source),
        target_format="docx",
        options={"template_name": str(template)},
    )

    result = MdToDocxConverter().convert(context)

    assert result.success is True
    output_path = Path(result.artifacts[0].staging_path)
    document = Document(str(output_path))
    assert document.core_properties.title == "Metadata Title"
    assert document.core_properties.subject == "Metadata Subject"
    with zipfile.ZipFile(output_path) as archive:
        assert archive.namelist().count("docProps/core.xml") == 1
        core_xml = archive.read("docProps/core.xml")
    assert b'xmlns:dcterms="http://purl.org/dc/terms/"' in core_xml
    assert b'xsi:type="dcterms:W3CDTF"' in core_xml
    assert b'xmlns:ns1="http://purl.org/dc/terms/"' not in core_xml


def test_quote_style_config_maps_depth_to_semantic_template_style_level() -> None:
    config = FakeConfigView(
        {
            "document": {
                "style": {
                    "quote": {
                        "md_to_docx": {
                            "level_1_style_key": "quote_3",
                            "level_2_style_key": "invalid-localized-name",
                        }
                    }
                }
            }
        }
    )

    levels = _resolve_quote_style_levels(config)

    assert levels[1] == 3
    assert levels[2] == 2
    assert levels[9] == 9


def test_code_and_formula_style_config_consumes_stable_semantic_keys() -> None:
    config = FakeConfigView(
        {
            "document": {
                "style": {
                    "code": {
                        "md_to_docx": {
                            "inline_code_style_key": "inline_code",
                            "code_block_style_key": "invalid-localized-name",
                        }
                    },
                    "formula": {
                        "md_to_docx": {
                            "inline_formula_style_key": "inline_formula",
                            "formula_block_style_key": "formula_block",
                        }
                    },
                }
            }
        }
    )

    keys = _resolve_template_style_keys(config)

    assert keys == {
        "inline_code": "inline_code",
        "code_block": "invalid-localized-name",
        "inline_formula": "inline_formula",
        "formula_block": "formula_block",
    }


@pytest.mark.parametrize("template_name", DOCX_TEMPLATES)
def test_bundled_templates_apply_localized_multilevel_quote_styles_without_duplicate_indent(
    tmp_path: Path,
    template_name: str,
) -> None:
    source = tmp_path / "nested-quotes.md"
    source.write_text(
        "> Level 1 quote.\n>> Level 2 quote.\n>>> Level 3 quote.\n",
        encoding="utf-8",
    )
    context, _workspace = make_context(
        str(source),
        target_format="docx",
        options={"template_name": str(PROJECT_ROOT / "templates" / template_name)},
    )

    result = MdToDocxConverter().convert(context)

    assert result.success is True
    document = Document(result.artifacts[0].staging_path)
    quote_paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.endswith("quote.")]
    assert [paragraph.text for paragraph in quote_paragraphs] == [
        "Level 1 quote.",
        "Level 2 quote.",
        "Level 3 quote.",
    ]
    for level, paragraph in enumerate(quote_paragraphs, start=1):
        direct_p_pr = paragraph._p.find(qn("w:pPr"))
        assert direct_p_pr is not None, template_name
        assert direct_p_pr.findall(qn("w:ind")) == [], template_name

        style = paragraph.style
        assert style is not None, template_name
        style_element = style._element
        assert style_element is not None, template_name
        style_p_pr = style_element.find(qn("w:pPr"))
        assert style_p_pr is not None, template_name
        style_indent = style_p_pr.find(qn("w:ind"))
        style_shading = style_p_pr.find(qn("w:shd"))
        style_borders = style_p_pr.find(qn("w:pBdr"))
        assert style_indent is not None, template_name
        assert style_shading is not None, template_name
        assert style_borders is not None, template_name
        assert style_indent.get(qn("w:left")) == str(480 + (level - 1) * 240), template_name
        assert style_indent.get(qn("w:right")) == "480", template_name
        assert style_shading.get(qn("w:fill")) == "F5F5F5", template_name
        assert style_borders.find(qn("w:left")) is not None, template_name
        text_run = next(run for run in paragraph.runs if run.text.strip())
        assert text_run.font.name is None, template_name
        assert text_run.font.size is None, template_name


@pytest.mark.parametrize("template_name", DOCX_TEMPLATES)
def test_bundled_templates_drive_body_code_formula_list_and_table_output_styles(
    tmp_path: Path,
    template_name: str,
) -> None:
    source = tmp_path / "semantic-styles.md"
    source.write_text(
        "Body text with `inline code` and $x+1$.\n\n"
        "```text\nalpha\nbeta\n```\n\n"
        "$$y=x^2$$\n\n"
        "- List item\n\n"
        "| Header |\n| --- |\n| Cell |\n",
        encoding="utf-8",
    )
    template_path = PROJECT_ROOT / "templates" / template_name
    template = Document(str(template_path))
    placeholder = find_body_placeholder(template)
    assert placeholder is not None, template_name
    expected_body_font = extract_body_font(template)

    context, _workspace = make_context(
        str(source),
        target_format="docx",
        options={"template_name": str(template_path)},
    )
    result = MdToDocxConverter().convert(context)

    assert result.success is True
    document = Document(result.artifacts[0].staging_path)
    body = next(paragraph for paragraph in document.paragraphs if paragraph.text.startswith("Body text"))
    code = next(
        paragraph for paragraph in document.paragraphs if _paragraph_text_including_sdt(paragraph).startswith("alpha")
    )
    list_item = next(paragraph for paragraph in document.paragraphs if paragraph.text == "List item")
    block_formula = next(
        paragraph
        for paragraph in document.paragraphs
        if not paragraph.text and paragraph._p.find(f".//{qn('m:oMath')}") is not None
    )

    body_style = body.style
    assert body_style is not None, template_name
    assert body_style.style_id == "DocWenBodyParagraph", template_name
    assert body_style.font.name == expected_body_font["name"], template_name
    assert body_style.font.size == expected_body_font["size"], template_name
    body_style_element = body_style._element
    assert body_style_element is not None, template_name
    body_style_r_pr = body_style_element.find(qn("w:rPr"))
    assert body_style_r_pr is not None, template_name
    body_style_fonts = body_style_r_pr.find(qn("w:rFonts"))
    assert body_style_fonts is not None, template_name
    assert body_style_fonts.get(qn("w:eastAsia")) == expected_body_font["east_asia"], template_name

    inline_code_run = next(run for run in body.runs if run.text == "inline code")
    inline_code_style = inline_code_run.style
    assert inline_code_style is not None, template_name
    assert inline_code_style.style_id == "DocWenInlineCode", template_name
    inline_math_styles = {
        style.get(qn("w:val")) for style in body._p.findall(f".//{qn('m:r')}/{qn('w:rPr')}/{qn('w:rStyle')}")
    }
    assert inline_math_styles == {"DocWenInlineFormula"}, template_name

    code_style = code.style
    assert code_style is not None, template_name
    assert code_style.style_id == "DocWenCodeBlock", template_name
    assert code._p.get_or_add_pPr().findall(qn("w:shd")) == [], template_name
    block_formula_style = block_formula.style
    assert block_formula_style is not None, template_name
    assert block_formula_style.style_id == "DocWenFormulaBlock", template_name

    list_item_style = list_item.style
    assert list_item_style is not None, template_name
    assert list_item_style.style_id == "DocWenListBlock", template_name
    assert list_item._p.get_or_add_pPr().findall(qn("w:ind")) == [], template_name
    assert list_item.runs[0].font.name is None, template_name
    assert list_item.runs[0].font.size is None, template_name

    table = document.tables[0]
    table_style = table.style
    assert table_style is not None, template_name
    assert table_style.style_id == "DocWenThreeLineTable", template_name
    header = table.cell(0, 0).paragraphs[0]
    cell = table.cell(1, 0).paragraphs[0]
    header_style = header.style
    cell_style = cell.style
    assert header_style is not None, template_name
    assert cell_style is not None, template_name
    assert header_style.style_id == "DocWenTableHeader", template_name
    assert cell_style.style_id == "DocWenTableContent", template_name
    assert next(run for run in header.runs if run.text).bold is None, template_name


def test_comprehensive_sample_alias_fallback_reaches_xlsx_and_csv_template_chain(tmp_path: Path) -> None:
    source = PROJECT_ROOT / "samples" / "sample.md"
    template = PROJECT_ROOT / "templates" / "English Sample Sheet Template.xlsx"

    xlsx_context, _xlsx_workspace = make_context(
        str(source),
        target_format="xlsx",
        options={"template_name": str(template)},
    )
    xlsx_result = MdToXlsxConverter().convert(xlsx_context)
    assert xlsx_result.success is True
    workbook = load_workbook(xlsx_result.artifacts[0].staging_path)
    try:
        expected_cells = FIXTURE["xlsx_projection"]["cells"]
        assert workbook.sheetnames == FIXTURE["xlsx_projection"]["sheet_names"]
        assert workbook["Sheet1"]["B1"].value == expected_cells["B1"]
        assert workbook["Sheet1"]["B2"].value == expected_cells["B2"]
    finally:
        workbook.close()

    csv_context, _csv_workspace = make_context(
        str(source),
        target_format="csv",
        options={"template_name": str(template)},
    )
    csv_result = MdToCsvConverter().convert(csv_context)
    assert csv_result.success is True
    assert len(csv_result.artifacts) == 3
    with Path(csv_result.artifacts[0].staging_path).open(encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.reader(handle))
    assert rows == FIXTURE["csv_projection"]["sheet1_rows"]
    assert all(Path(artifact.staging_path).read_bytes() == b"" for artifact in csv_result.artifacts[1:])


def test_xlsx_template_title_falls_back_to_source_stem_without_aliases(tmp_path: Path) -> None:
    source = tmp_path / "filename-fallback.md"
    source.write_text("| Food | Price |\n| --- | --- |\n| Pear | 4 |\n", encoding="utf-8")
    template = PROJECT_ROOT / "templates" / "English Sample Sheet Template.xlsx"
    context, _workspace = make_context(
        str(source),
        target_format="xlsx",
        options={"template_name": str(template)},
    )

    result = MdToXlsxConverter().convert(context)

    assert result.success is True
    workbook = load_workbook(result.artifacts[0].staging_path)
    try:
        assert workbook["Sheet1"]["B1"].value == "filename-fallback"
    finally:
        workbook.close()
