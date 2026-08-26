"""Focused source tests for the document-semantics v1 first slice."""

from __future__ import annotations

from zipfile import ZipFile

import pytest
from docx import Document
from docx.oxml.ns import qn

from docwen_core.docx_bookmarks import build_docx_bookmark_inventory
from docwen_core.docx_parsing.document_semantics import (
    extract_semantic_caption,
    extract_semantic_table_metadata,
    render_semantic_reference_text,
)
from docwen_plugin_markdown.document_semantics import analyze_document_semantics
from docwen_plugin_markdown.mistune_extensions import parse_markdown_text
from docwen_plugin_markdown.renderer import MdToDocxRenderer

pytestmark = pytest.mark.contract


def _analyze(markdown: str):
    return analyze_document_semantics(parse_markdown_text(markdown))


def test_caption_target_reference_and_citation_are_structured() -> None:
    markdown = """Figure: Blue square {#fig-blue}

![Synthetic blue square](pixel.png)

See @fig-blue and preserve [@smith2025; @wang2024].
"""
    analysis = _analyze(markdown)

    assert not analysis.has_errors
    assert analysis.oracle_projection["schema"] == "docwen.document_semantics.v1"
    figure, paragraph = analysis.oracle_projection["blocks"]
    assert figure == {
        "type": "figure",
        "caption": {
            "kind": "figure",
            "content": [{"type": "text", "value": "Blue square"}],
            "source_form": "canonical",
        },
        "image": {"source": "pixel.png", "alt": "Synthetic blue square"},
        "target_id": "fig-blue",
    }
    assert [item["type"] for item in paragraph["content"]] == [
        "text",
        "cross_reference",
        "text",
        "citation",
        "text",
    ]
    assert [(item.level, item.code) for item in analysis.diagnostics] == [
        ("warning", "interop.citation.processor_unavailable")
    ]


def test_equation_and_listing_oracle_preserve_authored_payloads() -> None:
    analysis = _analyze(
        """Equation: Euler

$$x^2$$

Listing: Example

```python extra-info
print("x")
```
"""
    )

    assert not analysis.has_errors
    assert analysis.oracle_projection["blocks"] == [
        {
            "type": "equation",
            "caption": {
                "kind": "equation",
                "content": [{"type": "text", "value": "Euler"}],
                "source_form": "canonical",
            },
            "latex": "x^2",
        },
        {
            "type": "listing",
            "caption": {
                "kind": "listing",
                "content": [{"type": "text", "value": "Example"}],
                "source_form": "canonical",
            },
            "code": 'print("x")',
            "language": "python",
        },
    ]


def test_invalid_semantics_diagnostics_are_stable_and_ordered() -> None:
    markdown = """Figure: Wrong kind {#fig-wrong-kind}

| A | B |
|---|---|
| 1 | 2 |

: Broken binding {#fig-broken}

This paragraph breaks caption binding.

![Synthetic image](pixel.png)

Figure: First duplicate {#fig-repeat}

![Synthetic image](pixel.png)

Figure: Second duplicate {#fig-repeat}

![Synthetic image](pixel.png)

See @tbl-missing.

| A | < |
|---|---|
| ^ | value |
{header-rows=3 header-cols=0 repeat-header=true}
"""
    analysis = _analyze(markdown)

    assert analysis.has_errors
    assert [(item.level, item.code, item.message) for item in analysis.diagnostics] == [
        ("error", "interop.caption.kind_mismatch", "Figure caption is followed by a table."),
        (
            "warning",
            "interop.caption.binding_broken",
            "An intervening paragraph breaks shorthand caption binding.",
        ),
        ("error", "interop.target.duplicate", "Target fig-repeat occurs more than once in the document."),
        ("error", "interop.reference.missing", "Target tbl-missing does not exist."),
        (
            "error",
            "interop.table.merge_non_rectangular",
            "Merge markers do not cover a complete rectangle.",
        ),
        ("error", "interop.table.attribute_invalid", "header-rows exceeds the table row count."),
    ]


def test_renderer_writes_native_table_fields_and_helpers_recover_them(tmp_path) -> None:
    markdown = """: Sales channels {#tbl-sales}

| Region | Sales | < | Total |
|---|---:|---:|---:|
| ^ | Online | Retail | ^ |
| North | 10 | 12 | 22 |
| South | 8 | 9 | 17 |
{header-rows=2 header-cols=1 repeat-header=true}

See @tbl-sales.
"""
    analysis = _analyze(markdown)
    document = Document()
    renderer = MdToDocxRenderer(document)
    renderer.render(analysis.ast)
    output = tmp_path / "semantic-table.docx"
    document.save(output)

    with ZipFile(output) as package:
        xml = package.read("word/document.xml").decode("utf-8")
    assert "SEQ Table" in xml
    assert "REF _DW_" in xml
    assert "w:gridSpan" in xml
    assert "w:vMerge" in xml
    assert xml.count("w:tblHeader") == 2

    loaded = Document(output)
    inventory = build_docx_bookmark_inventory(loaded)
    caption = next(
        extracted
        for paragraph in loaded.paragraphs
        if (extracted := extract_semantic_caption(paragraph._p, bookmark_inventory=inventory)) is not None
    )
    assert caption.target_id == "tbl-sales"
    assert caption.source_form == "shorthand"
    assert extract_semantic_table_metadata(loaded.tables[0]._tbl).repeat_header == "always"
    reference = next(
        extracted
        for paragraph in loaded.paragraphs
        if (extracted := render_semantic_reference_text(paragraph._p, bookmark_inventory=inventory)) is not None
    )
    assert reference == "See @tbl-sales."


def test_legacy_caption_extractor_ignores_unbalanced_target_bookmark() -> None:
    markdown = """Table: Safe failure {#tbl-safe}

| A |
|---|
| value |
"""
    analysis = _analyze(markdown)
    document = Document()
    MdToDocxRenderer(document).render(analysis.ast)
    caption = next(paragraph for paragraph in document.paragraphs if "SEQ Table" in paragraph._p.xml)
    target_start = next(
        item for item in caption._p.iter(qn("w:bookmarkStart")) if (item.get(qn("w:name")) or "").startswith("_DW_")
    )
    target_id = target_start.get(qn("w:id"))
    target_end = next(item for item in caption._p.iter(qn("w:bookmarkEnd")) if item.get(qn("w:id")) == target_id)
    target_parent = target_end.getparent()
    assert target_parent is not None
    target_parent.remove(target_end)
    inventory = build_docx_bookmark_inventory(document)

    assert extract_semantic_caption(caption._p, bookmark_inventory=inventory) is None
