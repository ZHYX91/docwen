"""Tests for Markdown→DOCX formula rendering adapter.

Covers F-F1-008 (inline formula in mixed paragraphs via AST),
F-F1-034 (formula block OMML).
"""

from __future__ import annotations

import pytest
from docx import Document

pytestmark = pytest.mark.contract


def _docx_xml(doc) -> str:
    return doc._element.xml


class TestFormulaRenderer:
    """Unit tests for standalone formula renderer functions."""

    def test_block_formula_creates_omath(self):
        from docwen_plugin_markdown.to_docx.formula_renderer import (
            render_block_formula,
        )

        doc = Document()
        render_block_formula(doc, r"E=mc^2")
        assert "oMath" in _docx_xml(doc)

    def test_inline_formula_embeds_omath_in_paragraph(self):
        from docwen_plugin_markdown.to_docx.formula_renderer import (
            render_inline_formula,
        )

        doc = Document()
        p = doc.add_paragraph("Before ")
        render_inline_formula(p, r"x_1")
        p.add_run(" after")
        xml = _docx_xml(doc)
        assert "oMath" in xml
        assert "Before" in xml

    def test_empty_formula_noops(self):
        from docwen_plugin_markdown.to_docx.formula_renderer import (
            render_block_formula,
        )

        doc = Document()
        render_block_formula(doc, "")
        assert "oMath" not in _docx_xml(doc)

    def test_block_formula_centered_paragraph(self):
        """Block formulas are placed in a centered paragraph (F-F1-034)."""
        from docwen_plugin_markdown.to_docx.formula_renderer import (
            render_block_formula,
        )

        doc = Document()
        render_block_formula(doc, r"\alpha + \beta")
        xml = _docx_xml(doc)
        assert "oMath" in xml
        # Centered alignment should be present (WD_ALIGN_PARAGRAPH.CENTER = 1)
        assert "jc" in xml

    def test_inline_formula_fallback_for_invalid_latex(self):
        """Invalid LaTeX is rendered as plain text rather than crashing."""
        from docwen_plugin_markdown.to_docx.formula_renderer import (
            render_inline_formula,
        )

        doc = Document()
        p = doc.add_paragraph()
        # Intentionally malformed LaTeX
        render_inline_formula(p, r"\invalid{")
        _docx_xml(doc)
        # Either we get OMML or plain text; should not crash
        assert len(doc.paragraphs) >= 1


class TestFormulaViaFullPipeline:
    """Formula rendering exercised through the full MD→DOCX pipeline.

    These tests verify that inline formulas within mixed paragraphs are
    not silently dropped (F-F1-008).
    """

    def test_inline_formula_in_paragraph_preserved(self):
        """Inline formula inside a paragraph is rendered to OMML."""
        from docwen_plugin_markdown.mistune_extensions import parse_markdown_text
        from docwen_plugin_markdown.renderer import MdToDocxRenderer

        md = "Some text with $x^2 + y^2$ formula inline."
        ast = parse_markdown_text(md)
        doc = Document()
        renderer = MdToDocxRenderer(doc)
        renderer.render(ast)
        xml = _docx_xml(doc)
        assert "Some text with" in xml
        assert "oMath" in xml, "Inline formula $x^2+y^2$ was dropped — F-F1-008 regression"
        assert "formula inline" in xml

    def test_multiple_inline_formulas_same_paragraph(self):
        """Multiple inline formulas in one paragraph all render."""
        from docwen_plugin_markdown.mistune_extensions import parse_markdown_text
        from docwen_plugin_markdown.renderer import MdToDocxRenderer

        md = "First $a+b$ then $c+d$ in one paragraph."
        ast = parse_markdown_text(md)
        doc = Document()
        renderer = MdToDocxRenderer(doc)
        renderer.render(ast)
        xml = _docx_xml(doc)
        # Should have two oMath elements
        count = xml.count("<m:oMath")
        assert count >= 2, f"Expected >=2 oMath elements, got {count} — F-F1-008: multiple inline formulas dropped"

    def test_inline_formula_with_formatting(self):
        """Inline formula alongside bold/italic text (mixed paragraph)."""
        from docwen_plugin_markdown.mistune_extensions import parse_markdown_text
        from docwen_plugin_markdown.renderer import MdToDocxRenderer

        md = "**Bold** and $E=mc^2$ and *italic*."
        ast = parse_markdown_text(md)
        doc = Document()
        renderer = MdToDocxRenderer(doc)
        renderer.render(ast)
        xml = _docx_xml(doc)
        assert "Bold" in xml
        assert "italic" in xml
        assert "oMath" in xml, "Formula $E=mc^2$ dropped when mixed with bold/italic"

    def test_block_formula_via_pipeline(self):
        """Block LaTeX formula $$...$$ renders to OMML."""
        from docwen_plugin_markdown.mistune_extensions import parse_markdown_text
        from docwen_plugin_markdown.renderer import MdToDocxRenderer

        md = "Before\n\n$$\nE = mc^2\n$$\n\nAfter"
        ast = parse_markdown_text(md)
        doc = Document()
        renderer = MdToDocxRenderer(doc)
        renderer.render(ast)
        xml = _docx_xml(doc)
        assert "oMath" in xml, "Block formula $$E=mc^2$$ was dropped"

    def test_single_line_block_formula_is_display_math_without_literal_dollars(self):
        """Standalone ``$$...$$`` matches the reference display-math syntax."""
        from docwen_plugin_markdown.mistune_extensions import parse_markdown_text
        from docwen_plugin_markdown.renderer import MdToDocxRenderer

        ast = parse_markdown_text(r"$$x = \frac{-b \pm \sqrt{b^2 - 4ac}}{2a}$$")
        assert ast == [
            {
                "type": "block_math",
                "raw": r"x = \frac{-b \pm \sqrt{b^2 - 4ac}}{2a}",
            }
        ]

        doc = Document()
        MdToDocxRenderer(doc).render(ast)
        xml = _docx_xml(doc)
        assert "oMath" in xml
        assert "$" not in xml

    def test_single_line_block_formula_syntax_inside_fence_remains_code(self):
        """The block rule must not rewrite literal display syntax in code."""
        from docwen_plugin_markdown.mistune_extensions import parse_markdown_text

        ast = parse_markdown_text("```text\n$$x+y$$\n```\n")

        assert ast[0]["type"] == "block_code"
        assert ast[0]["raw"] == "$$x+y$$\n"

    def test_inline_latex_preserved(self):
        """Inline LaTeX $...$ from mistune math plugin produces OMML."""
        from docwen_plugin_markdown.mistune_extensions import parse_markdown_text
        from docwen_plugin_markdown.renderer import MdToDocxRenderer

        md = "Value is $\\frac{1}{2}$ of total."
        ast = parse_markdown_text(md)
        doc = Document()
        renderer = MdToDocxRenderer(doc)
        renderer.render(ast)
        xml = _docx_xml(doc)
        assert "oMath" in xml, "Inline LaTeX $\\frac{1}{2}$ was dropped — F-F1-008 regression"
        assert "of total" in xml
