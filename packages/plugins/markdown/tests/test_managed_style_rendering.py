"""Renderer contracts for request-owned managed DOCX styles."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.styles.style import CharacterStyle, ParagraphStyle

from docwen_plugin_markdown.document_semantics import analyze_document_semantics
from docwen_plugin_markdown.mistune_extensions import parse_markdown_text
from docwen_plugin_markdown.renderer import MdToDocxRenderer
from docwen_plugin_markdown.renderer_inlines import (
    add_hyperlink,
    project_character_style_properties,
)
from docwen_plugin_markdown.template_filler import fill_template
from docwen_plugin_markdown.template_utils import BodyParagraphFormat, extract_body_paragraph_format
from docwen_plugin_markdown.to_docx.managed_styles import (
    complete_managed_styles,
    validate_managed_style_package,
)
from docwen_runtime.config.document_styles import build_document_style_catalog

from .conftest import PROJECT_ROOT

pytestmark = pytest.mark.contract


def _catalog():
    return build_document_style_catalog(
        {"gui": {"language": {"locale": "en_US"}}},
        locales_dir=PROJECT_ROOT / "i18n" / "locales",
    )


def _save_reopen(document, tmp_path: Path, name: str):
    path = tmp_path / name
    document.save(path)
    return path, Document(str(path))


def _paragraph_style_id(paragraph) -> str | None:
    style = paragraph._p.find(f"{qn('w:pPr')}/{qn('w:pStyle')}")
    return style.get(qn("w:val")) if style is not None else None


def _run_text(run) -> str:
    return "".join(text.text or "" for text in run.iter(qn("w:t")))


def test_managed_body_default_and_custom_placeholder_survive_reopen(tmp_path: Path) -> None:
    document, bindings = complete_managed_styles(Document(), _catalog())
    rendered = MdToDocxRenderer(document, body_style=document.styles["Normal"], managed_styles=bindings).render(
        parse_markdown_text("Default body.\n")
    )
    fill_template(document, {}, rendered, None)
    _path, reopened = _save_reopen(document, tmp_path, "default-body.docx")
    assert _paragraph_style_id(next(p for p in reopened.paragraphs if p.text == "Default body.")) == (
        "DocWenBodyParagraph"
    )

    template = Document()
    user_style = template.styles.add_style("UserBody", WD_STYLE_TYPE.PARAGRAPH)
    placeholder = template.add_paragraph("{{ body }}", style=user_style.name)
    placeholder.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    placeholder.paragraph_format.left_indent = Inches(0.3)
    placeholder.paragraph_format.space_after = Pt(17)
    placeholder.paragraph_format.keep_with_next = True
    template, bindings = complete_managed_styles(template, _catalog())
    placeholder = next(p for p in template.paragraphs if p.text == "{{ body }}")
    body_format = extract_body_paragraph_format(template)
    assert body_format is not None
    rendered = MdToDocxRenderer(
        template,
        body_style=placeholder.style,
        body_paragraph_format=body_format,
        managed_styles=bindings,
    ).render(parse_markdown_text("Custom body.\n"))
    fill_template(template, {}, rendered, placeholder)
    _path, reopened = _save_reopen(template, tmp_path, "custom-body.docx")
    body = next(p for p in reopened.paragraphs if p.text == "Custom body.")
    assert _paragraph_style_id(body) == "UserBody"
    p_pr = body._p.find(qn("w:pPr"))
    assert p_pr is not None
    alignment = p_pr.find(qn("w:jc"))
    indent = p_pr.find(qn("w:ind"))
    spacing = p_pr.find(qn("w:spacing"))
    assert alignment is not None and alignment.get(qn("w:val")) == "right"
    assert indent is not None and indent.get(qn("w:left")) == "432"
    assert spacing is not None and spacing.get(qn("w:after")) == "340"
    assert p_pr.find(qn("w:keepNext")) is not None


def test_managed_renderer_emits_stable_semantic_styles_and_rich_links(tmp_path: Path) -> None:
    document, bindings = complete_managed_styles(Document(), _catalog())
    inline_code = bindings.get("inline_code")
    inline_element = inline_code._element
    assert inline_element is not None
    inline_r_pr = inline_element.get_or_add_rPr()
    fonts = inline_r_pr.get_or_add_rFonts()
    for slot, value in (("ascii", "UserMono"), ("hAnsi", "UserMono"), ("eastAsia", "UserCJK"), ("cs", "UserCS")):
        fonts.set(qn(f"w:{slot}"), value)
    shading = inline_r_pr.find(qn("w:shd"))
    assert shading is not None
    shading.set(qn("w:fill"), "ABCDEF")
    grid = bindings.get("table_grid")
    grid_element = grid._element
    assert grid_element is not None
    grid_properties = grid_element.find(qn("w:tblPr"))
    if grid_properties is None:
        grid_properties = OxmlElement("w:tblPr")
        grid_element.append(grid_properties)
    grid_properties.append(OxmlElement("w:tblLayout"))

    analysis = analyze_document_semantics(
        parse_markdown_text(
            """Body [plain 😀 **bold** *italic* ~~strike~~ `code` ==mark== H~sub~ x^sup^ $x$](https://example.com).

Figure: Figure caption

![](missing.png)

Table: Table caption contains literal SEQ Figure

|  | Header |
|---|---|
|  | Body |

Equation: Equation caption

$$x^2$$

Listing: Listing caption

```python
print('x')
```
"""
        )
    )
    assert not analysis.has_errors
    body_format = BodyParagraphFormat(
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        left_indent=Inches(0.4),
        right_indent=None,
        first_line_indent=None,
        space_before=None,
        space_after=Pt(19),
        line_spacing=None,
        line_spacing_rule=None,
        keep_with_next=True,
        keep_together=None,
        page_break_before=None,
        widow_control=None,
    )
    renderer = MdToDocxRenderer(
        document,
        body_style=document.styles["Normal"],
        body_paragraph_format=body_format,
        table_style_name="Definitely Missing Style",
        managed_styles=bindings,
        source_file_path=str(tmp_path / "source.md"),
    )
    rendered = renderer.render(analysis.ast)
    fill_template(document, {}, rendered, None)
    path, reopened = _save_reopen(document, tmp_path, "managed-rendering.docx")
    validate_managed_style_package(path.read_bytes(), _catalog())

    assert _paragraph_style_id(next(p for p in reopened.paragraphs if p.text.startswith("Body plain"))) == (
        "DocWenBodyParagraph"
    )
    image = next(p for p in reopened.paragraphs if _paragraph_style_id(p) == "DocWenImageParagraph")
    assert _paragraph_style_id(image) == "DocWenImageParagraph"
    image_p_pr = image._p.find(qn("w:pPr"))
    assert image_p_pr is not None
    for direct_tag in ("jc", "ind", "spacing", "keepNext"):
        assert image_p_pr.find(qn(f"w:{direct_tag}")) is None

    caption_styles = {
        kind: _paragraph_style_id(next(p for p in reopened.paragraphs if f"SEQ {kind.title()}" in p._p.xml))
        for kind in ("figure", "table", "equation", "listing")
    }
    assert caption_styles == {
        "figure": "DocWenFigureCaption",
        "table": "DocWenTableCaption",
        "equation": "DocWenEquationCaption",
        "listing": "DocWenCodeBlockCaption",
    }
    assert "print('x')" in [p.text for p in reopened.paragraphs]

    table = reopened.tables[0]
    assert table.style is not None and table.style.style_id == "DocWenTableGrid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            expected = "DocWenTableHeader" if row_index == 0 else "DocWenTableContent"
            assert all(_paragraph_style_id(paragraph) == expected for paragraph in cell.paragraphs)
            assert all(run.bold is not True for paragraph in cell.paragraphs for run in paragraph.runs)

    body = next(p for p in reopened.paragraphs if p.text.startswith("Body plain"))
    hyperlink = body._p.find(qn("w:hyperlink"))
    assert hyperlink is not None
    assert "".join(item.text or "" for item in hyperlink.iter(qn("w:t"))) == (
        "plain 😀 bold italic strike code mark Hsub xsup $x$"
    )
    for run in hyperlink.findall(qn("w:r")):
        r_pr = run.find(qn("w:rPr"))
        assert r_pr is not None
        styles = r_pr.findall(qn("w:rStyle"))
        assert len(styles) == 1
        assert styles[0].get(qn("w:val")) == "Hyperlink"
        ranks = [
            (
                "rStyle",
                "rFonts",
                "b",
                "bCs",
                "i",
                "iCs",
                "strike",
                "color",
                "sz",
                "szCs",
                "highlight",
                "u",
                "shd",
                "vertAlign",
            ).index(child.tag.rsplit("}", 1)[-1])
            for child in r_pr
            if child.tag.rsplit("}", 1)[-1]
            in {
                "rStyle",
                "rFonts",
                "b",
                "bCs",
                "i",
                "iCs",
                "strike",
                "color",
                "sz",
                "szCs",
                "highlight",
                "u",
                "shd",
                "vertAlign",
            }
        ]
        assert ranks == sorted(ranks)
    runs_by_text = {
        "".join(text.text or "" for text in run.iter(qn("w:t"))): run for run in hyperlink.findall(qn("w:r"))
    }
    assert runs_by_text["bold"].find(f"{qn('w:rPr')}/{qn('w:b')}") is not None
    assert runs_by_text["italic"].find(f"{qn('w:rPr')}/{qn('w:i')}") is not None
    assert runs_by_text["strike"].find(f"{qn('w:rPr')}/{qn('w:strike')}") is not None
    mark = runs_by_text["mark"].find(f"{qn('w:rPr')}/{qn('w:highlight')}")
    assert mark is not None and mark.get(qn("w:val")) == "yellow"
    superscript = runs_by_text["sup"].find(f"{qn('w:rPr')}/{qn('w:vertAlign')}")
    subscript = runs_by_text["sub"].find(f"{qn('w:rPr')}/{qn('w:vertAlign')}")
    assert superscript is not None and superscript.get(qn("w:val")) == "superscript"
    assert subscript is not None and subscript.get(qn("w:val")) == "subscript"
    assert runs_by_text["$x$"].find(f".//{qn('m:oMath')}") is None
    code_run = next(
        run for run in hyperlink.findall(qn("w:r")) if "".join(t.text or "" for t in run.iter(qn("w:t"))) == "code"
    )
    code_properties = code_run.find(qn("w:rPr"))
    assert code_properties is not None
    code_fonts = code_properties.find(qn("w:rFonts"))
    assert code_fonts is not None
    assert code_fonts.get(qn("w:ascii")) == "UserMono"
    assert code_fonts.get(qn("w:hAnsi")) == "UserMono"
    assert code_fonts.get(qn("w:eastAsia")) == "UserCJK"
    assert code_fonts.get(qn("w:cs")) == "UserCS"
    code_shading = code_properties.find(qn("w:shd"))
    assert code_shading is not None and code_shading.get(qn("w:fill")) == "ABCDEF"
    emoji_run = next(
        run for run in hyperlink.findall(qn("w:r")) if "😀" in "".join(t.text or "" for t in run.iter(qn("w:t")))
    )
    emoji_fonts = emoji_run.find(f"{qn('w:rPr')}/{qn('w:rFonts')}")
    assert emoji_fonts is not None
    assert {emoji_fonts.get(qn(f"w:{slot}")) for slot in ("ascii", "hAnsi", "eastAsia", "cs")} == {"Segoe UI Emoji"}


def test_no_placeholder_batch_order_section_and_legacy_hyperlink_api(tmp_path: Path) -> None:
    document, bindings = complete_managed_styles(Document(), _catalog())
    document.add_paragraph("PREFACE")
    analysis = analyze_document_semantics(
        parse_markdown_text(
            """Equation: E

$$x$$

Listing: L

```
code
```

AFTER
"""
        )
    )
    rendered = MdToDocxRenderer(document, managed_styles=bindings).render(analysis.ast)
    fill_template(document, {}, rendered, None)
    path, reopened = _save_reopen(document, tmp_path, "no-placeholder.docx")
    body = reopened.element.find(qn("w:body"))
    assert body is not None
    body_children = list(body)
    assert body_children[-1].tag == qn("w:sectPr")
    texts = [p.text for p in reopened.paragraphs]
    assert texts[0] == "PREFACE"
    assert next(index for index, p in enumerate(reopened.paragraphs) if "SEQ Equation" in p._p.xml) < next(
        index for index, p in enumerate(reopened.paragraphs) if p._p.find(f".//{qn('m:oMath')}") is not None
    )
    assert next(index for index, p in enumerate(reopened.paragraphs) if "SEQ Listing" in p._p.xml) < texts.index("code")
    assert texts[-1] == "AFTER"
    validate_managed_style_package(path.read_bytes(), _catalog())

    legacy = Document()
    paragraph = legacy.add_paragraph()
    add_hyperlink(paragraph, "https://example.com", text="standalone")
    assert "standalone" in paragraph._p.xml
    MdToDocxRenderer(legacy).add_hyperlink(paragraph, "https://example.com", text="renderer")
    assert "renderer" in paragraph._p.xml

    sentinel = object()
    positional = MdToDocxRenderer(
        legacy,
        None,
        None,
        None,
        "Consolas",
        "E7E6E6",
        "remove",
        "remove",
        "Table Grid",
        None,
        None,
        None,
        sentinel,
    )
    assert positional._cancellation is sentinel


@pytest.mark.parametrize("mode", ("remove", "apply"))
def test_heading_link_format_mode_matches_non_link_semantics(mode: str) -> None:
    document, bindings = complete_managed_styles(Document(), _catalog())
    renderer = MdToDocxRenderer(document, managed_styles=bindings, heading_formatting_mode=mode)
    renderer.render(
        [
            {
                "type": "heading",
                "attrs": {"level": 1},
                "children": [
                    {"type": "text", "raw": "plain "},
                    {
                        "type": "link",
                        "attrs": {"url": "https://example.com"},
                        "children": [
                            {"type": "text", "raw": "plain"},
                            {"type": "strong", "children": [{"type": "text", "raw": "bold"}]},
                            {"type": "codespan", "raw": "code"},
                            {"type": "highlight", "raw": "mark"},
                            {"type": "superscript", "raw": "sup"},
                        ],
                    },
                ],
            }
        ]
    )
    hyperlink = document.paragraphs[0]._p.find(qn("w:hyperlink"))
    assert hyperlink is not None
    runs = {"".join(text.text or "" for text in run.iter(qn("w:t"))): run for run in hyperlink.findall(qn("w:r"))}
    plain_properties = runs["plain"].find(qn("w:rPr"))
    bold_properties = runs["bold"].find(qn("w:rPr"))
    code_properties = runs["code"].find(qn("w:rPr"))
    assert plain_properties is not None and bold_properties is not None and code_properties is not None
    if mode == "apply":
        plain_bold = plain_properties.find(qn("w:b"))
        plain_italic = plain_properties.find(qn("w:i"))
        assert plain_bold is not None and plain_bold.get(qn("w:val")) == "0"
        assert plain_italic is not None and plain_italic.get(qn("w:val")) == "0"
        assert bold_properties.find(qn("w:b")) is not None
    else:
        assert plain_properties.find(qn("w:b")) is None
        assert plain_properties.find(qn("w:i")) is None
        assert bold_properties.find(qn("w:b")) is None
    assert code_properties.find(qn("w:rFonts")) is not None
    assert runs["mark"].find(f"{qn('w:rPr')}/{qn('w:highlight')}") is not None
    assert runs["sup"].find(f"{qn('w:rPr')}/{qn('w:vertAlign')}") is not None


def test_inline_code_projection_distinguishes_empty_style_and_merges_inherited_properties() -> None:
    document, bindings = complete_managed_styles(Document(), _catalog())
    inline_code = bindings.get("inline_code")
    assert isinstance(inline_code, CharacterStyle)
    inline_element = inline_code._element
    assert inline_element is not None
    inline_r_pr = inline_element.get_or_add_rPr()
    for child in list(inline_r_pr):
        inline_r_pr.remove(child)

    paragraph = document.add_paragraph()
    renderer = MdToDocxRenderer(document, managed_styles=bindings)
    renderer.add_hyperlink(
        paragraph,
        "https://example.com",
        [{"type": "codespan", "raw": "empty-style"}],
    )
    link_run = paragraph._p.find(f"{qn('w:hyperlink')}/{qn('w:r')}")
    assert link_run is not None
    link_properties = link_run.find(qn("w:rPr"))
    assert link_properties is not None
    assert link_properties.find(qn("w:rFonts")) is None
    assert link_properties.find(qn("w:shd")) is None

    base = document.styles.add_style("UserCodeBase", WD_STYLE_TYPE.CHARACTER)
    assert isinstance(base, CharacterStyle)
    base_element = base._element
    assert base_element is not None
    base_r_pr = base_element.get_or_add_rPr()
    base_fonts = OxmlElement("w:rFonts")
    for slot, value in (("hAnsi", "BaseHAnsi"), ("eastAsia", "BaseCJK"), ("cs", "BaseCS")):
        base_fonts.set(qn(f"w:{slot}"), value)
    base_r_pr.append(base_fonts)
    base_lang = OxmlElement("w:lang")
    base_lang.set(qn("w:val"), "en-US")
    base_lang.set(qn("w:eastAsia"), "zh-CN")
    base_r_pr.append(base_lang)
    base_r_pr.append(OxmlElement("w:b"))
    inline_code.base_style = base
    leaf_fonts = OxmlElement("w:rFonts")
    leaf_fonts.set(qn("w:ascii"), "LeafAscii")
    inline_r_pr.append(leaf_fonts)
    leaf_lang = OxmlElement("w:lang")
    leaf_lang.set(qn("w:bidi"), "ar-SA")
    inline_r_pr.append(leaf_lang)

    bold_paragraph = document.styles.add_style("UserBoldParagraph", WD_STYLE_TYPE.PARAGRAPH)
    assert isinstance(bold_paragraph, ParagraphStyle)
    bold_element = bold_paragraph._element
    assert bold_element is not None
    bold_element.get_or_add_rPr().append(OxmlElement("w:b"))
    paragraph.style = bold_paragraph
    projection = project_character_style_properties(inline_code, paragraph=paragraph)
    assert projection is not None
    by_tag = {child.tag: child for child in projection}
    projected_fonts = by_tag[qn("w:rFonts")]
    assert projected_fonts.get(qn("w:ascii")) == "LeafAscii"
    assert projected_fonts.get(qn("w:hAnsi")) == "BaseHAnsi"
    assert projected_fonts.get(qn("w:eastAsia")) == "BaseCJK"
    assert projected_fonts.get(qn("w:cs")) == "BaseCS"
    projected_lang = by_tag[qn("w:lang")]
    assert projected_lang.get(qn("w:val")) == "en-US"
    assert projected_lang.get(qn("w:eastAsia")) == "zh-CN"
    assert projected_lang.get(qn("w:bidi")) == "ar-SA"
    projected_bold = by_tag[qn("w:b")]
    assert projected_bold.get(qn("w:val")) == "0"


def test_hyperlink_break_fallback_relation_cleanup_and_outer_script_precedence(monkeypatch) -> None:
    children = [
        {"type": "text", "raw": "a"},
        {"type": "linebreak"},
        {"type": "text", "raw": "b"},
        {
            "type": "subscript",
            "children": [{"type": "superscript", "children": [{"type": "text", "raw": "outer-sub"}]}],
        },
    ]
    document = Document()
    valid = document.add_paragraph()
    add_hyperlink(valid, "https://example.com", children=children)
    valid_link = valid._p.find(qn("w:hyperlink"))
    assert valid_link is not None and valid_link.find(f".//{qn('w:br')}") is not None
    outer_sub = next(run for run in valid_link.findall(qn("w:r")) if _run_text(run) == "outer-sub")
    vertical = outer_sub.find(f"{qn('w:rPr')}/{qn('w:vertAlign')}")
    assert vertical is not None and vertical.get(qn("w:val")) == "subscript"

    invalid = document.add_paragraph()
    add_hyperlink(invalid, "javascript:alert(1)", children=children)
    assert invalid._p.find(qn("w:hyperlink")) is None
    assert invalid._p.find(f".//{qn('w:br')}") is not None

    import docwen_plugin_markdown.renderer_inlines as renderer_inlines

    original_properties = renderer_inlines._hyperlink_run_properties

    def fail_properties(*args, **kwargs):
        raise RuntimeError("injected link construction failure")

    monkeypatch.setattr(renderer_inlines, "_hyperlink_run_properties", fail_properties)
    failed = document.add_paragraph()
    MdToDocxRenderer(document).add_hyperlink(failed, "https://failure.example", children)
    assert failed._p.find(qn("w:hyperlink")) is None
    assert failed._p.find(f".//{qn('w:br')}") is not None
    monkeypatch.setattr(renderer_inlines, "_hyperlink_run_properties", original_properties)

    class FailingElement:
        def append(self, _element) -> None:
            raise RuntimeError("injected append failure")

    class FakePart:
        def __init__(self, *, reused: bool) -> None:
            self.relationship_id = "rId9"
            self.rels = {self.relationship_id: object()} if reused else {}

        def relate_to(self, _target, _relationship_type, *, is_external: bool) -> str:
            assert is_external is True
            self.rels.setdefault(self.relationship_id, object())
            return self.relationship_id

        def drop_rel(self, relationship_id: str) -> None:
            self.rels.pop(relationship_id, None)

    class FakeParagraph:
        def __init__(self, part: FakePart) -> None:
            self.part = part
            self._p = FailingElement()

    for reused in (False, True):
        part = FakePart(reused=reused)
        add_hyperlink(
            FakeParagraph(part),
            "https://same.example",
            text="failure",
            fallback_renderer=lambda _paragraph, _children: None,
        )
        assert (part.relationship_id in part.rels) is reused


def test_direct_render_keeps_all_caption_objects_adjacent_and_managed_hr_styles(tmp_path: Path) -> None:
    document, bindings = complete_managed_styles(Document(), _catalog())
    analysis = analyze_document_semantics(
        parse_markdown_text(
            """Figure: F

![](missing.png)

Table: T

| H |
|---|
| B |

Equation: E

$$x$$

Listing: L

```python
code
```
"""
        )
    )
    MdToDocxRenderer(document, managed_styles=bindings, source_file_path=str(tmp_path / "source.md")).render(
        analysis.ast
    )
    body = document.element.find(qn("w:body"))
    assert body is not None
    children = list(body)
    figure = next(p._p for p in document.paragraphs if _paragraph_style_id(p) == "DocWenImageParagraph")
    figure_caption = next(p._p for p in document.paragraphs if "SEQ Figure" in p._p.xml)
    table_caption = next(p._p for p in document.paragraphs if "SEQ Table" in p._p.xml)
    table = document.tables[0]._tbl
    equation_caption = next(p._p for p in document.paragraphs if "SEQ Equation" in p._p.xml)
    equation = next(p._p for p in document.paragraphs if p._p.find(f".//{qn('m:oMath')}") is not None)
    listing_caption = next(p._p for p in document.paragraphs if "SEQ Listing" in p._p.xml)
    listing = next(p._p for p in document.paragraphs if p.text == "code")
    assert children.index(figure_caption) == children.index(figure) + 1
    assert children.index(table) == children.index(table_caption) + 1
    assert children.index(equation) == children.index(equation_caption) + 1
    assert children.index(listing) == children.index(listing_caption) + 1

    hr_document, hr_bindings = complete_managed_styles(Document(), _catalog())
    rendered_rules = MdToDocxRenderer(
        hr_document,
        managed_styles=hr_bindings,
        hr_actions={
            "dash": "horizontal_rule_1",
            "asterisk": "horizontal_rule_2",
            "underscore": "horizontal_rule_3",
        },
    ).render(
        [
            {"type": "thematic_break", "_hr_marker": "dash"},
            {"type": "thematic_break", "_hr_marker": "asterisk"},
            {"type": "thematic_break", "_hr_marker": "underscore"},
        ]
    )
    assert [_paragraph_style_id(paragraph) for paragraph in rendered_rules] == [
        "DocWenHorizontalRule1",
        "DocWenHorizontalRule2",
        "DocWenHorizontalRule3",
    ]
    assert all(paragraph._p.find(f"{qn('w:pPr')}/{qn('w:pBdr')}") is None for paragraph in rendered_rules)
