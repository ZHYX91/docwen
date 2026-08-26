"""OOXML-level tests for MD → DOCX break elements.

Validates that page breaks, section breaks, and horizontal rules produce
correct WordprocessingML structures in the output DOCX.
"""

from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from docwen_plugin_markdown.ast_transforms import annotate_ast_with_hr_attachments
from docwen_plugin_markdown.mistune_extensions import parse_markdown_text
from docwen_plugin_markdown.renderer import MdToDocxRenderer

pytestmark = pytest.mark.contract
from docwen_plugin_markdown.to_docx.breaks import (
    append_horizontal_rule_to_paragraph,
    append_page_break_to_paragraph,
    append_section_break_to_paragraph,
    insert_horizontal_rule,
    insert_page_break,
    insert_section_break,
)

# ── Namespace ────────────────────────────────────────────────────────────
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# ── Shared Markdown snippets ─────────────────────────────────────────────

MD_WITH_HR = "# Title\n\nSome text.\n\n---\n\nMore text."
MD_TWO_HRS = "# Doc\n\n---\n\nMiddle.\n\n---\n\nAfter."
MD_NO_HR = "# Just a heading\n\nWith a paragraph."


# ═════════════════════════════════════════════════════════════════════════
# Low-level unit tests — new-paragraph insertions
# ═════════════════════════════════════════════════════════════════════════


class TestInsertPageBreak:
    """F-F1-001: insert_page_break creates proper OOXML page break."""

    def test_creates_paragraph_with_page_break(self):
        doc = Document()
        p = insert_page_break(doc)

        # Must have added exactly one paragraph to the document
        assert len(doc.paragraphs) == 1

        # The returned paragraph OOXML must match what was inserted
        doc_para = doc.paragraphs[0]
        assert doc_para._p is p._p, "Returned paragraph does not match document paragraph"

        # OOXML: contain <w:br w:type="page"/>
        br_elements = p._p.findall(f"{{{_W}}}r/{{{_W}}}br")
        assert len(br_elements) >= 1, "Expected at least one w:br element"

        br = br_elements[0]
        assert br.get(qn("w:type")) == "page", f"Expected w:type='page', got {br.get(qn('w:type'))}"

    def test_page_break_is_empty_paragraph(self):
        """Page break paragraph should have no visible text."""
        doc = Document()
        p = insert_page_break(doc)
        assert p.text == ""


class TestInsertSectionBreak:
    """F-F1-002: insert_section_break creates proper OOXML section break."""

    def test_creates_paragraph_with_sectpr(self):
        doc = Document()
        p = insert_section_break(doc, section_type="next")

        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0]._p is p._p

        # OOXML: <w:pPr><w:sectPr> must exist
        sectPr_list = p._p.findall(f"{{{_W}}}pPr/{{{_W}}}sectPr")
        assert len(sectPr_list) == 1, "Expected exactly one w:sectPr in paragraph pPr"

    def test_section_type_is_nextPage_by_default(self):
        doc = Document()
        p = insert_section_break(doc)
        sectPr = p._p.find(f"{{{_W}}}pPr/{{{_W}}}sectPr")
        assert sectPr is not None, "Missing sectPr in paragraph"
        stype = sectPr.find(f"{{{_W}}}type")
        assert stype is not None, "Missing w:type in sectPr"
        assert stype.get(qn("w:val")) == "nextPage"

    @pytest.mark.parametrize(
        "short,expected_ooXML",
        [
            ("next", "nextPage"),
            ("continuous", "continuous"),
            ("even", "evenPage"),
            ("odd", "oddPage"),
        ],
    )
    def test_section_type_mapping(self, short, expected_ooXML):
        doc = Document()
        p = insert_section_break(doc, section_type=short)
        sectPr = p._p.find(f"{{{_W}}}pPr/{{{_W}}}sectPr")
        assert sectPr is not None, "Missing sectPr"
        stype = sectPr.find(f"{{{_W}}}type")
        assert stype is not None, "Missing type"
        assert stype.get(qn("w:val")) == expected_ooXML

    def test_copies_page_settings(self):
        """Section break sectPr must contain page settings from the document."""
        doc = Document()
        p = insert_section_break(doc)
        sectPr = p._p.find(f"{{{_W}}}pPr/{{{_W}}}sectPr")

        # At minimum pgSz and pgMar should be present
        assert sectPr is not None, "Missing sectPr"
        pgSz = sectPr.find(f"{{{_W}}}pgSz")
        pgMar = sectPr.find(f"{{{_W}}}pgMar")
        assert pgSz is not None, "Missing pgSz in sectPr"
        assert pgMar is not None, "Missing pgMar in sectPr"


class TestInsertHorizontalRule:
    """F-F1-003: insert_horizontal_rule creates bottom-border paragraph."""

    def test_creates_paragraph_with_bottom_border(self):
        doc = Document()
        p = insert_horizontal_rule(doc, variant="1")

        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0]._p is p._p

        # OOXML: <w:pPr><w:pBdr><w:bottom w:val="single" .../></w:pBdr></w:pPr>
        bottom_list = p._p.findall(f"{{{_W}}}pPr/{{{_W}}}pBdr/{{{_W}}}bottom")
        assert len(bottom_list) == 1, "Expected exactly one w:bottom border element"
        bottom = bottom_list[0]
        assert bottom.get(qn("w:val")) == "single"

    @pytest.mark.parametrize(
        "variant,expected_sz",
        [("1", "4"), ("2", "8"), ("3", "12")],
    )
    def test_variant_thickness(self, variant, expected_sz):
        doc = Document()
        p = insert_horizontal_rule(doc, variant=variant)
        bottom = p._p.find(f"{{{_W}}}pPr/{{{_W}}}pBdr/{{{_W}}}bottom")
        assert bottom is not None, f"Variant {variant}: bottom border not found"
        assert bottom.get(qn("w:sz")) == expected_sz, (
            f"Variant {variant}: expected sz={expected_sz}, got {bottom.get(qn('w:sz'))}"
        )

    def test_hr_paragraph_is_empty(self):
        """Horizontal rule paragraph should have no visible text."""
        doc = Document()
        p = insert_horizontal_rule(doc)
        assert p.text == ""

    def test_unknown_variant_falls_back_to_default(self):
        """Unknown variant should not crash; falls back to variant "1"."""
        doc = Document()
        p = insert_horizontal_rule(doc, variant="99")
        bottom = p._p.find(f"{{{_W}}}pPr/{{{_W}}}pBdr/{{{_W}}}bottom")
        # Should fall back to default sz="4"
        assert bottom is not None, "Missing bottom border"
        assert bottom.get(qn("w:sz")) == "4"


# ═════════════════════════════════════════════════════════════════════════
# Low-level unit tests — attach-to-previous variants
# ═════════════════════════════════════════════════════════════════════════


class TestAppendPageBreakToParagraph:
    """F-F1-004: append_page_break_to_paragraph."""

    def test_appends_br_to_existing_paragraph(self):
        doc = Document()
        p = doc.add_paragraph("Some text before break.")
        original_run_count = len(p._p.findall(f"{{{_W}}}r"))

        append_page_break_to_paragraph(p)

        # Should have one more run (the break run)
        runs_after = p._p.findall(f"{{{_W}}}r")
        assert len(runs_after) == original_run_count + 1

        # The last run should contain <w:br w:type="page"/>
        last_run = runs_after[-1]
        br_list = last_run.findall(f"{{{_W}}}br")
        assert len(br_list) == 1
        assert br_list[0].get(qn("w:type")) == "page"


class TestAppendSectionBreakToParagraph:
    """F-F1-005: append_section_break_to_paragraph."""

    def test_adds_sectpr_to_paragraph_ppr(self):
        doc = Document()
        p = doc.add_paragraph("Text before section break.")

        # No sectPr before
        assert p._p.find(f"{{{_W}}}pPr/{{{_W}}}sectPr") is None

        append_section_break_to_paragraph(p, doc, section_type="continuous")

        sectPr = p._p.find(f"{{{_W}}}pPr/{{{_W}}}sectPr")
        assert sectPr is not None, "sectPr was not added to paragraph pPr"
        stype = sectPr.find(f"{{{_W}}}type")
        assert stype is not None
        assert stype.get(qn("w:val")) == "continuous"


class TestAppendHorizontalRuleToParagraph:
    """F-F1-006: append_horizontal_rule_to_paragraph."""

    def test_adds_bottom_border_to_existing_paragraph(self):
        doc = Document()
        p = doc.add_paragraph("Some text.")

        # No pBdr before
        assert p._p.find(f"{{{_W}}}pPr/{{{_W}}}pBdr") is None

        append_horizontal_rule_to_paragraph(p, variant="2")

        bottom = p._p.find(f"{{{_W}}}pPr/{{{_W}}}pBdr/{{{_W}}}bottom")
        assert bottom is not None, "Bottom border was not added"
        assert bottom.get(qn("w:val")) == "single"
        assert bottom.get(qn("w:sz")) == "8"  # variant 2


# ═════════════════════════════════════════════════════════════════════════
# Renderer integration tests — full AST → DOCX pipeline
# ═════════════════════════════════════════════════════════════════════════


class TestRendererThematicBreak:
    """_handle_thematic_break produces real horizontal rules in the DOCX."""

    def test_hr_annotation_restores_source_markers_by_order(self):
        """Mistune normalizes HR tokens, so source-order annotation preserves marker kind."""
        md = "a\n\n---\n\nb\n\n***\n\nc\n\n___\n\nd"
        ast = parse_markdown_text(md)

        annotate_ast_with_hr_attachments(ast, set(), md)

        markers = [node.get("_hr_marker") for node in ast if node.get("type") == "thematic_break"]
        assert markers == ["dash", "asterisk", "underscore"]

    def test_thematic_break_produces_horizontal_rule_not_text(self):
        """A Markdown '---' must become a bottom-border paragraph, not text."""
        ast = parse_markdown_text(MD_WITH_HR)
        doc = Document()
        renderer = MdToDocxRenderer(doc)
        renderer.render(ast)

        # Find the thematic-break paragraph — it should have a bottom border
        hr_found = False
        for p in doc.paragraphs:
            bottom = p._p.find(f"{{{_W}}}pPr/{{{_W}}}pBdr/{{{_W}}}bottom")
            if bottom is not None:
                hr_found = True
                assert bottom.get(qn("w:val")) == "single"
                # Must NOT be the fake "─" * 60 text
                assert "─" not in p.text
                break

        assert hr_found, "No horizontal-rule paragraph (with w:bottom border) found in output"

    def test_thematic_break_paragraph_has_no_text(self):
        """HR paragraph should be visually empty (no runs with text)."""
        ast = parse_markdown_text(MD_WITH_HR)
        doc = Document()
        renderer = MdToDocxRenderer(doc)
        renderer.render(ast)

        for p in doc.paragraphs:
            bottom = p._p.find(f"{{{_W}}}pPr/{{{_W}}}pBdr/{{{_W}}}bottom")
            if bottom is not None:
                assert p.text == "", f"HR paragraph should have no text, got: {p.text!r}"
                return

        pytest.fail("No HR paragraph found")

    def test_two_thematic_breaks_produce_two_hrs(self):
        """Each '---' should become its own HR paragraph."""
        ast = parse_markdown_text(MD_TWO_HRS)
        doc = Document()
        renderer = MdToDocxRenderer(doc)
        renderer.render(ast)

        hr_count = 0
        for p in doc.paragraphs:
            bottom = p._p.find(f"{{{_W}}}pPr/{{{_W}}}pBdr/{{{_W}}}bottom")
            if bottom is not None:
                hr_count += 1

        assert hr_count == 2, f"Expected 2 HR paragraphs, got {hr_count}"

    def test_no_false_positives_for_normal_paragraphs(self):
        """Ordinary paragraphs must not have bottom borders."""
        ast = parse_markdown_text(MD_NO_HR)
        doc = Document()
        renderer = MdToDocxRenderer(doc)
        renderer.render(ast)

        for p in doc.paragraphs:
            bottom = p._p.find(f"{{{_W}}}pPr/{{{_W}}}pBdr/{{{_W}}}bottom")
            assert bottom is None, f"Non-HR paragraph should not have bottom border. Text: {p.text!r}"


# ═════════════════════════════════════════════════════════════════════════
# User-path integration — convert Markdown file → inspect DOCX
# ═════════════════════════════════════════════════════════════════════════


class TestUserPathBreakInDocx:
    """End-to-end path: Markdown → DOCX via converter, verify breaks in ZIP."""

    def test_page_break_function_usable_in_pipeline(self):
        """Simulates a programmatic use: page break inserted into DOCX."""
        doc = Document()
        doc.add_paragraph("Page 1 content.")
        insert_page_break(doc)
        doc.add_paragraph("Page 2 content.")

        # Verify structure
        assert len(doc.paragraphs) == 3
        # Middle paragraph has the break
        br = doc.paragraphs[1]._p.find(f"{{{_W}}}r/{{{_W}}}br")
        assert br is not None
        assert br.get(qn("w:type")) == "page"

    def test_section_break_function_usable_in_pipeline(self):
        """Simulates a programmatic use: section break inserted into DOCX."""
        doc = Document()
        doc.add_paragraph("Section A.")
        insert_section_break(doc, section_type="next")
        doc.add_paragraph("Section B.")

        # Verify structure
        sectPr = doc.paragraphs[1]._p.find(f"{{{_W}}}pPr/{{{_W}}}sectPr")
        assert sectPr is not None
        stype = sectPr.find(f"{{{_W}}}type")
        assert stype is not None, "Missing w:type in sectPr"
        assert stype.get(qn("w:val")) == "nextPage"

    def test_roundtrip_hr_preserves_structure(self):
        """Horizontal rules survive a DOCX save/reload roundtrip."""
        doc = Document()
        doc.add_paragraph("Before HR.")
        insert_horizontal_rule(doc, variant="1")
        doc.add_paragraph("After HR.")

        # Save and re-read
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
            tmp_path = tf.name
        try:
            doc.save(tmp_path)

            # Must be valid ZIP
            assert zipfile.is_zipfile(tmp_path)

            doc2 = Document(tmp_path)
            hr_found = False
            for p in doc2.paragraphs:
                bottom = p._p.find(f"{{{_W}}}pPr/{{{_W}}}pBdr/{{{_W}}}bottom")
                if bottom is not None:
                    hr_found = True
                    break
            assert hr_found, "HR not preserved after DOCX roundtrip"
        finally:
            Path(tmp_path).unlink(missing_ok=True)
