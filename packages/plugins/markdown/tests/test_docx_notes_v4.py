"""v4 existing-note collision, rich-body, relationship, and atomicity contracts."""

from __future__ import annotations

import base64
import hashlib
from copy import deepcopy
from pathlib import Path
from typing import TYPE_CHECKING, Any, cast
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import lxml.etree as etree
import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter
from docwen_plugin_markdown.to_docx.notes import (
    WML_NS,
    NoteContext,
    NoteWritebackError,
    prepare_note_context_for_document,
    write_notes_to_docx,
)

from .conftest import make_context

if TYPE_CHECKING:
    from docx.document import Document as DocumentObject

pytestmark = pytest.mark.contract

_RELS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
_OFFICE_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_FOOTNOTES_REL_TYPE = f"{_OFFICE_REL_NS}/footnotes"
_ENDNOTES_REL_TYPE = f"{_OFFICE_REL_NS}/endnotes"
_HYPERLINK_REL_TYPE = f"{_OFFICE_REL_NS}/hyperlink"
_IMAGE_REL_TYPE = f"{_OFFICE_REL_NS}/image"
_FOOTNOTES_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
_ENDNOTES_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml"
_PNG = base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII=")


def _note_children(text: str) -> list[list[dict[str, str]]]:
    return [[{"type": "text", "raw": text}]]


def _append_footnote_reference(document: DocumentObject, note_id: int) -> None:
    paragraph = document.paragraphs[0] if document.paragraphs else document.add_paragraph()
    run = OxmlElement("w:r")
    reference = OxmlElement("w:footnoteReference")
    reference.set(qn("w:id"), str(note_id))
    run.append(reference)
    paragraph._p.append(run)


def _append_endnote_reference(document: DocumentObject, note_id: int) -> None:
    paragraph = document.paragraphs[0] if document.paragraphs else document.add_paragraph()
    run = OxmlElement("w:r")
    reference = OxmlElement("w:endnoteReference")
    reference.set(qn("w:id"), str(note_id))
    run.append(reference)
    paragraph._p.append(run)


def _rich_footnotes_xml(damage: str | None) -> bytes:
    root = etree.Element(f"{{{WML_NS}}}footnotes", nsmap={"w": WML_NS, "r": _OFFICE_REL_NS})
    separator = etree.SubElement(root, f"{{{WML_NS}}}footnote")
    separator.set(f"{{{WML_NS}}}type", "separator")
    separator.set(f"{{{WML_NS}}}id", "-1")
    etree.SubElement(
        etree.SubElement(etree.SubElement(separator, f"{{{WML_NS}}}p"), f"{{{WML_NS}}}r"), f"{{{WML_NS}}}separator"
    )
    continuation = etree.SubElement(root, f"{{{WML_NS}}}footnote")
    continuation.set(f"{{{WML_NS}}}type", "continuationSeparator")
    continuation.set(f"{{{WML_NS}}}id", "0")
    etree.SubElement(
        etree.SubElement(etree.SubElement(continuation, f"{{{WML_NS}}}p"), f"{{{WML_NS}}}r"),
        f"{{{WML_NS}}}continuationSeparator",
    )
    rich = etree.SubElement(root, f"{{{WML_NS}}}footnote")
    rich.set(f"{{{WML_NS}}}id", "1")
    paragraph = etree.SubElement(rich, f"{{{WML_NS}}}p")
    hyperlink = etree.SubElement(paragraph, f"{{{WML_NS}}}hyperlink")
    hyperlink.set(f"{{{_OFFICE_REL_NS}}}id", "rId3")
    text = etree.SubElement(etree.SubElement(hyperlink, f"{{{WML_NS}}}r"), f"{{{WML_NS}}}t")
    text.text = "existing hyperlink"
    drawing = etree.SubElement(etree.SubElement(paragraph, f"{{{WML_NS}}}r"), f"{{{WML_NS}}}drawing")
    blip = etree.SubElement(
        drawing,
        "{http://schemas.openxmlformats.org/drawingml/2006/main}blip",
    )
    blip.set(f"{{{_OFFICE_REL_NS}}}embed", "rId7")
    table = etree.SubElement(rich, f"{{{WML_NS}}}tbl")
    row = etree.SubElement(table, f"{{{WML_NS}}}tr")
    cell = etree.SubElement(row, f"{{{WML_NS}}}tc")
    cell_text = etree.SubElement(
        etree.SubElement(etree.SubElement(cell, f"{{{WML_NS}}}p"), f"{{{WML_NS}}}r"),
        f"{{{WML_NS}}}t",
    )
    cell_text.text = "existing table"

    if damage == "duplicate_positive":
        root.append(deepcopy(rich))
    elif damage == "malformed_id":
        rich.set(f"{{{WML_NS}}}id", "not-an-integer")
    elif damage == "missing_reserved":
        root.remove(continuation)
    elif damage == "reserved_kind":
        separator.set(f"{{{WML_NS}}}type", "continuationSeparator")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _footnote_relationships_xml(damage: str | None) -> bytes:
    nsmap = {None: _RELS_NS}
    root = etree.Element(f"{{{_RELS_NS}}}Relationships", nsmap=cast(Any, nsmap))
    hyperlink = etree.SubElement(root, f"{{{_RELS_NS}}}Relationship")
    hyperlink.set("Id", "rId3")
    hyperlink.set("Type", _HYPERLINK_REL_TYPE)
    hyperlink.set("Target", "https://example.test/existing-note")
    hyperlink.set("TargetMode", "External")
    if damage != "missing_relationship":
        image = etree.SubElement(root, f"{{{_RELS_NS}}}Relationship")
        image.set("Id", "rId7")
        image.set("Type", _IMAGE_REL_TYPE)
        image.set("Target", "media/existing-note.png")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _inject_existing_footnote_graph(path: Path, *, damage: str | None = None) -> None:
    document = Document()
    document.add_paragraph("Existing body")
    _append_footnote_reference(document, 1)
    document.save(str(path))

    members: list[tuple[ZipInfo, bytes]] = []
    with ZipFile(path, "r") as archive:
        for item in archive.infolist():
            payload = archive.read(item.filename)
            if item.filename == "word/_rels/document.xml.rels":
                root = etree.fromstring(payload)
                used = {rel.get("Id") for rel in root.findall(f"{{{_RELS_NS}}}Relationship")}
                index = 1
                while f"rId{index}" in used:
                    index += 1
                relationship = etree.SubElement(root, f"{{{_RELS_NS}}}Relationship")
                relationship.set("Id", f"rId{index}")
                relationship.set("Type", _FOOTNOTES_REL_TYPE)
                relationship.set("Target", "footnotes.xml")
                payload = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            elif item.filename == "[Content_Types].xml":
                root = etree.fromstring(payload)
                override = etree.SubElement(root, f"{{{_CT_NS}}}Override")
                override.set("PartName", "/word/footnotes.xml")
                override.set("ContentType", _FOOTNOTES_CT)
                if not any(
                    item.get("Extension", "").casefold() == "png" for item in root.findall(f"{{{_CT_NS}}}Default")
                ):
                    default = etree.SubElement(root, f"{{{_CT_NS}}}Default")
                    default.set("Extension", "png")
                    default.set("ContentType", "image/png")
                payload = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            members.append((item, payload))
    with ZipFile(path, "w", ZIP_DEFLATED) as archive:
        for item, payload in members:
            archive.writestr(item, payload)
        archive.writestr("word/footnotes.xml", _rich_footnotes_xml(damage))
        archive.writestr("word/_rels/footnotes.xml.rels", _footnote_relationships_xml(damage))
        archive.writestr("word/media/existing-note.png", _PNG)


def _inject_existing_endnote_graph(path: Path) -> None:
    document = Document()
    document.add_paragraph("Existing endnote body")
    _append_endnote_reference(document, 1)
    document.save(str(path))

    endnotes = etree.Element(f"{{{WML_NS}}}endnotes", nsmap={"w": WML_NS, "r": _OFFICE_REL_NS})
    for note_id, note_type in ((-1, "separator"), (0, "continuationSeparator")):
        note = etree.SubElement(endnotes, f"{{{WML_NS}}}endnote")
        note.set(f"{{{WML_NS}}}id", str(note_id))
        note.set(f"{{{WML_NS}}}type", note_type)
        marker = "separator" if note_id == -1 else "continuationSeparator"
        etree.SubElement(
            etree.SubElement(etree.SubElement(note, f"{{{WML_NS}}}p"), f"{{{WML_NS}}}r"),
            f"{{{WML_NS}}}{marker}",
        )
    existing = etree.SubElement(endnotes, f"{{{WML_NS}}}endnote")
    existing.set(f"{{{WML_NS}}}id", "1")
    paragraph = etree.SubElement(existing, f"{{{WML_NS}}}p")
    hyperlink = etree.SubElement(paragraph, f"{{{WML_NS}}}hyperlink")
    hyperlink.set(f"{{{_OFFICE_REL_NS}}}id", "rId4")
    etree.SubElement(etree.SubElement(hyperlink, f"{{{WML_NS}}}r"), f"{{{WML_NS}}}t").text = "existing endnote"
    gap_witness = etree.SubElement(endnotes, f"{{{WML_NS}}}endnote")
    gap_witness.set(f"{{{WML_NS}}}id", "3")
    etree.SubElement(
        etree.SubElement(etree.SubElement(gap_witness, f"{{{WML_NS}}}p"), f"{{{WML_NS}}}r"),
        f"{{{WML_NS}}}t",
    ).text = "gap witness"
    endnotes_xml = etree.tostring(endnotes, xml_declaration=True, encoding="UTF-8", standalone=True)

    nsmap = {None: _RELS_NS}
    relationships = etree.Element(f"{{{_RELS_NS}}}Relationships", nsmap=cast(Any, nsmap))
    relationship = etree.SubElement(relationships, f"{{{_RELS_NS}}}Relationship")
    relationship.set("Id", "rId4")
    relationship.set("Type", _HYPERLINK_REL_TYPE)
    relationship.set("Target", "https://example.test/existing-endnote")
    relationship.set("TargetMode", "External")
    relationship_xml = etree.tostring(relationships, xml_declaration=True, encoding="UTF-8", standalone=True)

    members: list[tuple[ZipInfo, bytes]] = []
    with ZipFile(path, "r") as archive:
        for item in archive.infolist():
            payload = archive.read(item.filename)
            if item.filename == "word/_rels/document.xml.rels":
                root = etree.fromstring(payload)
                used = {rel.get("Id") for rel in root.findall(f"{{{_RELS_NS}}}Relationship")}
                index = 1
                while f"rId{index}" in used:
                    index += 1
                owner = etree.SubElement(root, f"{{{_RELS_NS}}}Relationship")
                owner.set("Id", f"rId{index}")
                owner.set("Type", _ENDNOTES_REL_TYPE)
                owner.set("Target", "endnotes.xml")
                payload = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            elif item.filename == "[Content_Types].xml":
                root = etree.fromstring(payload)
                override = etree.SubElement(root, f"{{{_CT_NS}}}Override")
                override.set("PartName", "/word/endnotes.xml")
                override.set("ContentType", _ENDNOTES_CT)
                payload = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            members.append((item, payload))
    with ZipFile(path, "w", ZIP_DEFLATED) as archive:
        for item, payload in members:
            archive.writestr(item, payload)
        archive.writestr("word/endnotes.xml", endnotes_xml)
        archive.writestr("word/_rels/endnotes.xml.rels", relationship_xml)


def _existing_graph_snapshot(path: Path) -> tuple[bytes, tuple[tuple[str, str, str, str], ...], bytes]:
    with ZipFile(path, "r") as archive:
        footnotes = etree.fromstring(archive.read("word/footnotes.xml"))
        existing = next(
            item for item in footnotes.findall(f"{{{WML_NS}}}footnote") if item.get(f"{{{WML_NS}}}id") == "1"
        )
        relationships = etree.fromstring(archive.read("word/_rels/footnotes.xml.rels"))
        rel_records = tuple(
            sorted(
                (
                    item.get("Id", ""),
                    item.get("Type", ""),
                    item.get("Target", ""),
                    item.get("TargetMode", ""),
                )
                for item in relationships.findall(f"{{{_RELS_NS}}}Relationship")
            )
        )
        return (
            etree.tostring(existing, method="c14n", exclusive=True),
            rel_records,
            archive.read("word/media/existing-note.png"),
        )


def _footnote_ids(path: Path) -> set[int]:
    with ZipFile(path, "r") as archive:
        root = etree.fromstring(archive.read("word/footnotes.xml"))
    return {int(item.get(f"{{{WML_NS}}}id", "0")) for item in root.findall(f"{{{WML_NS}}}footnote")}


def _body_footnote_ids(path: Path) -> set[int]:
    with ZipFile(path, "r") as archive:
        root = etree.fromstring(archive.read("word/document.xml"))
    return {int(item.get(f"{{{WML_NS}}}id", "0")) for item in root.iter(f"{{{WML_NS}}}footnoteReference")}


def _endnote_ids(path: Path) -> set[int]:
    with ZipFile(path, "r") as archive:
        root = etree.fromstring(archive.read("word/endnotes.xml"))
    return {int(item.get(f"{{{WML_NS}}}id", "0")) for item in root.findall(f"{{{WML_NS}}}endnote")}


def _body_endnote_ids(path: Path) -> set[int]:
    with ZipFile(path, "r") as archive:
        root = etree.fromstring(archive.read("word/document.xml"))
    return {int(item.get(f"{{{WML_NS}}}id", "0")) for item in root.iter(f"{{{WML_NS}}}endnoteReference")}


def test_existing_rich_note_graph_survives_collision_free_append_and_save_reopen(tmp_path: Path) -> None:
    path = tmp_path / "rich-existing.docx"
    _inject_existing_footnote_graph(path)
    before = _existing_graph_snapshot(path)

    document = Document(str(path))
    context = NoteContext()
    context._footnote_children = {"new": _note_children("new note")}
    prepare_note_context_for_document(document, context)
    reference = context.create_footnote_ref_run("new")
    assert reference is not None
    assert context.get_footnote_word_id("new") == 2
    document.paragraphs[0]._p.append(reference)
    document.save(str(path))
    write_notes_to_docx(str(path), context)

    assert _footnote_ids(path) == {-1, 0, 1, 2}
    assert _body_footnote_ids(path) == {1, 2}
    assert _existing_graph_snapshot(path) == before

    reopened_path = tmp_path / "rich-reopened.docx"
    Document(str(path)).save(str(reopened_path))
    assert _footnote_ids(reopened_path) == {-1, 0, 1, 2}
    assert _body_footnote_ids(reopened_path) == {1, 2}
    assert _existing_graph_snapshot(reopened_path) == before


def test_existing_endnote_domain_allocates_lowest_gap_and_preserves_relationship(tmp_path: Path) -> None:
    path = tmp_path / "existing-endnotes.docx"
    _inject_existing_endnote_graph(path)
    with ZipFile(path, "r") as archive:
        before_relationships = etree.fromstring(archive.read("word/_rels/endnotes.xml.rels"))

    document = Document(str(path))
    context = NoteContext()
    context._endnote_children = {"new": _note_children("new endnote")}
    prepare_note_context_for_document(document, context)
    reference = context.create_endnote_ref_run("new")
    assert reference is not None
    assert context.get_endnote_word_id("new") == 2
    document.paragraphs[0]._p.append(reference)
    document.save(str(path))
    write_notes_to_docx(str(path), context)

    assert _endnote_ids(path) == {-1, 0, 1, 2, 3}
    assert _body_endnote_ids(path) == {1, 2}
    with ZipFile(path, "r") as archive:
        after_relationships = etree.fromstring(archive.read("word/_rels/endnotes.xml.rels"))
    assert etree.tostring(after_relationships, method="c14n", exclusive=True) == etree.tostring(
        before_relationships,
        method="c14n",
        exclusive=True,
    )


def test_converter_preserves_existing_graph_and_allocates_matching_body_and_part_id(tmp_path: Path) -> None:
    template = tmp_path / "rich-template.docx"
    _inject_existing_footnote_graph(template)
    template_hash = hashlib.sha256(template.read_bytes()).hexdigest()
    before = _existing_graph_snapshot(template)
    source = tmp_path / "source.md"
    source.write_text("Body with a new note[^new].\n\n[^new]: New body.\n", encoding="utf-8")
    context, _workspace = make_context(
        str(source),
        target_format="docx",
        options={"template_name": str(template)},
    )

    result = MdToDocxConverter().convert(context)

    assert result.success, result.error
    output = Path(result.artifacts[0].staging_path)
    assert _footnote_ids(output) == {-1, 0, 1, 2}
    assert _body_footnote_ids(output) == {1, 2}
    assert _existing_graph_snapshot(output) == before
    assert hashlib.sha256(template.read_bytes()).hexdigest() == template_hash
    styles = {style.style_id for style in Document(str(output)).styles}
    assert {"FootnoteText", "FootnoteReference", "EndnoteText", "EndnoteReference"} <= styles


@pytest.mark.parametrize(
    "damage",
    ("duplicate_positive", "malformed_id", "missing_reserved", "reserved_kind", "missing_relationship"),
)
def test_converter_fails_closed_before_artifact_for_malformed_existing_note_graph(
    tmp_path: Path,
    damage: str,
) -> None:
    template = tmp_path / f"damaged-{damage}.docx"
    _inject_existing_footnote_graph(template, damage=damage)
    template_hash = hashlib.sha256(template.read_bytes()).hexdigest()
    source = tmp_path / "source.md"
    source.write_text("Body[^new].\n\n[^new]: New note.\n", encoding="utf-8")
    context, workspace = make_context(
        str(source),
        target_format="docx",
        options={"template_name": str(template)},
    )

    result = MdToDocxConverter().convert(context)

    assert result.success is False
    assert result.artifacts == []
    assert result.error is not None
    assert result.error.diagnostic_code == "MD2DOCX-NOTE-PART-INVALID"
    assert [item.code for item in result.diagnostics] == ["MD2DOCX-NOTE-PART-INVALID"]
    assert list(Path(workspace.staging_dir).glob("*.docx")) == []
    assert hashlib.sha256(template.read_bytes()).hexdigest() == template_hash


def test_unprepared_collision_fails_without_replacing_request_owned_package(tmp_path: Path) -> None:
    path = tmp_path / "collision.docx"
    _inject_existing_footnote_graph(path)
    document = Document(str(path))
    context = NoteContext()
    context._footnote_children = {"new": _note_children("new note")}
    reference = context.create_footnote_ref_run("new")
    assert reference is not None
    document.paragraphs[0]._p.append(reference)
    document.save(str(path))
    before = hashlib.sha256(path.read_bytes()).hexdigest()

    with pytest.raises(NoteWritebackError) as raised:
        write_notes_to_docx(str(path), context)

    assert raised.value.diagnostic_code == "MD2DOCX-NOTE-PART-INVALID"
    assert hashlib.sha256(path.read_bytes()).hexdigest() == before
