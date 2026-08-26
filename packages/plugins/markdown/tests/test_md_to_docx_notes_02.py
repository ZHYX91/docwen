"""Focused tests split from test_md_to_docx_notes.py."""

from __future__ import annotations

from ._md_to_docx_notes_support import (
    MD_ENDNOTE_ONLY,
    MD_FOOTNOTE_ONLY,
    MD_MULTI_FOOTNOTES,
    MD_NO_NOTES,
    MD_WITH_FOOTNOTES_ONLY,
    WML_NS,
    Document,
    MdToDocxRenderer,
    NoteContext,
    Path,
    _list_zip_entries,
    _mk_children,
    _q,
    _read_zip_entry,
    _render_plain,
    etree,
    extract_notes_from_ast,
    parse_markdown_text,
    process_md_body_with_notes,
    pytest,
)

pytestmark = pytest.mark.contract


class TestRendererFoundation:
    """F-F1-027: MdToDocxRenderer handles footnote_ref AST nodes."""

    def test_foundation_footnote_ref_creates_ooxml_reference(self):
        """[^id] in markdown produces a Word footnote reference, not literal text."""
        md = "# Test\n\nText [^1] here.\n\n[^1]: A footnote."

        cleaned_ast, note_ctx = process_md_body_with_notes(md)

        doc = Document()
        renderer = MdToDocxRenderer(doc, note_ctx=note_ctx)
        renderer.render(cleaned_ast)

        # The footnote_ref should NOT appear as literal [^1] text
        body_text = _render_plain(doc)
        assert "[^1]" not in body_text

        # The footnote definition should NOT appear in body
        assert "A footnote." not in body_text

        # The surrounding text should be present
        assert "Text" in body_text
        assert "here" in body_text

        # There should be OOXML footnote references in the document
        # Check by examining paragraph XML
        ref_count = 0
        for p in doc.paragraphs:
            for elem in p._p.iter():
                if elem.tag == _q("footnoteReference"):
                    ref_count += 1
        assert ref_count == 1, f"Expected 1 footnoteReference, got {ref_count}"

    def test_foundation_endnote_ref_creates_ooxml_reference(self):
        """[^endnote:id] produces an endnote reference."""
        md = "# Test\n\nEndnote[^endnote:e1] ref.\n\n[^endnote:e1]: End def."

        cleaned_ast, note_ctx = process_md_body_with_notes(md)

        doc = Document()
        renderer = MdToDocxRenderer(doc, note_ctx=note_ctx)
        renderer.render(cleaned_ast)

        body_text = _render_plain(doc)
        assert "End def." not in body_text

        en_refs = 0
        for p in doc.paragraphs:
            for elem in p._p.iter():
                if elem.tag == _q("endnoteReference"):
                    en_refs += 1
        assert en_refs == 1

    def test_foundation_undefined_note_falls_back_to_text(self):
        """Undefined [^id] renders as literal text in the output."""
        # This markdown has a [^ref] that has NO definition.
        # Mistune will NOT parse it as footnote_ref (since no
        # definition exists). The literal '[^ref]' text remains.
        md = "# Test\n\nUndefined [^ref] text."

        ast = parse_markdown_text(md)
        cleaned_ast, note_ctx = extract_notes_from_ast(ast)

        assert not note_ctx.has_notes

        doc = Document()
        renderer = MdToDocxRenderer(doc, note_ctx=note_ctx)
        renderer.render(cleaned_ast)

        body_text = _render_plain(doc)
        # Without a definition, mistune treats [^ref] as literal text
        assert "[^ref]" in body_text

    def test_foundation_multiple_footnotes_in_one_paragraph(self):
        """Multiple [^id] references in one paragraph all produce ref runs."""
        md = "# M\n\nA[^1] B[^2] C[^3].\n\n[^1]: one\n[^2]: two\n[^3]: three"

        ast = parse_markdown_text(md)
        cleaned_ast, note_ctx = extract_notes_from_ast(ast)

        doc = Document()
        renderer = MdToDocxRenderer(doc, note_ctx=note_ctx)
        renderer.render(cleaned_ast)

        ref_count = 0
        for p in doc.paragraphs:
            for elem in p._p.iter():
                if elem.tag == _q("footnoteReference"):
                    ref_count += 1
        assert ref_count == 3, f"Expected 3 footnote references, got {ref_count}"

    def test_foundation_note_body_elements_available(self):
        """NoteContext stores created body elements for downstream write-back."""
        md = "# T\n\n[^a] text.\n\n[^a]: note content"

        ast = parse_markdown_text(md)
        cleaned_ast, note_ctx = extract_notes_from_ast(ast)

        doc = Document()
        renderer = MdToDocxRenderer(doc, note_ctx=note_ctx)
        renderer.render(cleaned_ast)

        # Renderer should have triggered body element creation via
        # NoteContext when creating the reference run
        assert len(note_ctx.footnote_elements) == 1

        elem = note_ctx.footnote_elements[0]
        assert elem.tag == _q("footnote")
        assert elem.get(_q("id")) == "1"

    def test_foundation_no_zip_operations_in_model(self):
        """NoteContext and OOXML factories do not open/manipulate ZIP files."""
        ctx = NoteContext()
        ctx._footnote_children = {"n": _mk_children("content")}
        ctx.get_footnote_word_id("n")

        # The model is pure in-memory; no ZIP I/O
        # Verification: the code path above does not raise,
        # and no file-system side effects occur.

    def test_foundation_skip_footnotes_node_in_body(self):
        """The footnotes/footnote_item nodes are skipped, not rendered in body."""
        md = "# T\n\nText[^f].\n\n[^f]: hidden definition"

        ast = parse_markdown_text(md)
        cleaned_ast, note_ctx = extract_notes_from_ast(ast)

        doc = Document()
        renderer = MdToDocxRenderer(doc, note_ctx=note_ctx)
        renderer.render(cleaned_ast)

        # Definition text must NOT leak into body
        body_text = _render_plain(doc)
        assert "hidden definition" not in body_text

    def test_foundation_end_to_end_md_to_docx_with_notes(self):
        """Full end-to-end: MD with notes → DOCX with proper references."""
        from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter

        from .conftest import make_context, write_temp_md

        md_path = write_temp_md(MD_WITH_FOOTNOTES_ONLY)
        ctx, _workspace = make_context(md_path, target_format="docx")

        converter = MdToDocxConverter()
        result = converter.convert(ctx)

        assert result.success, f"Conversion failed: {result.error}"
        output_path = Path(result.artifacts[0].staging_path)
        assert output_path.exists()

        # Open and inspect
        from docx import Document

        doc = Document(str(output_path))

        # Footnote references present
        ref_count = 0
        for p in doc.paragraphs:
            for elem in p._p.iter():
                if elem.tag == _q("footnoteReference"):
                    ref_count += 1
        assert ref_count == 1, f"Expected 1 footnoteReference, got {ref_count}"

        # No literal [^a] in output
        body = _render_plain(doc)
        assert "[^a]" not in body

    @pytest.mark.parametrize(
        "markdown",
        [
            "Missing[^x].\n",
            "A[^x].\n\n[^x]: one\n[^x]: two\n",
            "A[^x] B[^footnote:x].\n\n[^x]: one\n[^footnote:x]: two\n",
        ],
    )
    def test_full_converter_rejects_invalid_note_graph_without_rewriting_source(self, markdown: str):
        from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter

        from .conftest import make_context, write_temp_md

        md_path = write_temp_md(markdown)
        source_path = Path(md_path)
        original = source_path.read_bytes()
        ctx, _workspace = make_context(md_path, target_format="docx")

        result = MdToDocxConverter().convert(ctx)

        assert not result.success
        assert result.error is not None
        assert result.error.diagnostic_code == "MD2DOCX-NOTE-SYNTAX-INVALID"
        assert [item.code for item in result.diagnostics] == ["MD2DOCX-NOTE-SYNTAX-INVALID"]
        assert source_path.read_bytes() == original


class TestWritebackFootnotes:
    """F-F1-028: write_notes_to_docx writes footnote elements into
    word/footnotes.xml of the DOCX ZIP."""

    def test_writeback_footnote_part_created(self):
        """A fresh DOCX (no pre-existing footnotes part) gets
        word/footnotes.xml with separator + user content."""
        from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter

        from .conftest import make_context, write_temp_md

        md_path = write_temp_md(MD_FOOTNOTE_ONLY)
        ctx, _workspace = make_context(md_path, target_format="docx")

        converter = MdToDocxConverter()
        result = converter.convert(ctx)
        assert result.success, f"Conversion failed: {result.error}"
        output_path = str(Path(result.artifacts[0].staging_path))

        entries = _list_zip_entries(output_path)
        assert "word/footnotes.xml" in entries, f"Expected word/footnotes.xml, got: {sorted(entries)}"

    def test_writeback_footnote_content_correct(self):
        """The footnotes.xml part contains the separator, continuation,
        and user footnote with correct content."""
        from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter

        from .conftest import make_context, write_temp_md

        md_path = write_temp_md(MD_FOOTNOTE_ONLY)
        ctx, _workspace = make_context(md_path, target_format="docx")

        converter = MdToDocxConverter()
        result = converter.convert(ctx)
        output_path = str(Path(result.artifacts[0].staging_path))

        fn_xml = _read_zip_entry(output_path, "word/footnotes.xml")
        root = etree.fromstring(fn_xml)

        # Should have separator (id=-1), continuation (id=0), and user (id=1)
        footnotes = root.findall(f"{{{WML_NS}}}footnote")
        assert len(footnotes) >= 3, f"Expected >=3 footnote elements, got {len(footnotes)}"

        ids = {fn.get(_q("id")) for fn in footnotes}
        assert "-1" in ids, "Missing separator footnote (id=-1)"
        assert "0" in ids, "Missing continuation separator (id=0)"
        assert "1" in ids, "Missing user footnote (id=1)"

        # Verify user footnote content
        user_fn = root.find(f"{{{WML_NS}}}footnote[@{{{WML_NS}}}id='1']")
        assert user_fn is not None
        texts = [t.text or "" for t in user_fn.iter(f"{{{WML_NS}}}t")]
        combined = "".join(texts)
        assert "Footnote body text here" in combined, f"Expected footnote body in: {combined}"

    def test_writeback_footnote_ref_in_body(self):
        """Body paragraphs contain footnoteReference OOXML elements."""
        from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter

        from .conftest import make_context, write_temp_md

        md_path = write_temp_md(MD_FOOTNOTE_ONLY)
        ctx, _workspace = make_context(md_path, target_format="docx")

        converter = MdToDocxConverter()
        result = converter.convert(ctx)
        output_path = str(Path(result.artifacts[0].staging_path))

        doc_xml = _read_zip_entry(output_path, "word/document.xml")
        root = etree.fromstring(doc_xml)
        refs = root.findall(f".//{{{WML_NS}}}footnoteReference")
        assert len(refs) >= 1, f"Expected >=1 footnoteReference, got {len(refs)}"

    def test_writeback_docx_openable_by_python_docx(self):
        """The generated DOCX can be opened by python-docx without error."""
        from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter

        from .conftest import make_context, write_temp_md

        md_path = write_temp_md(MD_FOOTNOTE_ONLY)
        ctx, _workspace = make_context(md_path, target_format="docx")

        converter = MdToDocxConverter()
        result = converter.convert(ctx)
        output_path = str(Path(result.artifacts[0].staging_path))

        # This must not raise
        from docx import Document

        doc = Document(output_path)
        assert doc is not None

        # Body text preserved (no def leak)
        body = _render_plain(doc)
        assert "Footnote body text here" not in body, "Footnote definition leaked into body"
        assert "Para with a footnote" in body

    def test_writeback_multiple_footnotes(self):
        """Multiple footnotes all get written to the part with correct IDs."""
        from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter

        from .conftest import make_context, write_temp_md

        md_path = write_temp_md(MD_MULTI_FOOTNOTES)
        ctx, _workspace = make_context(md_path, target_format="docx")

        converter = MdToDocxConverter()
        result = converter.convert(ctx)
        output_path = str(Path(result.artifacts[0].staging_path))

        fn_xml = _read_zip_entry(output_path, "word/footnotes.xml")
        root = etree.fromstring(fn_xml)

        footnotes = root.findall(f"{{{WML_NS}}}footnote")
        # separator + continuation + 3 user = 5
        assert len(footnotes) == 5, f"Expected 5 footnotes, got {len(footnotes)}"

        ids = {fn.get(_q("id")) for fn in footnotes}
        assert {"-1", "0", "1", "2", "3"} == ids, f"Unexpected IDs: {ids}"

        # Each user footnote has content
        for expected_id in ("1", "2", "3"):
            fn = root.find(f"{{{WML_NS}}}footnote[@{{{WML_NS}}}id='{expected_id}']")
            assert fn is not None, f"Missing footnote id={expected_id}"
            texts = [t.text or "" for t in fn.iter(f"{{{WML_NS}}}t")]
            assert any("note" in t.lower() or "multiline" in t.lower() for t in texts), (
                f"Footnote {expected_id} has no content"
            )

    def test_writeback_no_notes_no_parts(self):
        """When MD has no notes, footnotes.xml / endnotes.xml are NOT added."""
        from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter

        from .conftest import make_context, write_temp_md

        md_path = write_temp_md(MD_NO_NOTES)
        ctx, _workspace = make_context(md_path, target_format="docx")

        converter = MdToDocxConverter()
        result = converter.convert(ctx)
        output_path = str(Path(result.artifacts[0].staging_path))

        entries = _list_zip_entries(output_path)
        # WORD should not create these by default
        assert "word/footnotes.xml" not in entries
        assert "word/endnotes.xml" not in entries


class TestWritebackEndnotes:
    """F-F1-028: write_notes_to_docx writes endnote elements into
    word/endnotes.xml of the DOCX ZIP."""

    def test_writeback_endnote_part_created(self):
        """A fresh DOCX gets word/endnotes.xml with separator + user content."""
        from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter

        from .conftest import make_context, write_temp_md

        md_path = write_temp_md(MD_ENDNOTE_ONLY)
        ctx, _workspace = make_context(md_path, target_format="docx")

        converter = MdToDocxConverter()
        result = converter.convert(ctx)
        output_path = str(Path(result.artifacts[0].staging_path))

        entries = _list_zip_entries(output_path)
        assert "word/endnotes.xml" in entries, f"Expected word/endnotes.xml, got: {sorted(entries)}"

    def test_writeback_endnote_content_correct(self):
        """The endnotes.xml part has separator, continuation, and user endnote."""
        from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter

        from .conftest import make_context, write_temp_md

        md_path = write_temp_md(MD_ENDNOTE_ONLY)
        ctx, _workspace = make_context(md_path, target_format="docx")

        converter = MdToDocxConverter()
        result = converter.convert(ctx)
        output_path = str(Path(result.artifacts[0].staging_path))

        en_xml = _read_zip_entry(output_path, "word/endnotes.xml")
        root = etree.fromstring(en_xml)

        endnotes = root.findall(f"{{{WML_NS}}}endnote")
        assert len(endnotes) >= 3, f"Expected >=3 endnote elements, got {len(endnotes)}"

        ids = {en.get(_q("id")) for en in endnotes}
        assert "-1" in ids
        assert "0" in ids
        assert "1" in ids

        user_en = root.find(f"{{{WML_NS}}}endnote[@{{{WML_NS}}}id='1']")
        assert user_en is not None
        texts = [t.text or "" for t in user_en.iter(f"{{{WML_NS}}}t")]
        combined = "".join(texts)
        assert "Endnote body text here" in combined

    def test_writeback_endnote_ref_in_body(self):
        """Body paragraphs contain endnoteReference OOXML elements."""
        from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter

        from .conftest import make_context, write_temp_md

        md_path = write_temp_md(MD_ENDNOTE_ONLY)
        ctx, _workspace = make_context(md_path, target_format="docx")

        converter = MdToDocxConverter()
        result = converter.convert(ctx)
        output_path = str(Path(result.artifacts[0].staging_path))

        doc_xml = _read_zip_entry(output_path, "word/document.xml")
        root = etree.fromstring(doc_xml)
        refs = root.findall(f".//{{{WML_NS}}}endnoteReference")
        assert len(refs) >= 1, f"Expected >=1 endnoteReference, got {len(refs)}"

    def test_writeback_endnote_docx_openable(self):
        """Generated DOCX with endnotes is openable by python-docx."""
        from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter

        from .conftest import make_context, write_temp_md

        md_path = write_temp_md(MD_ENDNOTE_ONLY)
        ctx, _workspace = make_context(md_path, target_format="docx")

        converter = MdToDocxConverter()
        result = converter.convert(ctx)
        output_path = str(Path(result.artifacts[0].staging_path))

        from docx import Document

        doc = Document(output_path)
        assert doc is not None
        body = _render_plain(doc)
        assert "Endnote body text here" not in body
