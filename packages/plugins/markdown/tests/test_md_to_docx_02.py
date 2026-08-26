"""Focused tests split from test_md_to_docx.py."""

from __future__ import annotations

from ._md_to_docx_support import (
    SAMPLE_MD_CONTENT,
    FakeConfigView,
    MdToDocxConverter,
    Path,
    make_context,
    pytest,
    write_temp_md,
    zipfile,
)

pytestmark = pytest.mark.contract


class TestMdToDocxGolden:
    """Golden-level tests for MD → DOCX conversion."""

    @staticmethod
    def _verify_docx_structure(output_path: Path):
        """Verify the file is a well-formed DOCX (ZIP with expected entries)."""
        import zipfile

        assert zipfile.is_zipfile(output_path), "Output is not a valid ZIP file"
        with zipfile.ZipFile(output_path, "r") as zf:
            names = zf.namelist()
            assert "[Content_Types].xml" in names, "Missing [Content_Types].xml"
            assert any("word/document.xml" in n for n in names), "Missing word/document.xml"

    @staticmethod
    def _verify_docx_content(output_path: Path):
        """Verify the DOCX contains expected heading and paragraph structure."""
        from docx import Document

        doc = Document(str(output_path))

        paragraphs = doc.paragraphs
        assert len(paragraphs) >= 5, f"Expected at least 5 paragraphs, got {len(paragraphs)}"

        # Check heading styles
        heading_styles = [p.style.name for p in paragraphs if p.style and p.style.name and "Heading" in p.style.name]
        assert len(heading_styles) >= 2, f"Expected at least 2 headings, got {len(heading_styles)}"

    def test_yaml_placeholder_fill_preserves_surrounding_text_and_replaces_duplicates(self) -> None:
        from docx import Document

        from docwen_plugin_markdown.template_filler import fill_yaml_placeholders
        from docwen_plugin_markdown.template_utils import scan_placeholders

        document = Document()
        paragraph = document.add_paragraph()
        prefix = paragraph.add_run("Prefix ")
        prefix.bold = True
        paragraph.add_run("{{na")
        paragraph.add_run("me}}")
        suffix = paragraph.add_run(" suffix")
        suffix.italic = True
        paragraph.add_run(" / {{name}}")
        repeated = document.add_paragraph("Again {{ name }}!")
        table_paragraph = document.add_table(rows=1, cols=1).cell(0, 0).paragraphs[0]
        table_paragraph.add_run("Cell {{name}} tail")

        placeholder_map = scan_placeholders(document)

        assert len(placeholder_map["name"]) == 3
        fill_yaml_placeholders(document, {"name": "Alice"}, placeholder_map)
        assert paragraph.text == "Prefix Alice suffix / Alice"
        assert repeated.text == "Again Alice!"
        assert table_paragraph.text == "Cell Alice tail"
        assert prefix.bold is True
        assert suffix.italic is True

    def test_body_placeholder_direct_format_is_projected_only_to_body_paragraphs(self, tmp_path: Path) -> None:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml.ns import qn
        from docx.shared import Twips

        template_path = tmp_path / "direct-body-format.docx"
        template = Document()
        placeholder = template.add_paragraph("{{body}}")
        paragraph_format = placeholder.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph_format.left_indent = Twips(620)
        paragraph_format.right_indent = Twips(340)
        paragraph_format.first_line_indent = Twips(180)
        paragraph_format.space_before = Twips(240)
        paragraph_format.space_after = Twips(300)
        paragraph_format.line_spacing = 1.5
        paragraph_format.keep_together = True
        paragraph_format.keep_with_next = True
        paragraph_format.page_break_before = True
        paragraph_format.widow_control = False
        template.save(str(template_path))
        markdown_path = write_temp_md("# Heading\n\nBody one.\n\nBody two.\n\n- List item\n")
        context, _workspace = make_context(
            markdown_path,
            target_format="docx",
            options={"template_name": str(template_path)},
        )

        result = MdToDocxConverter().convert(context)

        assert result.success, result.error
        output = Document(result.artifacts[0].staging_path)
        for text in ("Body one.", "Body two."):
            body = next(paragraph for paragraph in output.paragraphs if paragraph.text == text)
            actual = body.paragraph_format
            assert actual.alignment == WD_ALIGN_PARAGRAPH.RIGHT
            assert actual.left_indent == Twips(620)
            assert actual.right_indent == Twips(340)
            assert actual.first_line_indent == Twips(180)
            assert actual.space_before == Twips(240)
            assert actual.space_after == Twips(300)
            assert actual.line_spacing == 1.5
            assert actual.line_spacing_rule == WD_LINE_SPACING.ONE_POINT_FIVE
            assert actual.keep_together is True
            assert actual.keep_with_next is True
            assert actual.page_break_before is True
            assert actual.widow_control is False

        heading = next(paragraph for paragraph in output.paragraphs if paragraph.text == "Heading")
        list_item = next(paragraph for paragraph in output.paragraphs if paragraph.text == "List item")
        assert heading.paragraph_format.alignment is None
        assert list_item.paragraph_format.alignment is None
        assert list_item._p.find(f"{qn('w:pPr')}/{qn('w:numPr')}") is not None

    def test_loads_content_validated_docx_template_with_wrong_suffix(self, tmp_path: Path) -> None:
        from docx import Document

        template_path = tmp_path / "renamed-template.xlsx"
        template = Document()
        template.add_paragraph("{{正文}}")
        template.save(str(template_path))
        md_path = write_temp_md("# Wrong suffix template\n\nBody")
        context, _workspace = make_context(
            md_path,
            target_format="docx",
            options={"template_name": str(template_path)},
        )

        result = MdToDocxConverter().convert(context)

        assert result.success is True
        output = Document(result.artifacts[0].staging_path)
        assert "Body" in "\n".join(paragraph.text for paragraph in output.paragraphs)

    def test_converts_sample_md_to_docx(self):
        """Convert a comprehensive Markdown sample to DOCX and verify output."""
        md_path = write_temp_md(SAMPLE_MD_CONTENT)
        ctx, workspace = make_context(md_path, target_format="docx")

        converter = MdToDocxConverter()
        result = converter.convert(ctx)

        assert result.success is True, f"Conversion failed: {result.error.message if result.error else 'unknown'}"
        assert len(result.artifacts) == 1
        artifact = result.artifacts[0]
        assert artifact.kind == "primary"
        assert artifact.is_primary is True
        assert artifact.media_type == ("application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        assert artifact.suggested_name.endswith(".docx")

        # Verify the DOCX file was written and is non-empty
        output_path = Path(artifact.staging_path)
        assert output_path.exists(), f"Output file missing: {output_path}"
        assert output_path.stat().st_size > 0, "Output file is empty"

        # Verify it's a valid DOCX (ZIP-based format with expected structure)
        self._verify_docx_structure(output_path)

        # Verify content: headings, paragraphs, lists, tables, etc.
        self._verify_docx_content(output_path)

        # Verify metrics
        assert result.metrics.duration_ms > 0
        assert result.metrics.input_bytes > 0
        assert result.metrics.output_bytes > 0

        # Verify the artifact was registered in workspace
        assert len(workspace.registered_artifacts) == 1

    def test_docx_contains_headings(self):
        """DOCX output contains the headings from the Markdown."""
        md_path = write_temp_md("# Test Heading\n\nSome text.")
        ctx, _workspace = make_context(md_path, target_format="docx")

        converter = MdToDocxConverter()
        result = converter.convert(ctx)

        assert result.success
        output_path = Path(result.artifacts[0].staging_path)
        assert output_path.exists()

        from docx import Document

        doc = Document(str(output_path))
        texts = [p.text for p in doc.paragraphs]
        assert any("Test Heading" in t for t in texts), f"Heading 'Test Heading' not found in DOCX paragraphs: {texts}"

    def test_docx_contains_paragraph_text(self):
        """DOCX output preserves paragraph text including inline formatting."""
        md_path = write_temp_md("# Title\n\nSome **bold** and *italic* text.")
        ctx, _workspace = make_context(md_path, target_format="docx")

        converter = MdToDocxConverter()
        result = converter.convert(ctx)

        assert result.success
        output_path = Path(result.artifacts[0].staging_path)

        from docx import Document

        doc = Document(str(output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "bold" in full_text, f"'bold' not found in: {full_text}"
        assert "italic" in full_text, f"'italic' not found in: {full_text}"

    def test_docx_contains_table(self):
        """DOCX output preserves Markdown tables."""
        md_content = "# Tables\n\n| A | B |\n|---|---|\n| 1 | 2 |\n"
        md_path = write_temp_md(md_content)
        ctx, _workspace = make_context(md_path, target_format="docx")

        converter = MdToDocxConverter()
        result = converter.convert(ctx)

        assert result.success
        output_path = Path(result.artifacts[0].staging_path)

        from docx import Document

        doc = Document(str(output_path))
        assert len(doc.tables) >= 1, "No tables found in DOCX output"

        table = doc.tables[0]
        # Check header row
        header_cells = [cell.text for cell in table.rows[0].cells]
        assert "A" in header_cells
        assert "B" in header_cells

    def test_docx_contains_lists(self):
        """DOCX output preserves Markdown lists."""
        md_content = "# Lists\n\n- Item A\n- Item B\n"
        md_path = write_temp_md(md_content)
        ctx, _workspace = make_context(md_path, target_format="docx")

        converter = MdToDocxConverter()
        result = converter.convert(ctx)

        assert result.success
        output_path = Path(result.artifacts[0].staging_path)

        from docx import Document

        doc = Document(str(output_path))
        texts = [p.text for p in doc.paragraphs]
        combined = " ".join(texts)
        assert "Item A" in combined, f"List item not found in: {combined}"

    def test_conversion_diagnostics(self):
        """Conversion result includes success diagnostics."""
        md_path = write_temp_md("# Title\n\nContent.")
        ctx, _workspace = make_context(md_path, target_format="docx")

        converter = MdToDocxConverter()
        result = converter.convert(ctx)

        assert result.success
        assert len(result.diagnostics) >= 1
        assert result.diagnostics[0].level == "info"
        assert result.diagnostics[0].code == "MD2DOCX-OK"

    def test_md_to_docx_runs_configured_yaml_field_processors(self, tmp_path, monkeypatch):
        """Configured field processors mutate YAML before template/metadata fill."""
        module_dir = tmp_path / "mods"
        module_dir.mkdir()
        (module_dir / "metadata_title_processor.py").write_text(
            "def process_yaml(data):\n    data['标题'] = 'Processed Title'\n",
            encoding="utf-8",
        )
        monkeypatch.syspath_prepend(str(module_dir))

        md_path = write_temp_md("---\n标题: Raw Title\n---\n# Body\n")
        ctx, _workspace = make_context(md_path, target_format="docx")
        ctx._config = FakeConfigView(
            {
                "gui": {"language": {"locale": "zh_CN"}},
                "field_processors": {
                    "settings": {"order": ["metadata"]},
                    "processors": {
                        "metadata": {
                            "module": "metadata_title_processor",
                            "locales": ["zh_CN"],
                            "enabled": True,
                        }
                    },
                },
            }
        )

        result = MdToDocxConverter().convert(ctx)

        assert result.success is True
        with zipfile.ZipFile(Path(result.artifacts[0].staging_path), "r") as zf:
            assert zf.namelist().count("docProps/core.xml") == 1
            core_xml = zf.read("docProps/core.xml").decode("utf-8")
        assert "Processed Title" in core_xml

    def test_md_to_docx_applies_gongwen_placeholder_cleanup_rules(self, tmp_path):
        """Gongwen field processors also restore old template cleanup rules."""
        from docx import Document

        template_path = tmp_path / "gongwen-template.docx"
        template = Document()
        template.add_paragraph("{{正文}}")
        template.add_paragraph("{{附件说明}}")
        table = template.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = "字段"
        table.rows[0].cells[1].text = "内容"
        table.rows[1].cells[0].text = "抄送机关"
        table.rows[1].cells[1].text = "{{抄送机关}}"
        table.rows[2].cells[0].text = "版记"
        table.rows[2].cells[1].text = "{{印发机关}}{{印发日期}}"
        template.save(template_path)

        md_path = write_temp_md(
            "---\n"
            "标题: 公文标题\n"
            "附件说明:\n"
            "  - 1. 材料清单\n"
            "  - （二）办理依据\n"
            "抄送机关: []\n"
            "印发机关: ''\n"
            "印发日期: ''\n"
            "---\n"
            "# 正文标题\n"
            "正文内容。\n"
        )
        ctx, _workspace = make_context(
            md_path,
            target_format="docx",
            options={"template_name": str(template_path)},
        )
        ctx._config = FakeConfigView(
            {
                "gui": {"language": {"locale": "zh_CN"}},
                "field_processors": {
                    "settings": {"order": ["gongwen"]},
                    "processors": {
                        "gongwen": {
                            "module": "docwen_plugin_markdown.field_processors.gongwen",
                            "locales": ["zh_CN"],
                            "enabled": True,
                        }
                    },
                },
            }
        )

        result = MdToDocxConverter().convert(ctx)

        assert result.success is True
        output_path = Path(result.artifacts[0].staging_path)
        doc = Document(str(output_path))
        attachment_paragraphs = [
            paragraph
            for paragraph in doc.paragraphs
            if "附件：1. 材料清单" in paragraph.text or "2. 办理依据" in paragraph.text
        ]
        document_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        assert len(attachment_paragraphs) == 2
        assert "附件：1. 材料清单" in document_text
        assert "2. 办理依据" in document_text
        assert "['附件" not in document_text
        for paragraph in attachment_paragraphs:
            left_indent = paragraph.paragraph_format.left_indent
            first_line_indent = paragraph.paragraph_format.first_line_indent
            assert left_indent is not None
            assert left_indent > 0
            assert first_line_indent is not None
            assert first_line_indent < 0
        assert doc.tables
        table_text = "\n".join(cell.text for row in doc.tables[0].rows for cell in row.cells)
        assert "字段" in table_text
        assert "抄送机关" not in table_text
        assert "版记" not in table_text
        assert "{{" not in table_text

    def test_md_to_docx_gongwen_attachment_in_table_falls_back_to_yaml_fill(self, tmp_path):
        """Attachment special handling should not block generic table-cell fill."""
        from docx import Document

        template_path = tmp_path / "gongwen-table-attachment-template.docx"
        template = Document()
        template.add_paragraph("{{正文}}")
        table = template.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "附件"
        table.rows[0].cells[1].text = "{{附件说明}}"
        template.save(template_path)

        md_path = write_temp_md(
            "---\n标题: 公文标题\n附件说明:\n  - 1. 材料清单\n  - （二）办理依据\n---\n# 正文标题\n正文内容。\n"
        )
        ctx, _workspace = make_context(
            md_path,
            target_format="docx",
            options={"template_name": str(template_path)},
        )
        ctx._config = FakeConfigView(
            {
                "gui": {"language": {"locale": "zh_CN"}},
                "field_processors": {
                    "settings": {"order": ["gongwen"]},
                    "processors": {
                        "gongwen": {
                            "module": "docwen_plugin_markdown.field_processors.gongwen",
                            "locales": ["zh_CN"],
                            "enabled": True,
                        }
                    },
                },
            }
        )

        result = MdToDocxConverter().convert(ctx)

        assert result.success is True
        doc = Document(str(result.artifacts[0].staging_path))
        table_text = "\n".join(cell.text for row in doc.tables[0].rows for cell in row.cells)
        assert "附件：1. 材料清单" in table_text
        assert "2. 办理依据" in table_text
        assert "['附件" not in table_text
        assert "{{附件说明}}" not in table_text

    @pytest.mark.parametrize(
        ("config_separator", "expected"),
        [
            (None, "甲、乙、丙"),
            (", ", "甲, 乙, 丙"),
            ("", "甲乙丙"),
            (" + ", "甲 + 乙 + 丙"),
        ],
    )
    def test_md_to_docx_yaml_list_separator_is_request_scoped_and_exact(
        self,
        tmp_path: Path,
        config_separator: str | None,
        expected: str,
    ) -> None:
        from docx import Document

        template_path = tmp_path / "yaml-list-template.docx"
        template = Document()
        template.add_paragraph("{{authors}}")
        template.add_paragraph("{{ignored}}")
        template.add_paragraph("{{正文}}")
        template.save(str(template_path))

        md_path = write_temp_md(
            "---\n"
            "authors:\n"
            "  - 甲\n"
            "  - [乙, null]\n"
            "  - ''\n"
            "  - 'null'\n"
            "  - 丙\n"
            "ignored: ['null', 'None', null, '']\n"
            "---\n"
            "正文。\n"
        )
        options = {"template_name": str(template_path)}
        config_values = (
            {} if config_separator is None else {"conversion": {"md_to_docx": {"list_separator": config_separator}}}
        )
        ctx, _workspace = make_context(
            md_path,
            target_format="docx",
            options=options,
            config_values=config_values,
        )

        result = MdToDocxConverter().convert(ctx)

        assert result.success is True
        output = Document(str(result.artifacts[0].staging_path))
        paragraphs = [paragraph.text for paragraph in output.paragraphs]
        assert expected in paragraphs
        assert "{{ignored}}" not in paragraphs
