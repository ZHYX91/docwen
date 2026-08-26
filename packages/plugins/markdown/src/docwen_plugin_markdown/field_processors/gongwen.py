"""Gongwen YAML field normalization for MD to DOCX templates."""

from __future__ import annotations

import copy
import datetime as _dt
import re
from collections.abc import Mapping
from typing import Any

from docx.shared import Pt

ATTACH_NUM_PATTERN = re.compile(
    r"^[一二三四五六七八九十㈠㈡㈢㈣㈤㈥㈦㈧㈨㈩]+、|"
    r"^（[一二三四五六七八九十]+）|"
    r"^\d+[\.．]\s*|"
    r"^[０１２３４５６７８９]+[\.．]\s*|"
    r"^（\d+）\s*|"
    r"^（[０１２３４５６７８９]+）\s*|"
    r"^[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳㉑㉒㉓㉔㉕㉖㉗㉘㉙㉚㉛㉜㉝㉞㉟㊱㊲㊳㊴㊵㊶㊷㊸㊹㊺㊻㊼㊽㊾㊿⓪⓵⓶⓷⓸⓹⓺⓻⓼⓽⓾⓿❶❷❸❹❺❻❼❽❾❿⓫⓬⓭⓮⓯⓰⓱⓲⓳⓴]\s*"
)

PLACEHOLDER_RULES: dict[str, list[list[str]]] = {
    "delete_paragraph_if_empty": [
        ["密级和保密期限"],
        ["紧急程度"],
        ["发文字号"],
        ["公开方式"],
        ["主送机关"],
        ["附注"],
        ["抄送机关"],
        ["附件说明"],
        ["份号", "发文字号"],
    ],
    "delete_cell_if_empty": [],
    "delete_row_if_empty": [
        ["抄送机关"],
        ["印发机关", "印发日期"],
    ],
    "delete_table_if_empty": [],
}

SPECIAL_PLACEHOLDER_HANDLERS = {
    "附件说明": "process_attachment_placeholder",
}


def process_yaml(data: dict[str, Any]) -> None:
    """Apply Gongwen field semantics in place."""
    _process_attachment_description(data)
    _process_cc_orgs(data)
    _process_special_fields(data)


def process_attachment_placeholder(
    doc: Any,
    yaml_data: dict[str, Any],
    *,
    placeholder_paragraphs: list[Any] | None = None,
) -> bool:
    """Replace every original body ``{{附件说明}}`` placeholder."""
    pattern = re.compile(r"\{\{\s*附件说明\s*\}\}")
    body_elements = {paragraph._element for paragraph in doc.paragraphs}
    candidates = doc.paragraphs if placeholder_paragraphs is None else placeholder_paragraphs
    targets = [
        paragraph
        for paragraph in candidates
        if paragraph._element in body_elements and pattern.search(paragraph.text) is not None
    ]
    if not targets:
        return False

    attachments = yaml_data.get("附件说明")
    if _is_empty(attachments):
        for target in targets:
            _remove_paragraph(target)
        return True
    if not isinstance(attachments, (list, tuple)):
        attachments = [attachments]
    attachments = [item for item in attachments if not _is_empty(item)]

    for target_para in targets:
        base_style = target_para.style
        base_rpr = None
        if target_para.runs:
            base_rpr = target_para.runs[0]._r.rPr

        char_width = (
            target_para.paragraph_format.left_indent / 2 if target_para.paragraph_format.left_indent else Pt(16)
        )
        left_indent = target_para.paragraph_format.left_indent or 2 * char_width
        is_single = len(attachments) == 1
        first_line_indent = -int(3 * char_width if is_single else 4.5 * char_width)

        parent = target_para._element.getparent()
        if parent is None:
            continue
        index = parent.index(target_para._element)

        for offset, line in enumerate(attachments):
            new_para = doc.add_paragraph(style=base_style)
            new_run = new_para.add_run(str(line))
            if base_rpr is not None:
                new_rpr = new_run._r.get_or_add_rPr()
                for child in base_rpr:
                    new_rpr.append(copy.deepcopy(child))
            new_para.paragraph_format.left_indent = left_indent
            new_para.paragraph_format.first_line_indent = first_line_indent
            new_element = new_para._element
            new_element.getparent().remove(new_element)
            parent.insert(index + offset, new_element)

        parent.remove(target_para._element)
    return True


def _process_attachment_description(data: dict[str, Any]) -> None:
    if "附件说明" not in data:
        return

    attachments = data["附件说明"]
    if _is_empty(attachments):
        data["附件说明"] = []
        return

    if not isinstance(attachments, list):
        attachments = [attachments]
    attachments = [item for item in attachments if not _is_empty(item)]

    cleaned_attachments: list[str] = []
    for item in attachments:
        content = str(item).strip() if item is not None else ""
        cleaned = ATTACH_NUM_PATTERN.sub("", _convert_to_halfwidth(content)).strip()
        if cleaned == "" and content != "":
            cleaned = content
        cleaned_attachments.append(cleaned)

    formatted: list[str] = []
    for index, content in enumerate(cleaned_attachments, 1):
        if len(cleaned_attachments) == 1:
            formatted.append(f"附件：{content}")
        elif index == 1:
            formatted.append(f"附件：{index}. {content}")
        else:
            indent = "\u3000\u3000\u3000" if index < 10 else "\u3000\u3000 "
            formatted.append(f"{indent}{index}. {content}")

    data["附件说明"] = formatted


def _remove_paragraph(para: Any) -> None:
    parent = para._element.getparent()
    if parent is not None:
        parent.remove(para._element)


def _process_cc_orgs(data: dict[str, Any]) -> None:
    if "抄送机关" not in data:
        return
    cc_orgs = data["抄送机关"]
    if _is_empty(cc_orgs):
        data["抄送机关"] = ""
        return
    if not isinstance(cc_orgs, list):
        cc_orgs = [cc_orgs]
    valid_orgs = [_format_display_value(org).strip() for org in cc_orgs if not _is_empty(org)]
    data["抄送机关"] = "，".join(valid_orgs)


def _process_special_fields(data: dict[str, Any]) -> None:
    if "附注" in data:
        data["附注"] = _process_notes(data["附注"])
    if "印发日期" in data:
        data["印发日期"] = _format_date(data["印发日期"], suffix="印发")
    if "成文日期" in data:
        data["成文日期"] = _format_date(data["成文日期"])


def _process_notes(notes: Any) -> str:
    if not notes:
        return ""
    text = str(notes)
    if re.match(r"^[（(].*?[)）]$", text):
        return text[1:-1]
    return text


def _format_date(value: Any, suffix: str = "") -> str:
    if not value:
        return ""
    if isinstance(value, (_dt.date, _dt.datetime)):
        date_obj = value
    else:
        date_obj = None
        for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y年%m月%d日", "%Y.%m.%d", "%Y年%m月%d号"):
            try:
                date_obj = _dt.datetime.strptime(str(value), fmt)
                break
            except ValueError:
                continue
    if date_obj:
        formatted = f"{date_obj.year}年{date_obj.month}月{date_obj.day}日"
        return f"{formatted}{suffix}"
    return str(value)


def _is_empty(value: Any) -> bool:
    if value is None:
        return True
    if isinstance(value, str):
        return value.strip() == ""
    if isinstance(value, Mapping):
        return all(_is_empty(item) for item in value.values())
    if isinstance(value, (list, tuple, set)):
        return all(_is_empty(item) for item in value)
    return False


def _format_display_value(value: Any) -> str:
    if isinstance(value, (list, tuple, set)):
        return "，".join(_format_display_value(item) for item in value if not _is_empty(item))
    return str(value)


def _convert_to_halfwidth(text: str) -> str:
    full_to_half = {chr(0xFF10 + index): chr(0x30 + index) for index in range(10)}
    return "".join(full_to_half.get(char, char) for char in text)
