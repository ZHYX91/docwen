"""MdToDocxRenderer — renders mistune Markdown AST to python-docx paragraphs.

The renderer takes a python-docx ``Document`` and a ``body_font`` dict,
walks the mistune AST, and returns a **list of Paragraph objects** created
on that document.

Paragraphs are created but not positioned — the :mod:`template_filler`
module relocates them to the body placeholder position after rendering.

This decoupling is key: the renderer handles *what* to render, and the
filler handles *where* to place it.
"""

from __future__ import annotations

import logging
from collections import Counter
from pathlib import Path
from typing import TYPE_CHECKING, Any

from docx.shared import Pt

from docwen_core.docx_semantics import (
    DocxSemanticRenderer,
    apply_semantic_table_roles,
)
from docwen_core.models.resolved_numbering import (
    RESOLVED_DOCUMENT_SCHEMA,
    ResolvedDocumentTarget,
)
from docwen_core.models.semantic_document import (
    SemanticCaption,
    SemanticReference,
)
from docwen_core.text.heading_merge import HALFWIDTH_HEADING_MERGE_PUNCTUATION
from docwen_plugin_markdown.renderer_inlines import (
    extract_text_content,
    project_character_style_properties,
    render_formula,
    render_inlines,
    serialize_inlines_to_markdown,
)
from docwen_plugin_markdown.renderer_utils import (
    add_word_soft_break,
    apply_header_row_bottom_border,
    apply_list_indent,
    apply_paragraph_shading,
    apply_quote_style,
    apply_run_east_asian_font,
    apply_table_grid_borders,
    apply_three_line_table_borders,
    enable_table_header_row_formatting,
    resolve_code_block_style,
    resolve_formula_block_style,
    resolve_inline_code_style,
    resolve_inline_formula_style,
    resolve_list_block_style,
    resolve_quote_style,
    resolve_table_content_style,
    resolve_table_style,
    set_table_left_indent,
)
from docwen_plugin_markdown.template_utils import BodyParagraphFormat
from docwen_plugin_markdown.to_docx.numbering import (
    DocxListNumbering,
    apply_list_to_paragraph,
)

if TYPE_CHECKING:
    from docwen_core.docx_resolved_numbering import ResolvedNumberingDocxSession
    from docwen_core.docx_semantics_v3 import DocxSemanticsV3Session
    from docwen_plugin_markdown.to_docx.managed_styles import ManagedStyleBindings

logger = logging.getLogger(__name__)


_RESOLVED_TARGET_KEY = "_docwen_resolved_v4_target"
_RESOLVED_CAPTION_CHILDREN_KEY = "_docwen_resolved_v4_caption_children"


def _contains_request_semantics(nodes: list[dict[str, Any]]) -> bool:
    for node in nodes:
        if node.get("schema") in {"docwen.markdown_semantics.v3", RESOLVED_DOCUMENT_SCHEMA}:
            return True
        children = node.get("children")
        if isinstance(children, list) and _contains_request_semantics(children):
            return True
    return False


class MdToDocxRenderer:
    """Renders a mistune Markdown AST to a list of python-docx Paragraphs.

    The renderer creates paragraphs on ``self._doc`` and returns them
    as a list. The filler (:func:`~template_filler.fill_template`) later
    relocates them to the template's body placeholder position.

    Typical usage::

        doc = resolve_template(template_name)
        body_font = extract_body_font(doc)
        renderer = MdToDocxRenderer(doc, body_font)
        paragraphs = renderer.render(ast)
        fill_template(doc, yaml_dict, paragraphs, placeholder_para)
    """

    def __init__(
        self,
        doc,
        body_font: dict[str, Any] | None = None,
        body_style=None,
        body_paragraph_format: BodyParagraphFormat | None = None,
        code_font: str = "Consolas",
        code_bg_color: str = "E7E6E6",
        heading_formatting_mode: str = "remove",
        table_header_formatting_mode: str = "remove",
        table_style_name: str = "Table Grid",
        table_style_key: str | None = None,
        quote_style_levels: dict[int, int] | None = None,
        template_style_keys: dict[str, str] | None = None,
        cancellation=None,
        note_ctx: Any = None,
        formatting_mode: str = "full",
        hr_mapping: str | None = None,
        hr_actions: dict[str, str] | None = None,
        hr_attachments: set[int] | None = None,
        source_file_path: str | None = None,
        declared_resource_resolver=None,
        managed_styles: ManagedStyleBindings | None = None,
        semantic_v3_session: DocxSemanticsV3Session | None = None,
        resolved_numbering_session: ResolvedNumberingDocxSession | None = None,
        resolved_image_urls: tuple[str, ...] | None = None,
        **kwargs: Any,
    ) -> None:
        """
        Args:
            doc: python-docx ``Document`` to create paragraphs on.
            body_font: Font dict with keys ``name``, ``east_asia``,
                ``size``. If ``None``, defaults are used.
            body_style: Template-owned paragraph style from the body
                placeholder. If absent, new paragraphs use ``Normal``.
            code_font: Monospace font for code spans and code blocks.
            code_bg_color: Background color (RRGGBB) for code elements.
            heading_formatting_mode: ``"apply"``, ``"keep"``, or
                ``"remove"``. Controls inline formatting in headings.
            table_header_formatting_mode: ``"apply"``, ``"keep"``, or
                ``"remove"``. ``"remove"`` lets the table style own header
                emphasis rather than duplicating Markdown inline formatting.
            table_style_name: Name of the Word table style to apply.
            table_style_key: Builtin table style key, when one was selected.
            quote_style_levels: Mapping from Markdown quote depth to the
                semantic quote style level declared by document config.
            template_style_keys: Configured semantic keys for code and
                formula styles. Unknown keys deliberately use direct-format
                fallbacks rather than guessing localized names.
            managed_styles: Stable request-owned styles produced by the
                pre-render OOXML completion pass. Direct renderer callers may
                omit this and retain the legacy format-discovery fallback.
            cancellation: Optional cancellation token with ``check()``.
            note_ctx: Optional ``NoteContext`` for footnote/endnote support.
            formatting_mode: ``"full"`` (preserve all), ``"minimal"`` (strip
                all inline formatting).
            hr_mapping: Optional; ``None`` for horizontal rule, or
                ``"dash"``/``"asterisk"``/``"underscore"`` for centered
                text separator fallback.
            hr_actions: Optional marker-action mapping for thematic breaks.
                Keys are ``"dash"``, ``"asterisk"``, ``"underscore"``;
                values are ``"page_break"``, ``"section_break"``,
                ``"horizontal_rule_1"``/``"2"``/``"3"``, or ``"ignore"``.
            hr_attachments: Optional set of 0-based AST node indexes where
                HR should attach to the preceding paragraph.
        """
        self._doc = doc
        self._body_font = body_font or {
            "name": "Calibri",
            "east_asia": "宋体",
            "size": Pt(10.5),
            "bold": False,
            "italic": False,
            "underline": False,
        }
        self._body_style = body_style
        self._body_paragraph_format = body_paragraph_format
        self._code_font = code_font
        self._code_bg_color = code_bg_color
        self._heading_formatting_mode = heading_formatting_mode
        self._table_header_formatting_mode = table_header_formatting_mode
        self._table_style_name = table_style_name
        self._table_style_key = table_style_key
        self._managed_styles = managed_styles
        style_keys = template_style_keys or {
            "code_block": "code_block",
            "inline_code": "inline_code",
            "formula_block": "formula_block",
            "inline_formula": "inline_formula",
        }
        stable_style = managed_styles.get if managed_styles is not None else None
        self._code_block_style = (
            stable_style("code_block")
            if stable_style is not None and style_keys.get("code_block") == "code_block"
            else resolve_code_block_style(doc)
            if style_keys.get("code_block") == "code_block"
            else None
        )
        self.inline_code_style = (
            stable_style("inline_code")
            if stable_style is not None and style_keys.get("inline_code") == "inline_code"
            else resolve_inline_code_style(doc)
            if style_keys.get("inline_code") == "inline_code"
            else None
        )
        self._formula_block_style = (
            stable_style("formula_block")
            if stable_style is not None and style_keys.get("formula_block") == "formula_block"
            else resolve_formula_block_style(doc)
            if style_keys.get("formula_block") == "formula_block"
            else None
        )
        self.inline_formula_style = (
            stable_style("inline_formula")
            if stable_style is not None and style_keys.get("inline_formula") == "inline_formula"
            else resolve_inline_formula_style(doc)
            if style_keys.get("inline_formula") == "inline_formula"
            else None
        )
        self.inline_formula_style_reference = (
            self.inline_formula_style.style_id if self.inline_formula_style is not None else None
        )
        self._list_block_style = (
            stable_style("list_block") if stable_style is not None else resolve_list_block_style(doc)
        )
        self._table_content_style = (
            stable_style("table_content") if stable_style is not None else resolve_table_content_style(doc)
        )
        self._table_header_style = stable_style("table_header") if stable_style is not None else None
        self._body_managed_style = stable_style("body_paragraph") if stable_style is not None else None
        self._image_paragraph_style = stable_style("image_paragraph") if stable_style is not None else None
        self._table_grid_style = stable_style("table_grid") if stable_style is not None else None
        self._hyperlink_style_id = stable_style("hyperlink").style_id if stable_style is not None else None
        self._caption_styles = (
            {
                "figure": stable_style("figure_caption"),
                "table": stable_style("table_caption"),
                "equation": stable_style("equation_caption"),
                "code_block": stable_style("code_block_caption"),
                "listing": stable_style("code_block_caption"),
            }
            if stable_style is not None
            else {}
        )
        self._horizontal_rule_styles = (
            {str(level): stable_style(f"horizontal_rule_{level}") for level in range(1, 4)}
            if stable_style is not None
            else {}
        )
        self._template_table_style = (
            stable_style(table_style_key)
            if stable_style is not None and table_style_key in {"three_line_table", "table_grid"}
            else resolve_table_style(doc, table_style_key or "")
        )
        self._quote_style_levels = quote_style_levels or {level: level for level in range(1, 10)}
        self._quote_styles = {
            level: stable_style(f"quote_{level}") if stable_style is not None else resolve_quote_style(doc, level)
            for level in set(self._quote_style_levels.values())
            if 1 <= level <= 9
        }
        self._quote_depth = 0
        self._cancellation = cancellation
        self._note_ctx = note_ctx
        self._formatting_mode = formatting_mode
        self._hr_mapping = hr_mapping
        self._hr_actions = hr_actions or {}
        self._hr_attachments = hr_attachments or set()
        self._source_dir = Path(source_file_path).resolve().parent if source_file_path else None
        self._declared_resource_resolver = declared_resource_resolver
        self._semantic_v3_session = semantic_v3_session
        self._resolved_numbering_session = resolved_numbering_session
        if semantic_v3_session is not None and resolved_numbering_session is not None:
            raise ValueError("v3 and resolved-v4 DOCX sessions are mutually exclusive")
        # Both routes share the already-frozen, profile-free ordinary-anchor
        # and fenced-source physical carriers.  The resolved session never
        # receives v3 target/number/reference/Citation markers; it only
        # implements these two carrier bind methods.
        self._source_carrier_session = semantic_v3_session or resolved_numbering_session
        self._resolved_image_inventory = Counter(resolved_image_urls) if resolved_image_urls is not None else None

        # Heading counters for optional numbering.
        self._h_counters: list[int] = [0] * 9

        # List depth tracking for table indent
        self._list_depth: int = 0

        # Node counter for cancellation throttling
        self._node_count: int = 0

        # Heading merge indices computed before rendering.
        self._merge_indices: set[int] = set()

        # Word-native list numbering context (collects definitions during render)
        self._list_numbering = DocxListNumbering(doc)

        self._semantic_renderer = DocxSemanticRenderer(doc)

    # ── Public accessors ────────────────────────────────────────────────

    @property
    def document(self):
        """Return the underlying python-docx ``Document``."""
        return self._doc

    @property
    def list_numbering(self) -> DocxListNumbering:
        """Return the numbering context for post-save writeback.

        The converter calls :func:`write_numbering_to_docx` with this
        object after ``Document.save()`` to inject Word-native numbering
        definitions into the DOCX ZIP.
        """
        return self._list_numbering

    # ── Main entry point ─────────────────────────────────────────────────

    def render(self, ast_nodes: list[dict[str, Any]]) -> list:
        """Walk the mistune AST and return a list of Paragraph objects.

        Supports heading merge: when a heading has ``_merge=True``, the
        next paragraph node is consumed as merged body content and skipped
        in the normal iteration.

        Supports HR attachment: when a thematic_break has ``_attach_to_prev``,
        the HR content is merged into the preceding paragraph rather than
        creating a standalone paragraph.

        Args:
            ast_nodes: List of mistune AST token dicts.

        Returns:
            List of python-docx ``Paragraph`` objects created on
            ``self._doc``.
        """
        # Resolve note styles from the document once per render
        if self._note_ctx is not None:
            self._note_ctx.resolve_note_styles(self._doc)

        paragraphs: list = []
        skip_merged = False
        for i, node in enumerate(ast_nodes):
            if skip_merged:
                skip_merged = False
                continue

            self._node_count += 1
            if self._node_count % 50 == 0 and self._cancellation is not None:
                self._cancellation.check()

            # Heading merge: pass the next node as context
            if node.get("_merge") and node.get("type") == "heading":
                next_node = ast_nodes[i + 1] if i + 1 < len(ast_nodes) else None
                result = self._handle_heading(node, next_node=next_node)
                if next_node and next_node.get("_merged_into_heading"):
                    skip_merged = True
            # HR attachment: merge into the preceding paragraph
            elif node.get("_attach_to_prev") and node.get("type") == "thematic_break":
                if paragraphs:
                    # For list results, append to the last element
                    last = paragraphs[-1][-1] if isinstance(paragraphs[-1], list) and paragraphs[-1] else paragraphs[-1]
                    self._append_hr_to_paragraph(last, node)
                # Don't append a standalone paragraph for this HR
                result = None
            else:
                result = self._dispatch(node)

            if result is not None:
                if isinstance(result, list):
                    paragraphs.extend(result)
                else:
                    paragraphs.append(result)
        if self._resolved_image_inventory is not None and any(self._resolved_image_inventory.values()):
            raise ValueError("not every authenticated resolved-v4 image occurrence was rendered")
        return paragraphs

    # ── Node dispatcher ─────────────────────────────────────────────────

    def _dispatch(self, node: dict[str, Any]):
        """Dispatch a single AST node to the appropriate handler.

        This is a separate method for testability and override.
        """
        ntype = node.get("type", "")
        handler = getattr(self, f"_handle_{ntype}", None)
        if handler:
            return handler(node)
        return None

    def _bind_v3_ordinary_anchor(self, elements: tuple[Any, ...], node: dict[str, Any]) -> None:
        """Bind one source-authenticated ordinary owner and its direct parent."""

        if self._source_carrier_session is None:
            return
        anchor = node.get("_docwen_v3_ordinary_anchor")
        if not isinstance(anchor, dict):
            return
        self._source_carrier_session.bind_ordinary_anchor(
            elements,
            anchor,
            direct_parent_source_id=node.get("_docwen_v3_ordinary_anchor_parent_source_id"),
        )

    def _resolved_target(self, node: dict[str, Any], kind: str | None = None) -> ResolvedDocumentTarget | None:
        target = node.get(_RESOLVED_TARGET_KEY)
        if target is None:
            return None
        if self._resolved_numbering_session is None:
            raise ValueError("resolved-v4 target requires a request-owned DOCX session")
        if not isinstance(target, ResolvedDocumentTarget) or (kind is not None and target.kind != kind):
            raise ValueError("resolved-v4 AST target kind is not typed or does not match its owner")
        return target

    def _resolved_caption_target(self, node: dict[str, Any]) -> ResolvedDocumentTarget | None:
        target = self._resolved_target(node)
        if target is not None and target.kind == "heading":
            raise ValueError("resolved-v4 caption owner carries a Heading target")
        return target

    def _create_resolved_caption(
        self,
        node: dict[str, Any],
        *,
        object_elements: tuple[Any, ...],
    ):
        target = self._resolved_caption_target(node)
        if target is None:
            return None
        kind = target.kind
        assert self._resolved_numbering_session is not None
        children = node.get(_RESOLVED_CAPTION_CHILDREN_KEY)
        if not isinstance(children, list):
            raise ValueError("resolved-v4 caption lost its range-bound inline children")
        caption = self._style_semantic_caption(self._doc.add_paragraph(), kind)
        if _contains_request_semantics(children):
            render_inlines(
                caption,
                children,
                code_font=self._code_font,
                code_bg_color=self._code_bg_color,
                renderer_instance=self,
            )
            self._resolved_numbering_session.bind_rendered_caption(
                caption,
                object_elements,
                source_start=target.source_start,
                source_end=target.source_end,
                kind=kind,
            )
        else:
            self._resolved_numbering_session.bind_caption(
                caption,
                object_elements,
                source_start=target.source_start,
                source_end=target.source_end,
                kind=kind,
            )
        return caption

    # ═══════════════════════════════════════════════════════════════════
    # Block-level handlers
    # ═══════════════════════════════════════════════════════════════════

    def _handle_paragraph(self, node: dict[str, Any]):
        """Render a ``<p>`` paragraph.

        Respects ``self._formatting_mode``:
        - ``"full"`` (default): preserve all inline formatting.
        - ``"minimal"``: strip all inline formatting, keep text only.
        """
        # Skip if this paragraph was merged into a heading
        if node.get("_merged_into_heading"):
            return None

        children = node.get("children", [])
        image_only = self._is_image_only_paragraph(children)
        body_style = self._paragraph_style(image_only=image_only)
        p = self._doc.add_paragraph(style=body_style)
        if not image_only and self._body_paragraph_format is not None:
            self._body_paragraph_format.apply_to(p)

        if children:
            self._render_body_inline_children(p, children)
        resolved_target = self._resolved_caption_target(node)
        if resolved_target is not None:
            self._bind_v3_ordinary_anchor((p._p,), node)
            caption = self._create_resolved_caption(
                node,
                object_elements=(p._p,),
            )
            assert caption is not None
            if resolved_target.kind == "figure":
                p._p.addnext(caption._p)
                return [p, caption]
            p._p.addprevious(caption._p)
            return [caption, p]
        v3_target = node.get("_docwen_v3_caption_target")
        if v3_target is not None:
            caption = self._create_v3_caption(v3_target)
            if v3_target["kind"] == "figure":
                p._p.addnext(caption._p)
            else:
                p._p.addprevious(caption._p)
            if self._semantic_v3_session is not None:
                anchor = node.get("_docwen_v3_ordinary_anchor")
                if anchor is not None:
                    self._bind_v3_ordinary_anchor((p._p,), node)
                self._semantic_v3_session.bind_caption(caption, (p._p,), v3_target)
            return [p, caption] if v3_target["kind"] == "figure" else [caption, p]
        if self._source_carrier_session is not None and node.get("_docwen_v3_ordinary_anchor") is not None:
            self._bind_v3_ordinary_anchor((p._p,), node)
        return p

    def _handle_heading(self, node: dict[str, Any], next_node: dict[str, Any] | None = None):
        """Render a heading (levels 1-9; levels 7-9 are a DocWen extension).

        Heading merge: when *next_node* is a paragraph marked with
        ``_merged_into_heading``, append its runs to this same Word paragraph.
        The paragraph retains its Heading style while the appended runs receive
        the template body font as direct formatting, matching classic DocWen's
        mixed heading/body paragraph contract.
        """
        level = max(1, min(node.get("attrs", {}).get("level", 1), 9))
        self._h_counters[level - 1] += 1
        for lx in range(level, 9):
            self._h_counters[lx] = 0

        heading = self._doc.add_heading(level=level)
        children = node.get("children", [])

        hmode = self._heading_formatting_mode
        if hmode == "remove":
            render_inlines(
                heading,
                children,
                code_font=self._code_font,
                code_bg_color=self._code_bg_color,
                strip_formatting=True,
                renderer_instance=self,
            )
        elif hmode == "apply":
            render_inlines(
                heading,
                children,
                code_font=self._code_font,
                code_bg_color=self._code_bg_color,
                override_style=True,
                renderer_instance=self,
            )
        elif hmode == "keep":
            if _contains_request_semantics(children):
                render_inlines(
                    heading,
                    children,
                    code_font=self._code_font,
                    code_bg_color=self._code_bg_color,
                    renderer_instance=self,
                )
            else:
                text = serialize_inlines_to_markdown(children)
                if text:
                    heading.add_run(text)
        else:
            render_inlines(
                heading,
                children,
                code_font=self._code_font,
                code_bg_color=self._code_bg_color,
                override_style=False,
                renderer_instance=self,
            )

        # ── Heading merge: append body runs to this heading paragraph ──
        if next_node and next_node.get("_merged_into_heading"):
            body_children = next_node.get("children", [])
            if body_children:
                body_run_start = len(heading.runs)
                if heading.text.rstrip().endswith(tuple(HALFWIDTH_HEADING_MERGE_PUNCTUATION)):
                    heading.add_run(" ")
                self._render_body_inline_children(heading, body_children)
                self._apply_body_font_to_runs(
                    heading,
                    start_index=body_run_start,
                    reset_inherited_emphasis=True,
                )

        resolved_target = self._resolved_target(node, "heading")
        if resolved_target is not None:
            assert self._resolved_numbering_session is not None
            if resolved_target.heading_level != level:
                raise ValueError("resolved-v4 Heading level differs from the rendered ATX level")
            # Bind only after any body merge so Core authenticates the final
            # authored-title prefix and permitted rendered suffix together.
            self._resolved_numbering_session.bind_heading(
                heading,
                source_start=resolved_target.source_start,
                source_end=resolved_target.source_end,
            )

        if self._semantic_v3_session is not None and node.get("_docwen_v3_heading_target") is not None:
            self._semantic_v3_session.bind_heading(
                heading,
                node["_docwen_v3_heading_target"],
            )

        return heading

    def _handle_block_code(self, node: dict[str, Any]):
        """Render a code block as a single paragraph with ``<w:br>`` soft breaks."""
        fenced_record = node.get("_docwen_v3_fenced_source")
        logical_body = node.get("_docwen_v3_fenced_body")
        source_text = logical_body if isinstance(logical_body, str) else (node.get("raw", "") or node.get("text", ""))
        text = str(source_text).rstrip("\r\n")
        caption_data = node.get("_document_semantics_caption")
        if not text and caption_data is None and fenced_record is None:
            return None
        p = self._doc.add_paragraph(style=self._code_block_style)
        if self._code_block_style is None:
            apply_paragraph_shading(p, self._code_bg_color)
        for i, code_line in enumerate(text.split("\n")):
            if i > 0:
                add_word_soft_break(p)
            run = p.add_run(code_line if code_line else " ")
            if self._code_block_style is None:
                run.font.name = self._code_font
                apply_run_east_asian_font(run, self._code_font)
        if self._source_carrier_session is not None and fenced_record is not None:
            if not isinstance(logical_body, str):
                raise ValueError("a fenced source record requires its authenticated logical body")
            self._source_carrier_session.bind_fenced_source(
                p,
                fenced_record,
                logical_body=logical_body,
            )
        resolved_target = self._resolved_caption_target(node)
        if resolved_target is not None:
            self._bind_v3_ordinary_anchor((p._p,), node)
            caption = self._create_resolved_caption(
                node,
                object_elements=(p._p,),
            )
            assert caption is not None
            if resolved_target.kind == "figure":
                p._p.addnext(caption._p)
                return [p, caption]
            p._p.addprevious(caption._p)
            return [caption, p]
        v3_target = node.get("_docwen_v3_caption_target")
        if v3_target is None:
            if self._source_carrier_session is not None and node.get("_docwen_v3_ordinary_anchor") is not None:
                self._bind_v3_ordinary_anchor((p._p,), node)
            if caption_data is None:
                return p
        if v3_target is not None:
            caption = self._create_v3_caption(v3_target)
            if v3_target["kind"] == "figure":
                p._p.addnext(caption._p)
            else:
                p._p.addprevious(caption._p)
            if self._semantic_v3_session is not None:
                anchor = node.get("_docwen_v3_ordinary_anchor")
                if anchor is not None:
                    self._bind_v3_ordinary_anchor((p._p,), node)
                self._semantic_v3_session.bind_caption(caption, (p._p,), v3_target)
            return [p, caption] if v3_target["kind"] == "figure" else [caption, p]
        if caption_data is None:
            return p
        caption = self._create_semantic_caption(
            caption_data,
            number=int(node.get("_document_semantics_number", 1)),
        )
        p._p.addprevious(caption._p)
        return [caption, p]

    def _handle_block_quote(self, node: dict[str, Any]):
        """Render a blockquote.

        Creates one paragraph per child using the template's localized,
        level-specific quote style.  Nested quote results keep the style
        applied by their own depth rather than accumulating duplicate
        ``w:ind`` elements from every ancestor.
        """
        children = node.get("children", [])
        paragraphs: list = []
        self._quote_depth += 1
        try:
            for child in children:
                result = self._dispatch(child)
                if result is None:
                    continue
                items = result if isinstance(result, list) else [result]
                if child.get("type") != "block_quote" and self._quote_depth <= 9:
                    style_level = self._quote_style_levels.get(self._quote_depth, self._quote_depth)
                    style = self._quote_styles.get(style_level)
                    for item in items:
                        if hasattr(item, "_p"):
                            apply_quote_style(item, style=style, level=style_level)
                            if self._managed_styles is None:
                                self._apply_body_font_to_runs(item)
                paragraphs.extend(items)
        finally:
            self._quote_depth -= 1
        if (
            self._source_carrier_session is not None
            and node.get("_docwen_v3_ordinary_anchor") is not None
            and paragraphs
        ):
            self._bind_v3_ordinary_anchor(
                tuple(item._p if hasattr(item, "_p") else item._element for item in paragraphs),
                node,
            )
        return paragraphs if paragraphs else None

    def _handle_list(self, node: dict[str, Any]):
        """Render a (possibly nested) list with Word-native numbering.

        Pre-analyzes the list tree to determine ordered/unordered type at
        each nesting depth, creates a ``w:abstractNum`` + ``w:num``
        definition, then renders paragraphs with ``w:numPr`` applied.

        Each call to ``_handle_list`` produces one list group with its own
        ``numId``.  Nested sub-lists share the same ``numId`` but use
        different ``ilvl`` values.
        """
        # 1. Analyze list structure → {depth: 'ordered'|'unordered'}
        level_types: dict[int, str] = {}
        self._analyze_list_levels(node, level_types, depth=0)

        # 2. Create numbering definition → numId
        num_id = self._list_numbering.create_list_definition(level_types)

        # 3. Render items with numId context
        children = node.get("children", [])
        paragraphs: list = []
        for child in children:
            item_paras = self._render_list_item(child, 0, num_id=num_id)
            if item_paras:
                paragraphs.extend(item_paras)
        if (
            self._source_carrier_session is not None
            and node.get("_docwen_v3_ordinary_anchor") is not None
            and paragraphs
        ):
            self._bind_v3_ordinary_anchor(
                tuple(item._p if hasattr(item, "_p") else item._element for item in paragraphs),
                node,
            )
        return paragraphs if paragraphs else None

    def _analyze_list_levels(
        self,
        node: dict[str, Any],
        level_types: dict[int, str],
        depth: int = 0,
    ) -> None:
        """Recursively collect list type per nesting depth.

        Populates *level_types* in-place.  When the same depth has
        conflicting types (both ordered and unordered), ``"ordered"`` wins.

        Mistune v3 uses ``type: "list"`` for all lists and distinguishes
        ordered vs unordered via ``node["attrs"]["ordered"]`` (bool).
        """
        if node.get("type") != "list":
            return

        is_ordered = bool(node.get("attrs", {}).get("ordered", False))
        list_type = "ordered" if is_ordered else "unordered"
        if depth not in level_types:
            level_types[depth] = list_type
        elif level_types[depth] != list_type:
            # Mixed ordered + unordered at same depth → ordered wins
            level_types[depth] = "ordered"

        for item in node.get("children", []):
            if item.get("type") == "list_item":
                for child in item.get("children", []):
                    if child.get("type") == "list":
                        self._analyze_list_levels(child, level_types, depth + 1)

    def _render_list_item(
        self,
        node: dict[str, Any],
        depth: int = 0,
        num_id: str | None = None,
    ) -> list:
        """Render a single list item (or nested list) at given depth.

        When *num_id* is provided, each list-item paragraph receives
        ``w:numPr`` via :func:`apply_list_to_paragraph` so Word treats it
        as a native numbered/bulleted list.
        """
        ntype = node.get("type", "")
        paragraphs: list = []

        if ntype == "list_item":
            children = node.get("children", [])
            for child in children:
                ctype = child.get("type", "")
                if ctype in ("paragraph", "block_text"):
                    p = self._doc.add_paragraph(style=self._list_block_style)
                    if num_id is not None:
                        apply_list_to_paragraph(p, num_id, depth)
                    render_inlines(
                        p,
                        child.get("children", []),
                        code_font=self._code_font,
                        code_bg_color=self._code_bg_color,
                        renderer_instance=self,
                    )
                    if self._managed_styles is None:
                        self._apply_body_font_to_runs(p)
                    paragraphs.append(p)
                elif ctype in ("list", "bullet_list", "ordered_list"):
                    sub_items = self._render_list_item(child, depth + 1, num_id=num_id)
                    paragraphs.extend(sub_items)
                elif ctype == "block_code":
                    rendered = self._handle_block_code(child)
                    if rendered is not None:
                        items = rendered if isinstance(rendered, list) else [rendered]
                        for item in items:
                            apply_list_indent(item, depth)
                        paragraphs.extend(items)
                elif ctype == "block_quote":
                    items = self._handle_block_quote(child)
                    if items:
                        for item_index, item in enumerate(items):
                            if item_index == 0 and num_id is not None and hasattr(item, "_p"):
                                apply_list_to_paragraph(item, num_id, depth)
                            else:
                                apply_list_indent(item, depth)
                        paragraphs.extend(items)
                elif ctype == "table":
                    table = self._handle_table(child)
                    if table is not None:
                        if depth > 0:
                            indent_twips = 360 * (depth + 1)
                            set_table_left_indent(table, indent_twips)
                        paragraphs.append(table)
                else:
                    prev_depth = self._list_depth
                    self._list_depth = depth + 1
                    try:
                        result = self._dispatch(child)
                        if result is not None:
                            if isinstance(result, list):
                                for r in result:
                                    apply_list_indent(r, depth) if hasattr(r, "_p") else None
                                paragraphs.extend(result)
                            else:
                                apply_list_indent(result, depth) if hasattr(result, "_p") else None
                                paragraphs.append(result)
                    finally:
                        self._list_depth = prev_depth
            if (
                self._source_carrier_session is not None
                and node.get("_docwen_v3_ordinary_anchor") is not None
                and paragraphs
            ):
                self._bind_v3_ordinary_anchor(
                    tuple(item._p if hasattr(item, "_p") else item._element for item in paragraphs),
                    node,
                )
        elif ntype in ("list", "bullet_list", "ordered_list"):
            for child in node.get("children", []):
                sub_items = self._render_list_item(child, depth, num_id=num_id)
                paragraphs.extend(sub_items)

        return paragraphs

    def _handle_thematic_break(self, node: dict[str, Any]):
        """Render a horizontal rule (thematic break).

        When ``self._hr_mapping`` is set to ``"dash"``, ``"asterisk"``,
        or ``"underscore"``, a centered text separator is used instead
        of a Word horizontal rule border.

        When the node has ``_attach_to_prev = True`` (HR attachment mode),
        the caller's :meth:`render` loop merges the HR into the preceding
        paragraph instead of creating a standalone paragraph — this method
        is not called directly in that case.
        """
        hr_text = self._hr_mapping_text()

        if hr_text:
            p = self._doc.add_paragraph()
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(hr_text * 5)
            if self._body_font:
                run.font.name = self._body_font.get("name")
                run.font.size = self._body_font.get("size")
            return p

        from docwen_plugin_markdown.to_docx.breaks import (
            insert_horizontal_rule,
            insert_page_break,
            insert_section_break,
        )

        action = self._resolve_hr_action(node)
        if action == "ignore":
            return None
        if action == "page_break":
            return insert_page_break(self._doc)
        if action == "section_break":
            return insert_section_break(self._doc, section_type="next")
        if action.startswith("horizontal_rule_"):
            variant = action.rsplit("_", 1)[-1]
            return insert_horizontal_rule(
                self._doc,
                variant=variant,
                style=self._horizontal_rule_styles.get(variant),
            )
        return insert_horizontal_rule(self._doc, variant="1", style=self._horizontal_rule_styles.get("1"))

    def _append_hr_to_paragraph(self, paragraph, node: dict[str, Any]) -> None:
        """Append HR separator content into an existing paragraph.

        Adds a soft break followed by a centered dash/asterisk/underscore
        text. This is used for attached HRs that merge into the preceding
        paragraph rather than creating a standalone block.
        """
        hr_text = self._hr_mapping_text()
        if hr_text:
            add_word_soft_break(paragraph)
            run = paragraph.add_run(hr_text * 5)
            if self._body_font:
                run.font.name = self._body_font.get("name")
                run.font.size = self._body_font.get("size")
            return

        from docwen_plugin_markdown.to_docx.breaks import (
            append_horizontal_rule_to_paragraph,
            append_page_break_to_paragraph,
            append_section_break_to_paragraph,
        )

        action = self._resolve_hr_action(node)
        if action == "ignore":
            return
        if action == "page_break":
            append_page_break_to_paragraph(paragraph)
        elif action == "section_break":
            append_section_break_to_paragraph(paragraph, self._doc, section_type="next")
        elif action.startswith("horizontal_rule_"):
            append_horizontal_rule_to_paragraph(paragraph, variant=action.rsplit("_", 1)[-1])
        else:
            append_horizontal_rule_to_paragraph(paragraph, variant="1")

    def _hr_mapping_text(self) -> str:
        return {
            "dash": "—",
            "asterisk": "*",
            "underscore": "_",
        }.get(self._hr_mapping or "", "")

    def _resolve_hr_action(self, node: dict[str, Any]) -> str:
        marker = str(node.get("_hr_marker") or "underscore")
        return self._hr_actions.get(marker, "horizontal_rule_1")

    def _handle_block_html(self, node: dict[str, Any]):
        """Skip raw HTML blocks in DOCX output."""
        return None

    def _handle_block_latex(self, node: dict[str, Any]):
        """Render block LaTeX as OMML equation."""
        render_formula(self._doc, node, block_style=self._formula_block_style)
        # Return the last paragraph (created by render_formula)
        paragraph = self._doc.paragraphs[-1] if self._doc.paragraphs else None
        if paragraph is None:
            return None
        resolved_target = self._resolved_caption_target(node)
        if resolved_target is not None:
            _strip_serialization_only_whitespace(paragraph._p)
            self._bind_v3_ordinary_anchor((paragraph._p,), node)
            caption = self._create_resolved_caption(
                node,
                object_elements=(paragraph._p,),
            )
            assert caption is not None
            if resolved_target.kind == "figure":
                paragraph._p.addnext(caption._p)
                return [paragraph, caption]
            paragraph._p.addprevious(caption._p)
            return [caption, paragraph]
        v3_target = node.get("_docwen_v3_caption_target")
        if v3_target is not None:
            caption = self._create_v3_caption(v3_target)
            if v3_target["kind"] == "figure":
                paragraph._p.addnext(caption._p)
            else:
                paragraph._p.addprevious(caption._p)
            if self._semantic_v3_session is not None:
                anchor = node.get("_docwen_v3_ordinary_anchor")
                if anchor is not None:
                    self._bind_v3_ordinary_anchor((paragraph._p,), node)
                self._semantic_v3_session.bind_caption(caption, (paragraph._p,), v3_target)
            return [paragraph, caption] if v3_target["kind"] == "figure" else [caption, paragraph]
        if self._source_carrier_session is not None and node.get("_docwen_v3_ordinary_anchor") is not None:
            self._bind_v3_ordinary_anchor((paragraph._p,), node)
        caption_data = node.get("_document_semantics_caption")
        if caption_data is None:
            return paragraph
        caption = self._create_semantic_caption(
            caption_data,
            number=int(node.get("_document_semantics_number", 1)),
        )
        paragraph._p.addprevious(caption._p)
        return [caption, paragraph]

    def _handle_block_math(self, node: dict[str, Any]):
        """Handle block_math from mistune's math plugin."""
        return self._handle_block_latex(node)

    def _handle_blank_line(self, node: dict[str, Any]):
        """Blank lines between blocks — no-op in DOCX."""
        return None

    def _handle_semantic_figure(self, node: dict[str, Any]):
        """Render a figure object followed by a Word-native caption."""

        object_node = node["object"]
        figure = self._doc.add_paragraph(style=self._image_paragraph_style or self._paragraph_style(image_only=True))
        self._render_body_inline_children(figure, object_node.get("children", []))
        caption = self._create_semantic_caption(
            node["_document_semantics_caption"],
            number=int(node.get("_document_semantics_number", 1)),
        )
        return [item for item in (figure, caption) if item is not None]

    def _handle_table(self, node: dict[str, Any]):
        """Render a Markdown table.

        Creates a python-docx table with inline formatting preserved
        in cells using ``_render_inlines``.
        """
        children = node.get("children", [])
        all_rows = self._build_table_rows(children)
        semantics = node.get("_document_semantics_table")
        rows = int(semantics["row_count"]) if semantics is not None else len(all_rows)
        if rows == 0:
            return None
        cols = (
            int(semantics["column_count"]) if semantics is not None else max((len(row) for row in all_rows), default=0)
        )
        if cols == 0:
            return None

        table = self._doc.add_table(rows=rows, cols=cols)
        is_three_line_table = self._is_three_line_table()
        self._apply_table_style(table)

        if self._list_depth > 0:
            indent_twips = 360 * self._list_depth
            set_table_left_indent(table, indent_twips)

        header_row_count = int(semantics["header_rows"]) if semantics is not None else self._count_header_rows(children)
        if header_row_count > 0:
            if semantics is None:
                enable_table_header_row_formatting(table)
            if is_three_line_table and self._managed_styles is None:
                for header_row_idx in range(min(header_row_count, len(table.rows))):
                    apply_header_row_bottom_border(table.rows[header_row_idx])

        if semantics is not None:
            apply_semantic_table_roles(
                table,
                header_rows=header_row_count,
                header_columns=int(semantics["header_columns"]),
                repeat_header=str(semantics["repeat_header"]),
            )
            render_cells = [(int(anchor["row"]), int(anchor["column"]), anchor) for anchor in semantics["anchors"]]
        else:
            render_cells = [
                (row_idx, col_idx, cell_node)
                for row_idx, row_cells in enumerate(all_rows)
                for col_idx, cell_node in enumerate(row_cells)
            ]

        for row_idx, col_idx, cell_node in render_cells:
            if col_idx >= cols:
                continue
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            # Render inline content into the cell's first paragraph
            cell_children = cell_node.get("children", [])
            p = cell.paragraphs[0]
            is_header_cell = cell_node.get("role") != "data" if semantics is not None else row_idx < header_row_count
            cell_style = self._table_header_style if is_header_cell else self._table_content_style
            if cell_style is not None:
                p.style = cell_style
            if cell_children:
                if (
                    is_header_cell
                    and self._table_header_formatting_mode == "keep"
                    and not _contains_request_semantics(cell_children)
                ):
                    p.add_run(serialize_inlines_to_markdown(cell_children))
                elif is_header_cell and self._table_header_formatting_mode == "remove":
                    render_inlines(
                        p,
                        cell_children,
                        code_font=self._code_font,
                        code_bg_color=self._code_bg_color,
                        strip_formatting=True,
                        renderer_instance=self,
                    )
                else:
                    self._render_body_inline_children(p, cell_children)
                if is_header_cell and self._managed_styles is None:
                    for run in p.runs:
                        if run.text:
                            run.bold = True

        if semantics is not None:
            for anchor in semantics["anchors"]:
                row_span = int(anchor["row_span"])
                column_span = int(anchor["column_span"])
                if row_span == 1 and column_span == 1:
                    continue
                row = int(anchor["row"])
                column = int(anchor["column"])
                table.cell(row, column).merge(table.cell(row + row_span - 1, column + column_span - 1))

        resolved_target = self._resolved_caption_target(node)
        if resolved_target is not None:
            self._bind_v3_ordinary_anchor((table._element,), node)
            caption = self._create_resolved_caption(
                node,
                object_elements=(table._element,),
            )
            assert caption is not None
            if resolved_target.kind == "figure":
                table._element.addnext(caption._p)
                return [table, caption]
            table._element.addprevious(caption._p)
            return [caption, table]

        v3_target = node.get("_docwen_v3_caption_target")
        if v3_target is not None:
            caption = self._create_v3_caption(v3_target)
            if v3_target["kind"] == "figure":
                table._element.addnext(caption._p)
            else:
                table._element.addprevious(caption._p)
            if self._semantic_v3_session is not None:
                anchor = node.get("_docwen_v3_ordinary_anchor")
                if anchor is not None:
                    self._bind_v3_ordinary_anchor((table._element,), node)
                self._semantic_v3_session.bind_caption(caption, (table._element,), v3_target)
            return [table, caption] if v3_target["kind"] == "figure" else [caption, table]
        if self._source_carrier_session is not None and node.get("_docwen_v3_ordinary_anchor") is not None:
            self._bind_v3_ordinary_anchor((table._element,), node)
        caption_data = node.get("_document_semantics_caption")
        if caption_data is not None:
            source_form = str(caption_data.get("source_form") or "imported")
            caption = self._semantic_renderer.render_caption_for_table(
                table,
                self._semantic_caption_model(
                    caption_data,
                    number=int(node.get("_document_semantics_number", 1)),
                ),
                source_form=source_form,
            )
            return [self._style_semantic_caption(caption, "table"), table]

        return table

    def _create_semantic_caption(self, caption: dict[str, Any], *, number: int):
        """Create a visible caption backed by a native ``SEQ`` field."""

        paragraph = self._semantic_renderer.render_caption(
            self._semantic_caption_model(caption, number=number),
            source_form=str(caption.get("source_form") or "imported"),
        )
        return self._style_semantic_caption(paragraph, str(caption["kind"]))

    def _create_v3_caption(self, target: dict[str, Any]):
        """Create one v3 caption without using the superseded ID grammar."""

        from docwen_core.docx_semantics_v3 import append_complex_field

        kind = str(target["kind"])
        counter = {
            "figure": "Figure",
            "table": "Table",
            "equation": "Equation",
            "code_block": "Code",
        }[kind]
        paragraph = self._doc.add_paragraph()
        paragraph.add_run(f"{counter} ")
        append_complex_field(
            paragraph,
            instruction=f" SEQ {counter} \\* ARABIC ",
            cached_result=str(target["number"]),
        )
        title = str(target.get("title") or "")
        if title:
            paragraph.add_run(f": {title}")
        return self._style_semantic_caption(paragraph, kind)

    def _semantic_caption_model(self, caption: dict[str, Any], *, number: int) -> SemanticCaption:
        kind = str(caption["kind"])
        target_id = caption.get("target_id")
        return SemanticCaption(
            kind=kind,  # type: ignore[arg-type]
            target_id=str(target_id) if target_id else None,
            cached_number=str(number),
            label=kind.title(),
            content=str(caption["content"]),
        )

    def _style_semantic_caption(self, paragraph, kind: str):
        managed_style = self._caption_styles.get(kind)
        if managed_style is not None:
            paragraph.style = managed_style
            return paragraph
        paragraph.style = self._body_style
        if self._body_paragraph_format is not None:
            self._body_paragraph_format.apply_to(paragraph)
        self._apply_body_font_to_runs(paragraph)
        return paragraph

    def render_semantic_reference(self, parent, node: dict[str, Any]) -> None:
        """Render a canonical cross-reference as a native ``REF`` field."""

        if node.get("schema") == RESOLVED_DOCUMENT_SCHEMA:
            if self._resolved_numbering_session is None:
                raise ValueError("resolved-v4 reference requires a request-owned DOCX session")
            self._resolved_numbering_session.render_reference(
                parent,
                source_start=int(node["source_start"]),
                source_end=int(node["source_end"]),
            )
            return
        if node.get("schema") == "docwen.markdown_semantics.v3":
            if self._semantic_v3_session is None:
                raise ValueError("v3 semantic reference requires a request-owned DOCX session")
            self._semantic_v3_session.render_reference(parent, node)
            return

        self._semantic_renderer.render_reference(
            parent,
            SemanticReference(
                target_id=str(node["target_id"]),
                cached_result=str(node.get("cached_result", "?")),
            ),
        )

    def render_semantic_citation(self, parent, node: dict[str, Any]) -> None:
        """Render one provider-resolved Citation from its authenticated range."""

        if node.get("schema") != RESOLVED_DOCUMENT_SCHEMA or self._resolved_numbering_session is None:
            # Historical source routes do not own a v4 Citation resolver.
            parent.add_run(str(node.get("raw", "")))
            return
        self._resolved_numbering_session.render_citation(
            parent,
            source_start=int(node["source_start"]),
            source_end=int(node["source_end"]),
        )

    def _render_body_inline_children(self, paragraph, children: list[dict[str, Any]]) -> None:
        has_request_semantics = _contains_request_semantics(children)
        if self._formatting_mode == "minimal" and not has_request_semantics:
            text = extract_text_content(children)
            if text:
                paragraph.add_run(text)
        elif self._formatting_mode == "keep" and not has_request_semantics:
            text = serialize_inlines_to_markdown(children)
            if text:
                paragraph.add_run(text)
        else:
            render_inlines(
                paragraph,
                children,
                code_font=self._code_font,
                code_bg_color=self._code_bg_color,
                renderer_instance=self,
            )

    def _paragraph_style(self, *, image_only: bool):
        if image_only and self._image_paragraph_style is not None:
            return self._image_paragraph_style
        if self._body_managed_style is None:
            return self._body_style
        if self._body_style is None or self._body_style.style_id in {"Normal", "normal"}:
            return self._body_managed_style
        return self._body_style

    @staticmethod
    def _is_image_only_paragraph(children: list[dict[str, Any]]) -> bool:
        meaningful = [
            child
            for child in children
            if child.get("type") not in {"softbreak", "linebreak"}
            and (child.get("type") != "text" or str(child.get("raw", "") or child.get("text", "")).strip())
        ]
        return len(meaningful) == 1 and meaningful[0].get("type") == "image"

    def _is_three_line_table(self) -> bool:
        key = (self._table_style_key or "").strip().lower()
        if key == "three_line_table":
            return True
        return self._table_style_name.strip().lower() in {"three line table", "三线表"}

    def _apply_table_style(self, table) -> None:
        if self._template_table_style is not None:
            table.style = self._template_table_style
            return

        try:
            table.style = self._table_style_name
            return
        except KeyError:
            logger.warning("Table style %r is unavailable; using a compatible fallback.", self._table_style_name)

        if self._is_three_line_table():
            apply_three_line_table_borders(table)
            return

        if self._table_grid_style is not None:
            table.style = self._table_grid_style
            return

        stable_grid = resolve_table_style(self._doc, "table_grid")
        if stable_grid is not None:
            table.style = stable_grid
            return

        try:
            table.style = "Table Grid"
        except KeyError:
            logger.warning("Table Grid style is unavailable; applying direct grid borders.")
            apply_table_grid_borders(table)

    def _apply_body_font_to_runs(
        self,
        paragraph,
        *,
        start_index: int = 0,
        reset_inherited_emphasis: bool = False,
    ) -> None:
        """Apply the body placeholder font while preserving explicit inline styles.

        ``reset_inherited_emphasis`` is used by mixed Heading paragraphs so
        unformatted body runs do not inherit the Heading style's bold/italic/
        underline values.  Explicit Markdown emphasis remains untouched.
        """
        inline_code_style_id = self.inline_code_style.style_id if self.inline_code_style is not None else None
        for run in paragraph.runs[start_index:]:
            if inline_code_style_id and run.style is not None and run.style.style_id == inline_code_style_id:
                continue
            if run.font.name is None:
                run.font.name = self._body_font.get("name")
                east_asia = self._body_font.get("east_asia")
                if east_asia:
                    apply_run_east_asian_font(run, east_asia)
            if run.font.size is None:
                run.font.size = self._body_font.get("size")
            if reset_inherited_emphasis:
                if run.bold is None:
                    run.bold = self._body_font.get("bold", False)
                if run.italic is None:
                    run.italic = self._body_font.get("italic", False)
                if run.underline is None:
                    run.underline = self._body_font.get("underline", False)

    def _handle_table_head(self, node: dict[str, Any]):
        return None  # handled inside _handle_table

    def _handle_table_body(self, node: dict[str, Any]):
        return None  # handled inside _handle_table

    def _handle_table_row(self, node: dict[str, Any]):
        return None  # handled inside _handle_table

    def _handle_table_cell(self, node: dict[str, Any]):
        return None  # handled inside _handle_table

    def _build_table_rows(self, children: list[dict[str, Any]]) -> list[list[dict[str, Any]]]:
        """Extract rows of cell nodes from a table AST.

        Handles both direct-cell (mistune table_head) and
        table_row-wrapped formats.
        """
        all_rows: list[list[dict[str, Any]]] = []
        for child in children:
            ctype = child.get("type", "")
            if ctype in ("table_head", "table_body"):
                sub = child.get("children", [])
                if sub and sub[0].get("type") == "table_cell":
                    all_rows.append(sub)
                else:
                    for row_node in sub:
                        if row_node.get("type") == "table_row":
                            all_rows.append(row_node.get("children", []))
        return all_rows

    @staticmethod
    def _count_header_rows(children: list[dict[str, Any]]) -> int:
        """Count header rows in a table AST."""
        count = 0
        for child in children:
            if child.get("type") == "table_head":
                sub = child.get("children", [])
                if sub and sub[0].get("type") == "table_cell":
                    count += 1
                else:
                    for sn in sub:
                        if sn.get("type") == "table_row":
                            count += 1
        return count

    # ── Inline rendering proxy (for subclass access) ────────────────────

    def render_inlines(
        self,
        parent,
        children: list[dict[str, Any]],
        override_style: bool = False,
        strip_formatting: bool = False,
    ) -> None:
        """Proxy to the standalone ``render_inlines`` with renderer config."""
        render_inlines(
            parent,
            children,
            code_font=self._code_font,
            code_bg_color=self._code_bg_color,
            override_style=override_style,
            strip_formatting=strip_formatting,
            renderer_instance=self,
        )

    def _add_plain_run(self, paragraph, text: str) -> None:
        """Add a plain text run using the body font."""
        run = paragraph.add_run(text)
        if self._body_font:
            run.font.name = self._body_font.get("name")
            run.font.size = self._body_font.get("size")
        return run

    def add_hyperlink(
        self,
        paragraph,
        url: str,
        text: list[dict[str, Any]] | str,
        *,
        override_style: bool = False,
        strip_formatting: bool = False,
    ) -> None:
        """Delegate to standalone hyperlink helper."""
        from docwen_plugin_markdown.renderer_inlines import add_hyperlink

        link_children = text if isinstance(text, list) else [{"type": "text", "raw": text}]

        def render_plain(target_paragraph, target_children: list[dict[str, Any]]) -> None:
            render_inlines(
                target_paragraph,
                target_children,
                code_font=self._code_font,
                code_bg_color=self._code_bg_color,
                override_style=override_style,
                strip_formatting=strip_formatting,
                renderer_instance=self,
            )

        add_hyperlink(
            paragraph,
            url,
            children=link_children,
            source_dir=self._source_dir,
            style_reference=self._hyperlink_style_id,
            override_style=override_style,
            strip_formatting=strip_formatting,
            code_font=self._code_font,
            code_bg_color=self._code_bg_color,
            inline_code_properties=project_character_style_properties(self.inline_code_style, paragraph=paragraph),
            fallback_renderer=render_plain,
        )

    def embed_image(
        self,
        paragraph,
        src: str,
        alt: str = "",
        width: int | None = None,
        height: int | None = None,
    ) -> None:
        """Delegate to standalone image embedding helper."""
        from docwen_plugin_markdown.renderer_inlines import embed_image

        resolved_v4 = self._resolved_image_inventory is not None
        if resolved_v4:
            assert self._resolved_image_inventory is not None
            if self._resolved_image_inventory[src] <= 0:
                raise ValueError("renderer received an image outside the authenticated resolved-v4 inventory")
            self._resolved_image_inventory[src] -= 1
        embed_image(
            paragraph,
            src,
            source_dir=None if resolved_v4 else self._source_dir,
            declared_resource_resolver=None if resolved_v4 else self._declared_resource_resolver,
            alt=alt,
            width=width,
            height=height,
            fail_closed=resolved_v4,
        )

    def _render_footnote_ref(self, parent, node: dict[str, Any]) -> None:
        """Render a ``footnote_ref`` AST node as a Word footnote/endnote reference.

        Delegates to the ``NoteContext`` if available; otherwise falls back
        to plain text ``[^{key}]``.
        """
        if self._note_ctx is None:
            parent.add_run(f"[^{node.get('raw', '')}]")
            return

        key: str = node.get("raw", "")
        if not key:
            return

        if key.upper().startswith("ENDNOTE-"):
            ref_run = self._note_ctx.create_endnote_ref_run(key)
        else:
            ref_run = self._note_ctx.create_footnote_ref_run(key)

        if ref_run is not None:
            parent._p.append(ref_run)
        else:
            parent.add_run(f"[^{key}]")


def _strip_serialization_only_whitespace(element: Any) -> None:
    """Remove XSLT indentation that cannot survive a DOCX save/reopen.

    Formula transforms may pretty-print MathML/OMML nodes.  A whitespace-only
    text value is indentation only when it belongs to a container element.
    Leaf text (including ``m:t``, ``w:t``, and ``w:instrText``) remains
    authored/semantic payload and is never removed.  Whitespace-only tails are
    XML indentation between sibling elements.
    """

    for item in element.iter():
        if len(item) > 0 and isinstance(item.text, str) and item.text.isspace():
            item.text = None
        if isinstance(item.tail, str) and item.tail.isspace():
            item.tail = None
