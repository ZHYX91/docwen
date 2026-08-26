"""Thin adapter: Markdown formula AST nodes → DOCX OMML elements."""

from __future__ import annotations

from typing import Any

from docx.oxml.ns import qn


def _latex_to_omml(latex: str) -> Any | None:
    """Convert LaTeX string to lxml OMML element via core formula pipeline."""
    import lxml.etree as etree

    from docwen_core.formula import latex_to_mathml, mathml_to_omml

    if not latex.strip():
        return None
    mathml = latex_to_mathml(latex)
    if not mathml:
        return None
    omml = mathml_to_omml(mathml)
    if not omml:
        return None
    return etree.fromstring(omml.encode("utf-8"))


def render_block_formula(document, latex: str, *, block_style=None) -> None:
    """Insert a centered paragraph with an OMML block equation."""
    omml = _latex_to_omml(latex)
    if omml is None:
        document.add_paragraph(latex, style=block_style)
        return
    p = document.add_paragraph(style=block_style)
    if block_style is None:
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except ImportError:
            p.alignment = 1  # CENTER
    p._p.append(omml)


def _apply_character_style(omml: Any, style_reference: str | None) -> None:
    """Apply a Word character style to formula runs and control properties."""
    if not style_reference:
        return

    import lxml.etree as etree

    math_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    word_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    targets = list(omml.iter(f"{{{math_ns}}}r")) + list(omml.iter(f"{{{math_ns}}}ctrlPr"))
    for target in targets:
        r_pr = target.find(qn("w:rPr"))
        if r_pr is None:
            r_pr = etree.Element(f"{{{word_ns}}}rPr")
            target.insert(0, r_pr)
        r_style = r_pr.find(qn("w:rStyle"))
        if r_style is None:
            r_style = etree.Element(f"{{{word_ns}}}rStyle")
            r_pr.insert(0, r_style)
        r_style.set(qn("w:val"), style_reference)


def render_inline_formula(paragraph, latex: str, style_reference: str | None = None) -> None:
    """Append an OMML inline equation to an existing paragraph."""
    omml = _latex_to_omml(latex)
    if omml is None:
        run = paragraph.add_run(latex)
        if style_reference:
            run.style = style_reference
        return
    _apply_character_style(omml, style_reference)
    paragraph._p.append(omml)
