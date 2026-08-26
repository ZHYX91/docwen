"""Break element insertion for MD → DOCX rendering.

Provides functions to insert page breaks, section breaks, and horizontal
rules into a python-docx Document, using OOXML-level element construction.

Two usage modes are supported:

* **New-paragraph mode** — ``insert_page_break``, ``insert_section_break``,
  ``insert_horizontal_rule`` create a standalone paragraph containing the
  break element.

* **Attach-to-previous mode** — ``append_page_break_to_paragraph``,
  ``append_section_break_to_paragraph``, ``append_horizontal_rule_to_paragraph``
  add the break element to an existing paragraph.

These functions close findings F-F1-001 through F-F1-005 (and F-F1-006).
"""

from __future__ import annotations

import copy
import logging
from typing import TYPE_CHECKING

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)

# ── Namespace constants ──────────────────────────────────────────────────
_WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# ── Section type mapping ─────────────────────────────────────────────────
_SECTION_TYPE_MAP: dict[str, str] = {
    "next": "nextPage",
    "continuous": "continuous",
    "even": "evenPage",
    "odd": "oddPage",
}

# ── Horizontal rule border widths (eighths of a point) ──────────────────
_HR_BORDER_SZ: dict[str, str] = {
    "1": "4",  # 0.5 pt  — fine
    "2": "8",  # 1.0 pt  — medium
    "3": "12",  # 1.5 pt  — thick
}


# ═════════════════════════════════════════════════════════════════════════
# New-paragraph insertions
# ═════════════════════════════════════════════════════════════════════════


def insert_page_break(doc) -> Paragraph:
    """Insert a page break as a new paragraph.

    Creates::

        <w:p>
            <w:r>
                <w:br w:type="page"/>
            </w:r>
        </w:p>

    Args:
        doc: python-docx ``Document`` object.

    Returns:
        The newly created ``Paragraph``.
    """
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
    logger.debug("Created page-break paragraph")
    return p


def insert_section_break(doc, section_type: str = "next") -> Paragraph:
    """Insert a section break as a new paragraph.

    The paragraph's ``pPr`` receives a ``w:sectPr`` element containing
    the section type and page settings copied from the document's
    default section.

    Supported *section_type* values:

    =============  ===========
    Short name     OOXML value
    =============  ===========
    ``"next"``     ``nextPage``
    ``"continuous"`` ``continuous``
    ``"even"``     ``evenPage``
    ``"odd"``      ``oddPage``
    =============  ===========

    Args:
        doc: python-docx ``Document`` object.
        section_type: One of ``"next"``, ``"continuous"``, ``"even"``,
            ``"odd"``.  Defaults to ``"next"``.

    Returns:
        The newly created ``Paragraph`` containing the section break.
    """
    word_type = _SECTION_TYPE_MAP.get(section_type, "nextPage")

    p = doc.add_paragraph()
    pPr = _get_or_add_pPr(p._p)

    sectPr = OxmlElement("w:sectPr")

    # Section type
    stype = OxmlElement("w:type")
    stype.set(qn("w:val"), word_type)
    sectPr.append(stype)

    _copy_page_settings(doc, sectPr)
    pPr.append(sectPr)

    logger.debug(f"Created section-break paragraph: type={word_type}")
    return p


def insert_horizontal_rule(doc, variant: str = "1", *, style=None) -> Paragraph:
    """Insert a horizontal rule as a new paragraph with a bottom border.

    Three thickness variants are available:

    =========  ===========
    *variant*  Border width
    =========  ===========
    ``"1"``    0.5 pt (fine)
    ``"2"``    1.0 pt (medium)
    ``"3"``    1.5 pt (thick)
    =========  ===========

    Args:
        doc: python-docx ``Document`` object.
        variant: One of ``"1"``, ``"2"``, ``"3"``.  Defaults to ``"1"``.

    Returns:
        The newly created ``Paragraph`` with a bottom border.
    """
    p = doc.add_paragraph(style=style)
    if style is None:
        _apply_bottom_border(p, variant)
    logger.debug(f"Created horizontal-rule paragraph: variant={variant}")
    return p


# ═════════════════════════════════════════════════════════════════════════
# Attach-to-previous helpers
# ═════════════════════════════════════════════════════════════════════════


def append_page_break_to_paragraph(paragraph) -> None:
    """Append a page break run to an existing paragraph.

    Adds ``<w:r><w:br w:type="page"/></w:r>`` to the end of *paragraph*.

    Args:
        paragraph: A python-docx ``Paragraph`` object.
    """
    run_el = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run_el.append(br)
    paragraph._p.append(run_el)
    logger.debug("Appended page break to paragraph")


def append_section_break_to_paragraph(paragraph, doc, section_type: str = "next") -> None:
    """Add a section break to an existing paragraph's ``pPr``.

    Inserts ``<w:sectPr>`` into the paragraph properties of *paragraph*.

    Args:
        paragraph: A python-docx ``Paragraph`` object.
        doc: python-docx ``Document`` object (for copying page settings).
        section_type: One of ``"next"``, ``"continuous"``, ``"even"``,
            ``"odd"``.  Defaults to ``"next"``.
    """
    word_type = _SECTION_TYPE_MAP.get(section_type, "nextPage")

    pPr = _get_or_add_pPr(paragraph._p)
    sectPr = OxmlElement("w:sectPr")

    stype = OxmlElement("w:type")
    stype.set(qn("w:val"), word_type)
    sectPr.append(stype)

    _copy_page_settings(doc, sectPr)
    pPr.append(sectPr)

    logger.debug(f"Appended section break to paragraph: type={word_type}")


def append_horizontal_rule_to_paragraph(paragraph, variant: str = "1") -> None:
    """Add a bottom border to an existing paragraph.

    Creates a horizontal-rule visual effect by adding ``<w:pBdr>``
    with a ``<w:bottom>`` border to the paragraph properties.

    Args:
        paragraph: A python-docx ``Paragraph`` object.
        variant: One of ``"1"``, ``"2"``, ``"3"``.  Defaults to ``"1"``.
    """
    _apply_bottom_border(paragraph, variant)
    logger.debug(f"Appended horizontal-rule border to paragraph: variant={variant}")


# ═════════════════════════════════════════════════════════════════════════
# Internal helpers
# ═════════════════════════════════════════════════════════════════════════


def _get_or_add_pPr(p_elem):
    """Return the ``w:pPr`` child of *p_elem*, creating it if absent.

    The ``pPr`` element is inserted as the first child so it precedes
    any ``w:r`` runs.
    """
    # Use the namespace-qualified tag for find
    tag = f"{{{_WORD_NS}}}pPr"
    pPr = p_elem.find(tag)
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_elem.insert(0, pPr)
    return pPr


def _copy_page_settings(doc, target_sectPr) -> None:
    """Copy page-layout settings from the document's last ``sectPr``.

    Copies ``pgSz``, ``pgMar``, ``cols``, and ``docGrid`` elements
    from the document body's default section properties into
    *target_sectPr* so that Word / WPS correctly recognises the
    section break type.
    """
    try:
        body = doc.element.body
        doc_sectPr = body.find(f"{{{_WORD_NS}}}sectPr")
        if doc_sectPr is None:
            logger.warning("No default sectPr found in document body; section break will use Word defaults.")
            return

        for tag in ("pgSz", "pgMar", "cols", "docGrid"):
            source = doc_sectPr.find(f"{{{_WORD_NS}}}{tag}")
            if source is not None:
                target_sectPr.append(copy.deepcopy(source))
                logger.debug(f"Copied page setting: {tag}")
    except Exception as exc:
        logger.warning(f"Failed to copy page settings: {exc}; section break will use Word defaults.")


def _apply_bottom_border(paragraph_or_paragraph, variant: str = "1") -> None:
    """Add a ``w:pBdr/w:bottom`` element to the paragraph's ``pPr``.

    *paragraph_or_paragraph* may be a python-docx ``Paragraph`` object or
    a raw lxml ``<w:p>`` element.
    """
    sz = _HR_BORDER_SZ.get(variant, "4")

    p_elem = paragraph_or_paragraph._p if hasattr(paragraph_or_paragraph, "_p") else paragraph_or_paragraph
    pPr = _get_or_add_pPr(p_elem)

    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)
