"""Template loading helpers for MD→DOCX conversion.

The runtime owns template discovery and content validation. This plugin only
loads the resolved path it receives, or creates the built-in fallback when no
template was selected.
"""

# pyright: reportGeneralTypeIssues=false

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Emu, Mm, Pt

from docwen_plugin_markdown.yaml_processor import BODY_PLACEHOLDER_ALIASES

# ── Built-in fallback template ──────────────────────────────────────────────

_BUILTIN_PLACEHOLDER = "{{正文}}"

# Default theme font fallback values: 宋体 (SimSun) for East-Asian,
# Calibri for Latin, 10.5 pt (五号) default size.
_DEFAULT_BODY_FONT: dict[str, Any] = {
    "name": "Calibri",
    "east_asia": "宋体",
    "size": Pt(10.5),
    "bold": False,
    "italic": False,
    "underline": False,
}


class TemplatePackageError(ValueError):
    """A stable invalid-input failure for a malformed template package."""

    diagnostic_code = "MD2DOCX-TEMPLATE-PACKAGE-INVALID"


@dataclass(frozen=True, slots=True)
class BodyParagraphFormat:
    """Direct paragraph-format contract projected from ``{{body}}``."""

    alignment: Any
    left_indent: Any
    right_indent: Any
    first_line_indent: Any
    space_before: Any
    space_after: Any
    line_spacing: Any
    line_spacing_rule: Any
    keep_together: bool | None
    keep_with_next: bool | None
    page_break_before: bool | None
    widow_control: bool | None

    def apply_to(self, paragraph: Any) -> None:
        paragraph_format = paragraph.paragraph_format
        for name in (
            "alignment",
            "left_indent",
            "right_indent",
            "first_line_indent",
            "space_before",
            "space_after",
            "line_spacing",
            "line_spacing_rule",
            "keep_together",
            "keep_with_next",
            "page_break_before",
            "widow_control",
        ):
            value = getattr(self, name)
            if value is not None:
                setattr(paragraph_format, name, value)


def _create_builtin_template() -> Document:
    """Create a minimal A4 document with ``{{正文}}`` as body placeholder.

    The built-in template:
    - A4 page (21 cm × 29.7 cm)
    - 2.54 cm margins on all sides
    - A single paragraph containing ``{{正文}}``
    """
    doc = Document()

    # ── Page setup ──────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)

    # Margins: 2.54 cm = 1 inch = 1440 twips ≈ 914400 EMUs
    margin = Emu(914400)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin

    # ── Placeholder paragraph ───────────────────────────────────────
    p = doc.add_paragraph()
    run = p.add_run(_BUILTIN_PLACEHOLDER)

    # Set default font properties on the run
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    _set_run_east_asian_font(run, "宋体")

    return doc


def _set_run_east_asian_font(run, font_name: str) -> None:
    """Set the East Asian font on a run via ``<w:rFonts>``."""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


# ═══════════════════════════════════════════════════════════════════════════
# Template loading
# ═══════════════════════════════════════════════════════════════════════════


def load_template(template_path: str | Path | None = None) -> Document:
    """Load a template file or return the built-in fallback.

    Args:
        template_path: Path to a ``.docx`` file to use as template.
            ``None`` returns the built-in A4 document.

    Returns:
        A python-docx ``Document`` opened from *template_path* or
        the built-in fallback.
    """
    if template_path is not None:
        path = Path(template_path)
        try:
            with ZipFile(path, "r") as archive:
                names = archive.namelist()
                if len(names) != len(set(names)):
                    raise TemplatePackageError("DOCX template contains duplicate ZIP members.")
        except TemplatePackageError:
            raise
        except (BadZipFile, OSError) as exc:
            raise TemplatePackageError(f"DOCX template package is unreadable: {exc}") from exc
        with path.open("rb") as stream:
            return Document(stream)
    return _create_builtin_template()


def resolve_template(
    template_name: str | Path | None = None,
) -> Document:
    """Load a runtime-resolved template path or the built-in fallback."""
    if template_name is None or not str(template_name).strip():
        return _create_builtin_template()
    return load_template(Path(template_name).expanduser().resolve(strict=True))


# ═══════════════════════════════════════════════════════════════════════════
# Placeholder scanning
# ═══════════════════════════════════════════════════════════════════════════


def find_body_placeholder(doc: Document) -> Any | None:
    """Locate the ``{{正文}}`` or ``{{body}}`` paragraph in *doc*.

    Returns the python-docx ``Paragraph`` containing the placeholder,
    or ``None`` if no placeholder is found.
    """
    for p in doc.paragraphs:
        key = _whole_paragraph_placeholder_key(p.text)
        if key in BODY_PLACEHOLDER_ALIASES:
            return p
    return None


def extract_body_font(doc: Document) -> dict[str, Any]:
    """Extract effective font settings from the body placeholder.

    Returns effective ``name``, ``east_asia``, ``size`` (a ``Pt`` value),
    ``bold``, ``italic``, and ``underline`` values.  A distributed template
    normally keeps these values on
    the placeholder's paragraph style rather than on the literal
    ``{{body}}`` run, so direct run properties take precedence and the
    paragraph style supplies the inherited values.  Falls back to the
    default theme font (宋体/Calibri/10.5 pt) when neither source defines a
    value.
    """
    placeholder = find_body_placeholder(doc)
    if placeholder is None:
        return dict(_DEFAULT_BODY_FONT)

    for run in placeholder.runs:
        if run.text and run.text.strip():
            font = run.font
            style = placeholder.style
            return {
                "name": font.name or _style_font_name(style) or _DEFAULT_BODY_FONT["name"],
                "east_asia": _get_east_asian_font(run, style),
                "size": font.size or _style_font_size(style) or _DEFAULT_BODY_FONT["size"],
                "bold": _effective_font_property(font, style, "bold", False),
                "italic": _effective_font_property(font, style, "italic", False),
                "underline": _effective_font_property(font, style, "underline", False),
            }

    return dict(_DEFAULT_BODY_FONT)


def extract_body_style(doc: Document) -> Any | None:
    """Return the paragraph style owned by the body placeholder.

    The placeholder is the template's authoritative body-format contract;
    using its style object avoids duplicating localized style names or
    unstable style IDs in plugin code.
    """
    placeholder = find_body_placeholder(doc)
    return placeholder.style if placeholder is not None else None


def extract_body_paragraph_format(doc: Document) -> BodyParagraphFormat | None:
    """Snapshot direct body formatting without copying unrelated OOXML."""
    placeholder = find_body_placeholder(doc)
    if placeholder is None:
        return None
    paragraph_format = placeholder.paragraph_format
    return BodyParagraphFormat(
        alignment=paragraph_format.alignment,
        left_indent=paragraph_format.left_indent,
        right_indent=paragraph_format.right_indent,
        first_line_indent=paragraph_format.first_line_indent,
        space_before=paragraph_format.space_before,
        space_after=paragraph_format.space_after,
        line_spacing=paragraph_format.line_spacing,
        line_spacing_rule=paragraph_format.line_spacing_rule,
        keep_together=paragraph_format.keep_together,
        keep_with_next=paragraph_format.keep_with_next,
        page_break_before=paragraph_format.page_break_before,
        widow_control=paragraph_format.widow_control,
    )


def _style_font_name(style: Any) -> str | None:
    """Return the first explicitly defined Latin font in a style chain."""
    current = style
    while current is not None:
        name = current.font.name
        if name:
            return name
        current = current.base_style
    return None


def _style_font_size(style: Any):
    """Return the first explicitly defined font size in a style chain."""
    current = style
    while current is not None:
        size = current.font.size
        if size is not None:
            return size
        current = current.base_style
    return None


def _effective_font_property(font: Any, style: Any, property_name: str, default: Any) -> Any:
    """Resolve a boolean-like run font property through its style chain."""

    direct = getattr(font, property_name, None)
    if direct is not None:
        return direct
    current = style
    while current is not None:
        value = getattr(current.font, property_name, None)
        if value is not None:
            return value
        current = current.base_style
    return default


def _get_east_asian_font(run, style: Any = None) -> str:
    """Read ``w:eastAsia`` from a run, then its paragraph style chain."""
    try:
        rPr = run._r.find(qn("w:rPr"))
        if rPr is not None:
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is not None:
                ea = rFonts.get(qn("w:eastAsia"))
                if ea:
                    return ea
    except Exception:
        pass

    current = style
    while current is not None:
        try:
            rPr = current._element.find(qn("w:rPr"))
            rFonts = rPr.find(qn("w:rFonts")) if rPr is not None else None
            if rFonts is not None:
                ea = rFonts.get(qn("w:eastAsia"))
                if ea:
                    return ea
        except Exception:
            pass
        current = current.base_style
    return _DEFAULT_BODY_FONT["east_asia"]


def scan_placeholders(doc: Document) -> dict[str, list[Any]]:
    """Walk *doc* and find all ``{{key}}`` occurrences.

    Returns:
        ``{key: [paragraph, ...]}`` mapping for every placeholder
        occurrence (without braces). A paragraph is listed once per key even
        when that key occurs more than once in the paragraph.
    """
    pattern = re.compile(r"\{\{\s*([^{}\r\n]+?)\s*\}\}")
    placeholders: dict[str, list[Any]] = {}
    for p in _iter_document_paragraphs(doc):
        keys = dict.fromkeys(match.group(1).strip() for match in pattern.finditer(p.text))
        for key in keys:
            placeholders.setdefault(key, []).append(p)
    return placeholders


def _whole_paragraph_placeholder_key(text: str) -> str | None:
    match = re.fullmatch(r"\s*\{\{\s*([^{}\r\n]+?)\s*\}\}\s*", text or "")
    if match is None:
        return None
    return match.group(1).strip()


def _iter_document_paragraphs(doc: Document):
    # Keep the element wrappers themselves alive while traversing. Using
    # ``id(paragraph._element)`` is unsafe here because lxml may release a
    # wrapper between cells and reuse its Python object id, which silently
    # drops a later, distinct table paragraph from placeholder scanning.
    seen: set[Any] = set()

    for paragraph in doc.paragraphs:
        seen.add(paragraph._element)
        yield paragraph

    for table in doc.tables:
        yield from _iter_table_paragraphs(table, seen)


def _iter_table_paragraphs(table, seen: set[Any]):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                element = paragraph._element
                if element in seen:
                    continue
                seen.add(element)
                yield paragraph
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table, seen)
