"""Inline rendering utilities for MD→DOCX conversion.

Provides:
- ``render_inlines()`` — dispatch inline AST nodes to runs
- ``render_formula()`` — inline/block OMML formula
- ``handle_footnote_ref()`` — footnote/endnote reference
- ``extract_text_content()`` — plain text extraction
- ``add_hyperlink()`` — clickable hyperlink
- ``embed_image()`` — inline image embedding
"""

from __future__ import annotations

import logging
import re
from collections.abc import Callable, Iterable
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path, PureWindowsPath
from typing import Any
from urllib.parse import quote, unquote, urlsplit, urlunsplit

from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt

from docwen_plugin_markdown.renderer_utils import (
    add_word_soft_break,
    apply_run_east_asian_font,
    apply_run_shading,
)

logger = logging.getLogger(__name__)


def render_inlines(
    parent,
    children: list[dict[str, Any]],
    code_font: str = "Consolas",
    code_bg_color: str = "E7E6E6",
    override_style: bool = False,
    strip_formatting: bool = False,
    renderer_instance: Any = None,
) -> None:
    """Render inline children into a paragraph or heading element.

    When *override_style* is True (used in heading contexts with
    ``heading_formatting_mode="apply"``), unmarked text runs are
    explicitly set to ``bold=False`` / ``italic=False`` so the heading
    style's defaults do not force bold on plain text.

    When *strip_formatting* is True (used with
    ``heading_formatting_mode="remove"``), bold/italic/strikethrough
    formatting is omitted from all runs, but semantic nodes such as
    formulas, links, images, and code spans are still rendered.

    Args:
        parent: A python-docx ``Paragraph`` or ``Heading`` object.
        children: List of mistune AST inline node dicts.
        code_font: Monospace font name for code spans.
        code_bg_color: Background color for code spans.
        override_style: If True, set bold=False + italic=False on plain text.
        strip_formatting: If True, omit bold/italic/strikethrough formatting.
        renderer_instance: Optional renderer instance with access to
            ``_embed_image``, ``_add_hyperlink``, etc.
    """
    for child in children:
        ctype = child.get("type", "")
        if ctype == "text":
            _render_text(parent, child, override_style)
        elif ctype == "strong":
            _render_strong(
                parent,
                child,
                strip_formatting,
                code_font=code_font,
                code_bg_color=code_bg_color,
                override_style=override_style,
                renderer_instance=renderer_instance,
            )
        elif ctype == "emphasis":
            _render_emphasis(
                parent,
                child,
                strip_formatting,
                code_font=code_font,
                code_bg_color=code_bg_color,
                override_style=override_style,
                renderer_instance=renderer_instance,
            )
        elif ctype == "strikethrough":
            _render_strikethrough(
                parent,
                child,
                strip_formatting,
                code_font=code_font,
                code_bg_color=code_bg_color,
                override_style=override_style,
                renderer_instance=renderer_instance,
            )
        elif ctype == "codespan":
            _render_codespan(parent, child, code_font, code_bg_color, renderer_instance)
        elif ctype == "link":
            _render_link(
                parent,
                child,
                renderer_instance,
                override_style=override_style,
                strip_formatting=strip_formatting,
            )
        elif ctype == "image":
            _render_image(parent, child, renderer_instance)
        elif ctype in ("linebreak", "softbreak"):
            add_word_soft_break(parent)
        elif ctype in ("inline_math", "inline_latex"):
            _render_inline_math(parent, child, renderer_instance)
        elif ctype == "inline_html":
            _render_inline_html(parent, child)
        elif ctype in ("highlight", "mark"):
            _render_highlight(parent, child)
        elif ctype == "superscript":
            _render_superscript(parent, child)
        elif ctype == "subscript":
            _render_subscript(parent, child)
        elif ctype in ("underline", "insert"):
            _render_underline(parent, child)
        elif ctype == "footnote_ref":
            _render_footnote_ref(parent, child, renderer_instance)
        elif ctype == "semantic_cross_reference":
            if renderer_instance is not None:
                renderer_instance.render_semantic_reference(parent, child)
            else:
                parent.add_run(f"@{child.get('target_id', '')}")
        elif ctype == "semantic_citation":
            if renderer_instance is not None:
                renderer_instance.render_semantic_citation(parent, child)
            else:
                parent.add_run(child.get("raw", ""))
        else:
            # Render nested children
            nested = child.get("children", [])
            if nested:
                render_inlines(
                    parent,
                    nested,
                    code_font=code_font,
                    code_bg_color=code_bg_color,
                    override_style=override_style,
                    strip_formatting=strip_formatting,
                    renderer_instance=renderer_instance,
                )


def _render_text(parent, child: dict[str, Any], override_style: bool) -> None:
    """Render a plain text node, splitting emoji for dedicated font rendering."""
    text = child.get("raw", "") or child.get("text", "")
    emoji_parts = _split_text_by_emoji(text)
    for part_text, is_emoji in emoji_parts:
        if not part_text:
            continue
        run = parent.add_run(part_text)
        if override_style:
            run.bold = False
            run.italic = False
        if is_emoji:
            _set_emoji_font(run)


def _render_strong(
    parent,
    child: dict[str, Any],
    strip_formatting: bool,
    *,
    code_font: str,
    code_bg_color: str,
    override_style: bool,
    renderer_instance: Any,
) -> None:
    """Render a **bold** node."""
    _render_formatted_container(
        parent,
        child,
        strip_formatting=strip_formatting,
        code_font=code_font,
        code_bg_color=code_bg_color,
        override_style=override_style,
        renderer_instance=renderer_instance,
        property_name="bold",
    )


def _render_emphasis(
    parent,
    child: dict[str, Any],
    strip_formatting: bool,
    *,
    code_font: str,
    code_bg_color: str,
    override_style: bool,
    renderer_instance: Any,
) -> None:
    """Render an *italic* node."""
    _render_formatted_container(
        parent,
        child,
        strip_formatting=strip_formatting,
        code_font=code_font,
        code_bg_color=code_bg_color,
        override_style=override_style,
        renderer_instance=renderer_instance,
        property_name="italic",
    )


def _render_strikethrough(
    parent,
    child: dict[str, Any],
    strip_formatting: bool,
    *,
    code_font: str,
    code_bg_color: str,
    override_style: bool,
    renderer_instance: Any,
) -> None:
    """Render a ~~strikethrough~~ node."""
    _render_formatted_container(
        parent,
        child,
        strip_formatting=strip_formatting,
        code_font=code_font,
        code_bg_color=code_bg_color,
        override_style=override_style,
        renderer_instance=renderer_instance,
        property_name="strike",
    )


def _render_formatted_container(
    parent,
    child: dict[str, Any],
    *,
    strip_formatting: bool,
    code_font: str,
    code_bg_color: str,
    override_style: bool,
    renderer_instance: Any,
    property_name: str,
) -> None:
    children = child.get("children")
    if not isinstance(children, list) or not children:
        text = str(child.get("raw", "") or child.get("text", ""))
        run = parent.add_run(text)
        if not strip_formatting:
            if property_name == "strike":
                run.font.strike = True
            else:
                setattr(run, property_name, True)
        return

    first_run = len(parent.runs)
    render_inlines(
        parent,
        children,
        code_font=code_font,
        code_bg_color=code_bg_color,
        override_style=override_style,
        strip_formatting=strip_formatting,
        renderer_instance=renderer_instance,
    )
    if strip_formatting:
        return
    for run in parent.runs[first_run:]:
        if property_name == "strike":
            run.font.strike = True
        else:
            setattr(run, property_name, True)


def _render_codespan(
    parent,
    child: dict[str, Any],
    code_font: str,
    code_bg_color: str,
    renderer_instance: Any,
) -> None:
    """Render an inline ``code`` span."""
    text = child.get("raw", "") or child.get("text", "")
    run = parent.add_run(text)
    inline_style = getattr(renderer_instance, "inline_code_style", None)
    if inline_style is not None:
        run.style = inline_style
        return
    run.font.name = code_font
    apply_run_shading(run, code_bg_color)


def _render_link(
    parent,
    child: dict[str, Any],
    renderer_instance: Any,
    *,
    override_style: bool,
    strip_formatting: bool,
) -> None:
    """Render a ``[text](url)`` link."""
    children = child.get("children", [])
    url = child.get("attrs", {}).get("url", "")
    if url and renderer_instance and hasattr(renderer_instance, "add_hyperlink"):
        renderer_instance.add_hyperlink(
            parent,
            url,
            children,
            override_style=override_style,
            strip_formatting=strip_formatting,
        )
    elif url:
        add_hyperlink(
            parent,
            url,
            children=children,
            override_style=override_style,
            strip_formatting=strip_formatting,
        )
    else:
        render_inlines(
            parent,
            children,
            code_font=getattr(renderer_instance, "_code_font", "Consolas"),
            code_bg_color=getattr(renderer_instance, "_code_bg_color", "E7E6E6"),
            override_style=override_style,
            strip_formatting=strip_formatting,
            renderer_instance=renderer_instance,
        )


def _render_image(parent, child: dict[str, Any], renderer_instance: Any) -> None:
    """Render an ``![alt](src)`` image."""
    attrs = child.get("attrs", {})
    src = attrs.get("url", "")
    alt = extract_text_content(child.get("children", []))
    width, height = _parse_docwen_image_size(attrs.get("title", ""))
    if renderer_instance and hasattr(renderer_instance, "embed_image"):
        renderer_instance.embed_image(parent, src, alt=alt, width=width, height=height)
    else:
        embed_image(parent, src, alt=alt, width=width, height=height)


def _parse_docwen_image_size(title: str) -> tuple[int | None, int | None]:
    match = re.fullmatch(r"docwen-size=(\d*)x(\d*)", title or "")
    if match is None:
        return None, None
    width_text, height_text = match.groups()
    width = int(width_text) if width_text else None
    height = int(height_text) if height_text else None
    return width, height


def _render_inline_math(parent, child: dict[str, Any], renderer_instance: Any) -> None:
    """Render inline math as OMML."""
    from docwen_plugin_markdown.to_docx.formula_renderer import render_inline_formula

    raw = child.get("raw", "") or child.get("text", "")
    if raw.strip():
        style_reference = getattr(renderer_instance, "inline_formula_style_reference", None)
        render_inline_formula(parent, raw.strip(), style_reference=style_reference)


def _render_inline_html(parent, child: dict[str, Any]) -> None:
    """Render inline HTML — strip or forward."""
    text = child.get("raw", "") or child.get("text", "")
    if text.strip():
        parent.add_run(text)


def _render_highlight(parent, child: dict[str, Any]) -> None:
    """Render ``==highlight==`` as Word highlighted text."""
    text = child.get("raw", "") or extract_text_content(child.get("children", []))
    run = parent.add_run(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _render_superscript(parent, child: dict[str, Any]) -> None:
    """Render ``^superscript^`` as superscript text."""
    text = child.get("raw", "") or extract_text_content(child.get("children", []))
    run = parent.add_run(text)
    run.font.superscript = True


def _render_subscript(parent, child: dict[str, Any]) -> None:
    """Render ``~subscript~`` as subscript text."""
    text = child.get("raw", "") or extract_text_content(child.get("children", []))
    run = parent.add_run(text)
    run.font.subscript = True


def _render_underline(parent, child: dict[str, Any]) -> None:
    """Render ``<u>underline</u>`` as underlined text."""
    text = child.get("raw", "") or extract_text_content(child.get("children", []))
    parent.add_run(text)
    parent.runs[-1].underline = True


def _render_footnote_ref(parent, child: dict[str, Any], renderer_instance: Any) -> None:
    """Render a footnote reference via the renderer instance if available."""
    if renderer_instance and hasattr(renderer_instance, "_note_ctx"):
        renderer_instance._render_footnote_ref(parent, child)
    else:
        parent.add_run(f"[^{child.get('raw', '')}]")


# ═══════════════════════════════════════════════════════════════════════════
# Formula rendering (block + inline)
# ═══════════════════════════════════════════════════════════════════════════


def render_formula(doc, node: dict[str, Any], *, block_style=None) -> None:
    """Render a block-level formula as OMML in a new paragraph.

    Args:
        doc: python-docx ``Document``.
        node: Mistune AST node for block_latex or block_math.
    """
    from docwen_plugin_markdown.to_docx.formula_renderer import render_block_formula

    raw = node.get("text", "") or node.get("raw", "")
    render_block_formula(doc, raw.strip(), block_style=block_style)


# ═══════════════════════════════════════════════════════════════════════════
# Hyperlink and image helpers
# ═══════════════════════════════════════════════════════════════════════════


def add_hyperlink(
    paragraph,
    url: str,
    text: str | None = None,
    *,
    children: list[dict[str, Any]] | None = None,
    source_dir: Path | None = None,
    style_reference: str | None = None,
    override_style: bool = False,
    strip_formatting: bool = False,
    code_font: str = "Consolas",
    code_bg_color: str = "E7E6E6",
    inline_code_properties: tuple[Any, ...] | None = None,
    fallback_renderer: Callable[[Any, list[dict[str, Any]]], None] | None = None,
) -> None:
    """Add a clickable hyperlink to a DOCX paragraph.

    Creates an OOXML ``<w:hyperlink>`` element with the target URL
    and relationship. Falls back to honest plain text if the URL is invalid
    or the hyperlink creation fails.
    """
    link_children = children if children is not None else [{"type": "text", "raw": text or ""}]
    target = _resolve_hyperlink_target(url, source_dir)
    if not target or not target.lower().startswith(("http://", "https://", "mailto:", "ftp://", "file://")):
        if fallback_renderer is not None:
            fallback_renderer(paragraph, link_children)
            return
        _append_plain_link_runs(
            paragraph,
            link_children,
            override_style=override_style,
            strip_formatting=strip_formatting,
            code_font=code_font,
            code_bg_color=code_bg_color,
            inline_code_properties=inline_code_properties,
        )
        return

    try:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml.shared import OxmlElement  # pyright: ignore[reportPrivateImportUsage]

        hyperlink = OxmlElement("w:hyperlink")

        for segment in _link_segments(link_children):
            for segment_text, is_emoji in _split_link_segment(segment):
                if not segment_text:
                    continue
                r = OxmlElement("w:r")
                r_pr, line_break = _hyperlink_run_properties(
                    segment,
                    is_emoji=is_emoji,
                    style_reference=style_reference,
                    override_style=override_style,
                    strip_formatting=strip_formatting,
                    code_font=code_font,
                    code_bg_color=code_bg_color,
                    inline_code_properties=inline_code_properties,
                )
                r.append(r_pr)
                if line_break is not None:
                    r.append(line_break)
                else:
                    t = OxmlElement("w:t")
                    t.set(qn("xml:space"), "preserve")
                    t.text = segment_text
                    r.append(t)
                hyperlink.append(r)

        part = paragraph.part
        existing_relationship_ids = set(part.rels)
        r_id = part.relate_to(target, RT.HYPERLINK, is_external=True)
        try:
            hyperlink.set(qn("r:id"), r_id)
            paragraph._p.append(hyperlink)
        except Exception:
            if r_id not in existing_relationship_ids:
                part.drop_rel(r_id)
            raise
    except Exception:
        if fallback_renderer is not None:
            fallback_renderer(paragraph, link_children)
            return
        _append_plain_link_runs(
            paragraph,
            link_children,
            override_style=override_style,
            strip_formatting=strip_formatting,
            code_font=code_font,
            code_bg_color=code_bg_color,
            inline_code_properties=inline_code_properties,
        )


@dataclass(frozen=True, slots=True)
class _LinkSegment:
    text: str
    bold: bool = False
    italic: bool = False
    strike: bool = False
    underline: bool = False
    code: bool = False
    highlight: bool = False
    superscript: bool = False
    subscript: bool = False
    formula: bool = False
    line_break: bool = False

    @property
    def is_plain(self) -> bool:
        return not any(
            (
                self.bold,
                self.italic,
                self.strike,
                self.underline,
                self.code,
                self.highlight,
                self.superscript,
                self.subscript,
                self.formula,
                self.line_break,
            )
        )


def _link_segments(
    children: list[dict[str, Any]],
    *,
    bold: bool = False,
    italic: bool = False,
    strike: bool = False,
    underline: bool = False,
    code: bool = False,
    highlight: bool = False,
    superscript: bool = False,
    subscript: bool = False,
    formula: bool = False,
) -> list[_LinkSegment]:
    segments: list[_LinkSegment] = []
    for child in children:
        child_type = child.get("type", "")
        nested = child.get("children", [])
        if child_type in {"linebreak", "softbreak"}:
            segments.append(
                _LinkSegment(
                    " ",
                    bold,
                    italic,
                    strike,
                    underline,
                    code,
                    highlight,
                    superscript,
                    subscript,
                    formula,
                    line_break=True,
                )
            )
        elif nested:
            segments.extend(
                _link_segments(
                    nested,
                    bold=bold or child_type == "strong",
                    italic=italic or child_type == "emphasis",
                    strike=strike or child_type == "strikethrough",
                    underline=underline or child_type in {"underline", "insert"},
                    code=code or child_type == "codespan",
                    highlight=highlight or child_type in {"highlight", "mark"},
                    superscript=superscript or (child_type == "superscript" and not subscript),
                    subscript=subscript or (child_type == "subscript" and not superscript),
                    formula=formula or child_type in {"inline_math", "inline_latex"},
                )
            )
        else:
            value = child.get("raw", "") or child.get("text", "")
            if child_type in {"inline_math", "inline_latex"}:
                value = f"${value}$"
            segments.append(
                _LinkSegment(
                    str(value),
                    bold,
                    italic,
                    strike,
                    underline,
                    code or child_type == "codespan",
                    highlight or child_type in {"highlight", "mark"},
                    superscript or (child_type == "superscript" and not subscript),
                    subscript or (child_type == "subscript" and not superscript),
                    formula or child_type in {"inline_math", "inline_latex"},
                )
            )
    return segments


_RUN_PROPERTY_ORDER = (
    "rStyle",
    "rFonts",
    "b",
    "bCs",
    "i",
    "iCs",
    "caps",
    "smallCaps",
    "strike",
    "dstrike",
    "outline",
    "shadow",
    "emboss",
    "imprint",
    "noProof",
    "snapToGrid",
    "vanish",
    "webHidden",
    "color",
    "spacing",
    "w",
    "kern",
    "position",
    "sz",
    "szCs",
    "highlight",
    "u",
    "effect",
    "bdr",
    "shd",
    "fitText",
    "vertAlign",
    "rtl",
    "cs",
    "em",
    "lang",
    "eastAsianLayout",
    "specVanish",
    "oMath",
)
_RUN_PROPERTY_RANK = {qn(f"w:{name}"): index for index, name in enumerate(_RUN_PROPERTY_ORDER)}


def project_character_style_properties(style: Any, *, paragraph: Any | None = None) -> tuple[Any, ...] | None:
    """Freeze effective character-style properties for a hyperlink run.

    A hyperlink run can reference only one character style. Code labels keep
    ``Hyperlink`` as that identity and project the request-owned Inline Code
    style chain as direct properties. Later (leaf) definitions win by QName.
    """

    if style is None:
        return None
    chain: list[Any] = []
    seen: set[str] = set()
    current = style
    while current is not None:
        style_id = str(getattr(current, "style_id", ""))
        if style_id in seen:
            break
        seen.add(style_id)
        chain.append(current)
        current = getattr(current, "base_style", None)
    properties: dict[str, Any] = {}
    declared_toggles: set[str] = set()
    for item in reversed(chain):
        r_pr = getattr(getattr(item, "_element", None), "rPr", None)
        if r_pr is None:
            continue
        for child in r_pr:
            if child.tag == qn("w:rStyle"):
                continue
            if child.tag in _TOGGLE_PROPERTY_QNAMES:
                declared_toggles.add(child.tag)
                continue
            existing = properties.get(child.tag)
            if existing is not None and child.tag in _ATTRIBUTE_MERGED_PROPERTY_QNAMES:
                merged = deepcopy(existing)
                for name, value in child.attrib.items():
                    merged.set(name, value)
                properties[child.tag] = merged
            else:
                properties[child.tag] = deepcopy(child)
    if paragraph is not None:
        for tag in declared_toggles:
            element = OxmlElement(f"w:{tag.rsplit('}', 1)[-1]}")
            if not _effective_toggle_value(paragraph, style, tag):
                element.set(qn("w:val"), "0")
            properties[tag] = element
    return tuple(_ordered_run_properties(properties.values()))


def _effective_toggle_value(paragraph: Any, character_style: Any, tag: str) -> bool:
    value = False
    styles_element = getattr(getattr(paragraph.part.document, "styles", None), "element", None)
    if styles_element is not None:
        default = styles_element.find(f"{qn('w:docDefaults')}/{qn('w:rPrDefault')}/{qn('w:rPr')}/{tag}")
        value = _on_off_value(default)
    for style in (
        *_style_chain_root_first(getattr(paragraph, "style", None)),
        *_style_chain_root_first(character_style),
    ):
        element = getattr(style, "_element", None)
        r_pr = getattr(element, "rPr", None)
        declaration = r_pr.find(tag) if r_pr is not None else None
        if _on_off_value(declaration):
            value = not value
    return value


def _style_chain_root_first(style: Any) -> tuple[Any, ...]:
    chain: list[Any] = []
    seen: set[int] = set()
    current = style
    while current is not None and id(current) not in seen:
        seen.add(id(current))
        chain.append(current)
        current = getattr(current, "base_style", None)
    return tuple(reversed(chain))


def _on_off_value(element: Any | None) -> bool:
    if element is None:
        return False
    value = str(element.get(qn("w:val"), "1")).strip().lower()
    return value not in {"0", "false", "off", "no"}


def _split_link_segment(segment: _LinkSegment) -> list[tuple[str, bool]]:
    if segment.is_plain:
        return _split_text_by_emoji(segment.text)
    return [(segment.text, False)]


def _hyperlink_run_properties(
    segment: _LinkSegment,
    *,
    is_emoji: bool,
    style_reference: str | None,
    override_style: bool,
    strip_formatting: bool,
    code_font: str,
    code_bg_color: str,
    inline_code_properties: tuple[Any, ...] | None,
):
    from docx.oxml.shared import OxmlElement  # pyright: ignore[reportPrivateImportUsage]

    properties: dict[str, Any] = {}
    if segment.code:
        for child in inline_code_properties or ():
            if child.tag != qn("w:rStyle"):
                properties[child.tag] = deepcopy(child)
        if inline_code_properties is None:
            fonts = OxmlElement("w:rFonts")
            for attribute in ("ascii", "hAnsi", "eastAsia"):
                fonts.set(qn(f"w:{attribute}"), code_font)
            properties[fonts.tag] = fonts
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), code_bg_color)
            properties[shading.tag] = shading
    if style_reference:
        style = OxmlElement("w:rStyle")
        style.set(qn("w:val"), style_reference)
        properties[style.tag] = style
    if is_emoji:
        fonts = OxmlElement("w:rFonts")
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn(f"w:{attribute}"), "Segoe UI Emoji")
        properties[fonts.tag] = fonts

    resolved_bold = (
        None if strip_formatting else True if segment.bold else False if override_style and segment.is_plain else None
    )
    resolved_italic = (
        None if strip_formatting else True if segment.italic else False if override_style and segment.is_plain else None
    )
    for enabled, tag in (
        (resolved_bold, "b"),
        (resolved_italic, "i"),
        (None if strip_formatting else segment.strike or None, "strike"),
        (segment.underline or None, "u"),
    ):
        if enabled is None:
            continue
        element = OxmlElement(f"w:{tag}")
        if not enabled:
            element.set(qn("w:val"), "0")
        elif tag == "u":
            element.set(qn("w:val"), "single")
        properties[element.tag] = element
    if segment.highlight:
        element = OxmlElement("w:highlight")
        element.set(qn("w:val"), "yellow")
        properties[element.tag] = element
    if segment.superscript or segment.subscript:
        element = OxmlElement("w:vertAlign")
        element.set(qn("w:val"), "superscript" if segment.superscript else "subscript")
        properties[element.tag] = element
    if segment.line_break:
        element = OxmlElement("w:br")
        properties[element.tag] = element

    r_pr = OxmlElement("w:rPr")
    line_break = properties.pop(qn("w:br"), None)
    for child in _ordered_run_properties(properties.values()):
        r_pr.append(child)
    return r_pr, line_break


def _ordered_run_properties(properties: Iterable[Any]) -> list[Any]:
    return sorted(properties, key=lambda child: _RUN_PROPERTY_RANK.get(child.tag, len(_RUN_PROPERTY_RANK)))


def _append_plain_link_runs(
    paragraph,
    children: list[dict[str, Any]],
    *,
    override_style: bool = False,
    strip_formatting: bool = False,
    code_font: str = "Consolas",
    code_bg_color: str = "E7E6E6",
    inline_code_properties: tuple[Any, ...] | None = None,
) -> None:
    for segment in _link_segments(children):
        if segment.line_break:
            add_word_soft_break(paragraph)
            continue
        for segment_text, is_emoji in _split_link_segment(segment):
            if not segment_text:
                continue
            run = paragraph.add_run(segment_text)
            run.bold = (
                None
                if strip_formatting
                else True
                if segment.bold
                else False
                if override_style and segment.is_plain
                else None
            )
            run.italic = (
                None
                if strip_formatting
                else True
                if segment.italic
                else False
                if override_style and segment.is_plain
                else None
            )
            run.font.strike = None if strip_formatting else segment.strike or None
            run.underline = segment.underline or None
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW if segment.highlight else None
            run.font.superscript = segment.superscript or None
            run.font.subscript = segment.subscript or None
            if segment.code:
                if inline_code_properties is not None:
                    r_pr = run._r.get_or_add_rPr()
                    for child in inline_code_properties:
                        existing = r_pr.find(child.tag)
                        if existing is not None:
                            r_pr.remove(existing)
                        r_pr.append(deepcopy(child))
                else:
                    run.font.name = code_font
                    apply_run_east_asian_font(run, code_font)
                    apply_run_shading(run, code_bg_color)
            if is_emoji:
                _set_emoji_font(run)
            r_pr = run._r.rPr
            if r_pr is not None:
                children_in_order = list(r_pr)
                for child in children_in_order:
                    r_pr.remove(child)
                for child in _ordered_run_properties(children_in_order):
                    r_pr.append(child)


_TOGGLE_PROPERTY_QNAMES = {
    qn(f"w:{name}")
    for name in (
        "b",
        "bCs",
        "i",
        "iCs",
        "caps",
        "smallCaps",
        "strike",
        "dstrike",
        "outline",
        "shadow",
        "emboss",
        "imprint",
        "vanish",
        "webHidden",
        "rtl",
        "cs",
        "specVanish",
        "oMath",
    )
}
_ATTRIBUTE_MERGED_PROPERTY_QNAMES = {qn("w:rFonts"), qn("w:lang")}


def _resolve_hyperlink_target(url: str, source_dir: Path | None) -> str:
    if not url:
        return ""
    if url.startswith("#"):
        return url

    # ``urlsplit("C:/...")`` treats the drive letter as a URI scheme and
    # drops it from ``parsed.path``.  Handle drive-absolute Windows paths
    # before URI parsing so cross-drive links remain faithful on every host.
    if re.match(r"^[A-Za-z]:[/\\]", url):
        path_and_query, separator, fragment = url.partition("#")
        path_text, query_separator, query = path_and_query.partition("?")
        target = PureWindowsPath(unquote(path_text)).as_uri()
        if query_separator:
            target += f"?{query}"
        if separator:
            target += f"#{quote(unquote(fragment), safe='/-._~')}"
        return target

    parsed = urlsplit(url)
    normalized_scheme = parsed.scheme.lower()
    if normalized_scheme in {"http", "https", "mailto", "ftp", "file"}:
        return urlunsplit(
            (
                normalized_scheme,
                parsed.netloc,
                parsed.path,
                parsed.query,
                parsed.fragment,
            )
        )
    if parsed.scheme:
        return url
    if parsed.netloc:
        # Protocol-relative URLs are not local filesystem paths.  Keep the
        # source value so the caller can honestly downgrade unsupported input.
        return url
    if source_dir is None:
        return url
    local_path = Path(unquote(parsed.path))
    if not local_path.is_absolute():
        local_path = source_dir / local_path
    target = local_path.resolve().as_uri()
    if parsed.query:
        target += f"?{parsed.query}"
    if parsed.fragment:
        target += f"#{quote(unquote(parsed.fragment), safe='/-._~')}"
    return target


def embed_image(
    paragraph,
    src: str,
    *,
    source_dir: Path | None = None,
    alt: str = "",
    width: int | None = None,
    height: int | None = None,
    declared_resource_resolver=None,
    fail_closed: bool = False,
) -> None:
    """Embed an image file into a DOCX paragraph.

    Reads the image file at ``src`` and embeds it as an inline drawing
    in the paragraph. Falls back to a text placeholder if the image
    cannot be read.
    """
    if not src:
        return
    try:
        path = (
            Path(declared_resource_resolver.resolve(src))
            if declared_resource_resolver is not None
            else _resolve_image_path(src, source_dir)
        )
        if not path.is_file():
            if fail_closed:
                raise ValueError("authenticated resolved-v4 image staging path is missing")
            paragraph.add_run(f"[Image: {path.name}]").italic = True
            return

        run = paragraph.add_run()
        width_emu, height_emu = _calculate_image_extent(
            path,
            width=width,
            height=height,
            usable_width_emu=_paragraph_usable_width_emu(paragraph),
        )
        inline_shape = run.add_picture(str(path), width=width_emu, height=height_emu)
        doc_pr = inline_shape._inline.docPr
        if alt:
            doc_pr.set("descr", alt)
        source_name = PureWindowsPath(unquote(urlsplit(src).path)).name
        if source_name:
            # ``title`` is a standard DrawingML property.  Keeping only the
            # basename avoids leaking an absolute source path while allowing
            # semantic round-trips to materialize a stable resource name.
            doc_pr.set("title", source_name)
        paragraph.add_run("\n")
    except Exception:
        if fail_closed:
            raise
        paragraph.add_run(f"[Image: {Path(src).name}]").italic = True


def _resolve_image_path(src: str, source_dir: Path | None) -> Path:
    if re.match(r"^[A-Za-z]:[/\\]", src):
        return Path(unquote(src)).resolve()
    parsed = urlsplit(src)
    if parsed.scheme.lower() == "file":
        file_path = unquote(parsed.path)
        if parsed.netloc and parsed.netloc.lower() != "localhost":
            file_path = f"//{parsed.netloc}{file_path}"
        elif re.match(r"^/[A-Za-z]:/", file_path):
            file_path = file_path[1:]
        return Path(file_path)
    if src.startswith("//") or parsed.scheme:
        raise ValueError(f"remote image target is not a local path: {src}")
    path = Path(unquote(parsed.path))
    if not path.is_absolute() and source_dir is not None:
        path = source_dir / path
    return path.resolve()


def _paragraph_usable_width_emu(paragraph) -> int | None:
    """Return the page/body or table-cell width available to an image."""
    try:
        document = paragraph.part.document
        section = document.sections[0]
        page_usable = int(section.page_width - section.left_margin - section.right_margin)
        parent = getattr(paragraph, "_parent", None)
        cell_width = getattr(parent, "width", None)
        if cell_width is not None and int(cell_width) > 0:
            return max(0, int(cell_width) - (2 * int(Pt(5.4))))
        return page_usable
    except Exception as exc:
        logger.debug("Could not determine paragraph image width: %s", exc)
        return None


def _calculate_image_extent(
    path: Path,
    *,
    width: int | None,
    height: int | None,
    usable_width_emu: int | None,
) -> tuple[Emu | None, Emu | None]:
    """Match reference DPI-aware sizing and available-width clamping."""
    try:
        from PIL import Image

        with Image.open(path) as image:
            original_width_px, original_height_px = image.size
            dpi = image.info.get("dpi")
            dpi_x = (
                float(dpi[0])
                if isinstance(dpi, tuple) and dpi and isinstance(dpi[0], (int, float)) and dpi[0] > 0
                else 96.0
            )
    except Exception:
        return None, None

    px_to_emu = 914400.0 / dpi_x
    original_width_emu = int(original_width_px * px_to_emu)
    original_height_emu = int(original_height_px * px_to_emu)
    target_width_emu = int(width * px_to_emu) if width is not None else original_width_emu
    if usable_width_emu is not None and usable_width_emu > 0 and target_width_emu > usable_width_emu:
        target_width_emu = usable_width_emu
    if height is not None:
        target_height_emu = int(height * px_to_emu)
    elif original_width_emu > 0:
        target_height_emu = int(target_width_emu * (original_height_emu / original_width_emu))
    else:
        target_height_emu = original_height_emu
    return Emu(target_width_emu), Emu(target_height_emu)


# ═══════════════════════════════════════════════════════════════════════════
# Plain text extraction
# ═══════════════════════════════════════════════════════════════════════════


def extract_text_content(children: list[dict[str, Any]]) -> str:
    """Extract plain text from inline AST children.

    Args:
        children: List of mistune AST inline node dicts.

    Returns:
        Concatenated plain text.
    """
    parts: list[str] = []
    for child in children:
        ctype = child.get("type", "")
        if ctype == "text":
            parts.append(child.get("raw", "") or child.get("text", ""))
        elif ctype == "semantic_cross_reference":
            parts.append(f"@{child.get('target_id', '')}")
        elif ctype == "semantic_citation":
            parts.append(child.get("raw", ""))
        elif ctype in ("strong", "emphasis", "codespan", "link"):
            parts.append(child.get("raw", "") or extract_text_content(child.get("children", [])))
        elif ctype in ("linebreak", "softbreak"):
            parts.append(" ")
        else:
            nested = child.get("children", [])
            if nested:
                parts.append(extract_text_content(nested))
    return "".join(parts)


def serialize_inlines_to_markdown(children: list[dict[str, Any]]) -> str:
    """Best-effort reconstruction of inline Markdown markers from AST nodes.

    Mistune normalizes inline formatting before the renderer sees it.
    ``formatting_mode="keep"`` displays the original Markdown markers, so the
    renderer reconstructs the user-visible marker form for the
    inline constructs DocWen supports.
    """
    parts: list[str] = []
    for child in children:
        ctype = child.get("type", "")
        nested = child.get("children", [])
        text = child.get("raw", "") or child.get("text", "")
        if ctype == "text":
            parts.append(text)
        elif ctype == "strong":
            parts.append(f"**{serialize_inlines_to_markdown(nested) or text}**")
        elif ctype == "emphasis":
            parts.append(f"*{serialize_inlines_to_markdown(nested) or text}*")
        elif ctype == "strikethrough":
            parts.append(f"~~{serialize_inlines_to_markdown(nested) or text}~~")
        elif ctype == "codespan":
            parts.append(f"`{text}`")
        elif ctype == "link":
            url = child.get("attrs", {}).get("url", "")
            parts.append(f"[{serialize_inlines_to_markdown(nested) or text}]({url})")
        elif ctype == "image":
            url = child.get("attrs", {}).get("url", "")
            alt = serialize_inlines_to_markdown(nested) or text
            parts.append(f"![{alt}]({url})")
        elif ctype in ("linebreak", "softbreak"):
            parts.append("\n")
        elif ctype in ("inline_math", "inline_latex"):
            parts.append(f"${text}$")
        elif ctype == "inline_html":
            parts.append(text)
        elif ctype in ("highlight", "mark"):
            parts.append(f"=={serialize_inlines_to_markdown(nested) or text}==")
        elif ctype == "superscript":
            parts.append(f"^{serialize_inlines_to_markdown(nested) or text}^")
        elif ctype == "subscript":
            parts.append(f"~{serialize_inlines_to_markdown(nested) or text}~")
        elif ctype in ("underline", "insert"):
            parts.append(f"<u>{serialize_inlines_to_markdown(nested) or text}</u>")
        elif ctype == "footnote_ref":
            parts.append(f"[^{text}]")
        elif ctype == "semantic_cross_reference":
            parts.append(f"@{child.get('target_id', '')}")
        elif ctype == "semantic_citation":
            parts.append(child.get("raw", ""))
        elif nested:
            parts.append(serialize_inlines_to_markdown(nested))
        else:
            parts.append(text)
    return "".join(parts)


# ═══════════════════════════════════════════════════════════════════════════
# Emoji splitting
# ═══════════════════════════════════════════════════════════════════════════


def _split_text_by_emoji(text: str) -> list[tuple[str, bool]]:
    """Split text into ``(segment, is_emoji)`` parts.

    Uses the ``emoji`` library if available; otherwise returns the
    whole text as a single non-emoji segment.
    """
    if not text:
        return []
    try:
        import emoji as _emoji_pkg
    except ImportError:
        return [(text, False)]
    import warnings

    with warnings.catch_warnings():
        warnings.simplefilter("ignore", DeprecationWarning)
        emoji_data = _emoji_pkg.emoji_list(text)
    if not emoji_data:
        return [(text, False)]
    parts: list[tuple[str, bool]] = []
    cursor = 0
    for em in emoji_data:
        start = em["match_start"]
        end = em["match_end"]
        if start > cursor:
            parts.append((text[cursor:start], False))
        parts.append((text[start:end], True))
        cursor = end
    if cursor < len(text):
        parts.append((text[cursor:], False))
    return parts


def _set_emoji_font(run) -> None:
    """Set Segoe UI Emoji font on all font slots of a run."""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.makeelement(qn("w:rFonts"), {})
    rFonts.set(qn("w:ascii"), "Segoe UI Emoji")
    rFonts.set(qn("w:hAnsi"), "Segoe UI Emoji")
    rFonts.set(qn("w:eastAsia"), "Segoe UI Emoji")
    rFonts.set(qn("w:cs"), "Segoe UI Emoji")
    rPr.insert(0, rFonts)
