"""Real DOCX → Markdown converter using python-docx.

This module implements the core conversion logic for ROUTE-DOC-001
(document → md).  It reads a .docx file, extracts its content
(paragraphs, headings, tables, images), and produces a Markdown
document in the staging directory.

The converter:
- Only writes to staging via ``WorkspaceHandle``.
- Checks cancellation before expensive operations.
- Reports progress through ``ProgressSink``.
- Returns a ``ConversionResult`` with ``ArtifactManifest`` entries.
"""

from __future__ import annotations

import os
import re
import time
import uuid
from pathlib import Path
from threading import RLock
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from docwen_core.protocols.execution_context import PluginExecutionContext

from docwen_core.docx_bookmarks import DocxBookmarkInventory, build_docx_bookmark_inventory
from docwen_core.docx_parsing.document_semantics import (
    extract_object_target,
    extract_semantic_caption,
    extract_semantic_table_metadata,
    render_semantic_reference_text,
)
from docwen_core.docx_parsing.format_features import (
    DocxMarkdownFormattingConfig,
    DocxMarkdownSyntaxConfig,
)
from docwen_core.docx_parsing.textbox_extraction import extract_textbox_paragraphs
from docwen_core.docx_resolved_numbering_recovery import ResolvedNumberingV4Recovery
from docwen_core.docx_semantics_v3 import DocxSemanticsV3Recovery
from docwen_core.export_semantics import MarkdownExportSemantics
from docwen_core.formula.constants import OMML_NS
from docwen_core.text.heading_numbering import (
    HeadingFormatter,
    NumberingSchemeResolutionError,
    resolve_heading_numbering_scheme,
)
from docwen_core.text.image_markdown import build_image_ocr_sidecar
from docwen_core.text.ocr import format_ocr_best_effort_warning
from docwen_plugin_document.to_markdown.formula_extractor import (
    extract_formula_from_element,
)
from docwen_plugin_document.to_markdown.request_policy import (
    DocxMarkdownRequestPolicy,
    build_docx_markdown_request_policy,
)

_LIST_CONTEXT_STYLE_NAMES = {"List Paragraph", "列表段落", "ListParagraph", "List Block", "列表块", "ListBlock"}
_MIN_EXTRA_LIST_INDENT_TWIPS = 300
_COMMONMARK_LIST_PREFIX_RE = re.compile(r"^(?:[*+-]|[0-9]{1,9}[.)]) {1,4}$")
_AUTHENTICATED_FENCED_CLOSER_RE = re.compile(r"^(?P<prefix>[ \t>+*0-9.)\[\]xX-]*?)(?:`{3,}|~{3,})[ \t]*$")


def _report_ocr_best_effort(progress: Any, status: object, *, location: str) -> None:
    """Report one safe, request-visible warning for a fallible OCR outcome."""
    message = format_ocr_best_effort_warning(status)
    if message is None:
        return
    progress.report_diagnostic(
        "warning",
        message,
        code="OCR-BEST-EFFORT",
        location=location,
    )


def _format_main_ocr_blockquote(ocr_text: str, title: str) -> str:
    """Format request-owned OCR presentation for main-Markdown placement."""
    lines: list[str] = []
    if title:
        lines.append(f"> {title}")
    lines.extend(f"> {line}" for line in ocr_text.splitlines())
    return "\n".join(lines)


def _authenticated_fenced_closing_prefix(authored: str) -> str:
    """Recover a present closer's exact container prefix from proven bytes."""

    physical_lines = authored.splitlines()
    if not physical_lines:
        raise ValueError("nested fenced anchor has no authenticated closing line")
    match = _AUTHENTICATED_FENCED_CLOSER_RE.fullmatch(physical_lines[-1])
    if match is None:
        raise ValueError("nested fenced anchor requires an authenticated present closer")
    return match.group("prefix")


def _nested_ordinary_anchor_marker_prefix(outer_groups: tuple[Any, ...]) -> str:
    """Return the canonical Markdown path for one nested post-block marker."""

    prefix = ""
    for group in reversed(outer_groups):
        kind = group.anchor.block_kind
        if kind in {"block_quote", "callout"}:
            prefix += "> "
        elif kind in {"list", "list_item"}:
            prefix += "  "
        else:
            raise ValueError("nested ordinary anchor has an unsupported source container")
    return prefix


class DocxToMarkdownConverter:
    """Convert a DOCX file to Markdown.

    Uses ``python-docx`` to parse the document structure and produces
    CommonMark-compatible output.  Images embedded in the DOCX are
    extracted to the staging directory.
    """

    # Configurable separators (L1+L2)
    _page_break_separator: str = "---"
    _section_break_separator: str = "***"
    _horizontal_rule_separator: str = "___"
    _preserve_formatting: bool = True
    _preserve_heading_formatting: bool = False
    _preserve_table_header_formatting: bool = False
    _unordered_list_marker_type: str = "dash"
    _list_indent_spaces: int = 4

    def __init__(self) -> None:
        """Create an isolated converter instance for one request."""
        self._request_policy = DocxMarkdownRequestPolicy(
            formatting=DocxMarkdownFormattingConfig(),
            syntax=DocxMarkdownSyntaxConfig(),
            style_detector=None,
            export=MarkdownExportSemantics(),
            ocr_blockquote_title="",
        )
        self._request_policy_resolved = False
        self._syntax_config = self._request_policy.syntax
        self._note_extractor: Any | None = None
        self._semantic_bookmark_inventory: DocxBookmarkInventory | None = None
        self._semantic_v3_recovery = DocxSemanticsV3Recovery()
        self._resolved_v4_recovery: ResolvedNumberingV4Recovery | None = None
        self._resolved_v4_diagnostics: list[tuple[str, str, str]] = []
        self._pending_artifacts: list[Any] = []
        self._pending_primary_path: str | None = None
        self._conversion_lock = RLock()

    def _syntax_for_rendering(self) -> DocxMarkdownSyntaxConfig:
        """Return the request syntax, or the deterministic constructor default."""
        return self._syntax_config

    def _apply_request_policy(self, policy: DocxMarkdownRequestPolicy) -> None:
        """Install one resolved policy before document parsing begins."""
        self._request_policy = policy
        self._request_policy_resolved = True
        self._syntax_config = policy.syntax
        self._page_break_separator = policy.export.page_break_separator
        self._section_break_separator = policy.export.section_break_separator
        self._horizontal_rule_separator = policy.export.horizontal_rule_separator
        self._preserve_formatting = policy.formatting.preserve_formatting
        self._preserve_heading_formatting = policy.formatting.preserve_heading_formatting
        self._preserve_table_header_formatting = policy.formatting.preserve_table_header_formatting
        self._unordered_list_marker_type = policy.syntax.unordered_list
        self._list_indent_spaces = policy.syntax.indent_spaces

    # ── Public entry point ──────────────────────────────────────────

    def convert(self, context: PluginExecutionContext) -> Any:
        """Serialize use of request-owned state on one converter instance."""
        with self._conversion_lock:
            self._discard_pending_artifacts()
            return self._convert_once(context)

    def _convert_once(self, context: PluginExecutionContext) -> Any:
        """Run the DOCX → Markdown conversion.

        Args:
            context: The plugin execution context providing workspace,
                     config, progress, cancellation, and logger.

        Returns:
            ``ConversionResult`` with staging artifacts.
        """
        self._resolved_v4_diagnostics.clear()
        from docwen_core.models.result import (
            ConversionDiagnostic,
            ConversionErrorInfo,
            ConversionMetrics,
            ConversionResult,
        )

        t_start = time.monotonic()
        task_id = context.request.request_id
        input_path = context.workspace.input_path
        # 1. Check cancellation before starting
        context.cancellation.check()

        # 2. Report start
        context.progress.report_progress(0.0, "Starting DOCX → Markdown conversion")
        context.logger.info(f"DOCX→MD: reading {input_path}")

        # 3. Read options
        options = context.request.options

        # Resolve the authoritative request snapshot once before parsing.
        policy = build_docx_markdown_request_policy(context, options)
        self._apply_request_policy(policy)

        # remove_numbering: strip heading numbering prefixes when True
        remove_numbering = options.get("remove_numbering", True)

        # add_numbering: prepend scheme-based numbering
        add_numbering = options.get("add_numbering", False)
        numbering_scheme = options.get("numbering_scheme", "")

        heading_formatter = None
        if add_numbering:
            try:
                scheme_config = resolve_heading_numbering_scheme(
                    numbering_scheme,
                    context.numbering_registry,
                )
            except NumberingSchemeResolutionError as exc:
                return ConversionResult(
                    task_id=task_id,
                    success=False,
                    error=ConversionErrorInfo(
                        error_type=exc.error_type,
                        message=str(exc),
                        diagnostic_code=exc.diagnostic_code,
                    ),
                    diagnostics=[
                        ConversionDiagnostic(
                            level="error",
                            message=str(exc),
                            code=exc.diagnostic_code,
                        )
                    ],
                    metrics=ConversionMetrics(
                        duration_ms=(time.monotonic() - t_start) * 1000.0,
                    ),
                )
            heading_formatter = HeadingFormatter(scheme_config)

        # 4. Parse the DOCX (standard path)
        context.cancellation.check()
        context.progress.report_progress(10.0, "Parsing DOCX structure...")

        try:
            from docx import Document as _StandardDocument

            _doc = _StandardDocument(input_path)
            metadata, skip_indices = self._extract_title_metadata(_doc, input_path)
            locale = str(options.get("locale", "en"))
            yaml_header = self._build_yaml_header(
                metadata,
                locale=locale,
                yaml_key_labels=options.get("yaml_key_labels"),
            )

            markdown_content, stats = self._parse_docx(
                input_path=input_path,
                context=context,
                remove_numbering=remove_numbering,
                heading_formatter=heading_formatter,
                skip_indices=skip_indices,
            )

            # Prepend YAML front matter to the markdown output
            if yaml_header:
                markdown_content = yaml_header + markdown_content
        except Exception as exc:
            self._discard_pending_artifacts()
            context.logger.error(f"DOCX→MD conversion failed: {exc}")
            from docwen_core.models.result import ConversionErrorInfo

            return ConversionResult(
                task_id=task_id,
                success=False,
                error=ConversionErrorInfo(
                    error_type="conversion_failed",
                    message=str(exc),
                    diagnostic_code="DOCX2MD-PARSE-ERROR",
                ),
                diagnostics=[
                    ConversionDiagnostic(
                        level="error",
                        message=f"Failed to parse DOCX: {exc}",
                        code="DOCX2MD-PARSE-ERROR",
                    ),
                ],
            )

        # 5. Write the markdown to staging
        context.cancellation.check()
        context.progress.report_progress(80.0, "Writing Markdown to staging...")

        output_path = context.workspace.create_artifact_path("primary", ".md")
        self._pending_primary_path = output_path
        try:
            with open(output_path, "w", encoding="utf-8", newline="") as f:
                f.write(markdown_content)
        except OSError as exc:
            Path(output_path).unlink(missing_ok=True)
            self._discard_pending_artifacts()
            context.logger.error(f"DOCX→MD write failed: {exc}")
            return ConversionResult(
                task_id=task_id,
                success=False,
                error=ConversionErrorInfo(
                    error_type="conversion_failed",
                    message=f"Failed to write output file: {exc}",
                    diagnostic_code="DOCX2MD-WRITE-ERROR",
                ),
                diagnostics=[
                    ConversionDiagnostic(
                        level="error",
                        message=f"File write error at {output_path}: {exc}",
                        code="DOCX2MD-WRITE-ERROR",
                    ),
                ],
            )

        # 6. Build artifact manifest
        input_basename = os.path.basename(input_path)
        suggested_name = input_basename.rsplit(".", 1)[0] + ".md"

        from docwen_core.models.artifact import ArtifactManifest

        artifact = ArtifactManifest(
            artifact_id=str(uuid.uuid4()),
            kind="primary",
            staging_path=output_path,
            suggested_name=suggested_name,
            media_type="text/markdown",
            metadata={
                "paragraph_count": stats.get("paragraphs", 0),
                "heading_count": stats.get("headings", 0),
                "table_count": stats.get("tables", 0),
                "image_count": stats.get("images", 0),
            },
            is_primary=True,
        )
        note_definition_loss_count = stats.get("note_definition_loss_count", 0)
        if note_definition_loss_count:
            artifact.metadata["note_definition_loss_count"] = note_definition_loss_count
        image_owner_resource_omitted_count = stats.get("image_owner_resource_omitted_count", 0)
        if image_owner_resource_omitted_count:
            artifact.metadata["image_owner_resource_omitted_count"] = image_owner_resource_omitted_count
        preexisting_artifacts = list(getattr(context.workspace, "registered_artifacts", []))
        pending_artifacts = list(self._pending_artifacts)

        # 7. Report completion
        duration_ms = (time.monotonic() - t_start) * 1000.0
        all_diagnostics = [
            ConversionDiagnostic(
                level="info",
                message=(
                    f"Converted DOCX to Markdown: {stats['paragraphs']} paragraphs, "
                    f"{stats['headings']} headings, {stats['tables']} tables"
                ),
                code="DOCX2MD-OK",
            ),
        ]
        if note_definition_loss_count:
            footnote_loss_count = stats.get("footnote_definition_loss_count", 0)
            endnote_loss_count = stats.get("endnote_definition_loss_count", 0)
            all_diagnostics.append(
                ConversionDiagnostic(
                    level="warning",
                    message=(
                        "Some referenced note definitions could not be read from the DOCX "
                        f"(footnotes={footnote_loss_count}, endnotes={endnote_loss_count}). "
                        "The Markdown body was preserved, but those definitions are missing."
                    ),
                    code="DOCX2MD-NOTE-DEFINITION-LOSS",
                    location="word/footnotes.xml or word/endnotes.xml",
                )
            )
        all_diagnostics.extend(
            ConversionDiagnostic(
                level="warning",
                message=message,
                code=code,
                location=location,
            )
            for code, message, location in self._resolved_v4_diagnostics
        )
        if image_owner_resource_omitted_count:
            all_diagnostics.append(
                ConversionDiagnostic(
                    level="warning",
                    message=(
                        f"Recovered {image_owner_resource_omitted_count} authenticated image owner(s) "
                        "with the exact resource-less carrier because preserve_resources is false."
                    ),
                    code="DOCX2MD-IMAGE-OWNER-RESOURCE-OMITTED",
                    artifact_id=artifact.artifact_id,
                )
            )

        result = ConversionResult(
            task_id=task_id,
            success=True,
            artifacts=[artifact, *preexisting_artifacts, *pending_artifacts],
            diagnostics=all_diagnostics,
            error=None,
            metrics=ConversionMetrics(
                duration_ms=duration_ms,
                input_bytes=os.path.getsize(input_path) if os.path.isfile(input_path) else 0,
                output_bytes=len(markdown_content.encode("utf-8")),
                extra=stats,
            ),
        )
        try:
            context.progress.report_artifact_ready(artifact.artifact_id, suggested_name)
            context.progress.report_progress(100.0, "Conversion complete")
            context.logger.info(
                f"DOCX→MD complete: {stats['paragraphs']} paragraphs, "
                f"{stats['headings']} headings, {stats['tables']} tables, "
                f"{stats['images']} images"
            )
        except Exception:
            Path(output_path).unlink(missing_ok=True)
            self._discard_pending_artifacts()
            raise
        context.workspace.add_artifact(artifact)
        for pending in pending_artifacts:
            context.workspace.add_artifact(pending)
        self._pending_artifacts.clear()
        self._pending_primary_path = None
        return result

    # ── Internal parsing ────────────────────────────────────────────

    def _parse_docx(
        self,
        input_path: str,
        context: PluginExecutionContext,
        remove_numbering: bool = False,
        heading_formatter: Any = None,
        skip_indices: set | None = None,
    ) -> tuple[str, dict[str, int]]:
        """Parse a DOCX file and return (markdown_text, stats_dict).

        Iterates over the document body in document order, processing
        paragraphs (including those inside SDT containers) and tables.
        Uses python-docx Paragraph objects for correct text extraction.

        Images embedded in the DOCX are extracted to request-owned staging
        when recognition or preservation needs them. Only
        ``to_md_keep_images=True`` exports those bytes as image artifacts.

        Args:
            skip_indices: Optional set of paragraph indices to skip in the
                body — used to avoid duplicating Title/Subtitle paragraphs
                that were already consumed as YAML front matter metadata.
        """
        if not self._request_policy_resolved:
            self._apply_request_policy(build_docx_markdown_request_policy(context, context.request.options))

        skip = skip_indices or set()

        from docx import Document

        doc = Document(input_path)
        self._resolved_v4_recovery = ResolvedNumberingV4Recovery.load_if_present(input_path, doc)
        if self._resolved_v4_recovery is None:
            self._semantic_v3_recovery = DocxSemanticsV3Recovery.load(input_path, doc)
        else:
            self._semantic_v3_recovery = self._resolved_v4_recovery
            diagnostic = self._resolved_v4_recovery.source_recovery_diagnostic
            self._resolved_v4_diagnostics.append((diagnostic.code, diagnostic.message, "DOCX package"))
            context.progress.report_diagnostic(
                "warning",
                diagnostic.message,
                code=diagnostic.code,
                location="DOCX package",
            )
        lines: list[str] = []
        exact_fenced_fragments: dict[str, str] = {}

        def _append_exact_fenced_fragment(authored: str) -> None:
            token = f"\0DOCWEN-FENCED-SOURCE-{len(exact_fenced_fragments):08d}\0"
            if token in authored:
                raise ValueError("fenced source collides with the internal recovery token")
            exact_fenced_fragments[token] = authored
            lines.extend((token, ""))

        stats: dict[str, int] = {
            "paragraphs": 0,
            "headings": 0,
            "tables": 0,
            "images": 0,
            "image_owner_resource_omitted_count": 0,
        }

        # Extract textbox content (C1) BEFORE body iteration
        textbox_paragraphs = extract_textbox_paragraphs(doc)
        body_textboxes_by_anchor: dict[int, list[Any]] = {}
        trailing_textbox_paragraphs: list[Any] = []
        for textbox_paragraph in textbox_paragraphs:
            if textbox_paragraph.source == "textbox" and textbox_paragraph.anchor_index is not None:
                body_textboxes_by_anchor.setdefault(textbox_paragraph.anchor_index, []).append(textbox_paragraph)
            else:
                trailing_textbox_paragraphs.append(textbox_paragraph)

        def _append_textboxes(items: list[Any]) -> None:
            for textbox_paragraph in items:
                textbox_text = self._render_extracted_paragraph(textbox_paragraph)
                if textbox_text:
                    lines.append(textbox_text)
                    lines.append("")
                    stats["paragraphs"] += 1

        # Note extractor for inline references and definitions block
        from docwen_plugin_document.shared.note_extraction import NoteExtractor

        self._note_extractor = NoteExtractor(doc, input_path)

        # List detection and numbering infrastructure
        from docwen_plugin_document.shared.list_processing import ListCounterManager
        from docwen_plugin_document.shared.numbering_index import NumberingIndex

        numbering_index = NumberingIndex(doc)
        list_counter = ListCounterManager()

        # Border group tracker for paragraph-border → horizontal-rule
        from docwen_core.docx_parsing.break_utils import BorderGroupTracker

        border_tracker = BorderGroupTracker(separator=self._horizontal_rule_separator)

        # Code-block accumulator for progressive code-block detection
        from docwen_core.docx_parsing.format_features import CodeBlockAccumulator

        code_block_acc = CodeBlockAccumulator(indent_spaces=self._list_indent_spaces)

        options = context.request.options
        preserve_resources = bool(options.get("to_md_keep_images", True))
        recognize_text = bool(options.get("to_md_enable_ocr", False))
        process_images = preserve_resources or recognize_text
        heading_cleanup_rules = getattr(context, "heading_cleanup_rules", ()) or ()

        style_detector_config = self._request_policy.style_detector
        export_modes = self._request_policy.resolve_export_modes()

        # OCR placement and per-document image counter for stable sidecar naming
        ocr_placement: str = export_modes["ocr_placement_mode"]
        table_merge_strategy: str = export_modes["table_merge_export_strategy"]
        ocr_language = str(options.get("ocr_language") or "auto")
        current_locale = str(options.get("locale") or "zh_CN")
        _img_seq = 0
        _main_stem = os.path.splitext(os.path.basename(str(input_path)))[0]

        # Build a lookup from XML paragraph element to Paragraph object
        # so we can use python-docx's high-level API for text extraction
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph

        para_by_element: dict[int, Any] = {}
        for p_element in doc.element.body.iter(qn("w:p")):  # pyright: ignore[reportAttributeAccessIssue]
            # ``Document.paragraphs`` excludes table-cell and SDT-contained
            # paragraphs. Wrapping every body paragraph against the document
            # part keeps relationships/styles available to the shared run
            # renderer without introducing a second extraction path.
            para_by_element[id(p_element)] = Paragraph(p_element, doc)

        body_elements = self._semantic_v3_recovery.logical_body_elements(doc)
        if preserve_resources and export_modes["image_extraction_mode"] == "omit":
            for element in body_elements:
                caption = self._semantic_v3_recovery.caption_for_object(element)
                anchor_groups = self._semantic_v3_recovery.ordinary_anchor_groups(element)
                if (caption is not None and caption.kind == "figure") or any(
                    group.anchor.block_kind == "image" for group in anchor_groups
                ):
                    raise ValueError("image_mode=omit cannot preserve an authenticated DOCX image owner")
        semantic_bookmark_inventory = build_docx_bookmark_inventory(doc)
        self._semantic_bookmark_inventory = semantic_bookmark_inventory
        semantic_caption_by_object_index: dict[int, Any] = {}
        semantic_table_captions: dict[str, tuple[int, Any]] = {}
        semantic_caption_indices: set[int] = set()
        for caption_index, element in enumerate(body_elements):
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
            if tag != "p":
                continue
            caption = extract_semantic_caption(
                element,
                bookmark_inventory=semantic_bookmark_inventory,
            )
            if caption is None:
                continue
            object_index = None
            if caption.kind == "table" and caption.target_id is not None:
                semantic_table_captions[caption.target_id] = (caption_index, caption)
            elif caption.kind == "figure" and caption_index > 0:
                previous = body_elements[caption_index - 1]
                previous_tag = previous.tag.split("}")[-1] if "}" in previous.tag else previous.tag
                if previous_tag == "p" and (
                    previous.find(f".//{qn('w:drawing')}") is not None
                    or previous.find(f".//{qn('w:pict')}") is not None
                ):
                    object_index = caption_index - 1
            if object_index is not None:
                semantic_caption_by_object_index[object_index] = caption
                semantic_caption_indices.add(caption_index)

        for object_index, element in enumerate(body_elements):
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
            if tag != "tbl":
                continue
            target_id = extract_object_target(
                element,
                bookmark_inventory=semantic_bookmark_inventory,
            )
            if target_id is None or target_id not in semantic_table_captions:
                continue
            caption_index, caption = semantic_table_captions[target_id]
            semantic_caption_by_object_index[object_index] = caption
            semantic_caption_indices.add(caption_index)

        total_elements = len(body_elements) if body_elements else 1
        active_list_context_level: int | None = None

        for idx, child in enumerate(body_elements):
            # Check cancellation periodically
            if idx % 10 == 0:
                context.cancellation.check()
                progress = 10.0 + 60.0 * (idx / max(total_elements, 1))
                context.progress.report_progress(progress, f"Processing element {idx + 1}/{total_elements}")

            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if self._semantic_v3_recovery.is_caption_element(child):
                _append_textboxes(body_textboxes_by_anchor.get(idx, []))
                continue

            v3_caption = self._semantic_v3_recovery.caption_for_object(child)
            if v3_caption is not None:
                keyword = {
                    "figure": "Figure",
                    "table": "Table",
                    "equation": "Equation",
                    "code_block": "Code",
                }[v3_caption.kind]
                declaration = f"{keyword}:"
                if v3_caption.title:
                    declaration += f" {v3_caption.title}"
                if v3_caption.source_id is not None:
                    declaration += f" ^{v3_caption.source_id}"
                lines.extend((declaration, ""))

            if tag == "p":
                if idx in semantic_caption_indices:
                    _append_textboxes(body_textboxes_by_anchor.get(idx, []))
                    continue
                # Skip paragraphs that were consumed as Title/Subtitle metadata
                if idx in skip:
                    _append_textboxes(body_textboxes_by_anchor.get(idx, []))
                    continue
                para = para_by_element.get(id(child))
                list_item_level = self._paragraph_list_item_level(para, numbering_index)
                continuation_list_level = (
                    active_list_context_level
                    if (
                        list_item_level is None
                        and active_list_context_level is not None
                        and self._is_list_continuation_paragraph(para)
                    )
                    else None
                )
                semantic_caption = semantic_caption_by_object_index.get(idx)
                if semantic_caption is not None:
                    declaration = (
                        ":" if semantic_caption.source_form == "shorthand" else f"{semantic_caption.kind.title()}:"
                    )
                    lines.append(f"{declaration} {semantic_caption.content} {{#{semantic_caption.target_id}}}")
                    lines.append("")
                exact_fenced_source = self._semantic_v3_recovery.render_fenced_source(child)
                if exact_fenced_source is not None:
                    if code_block_acc.in_code_block:
                        code_markdown = code_block_acc.finalize()
                        if code_markdown:
                            lines.extend((code_markdown, ""))
                    _append_exact_fenced_fragment(exact_fenced_source)
                    anchor_groups = self._semantic_v3_recovery.ordinary_anchor_groups(child)
                    nested_prefix: str | None = None
                    for group_position, group in enumerate(anchor_groups):
                        if (
                            group.anchor.block_kind not in {"paragraph", "image", "list_item"}
                            and group.index == len(group.elements) - 1
                        ):
                            if group_position < len(anchor_groups) - 1:
                                if nested_prefix is None:
                                    nested_prefix = _authenticated_fenced_closing_prefix(exact_fenced_source)
                                if not nested_prefix:
                                    raise ValueError("nested fenced anchor has no authenticated container prefix")
                                blank_container = nested_prefix.rstrip(" \t")
                                if blank_container:
                                    if not lines or lines[-1] != "":
                                        raise ValueError("nested fenced anchor lost its canonical separator slot")
                                    lines[-1] = blank_container
                                lines.extend((f"{nested_prefix}^{group.anchor.source_id}", ""))
                            else:
                                lines.extend((f"^{group.anchor.source_id}", ""))
                    stats["paragraphs"] += 1
                    if list_item_level is not None:
                        active_list_context_level = list_item_level
                    elif continuation_list_level is None:
                        active_list_context_level = None
                    _append_textboxes(body_textboxes_by_anchor.get(idx, []))
                    continue
                para_lines, para_stats = self._process_paragraph(
                    child,
                    para_by_element,
                    remove_numbering=remove_numbering,
                    heading_cleanup_rules=heading_cleanup_rules,
                    heading_formatter=heading_formatter,
                    numbering_index=numbering_index,
                    list_counter=list_counter,
                    border_tracker=border_tracker,
                    code_block_acc=code_block_acc,
                    style_detector_config=style_detector_config,
                    continuation_list_level=continuation_list_level,
                    diagnostic_sink=context.progress,
                    diagnostic_location=f"document.xml body element {idx}",
                )
                anchor_groups = self._semantic_v3_recovery.ordinary_anchor_groups(child)
                anchor_group = anchor_groups[0] if anchor_groups else None
                preserve_image_owner = (v3_caption is not None and v3_caption.kind == "figure") or (
                    anchor_group is not None and anchor_group.anchor.block_kind == "image"
                )
                img_refs: list[str] = []
                # Recognition and resource preservation are independent. OCR
                # may use request-owned temporary image bytes without exporting
                # those bytes as Bundle resources.
                if process_images or preserve_image_owner:
                    img_refs, infos = self._process_images_for_output(
                        child,
                        doc,
                        context,
                        prefer_source_names=(
                            semantic_caption is not None or (v3_caption is not None and v3_caption.kind == "figure")
                        ),
                        preserve_resources=preserve_resources,
                        preserve_owner=preserve_image_owner,
                    )
                    if preserve_image_owner and len(infos) != 1:
                        raise ValueError("an authenticated image owner must recover exactly one image")
                    img_count = len(infos) if preserve_resources else 0
                    if preserve_image_owner and not preserve_resources:
                        stats["image_owner_resource_omitted_count"] += len(infos)
                    try:
                        # ── OCR per image ─────────────────────────────
                        if recognize_text and infos:
                            img_refs, _img_seq = self._ocr_per_image(
                                img_refs,
                                infos,
                                _main_stem,
                                _img_seq,
                                ocr_placement,
                                context,
                                ocr_language=ocr_language,
                                current_locale=current_locale,
                                retain_image_owner=preserve_image_owner,
                            )
                    finally:
                        if not preserve_resources:
                            self._discard_unpreserved_images(infos)
                    # Merge exactly once after OCR has either preserved or
                    # replaced each request-owned image reference.
                    para_lines = self._merge_img_refs_into_lines(para_lines, img_refs)
                    stats["images"] += img_count
                if anchor_group is not None and anchor_group.anchor.block_kind == "image":
                    if not img_refs:
                        raise ValueError(
                            "an authenticated image anchor cannot be recovered when its image resource is omitted"
                        )
                    para_lines = self._append_anchor_to_first_image_reference(
                        para_lines,
                        anchor_group.anchor.source_id,
                    )
                lines.extend(para_lines)
                for group_position, group in enumerate(anchor_groups):
                    if (
                        group.anchor.block_kind not in {"paragraph", "image", "list_item"}
                        and group.index == len(group.elements) - 1
                    ):
                        if code_block_acc.in_code_block:
                            code_markdown = code_block_acc.finalize()
                            if code_markdown:
                                lines.extend((code_markdown, ""))
                        marker_prefix = _nested_ordinary_anchor_marker_prefix(
                            anchor_groups[group_position + 1 :],
                        )
                        blank_container = marker_prefix.rstrip()
                        if ">" in blank_container:
                            lines.append(blank_container)
                        lines.extend((f"{marker_prefix}^{group.anchor.source_id}", ""))
                stats["paragraphs"] += para_stats.get("paragraphs", 0)
                stats["headings"] += para_stats.get("headings", 0)
                if list_item_level is not None:
                    active_list_context_level = list_item_level
                elif continuation_list_level is None:
                    active_list_context_level = None
            elif tag == "tbl":
                semantic_caption = semantic_caption_by_object_index.get(idx)
                if semantic_caption is not None:
                    declaration = (
                        ":" if semantic_caption.source_form == "shorthand" else f"{semantic_caption.kind.title()}:"
                    )
                    lines.append(f"{declaration} {semantic_caption.content} {{#{semantic_caption.target_id}}}")
                    lines.append("")
                effective_table_merge_strategy = "marker" if semantic_caption is not None else table_merge_strategy
                table_list_context_level = (
                    active_list_context_level
                    if active_list_context_level is not None and self._table_has_extra_list_indent(child)
                    else None
                )
                if process_images:
                    tbl_lines, tbl_images, _img_seq = self._process_table_for_output(
                        child,
                        para_by_element=para_by_element,
                        table_merge_strategy=effective_table_merge_strategy,
                        doc=doc,
                        context=context,
                        enable_ocr=recognize_text,
                        preserve_resources=preserve_resources,
                        main_stem=_main_stem,
                        img_seq=_img_seq,
                        ocr_placement=ocr_placement,
                        ocr_language=ocr_language,
                        current_locale=current_locale,
                    )
                else:
                    tbl_lines, tbl_images = self._process_table(
                        child,
                        para_by_element,
                        table_merge_strategy=effective_table_merge_strategy,
                    )
                if semantic_caption is not None:
                    metadata = extract_semantic_table_metadata(child)
                    attributes = [
                        f"header-rows={metadata.header_rows}",
                        f"header-cols={metadata.header_columns}",
                    ]
                    if metadata.repeat_header == "always":
                        attributes.append("repeat-header=true")
                    elif metadata.repeat_header == "never":
                        attributes.append("repeat-header=false")
                    attribute_line = "{" + " ".join(attributes) + "}"
                    if tbl_lines and tbl_lines[-1] == "":
                        tbl_lines.insert(-1, attribute_line)
                    else:
                        tbl_lines.extend([attribute_line, ""])
                if table_list_context_level is not None:
                    table_indent = " " * self._list_indent_spaces * (table_list_context_level + 1)
                    tbl_lines = [f"{table_indent}{line}" if line else line for line in tbl_lines]
                anchor_groups = self._semantic_v3_recovery.ordinary_anchor_groups(child)
                for group_position, group in enumerate(anchor_groups):
                    if group.index == len(group.elements) - 1:
                        marker_prefix = _nested_ordinary_anchor_marker_prefix(
                            anchor_groups[group_position + 1 :],
                        )
                        blank_container = marker_prefix.rstrip()
                        if ">" in blank_container:
                            tbl_lines.append(blank_container)
                        tbl_lines.extend(("", f"{marker_prefix}^{group.anchor.source_id}", ""))
                lines.extend(tbl_lines)
                stats["tables"] += 1
                stats["images"] += tbl_images
                if table_list_context_level is None:
                    active_list_context_level = None
            elif tag == "sdt":
                active_list_context_level = None
                # Structured document tag — process inner paragraphs
                sdt_lines, sdt_stats = self._process_sdt(
                    child,
                    para_by_element,
                    remove_numbering=remove_numbering,
                    heading_cleanup_rules=heading_cleanup_rules,
                    numbering_index=numbering_index,
                    list_counter=list_counter,
                    border_tracker=border_tracker,
                    code_block_acc=code_block_acc,
                    style_detector_config=style_detector_config,
                    heading_formatter=heading_formatter,
                    table_merge_strategy=table_merge_strategy,
                    diagnostic_sink=context.progress,
                )
                if process_images:
                    img_refs, infos = self._process_images_for_output(
                        child,
                        doc,
                        context,
                        preserve_resources=preserve_resources,
                    )
                    img_count = len(infos) if preserve_resources else 0
                    try:
                        # ── OCR per image ─────────────────────────────
                        if recognize_text and infos:
                            img_refs, _img_seq = self._ocr_per_image(
                                img_refs,
                                infos,
                                _main_stem,
                                _img_seq,
                                ocr_placement,
                                context,
                                ocr_language=ocr_language,
                                current_locale=current_locale,
                            )
                    finally:
                        if not preserve_resources:
                            self._discard_unpreserved_images(infos)
                    # Merge image refs into SDT output
                    sdt_lines = self._merge_img_refs_into_lines(sdt_lines, img_refs)
                    stats["images"] += img_count
                lines.extend(sdt_lines)
                stats["paragraphs"] += sdt_stats.get("paragraphs", 0)
                stats["headings"] += sdt_stats.get("headings", 0)
                stats["tables"] += sdt_stats.get("tables", 0)
                stats["images"] += sdt_stats.get("images", 0)

            _append_textboxes(body_textboxes_by_anchor.get(idx, []))

        # Header/footer and any unanchored textbox content has no body position.
        _append_textboxes(trailing_textbox_paragraphs)

        # Finalize code block accumulator — emit closing fence if open
        cb_final = code_block_acc.finalize()
        if cb_final:
            lines.append(cb_final)
            lines.append("")

        # Finalize border tracker — emit closing separator if still in group
        final_sep = border_tracker.finalize()
        if final_sep:
            lines.append(final_sep)
            lines.append("")

        # Normalize: smart content joining (M3)
        #  - Consecutive list items → single newline (no blank)
        #  - Consecutive blockquotes → single newline (no blank)
        #  - Different block types → blank line separator
        import re as _re

        def _detect_block_type(ln: str) -> str | None:
            """Return ``"list"``, ``"quote"``, ``"other"``, or ``None`` (blank)."""
            stripped = ln.strip()
            if not stripped:
                return None
            authored_fence = exact_fenced_fragments.get(ln)
            if authored_fence is not None:
                first_line = authored_fence.splitlines()[0].strip()
                if first_line.startswith(">"):
                    return "quote"
                if _re.match(r"^(?:[-*+] |\d+[.)] )", first_line):
                    return "list"
                return "other"
            if stripped.startswith(">"):
                return "quote"
            if _re.match(r"^(\s*[-*+] |\s*\d+[.)] )", stripped):
                return "list"
            return "other"

        normalized: list[str] = []
        prev_blank = False
        for line in lines:
            is_blank = line.strip() == ""

            if is_blank:
                prev_blank = True
                continue  # defer blank — decide when next non-blank arrives

            # Non-blank line
            block_type = _detect_block_type(line)

            if prev_blank and normalized:
                # Get the type of the last emitted non-blank line
                last_type = _detect_block_type(normalized[-1])
                if last_type != block_type or block_type not in ("list", "quote"):
                    normalized.append("")  # different types → blank separator

            normalized.append(line)
            prev_blank = False

        # Append note definitions block if notes exist
        definitions = self._note_extractor.build_definitions_block()
        if definitions:
            normalized.append("")
            normalized.extend(definitions.splitlines())

        note_losses = self._note_extractor.definition_loss_counts()
        note_definition_loss_count = sum(note_losses.values())
        if note_definition_loss_count:
            stats["footnote_definition_loss_count"] = note_losses["footnotes"]
            stats["endnote_definition_loss_count"] = note_losses["endnotes"]
            stats["note_definition_loss_count"] = note_definition_loss_count

        markdown = "\n".join(normalized).strip() + "\n"
        for token, authored in exact_fenced_fragments.items():
            needle = token + "\n"
            if markdown.count(needle) != 1:
                raise ValueError("fenced source recovery token lost its unique output position")
            markdown = markdown.replace(needle, authored, 1)
        if any(token in markdown for token in exact_fenced_fragments):
            raise ValueError("fenced source recovery left an internal token in output")
        return markdown, stats

    @staticmethod
    def _format_yaml_value(value: Any) -> str:
        """Format a value as a YAML-safe string.

        Wraps values in single quotes when they contain characters that
        YAML parsers would misinterpret (leading ``[``, ``{``, ``#``,
        ``&``, ``*``, ``!``, ``|``, ``>``, ``%``, ``@``, ````` ``,
        ``'``, ``"``, ``:``, or trailing ``#`` with preceding space).
        Handles None/empty gracefully.

        Used by ``_build_yaml_header`` for safe YAML front matter.
        """
        if value is None:
            return ""
        s = str(value).replace("\n", " ").replace("\r", "")
        s = s.strip()
        if not s:
            return ""

        special_lead = {"[", "{", "#", "&", "*", "!", "|", ">", "%", "@", "`"}
        needs_quote = (
            s[0] in special_lead
            or (len(s) > 1 and s[0] == "-" and (s[1].isspace() or s[1] == " "))
            or "'" in s
            or '"' in s
            or ": " in s
            or s.endswith(":")
            or " #" in s
            or s.lower() in ("true", "false", "null", "yes", "no", "on", "off", "~")
            or s.isdigit()
            or (s.startswith("0") and len(s) > 1 and s[1].isdigit())
        )

        if needs_quote:
            if '"' not in s:
                return f'"{s}"'
            escaped = s.replace("'", "''")
            return f"'{escaped}'"
        return s

    @staticmethod
    def _build_yaml_header(
        metadata: dict,
        locale: str = "en",
        yaml_key_labels: object | None = None,
    ) -> str:
        """Build a YAML front matter string from metadata dict.

        Produces a complete YAML front matter block (``---`` delimiters,
        ``aliases:`` list, ``title:`` / ``subtitle:`` fields, closing
        ``---``) suitable for prepending to Markdown output.

        Args:
            metadata: Dict with keys ``aliases`` (list), ``title`` (str),
                ``subtitle`` (str).
            locale: Locale string for localized YAML key names.
            yaml_key_labels: Optional pre-resolved labels from the application
                edge, e.g. ``{"title": "Titel", "subtitle": "Untertitel"}``.

        Returns:
            A YAML front matter string ending with ``\\n``, or an empty
            string if the metadata is empty/trivial.
        """
        aliases = metadata.get("aliases") or []
        title = metadata.get("title") or ""
        subtitle = metadata.get("subtitle") or ""

        parts: list[str] = ["---"]

        # L3: Locale-aware YAML key names.  Plugins consume resolved labels
        # instead of loading runtime resources themselves.
        if locale in ("zh_CN", "zh-CN", "zh"):
            key_aliases = "别名"
            key_title = "标题"
            key_subtitle = "副标题"
        else:
            key_aliases = "aliases"
            key_title = "title"
            key_subtitle = "subtitle"
        if isinstance(yaml_key_labels, dict):
            label_title = yaml_key_labels.get("title")
            label_subtitle = yaml_key_labels.get("subtitle")
            label_aliases = yaml_key_labels.get("aliases")
            if isinstance(label_title, str) and label_title.strip():
                key_title = label_title.strip()
            if isinstance(label_subtitle, str) and label_subtitle.strip():
                key_subtitle = label_subtitle.strip()
            if isinstance(label_aliases, str) and label_aliases.strip():
                key_aliases = label_aliases.strip()

        if aliases:
            parts.append(f"{key_aliases}:")
            for alias in aliases:
                safe = DocxToMarkdownConverter._format_yaml_value(alias)
                parts.append(f"  - {safe}")
        else:
            parts.append(f"{key_aliases}: []")

        safe_title = DocxToMarkdownConverter._format_yaml_value(title)
        parts.append(f"{key_title}: {safe_title}" if safe_title else f"{key_title}: ")

        safe_subtitle = DocxToMarkdownConverter._format_yaml_value(subtitle)
        parts.append(f"{key_subtitle}: {safe_subtitle}" if safe_subtitle else f"{key_subtitle}: ")

        parts.append("---")
        parts.append("")
        return "\n".join(parts) + "\n"

    @staticmethod
    def _extract_title_metadata(
        doc: Any,
        docx_path: str,
    ) -> tuple[dict, set]:
        """Extract title/subtitle metadata from a python-docx Document.

        Scans paragraphs for Word "Title" and "Subtitle" styles at the
        start of the document.  Collects consecutive matches, merges
        multi-paragraph titles, and falls back to the filename stem
        when no Title style is found.

        Args:
            doc: A python-docx ``Document`` object.
            docx_path: Absolute path to the DOCX file (for filename
                fallback).

        Returns:
            A tuple of ``(metadata, skip_indices)`` where ``metadata``
            is a dict with keys ``aliases``, ``title``, ``subtitle``
            and ``skip_indices`` is a set of paragraph indices to skip
            in the body rendering.
        """
        import os as _os

        metadata: dict = {"aliases": [], "title": "", "subtitle": ""}
        title_indices: list[int] = []
        subtitle_indices: list[int] = []
        started_body = False  # True once we have seen a non-title/subtitle paragraph

        for i, para in enumerate(doc.paragraphs):
            style_name = para.style.name if para.style else ""

            if started_body:
                break  # Title/Subtitle only meaningful at document start

            if style_name == "Title":
                title_indices.append(i)
            elif style_name == "Subtitle":
                subtitle_indices.append(i)
            else:
                # First non-title, non-subtitle paragraph starts body content
                started_body = True
                # Do NOT break yet — we still need to process the accumulated
                # title/subtitle indices, which is done below the loop.

        # Merge title paragraphs
        if title_indices:
            title = "".join(doc.paragraphs[i].text.replace("\n", "").replace("\r", "").strip() for i in title_indices)
            metadata["title"] = title
            metadata["aliases"].append(title)
        else:
            filename = _os.path.basename(docx_path)
            stem = filename.rsplit(".", 1)[0]
            metadata["title"] = stem
            metadata["aliases"].append(stem)

        # Merge subtitle paragraphs
        if subtitle_indices:
            subtitle = "".join(
                doc.paragraphs[i].text.replace("\n", "").replace("\r", "").strip() for i in subtitle_indices
            )
            metadata["subtitle"] = subtitle

        skip_indices: set = set(title_indices) | set(subtitle_indices)
        return metadata, skip_indices

    def _process_sdt(
        self,
        sdt_element: Any,
        para_by_element: dict[int, Any],
        numbering_index: Any = None,
        list_counter: Any = None,
        border_tracker: Any = None,
        code_block_acc: Any = None,
        style_detector_config: Any = None,
        heading_formatter: Any = None,
        table_merge_strategy: str = "fill",
        remove_numbering: bool = False,
        heading_cleanup_rules: Any = (),
        diagnostic_sink: Any = None,
    ) -> tuple[list[str], dict[str, int]]:
        """Process a structured document tag (SDT) by recursing into its content.

        SDT elements wrap paragraphs, tables, and other content with
        metadata.  We extract the inner paragraphs and process them.
        """

        lines: list[str] = []
        stats: dict[str, int] = {"paragraphs": 0, "headings": 0, "tables": 0, "images": 0}
        active_list_context_level: int | None = None

        for inner_index, inner in enumerate(sdt_element):
            tag = inner.tag.split("}")[-1] if "}" in inner.tag else inner.tag

            if tag == "p":
                para = para_by_element.get(id(inner))
                list_item_level = self._paragraph_list_item_level(para, numbering_index)
                continuation_list_level = (
                    active_list_context_level
                    if (
                        list_item_level is None
                        and active_list_context_level is not None
                        and self._is_list_continuation_paragraph(para)
                    )
                    else None
                )
                para_lines, para_stats = self._process_paragraph(
                    inner,
                    para_by_element,
                    remove_numbering=remove_numbering,
                    heading_cleanup_rules=heading_cleanup_rules,
                    numbering_index=numbering_index,
                    list_counter=list_counter,
                    border_tracker=border_tracker,
                    code_block_acc=code_block_acc,
                    style_detector_config=style_detector_config,
                    heading_formatter=heading_formatter,
                    continuation_list_level=continuation_list_level,
                    diagnostic_sink=diagnostic_sink,
                    diagnostic_location=f"structured document paragraph {inner_index}",
                )
                lines.extend(para_lines)
                stats["paragraphs"] += para_stats.get("paragraphs", 0)
                stats["headings"] += para_stats.get("headings", 0)
                if list_item_level is not None:
                    active_list_context_level = list_item_level
                elif continuation_list_level is None:
                    active_list_context_level = None
            elif tag == "tbl":
                table_list_context_level = (
                    active_list_context_level
                    if active_list_context_level is not None and self._table_has_extra_list_indent(inner)
                    else None
                )
                # Tables inside SDT — process inline instead of silently dropping
                tbl_lines, tbl_images = self._process_table(
                    inner,
                    para_by_element=para_by_element,
                    table_merge_strategy=table_merge_strategy,
                )
                if table_list_context_level is not None:
                    table_indent = " " * self._list_indent_spaces * (table_list_context_level + 1)
                    tbl_lines = [f"{table_indent}{line}" if line else line for line in tbl_lines]
                lines.extend(tbl_lines)
                stats["tables"] += 1
                stats["images"] += tbl_images
                if table_list_context_level is None:
                    active_list_context_level = None
            elif tag == "sdtContent" or tag == "sdt":
                active_list_context_level = None
                # Recurse into nested SDT
                nested_lines, nested_stats = self._process_sdt(
                    inner,
                    para_by_element,
                    remove_numbering=remove_numbering,
                    heading_cleanup_rules=heading_cleanup_rules,
                    numbering_index=numbering_index,
                    list_counter=list_counter,
                    border_tracker=border_tracker,
                    code_block_acc=code_block_acc,
                    style_detector_config=style_detector_config,
                    heading_formatter=heading_formatter,
                    table_merge_strategy=table_merge_strategy,
                    diagnostic_sink=diagnostic_sink,
                )
                lines.extend(nested_lines)
                stats["paragraphs"] += nested_stats.get("paragraphs", 0)
                stats["headings"] += nested_stats.get("headings", 0)
                stats["tables"] += nested_stats.get("tables", 0)
                stats["images"] += nested_stats.get("images", 0)

        return lines, stats

    # ── Paragraph processing ────────────────────────────────────────

    def _process_paragraph(
        self,
        para_element: Any,
        para_by_element: dict[int, Any],
        remove_numbering: bool = False,
        heading_formatter: Any = None,
        numbering_index: Any = None,
        list_counter: Any = None,
        border_tracker: Any = None,
        code_block_acc: Any = None,
        style_detector_config: Any = None,
        continuation_list_level: int | None = None,
        heading_cleanup_rules: Any = (),
        diagnostic_sink: Any = None,
        diagnostic_location: str = "",
    ) -> tuple[list[str], dict[str, int]]:
        """Process a single ``<w:p>`` element.

        Uses python-docx's ``Paragraph`` object when available for
        correct text extraction (handles field codes, SDT content,
        and complex run structures).  Integrates break detection,
        style-driven code/quote blocks, gray-shading fallback, and
        pStyle-based numbering resolution.

        Returns (list_of_markdown_lines, stats_dict).
        """

        lines: list[str] = []
        stats: dict[str, int] = {"paragraphs": 0, "headings": 0}

        # Get the Paragraph object for proper text extraction
        para = para_by_element.get(id(para_element))
        paragraph_text = (
            para.text.strip() if para is not None else self._extract_paragraph_text_raw(para_element).strip()
        )
        semantic_bookmark_inventory = self._semantic_bookmark_inventory
        semantic_reference_text = self._semantic_v3_recovery.render_paragraph_text(para_element)
        if semantic_reference_text is None and semantic_bookmark_inventory is not None:
            semantic_reference_text = render_semantic_reference_text(
                para_element,
                bookmark_inventory=semantic_bookmark_inventory,
            )
        if semantic_reference_text is not None:
            paragraph_text = semantic_reference_text.strip()
        has_visible_wrapper = False
        if para is not None:
            w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            has_visible_wrapper = any(
                para._p.find(f".//{{{w_ns}}}{tag}") is not None for tag in ("ins", "fldSimple", "hyperlink")
            )
        if para is not None and (not paragraph_text or has_visible_wrapper):
            # python-docx Paragraph.text can be empty when all visible text is
            # wrapped by w:ins, w:fldSimple, or another XML container.  It can
            # also be only partially populated when a direct run is followed
            # by one of those wrappers.  The shared renderer owns the accepted
            # revision/field policy, so prefer its complete visible projection
            # whenever such a wrapper is present.
            rendered_text = self._extract_paragraph_text_formatted(
                para,
                preserve_formatting=False,
                style_detector_config=style_detector_config,
            ).strip()
            if rendered_text:
                paragraph_text = rendered_text

        # ── 0. Page/section break detection ─────────────────────────
        from docwen_core.docx_parsing.break_utils import (
            detect_all_breaks,
            extract_paragraph_border_info,
        )

        all_breaks: list[tuple[str, str]] = []
        if para is not None:
            all_breaks = detect_all_breaks(para)
        has_page_break = any(b[0] == "page" for b in all_breaks)
        has_section_break = any(b[0] == "section" for b in all_breaks)

        # ── 1. Border group tracking ────────────────────────────────
        if para is not None and border_tracker is not None:
            border_info = extract_paragraph_border_info(para)
            for sep in border_tracker.process_paragraph(border_info):
                lines.append(sep)
                lines.append("")

        # ── 2. Get paragraph style name ─────────────────────────────
        style_name = ""
        if para is not None and para.style is not None:
            style_name = para.style.name or ""

        # ── 3. Style-driven code-block / quote detection ────────────
        from docwen_core.docx_parsing.format_features import (
            StyleDetectorConfig,
            detect_full_paragraph_run_style,
            detect_paragraph_style_type,
            has_paragraph_gray_shading,
        )

        style_type, style_value = (
            detect_paragraph_style_type(para, config=style_detector_config) if para is not None else (None, None)
        )
        paragraph_list_context_level = self._paragraph_list_context_level(para, numbering_index)
        if paragraph_list_context_level == 0 and continuation_list_level is not None:
            paragraph_list_context_level = continuation_list_level + 1

        def _append_block_separator(separator: str, *, indent: str = "") -> None:
            """Emit a thematic separator as a distinct Markdown block."""
            if not separator:
                return
            if lines and lines[-1] != "":
                lines.append("")
            lines.append(f"{indent}{separator}" if indent else separator)
            lines.append("")

        def _accumulate_code_paragraph() -> None:
            """Accumulate one code paragraph, closing fences at page breaks."""
            if has_page_break and para is not None:
                code_parts = self._render_page_break_parts(
                    para,
                    preserve_formatting=False,
                    style_detector_config=style_detector_config,
                )
                for part_type, part in code_parts:
                    if part_type == "separator":
                        if code_block_acc is not None and code_block_acc.in_code_block:
                            code_markdown = code_block_acc.finalize()
                            if code_markdown:
                                lines.append(code_markdown)
                                lines.append("")
                        code_indent = " " * self._list_indent_spaces * paragraph_list_context_level
                        _append_block_separator(part, indent=code_indent)
                        continue
                    if code_block_acc is not None:
                        if not code_block_acc.in_code_block:
                            code_block_acc.start(list_level=paragraph_list_context_level)
                        code_block_acc.add_line(part)
                return

            if code_block_acc is not None:
                if not code_block_acc.in_code_block:
                    code_block_acc.start(list_level=paragraph_list_context_level)
                code_block_acc.add_line(paragraph_text)

        def _finish_code_at_section_break() -> None:
            if not has_section_break or not self._section_break_separator:
                return
            if code_block_acc is not None and code_block_acc.in_code_block:
                code_markdown = code_block_acc.finalize()
                if code_markdown:
                    lines.append(code_markdown)
                    lines.append("")
            _append_block_separator(self._section_break_separator)

        # 3a. Code block (style-based)
        if style_type == "code_block":
            _accumulate_code_paragraph()
            _finish_code_at_section_break()
            return lines, stats

        # 3b. Shading-based code block (fallback)
        effective_style_config = style_detector_config or StyleDetectorConfig()
        if para is not None and has_paragraph_gray_shading(
            para,
            wps_enabled=effective_style_config.wps_shading_enabled,
            word_enabled=effective_style_config.word_shading_enabled,
        ):
            _accumulate_code_paragraph()
            _finish_code_at_section_break()
            return lines, stats

        # 3c. A paragraph whose non-empty runs all carry one configured
        # character style is promoted to the corresponding block form.
        if para is not None and style_type is None:
            run_style_type, run_style_value = detect_full_paragraph_run_style(
                para,
                config=effective_style_config,
            )
            if run_style_type == "code_block":
                _accumulate_code_paragraph()
                _finish_code_at_section_break()
                return lines, stats
            if run_style_type == "quote":
                style_type, style_value = run_style_type, run_style_value

        # 3d. Quote block
        if style_type == "quote":
            # Close any open code block first
            if code_block_acc is not None and code_block_acc.in_code_block:
                cb_md = code_block_acc.finalize()
                if cb_md:
                    lines.append(cb_md)
                    lines.append("")
            quote_parts = (
                self._render_page_break_parts(
                    para,
                    preserve_formatting=self._preserve_formatting,
                    style_detector_config=style_detector_config,
                )
                if has_page_break and para is not None
                else [("text", paragraph_text)]
            )
            level = style_value if isinstance(style_value, int) else 1
            prefix = ">" * max(level, 1)
            quote_indent = " " * self._list_indent_spaces * paragraph_list_context_level
            anchor_groups = self._semantic_v3_recovery.ordinary_anchor_groups(para_element)
            owner_kinds = tuple(group.anchor.block_kind for group in anchor_groups)
            list_inside_quote = owner_kinds[:2] in {
                ("list", "block_quote"),
                ("list", "callout"),
            }
            quote_inside_list = owner_kinds[:2] in {
                ("block_quote", "list"),
                ("callout", "list"),
            }
            nested_num_id: str | None = None
            nested_level = 0
            nested_list_type: str | None = None
            if list_inside_quote or quote_inside_list:
                from docwen_plugin_document.shared.list_processing import (
                    detect_list_item,
                    format_list_marker,
                )

                nested_num_id, nested_ilvl, nested_list_type = detect_list_item(para, numbering_index)
                nested_level = nested_ilvl if isinstance(nested_ilvl, int) else 0

                def _nested_list_marker() -> str:
                    if nested_list_type == "ordered" and nested_num_id is not None and list_counter is not None:
                        return format_list_marker(
                            "ordered",
                            list_counter.next(nested_num_id, nested_level),
                        )
                    return format_list_marker("bullet", 0, self._unordered_list_marker_type)

            for part_type, part in quote_parts:
                if part_type == "separator":
                    _append_block_separator(part, indent=quote_indent)
                    continue
                if not part:
                    continue
                for line_index, qline in enumerate(part.splitlines()):
                    if list_inside_quote:
                        marker = _nested_list_marker() if line_index == 0 else ""
                        list_indent = " " * self._list_indent_spaces * nested_level
                        list_prefix = f"{list_indent}{marker} " if marker else f"{list_indent}  "
                        lines.append(f"{prefix} {list_prefix}{qline}".rstrip())
                    elif quote_inside_list:
                        if nested_list_type is not None and line_index == 0:
                            marker = _nested_list_marker()
                            list_indent = " " * self._list_indent_spaces * nested_level
                            lines.append(f"{list_indent}{marker} {prefix} {qline}".rstrip())
                        else:
                            lines.append(f"  {prefix} {qline}".rstrip())
                    else:
                        quote_line = f"{prefix} {qline}" if qline else prefix
                        lines.append(f"{quote_indent}{quote_line}" if quote_indent else quote_line)
                lines.append("")
            if has_section_break and self._section_break_separator:
                _append_block_separator(self._section_break_separator)
            return lines, stats

        # ── 4. Close code block if exiting to normal content ────────
        if code_block_acc is not None and code_block_acc.in_code_block:
            cb_md = code_block_acc.finalize()
            if cb_md:
                lines.append(cb_md)
                lines.append("")

        # ── 5. Heading detection ────────────────────────────────────
        outline_lvl = self._get_outline_level(
            para_element,
            paragraph=para,
            diagnostic_sink=diagnostic_sink,
            diagnostic_location=diagnostic_location,
        )
        is_title_style = style_name in ("Title", "Subtitle")
        style_heading_level = self._heading_level_from_style(style_name)
        is_heading = outline_lvl is not None or style_heading_level is not None or is_title_style

        # Determine heading level for Title/Subtitle styles
        heading_level = 1
        if is_title_style and outline_lvl is None:
            heading_level = 1 if style_name == "Title" else 2
        elif is_heading and outline_lvl is not None:
            heading_level = max(1, min(9, outline_lvl))
        elif style_heading_level is not None:
            heading_level = style_heading_level

        # ── 6. Extract text ─────────────────────────────────────────
        text = paragraph_text
        source_anchor = self._semantic_v3_recovery.source_anchor(para_element)

        # ── 6a. Formula extraction (H2) — OMML → LaTeX ──────────────
        _MATH_NS = OMML_NS["m"]
        _formula_text: str | None = None
        _has_formula = (
            para_element.find(f".//{{{_MATH_NS}}}oMath") is not None
            or para_element.find(f".//{{{_MATH_NS}}}oMathPara") is not None
        )
        if _has_formula:
            _formula_text = self._build_paragraph_text_with_formulas(
                para_element,
                para,
                preserve_formatting=(self._preserve_heading_formatting if is_heading else self._preserve_formatting),
            )

        if not text and _formula_text is None:
            # Empty paragraph — emit page/section break separators if any
            if has_page_break and self._page_break_separator:
                _append_block_separator(self._page_break_separator)
            if has_section_break and self._section_break_separator:
                _append_block_separator(self._section_break_separator)
            if not is_heading and not lines:
                return [""], stats
            return lines, stats

        # ── 7. Heading output ───────────────────────────────────────
        if is_heading:
            prefix = "#" * heading_level
            if _formula_text is not None:
                display_text = _formula_text
            elif self._preserve_heading_formatting and para is not None:
                display_text = self._extract_paragraph_text_formatted(
                    para,
                    preserve_formatting=True,
                    style_detector_config=style_detector_config,
                )
            else:
                display_text = text

            heading_parts: list[tuple[str, str]] | None = None
            heading_text_index: int | None = None
            if has_page_break and para is not None:
                heading_parts = self._render_page_break_parts(
                    para,
                    preserve_formatting=self._preserve_heading_formatting,
                    style_detector_config=style_detector_config,
                )
                heading_text_index = next(
                    (index for index, (part_type, _part) in enumerate(heading_parts) if part_type == "text"),
                    None,
                )
                if heading_text_index is not None:
                    display_text = heading_parts[heading_text_index][1]

            resolved_v4_heading = (
                self._resolved_v4_recovery.heading_import(para)
                if self._resolved_v4_recovery is not None and para is not None
                else None
            )
            if resolved_v4_heading is not None:
                for diagnostic in resolved_v4_heading.diagnostics:
                    self._resolved_v4_diagnostics.append((diagnostic.code, diagnostic.message, diagnostic_location))
                    if diagnostic_sink is not None:
                        diagnostic_sink.report_diagnostic(
                            "warning",
                            diagnostic.message,
                            code=diagnostic.code,
                            location=diagnostic_location,
                        )
            else:
                # Historical pStyle/add/remove behavior is isolated from an
                # authenticated resolved-v4 package.  A v4 Word list number is
                # a separate proven fact and is never inserted into or parsed
                # from the authored Heading text.
                if (
                    numbering_index is not None
                    and list_counter is not None
                    and para is not None
                    and para.style is not None
                ):
                    style_id = getattr(para.style, "style_id", None)
                    if isinstance(style_id, str) and style_id:
                        level_info = numbering_index.lookup_by_style_id(style_id)
                        if level_info is not None:
                            num_id_key = f"abs_{level_info.abstract_num_id}"
                            counter_val = (
                                list_counter.next(num_id_key, level_info.ilvl, start=level_info.start)
                                if list_counter
                                else level_info.start
                            )
                            parent_counters = list_counter.snapshot(num_id_key) if list_counter else None
                            num_prefix = numbering_index.render_numbering_text(
                                level_info,
                                counter_val,
                                parent_counters=parent_counters,
                            )
                            if num_prefix and not remove_numbering:
                                display_text = f"{num_prefix}{display_text}"

                if remove_numbering:
                    from docwen_core.text.heading_numbering import strip_heading_prefix

                    numbering, plain_without_numbering = strip_heading_prefix(
                        text,
                        rules=heading_cleanup_rules,
                    )
                    if numbering:
                        marker_index = display_text.find(numbering)
                        if marker_index >= 0:
                            display_text = display_text[:marker_index] + display_text[marker_index + len(numbering) :]
                        else:
                            display_text = plain_without_numbering
                    else:
                        _, display_text = strip_heading_prefix(
                            display_text,
                            rules=heading_cleanup_rules,
                        )

                # Add new numbering after potential removal on the isolated
                # historical route only.
                if heading_formatter is not None:
                    display_text = heading_formatter.format_heading(display_text, heading_level)

            if source_anchor is not None:
                if source_anchor.owner_kind != "semantic_target" or source_anchor.block_kind != "heading":
                    raise ValueError("a v3 Heading paragraph has non-Heading source ownership")
                display_text = f"{display_text} ^{source_anchor.source_id}"

            if heading_parts is not None:
                if heading_text_index is not None:
                    heading_parts[heading_text_index] = ("text", display_text)
                for part_type, part in heading_parts:
                    if part_type == "separator":
                        _append_block_separator(part)
                    elif part:
                        lines.append(f"{prefix} {part}")
                        lines.append("")
            else:
                lines.append(f"{prefix} {display_text}")
                lines.append("")
            if has_section_break and self._section_break_separator:
                _append_block_separator(self._section_break_separator)
            stats["headings"] += 1
        else:
            # ── 8. Regular paragraph ────────────────────────────────
            from docwen_plugin_document.shared.list_processing import (
                detect_list_item,
                format_list_marker,
            )

            num_id, ilvl, list_type = None, None, None
            if para is not None and numbering_index is not None:
                num_id, ilvl, list_type = detect_list_item(para, numbering_index)

            if _formula_text is not None:
                display = _formula_text
            elif semantic_reference_text is not None:
                display = semantic_reference_text
            else:
                formatted_text = self._extract_paragraph_text_formatted(
                    para if para is not None else para_element,
                    preserve_formatting=self._preserve_formatting,
                    style_detector_config=style_detector_config,
                )
                display = formatted_text if formatted_text.strip() else text
            if source_anchor is not None:
                if source_anchor.owner_kind == "semantic_target" and source_anchor.block_kind in {
                    "figure",
                    "equation",
                    "code_block",
                }:
                    pass
                elif source_anchor.owner_kind != "ordinary_anchor":
                    raise ValueError("a v3 ordinary paragraph has invalid source ownership")
                elif source_anchor.block_kind == "paragraph":
                    display = f"{display} ^{source_anchor.source_id}"
                elif source_anchor.block_kind == "list_item":
                    anchor_group = self._semantic_v3_recovery.ordinary_anchor_group(para_element)
                    if anchor_group is None or anchor_group.index == 0:
                        display = f"{display} ^{source_anchor.source_id}"

            page_parts = (
                self._render_page_break_parts(
                    para,
                    preserve_formatting=self._preserve_formatting,
                    style_detector_config=style_detector_config,
                )
                if has_page_break and para is not None
                else None
            )

            # Handle pStyle-based numbering (abs_ num_id from pStyle fallback)
            if list_type and num_id and num_id.startswith("abs_") and numbering_index is not None:
                abs_id = num_id[4:]  # strip "abs_" prefix
                level_info = numbering_index.lookup_by_abstract(abs_id, ilvl if isinstance(ilvl, int) else 0)
                if level_info is not None and not remove_numbering:
                    num_id_key = f"abs_{abs_id}"
                    counter_val = (
                        list_counter.next(num_id_key, level_info.ilvl, start=level_info.start)
                        if list_counter
                        else level_info.start
                    )
                    parent_counters = list_counter.snapshot(num_id_key) if list_counter else None
                    pstyle_prefix = numbering_index.render_numbering_text(
                        level_info,
                        counter_val,
                        parent_counters=parent_counters,
                    )
                    first_prefix = pstyle_prefix
                else:
                    first_prefix = ""
                if page_parts is not None:
                    first_text = True
                    continuation_indent = (
                        " " * len(first_prefix)
                        if first_prefix and _COMMONMARK_LIST_PREFIX_RE.fullmatch(first_prefix)
                        else ""
                    )
                    for part_type, part in page_parts:
                        if part_type == "separator":
                            _append_block_separator(part, indent=continuation_indent if not first_text else "")
                        elif first_text:
                            lines.append(f"{first_prefix}{part}")
                            first_text = False
                        else:
                            lines.append(f"{continuation_indent}{part}")
                else:
                    lines.append(f"{first_prefix}{display}")
            elif list_type and num_id and list_counter is not None:
                # Standard list item — indent and prefix with marker
                level = ilvl if isinstance(ilvl, int) else 0
                indent = " " * self._list_indent_spaces * level
                if list_type == "ordered":
                    ordinal = list_counter.next(num_id, level)
                    marker = format_list_marker("ordered", ordinal)
                else:
                    marker = format_list_marker("bullet", 0, self._unordered_list_marker_type)
                if page_parts is not None:
                    first_text = True
                    content_indent = " " * max(self._list_indent_spaces, len(marker) + 1)
                    continuation_indent = f"{indent}{content_indent}"
                    for part_type, part in page_parts:
                        if part_type == "separator":
                            _append_block_separator(part, indent=continuation_indent if not first_text else indent)
                        elif first_text:
                            lines.append(f"{indent}{marker} {part}")
                            first_text = False
                        else:
                            lines.append(f"{continuation_indent}{part}")
                else:
                    lines.append(f"{indent}{marker} {display}")
            elif continuation_list_level is not None:
                indent = " " * self._list_indent_spaces * (continuation_list_level + 1)
                if page_parts is not None:
                    for part_type, part in page_parts:
                        if part_type == "separator":
                            _append_block_separator(part, indent=indent)
                        else:
                            lines.append(f"{indent}{part}")
                else:
                    lines.append(f"{indent}{display}")
            else:
                # Normal paragraph — reset list counter continuity
                if list_counter is not None:
                    list_counter.reset()

                # Page break text splitting
                if page_parts is not None:
                    for part_type, part in page_parts:
                        if part_type == "separator":
                            _append_block_separator(part)
                        else:
                            lines.append(part)
                else:
                    lines.append(display)

            if has_section_break and self._section_break_separator:
                _append_block_separator(self._section_break_separator)
            if not lines or lines[-1] != "":
                lines.append("")
            stats["paragraphs"] += 1

        return lines, stats

    def _render_page_break_parts(
        self,
        para: Any,
        *,
        preserve_formatting: bool | None = None,
        style_detector_config: Any = None,
    ) -> list[tuple[str, str]]:
        """Render page-separated paragraph parts without losing rich content."""
        from docwen_plugin_document.shared.markdown_runs import (
            render_paragraph_runs_split_on_page_breaks,
        )

        standalone_formula = not (getattr(para, "text", "") or "").strip()

        def _render_formula(child: Any) -> str | None:
            tag = child.tag.split("}")[-1] if "}" in (child.tag or "") else (child.tag or "")
            return extract_formula_from_element(
                child,
                block=(tag == "oMathPara" or standalone_formula),
            )

        segments = render_paragraph_runs_split_on_page_breaks(
            para,
            note_extractor=self._note_extractor,
            preserve_formatting=(self._preserve_formatting if preserve_formatting is None else preserve_formatting),
            syntax_config=self._syntax_for_rendering(),
            style_detector_config=style_detector_config,
            math_renderer=_render_formula,
        )
        parts: list[tuple[str, str]] = []
        for index, segment in enumerate(segments):
            text = segment.strip()
            if text:
                parts.append(("text", text))
            if index < len(segments) - 1 and self._page_break_separator and (not parts or parts[-1][0] != "separator"):
                parts.append(("separator", self._page_break_separator))

        if not parts:
            return [("text", getattr(para, "text", ""))]
        return parts

    def _paragraph_list_item_level(self, para: Any, numbering_index: Any = None) -> int | None:
        """Return the Word list item level for the paragraph itself."""
        if para is None:
            return None
        from docwen_plugin_document.shared.list_processing import detect_list_item

        _num_id, ilvl, list_type = detect_list_item(para, numbering_index)
        if list_type and isinstance(ilvl, int):
            return ilvl
        return None

    def _paragraph_list_context_level(self, para: Any, numbering_index: Any = None) -> int:
        """Return Markdown block indentation level for a list-contained paragraph."""
        item_level = self._paragraph_list_item_level(para, numbering_index)
        return item_level + 1 if item_level is not None else 0

    def _is_list_continuation_paragraph(self, para: Any) -> bool:
        """Return True when a paragraph should stay inside the previous list item."""
        if para is None:
            return False
        style_name = ""
        if getattr(para, "style", None) is not None:
            style_name = getattr(para.style, "name", "") or ""
        return style_name in _LIST_CONTEXT_STYLE_NAMES or self._has_extra_list_indent(para)

    def _has_extra_list_indent(self, para: Any) -> bool:
        current_left = self._indent_twips(getattr(getattr(para, "paragraph_format", None), "left_indent", None))
        if current_left is None:
            return False
        style = getattr(para, "style", None)
        style_left = self._indent_twips(getattr(getattr(style, "paragraph_format", None), "left_indent", None)) or 0
        return current_left - style_left >= _MIN_EXTRA_LIST_INDENT_TWIPS

    def _table_has_extra_list_indent(self, tbl_element: Any) -> bool:
        try:
            from docx.oxml.ns import qn

            tbl_pr = tbl_element.find(qn("w:tblPr"))
            if tbl_pr is None:
                return False
            tbl_ind = tbl_pr.find(qn("w:tblInd"))
            if tbl_ind is None:
                return False
            indent_value = tbl_ind.get(qn("w:w")) or tbl_ind.get("w:w")
            if indent_value is None:
                return False
            return int(indent_value) >= _MIN_EXTRA_LIST_INDENT_TWIPS
        except Exception:
            return False

    @staticmethod
    def _indent_twips(indent: Any) -> int | None:
        if indent is None:
            return None
        twips = getattr(indent, "twips", None)
        if twips is not None:
            return int(twips)
        try:
            return int(indent)
        except (TypeError, ValueError):
            return None

    def _render_extracted_paragraph(self, tb_para: Any) -> str:
        """Render an extracted textbox paragraph as Markdown text.

        ``ExtractedParagraph`` carries only plain text (no XML element),
        so this simply returns the stripped text.
        """
        text = tb_para.text.strip() if tb_para.text else ""
        return text

    def _build_paragraph_text_with_formulas(
        self,
        para_element: Any,
        para: Any = None,
        *,
        preserve_formatting: bool = True,
    ) -> str:
        """Build paragraph text with OMML formulas converted to ``$...$`` LaTeX.

        Walks the paragraph XML children in document order, converting
        ``<m:oMath>`` (inline) and ``<m:oMathPara>`` (display) elements
        to their LaTeX representation via ``extract_formula_from_element``.
        Non-math children are rendered with standard run formatting (bold,
        italic, hyperlinks, inline note references, etc.).
        """
        from docwen_plugin_document.shared.markdown_runs import (
            append_formatted_run_text,
            resolve_hyperlink_target,
        )

        _MATH_NS = OMML_NS["m"]
        elem = getattr(para, "_p", para_element) if para is not None else para_element
        paragraph_text = (
            para.text.strip() if para is not None else self._extract_paragraph_text_raw(para_element).strip()
        )
        standalone_formula = not paragraph_text
        parts: list[str] = []

        def _process_children(parent: Any) -> None:
            for child in parent:
                _process_child(child)

        def _process_alternate_content(element: Any) -> None:
            """Render one effective ``mc:AlternateContent`` branch.

            Word commonly stores the same formula in a modern ``Choice`` and
            a compatibility ``Fallback``.  Rendering both duplicates content,
            while skipping the wrapper loses the formula entirely.  Prefer
            the first Choice that produces output, then use Fallback only when
            every Choice is empty or unsupported by this renderer.
            """
            branches: dict[str, list[Any]] = {"Choice": [], "Fallback": []}
            for branch in element:
                branch_tag = branch.tag.split("}")[-1] if "}" in (branch.tag or "") else (branch.tag or "")
                if branch_tag in branches:
                    branches[branch_tag].append(branch)

            for branch_kind in ("Choice", "Fallback"):
                for branch in branches[branch_kind]:
                    previous_parts = list(parts)
                    _process_children(branch)
                    if parts != previous_parts:
                        return

        def _process_child(child: Any) -> None:
            tag = child.tag.split("}")[-1] if "}" in (child.tag or "") else (child.tag or "")

            if tag == "oMath":
                # Some legacy Word producers emit display math as a bare
                # ``oMath`` in an otherwise empty paragraph instead of an
                # ``oMathPara`` wrapper.  Both old DocWen converters treated
                # that shape as block math; mixed-content paragraphs remain
                # inline.
                latex = extract_formula_from_element(child, block=standalone_formula)
                if latex:
                    parts.append(latex)
            elif tag == "oMathPara":
                latex = extract_formula_from_element(child, block=True)
                if latex:
                    parts.append(latex)
            elif tag == "r":
                _handle_run(child)
            elif tag == "hyperlink":
                url = resolve_hyperlink_target(para, child) if para is not None else None
                if url:
                    previous_parts = list(parts)
                    parts.append("[")
                    opener_parts = list(parts)
                _process_children(child)
                if url:
                    if parts == opener_parts:
                        parts[:] = previous_parts
                    else:
                        parts.append(f"]({url})")
            elif tag in ("ins", "moveTo", "fldSimple", "smartTag", "sdt", "sdtContent", "customXml"):
                _process_children(child)
            elif tag in ("del", "moveFrom"):
                pass
            elif tag == "AlternateContent":
                _process_alternate_content(child)

        def _handle_run(run: Any) -> None:
            from docx.oxml.ns import qn

            br = run.find(qn("w:br"))
            if br is not None:
                parts.append("\n")
                return
            tab = run.find(qn("w:tab"))
            if tab is not None:
                parts.append("\t")
                return

            # Note references (before text check — ref runs may lack w:t)
            _ne = self._note_extractor
            fn_ref = run.find(qn("w:footnoteReference"))
            if fn_ref is not None and _ne is not None:
                w_id = fn_ref.get(qn("w:id"))
                if w_id is not None:
                    parts.append(_ne.get_reference_text("footnote", int(w_id)))
                    return

            en_ref = run.find(qn("w:endnoteReference"))
            if en_ref is not None and _ne is not None:
                w_id = en_ref.get(qn("w:id"))
                if w_id is not None:
                    parts.append(_ne.get_reference_text("endnote", int(w_id)))
                    return

            t = run.find(qn("w:t"))
            if t is None or t.text is None:
                return
            if not preserve_formatting:
                parts.append(t.text)
                return
            append_formatted_run_text(
                parts,
                t.text,
                run,
                syntax_config=self._syntax_for_rendering(),
            )

        _process_children(elem)
        text = "".join(parts)

        return text

    def _extract_paragraph_text_raw(self, para_element: Any) -> str:
        """Extract text from raw XML paragraph element (fallback)."""
        from docx.oxml.ns import qn

        parts: list[str] = []
        for run in para_element.findall(qn("w:r")):
            t = run.find(qn("w:t"))
            if t is not None and t.text:
                parts.append(t.text)
        return "".join(parts)

    def _extract_paragraph_text_formatted(
        self,
        para_element: Any,
        *,
        preserve_formatting: bool = True,
        style_detector_config: Any = None,
    ) -> str:
        """Extract text with inline formatting markers.

        Delegates to shared ``render_paragraph_runs`` for formatting
        and hyperlink target resolution.
        """
        from docwen_plugin_document.shared.markdown_runs import render_paragraph_runs

        return render_paragraph_runs(
            para_element,
            note_extractor=self._note_extractor,
            preserve_formatting=preserve_formatting,
            syntax_config=self._syntax_for_rendering(),
            style_detector_config=style_detector_config,
        )

    def _process_images_for_output(
        self,
        element: Any,
        doc: Any,
        context: Any,
        *,
        prefer_source_names: bool = False,
        preserve_resources: bool = True,
        preserve_owner: bool = False,
    ) -> tuple[list[str], list[Any]]:
        """Extract images from *element* and produce Markdown refs.

        Uses the shared ``extract_images_from_element`` helper and the
        core ``generate_image_markdown`` function so that output respects
        request-owned image mode and link-style policy.

        Returns (refs, infos) so the caller can attach OCR text.
        """
        from docwen_core.docx_parsing.image_extraction import (
            extract_images_from_element,
        )
        from docwen_core.text.image_markdown import generate_image_markdown

        output_dir = str(context.workspace.staging_dir)
        infos = extract_images_from_element(
            element,
            doc.part.related_parts,
            output_dir,
            name_prefix="docx-image",
            prefer_source_names=prefer_source_names,
        )
        export_modes = self._request_policy.resolve_export_modes()
        image_mode = export_modes["image_extraction_mode"]
        if image_mode not in {"file", "base64", "embed", "omit"}:
            image_mode = "file"
        image_link_style = "markdown_embed" if prefer_source_names else self._request_policy.image_link_style

        if preserve_resources:
            refs = [
                generate_image_markdown(
                    image_path=self._image_markdown_path(info, image_mode),
                    image_mode=image_mode,
                    image_link_style=image_link_style,
                    alt_text=info.alt,
                    export_semantics=self._request_policy.export,
                )
                + "\n"
                for info in infos
            ]
        elif preserve_owner:
            refs = ["![image omitted]()\n" for _info in infos]
        else:
            refs = ["" for _info in infos]
        if preserve_resources:
            for info in infos:
                self._register_image_artifact(context, info)
        return refs, infos

    @staticmethod
    def _discard_unpreserved_images(infos: list[Any]) -> None:
        """Remove request-owned extraction bytes after OCR-only use."""
        from pathlib import Path

        for info in infos:
            Path(info.path).unlink(missing_ok=True)

    @staticmethod
    def _image_markdown_path(info: Any, image_mode: str) -> str:
        """Return the path/target that should appear in Markdown."""
        from pathlib import Path

        if image_mode == "base64":
            return info.path
        return Path(info.path).name

    def _register_image_artifact(self, context: Any, info: Any) -> None:
        """Stage an image manifest until the primary Markdown is durable."""
        from pathlib import Path

        from docwen_core.models.artifact import ARTIFACT_KIND_IMAGE, ArtifactManifest

        registered_paths = {
            str(Path(artifact.staging_path).resolve())
            for artifact in [
                *getattr(context.workspace, "registered_artifacts", []),
                *self._pending_artifacts,
            ]
        }
        image_path = str(Path(info.path).resolve())
        if image_path in registered_paths:
            return

        suggested_name = Path(info.path).name
        self._pending_artifacts.append(
            ArtifactManifest(
                artifact_id=str(uuid.uuid4()),
                kind=ARTIFACT_KIND_IMAGE,
                staging_path=info.path,
                suggested_name=suggested_name,
                media_type=info.content_type or "application/octet-stream",
                metadata={
                    "source_format": "docx",
                    "rel_id": info.rel_id,
                    "alt": info.alt,
                },
                is_primary=False,
            )
        )

    def _discard_pending_artifacts(self) -> None:
        """Delete uncommitted sidecars and clear their request-local manifests."""

        from pathlib import Path

        for artifact in self._pending_artifacts:
            Path(artifact.staging_path).unlink(missing_ok=True)
        self._pending_artifacts.clear()
        if self._pending_primary_path is not None:
            Path(self._pending_primary_path).unlink(missing_ok=True)
            self._pending_primary_path = None

    def _extract_paragraph_images(
        self,
        element: Any,
        doc: Any,
        context: Any,
    ) -> tuple[list[str], int]:
        """Extract embedded images from any XML element.

        Delegates to the shared ``extract_images_from_element`` helper.
        Writes images to the staging directory and returns Markdown refs.

        Returns (list_of_markdown_image_refs, image_count).
        """
        from docwen_core.docx_parsing.image_extraction import (
            extract_images_from_element,
        )

        # Use staging directory for extraction
        output_dir = str(context.workspace.staging_dir)
        infos = extract_images_from_element(
            element,
            doc.part.related_parts,
            output_dir,
            name_prefix="docx-image",
        )

        refs = [f"![{info.alt}]({info.path})\n" for info in infos]
        return refs, len(infos)

    def _ocr_per_image(
        self,
        img_refs: list[str],
        infos: list[Any],
        main_stem: str,
        img_seq: int,
        ocr_placement: str,
        context: Any,
        *,
        ocr_language: str = "auto",
        current_locale: str = "zh_CN",
        retain_image_owner: bool = False,
    ) -> tuple[list[str], int]:
        """Run OCR on extracted images, producing inline blockquotes or sidecars.

        Returns updated *img_refs* (with OCR blockquotes or .md links) and the
        new *img_seq* value.
        """
        from docwen_core.detection import detect_content_format
        from docwen_core.text.ocr import run_ocr_outcome

        md_style = self._request_policy.export.md_file_link_style
        for i, info in enumerate(infos):
            outcome = run_ocr_outcome(
                info.path,
                source_format=detect_content_format(str(info.path)).format,
                ocr_language=ocr_language,
                current_locale=current_locale,
            )
            _report_ocr_best_effort(
                context.progress,
                outcome.status,
                location=str(info.path),
            )
            ocr_text = outcome.recognized_text
            if not ocr_text:
                continue

            if ocr_placement == "image_md":
                img_seq += 1
                sidecar_stem = f"{main_stem}__img_{img_seq:03d}_ocr"
                sidecar_text, repl_link = build_image_ocr_sidecar(
                    sidecar_stem=sidecar_stem,
                    source_format="docx",
                    image_markdown="" if retain_image_owner else img_refs[i],
                    ocr_text=ocr_text,
                    md_link_style=md_style,
                )
                # Write sidecar to staging
                from pathlib import Path

                sidecar_path = context.workspace.create_artifact_path("auxiliary", ".md")
                sidecar_file = Path(sidecar_path)
                try:
                    sidecar_file.write_text(sidecar_text, encoding="utf-8")
                except BaseException:
                    sidecar_file.unlink(missing_ok=True)
                    raise
                # Register artifact
                import uuid

                from docwen_core.markdown_utils import sanitize_filename
                from docwen_core.models.artifact import ArtifactManifest

                ocr_md_name = sanitize_filename(f"{sidecar_stem}.md")
                sidecar_artifact = ArtifactManifest(
                    artifact_id=str(uuid.uuid4()),
                    kind="auxiliary",
                    staging_path=sidecar_path,
                    suggested_name=ocr_md_name,
                    media_type="text/markdown",
                    metadata={"source_format": "docx", "ocr": True},
                    is_primary=False,
                )
                self._pending_artifacts.append(sidecar_artifact)
                if not retain_image_owner:
                    img_refs[i] = repl_link + "\n"
            else:
                # The configured title is specifically a main-Markdown
                # presentation fragment.  Preserve its Markdown markup;
                # image_md sidecars intentionally remain title-free.
                ocr_block = _format_main_ocr_blockquote(
                    ocr_text,
                    self._request_policy.ocr_blockquote_title,
                )
                img_refs[i] = img_refs[i].rstrip("\n") + f"\n{ocr_block}\n\n"

        return img_refs, img_seq

    @staticmethod
    def _merge_img_refs_into_lines(lines: list[str], img_refs: list[str]) -> list[str]:
        """Merge image reference strings into output lines.

        Appends each img_ref to the last non-empty line, or creates
        a new line if all lines are empty.  Used by paragraph, table,
        and SDT branches so the merge behaviour stays consistent.
        """
        if not img_refs:
            return lines

        merged_refs = "".join(img_refs)
        if not merged_refs:
            return lines
        for index in range(len(lines) - 1, -1, -1):
            if lines[index].strip():
                merged = list(lines)
                merged[index] += merged_refs
                return merged

        if lines:
            return [merged_refs, *lines]
        return [merged_refs, ""]

    @staticmethod
    def _append_inline_anchor(lines: list[str], source_id: str) -> list[str]:
        """Attach an authenticated inline anchor to the last visible line."""

        output = list(lines)
        for index in range(len(output) - 1, -1, -1):
            if output[index].strip():
                content = output[index].rstrip("\r\n")
                line_ending = output[index][len(content) :]
                output[index] = f"{content} ^{source_id}{line_ending}"
                return output
        return [f"^{source_id}", *output]

    @staticmethod
    def _append_anchor_to_first_image_reference(lines: list[str], source_id: str) -> list[str]:
        """Attach an image-owned ID before any following OCR presentation."""

        output = list(lines)
        for index, line in enumerate(output):
            segments = line.splitlines(keepends=True)
            for segment_index, segment in enumerate(segments):
                if "![[" not in segment and "![" not in segment:
                    continue
                content = segment.rstrip("\r\n")
                line_ending = segment[len(content) :]
                segments[segment_index] = f"{content} ^{source_id}{line_ending}"
                output[index] = "".join(segments)
                return output
        raise ValueError("an authenticated image anchor has no visible image reference")

    @staticmethod
    def _image_ext_from_content_type(content_type: str) -> str:
        """Map MIME content type to a file extension."""
        type_map = {
            "image/png": "png",
            "image/jpeg": "jpg",
            "image/gif": "gif",
            "image/bmp": "bmp",
            "image/tiff": "tiff",
            "image/x-wmf": "wmf",
            "image/x-emf": "emf",
            "image/svg+xml": "svg",
        }
        return type_map.get(content_type, "png")

    def _get_outline_level(
        self,
        para_element: Any,
        *,
        paragraph: Any = None,
        diagnostic_sink: Any = None,
        diagnostic_location: str = "",
    ) -> int | None:
        """Get the effective heading level, including inherited style values.

        Word stores the outline level for built-in Heading styles in
        ``styles.xml`` rather than directly on most paragraphs.  Check direct
        paragraph formatting first, then walk the paragraph style's
        ``basedOn`` chain.  Word's ``outlineLvl=9`` sentinel means body text
        and is deliberately excluded.
        """
        try:
            from docx.oxml.ns import qn

            body_text_outline = -1

            def _level_from_properties(properties: Any) -> int | None:
                if properties is None:
                    return None
                outline = properties.find(qn("w:outlineLvl"))
                if outline is not None:
                    val = outline.get(qn("w:val"))
                    if val is not None:
                        level = int(val)
                        if 0 <= level <= 8:
                            return level + 1  # 0-based → 1-based
                        if level == 9:
                            return body_text_outline

                return None

            direct_level = _level_from_properties(para_element.find(qn("w:pPr")))
            if direct_level == body_text_outline:
                return None
            if direct_level is not None:
                return direct_level

            style = getattr(paragraph, "style", None)
            visited: set[str] = set()
            while style is not None:
                style_id = str(getattr(style, "style_id", ""))
                if style_id in visited:
                    break
                visited.add(style_id)
                style_element = getattr(style, "element", None)
                # Some callers supply lightweight style doubles carrying only
                # a name/id.  Do not interpret their dynamic attributes as
                # real WordprocessingML properties.
                if not isinstance(getattr(style_element, "tag", None), str):
                    break
                style_level = _level_from_properties(getattr(style_element, "pPr", None))
                if style_level == body_text_outline:
                    return None
                if style_level is not None:
                    return style_level
                style = getattr(style, "base_style", None)
        except Exception as exc:
            if diagnostic_sink is not None:
                diagnostic_sink.report_diagnostic(
                    "warning",
                    (
                        "Could not parse a Word outline level; kept the paragraph "
                        f"and used style fallback after {type(exc).__name__}."
                    ),
                    code="DOCX2MD-OUTLINE-FALLBACK",
                    location=diagnostic_location,
                )
        return None

    @staticmethod
    def _heading_level_from_style(style_name: str) -> int | None:
        """Extract an exact 1-9 level from ``Heading N``/``HeadingN``."""
        match = re.fullmatch(r"heading\s*([1-9])", style_name.strip(), flags=re.IGNORECASE)
        return int(match.group(1)) if match is not None else None

    # ── Table processing ────────────────────────────────────────────

    def _get_cell_text(
        self,
        cell: Any,
        para_by_element: dict[int, Any] | None,
        _depth: int = 0,
        table_merge_strategy: str = "fill",
        preserve_formatting: bool = True,
        paragraph_image_renderer: Any = None,
    ) -> str:
        """Extract formatted text from a table cell (``<w:tc>``).

        Delegates to ``render_paragraph_runs`` for formatting, handles
        formula content (OMML → LaTeX) and nested tables (``<w:tbl>``).

        Args:
            cell: The ``<w:tc>`` lxml element.
            para_by_element: Lookup from element id to python-docx Paragraph.
            _depth: Recursion depth for nested tables (max 2).
        """
        from docwen_plugin_document.shared.markdown_runs import render_paragraph_runs

        cell_text_parts: list[str] = []

        # Preserve direct child order so an image remains associated with the
        # exact paragraph/cell that owned it instead of drifting to the last
        # Markdown row of the whole table.
        for child in cell:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "tbl" and _depth < 2:
                nested_lines, _ = self._process_table(
                    child,
                    para_by_element=para_by_element,
                    _depth=_depth + 1,
                    table_merge_strategy=table_merge_strategy,
                    paragraph_image_renderer=paragraph_image_renderer,
                )
                if nested_lines:
                    for nl in nested_lines:
                        if nl.strip():
                            cell_text_parts.append("  " + nl)
                continue
            if tag != "p":
                continue
            para_elem = child
            para = (para_by_element or {}).get(id(para_elem))
            _MATH_NS = OMML_NS["m"]
            if (
                para_elem.find(f".//{{{_MATH_NS}}}oMath") is not None
                or para_elem.find(f".//{{{_MATH_NS}}}oMathPara") is not None
            ):
                formula_text = self._build_paragraph_text_with_formulas(
                    para_elem,
                    para,
                    preserve_formatting=preserve_formatting,
                )
                if formula_text.strip():
                    cell_text_parts.append(formula_text.strip())
            elif para is not None:
                formatted = render_paragraph_runs(
                    para,
                    note_extractor=self._note_extractor,
                    preserve_formatting=preserve_formatting,
                    syntax_config=self._syntax_for_rendering(),
                    style_detector_config=self._request_policy.style_detector,
                )
                if formatted.strip():
                    cell_text_parts.append(formatted.strip())
            else:
                text = self._extract_paragraph_text_raw(para_elem)
                if text.strip():
                    cell_text_parts.append(text.strip())
            if paragraph_image_renderer is not None:
                cell_text_parts.extend(
                    ref.strip() for ref in paragraph_image_renderer(para_elem) if isinstance(ref, str) and ref.strip()
                )
        # Markdown tables cannot contain literal newlines. Preserve distinct
        # Word paragraphs (and any multi-line nested content) as HTML breaks,
        # matching the readable old-system table-cell contract.
        from docwen_core.links import escape_unescaped_pipes

        cell_text = "<br>".join(
            part.replace("\r\n", "<br>").replace("\n", "<br>").replace("\r", "<br>") for part in cell_text_parts if part
        )
        return escape_unescaped_pipes(cell_text)

    def _process_table(
        self,
        tbl_element: Any,
        para_by_element: dict[int, Any] | None = None,
        _depth: int = 0,
        table_merge_strategy: str = "fill",
        paragraph_image_renderer: Any = None,
    ) -> tuple[list[str], int]:
        """Process a ``<w:tbl>`` element into a Markdown table.

        Detects ``w:gridSpan`` (colspan) and ``w:vMerge`` (rowspan) to
        build a merge-aware semantic grid via ``build_table_semantic_grid``,
        then renders with ``render_table_semantic_grid``.

        Args:
            tbl_element: The ``<w:tbl>`` lxml element.
            para_by_element: Lookup from element id to python-docx Paragraph.
            _depth: Recursion depth for nested tables (max 2 guards loops).

        Returns (list_of_markdown_lines, image_count).
        """
        from docwen_core.docx_parsing.table_extraction import (
            markdown_table_lines,
            render_docx_table_rows,
        )

        table_metadata = extract_semantic_table_metadata(tbl_element)
        structural = table_metadata.header_rows > 1 or table_metadata.header_columns > 0
        resolved_round_trip = self._resolved_v4_recovery is not None or bool(self._resolved_v4_diagnostics)
        use_merge_markers = table_merge_strategy == "marker" or structural or resolved_round_trip
        rendered = render_docx_table_rows(
            tbl_element,
            cell_text_resolver=lambda cell, row_index, _virtual_col: self._get_cell_text(
                cell,
                para_by_element,
                _depth=_depth,
                table_merge_strategy=table_merge_strategy,
                preserve_formatting=(
                    self._preserve_table_header_formatting if row_index == 0 else self._preserve_formatting
                ),
                paragraph_image_renderer=paragraph_image_renderer,
            ),
            strategy="marker" if use_merge_markers else table_merge_strategy,
            # Once the Structural Tables dialect is recognized, uncovered
            # literal marker cells must always be escaped.  In fill mode the
            # repeated values are literals too; in marker mode covered cells
            # remain the structural carriers.
            escape_literal_merge_markers=True,
        )
        lines = markdown_table_lines(
            rendered,
            header_rows=table_metadata.header_rows if structural else 1,
            header_columns=table_metadata.header_columns if structural else 0,
        )
        if lines:
            lines.append("")
        return lines, 0

    def _process_table_for_output(
        self,
        tbl_element: Any,
        *,
        para_by_element: dict[int, Any] | None,
        table_merge_strategy: str,
        doc: Any,
        context: Any,
        enable_ocr: bool,
        preserve_resources: bool,
        main_stem: str,
        img_seq: int,
        ocr_placement: str,
        ocr_language: str,
        current_locale: str,
    ) -> tuple[list[str], int, int]:
        """Render a table while projecting each image into its owning cell."""
        image_count = 0

        def _render_paragraph_images(paragraph_element: Any) -> list[str]:
            nonlocal image_count, img_seq
            refs, infos = self._process_images_for_output(
                paragraph_element,
                doc,
                context,
                preserve_resources=preserve_resources,
            )
            image_count += len(infos) if preserve_resources else 0
            try:
                if enable_ocr and infos:
                    refs, img_seq = self._ocr_per_image(
                        refs,
                        infos,
                        main_stem,
                        img_seq,
                        ocr_placement,
                        context,
                        ocr_language=ocr_language,
                        current_locale=current_locale,
                    )
            finally:
                if not preserve_resources:
                    self._discard_unpreserved_images(infos)
            return refs

        lines, _unused_image_count = self._process_table(
            tbl_element,
            para_by_element=para_by_element,
            table_merge_strategy=table_merge_strategy,
            paragraph_image_renderer=_render_paragraph_images,
        )
        return lines, image_count, img_seq
