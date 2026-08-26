import pytest
from docx import Document
from docx.oxml.ns import qn

from docwen_core.docx_parsing.table_extraction import (
    render_docx_table_rows,
)

pytestmark = pytest.mark.unit


def _plain_text(cell, _row: int, _column: int) -> str:
    return "".join(node.text or "" for node in cell.iter(qn("w:t")))


def test_shared_table_rows_use_consumer_cell_resolver():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"

    rows = render_docx_table_rows(table._tbl, cell_text_resolver=_plain_text)

    assert rows == [["A", "B"]]


def test_shared_table_rows_project_merged_cells_once():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "merged"
    table.cell(0, 0).merge(table.cell(0, 1))

    rows = render_docx_table_rows(
        table._tbl,
        cell_text_resolver=_plain_text,
        strategy="marker",
    )

    assert rows == [["merged", "<"]]
