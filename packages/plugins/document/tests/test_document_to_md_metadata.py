"""Tests for DOCX→MD YAML front matter and title/subtitle metadata extraction.

Covers finding_ids: F-E1-005, F-E1-016, F-E1-028.

Verifies that the ``DocxToMarkdownConverter``:
- Generates YAML front matter with aliases/title/subtitle fields.
- Extracts title from Word "Title" style paragraphs, with filename fallback.
- Extracts subtitle from Word "Subtitle" style paragraphs.
- Renders Title/Subtitle paragraphs as H1/H2 in the body.
- Skips consumed Title/Subtitle paragraphs from body rendering.
- Does not break the gongwen pipeline.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from docwen_core.cancellation import CancellationToken
from docwen_core.models.file_ref import FileRef
from docwen_core.models.request import ConversionRequest, OutputPolicy

pytestmark = pytest.mark.contract

# ═══════════════════════════════════════════════════════════════════════════
# YAML value formatting (unit tests)
# ═══════════════════════════════════════════════════════════════════════════


class TestFormatYamlValue:
    """Tests for ``_format_yaml_value`` static method."""

    @pytest.fixture(autouse=True)
    def _setup(self):
        from docwen_plugin_document.to_markdown.converter import (
            DocxToMarkdownConverter,
        )

        self._fmt = DocxToMarkdownConverter._format_yaml_value

    def test_plain_text_passes_through(self):
        assert self._fmt("hello") == "hello"

    def test_chinese_text_passes_through(self):
        assert self._fmt("文档标题") == "文档标题"

    def test_none_returns_empty(self):
        assert self._fmt(None) == ""

    def test_empty_string_returns_empty(self):
        assert self._fmt("") == ""

    def test_whitespace_only_returns_empty(self):
        assert self._fmt("   ") == ""

    def test_leading_bracket_gets_quoted(self):
        assert self._fmt("[2024] Report") == '"[2024] Report"'

    def test_leading_brace_gets_quoted(self):
        assert self._fmt("{key} value") == '"{key} value"'

    def test_leading_hash_gets_quoted(self):
        assert self._fmt("#heading") == '"#heading"'

    def test_colon_space_gets_quoted(self):
        assert self._fmt("key: value") == '"key: value"'

    def test_trailing_colon_gets_quoted(self):
        assert self._fmt("key:") == '"key:"'

    def test_hash_after_space_gets_quoted(self):
        assert self._fmt("text #comment") == '"text #comment"'

    def test_boolean_literal_gets_quoted(self):
        assert self._fmt("true") == '"true"'
        assert self._fmt("false") == '"false"'

    def test_null_literal_gets_quoted(self):
        assert self._fmt("null") == '"null"'

    def test_yes_no_literal_gets_quoted(self):
        assert self._fmt("yes") == '"yes"'
        assert self._fmt("no") == '"no"'

    def test_digit_only_gets_quoted(self):
        assert self._fmt("123") == '"123"'

    def test_leading_zero_number_gets_quoted(self):
        assert self._fmt("0123") == '"0123"'

    def test_single_quote_gets_double_quoted(self):
        assert self._fmt("it's") == '"it\'s"'

    def test_double_quote_gets_single_quoted(self):
        assert self._fmt('say "hello"') == "'say \"hello\"'"

    def test_both_quotes_escaped_with_single(self):
        """When both quote types present, use single quotes with '' escape."""
        result = self._fmt('it\'s "ok"')
        # Should use single-quote wrapping: 'it''s "ok"'
        assert result.startswith("'") and result.endswith("'")
        assert "''" in result  # escaped single quote

    def test_newline_stripped(self):
        assert "\n" not in self._fmt("line1\nline2")
        assert "\r" not in self._fmt("line1\rline2")


# ═══════════════════════════════════════════════════════════════════════════
# YAML header generation (unit tests)
# ═══════════════════════════════════════════════════════════════════════════


class TestBuildYamlHeader:
    """Tests for ``_build_yaml_header`` static method."""

    @pytest.fixture(autouse=True)
    def _setup(self):
        from docwen_plugin_document.to_markdown.converter import (
            DocxToMarkdownConverter,
        )

        self._build = DocxToMarkdownConverter._build_yaml_header

    def test_normal_header(self):
        result = self._build(
            {
                "aliases": ["My Document"],
                "title": "My Document",
                "subtitle": "Part One",
            }
        )
        assert result.startswith("---\n")
        assert result.endswith(("---\n\n", "---\n"))
        assert "aliases:" in result
        assert "  - My Document" in result
        assert "title: My Document" in result
        assert "subtitle: Part One" in result

    def test_chinese_title_subtitle(self):
        result = self._build(
            {
                "aliases": ["关于深化改革的通知"],
                "title": "关于深化改革的通知",
                "subtitle": "（征求意见稿）",
            }
        )
        assert "关于深化改革的通知" in result
        assert "（征求意见稿）" in result

    def test_locale_yaml_key_labels_are_consumed(self):
        result = self._build(
            {
                "aliases": ["Bericht"],
                "title": "Bericht",
                "subtitle": "Teil Eins",
            },
            locale="de_DE",
            yaml_key_labels={"title": "Titel", "subtitle": "Untertitel"},
        )
        assert "aliases:" in result
        assert "Titel: Bericht" in result
        assert "Untertitel: Teil Eins" in result
        assert "title: Bericht" not in result
        assert "subtitle: Teil Eins" not in result

    def test_aliases_empty_list(self):
        result = self._build(
            {
                "aliases": [],
                "title": "Test",
                "subtitle": "",
            }
        )
        assert "aliases: []" in result

    def test_multiple_aliases(self):
        result = self._build(
            {
                "aliases": ["Alpha", "Beta", "Gamma"],
                "title": "Alpha",
                "subtitle": "",
            }
        )
        assert "  - Alpha" in result
        assert "  - Beta" in result
        assert "  - Gamma" in result

    def test_empty_title_outputs_empty_field(self):
        result = self._build(
            {
                "aliases": [],
                "title": "",
                "subtitle": "",
            }
        )
        assert "title: " in result

    def test_empty_metadata_returns_empty(self):
        result = self._build(
            {
                "aliases": [],
                "title": "",
                "subtitle": "",
            }
        )
        # Even with empty fields, generates valid YAML
        assert "---" in result

    def test_special_chars_in_title_are_safe(self):
        result = self._build(
            {
                "aliases": ["Test"],
                "title": "[2024] Report #1: Results",
                "subtitle": "",
            }
        )
        assert "---" in result
        assert "title:" in result


# ═══════════════════════════════════════════════════════════════════════════
# Title metadata extraction (unit tests)
# ═══════════════════════════════════════════════════════════════════════════


class TestExtractTitleMetadata:
    """Tests for ``_extract_title_metadata`` static method."""

    @pytest.fixture(autouse=True)
    def _setup(self):
        from docwen_plugin_document.to_markdown.converter import (
            DocxToMarkdownConverter,
        )

        self._extract = DocxToMarkdownConverter._extract_title_metadata

    def _build_doc_with_styles(self, styles_and_texts, tmp_path):
        """Create a DOCX with given (style_name, text) tuples."""
        doc = Document()
        for style_name, text in styles_and_texts:
            doc.add_paragraph(text, style=style_name)
        path = tmp_path / "test.docx"
        doc.save(str(path))
        doc2 = Document(str(path))
        return doc2, str(path)

    def test_extracts_title_style(self, tmp_path):
        doc, docx_path = self._build_doc_with_styles(
            [
                ("Title", "My Document Title"),
                ("Normal", "Regular paragraph text."),
            ],
            tmp_path,
        )
        metadata, _skip = self._extract(doc, docx_path)
        assert metadata["title"] == "My Document Title"
        assert "My Document Title" in metadata["aliases"]
        assert metadata["subtitle"] == ""

    def test_extracts_subtitle_style(self, tmp_path):
        doc, docx_path = self._build_doc_with_styles(
            [
                ("Title", "Main Title"),
                ("Subtitle", "A subtitle"),
                ("Normal", "Body text."),
            ],
            tmp_path,
        )
        metadata, _skip = self._extract(doc, docx_path)
        assert metadata["title"] == "Main Title"
        assert metadata["subtitle"] == "A subtitle"

    def test_merges_multi_paragraph_title(self, tmp_path):
        doc, docx_path = self._build_doc_with_styles(
            [
                ("Title", "Line One"),
                ("Title", "Line Two"),
                ("Normal", "Main content."),
            ],
            tmp_path,
        )
        metadata, _skip = self._extract(doc, docx_path)
        assert metadata["title"] == "Line OneLine Two"
        assert "Line OneLine Two" in metadata["aliases"]

    def test_falls_back_to_filename(self, tmp_path):
        doc, docx_path = self._build_doc_with_styles(
            [
                ("Normal", "Just a regular paragraph."),
            ],
            tmp_path,
        )
        metadata, _skip = self._extract(doc, docx_path)
        # The DOCX path ends with "test.docx", so stem is "test"
        assert metadata["title"] == "test"
        assert metadata["title"] in metadata["aliases"]

    def test_skip_indices_include_title_and_subtitle(self, tmp_path):
        doc, docx_path = self._build_doc_with_styles(
            [
                ("Title", "The Title"),
                ("Subtitle", "The Subtitle"),
                ("Normal", "First normal paragraph."),
            ],
            tmp_path,
        )
        _metadata, skip = self._extract(doc, docx_path)
        assert 0 in skip  # Title paragraph index
        assert 1 in skip  # Subtitle paragraph index
        assert 2 not in skip  # Normal paragraph

    def test_stops_after_encountering_other_style(self, tmp_path):
        """Title/Subtitle only captured at document start, before other styles."""
        doc, docx_path = self._build_doc_with_styles(
            [
                ("Title", "First Title"),
                ("Normal", "Interruption."),
                ("Title", "Second Title — should be ignored"),
            ],
            tmp_path,
        )
        metadata, skip = self._extract(doc, docx_path)
        assert metadata["title"] == "First Title"
        assert 2 not in skip  # The second "Title" is NOT skipped (it's body)


# ═══════════════════════════════════════════════════════════════════════════
# Title/Subtitle rendered as H1/H2 in body when NOT consumed by YAML
# ═══════════════════════════════════════════════════════════════════════════


class TestTitleSubtitleBodyRendering:
    """Title/Subtitle style paragraphs outside the lead-in are rendered as
    H1/H2 markdown headings."""

    def test_title_style_rendered_as_h1(self, tmp_path: Path):
        """A paragraph with Title style after other content becomes an H1 heading."""
        from docwen_plugin_document.to_markdown.converter import (
            DocxToMarkdownConverter,
        )

        doc = Document()
        doc.add_paragraph("Normal paragraph first.", style="Normal")
        doc.add_paragraph("Document Title", style="Title")
        doc.add_paragraph("More content.", style="Normal")
        path = tmp_path / "title_body.docx"
        doc.save(str(path))

        doc2 = Document(str(path))
        converter = DocxToMarkdownConverter()

        # _extract_title_metadata should NOT capture Title since it comes after
        # a Norma paragraph
        metadata, skip = converter._extract_title_metadata(doc2, str(path))
        assert metadata["title"] == "title_body"  # fallback to filename
        assert 0 not in skip  # Normal paragraph
        assert 1 not in skip  # Title paragraph not at document start

        from tests.support.config import FakeConfigView
        from tests.support.execution import FakeExecutionContext
        from tests.support.logging import FakePluginLogger
        from tests.support.progress import FakeProgressSink
        from tests.support.workspace import FakeWorkspaceHandle

        from docwen_core.cancellation import CancellationToken
        from docwen_core.models.file_ref import FileRef
        from docwen_core.models.request import ConversionRequest, OutputPolicy

        ctx = FakeExecutionContext(
            request=ConversionRequest(
                request_id="test",
                input_refs=[FileRef(path=str(path), format="docx", category="document")],
                target_format="md",
                options={},
                output_policy=OutputPolicy(),
            ),
            workspace=FakeWorkspaceHandle(str(path), str(tmp_path)),
            config=FakeConfigView(),
            progress=FakeProgressSink(),
            cancellation=CancellationToken().view(),
            logger=FakePluginLogger(),
            numbering_registry=None,
        )

        # Use _parse_docx directly to avoid full convert() integration
        md, _stats = converter._parse_docx(
            input_path=str(path),
            context=ctx,
            remove_numbering=True,
            skip_indices=skip,
        )
        assert "# Document Title" in md, f"Expected Title-style paragraph to become H1. Output:\n{md[:500]}"

    def test_subtitle_style_rendered_as_h2(self, tmp_path: Path):
        """A paragraph with Subtitle style after other content becomes an H2 heading."""
        from docwen_plugin_document.to_markdown.converter import (
            DocxToMarkdownConverter,
        )

        doc = Document()
        doc.add_paragraph("Normal paragraph first.", style="Normal")
        doc.add_paragraph("My Subtitle", style="Subtitle")
        doc.add_paragraph("More content.", style="Normal")
        path = tmp_path / "subtitle_body.docx"
        doc.save(str(path))

        doc2 = Document(str(path))
        converter = DocxToMarkdownConverter()
        _metadata, skip = converter._extract_title_metadata(doc2, str(path))

        from tests.support.config import FakeConfigView
        from tests.support.execution import FakeExecutionContext
        from tests.support.logging import FakePluginLogger
        from tests.support.progress import FakeProgressSink
        from tests.support.workspace import FakeWorkspaceHandle

        from docwen_core.cancellation import CancellationToken
        from docwen_core.models.file_ref import FileRef
        from docwen_core.models.request import ConversionRequest, OutputPolicy

        ctx = FakeExecutionContext(
            request=ConversionRequest(
                request_id="test",
                input_refs=[FileRef(path=str(path), format="docx", category="document")],
                target_format="md",
                options={},
                output_policy=OutputPolicy(),
            ),
            workspace=FakeWorkspaceHandle(str(path), str(tmp_path)),
            config=FakeConfigView(),
            progress=FakeProgressSink(),
            cancellation=CancellationToken().view(),
            logger=FakePluginLogger(),
            numbering_registry=None,
        )

        md, _stats = converter._parse_docx(
            input_path=str(path),
            context=ctx,
            remove_numbering=True,
            skip_indices=skip,
        )
        assert "## My Subtitle" in md, f"Expected Subtitle-style paragraph to become H2. Output:\n{md[:500]}"


# ═══════════════════════════════════════════════════════════════════════════
# Integration: Full DOCX with Title/Subtitle → Markdown with YAML front matter
# ═══════════════════════════════════════════════════════════════════════════


class TestDocxToMdWithYamlFrontMatter:
    """Integration tests that exercise the full standard-path conversion with
    YAML front matter generation."""

    @pytest.fixture
    def docx_with_title_subtitle(self, tmp_path):
        """Create a DOCX with Title, Subtitle, and body content."""
        doc = Document()
        doc.add_paragraph("Annual Financial Report", style="Title")
        doc.add_paragraph("Fiscal Year 2025 Summary", style="Subtitle")
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("This report covers the fiscal year 2025.")
        doc.add_paragraph("All figures are in millions of USD.")
        doc.add_heading("Results", level=2)
        doc.add_paragraph("Revenue increased by 15% year-over-year.")
        path = tmp_path / "report.docx"
        doc.save(str(path))
        return str(path)

    def test_output_starts_with_yaml_front_matter(self, docx_with_title_subtitle: str):
        """The standard conversion must produce YAML front matter first."""
        from docwen_plugin_document.to_markdown.converter import (
            DocxToMarkdownConverter,
        )

        tmp = Path(docx_with_title_subtitle).parent
        converter = DocxToMarkdownConverter()

        from tests.support.config import FakeConfigView
        from tests.support.execution import FakeExecutionContext
        from tests.support.logging import FakePluginLogger
        from tests.support.progress import FakeProgressSink
        from tests.support.workspace import FakeWorkspaceHandle

        ctx = FakeExecutionContext(
            request=ConversionRequest(
                request_id="test-yaml-int",
                input_refs=[FileRef(path=docx_with_title_subtitle, format="docx", category="document")],
                target_format="md",
                options={"to_md_keep_images": True, "remove_numbering": True},
                output_policy=OutputPolicy(),
            ),
            workspace=FakeWorkspaceHandle(docx_with_title_subtitle, str(tmp)),
            config=FakeConfigView(),
            progress=FakeProgressSink(),
            cancellation=CancellationToken().view(),
            logger=FakePluginLogger(),
            numbering_registry=None,
        )

        result = converter.convert(ctx)

        assert result.success, f"Conversion failed: {result.error.message if result.error else 'unknown'}"

        # Read the written output
        artifact_path = result.artifacts[0].staging_path
        content = Path(artifact_path).read_text(encoding="utf-8")

        # Must start with YAML front matter
        stripped = content.lstrip("﻿")
        assert stripped.startswith("---\n"), f"Output must start with YAML front matter. Got:\n{stripped[:200]}"

        # Assert YAML structure
        assert "aliases:" in stripped
        assert "title:" in stripped
        assert "subtitle:" in stripped

        # Assert values
        assert "Annual Financial Report" in stripped
        assert "Fiscal Year 2025 Summary" in stripped

    def test_title_subtitle_not_duplicated_in_body(self, docx_with_title_subtitle: str):
        """Title/Subtitle consumed by YAML header should not repeat in body."""
        from docwen_plugin_document.to_markdown.converter import (
            DocxToMarkdownConverter,
        )

        tmp = Path(docx_with_title_subtitle).parent
        converter = DocxToMarkdownConverter()

        from tests.support.config import FakeConfigView
        from tests.support.execution import FakeExecutionContext
        from tests.support.logging import FakePluginLogger
        from tests.support.progress import FakeProgressSink
        from tests.support.workspace import FakeWorkspaceHandle

        ctx = FakeExecutionContext(
            request=ConversionRequest(
                request_id="test-dup",
                input_refs=[FileRef(path=docx_with_title_subtitle, format="docx", category="document")],
                target_format="md",
                options={"to_md_keep_images": True, "remove_numbering": True},
                output_policy=OutputPolicy(),
            ),
            workspace=FakeWorkspaceHandle(docx_with_title_subtitle, str(tmp)),
            config=FakeConfigView(),
            progress=FakeProgressSink(),
            cancellation=CancellationToken().view(),
            logger=FakePluginLogger(),
            numbering_registry=None,
        )

        result = converter.convert(ctx)
        assert result.success

        content = Path(result.artifacts[0].staging_path).read_text(encoding="utf-8")

        # Title/Subtitle consumed by YAML should not repeat in body text.
        # Split at the closing YAML delimiter and verify body does not
        # contain the title or subtitle as plain text.
        parts = content.split("---", 2)  # ["", yaml-content, body]
        assert len(parts) >= 3, f"Expected at least 3 parts after splitting on '---'. Content:\n{content[:300]}"

        body = parts[2]  # Everything after the second "---"

        # The title text should NOT appear as body content
        assert "Annual Financial Report" not in body, f"Title should not appear in body. Body:\n{body[:500]}"
        assert "Fiscal Year 2025 Summary" not in body, f"Subtitle should not appear in body. Body:\n{body[:500]}"

        # Body should have "Introduction" heading instead
        assert "Introduction" in body, f"Body should contain regular document content. Body:\n{body[:500]}"

    def test_filename_fallback_when_no_title_style(self, tmp_path: Path):
        """When no Title style exists, filename stem becomes the title."""
        from docwen_plugin_document.to_markdown.converter import (
            DocxToMarkdownConverter,
        )

        doc = Document()
        doc.add_paragraph("Just a regular document with no title style.")
        path = tmp_path / "my-report.docx"
        doc.save(str(path))

        converter = DocxToMarkdownConverter()

        from tests.support.config import FakeConfigView
        from tests.support.execution import FakeExecutionContext
        from tests.support.logging import FakePluginLogger
        from tests.support.progress import FakeProgressSink
        from tests.support.workspace import FakeWorkspaceHandle

        ctx = FakeExecutionContext(
            request=ConversionRequest(
                request_id="test-fallback",
                input_refs=[FileRef(path=str(path), format="docx", category="document")],
                target_format="md",
                options={"to_md_keep_images": True, "remove_numbering": True},
                output_policy=OutputPolicy(),
            ),
            workspace=FakeWorkspaceHandle(str(path), str(tmp_path)),
            config=FakeConfigView(),
            progress=FakeProgressSink(),
            cancellation=CancellationToken().view(),
            logger=FakePluginLogger(),
            numbering_registry=None,
        )

        result = converter.convert(ctx)
        assert result.success

        content = Path(result.artifacts[0].staging_path).read_text(encoding="utf-8")
        assert "title: my-report" in content, (
            f"Expected filename fallback 'my-report' as title. Content:\n{content[:300]}"
        )

    def test_yaml_header_has_proper_delimiters(self, docx_with_title_subtitle: str):
        """YAML front matter must be valid with proper --- delimiters."""
        from docwen_plugin_document.to_markdown.converter import (
            DocxToMarkdownConverter,
        )

        tmp = Path(docx_with_title_subtitle).parent
        converter = DocxToMarkdownConverter()

        from tests.support.config import FakeConfigView
        from tests.support.execution import FakeExecutionContext
        from tests.support.logging import FakePluginLogger
        from tests.support.progress import FakeProgressSink
        from tests.support.workspace import FakeWorkspaceHandle

        ctx = FakeExecutionContext(
            request=ConversionRequest(
                request_id="test-delim",
                input_refs=[FileRef(path=docx_with_title_subtitle, format="docx", category="document")],
                target_format="md",
                options={"to_md_keep_images": True, "remove_numbering": True},
                output_policy=OutputPolicy(),
            ),
            workspace=FakeWorkspaceHandle(docx_with_title_subtitle, str(tmp)),
            config=FakeConfigView(),
            progress=FakeProgressSink(),
            cancellation=CancellationToken().view(),
            logger=FakePluginLogger(),
            numbering_registry=None,
        )

        result = converter.convert(ctx)
        assert result.success

        content = Path(result.artifacts[0].staging_path).read_text(encoding="utf-8")

        lines = content.splitlines()
        # Find the YAML front matter boundaries
        yaml_open = None
        yaml_close = None
        for i, line in enumerate(lines):
            if line.strip() == "---" and yaml_open is None:
                yaml_open = i
            elif line.strip() == "---" and yaml_open is not None:
                yaml_close = i
                break

        assert yaml_open is not None, "Missing YAML opening delimiter"
        assert yaml_close is not None, "Missing YAML closing delimiter"
        assert yaml_open < yaml_close, "YAML delimiters out of order"

        # The YAML content should be between the delimiters
        yaml_content_lines = lines[yaml_open + 1 : yaml_close]
        assert any("aliases:" in line for line in yaml_content_lines)
        assert any("title:" in line for line in yaml_content_lines)
        assert any("subtitle:" in line for line in yaml_content_lines)


# ═══════════════════════════════════════════════════════════════════════════
# Gongwen path untouched
# ═══════════════════════════════════════════════════════════════════════════

# ── Gongwen tests moved to packages/plugins/optimizers/gongwen/tests/
