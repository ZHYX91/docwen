"""Request-scoped policy contracts for the production DOCX -> Markdown route.

These tests deliberately exercise ``DocumentPlugin.convert`` and inspect the
written Markdown artifact. Every DOCX formatting decision is projected from
the execution-context snapshot and remains owned by that request.
"""

from __future__ import annotations

import base64
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from threading import Barrier, Event
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from tests.support.config import FakeConfigView
from tests.support.execution import FakeExecutionContext
from tests.support.logging import FakePluginLogger
from tests.support.progress import FakeProgressSink
from tests.support.workspace import FakeWorkspaceHandle

from docwen_core.cancellation import CancellationToken
from docwen_core.models.file_ref import FileRef
from docwen_core.models.request import ConversionRequest, OutputPolicy
from docwen_plugin_document.plugin import DocumentPlugin
from docwen_plugin_document.to_markdown.converter import DocxToMarkdownConverter
from docwen_plugin_document.to_markdown.request_policy import (
    build_docx_markdown_request_policy,
)

pytestmark = pytest.mark.unit

_TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=="
)

_FOOTNOTES_XML = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>
  <w:footnote w:id="1"><w:p><w:r><w:footnoteRef/></w:r><w:r><w:t>request note</w:t></w:r></w:p></w:footnote>
</w:footnotes>
"""


def _request_policy(
    *,
    preserve_formatting: bool = True,
    page_break: str = "___",
) -> dict[str, Any]:
    """Return an execution-context snapshot with every DOCX policy owner."""
    return {
        "conversion": {
            "docx_to_md": {
                "preserve_formatting": preserve_formatting,
                "preserve_heading_formatting": True,
                "preserve_table_header_formatting": True,
            },
            "syntax": {
                "bold": "underscore",
                "italic": "underscore",
                "strikethrough": "html",
                "highlight": "html",
                "superscript": "extended",
                "subscript": "extended",
                "unordered_list": "plus",
                "indent_spaces": 2,
            },
            "horizontal_rule": {
                "docx_to_md": {
                    "page_break": page_break,
                    "section_break": "---",
                    "horizontal_rule": "***",
                }
            },
            # Compatibility owner consumed by MarkdownExportSemantics.
            "table_merge_export_strategy": "empty",
        },
        "document": {
            "to_md_table_merge_export_strategy": "empty",
            "style": {
                "code": {
                    "docx_to_md": {
                        "paragraph_style_aliases": ["Policy Zeta"],
                        "fuzzy_match_enabled": False,
                    }
                },
                "quote": {
                    "docx_to_md": {
                        "level_style_aliases": {"Request Aside": 2},
                        "paragraph_style_aliases": [],
                        "fuzzy_match_enabled": False,
                    }
                },
            },
        },
        # The request-owned export section is the sole image/OCR mode owner.
        "export": {
            "to_md_image_extraction_mode": "omit",
            "to_md_ocr_placement_mode": "main_md",
        },
        "link": {
            "format": {
                "image_link_style": "markdown_embed",
                "md_file_link_style": "markdown_link",
            }
        },
    }


def _context(
    tmp_path: Path,
    input_path: Path,
    *,
    request_id: str,
    config: dict[str, Any],
    options: dict[str, Any] | None = None,
    ocr_blockquote_title: str = "",
) -> FakeExecutionContext:
    staging = tmp_path / f"staging-{request_id}"
    staging.mkdir()
    return FakeExecutionContext(
        request=ConversionRequest(
            request_id=request_id,
            input_refs=[
                FileRef(
                    path=str(input_path),
                    format="docx",
                    category="document",
                    size_bytes=input_path.stat().st_size,
                )
            ],
            target_format="md",
            options=options or {},
            output_policy=OutputPolicy(),
            config_snapshot=dict(config),
        ),
        workspace=FakeWorkspaceHandle(str(input_path), str(staging)),
        config=FakeConfigView(config),
        progress=FakeProgressSink(),
        cancellation=CancellationToken().view(),
        logger=FakePluginLogger(),
        numbering_registry=None,
        ocr_blockquote_title=ocr_blockquote_title,
    )


def _markdown_from_result(result: Any) -> str:
    assert result.success, result.error
    primary = next(artifact for artifact in result.artifacts if artifact.is_primary)
    return Path(primary.staging_path).read_text(encoding="utf-8")


class _ParagraphStyleProbe:
    def __init__(self, name: str) -> None:
        self.style = type("_Style", (), {"name": name})()


def _add_footnote_part(path: Path, *, note_text: str = "request note") -> None:
    """Add the smallest note part needed by the converter's ZIP fallback."""
    rewritten = path.with_name(f"{path.stem}-with-note.docx")
    footnotes_xml = _FOOTNOTES_XML.replace(b"request note", note_text.encode("utf-8"))
    with ZipFile(path, "r") as source, ZipFile(rewritten, "w", ZIP_DEFLATED) as target:
        for item in source.infolist():
            target.writestr(item, source.read(item.filename))
        target.writestr("word/footnotes.xml", footnotes_xml)
    rewritten.replace(path)


def _add_malformed_note_part(path: Path, part_name: str) -> None:
    """Add a present but malformed note part for loss-semantics testing."""
    rewritten = path.with_name(f"{path.stem}-with-malformed-note.docx")
    with ZipFile(path, "r") as source, ZipFile(rewritten, "w", ZIP_DEFLATED) as target:
        for item in source.infolist():
            target.writestr(item, source.read(item.filename))
        target.writestr(f"word/{part_name}.xml", f"<w:{part_name}><broken".encode())
    rewritten.replace(path)


def _build_malformed_note_probe_docx(tmp_path: Path, note_type: str) -> Path:
    """Build a readable DOCX whose referenced note definitions are corrupt."""
    part_name = f"{note_type}s"
    doc = Document()
    paragraph = doc.add_paragraph("body note")
    reference_run = paragraph.add_run()
    reference = OxmlElement(f"w:{note_type}Reference")
    reference.set(qn("w:id"), "1")
    reference_run._r.append(reference)  # pyright: ignore[reportPrivateUsage]
    path = tmp_path / f"malformed-{note_type}.docx"
    doc.save(str(path))
    _add_malformed_note_part(path, part_name)
    return path


def _build_policy_probe_docx(tmp_path: Path, *, name: str = "policy-probe.docx") -> Path:
    """Build a small real DOCX covering all observable policy families."""
    image_path = tmp_path / f"{Path(name).stem}.png"
    image_path.write_bytes(_TINY_PNG)

    doc = Document()
    doc.styles.add_style("Policy Zeta", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("Request Aside", WD_STYLE_TYPE.PARAGRAPH)

    inline = doc.add_paragraph()
    bold = inline.add_run("request bold")
    bold.bold = True
    inline.add_run(" and ")
    italic = inline.add_run("request italic")
    italic.italic = True

    heading = doc.add_heading(level=2)
    heading_bold = heading.add_run("request heading")
    heading_bold.bold = True

    list_para = doc.add_paragraph("nested request item")
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "1")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "2")
    num_pr.extend((ilvl, num_id))
    list_para._p.get_or_add_pPr().append(num_pr)  # pyright: ignore[reportPrivateUsage]

    break_para = doc.add_paragraph("before request break")
    break_para.add_run().add_break(WD_BREAK.PAGE)
    break_para.add_run("after request break")

    doc.add_paragraph("print('request policy')", style="Policy Zeta")
    doc.add_paragraph("request aside", style="Request Aside")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = ""
    header = table.cell(0, 0).paragraphs[0].add_run("Merged Header")
    header.bold = True
    table.cell(0, 1).text = "Top"
    table.cell(1, 1).text = "Bottom"
    table.cell(0, 0).merge(table.cell(1, 0))

    image_para = doc.add_paragraph("request image ")
    image_para.add_run().add_picture(str(image_path))

    note_para = doc.add_paragraph("note anchor")
    note_ref_run = note_para.add_run()
    note_ref = OxmlElement("w:footnoteReference")
    note_ref.set(qn("w:id"), "1")
    note_ref_run._r.append(note_ref)  # pyright: ignore[reportPrivateUsage]

    path = tmp_path / name
    doc.save(str(path))
    _add_footnote_part(path)
    return path


def _build_concurrency_probe_docx(
    tmp_path: Path,
    *,
    name: str,
    note_text: str | None = None,
) -> Path:
    doc = Document()
    formatted = doc.add_paragraph()
    run = formatted.add_run("thread bold")
    run.bold = True
    break_paragraph = doc.add_paragraph("thread head")
    break_paragraph.add_run().add_break(WD_BREAK.PAGE)
    break_paragraph.add_run("thread tail")
    if note_text is not None:
        note_paragraph = doc.add_paragraph("thread note")
        note_ref_run = note_paragraph.add_run()
        note_ref = OxmlElement("w:footnoteReference")
        note_ref.set(qn("w:id"), "1")
        note_ref_run._r.append(note_ref)  # pyright: ignore[reportPrivateUsage]
    path = tmp_path / name
    doc.save(str(path))
    if note_text is not None:
        _add_footnote_part(path, note_text=note_text)
    return path


def _build_rich_page_break_probe_docx(tmp_path: Path, *, name: str) -> Path:
    doc = Document()
    paragraph = doc.add_paragraph()
    before = paragraph.add_run("rich before")
    before.bold = True
    before.add_break(WD_BREAK.PAGE)
    after = paragraph.add_run("rich after")
    after.italic = True
    note_run = paragraph.add_run()
    note_ref = OxmlElement("w:footnoteReference")
    note_ref.set(qn("w:id"), "1")
    note_run._r.append(note_ref)  # pyright: ignore[reportPrivateUsage]

    path = tmp_path / name
    doc.save(str(path))
    _add_footnote_part(path, note_text="page break note")
    return path


def _build_formula_probe_docx(tmp_path: Path) -> Path:
    doc = Document()
    paragraph = doc.add_paragraph()
    bold = paragraph.add_run("formula bold")
    bold.bold = True
    paragraph.add_run(" ")
    omath = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = "x"
    math_run.append(math_text)
    omath.append(math_run)
    paragraph._p.append(omath)  # pyright: ignore[reportPrivateUsage]
    paragraph.add_run(" ")
    italic = paragraph.add_run("formula italic")
    italic.italic = True
    path = tmp_path / "formula-policy.docx"
    doc.save(str(path))
    return path


def _build_recursive_policy_probe_docx(tmp_path: Path) -> Path:
    doc = Document()

    sdt_paragraph = doc.add_paragraph()
    sdt_bold = sdt_paragraph.add_run("sdt bold")
    sdt_bold.bold = True
    body = doc._element.body  # pyright: ignore[reportPrivateUsage]
    body.remove(sdt_paragraph._p)  # pyright: ignore[reportPrivateUsage]
    sdt = OxmlElement("w:sdt")
    sdt_content = OxmlElement("w:sdtContent")
    sdt_content.append(sdt_paragraph._p)  # pyright: ignore[reportPrivateUsage]
    sdt.append(sdt_content)
    body.insert(len(body) - 1, sdt)

    outer = doc.add_table(rows=1, cols=1)
    nested = outer.cell(0, 0).add_table(rows=1, cols=1)
    nested_paragraph = nested.cell(0, 0).paragraphs[0]
    nested_bold = nested_paragraph.add_run("nested bold")
    nested_bold.bold = True

    path = tmp_path / "recursive-policy.docx"
    doc.save(str(path))
    return path


__all__ = (
    "Any",
    "Barrier",
    "Document",
    "DocumentPlugin",
    "DocxToMarkdownConverter",
    "Event",
    "FakeProgressSink",
    "Path",
    "ThreadPoolExecutor",
    "_ParagraphStyleProbe",
    "_build_concurrency_probe_docx",
    "_build_formula_probe_docx",
    "_build_malformed_note_probe_docx",
    "_build_policy_probe_docx",
    "_build_recursive_policy_probe_docx",
    "_build_rich_page_break_probe_docx",
    "_context",
    "_markdown_from_result",
    "_request_policy",
    "build_docx_markdown_request_policy",
    "pytest",
    "pytestmark",
)
