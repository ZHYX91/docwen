from pathlib import Path

import pytest
from docx import Document

from docwen_core.docx_parsing.image_extraction import (
    _content_type_to_ext,
    extract_images_from_element,
)

pytestmark = pytest.mark.unit


def test_content_type_filters_office_vector_formats():
    assert _content_type_to_ext("image/png") == ".png"
    assert _content_type_to_ext("image/jpeg") == ".jpg"
    assert _content_type_to_ext("image/x-emf") is None
    assert _content_type_to_ext("image/x-wmf") is None


def test_extract_images_from_element_accepts_table_element(tmp_path: Path):
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    para = table.cell(0, 0).paragraphs[0]
    run = para.add_run()
    png = tmp_path / "tiny.png"
    png.write_bytes(
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR"
        b"\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde"
        b"\x00\x00\x00\x0cIDATx\x9cc```\x00\x00\x00\x04\x00\x01\xf6\x178U"
        b"\x00\x00\x00\x00IEND\xaeB`\x82"
    )
    run.add_picture(str(png))

    infos = extract_images_from_element(
        table._tbl,
        doc.part.related_parts,
        str(tmp_path),
        name_prefix="test-image",
    )

    assert len(infos) == 1
    assert infos[0].path.endswith(".png")
    assert Path(infos[0].path).exists()


def test_partial_image_write_removes_every_file_from_the_failed_extraction(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    doc = Document()
    paragraph = doc.add_paragraph()
    png = tmp_path / "tiny.png"
    png.write_bytes(
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR"
        b"\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde"
        b"\x00\x00\x00\x0cIDATx\x9cc```\x00\x00\x00\x04\x00\x01\xf6\x178U"
        b"\x00\x00\x00\x00IEND\xaeB`\x82"
    )
    paragraph.add_run().add_picture(str(png))
    paragraph.add_run().add_picture(str(png))
    real_write_bytes = Path.write_bytes
    writes = 0

    def fail_second_write(path: Path, content: bytes) -> int:
        nonlocal writes
        if path.name.startswith("partial-image"):
            writes += 1
            if writes == 2:
                real_write_bytes(path, b"partial")
                raise OSError("synthetic partial image write")
        return real_write_bytes(path, content)

    monkeypatch.setattr(Path, "write_bytes", fail_second_write)
    with pytest.raises(OSError, match="partial image write"):
        extract_images_from_element(
            paragraph._p,
            doc.part.related_parts,
            str(tmp_path),
            name_prefix="partial-image",
        )

    assert list(tmp_path.glob("partial-image*")) == []
