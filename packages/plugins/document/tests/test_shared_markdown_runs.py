import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docwen_core.docx_parsing.format_features import (
    DocxMarkdownSyntaxConfig,
    StyleDetectorConfig,
)
from docwen_plugin_document.shared.markdown_runs import (
    append_formatted_run_text,
    render_paragraph_runs,
    render_paragraph_runs_split_on_page_breaks,
    resolve_hyperlink_target,
)

pytestmark = pytest.mark.unit


def test_resolve_hyperlink_target_returns_none_for_missing_rel():
    doc = Document()
    para = doc.add_paragraph("plain")

    class Link:
        def get(self, key, default=None):
            return "rId999"

    assert resolve_hyperlink_target(para, Link()) is None


def test_resolve_hyperlink_target_resolves_real_url():
    doc = Document()
    para = doc.add_paragraph()
    rel_id = para.part.relate_to(
        "https://example.com/path",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    # Simulate a hyperlink XML element with r:id=rel_id
    from lxml import etree

    hl = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink")
    hl.set(
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id",
        rel_id,
    )

    result = resolve_hyperlink_target(para, hl)
    assert result == "https://example.com/path"


@pytest.mark.parametrize(
    ("apply_format", "expected"),
    [
        (lambda run: setattr(run, "bold", True), "__value__"),
        (lambda run: setattr(run, "italic", True), "_value_"),
        (lambda run: setattr(run, "underline", True), "<u>value</u>"),
        (lambda run: setattr(run.font, "strike", True), "<del>value</del>"),
        (lambda run: setattr(run.font, "superscript", True), "^value^"),
        (lambda run: setattr(run.font, "subscript", True), "~value~"),
        (
            lambda run: setattr(run.font, "highlight_color", WD_COLOR_INDEX.YELLOW),
            "<mark>value</mark>",
        ),
    ],
)
def test_append_formatted_run_text_honors_explicit_syntax_config(apply_format, expected):
    doc = Document()
    run = doc.add_paragraph().add_run("value")
    apply_format(run)
    syntax = DocxMarkdownSyntaxConfig(
        bold="underscore",
        italic="underscore",
        strikethrough="html",
        highlight="html",
        superscript="extended",
        subscript="extended",
    )
    parts: list[str] = []

    append_formatted_run_text(parts, run.text, run._r, syntax_config=syntax)

    assert "".join(parts) == expected


def test_append_formatted_run_text_preserves_gray_shading_as_inline_code():
    doc = Document()
    run = doc.add_paragraph().add_run("value")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "D9D9D9")
    run._r.get_or_add_rPr().append(shading)
    parts: list[str] = []

    append_formatted_run_text(
        parts,
        run.text,
        run._r,
        syntax_config=DocxMarkdownSyntaxConfig(),
    )

    assert "".join(parts) == "`value`"


def test_render_paragraph_runs_coalesces_adjacent_runs_through_shared_helper():
    doc = Document()
    paragraph = doc.add_paragraph()
    first = paragraph.add_run("Hello ")
    first.bold = True
    second = paragraph.add_run("World")
    second.bold = True

    rendered = render_paragraph_runs(
        paragraph,
        syntax_config=DocxMarkdownSyntaxConfig(bold="underscore"),
    )

    assert rendered == "__Hello World__"


def test_page_break_segments_preserve_formatting_inside_the_same_run():
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Before")
    run.bold = True
    run.add_break(WD_BREAK.PAGE)
    run.add_text("After")

    rendered = render_paragraph_runs_split_on_page_breaks(
        paragraph,
        syntax_config=DocxMarkdownSyntaxConfig(),
    )

    assert rendered == ["**Before**", "**After**"]


def test_page_break_segments_preserve_formula_callback_order():
    doc = Document()
    paragraph = doc.add_paragraph("Before")
    formula = OxmlElement("m:oMath")
    formula_text = OxmlElement("m:t")
    formula_text.text = "x"
    formula.append(formula_text)
    paragraph._p.append(formula)
    break_run = paragraph.add_run()
    break_run.add_break(WD_BREAK.PAGE)
    paragraph.add_run("After")

    rendered = render_paragraph_runs_split_on_page_breaks(
        paragraph,
        syntax_config=DocxMarkdownSyntaxConfig(),
        math_renderer=lambda _element: "$x$",
    )

    assert rendered == ["Before$x$", "After"]


@pytest.mark.parametrize("wrapper_tag", ("w:ins", "w:moveTo"))
def test_page_break_segments_include_accepted_revision_text(wrapper_tag: str):
    doc = Document()
    paragraph = doc.add_paragraph()
    wrapper = OxmlElement(wrapper_tag)
    run = OxmlElement("w:r")
    before = OxmlElement("w:t")
    before.text = "Before"
    run.append(before)
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")
    run.append(page_break)
    after = OxmlElement("w:t")
    after.text = "After"
    run.append(after)
    wrapper.append(run)
    paragraph._p.append(wrapper)

    rendered = render_paragraph_runs_split_on_page_breaks(
        paragraph,
        syntax_config=DocxMarkdownSyntaxConfig(),
    )

    assert rendered == ["Before", "After"]


def test_page_break_segments_retain_empty_segment_between_adjacent_breaks():
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Before")
    run.add_break(WD_BREAK.PAGE)
    run.add_break(WD_BREAK.PAGE)
    run.add_text("After")

    rendered = render_paragraph_runs_split_on_page_breaks(
        paragraph,
        syntax_config=DocxMarkdownSyntaxConfig(),
    )

    assert rendered == ["Before", "", "After"]


def test_page_break_segments_preserve_note_references():
    class NoteExtractor:
        def get_reference_text(self, note_type, note_id):
            return f"[^{note_type}-{note_id}]"

    doc = Document()
    paragraph = doc.add_paragraph("Before")
    reference_run = paragraph.add_run()
    reference = OxmlElement("w:footnoteReference")
    reference.set(qn("w:id"), "7")
    reference_run._r.append(reference)
    reference_run.add_break(WD_BREAK.PAGE)
    paragraph.add_run("After")

    rendered = render_paragraph_runs_split_on_page_breaks(
        paragraph,
        note_extractor=NoteExtractor(),
        syntax_config=DocxMarkdownSyntaxConfig(),
    )

    assert rendered == ["Before[^footnote-7]", "After"]


def test_page_break_segments_wrap_each_hyperlink_side_independently():
    doc = Document()
    paragraph = doc.add_paragraph()
    rel_id = paragraph.part.relate_to(
        "https://example.com/page-aware",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run_properties.append(OxmlElement("w:b"))
    run.append(run_properties)
    before = OxmlElement("w:t")
    before.text = "Before"
    run.append(before)
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")
    run.append(page_break)
    after = OxmlElement("w:t")
    after.text = "After"
    run.append(after)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)

    rendered = render_paragraph_runs_split_on_page_breaks(
        paragraph,
        syntax_config=DocxMarkdownSyntaxConfig(),
    )

    assert rendered == [
        "[**Before**](https://example.com/page-aware)",
        "[**After**](https://example.com/page-aware)",
    ]


def test_regular_paragraph_renderer_keeps_page_break_as_line_break():
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Before")
    run.add_break(WD_BREAK.PAGE)
    run.add_text("After")

    rendered = render_paragraph_runs(
        paragraph,
        syntax_config=DocxMarkdownSyntaxConfig(),
    )

    assert rendered == "Before\nAfter"


@pytest.mark.parametrize(
    ("style_name", "config_field", "expected"),
    [("HTML Code", "code", "`value`"), ("Quote Char", "quote", "value")],
)
def test_render_paragraph_runs_consumes_character_style_aliases(style_name, config_field, expected):
    doc = Document()
    try:
        doc.styles[style_name]
    except KeyError:
        doc.styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("value")
    run.style = style_name
    config = StyleDetectorConfig(
        code_character_style_names=frozenset({style_name}) if config_field == "code" else frozenset(),
        quote_character_style_names=frozenset({style_name}) if config_field == "quote" else frozenset(),
    )

    rendered = render_paragraph_runs(
        paragraph,
        syntax_config=DocxMarkdownSyntaxConfig(),
        style_detector_config=config,
    )

    assert rendered == expected


def test_render_paragraph_runs_recognizes_associated_character_style():
    doc = Document()
    style_name = "Runtime Code Alias Char"
    doc.styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("value")
    run.style = style_name

    rendered = render_paragraph_runs(
        paragraph,
        syntax_config=DocxMarkdownSyntaxConfig(),
        style_detector_config=StyleDetectorConfig(code_block_style_fragments=("Runtime Code Alias",)),
    )

    assert rendered == "`value`"


@pytest.mark.parametrize(
    ("tag", "value"),
    [("w:b", "0"), ("w:i", "false"), ("w:strike", "off"), ("w:u", "none")],
)
def test_explicitly_disabled_run_properties_do_not_emit_markdown(tag, value):
    doc = Document()
    run = doc.add_paragraph().add_run("value")
    property_element = OxmlElement(tag)
    property_element.set(qn("w:val"), value)
    run._r.get_or_add_rPr().append(property_element)

    rendered = render_paragraph_runs(
        doc.paragraphs[0],
        syntax_config=DocxMarkdownSyntaxConfig(),
    )

    assert rendered == "value"


@pytest.mark.parametrize(
    ("fill", "value", "config"),
    [
        ("D9D9D9", "clear", StyleDetectorConfig(wps_shading_enabled=False)),
        ("FFFFFF", "pct15", StyleDetectorConfig(word_shading_enabled=False)),
    ],
)
def test_run_shading_switches_can_disable_each_producer(fill, value, config):
    doc = Document()
    run = doc.add_paragraph().add_run("value")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), value)
    run._r.get_or_add_rPr().append(shading)

    rendered = render_paragraph_runs(
        doc.paragraphs[0],
        syntax_config=DocxMarkdownSyntaxConfig(),
        style_detector_config=config,
    )

    assert rendered == "value"
