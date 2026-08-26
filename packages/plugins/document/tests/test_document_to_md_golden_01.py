"""Focused tests split from test_document_to_md_golden.py."""

from __future__ import annotations

from ._document_to_md_golden_support import (
    _OLD_SYSTEM_COMPREHENSIVE_DOCX,
    Path,
    _build_runtime_pipeline,
    _document_node_root,
    _heading_counts,
    _load_docx_to_md_old_system_fixture,
    _run_conversion,
    os,
    pytest,
    re,
)

pytestmark = [pytest.mark.golden, pytest.mark.contract]


class TestDocxToMdGolden:
    """Golden parity tests for ROUTE-DOC-001: document → md.

    Uses a programmatically-generated sample DOCX with known content
    (headings, paragraphs with formatting, tables) for structural
    validation, plus a real template for integration smoke testing.
    """

    @pytest.fixture
    def pipeline(self):
        """Build the full runtime pipeline with the real DocumentPlugin."""
        plugin, task_mgr, ws_mgr, ws_root = _build_runtime_pipeline()
        yield plugin, task_mgr, ws_mgr
        ws_mgr.cleanup_all()
        import shutil

        shutil.rmtree(ws_root, ignore_errors=True)

    @staticmethod
    def _markdown_table_cells(line: str) -> list[str]:
        row = line.strip()
        if row.startswith("|"):
            row = row[1:]
        if row.endswith("|"):
            row = row[:-1]
        return [cell.replace(r"\|", "|").strip() for cell in re.split(r"(?<!\\)\|", row)]

    @staticmethod
    def _verify_tables_well_formed(content: str) -> None:
        """Verify that Markdown tables in content are well-formed and have data rows."""
        lines = content.splitlines()
        in_table = False
        separator_seen = False
        data_rows_seen = 0
        column_count = 0

        for line in lines:
            stripped = line.strip()
            is_table_line = stripped.startswith("|") and stripped.endswith("|")

            if is_table_line:
                if not in_table:
                    in_table = True
                    separator_seen = False
                    data_rows_seen = 0
                    column_count = len(TestDocxToMdGolden._markdown_table_cells(stripped))
                else:
                    if not separator_seen:
                        separator_seen = True
                        parts = TestDocxToMdGolden._markdown_table_cells(stripped)
                        assert all(set(p) <= {"-", ":", " "} and "-" in p for p in parts), (
                            f"Malformed table separator: {stripped}"
                        )
                    else:
                        cols = len(TestDocxToMdGolden._markdown_table_cells(stripped))
                        assert cols == column_count, f"Table column count mismatch: expected {column_count}, got {cols}"
                        data_rows_seen += 1
            else:
                if in_table and stripped == "":
                    in_table = False
                    separator_seen = False
                    assert data_rows_seen >= 1, f"Table should have at least 1 data row, got {data_rows_seen}"

    def test_conversion_succeeds_with_sample_docx(self, pipeline, sample_docx_path, tmp_path) -> None:
        """GOLDEN-002: The plugin must successfully convert a DOCX to Markdown."""
        _plugin, task_mgr, _ws_mgr = pipeline

        output_dir = tmp_path / "output"
        output_dir.mkdir()

        result = _run_conversion(task_mgr, sample_docx_path, output_dir)

        assert result.success, f"Conversion failed: {result.error.message if result.error else 'unknown'}"
        assert len(result.artifacts) >= 1
        assert result.artifacts[0].is_primary is True
        assert result.artifacts[0].kind == "primary"

        # Verify output file at final location
        final_path = result.artifacts[0].staging_path
        assert os.path.isfile(final_path), f"Output file not found: {final_path}"

        # Content assertions: output must contain expected text
        content = Path(final_path).read_text(encoding="utf-8")
        assert len(content.strip()) > 0, "Output should be non-empty"
        assert "# Test Document" in content, "Should contain h1 heading text"
        assert "test document used for golden parity testing" in content.lower()
        assert "Section One" in content, "Should contain h2 heading text"

    def test_docx_to_md_matches_old_system_semantic_fixture(self, pipeline, sample_docx_path, tmp_path) -> None:
        """GOLDEN-002: Current output preserves old-system DOCX→MD core semantics."""
        _plugin, task_mgr, ws_mgr = pipeline
        fixture = _load_docx_to_md_old_system_fixture()
        assert fixture["golden_id"] == "GOLDEN-002"

        output_dir = tmp_path / "output_old_system_semantics"
        output_dir.mkdir()

        result = _run_conversion(
            task_mgr,
            sample_docx_path,
            output_dir,
            to_md_keep_images=False,
            remove_numbering=True,
            yaml_key_labels={"title": "Titel", "subtitle": "Untertitel"},
        )
        assert result.success, f"Conversion failed: {result.error.message if result.error else 'unknown'}"
        assert len(result.artifacts) == 2

        artifact = result.artifacts[0]
        artifact_path = Path(artifact.staging_path)
        node_root = _document_node_root(artifact_path, output_dir)
        assert artifact_path.name == f"{node_root.name}.md"
        assert artifact_path.exists()
        assert any(d.code == "FINALIZER_DONE" for d in result.diagnostics)

        content = artifact_path.read_text(encoding="utf-8")
        assert str(Path(ws_mgr.root_dir)) not in content
        assert f"Titel: {sample_docx_path.stem}" in content
        assert "Untertitel:" in content
        assert f"title: {sample_docx_path.stem}" not in content
        assert "subtitle:" not in content

        for expected in fixture["expected_markdown_semantics"]["contains"]:
            assert expected in content
        assert _heading_counts(content) == fixture["expected_markdown_semantics"]["heading_counts"]

        current = fixture["projects"]["docwen-current"]
        assert artifact.media_type == current["artifact_media_type"]
        assert artifact.metadata["source_suggested_name"] == current["suggested_name"]
        for key, value in current["metadata"].items():
            assert artifact.metadata[key] == value

        for project_name in ("docwen-ref-tk", "docwen-ref-pyside6", "docwen-current"):
            project = fixture["projects"][project_name]
            assert project["success"] is True
            assert project["markdown_contains_all_expected_semantics"] is True

    def test_old_system_comprehensive_docx_preserves_rules_and_formula_block(self, pipeline, tmp_path) -> None:
        """The checked-in old DOCX artifact keeps its Word rules and display math."""
        _plugin, task_mgr, _ws_mgr = pipeline
        output_dir = tmp_path / "old_system_comprehensive"
        output_dir.mkdir()

        result = _run_conversion(
            task_mgr,
            _OLD_SYSTEM_COMPREHENSIVE_DOCX,
            output_dir,
            remove_numbering=False,
            yaml_key_labels={"aliases": "aliases", "title": "标题", "subtitle": "副标题"},
        )

        assert result.success, result.error.message if result.error else "conversion failed"
        assert any(item.code == "FINALIZER_DONE" for item in result.diagnostics)
        content = Path(result.artifacts[0].staging_path).read_text(encoding="utf-8")
        assert content.splitlines().count("___") == 3
        assert "$$a^{2}+b^{2}=c^{2}$$" in content
        assert "$a^{2}+b^{2}=c^{2}$" not in content.splitlines()

    def test_body_text_outline_sentinel_stays_body_through_finalizer(self, pipeline, tmp_path) -> None:
        """Word outlineLvl=9 must not turn every Normal paragraph into an H6."""
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        source = tmp_path / "body_outline_sentinel.docx"
        document = Document()
        paragraphs = []
        for text in ("First ordinary body", "Second ordinary body"):
            paragraph = document.add_paragraph(text)
            paragraphs.append(paragraph)
            outline = OxmlElement("w:outlineLvl")
            outline.set(qn("w:val"), "9")
            paragraph._p.get_or_add_pPr().append(outline)
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), "0")
        num_pr.extend((ilvl, num_id))
        paragraphs[1]._p.get_or_add_pPr().append(num_pr)
        document.save(source)

        _plugin, task_mgr, _ws_mgr = pipeline
        output_dir = tmp_path / "body_outline_output"
        output_dir.mkdir()
        result = _run_conversion(task_mgr, source, output_dir, remove_numbering=False)

        assert result.success, result.error.message if result.error else "conversion failed"
        artifact = result.artifacts[0]
        content = Path(artifact.staging_path).read_text(encoding="utf-8")
        assert artifact.metadata["paragraph_count"] == 2
        assert artifact.metadata["heading_count"] == 0
        assert "First ordinary body" in content
        assert "Second ordinary body" in content
        assert "- Second ordinary body" not in content
        assert not any(line.startswith("#") for line in content.splitlines())
        assert any(item.code == "FINALIZER_DONE" for item in result.diagnostics)

    def test_malformed_outline_level_uses_style_fallback_with_diagnostic(self, pipeline, tmp_path) -> None:
        """Invalid outlineLvl keeps content and emits a stable warning."""
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        source = tmp_path / "malformed_outline.docx"
        document = Document()
        paragraph = document.add_paragraph("Retained body content")
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), "not-an-integer")
        paragraph._p.get_or_add_pPr().append(outline)
        document.save(source)

        _plugin, task_mgr, _ws_mgr = pipeline
        output_dir = tmp_path / "malformed_outline_output"
        output_dir.mkdir()
        result = _run_conversion(task_mgr, source, output_dir)

        assert result.success, result.error.message if result.error else "conversion failed"
        content = Path(result.artifacts[0].staging_path).read_text(encoding="utf-8")
        assert "Retained body content" in content
        diagnostic = next(item for item in result.diagnostics if item.code == "DOCX2MD-OUTLINE-FALLBACK")
        assert diagnostic.level == "warning"
        assert diagnostic.location == "document.xml body element 0"

    def test_output_is_valid_markdown(self, pipeline, sample_docx_path, tmp_path) -> None:
        """GOLDEN-002: Output must be valid Markdown with structural markers."""
        _plugin, task_mgr, _ws_mgr = pipeline

        output_dir = tmp_path / "output_md"
        output_dir.mkdir()

        result = _run_conversion(task_mgr, sample_docx_path, output_dir)
        assert result.success

        content = Path(result.artifacts[0].staging_path).read_text(encoding="utf-8")

        # Must be non-empty
        assert len(content.strip()) > 0, "Output Markdown is empty"

        # Must contain markdown headings
        assert "# " in content, (
            f"Output should contain markdown headings (e.g., '# Test Document'). Content:\n{content[:200]}"
        )

        # Must have meaningful text content
        lines = [line for line in content.splitlines() if line.strip()]
        assert len(lines) >= 5, f"Expected at least 5 non-empty lines, got {len(lines)}"

    def test_heading_structure_preserved(self, pipeline, sample_docx_path, tmp_path) -> None:
        """GOLDEN-002: Heading levels from DOCX styles must be preserved."""
        _plugin, task_mgr, _ws_mgr = pipeline

        output_dir = tmp_path / "output_headings"
        output_dir.mkdir()

        result = _run_conversion(task_mgr, sample_docx_path, output_dir)
        assert result.success

        content = Path(result.artifacts[0].staging_path).read_text(encoding="utf-8")

        # Count headings by level
        heading_counts: dict[int, int] = {}
        for line in content.splitlines():
            stripped = line.strip()
            if stripped.startswith("#"):
                level = 0
                for ch in stripped:
                    if ch == "#":
                        level += 1
                    else:
                        break
                if 1 <= level <= 6 and level < len(stripped) and stripped[level] == " ":
                    heading_counts[level] = heading_counts.get(level, 0) + 1

        total_headings = sum(heading_counts.values())
        assert total_headings >= 2, (
            f"Expected at least 2 headings, found {total_headings}. "
            f"Headings by level: {heading_counts}\nContent:\n{content[:300]}"
        )

        # Should have "Test Document" as level-1 heading
        assert heading_counts.get(1, 0) >= 1, (
            f"Expected at least one level-1 heading (h1='Test Document'). Got: {heading_counts}"
        )

        # Should have level-2 headings
        assert heading_counts.get(2, 0) >= 2, f"Expected at least two level-2 headings. Got: {heading_counts}"

        # Verify actual heading text content (not just level counts)
        assert "# Test Document" in content, f"h1 heading text 'Test Document' not found. Content:\n{content[:300]}"
        assert "## Section One" in content, f"h2 heading text 'Section One' not found. Content:\n{content[:300]}"
        assert "## Table Section" in content, f"h2 heading text 'Table Section' not found. Content:\n{content[:300]}"

    def test_table_structure_preserved(self, pipeline, sample_docx_path, tmp_path) -> None:
        """GOLDEN-002: Tables must be preserved as well-formed Markdown tables."""
        _plugin, task_mgr, _ws_mgr = pipeline

        output_dir = tmp_path / "output_tables"
        output_dir.mkdir()

        result = _run_conversion(task_mgr, sample_docx_path, output_dir)
        assert result.success

        content = Path(result.artifacts[0].staging_path).read_text(encoding="utf-8")

        # Must have table markers
        assert "|" in content, f"Output should contain table markers. Content:\n{content[:300]}"

        # Check for table header with known column names
        assert "Name" in content, "Table header 'Name' should be preserved"
        assert "Value" in content, "Table header 'Value' should be preserved"
        assert "Description" in content, "Table header 'Description' should be preserved"

        # Check for data cells
        assert "Alpha" in content, "Table data 'Alpha' should be preserved"
        assert "Beta" in content, "Table data 'Beta' should be preserved"

        self._verify_tables_well_formed(content)

    def test_table_well_formed_helper_allows_escaped_pipes(self) -> None:
        """Markdown table validation should not count escaped pipes as columns."""
        content = "\n".join(
            [
                "| Header | Value |",
                "| --- | --- |",
                r"| A\|B | C |",
                "",
            ]
        )

        self._verify_tables_well_formed(content)

    def test_docx_table_merge_strategy_empty_matches_old_system_blank_repeat(self, pipeline, tmp_path) -> None:
        """GOLDEN-002: DOCX merged cells can render old-system blank covered cells."""
        from docx import Document

        _plugin, task_mgr, _ws_mgr = pipeline

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "H"
        table.cell(0, 1).text = ""
        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(1, 0).text = "A"
        table.cell(1, 1).text = "B"
        docx_path = tmp_path / "docx-merged-table.docx"
        doc.save(docx_path)

        output_dir = tmp_path / "output_tables_empty_merge"
        output_dir.mkdir()

        result = _run_conversion(
            task_mgr,
            docx_path,
            output_dir,
            to_md_keep_images=False,
            table_merge_strategy="empty",
        )
        assert result.success, f"Conversion failed: {result.error.message if result.error else 'unknown'}"

        content = Path(result.artifacts[0].staging_path).read_text(encoding="utf-8")
        assert "| H |  |" in content
        assert "| --- | --- |" in content
        assert "| A | B |" in content

    def test_table_cell_paragraph_breaks_and_formatting_survive_finalizer(self, pipeline, tmp_path) -> None:
        """Real form tables keep paragraph boundaries and inline formatting."""
        from docx import Document

        _plugin, task_mgr, _ws_mgr = pipeline

        doc = Document()
        table = doc.add_table(rows=2, cols=1)
        table.cell(0, 0).text = "Header"
        body_cell = table.cell(1, 0)
        underlined = body_cell.paragraphs[0].add_run("Underlined")
        underlined.underline = True
        body_cell.add_paragraph("Second paragraph")
        docx_path = tmp_path / "docx-table-paragraph-formatting.docx"
        doc.save(docx_path)

        output_dir = tmp_path / "output_table_paragraph_formatting"
        output_dir.mkdir()
        result = _run_conversion(
            task_mgr,
            docx_path,
            output_dir,
            preserve_formatting=True,
        )

        assert result.success, result.error.message if result.error else "conversion failed"
        content = Path(result.artifacts[0].staging_path).read_text(encoding="utf-8")
        assert "| <u>Underlined</u><br>Second paragraph |" in content
        assert any(item.code == "FINALIZER_DONE" for item in result.diagnostics)

    def test_table_cell_pipes_are_escaped_through_the_real_converter(self, pipeline, tmp_path) -> None:
        """Cell pipes stay data while backslashes and paragraph breaks survive."""
        from docx import Document

        _plugin, task_mgr, _ws_mgr = pipeline

        doc = Document()
        table = doc.add_table(rows=2, cols=1)
        table.cell(0, 0).text = "Header | value"
        body_cell = table.cell(1, 0)
        body_cell.text = r"Path C:\Temp | ready; existing \| token"
        body_cell.add_paragraph("Second paragraph")
        docx_path = tmp_path / "docx-table-pipe-escaping.docx"
        doc.save(docx_path)

        output_dir = tmp_path / "output_table_pipe_escaping"
        output_dir.mkdir()
        result = _run_conversion(task_mgr, docx_path, output_dir)

        assert result.success, result.error.message if result.error else "conversion failed"
        content = Path(result.artifacts[0].staging_path).read_text(encoding="utf-8")
        assert "| Header \\| value |" in content
        assert r"| Path C:\Temp \| ready; existing \| token<br>Second paragraph |" in content

    def test_drawingml_textbox_paragraph_is_counted_through_finalizer(self, pipeline, tmp_path) -> None:
        """Exported textbox paragraphs contribute to artifact observability."""
        from docx import Document
        from lxml import etree

        from docwen_core.docx_parsing.xml_ns import NS_W, NS_WPS

        _plugin, task_mgr, _ws_mgr = pipeline
        doc = Document()
        doc.add_paragraph("Before textbox")
        outer_para = doc.add_paragraph()._p
        run = etree.SubElement(outer_para, f"{{{NS_W}}}r")
        drawing = etree.SubElement(run, f"{{{NS_W}}}drawing")
        textbox = etree.SubElement(drawing, f"{{{NS_WPS}}}txbx")
        content = etree.SubElement(textbox, f"{{{NS_W}}}txbxContent")
        inner_para = etree.SubElement(content, f"{{{NS_W}}}p")
        inner_run = etree.SubElement(inner_para, f"{{{NS_W}}}r")
        text = etree.SubElement(inner_run, f"{{{NS_W}}}t")
        text.text = "Inside textbox"
        doc.add_paragraph("After textbox")
        docx_path = tmp_path / "drawingml-textbox.docx"
        doc.save(docx_path)

        output_dir = tmp_path / "output_drawingml_textbox"
        output_dir.mkdir()
        result = _run_conversion(task_mgr, docx_path, output_dir)

        assert result.success, result.error.message if result.error else "conversion failed"
        artifact = next(item for item in result.artifacts if item.is_primary)
        markdown = Path(artifact.staging_path).read_text(encoding="utf-8")
        assert markdown.index("Before textbox") < markdown.index("Inside textbox") < markdown.index("After textbox")
        assert artifact.metadata["paragraph_count"] == 3
        assert any(item.code == "FINALIZER_DONE" for item in result.diagnostics)

    def test_tracked_insertion_reaches_final_artifact(self, pipeline, tmp_path) -> None:
        """A paragraph made only of revision wrappers must not be treated as empty."""
        from docx import Document
        from docx.oxml import OxmlElement

        _plugin, task_mgr, _ws_mgr = pipeline
        doc = Document()
        para = doc.add_paragraph()

        inserted = OxmlElement("w:ins")
        inserted_run = OxmlElement("w:r")
        inserted_text = OxmlElement("w:t")
        inserted_text.text = "Accepted insertion"
        inserted_run.append(inserted_text)
        inserted.append(inserted_run)
        para._p.append(inserted)

        deleted = OxmlElement("w:del")
        deleted_run = OxmlElement("w:r")
        deleted_text = OxmlElement("w:delText")
        deleted_text.text = "Rejected deletion"
        deleted_run.append(deleted_text)
        deleted.append(deleted_run)
        para._p.append(deleted)

        docx_path = tmp_path / "tracked-insertion.docx"
        doc.save(docx_path)
        output_dir = tmp_path / "output_tracked_insertion"
        output_dir.mkdir()

        result = _run_conversion(task_mgr, docx_path, output_dir)

        assert result.success, result.error.message if result.error else "conversion failed"
        artifact = next(item for item in result.artifacts if item.is_primary)
        markdown = Path(artifact.staging_path).read_text(encoding="utf-8")
        assert "Accepted insertion" in markdown
        assert "Rejected deletion" not in markdown
        assert artifact.metadata["paragraph_count"] == 1
        assert any(item.code == "FINALIZER_DONE" for item in result.diagnostics)
