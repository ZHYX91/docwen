"""V4 DOCX fidelity controls remain independent at the plugin boundary."""

from __future__ import annotations

import base64
import builtins
import hashlib
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from tests.support.config import FakeConfigView
from tests.support.execution import FakeExecutionContext
from tests.support.logging import FakePluginLogger
from tests.support.progress import FakeProgressSink
from tests.support.workspace import FakeWorkspaceHandle

from docwen_core.cancellation import CancellationToken
from docwen_core.docx_semantics_v3 import (
    CaptionStyleBindingV3,
    DocxSemanticsV3Session,
    append_complex_field,
)
from docwen_core.models.file_ref import FileRef
from docwen_core.models.request import ConversionRequest, OutputPolicy
from docwen_plugin_document.plugin import DocumentPlugin

pytestmark = pytest.mark.unit

_TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=="
)


def _probe_docx(tmp_path: Path) -> Path:
    image_path = tmp_path / "probe.png"
    image_path.write_bytes(_TINY_PNG)
    doc = Document()
    doc.add_paragraph("before image").add_run().add_picture(str(image_path))
    path = tmp_path / "probe.docx"
    doc.save(str(path))
    return path


def _probe_docx_with_two_images(tmp_path: Path) -> Path:
    image_path = tmp_path / "two-probe.png"
    image_path.write_bytes(_TINY_PNG)
    doc = Document()
    paragraph = doc.add_paragraph("two images")
    paragraph.add_run().add_picture(str(image_path))
    paragraph.add_run().add_picture(str(image_path))
    path = tmp_path / "two-probe.docx"
    doc.save(str(path))
    return path


def _authenticated_owner_docx(tmp_path: Path, owner_kind: str) -> Path:
    image_path = tmp_path / f"authenticated-{owner_kind}.png"
    image_path.write_bytes(_TINY_PNG)
    document = Document()
    source = (
        "Figure: Authenticated figure ^figure-owner\n\n![pixel](pixel.png)\n"
        if owner_kind == "figure"
        else "![pixel](pixel.png) ^ordinary-owner\n"
    )
    bindings: list[CaptionStyleBindingV3] = []
    styles: dict[str, Any] = {}
    for key, style_id, visible_name in (
        ("figure_caption", "DocWenFigureCaption", "Figure Caption"),
        ("table_caption", "DocWenTableCaption", "Table Caption"),
        ("equation_caption", "DocWenEquationCaption", "Equation Caption"),
        ("code_block_caption", "DocWenCodeBlockCaption", "Code Block Caption"),
    ):
        style = document.styles.add_style(visible_name, WD_STYLE_TYPE.PARAGRAPH)
        style._element.set(qn("w:styleId"), style_id)  # pyright: ignore[reportAttributeAccessIssue]
        bindings.append(CaptionStyleBindingV3(key, style_id, visible_name))  # type: ignore[arg-type]
        styles[key] = style
    session = DocxSemanticsV3Session(
        document,
        source_sha256=hashlib.sha256(source.encode()).hexdigest(),
        caption_style_bindings=tuple(bindings),
    )
    owner = document.add_paragraph()
    owner.add_run().add_picture(str(image_path))
    if owner_kind == "figure":
        caption = document.add_paragraph(style=styles["figure_caption"])
        caption.add_run("Figure ")
        append_complex_field(caption, instruction=" SEQ Figure \\* ARABIC ", cached_result="1")
        caption.add_run(": Authenticated figure")
        session.bind_caption(
            caption,
            (owner._p,),
            {
                "kind": "figure",
                "id": "figure-owner",
                "number": "1",
                "title": "Authenticated figure",
            },
        )
    else:
        session.bind_ordinary_anchor(
            (owner._p,),
            {"block_kind": "image", "id": "ordinary-owner"},
        )
    session.finalize_document()
    path = tmp_path / f"authenticated-{owner_kind}.docx"
    document.save(str(path))
    session.write_package(path)
    session.prove_package(path)
    return path


def _context(
    tmp_path: Path,
    input_path: Path,
    *,
    options: dict[str, Any],
) -> FakeExecutionContext:
    staging = tmp_path / "staging"
    staging.mkdir()
    config = {
        "export": {
            "to_md_image_extraction_mode": "file",
            "to_md_ocr_placement_mode": "main_md",
        },
        "conversion": {"ocr_output": {"show_blockquote_title": False}},
        "link": {
            "format": {
                "image_link_style": "markdown_embed",
                "md_file_link_style": "markdown_link",
            }
        },
    }
    return FakeExecutionContext(
        request=ConversionRequest(
            request_id="docx-fidelity-v4",
            input_refs=[
                FileRef(
                    path=str(input_path),
                    format="docx",
                    category="document",
                    size_bytes=input_path.stat().st_size,
                )
            ],
            target_format="md",
            options=options,
            output_policy=OutputPolicy(),
            config_snapshot=dict(config),
        ),
        workspace=FakeWorkspaceHandle(str(input_path), str(staging)),
        config=FakeConfigView(config),
        progress=FakeProgressSink(),
        cancellation=CancellationToken().view(),
        logger=FakePluginLogger(),
        numbering_registry=None,
        ocr_blockquote_title="",
    )


def _markdown_texts(result: Any) -> list[str]:
    return [
        Path(artifact.staging_path).read_text(encoding="utf-8")
        for artifact in result.artifacts
        if artifact.media_type == "text/markdown"
    ]


@pytest.mark.parametrize("recognize_text", [False, True])
@pytest.mark.parametrize("preserve_resources", [False, True])
@pytest.mark.parametrize("ocr_placement", ["main_md", "image_md"])
def test_docx_recognition_and_resource_preservation_are_independent(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    *,
    recognize_text: bool,
    preserve_resources: bool,
    ocr_placement: str,
) -> None:
    import docwen_core.text.ocr as ocr

    input_path = _probe_docx(tmp_path)
    context = _context(
        tmp_path,
        input_path,
        options={
            "to_md_enable_ocr": recognize_text,
            "to_md_keep_images": preserve_resources,
            "image_mode": "file",
            "ocr_placement": ocr_placement,
        },
    )
    observed_ocr_inputs: list[Path] = []

    def _run_ocr(path: str, **_kwargs: Any) -> ocr.OcrOutcome:
        observed = Path(path)
        assert observed.is_file()
        observed_ocr_inputs.append(observed)
        return ocr.OcrOutcome(ocr.OcrStatus.SUCCESS, text="recognized v4 text")

    monkeypatch.setattr(ocr, "run_ocr_outcome", _run_ocr)
    result = DocumentPlugin().convert(context)

    assert result.success, result.error
    images = [artifact for artifact in result.artifacts if artifact.kind == "image"]
    ocr_fragments = [
        artifact
        for artifact in result.artifacts
        if artifact.kind == "auxiliary" and artifact.metadata.get("ocr") is True
    ]
    markdown_texts = _markdown_texts(result)
    primary_text = next(
        Path(artifact.staging_path).read_text(encoding="utf-8") for artifact in result.artifacts if artifact.is_primary
    )

    assert len(observed_ocr_inputs) == int(recognize_text)
    assert len(images) == int(preserve_resources)
    assert len(ocr_fragments) == int(recognize_text and ocr_placement == "image_md")
    assert ("recognized v4 text" in primary_text) is (recognize_text and ocr_placement == "main_md")
    assert any("recognized v4 text" in text for text in markdown_texts) is recognize_text
    assert any("docx-image" in text for text in markdown_texts) is preserve_resources
    primary = next(artifact for artifact in result.artifacts if artifact.is_primary)
    assert primary.metadata["image_count"] == int(preserve_resources)

    if observed_ocr_inputs:
        assert observed_ocr_inputs[0].exists() is preserve_resources


@pytest.mark.parametrize("image_mode", ["omit", "base64"])
def test_docx_image_presentation_never_overrides_resource_preservation(
    tmp_path: Path,
    *,
    image_mode: str,
) -> None:
    input_path = _probe_docx(tmp_path)
    context = _context(
        tmp_path,
        input_path,
        options={
            "to_md_enable_ocr": False,
            "to_md_keep_images": True,
            "image_mode": image_mode,
        },
    )

    result = DocumentPlugin().convert(context)

    assert result.success, result.error
    images = [artifact for artifact in result.artifacts if artifact.kind == "image"]
    assert len(images) == 1
    assert Path(images[0].staging_path).is_file()
    primary_text = next(
        Path(artifact.staging_path).read_text(encoding="utf-8") for artifact in result.artifacts if artifact.is_primary
    )
    assert ("<!-- image omitted:" in primary_text) is (image_mode == "omit")
    assert ("data:image/png;base64," in primary_text) is (image_mode == "base64")


@pytest.mark.parametrize("owner_kind", ["figure", "ordinary"])
def test_authenticated_image_owner_omit_fails_before_staging(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    owner_kind: str,
) -> None:
    import docwen_core.text.ocr as ocr

    input_path = _authenticated_owner_docx(tmp_path, owner_kind)
    context = _context(
        tmp_path,
        input_path,
        options={
            "to_md_enable_ocr": True,
            "to_md_keep_images": True,
            "image_mode": "omit",
        },
    )
    observed_ocr_inputs: list[str] = []
    monkeypatch.setattr(
        ocr,
        "run_ocr_outcome",
        lambda path, **_kwargs: observed_ocr_inputs.append(path),
    )

    result = DocumentPlugin().convert(context)

    assert not result.success
    assert result.error is not None
    assert result.error.diagnostic_code == "DOCX2MD-PARSE-ERROR"
    assert result.error.message == "image_mode=omit cannot preserve an authenticated DOCX image owner"
    assert result.artifacts == []
    assert context.workspace.registered_artifacts == []
    assert list(Path(context.workspace.staging_dir).iterdir()) == []
    assert observed_ocr_inputs == []


@pytest.mark.parametrize("owner_kind", ["figure", "ordinary"])
@pytest.mark.parametrize(
    ("image_mode", "preserve_resources"),
    [("base64", True), ("omit", False)],
)
def test_authenticated_image_owner_survives_owner_preserving_non_file_modes(
    tmp_path: Path,
    owner_kind: str,
    image_mode: str,
    preserve_resources: bool,
) -> None:
    input_path = _authenticated_owner_docx(tmp_path, owner_kind)
    context = _context(
        tmp_path,
        input_path,
        options={
            "to_md_enable_ocr": False,
            "to_md_keep_images": preserve_resources,
            "image_mode": image_mode,
        },
    )

    result = DocumentPlugin().convert(context)

    assert result.success, result.error
    primary_text = next(
        Path(artifact.staging_path).read_text(encoding="utf-8") for artifact in result.artifacts if artifact.is_primary
    )
    images = [artifact for artifact in result.artifacts if artifact.kind == "image"]
    assert len(images) == int(preserve_resources)
    assert "<!-- image omitted:" not in primary_text
    assert ("data:image/png;base64," in primary_text) is preserve_resources
    assert ("![image omitted]()" in primary_text) is (not preserve_resources)
    if owner_kind == "figure":
        assert "Figure: Authenticated figure ^figure-owner" in primary_text
    else:
        assert "^ordinary-owner" in primary_text


def test_primary_write_failure_rolls_back_pending_image_artifacts(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    input_path = _probe_docx(tmp_path)
    context = _context(
        tmp_path,
        input_path,
        options={
            "to_md_enable_ocr": False,
            "to_md_keep_images": True,
            "image_mode": "file",
        },
    )
    real_open = builtins.open

    def _fail_primary(path: Any, mode: str = "r", *args: Any, **kwargs: Any):
        if str(path).endswith(".md") and "w" in mode:
            raise OSError("synthetic primary write failure")
        return real_open(path, mode, *args, **kwargs)

    monkeypatch.setattr(builtins, "open", _fail_primary)
    result = DocumentPlugin().convert(context)

    assert not result.success
    assert result.artifacts == []
    assert context.workspace.registered_artifacts == []
    assert list(Path(context.workspace.staging_dir).iterdir()) == []


def test_sidecar_write_failure_rolls_back_pending_image_artifacts(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import docwen_core.text.ocr as ocr

    input_path = _probe_docx(tmp_path)
    context = _context(
        tmp_path,
        input_path,
        options={
            "to_md_enable_ocr": True,
            "to_md_keep_images": True,
            "image_mode": "file",
            "ocr_placement": "image_md",
        },
    )
    monkeypatch.setattr(
        ocr,
        "run_ocr_outcome",
        lambda *_args, **_kwargs: ocr.OcrOutcome(ocr.OcrStatus.SUCCESS, text="recognized"),
    )
    real_write_text = Path.write_text

    def _fail_sidecar(path: Path, *args: Any, **kwargs: Any) -> int:
        if path.suffix == ".md":
            real_write_text(path, "partial", encoding="utf-8")
            raise OSError("synthetic sidecar write failure")
        return real_write_text(path, *args, **kwargs)

    monkeypatch.setattr(Path, "write_text", _fail_sidecar)
    result = DocumentPlugin().convert(context)

    assert not result.success
    assert result.artifacts == []
    assert context.workspace.registered_artifacts == []
    assert list(Path(context.workspace.staging_dir).iterdir()) == []


def test_partial_second_image_write_leaves_staging_empty(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    input_path = _probe_docx_with_two_images(tmp_path)
    context = _context(
        tmp_path,
        input_path,
        options={
            "to_md_enable_ocr": False,
            "to_md_keep_images": True,
            "image_mode": "file",
        },
    )
    real_write_bytes = Path.write_bytes
    image_writes = 0

    def _fail_second_image(path: Path, content: bytes) -> int:
        nonlocal image_writes
        if path.parent == Path(context.workspace.staging_dir) and path.suffix == ".png":
            image_writes += 1
            if image_writes == 2:
                real_write_bytes(path, b"partial")
                raise OSError("synthetic partial extraction write")
        return real_write_bytes(path, content)

    monkeypatch.setattr(Path, "write_bytes", _fail_second_image)
    result = DocumentPlugin().convert(context)

    assert not result.success
    assert result.artifacts == []
    assert context.workspace.registered_artifacts == []
    assert list(Path(context.workspace.staging_dir).iterdir()) == []
