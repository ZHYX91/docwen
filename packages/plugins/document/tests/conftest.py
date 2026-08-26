"""Shared fixtures for docx plugin tests."""

from __future__ import annotations

import sys
from pathlib import Path

import pytest

PROJECT_ROOT = Path(__file__).resolve().parents[4]

LOCAL_SRC_PATHS = [
    PROJECT_ROOT,
    PROJECT_ROOT / "packages" / "core" / "src",
    PROJECT_ROOT / "packages" / "runtime" / "src",
    PROJECT_ROOT / "packages" / "plugins" / "document" / "src",
]

for path in reversed(LOCAL_SRC_PATHS):
    if str(path) not in sys.path:
        sys.path.insert(0, str(path))


@pytest.fixture
def sample_docx_path(tmp_path: Path) -> Path:
    """Create a sample DOCX file with known content for golden testing.

    Contains:
    - Heading 1: "Test Document"
    - Heading 2: "Section One"
    - Regular paragraph with bold text
    - Regular paragraph with italic text
    - A simple 3-row table
    - Plain paragraph

    Returns the path to the created DOCX file.
    """
    from docx import Document

    doc = Document()

    # Title / Heading 1
    doc.add_heading("Test Document", level=1)

    # Intro paragraph
    doc.add_paragraph("This is a test document used for golden parity testing.")

    # Heading 2
    doc.add_heading("Section One", level=2)

    # Bold paragraph
    p_bold = doc.add_paragraph()
    run_b = p_bold.add_run("This text is bold.")
    run_b.bold = True
    p_bold.add_run(" This text is normal.")

    # Italic paragraph
    p_italic = doc.add_paragraph()
    run_i = p_italic.add_run("This text is italic.")
    run_i.italic = True

    # Plain paragraph
    doc.add_paragraph("A plain paragraph with no special formatting.")

    # Heading 3
    doc.add_heading("Table Section", level=2)

    # Table
    table = doc.add_table(rows=3, cols=3, style="Table Grid")
    # Header row
    for j, header in enumerate(["Name", "Value", "Description"]):
        table.rows[0].cells[j].text = header
    # Data rows
    table.rows[1].cells[0].text = "Alpha"
    table.rows[1].cells[1].text = "100"
    table.rows[1].cells[2].text = "First item"
    table.rows[2].cells[0].text = "Beta"
    table.rows[2].cells[1].text = "200"
    table.rows[2].cells[2].text = "Second item"

    # Save
    output_path = tmp_path / "golden_test.docx"
    doc.save(str(output_path))

    return output_path
