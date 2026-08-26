"""Focused tests split from test_document_to_md_golden.py."""

from __future__ import annotations

from ._document_to_md_golden_support import (
    ConversionRequest,
    FileRef,
    OutputPolicy,
    Path,
    _build_runtime_pipeline,
    _document_node_root,
    _run_conversion,
    os,
    pytest,
    re,
)

pytestmark = [pytest.mark.golden, pytest.mark.contract]


class TestDocxToMdGolden:
    """Golden parity tests for ROUTE-DOC-001: document → md.

    Uses a programmatically-generated sample DOCX with known content
    (headings, paragraphs with formatting, tables) for structural
    validation, plus a real template for integration smoke testing.
    """

    @pytest.fixture
    def pipeline(self):
        """Build the full runtime pipeline with the real DocumentPlugin."""
        plugin, task_mgr, ws_mgr, ws_root = _build_runtime_pipeline()
        yield plugin, task_mgr, ws_mgr
        ws_mgr.cleanup_all()
        import shutil

        shutil.rmtree(ws_root, ignore_errors=True)

    @staticmethod
    def _markdown_table_cells(line: str) -> list[str]:
        row = line.strip()
        if row.startswith("|"):
            row = row[1:]
        if row.endswith("|"):
            row = row[:-1]
        return [cell.replace(r"\|", "|").strip() for cell in re.split(r"(?<!\\)\|", row)]

    @staticmethod
    def _verify_tables_well_formed(content: str) -> None:
        """Verify that Markdown tables in content are well-formed and have data rows."""
        lines = content.splitlines()
        in_table = False
        separator_seen = False
        data_rows_seen = 0
        column_count = 0

        for line in lines:
            stripped = line.strip()
            is_table_line = stripped.startswith("|") and stripped.endswith("|")

            if is_table_line:
                if not in_table:
                    in_table = True
                    separator_seen = False
                    data_rows_seen = 0
                    column_count = len(TestDocxToMdGolden._markdown_table_cells(stripped))
                else:
                    if not separator_seen:
                        separator_seen = True
                        parts = TestDocxToMdGolden._markdown_table_cells(stripped)
                        assert all(set(p) <= {"-", ":", " "} and "-" in p for p in parts), (
                            f"Malformed table separator: {stripped}"
                        )
                    else:
                        cols = len(TestDocxToMdGolden._markdown_table_cells(stripped))
                        assert cols == column_count, f"Table column count mismatch: expected {column_count}, got {cols}"
                        data_rows_seen += 1
            else:
                if in_table and stripped == "":
                    in_table = False
                    separator_seen = False
                    assert data_rows_seen >= 1, f"Table should have at least 1 data row, got {data_rows_seen}"

    def test_paragraph_text_preserved(self, pipeline, sample_docx_path, tmp_path) -> None:
        """GOLDEN-002: Paragraph text must be preserved in output."""
        _plugin, task_mgr, _ws_mgr = pipeline

        output_dir = tmp_path / "output_para"
        output_dir.mkdir()

        result = _run_conversion(task_mgr, sample_docx_path, output_dir)
        assert result.success

        content = Path(result.artifacts[0].staging_path).read_text(encoding="utf-8")

        assert "test document used for golden parity testing" in content.lower()
        assert "plain paragraph" in content.lower()

    def test_bold_text_has_markers(self, pipeline, sample_docx_path, tmp_path) -> None:
        """GOLDEN-002: Bold text should be marked with **...** in output."""
        _plugin, task_mgr, _ws_mgr = pipeline

        output_dir = tmp_path / "output_bold"
        output_dir.mkdir()

        result = _run_conversion(task_mgr, sample_docx_path, output_dir)
        assert result.success

        content = Path(result.artifacts[0].staging_path).read_text(encoding="utf-8")

        # Bold text should be marked
        assert "**This text is bold.**" in content or "**This text is bold" in content, (
            f"Bold text markers not found. Content:\n{content[:300]}"
        )

    def test_italic_text_has_markers(self, pipeline, sample_docx_path, tmp_path) -> None:
        """GOLDEN-002: Italic text should be marked with *...* in output."""
        _plugin, task_mgr, _ws_mgr = pipeline

        output_dir = tmp_path / "output_italic"
        output_dir.mkdir()

        result = _run_conversion(task_mgr, sample_docx_path, output_dir)
        assert result.success

        content = Path(result.artifacts[0].staging_path).read_text(encoding="utf-8")

        # Italic text should be marked
        assert "*This text is italic.*" in content or "*This text is italic" in content, (
            f"Italic text markers not found. Content:\n{content[:300]}"
        )

    def test_docx_to_md_merges_adjacent_same_style_runs(self, pipeline, tmp_path) -> None:
        """GOLDEN-002: Split DOCX runs render as one Markdown span per style."""
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        _plugin, task_mgr, _ws_mgr = pipeline

        def add_gray_shading(run) -> None:
            rpr = run._r.get_or_add_rPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "D9D9D9")
            shd.set(qn("w:val"), "clear")
            rpr.append(shd)

        doc = Document()
        bold_para = doc.add_paragraph()
        bold_run_1 = bold_para.add_run("Hello ")
        bold_run_1.bold = True
        bold_run_2 = bold_para.add_run("World")
        bold_run_2.bold = True

        code_para = doc.add_paragraph()
        code_run_1 = code_para.add_run("ab")
        add_gray_shading(code_run_1)
        code_run_2 = code_para.add_run("cd")
        add_gray_shading(code_run_2)
        code_para.add_run("x")

        docx_path = tmp_path / "split-runs.docx"
        doc.save(docx_path)

        output_dir = tmp_path / "output_split_runs"
        output_dir.mkdir()

        result = _run_conversion(task_mgr, docx_path, output_dir, to_md_keep_images=False)
        assert result.success, f"Conversion failed: {result.error.message if result.error else 'unknown'}"

        content = Path(result.artifacts[0].staging_path).read_text(encoding="utf-8")
        assert "**Hello World**" in content
        assert "**Hello ****World**" not in content
        assert "`abcd`x" in content
        assert "`ab``cd`x" not in content

    def test_conversion_metrics_meaningful(self, pipeline, sample_docx_path, tmp_path) -> None:
        """GOLDEN-002: ConversionResult must contain meaningful metrics."""
        _plugin, task_mgr, _ws_mgr = pipeline

        output_dir = tmp_path / "output_metrics"
        output_dir.mkdir()

        result = _run_conversion(task_mgr, sample_docx_path, output_dir)
        assert result.success

        assert result.metrics.duration_ms >= 0
        assert result.metrics.input_bytes > 0, "Input bytes should be > 0"
        assert result.metrics.output_bytes > 0, "Output bytes should be > 0"

        # Stats are on the artifact metadata (plugin → artifact manifest)
        artifact = result.artifacts[0]
        meta = artifact.metadata
        assert meta.get("paragraph_count", 0) >= 3, f"Expected at least 3 paragraphs, got artifact metadata={meta}"
        assert meta.get("heading_count", 0) >= 2, f"Expected at least 2 headings, got artifact metadata={meta}"
        assert meta.get("table_count", 0) >= 1, f"Expected at least 1 table, got artifact metadata={meta}"

    def test_artifact_metadata_complete(self, pipeline, sample_docx_path, tmp_path) -> None:
        """GOLDEN-002: ArtifactManifest must carry complete metadata with correct values."""
        _plugin, task_mgr, _ws_mgr = pipeline

        output_dir = tmp_path / "output_artifact_meta"
        output_dir.mkdir()

        result = _run_conversion(task_mgr, sample_docx_path, output_dir)
        assert result.success

        artifact = result.artifacts[0]
        assert artifact.artifact_id
        assert artifact.kind == "primary"
        assert artifact.suggested_name.endswith(".md")
        assert artifact.media_type == "text/markdown"
        assert artifact.is_primary is True

        # Verify metadata keys exist with meaningful values
        assert artifact.metadata.get("paragraph_count", 0) >= 3, f"Expected >=3 paragraphs, got {artifact.metadata}"
        assert artifact.metadata.get("heading_count", 0) >= 2, f"Expected >=2 headings, got {artifact.metadata}"
        assert artifact.metadata.get("table_count", 0) >= 1, f"Expected >=1 table, got {artifact.metadata}"
        assert "image_count" in artifact.metadata

    def test_docx_images_are_finalized_as_relative_markdown_artifacts(self, pipeline, tmp_path) -> None:
        """GOLDEN-002: Extracted DOCX images must survive runtime finalization."""
        from docx import Document

        _plugin, task_mgr, _ws_mgr = pipeline

        image_path = tmp_path / "tiny.png"
        image_path.write_bytes(
            b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR"
            b"\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde"
            b"\x00\x00\x00\x0cIDATx\x9cc```\x00\x00\x00\x04\x00\x01\xf6\x178U"
            b"\x00\x00\x00\x00IEND\xaeB`\x82"
        )

        doc = Document()
        doc.add_heading("Image Artifact Probe", level=1)
        run = doc.add_paragraph("Image: ").add_run()
        run.add_picture(str(image_path))
        docx_path = tmp_path / "image_artifact_probe.docx"
        doc.save(docx_path)

        output_dir = tmp_path / "output_images"
        output_dir.mkdir()

        result = _run_conversion(
            task_mgr,
            docx_path,
            output_dir,
            to_md_keep_images=True,
            image_mode="file",
        )

        assert result.success
        markdown_artifact = next(artifact for artifact in result.artifacts if artifact.is_primary)
        image_artifacts = [artifact for artifact in result.artifacts if artifact.kind == "image"]
        assert markdown_artifact.metadata["image_count"] == 1
        assert len(image_artifacts) == 1
        node_root = _document_node_root(Path(markdown_artifact.staging_path), output_dir)
        assert _document_node_root(Path(image_artifacts[0].staging_path), output_dir) == node_root
        assert Path(image_artifacts[0].staging_path).is_file()

        content = Path(markdown_artifact.staging_path).read_text(encoding="utf-8")
        assert image_artifacts[0].suggested_name in content
        assert str(output_dir) not in content
        assert "docwen_ws_" not in content

    def test_image_only_docx_paragraph_is_finalized_and_linked(self, pipeline, tmp_path) -> None:
        """A picture-only paragraph must not create an unreferenced final artifact."""
        from docx import Document

        _plugin, task_mgr, _ws_mgr = pipeline

        image_path = tmp_path / "image-only.png"
        image_path.write_bytes(
            b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR"
            b"\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde"
            b"\x00\x00\x00\x0cIDATx\x9cc```\x00\x00\x00\x04\x00\x01\xf6\x178U"
            b"\x00\x00\x00\x00IEND\xaeB`\x82"
        )

        doc = Document()
        doc.add_heading("Image-only Artifact Probe", level=1)
        doc.add_paragraph().add_run().add_picture(str(image_path))
        docx_path = tmp_path / "image_only_artifact_probe.docx"
        doc.save(docx_path)

        output_dir = tmp_path / "output_image_only"
        output_dir.mkdir()
        result = _run_conversion(
            task_mgr,
            docx_path,
            output_dir,
            to_md_keep_images=True,
            image_mode="file",
        )

        assert result.success
        markdown_artifact = next(artifact for artifact in result.artifacts if artifact.is_primary)
        image_artifacts = [artifact for artifact in result.artifacts if artifact.kind == "image"]
        assert markdown_artifact.metadata["image_count"] == 1
        assert len(image_artifacts) == 1

        content = Path(markdown_artifact.staging_path).read_text(encoding="utf-8")
        assert content.count(image_artifacts[0].suggested_name) == 1
        assert Path(image_artifacts[0].staging_path).is_file()
        assert str(output_dir) not in content
        assert "docwen_ws_" not in content

    def test_docx_image_mode_uses_request_export_semantics(self, pipeline, tmp_path) -> None:
        """DOCX→MD inherits its admitted Export mode when image_mode is omitted."""
        from docx import Document

        _plugin, task_mgr, _ws_mgr = pipeline

        image_path = tmp_path / "tiny.png"
        image_path.write_bytes(
            b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR"
            b"\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde"
            b"\x00\x00\x00\x0cIDATx\x9cc```\x00\x00\x00\x04\x00\x01\xf6\x178U"
            b"\x00\x00\x00\x00IEND\xaeB`\x82"
        )

        doc = Document()
        doc.add_heading("Export Fallback Probe", level=1)
        run = doc.add_paragraph("Image: ").add_run()
        run.add_picture(str(image_path))
        docx_path = tmp_path / "export_fallback_probe.docx"
        doc.save(docx_path)

        output_dir = tmp_path / "output_export_fallback"
        output_dir.mkdir()

        result = _run_conversion(
            task_mgr,
            docx_path,
            output_dir,
            config_snapshot={"export": {"to_md_image_extraction_mode": "base64"}},
            to_md_keep_images=True,
        )

        assert result.success
        markdown_artifact = next(artifact for artifact in result.artifacts if artifact.is_primary)
        content = Path(markdown_artifact.staging_path).read_text(encoding="utf-8")
        assert "data:image/png;base64," in content
        assert sum(artifact.kind == "image" for artifact in result.artifacts) == 1

    def test_conversion_succeeds_with_real_template(self, pipeline, tmp_path) -> None:
        """GOLDEN-002: Conversion should succeed even with template-only DOCX."""
        _plugin, task_mgr, _ws_mgr = pipeline

        # Find the real template
        template = Path(__file__).resolve().parents[4] / "templates" / "简体中文通用模板.docx"
        if not template.exists():
            template = Path(__file__).resolve().parents[4] / "templates" / "English General Template.docx"
        if not template.exists():
            pytest.skip("No template file available")

        output_dir = tmp_path / "output_template"
        output_dir.mkdir()

        result = _run_conversion(task_mgr, template, output_dir)
        assert result.success, f"Template conversion failed: {result.error.message if result.error else 'unknown'}"
        assert len(result.artifacts) >= 1
        assert os.path.isfile(result.artifacts[0].staging_path)

        # Content assertions: real template should produce non-empty output
        content = Path(result.artifacts[0].staging_path).read_text(encoding="utf-8")
        assert len(content.strip()) > 0, "Template output should be non-empty"
        lines = [line for line in content.splitlines() if line.strip()]
        assert len(lines) >= 1, "Template output should have at least 1 non-empty line"

    def test_cancellation_before_start(self, pipeline, sample_docx_path, tmp_path) -> None:
        """Cancelling before conversion should produce a cancelled result."""
        _plugin, task_mgr, _ws_mgr = pipeline

        output_dir = tmp_path / "output_cancel"
        output_dir.mkdir()

        request = ConversionRequest(
            request_id="golden-002-cancel",
            input_refs=[
                FileRef(
                    path=str(sample_docx_path),
                    format="docx",
                    category="document",
                    size_bytes=sample_docx_path.stat().st_size,
                )
            ],
            target_format="md",
            output_policy=OutputPolicy(output_dir=str(output_dir)),
        )

        task_mgr.cancel("golden-002-cancel")
        result = task_mgr.execute_single(request)

        assert result.success is False
        assert result.error is not None
        assert result.error.error_type == "cancelled", (
            f"Expected error_type='cancelled', got '{result.error.error_type}'"
        )

    def test_enex_to_md_not_handled_by_document(self, pipeline, tmp_path) -> None:
        """ENEX→MD has been migrated to markup plugin — document must reject it."""
        _plugin, task_mgr, _ws_mgr = pipeline

        output_dir = tmp_path / "output_notimpl"
        output_dir.mkdir()

        dummy_enex = tmp_path / "dummy.enex"
        dummy_enex.write_text("not valid xml", encoding="utf-8")

        request = ConversionRequest(
            request_id="notimpl-test",
            input_refs=[
                FileRef(
                    path=str(dummy_enex),
                    format="enex",
                    category="markup",
                    size_bytes=dummy_enex.stat().st_size,
                )
            ],
            target_format="md",
            output_policy=OutputPolicy(output_dir=str(output_dir)),
        )
        result = task_mgr.execute_single(request)

        # DocumentPlugin does not handle enex→md — route migrated to markup plugin.
        assert result.success is False
        assert result.error is not None
        assert "No plugin found" in result.error.message

    def test_html_to_md_not_handled_by_document(self, pipeline, tmp_path) -> None:
        """HTML→MD has been migrated to markup plugin — document must reject it."""
        _plugin, task_mgr, _ws_mgr = pipeline

        output_dir = tmp_path / "output_notimpl_html"
        output_dir.mkdir()

        html_path = tmp_path / "minimal.html"
        html_path.write_text("<html><body><p>Hello World</p></body></html>", encoding="utf-8")

        request = ConversionRequest(
            request_id="notimpl-html",
            input_refs=[
                FileRef(
                    path=str(html_path),
                    format="html",
                    category="markup",
                    size_bytes=html_path.stat().st_size,
                )
            ],
            target_format="md",
            output_policy=OutputPolicy(output_dir=str(output_dir)),
        )
        result = task_mgr.execute_single(request)

        # DocumentPlugin does not handle html→md — route migrated to markup plugin.
        assert result.success is False
        assert result.error is not None
        assert "No plugin found" in result.error.message

    def test_ppt_to_md_not_handled_by_document(self, pipeline, tmp_path) -> None:
        """PPT→MD has been migrated to presentation plugin — document must reject it."""
        _plugin, task_mgr, _ws_mgr = pipeline

        output_dir = tmp_path / "output_notimpl_pptx"
        output_dir.mkdir()

        dummy_ppt = tmp_path / "dummy.ppt"
        dummy_ppt.write_bytes(b"")

        request = ConversionRequest(
            request_id="notimpl-pptx",
            input_refs=[
                FileRef(
                    path=str(dummy_ppt),
                    format="ppt",
                    category="presentation",
                    size_bytes=0,
                )
            ],
            target_format="md",
            output_policy=OutputPolicy(output_dir=str(output_dir)),
        )
        result = task_mgr.execute_single(request)

        # DocumentPlugin does not handle ppt→md — route migrated to presentation plugin.
        assert result.success is False
        assert result.error is not None
        assert "No plugin found" in result.error.message

    def test_smart_converter_route_rejects_missing_input(self, pipeline, tmp_path) -> None:
        """SmartDoc should reject empty or missing source files before bridge calls."""
        _plugin, task_mgr, _ws_mgr = pipeline

        output_dir = tmp_path / "output_notimpl_smart"
        output_dir.mkdir()

        request = ConversionRequest(
            request_id="notimpl-smart",
            input_refs=[
                FileRef(
                    path=str(tmp_path / "dummy.odt"),
                    format="odt",
                    category="document",
                    size_bytes=0,
                )
            ],
            target_format="docx",
            output_policy=OutputPolicy(output_dir=str(output_dir)),
        )
        result = task_mgr.execute_single(request)

        assert result.success is False
        assert result.error is not None
        assert result.error.error_type == "invalid_input"
        assert "odt" in result.error.message.lower()
