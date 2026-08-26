import tomllib
from pathlib import Path
from typing import cast
from unittest.mock import MagicMock

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.styles.style import ParagraphStyle

from docwen_core.docx_parsing.format_features import (
    CodeBlockAccumulator,
    detect_paragraph_style_type,
    extract_alignment,
    extract_font_info,
    extract_outline_level,
    has_gray_shading,
    has_paragraph_gray_shading,
)

pytestmark = pytest.mark.unit

# ── Existing helpers ───────────────────────────────────────────────────


def _set_east_asia_font(owner, name: str) -> None:
    r_pr = owner._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), name)


def _set_east_asia_theme(owner, token: str) -> None:
    r_pr = owner._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsiaTheme"), token)


def _set_custom_hans_theme(doc, typeface: str) -> None:
    theme_part = next(rel.target_part for rel in doc.part.rels.values() if rel.reltype.endswith("/theme"))
    theme_part._blob = f"""\
<a:theme xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" name="RA1">
  <a:themeElements>
    <a:fontScheme name="RA1">
      <a:majorFont>
        <a:latin typeface="Latin Major"/><a:ea typeface=""/><a:cs typeface=""/>
        <a:font script="Hans" typeface="{typeface}"/>
      </a:majorFont>
      <a:minorFont><a:latin typeface="Latin Minor"/><a:ea typeface=""/><a:cs typeface=""/></a:minorFont>
    </a:fontScheme>
  </a:themeElements>
</a:theme>
""".encode()
    theme_font_lang = doc.settings.element.find(qn("w:themeFontLang"))
    assert theme_font_lang is not None
    theme_font_lang.set(qn("w:eastAsia"), "zh-CN")


def test_extract_font_info_from_first_run():
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("hello")
    run.font.name = "SimSun"
    run.font.size = Pt(12)

    assert extract_font_info(para) == ("SimSun", 12.0)


def test_extract_font_info_prefers_east_asia_for_cjk_run():
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("公文标题")
    _set_east_asia_font(run, "方正小标宋简体")
    run.font.size = Pt(22)

    assert extract_font_info(para) == ("方正小标宋简体", 22.0)


def test_extract_font_info_ignores_decorative_leading_latin_run():
    doc = Document()
    para = doc.add_paragraph()
    leading = para.add_run("prefix ")
    leading.font.name = "Calibri"
    leading.font.size = Pt(8)
    title = para.add_run("公文标题")
    _set_east_asia_font(title, "方正小标宋简体")
    title.font.size = Pt(22)

    assert extract_font_info(para) == ("方正小标宋简体", 22.0)


def test_extract_font_info_cascades_name_and_size_independently():
    doc = Document()
    base = cast(ParagraphStyle, doc.styles.add_style("RA1 Base", WD_STYLE_TYPE.PARAGRAPH))
    _set_east_asia_font(base, "方正小标宋简体")
    base.font.size = Pt(18)
    child = cast(ParagraphStyle, doc.styles.add_style("RA1 Child", WD_STYLE_TYPE.PARAGRAPH))
    child.base_style = base

    para = doc.add_paragraph(style=child)
    run = para.add_run("公文标题")
    run.font.size = Pt(22)
    assert extract_font_info(para) == ("方正小标宋简体", 22.0)

    para = doc.add_paragraph(style=child)
    run = para.add_run("公文标题")
    _set_east_asia_font(run, "华文中宋")
    assert extract_font_info(para) == ("华文中宋", 18.0)


def test_extract_font_info_falls_back_to_document_defaults():
    doc = Document()
    styles = doc.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    assert doc_defaults is not None
    r_pr_default = doc_defaults.find(qn("w:rPrDefault"))
    assert r_pr_default is not None
    r_pr = r_pr_default.find(qn("w:rPr"))
    assert r_pr is not None
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), "默认宋体")
    size = r_pr.find(qn("w:sz"))
    if size is None:
        size = OxmlElement("w:sz")
        r_pr.append(size)
    size.set(qn("w:val"), "44")

    para = doc.add_paragraph("公文标题")
    assert extract_font_info(para) == ("默认宋体", 22.0)


def test_extract_font_info_resolves_east_asia_theme_by_language():
    doc = Document()
    _set_custom_hans_theme(doc, "主题小标宋")
    para = doc.add_paragraph()
    run = para.add_run("公文标题")
    _set_east_asia_theme(run, "majorEastAsia")
    run.font.size = Pt(22)

    assert extract_font_info(para) == ("主题小标宋", 22.0)


def test_extract_font_info_broken_theme_falls_back_to_style():
    doc = Document()
    style = cast(ParagraphStyle, doc.styles.add_style("RA1 Theme Fallback", WD_STYLE_TYPE.PARAGRAPH))
    _set_east_asia_font(style, "样式小标宋")
    style.font.size = Pt(22)
    para = doc.add_paragraph(style=style)
    run = para.add_run("公文标题")
    _set_east_asia_theme(run, "majorEastAsia")
    theme_part = next(rel.target_part for rel in doc.part.rels.values() if rel.reltype.endswith("/theme"))
    theme_part._blob = b"<broken"

    assert extract_font_info(para) == ("样式小标宋", 22.0)


def test_extract_font_info_falls_back_to_normal_style_components():
    doc = Document()
    normal = cast(ParagraphStyle, doc.styles["Normal"])
    _set_east_asia_font(normal, "Normal 小标宋")
    normal.font.size = Pt(19)
    custom = cast(ParagraphStyle, doc.styles.add_style("RA1 Unbased", WD_STYLE_TYPE.PARAGRAPH))

    para = doc.add_paragraph(style=custom)
    para.add_run("公文标题")

    assert extract_font_info(para) == ("Normal 小标宋", 19.0)


def test_extract_font_info_combines_custom_and_normal_components_both_directions():
    doc = Document()
    normal = cast(ParagraphStyle, doc.styles["Normal"])
    _set_east_asia_font(normal, "Normal 宋体")
    normal.font.size = Pt(17)

    custom_name = cast(ParagraphStyle, doc.styles.add_style("RA1 Name Only", WD_STYLE_TYPE.PARAGRAPH))
    _set_east_asia_font(custom_name, "自定义宋体")
    paragraph = doc.add_paragraph(style=custom_name)
    paragraph.add_run("公文标题")
    assert extract_font_info(paragraph) == ("自定义宋体", 17.0)

    custom_size = cast(ParagraphStyle, doc.styles.add_style("RA1 Size Only", WD_STYLE_TYPE.PARAGRAPH))
    custom_size.font.size = Pt(21)
    paragraph = doc.add_paragraph(style=custom_size)
    paragraph.add_run("公文标题")
    assert extract_font_info(paragraph) == ("Normal 宋体", 21.0)


def test_extract_alignment_name():
    doc = Document()
    para = doc.add_paragraph("center")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    assert extract_alignment(para) == "CENTER"


def test_extract_outline_level_from_heading_style():
    doc = Document()
    para = doc.add_paragraph("Heading")
    para.style = "Heading 2"

    assert extract_outline_level(para) == 1


def test_extract_outline_level_treats_level_nine_as_body_text():
    """OOXML outline level 9 is Word's body-text sentinel, not Heading 10."""
    doc = Document()
    para = doc.add_paragraph("ordinary body")
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), "9")
    para._p.get_or_add_pPr().append(outline)

    assert extract_outline_level(para) is None


# ── Style type detection ───────────────────────────────────────────────


def test_detect_code_block_style_english():
    para = MagicMock()
    para.style.name = "Code Block"

    stype, sval = detect_paragraph_style_type(para)
    assert stype == "code_block"
    assert sval is True


def test_detect_code_block_style_chinese():
    para = MagicMock()
    para.style.name = "代码块"

    stype, _sval = detect_paragraph_style_type(para)
    assert stype == "code_block"


def test_detect_quote_style_with_level():
    para = MagicMock()
    para.style.name = "Quote 3"

    stype, sval = detect_paragraph_style_type(para)
    assert stype == "quote"
    assert sval == 3


def test_detect_quote_style_chinese():
    para = MagicMock()
    para.style.name = "引用 2"

    stype, sval = detect_paragraph_style_type(para)
    assert stype == "quote"
    assert sval == 2


def test_detect_style_normal_paragraph():
    para = MagicMock()
    para.style.name = "Normal"

    stype, sval = detect_paragraph_style_type(para)
    assert stype is None
    assert sval is None


def test_detect_style_no_style():
    para = MagicMock()
    para.style.name = None

    stype, sval = detect_paragraph_style_type(para)
    assert stype is None
    assert sval is None


def test_detect_generic_quote_style():
    para = MagicMock()
    para.style.name = "Intense Quote"

    stype, sval = detect_paragraph_style_type(para)
    assert stype == "quote"
    assert sval == 1


def test_every_bundled_locale_template_style_is_detected_from_real_docx() -> None:
    repo_root = Path(__file__).resolve().parents[4]
    locale_paths = sorted((repo_root / "i18n" / "locales").glob("*.toml"))
    template_paths = sorted((repo_root / "templates").glob("*.docx"))
    assert len(locale_paths) == 11
    assert len(template_paths) == 11

    template_documents = [(path, Document(str(path))) for path in template_paths]
    template_style_names = {path: {style.name for style in document.styles} for path, document in template_documents}

    for locale_path in locale_paths:
        locale_table = tomllib.loads(locale_path.read_text(encoding="utf-8"))
        styles = locale_table["styles"]
        code_name = styles["code_block"]
        matching_templates = [
            (path, document) for path, document in template_documents if code_name in template_style_names[path]
        ]
        assert len(matching_templates) == 1, (locale_path.name, code_name, matching_templates)
        template_path, document = matching_templates[0]

        code_paragraph = document.add_paragraph("code", style=code_name)
        assert detect_paragraph_style_type(code_paragraph) == ("code_block", True), (
            locale_path.name,
            template_path.name,
            code_name,
        )

        for level in range(1, 10):
            quote_name = styles[f"quote_{level}"]
            assert quote_name in template_style_names[template_path], (
                locale_path.name,
                template_path.name,
                quote_name,
            )
            quote_paragraph = document.add_paragraph("quote", style=quote_name)
            assert detect_paragraph_style_type(quote_paragraph) == ("quote", level), (
                locale_path.name,
                template_path.name,
                quote_name,
            )


# ── Gray shading detection ─────────────────────────────────────────────


def test_has_gray_shading_wps_fill():
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("code")
    # Inject gray shading into the run
    rPr = OxmlElement("w:rPr")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "D9D9D9")
    rPr.append(shd)
    run._r.insert(0, rPr)

    assert has_gray_shading(run) is True


def test_has_gray_shading_no_shading():
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("plain")

    assert has_gray_shading(run) is False


def test_has_gray_shading_word_pattern_fill():
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("shaded")
    rPr = OxmlElement("w:rPr")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "pct15")
    shd.set(qn("w:fill"), "FFFFFF")
    rPr.append(shd)
    run._r.insert(0, rPr)

    assert has_gray_shading(run) is True


def test_has_gray_shading_not_gray():
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("yellow highlight")
    rPr = OxmlElement("w:rPr")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "FFFF00")  # yellow, not gray
    rPr.append(shd)
    run._r.insert(0, rPr)

    assert has_gray_shading(run) is False


def test_has_paragraph_gray_shading():
    doc = Document()
    para = doc.add_paragraph("code block para")
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "BFBFBF")
    pPr.append(shd)

    assert has_paragraph_gray_shading(para) is True


def test_has_paragraph_gray_shading_none():
    doc = Document()
    para = doc.add_paragraph("normal para")

    assert has_paragraph_gray_shading(para) is False


# ── CodeBlockAccumulator ───────────────────────────────────────────────


def test_code_block_accumulator_basic():
    acc = CodeBlockAccumulator()
    acc.start()
    acc.add_line("def foo():")
    acc.add_line("    return 42")
    result = acc.finalize()

    assert result is not None
    assert "```" in result
    assert "def foo()" in result
    assert "return 42" in result
    assert not acc.in_code_block


def test_code_block_accumulator_empty():
    acc = CodeBlockAccumulator()
    acc.start()
    result = acc.finalize()

    assert result is None
    assert not acc.in_code_block


def test_code_block_accumulator_not_started():
    acc = CodeBlockAccumulator()
    result = acc.finalize()

    assert result is None


def test_code_block_accumulator_with_indent():
    acc = CodeBlockAccumulator(indent_spaces=4)
    acc.start(list_level=2)
    acc.add_line("code line")
    result = acc.finalize()

    assert result is not None
    # Should have 8-space indent (4 * 2)
    assert "        ```" in result
    assert "        code line" in result


def test_code_block_accumulator_reuse():
    acc = CodeBlockAccumulator()
    acc.start()
    acc.add_line("first block")
    r1 = acc.finalize()
    assert r1 is not None
    assert "first block" in r1

    acc.start()
    acc.add_line("second block")
    r2 = acc.finalize()
    assert r2 is not None
    assert "second block" in r2
