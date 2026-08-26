import pytest
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docwen_core.docx_parsing.break_utils import (
    BorderGroupTracker,
    detect_all_breaks,
    detect_horizontal_rule,
    detect_page_break,
    detect_page_break_in_run,
    detect_section_break,
    extract_paragraph_border_info,
    split_paragraph_by_page_breaks,
)

pytestmark = pytest.mark.unit

# ── Page break detection ───────────────────────────────────────────────


def test_detect_page_break_in_paragraph():
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("before")
    run.add_break(WD_BREAK.PAGE)
    para.add_run("after")

    assert detect_page_break(para) is True


def test_detect_page_break_false_for_normal_paragraph():
    doc = Document()
    para = doc.add_paragraph("just text")

    assert detect_page_break(para) is False


def test_detect_page_break_in_run():
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("before")
    run.add_break(WD_BREAK.PAGE)
    para.add_run("after")

    found = False
    for r in para.runs:
        if detect_page_break_in_run(r):
            found = True
    assert found


def test_detect_page_break_in_run_false_for_normal():
    doc = Document()
    para = doc.add_paragraph("normal")

    for r in para.runs:
        assert detect_page_break_in_run(r) is False


# ── Section break detection ────────────────────────────────────────────


def test_detect_section_break_returns_none_for_normal_paragraph():
    doc = Document()
    para = doc.add_paragraph("normal")

    assert detect_section_break(para) is None


def test_detect_section_break_with_sectPr():
    """Paragraph that carries a w:sectPr child (as Word does for last section)."""
    doc = Document()
    para = doc.add_paragraph("last para")
    # Inject a sectPr element into the paragraph
    sect_pr = OxmlElement("w:sectPr")
    para._p.append(sect_pr)

    result = detect_section_break(para)
    assert result is not None
    assert result[0] == "nextPage"


# ── detect_all_breaks ──────────────────────────────────────────────────


def test_detect_all_breaks_combined():
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("before")
    run.add_break(WD_BREAK.PAGE)
    para.add_run("after")
    # Add a sectPr
    sect_pr = OxmlElement("w:sectPr")
    para._p.append(sect_pr)

    breaks = detect_all_breaks(para)
    types = [b[0] for b in breaks]
    assert "page" in types
    assert "section" in types


def test_detect_all_breaks_none():
    doc = Document()
    para = doc.add_paragraph("nothing")
    assert detect_all_breaks(para) == []


# ── Horizontal rule detection ──────────────────────────────────────────


def test_detect_horizontal_rule_with_bottom_border_empty():
    doc = Document()
    para = doc.add_paragraph("")  # empty text
    # Add a bottom border to the paragraph
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    pBdr.append(bottom)
    pPr.append(pBdr)

    assert detect_horizontal_rule(para) is True


def test_detect_horizontal_rule_border_with_text():
    doc = Document()
    para = doc.add_paragraph("has text")
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Not empty → not a horizontal rule
    assert detect_horizontal_rule(para) is False


def test_extract_paragraph_border_info():
    doc = Document()
    para = doc.add_paragraph("text")
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for direction in ("top", "bottom"):
        elem = OxmlElement(f"w:{direction}")
        elem.set(qn("w:val"), "single")
        pBdr.append(elem)
    pPr.append(pBdr)

    info = extract_paragraph_border_info(para)
    assert info.get("top") == "single"
    assert info.get("bottom") == "single"
    assert "left" not in info


# ── Border group tracker ───────────────────────────────────────────────


def test_border_group_tracker_initial_state():
    tracker = BorderGroupTracker()
    assert not tracker.is_in_group


def test_border_group_tracker_open_and_close():
    tracker = BorderGroupTracker()
    # Simulate a paragraph with top border (opens group)
    result = tracker.process_paragraph(border_info={"top": "single"})
    assert tracker.is_in_group
    assert "---" in result

    # Simulate a paragraph with bottom border (closes group)
    result = tracker.process_paragraph(border_info={"bottom": "single"})
    assert not tracker.is_in_group


def test_border_group_tracker_between_border():
    tracker = BorderGroupTracker()
    tracker.process_paragraph(border_info={"top": "single"})
    assert tracker.is_in_group

    # Between border inside group
    result = tracker.process_paragraph(border_info={"between": "single"})
    assert "---" in result
    assert tracker.is_in_group


def test_border_group_tracker_bottom_only_rule_emits_before_following_paragraph():
    """A Word auto-rule is commonly an empty paragraph with only a bottom border."""
    tracker = BorderGroupTracker(separator="___")

    assert tracker.process_paragraph(border_info={"bottom": "single"}) == []
    assert tracker.is_in_group
    assert tracker.process_paragraph(border_info={}) == ["___"]
    assert not tracker.is_in_group


def test_border_group_tracker_bottom_only_rule_finalizes_at_document_end():
    tracker = BorderGroupTracker(separator="___")

    assert tracker.process_paragraph(border_info={"bottom": "single"}) == []
    assert tracker.finalize() == "___"
    assert not tracker.is_in_group


def test_border_group_tracker_finalize():
    tracker = BorderGroupTracker()
    tracker.process_paragraph(border_info={"top": "single"})
    assert tracker.is_in_group

    result = tracker.finalize()
    assert result == "---"
    assert not tracker.is_in_group


def test_border_group_tracker_empty_separator_omits_output():
    tracker = BorderGroupTracker(separator="")

    result = tracker.process_paragraph(border_info={"top": "single"})
    assert result == []
    assert tracker.is_in_group

    assert tracker.finalize() is None
    assert not tracker.is_in_group


# ── Page-break text splitting ──────────────────────────────────────────


def test_split_paragraph_by_page_breaks_simple():
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("first part")
    run = para.add_run(" ")
    run.add_break(WD_BREAK.PAGE)
    para.add_run("second part")

    result = split_paragraph_by_page_breaks(para)
    # Should contain text before and after the page break, with "---" between
    joined = "".join(result)
    assert "first part" in joined
    assert "---" in result
    assert "second part" in joined


def test_split_paragraph_by_page_breaks_no_break():
    doc = Document()
    para = doc.add_paragraph("plain text, no breaks")

    result = split_paragraph_by_page_breaks(para)
    assert len(result) == 1
    assert "plain text" in result[0]


def test_split_paragraph_by_page_breaks_collapse_adjacent():
    """Adjacent page breaks should collapse to a single separator."""
    doc = Document()
    para = doc.add_paragraph()
    run1 = para.add_run(" ")
    run1.add_break(WD_BREAK.PAGE)
    run2 = para.add_run(" ")
    run2.add_break(WD_BREAK.PAGE)
    para.add_run("after")

    result = split_paragraph_by_page_breaks(para)
    # Should not have consecutive "---"
    sep_count = sum(1 for r in result if r == "---")
    assert sep_count == 1


def test_split_paragraph_by_page_breaks_empty_separator_omits_token():
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("first part")
    run = para.add_run(" ")
    run.add_break(WD_BREAK.PAGE)
    para.add_run("second part")

    result = split_paragraph_by_page_breaks(para, separator="")
    assert result == ["first part", "second part"]
