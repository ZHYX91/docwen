"""Tests for numbering helpers and NumberingIndex methods.

Covers:
- number_to_chinese, number_to_circled (via core)
- _format_counter
- render_numbering_text
- lookup_by_abstract, lookup_by_num_id
"""

from pathlib import Path
from typing import Any, cast

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.styles.style import ParagraphStyle

from docwen_core.text.numbering import (
    number_to_chinese,
    number_to_circled,
    number_to_letter_lower,
    number_to_roman_lower,
)
from docwen_plugin_document.shared.numbering_index import (
    NumberingIndex,
    NumberingLevel,
    _format_counter,
)

pytestmark = pytest.mark.contract

# ── Standalone number helpers ────────────────────────────────────────────


def test_number_helpers_cover_chinese_and_circled():
    assert number_to_chinese(1) == "一"
    assert number_to_chinese(10) == "十"
    assert number_to_chinese(12) == "十二"
    assert number_to_circled(1) == "①"
    assert number_to_circled(10) == "⑩"


def test_number_to_letter_lower():
    assert number_to_letter_lower(1) == "a"
    assert number_to_letter_lower(26) == "z"
    assert number_to_letter_lower(27) == "aa"


def test_number_to_roman_lower():
    assert number_to_roman_lower(1) == "i"
    assert number_to_roman_lower(4) == "iv"
    assert number_to_roman_lower(9) == "ix"
    assert number_to_roman_lower(10) == "x"


# ── _format_counter ──────────────────────────────────────────────────────


@pytest.mark.parametrize(
    "n,fmt,expected",
    [
        (1, "decimal", "1"),
        (3, "decimal", "3"),
        (1, "chineseCountingThousand", "一"),
        (3, "chineseCounting", "三"),
        (1, "lowerLetter", "a"),
        (2, "upperLetter", "B"),
        (1, "lowerRoman", "i"),
        (4, "upperRoman", "IV"),
        (1, "decimalEnclosedCircle", "①"),
        (5, "decimalEnclosedCircle", "⑤"),
        (99, "unknownFmt", "99"),
    ],
)
def test_format_counter(n, fmt, expected):
    assert _format_counter(n, fmt) == expected


def test_format_counter_edge_cases():
    assert _format_counter(0, "decimal") == "0"
    # Core's number_to_chinese returns "" for n <= 0
    assert _format_counter(0, "chineseCountingThousand") == ""
    # Core's number_to_letter_lower returns "" for n <= 0
    assert _format_counter(0, "lowerLetter") == ""
    assert _format_counter(10, "chineseCountingThousand") == "十"
    assert _format_counter(11, "chineseCountingThousand") == "十一"
    assert _format_counter(20, "chineseCountingThousand") == "二十"


# ── Test fixture: a pre-populated NumberingIndex ─────────────────────────


def _make_index() -> NumberingIndex:
    idx = NumberingIndex.__new__(NumberingIndex)
    idx._num_to_abstract = {"10": "0", "20": "1"}
    idx._abstract_num_style_links = {}
    idx._abstract_levels = {
        "0": {
            0: {"numFmt": "chineseCountingThousand", "lvlText": "%1、", "pStyle": ""},
            1: {"numFmt": "decimal", "lvlText": "（%2）", "pStyle": ""},
        },
        "1": {
            0: {"numFmt": "decimal", "lvlText": "%1.", "pStyle": "1Heading1"},
        },
    }
    return idx


def _append_numbering_level(
    abstract_num: object,
    *,
    ilvl: int,
    start: str | None,
    suff: str | None,
    lvl_text: str,
    p_style: str,
) -> None:
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), str(ilvl))
    if start is not None:
        start_element = OxmlElement("w:start")
        start_element.set(qn("w:val"), start)
        level.append(start_element)
    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), "decimal")
    level.append(num_format)
    style = OxmlElement("w:pStyle")
    style.set(qn("w:val"), p_style)
    level.append(style)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), lvl_text)
    level.append(text)
    if suff is not None:
        suffix = OxmlElement("w:suff")
        suffix.set(qn("w:val"), suff)
        level.append(suffix)
    abstract_num.append(level)  # type: ignore[attr-defined]


def _install_numbering_fixture(doc: Any) -> tuple[str, str, str]:
    """Install a real multi-level numbering definition in numbering.xml."""
    level_two = doc.styles.add_style("Level Two Numbered", WD_STYLE_TYPE.PARAGRAPH)
    level_three = doc.styles.add_style("Level Three Numbered", WD_STYLE_TYPE.PARAGRAPH)

    abstract_id = "91001"
    num_id = "92001"
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), abstract_id)
    _append_numbering_level(
        abstract_num,
        ilvl=0,
        start="3",
        suff="nothing",
        lvl_text="%1.",
        p_style=doc.styles["Heading 1"].style_id,
    )
    _append_numbering_level(
        abstract_num,
        ilvl=1,
        start="4",
        suff="space",
        lvl_text="%1.%2)",
        p_style=level_two.style_id,
    )
    _append_numbering_level(
        abstract_num,
        ilvl=2,
        start="5",
        suff="tab",
        lvl_text="%1.%2.%3)",
        p_style=level_three.style_id,
    )
    _append_numbering_level(
        abstract_num,
        ilvl=3,
        start="invalid",
        suff=None,
        lvl_text="%4.",
        p_style="InvalidStartStyle",
    )

    numbering_root = doc.part.numbering_part.element
    first_num_index = next(
        (index for index, child in enumerate(numbering_root) if child.tag == qn("w:num")),
        len(numbering_root),
    )
    numbering_root.insert(first_num_index, abstract_num)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num.append(abstract_ref)
    numbering_root.append(num)
    return num_id, level_two.style_id, level_three.style_id


def _inject_numpr(paragraph: object, *, num_id: str, ilvl: int) -> None:
    paragraph_xml = paragraph._p  # type: ignore[attr-defined]
    properties = paragraph_xml.get_or_add_pPr()
    num_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), str(ilvl))
    identifier = OxmlElement("w:numId")
    identifier.set(qn("w:val"), num_id)
    num_properties.extend((level, identifier))
    properties.append(num_properties)


# ── lookup_by_abstract ───────────────────────────────────────────────────


def test_lookup_by_abstract_finds_existing():
    idx = _make_index()
    result = idx.lookup_by_abstract("0", 0)
    assert result is not None
    assert result.num_fmt == "chineseCountingThousand"
    assert result.lvl_text == "%1、"
    assert result.num_id == "10"  # mapped from _num_to_abstract
    assert result.abstract_num_id == "0"
    assert result.ilvl == 0


def test_lookup_by_abstract_missing_abstract_num_id():
    idx = _make_index()
    result = idx.lookup_by_abstract("99", 0)
    assert result is None


def test_lookup_by_abstract_missing_ilvl():
    idx = _make_index()
    result = idx.lookup_by_abstract("0", 9)
    assert result is None


# ── lookup_by_num_id ─────────────────────────────────────────────────────


def test_lookup_by_num_id_delegates_to_lookup():
    idx = _make_index()
    result = idx.lookup_by_num_id("10", 0)
    assert result is not None
    assert result.num_fmt == "chineseCountingThousand"
    assert result.num_id == "10"


def test_lookup_by_num_id_missing():
    idx = _make_index()
    result = idx.lookup_by_num_id("99", 0)
    assert result is None


# ── lookup_by_style_id ───────────────────────────────────────────────────


def test_lookup_by_style_id_finds_pstyle_match():
    idx = _make_index()
    result = idx.lookup_by_style_id("1Heading1")
    assert result is not None
    assert result.num_id == "20"
    assert result.abstract_num_id == "1"
    assert result.ilvl == 0
    assert result.num_fmt == "decimal"
    assert result.lvl_text == "%1."
    assert result.p_style == "1Heading1"


def test_lookup_by_style_id_no_match():
    idx = _make_index()
    result = idx.lookup_by_style_id("NonExistent")
    assert result is None


# ── render_numbering_text ────────────────────────────────────────────────


def test_render_numbering_text_simple_placeholder():
    idx = _make_index()
    level = NumberingLevel(
        num_id="10",
        abstract_num_id="0",
        ilvl=0,
        num_fmt="chineseCountingThousand",
        lvl_text="%1、",
    )
    result = idx.render_numbering_text(level, 1)
    assert result == "一、 "


def test_render_numbering_text_counter_value():
    idx = _make_index()
    level = NumberingLevel(
        num_id="10",
        abstract_num_id="0",
        ilvl=0,
        num_fmt="chineseCountingThousand",
        lvl_text="%1、",
    )
    result = idx.render_numbering_text(level, 5)
    assert result == "五、 "


def test_render_numbering_text_multi_level():
    idx = _make_index()
    level = NumberingLevel(
        num_id="10",
        abstract_num_id="0",
        ilvl=1,
        num_fmt="decimal",
        lvl_text="%1.%2",
    )
    result = idx.render_numbering_text(level, 3, parent_counters={0: 1})
    # %1 sibling (ilvl 0) has chineseCountingThousand fmt → 一
    # %2 current (ilvl 1) has decimal fmt → 3
    assert result == "一.3 "


def test_render_numbering_text_empty_lvltext_fallback():
    """Empty lvlText falls back to preview_numbering_text."""
    idx = _make_index()
    level = NumberingLevel(
        num_id="10",
        abstract_num_id="0",
        ilvl=0,
        num_fmt="decimal",
        lvl_text="",
    )
    result = idx.render_numbering_text(level, 7)
    assert result == "7 "


def test_render_numbering_text_missing_sibling_level():
    """When a sibling level doesn't exist, counter starts at 1."""
    idx = _make_index()
    # Level with %2 placeholder but ilvl=0 — sibling at ilvl=1 doesn't exist
    level = NumberingLevel(
        num_id="20",
        abstract_num_id="1",
        ilvl=0,
        num_fmt="decimal",
        lvl_text="%1.%2",
    )
    result = idx.render_numbering_text(level, 1)
    # %1 → current ilvl 0 = 1 → "1", %2 → sibling ilvl 1 doesn't exist → val=1 → "1"
    assert result == "1.1 "


def test_render_numbering_text_all_nine_placeholders():
    """lvlText can contain %1 through %9 — all resolve with appropriate fallback."""
    idx = _make_index()
    level = NumberingLevel(
        num_id="10",
        abstract_num_id="0",
        ilvl=0,
        num_fmt="decimal",
        lvl_text="%1.%2.%3.%4.%5.%6.%7.%8.%9",
    )
    # %1 → current ilvl 0 → counter_value=2 → sibling fmt is chineseCountingThousand → "二"
    # %2-%9 → siblings at ilvl 1-8 don't exist (only ilvl 1 exists, rest don't) → val=1, fmt fallback → "1"
    result = idx.render_numbering_text(level, 2)
    assert result == "二.1.1.1.1.1.1.1.1 "


def test_render_numbering_text_sibling_from_parent_counters():
    """Parent counters dict provides sibling values for non-current levels."""
    idx = _make_index()
    # Use a level with multiple placeholders where parent counters supply values
    level = NumberingLevel(
        num_id="10",
        abstract_num_id="0",
        ilvl=1,
        num_fmt="decimal",
        lvl_text="%1.%2",
    )
    result = idx.render_numbering_text(level, 4, parent_counters={0: 2, 1: 4})
    # %1 sibling (ilvl 0) has chineseCountingThousand fmt → 二
    # %2 current (ilvl 1) has decimal fmt → 4
    assert result == "二.4 "


def test_real_numbering_xml_roundtrip_preserves_start_suffix_and_style(tmp_path: Path) -> None:
    doc = Document()
    num_id, level_two_style, level_three_style = _install_numbering_fixture(doc)
    source = tmp_path / "numbering-contract.docx"
    doc.save(str(source))

    reopened = Document(str(source))
    index = NumberingIndex(reopened)
    level_zero = index.lookup(num_id, 0)
    level_one = index.lookup(num_id, 1)
    level_two = index.lookup(num_id, 2)
    invalid_start = index.lookup(num_id, 3)

    assert level_zero is not None
    assert (level_zero.start, level_zero.suff, level_zero.p_style) == (3, "nothing", "Heading1")
    assert index.render_numbering_text(level_zero, level_zero.start) == "3."
    assert level_one is not None
    assert (level_one.start, level_one.suff, level_one.p_style) == (4, "space", level_two_style)
    assert index.render_numbering_text(level_one, level_one.start) == "3.4) "
    assert level_two is not None
    assert (level_two.start, level_two.suff, level_two.p_style) == (5, "tab", level_three_style)
    assert index.render_numbering_text(level_two, level_two.start) == "3.4.5) "
    assert invalid_start is not None
    assert (invalid_start.start, invalid_start.suff) == (1, "tab")
    assert index.render_numbering_text(invalid_start, invalid_start.start) == "1. "
    assert index.lookup_by_style_id(level_two_style) == level_one


def test_real_numbering_xml_drives_exact_markdown_sequence(tmp_path: Path) -> None:
    from docwen_plugin_document.shared.list_processing import ListCounterManager
    from docwen_plugin_document.to_markdown.converter import DocxToMarkdownConverter

    doc = Document()
    num_id, level_two_style, level_three_style = _install_numbering_fixture(doc)
    styles_by_id = {
        style.style_id: cast(ParagraphStyle, style) for style in doc.styles if style.type == WD_STYLE_TYPE.PARAGRAPH
    }
    doc.add_paragraph("Direct heading", style="Heading 1")
    doc.add_paragraph("Second heading", style="Heading 1")
    doc.add_paragraph("Child", style=styles_by_id[level_two_style])
    doc.add_paragraph("Child again", style=styles_by_id[level_two_style])
    doc.add_paragraph("Deep", style=styles_by_id[level_three_style])
    direct = doc.add_paragraph("Direct list")
    _inject_numpr(direct, num_id=num_id, ilvl=1)
    source = tmp_path / "numbering-sequence.docx"
    doc.save(str(source))

    reopened = Document(str(source))
    index = NumberingIndex(reopened)
    counter = ListCounterManager()
    converter = DocxToMarkdownConverter()
    paragraphs = list(reopened.paragraphs)
    para_by_element = {id(paragraph._element): paragraph for paragraph in paragraphs}
    rendered: list[str] = []
    for paragraph in paragraphs:
        lines, _stats = converter._process_paragraph(
            paragraph._element,
            para_by_element=para_by_element,
            numbering_index=index,
            list_counter=counter,
        )
        rendered.extend(lines)

    assert rendered[::2] == [
        "# 3.Direct heading",
        "# 4.Second heading",
        "4.4) Child",
        "4.5) Child again",
        "4.5.5) Deep",
        "    1. Direct list",
    ]
    assert rendered[1::2] == [""] * 6
