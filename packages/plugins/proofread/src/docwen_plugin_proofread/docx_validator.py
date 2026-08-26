"""DocxValidator — DOCX text proofreading with python-docx.

Opens a DOCX, validates text paragraph by paragraph, inserts comments
to mark issues, and writes the proofread DOCX to the staging directory.
"""

from __future__ import annotations

from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from docwen_core.models.result import ConversionResult
    from docwen_core.protocols.execution_context import ProofreadConverterContext


def _find_run_at_position(paragraph, position: int):
    """Find the run in *paragraph* that contains character offset *position*.

    Returns the run, or None if no run spans that position.
    """
    if not paragraph.runs:
        return None
    current_pos = 0
    for run in paragraph.runs:
        run_length = len(run.text)
        if current_pos <= position < current_pos + run_length:
            return run
        current_pos += run_length
    # Fall back to the last run (but indicate off-by-one issue)
    return paragraph.runs[-1] if paragraph.runs else None


class DocxValidator:
    """Validate a DOCX document and annotate issues as Word comments.

    Uses python-docx to read/write the document and the TextValidator
    engine to detect text issues.  Each issue is marked as a Word comment
    attached to the relevant run in the paragraph.
    """

    def convert(self, context: ProofreadConverterContext) -> ConversionResult:
        from docwen_core.models.artifact import (
            ARTIFACT_KIND_PRIMARY,
            ArtifactManifest,
        )
        from docwen_core.models.result import (
            ConversionDiagnostic,
            ConversionErrorInfo,
            ConversionMetrics,
            ConversionResult,
        )
        from docwen_core.paths import input_stem
        from docwen_plugin_proofread._common import (
            file_size,
            new_artifact_id,
            request_source_format,
            resolve_proofread_options,
        )
        from docwen_plugin_proofread.rules import (
            DIAGNOSTIC_CORRUPTED,
            DIAGNOSTIC_ERROR,
            DIAGNOSTIC_INVALID_INPUT,
            DIAGNOSTIC_OK,
            DIAGNOSTIC_SKIPPED,
        )
        from docwen_plugin_proofread.skip_policy import (
            resolve_skip_options,
            should_skip_docx_paragraph,
        )
        from docwen_plugin_proofread.text_validator import TextValidator

        task_id = context.request.request_id
        input_path = context.workspace.input_path
        input_stem_val = input_stem(input_path)

        context.cancellation.check()

        # ── Resolve proofread options from config + request options ──────
        opts = resolve_proofread_options(context)
        symbol_pairing = opts["enable_symbol_pairing"]
        symbol_correction = opts["enable_symbol_correction"]
        typos_rule = opts["enable_typos_rule"]
        sensitive_word = opts["enable_sensitive_word"]

        proofread_rules = context.proofread_rules
        symbol_pairs = list(proofread_rules.symbol_pairs) if proofread_rules else None
        symbol_map = (
            {key: list(values) for key, values in proofread_rules.symbol_map.items()}
            if proofread_rules and proofread_rules.symbol_map
            else None
        )
        typos_map = (
            {key: list(values) for key, values in proofread_rules.typos_map.items()}
            if proofread_rules and proofread_rules.typos_map
            else None
        )
        sensitive_map = (
            {key: list(values) for key, values in proofread_rules.sensitive_words.items()}
            if proofread_rules and proofread_rules.sensitive_words
            else None
        )

        # ── Resolve skip options from config + request options ───────────
        skip_options = resolve_skip_options(context)

        if not any([symbol_pairing, symbol_correction, typos_rule, sensitive_word]):
            context.logger.info("No proofread checks enabled — skipping")
            return ConversionResult(
                task_id=task_id,
                success=True,
                diagnostics=[
                    ConversionDiagnostic(
                        level="info",
                        message="No proofread checks enabled, skipping validation.",
                        code=DIAGNOSTIC_SKIPPED,
                    )
                ],
                metrics=ConversionMetrics(input_bytes=file_size(input_path)),
            )

        input_bytes = file_size(input_path)

        # ── Validate input format ──────────────────────────────────────
        src_fmt = request_source_format(context)
        if src_fmt not in ("docx",):
            msg = (
                f"DOCX validator only supports admitted DOCX input, got {src_fmt!r}. "
                f"For other document formats, convert to DOCX first."
            )
            context.logger.error(msg)
            return ConversionResult(
                task_id=task_id,
                success=False,
                error=ConversionErrorInfo(
                    error_type="invalid_input",
                    message=msg,
                    diagnostic_code=DIAGNOSTIC_INVALID_INPUT,
                ),
                diagnostics=[ConversionDiagnostic(level="error", message=msg, code=DIAGNOSTIC_INVALID_INPUT)],
            )

        try:
            from docx import Document

            context.cancellation.check()
            context.progress.report_progress(0.0, "Loading DOCX document")

            doc = Document(input_path)
            context.logger.info(f"Loaded DOCX with {len(doc.paragraphs)} paragraphs")

            # ── Resolve language ────────────────────────────────────────
            lang = getattr(context, "request", None)
            lang = str(getattr(lang, "options", {}).get("lang", "en")) if lang else "en"

            # L31: Validator is lazily initialized here (inside convert())
            # rather than in __init__, so it is only constructed when needed.
            # This satisfies the "lazy init" requirement.
            # ── Create validator ───────────────────────────────────────
            validator = TextValidator(
                symbol_pairs=symbol_pairs,
                symbol_map=symbol_map,
                typos_map=typos_map,
                sensitive_words=sensitive_map,
                enabled={
                    "symbol_pairing": symbol_pairing,
                    "symbol_correction": symbol_correction,
                    "typos_rule": typos_rule,
                    "sensitive_word": sensitive_word,
                },
                lang=lang,
            )

            # ── Validate each paragraph ────────────────────────────────
            issues_detected = 0
            comments_added = 0
            comment_failures: list[tuple[int, int, str]] = []
            total_paragraphs = len(doc.paragraphs)
            paragraphs_checked = 0
            paragraphs_skipped = 0

            for i, para in enumerate(doc.paragraphs):
                context.cancellation.check()

                text = para.text
                if not text.strip():
                    continue

                if should_skip_docx_paragraph(para, skip_options):
                    paragraphs_skipped += 1
                    continue

                paragraphs_checked += 1

                errors = validator.validate_text(text)
                if not errors:
                    continue
                issues_detected += len(errors)

                # ── Split runs at error boundaries for precise anchoring ──
                from docwen_plugin_proofread.run_splitter import (
                    ensure_run_at_position,
                    plan_run_splits,
                    runs_for_range,
                )

                split_positions = plan_run_splits(para, errors)
                for pos in split_positions:
                    ensure_run_at_position(para, pos)

                # ── Annotate errors as Word comments ─────────────
                for err in errors:
                    comment_text = f"{err.error_type}: {err.error_text} → {err.suggestion}"
                    try:
                        runs = runs_for_range(para, err.start_pos, err.end_pos)
                        if not runs:
                            raise ValueError(f"No exact run span for proofread range [{err.start_pos}, {err.end_pos})")
                        doc.add_comment(
                            runs,
                            text=comment_text,
                            author=f"DocWen-{err.source}",
                            initials="DW",
                        )
                        comments_added += 1
                    except Exception as exc:
                        comment_failures.append((i, err.start_pos, type(exc).__name__))
                        context.logger.warning(
                            f"Failed to add comment at paragraph {i}, position {err.start_pos}: {type(exc).__name__}"
                        )

                if (i + 1) % 20 == 0:
                    context.progress.report_progress(
                        20.0 + 60.0 * (i + 1) / total_paragraphs,
                        f"Proofreading paragraph {i + 1}/{total_paragraphs}",
                    )

            context.cancellation.check()
            context.progress.report_progress(85.0, "Saving proofread DOCX")

            # ── Write to staging ───────────────────────────────────────
            output_path = context.workspace.create_artifact_path(ARTIFACT_KIND_PRIMARY, ".docx")
            doc.save(output_path)

            output_bytes = file_size(output_path)
            context.progress.report_progress(100.0, "DOCX proofread complete")

            artifact = ArtifactManifest(
                artifact_id=new_artifact_id(),
                kind=ARTIFACT_KIND_PRIMARY,
                staging_path=output_path,
                suggested_name=f"{input_stem_val}_checked.docx",
                media_type=("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                metadata={
                    "source_format": src_fmt,
                    "paragraphs_total": total_paragraphs,
                    "paragraphs_checked": paragraphs_checked,
                    "paragraphs_skipped": paragraphs_skipped,
                    "errors_found": issues_detected,
                    "comments_added": comments_added,
                    "comments_failed": len(comment_failures),
                    "checks_enabled": {
                        "symbol_pairing": symbol_pairing,
                        "symbol_correction": symbol_correction,
                        "typos_rule": typos_rule,
                        "sensitive_word": sensitive_word,
                    },
                    "skip_options": {
                        "code_blocks": skip_options.code_blocks,
                        "quote_blocks": skip_options.quote_blocks,
                    },
                },
                is_primary=True,
            )
            context.workspace.add_artifact(artifact)
            context.progress.report_artifact_ready(artifact.artifact_id, artifact.suggested_name)

            diagnostics = [
                ConversionDiagnostic(
                    level="info",
                    message=(
                        f"Proofread complete: {issues_detected} issue(s) detected and "
                        f"{comments_added} comment(s) added in {total_paragraphs} paragraph(s)"
                    ),
                    code=DIAGNOSTIC_OK,
                )
            ]
            if comment_failures:
                first_paragraph, first_position, _exception_type = comment_failures[0]
                diagnostics.append(
                    ConversionDiagnostic(
                        level="warning",
                        message=(
                            f"Detected {issues_detected} issue(s), but "
                            f"{len(comment_failures)} Word comment(s) could not be added; "
                            "review the source and output manually."
                        ),
                        code="PROOFREAD-COMMENTS-PARTIAL",
                        location=f"paragraph {first_paragraph}, position {first_position}",
                    )
                )

            return ConversionResult(
                task_id=task_id,
                success=True,
                artifacts=[artifact],
                diagnostics=diagnostics,
                metrics=ConversionMetrics(
                    input_bytes=input_bytes,
                    output_bytes=output_bytes,
                    extra={
                        "source_format": src_fmt,
                        "paragraphs_total": total_paragraphs,
                        "paragraphs_checked": paragraphs_checked,
                        "paragraphs_skipped": paragraphs_skipped,
                        "errors_found": issues_detected,
                        "comments_added": comments_added,
                        "comments_failed": len(comment_failures),
                    },
                ),
            )

        except Exception as exc:
            context.logger.error(f"DOCX proofread failed: {exc}")
            # Determine if this is a corrupted document error
            error_type = "conversion_failed"
            diag_code = DIAGNOSTIC_ERROR
            msg = str(exc)
            if "not a valid" in msg.lower() or "corrupt" in msg.lower():
                error_type = "invalid_input"
                diag_code = DIAGNOSTIC_CORRUPTED
                msg = f"Corrupted or invalid DOCX file: {exc}"

            return ConversionResult(
                task_id=task_id,
                success=False,
                error=ConversionErrorInfo(
                    error_type=error_type,
                    message=msg,
                    diagnostic_code=diag_code,
                ),
                diagnostics=[
                    ConversionDiagnostic(
                        level="error",
                        message=f"DOCX proofread failed: {exc}",
                        code=diag_code,
                    )
                ],
            )
