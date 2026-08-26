"""Focused tests split from test_proofread_plugin.py."""

from __future__ import annotations

from ._proofread_plugin_support import (
    Path,
    _build_fake_context,
    _create_test_docx,
    _expected_text_issues_from_fixture,
    _load_proofread_old_system_fixture,
    _load_real_dictionary_rules,
    _normalize_text_error,
    base64,
    json,
    pytest,
    zipfile,
)

pytestmark = pytest.mark.golden


@pytest.mark.contract
def test_docx_comment_failures_preserve_issue_count_and_emit_warning(tmp_path, monkeypatch) -> None:
    """Detected issues stay visible when Word comment insertion degrades."""
    from docx.document import Document as DocumentClass

    from docwen_plugin_proofread.docx_validator import DocxValidator

    source = tmp_path / "comment-failure.docx"
    staging = tmp_path / "staging"
    staging.mkdir()
    _create_test_docx(str(source), ["价格是１２３元。"])

    def _fail_add_comment(self, *_args, **_kwargs):
        raise RuntimeError("comment part unavailable")

    monkeypatch.setattr(DocumentClass, "add_comment", _fail_add_comment)
    context = _build_fake_context(
        str(source),
        str(staging),
        target_format="docx",
        action_name="validate",
        source_format="docx",
        options={
            "enable_symbol_correction": True,
            "enable_symbol_pairing": False,
            "enable_typos_rule": False,
            "enable_sensitive_word": False,
        },
    )

    result = DocxValidator().convert(context)

    assert result.success is True
    artifact = result.artifacts[0]
    assert artifact.metadata["errors_found"] == 3
    assert artifact.metadata["comments_added"] == 0
    assert artifact.metadata["comments_failed"] == 3
    assert result.metrics.extra["errors_found"] == 3
    warning = next(item for item in result.diagnostics if item.code == "PROOFREAD-COMMENTS-PARTIAL")
    assert warning.level == "warning"
    assert "3 Word comment(s)" in warning.message


@pytest.mark.contract
@pytest.mark.parametrize(
    "run_texts",
    [("AASecret42ZZ",), ("AA", "Se", "cret", "42", "ZZ")],
    ids=["single-run-interior", "cross-run"],
)
def test_docx_validator_saved_comment_anchor_covers_exact_half_open_error_range(
    tmp_path: Path,
    run_texts: tuple[str, ...],
) -> None:
    from docx import Document

    from docwen_core.models.proofread import ProofreadRules
    from docwen_plugin_proofread.anchor_report import extract_occurrences_from_document_xml, read_docx_part
    from docwen_plugin_proofread.docx_validator import DocxValidator

    source = tmp_path / "exact-anchor.docx"
    staging = tmp_path / "staging"
    staging.mkdir()
    document = Document()
    paragraph = document.add_paragraph()
    for text in run_texts:
        paragraph.add_run(text)
    document.save(str(source))

    context = _build_fake_context(
        str(source),
        str(staging),
        target_format="docx",
        action_name="validate",
        source_format="docx",
        options={
            "enable_symbol_pairing": False,
            "enable_symbol_correction": False,
            "enable_typos_rule": False,
            "enable_sensitive_word": True,
        },
        proofread_rules=ProofreadRules(sensitive_words={"Secret42": ()}),
    )

    result = DocxValidator().convert(context)

    assert result.success is True
    assert result.artifacts[0].metadata["comments_added"] == 1
    document_xml = read_docx_part(Path(result.artifacts[0].staging_path), "word/document.xml")
    assert document_xml is not None
    occurrences, diagnostics = extract_occurrences_from_document_xml(document_xml, 20, False)
    assert diagnostics.cross_paragraph == []
    assert diagnostics.end_without_start_ids == []
    assert diagnostics.start_without_end_ids == []
    assert [(item.start, item.end, item.covered_text) for item in occurrences] == [(2, 10, "Secret42")]


@pytest.mark.contract
def test_docx_validator_preserves_drawing_and_anchors_adjacent_text_exactly(tmp_path: Path) -> None:
    from docx import Document
    from docx.shared import Inches

    from docwen_core.models.proofread import ProofreadRules
    from docwen_plugin_proofread.anchor_report import (
        extract_comment_texts_from_comments_xml,
        extract_occurrences_from_document_xml,
        read_docx_part,
    )
    from docwen_plugin_proofread.docx_validator import DocxValidator

    image_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4nGNgYAAAAAMAAWgmWQ0AAAAASUVORK5CYII="
    )
    image = tmp_path / "pixel.png"
    image.write_bytes(image_bytes)
    source = tmp_path / "drawing-anchor.docx"
    staging = tmp_path / "staging"
    staging.mkdir()
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("含图：")
    paragraph.add_run().add_picture(str(image), width=Inches(0.2))
    paragraph.add_run("才料（")
    document.save(str(source))

    context = _build_fake_context(
        str(source),
        str(staging),
        target_format="docx",
        action_name="validate",
        source_format="docx",
        options={
            "enable_symbol_pairing": True,
            "enable_symbol_correction": False,
            "enable_typos_rule": True,
            "enable_sensitive_word": False,
        },
        proofread_rules=ProofreadRules(
            symbol_pairs=(("（", "）"),),
            typos_map={"材料": ("才料",)},
        ),
    )

    result = DocxValidator().convert(context)

    assert result.success is True
    assert result.artifacts[0].metadata["comments_added"] == 2
    output = Path(result.artifacts[0].staging_path)
    with zipfile.ZipFile(output) as archive:
        media = [archive.read(name) for name in archive.namelist() if name.startswith("word/media/")]
    assert media == [image_bytes]
    document_xml = read_docx_part(output, "word/document.xml")
    comments_xml = read_docx_part(output, "word/comments.xml")
    assert document_xml is not None and comments_xml is not None
    occurrences, diagnostics = extract_occurrences_from_document_xml(document_xml, 20, False)
    comments = extract_comment_texts_from_comments_xml(comments_xml, False)
    assert not diagnostics.cross_paragraph
    assert not diagnostics.end_without_start_ids
    assert not diagnostics.start_without_end_ids
    covered_by_comment = {comments[item.comment_id]: item.covered_text for item in occurrences}
    assert any("才料" in text and covered == "才料" for text, covered in covered_by_comment.items())
    assert any("（" in text and covered == "（" for text, covered in covered_by_comment.items())


@pytest.mark.contract
class TestOldSystemProofreadFixture:
    def test_text_validator_matches_old_system_semantic_fixture(self) -> None:
        """Current text engine preserves the old systems' normalized proofread semantics."""
        from docwen_plugin_proofread.text_validator import TextValidator

        fixture = _load_proofread_old_system_fixture()
        rules = fixture["rules"]
        validator = TextValidator(
            symbol_pairs=[tuple(pair) for pair in rules["symbol_pairs"]],
            symbol_map={key: list(values) for key, values in rules["symbol_map"].items()},
            typos_map={key: list(values) for key, values in rules["typos_map"].items()},
            sensitive_words={key: list(values) for key, values in rules["sensitive_words"].items()},
            enabled={
                "symbol_pairing": rules["checks_enabled"]["symbol_pairing"],
                "symbol_correction": rules["checks_enabled"]["symbol_correction"],
                "typos_rule": rules["checks_enabled"]["typos_rule"],
                "sensitive_word": rules["checks_enabled"]["sensitive_word"],
            },
            lang="en",
        )

        actual = [_normalize_text_error(error) for error in validator.validate_text(fixture["input_text"])]

        assert actual == _expected_text_issues_from_fixture(fixture)

    def test_text_validator_preserves_pyside6_multi_pattern_regressions(self) -> None:
        """Current text engine keeps the PySide6 multi-pattern scan improvements."""
        from docwen_plugin_proofread.text_validator import TextValidator

        fixture = _load_proofread_old_system_fixture()

        for case in fixture["old_pyside6_multi_pattern_regression_cases"]:
            rules = case["rules"]
            validator = TextValidator(
                symbol_pairs=[tuple(pair) for pair in rules.get("symbol_pairs", [])],
                symbol_map={key: list(values) for key, values in rules.get("symbol_map", {}).items()},
                typos_map={key: list(values) for key, values in rules.get("typos_map", {}).items()},
                sensitive_words={key: list(values) for key, values in rules.get("sensitive_words", {}).items()},
                enabled=case["checks_enabled"],
                lang="en",
            )

            actual = [_normalize_text_error(error) for error in validator.validate_text(case["input_text"])]

            assert actual == case["expected_issues"], case["case_id"]

    def test_markdown_sanitizer_preserves_old_system_text_slicer_semantics(self) -> None:
        """Current Markdown sanitizer keeps the old md_spell/text_slicer behavior."""
        from docwen_plugin_proofread.md_validator import _sanitize_markdown

        fixture = _load_proofread_old_system_fixture()
        cases = {case["case_id"]: case for case in fixture["old_system_markdown_sanitizer_cases"]}

        visible_case = cases["visible_link_text_preserved_markers_and_code_blanked"]
        sanitized = _sanitize_markdown(visible_case["input_markdown"])
        for text in visible_case["sanitized_contains"]:
            assert text in sanitized.sanitized_text
        for text in visible_case["sanitized_absent"]:
            assert text not in sanitized.sanitized_text

        position_case = cases["link_text_offset_uses_original_file_line_numbers"]
        sanitized = _sanitize_markdown(position_case["input_markdown"])
        offset = sanitized.sanitized_text.index(position_case["needle"])
        line, col = sanitized.offset_to_contract_line_col(offset)
        assert (line, col) == (position_case["expected_line"] - 1, position_case["expected_col"] - 1)

    def test_real_dictionary_official_paragraph_matches_pyside6_projection(self, tmp_path: Path) -> None:
        """A file-backed typo dictionary preserves the corrected PySide6 projection."""
        from docwen_plugin_proofread.text_validator import TextValidator

        fixture = _load_proofread_old_system_fixture()["real_dictionary_official_docx_probe"]
        rules = _load_real_dictionary_rules(tmp_path / "user-config")
        assert rules.typos_map == {"分隔线": ("分割线",)}

        validator = TextValidator(
            symbol_pairs=[],
            symbol_map={},
            typos_map={key: list(values) for key, values in rules.typos_map.items()},
            sensitive_words={},
            enabled={
                "symbol_pairing": False,
                "symbol_correction": False,
                "typos_rule": True,
                "sensitive_word": False,
            },
            lang="en",
        )
        errors = validator.validate_text(fixture["markdown_projection"]["offending_paragraph_text"])

        assert [
            {
                "rule_key": "typo",
                "start_pos": error.start_pos,
                "end_pos": error.end_pos,
                "error_text": error.error_text,
                "suggestion": error.suggestion,
            }
            for error in errors
        ] == [
            {
                "rule_key": fixture["expected_normalized_issue"]["rule_key"],
                "start_pos": fixture["expected_normalized_issue"]["col_start"] - 1,
                "end_pos": fixture["expected_normalized_issue"]["col_end"],
                "error_text": fixture["expected_normalized_issue"]["error_text"],
                "suggestion": fixture["expected_normalized_issue"]["suggestion"],
            }
        ]

    def test_real_dictionary_official_paragraph_reaches_docx_and_markdown_carriers(self, tmp_path: Path) -> None:
        """The file-backed rule reaches DOCX comments and Markdown line/column JSON."""
        from docwen_plugin_proofread.anchor_report import _extract_comments
        from docwen_plugin_proofread.docx_validator import DocxValidator
        from docwen_plugin_proofread.md_validator import MarkdownValidator

        fixture = _load_proofread_old_system_fixture()["real_dictionary_official_docx_probe"]
        expected = fixture["expected_normalized_issue"]
        paragraph = fixture["markdown_projection"]["offending_paragraph_text"]
        rules = _load_real_dictionary_rules(tmp_path / "user-config")
        options = {
            "enable_symbol_pairing": False,
            "enable_symbol_correction": False,
            "enable_typos_rule": True,
            "enable_sensitive_word": False,
        }

        docx_path = tmp_path / "official-shape.docx"
        _create_test_docx(str(docx_path), ["占位"] * 79 + [paragraph])
        docx_staging = tmp_path / "docx-staging"
        docx_staging.mkdir()
        docx_context = _build_fake_context(
            str(docx_path),
            str(docx_staging),
            target_format="document",
            action_name="validate",
            source_format="docx",
            options=options,
            proofread_rules=rules,
        )
        docx_result = DocxValidator().convert(docx_context)
        comments = _extract_comments(Path(docx_result.artifacts[0].staging_path))
        assert docx_result.success is True
        assert len(comments) == 1
        assert comments[0].anchor_paragraph_index == expected["docx_anchor_paragraph_index"]
        assert expected["error_text"] in comments[0].text
        assert expected["suggestion"] in comments[0].text

        markdown_path = tmp_path / "official-body.md"
        markdown_path.write_text("\n".join(["占位"] * 72 + [paragraph]) + "\n", encoding="utf-8")
        markdown_staging = tmp_path / "markdown-staging"
        markdown_staging.mkdir()
        markdown_context = _build_fake_context(
            str(markdown_path),
            str(markdown_staging),
            target_format="markdown",
            action_name="validate",
            source_format="markdown",
            options=options,
            proofread_rules=rules,
        )
        markdown_result = MarkdownValidator().convert(markdown_context)
        report = json.loads(Path(markdown_result.artifacts[0].staging_path).read_text(encoding="utf-8"))
        assert markdown_result.success is True
        assert len(report["issues"]) == 1
        issue = report["issues"][0]
        assert {
            "range": issue["range"],
            "error_text": issue["error_text"],
            "suggestion": issue["suggestion"],
            "error_type": issue["error_type"],
            "source": issue["source"],
            "rule_key": issue["rule_key"],
        } == {
            "range": {
                "start": {
                    "offset": issue["range"]["start"]["offset"],
                    "line": expected["line"] - 1,
                    "column": expected["col_start"] - 1,
                },
                "end": {
                    "offset": issue["range"]["end"]["offset"],
                    "line": expected["line"] - 1,
                    "column": expected["col_end"],
                },
            },
            "error_text": expected["error_text"],
            "suggestion": expected["suggestion"],
            "error_type": "Typo",
            "source": "typo",
            "rule_key": "typo",
        }
        assert issue["matched_text"] == expected["error_text"]
        assert issue["fix"] == {
            "kind": "replace_text",
            "replacement": expected["suggestion"],
            "applicable": True,
        }
