"""Focused tests split from test_proofread_plugin.py."""

from __future__ import annotations

from ._proofread_plugin_support import (
    _PROJECT_ROOT,
    Any,
    Path,
    build_document_style_catalog,
    dataclass,
    field,
    pytest,
)

pytestmark = pytest.mark.golden


@pytest.mark.integration
class TestProofreadConfigWiring:
    """Verify the proofread plugin reads config from context.config (config wiring).

    These tests exercise the real ``RuntimeExecutionContext`` → ``ReadOnlyConfigView``
    path that plugins receive in production, not ad-hoc fake objects.
    """

    @staticmethod
    def _make_context(
        engine_config: dict[str, Any] | None = None,
        request_options: dict[str, Any] | None = None,
        extra_config_snapshot: dict[str, Any] | None = None,
    ) -> Any:
        """Create a ``RuntimeExecutionContext`` with the given engine config.

        This mirrors the production path: ConfigLoader.as_dict() →
        adapter config_snapshot → RuntimeExecutionContext → ReadOnlyConfigView.
        """
        from dataclasses import dataclass, field

        from docwen_core.cancellation import CancellationToken
        from docwen_core.export_semantics import MarkdownExportSemantics
        from docwen_core.models.file_ref import FileRef
        from docwen_core.models.request import ConversionRequest, OutputPolicy
        from docwen_runtime._execution_context import RuntimeExecutionContext

        config_snapshot: dict[str, Any] = dict(extra_config_snapshot or {})
        if engine_config is not None:
            config_snapshot["proofread"] = {"engine": engine_config}

        @dataclass
        class _MinimalWorkspace:
            input_path: str = "/tmp/test.md"
            staging_dir: str = "/tmp/staging"
            _counter: list[int] = field(default_factory=lambda: [0])

            def create_artifact_path(self, kind: str, suffix: str) -> str:
                self._counter[0] += 1
                return f"/tmp/staging/{kind}_{self._counter[0]}{suffix}"

            def add_artifact(self, manifest: Any) -> None:
                pass

        req = ConversionRequest(
            request_id="test-pr-001",
            input_refs=[FileRef(path="/tmp/test.md", format="markdown", category="text")],
            target_format="markdown",
            action_name="validate",
            options=dict(request_options or {}),
            output_policy=OutputPolicy(),
            config_snapshot=config_snapshot,
        )

        return RuntimeExecutionContext(
            request=req,
            workspace=_MinimalWorkspace(),  # type: ignore[arg-type]
            config_snapshot=config_snapshot,
            cancellation_token=CancellationToken(),
            document_style_catalog=build_document_style_catalog(
                config_snapshot,
                request_options=request_options,
                locales_dir=_PROJECT_ROOT / "i18n" / "locales",
            ),
            markdown_export_semantics=MarkdownExportSemantics.from_config_snapshot(config_snapshot),
        )

    def test_resolve_proofread_options_reads_from_config(self) -> None:
        """resolve_proofread_options must read engine settings from the
        real ReadOnlyConfigView (via RuntimeExecutionContext)."""
        from docwen_plugin_proofread._common import resolve_proofread_options

        ctx = self._make_context(
            engine_config={
                "enable_symbol_pairing": False,
                "enable_symbol_correction": True,
                "enable_typos_rule": False,
                "enable_sensitive_word": True,
            }
        )

        opts = resolve_proofread_options(ctx)
        assert opts["enable_symbol_pairing"] is False
        assert opts["enable_symbol_correction"] is True
        assert opts["enable_typos_rule"] is False
        assert opts["enable_sensitive_word"] is True

    def test_request_options_override_config_defaults(self) -> None:
        """Request options must take precedence over config defaults."""
        from docwen_plugin_proofread._common import resolve_proofread_options

        ctx = self._make_context(
            engine_config={
                "enable_symbol_pairing": True,
                "enable_typos_rule": True,
            },
            request_options={
                "enable_symbol_pairing": False,  # override
            },
        )

        opts = resolve_proofread_options(ctx)
        # Request overrides config
        assert opts["enable_symbol_pairing"] is False
        # Not in request → falls back to config
        assert opts["enable_typos_rule"] is True
        # Not in either → defaults
        assert opts["enable_symbol_correction"] is True

    def test_extra_options_take_highest_precedence(self) -> None:
        """Explicit extra_options must override both config and request options."""
        from docwen_plugin_proofread._common import resolve_proofread_options

        ctx = self._make_context(
            engine_config={"enable_symbol_pairing": True},
            request_options={"enable_symbol_pairing": False},
        )

        opts = resolve_proofread_options(ctx, extra_options={"enable_symbol_pairing": True})
        # Extra options override everything
        assert opts["enable_symbol_pairing"] is True

    def test_missing_engine_config_uses_defaults(self) -> None:
        """When engine key is missing from config, defaults are used."""
        from docwen_plugin_proofread._common import resolve_proofread_options

        ctx = self._make_context()  # No engine_config → no "engine" key

        opts = resolve_proofread_options(ctx)
        assert opts["enable_symbol_pairing"] is True  # default
        assert opts["enable_symbol_correction"] is True  # default
        assert opts["enable_typos_rule"] is True  # default
        assert opts["enable_sensitive_word"] is True  # default (matches DEFAULT_PROOFREAD_SETTINGS_TOML)

    def test_config_loader_to_context_user_path(self) -> None:
        """End-to-end: ConfigLoader → config_snapshot → RuntimeExecutionContext
        → resolve_proofread_options.  This is the production path."""
        import tempfile

        from docwen_plugin_proofread._common import resolve_proofread_options
        from docwen_runtime.config.loader import ConfigLoader

        with tempfile.TemporaryDirectory() as tmpdir:
            loader = ConfigLoader(
                base_dir=Path(__file__).resolve().parent.parent.parent.parent.parent / "configs",
                user_dir=Path(tmpdir),
            )
            config_dict = loader.config.as_dict()

            # Simulate the production pipeline:
            # ConfigLoader → adapter config_snapshot → RuntimeExecutionContext
            ctx = self._make_context(
                extra_config_snapshot=config_dict,
                # config_dict already has "proofread.engine" section from
                # the three-layer merge of base configs
            )

            opts = resolve_proofread_options(ctx)
            # All defaults from DEFAULT_PROOFREAD_SETTINGS_TOML are True
            assert opts["enable_symbol_pairing"] is True
            assert opts["enable_symbol_correction"] is True
            assert opts["enable_typos_rule"] is True
            assert opts["enable_sensitive_word"] is True

    def test_docx_validator_uses_resolve_proofread_options(self) -> None:
        """DocxValidator.convert must use resolve_proofread_options
        via a real RuntimeExecutionContext."""
        import os
        import tempfile
        from pathlib import Path

        from docwen_core.cancellation import CancellationToken
        from docwen_core.export_semantics import MarkdownExportSemantics
        from docwen_core.models.file_ref import FileRef
        from docwen_core.models.request import ConversionRequest, OutputPolicy
        from docwen_plugin_proofread.docx_validator import DocxValidator
        from docwen_runtime._execution_context import RuntimeExecutionContext

        with tempfile.TemporaryDirectory() as staging:
            docx_path = os.path.join(staging, "test.docx")
            from docx import Document

            doc = Document()
            doc.add_paragraph("Clean text.")
            doc.save(docx_path)

            @dataclass
            class _RealishWorkspace:
                input_path: str
                staging_dir: str
                _counter: list[int] = field(default_factory=lambda: [0])
                _artifacts: list[Any] = field(default_factory=list)

                def create_artifact_path(self, kind: str, suffix: str) -> str:
                    self._counter[0] += 1
                    return str(Path(self.staging_dir) / f"{kind}_{self._counter[0]}{suffix}")

                def add_artifact(self, manifest: Any) -> None:
                    self._artifacts.append(manifest)

            ctx = RuntimeExecutionContext(
                request=ConversionRequest(
                    request_id="test-config-001",
                    input_refs=[FileRef(path=docx_path, format="docx", category="document")],
                    target_format="document",
                    action_name="validate",
                    options={},  # No request options — must use config
                    output_policy=OutputPolicy(),
                    config_snapshot={
                        "engine": {
                            "enable_symbol_pairing": True,
                            "enable_symbol_correction": False,
                            "enable_typos_rule": False,
                            "enable_sensitive_word": False,
                        },
                    },
                ),
                workspace=_RealishWorkspace(docx_path, staging),  # type: ignore[arg-type]
                config_snapshot={
                    "engine": {
                        "enable_symbol_pairing": True,
                        "enable_symbol_correction": False,
                        "enable_typos_rule": False,
                        "enable_sensitive_word": False,
                    },
                },
                cancellation_token=CancellationToken(),
                document_style_catalog=build_document_style_catalog(
                    {},
                    locales_dir=_PROJECT_ROOT / "i18n" / "locales",
                ),
                markdown_export_semantics=MarkdownExportSemantics(),
            )

            result = DocxValidator().convert(ctx)
            # Should succeed — only symbol_pairing is enabled
            assert result.success is True

    def test_md_validator_uses_resolve_proofread_options(self) -> None:
        """MarkdownValidator.convert must use resolve_proofread_options
        via a real RuntimeExecutionContext."""
        import os
        import tempfile
        from pathlib import Path

        from docwen_core.cancellation import CancellationToken
        from docwen_core.export_semantics import MarkdownExportSemantics
        from docwen_core.models.file_ref import FileRef
        from docwen_core.models.request import ConversionRequest, OutputPolicy
        from docwen_plugin_proofread.md_validator import MarkdownValidator
        from docwen_runtime._execution_context import RuntimeExecutionContext

        with tempfile.TemporaryDirectory() as staging:
            md_path = os.path.join(staging, "test.md")
            Path(md_path).write_text("Clean markdown text.", encoding="utf-8")

            @dataclass
            class _RealishWorkspace:
                input_path: str
                staging_dir: str
                _counter: list[int] = field(default_factory=lambda: [0])
                _artifacts: list[Any] = field(default_factory=list)

                def create_artifact_path(self, kind: str, suffix: str) -> str:
                    self._counter[0] += 1
                    return str(Path(self.staging_dir) / f"{kind}_{self._counter[0]}{suffix}")

                def add_artifact(self, manifest: Any) -> None:
                    self._artifacts.append(manifest)

            ctx = RuntimeExecutionContext(
                request=ConversionRequest(
                    request_id="test-config-md-001",
                    input_refs=[FileRef(path=md_path, format="markdown", category="text")],
                    target_format="markdown",
                    action_name="validate",
                    options={},  # No request options — must use config
                    output_policy=OutputPolicy(),
                    config_snapshot={
                        "engine": {
                            "enable_symbol_pairing": True,
                            "enable_symbol_correction": False,
                            "enable_typos_rule": False,
                            "enable_sensitive_word": False,
                        },
                    },
                ),
                workspace=_RealishWorkspace(md_path, staging),  # type: ignore[arg-type]
                config_snapshot={
                    "engine": {
                        "enable_symbol_pairing": True,
                        "enable_symbol_correction": False,
                        "enable_typos_rule": False,
                        "enable_sensitive_word": False,
                    },
                },
                cancellation_token=CancellationToken(),
                document_style_catalog=build_document_style_catalog(
                    {},
                    locales_dir=_PROJECT_ROOT / "i18n" / "locales",
                ),
                markdown_export_semantics=MarkdownExportSemantics(),
            )

            result = MarkdownValidator().convert(ctx)
            assert result.success is True


@pytest.mark.contract
class TestSkipPolicyIntegration:
    def test_docx_validator_skips_code_style_paragraphs(self, tmp_path):
        from types import SimpleNamespace

        from docx import Document

        from docwen_plugin_proofread.docx_validator import DocxValidator

        input_path = tmp_path / "input.docx"
        staging_dir = tmp_path / "staging"
        staging_dir.mkdir()

        doc = Document()
        code_style = doc.styles.add_style("Code Block", 1)
        if code_style is not None:
            doc.add_paragraph("（", style=str(code_style.name))
        else:
            doc.add_paragraph("（")
        doc.add_paragraph("（")
        doc.save(str(input_path))

        artifacts = []

        class Workspace:
            def __init__(self):
                self.input_path = str(input_path)
                self.staging_dir = str(staging_dir)

            def create_artifact_path(self, kind, suffix):
                return str(staging_dir / f"artifact{suffix}")

            def add_artifact(self, artifact):
                artifacts.append(artifact)

        context = SimpleNamespace(
            request=SimpleNamespace(
                request_id="req",
                options={},
                input_refs=[SimpleNamespace(path=str(input_path), format="docx")],
            ),
            workspace=Workspace(),
            config=SimpleNamespace(
                get=lambda self, k, d=None: d,
                get_plugin_config=lambda self, pid: {},
            ),
            cancellation=SimpleNamespace(check=lambda: None),
            progress=SimpleNamespace(
                report_progress=lambda *args, **kwargs: None,
                report_diagnostic=lambda *args, level=None, m=None, c=None, loc=None, **kwargs: None,
                report_artifact_ready=lambda *args, **kwargs: None,
            ),
            logger=SimpleNamespace(
                debug=lambda *args, **kwargs: None,
                info=lambda *args, **kwargs: None,
                error=lambda *args, **kwargs: None,
                warning=lambda *args, **kwargs: None,
            ),
            proofread_rules=None,
        )

        result = DocxValidator().convert(context)  # type: ignore[arg-type]

        assert result.success is True
        assert artifacts[0].metadata["paragraphs_checked"] == 1
        assert artifacts[0].metadata["paragraphs_skipped"] == 1
        assert artifacts[0].metadata["skip_options"] == {
            "code_blocks": True,
            "quote_blocks": False,
        }
        assert result.metrics.extra["paragraphs_skipped"] == 1
