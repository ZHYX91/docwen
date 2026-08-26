"""Tests for run_splitter — run boundary splitting for precise comment anchoring.

These tests use python-docx to create real paragraphs and verify that
``plan_run_splits`` and ``ensure_run_at_position`` correctly split runs
at given character positions.
"""

from __future__ import annotations

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docwen_plugin_proofread.run_splitter import (
    ensure_run_at_position,
    plan_run_splits,
    runs_for_range,
)
from docwen_plugin_proofread.text_validator import TextError

pytestmark = pytest.mark.unit

# ═══════════════════════════════════════════════════════════════════════
# Fixtures
# ═══════════════════════════════════════════════════════════════════════


@pytest.fixture
def doc():
    """A fresh empty document."""
    return Document()


@pytest.fixture
def simple_para(doc):
    """A paragraph with a single run containing 10 chars."""
    p = doc.add_paragraph("ABCDEFGHIJ")
    assert len(p.runs) == 1
    assert p.text == "ABCDEFGHIJ"
    return p


def _make_multi_run_para(doc, texts, para_text=None):
    """Create a paragraph with multiple runs, each with given *texts*.

    Returns the paragraph.  Its ``.text`` will be the concatenation.
    """
    p = doc.add_paragraph()
    # Remove the empty run that add_paragraph creates
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    for t in texts:
        run = p.add_run(t)
        assert run.text == t
    assert p.text == para_text or "".join(texts)
    return p


# ═══════════════════════════════════════════════════════════════════════
# plan_run_splits
# ═══════════════════════════════════════════════════════════════════════


class TestPlanRunSplits:
    def test_no_errors_returns_empty(self, simple_para):
        assert plan_run_splits(simple_para, []) == []

    def test_error_start_at_run_boundary_still_splits_its_end(self, simple_para):
        """A start boundary does not make the half-open end boundary optional."""
        errors = [TextError(0, 1, "A", "B", "typo", "typo")]
        assert plan_run_splits(simple_para, errors) == [1]

    def test_error_inside_run_returns_position(self, simple_para):
        """Both half-open boundaries inside one run require splits."""
        errors = [TextError(3, 4, "D", "E", "typo", "typo")]
        assert plan_run_splits(simple_para, errors) == [3, 4]

    def test_multiple_errors_inside_run(self, simple_para):
        """Two errors inside a run contribute both start and end boundaries."""
        errors = [
            TextError(2, 3, "C", "D", "typo", "typo"),
            TextError(5, 6, "F", "G", "typo", "typo"),
        ]
        assert plan_run_splits(simple_para, errors) == [2, 3, 5, 6]

    def test_errors_in_multi_run_para(self, doc):
        """Errors in a paragraph with multiple runs."""
        p = _make_multi_run_para(doc, ["0123", "4567", "89AB"])
        # Runs: [0-3], [4-7], [8-11]
        errors = [
            TextError(2, 3, "2", "x", "typo", "typo"),  # inside run 0
            TextError(6, 7, "6", "x", "typo", "typo"),  # inside run 1
            TextError(10, 11, "A", "x", "typo", "typo"),  # inside run 2
        ]
        assert plan_run_splits(p, errors) == [2, 3, 6, 7, 10, 11]

    def test_error_at_last_run_boundary(self, doc):
        """An error at a run start still splits its end inside that run."""
        p = _make_multi_run_para(doc, ["ABC", "DEF"])
        # Runs: [0-2], [3-5]. Error at 3 (start of run 1) — no split.
        errors = [TextError(3, 4, "D", "x", "typo", "typo")]
        assert plan_run_splits(p, errors) == [4]

    def test_position_beyond_runs(self, simple_para):
        """Position beyond paragraph length — no split (defensive)."""
        errors = [TextError(99, 100, "?", "!", "typo", "typo")]
        assert plan_run_splits(simple_para, errors) == []

    def test_empty_paragraph_no_runs(self, doc):
        """Empty paragraph has no runs — empty split list."""
        p = doc.add_paragraph("")
        assert len(p.runs) == 0
        errors = [TextError(0, 1, "?", "!", "typo", "typo")]
        assert plan_run_splits(p, errors) == []

    def test_deduplicates_positions(self, simple_para):
        """Multiple errors at the same position produce one split."""
        errors = [
            TextError(3, 4, "D", "E", "typo", "typo"),
            TextError(3, 5, "DE", "EF", "symbol", "symbol"),
        ]
        assert plan_run_splits(simple_para, errors) == [3, 4, 5]


# ═══════════════════════════════════════════════════════════════════════
# ensure_run_at_position
# ═══════════════════════════════════════════════════════════════════════


def _same_run(a, b):
    """Check if two Run objects wrap the same underlying XML element."""
    return a._r is b._r


class TestEnsureRunAtPosition:
    def test_position_at_start_returns_same_run(self, simple_para):
        """Position 0 (start of first run) — no split, returns same run."""
        result = ensure_run_at_position(simple_para, 0)
        assert _same_run(result, simple_para.runs[0])
        assert simple_para.text == "ABCDEFGHIJ"
        assert len(simple_para.runs) == 1

    def test_position_in_middle_splits_run(self, simple_para):
        """Split at position 3: [ABC][DEFGHIJ]."""
        result = ensure_run_at_position(simple_para, 3)
        assert len(simple_para.runs) == 2
        assert simple_para.runs[0].text == "ABC"
        assert simple_para.runs[1].text == "DEFGHIJ"
        assert _same_run(result, simple_para.runs[1])
        # Total text unchanged
        assert simple_para.text == "ABCDEFGHIJ"

    def test_position_at_end_no_split(self, simple_para):
        """Position equal to total length — no split, returns last run."""
        result = ensure_run_at_position(simple_para, 10)
        assert len(simple_para.runs) == 1
        assert _same_run(result, simple_para.runs[0])

    def test_position_beyond_ranges_no_split(self, simple_para):
        """Position beyond total length — returns last run."""
        result = ensure_run_at_position(simple_para, 99)
        assert _same_run(result, simple_para.runs[0])

    def test_split_preserves_formatting(self, doc):
        """After split, both runs should retain run properties (rPr)."""
        p = doc.add_paragraph()
        run = p.add_run("BoldItalic")
        # Add bold + italic via XML manipulation
        rPr = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        rPr.append(b)
        i = OxmlElement("w:i")
        rPr.append(i)
        run._r.insert(0, rPr)

        ensure_run_at_position(p, 4)
        assert len(p.runs) == 2
        # Both runs should have rPr
        assert p.runs[0]._r.find(qn("w:rPr")) is not None
        assert p.runs[1]._r.find(qn("w:rPr")) is not None
        assert p.runs[0].text == "Bold"
        assert p.runs[1].text == "Italic"
        assert p.text == "BoldItalic"

    def test_multiple_runs_split_middle_run(self, doc):
        """In a 3-run paragraph, split inside the middle run."""
        p = _make_multi_run_para(doc, ["AAA", "BBB", "CCC"])
        # Runs: [0-2]=AAA, [3-5]=BBB, [6-8]=CCC
        # Position 4 is inside BBB at offset 1 -> B|BB
        result = ensure_run_at_position(p, 4)
        assert len(p.runs) == 4
        assert p.runs[0].text == "AAA"
        assert p.runs[1].text == "B"  # before split (BBB[:1])
        assert p.runs[2].text == "BB"  # after split (BBB[1:])
        assert p.runs[3].text == "CCC"
        assert p.text == "AAABBBCCC"
        assert _same_run(result, p.runs[2])

    def test_idempotent_double_split(self, simple_para):
        """Splitting the same position twice should be a no-op for the second."""
        ensure_run_at_position(simple_para, 3)
        assert len(simple_para.runs) == 2
        result = ensure_run_at_position(simple_para, 3)
        # Should not create a third run
        assert len(simple_para.runs) == 2
        # Should return the run starting at position 3
        assert _same_run(result, simple_para.runs[1])

    def test_multiple_splits_in_single_run(self, simple_para):
        """Split a single run at two different positions."""
        ensure_run_at_position(simple_para, 3)  # ABC|DEFGHIJ
        assert len(simple_para.runs) == 2
        ensure_run_at_position(simple_para, 7)  # ABC|DEFG|HIJ
        assert len(simple_para.runs) == 3
        assert simple_para.runs[0].text == "ABC"
        assert simple_para.runs[1].text == "DEFG"
        assert simple_para.runs[2].text == "HIJ"
        assert simple_para.text == "ABCDEFGHIJ"

    def test_plan_then_ensure_workflow(self, simple_para):
        """Full workflow: plan splits then apply them."""
        errors = [
            TextError(2, 3, "C", "D", "typo", "typo"),
            TextError(5, 6, "F", "G", "typo", "typo"),
        ]
        positions = plan_run_splits(simple_para, errors)
        assert positions == [2, 3, 5, 6]
        for pos in positions:
            ensure_run_at_position(simple_para, pos)
        assert len(simple_para.runs) == 5
        assert simple_para.runs[0].text == "AB"
        assert simple_para.runs[1].text == "C"
        assert simple_para.runs[2].text == "DE"
        assert simple_para.runs[3].text == "F"
        assert simple_para.runs[4].text == "GHIJ"
        assert simple_para.text == "ABCDEFGHIJ"

    def test_empty_paragraph_adds_run(self, doc):
        """Empty paragraph with no runs should create one."""
        p = doc.add_paragraph("")
        # Forcefully remove the auto-created run
        for r in list(p.runs):
            r._r.getparent().remove(r._r)
        assert len(p.runs) == 0
        result = ensure_run_at_position(p, 0)
        assert len(p.runs) == 1
        assert _same_run(result, p.runs[0])

    def test_split_in_multi_run_then_find_positions(self, doc):
        """After splitting, _find_run_at_position finds correct runs."""
        p = _make_multi_run_para(doc, ["Hello ", "World!", "Foo"])

        # "Hello " is 6 chars (positions 0-5)
        # "World!" is 6 chars (positions 6-11)
        # Position 7 is inside World! at offset 1 -> W|orld!
        ensure_run_at_position(p, 7)

        # Now verify runs
        assert len(p.runs) == 4
        assert p.runs[0].text == "Hello "
        assert p.runs[1].text == "W"
        assert p.runs[2].text == "orld!"
        assert p.runs[3].text == "Foo"

        # _find_run_at_position should find the correct runs
        from docwen_plugin_proofread.docx_validator import _find_run_at_position

        assert _same_run(_find_run_at_position(p, 0), p.runs[0])
        assert _same_run(_find_run_at_position(p, 5), p.runs[0])  # last char of Hello
        assert _same_run(_find_run_at_position(p, 6), p.runs[1])  # start of split (W)
        assert _same_run(_find_run_at_position(p, 7), p.runs[2])  # second split piece
        assert _same_run(_find_run_at_position(p, 11), p.runs[2])  # end of orld!
        assert _same_run(_find_run_at_position(p, 12), p.runs[3])  # Foo

    def test_special_ooxml_run_fails_closed_without_mutating_content(self, doc, tmp_path):
        """Tabs and breaks must never be flattened or duplicated by splitting."""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_text("AB")
        run.add_tab()
        run.add_text("CD")
        run.add_break()
        run.add_text("EF")
        before_text = paragraph.text
        before_xml = run._r.xml

        result = ensure_run_at_position(paragraph, 3)

        assert _same_run(result, run)
        assert len(paragraph.runs) == 1
        assert paragraph.text == before_text == "AB\tCD\nEF"
        assert run._r.xml == before_xml

        output = tmp_path / "special-content.docx"
        doc.save(output)
        reloaded = Document(output)
        reloaded_run = reloaded.paragraphs[0].runs[0]
        assert reloaded.paragraphs[0].text == before_text
        assert [child.tag for child in reloaded_run._r.iterchildren()] == [
            qn("w:t"),
            qn("w:tab"),
            qn("w:t"),
            qn("w:br"),
            qn("w:t"),
        ]
        assert [child.text for child in reloaded_run._r.iterchildren() if child.tag == qn("w:t")] == ["AB", "CD", "EF"]


class TestRunsForRange:
    def test_single_run_middle_range_becomes_exact_span(self, simple_para):
        error = TextError(2, 8, "CDEFGH", "x", "typo", "typo")
        for position in plan_run_splits(simple_para, [error]):
            ensure_run_at_position(simple_para, position)

        selected = runs_for_range(simple_para, error.start_pos, error.end_pos)

        assert [run.text for run in selected] == ["CDEFGH"]

    def test_cross_run_range_returns_all_runs_between_exact_boundaries(self, doc):
        paragraph = _make_multi_run_para(doc, ["AA", "Se", "cret", "42", "ZZ"])

        selected = runs_for_range(paragraph, 2, 10)

        assert [run.text for run in selected] == ["Se", "cret", "42"]

    @pytest.mark.parametrize(("start", "end"), [(-1, 1), (1, 1), (2, 1), (0, 11)])
    def test_invalid_ranges_fail_closed(self, simple_para, start, end):
        assert runs_for_range(simple_para, start, end) == []

    def test_unsplit_interior_boundary_fails_closed(self, simple_para):
        assert runs_for_range(simple_para, 2, 8) == []
