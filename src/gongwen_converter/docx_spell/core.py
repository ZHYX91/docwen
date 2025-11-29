"""
DOCX文档错别字检查核心处理模块

本模块是公文转换器的校对检查引擎，负责对Word文档进行文本校对和批注添加。
主要功能包括：
- 支持多种校对规则（标点配对、符号校对、错别字校对、敏感词匹配）
- 在DOCX文档中添加详细的批注说明
- 支持单文件处理和批量处理模式
- 自动生成带时间戳的输出文件
- 提供灵活的错别字检查选项配置

采用模块化设计，将文本验证、错误分组、批注添加等功能分离，确保代码的可维护性。
"""

import os
import datetime
import logging
from docx import Document
from .spell_checker import TextValidator
from .utils import plan_run_splits, rebuild_paragraph_with_splits
from gongwen_converter.utils.path_utils import generate_output_path

# 配置日志
logger = logging.getLogger(__name__)

from typing import Callable, Optional
import threading

def process_docx(
    docx_path: str,
    output_path: str = None,
    output_dir: str = None,
    spell_check_options: int = None,
    progress_callback: Optional[Callable[[str], None]] = None,
    cancel_event: Optional[threading.Event] = None
) -> str:
    """
    处理DOCX文件，添加错别字批注
    
    参数:
        docx_path: 输入DOCX文件路径
        output_path: 输出文件路径（可选）
        output_dir: 输出目录（可选，与output_path二选一）
        spell_check_options: 错别字检查选项
            None: 使用配置文件默认设置
            0: 不进行错别字检查
            1-15: 按位组合的校对规则（1:标点配对, 2:符号校对, 4:错别字校对, 8:敏感词匹配）
        progress_callback: 进度回调函数 (可选)
        
    返回:
        str: 处理后的文件路径
        
    详细说明:
        1. 如果没有提供输出路径，自动生成带时间戳的输出路径
        2. 格式: 原文件名_YYYYMMDD_HHMMSS.docx
        3. 确保不会覆盖原始文件
        4. 支持自定义错别字检查规则配置
    """
    logger.info("=" * 80)
    logger.info(f"开始处理DOCX文档: {docx_path}")
    
    # 验证输入文件是否存在
    if not os.path.exists(docx_path):
        logger.error(f"输入文件不存在: {docx_path}")
        raise FileNotFoundError(f"输入文件不存在: {docx_path}")
    
    # 创建输出路径（如果未提供）
    if not output_path:
        # 使用统一的路径生成方法
        output_path = generate_output_path(
            docx_path,
            output_dir=output_dir,
            section="",
            add_timestamp=True,
            description="checked",
            file_type="docx"
        )
        logger.info(f"自动生成输出路径: {output_path}")
    
    try:
        if progress_callback:
            progress_callback("正在加载文档...")
        logger.info("加载DOCX文档...")
        # 使用docx库加载文档
        doc = Document(docx_path)
        logger.info(f"成功加载文档，包含 {len(doc.paragraphs)} 个段落")
        
        # 根据选项决定是否进行错别字检查
        if spell_check_options == 0:
            logger.info("用户选择不进行错别字检查，直接保存文档")
            doc.save(output_path)
            logger.info(f"文档保存成功: {output_path} (大小: {os.path.getsize(output_path)} 字节)")
            return output_path

        # 初始化文本验证器
        if progress_callback:
            progress_callback("正在初始化校对引擎...")
        logger.info("初始化文本验证器...")
        validator = create_validator_with_options(spell_check_options)
        logger.info("文本验证器初始化完成")
        
        # 处理文档中的所有段落
        total_errors = 0
        total_paragraphs = len(doc.paragraphs)
        
        logger.info(f"开始处理 {total_paragraphs} 个段落...")
        for i, para in enumerate(doc.paragraphs):
            if cancel_event and cancel_event.is_set():
                logger.info("操作被用户取消")
                return None
            # 跳过空段落
            if not para.text.strip():
                logger.debug(f"跳过空段落: {i+1}/{total_paragraphs}")
                continue
                
            # 处理当前段落
            errors_found = process_paragraph(i, para, doc, validator)
            total_errors += errors_found
            
            # 每处理10个段落记录一次进度
            if (i + 1) % 10 == 0 or (i + 1) == total_paragraphs:
                logger.info(f"进度: {i+1}/{total_paragraphs} 段落 | 累计错误: {total_errors}")
                if progress_callback:
                    progress_callback(f"正在校对... {i+1}/{total_paragraphs}")
        
        # 保存处理后的文档
        if progress_callback:
            progress_callback("正在保存校对结果...")
        logger.info(f"文档处理完成，共发现 {total_errors} 个错误，准备保存...")
        doc.save(output_path)
        logger.info(f"文档保存成功: {output_path} (大小: {os.path.getsize(output_path)} 字节)")
        
        return output_path
        
    except Exception as e:
        logger.error(f"处理文档失败: {str(e)}", exc_info=True)
        raise RuntimeError(f"处理文档失败: {str(e)}") from e
    
def create_validator_with_options(spell_check_options: int = None) -> TextValidator:
    """
    根据校对选项创建文本验证器
    
    参数:
        spell_check_options: 校对选项
            None: 使用配置文件默认设置
            1-15: 按位组合的校对规则（1:标点配对, 2:符号校对, 4:错别字校对, 8:敏感词匹配）
            
    返回:
        TextValidator: 配置好的文本验证器
    """
    # 如果没有提供选项，使用默认配置创建验证器
    if spell_check_options is None:
        logger.info("使用配置文件默认设置初始化验证器")
        return TextValidator()
    
    # 解析选项值
    symbol_pairing_enabled = bool(spell_check_options & 1)  # 第一位：标点配对
    symbol_correction_enabled = bool(spell_check_options & 2)  # 第二位：符号校对
    typos_rule_enabled = bool(spell_check_options & 4)  # 第三位：错别字校对
    sensitive_word_enabled = bool(spell_check_options & 8)  # 第四位：敏感词匹配
    
    logger.info("使用校对规则初始化验证器:")
    logger.info(f"  标点配对: {symbol_pairing_enabled}")
    logger.info(f"  符号校对: {symbol_correction_enabled}")
    logger.info(f"  错别字校对: {typos_rule_enabled}")
    logger.info(f"  敏感词匹配: {sensitive_word_enabled}")
    
    # 创建验证器实例并覆盖配置
    return TextValidator(
        symbol_pairing=symbol_pairing_enabled,
        symbol_correction=symbol_correction_enabled,
        typos_rule=typos_rule_enabled,
        sensitive_word=sensitive_word_enabled
    )

def batch_process_docx(directory_path: str, output_dir: str = None) -> list:
    """
    批量处理目录中的所有DOCX文件（添加时间戳后缀）
    
    参数:
        directory_path: 包含DOCX文件的目录
        output_dir: 输出目录（可选）
        
    返回:
        list: 处理成功的文件路径列表
    """
    logger.info(f"开始批量处理目录: {directory_path}")
    
    # 验证输入目录
    if not os.path.isdir(directory_path):
        logger.error(f"输入目录不存在或不是目录: {directory_path}")
        raise NotADirectoryError(f"输入目录不存在或不是目录: {directory_path}")
    
    # 设置输出目录
    if not output_dir:
        # 使用统一的时间戳格式
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = os.path.join(directory_path, f"checked_docs_{timestamp}")
        logger.info(f"自动设置输出目录: {output_dir}")
    
    # 创建输出目录
    os.makedirs(output_dir, exist_ok=True)
    logger.info(f"确保输出目录存在: {output_dir}")
    
    # 查找所有DOCX文件
    docx_files = []
    for root, _, files in os.walk(directory_path):
        for file in files:
            if file.lower().endswith('.docx'):
                docx_files.append(os.path.join(root, file))
    
    logger.info(f"找到 {len(docx_files)} 个DOCX文件")
    
    # 处理每个文件
    processed_files = []
    for i, file_path in enumerate(docx_files):
        try:
            logger.info(f"处理文件 {i+1}/{len(docx_files)}: {file_path}")
            
            # 生成输出路径（使用统一方法）
            output_path = generate_output_path(
                file_path,
                output_dir=output_dir,
                section="",
                add_timestamp=True,
                description="checked",
                file_type="docx"
            )
            
            # 处理文档
            result = process_docx(file_path, output_path)
            processed_files.append(result)
            logger.info(f"文件处理完成: {result}")
            
        except Exception as e:
            logger.error(f"处理文件失败: {file_path}, 错误: {str(e)}")
    
    logger.info(f"批量处理完成，成功处理 {len(processed_files)}/{len(docx_files)} 个文件")
    return processed_files

def process_paragraph(para_index: int, paragraph, doc, validator) -> int:
    """
    处理单个段落，添加错别字批注
    
    参数:
        para_index: 段落索引（从0开始）
        paragraph: 段落对象
        doc: Document对象
        validator: TextValidator实例
        
    返回:
        int: 本段落发现的错误数量
        
    详细说明:
        1. 获取段落文本
        2. 使用验证器检测文本中的错误
        3. 一次性重建段落（标注所有错误位置）
        4. 为每个错误添加批注
    """
    text = paragraph.text
    if not text.strip():
        logger.debug(f"跳过空段落: {para_index+1}")
        return 0
    
    try:
        # 记录段落信息
        para_info = f"段落 {para_index+1}: '{text[:30]}...'" if len(text) > 30 else f"段落 {para_index+1}: '{text}'"
        logger.debug(f"处理 {para_info}")
        
        # 执行文本校对
        logger.debug("开始文本校对...")
        errors = validator.validate_text(text)
        
        if not errors:
            logger.debug("未发现错误")
            return 0
        
        logger.info(f"{para_info} 发现 {len(errors)} 个错误")
        
        # 一次性处理所有错误
        error_count = add_comments_for_errors(paragraph, doc, errors)
        
        return error_count
        
    except Exception as e:
        logger.error(f"处理段落失败: {str(e)}", exc_info=True)
        return 0

def add_comments_for_errors(paragraph, doc, errors) -> int:
    """
    为段落中的所有错误添加批注（一次性重建）
    
    参数:
        paragraph: 段落对象
        doc: Document对象
        errors: 错误列表
        
    返回:
        int: 成功添加的批注数量
        
    详细说明:
        1. 规划所有错误的run拆分方案
        2. 一次性重建段落（标注所有错误位置）
        3. 为每个错误run添加批注
        
    改进：
        - 只重建一次段落，避免批注丢失
        - 所有错误都被精确标注
        - 逻辑清晰，易于维护
    """
    try:
        logger.info(f"开始为 {len(errors)} 个错误添加批注")
        
        # 1. 规划所有错误的run拆分方案
        split_plan = plan_run_splits(paragraph, errors)
        
        # 2. 一次性重建段落，获取所有错误run
        new_paragraph, error_runs = rebuild_paragraph_with_splits(paragraph, split_plan, doc)
        
        if not error_runs:
            logger.warning("重建段落后未找到错误run，无法添加批注")
            return 0
        
        # 3. 为每个错误run添加批注
        success_count = 0
        for idx, error in enumerate(errors):
            if idx not in error_runs:
                logger.warning(f"未找到错误 {idx} 对应的run")
                continue
            
            error_run = error_runs[idx]
            comment_text = f"{error.error_type}：{error.error_text} → {error.suggestion}"
            
            try:
                doc.add_comment(
                    error_run,
                    text=comment_text,
                    author=f"公文校对-{error.source}",
                    initials="Sys"
                )
                logger.info(f"成功添加批注 [{idx+1}/{len(errors)}] (来源: {error.source})")
                success_count += 1
            except Exception as e:
                logger.error(f"添加批注失败 [{idx+1}/{len(errors)}]: {str(e)}")
        
        logger.info(f"批注添加完成: {success_count}/{len(errors)}")
        return success_count
        
    except Exception as e:
        logger.error(f"批量添加批注失败: {str(e)}", exc_info=True)
        return 0


# 模块测试代码
if __name__ == "__main__":
    # 配置日志
    logging.basicConfig(
        level=logging.DEBUG,
        format="%(asctime)s [%(levelname)s] %(name)s - %(message)s",
        handlers=[
            logging.StreamHandler(),
            logging.FileHandler("docx_spell_core_test.log")
        ]
    )
    
    logger.info("docx_spell.core模块测试开始")
    
    # 测试单文件处理
    test_file = "test.docx"
    if os.path.exists(test_file):
        logger.info(f"测试单文件处理: {test_file}")
        try:
            # 测试错别字 (1+2+4=7)
            result = process_docx(test_file, spell_check_options=7)
            logger.info(f"测试成功，输出文件: {result}")
        except Exception as e:
            logger.error(f"测试失败: {str(e)}")
    else:
        logger.warning(f"测试文件不存在: {test_file}")
    
    # 测试批量处理
    test_dir = "test_docs"
    if os.path.isdir(test_dir):
        logger.info(f"测试批量处理: {test_dir}")
        try:
            results = batch_process_docx(test_dir)
            logger.info(f"批量处理完成，成功处理 {len(results)} 个文件")
        except Exception as e:
            logger.error(f"批量处理失败: {str(e)}")
    else:
        logger.warning(f"测试目录不存在: {test_dir}")
    
    logger.info("docx_spell.core模块测试结束")
