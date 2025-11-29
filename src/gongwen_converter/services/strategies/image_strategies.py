"""
图片文件转换策略
支持图片转PDF、Markdown、DOCX等格式
"""

import os
import logging
import tempfile
import shutil
from typing import Dict, Any, Callable, Optional, TYPE_CHECKING
from .base_strategy import BaseStrategy
from gongwen_converter.services.result import ConversionResult
from . import register_conversion, register_action, CATEGORY_IMAGE
from gongwen_converter.utils.path_utils import generate_output_path

if TYPE_CHECKING:
    from docx import Document

logger = logging.getLogger(__name__)


def get_image_format_description(actual_format: str) -> str:
    """
    根据图片格式生成描述（使用真实格式）
    
    参数:
        actual_format: 图片的真实格式（如 'png', 'jpg', 'bmp'）
        
    返回:
        str: 格式描述，如 'fromPng', 'fromJpg' 等
    """
    format_map = {
        'png': 'fromPng',
        'jpg': 'fromJpg',
        'jpeg': 'fromJpeg',
        'tif': 'fromTif',
        'tiff': 'fromTiff',
        'gif': 'fromGif',
        'webp': 'fromWebp',
        'bmp': 'fromBmp',
        'heic': 'fromHeic',
        'heif': 'fromHeif',
    }
    
    return format_map.get(actual_format.lower() if actual_format else '', 'fromImage')


def is_multipage_tiff(file_path: str) -> bool:
    """
    检测TIFF是否为多页
    
    参数:
        file_path: 图片文件路径
        
    返回:
        bool: 是否为多页TIFF
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    if ext not in ['.tif', '.tiff']:
        return False
    
    try:
        from PIL import Image
        with Image.open(file_path) as img:
            # 尝试获取帧数
            n_frames = getattr(img, 'n_frames', 1)
            return n_frames > 1
    except Exception as e:
        logger.warning(f"检测TIFF页数失败: {e}")
        return False


def extract_tiff_pages(
    file_path: str,
    output_dir: str = None,
    actual_format: str = 'tiff',
    progress_callback: Optional[Callable[[str], None]] = None
) -> list:
    """
    提取TIFF每一页为独立的PNG文件，使用标准化命名规则
    
    文件命名格式：{原始基础名}_page{N}_{时间戳}_from{格式}.png
    示例：报告_page1_20250109_123000_fromTiff.png
    
    参数:
        file_path: TIFF文件路径
        output_dir: 输出目录（必需）。拆分的PNG文件将保存到此目录
        actual_format: 实际文件格式（默认'tiff'），用于生成描述信息
        progress_callback: 进度回调函数（可选）
        
    返回:
        list: 文件路径列表 [(页码, 文件路径), ...]
        
    注意:
        - 所有页面共享同一个时间戳，确保同一批次拆分的文件时间戳一致
        - 为避免时间戳冲突，在循环外生成统一时间戳，然后手动构建文件名
    """
    from PIL import Image
    import datetime
    
    if not output_dir:
        raise ValueError("extract_tiff_pages 需要指定 output_dir 参数")
    
    temp_files = []
    
    try:
        with Image.open(file_path) as img:
            n_frames = getattr(img, 'n_frames', 1)
            logger.info(f"TIFF文件共 {n_frames} 页，开始拆分")
            
            # 步骤1：生成统一时间戳（在循环外）
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            logger.debug(f"生成统一时间戳: {timestamp}")
            
            # 步骤2：提取原始文件名（清理旧时间戳）
            import re
            base_name = os.path.splitext(os.path.basename(file_path))[0]
            
            # 移除旧时间戳和描述
            timestamp_pattern = r'(_\d{8}_\d{6})(?:.*)?$'
            match = re.search(timestamp_pattern, base_name)
            if match:
                base_name_clean = base_name[:match.start()]
            else:
                base_name_clean = base_name
            
            logger.debug(f"清理后的基础文件名: {base_name_clean}")
            
            # 步骤3：生成描述
            description = f'from{actual_format.capitalize()}'
            
            # 步骤4：遍历每一页
            for i in range(n_frames):
                img.seek(i)
                page_num = i + 1
                
                # 更新进度
                if progress_callback:
                    progress_callback(f"提取第{page_num}页")
                
                # 手动构建文件名（复用统一时间戳）
                # 格式：{原始名}_page{N}_{时间戳}_{描述}.png
                filename = f"{base_name_clean}_page{page_num}_{timestamp}_{description}.png"
                page_path = os.path.join(output_dir, filename)
                
                logger.debug(f"第 {page_num} 页输出路径: {os.path.basename(page_path)}")
                
                # 转换并保存当前帧
                # RGBA模式转为RGB（白色背景）
                frame = img.copy()
                if frame.mode == 'RGBA':
                    background = Image.new('RGB', frame.size, (255, 255, 255))
                    background.paste(frame, mask=frame.split()[3])
                    frame = background
                    logger.debug(f"第 {page_num} 页: RGBA转RGB（白色背景）")
                elif frame.mode != 'RGB':
                    frame = frame.convert('RGB')
                    logger.debug(f"第 {page_num} 页: {frame.mode}转RGB")
                
                # 保存为PNG格式
                frame.save(page_path)
                temp_files.append((page_num, page_path))
                logger.info(f"✓ 第 {page_num} 页提取成功: {os.path.basename(page_path)}")
        
        logger.info(f"TIFF拆分完成，共提取 {len(temp_files)} 页")
        return temp_files
    
    except Exception as e:
        # 清理已创建的文件
        logger.error(f"TIFF拆分失败，清理已创建的文件")
        for _, temp_path in temp_files:
            try:
                if os.path.exists(temp_path):
                    os.remove(temp_path)
                    logger.debug(f"已删除: {temp_path}")
            except Exception as cleanup_error:
                logger.warning(f"清理文件失败: {temp_path}, 错误: {cleanup_error}")
        
        logger.error(f"提取TIFF页面失败: {e}", exc_info=True)
        raise


@register_conversion(CATEGORY_IMAGE, 'pdf')
class ImageToPdfStrategy(BaseStrategy):
    """
    将图片文件转换为PDF格式的策略
    
    支持三种质量模式：
    1. original - 原图嵌入（推荐）：保持原始分辨率，无损嵌入
    2. a4 - 适合A4纸：自动判断横向/纵向，适配210×297mm
    3. a3 - 适合A3纸：自动判断横向/纵向，适配297×420mm
    
    支持的输入格式：
    - 直接支持：JPG, PNG, TIFF, GIF, WEBP
    - 预处理后支持：BMP, HEIC (已通过批量列表自动转换为PNG)
    
    特殊处理：
    - 多页TIFF：自动检测并处理每一帧
    """
    
    def execute(
        self,
        file_path: str,
        options: Optional[Dict[str, Any]] = None,
        progress_callback: Optional[Callable[[str], None]] = None
    ) -> ConversionResult:
        """
        执行图片到PDF的转换
        
        Args:
            file_path: 输入图片文件路径
            options: 转换选项，包含quality_mode参数
            progress_callback: 进度更新回调函数
            
        Returns:
            ConversionResult: 转换结果对象
        """
        try:
            if progress_callback:
                progress_callback("正在转换为PDF...")
            
            import img2pdf
            from PIL import Image
            
            # 获取质量模式，默认为'original'
            quality_mode = options.get('quality_mode', 'original') if options else 'original'
            actual_format = options.get('actual_format') if options else None  # 从options中提取actual_format
            logger.info(f"图片转PDF质量模式: {quality_mode}")
            
            from gongwen_converter.utils.workspace_manager import get_output_directory
            output_dir = get_output_directory(file_path)
            
            # 使用标准的临时目录
            with tempfile.TemporaryDirectory() as temp_dir:
                # 步骤1：创建输入副本 input.{ext}
                from gongwen_converter.utils.workspace_manager import prepare_input_file
                temp_input = prepare_input_file(file_path, temp_dir, actual_format or 'jpg')
                logger.debug(f"已创建输入副本: {os.path.basename(temp_input)}")
                
                # 步骤2：检测并预处理特殊格式（从副本处理）
                # 返回处理后的文件路径和是否是转换后的中间文件
                file_to_convert = self._preprocess_image(temp_input, temp_dir, None, actual_format)
                
                # 记录中间文件（如果有格式转换）
                intermediate_file = None
                if actual_format in ['bmp', 'heic', 'heif'] and file_to_convert != temp_input:
                    intermediate_file = file_to_convert
                    logger.debug(f"记录中间文件: {os.path.basename(intermediate_file)}")
                
                # 读取图片信息，判断方向
                with Image.open(file_to_convert) as img:
                    width, height = img.size
                    is_landscape = width > height
                    logger.debug(f"图片尺寸: {width}x{height}, 横向: {is_landscape}")
                
                # 根据质量模式设置PDF布局
                layout_fun = self._create_layout(quality_mode, is_landscape)
                
                # 转换为PDF
                logger.debug(f"开始转换图片为PDF: {file_to_convert}")
                if layout_fun is None:
                    # 原图嵌入模式：不传递layout_fun参数
                    pdf_bytes = img2pdf.convert(file_to_convert)
                else:
                    # A4/A3模式：传递layout_fun参数
                    pdf_bytes = img2pdf.convert(file_to_convert, layout_fun=layout_fun)
                
                # 生成输出文件名
                description = get_image_format_description(actual_format) if actual_format else "fromImage"
                output_filename = os.path.basename(
                    generate_output_path(
                        file_path,
                        section="",
                        add_timestamp=True,
                        description=description,
                        file_type="pdf"
                    )
                )
                
                # 保存PDF到临时目录
                temp_output = os.path.join(temp_dir, output_filename)
                with open(temp_output, 'wb') as f:
                    f.write(pdf_bytes)
                
                # 准备最终输出路径
                output_path = os.path.join(output_dir, output_filename)
                
                # 处理中间文件保留
                should_keep = self._should_keep_intermediates()
                if should_keep:
                    # 保留中间文件（排除输入副本）
                    logger.info("保留中间文件，移动规范命名的文件到输出目录")
                    for filename in os.listdir(temp_dir):
                        # 排除输入副本文件
                        if filename.startswith('input.'):
                            logger.debug(f"跳过输入副本: {filename}")
                            continue
                        src = os.path.join(temp_dir, filename)
                        if os.path.isfile(src):
                            dst = os.path.join(output_dir, filename)
                            shutil.move(src, dst)
                            logger.debug(f"保留中间文件: {filename}")
                else:
                    # 只移动最终PDF文件
                    logger.debug("清理中间文件，只移动最终文件")
                    shutil.move(temp_output, output_path)
                
                logger.info(f"成功转换为PDF: {output_path}")
            
            if progress_callback:
                progress_callback("转换完成。")
            
            return ConversionResult(
                success=True,
                output_path=output_path,
                message="转换为PDF成功。"
            )
        
        except Exception as e:
            logger.error(f"图片转PDF失败: {e}", exc_info=True)
            return ConversionResult(
                success=False,
                message=f"转换失败: {e}",
                error=e
            )
    
    def _preprocess_image(self, file_path: str, temp_dir: str, cancel_event=None, actual_format: str = None) -> str:
        """
        预处理图片：将非标准格式（BMP/HEIC/HEIF）转换为PNG
        
        Args:
            file_path: 原始文件路径
            temp_dir: 临时目录路径（必需）
            cancel_event: 取消事件（可选）
            actual_format: 实际文件格式（可选，如果不提供则自动检测）
            
        Returns:
            str: 处理后的文件路径
            
        说明:
            - 使用actual_format参数避免重复检测文件格式
            - 如果输入已经是标准格式（JPG/PNG/TIFF等），直接返回原路径
            - 如果是BMP/HEIC/HEIF，转换为PNG并保存到临时目录
            - 转换后的文件由 TempFileManager 自动管理清理
        """
        # 如果没有提供actual_format，则检测
        if actual_format is None:
            from gongwen_converter.utils.file_type_utils import detect_actual_file_format
            actual_format = detect_actual_file_format(file_path)
            logger.debug(f"自动检测图片文件格式: {actual_format}")
        else:
            logger.debug(f"使用传入的文件格式: {actual_format}")
        
        # 标准格式：直接使用，无需转换
        if actual_format in ['jpeg', 'jpg', 'png', 'tiff', 'tif', 'gif', 'webp']:
            logger.debug(f"文件已是标准图片格式({actual_format})，无需转换: {file_path}")
            return file_path
        
        # BMP格式需要转换为PNG
        if actual_format == 'bmp':
            logger.info(f"检测到BMP格式，转换为PNG: {os.path.basename(file_path)}")
            
            try:
                from gongwen_converter.converter.formats.image import bmp_to_png
                
                # 调用统一的BMP转PNG函数，传递temp_dir参数
                converted_path = bmp_to_png(file_path, output_dir=temp_dir)
                logger.info(f"BMP转PNG成功: {os.path.basename(converted_path)}")
                return converted_path
                
            except Exception as e:
                logger.error(f"BMP转PNG失败: {e}")
                raise RuntimeError(f"BMP转PNG失败: {e}")
        
        # HEIC/HEIF格式需要转换为PNG
        elif actual_format in ['heic', 'heif']:
            logger.info(f"检测到{actual_format.upper()}格式，转换为PNG: {os.path.basename(file_path)}")
            
            try:
                from gongwen_converter.converter.formats.image import heic_to_png
                
                # 调用统一的HEIC转PNG函数，传递temp_dir参数
                converted_path = heic_to_png(file_path, output_dir=temp_dir)
                logger.info(f"{actual_format.upper()}转PNG成功: {os.path.basename(converted_path)}")
                return converted_path
                
            except Exception as e:
                logger.error(f"{actual_format.upper()}转PNG失败: {e}")
                raise RuntimeError(f"{actual_format.upper()}转PNG失败: {e}")
        
        # 其他不支持的格式：直接返回原路径
        logger.warning(f"不支持的图片格式: {actual_format}，尝试直接使用")
        return file_path
    
    def _create_layout(self, quality_mode: str, is_landscape: bool):
        """
        根据质量模式创建PDF布局函数
        
        Args:
            quality_mode: 质量模式 ('original', 'a4', 'a3')
            is_landscape: 是否为横向图片
            
        Returns:
            布局函数或None
        """
        import img2pdf
        
        if quality_mode == 'original':
            # 原图嵌入模式：不设置页面尺寸
            logger.debug("使用原图嵌入模式")
            return None
        
        elif quality_mode == 'a4':
            # A4纸张尺寸：210×297mm
            if is_landscape:
                # 横向：宽297mm，高210mm
                pagesize = (img2pdf.mm_to_pt(297), img2pdf.mm_to_pt(210))
                logger.debug("使用A4横向布局")
            else:
                # 纵向：宽210mm，高297mm
                pagesize = (img2pdf.mm_to_pt(210), img2pdf.mm_to_pt(297))
                logger.debug("使用A4纵向布局")
            
            return img2pdf.get_layout_fun(pagesize)
        
        elif quality_mode == 'a3':
            # A3纸张尺寸：297×420mm
            if is_landscape:
                # 横向：宽420mm，高297mm
                pagesize = (img2pdf.mm_to_pt(420), img2pdf.mm_to_pt(297))
                logger.debug("使用A3横向布局")
            else:
                # 纵向：宽297mm，高420mm
                pagesize = (img2pdf.mm_to_pt(297), img2pdf.mm_to_pt(420))
                logger.debug("使用A3纵向布局")
            
            return img2pdf.get_layout_fun(pagesize)
        
        else:
            logger.warning(f"未知的质量模式: {quality_mode}，使用原图嵌入")
            return None
    
    @staticmethod
    def _should_keep_intermediates() -> bool:
        """判断是否应该保留中间文件"""
        try:
            from gongwen_converter.config.config_manager import config_manager
            return config_manager.get_save_intermediate_files()
        except Exception as e:
            logger.warning(f"读取中间文件配置失败: {e}，使用默认设置（不保存中间文件）")
            return False


@register_conversion(CATEGORY_IMAGE, 'md')
class ImageToMarkdownStrategy(BaseStrategy):
    """
    将图片文件转换为Markdown格式的策略（含OCR）- 新版本
    
    输出结构：
    文件夹名_时间戳_fromFormat/
    ├── 图片文件（原始或转换后）
    └── 文件夹名_时间戳_fromFormat.md
    
    MD格式：
    ![图片文件名](./图片文件名)
    
    OCR识别内容
    
    多页TIFF会拆分为多个PNG，每个PNG对应一个图片链接和OCR代码块
    """
    
    def execute(
        self,
        file_path: str,
        options: Optional[Dict[str, Any]] = None,
        progress_callback: Optional[Callable[[str], None]] = None
    ) -> ConversionResult:
        """
        执行图片到Markdown的转换（新统一流程）
        
        Args:
            file_path: 输入图片文件路径
            options: 转换选项，包含actual_format
            progress_callback: 进度更新回调函数
            
        Returns:
            ConversionResult: 转换结果对象
        """
        try:
            actual_format = options.get('actual_format') if options else None
            
            from gongwen_converter.utils.workspace_manager import get_output_directory
            output_dir = get_output_directory(file_path)
            
            original_basename = os.path.splitext(os.path.basename(file_path))[0]
            
            with tempfile.TemporaryDirectory() as temp_dir:
                # === 步骤1：创建输入副本 input.{ext} ===
                if progress_callback:
                    progress_callback("准备转换...")
                
                from gongwen_converter.utils.workspace_manager import prepare_input_file
                temp_image = prepare_input_file(file_path, temp_dir, actual_format or 'jpg')
                logger.debug(f"已创建输入副本: {os.path.basename(temp_image)}")
                
                # === 步骤2：生成统一的basename ===
                description = get_image_format_description(actual_format) if actual_format else 'fromImage'
                output_path_with_ext = generate_output_path(
                    file_path, None, "", True, description, 'md'
                )
                basename = os.path.splitext(os.path.basename(output_path_with_ext))[0]
                logger.info(f"统一basename: {basename}")
                
                # === 步骤3：创建子文件夹 ===
                folder_path = os.path.join(temp_dir, basename)
                os.makedirs(folder_path)
                logger.debug(f"子文件夹: {folder_path}")
                
                # === 步骤4：处理图片到子文件夹 ===
                images_in_folder = []  # [(图片路径, 图片文件名), ...]
                
                if is_multipage_tiff(file_path):
                    # 情况1：多页TIFF - 使用标准化命名直接提取到子文件夹
                    logger.info("检测到多页TIFF，使用标准化命名拆分")
                    pages = extract_tiff_pages(
                        file_path=temp_image,
                        output_dir=folder_path,
                        actual_format=actual_format or 'tiff',
                        progress_callback=progress_callback
                    )
                    images_in_folder = [(path, os.path.basename(path)) for _, path in pages]
                    logger.debug(f"多页TIFF拆分完成，共 {len(pages)} 页")
                
                elif actual_format in ['heic', 'heif', 'bmp']:
                    # 情况2：需要转换格式 - 直接转换到子文件夹
                    if progress_callback:
                        progress_callback(f"转换{actual_format.upper()}为PNG...")
                    
                    logger.info(f"{actual_format.upper()}转换为PNG")
                    image_filename = f'{original_basename}.png'
                    final_image_path = os.path.join(folder_path, image_filename)
                    
                    from gongwen_converter.converter.formats.image import heic_to_png, bmp_to_png
                    if actual_format in ['heic', 'heif']:
                        heic_to_png(temp_image, output_path=final_image_path)
                    else:  # bmp
                        bmp_to_png(temp_image, output_path=final_image_path)
                    
                    images_in_folder = [(final_image_path, image_filename)]
                
                else:
                    # 情况3：标准格式 - 移动到子文件夹
                    if progress_callback:
                        progress_callback("处理图片...")
                    
                    logger.info(f"标准{actual_format.upper()}格式")
                    image_filename = f'{original_basename}.{actual_format}'
                    final_image_path = os.path.join(folder_path, image_filename)
                    shutil.move(temp_image, final_image_path)
                    
                    images_in_folder = [(final_image_path, image_filename)]
                
                # === 步骤5：统一OCR子文件夹的所有图片并生成MD内容 ===
                md_content = ""
                
                try:
                    from gongwen_converter.utils.ocr_utils import extract_text_simple
                    ocr_available = True
                except ImportError:
                    logger.warning("PaddleOCR未安装，跳过文字识别")
                    ocr_available = False
                
                # 获取导出选项
                extract_image = options.get('extract_image', True) if options else True
                extract_ocr = options.get('extract_ocr', True) if options else True
                
                logger.info(f"导出选项 - 提取图片: {extract_image}, OCR: {extract_ocr}")
                
                # 获取链接格式配置
                from gongwen_converter.config.config_manager import config_manager
                from gongwen_converter.utils.markdown_utils import format_image_link
                
                link_settings = config_manager.get_markdown_link_style_settings()
                image_format = link_settings.get("image_link_format", "wiki")
                image_embed = link_settings.get("image_embed", True)
                
                for img_path, img_filename in images_in_folder:
                    # 根据选项添加图片链接
                    if extract_image:
                        image_link = format_image_link(img_filename, image_format, image_embed)
                        md_content += f"{image_link}\n\n"
                    
                # 根据选项进行OCR识别
                if extract_ocr:
                    if ocr_available:
                        if progress_callback:
                            if len(images_in_folder) > 1:
                                # 多页时显示页码
                                page_info = img_filename.replace(f'{original_basename}_page', '').replace('.png', '')
                                progress_callback(f"识别第{page_info}页")
                            else:
                                progress_callback("识别文字中...")
                        
                        try:
                            ocr_text = extract_text_simple(img_path)
                            if ocr_text:
                                md_content += f"{ocr_text}\n\n"
                                logger.debug(f"OCR识别成功: {img_filename}")
                        except Exception as e:
                            logger.warning(f"OCR识别失败: {e}")
                
                # === 步骤6：保存MD文件 ===
                if progress_callback:
                    progress_callback("生成Markdown文件...")
                
                md_filename = f'{basename}.md'
                md_path = os.path.join(folder_path, md_filename)
                with open(md_path, 'w', encoding='utf-8') as f:
                    f.write(md_content)
                logger.info(f"MD文件已生成: {md_path}")
                
                # === 步骤7：移动子文件夹到目标目录 ===
                final_folder = os.path.join(output_dir, basename)
                if os.path.exists(final_folder):
                    shutil.rmtree(final_folder)
                shutil.move(folder_path, final_folder)
                logger.info(f"子文件夹已移动到: {final_folder}")
                
                if progress_callback:
                    progress_callback("转换完成。")
                
                return ConversionResult(
                    success=True,
                    output_path=os.path.join(final_folder, md_filename),
                    message="转换为Markdown成功。"
                )
        
        except Exception as e:
            logger.error(f"图片转Markdown失败: {e}", exc_info=True)
            return ConversionResult(
                success=False,
                message=f"转换失败: {e}",
                error=e
            )


@register_conversion(CATEGORY_IMAGE, 'docx')
class ImageToDocxStrategy(BaseStrategy):
    """
    将图片文件嵌入到DOCX文档的策略（含OCR）
    
    功能：
    - 图片居中对齐
    - OCR识别文字
    - 保持原图字号大小
    - 多页TIFF：每一页分别插入，每页图片下方是该页OCR内容
    """
    
    def execute(
        self,
        file_path: str,
        options: Optional[Dict[str, Any]] = None,
        progress_callback: Optional[Callable[[str], None]] = None
    ) -> ConversionResult:
        """
        执行图片到DOCX的转换（含OCR+居中+字号）
        
        支持保留中间文件（当配置启用时）：
        - 多页TIFF拆分的PNG文件
        - HEIC/BMP转换的PNG文件
        
        Args:
            file_path: 输入图片文件路径
            options: 转换选项，包含actual_format
            progress_callback: 进度更新回调函数
            
        Returns:
            ConversionResult: 转换结果对象
        """
        import uuid
        
        try:
            if progress_callback:
                progress_callback("正在转换为DOCX...")
            
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from PIL import Image
            
            # 获取actual_format和输出目录
            actual_format = options.get('actual_format') if options else None
            
            from gongwen_converter.utils.workspace_manager import get_output_directory
            output_dir = get_output_directory(file_path)
            
            # 检查是否需要保留中间文件
            should_keep = self._should_keep_intermediates()
            logger.info(f"保留中间文件配置: {should_keep}")
            
            # 创建临时目录（不使用自动清理的TemporaryDirectory）
            temp_dir = os.path.join(
                tempfile.gettempdir(),
                f"gongwen_temp_{uuid.uuid4().hex[:8]}"
            )
            os.makedirs(temp_dir, exist_ok=True)
            logger.debug(f"创建临时目录: {temp_dir}")
            
            try:
                from gongwen_converter.utils.workspace_manager import prepare_input_file
                temp_input = prepare_input_file(file_path, temp_dir, actual_format or 'jpg')
                logger.debug(f"已创建输入副本: {os.path.basename(temp_input)}")
                
                # 创建新文档
                doc = Document()
                
                # 用于记录中间文件
                intermediate_files = []
                
                # 检测是否为多页TIFF
                if is_multipage_tiff(temp_input):
                    logger.info(f"检测到多页TIFF，使用标准化命名拆分")
                    
                    # 提取TIFF每一页到临时目录，使用标准化命名
                    temp_pages = extract_tiff_pages(
                        file_path=temp_input,
                        output_dir=temp_dir,
                        actual_format=actual_format or 'tiff',
                        progress_callback=progress_callback
                    )
                    
                    # 记录中间文件
                    intermediate_files.extend([path for _, path in temp_pages])
                    
                    # 处理每一页
                    for page_num, temp_page_path in temp_pages:
                        if progress_callback:
                            progress_callback(f"处理第{page_num}页")
                        
                        # 插入该页图片
                        self._add_image_to_doc(doc, temp_page_path)
                        
                        # OCR识别该页
                        self._add_ocr_to_doc(doc, temp_page_path, progress_callback)
                        
                        # 如果不是最后一页，添加分页符
                        if page_num < len(temp_pages):
                            doc.add_page_break()
                        
                        logger.info(f"✓ 第 {page_num} 页处理完成")
                
                elif actual_format in ['heic', 'heif', 'bmp']:
                    # 需要格式转换的图片
                    logger.info(f"检测到{actual_format.upper()}格式，转换为PNG")
                    
                    if progress_callback:
                        progress_callback(f"转换{actual_format.upper()}为PNG...")
                    
                    # 使用标准化命名生成PNG文件
                    description = get_image_format_description(actual_format)
                    converted_filename = os.path.basename(
                        generate_output_path(
                            input_path=file_path,
                            output_dir=temp_dir,
                            section="",
                            add_timestamp=True,
                            description=description,
                            file_type='png'
                        )
                    )
                    converted_path = os.path.join(temp_dir, converted_filename)
                    
                    # 执行转换
                    from gongwen_converter.converter.formats.image import heic_to_png, bmp_to_png
                    if actual_format in ['heic', 'heif']:
                        heic_to_png(temp_input, output_path=converted_path)
                    else:  # bmp
                        bmp_to_png(temp_input, output_path=converted_path)
                    
                    logger.info(f"✓ {actual_format.upper()}转PNG成功: {os.path.basename(converted_path)}")
                    
                    # 记录中间文件
                    intermediate_files.append(converted_path)
                    
                    # 处理转换后的图片
                    self._add_image_to_doc(doc, converted_path)
                    self._add_ocr_to_doc(doc, converted_path, progress_callback)
                
                else:
                    # 标准格式图片，直接处理
                    logger.info(f"标准{actual_format.upper() if actual_format else ''}格式，直接处理")
                    self._add_image_to_doc(doc, temp_input)
                    self._add_ocr_to_doc(doc, temp_input, progress_callback)
                
                # 生成输出路径
                description = get_image_format_description(actual_format) if actual_format else "fromImage"
                output_path = generate_output_path(
                    file_path,
                    section="",
                    add_timestamp=True,
                    description=description,
                    file_type="docx"
                )
                
                # 保存文档
                doc.save(output_path)
                logger.info(f"✓ 成功生成DOCX文件: {os.path.basename(output_path)}")
                
                # 处理中间文件
                if should_keep and intermediate_files:
                    logger.info(f"保留 {len(intermediate_files)} 个中间文件到输出目录")
                    for intermediate_file in intermediate_files:
                        if os.path.exists(intermediate_file):
                            filename = os.path.basename(intermediate_file)
                            dest_path = os.path.join(output_dir, filename)
                            shutil.move(intermediate_file, dest_path)
                            logger.info(f"✓ 保留中间文件: {filename}")
                else:
                    if intermediate_files:
                        logger.debug(f"清理 {len(intermediate_files)} 个中间文件")
                
            finally:
                # 清理临时目录
                try:
                    if os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir, ignore_errors=True)
                        logger.debug(f"临时目录已清理: {temp_dir}")
                except Exception as cleanup_error:
                    logger.warning(f"清理临时目录失败: {cleanup_error}")
            
            if progress_callback:
                progress_callback("转换完成。")
            
            return ConversionResult(
                success=True,
                output_path=output_path,
                message="转换为DOCX成功（含OCR）。"
            )
        
        except Exception as e:
            logger.error(f"图片转DOCX失败: {e}", exc_info=True)
            return ConversionResult(
                success=False,
                message=f"转换失败: {e}",
                error=e
            )
    
    @staticmethod
    def _should_keep_intermediates() -> bool:
        """
        判断是否应该保留中间文件
        
        Returns:
            bool: True表示保留，False表示不保留
        """
        try:
            from gongwen_converter.config.config_manager import config_manager
            intermediate_settings = config_manager.get_intermediate_files_settings()
            return intermediate_settings.get("save_to_output", False)
        except Exception as e:
            logger.warning(f"读取中间文件配置失败: {e}，使用默认设置（不保存）")
            return False
    
    def _add_image_to_doc(self, doc: 'Document', image_path: str) -> None:
        """
        将图片添加到文档中（居中对齐）
        
        Args:
            doc: Document对象
            image_path: 图片路径
        """
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from PIL import Image
        
        # 获取图片尺寸
        with Image.open(image_path) as img:
            width_px, height_px = img.size
            logger.debug(f"图片尺寸: {width_px}x{height_px} 像素")
        
        # 计算合适的显示宽度（最大6英寸，约A4纸宽度）
        max_width_inches = 6.0
        
        # 如果图片宽度超过最大宽度，按比例缩放
        if width_px > max_width_inches * 96:  # 假设96 DPI
            display_width = max_width_inches
        else:
            display_width = width_px / 96  # 转换为英寸
        
        # 添加图片到文档（居中对齐）
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(display_width))
        
        logger.debug(f"图片已添加并居中对齐: {image_path}")
    
    def _add_ocr_to_doc(self, doc: 'Document', image_path: str, progress_callback: Optional[Callable[[str], None]] = None) -> None:
        """
        对图片进行OCR识别并将结果添加到文档
        
        Args:
            doc: Document对象
            image_path: 图片路径
            progress_callback: 进度回调
        """
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        try:
            if progress_callback:
                progress_callback("识别文字中...")
            
            from gongwen_converter.utils.ocr_utils import extract_text_with_sizes
            
            ocr_blocks = extract_text_with_sizes(image_path)
            
            if ocr_blocks:
                logger.info(f"OCR识别成功，添加 {len(ocr_blocks)} 行文字")
                
                # 添加每行OCR文字（居中对齐，保持字号）
                for block in ocr_blocks:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(block['text'])
                    run.font.size = Pt(block['font_size'])
                    logger.debug(f"添加文字: {block['text']} (字号: {block['font_size']}pt)")
            else:
                logger.info("OCR未识别到文字，仅保留图片")
                
        except ImportError as e:
            logger.warning(f"PaddleOCR未安装，跳过文字识别: {e}")
            logger.info("如需OCR功能，请安装: pip install paddleocr paddlepaddle")
        except Exception as e:
            logger.warning(f"OCR识别失败，仅保留图片: {e}")


@register_conversion(CATEGORY_IMAGE, CATEGORY_IMAGE)
class ImageFormatConversionStrategy(BaseStrategy):
    """
    通用图片格式转换策略
    
    支持以下格式的相互转换：
    - JPEG/JPG ↔ PNG ↔ GIF ↔ BMP ↔ TIFF ↔ WebP
    
    使用 Pillow 库进行格式转换，自动处理：
    - 透明度转换（RGBA → RGB，白色背景）
    - 颜色模式转换（P, LA 等 → RGB）
    - 多页TIFF（仅保留第一页）
    """
    
    def execute(
        self,
        file_path: str,
        options: Optional[Dict[str, Any]] = None,
        progress_callback: Optional[Callable[[str], None]] = None
    ) -> ConversionResult:
        """
        执行图片格式转换
        
        标准化工作流程：
        1. 在临时目录创建输入文件副本（使用 prepare_input_file）
        2. 从副本读取并转换格式
        3. 保存到最终输出目录
        4. 自动清理临时目录
        
        多页TIFF处理：
        - 检测到多页TIFF时，拆分每一页并分别转换
        - 每页使用标准化命名：原始名_page{N}_{时间戳}_from{格式}.{目标格式}
        - 所有页面共享同一时间戳
        
        Args:
            file_path: 输入图片文件路径
            options: 转换选项，包含actual_format和target_format
            progress_callback: 进度更新回调函数
            
        Returns:
            ConversionResult: 转换结果对象
        """
        try:
            from PIL import Image
            
            # 获取实际格式
            actual_format = options.get('actual_format') if options else None
            if not actual_format:
                from gongwen_converter.utils.file_type_utils import detect_actual_file_format
                actual_format = detect_actual_file_format(file_path)
            
            # 从options中获取目标格式
            target_format = options.get('target_format') if options else None
            
            if not target_format:
                logger.warning("未能从选项中获取目标格式")
                return ConversionResult(
                    success=False,
                    message="转换失败：未指定目标格式"
                )
            
            logger.info(f"图片格式转换: {actual_format.upper()} → {target_format.upper()}")
            
            if progress_callback:
                progress_callback(f"转换{actual_format.upper()}为{target_format.upper()}...")
            
            from gongwen_converter.utils.workspace_manager import get_output_directory
            output_dir = get_output_directory(file_path)
            
            # 保存options到实例变量，供子方法使用
            self._current_options = options
            
            # 使用临时目录
            with tempfile.TemporaryDirectory() as temp_dir:
                # 创建输入副本
                from gongwen_converter.utils.workspace_manager import prepare_input_file
                temp_input = prepare_input_file(file_path, temp_dir, actual_format)
                logger.debug(f"已创建输入副本: {os.path.basename(temp_input)}")
                
                # 检测是否为多页TIFF
                if is_multipage_tiff(temp_input):
                    # 多页TIFF：拆分每一页并分别转换
                    return self._convert_multipage_tiff(
                        temp_input, file_path, actual_format, target_format,
                        output_dir, progress_callback
                    )
                else:
                    # 单页图片：直接转换
                    return self._convert_single_image(
                        temp_input, file_path, actual_format, target_format,
                        output_dir, progress_callback
                    )
        
        except Exception as e:
            logger.error(f"图片格式转换失败: {e}", exc_info=True)
            return ConversionResult(
                success=False,
                message=f"转换失败: {e}",
                error=e
            )
    
    def _convert_single_image(
        self,
        temp_input: str,
        original_path: str,
        actual_format: str,
        target_format: str,
        output_dir: str,
        progress_callback: Optional[Callable[[str], None]] = None
    ) -> ConversionResult:
        """
        转换单页图片（支持压缩选项）
        
        Args:
            temp_input: 临时输入文件路径
            original_path: 原始文件路径
            actual_format: 实际格式
            target_format: 目标格式
            output_dir: 输出目录
            progress_callback: 进度回调
            
        Returns:
            ConversionResult: 转换结果
        """
        # 使用新的convert_image函数，支持压缩选项
        from gongwen_converter.converter.formats.image import convert_image
        
        # 策略层负责生成输出路径（使用原始文件路径）
        description = get_image_format_description(actual_format)
        output_path = generate_output_path(
            original_path,  # 使用原始文件路径生成文件名
            output_dir=output_dir,
            section="",
            add_timestamp=True,
            description=description,
            file_type=target_format
        )
        
        # 准备压缩选项
        compress_options = {}
        
        # 从父方法的options中提取压缩选项（通过self临时存储）
        if hasattr(self, '_current_options') and self._current_options:
            compress_mode = self._current_options.get('compress_mode', 'lossless')
            compress_options['compress_mode'] = compress_mode
            
            if compress_mode == 'limit_size':
                compress_options['size_limit'] = self._current_options.get('size_limit')
                compress_options['size_unit'] = self._current_options.get('size_unit', 'KB')
                logger.info(f"使用压缩模式: {compress_options['size_limit']}{compress_options['size_unit']}")
            else:
                logger.info("使用最高质量模式")
        else:
            logger.info("未提供压缩选项，使用默认最高质量模式")
            compress_options['compress_mode'] = 'lossless'
        
        try:
            # 调用convert_image函数执行转换
            result_path = convert_image(
                source_path=temp_input,    # 临时副本用于读取
                target_format=target_format,
                output_path=output_path,   # 基于原始文件名的输出路径
                options=compress_options
            )
            
            logger.info(f"✓ 格式转换成功: {os.path.basename(result_path)}")
            
            if progress_callback:
                progress_callback("转换完成。")
            
            # 标准化扩展名用于消息
            file_ext = self._normalize_extension(target_format)
            
            return ConversionResult(
                success=True,
                output_path=result_path,
                message=f"已成功转换为{file_ext.upper()}"
            )
        
        except Exception as e:
            logger.error(f"图片转换失败: {e}", exc_info=True)
            return ConversionResult(
                success=False,
                message=f"转换失败: {e}",
                error=e
            )
    
    def _convert_multipage_tiff(
        self,
        temp_input: str,
        original_path: str,
        actual_format: str,
        target_format: str,
        output_dir: str,
        progress_callback: Optional[Callable[[str], None]] = None
    ) -> ConversionResult:
        """
        转换多页TIFF，拆分每一页并分别转换
        
        Args:
            temp_input: 临时输入文件路径
            original_path: 原始文件路径
            actual_format: 实际格式
            target_format: 目标格式
            output_dir: 输出目录
            progress_callback: 进度回调
            
        Returns:
            ConversionResult: 转换结果（返回第一页的路径）
        """
        from PIL import Image
        
        output_files = []
        
        # 标准化文件扩展名
        file_ext = self._normalize_extension(target_format)
        pil_format = self._get_pil_format(target_format)
        
        try:
            with Image.open(temp_input) as img:
                n_frames = getattr(img, 'n_frames', 1)
                logger.info(f"检测到多页TIFF（共{n_frames}页），开始拆分转换")
                
                for i in range(n_frames):
                    img.seek(i)
                    page_num = i + 1
                    
                    # 更新进度
                    if progress_callback:
                        progress_callback(f"转换第{page_num}页为{file_ext.upper()}...")
                    
                    # 生成当前页的输出路径
                    # 使用 generate_output_path，section参数标识页码
                    description = get_image_format_description(actual_format)
                    page_output_path = generate_output_path(
                        original_path,
                        output_dir=output_dir,
                        section=f'page{page_num}',
                        add_timestamp=True,
                        description=description,
                        file_type=file_ext
                    )
                    
                    logger.debug(f"第{page_num}页输出路径: {os.path.basename(page_output_path)}")
                    
                    # 复制当前帧
                    frame = img.copy()
                    
                    # 处理颜色模式
                    converted_frame = self._convert_image_mode(frame, target_format)
                    
                    # 保存为目标格式
                    save_kwargs = self._get_save_kwargs(target_format)
                    converted_frame.save(page_output_path, format=pil_format, **save_kwargs)
                    
                    output_files.append(page_output_path)
                    logger.info(f"✓ 第{page_num}页转换成功: {os.path.basename(page_output_path)}")
            
            logger.info(f"多页TIFF转换完成，共转换{len(output_files)}页")
            
            if progress_callback:
                progress_callback("转换完成。")
            
            # 返回第一页的路径作为主输出
            return ConversionResult(
                success=True,
                output_path=output_files[0] if output_files else None,
                message=f"已成功转换为{file_ext.upper()}（共{len(output_files)}页）"
            )
        
        except Exception as e:
            logger.error(f"多页TIFF转换失败: {e}", exc_info=True)
            # 清理已创建的文件
            for output_file in output_files:
                try:
                    if os.path.exists(output_file):
                        os.remove(output_file)
                        logger.debug(f"已删除: {output_file}")
                except Exception as cleanup_error:
                    logger.warning(f"清理文件失败: {output_file}, 错误: {cleanup_error}")
            raise
    
    def _convert_image_mode(self, img, target_format: str):
        """
        根据目标格式转换图片颜色模式
        
        Args:
            img: PIL Image对象
            target_format: 目标格式（jpeg, png, gif等）
            
        Returns:
            转换后的PIL Image对象
        """
        from PIL import Image
        
        target_format = target_format.lower()
        
        # JPEG不支持透明度，需要转为RGB
        if target_format in ['jpeg', 'jpg']:
            if img.mode in ('RGBA', 'LA', 'P'):
                logger.debug(f"JPEG不支持透明度，{img.mode} → RGB（白色背景）")
                # 创建白色背景
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                if img.mode in ('RGBA', 'LA'):
                    background.paste(img, mask=img.split()[-1])  # 使用alpha通道
                else:
                    background.paste(img)
                return background
            elif img.mode != 'RGB':
                logger.debug(f"{img.mode} → RGB")
                return img.convert('RGB')
            return img
        
        # GIF支持透明度但只有256色
        elif target_format == 'gif':
            if img.mode not in ('P', 'L'):
                logger.debug(f"{img.mode} → P（256色）")
                return img.convert('P')
            return img
        
        # BMP不支持透明度
        elif target_format == 'bmp':
            if img.mode in ('RGBA', 'LA', 'P'):
                logger.debug(f"BMP不支持透明度，{img.mode} → RGB（白色背景）")
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                if img.mode in ('RGBA', 'LA'):
                    background.paste(img, mask=img.split()[-1])
                else:
                    background.paste(img)
                return background
            elif img.mode != 'RGB':
                logger.debug(f"{img.mode} → RGB")
                return img.convert('RGB')
            return img
        
        # PNG和WebP支持透明度，保持RGBA
        elif target_format in ['png', 'webp']:
            if img.mode in ('P', 'L', 'LA'):
                logger.debug(f"{img.mode} → RGBA")
                return img.convert('RGBA')
            return img
        
        # TIFF支持多种模式，保持原样或转RGB
        elif target_format in ['tiff', 'tif']:
            if img.mode in ('P', 'LA'):
                logger.debug(f"{img.mode} → RGB")
                return img.convert('RGB')
            return img
        
        # 默认：保持原样
        return img
    
    def _normalize_extension(self, format_name: str) -> str:
        """
        标准化文件扩展名（统一使用简短格式）
        
        Args:
            format_name: 格式名称
            
        Returns:
            标准化的扩展名
        """
        format_map = {
            'jpeg': 'jpg',
            'jpg': 'jpg',
            'tiff': 'tif',
            'tif': 'tif',
        }
        return format_map.get(format_name.lower(), format_name.lower())
    
    def _get_pil_format(self, format_name: str) -> str:
        """
        获取PIL库使用的格式名称
        
        Args:
            format_name: 格式名称
            
        Returns:
            PIL格式名称（大写）
        """
        # PIL使用JPEG和TIFF作为格式名称
        format_map = {
            'jpeg': 'JPEG',
            'jpg': 'JPEG',
            'tiff': 'TIFF',
            'tif': 'TIFF',
            'png': 'PNG',
            'gif': 'GIF',
            'bmp': 'BMP',
            'webp': 'WEBP',
        }
        return format_map.get(format_name.lower(), format_name.upper())
    
    def _get_save_kwargs(self, target_format: str) -> dict:
        """
        获取保存图片时的参数
        
        Args:
            target_format: 目标格式
            
        Returns:
            保存参数字典
        """
        target_format = target_format.lower()
        
        if target_format in ['jpeg', 'jpg']:
            return {'quality': 95, 'optimize': True}
        elif target_format == 'png':
            return {'optimize': True}
        elif target_format == 'webp':
            return {'quality': 95, 'method': 6}  # method=6 最高质量
        elif target_format in ['tiff', 'tif']:
            return {'compression': 'tiff_lzw'}  # 使用LZW无损压缩
        else:
            return {}


@register_action("merge_images_to_tiff")
class MergeImagesToTiffStrategy(BaseStrategy):
    """
    将多个图片合并为单个多页TIFF文件的策略
    
    支持两种透明处理模式：
    1. smart（推荐）：智能判断
       - 如果所有图片都有透明通道，转为RGBA保留透明效果
       - 否则统一转为RGB，透明变白
    2. RGB：统一转为RGB模式，透明背景变白
    
    功能特点：
    - 按批量列表顺序合并（从上到下）
    - 保持原始尺寸
    - 输出到第一个文件所在目录
    """
    
    def execute(
        self,
        file_path: str,
        options: Optional[Dict[str, Any]] = None,
        progress_callback: Optional[Callable[[str], None]] = None
    ) -> ConversionResult:
        """
        执行图片合并为TIFF
        
        Args:
            file_path: 第一个图片文件路径（用于确定输出目录）
            options: 合并选项
                - mode: "smart" 或 "RGB"（默认"smart"）
                - file_list: 要合并的文件路径列表（必需）
            progress_callback: 进度更新回调函数
            
        Returns:
            ConversionResult: 转换结果对象
        """
        try:
            from PIL import Image
            
            # 获取选项
            mode = options.get('mode', 'smart') if options else 'smart'
            file_list = options.get('file_list', []) if options else []
            
            if not file_list:
                return ConversionResult(
                    success=False,
                    message="没有要合并的文件"
                )
            
            logger.info(f"开始合并 {len(file_list)} 个图片为TIFF，模式: {mode}")
            
            if progress_callback:
                progress_callback(f"准备合并 {len(file_list)} 个图片...")
            
            # 输出目录为选中文件所在目录
            from gongwen_converter.utils.workspace_manager import get_output_directory
            output_dir = get_output_directory(file_path)
            
            # 步骤1：加载所有图片
            images = []
            has_alpha_list = []
            
            for idx, img_path in enumerate(file_list, 1):
                if progress_callback:
                    progress_callback(f"加载第 {idx}/{len(file_list)} 个图片...")
                
                try:
                    img = Image.open(img_path)
                    # 检测是否有透明通道
                    has_alpha = img.mode in ['RGBA', 'LA', 'PA']
                    has_alpha_list.append(has_alpha)
                    
                    images.append(img)
                    logger.debug(f"✓ 加载图片 {idx}: {os.path.basename(img_path)} (模式: {img.mode}, 透明: {has_alpha})")
                except Exception as e:
                    logger.warning(f"跳过无法加载的图片: {img_path}, 错误: {e}")
                    continue
            
            if not images:
                return ConversionResult(
                    success=False,
                    message="没有成功加载任何图片"
                )
            
            if progress_callback:
                progress_callback("确定颜色模式...")
            
            # 步骤2：确定目标颜色模式
            if mode == 'smart':
                # 智能模式：如果所有图片都有透明，用RGBA；否则用RGB
                all_have_alpha = all(has_alpha_list)
                target_mode = 'RGBA' if all_have_alpha else 'RGB'
                logger.info(f"智能模式 - 所有图片透明: {all_have_alpha}, 目标模式: {target_mode}")
            else:
                # RGB模式：统一转RGB
                target_mode = 'RGB'
                logger.info(f"RGB模式 - 目标模式: {target_mode}")
            
            if progress_callback:
                progress_callback(f"转换为{target_mode}模式...")
            
            # 步骤3：转换所有图片到目标模式
            converted_images = []
            for idx, img in enumerate(images, 1):
                if img.mode != target_mode:
                    if target_mode == 'RGB' and img.mode in ['RGBA', 'LA', 'PA']:
                        # 转RGB：创建白色背景
                        logger.debug(f"图片 {idx}: {img.mode} → RGB (白色背景)")
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        if img.mode == 'PA':
                            img = img.convert('RGBA')
                        if img.mode in ['RGBA', 'LA']:
                            background.paste(img, mask=img.split()[-1])
                        else:
                            background.paste(img)
                        converted_images.append(background)
                    else:
                        # 其他模式转换
                        logger.debug(f"图片 {idx}: {img.mode} → {target_mode}")
                        converted_images.append(img.convert(target_mode))
                else:
                    logger.debug(f"图片 {idx}: 已是{target_mode}模式")
                    converted_images.append(img)
            
            if progress_callback:
                progress_callback("生成TIFF文件...")
            
            # 步骤4：生成输出文件名（使用选中文件所在目录）
            from datetime import datetime
            from gongwen_converter.utils.workspace_manager import get_output_directory
            selected_file = options.get("selected_file", file_path) if options else file_path
            output_dir = get_output_directory(selected_file)
            logger.info(f"输出目录基于: {os.path.basename(selected_file)}")
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"合并TIF_{timestamp}.tif"
            output_path = os.path.join(output_dir, output_filename)
            
            # 步骤5：保存为多页TIFF
            logger.info(f"保存多页TIFF文件: {output_filename}")
            converted_images[0].save(
                output_path,
                save_all=True,
                append_images=converted_images[1:] if len(converted_images) > 1 else [],
                compression='tiff_lzw'  # 使用LZW无损压缩
            )
            
            logger.info(f"✓ 成功合并 {len(converted_images)} 个图片为TIFF: {output_filename}")
            
            if progress_callback:
                progress_callback("合并完成。")
            
            return ConversionResult(
                success=True,
                output_path=output_path,
                message=f"已成功合并 {len(converted_images)} 个图片为TIFF"
            )
        
        except Exception as e:
            logger.error(f"合并图片为TIFF失败: {e}", exc_info=True)
            return ConversionResult(
                success=False,
                message=f"合并失败: {e}",
                error=e
            )


# 模块测试
if __name__ == "__main__":
    logging.basicConfig(
        level=logging.DEBUG,
        format="%(asctime)s [%(levelname)s] %(name)s:%(lineno)d - %(message)s"
    )
    
    logger.info("图片转换策略模块测试开始")
    
    # 测试文件路径（需要实际存在的文件）
    test_image = "test.jpg"
    
    if os.path.exists(test_image):
        # 测试图片转PDF（原图嵌入）
        try:
            logger.info(f"测试图片转PDF（原图嵌入）: {test_image}")
            strategy = ImageToPdfStrategy()
            result = strategy.execute(test_image, {'quality_mode': 'original'})
            if result.success:
                logger.info(f"测试成功! 输出文件: {result.output_path}")
            else:
                logger.error(f"测试失败: {result.message}")
        except Exception as e:
            logger.error(f"测试异常: {e}")
        
        # 测试图片转Markdown
        try:
            logger.info(f"测试图片转Markdown: {test_image}")
            strategy = ImageToMarkdownStrategy()
            result = strategy.execute(test_image)
            if result.success:
                logger.info(f"测试成功! 输出文件: {result.output_path}")
            else:
                logger.error(f"测试失败: {result.message}")
        except Exception as e:
            logger.error(f"测试异常: {e}")
        
        # 测试图片转DOCX
        try:
            logger.info(f"测试图片转DOCX: {test_image}")
            strategy = ImageToDocxStrategy()
            result = strategy.execute(test_image)
            if result.success:
                logger.info(f"测试成功! 输出文件: {result.output_path}")
            else:
                logger.error(f"测试失败: {result.message}")
        except Exception as e:
            logger.error(f"测试异常: {e}")
    else:
        logger.warning(f"测试文件不存在: {test_image}")
    
    logger.info("图片转换策略模块测试结束")
