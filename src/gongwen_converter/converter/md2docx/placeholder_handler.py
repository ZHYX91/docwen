"""
占位符处理器模块
实现DOCX文档中占位符的识别、标记和处理功能
使用隐藏文本标记方案替代直接XML操作，提高健壮性
"""

import re
import logging
import copy
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Pt
from gongwen_converter.utils import docx_utils
from gongwen_converter.utils.text_utils import format_display_value
from gongwen_converter.utils.validation_utils import is_value_empty

# 配置日志
logger = logging.getLogger(__name__)

# 特殊标记格式
SPECIAL_MARKER_PREFIX = "%%SPECIAL_PH:"
SPECIAL_MARKER_SUFFIX = "%%"

# 特殊处理的占位符
SPECIAL_PLACEHOLDERS = {
    "正文": "{{正文}}",
    "附件说明": "{{附件说明}}"
}
SPECIAL_PLACEHOLDER_LIST = list(SPECIAL_PLACEHOLDERS.values())

# 占位符处理规则配置
# 
# 配置结构说明：
# - 每个规则类型包含若干"组"
# - 每个"组"是一个字段列表（单字段或多字段）
# - 只有当组内所有字段都为空（不存在或值为空）时，才执行对应操作
#
# 支持5种处理方式：
# 1. 默认行为：不在任何规则中的占位符，直接替换为空字符串
# 2. delete_paragraph_if_empty: 删除占位符所在段落（仅在段落中生效）
# 3. delete_cell_if_empty: 清空占位符所在单元格（仅在表格中生效，检查单元格级别）
# 4. delete_row_if_empty: 删除占位符所在表格行（仅在表格中生效，检查整行级别）
# 5. delete_table_if_empty: 删除占位符所在整个表格（保留用于未来扩展）
#
# 示例：
# - ["密级和保密期限"]：单字段组，该字段为空时删除段落
# - ["印发机关", "印发日期"]：多字段组，两个字段都为空，且在同一行时，才删除表格行
PLACEHOLDER_RULES = {
    "delete_paragraph_if_empty": [
        ["密级和保密期限"],  # 单字段组
        ["紧急程度"],
        ["发文字号"],
        ["公开方式"],
        ["主送机关"],
        ["附注"],
        ["抄送机关"],
        ["附件说明"],
        ["份号", "发文字号"],  # 多字段组：两个字段都为空，且在同一段，才删除段落
    ],
    
    "delete_cell_if_empty": [
        # 清空单元格规则：组内所有字段都在同一单元格中且都为空时，清空该单元格
        # 示例：["字段1", "字段2"] - 两个字段都在同一单元格且都为空时清空单元格
    ],
    
    "delete_row_if_empty": [
        ["抄送机关"],           # 单字段组
        ["印发机关", "印发日期"],  # 多字段组：两个字段都为空且在同一行，才删除表格行
    ],
    
    "delete_table_if_empty": [
        # 注意：当表格所有行都被delete_row_if_empty规则删除时，
        # 会自动触发删除整个表格的逻辑（在docx_processor.py中实现）
        # 此配置保留用于未来可能的明确指定场景
    ],
}

def mark_special_placeholders(doc):
    """
    标记特殊占位符段落 - 使用隐藏文本标记方案
    在包含特殊占位符的段落末尾添加隐藏标记Run
    """
    logger.info("开始标记特殊占位符段落 (使用隐藏文本标记方案)...")
    
    # 获取特殊占位符列表（确保是字符串）
    special_texts = SPECIAL_PLACEHOLDER_LIST
    logger.debug(f"特殊占位符列表: {special_texts}")
    
    # 遍历所有段落
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if not text:
            continue
            
        # 检查段落是否包含任何特殊占位符
        for ph_text in special_texts:
            if ph_text in text:
                # 获取占位符类型名称
                ph_type = None
                for key, value in SPECIAL_PLACEHOLDERS.items():
                    if value == ph_text:
                        ph_type = key
                        break
                
                if not ph_type:
                    logger.warning(f"未找到占位符类型: {ph_text}")
                    continue
                
                # 检查是否已标记
                if is_special_marked(paragraph):
                    logger.debug(f"段落已标记，跳过: {text[:30]}...")
                    continue
                
                # 添加隐藏标记Run
                marker_text = f"{SPECIAL_MARKER_PREFIX}{ph_type}{SPECIAL_MARKER_SUFFIX}"
                marker_run = paragraph.add_run()
                marker_run.text = marker_text
                
                # 设置隐藏样式（白色字体，极小字号）
                font = marker_run.font
                font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # 白色
                font.size = Pt(1)  # 极小字号
                
                logger.debug(f"添加隐藏标记: {marker_text} -> 段落: {text[:30]}...")
                break  # 一个段落只标记一次
    
    logger.info("特殊占位符标记完成")

def is_special_marked(paragraph):
    """
    检查段落是否被标记为特殊占位符
    通过检查段落末尾的隐藏标记Run
    
    参数:
        paragraph: 段落对象
        
    返回:
        bool: 如果是特殊标记段落返回True，否则False
    """
    # 检查段落是否有Run
    if not paragraph.runs:
        return False
    
    # 检查最后一个Run是否是隐藏标记
    last_run = paragraph.runs[-1]
    last_run_text = last_run.text.strip()
    
    # 检查是否符合标记格式
    is_marked = (
        last_run_text.startswith(SPECIAL_MARKER_PREFIX) and 
        last_run_text.endswith(SPECIAL_MARKER_SUFFIX)
    )
    
    if is_marked:
        logger.debug(f"检测到隐藏标记: {last_run_text}")
    
    return is_marked

def get_special_mark_type(paragraph):
    """
    获取特殊标记的类型
    
    参数:
        paragraph: 段落对象
        
    返回:
        str: 标记类型名称，如果不是特殊标记返回None
    """
    if not is_special_marked(paragraph):
        return None
    
    # 提取标记类型
    last_run_text = paragraph.runs[-1].text.strip()
    mark_type = last_run_text.replace(SPECIAL_MARKER_PREFIX, "").replace(SPECIAL_MARKER_SUFFIX, "")
    
    logger.debug(f"提取标记类型: {mark_type}")
    return mark_type

def remove_special_mark(paragraph):
    """
    移除特殊标记
    通过删除标记Run实现
    
    参数:
        paragraph: 段落对象
    """
    if not is_special_marked(paragraph):
        return
    
    try:
        # 获取标记Run
        last_run = paragraph.runs[-1]
        
        # 检查是否是标记Run
        last_run_text = last_run.text.strip()
        if (
            last_run_text.startswith(SPECIAL_MARKER_PREFIX) and 
            last_run_text.endswith(SPECIAL_MARKER_SUFFIX)
        ):
            # 删除标记Run
            p = last_run._element.getparent()
            p.remove(last_run._element)
            logger.debug("已移除隐藏标记")
    except Exception as e:
        logger.warning(f"移除标记失败: {str(e)}")

def extract_placeholder_keys(text: str):
    """
    从文本中提取所有占位符键名
    返回: 包含所有占位符键名的列表
    """
    # 使用正则表达式匹配所有 {{key}} 格式的占位符
    pattern = r'\{\{(\w+)\}\}'
    matches = re.findall(pattern, text)
    return list(set(matches))  # 去重后返回

def get_placeholder_rules(key: str):
    """
    查找占位符所属的所有规则和组
    
    工作原理：
    1. 遍历PLACEHOLDER_RULES中的所有规则类型
    2. 对每个规则类型，遍历其包含的所有组
    3. 如果占位符键在某个组中，将该规则类型和组添加到结果列表
    4. 返回所有匹配的规则列表
    
    参数:
        key: 占位符键名（如"密级和保密期限"）
        
    返回:
        list: [(rule_type, group), ...] 规则列表
              每个元组包含：
              - rule_type: 规则类型字符串
                * 'delete_paragraph_if_empty': 删除段落（仅在段落中生效）
                * 'delete_cell_if_empty': 清空单元格（仅在表格中生效）
                * 'delete_row_if_empty': 删除表格行（仅在表格中生效）
                * 'delete_table_if_empty': 删除整个表格
              - group: 包含该键的组（字段列表）
              
              如果不属于任何规则，返回空列表，表示使用默认行为（替换为空字符串）
    """
    rules = []
    
    # 遍历所有规则类型和组
    for rule_type, groups in PLACEHOLDER_RULES.items():
        for group in groups:
            if key in group:
                rules.append((rule_type, group))
                logger.debug(f"占位符 '{key}' 属于规则 '{rule_type}', 组: {group}")
    
    # 如果不属于任何规则，返回空列表
    if not rules:
        logger.debug(f"占位符 '{key}' 不属于任何特殊规则，将替换为空字符串")
    
    return rules

def check_group_all_empty(group: list, yaml_data: dict):
    """
    检查组内所有字段是否都为空
    
    工作原理：
    1. 遍历组内的每个字段
    2. 如果字段在YAML数据中存在：
       - 获取其值并检查是否为空（使用is_value_empty函数）
       - 如果不为空，立即返回False，表示组不是全空的
    3. 如果字段不存在，视为空值
    4. 只有当组内所有字段都为空时，才返回True
    
    判断"空"的标准（由is_value_empty函数实现）：
    - None值
    - 空字符串
    - 空列表
    - 仅包含空白字符的字符串
    
    参数:
        group: 字段名列表，如["印发机关", "印发日期"]
        yaml_data: YAML数据字典
        
    返回:
        bool: 如果组内所有字段都为空（不存在或值为空）则返回True
        
    示例:
        组["印发机关", "印发日期"]：
        - 如果两个字段都为空 → True
        - 如果任何一个字段有值 → False
    """
    if not group:
        return False
    
    # 遍历组内所有字段
    for field in group:
        # 检查字段是否存在
        if field in yaml_data:
            value = yaml_data[field]
            # 如果字段不为空，组就不是全空的
            if not is_value_empty(value):
                logger.debug(f"组中字段 '{field}' 不为空: {value}")
                return False
        else:
            # 字段不存在，视为空
            logger.debug(f"组中字段 '{field}' 不存在，视为空")
    
    # 所有字段都为空
    logger.debug(f"组 {group} 中所有字段都为空")
    return True

def process_paragraph_placeholders(paragraph, yaml_data, paragraphs_to_remove):
    """
    处理单个段落中的占位符（支持组处理）
    
    处理流程：
    1. 提取段落中所有占位符键名
    2. 第一遍遍历：检查是否需要删除段落
       - 查找每个占位符所属的规则和组
       - 使用processed_groups避免重复检查同一组
       - 如果是delete_paragraph_if_empty规则且组内全空，标记删除段落
    3. 如果需要删除段落，加入删除列表并返回
    4. 第二遍遍历：正常替换占位符
       - 从YAML数据获取值或使用空字符串
       - 调用replace_placeholder_in_paragraph进行替换
    
    示例：
    - 段落包含{{密级和保密期限}}，该字段为空 → 删除段落
    - 段落包含{{标题}}，有值 → 替换为实际值
    - 段落包含{{不存在的字段}} → 替换为空字符串
    
    参数:
        paragraph: 段落对象
        yaml_data: YAML数据字典
        paragraphs_to_remove: 需要删除的段落列表（输出参数）
    """
    text = paragraph.text
    logger.debug(f"处理段落: {text[:50]}...")
    
    # 步骤1：提取段落中的所有占位符键
    placeholder_keys = extract_placeholder_keys(text)
    if not placeholder_keys:
        logger.debug("段落中无占位符，跳过处理")
        return
    
    logger.debug(f"找到占位符键: {placeholder_keys}")
    
    # 用于跟踪已处理的组（避免重复检查同一组）
    # 例如：如果段落包含["印发机关", "印发日期"]，只需检查一次组
    processed_groups = set()
    paragraph_should_remove = False
    
    # 步骤2：第一遍遍历 - 检查是否需要删除段落
    for key in placeholder_keys:
        # 查找占位符所属的所有规则
        rules = get_placeholder_rules(key)
        
        # 只处理 delete_paragraph_if_empty 规则
        for rule_type, group in rules:
            if rule_type != 'delete_paragraph_if_empty':
                continue  # 跳过非段落规则
            
            # 跳过已处理的组（避免重复检查）
            group_tuple = tuple(group)
            if group_tuple in processed_groups:
                logger.debug(f"组 {group} 已处理，跳过")
                continue
            processed_groups.add(group_tuple)
            
            # 【关键】检查组内所有字段是否都在当前段落中
            # 只有组内所有字段都在同一段落，才判断是否删除
            if all(field in placeholder_keys for field in group):
                # 组内所有字段都在当前段落，检查是否都为空
                if check_group_all_empty(group, yaml_data):
                    logger.info(f"组 {group} 所有字段都在段落中且都为空，标记删除段落")
                    paragraph_should_remove = True
                    break  # 找到需要删除的原因，不再继续检查
            else:
                # 组内字段不全在当前段落，不删除
                logger.debug(f"组 {group} 的字段不全在当前段落中，跳过删除检查")
        
        if paragraph_should_remove:
            break  # 外层循环也要退出
    
    # 步骤3：如果需要删除段落，添加到移除列表并返回
    if paragraph_should_remove:
        paragraphs_to_remove.append(paragraph)
        logger.debug("段落已加入删除列表")
        return
    
    # 步骤4：第二遍遍历 - 正常替换占位符
    for key in placeholder_keys:
        placeholder = f'{{{{{key}}}}}'
        
        # 获取值并格式化
        if key in yaml_data:
            value = yaml_data[key]
            display_value = format_display_value(value)
            logger.debug(f"字段 '{key}' 的值: {display_value}")
        else:
            # 键不存在，替换为空字符串（默认行为）
            display_value = ""
            logger.debug(f"键 '{key}' 不存在，替换为空字符串")
        
        # 执行替换
        if placeholder in text:
            replace_placeholder_in_paragraph(paragraph, placeholder, display_value)
            logger.debug(f"成功替换占位符: {placeholder} -> {display_value}")

def process_table_cell_placeholders(paragraph, yaml_data, row, rows_to_remove, row_placeholder_keys=None):
    """
    处理表格单元格中的占位符（支持组处理）
    
    处理流程：
    1. 提取单元格段落中所有占位符键名
    2. 第一遍遍历：检查是否需要删除行或清空单元格
       - 查找每个占位符所属的规则和组
       - 使用processed_groups避免重复检查同一组
       - delete_row_if_empty规则：检查组内所有字段是否都在整行中，且都为空时删除整行（优先级高）
       - delete_paragraph_if_empty规则：组内全空时清空单元格内容
    3. 如果需要删除行，加入删除列表并返回True
    4. 如果需要清空单元格，清空所有内容并返回False
    5. 第二遍遍历：正常替换占位符
       - 从YAML数据获取值或使用空字符串
       - 调用replace_placeholder_in_paragraph进行替换
    
    注意事项：
    - delete_row_if_empty优先级最高，一旦触发立即返回
    - delete_row_if_empty检查组内所有字段是否都在整行中（跨单元格）
    - delete_paragraph_if_empty在表格中表现为清空单元格，而非删除段落
    - 如果表格所有行都被删除，在docx_processor.py中会自动删除整个表格
    
    示例：
    - 行包含{{抄送机关}}，该字段为空 → 删除整行
    - 行包含{{印发机关}}（单元格1）和{{印发日期}}（单元格2），两者都为空 → 删除整行
    - 行只有{{印发机关}}，印发日期不在该行 → 不删除
    - 单元格包含{{密级和保密期限}}，该字段为空 → 清空单元格
    - 单元格包含{{发文字号}}，有值 → 替换为实际值
    
    参数:
        paragraph: 表格单元格中的段落对象
        yaml_data: YAML数据字典
        row: 当前表格行对象
        rows_to_remove: 需要删除的行列表（输出参数）
        row_placeholder_keys: 整行的所有占位符键集合（可选，用于行级检查）
        
    返回:
        bool: 是否需要删除整行（True=删除行，False=不删除行）
    """
    text = paragraph.text
    logger.debug(f"处理表格单元格: {text[:50]}...")
    
    # 步骤1：提取段落中的所有占位符键
    placeholder_keys = extract_placeholder_keys(text)
    if not placeholder_keys:
        logger.debug("单元格中无占位符，跳过处理")
        return False
    
    logger.debug(f"找到占位符键: {placeholder_keys}")
    
    # 用于跟踪已处理的组（避免重复检查同一组）
    processed_groups = set()
    cell_should_clear = False
    row_should_remove = False
    
    # 步骤2：第一遍遍历 - 检查是否需要删除行或清空单元格
    for key in placeholder_keys:
        # 查找占位符所属的所有规则
        rules = get_placeholder_rules(key)
        
        # 只处理表格相关规则：delete_row_if_empty 和 delete_cell_if_empty
        for rule_type, group in rules:
            # 跳过非表格规则
            if rule_type not in ('delete_row_if_empty', 'delete_cell_if_empty'):
                continue
            
            # 跳过已处理的组（避免重复检查）
            group_tuple = tuple(group)
            if group_tuple in processed_groups:
                logger.debug(f"组 {group} 已处理，跳过")
                continue
            processed_groups.add(group_tuple)
            
            # 处理delete_row_if_empty规则（优先级高）
            if rule_type == 'delete_row_if_empty':
                # 【关键】检查组内所有字段是否都在整行中
                # 如果提供了row_placeholder_keys，使用整行检查；否则只检查当前单元格
                if row_placeholder_keys is not None:
                    # 方案B（宽松）：检查整行
                    if all(field in row_placeholder_keys for field in group):
                        # 组内所有字段都在整行中，检查是否都为空
                        if check_group_all_empty(group, yaml_data):
                            logger.info(f"组 {group} 所有字段都在整行中且都为空，标记删除行")
                            row_should_remove = True
                            break  # 找到需要删除行的原因，不再继续检查
                    else:
                        # 组内字段不全在整行中，不删除
                        logger.debug(f"组 {group} 的字段不全在整行中，跳过删除检查")
                else:
                    # 未提供整行占位符集合，降级为检查当前单元格（向后兼容）
                    logger.warning("未提供row_placeholder_keys，使用单元格级检查")
                    if check_group_all_empty(group, yaml_data):
                        logger.info(f"组 {group} 所有字段都为空，标记删除行")
                        row_should_remove = True
                        break
            
            # 处理delete_cell_if_empty规则（优先级中）
            elif rule_type == 'delete_cell_if_empty':
                # 【关键】检查组内所有字段是否都在当前单元格中
                # 只有组内所有字段都在同一单元格，才判断是否清空
                if all(field in placeholder_keys for field in group):
                    # 组内所有字段都在当前单元格，检查是否都为空
                    if check_group_all_empty(group, yaml_data):
                        logger.info(f"组 {group} 所有字段都在单元格中且都为空，标记清空单元格")
                        cell_should_clear = True
                else:
                    # 组内字段不全在当前单元格，不清空
                    logger.debug(f"组 {group} 的字段不全在当前单元格中，跳过清空检查")
        
        if row_should_remove:
            break  # 外层循环也要退出
    
    # 步骤3：如果行需要删除，添加到移除列表并返回True
    if row_should_remove:
        rows_to_remove.append(row)
        logger.debug("行已加入删除列表")
        return True
    
    # 步骤4：如果需要清空单元格，清空所有内容
    if cell_should_clear:
        # 获取单元格元素（tc元素）
        cell_element = paragraph._element.getparent()
        # 遍历单元格中的所有段落
        for para_element in cell_element.findall('.//w:p', docx_utils.NAMESPACES):
            # 遍历段落中的所有run
            for run_element in para_element.findall('.//w:r', docx_utils.NAMESPACES):
                # 遍历run中的所有文本元素
                for text_element in run_element.findall('.//w:t', docx_utils.NAMESPACES):
                    text_element.text = ""
        logger.debug("已清空单元格内容")
        return False
    
    # 步骤5：第二遍遍历 - 正常替换占位符
    for key in placeholder_keys:
        placeholder = f'{{{{{key}}}}}'
        
        # 获取值并格式化
        if key in yaml_data:
            value = yaml_data[key]
            display_value = format_display_value(value)
            logger.debug(f"字段 '{key}' 的值: {display_value}")
        else:
            # 键不存在，替换为空字符串（默认行为）
            display_value = ""
            logger.debug(f"键 '{key}' 不存在，替换为空字符串")
        
        # 执行替换
        if placeholder in text:
            replace_placeholder_in_paragraph(paragraph, placeholder, display_value)
            logger.debug(f"成功替换占位符: {placeholder} -> {display_value}")
    
    return False

def replace_placeholder_in_paragraph(paragraph, placeholder, value):
    """
    精确替换段落中的占位符（合并相同样式的run后处理）
    
    参数:
        paragraph: 段落对象
        placeholder: 占位符文本
        value: 要替换的值
    """
    logger.debug(f"开始处理段落占位符: {placeholder}")
    logger.debug(f"段落原始文本: {paragraph.text}")
    
    # 情况1：整个段落就是占位符（最高优先级）
    if paragraph.text and paragraph.text.strip() == placeholder.strip():
        logger.debug("情况1：整个段落就是占位符")
        paragraph.text = value
        return
    
    # 步骤2：合并相同样式的连续run
    logger.debug("步骤2：合并相同样式的连续run")
    merged_runs = merge_similar_runs(paragraph)
    
    # 情况3：在单个run内查找占位符
    logger.debug("情况3：在单个run内查找占位符")
    for i, run in enumerate(paragraph.runs):
        if placeholder in run.text:
            logger.debug(f"在run[{i}]中找到占位符: {run.text}")
            run.text = run.text.replace(placeholder, value)
            return
    
    # 情况4：处理跨多个run的占位符
    logger.debug("情况4：处理跨多个run的占位符")
    full_text = ''.join(run.text for run in paragraph.runs)
    if placeholder in full_text:
        logger.debug(f"在合并文本中找到跨run占位符: {full_text}")
        
        # 保存第一个run的样式（用于保持格式）
        first_run = paragraph.runs[0] if paragraph.runs else None
        
        # 清空所有现有run
        for run in paragraph.runs:
            run.text = ""
        
        # 创建新run并应用原样式
        new_text = full_text.replace(placeholder, value)
        new_run = paragraph.add_run(new_text)
        
        # 应用原始样式
        if first_run and first_run._element.rPr is not None:
            logger.debug("应用原始样式到新run")
            docx_utils.apply_run_style(new_run, first_run)
        
        logger.debug(f"跨run替换完成: {new_text}")
        return
    
    logger.warning(f"在段落中未找到占位符: {placeholder}")

def merge_similar_runs(paragraph):
    """
    合并相同样式的连续run
    
    参数:
        paragraph: 段落对象
        
    返回:
        list: 合并后的run列表
    """
    if not paragraph.runs or len(paragraph.runs) < 2:
        logger.debug("无需合并：少于2个run")
        return paragraph.runs
    
    logger.debug(f"原始run数量: {len(paragraph.runs)}")
    
    # 创建一个新的run列表用于合并
    merged_runs = []
    current_run = None
    
    for run in paragraph.runs:
        # 如果是第一个run，直接开始新组
        if current_run is None:
            current_run = {
                'text': run.text,
                'rPr': run._element.rPr
            }
            continue
        
        # 检查当前run是否与前一个样式相同
        if docx_utils.is_rPr_equal(current_run['rPr'], run._element.rPr):
            logger.debug(f"合并相同样式run: '{current_run['text']}' + '{run.text}'")
            current_run['text'] += run.text
        else:
            # 保存当前组并开始新组
            merged_runs.append(current_run)
            current_run = {
                'text': run.text,
                'rPr': run._element.rPr
            }
    
    # 添加最后一组
    if current_run is not None:
        merged_runs.append(current_run)
    
    logger.debug(f"合并后run数量: {len(merged_runs)}")
    
    # 清空原始run
    for run in paragraph.runs[:]:
        p = run._element.getparent()
        p.remove(run._element)
    
    # 创建新的合并后的run
    new_runs = []
    for merged in merged_runs:
        new_run = paragraph.add_run(merged['text'])
        if merged['rPr'] is not None:
            docx_utils.apply_run_style(new_run, merged['rPr'])
        new_runs.append(new_run)
    
    return new_runs

def find_run_at_position(paragraph, position):
    """
    在段落中查找包含指定文本位置的run
    
    参数:
        paragraph: 段落对象
        position: 文本位置
        
    返回:
        Run对象 或 None
    """
    current_pos = 0
    for run in paragraph.runs:
        run_length = len(run.text)
        if current_pos <= position < current_pos + run_length:
            return run
        current_pos += run_length
    return paragraph.runs[0] if paragraph.runs else None

def try_remove_element(element):
    """安全移除元素"""
    try:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)
    except Exception as e:
        logger.warning(f"移除元素失败: {str(e)}")

def process_attachment_description_placeholder(doc, yaml_data):
    """
    处理附件说明占位符
    根据附件数量自动调整悬挂缩进
    """
    logger.info("开始处理附件说明占位符...")
    
    # 获取占位符
    attach_desc_ph = SPECIAL_PLACEHOLDERS.get('附件说明', '{{附件说明}}')
    attach_desc_para = None
    
    # 查找占位符段落
    for para_idx, para in enumerate(doc.paragraphs):
        if attach_desc_ph in para.text:
            remove_special_mark(para)
            logger.debug(f"在第 {para_idx+1} 段找到附件说明占位符")
            attach_desc_para = para
            break
    
    if not attach_desc_para:
        logger.info("未找到附件说明占位符")
        return
    
    # 获取附件说明数据
    attachment_desc = yaml_data.get('附件说明')
    
    # 如果为空，删除占位符段落
    if is_value_empty(attachment_desc):
        logger.info("附件说明为空，删除占位符段落")
        try_remove_element(attach_desc_para._element)
        return
    
    # 确保是列表
    if not isinstance(attachment_desc, list):
        attachment_desc = [attachment_desc]
    
    # 保存原始段落样式和格式
    base_style = attach_desc_para.style
    base_rpr = None
    if attach_desc_para.runs:
        base_rpr = attach_desc_para.runs[0]._element.rPr
    
    # 【关键】从模板段落的左侧缩进计算字符宽度
    left_indent = attach_desc_para.paragraph_format.left_indent
    if left_indent:
        char_width = left_indent / 2  # 左侧缩进2字符，所以除以2得到单字符宽度
        logger.info(f"从模板获取字符宽度: {char_width} (左侧缩进: {left_indent})")
    else:
        # 如果模板未设置left_indent，使用默认值
        char_width = Pt(16)  # 默认按三号字计算
        logger.warning(f"模板未设置左侧缩进，使用默认字符宽度: {char_width}")
        # 设置默认左侧缩进
        left_indent = 2 * char_width
    
    # 获取插入位置
    parent = attach_desc_para._element.getparent()
    index = parent.index(attach_desc_para._element)
    
    # 判断是单附件还是多附件
    is_single = len(attachment_desc) == 1
    
    # 计算悬挂缩进
    # 注意：python-docx要求缩进值必须是整数，所以需要转换
    # 使用方案A（文本加空格）+ 悬挂缩进：
    #   - 单附件：悬挂缩进3字符
    #   - 多附件：所有行悬挂缩进4.5字符（第2行起文本中已加空格）
    
    if is_single:
        # 单附件：左侧2字符 + 悬挂3字符
        hanging_indent_single = int(3 * char_width)
        logger.info(f"单附件：left_indent = {left_indent}, hanging_indent = {hanging_indent_single}")
    else:
        # 多附件：左侧2字符 + 悬挂4.5字符（所有行相同）
        hanging_indent_multi = int(4.5 * char_width)
        logger.info(f"多附件：left_indent = {left_indent}, hanging_indent = {hanging_indent_multi}")
    
    # 插入附件说明行
    for i, line in enumerate(attachment_desc):
        new_p = doc.add_paragraph(style=base_style)
        new_run = new_p.add_run(str(line))
        
        # 应用基本run样式
        if base_rpr is not None:
            new_rPr = new_run._element.get_or_add_rPr()
            for child in base_rpr.getchildren():
                new_child = copy.deepcopy(child)
                new_rPr.append(new_child)
        
        # 设置段落格式：保持模板的左侧缩进，设置悬挂缩进
        pf = new_p.paragraph_format
        pf.left_indent = left_indent  # 保持模板的2字符左侧缩进
        
        # 设置悬挂缩进（通过负的first_line_indent实现）
        if is_single:
            # 单附件：悬挂3字符
            pf.first_line_indent = -hanging_indent_single
        else:
            # 多附件：所有行悬挂4.5字符
            pf.first_line_indent = -hanging_indent_multi
        
        # 插入到正确位置
        new_p._p.getparent().remove(new_p._p)
        parent.insert(index + i, new_p._p)
        logger.debug(f"在第 {index + i} 位置插入附件说明行: {line[:20]}...")
    
    # 删除原始占位符段落
    logger.debug("删除附件说明占位符段落")
    parent.remove(attach_desc_para._element)
    
    logger.info("附件说明占位符处理完成")


def _process_paragraph_images(paragraph, context='document'):
    """
    处理单个段落中的图片占位符（核心处理逻辑）
    
    参数:
        paragraph: 段落对象
        context: 上下文标识 ('document' 或 'table')
    
    返回:
        int: 成功插入的图片数量
    """
    import os
    from docx.shared import Inches
    
    text = paragraph.text
    if not text:
        return 0
    
    # 图片占位符正则：{{IMAGE:路径}}
    image_pattern = r'\{\{IMAGE:([^}]+)\}\}'
    matches = re.findall(image_pattern, text)
    if not matches:
        return 0
    
    logger.debug(f"在{context}段落中找到 {len(matches)} 个图片占位符")
    inserted_count = 0
    
    # 处理每个图片占位符
    for image_path in matches:
        placeholder = f"{{{{IMAGE:{image_path}}}}}"
        
        # 检查图片文件是否存在
        if not os.path.exists(image_path):
            logger.warning(f"图片文件不存在，跳过: {image_path}")
            # 删除占位符
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, "")
            continue
        
        # 检查文件扩展名
        ext = os.path.splitext(image_path)[1].lower()
        supported_formats = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.emf', '.wmf'}
        if ext not in supported_formats:
            logger.warning(f"不支持的图片格式: {ext}，跳过: {image_path}")
            # 删除占位符
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, "")
            continue
        
        try:
            # 清空段落原有内容（包括占位符）
            for run in paragraph.runs[:]:
                p = run._element.getparent()
                p.remove(run._element)
            
            # 插入图片（使用默认宽度，保持宽高比）
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(4.5))
            
            # 设置段落行距为单倍行距（避免图片被裁剪）
            from docx.enum.text import WD_LINE_SPACING
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            logger.debug("已设置图片段落为单倍行距")
            
            inserted_count += 1
            logger.info(f"成功插入图片 ({inserted_count}) [{context}]: {os.path.basename(image_path)}")
            
        except Exception as e:
            logger.error(f"插入图片失败: {image_path} | 错误: {str(e)}")
            # 保留占位符供手动处理
            paragraph.add_run(placeholder)
    
    return inserted_count


def process_image_placeholders(doc):
    """
    处理文档中的图片占位符 {{IMAGE:路径}}
    
    处理流程：
    1. 遍历所有文档段落和表格单元格段落
    2. 查找包含 {{IMAGE:路径}} 的段落
    3. 提取图片路径
    4. 检查图片文件是否存在
    5. 在段落中插入图片并设置上下环绕
    6. 删除占位符文本
    
    参数:
        doc: Document对象
    
    返回:
        int: 成功插入的图片数量
    """
    logger.info("开始处理图片占位符...")
    
    total_inserted = 0
    
    # 1. 处理文档段落中的图片
    logger.debug("处理文档段落中的图片...")
    for paragraph in doc.paragraphs:
        count = _process_paragraph_images(paragraph, 'document')
        total_inserted += count
    
    # 2. 处理表格单元格中的图片
    logger.debug("处理表格单元格中的图片...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    count = _process_paragraph_images(paragraph, 'table')
                    total_inserted += count
    
    logger.info(f"图片占位符处理完成 | 成功插入: {total_inserted} 个（文档+表格）")
    return total_inserted
