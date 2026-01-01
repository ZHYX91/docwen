"""
表格处理模块

本模块负责 Word 表格的创建、样式应用和格式设置。
包括表格样式检测、注入、以及 WPS 兼容性处理。

主要功能：
- create_word_table: 创建 Word 表格并应用样式
- get_or_inject_table_style: 根据用户配置获取或注入表格样式
- inject_table_style: 注入表格样式（三线表/网格表/自定义）
- inject_table_content_style: 注入表格内容段落样式

WPS 兼容性说明：
WPS 对表格样式的 tblStylePr（条件格式）支持不完整，无法正确渲染 firstRow 中定义的边框。
因此需要在单元格级别（tcPr/tcBorders）显式设置边框，作为"双保险"。
"""

import logging
from lxml import etree
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ..style.templates import (
    get_localized_three_line_table_template,
    get_localized_table_grid_template,
    get_localized_table_content_template,
    get_custom_table_grid_template,
)
from .formula_handler import process_paragraph_formulas, is_formula_supported
from gongwen_converter.i18n.style_resolver import StyleNameResolver

# 配置日志
logger = logging.getLogger(__name__)

# OOXML 命名空间
WORD_NAMESPACE = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NSMAP = {'w': WORD_NAMESPACE}


def create_word_table(doc, table_data: dict, fonts: dict = None):
    """
    创建 Word 表格
    
    根据 Markdown 解析出的表格数据创建 Word 表格，并应用适当的样式。
    
    参数:
        doc: python-docx Document 对象
        table_data: 表格数据字典，包含：
            - headers: 表头列表
            - rows: 数据行列表
        fonts: 字体格式配置（可选，用于单元格文字）
    
    返回:
        Table: python-docx 表格对象
    """
    from .text_handler import add_formatted_text_to_paragraph_with_breaks
    from gongwen_converter.config.config_manager import config_manager
    
    headers = table_data['headers']
    rows = table_data['rows']
    
    logger.debug(f"创建Word表格: {len(headers)}列 x {len(rows)+1}行（含表头）")
    
    # 创建表格（表头 + 数据行）
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    
    # 1. 根据用户配置获取或注入表格样式
    table_style_name = get_or_inject_table_style(doc)
    logger.info(f"使用表格样式: {table_style_name}")
    
    # 应用表格样式（优雅降级）
    try:
        table.style = table_style_name
        logger.info(f"成功应用表格样式: {table_style_name}")
    except KeyError:
        # 三线表样式失败时，降级到 Normal Table（100% 存在的默认样式）
        logger.warning(f"表格样式 '{table_style_name}' 应用失败，回退到 Normal Table")
        table.style = 'Normal Table'
    
    # 2. 注入并获取表格内容样式
    content_style_name = inject_table_content_style(doc)
    
    # 3. 启用首行条件格式（让表格样式的 firstRow 生效）
    _enable_first_row_formatting(table)
    
    # 获取格式处理模式配置
    formatting_mode = config_manager.get_md_to_docx_formatting_mode()
    table_header_formatting_mode = config_manager.get_md_to_docx_table_header_formatting_mode()
    
    # 4. 设置表头
    header_row = table.rows[0]
    for col_idx, header in enumerate(headers):
        cell = header_row.cells[col_idx]
        header_text = str(header) if header else ''
        
        # 清空默认段落
        cell.text = ''
        paragraph = cell.paragraphs[0]
        
        # 应用表格内容样式
        try:
            paragraph.style = content_style_name
        except Exception:
            pass  # 样式不存在时忽略
        
        # 使用表头专用的格式处理模式（支持行内代码和换行）
        add_formatted_text_to_paragraph_with_breaks(paragraph, header_text, None, table_header_formatting_mode, doc=doc)
        
        logger.debug(f"设置表头单元格 [0,{col_idx}]: {header}")
    
    # 5. WPS 兼容性修复：为首行每个单元格显式设置底部边框
    _apply_header_row_bottom_border(header_row)
    
    # 6. 填充数据行
    for row_idx, row_data in enumerate(rows):
        data_row = table.rows[row_idx + 1]  # +1 因为第0行是表头
        for col_idx, cell_value in enumerate(row_data):
            if col_idx < len(headers):  # 确保不超出列数
                cell = data_row.cells[col_idx]
                cell_text = str(cell_value) if cell_value else ''
                
                # 清空默认段落
                cell.text = ''
                paragraph = cell.paragraphs[0]
                
                # 应用表格内容样式
                try:
                    paragraph.style = content_style_name
                except Exception:
                    pass  # 样式不存在时忽略
                
                # 检查是否包含公式
                has_formula = is_formula_supported() and '$' in cell_text
                
                if has_formula:
                    # 使用公式处理器处理单元格内容
                    process_paragraph_formulas(paragraph, cell_text)
                    logger.debug(f"填充数据单元格(含公式) [{row_idx+1},{col_idx}]")
                else:
                    # 使用格式解析（支持行内代码和换行）
                    add_formatted_text_to_paragraph_with_breaks(paragraph, cell_text, fonts, formatting_mode, doc=doc)
                    logger.debug(f"填充数据单元格 [{row_idx+1},{col_idx}]: {cell_value[:30] if cell_value else ''}...")
    
    logger.info(f"Word表格创建完成: {len(headers)}列 x {len(rows)}行数据")
    return table


def get_or_inject_table_style(doc) -> str:
    """
    根据用户配置获取或注入表格样式
    
    逻辑：
    1. 读取用户配置的样式（内置或自定义）
    2. 检查该样式是否在模板中已存在
    3. 如果存在 → 直接使用
    4. 如果不存在 → 注入样式定义
    
    参数:
        doc: python-docx Document 对象
    
    返回:
        str: 实际使用的样式名称
    """
    from gongwen_converter.config.config_manager import config_manager
    
    style_resolver = StyleNameResolver()
    styles_element = doc.part.styles._element
    
    # 获取表格样式配置
    style_mode = config_manager.get_table_style_mode()
    
    if style_mode == "custom":
        # 自定义样式模式
        custom_name = config_manager.get_custom_table_style_name()
        if custom_name:
            # 检查模板中是否已有该样式
            existing_style = _find_table_style_by_name(styles_element, custom_name)
            if existing_style:
                logger.debug(f"模板中找到用户配置的自定义表格样式: {existing_style}")
                return existing_style
            # 不存在则注入
            return _inject_custom_table_style(doc, custom_name, styles_element)
        else:
            # 自定义名称为空，回退到内置三线表
            logger.warning("自定义表格样式名称为空，回退到三线表")
            style_mode = "builtin"
    
    # 内置样式模式
    builtin_key = config_manager.get_builtin_table_style_key()
    
    if builtin_key == "table_grid":
        # 检查网格表是否已存在
        if not style_resolver.should_inject(doc, "table_grid"):
            usable_name = style_resolver.get_usable_name(doc, "table_grid")
            if usable_name:
                logger.debug(f"模板中找到用户配置的网格表样式: {usable_name}")
                return usable_name
        return _inject_table_grid_style(doc, style_resolver, styles_element)
    else:
        # 检查三线表是否已存在
        if not style_resolver.should_inject(doc, "three_line_table"):
            usable_name = style_resolver.get_usable_name(doc, "three_line_table")
            if usable_name:
                logger.debug(f"模板中找到用户配置的三线表样式: {usable_name}")
                return usable_name
        return _inject_three_line_table_style(doc, style_resolver, styles_element)


def _find_table_style_by_name(styles_element, style_name: str) -> str:
    """
    在模板中按名称查找表格样式
    
    参数:
        styles_element: styles.xml 元素
        style_name: 要查找的样式名称
    
    返回:
        str: 找到的样式名称，未找到返回 None
    """
    try:
        for style in styles_element.findall('.//w:style', namespaces=NSMAP):
            style_type = style.get(f'{{{WORD_NAMESPACE}}}type', '')
            if style_type != 'table':
                continue
            
            name_elem = style.find('w:name', namespaces=NSMAP)
            name = name_elem.get(f'{{{WORD_NAMESPACE}}}val', '') if name_elem is not None else ''
            
            if name.lower() == style_name.lower():
                return name
        return None
    except Exception as e:
        logger.warning(f"查找表格样式失败: {e}")
        return None


def inject_table_style(doc) -> str:
    """
    向文档注入表格样式定义（支持内置样式和自定义样式）
    
    根据配置决定使用哪种表格样式：
    - builtin 模式：使用内置样式（三线表或网格表）
    - custom 模式：使用用户自定义的样式名称
    
    内置样式说明：
    - three_line_table: 三线表（学术论文常用）
    - table_grid: 网格表（办公文档常用）
    
    国际化说明：
    - 中文环境：注入 "三线表"/"网格表" 样式
    - 英文环境：注入 "Three Line Table"/"Table Grid" 样式
    
    参数:
        doc: python-docx Document 对象
    
    返回:
        str: 实际使用的样式名称
    """
    from gongwen_converter.config.config_manager import config_manager
    
    style_resolver = StyleNameResolver()
    styles_element = doc.part.styles._element
    
    # 获取表格样式配置
    style_mode = config_manager.get_table_style_mode()
    
    if style_mode == "custom":
        # 自定义样式模式
        custom_name = config_manager.get_custom_table_style_name()
        if custom_name:
            return _inject_custom_table_style(doc, custom_name, styles_element)
        else:
            # 自定义名称为空，回退到内置三线表
            logger.warning("自定义表格样式名称为空，回退到三线表")
            style_mode = "builtin"
    
    # 内置样式模式
    builtin_key = config_manager.get_builtin_table_style_key()
    
    if builtin_key == "table_grid":
        return _inject_table_grid_style(doc, style_resolver, styles_element)
    else:
        return _inject_three_line_table_style(doc, style_resolver, styles_element)


def _inject_three_line_table_style(doc, style_resolver, styles_element) -> str:
    """注入三线表样式"""
    style_key = "three_line_table"
    style_id = "ThreeLineTable"
    
    style_name = style_resolver.get_injection_name(style_key)
    logger.debug(f"三线表样式名: {style_name}")
    
    # 检查是否需要注入
    if not style_resolver.should_inject(doc, style_key):
        usable_name = style_resolver.get_usable_name(doc, style_key)
        if usable_name:
            logger.debug(f"三线表样式已存在，使用: {usable_name}")
            return usable_name
    
    # 检查 styleId 是否已存在
    try:
        existing = styles_element.find(f".//w:style[@w:styleId='{style_id}']", namespaces=NSMAP)
        if existing is not None:
            logger.debug(f"三线表样式 styleId 已存在，跳过注入")
            return style_name
    except Exception as e:
        logger.warning(f"检查三线表样式时出错: {e}")
    
    # 注入样式
    try:
        template = get_localized_three_line_table_template(style_name)
        style_element = etree.fromstring(template.encode('utf-8'))
        styles_element.append(style_element)
        logger.info(f"成功注入三线表样式: {style_name}")
        return style_name
    except Exception as e:
        logger.error(f"注入三线表样式失败: {e}")
        return "Table Grid"


def _inject_table_grid_style(doc, style_resolver, styles_element) -> str:
    """注入网格表样式"""
    style_key = "table_grid"
    style_id = "TableGrid"
    
    style_name = style_resolver.get_injection_name(style_key)
    logger.debug(f"网格表样式名: {style_name}")
    
    # 检查是否需要注入
    if not style_resolver.should_inject(doc, style_key):
        usable_name = style_resolver.get_usable_name(doc, style_key)
        if usable_name:
            logger.debug(f"网格表样式已存在，使用: {usable_name}")
            return usable_name
    
    # 检查 styleId 是否已存在
    try:
        existing = styles_element.find(f".//w:style[@w:styleId='{style_id}']", namespaces=NSMAP)
        if existing is not None:
            logger.debug(f"网格表样式 styleId 已存在，跳过注入")
            return style_name
    except Exception as e:
        logger.warning(f"检查网格表样式时出错: {e}")
    
    # 注入样式
    try:
        template = get_localized_table_grid_template(style_name)
        style_element = etree.fromstring(template.encode('utf-8'))
        styles_element.append(style_element)
        logger.info(f"成功注入网格表样式: {style_name}")
        return style_name
    except Exception as e:
        logger.error(f"注入网格表样式失败: {e}")
        return "Table Grid"


def _inject_custom_table_style(doc, custom_name: str, styles_element) -> str:
    """
    处理自定义表格样式
    
    逻辑：
    1. 检查模板中是否存在该样式名
    2. 如果存在，直接使用
    3. 如果不存在，创建一个网格表格式的同名样式
    """
    # 检查模板中是否已有该样式
    try:
        for style in styles_element.findall('.//w:style', namespaces=NSMAP):
            style_type = style.get(f'{{{WORD_NAMESPACE}}}type', '')
            if style_type != 'table':
                continue
            
            name_elem = style.find('w:name', namespaces=NSMAP)
            name = name_elem.get(f'{{{WORD_NAMESPACE}}}val', '') if name_elem is not None else ''
            
            if name.lower() == custom_name.lower():
                logger.debug(f"模板中找到自定义表格样式: {name}")
                return name
    except Exception as e:
        logger.warning(f"检查自定义表格样式时出错: {e}")
    
    # 模板中不存在，注入网格表格式的自定义样式
    try:
        template = get_custom_table_grid_template(custom_name)
        style_element = etree.fromstring(template.encode('utf-8'))
        styles_element.append(style_element)
        logger.info(f"成功注入自定义表格样式(网格表格式): {custom_name}")
        return custom_name
    except Exception as e:
        logger.error(f"注入自定义表格样式失败: {e}")
        return "Table Grid"


def inject_table_content_style(doc) -> str:
    """
    向文档注入表格内容段落样式（国际化版本）
    
    表格内容样式特点：
    - 五号字 (10.5pt)
    - 居中对齐
    - 无首行缩进
    - 单倍行距
    
    国际化说明：
    - 中文环境：注入 "表格内容" 样式
    - 英文环境：注入 "Table Content" 样式
    - 检测时：兼容识别所有语言版本的样式名
    
    参数:
        doc: python-docx Document 对象
    
    返回:
        str: 样式名称（国际化），失败时返回 'Normal'
    """
    style_resolver = StyleNameResolver()
    style_key = "table_content"
    style_id = "TableContent"
    
    # 获取国际化的样式名
    style_name = style_resolver.get_injection_name(style_key)
    logger.debug(f"表格内容样式名: {style_name}")
    
    # 检查是否需要注入（检查所有语言版本的样式）
    if not style_resolver.should_inject(doc, style_key):
        # 返回可用的样式名（按优先级：当前语言 > zh_CN > en_US）
        usable_name = style_resolver.get_usable_name(doc, style_key)
        if usable_name:
            logger.debug(f"表格内容样式已存在，使用: {usable_name}")
            return usable_name
    
    # 检查 styleId 是否已存在（避免重复注入）
    try:
        styles_element = doc.part.styles._element
        existing = styles_element.find(
            f".//w:style[@w:styleId='{style_id}']",
            namespaces=NSMAP
        )
        
        if existing is not None:
            logger.debug(f"表格内容样式 styleId 已存在，跳过注入: {style_id}")
            return style_name
    except Exception as e:
        logger.warning(f"检查表格内容样式时出错: {e}")
    
    # 注入样式（使用国际化模板）
    try:
        template = get_localized_table_content_template(style_name)
        style_element = etree.fromstring(template.encode('utf-8'))
        styles_element.append(style_element)
        logger.info(f"成功注入表格内容样式: {style_name}")
        return style_name
    except Exception as e:
        logger.error(f"注入表格内容样式失败: {e}")
        return 'Normal'


def _enable_first_row_formatting(table):
    """
    启用表格的首行条件格式
    
    通过设置 tblLook 元素的 firstRow 属性，让表格样式中定义的 firstRow 条件格式生效。
    
    生成 XML:
    <w:tblPr>
        <w:tblLook w:firstRow="1" w:lastRow="0" w:firstColumn="0" 
                   w:lastColumn="0" w:noHBand="0" w:noVBand="1"/>
    </w:tblPr>
    
    参数:
        table: python-docx Table 对象
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # 添加 tblLook 元素
    tblLook = OxmlElement('w:tblLook')
    tblLook.set(qn('w:firstRow'), '1')      # 启用首行条件格式
    tblLook.set(qn('w:lastRow'), '0')       # 禁用末行条件格式
    tblLook.set(qn('w:firstColumn'), '0')   # 禁用首列条件格式
    tblLook.set(qn('w:lastColumn'), '0')    # 禁用末列条件格式
    tblLook.set(qn('w:noHBand'), '0')       # 启用水平带状
    tblLook.set(qn('w:noVBand'), '1')       # 禁用垂直带状
    tblPr.append(tblLook)
    
    logger.debug("已启用表格首行条件格式 (tblLook.firstRow=1)")


def _apply_header_row_bottom_border(header_row):
    """
    为表头行的每个单元格显式设置底部边框（WPS 兼容性修复）
    
    背景说明：
    Word 和 WPS 对表格样式的 tblStylePr（条件格式）支持存在差异：
    - Word: 完整支持 firstRow 中定义的边框，表头底线正常显示
    - WPS: 对 tblStylePr 支持不完整，可能无法渲染 firstRow 中的边框
    
    解决方案：
    在单元格级别（tcPr/tcBorders）显式设置底部边框，作为"双保险"：
    - Word 中：样式边框和单元格边框都生效，效果一致
    - WPS 中：单元格边框生效，补偿样式边框的缺失
    
    生成 XML:
    <w:tc>
        <w:tcPr>
            <w:tcBorders>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            </w:tcBorders>
        </w:tcPr>
        ...
    </w:tc>
    
    参数:
        header_row: 表格的表头行对象（table.rows[0]）
    """
    WORD_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    
    for cell in header_row.cells:
        tc = cell._tc
        
        # 获取或创建单元格属性 <w:tcPr>
        tcPr = tc.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.insert(0, tcPr)
        
        # 检查是否已有 tcBorders 元素
        existing_borders = tcPr.find(f'{WORD_NS}tcBorders')
        if existing_borders is not None:
            # 已有边框设置，检查是否需要更新底部边框
            existing_bottom = existing_borders.find(f'{WORD_NS}bottom')
            if existing_bottom is not None:
                # 已有底部边框，跳过（避免覆盖用户设置）
                continue
            # 没有底部边框，添加
            tcBorders = existing_borders
        else:
            # 创建边框容器 <w:tcBorders>
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        
        # 创建底部边框
        # sz="4" 对应 0.5pt，与三线表样式一致
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')       # 0.5pt 细线
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000')  # 黑色
        tcBorders.append(bottom)
    
    logger.debug(f"已为表头行 {len(header_row.cells)} 个单元格设置底部边框（WPS兼容性修复）")
