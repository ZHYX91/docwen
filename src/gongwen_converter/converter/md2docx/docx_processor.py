"""
DOCX处理器模块
"""

import logging
import tempfile
from docx import Document
from .placeholder_handler import (
    mark_special_placeholders,
    is_special_marked,
    remove_special_mark,
    process_paragraph_placeholders,
    process_table_cell_placeholders,
    process_attachment_description_placeholder,
    try_remove_element
)
from . import xml_processor
from gongwen_converter.utils import docx_utils

# 配置日志
logger = logging.getLogger(__name__)

def load_document(template_path):
    """加载Word文档模板"""
    try:
        doc = Document(template_path)
        logger.info(f"成功加载模板文件: {template_path}")
        return doc
    except Exception as e:
        logger.error(f"加载模板文件失败: {template_path}, 错误: {str(e)}")
        raise

def replace_placeholders(doc, yaml_data, body_data):
    """
    主处理函数：替换文档中的所有占位符
    移除文本校对功能
    
    参数:
        doc: Document对象
        yaml_data: 字典，包含要替换的键值对
        body_data: 列表，包含处理后的正文段落数据
        
    返回:
        处理后的临时文件路径
    """
    logger.info("开始处理文档占位符...")
    
    # 第一步：标记特殊占位符段落
    mark_special_placeholders(doc)
    
    # 第二步：处理常规段落占位符
    paragraphs_to_remove = []
    for paragraph in doc.paragraphs:
        # 跳过已标记的特殊占位符
        if not is_special_marked(paragraph):
            logger.debug(f"处理段落: {paragraph.text[:50]}...")
            process_paragraph_placeholders(paragraph, yaml_data, paragraphs_to_remove)
    
    # 第三步：处理表格中的占位符
    rows_to_remove = []
    tables_to_remove = []  # 存储需要删除的表格
    for table in doc.tables:
        row_removal_flags = []  # 标记需要删除的行
        for row_idx, row in enumerate(table.rows):
            # 【新增】收集整行的所有占位符（用于行级删除检查）
            row_placeholder_keys = set()
            for cell in row.cells:
                for para in cell.paragraphs:
                    # 提取占位符键名并添加到整行集合
                    from .placeholder_handler import extract_placeholder_keys
                    cell_keys = extract_placeholder_keys(para.text)
                    row_placeholder_keys.update(cell_keys)
            
            logger.debug(f"表格第{row_idx+1}行的占位符: {row_placeholder_keys}")
            
            # 处理单元格，传入整行占位符集合
            row_should_remove = False  # 标记当前行是否需要删除
            for cell in row.cells:
                for para in cell.paragraphs:
                    # 处理单元格占位符，传入整行占位符集合，返回是否需要删除行
                    if process_table_cell_placeholders(para, yaml_data, row, rows_to_remove, row_placeholder_keys):
                        row_should_remove = True
                        break
                if row_should_remove:
                    break
            row_removal_flags.append(row_should_remove)
        
        # 检查表格是否所有行都需要删除
        # 如果所有行都需要删除，则删除整个表格（避免留下空表格框架）
        if row_removal_flags and all(row_removal_flags):
            logger.info(f"表格中所有行都需要删除，标记删除整个表格")
            tables_to_remove.append(table)
    
    # 第四步：处理公文正文占位符
    process_main_content(doc, body_data, yaml_data)

    # 第五步：处理附件说明占位符
    process_attachment_description_placeholder(doc, yaml_data)

    # 第六步：处理图片占位符
    from .placeholder_handler import process_image_placeholders
    process_image_placeholders(doc)

    # 第七步：删除需要移除的元素
    # 先删除表格（如果表格需要删除）
    for table in tables_to_remove:
        try:
            table._element.getparent().remove(table._element)
            logger.info(f"已删除整个表格（只剩一行且需要删除）")
        except Exception as e:
            logger.warning(f"删除表格失败: {str(e)}")
    
    # 删除需要移除的段落和表格行
    remove_marked_elements(paragraphs_to_remove, rows_to_remove)

    # 第八步：深度处理所有XML内容
    temp_path = save_and_process_temp_file(doc, yaml_data)

    return temp_path

def process_main_content(doc, body_data, yaml_data):
    """处理公文正文占位符"""
    for paragraph in doc.paragraphs:
        # 使用硬编码的特殊占位符
        from .placeholder_handler import SPECIAL_PLACEHOLDERS
        main_content_ph = SPECIAL_PLACEHOLDERS.get('公文正文', '{{公文正文}}')
        
        if main_content_ph in paragraph.text:
            # 移除特殊标记
            remove_special_mark(paragraph)
            
            base_style = paragraph.style
            
            if not paragraph.runs:
                paragraph.add_run(main_content_ph)

            # 提取格式信息
            fonts = docx_utils.extract_format_from_paragraph(paragraph)
            logger.info(f"正文格式提取结果 - 字体: {fonts['eastAsia']}, 字号: {fonts['sz']}")

            # 准备插入新内容
            parent = paragraph._element.getparent()
            index = parent.index(paragraph._element)
            
            # 插入处理后的内容
            for i, item in enumerate(body_data):
                item_type = item.get('type')
                level = item.get('level', 0)
                
                # 处理表格
                if item_type == 'table':
                    table_data = item.get('table_data')
                    word_table = create_word_table(doc, table_data)
                    # 插入表格到正确位置
                    parent.insert(index + i, word_table._element)
                    logger.info(f"插入Word表格: {len(table_data['headers'])}列 x {len(table_data['rows'])}行")
                    continue
                
                # 创建新段落
                if 1 <= level <= 5:
                    # 使用 add_heading 创建标题，自动激活样式
                    new_p = doc.add_heading('', level=level)
                    # 添加标题文本
                    title_run = new_p.add_run(item['text'])
                else:
                    new_p = doc.add_paragraph(style=base_style)
                    title_run = new_p.add_run(item['text'])

                
                # 添加正文内容（如果有）
                if item_type == 'heading_with_content':
                    content_run = new_p.add_run(item['content'])
                    # 直接调用docx_utils中的函数
                    docx_utils.apply_body_formatting(content_run, fonts)
                
                # 插入到正确位置
                new_p._p.getparent().remove(new_p._p)
                parent.insert(index + i, new_p._p)
            
            # 移除原始占位符段落
            parent.remove(paragraph._element)
            break

def remove_marked_elements(paragraphs_to_remove, rows_to_remove):
    """移除标记为需要删除的段落和表格行"""
    # 移除段落
    for para in paragraphs_to_remove:
        try_remove_element(para._element)
    
    # 移除表格行
    for row in rows_to_remove:
        try:
            parent = row._element.getparent()
            if parent is not None:
                parent.remove(row._element)
        except Exception as e:
            logger.warning(f"移除表格行失败: {str(e)}")

def create_word_table(doc, table_data):
    """
    创建Word表格
    
    参数:
        doc: Document对象
        table_data: 表格数据字典，包含：
            - headers: 表头列表
            - rows: 数据行列表
    
    返回:
        Table对象
    """
    headers = table_data['headers']
    rows = table_data['rows']
    
    logger.debug(f"创建Word表格: {len(headers)}列 x {len(rows)+1}行（含表头）")
    
    # 创建表格（表头 + 数据行）
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    
    # 设置表格样式为网格样式（优雅降级）
    try:
        table.style = 'Table Grid'
        logger.info("成功应用Table Grid样式")
    except KeyError:
        logger.warning("模板中不存在'Table Grid'样式，使用默认样式")
        # 不设置样式，使用Word默认表格样式
    
    # 设置表头
    header_row = table.rows[0]
    for col_idx, header in enumerate(headers):
        cell = header_row.cells[col_idx]
        cell.text = str(header)
        # 表头加粗
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
        logger.debug(f"设置表头单元格 [0,{col_idx}]: {header}")
    
    # 填充数据
    for row_idx, row_data in enumerate(rows):
        data_row = table.rows[row_idx + 1]  # +1 因为第0行是表头
        for col_idx, cell_value in enumerate(row_data):
            if col_idx < len(headers):  # 确保不超出列数
                cell = data_row.cells[col_idx]
                cell.text = str(cell_value) if cell_value else ''
                logger.debug(f"填充数据单元格 [{row_idx+1},{col_idx}]: {cell_value}")
    
    logger.info(f"Word表格创建完成: {len(headers)}列 x {len(rows)}行数据")
    return table


def save_and_process_temp_file(doc, yaml_data):
    """保存临时文件并进行深度处理"""
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_docx.name)
    temp_docx.close()
    
    logger.debug("开始深度处理DOCX文件中的占位符")

    xml_processor.process_docx_file(temp_docx.name, yaml_data)
    
    return temp_docx.name
