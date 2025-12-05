
"""
公文格式DOCX转MD转换器

实现公文专用的转换逻辑：
- 三轮公文元素识别
- YAML元数据提取（14个字段）
- 附件内容单独输出
- 支持图片提取和OCR
"""

import os
import re
import logging
import datetime
import tempfile
import threading
from typing import Callable, Optional
from docx import Document

# 导入处理模块
from .content_injector import process_document_with_special_content
from .element_scorer import create_element_scorer
from .image_processor import extract_images_from_docx, process_image_with_ocr

# 导入工具函数
from gongwen_converter.utils.docx_utils import extract_format_from_paragraph_style, get_effective_run_format
from gongwen_converter.utils.heading_utils import split_content_by_delimiters, add_markdown_heading, detect_heading_level
from gongwen_converter.utils.validation_utils import contains_chinese
from gongwen_converter.utils.date_utils import convert_date_format
from gongwen_converter.utils.text_utils import (
    remove_colon,
    extract_after_colon,
    remove_brackets,
    process_copy_to,
    process_attachment_item,
    is_pure_number
)

logger = logging.getLogger(__name__)


def convert_docx_to_md_gongwen(
    docx_path: str,
    extract_image: bool = True,
    extract_ocr: bool = False,
    progress_callback: Optional[Callable[[str], None]] = None,
    cancel_event: Optional[threading.Event] = None,
    output_folder: Optional[str] = None,
    original_file_path: Optional[str] = None
):
    """
    公文优化模式的DOCX转MD转换
    
    采用三轮识别策略提取公文元素，生成详细的YAML元数据。
    
    参数:
        docx_path: DOCX文件路径（可能是临时副本）
        extract_image: 是否保留图片（由GUI传入，默认True）
        extract_ocr: 是否进行OCR识别（由GUI传入，默认False）
        progress_callback: 进度回调函数 (可选)
        cancel_event: 取消事件 (可选)
        output_folder: 输出文件夹路径，用于保存图片 (可选)
        original_file_path: 原始文件路径（用于图片命名，可选）
    
    返回:
        dict: 转换结果字典
    """
    # 直接使用GUI传入的参数
    keep_images = extract_image
    enable_ocr = extract_ocr
    
    logger.info(f"开始转换DOCX文件: {docx_path}")
    logger.info(f"导出选项 - 提取图片: {keep_images}, OCR识别: {enable_ocr}")
    
    try:
        # 加载DOCX文档
        if progress_callback:
            progress_callback("正在加载文档...")
        doc = Document(docx_path)
        logger.info(f"成功加载DOCX文档, 包含 {len(doc.paragraphs)} 个段落")
        
        # 第一步：处理文本框和表格内容，插入到doc.paragraphs中
        if progress_callback:
            progress_callback("正在处理文本框和表格...")
        logger.info("开始处理文本框和表格内容...")
        # 使用临时目录处理文本框和表格
        import tempfile
        temp_dir = tempfile.gettempdir()
        modified_doc, extracted_path = process_document_with_special_content(doc, docx_path, temp_dir)
        
        if extracted_path:
            logger.info(f"提取文档已生成: {extracted_path}")
        else:
            logger.info("提取文档未生成（已禁用输出或生成失败）")
        
        # 使用处理后的文档继续后续流程
        doc = modified_doc
        logger.info(f"处理后的文档包含 {len(doc.paragraphs)} 个段落")
        
        # 创建评分器
        scorer = create_element_scorer()
        logger.info(f"初始化公文元素评分器: {scorer}")
        
        # 收集文档上下文信息
        context = {
            'total_paras': len(doc.paragraphs),
            'para_index': 0,  # 当前段落索引
            'para_non_empty': {},  # 记录非空段落
            'doc': doc  # 文档对象，用于访问其他段落
        }
        
        # 初始化YAML信息
        yaml_info = {
            "aliases": [],  # 标题列表
            "标题": "",
            "份号": "",
            "密级和保密期限": "",
            "紧急程度": "",
            "发文字号": "",
            "发文机关标志": "",
            "签发人": [],  # 签发人列表
            "发文机关署名": "",
            "成文日期": "",
            "印发日期": "",
            "主送机关": "",
            "附注": "",
            "印发机关": "",
            "抄送机关": [],  # 抄送机关列表
            "附件说明": [],  # 附件说明列表
            "公开方式": ""
        }
        
        # 需要跳过的段落索引    
        skip_para_indices = []  
        
        # 用于存储标题组信息
        title_groups = []  # 存储所有标题组 [(起始索引, 结束索引, 合并文本)]
        
        # 存储附件内容段落的索引
        attachment_content_indices = []
        
        # 初始化段落状态管理
        identified_paragraphs = set()  # 记录已确定元素的段落索引
        paragraph_elements = {}  # 记录每个段落的最终元素类型
        
        # 提前提取图片信息（用于第3轮识别空图片段落）
        images_info = []
        if (keep_images or enable_ocr) and output_folder:
            if progress_callback:
                progress_callback("正在提取图片...")
            
            try:
                from .image_processor import extract_images_from_docx
                # 优先使用原始文件路径进行图片命名，如未提供则使用docx_path
                path_for_naming = original_file_path or docx_path
                images_info = extract_images_from_docx(doc, output_folder, path_for_naming)
                logger.info(f"图片提取完成，共 {len(images_info)} 张")
            except Exception as e:
                logger.error(f"图片提取失败: {e}", exc_info=True)
                images_info = []
        elif (keep_images or enable_ocr) and not output_folder:
            logger.warning("配置启用了图片保留或OCR，但未提供output_folder参数，跳过图片提取")
        
        # 第1轮：识别大部分唯一性元素
        if progress_callback:
            progress_callback("第1轮：识别唯一性元素...")
        logger.info("开始第1轮识别：大部分唯一性元素")
        for idx, para in enumerate(doc.paragraphs):
            if cancel_event and cancel_event.is_set():
                logger.info("操作被用户取消")
                return False
            # 获取段落文本
            para_text = para.text.strip()
            context['para_index'] = idx
            context['para_non_empty'][idx] = bool(para_text)
            
            # 跳过空段落
            if not para_text:
                logger.debug(f"跳过空段落: {idx+1}")
                continue
            
            # 获取段落基础格式
            para_fonts = extract_format_from_paragraph_style(para)
            
            # 确保 para_fonts 不为 None
            if para_fonts is None:
                para_fonts = {}
                logger.warning(f"段落{idx+1}基础格式为空，使用空字典")
            
            # 获取有效Run格式
            effective_fonts = get_effective_run_format(para, para_fonts)
            
            # 确保 effective_fonts 不为 None
            if effective_fonts is None:
                effective_fonts = {}
                logger.warning(f"段落{idx+1}有效Run格式为空，使用空字典")
            
            # 创建用于评分的run字典
            run_dict = {
                'text': para_text,
                'fonts': effective_fonts,
                'is_chinese': contains_chinese(para_text)
            }
            
            # 第1轮：识别大部分唯一性元素
            best_element, best_score = scorer.score_round1_unique_elements(run_dict, context)
            
            if best_element and best_score >= 80:
                element_name = scorer.ELEMENT_TYPES.get(best_element, "")
                logger.info(f"第1轮段落 {idx+1} 识别结果: {element_name} (得分: {best_score})")
                
                # 更新上下文和记录
                scorer.update_context(best_element, idx)
                identified_paragraphs.add(idx)
                paragraph_elements[idx] = best_element
                
        
        # 第2轮：识别发文机关标志和印发机关（只处理未确定的段落）
        if progress_callback:
            progress_callback("第2轮：识别机关标志...")
        logger.info("开始第2轮识别：发文机关标志和印发机关")
        for idx, para in enumerate(doc.paragraphs):
            if cancel_event and cancel_event.is_set():
                logger.info("操作被用户取消")
                return False
            # 跳过已确定的段落和空段落
            if idx in identified_paragraphs:
                continue
                
            para_text = para.text.strip()
            if not para_text:
                continue
                
            context['para_index'] = idx
            
            # 获取段落格式信息
            para_fonts = extract_format_from_paragraph_style(para)
            if para_fonts is None:
                para_fonts = {}
            
            effective_fonts = get_effective_run_format(para, para_fonts)
            if effective_fonts is None:
                effective_fonts = {}
            
            run_dict = {
                'text': para_text,
                'fonts': effective_fonts,
                'is_chinese': contains_chinese(para_text)
            }
            
            # 第2轮：识别发文机关标志和印发机关
            best_element, best_score = scorer.score_round2_unique_elements(run_dict, context)
            
            if best_element and best_score >= 80:
                element_name = scorer.ELEMENT_TYPES.get(best_element, "")
                logger.info(f"第2轮段落 {idx+1} 识别结果: {element_name} (得分: {best_score})")
                
                # 更新上下文和记录
                scorer.update_context(best_element, idx)
                identified_paragraphs.add(idx)
                paragraph_elements[idx] = best_element
        
        # 第3轮：识别非唯一性元素（只处理前两轮未确定的段落）
        if progress_callback:
            progress_callback("第3轮：识别非唯一性元素...")
        logger.info("开始第3轮识别：非唯一性元素")
        for idx, para in enumerate(doc.paragraphs):
            if cancel_event and cancel_event.is_set():
                logger.info("操作被用户取消")
                return False
            # 跳过已确定的段落和空段落
            if idx in identified_paragraphs:
                continue
                
            para_text = para.text.strip()
            if not para_text:
                # 空段落，检查是否有图片
                has_image = any(img['para_index'] == idx for img in images_info)
                if has_image:
                    # 判断空图片段落应该归属哪个区域
                    if idx > scorer.last_known_element_index:
                        # 附件区域的空图片段落
                        paragraph_elements[idx] = 'attachment_content'
                        attachment_content_indices.append(idx)
                        skip_para_indices.append(idx)
                        logger.info(f"检测到附件区域的空图片段落 {idx+1}")
                    elif paragraph_elements.get(idx) is None:
                        # 正文区域的空图片段落，标记为一般内容（不添加到skip_para_indices，以便后续插入图片）
                        paragraph_elements[idx] = 'general_content'
                        logger.info(f"检测到正文区域的空图片段落 {idx+1}")
                continue
                
            context['para_index'] = idx
            
            # 获取段落格式信息
            para_fonts = extract_format_from_paragraph_style(para)
            if para_fonts is None:
                para_fonts = {}
            
            effective_fonts = get_effective_run_format(para, para_fonts)
            if effective_fonts is None:
                effective_fonts = {}
            
            run_dict = {
                'text': para_text,
                'fonts': effective_fonts,
                'is_chinese': contains_chinese(para_text)
            }
            
            # 第3轮：识别非唯一性元素
            best_element, best_score = scorer.score_round3_non_unique_elements(run_dict, context)
            
            if best_element and best_score > 0:
                element_name = scorer.ELEMENT_TYPES.get(best_element, "")
                logger.info(f"第3轮段落 {idx+1} 识别结果: {element_name} (得分: {best_score})")
                
                # 更新上下文和记录
                scorer.update_context(best_element, idx)
                paragraph_elements[idx] = best_element
                
                # 特别处理：识别附件内容
                if best_element == 'attachment_content':
                    logger.info(f"检测到附件内容段落 {idx+1}: '{para_text}'")
                    attachment_content_indices.append(idx)
                    skip_para_indices.append(idx)  # 附件内容不写入主报告
        
        # 提取关键信息到YAML（基于三轮识别的结果）
        for idx, best_element in paragraph_elements.items():
            para = doc.paragraphs[idx]
            para_text = para.text.strip()
            
            # 处理组合ID（份号+发文字号）
            if best_element == 'combined_id':
                # 拆分份号和发文字号
                parts = re.split(r'[\s\x20\t]+', para_text, 1)
                if len(parts) >= 2:
                    yaml_info['份号'] = parts[0]  # 开头是份号
                    yaml_info['发文字号'] = parts[1]  # 剩余部分是发文字号
                    logger.info(f"提取份号和发文字号: 份号={parts[0]}, 发文字号={parts[1]}")
                else:
                    yaml_info['发文字号'] = para_text  # 无法拆分则整个作为发文字号
                skip_para_indices.append(idx)

            # 标题处理（不在标题组中处理，因为标题组会单独处理）
            elif best_element == 'title' and not yaml_info['aliases']:
                # 这里不直接设置标题，等标题组处理完成后统一设置
                skip_para_indices.append(idx)
            
            # 份号处理
            elif best_element == 'copy_id' and not yaml_info['份号']:
                yaml_info['份号'] = para_text
                skip_para_indices.append(idx)
            
            # 密级和保密期限处理
            elif best_element == 'security' and not yaml_info['密级和保密期限']:
                yaml_info['密级和保密期限'] = para_text
                skip_para_indices.append(idx)
            
            # 紧急程度处理
            elif best_element == 'urgency' and not yaml_info['紧急程度']:
                yaml_info['紧急程度'] = para_text
                skip_para_indices.append(idx)
            
            # 发文字号处理
            elif best_element == 'doc_number' and not yaml_info['发文字号']:
                yaml_info['发文字号'] = para_text
                skip_para_indices.append(idx)
            
            # 发文机关标志处理
            elif best_element == 'issuing_authority_mark' and not yaml_info['发文机关标志']:
                yaml_info['发文机关标志'] = para_text
                skip_para_indices.append(idx)
            
            # 组合格式：发文字号+签发人
            elif best_element == 'combined_doc_number_signer':
                doc_num, signers = _extract_doc_number_and_signers(para_text)
                if doc_num and not yaml_info['发文字号']:
                    yaml_info['发文字号'] = doc_num
                if signers:
                    yaml_info['签发人'].extend(signers)
                skip_para_indices.append(idx)
                logger.info(f"提取发文字号+签发人: 发文字号={doc_num}, 签发人={signers}")
            
            # 独立签发人
            elif best_element == 'signer':
                signer = extract_after_colon(para_text)
                if signer:
                    yaml_info['签发人'].append(signer)
                    logger.info(f"提取签发人: {signer}")
                skip_para_indices.append(idx)
            
            # 后续签发人（纯人名或组合格式）
            elif best_element == 'signer_following':
                signers = _extract_signers_from_text(para_text)
                if signers:
                    yaml_info['签发人'].extend(signers)
                    logger.info(f"提取后续签发人: {signers}")
                skip_para_indices.append(idx)
            
            # 组合后续：发文字号+人名
            elif best_element == 'combined_doc_number_signer_following':
                doc_num, name = _extract_doc_number_and_name(para_text)
                if name:
                    yaml_info['签发人'].append(name)
                    logger.info(f"提取发文字号+人名: 发文字号={doc_num}, 签发人={name}")
                skip_para_indices.append(idx)
            
            # 发文机关署名处理
            elif best_element == 'issuing_authority_signature' and not yaml_info['发文机关署名']:
                yaml_info['发文机关署名'] = para_text
                skip_para_indices.append(idx)
            
            # 成文日期处理
            elif best_element == 'issue_date' and not yaml_info['成文日期']:
                yaml_info['成文日期'] = convert_date_format(para_text)
                skip_para_indices.append(idx)
            
            # 主送机关处理
            elif best_element == 'recipient' and not yaml_info['主送机关']:
                yaml_info['主送机关'] = remove_colon(para_text)
                skip_para_indices.append(idx)
            
            # 附注处理
            elif best_element == 'notes' and not yaml_info['附注']:
                yaml_info['附注'] = remove_brackets(para_text)
                skip_para_indices.append(idx)
            
            # 公开方式处理
            elif best_element == 'disclosure' and not yaml_info['公开方式']:
                yaml_info['公开方式'] = extract_after_colon(para_text)
                skip_para_indices.append(idx)
            
            # 抄送机关处理
            elif best_element == 'copy_to' and not yaml_info['抄送机关']:
                # 处理抄送机关文本
                copy_to_list = process_copy_to(para_text)
                yaml_info['抄送机关'] = copy_to_list
                skip_para_indices.append(idx)
            
            # 附件说明处理
            elif best_element == 'attachment' or best_element == 'attachment_following':
                # 处理附件项
                item = process_attachment_item(para_text)
                yaml_info['附件说明'].append(item)
                skip_para_indices.append(idx)
            
            # 印发机关处理
            elif best_element == 'printing_authority' and not yaml_info['印发机关']:
                yaml_info['印发机关'] = para_text
                skip_para_indices.append(idx)
            
            # 印发日期处理
            elif best_element == 'printing_date' and not yaml_info['印发日期']:
                # 去除"印发"字样并转换日期格式
                date_text = para_text.replace('印发', '').strip()
                yaml_info['印发日期'] = convert_date_format(date_text)
                skip_para_indices.append(idx)
        
        # 查找所有标题和标题后续部分
        current_title_indices = []
        for i in range(len(doc.paragraphs)):
            element = paragraph_elements.get(i)
            if element == 'title':
                if current_title_indices: # 如果上一组未结束，先存起来
                    title_groups.append(current_title_indices)
                current_title_indices = [i] # 开启新的一组
            elif element == 'title_following':
                if current_title_indices: # 只有当存在上一段标题时才添加
                    current_title_indices.append(i)
            else:
                if current_title_indices: # 遇到非标题，结束当前组
                    title_groups.append(current_title_indices)
                    current_title_indices = []
        if current_title_indices: # 处理文档结尾的最后一组标题
            title_groups.append(current_title_indices)

        # 合并标题文本
        if title_groups:
            first_title_indices = title_groups[0]
            full_title = "".join(
                doc.paragraphs[i].text.replace('\n', '').replace('\r', '').strip() 
                for i in first_title_indices
            )
            
            start_idx, end_idx = first_title_indices[0], first_title_indices[-1]
            logger.info(f"提取并合并标题: 段落 {start_idx+1}-{end_idx+1}, 文本: '{full_title}'")
            
            yaml_info['标题'] = full_title
            yaml_info['aliases'].append(full_title)
            
            # 标记所有标题相关段落以便跳过
            for i in first_title_indices:
                if i not in skip_para_indices:
                    skip_para_indices.append(i)
        else:
            logger.warning("未检测到任何标题组")
        
        # 生成主要部分Markdown内容
        if progress_callback:
            progress_callback("正在生成Markdown内容...")
        
        main_content = _generate_markdown_content(yaml_info, doc, skip_para_indices, images_info, keep_images, enable_ocr, output_folder)
        logger.info("主要Markdown内容生成完成")
        
        # 生成附件部分Markdown内容
        attachment_content = None
        if attachment_content_indices:
            if cancel_event and cancel_event.is_set():
                logger.info("操作被用户取消")
                return {
                    'success': False,
                    'main_content': None,
                    'attachment_content': None,
                    'metadata': yaml_info,
                    'error': '操作已取消'
                }
            attachment_content = _generate_attachment_content(doc, attachment_content_indices, docx_path, images_info, keep_images, enable_ocr, output_folder)
            logger.info("附件Markdown内容生成完成")
        else:
            logger.info("未检测到附件内容")
        
        return {
            'success': True,
            'main_content': main_content,
            'attachment_content': attachment_content,
            'metadata': yaml_info,
            'error': None
        }
        
    except Exception as e:
        error_msg = f"转换DOCX文件失败: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return {
            'success': False,
            'main_content': None,
            'attachment_content': None,
            'metadata': {},
            'error': error_msg
        }

def _generate_markdown_content(yaml_info: dict, doc, skip_para_indices: list, images_info: list = None, 
                               keep_images: bool = True, enable_ocr: bool = False, output_folder: str = None) -> str:
    """
    生成主要Markdown内容
    
    参数:
        yaml_info: YAML元数据字典
        doc: Document对象
        skip_para_indices: 需要跳过的段落索引列表
        images_info: 图片信息列表 (可选)
        keep_images: 是否保留图片 (默认True)
        enable_ocr: 是否启用OCR识别 (默认False)
        output_folder: 输出文件夹路径 (可选)
    
    返回:
        str: 完整的Markdown内容（包含YAML头部）
    """
    if images_info is None:
        images_info = []
    
    logger.info(f"生成Markdown内容 - 图片数量: {len(images_info)}, 保留图片: {keep_images}, OCR: {enable_ocr}")
    lines = []
    
    # 写入YAML头部
    lines.append("---")
    for key, value in yaml_info.items():
        # 特殊处理列表类型
        if isinstance(value, list):
            if value:
                lines.append(f"{key}:")
                for item in value:
                    # 检查列表中的每个元素是否为纯数字，如果是则用双引号包裹
                    if is_pure_number(item):
                        lines.append(f'  - "{item}"')
                    else:
                        lines.append(f"  - {item}")
            else:
                lines.append(f"{key}: []")
        else:
            # 检查单个值是否为纯数字，如果是则用双引号包裹
            if is_pure_number(value):
                lines.append(f'{key}: "{value}"')
            else:
                lines.append(f"{key}: {value}")
    lines.append("---")
    lines.append("")
    
    # 创建内容列表，用于存储正文区域的所有段落（包括小标题和正文文本）
    content_lines = []
    
    # 处理所有段落（跳过已提取到YAML的段落和附件内容）
    for idx, para in enumerate(doc.paragraphs):
        # 检查是否需要跳过
        if idx in skip_para_indices:
            logger.debug(f"跳过已提取到YAML的段落: {idx+1}")
            continue
        
        # 获取段落文本
        para_text = para.text.strip()
        
        # 处理空段落 - 先检查是否有图片
        if not para_text:
            # 检查当前空段落是否有图片，如果有则根据选项处理
            para_images = [img for img in images_info if img['para_index'] == idx]
            if para_images:
                for img in para_images:
                    image_link = process_image_with_ocr(img, keep_images, enable_ocr, output_folder)
                    content_lines.append(image_link)
                    logger.debug(f"在空段落 {idx + 1} 插入图片: {image_link}")
            else:
                logger.debug(f"跳过空段落: {idx+1}")
            continue
        
        # 检测标题级别和清理序号
        cleaned_text, heading_level = detect_heading_level(para_text)
        
        # 根据标题级别处理内容
        if heading_level > 0:
            # 分割内容
            content1, content2 = split_content_by_delimiters(cleaned_text)
            
            # 添加Markdown标题符号
            if content1:
                content1 = add_markdown_heading(content1, heading_level)
            
            # 组合段落内容
            if content1 and content2:
                # 内容1和内容2之间只有一个换行符
                paragraph_content = f"{content1}\n{content2}"
            elif content1:
                paragraph_content = content1
            else:
                paragraph_content = cleaned_text
        else:
            # 非小标题段落直接使用原始文本
            paragraph_content = para_text
        
        # 添加到内容列表
        content_lines.append(paragraph_content)
        
        # 检查当前段落是否有图片，如果有则根据选项处理
        para_images = [img for img in images_info if img['para_index'] == idx]
        for img in para_images:
            # 在段落后根据选项插入图片或图片md文件
            image_link = process_image_with_ocr(img, keep_images, enable_ocr, output_folder)
            content_lines.append(image_link)
            logger.debug(f"在段落 {idx + 1} 后插入图片: {image_link}")
    
    # 添加正文区域内容（段落之间用两个换行符分隔）
    if content_lines:
        lines.append("\n\n".join(content_lines))
    
    return "\n".join(lines)

def _generate_attachment_content(doc, content_indices: list, original_docx_path: str, images_info: list = None,
                                 keep_images: bool = True, enable_ocr: bool = False, output_folder: str = None) -> str:
    """
    生成附件Markdown内容
    
    参数:
        doc: Document对象
        content_indices: 附件内容段落的索引列表
        original_docx_path: 原始DOCX文件路径
        images_info: 图片信息列表 (可选)
        keep_images: 是否保留图片 (默认True)
        enable_ocr: 是否启用OCR识别 (默认False)
        output_folder: 输出文件夹路径 (可选)
    
    返回:
        str: 附件Markdown内容（包含YAML头部）
    """
    if images_info is None:
        images_info = []
    lines = []
    
    # 添加YAML头部
    lines.append("---")
    lines.append(f"来源文件: {os.path.basename(original_docx_path)}")
    lines.append(f"提取时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("类型: 附件内容")
    lines.append("---")
    lines.append("")
    
    # 提取附件内容
    attachment_lines = []
    for idx in content_indices:
        if idx < len(doc.paragraphs):
            para = doc.paragraphs[idx]
            para_text = para.text.strip()
            
            # 处理空段落 - 先检查是否有图片
            if not para_text:
                # 检查当前空段落是否有图片，如果有则根据选项处理
                para_images = [img for img in images_info if img['para_index'] == idx]
                if para_images:
                    for img in para_images:
                        image_link = process_image_with_ocr(img, keep_images, enable_ocr, output_folder)
                        attachment_lines.append(image_link)
                        logger.debug(f"在附件空段落 {idx + 1} 插入图片: {image_link}")
                continue
            
            # 检测标题级别和清理序号
            cleaned_text, heading_level = detect_heading_level(para_text)
            
            # 根据标题级别处理内容
            if heading_level > 0:
                # 分割内容
                content1, content2 = split_content_by_delimiters(cleaned_text)
                
                # 添加Markdown标题符号
                if content1:
                    content1 = add_markdown_heading(content1, heading_level)
                
                # 组合段落内容
                if content1 and content2:
                    paragraph_content = f"{content1}\n{content2}"
                elif content1:
                    paragraph_content = content1
                else:
                    paragraph_content = cleaned_text
            else:
                # 非小标题段落直接使用原始文本
                paragraph_content = para_text
            
            attachment_lines.append(paragraph_content)
            
            # 检查当前段落是否有图片，如果有则根据选项处理
            para_images = [img for img in images_info if img['para_index'] == idx]
            for img in para_images:
                # 在段落后根据选项插入图片或图片md文件
                image_link = process_image_with_ocr(img, keep_images, enable_ocr, output_folder)
                attachment_lines.append(image_link)
                logger.debug(f"在附件段落 {idx + 1} 后插入图片: {image_link}")
    
    # 添加正文内容
    if attachment_lines:
        lines.append("\n\n".join(attachment_lines))
    
    return "\n".join(lines)

def _extract_doc_number_and_signers(text: str) -> tuple:
    """
    从"发文字号+签发人"组合格式中提取信息
    格式：XX〔2024〕1号   签发人：张三
    
    返回:
        tuple: (发文字号, [签发人列表])
    """
    text = text.strip()
    
    # 分离"签发人："及其之前的内容
    match = re.match(r'^(.+?)[\s\u3000\t]+签发人[：:](.+)$', text)
    if match:
        doc_number = match.group(1).strip()
        signer_text = match.group(2).strip()
        
        # 提取签发人（可能多个，用顿号或空格分隔）
        signers = re.split(r'[、\s\u3000]+', signer_text)
        signers = [s.strip() for s in signers if s.strip()]
        
        logger.debug(f"提取发文字号和签发人: 发文字号='{doc_number}', 签发人={signers}")
        return doc_number, signers
    
    logger.debug(f"无法提取发文字号和签发人: '{text}'")
    return "", []

def _extract_signers_from_text(text: str) -> list:
    """
    从纯人名文本中提取多个签发人
    格式：李四、王五  或  李四 王五
    
    返回:
        list: 签发人列表
    """
    text = text.strip()
    
    # 按顿号或空格分割
    signers = re.split(r'[、\s\u3000]+', text)
    signers = [s.strip() for s in signers if s.strip()]
    
    # 验证每个都是人名格式（2-4个中文字符）
    valid_signers = []
    for signer in signers:
        if re.match(r'^[\u4e00-\u9fa5]{2,4}$', signer):
            valid_signers.append(signer)
        else:
            logger.warning(f"签发人格式不符合预期: '{signer}'")
    
    logger.debug(f"提取签发人: {valid_signers}")
    return valid_signers

def _extract_doc_number_and_name(text: str) -> tuple:
    """
    从"发文字号+人名"格式中提取信息（无"签发人："前缀）
    格式：XX〔2024〕2号   李四
    
    返回:
        tuple: (发文字号, 签发人)
    """
    text = text.strip()
    
    # 使用正则提取：发文字号 + 空白 + 人名（2-4个中文字符）
    # 需要匹配多种发文字号格式
    patterns = [
        r'^([\u4e00-\u9fa5A-Za-z0-9]+〔\d{4}〕\s*\d*\s*号)[\s\u3000\t]+([\u4e00-\u9fa5]{2,4})$',
        r'^([\u4e00-\u9fa5A-Za-z0-9]+\[\d{4}\]\s*\d*\s*号)[\s\u3000\t]+([\u4e00-\u9fa5]{2,4})$',
        r'^([\u4e00-\u9fa5A-Za-z0-9]+\(\d{4}\)\s*\d*\s*号)[\s\u3000\t]+([\u4e00-\u9fa5]{2,4})$',
        r'^([\u4e00-\u9fa5A-Za-z0-9]+\d{4}-\s*\d*\s*号)[\s\u3000\t]+([\u4e00-\u9fa5]{2,4})$',
    ]
    
    for pattern in patterns:
        match = re.match(pattern, text)
        if match:
            doc_number = match.group(1).strip()
            name = match.group(2).strip()
            logger.debug(f"提取发文字号和人名: 发文字号='{doc_number}', 人名='{name}'")
            return doc_number, name
    
    logger.debug(f"无法提取发文字号和人名: '{text}'")
    return "", ""
