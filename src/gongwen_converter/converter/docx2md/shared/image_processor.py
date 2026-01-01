"""
DOCX图片提取模块

从DOCX文档中提取内嵌图片并保存到指定文件夹
"""

import os
import logging
from typing import List, Dict, Optional, Callable
from gongwen_converter.utils.path_utils import generate_output_path

logger = logging.getLogger(__name__)

# 支持的图片格式（真实的外部图片）
SUPPORTED_IMAGE_FORMATS = {
    'jpeg', 'jpg', 'png', 'gif', 'tiff', 'tif', 'bmp'
}

# 忽略的格式（Office绘图形状）
IGNORED_FORMATS = {
    'emf', 'wmf'
}


def get_paragraph_images(para, images_info):
    """
    获取段落关联的图片列表（通过 paragraph 对象精确匹配）
    
    这是一个共享函数，可被所有文档转换器使用。
    
    参数:
        para: Word段落对象
        images_info: 图片信息列表
        
    返回:
        list: 与该段落关联的图片列表
    """
    para_images = []
    for img in images_info:
        img_para = img.get('paragraph')
        # 使用底层 XML 元素比较确保匹配准确性
        if img_para is not None and para._element is not None and img_para._element == para._element:
            para_images.append(img)
    return para_images


def extract_images_from_docx(
    doc, 
    output_folder: str, 
    original_file_path: str,
    progress_callback: Optional[Callable[[str], None]] = None
) -> List[Dict]:
    """
    从DOCX文档中递归提取所有图片（包括嵌套表格中的图片）
    
    参数:
        doc: Document对象（python-docx）
        output_folder: 图片保存的文件夹路径
        original_file_path: 原始DOCX文件路径（用于生成规范的图片文件名）
        progress_callback: 进度回调函数（可选）
        
    返回:
        List[Dict]: 图片信息列表，每个元素包含：
            - paragraph: 图片所在的Paragraph对象（关键用于匹配）
            - filename: 图片文件名
            - image_path: 图片完整路径
    """
    logger.info(f"开始递归提取DOCX中的图片到: {output_folder}")
    
    # 确保输出文件夹存在
    os.makedirs(output_folder, exist_ok=True)
    
    images_info = []
    
    # 使用可变对象存储计数器，以便在递归函数中修改
    counter = {'val': 1}
    
    # 内部递归处理函数
    def _process_container(container):
        # 1. 处理容器中的段落
        if hasattr(container, 'paragraphs'):
            for para in container.paragraphs:
                _process_paragraph_images(para, doc, output_folder, original_file_path, images_info, counter, progress_callback)
        
        # 2. 处理容器中的表格
        if hasattr(container, 'tables'):
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # 递归处理单元格（单元格也是一个容器，包含段落和可能的嵌套表格）
                        _process_container(cell)

    try:
        # 从文档根节点开始处理
        _process_container(doc)
    except Exception as e:
        logger.error(f"递归提取图片时出错: {e}", exc_info=True)
    
    logger.info(f"图片提取完成，共提取 {len(images_info)} 张图片")
    return images_info


def _process_paragraph_images(para, doc, output_folder, original_file_path, images_info, counter, progress_callback):
    """
    处理单个段落中的图片
    """
    try:
        drawings = para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
        
        for drawing in drawings:
            blips = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
            
            for blip in blips:
                embed_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                rId = blip.get(embed_attr)
                
                if not rId:
                    continue
                
                try:
                    # 注意：如果是在嵌套部分，doc.part 可能访问不到图片，需要从 para.part 访问
                    # python-docx 的 ElementProxy (如 Paragraph) 通常有 part 属性指向它所属的 Part (DocumentPart, HeaderPart 等)
                    part = para.part
                    image_part = part.related_parts[rId]
                    
                    content_type = image_part.content_type
                    ext = content_type.split('/')[-1] if '/' in content_type else 'png'
                    ext = ext.lower()
                    
                    if ext in SUPPORTED_IMAGE_FORMATS:
                        image_path = generate_output_path(
                            input_path=original_file_path,
                            output_dir=output_folder,
                            section=f"image{counter['val']}",
                            add_timestamp=True,
                            description="fromDocx",
                            file_type=ext
                        )
                        filename = os.path.basename(image_path)
                        
                        with open(image_path, 'wb') as f:
                            f.write(image_part.blob)
                        
                        logger.debug(f"提取图片: {filename}")
                        
                        # 记录图片信息，关键是存储 paragraph 对象用于后续匹配
                        images_info.append({
                            'paragraph': para,  # 关键：存储对象引用
                            'filename': filename,
                            'image_path': image_path,
                            # 兼容性字段（可选）
                            'para_index': -1 
                        })
                        
                        counter['val'] += 1
                        
                        if progress_callback:
                            progress_callback(f"正在提取图片：{counter['val']}")
                        
                    elif ext in IGNORED_FORMATS:
                        pass
                    else:
                        logger.warning(f"未知图片格式: {ext}")
                
                except KeyError:
                    logger.warning(f"无法找到关系ID为 {rId} 的图片部分")
                except Exception as e:
                    logger.error(f"提取图片时出错: {e}")
                    
    except Exception as e:
        logger.error(f"处理段落图片时出错: {e}")


def process_image_with_ocr(
    img: dict, 
    keep_images: bool, 
    enable_ocr: bool, 
    output_folder: str,
    progress_callback=None,
    current_index: int = 1,
    total_images: int = 1,
    cancel_event=None
) -> str:
    """
    根据选项处理图片，生成对应的Markdown链接
    
    这是一个共享函数，可被所有文种转换器使用。
    
    参数:
        img: 图片信息字典，包含 'filename', 'image_path', 'para_index' 等
        keep_images: 是否保留图片
        enable_ocr: 是否启用OCR识别
        output_folder: 输出文件夹路径
        progress_callback: 进度回调函数（可选）
    
    返回:
        str: Markdown链接（图片链接或图片md文件链接）
    """
    from gongwen_converter.config.config_manager import config_manager
    from gongwen_converter.utils.markdown_utils import format_image_link, format_md_file_link
    
    filename = img['filename']
    image_path = img['image_path']
    
    logger.debug(f"处理图片: {filename}, 保留图片: {keep_images}, OCR: {enable_ocr}")
    
    # 获取链接格式配置
    link_settings = config_manager.get_markdown_link_style_settings()
    image_link_style = link_settings.get("image_link_style", "wiki_embed")
    md_file_link_style = link_settings.get("md_file_link_style", "wiki_embed")
    
    # 场景1：只勾选提取图片
    if keep_images and not enable_ocr:
        logger.info(f"场景1：只保留图片 - {filename}")
        return format_image_link(filename, image_link_style)
    
    # 场景2：只勾选OCR
    elif not keep_images and enable_ocr:
        logger.info(f"场景2：只OCR识别 - {filename}")
        # 创建图片md文件（只包含OCR文本）
        md_filename = create_image_md_file(
            image_path, filename, output_folder, 
            include_image=False, include_ocr=True,
            progress_callback=progress_callback,
            current_index=current_index,
            total_images=total_images,
            cancel_event=cancel_event
        )
        return format_md_file_link(md_filename, md_file_link_style)
    
    # 场景3：两者都勾选
    elif keep_images and enable_ocr:
        logger.info(f"场景3：图片 + OCR - {filename}")
        # 创建图片md文件（包含图片链接和OCR文本）
        md_filename = create_image_md_file(
            image_path, filename, output_folder, 
            include_image=True, include_ocr=True,
            progress_callback=progress_callback,
            current_index=current_index,
            total_images=total_images,
            cancel_event=cancel_event
        )
        return format_md_file_link(md_filename, md_file_link_style)
    
    # 都不勾选（不应该出现这种情况，但作为兜底）
    else:
        logger.warning(f"都不勾选时不应该调用此函数 - {filename}")
        return format_image_link(filename, image_link_style)


def create_image_md_file(
    image_path: str, 
    image_filename: str, 
    output_folder: str, 
    include_image: bool, 
    include_ocr: bool,
    progress_callback: Optional[Callable[[str], None]] = None,
    current_index: int = 1,
    total_images: int = 1,
    cancel_event=None
) -> str:
    """
    创建图片的markdown文件
    
    这是一个共享函数，可被所有文种转换器使用。
    
    参数:
        image_path: 图片文件的完整路径
        image_filename: 图片文件名（如 'image_1.png'）
        output_folder: 输出文件夹路径
        include_image: 是否在md文件中包含图片链接
        include_ocr: 是否在md文件中包含OCR识别文本
        progress_callback: 进度回调函数（可选）
    
    返回:
        str: 创建的md文件名（如 'image_1.md'）
    """
    # 生成md文件名（与图片同名，扩展名改为.md）
    base_name = os.path.splitext(image_filename)[0]
    md_filename = f"{base_name}.md"
    md_path = os.path.join(output_folder, md_filename)
    
    try:
        lines = []
        
        # 包含图片链接（使用配置的链接格式）
        if include_image:
            from gongwen_converter.config.config_manager import config_manager
            from gongwen_converter.utils.markdown_utils import format_image_link
            link_settings = config_manager.get_markdown_link_style_settings()
            image_link_style = link_settings.get("image_link_style", "wiki_embed")
            lines.append(format_image_link(image_filename, image_link_style))
            lines.append("")  # 空行
        
        # 包含OCR识别文本
        if include_ocr:
            # 检查取消
            if cancel_event and cancel_event.is_set():
                logger.info("OCR识别被取消")
                return image_filename
            
            # 进度回调
            if progress_callback:
                progress_callback(f"正在OCR识别：{current_index}/{total_images}")
            
            logger.info(f"开始OCR识别: {image_filename}")
            from gongwen_converter.utils.ocr_utils import extract_text_simple
            # 传递cancel_event给OCR函数
            ocr_text = extract_text_simple(image_path, cancel_event)
            
            if ocr_text:
                lines.append(ocr_text)
                logger.info(f"OCR识别完成: {image_filename}, 识别出 {len(ocr_text)} 个字符")
        
        # 写入文件
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        
        logger.info(f"创建图片md文件: {md_filename}")
        return md_filename
        
    except Exception as e:
        logger.error(f"创建图片md文件失败: {md_filename}, 错误: {e}", exc_info=True)
        # 失败时返回图片链接作为兜底
        return image_filename


# 模块测试
if __name__ == "__main__":
    # 配置日志
    logging.basicConfig(
        level=logging.DEBUG,
        format="%(asctime)s [%(levelname)s] %(name)s:%(lineno)d - %(message)s"
    )
    
    logger.info("=== DOCX图片提取模块测试 ===\n")
    
    # 测试文件
    test_docx = "test.docx"
    
    if os.path.exists(test_docx):
        from docx import Document
        import tempfile
        
        try:
            doc = Document(test_docx)
            logger.info(f"成功加载文档: {test_docx}")
            
            with tempfile.TemporaryDirectory() as temp_dir:
                output_folder = os.path.join(temp_dir, "test_images")
                
                images_info = extract_images_from_docx(doc, output_folder, test_docx)
                
                logger.info(f"\n提取结果:")
                logger.info(f"  总计: {len(images_info)} 张图片")
                
                for img in images_info:
                    logger.info(f"  - {img['filename']} (段落 {img['para_index'] + 1})")
                    logger.info(f"    路径: {img['image_path']}")
                    logger.info(f"    文件存在: {os.path.exists(img['image_path'])}")
        
        except Exception as e:
            logger.error(f"测试失败: {e}", exc_info=True)
    else:
        logger.warning(f"测试文件不存在: {test_docx}")
    
    logger.info("\n=== 测试结束 ===")
