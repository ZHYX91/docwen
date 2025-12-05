"""
简化模式DOCX转MD转换器

实现普通Word文档的基础转换逻辑。

特性：
- 基于Word样式（Title、Subtitle、Heading 1-6）转换
- 提取Title/Subtitle到YAML元数据
- 不做公文元素识别
- 支持图片提取和OCR
- 支持文本框和表格处理

与公文模式的区别：
- 无三轮元素识别
- YAML只有2个字段（标题、副标题）
- 不处理附件内容
- 标题直接基于Word样式，不做内容分析
"""

import logging
import re
from docx import Document
from typing import Optional, Callable
import threading

from .content_injector import process_document_with_special_content
from gongwen_converter.utils.text_utils import is_pure_number

# 配置日志
logger = logging.getLogger(__name__)


def convert_docx_to_md_simple(
    docx_path: str,
    extract_image: bool = True,
    extract_ocr: bool = False,
    progress_callback: Optional[Callable[[str], None]] = None,
    cancel_event: Optional[threading.Event] = None,
    output_folder: Optional[str] = None,
    original_file_path: Optional[str] = None
):
    """
    简化模式的DOCX转MD转换
    
    适用于普通Word文档的基础转换，不做公文元素识别。
    
    参数:
        docx_path: DOCX文件路径（可能是临时副本）
        extract_image: 是否保留图片（由GUI传入，默认True）
        extract_ocr: 是否进行OCR识别（由GUI传入，默认False）
        progress_callback: 进度回调函数 (可选)
        cancel_event: 取消事件 (可选)
        output_folder: 输出文件夹路径，用于保存图片 (可选)
        original_file_path: 原始文件路径（用于图片命名，可选）
    
    返回:
        dict: 转换结果字典，包含以下键：
            - success: bool - 转换是否成功
            - main_content: str - 主要markdown内容（含YAML头部）
            - attachment_content: None - 简化模式不处理附件
            - metadata: dict - 提取的YAML元数据（标题、副标题）
            - error: str or None - 错误信息（如果失败）
    """
    # 直接使用GUI传入的参数
    keep_images = extract_image
    enable_ocr = extract_ocr
    
    logger.info(f"开始简化模式转换DOCX文件: {docx_path}")
    logger.info(f"导出选项 - 提取图片: {keep_images}, OCR识别: {enable_ocr}")
    
    try:
        # 1. 加载DOCX文档
        if progress_callback:
            progress_callback("正在加载文档...")
        doc = Document(docx_path)
        logger.info(f"成功加载DOCX文档, 包含 {len(doc.paragraphs)} 个段落")
        
        # 2. 处理文本框和表格内容
        if progress_callback:
            progress_callback("正在处理文本框和表格...")
        logger.info("开始处理文本框和表格内容...")
        
        import tempfile
        temp_dir = tempfile.gettempdir()
        modified_doc, extracted_path = process_document_with_special_content(doc, docx_path, temp_dir)
        
        if extracted_path:
            logger.info(f"提取文档已生成: {extracted_path}")
        else:
            logger.info("提取文档未生成（已禁用输出或生成失败）")
        
        # 使用处理后的文档继续后续流程
        doc = modified_doc
        logger.info(f"处理后的文档包含 {len(doc.paragraphs)} 个段落")
        
        # 3. 提取图片信息
        images_info = []
        if (keep_images or enable_ocr) and output_folder:
            if progress_callback:
                progress_callback("正在提取图片...")
            
            try:
                from .image_processor import extract_images_from_docx
                # 优先使用原始文件路径进行图片命名，如未提供则使用docx_path
                path_for_naming = original_file_path or docx_path
                images_info = extract_images_from_docx(doc, output_folder, path_for_naming)
                logger.info(f"图片提取完成，共 {len(images_info)} 张")
            except Exception as e:
                logger.error(f"图片提取失败: {e}", exc_info=True)
                images_info = []
        elif (keep_images or enable_ocr) and not output_folder:
            logger.warning("配置启用了图片保留或OCR，但未提供output_folder参数，跳过图片提取")
        
        # 4. 提取Title和Subtitle（只取第一个）
        if progress_callback:
            progress_callback("正在提取标题...")
        
        title = ""
        subtitle = ""
        title_indices = []  # 记录Title段落索引（用于跳过）
        subtitle_indices = []  # 记录Subtitle段落索引（用于跳过）
        
        for idx, para in enumerate(doc.paragraphs):
            if cancel_event and cancel_event.is_set():
                logger.info("操作被用户取消")
                return {
                    'success': False,
                    'main_content': None,
                    'attachment_content': None,
                    'metadata': {},
                    'error': '操作已取消'
                }
            
            style_name = para.style.name
            
            # 跳过没有样式的段落
            if not style_name:
                continue
            
            # 提取第一个Title
            if style_name == 'Title' and not title:
                title = para.text.strip()
                title_indices.append(idx)
                logger.info(f"检测到Title样式 (段落 {idx+1}): '{title}'")
            
            # 提取第一个Subtitle
            elif style_name == 'Subtitle' and not subtitle:
                subtitle = para.text.strip()
                subtitle_indices.append(idx)
                logger.info(f"检测到Subtitle样式 (段落 {idx+1}): '{subtitle}'")
            
            # 两个都找到就停止
            if title and subtitle:
                break
        
        # 5. 生成Markdown内容
        if progress_callback:
            progress_callback("正在生成Markdown内容...")
        
        skip_indices = title_indices + subtitle_indices
        main_content = _generate_markdown_content_simple(
            title=title,
            subtitle=subtitle,
            doc=doc,
            skip_indices=skip_indices,
            images_info=images_info,
            keep_images=keep_images,
            enable_ocr=enable_ocr,
            output_folder=output_folder
        )
        logger.info("Markdown内容生成完成")
        
        # 6. 返回结果
        metadata = {
            '标题': title,
            '副标题': subtitle
        }
        
        logger.info("简化模式转换成功")
        return {
            'success': True,
            'main_content': main_content,
            'attachment_content': None,  # 简化模式不处理附件
            'metadata': metadata,
            'error': None
        }
        
    except Exception as e:
        error_msg = f"简化模式转换DOCX文件失败: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return {
            'success': False,
            'main_content': None,
            'attachment_content': None,
            'metadata': {},
            'error': error_msg
        }


def _generate_markdown_content_simple(
    title: str,
    subtitle: str,
    doc,
    skip_indices: list,
    images_info: list = None,
    keep_images: bool = True,
    enable_ocr: bool = False,
    output_folder: str = None
) -> str:
    """
    生成简化模式的Markdown内容
    
    参数:
        title: 文档标题
        subtitle: 文档副标题
        doc: Document对象
        skip_indices: 需要跳过的段落索引列表（Title/Subtitle）
        images_info: 图片信息列表 (可选)
        keep_images: 是否保留图片 (默认True)
        enable_ocr: 是否启用OCR识别 (默认False)
        output_folder: 输出文件夹路径 (可选)
    
    返回:
        str: 完整的Markdown内容（包含YAML头部）
    """
    if images_info is None:
        images_info = []
    
    logger.info(f"生成简化模式Markdown - 标题: '{title}', 副标题: '{subtitle}'")
    lines = []
    
    # 1. 写入YAML头部
    lines.append("---")
    if title:
        # 检查是否为纯数字，如果是则用双引号包裹
        if is_pure_number(title):
            lines.append(f'标题: "{title}"')
        else:
            lines.append(f"标题: {title}")
    if subtitle:
        if is_pure_number(subtitle):
            lines.append(f'副标题: "{subtitle}"')
        else:
            lines.append(f"副标题: {subtitle}")
    lines.append("---")
    lines.append("")
    
    # 2. 处理正文区域段落（包括小标题和正文文本）
    content_lines = []
    
    for idx, para in enumerate(doc.paragraphs):
        # 跳过已提取到YAML的Title/Subtitle
        if idx in skip_indices:
            logger.debug(f"跳过已提取到YAML的段落: {idx+1}")
            continue
        
        # 获取段落文本
        para_text = para.text.strip()
        
        # 处理空段落 - 先检查是否有图片
        if not para_text:
            # 检查当前空段落是否有图片
            para_images = [img for img in images_info if img['para_index'] == idx]
            if para_images:
                for img in para_images:
                    image_link = _process_image_with_ocr(img, keep_images, enable_ocr, output_folder)
                    content_lines.append(image_link)
                    logger.debug(f"在空段落 {idx + 1} 插入图片: {image_link}")
            else:
                logger.debug(f"跳过空段落: {idx+1}")
            continue
        
        # 获取段落样式
        style_name = para.style.name
        
        # 检查是否为Heading样式（添加空值保护）
        if style_name and style_name.startswith('Heading'):
            # 提取级别：'Heading 1' → 1
            match = re.match(r'Heading (\d)', style_name)
            if match:
                level = int(match.group(1))
                # Heading 1-6 直接映射到 # 至 ######
                markdown_line = '#' * level + ' ' + para_text
                content_lines.append(markdown_line)
                logger.debug(f"段落 {idx+1} 转换为标题: Heading {level} → {'#' * level}")
            else:
                # 无法解析级别，当作普通段落
                content_lines.append(para_text)
                logger.warning(f"段落 {idx+1} 样式为 '{style_name}'，但无法解析级别")
        else:
            # 普通段落直接输出
            content_lines.append(para_text)
            logger.debug(f"段落 {idx+1} 作为普通段落: 样式='{style_name}'")
        
        # 检查当前段落是否有图片
        para_images = [img for img in images_info if img['para_index'] == idx]
        for img in para_images:
            # 在段落后插入图片
            image_link = _process_image_with_ocr(img, keep_images, enable_ocr, output_folder)
            content_lines.append(image_link)
            logger.debug(f"在段落 {idx + 1} 后插入图片: {image_link}")
    
    # 3. 添加正文内容（段落之间用两个换行符分隔）
    if content_lines:
        lines.append("\n\n".join(content_lines))
    
    return "\n".join(lines)


def _process_image_with_ocr(img: dict, keep_images: bool, enable_ocr: bool, output_folder: str) -> str:
    """
    根据选项处理图片，生成对应的Markdown链接
    
    参数:
        img: 图片信息字典，包含 'filename', 'image_path', 'para_index' 等
        keep_images: 是否保留图片
        enable_ocr: 是否启用OCR识别
        output_folder: 输出文件夹路径
    
    返回:
        str: Markdown链接（图片链接或图片md文件链接）
    """
    from gongwen_converter.config.config_manager import config_manager
    from gongwen_converter.utils.markdown_utils import format_image_link, format_md_file_link
    
    filename = img['filename']
    image_path = img['image_path']
    
    logger.debug(f"处理图片: {filename}, 保留图片: {keep_images}, OCR: {enable_ocr}")
    
    # 获取链接格式配置
    link_settings = config_manager.get_markdown_link_style_settings()
    image_format = link_settings.get("image_link_format", "wiki")
    image_embed = link_settings.get("image_embed", True)
    md_file_format = link_settings.get("md_file_link_format", "wiki")
    md_file_embed = link_settings.get("md_file_embed", True)
    
    # 场景1：只勾选提取图片
    if keep_images and not enable_ocr:
        logger.info(f"场景1：只保留图片 - {filename}")
        return format_image_link(filename, image_format, image_embed)
    
    # 场景2：只勾选OCR
    elif not keep_images and enable_ocr:
        logger.info(f"场景2：只OCR识别 - {filename}")
        # 创建图片md文件（只包含OCR文本）
        md_filename = _create_image_md_file(image_path, filename, output_folder, 
                                            include_image=False, include_ocr=True)
        return format_md_file_link(md_filename, md_file_format, md_file_embed)
    
    # 场景3：两者都勾选
    elif keep_images and enable_ocr:
        logger.info(f"场景3：图片 + OCR - {filename}")
        # 创建图片md文件（包含图片链接和OCR文本）
        md_filename = _create_image_md_file(image_path, filename, output_folder, 
                                            include_image=True, include_ocr=True)
        return format_md_file_link(md_filename, md_file_format, md_file_embed)
    
    # 都不勾选（不应该出现这种情况，但作为兜底）
    else:
        logger.warning(f"都不勾选时不应该调用此函数 - {filename}")
        return format_image_link(filename, image_format, image_embed)


def _create_image_md_file(image_path: str, image_filename: str, output_folder: str, 
                          include_image: bool, include_ocr: bool) -> str:
    """
    创建图片的markdown文件
    
    参数:
        image_path: 图片文件的完整路径
        image_filename: 图片文件名（如 'image_1.png'）
        output_folder: 输出文件夹路径
        include_image: 是否在md文件中包含图片链接
        include_ocr: 是否在md文件中包含OCR识别文本
    
    返回:
        str: 创建的md文件名（如 'image_1.md'）
    """
    import os
    
    # 生成md文件名（与图片同名，扩展名改为.md）
    base_name = os.path.splitext(image_filename)[0]
    md_filename = f"{base_name}.md"
    md_path = os.path.join(output_folder, md_filename)
    
    try:
        lines = []
        
        # 包含图片链接
        if include_image:
            lines.append(f"![{image_filename}]({image_filename})")
            lines.append("")  # 空行
        
        # 包含OCR识别文本
        if include_ocr:
            logger.info(f"开始OCR识别: {image_filename}")
            from gongwen_converter.utils.ocr_utils import extract_text_simple
            ocr_text = extract_text_simple(image_path)
            
            if ocr_text:
                lines.append(ocr_text)
                logger.info(f"OCR识别完成: {image_filename}, 识别出 {len(ocr_text)} 个字符")
        
        # 写入文件
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        
        logger.info(f"创建图片md文件: {md_filename}")
        return md_filename
        
    except Exception as e:
        logger.error(f"创建图片md文件失败: {md_filename}, 错误: {e}", exc_info=True)
        # 失败时返回图片链接作为兜底
        return image_filename


# 模块测试代码
if __name__ == "__main__":
    # 配置日志
    import logging
    logging.basicConfig(level=logging.DEBUG)
    logger.info("简化模式DOCX转MD测试")
    
    # 测试转换功能
    result = convert_docx_to_md_simple(
        docx_path="test.docx",
        extract_image=True,
        extract_ocr=False
    )
    
    if result['success']:
        print("转换成功!")
        print(f"标题: {result['metadata'].get('标题', '')}")
        print(f"副标题: {result['metadata'].get('副标题', '')}")
        print(f"内容长度: {len(result['main_content'])}")
    else:
        print(f"转换失败: {result['error']}")
