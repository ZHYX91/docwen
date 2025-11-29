"""
PDF转DOCX工具模块

提供PDF到DOCX的转换功能,支持多种转换工具的三级容错机制:
1. Microsoft Word COM接口
2. LibreOffice命令行
3. pdf2docx纯Python库(备选方案)

用于支持PDF灵活提取的ab/ac/abc组合:
- ab: 文本+图片 
- ac: 文本+OCR
- abc: 全功能提取

工作流程:
1. 使用外部工具将PDF转换为DOCX
2. 从DOCX中提取带结构的内容(保证位置关系)
3. 根据选项处理图片和OCR
4. ac组合特殊处理:删除图片只保留OCR文字
"""

import os
import logging
import subprocess
import threading
import tempfile
import shutil
from typing import Optional, Callable

logger = logging.getLogger(__name__)


# ==================== LibreOffice支持 ====================

def _check_libreoffice_available() -> bool:
    """
    检查系统中是否安装了LibreOffice
    
    返回:
        bool: True表示可用,False表示不可用
    """
    try:
        result = subprocess.run(
            ["soffice", "--version"],
            capture_output=True,
            timeout=5,
            text=True
        )
        if result.returncode == 0:
            logger.debug(f"检测到LibreOffice: {result.stdout.strip()}")
            return True
        return False
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False
    except Exception as e:
        logger.debug(f"检测LibreOffice时出错: {e}")
        return False


def _convert_pdf_with_libreoffice(
    pdf_path: str,
    output_dir: str,
    cancel_event: Optional[threading.Event] = None
) -> Optional[str]:
    """
    使用LibreOffice命令行将PDF转换为DOCX
    
    参数:
        pdf_path: PDF文件路径
        output_dir: 输出目录
        cancel_event: 取消事件(可选)
        
    返回:
        str: 转换后的DOCX文件路径,失败时返回None
    """
    try:
        if cancel_event and cancel_event.is_set():
            logger.info("LibreOffice转换被取消")
            return None
        
        cmd = [
            "soffice",
            "--headless",
            "--convert-to", "docx",
            "--outdir", output_dir,
            pdf_path
        ]
        
        logger.info(f"使用LibreOffice转换PDF: {os.path.basename(pdf_path)}")
        logger.debug(f"LibreOffice命令: {' '.join(cmd)}")
        
        # PDF转换可能需要更长的时间
        result = subprocess.run(
            cmd, 
            capture_output=True, 
            timeout=180,  # 3分钟超时
            text=True
        )
        
        if result.returncode == 0:
            base_name = os.path.splitext(os.path.basename(pdf_path))[0]
            output_file = os.path.join(output_dir, f"{base_name}.docx")
            
            if os.path.exists(output_file):
                logger.info(f"✓ LibreOffice转换成功: {os.path.basename(output_file)}")
                return output_file
            else:
                logger.error("LibreOffice执行成功但未找到输出文件")
                if result.stdout:
                    logger.debug(f"stdout: {result.stdout}")
                if result.stderr:
                    logger.debug(f"stderr: {result.stderr}")
                return None
        else:
            logger.error(f"LibreOffice转换失败(返回码: {result.returncode})")
            if result.stderr:
                logger.error(f"错误信息: {result.stderr}")
            return None
            
    except FileNotFoundError:
        logger.debug("LibreOffice未安装")
        return None
    except subprocess.TimeoutExpired:
        logger.error("LibreOffice转换超时(超过3分钟)")
        return None
    except Exception as e:
        logger.error(f"LibreOffice转换出错: {e}")
        return None


# ==================== Word COM接口支持 ====================

def _convert_pdf_with_word(
    pdf_path: str,
    output_path: str,
    cancel_event: Optional[threading.Event] = None
) -> Optional[str]:
    """
    使用Microsoft Word COM接口将PDF转换为DOCX
    
    参数:
        pdf_path: PDF文件路径
        output_path: 输出DOCX文件路径
        cancel_event: 取消事件(可选)
        
    返回:
        str: 转换后的DOCX文件路径,失败时返回None
    """
    try:
        import win32com.client
        import pythoncom
    except ImportError:
        logger.debug("pywin32未安装,无法使用Word COM接口")
        return None
    
    if cancel_event and cancel_event.is_set():
        logger.info("Word转换被取消")
        return None
    
    word_app = None
    doc = None
    
    try:
        pythoncom.CoInitialize()
        
        logger.info(f"使用Microsoft Word转换PDF: {os.path.basename(pdf_path)}")
        
        # 创建Word应用程序实例
        word_app = win32com.client.Dispatch("Word.Application")
        word_app.Visible = False
        word_app.DisplayAlerts = False
        
        # 打开PDF文件
        # OpenFormat参数: 0=默认, 6=RTF, 10=HTML, 等
        abs_pdf_path = os.path.abspath(pdf_path)
        abs_output_path = os.path.abspath(output_path)
        
        logger.debug(f"打开PDF: {abs_pdf_path}")
        doc = word_app.Documents.Open(
            abs_pdf_path,
            ReadOnly=True,
            ConfirmConversions=False,
            AddToRecentFiles=False
        )
        
        if cancel_event and cancel_event.is_set():
            logger.info("Word转换被取消")
            return None
        
        # 保存为DOCX格式
        # FileFormat: 12 = wdFormatXMLDocument (DOCX)
        logger.debug(f"保存为DOCX: {abs_output_path}")
        doc.SaveAs(abs_output_path, FileFormat=12)
        
        logger.info(f"✓ Word转换成功: {os.path.basename(output_path)}")
        return output_path
        
    except Exception as e:
        logger.debug(f"Word COM接口转换失败: {e}")
        return None
    finally:
        # 清理资源
        if doc:
            try:
                doc.Close(SaveChanges=False)
            except:
                pass
        if word_app:
            try:
                word_app.Quit()
            except:
                pass
        try:
            pythoncom.CoUninitialize()
        except:
            pass


# ==================== pdf2docx库支持 ====================

def _convert_pdf_with_pdf2docx(
    pdf_path: str,
    output_path: str,
    cancel_event: Optional[threading.Event] = None
) -> Optional[str]:
    """
    使用pdf2docx纯Python库将PDF转换为DOCX
    
    这是最后的备选方案,不需要外部软件。
    优点: 纯Python实现,跨平台
    缺点: 转换质量可能不如Word/LibreOffice
    
    参数:
        pdf_path: PDF文件路径
        output_path: 输出DOCX文件路径
        cancel_event: 取消事件(可选)
        
    返回:
        str: 转换后的DOCX文件路径,失败时返回None
    """
    try:
        from pdf2docx import Converter
    except ImportError:
        logger.debug("pdf2docx库未安装")
        return None
    
    if cancel_event and cancel_event.is_set():
        logger.info("pdf2docx转换被取消")
        return None
    
    try:
        logger.info(f"使用pdf2docx库转换PDF: {os.path.basename(pdf_path)}")
        
        cv = Converter(pdf_path)
        cv.convert(output_path)
        cv.close()
        
        if os.path.exists(output_path):
            logger.info(f"✓ pdf2docx转换成功: {os.path.basename(output_path)}")
            return output_path
        else:
            logger.error("pdf2docx执行完成但未找到输出文件")
            return None
            
    except Exception as e:
        logger.error(f"pdf2docx转换失败: {e}")
        return None


# ==================== 统一的PDF转DOCX接口 ====================

def convert_pdf_to_docx(
    pdf_path: str,
    output_path: str,
    cancel_event: Optional[threading.Event] = None,
    software_callback: Optional[Callable[[str], None]] = None
) -> Optional[str]:
    """
    将PDF转换为DOCX(三级容错 + 临时目录保护)
    
    改进:
    - 使用临时目录处理PDF副本
    - 保护原始PDF文件不被修改或锁定
    - 所有转换工具操作临时副本
    
    转换优先级:
    1. Microsoft Word COM接口 (质量最佳,但需要安装Word)
    2. LibreOffice命令行 (免费开源,跨平台)
    3. pdf2docx纯Python库 (备选方案,不需要外部软件)
    
    参数:
        pdf_path: PDF文件路径
        output_path: 输出DOCX文件路径
        cancel_event: 取消事件(可选)
        software_callback: 软件使用回调函数(可选),用于记录使用了哪个工具
        
    返回:
        str: 成功时返回output_path
        None: 失败时返回None
        
    异常:
        InterruptedError: 操作被用户取消
        
    使用示例:
        >>> docx_path = convert_pdf_to_docx("test.pdf", "test.docx")
        >>> if docx_path:
        ...     print(f"转换成功: {docx_path}")
    """
    if not os.path.exists(pdf_path):
        logger.error(f"PDF文件不存在: {pdf_path}")
        return None
    
    if cancel_event and cancel_event.is_set():
        logger.info("转换在开始前被取消")
        raise InterruptedError("操作已被取消")
    
    # 确保输出目录存在
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    
    logger.info(f"开始PDF转DOCX: {os.path.basename(pdf_path)}")
    logger.debug(f"输入: {pdf_path}")
    logger.debug(f"输出: {output_path}")
    
    # 使用临时目录处理转换
    with tempfile.TemporaryDirectory() as temp_dir:
        # 步骤1: 复制PDF到临时目录（保护原始文件）
        temp_pdf = os.path.join(temp_dir, "input.pdf")
        shutil.copy2(pdf_path, temp_pdf)
        logger.debug(f"已创建PDF副本: {os.path.basename(temp_pdf)}")
        
        # 步骤2: 在临时目录生成临时DOCX
        temp_docx = os.path.join(temp_dir, "temp_output.docx")
        
        # 步骤3: 三级容错转换（操作临时副本）
        # 第1级: 尝试Microsoft Word
        logger.info("[1/3] 尝试Microsoft Word")
        result = _convert_pdf_with_word(temp_pdf, temp_docx, cancel_event)
        if result and os.path.exists(result):
            logger.info("✓ 转换成功 [Microsoft Word]")
            if software_callback:
                software_callback("Microsoft Word")
            # 移动到最终位置
            shutil.move(result, output_path)
            return output_path
        
        if cancel_event and cancel_event.is_set():
            logger.info("转换被用户取消")
            raise InterruptedError("操作已被取消")
        
        # 第2级: 尝试LibreOffice
        logger.info("[2/3] 尝试LibreOffice")
        result = _convert_pdf_with_libreoffice(temp_pdf, temp_dir, cancel_event)
        if result and os.path.exists(result):
            logger.info("✓ 转换成功 [LibreOffice]")
            if software_callback:
                software_callback("LibreOffice")
            # 移动到最终位置
            shutil.move(result, output_path)
            return output_path
        
        if cancel_event and cancel_event.is_set():
            logger.info("转换被用户取消")
            raise InterruptedError("操作已被取消")
        
        # 第3级: 尝试pdf2docx库
        logger.info("[3/3] 尝试pdf2docx库")
        result = _convert_pdf_with_pdf2docx(temp_pdf, temp_docx, cancel_event)
        if result and os.path.exists(result):
            logger.info("✓ 转换成功 [pdf2docx库]")
            if software_callback:
                software_callback("pdf2docx")
            # 移动到最终位置
            shutil.move(result, output_path)
            return output_path
        
        # 所有方法都失败
        logger.error("✗ PDF转DOCX失败: 所有转换方法均不可用")
        _show_conversion_failed_dialog()
        return None
    # 临时目录自动清理


def _show_conversion_failed_dialog():
    """显示友好的转换失败提示对话框"""
    try:
        import tkinter as tk
        from tkinter import messagebox
        
        root = tk.Tk()
        root.withdraw()
        
        messagebox.showerror(
            "PDF转换失败",
            "无法将PDF转换为DOCX格式,可能的原因:\n\n"
            "❌ 未安装支持的转换工具\n"
            "   • Microsoft Word\n"
            "   • LibreOffice\n"
            "   • pdf2docx库\n\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "建议解决方案:\n\n"
            "1. 安装Microsoft Word(推荐,转换质量最佳)\n"
            "2. 安装LibreOffice(免费开源)\n"
            "3. 安装pdf2docx库:\n"
            "   pip install pdf2docx\n\n"
            "4. 或使用只提取文本(a)、只提取图片(b)、\n"
            "   只OCR(c)、图片+OCR(bc)等\n"
            "   不需要外部工具的组合"
        )
        
        root.destroy()
    except Exception as e:
        logger.error(f"显示错误对话框失败: {e}")


# ==================== 从DOCX提取内容 ====================

def extract_content_from_docx(
    docx_path: str,
    options,
    output_dir: str,
    cancel_event: Optional[threading.Event] = None
) -> dict:
    """
    从DOCX文件中提取内容(文本、图片、OCR)
    
    用于支持PDF灵活提取的ab/ac/abc组合,工作流程:
    1. 从DOCX中提取文本内容
    2. 从DOCX中提取图片
    3. 对图片进行OCR(如果需要)
    4. ac组合特殊处理:删除图片只保留OCR文字
    
    参数:
        docx_path: DOCX文件路径
        options: PDFExtractionOption提取选项
        output_dir: 输出目录
        cancel_event: 取消事件(可选)
        
    返回:
        dict: 提取结果，结构为:
            {
                'text': str,              # 提取的文本内容(如果选择了TEXT)
                'images': List[Dict],      # 提取的图片信息(如果选择了IMAGES)
                'ocr_results': List[str],  # OCR识别结果(如果选择了OCR)
                'method': str,             # 使用的处理方法: 'external_tool'
            }
    """
    try:
        from docx import Document
        from docx.oxml import parse_xml
        from docx.oxml.ns import qn
    except ImportError:
        logger.error("python-docx库未安装")
        raise ImportError("需要python-docx库: pip install python-docx")
    
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX文件不存在: {docx_path}")
    
    logger.info(f"开始从DOCX提取内容: {os.path.basename(docx_path)}")
    
    result = {
        'text': '',
        'images': [],
        'ocr_results': [],
        'method': 'external_tool'
    }
    
    try:
        doc = Document(docx_path)
        logger.info(f"DOCX文档加载成功,共{len(doc.paragraphs)}个段落")
        
        # 1. 提取文本(如果需要)
        if options.has_text():
            if cancel_event and cancel_event.is_set():
                raise InterruptedError("操作已被取消")
            
            logger.info("开始提取文本内容...")
            text_parts = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    text_parts.append(text)
            
            result['text'] = '\n\n'.join(text_parts)
            logger.info(f"文本提取完成,共{len(result['text'])}字符")
        
        # 2. 提取图片(如果需要IMAGES或OCR)
        if options.has_images() or options.has_ocr():
            if cancel_event and cancel_event.is_set():
                raise InterruptedError("操作已被取消")
            
            logger.info("开始提取图片...")
            
            # 创建图片保存目录
            docx_basename = os.path.splitext(os.path.basename(docx_path))[0]
            images_dir = os.path.join(output_dir, f"{docx_basename}_images")
            os.makedirs(images_dir, exist_ok=True)
            
            # 从DOCX中提取所有图片
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        
                        # 确定图片格式
                        image_ext = _get_image_extension(image_data)
                        
                        # 保存图片
                        image_count += 1
                        image_filename = f"image_{image_count}.{image_ext}"
                        image_path = os.path.join(images_dir, image_filename)
                        
                        with open(image_path, 'wb') as img_file:
                            img_file.write(image_data)
                        
                        # 获取图片尺寸
                        try:
                            from PIL import Image
                            with Image.open(image_path) as img:
                                width, height = img.size
                        except:
                            width, height = 0, 0
                        
                        result['images'].append({
                            'type': 'docx_embedded',
                            'path': image_path,
                            'width': width,
                            'height': height
                        })
                        
                        logger.debug(f"提取图片: {image_filename}, 尺寸: {width}x{height}")
                        
                    except Exception as e:
                        logger.warning(f"提取图片失败: {e}")
                        continue
            
            logger.info(f"图片提取完成,共{len(result['images'])}张")
        
        # 3. OCR识别(如果需要)
        if options.has_ocr():
            if cancel_event and cancel_event.is_set():
                raise InterruptedError("操作已被取消")
            
            logger.info("开始OCR识别...")
            
            for idx, img_info in enumerate(result['images'], start=1):
                image_path = img_info['path']
                logger.debug(f"OCR识别 {idx}/{len(result['images'])}: {os.path.basename(image_path)}")
                
                try:
                    from gongwen_converter.utils.ocr_utils import extract_text_simple
                    ocr_text = extract_text_simple(image_path)
                    
                    if ocr_text:
                        logger.debug(f"  识别到{len(ocr_text)}字符")
                    else:
                        logger.debug("  未识别到文字")
                    
                    result['ocr_results'].append(ocr_text)
                    
                except ImportError:
                    logger.warning("OCR功能不可用,请确保已安装PaddleOCR")
                    result['ocr_results'].append("")
                except Exception as e:
                    logger.error(f"  OCR识别失败: {e}")
                    result['ocr_results'].append("")
            
            logger.info(f"OCR识别完成,共{len(result['ocr_results'])}个结果")
        
        # 4. ac组合特殊处理:删除图片只保留OCR文字
        if options.has_text() and options.has_ocr() and not options.has_images():
            logger.info("ac组合:删除图片文件,只保留OCR文字")
            
            # 删除所有图片文件
            for img_info in result['images']:
                try:
                    if os.path.exists(img_info['path']):
                        os.remove(img_info['path'])
                        logger.debug(f"删除图片: {os.path.basename(img_info['path'])}")
                except Exception as e:
                    logger.warning(f"删除图片失败: {e}")
            
            # 清空图片列表
            result['images'] = []
            logger.info("图片文件已删除")
        
        return result
        
    except Exception as e:
        logger.error(f"从DOCX提取内容失败: {e}")
        raise


def _get_image_extension(image_data: bytes) -> str:
    """
    根据图片数据判断图片格式
    
    参数:
        image_data: 图片二进制数据
        
    返回:
        str: 图片扩展名(png, jpg, gif等)
    """
    # 检查文件头判断格式
    if image_data[:8] == b'\x89PNG\r\n\x1a\n':
        return 'png'
    elif image_data[:2] == b'\xff\xd8':
        return 'jpg'
    elif image_data[:6] in (b'GIF87a', b'GIF89a'):
        return 'gif'
    elif image_data[:2] == b'BM':
        return 'bmp'
    else:
        # 默认使用png
        return 'png'


# ==================== 公开API ====================

__all__ = [
    'convert_pdf_to_docx',
    'extract_content_from_docx',
]


# ==================== 模块测试 ====================

if __name__ == "__main__":
    # 配置日志
    logging.basicConfig(
        level=logging.DEBUG,
        format="%(asctime)s [%(levelname)s] %(name)s:%(lineno)d - %(message)s"
    )
    
    logger.info("=== PDF转DOCX工具测试 ===\n")
    
    # 测试文件
    test_pdf = "test.pdf"
    
    if not os.path.exists(test_pdf):
        logger.warning(f"测试文件不存在: {test_pdf}")
        logger.info("请放置一个名为test.pdf的文件进行测试")
    else:
        import tempfile
        
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = os.path.join(temp_dir, "test_output.docx")
            
            logger.info(f"测试PDF: {test_pdf}")
            logger.info(f"输出: {output_path}\n")
            
            def log_software(software_name):
                logger.info(f"使用的转换工具: {software_name}")
            
            result = convert_pdf_to_docx(
                test_pdf, 
                output_path,
                software_callback=log_software
            )
            
            if result:
                file_size = os.path.getsize(result)
                logger.info(f"\n✓ 测试成功!")
                logger.info(f"输出文件: {result}")
                logger.info(f"文件大小: {file_size:,} 字节")
            else:
                logger.error("\n✗ 测试失败!")
    
    logger.info("\n=== 测试结束 ===")
