"""
DOCX文档错别字检查核心处理模块

本模块是校对检查引擎，负责对Word文档进行文本校对和批注添加。
主要功能包括：
- 支持多种校对规则（标点配对、符号校对、错别字校对、敏感词匹配）
- 在DOCX文档中添加详细的批注说明
- 支持单文件处理和批量处理模式
- 自动生成带时间戳的输出文件
- 提供灵活的错别字检查选项配置

采用模块化设计，将文本验证、错误分组、批注添加等功能分离，确保代码的可维护性。
"""

import datetime
import logging
import tempfile
import threading
from collections.abc import Callable
from pathlib import Path

from docx import Document

from docwen.config.config_manager import config_manager
from docwen.proofread_keys import SENSITIVE_WORD, SYMBOL_CORRECTION, SYMBOL_PAIRING, TYPOS_RULE
from docwen.translation import t
from docwen.utils.docx_utils import NAMESPACES as DOCX_NAMESPACES
from docwen.utils.path_utils import generate_output_path

from .spell_checker import TextValidator
from .utils import find_run_at_position, plan_run_splits, rebuild_paragraph_with_splits

# 配置日志
logger = logging.getLogger(__name__)

NON_TEXT_XPATHS = [
    ".//w:drawing",
    ".//m:oMath",
    ".//w:oMath",
    ".//w:hyperlink",
    ".//w:br[@w:type or @w:clear]",
    ".//w:fldChar",
    ".//w:instrText",
    ".//w:bookmarkStart",
    ".//w:bookmarkEnd",
    ".//w:sdt",
    ".//w:footnoteReference",
    ".//w:endnoteReference",
    ".//w:commentReference",
    ".//w:sym",
    ".//w:ruby",
    ".//w:object",
]


def is_code_paragraph(paragraph) -> bool:
    """
    检测段落是否为代码块段落（复用 docx2md 的检测逻辑）

    参数:
        paragraph: docx段落对象

    返回:
        bool: 是否为代码段落
    """
    try:
        # 延迟导入，避免循环导入
        from docwen.converter.docx2md.shared.style_detector import detect_paragraph_style_type

        style_type, _ = detect_paragraph_style_type(paragraph, config_manager)
        return style_type == "code_block"
    except Exception as e:
        logger.warning(f"检测代码段落时出错: {e}")
        return False


def is_quote_paragraph(paragraph) -> bool:
    """
    检测段落是否为引用段落（复用 docx2md 的检测逻辑）

    参数:
        paragraph: docx段落对象

    返回:
        bool: 是否为引用段落
    """
    try:
        # 延迟导入，避免循环导入
        from docwen.converter.docx2md.shared.style_detector import detect_paragraph_style_type

        style_type, _ = detect_paragraph_style_type(paragraph, config_manager)
        return style_type == "quote"
    except Exception as e:
        logger.warning(f"检测引用段落时出错: {e}")
        return False


def has_non_text_content(paragraph) -> bool:
    element = paragraph._element
    try:
        return any(element.xpath(xpath, namespaces=DOCX_NAMESPACES) for xpath in NON_TEXT_XPATHS)
    except Exception as e:
        logger.warning(f"检测段落非文本节点时出错，保守降级为非文本段落: {e}")
        return True


def process_docx(
    docx_path: str,
    output_path: str | None = None,
    output_dir: str | None = None,
    proofread_options: dict[str, bool] | None = None,
    progress_callback: Callable[[str], None] | None = None,
    cancel_event: threading.Event | None = None,
) -> str | None:
    """
    处理DOCX文件，添加错别字批注

    参数:
        docx_path: 输入DOCX文件路径
        output_path: 输出文件路径（可选）
        output_dir: 输出目录（可选，与output_path二选一）
        proofread_options: 校对选项字典
            None: 使用配置文件默认设置
            {}: 空字典表示不进行校对
            {
                "symbol_pairing": True/False,     # 标点配对
                "symbol_correction": True/False,  # 符号校对
                "typos_rule": True/False,         # 错别字校对
                "sensitive_word": True/False      # 敏感词匹配
            }
        progress_callback: 进度回调函数 (可选)

    返回:
        str: 处理后的文件路径

    详细说明:
        1. 如果没有提供输出路径，自动生成带时间戳的输出路径
        2. 格式: 原文件名_YYYYMMDD_HHMMSS.docx
        3. 确保不会覆盖原始文件
        4. 支持自定义校对规则配置
    """
    logger.info("=" * 80)
    logger.info(f"开始处理DOCX文档: {docx_path}")

    # 验证输入文件是否存在
    if not Path(docx_path).exists():
        logger.error(f"输入文件不存在: {docx_path}")
        raise FileNotFoundError(f"输入文件不存在: {docx_path}")

    # 创建输出路径（如果未提供）
    if not output_path:
        # 使用统一的路径生成方法
        output_path = generate_output_path(
            docx_path, output_dir=output_dir, section="", add_timestamp=True, description="checked", file_type="docx"
        )
        logger.info(f"自动生成输出路径: {output_path}")

    try:
        if progress_callback:
            progress_callback(t("conversion.progress.loading_document"))
        logger.info("加载DOCX文档...")
        # 使用docx库加载文档
        doc = Document(docx_path)
        logger.info(f"成功加载文档，包含 {len(doc.paragraphs)} 个段落")

        # 根据选项决定是否进行校对
        # 空字典表示不进行校对
        if proofread_options is not None and not any(proofread_options.values()):
            logger.info("用户未启用任何校对规则，直接保存文档")
            from docwen.utils.workspace_manager import write_temp_file_then_finalize

            with tempfile.TemporaryDirectory() as temp_dir:
                saved_path, _ = write_temp_file_then_finalize(
                    temp_dir=temp_dir,
                    target_path=output_path,
                    original_input_file=docx_path,
                    writer=doc.save,
                )
            if not saved_path:
                raise OSError("文档保存失败")

            logger.info(f"文档保存成功: {saved_path} (大小: {Path(saved_path).stat().st_size} 字节)")
            return saved_path

        # 初始化文本验证器
        if progress_callback:
            progress_callback(t("conversion.progress.initializing_proofread"))
        logger.info("初始化文本验证器...")
        validator = create_validator_with_options(proofread_options)
        logger.info("文本验证器初始化完成")

        # 处理文档中的所有段落
        total_errors = 0
        total_paragraphs = len(doc.paragraphs)

        logger.info(f"开始处理 {total_paragraphs} 个段落...")
        for i, para in enumerate(doc.paragraphs):
            if cancel_event and cancel_event.is_set():
                logger.info("操作被用户取消")
                return None
            # 跳过空段落
            if not para.text.strip():
                logger.debug(f"跳过空段落: {i + 1}/{total_paragraphs}")
                continue

            # 处理当前段落
            errors_found = process_paragraph(i, para, doc, validator)
            total_errors += errors_found

            # 每处理10个段落记录一次进度
            if (i + 1) % 10 == 0 or (i + 1) == total_paragraphs:
                logger.info(f"进度: {i + 1}/{total_paragraphs} 段落 | 累计错误: {total_errors}")
                if progress_callback:
                    progress_callback(
                        t("conversion.progress.proofreading_progress", current=i + 1, total=total_paragraphs)
                    )

        # 保存处理后的文档
        if progress_callback:
            progress_callback(t("conversion.progress.saving_proofread_result"))
        logger.info(f"文档处理完成，共发现 {total_errors} 个错误，准备保存...")
        from docwen.utils.workspace_manager import write_temp_file_then_finalize

        with tempfile.TemporaryDirectory() as temp_dir:
            saved_path, _ = write_temp_file_then_finalize(
                temp_dir=temp_dir,
                target_path=output_path,
                original_input_file=docx_path,
                writer=doc.save,
            )
        if not saved_path:
            raise OSError("文档保存失败")

        logger.info(f"文档保存成功: {saved_path} (大小: {Path(saved_path).stat().st_size} 字节)")
        return saved_path

    except Exception as e:
        logger.error(f"处理文档失败: {e!s}", exc_info=True)
        raise RuntimeError(f"处理文档失败: {e!s}") from e


def create_validator_with_options(proofread_options: dict[str, bool] | None = None) -> TextValidator:
    """
    根据校对选项创建文本验证器

    参数:
        proofread_options: 校对选项字典
            None: 使用配置文件默认设置
            {
                "symbol_pairing": True/False,     # 标点配对
                "symbol_correction": True/False,  # 符号校对
                "typos_rule": True/False,         # 错别字校对
                "sensitive_word": True/False      # 敏感词匹配
            }

    返回:
        TextValidator: 配置好的文本验证器
    """
    # None 表示“未提供显式选择”，按配置默认创建验证器
    if proofread_options is None:
        logger.info("使用配置文件默认设置初始化验证器")
        return TextValidator()

    # 空字典表示“用户未启用任何规则”
    if not proofread_options:
        return TextValidator(
            symbol_pairing=False,
            symbol_correction=False,
            typos_rule=False,
            sensitive_word=False,
        )

    try:
        engine_config = config_manager.get_proofread_engine_config()
    except Exception:
        engine_config = {}

    defaults = {
        SYMBOL_PAIRING: engine_config.get("enable_symbol_pairing", True),
        SYMBOL_CORRECTION: engine_config.get("enable_symbol_correction", True),
        TYPOS_RULE: engine_config.get("enable_typos_rule", True),
        SENSITIVE_WORD: engine_config.get("enable_sensitive_word", True),
    }

    symbol_pairing_enabled = bool(proofread_options.get(SYMBOL_PAIRING, defaults[SYMBOL_PAIRING]))
    symbol_correction_enabled = bool(proofread_options.get(SYMBOL_CORRECTION, defaults[SYMBOL_CORRECTION]))
    typos_rule_enabled = bool(proofread_options.get(TYPOS_RULE, defaults[TYPOS_RULE]))
    sensitive_word_enabled = bool(proofread_options.get(SENSITIVE_WORD, defaults[SENSITIVE_WORD]))

    logger.info("使用校对规则初始化验证器:")
    logger.info(f"  标点配对: {symbol_pairing_enabled}")
    logger.info(f"  符号校对: {symbol_correction_enabled}")
    logger.info(f"  错别字校对: {typos_rule_enabled}")
    logger.info(f"  敏感词匹配: {sensitive_word_enabled}")

    # 创建验证器实例并覆盖配置
    return TextValidator(
        symbol_pairing=symbol_pairing_enabled,
        symbol_correction=symbol_correction_enabled,
        typos_rule=typos_rule_enabled,
        sensitive_word=sensitive_word_enabled,
    )


def batch_process_docx(directory_path: str, output_dir: str | None = None) -> list[str]:
    """
    批量处理目录中的所有DOCX文件（添加时间戳后缀）

    参数:
        directory_path: 包含DOCX文件的目录
        output_dir: 输出目录（可选）

    返回:
        list: 处理成功的文件路径列表
    """
    logger.info(f"开始批量处理目录: {directory_path}")

    # 验证输入目录
    directory = Path(directory_path)
    if not directory.is_dir():
        logger.error(f"输入目录不存在或不是目录: {directory_path}")
        raise NotADirectoryError(f"输入目录不存在或不是目录: {directory_path}")

    # 设置输出目录
    if not output_dir:
        # 使用统一的时间戳格式
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = str(directory / f"checked_docs_{timestamp}")
        logger.info(f"自动设置输出目录: {output_dir}")

    # 创建输出目录
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    logger.info(f"确保输出目录存在: {output_dir}")

    # 查找所有DOCX文件
    docx_files = [str(p) for p in directory.rglob("*.docx") if p.is_file()]

    logger.info(f"找到 {len(docx_files)} 个DOCX文件")

    # 处理每个文件
    processed_files = []
    for i, file_path in enumerate(docx_files):
        try:
            logger.info(f"处理文件 {i + 1}/{len(docx_files)}: {file_path}")

            # 生成输出路径（使用统一方法）
            output_path = generate_output_path(
                file_path,
                output_dir=output_dir,
                section="",
                add_timestamp=True,
                description="checked",
                file_type="docx",
                strict_dir=True,
            )

            # 处理文档
            result = process_docx(file_path, output_path)
            processed_files.append(result)
            logger.info(f"文件处理完成: {result}")

        except Exception as e:
            logger.error(f"处理文件失败: {file_path}, 错误: {e!s}")

    logger.info(f"批量处理完成，成功处理 {len(processed_files)}/{len(docx_files)} 个文件")
    return processed_files


def process_paragraph(para_index: int, paragraph, doc, validator) -> int:
    """
    处理单个段落，添加错别字批注

    参数:
        para_index: 段落索引（从0开始）
        paragraph: 段落对象
        doc: Document对象
        validator: TextValidator实例

    返回:
        int: 本段落发现的错误数量

    详细说明:
        1. 获取段落文本
        2. 检测并跳过代码块/引用段落（根据配置）
        3. 使用验证器检测文本中的错误
        4. 一次性重建段落（标注所有错误位置）
        5. 为每个错误添加批注
    """
    text = paragraph.text
    if not text.strip():
        logger.debug(f"跳过空段落: {para_index + 1}")
        return 0

    try:
        # 记录段落信息
        para_info = f"段落 {para_index + 1}: '{text[:30]}...'" if len(text) > 30 else f"段落 {para_index + 1}: '{text}'"
        logger.debug(f"处理 {para_info}")

        # 检测并跳过代码块（根据配置）
        if config_manager.is_skip_code_blocks_enabled() and is_code_paragraph(paragraph):
            logger.info(f"跳过代码段落: {para_info}")
            return 0

        # 检测并跳过引用块（根据配置）
        if config_manager.is_skip_quote_blocks_enabled() and is_quote_paragraph(paragraph):
            logger.info(f"跳过引用段落: {para_info}")
            return 0

        # 执行文本校对
        logger.debug("开始文本校对...")
        errors = validator.validate_text(text)

        if not errors:
            logger.debug("未发现错误")
            return 0

        logger.info(f"{para_info} 发现 {len(errors)} 个错误")

        # 一次性处理所有错误
        error_count = add_comments_for_errors(paragraph, doc, errors)

        return error_count

    except Exception as e:
        logger.error(f"处理段落失败: {e!s}", exc_info=True)
        return 0


def add_comments_for_errors(paragraph, doc, errors) -> int:
    """
    为段落中的所有错误添加批注（一次性重建）

    参数:
        paragraph: 段落对象
        doc: Document对象
        errors: 错误列表

    返回:
        int: 成功添加的批注数量

    详细说明:
        1. 规划所有错误的run拆分方案
        2. 一次性重建段落（标注所有错误位置）
        3. 为每个错误run添加批注

    改进：
        - 只重建一次段落，避免批注丢失
        - 所有错误都被精确标注
        - 逻辑清晰，易于维护
    """
    try:
        logger.info(f"开始为 {len(errors)} 个错误添加批注")

        if has_non_text_content(paragraph) and config_manager.is_skip_rebuild_on_non_text_enabled():
            if config_manager.is_log_skipped_enabled():
                try:
                    para_text = paragraph.text or ""
                    para_info = para_text[:30] + "..." if len(para_text) > 30 else para_text
                    logger.warning(f"段落包含非文本节点，跳过重建并降级批注: '{para_info}'")
                except Exception:
                    logger.warning("段落包含非文本节点，跳过重建并降级批注")

            success_count = 0
            for idx, error in enumerate(errors):
                error_run = find_run_at_position(paragraph, error.start_pos)
                if error_run is None:
                    error_run = paragraph.add_run("")

                comment_text = f"{error.error_type}：{error.error_text} → {error.suggestion}"
                try:
                    doc.add_comment(error_run, text=comment_text, author=f"DocWen-{error.source}", initials="Sys")
                    logger.info(f"成功添加批注 [{idx + 1}/{len(errors)}] (来源: {error.source})")
                    success_count += 1
                except Exception as e:
                    logger.error(f"添加批注失败 [{idx + 1}/{len(errors)}]: {e!s}")

            logger.info(f"批注添加完成: {success_count}/{len(errors)}")
            return success_count

        # 1. 规划所有错误的run拆分方案
        split_plan = plan_run_splits(paragraph, errors)

        # 2. 一次性重建段落，获取所有错误run
        _new_paragraph, error_runs = rebuild_paragraph_with_splits(paragraph, split_plan, doc)

        if not error_runs:
            logger.warning("重建段落后未找到错误run，无法添加批注")
            return 0

        # 3. 为每个错误run添加批注
        success_count = 0
        for idx, error in enumerate(errors):
            if idx not in error_runs:
                logger.warning(f"未找到错误 {idx} 对应的run")
                continue

            error_run = error_runs[idx]
            comment_text = f"{error.error_type}：{error.error_text} → {error.suggestion}"

            try:
                doc.add_comment(error_run, text=comment_text, author=f"DocWen-{error.source}", initials="Sys")
                logger.info(f"成功添加批注 [{idx + 1}/{len(errors)}] (来源: {error.source})")
                success_count += 1
            except Exception as e:
                logger.error(f"添加批注失败 [{idx + 1}/{len(errors)}]: {e!s}")

        logger.info(f"批注添加完成: {success_count}/{len(errors)}")
        return success_count

    except Exception as e:
        logger.error(f"批量添加批注失败: {e!s}", exc_info=True)
        return 0
