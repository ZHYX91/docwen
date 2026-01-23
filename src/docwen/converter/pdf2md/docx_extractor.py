"""
DOCX内容提取模块

从DOCX文件中提取内容（文本、图片、OCR），
用于支持PDF灵活提取的ab/ac/abc组合。

工作流程：
1. 使用外部工具将PDF转换为DOCX
2. 从DOCX中提取带结构的内容（保证位置关系）
3. 根据选项处理图片和OCR
4. ac组合特殊处理：删除图片只保留OCR文字
"""

import os
import logging
from typing import Optional
import threading

logger = logging.getLogger(__name__)


def extract_content_from_docx(
    docx_path: str,
    options,
    output_dir: str,
    cancel_event: Optional[threading.Event] = None
) -> dict:
    """
    从DOCX文件中提取内容（文本、图片、OCR）
    
    用于支持PDF灵活提取的ab/ac/abc组合，工作流程：
    1. 从DOCX中提取文本内容
    2. 从DOCX中提取图片
    3. 对图片进行OCR（如果需要）
    4. ac组合特殊处理：删除图片只保留OCR文字
    
    参数:
        docx_path: DOCX文件路径
        options: PDFExtractionOption提取选项
        output_dir: 输出目录
        cancel_event: 取消事件（可选）
        
    返回:
        dict: 提取结果，结构为:
            {
                'text': str,              # 提取的文本内容（如果选择了TEXT）
                'images': List[Dict],      # 提取的图片信息（如果选择了IMAGES）
                'ocr_results': List[str],  # OCR识别结果（如果选择了OCR）
                'method': str,             # 使用的处理方法: 'external_tool'
            }
    """
    try:
        from docx import Document
        from docx.oxml import parse_xml
        from docx.oxml.ns import qn
    except ImportError:
        logger.error("python-docx库未安装")
        raise ImportError("需要python-docx库: pip install python-docx")
    
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX文件不存在: {docx_path}")
    
    logger.info(f"开始从DOCX提取内容: {os.path.basename(docx_path)}")
    
    result = {
        'text': '',
        'images': [],
        'ocr_results': [],
        'method': 'external_tool'
    }
    
    try:
        doc = Document(docx_path)
        logger.info(f"DOCX文档加载成功，共{len(doc.paragraphs)}个段落")
        
        # 1. 提取文本（如果需要）
        if options.has_text():
            if cancel_event and cancel_event.is_set():
                raise InterruptedError("操作已被取消")
            
            logger.info("开始提取文本内容...")
            text_parts = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    text_parts.append(text)
            
            result['text'] = '\n\n'.join(text_parts)
            logger.info(f"文本提取完成，共{len(result['text'])}字符")
        
        # 2. 提取图片（如果需要IMAGES或OCR）
        if options.has_images() or options.has_ocr():
            if cancel_event and cancel_event.is_set():
                raise InterruptedError("操作已被取消")
            
            logger.info("开始提取图片...")
            
            # 创建图片保存目录
            docx_basename = os.path.splitext(os.path.basename(docx_path))[0]
            images_dir = os.path.join(output_dir, f"{docx_basename}_images")
            os.makedirs(images_dir, exist_ok=True)
            
            # 从DOCX中提取所有图片
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        
                        # 确定图片格式
                        image_ext = _get_image_extension(image_data)
                        
                        # 保存图片
                        image_count += 1
                        image_filename = f"image_{image_count}.{image_ext}"
                        image_path = os.path.join(images_dir, image_filename)
                        
                        with open(image_path, 'wb') as img_file:
                            img_file.write(image_data)
                        
                        # 获取图片尺寸
                        try:
                            from PIL import Image
                            with Image.open(image_path) as img:
                                width, height = img.size
                        except:
                            width, height = 0, 0
                        
                        result['images'].append({
                            'type': 'docx_embedded',
                            'path': image_path,
                            'width': width,
                            'height': height
                        })
                        
                        logger.debug(f"提取图片: {image_filename}, 尺寸: {width}x{height}")
                        
                    except Exception as e:
                        logger.warning(f"提取图片失败: {e}")
                        continue
            
            logger.info(f"图片提取完成，共{len(result['images'])}张")
        
        # 3. OCR识别（如果需要）
        if options.has_ocr():
            if cancel_event and cancel_event.is_set():
                raise InterruptedError("操作已被取消")
            
            logger.info("开始OCR识别...")
            
            for idx, img_info in enumerate(result['images'], start=1):
                image_path = img_info['path']
                logger.debug(f"OCR识别 {idx}/{len(result['images'])}: {os.path.basename(image_path)}")
                
                try:
                    from docwen.utils.ocr_utils import extract_text_simple
                    ocr_text = extract_text_simple(image_path)
                    
                    if ocr_text:
                        logger.debug(f"  识别到{len(ocr_text)}字符")
                    else:
                        logger.debug("  未识别到文字")
                    
                    result['ocr_results'].append(ocr_text)
                    
                except ImportError:
                    logger.warning("OCR功能不可用，请确保已安装PaddleOCR")
                    result['ocr_results'].append("")
                except Exception as e:
                    logger.error(f"  OCR识别失败: {e}")
                    result['ocr_results'].append("")
            
            logger.info(f"OCR识别完成，共{len(result['ocr_results'])}个结果")
        
        # 4. ac组合特殊处理：删除图片只保留OCR文字
        if options.has_text() and options.has_ocr() and not options.has_images():
            logger.info("ac组合：删除图片文件，只保留OCR文字")
            
            # 删除所有图片文件
            for img_info in result['images']:
                try:
                    if os.path.exists(img_info['path']):
                        os.remove(img_info['path'])
                        logger.debug(f"删除图片: {os.path.basename(img_info['path'])}")
                except Exception as e:
                    logger.warning(f"删除图片失败: {e}")
            
            # 清空图片列表
            result['images'] = []
            logger.info("图片文件已删除")
        
        return result
        
    except Exception as e:
        logger.error(f"从DOCX提取内容失败: {e}")
        raise


def _get_image_extension(image_data: bytes) -> str:
    """
    根据图片数据判断图片格式
    
    参数:
        image_data: 图片二进制数据
        
    返回:
        str: 图片扩展名（png, jpg, gif等）
    """
    # 检查文件头判断格式
    if image_data[:8] == b'\x89PNG\r\n\x1a\n':
        return 'png'
    elif image_data[:2] == b'\xff\xd8':
        return 'jpg'
    elif image_data[:6] in (b'GIF87a', b'GIF89a'):
        return 'gif'
    elif image_data[:2] == b'BM':
        return 'bmp'
    else:
        # 默认使用png
        return 'png'


# ==================== 公开API ====================

__all__ = [
    'extract_content_from_docx',
]
