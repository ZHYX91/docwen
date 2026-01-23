"""
分页/分节/分隔线处理模块

本模块负责处理 Markdown 中的水平线（---、***、___）转换为 Word 文档元素。
根据配置，水平线可以转换为：
- 分页符 (page break)
- 分节符 (section break)
- 分隔线 (horizontal rule style)

主要功能：
- insert_page_break: 插入分页符
- insert_section_break: 插入分节符
- insert_horizontal_rule: 插入分隔线（使用 Horizontal Rule 样式）
- append_*_to_paragraph: 在现有段落末尾添加分隔元素

分节符说明：
分节符必须包含完整的页面设置信息（pgSz, pgMar, cols, docGrid），
否则 Word/WPS 可能无法正确识别分节类型。
本模块会从文档末尾的 sectPr 复制页面设置。
"""

import copy
import logging
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 配置日志
logger = logging.getLogger(__name__)


def insert_page_break(doc):
    """
    插入分页符（创建包含分页符的新段落）
    
    生成 XML:
    <w:p>
        <w:r>
            <w:br w:type="page"/>
        </w:r>
    </w:p>
    
    参数:
        doc: python-docx Document 对象
        
    返回:
        Paragraph: 包含分页符的段落对象
    """
    # 创建空段落
    new_p = doc.add_paragraph()
    
    # 创建 run 元素
    run = new_p.add_run()
    
    # 创建分页符元素 <w:br w:type="page"/>
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)
    
    logger.debug("创建分页符段落")
    return new_p


def insert_section_break(doc, section_type: str = 'next'):
    """
    插入分节符（创建包含分节符的新段落）
    
    分节符通过在段落的 pPr 中添加 sectPr 元素实现。
    sectPr 必须包含完整的页面设置信息，否则 Word/WPS 可能无法正确识别分节类型。
    
    生成 XML:
    <w:p>
        <w:pPr>
            <w:sectPr>
                <w:type w:val="{type}"/>
                <w:pgSz .../>
                <w:pgMar .../>
                ...
            </w:sectPr>
        </w:pPr>
    </w:p>
    
    参数:
        doc: python-docx Document 对象
        section_type: 分节符类型
            - "next" → nextPage（下一页，最常用）
            - "continuous" → continuous（连续）
            - "even" → evenPage（偶数页）
            - "odd" → oddPage（奇数页）
        
    返回:
        Paragraph: 包含分节符的段落对象
    """
    # 映射分节符类型名称
    type_mapping = {
        'next': 'nextPage',
        'continuous': 'continuous',
        'even': 'evenPage',
        'odd': 'oddPage'
    }
    
    word_type = type_mapping.get(section_type, 'nextPage')
    
    # 创建空段落
    new_p = doc.add_paragraph()
    
    # 获取或创建段落属性 <w:pPr>
    pPr = new_p._p.get_or_add_pPr()
    
    # 创建分节属性 <w:sectPr>
    sectPr = OxmlElement('w:sectPr')
    
    # 创建分节类型 <w:type w:val="..."/>
    sect_type = OxmlElement('w:type')
    sect_type.set(qn('w:val'), word_type)
    sectPr.append(sect_type)
    
    # 从文档中复制页面设置信息
    _copy_page_settings_to_sectpr(doc, sectPr)
    
    # 将 sectPr 添加到段落属性中
    pPr.append(sectPr)
    
    logger.debug(f"创建分节符段落: 类型={word_type}")
    return new_p


def insert_horizontal_rule(doc, hr_num: str):
    """
    插入分隔线（应用 Horizontal Rule 1/2/3 段落样式，支持国际化）
    
    分隔线是通过 Horizontal Rule N 段落样式实现的视觉分隔元素。
    该样式由 style_injector 在转换前自动注入到模板中。
    
    Horizontal Rule 样式特点：
    - 空段落 + 底部边框
    - Horizontal Rule 1: 细实线 (0.5pt)
    - Horizontal Rule 2: 中等实线 (1pt)
    - Horizontal Rule 3: 粗实线 (1.5pt)
    
    参数:
        doc: python-docx Document 对象
        hr_num: 分隔线编号 ("1", "2", "3")
        
    返回:
        Paragraph: 应用了 Horizontal Rule N 样式的空段落对象
    """
    # 使用辅助函数查找分隔线样式（支持国际化）
    from ..style.helper import get_horizontal_rule_style_name
    style_name = get_horizontal_rule_style_name(doc, hr_num)
    
    if style_name:
        try:
            new_p = doc.add_paragraph(style=style_name)
            logger.debug(f"创建分隔线段落（{style_name} 样式）")
            return new_p
        except KeyError:
            logger.warning(f"分隔线样式 '{style_name}' 不存在，使用兼容方式")
    else:
        logger.debug(f"未找到分隔线样式 HorizontalRule {hr_num}，使用兼容方式")
    
    # 样式不存在时，创建普通段落并手动添加底部边框
    new_p = doc.add_paragraph()
    _apply_bottom_border_to_paragraph(new_p, hr_num)
    
    return new_p


def append_page_break_to_paragraph(para_elem):
    """
    在段落末尾添加分页符（与文本在同一段落）
    
    用于 attach_to='previous' 场景，将分页符附加到前一个段落末尾，
    而不是创建新段落。
    
    生成 XML: 在现有 <w:p> 末尾添加 <w:r><w:br w:type="page"/></w:r>
    
    参数:
        para_elem: 段落的 XML 元素 (lxml Element)
    """
    # 创建 run 元素
    r = OxmlElement('w:r')
    
    # 创建分页符元素 <w:br w:type="page"/>
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    
    # 添加到段落末尾
    para_elem.append(r)
    logger.debug("在段落末尾添加分页符（同段）")


def append_section_break_to_paragraph(para_elem, doc, section_type: str = 'next'):
    """
    在段落的 pPr 中添加分节符（与文本在同一段落）
    
    用于 attach_to='previous' 场景，将分节符附加到前一个段落，
    而不是创建新段落。
    
    生成 XML: 在段落的 <w:pPr> 中添加 <w:sectPr>...</w:sectPr>
    
    参数:
        para_elem: 段落的 XML 元素 (lxml Element)
        doc: python-docx Document 对象（用于复制页面设置）
        section_type: 分节符类型 ("next", "continuous", "even", "odd")
    """
    WORD_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    
    # 映射分节符类型名称
    type_mapping = {
        'next': 'nextPage',
        'continuous': 'continuous',
        'even': 'evenPage',
        'odd': 'oddPage'
    }
    word_type = type_mapping.get(section_type, 'nextPage')
    
    # 获取或创建段落属性 <w:pPr>
    pPr = para_elem.find(f'{WORD_NS}pPr')
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        # pPr 应该是段落的第一个子元素
        para_elem.insert(0, pPr)
    
    # 创建分节属性 <w:sectPr>
    sectPr = OxmlElement('w:sectPr')
    
    # 创建分节类型 <w:type w:val="..."/>
    sect_type_elem = OxmlElement('w:type')
    sect_type_elem.set(qn('w:val'), word_type)
    sectPr.append(sect_type_elem)
    
    # 从文档中复制页面设置信息
    _copy_page_settings_to_sectpr(doc, sectPr)
    
    # 将 sectPr 添加到段落属性中
    pPr.append(sectPr)
    logger.debug(f"在段落 pPr 中添加分节符（同段）: 类型={word_type}")


def append_horizontal_rule_to_paragraph(para_elem):
    """
    在段落添加底部边框（分隔线效果，与文本在同一段落）
    
    通过设置段落底部边框实现分隔线视觉效果。
    
    生成 XML: 在段落的 <w:pPr> 中添加 <w:pBdr><w:bottom .../></w:pBdr>
    
    参数:
        para_elem: 段落的 XML 元素 (lxml Element)
    """
    WORD_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    
    # 获取或创建段落属性 <w:pPr>
    pPr = para_elem.find(f'{WORD_NS}pPr')
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        # pPr 应该是段落的第一个子元素
        para_elem.insert(0, pPr)
    
    # 创建段落边框容器 <w:pBdr>
    pBdr = OxmlElement('w:pBdr')
    
    # 创建底部边框 <w:bottom w:val="single" w:sz="4" w:space="1" w:color="auto"/>
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')      # 边框宽度（半磅单位，4=0.5pt）
    bottom.set(qn('w:space'), '1')   # 边框与文本的间距（磅）
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    
    # 将 pBdr 添加到段落属性中
    pPr.append(pBdr)
    logger.debug("在段落 pPr 中添加底部边框（分隔线，同段）")


def _copy_page_settings_to_sectpr(doc, sectPr):
    """
    从文档末尾的 sectPr 复制页面设置到新的 sectPr 元素
    
    分节符必须包含完整的页面设置信息（pgSz, pgMar, cols, docGrid），
    否则 Word/WPS 可能无法正确识别分节类型。
    
    参数:
        doc: python-docx Document 对象
        sectPr: 要填充页面设置的 sectPr 元素
    """
    WORD_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    
    try:
        # 获取文档的 body 元素
        body = doc.element.body
        
        # 查找文档末尾的 sectPr（文档默认页面设置）
        doc_sectPr = body.find(f'{WORD_NS}sectPr')
        
        if doc_sectPr is not None:
            # 复制页面设置相关元素
            page_settings_tags = ['pgSz', 'pgMar', 'cols', 'docGrid']
            
            for tag in page_settings_tags:
                source_elem = doc_sectPr.find(f'{WORD_NS}{tag}')
                if source_elem is not None:
                    # 深拷贝元素并添加到新的 sectPr
                    new_elem = copy.deepcopy(source_elem)
                    sectPr.append(new_elem)
                    logger.debug(f"复制页面设置元素: {tag}")
        else:
            logger.warning("未找到文档的默认页面设置，分节符将使用默认值")
            
    except Exception as e:
        logger.warning(f"复制页面设置失败: {e}，分节符将使用默认值")


def _apply_bottom_border_to_paragraph(paragraph, hr_num: str):
    """
    为段落添加底部边框（分隔线兼容方式）
    
    当 Horizontal Rule 样式不存在时，使用此函数手动添加底部边框。
    
    参数:
        paragraph: python-docx Paragraph 对象
        hr_num: 分隔线编号 ("1", "2", "3")，决定边框粗细
    """
    # 边框粗细映射（半磅单位）
    sz_mapping = {
        '1': '4',   # 0.5pt
        '2': '8',   # 1pt
        '3': '12',  # 1.5pt
    }
    sz = sz_mapping.get(hr_num, '4')
    
    # 获取或创建段落属性
    pPr = paragraph._p.get_or_add_pPr()
    
    # 创建段落边框容器
    pBdr = OxmlElement('w:pBdr')
    
    # 创建底部边框
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    
    pPr.append(pBdr)
    logger.debug(f"为段落添加底部边框（兼容方式，sz={sz}）")
