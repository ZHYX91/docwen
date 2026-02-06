"""
表格处理模块

本模块负责 Word 表格的创建和样式应用。
样式注入已移至 style/injector.py，本模块只负责使用已注入的样式。

主要功能：
- create_word_table: 创建 Word 表格并应用样式
- get_table_style_name: 根据用户配置获取表格样式名
- get_table_content_style_name: 获取表格内容段落样式名

WPS 兼容性说明：
WPS 对表格样式的 tblStylePr（条件格式）支持不完整，无法正确渲染 firstRow 中定义的边框。
因此需要在单元格级别（tcPr/tcBorders）显式设置边框，作为"双保险"。
"""

import logging
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .formula_handler import process_paragraph_formulas, is_formula_supported
from docwen.i18n.style_resolver import StyleNameResolver

# 配置日志
logger = logging.getLogger(__name__)


def _apply_cell_alignment(paragraph, alignment: str):
    if alignment == 'default':
        return
    
    alignment_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
    }
    
    if alignment in alignment_map:
        paragraph.alignment = alignment_map[alignment]
        logger.debug(f"设置单元格对齐: {alignment}")


def create_word_table(doc, table_data: dict, fonts: dict = None):
    """
    创建 Word 表格
    
    根据 Markdown 解析出的表格数据创建 Word 表格，并应用适当的样式。
    
    参数:
        doc: python-docx Document 对象
        table_data: 表格数据字典，包含：
            - headers: 表头列表
            - rows: 数据行列表
        fonts: 字体格式配置（可选，用于单元格文字）
    
    返回:
        Table: python-docx 表格对象
    """
    from .text_handler import add_formatted_text_to_paragraph
    from docwen.config.config_manager import config_manager
    
    headers = table_data['headers']
    rows = table_data['rows']
    
    logger.debug(f"创建Word表格: {len(headers)}列 x {len(rows)+1}行（含表头）")
    
    # 创建表格（表头 + 数据行）
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    
    # 1. 根据用户配置获取表格样式名
    table_style_name = get_table_style_name()
    logger.info(f"使用表格样式: {table_style_name}")
    
    # 应用表格样式（优雅降级）
    try:
        table.style = table_style_name
        logger.info(f"成功应用表格样式: {table_style_name}")
    except KeyError:
        logger.warning(f"表格样式 '{table_style_name}' 应用失败，尝试回退到内置表格样式")
        fallback_applied = False
        resolver = StyleNameResolver()
        fallback_candidates = [
            resolver.get_injection_name("table_grid"),
            resolver.get_injection_name("three_line_table"),
        ]
        for fallback_name in fallback_candidates:
            if not fallback_name or fallback_name == table_style_name:
                continue
            try:
                table.style = fallback_name
                logger.info(f"成功回退到表格样式: {fallback_name}")
                fallback_applied = True
                break
            except KeyError:
                continue
        if not fallback_applied:
            logger.warning("未能应用任何回退表格样式，将使用模板默认表格样式")
    
    # 2. 获取表格内容样式名
    content_style_name = get_table_content_style_name()
    alignments = table_data.get('alignments', [])
    
    # 3. 启用首行条件格式（让表格样式的 firstRow 生效）
    _enable_first_row_formatting(table)
    
    # 获取格式处理模式配置
    formatting_mode = config_manager.get_md_to_docx_formatting_mode()
    table_header_formatting_mode = config_manager.get_md_to_docx_table_header_formatting_mode()
    
    # 4. 设置表头
    header_row = table.rows[0]
    for col_idx, header in enumerate(headers):
        cell = header_row.cells[col_idx]
        header_text = str(header) if header else ''
        
        # 清空默认段落
        cell.text = ''
        paragraph = cell.paragraphs[0]
        
        # 应用表格内容样式
        try:
            paragraph.style = content_style_name
        except Exception:
            pass  # 样式不存在时忽略
        
        if col_idx < len(alignments):
            _apply_cell_alignment(paragraph, alignments[col_idx])
        
        # 使用表头专用的格式处理模式（支持行内代码和换行）
        add_formatted_text_to_paragraph(paragraph, header_text, None, table_header_formatting_mode, doc=doc)
        
        logger.debug(f"设置表头单元格 [0,{col_idx}]: {header}")
    
    # 5. WPS 兼容性修复：为首行每个单元格显式设置底部边框
    _apply_header_row_bottom_border(header_row)
    
    # 6. 填充数据行
    # 注意：数据行不传入 fonts 参数，让表格内容样式完全控制字体和字号
    # Markdown 格式（加粗、斜体等）仍然正常解析生效
    for row_idx, row_data in enumerate(rows):
        data_row = table.rows[row_idx + 1]  # +1 因为第0行是表头
        for col_idx, cell_value in enumerate(row_data):
            if col_idx < len(headers):  # 确保不超出列数
                cell = data_row.cells[col_idx]
                cell_text = str(cell_value) if cell_value else ''
                
                # 清空默认段落
                cell.text = ''
                paragraph = cell.paragraphs[0]
                
                # 应用表格内容样式
                try:
                    paragraph.style = content_style_name
                except Exception:
                    pass  # 样式不存在时忽略
                
                if col_idx < len(alignments):
                    _apply_cell_alignment(paragraph, alignments[col_idx])
                
                # 检查是否包含公式
                has_formula = is_formula_supported() and '$' in cell_text
                
                if has_formula:
                    # 使用公式处理器处理单元格内容
                    process_paragraph_formulas(paragraph, cell_text)
                    logger.debug(f"填充数据单元格(含公式) [{row_idx+1},{col_idx}]")
                else:
                    # 使用格式解析（支持行内代码和换行）
                    # 传入 None 避免正文格式覆盖表格内容样式
                    add_formatted_text_to_paragraph(paragraph, cell_text, None, formatting_mode, doc=doc)
                    logger.debug(f"填充数据单元格 [{row_idx+1},{col_idx}]: {cell_value[:30] if cell_value else ''}...")
    
    logger.info(f"Word表格创建完成: {len(headers)}列 x {len(rows)}行数据")
    return table


def get_table_style_name() -> str:
    """
    根据用户配置获取表格样式名
    
    样式已在 style/injector.py 中注入，本函数只返回当前语言对应的样式名。
    
    返回:
        str: 表格样式名称（国际化）
    """
    from docwen.config.config_manager import config_manager
    
    style_resolver = StyleNameResolver()
    
    # 获取表格样式配置
    style_mode = config_manager.get_table_style_mode()
    
    if style_mode == "custom":
        # 自定义样式模式
        custom_name = config_manager.get_custom_table_style_name()
        if custom_name:
            logger.debug(f"使用自定义表格样式: {custom_name}")
            return custom_name
        else:
            # 自定义名称为空，回退到内置三线表
            logger.warning("自定义表格样式名称为空，回退到三线表")
            style_mode = "builtin"
    
    # 内置样式模式
    builtin_key = config_manager.get_builtin_table_style_key()
    
    if builtin_key == "table_grid":
        style_name = style_resolver.get_injection_name("table_grid")
        logger.debug(f"使用网格表样式: {style_name}")
        return style_name
    else:
        style_name = style_resolver.get_injection_name("three_line_table")
        logger.debug(f"使用三线表样式: {style_name}")
        return style_name


def get_table_content_style_name() -> str:
    """
    获取表格内容段落样式名
    
    样式已在 style/injector.py 中注入，本函数只返回当前语言对应的样式名。
    
    返回:
        str: 表格内容样式名称（国际化）
    """
    style_resolver = StyleNameResolver()
    style_name = style_resolver.get_injection_name("table_content")
    logger.debug(f"使用表格内容样式: {style_name}")
    return style_name


def _enable_first_row_formatting(table):
    """
    启用表格的首行条件格式
    
    通过设置 tblLook 元素的 firstRow 属性，让表格样式中定义的 firstRow 条件格式生效。
    
    生成 XML:
    <w:tblPr>
        <w:tblLook w:firstRow="1" w:lastRow="0" w:firstColumn="0" 
                   w:lastColumn="0" w:noHBand="0" w:noVBand="1"/>
    </w:tblPr>
    
    参数:
        table: python-docx Table 对象
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # 添加 tblLook 元素
    tblLook = OxmlElement('w:tblLook')
    tblLook.set(qn('w:firstRow'), '1')      # 启用首行条件格式
    tblLook.set(qn('w:lastRow'), '0')       # 禁用末行条件格式
    tblLook.set(qn('w:firstColumn'), '0')   # 禁用首列条件格式
    tblLook.set(qn('w:lastColumn'), '0')    # 禁用末列条件格式
    tblLook.set(qn('w:noHBand'), '0')       # 启用水平带状
    tblLook.set(qn('w:noVBand'), '1')       # 禁用垂直带状
    tblPr.append(tblLook)
    
    logger.debug("已启用表格首行条件格式 (tblLook.firstRow=1)")


def _apply_header_row_bottom_border(header_row):
    """
    为表头行的每个单元格显式设置底部边框（WPS 兼容性修复）
    
    背景说明：
    Word 和 WPS 对表格样式的 tblStylePr（条件格式）支持存在差异：
    - Word: 完整支持 firstRow 中定义的边框，表头底线正常显示
    - WPS: 对 tblStylePr 支持不完整，可能无法渲染 firstRow 中的边框
    
    解决方案：
    在单元格级别（tcPr/tcBorders）显式设置底部边框，作为"双保险"：
    - Word 中：样式边框和单元格边框都生效，效果一致
    - WPS 中：单元格边框生效，补偿样式边框的缺失
    
    生成 XML:
    <w:tc>
        <w:tcPr>
            <w:tcBorders>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            </w:tcBorders>
        </w:tcPr>
        ...
    </w:tc>
    
    参数:
        header_row: 表格的表头行对象（table.rows[0]）
    """
    WORD_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    
    for cell in header_row.cells:
        tc = cell._tc
        
        # 获取或创建单元格属性 <w:tcPr>
        tcPr = tc.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.insert(0, tcPr)
        
        # 检查是否已有 tcBorders 元素
        existing_borders = tcPr.find(f'{WORD_NS}tcBorders')
        if existing_borders is not None:
            # 已有边框设置，检查是否需要更新底部边框
            existing_bottom = existing_borders.find(f'{WORD_NS}bottom')
            if existing_bottom is not None:
                # 已有底部边框，跳过（避免覆盖用户设置）
                continue
            # 没有底部边框，添加
            tcBorders = existing_borders
        else:
            # 创建边框容器 <w:tcBorders>
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        
        # 创建底部边框
        # sz="4" 对应 0.5pt，与三线表样式一致
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')       # 0.5pt 细线
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000')  # 黑色
        tcBorders.append(bottom)
    
    logger.debug(f"已为表头行 {len(header_row.cells)} 个单元格设置底部边框（WPS兼容性修复）")
