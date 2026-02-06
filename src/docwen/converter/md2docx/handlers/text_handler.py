"""
文本格式化处理模块

负责将 Markdown 格式标记转换为 Word Run 格式：
- ***粗斜体*** → Bold + Italic
- **粗体** → Bold
- *斜体* → Italic
- `代码` → 灰色背景 + Consolas字体
- ~~删除线~~ → Strikethrough
- ==高亮== → 黄色高亮
- ^上标^ → Superscript
- ~下标~ → Subscript (支持 H~2~O 等化学式)
- <u>下划线</u> → Underline

主要组件：
- parse_markdown_formatting(): 解析MD格式标记
- apply_formats_to_run(): 应用格式到Run
- add_formatted_text_to_paragraph(): 添加格式化文本到段落
- add_formatted_text_to_heading(): 添加格式化文本到标题

依赖：
- python-docx: Word文档操作
- config_manager: 读取样式配置
"""

import re
import logging
import emoji
from docx.enum.text import WD_UNDERLINE, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# Emoji 字体名称（Windows 系统）
EMOJI_FONT = "Segoe UI Emoji"


def _split_text_by_emoji(text: str) -> list:
    """
    将文本按 emoji 和非 emoji 部分分离
    
    参数:
        text: 待分离的文本
    
    返回:
        list: 分离后的片段列表，每个元素为元组 (text, is_emoji)
              例如: [("⚠️", True), (" 文件未找到", False)]
    """
    if not text:
        return []
    
    result = []
    last_end = 0
    
    for item in emoji.emoji_list(text):
        start = item["match_start"]
        end = item["match_end"]
        emoji_char = item["emoji"]
        if start > last_end:
            result.append((text[last_end:start], False))
        result.append((emoji_char, True))
        last_end = end
    
    # 最后的非 emoji 部分
    if last_end < len(text):
        result.append((text[last_end:], False))
    
    # 如果没有 emoji，返回整个文本
    if not result:
        result.append((text, False))
    
    return result


def _add_run_with_emoji_support(paragraph, text: str, base_fonts: dict = None):
    """
    添加支持 emoji 字体的 run 到段落
    
    将文本按 emoji 和非 emoji 分离，分别创建 run：
    - emoji 部分：使用 Segoe UI Emoji 字体
    - 非 emoji 部分：使用 base_fonts 或默认字体
    
    参数:
        paragraph: Word 段落对象
        text: 要添加的文本
        base_fonts: 基础字体设置字典（可选，用于非 emoji 部分）
    
    返回:
        list: 创建的 run 列表（用于后续格式应用）
    """
    from docwen.utils.docx_utils import apply_body_formatting
    
    runs = []
    parts = _split_text_by_emoji(text)
    
    for part_text, is_emoji in parts:
        if not part_text:
            continue
        
        run = paragraph.add_run(part_text)
        
        if is_emoji:
            # emoji 部分：设置 emoji 字体
            run.font.name = EMOJI_FONT
            # 确保 rPr 存在
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:ascii'), EMOJI_FONT)
            rFonts.set(qn('w:hAnsi'), EMOJI_FONT)
            rFonts.set(qn('w:eastAsia'), EMOJI_FONT)
            rFonts.set(qn('w:cs'), EMOJI_FONT)
            logger.debug(f"设置 emoji 字体: {EMOJI_FONT}")
        else:
            # 非 emoji 部分：应用基础字体（如果提供）
            if base_fonts:
                apply_body_formatting(run, base_fonts)
        
        runs.append((run, is_emoji))
    
    return runs


# ==================================
#  Markdown格式解析 - MD转DOCX
# ==================================

def parse_markdown_formatting(text: str, mode: str = "apply") -> list:
    """
    解析Markdown格式标记，返回格式化片段列表
    
    参数:
        text: 包含Markdown格式标记的文本
        mode: 处理模式
            - "apply": 解析格式标记，返回带格式信息的片段列表
            - "keep": 保留原样，不解析格式标记
            - "remove": 清理格式标记，只返回纯文本
    
    返回:
        list: 格式化片段列表，每个元素为字典：
            {
                'text': '文本内容',
                'formats': ['bold', 'italic', ...]  # 格式列表
            }
    
    示例:
        >>> parse_markdown_formatting("这是**粗体**文本")
        [{'text': '这是', 'formats': []}, {'text': '粗体', 'formats': ['bold']}, {'text': '文本', 'formats': []}]
    """
    logger.debug(f"解析Markdown格式: mode={mode}, text_len={len(text)}")
    
    if mode == "keep":
        return [{'text': text, 'formats': []}]
    
    if mode == "remove":
        cleaned = _strip_all_format_marks(text)
        return [{'text': cleaned, 'formats': []}]
    
    # mode == "apply": 解析格式标记
    return _parse_with_code_priority(text)


def _strip_all_format_marks(text: str) -> str:
    """
    清除所有格式标记，只保留纯文本
    
    对于代码标记（`code`）：只清理外层反引号，保留内部内容原样
    对于其他格式标记：清理所有层级
    
    参数:
        text: 包含格式标记的文本
    
    返回:
        str: 清除标记后的纯文本
    
    示例:
        >>> _strip_all_format_marks("`**粗体**`")
        '**粗体**'  # 代码内容保留原样
        >>> _strip_all_format_marks("**粗体**")
        '粗体'  # 普通格式标记被清理
    """
    # 使用代码优先的策略：先保护代码内容，再清理其他格式
    code_pattern = re.compile(r'`([^`]+)`')
    
    # 分段处理：代码段只移除反引号，非代码段清理所有格式
    result_parts = []
    last_end = 0
    
    for match in code_pattern.finditer(text):
        # 处理代码之前的非代码文本：清理所有格式标记
        before_text = text[last_end:match.start()]
        if before_text:
            result_parts.append(_strip_non_code_format_marks(before_text))
        
        # 处理代码内容：只移除反引号，保留内部内容原样
        code_content = match.group(1)
        result_parts.append(code_content)
        
        last_end = match.end()
    
    # 处理最后的非代码文本
    if last_end < len(text):
        remaining = text[last_end:]
        result_parts.append(_strip_non_code_format_marks(remaining))
    
    # 如果没有代码，直接清理全部
    if not result_parts:
        return _strip_non_code_format_marks(text)
    
    return ''.join(result_parts)


def _strip_non_code_format_marks(text: str) -> str:
    """
    清除非代码部分的所有格式标记
    
    参数:
        text: 不包含代码标记的文本
    
    返回:
        str: 清除标记后的纯文本
    """
    result = text
    
    # 粗斜体（必须在粗体和斜体之前）
    result = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', result)
    result = re.sub(r'___(.+?)___', r'\1', result)
    # 粗体
    result = re.sub(r'\*\*(.+?)\*\*', r'\1', result)
    result = re.sub(r'__(.+?)__', r'\1', result)
    # 斜体
    result = re.sub(r'\*(.+?)\*', r'\1', result)
    result = re.sub(r'_(.+?)_', r'\1', result)
    # 删除线
    result = re.sub(r'~~(.+?)~~', r'\1', result)
    result = re.sub(r'<del>(.+?)</del>', r'\1', result)
    # 高亮
    result = re.sub(r'==(.+?)==', r'\1', result)
    result = re.sub(r'<mark>(.+?)</mark>', r'\1', result)
    # 上标
    result = re.sub(r'\^(.+?)\^', r'\1', result)
    result = re.sub(r'<sup>(.+?)</sup>', r'\1', result)
    # 下标
    result = re.sub(r'~([^~]+?)~', r'\1', result)
    result = re.sub(r'<sub>(.+?)</sub>', r'\1', result)
    # 下划线
    result = re.sub(r'<u>(.+?)</u>', r'\1', result)
    
    return result


def _parse_with_code_priority(text: str) -> list:
    """
    优先处理代码标记，代码内不解析其他格式
    
    支持的代码格式：
    1. 单反引号行内代码: `code`
    2. 双反引号行内代码: ``code`` (用于包含反引号的代码)
    3. 三反引号围栏代码: ```code``` (表格单元格中的围栏代码块)
    
    参数:
        text: 待解析的文本
    
    返回:
        list: 格式化片段列表
    """
    segments = []
    
    # 按优先级处理代码标记（长的优先）
    # 三反引号 > 双反引号 > 单反引号
    # 注意：这里的三反引号是表格单元格中的行内形式，如 ```code```
    code_pattern = re.compile(r'```([^`]+?)```|``([^`]+?)``|`([^`]+?)`')
    last_end = 0
    
    for match in code_pattern.finditer(text):
        # 代码之前的文本：解析格式
        before_text = text[last_end:match.start()]
        if before_text:
            segments.extend(_parse_formatting_marks(before_text))
        
        # 代码文本：不解析，保留原样，标记为code
        # 三反引号、双反引号、单反引号分别对应group(1)、group(2)、group(3)
        code_content = match.group(1) or match.group(2) or match.group(3)
        segments.append({'text': code_content, 'formats': ['code']})
        
        last_end = match.end()
    
    # 最后的文本：解析格式
    if last_end < len(text):
        remaining = text[last_end:]
        segments.extend(_parse_formatting_marks(remaining))
    
    # 如果没有代码，直接解析全部文本
    if not segments:
        segments = _parse_formatting_marks(text)
    
    return segments


def _parse_formatting_marks(text: str) -> list:
    """
    递归解析格式化文本，返回片段列表（不包含代码）
    
    采用"位置优先"策略：
    - 收集所有可能的格式匹配
    - 按起始位置排序，位置最早的先处理
    - 位置相同时，按匹配长度排序，最长的先处理（外层标签优先）
    
    这确保了嵌套格式的正确处理：
    - <u>下划线中有**粗体**</u> → <u> 位置 0，先处理
    - **<u>粗体下划线</u>** → ** 位置 0 且更长，先处理
    
    支持的格式标记：
    1. 粗斜体: ***text*** 或 ___text___
    2. 粗体: **text** 或 __text__
    3. 斜体: *text* 或 _text_
    4. 删除线: ~~text~~ 或 <del>text</del>
    5. 高亮: ==text== 或 <mark>text</mark>
    6. 上标: ^text^ 或 <sup>text</sup>
    7. 下标: ~text~ 或 <sub>text</sub>
    8. 下划线: <u>text</u>
    
    参数:
        text: 待解析的文本
    
    返回:
        list: 格式化片段列表
    """
    if not text:
        return []
    
    # 格式标记正则表达式
    format_patterns = [
        # 粗斜体（必须在粗体和斜体之前检测，避免被部分匹配）
        (re.compile(r'\*\*\*(.+?)\*\*\*'), 'bold_italic'),
        (re.compile(r'___(.+?)___'), 'bold_italic'),
        # 粗体
        (re.compile(r'\*\*(.+?)\*\*'), 'bold'),
        (re.compile(r'__(.+?)__'), 'bold'),
        # 删除线（必须在下标之前，因为 ~~ 包含 ~）
        (re.compile(r'~~(.+?)~~'), 'strikethrough'),
        (re.compile(r'<del>(.+?)</del>', re.IGNORECASE), 'strikethrough'),
        # 高亮
        (re.compile(r'==(.+?)=='), 'highlight'),
        (re.compile(r'<mark>(.+?)</mark>', re.IGNORECASE), 'highlight'),
        # 上标
        (re.compile(r'\^([^\^]+?)\^'), 'superscript'),
        (re.compile(r'<sup>(.+?)</sup>', re.IGNORECASE), 'superscript'),
        # 下标（排除连续 ~，支持化学式如 H~2~O）
        (re.compile(r'(?<!~)~([^~]+?)~(?!~)'), 'subscript'),
        (re.compile(r'<sub>(.+?)</sub>', re.IGNORECASE), 'subscript'),
        # 下划线
        (re.compile(r'<u>(.+?)</u>', re.IGNORECASE), 'underline'),
        # 斜体（放在最后，避免与粗体冲突）
        (re.compile(r'\*(.+?)\*'), 'italic'),
        # 下划线斜体：确保前后是空白、标点或字符串边界
        (re.compile(r'(?:^|(?<=\s)|(?<=[^\w]))_([^_]+?)_(?=\s|[^\w]|$)'), 'italic'),
    ]
    
    # 收集所有可能的匹配
    all_matches = []
    for pattern, format_type in format_patterns:
        for match in pattern.finditer(text):
            all_matches.append({
                'match': match,
                'format_type': format_type,
                'start': match.start(),
                'end': match.end(),
                'length': match.end() - match.start()
            })
    
    # 如果没有匹配，返回纯文本
    if not all_matches:
        return [{'text': text, 'formats': []}]
    
    # 按 (起始位置, -长度) 排序：位置最早的优先，位置相同则最长的优先
    all_matches.sort(key=lambda x: (x['start'], -x['length']))
    
    # 处理第一个匹配（位置最早且最长的）
    best_match = all_matches[0]
    match = best_match['match']
    format_type = best_match['format_type']
    
    segments = []
    
    # 匹配之前的文本
    before = text[:match.start()]
    if before:
        segments.extend(_parse_formatting_marks(before))
    
    # 匹配的格式化内容（递归处理可能的嵌套格式）
    inner_text = match.group(1)
    inner_segments = _parse_formatting_marks(inner_text)
    
    # 为内部片段添加当前格式
    for seg in inner_segments:
        if format_type not in seg['formats']:
            seg['formats'].append(format_type)
        segments.append(seg)
    
    # 匹配之后的文本
    after = text[match.end():]
    if after:
        segments.extend(_parse_formatting_marks(after))
    
    return segments


def apply_formats_to_run(run, formats: list, doc=None, code_font: str = "Consolas", code_bg_color: str = "D9D9D9", override_style: bool = False):
    """
    将格式列表应用到Word Run对象
    
    参数:
        run: Word Run对象
        formats: 格式列表，如 ['bold', 'italic', 'code']
        doc: Word Document对象（用于获取字符样式，可选）
        code_font: 代码字体名称
        code_bg_color: 代码背景颜色（十六进制）
        override_style: 是否覆盖样式默认格式（用于标题场景）
            - True: 未标记的格式显式设为False，覆盖样式的默认格式
            - False: 只设置标记的格式，未标记的继承样式
    """
    # 如果需要覆盖样式，先显式设置所有基础格式为False
    if override_style:
        if 'bold' not in formats:
            run.bold = False
        if 'italic' not in formats:
            run.italic = False
    
    for fmt in formats:
        if fmt == 'bold':
            run.bold = True
        elif fmt == 'italic':
            run.italic = True
        elif fmt == 'bold_italic':
            run.bold = True
            run.italic = True
        elif fmt == 'strikethrough':
            run.font.strike = True
        elif fmt == 'underline':
            run.underline = WD_UNDERLINE.SINGLE
        elif fmt == 'highlight':
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        elif fmt == 'superscript':
            run.font.superscript = True
        elif fmt == 'subscript':
            run.font.subscript = True
        elif fmt == 'code':
            # 优先尝试应用行内代码字符样式（支持国际化）
            style_applied = False
            if doc is not None:
                from ..style.helper import get_inline_code_style_name
                inline_code_style_name = get_inline_code_style_name(doc)
                if inline_code_style_name:
                    try:
                        run.style = doc.styles[inline_code_style_name]
                        style_applied = True
                        logger.debug(f"应用行内代码字符样式: {inline_code_style_name}")
                    except KeyError:
                        logger.debug(f"行内代码样式 {inline_code_style_name} 不存在，使用底纹兼容")
                else:
                    logger.debug("未找到行内代码样式，使用底纹兼容")
            
            # 如果样式不存在或未传入doc，使用底纹兼容
            if not style_applied:
                run.font.name = code_font
                run._element.rPr.rFonts.set(qn('w:eastAsia'), code_font)
                
                # 添加灰色背景（字符底纹）
                rPr = run._element.get_or_add_rPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), code_bg_color)
                rPr.append(shd)
    
    logger.debug(f"应用格式到Run: {formats}, override_style={override_style}")


def add_formatted_text_to_paragraph(paragraph, text: str, base_fonts: dict = None, mode: str = "apply", doc=None):
    """
    将带Markdown格式的文本添加到段落（用于正文）
    
    支持 emoji 字符的正确显示（自动分离 emoji 并设置专用字体）
    支持 <br> 标签转换为 Word 软回车
    
    参数:
        paragraph: Word段落对象
        text: 包含Markdown格式标记的文本（支持 <br> 换行）
        base_fonts: 基础字体设置字典（可选）
        mode: 格式处理模式 ("apply", "keep", "remove")
        doc: Word Document对象（用于获取字符样式，可选）
    """
    from docwen.utils.docx_utils import apply_body_formatting
    from docwen.config.config_manager import config_manager
    
    # 获取代码配置
    code_font = config_manager.get_code_font()
    code_bg_color = config_manager.get_code_background_color()
    
    # 解析格式
    segments = parse_markdown_formatting(text, mode)
    
    br_pattern = re.compile(r'<br\s*/?>', re.IGNORECASE)

    emitted_any = False
    run_count = 0

    def add_soft_break():
        run = paragraph.add_run()
        br = OxmlElement('w:br')
        run._r.append(br)

    def add_text_with_formats(seg_text: str, seg_formats: list):
        nonlocal emitted_any, run_count
        if not seg_text:
            return

        emoji_parts = _split_text_by_emoji(seg_text)
        for part_text, is_emoji in emoji_parts:
            if not part_text:
                continue

            run = paragraph.add_run(part_text)
            run_count += 1
            emitted_any = True

            if seg_formats and (not is_emoji or 'code' not in seg_formats):
                apply_formats_to_run(
                    run,
                    seg_formats,
                    doc=doc,
                    code_font=code_font,
                    code_bg_color=code_bg_color,
                    override_style=False
                )

            if is_emoji:
                run.font.name = EMOJI_FONT
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:ascii'), EMOJI_FONT)
                rFonts.set(qn('w:hAnsi'), EMOJI_FONT)
                rFonts.set(qn('w:eastAsia'), EMOJI_FONT)
                rFonts.set(qn('w:cs'), EMOJI_FONT)
                logger.debug(f"设置 emoji 字体: {part_text}")
            else:
                if base_fonts and 'code' not in seg_formats:
                    merged_fonts = base_fonts.copy()
                    if 'bold' in seg_formats or 'bold_italic' in seg_formats:
                        merged_fonts['b'] = True
                    if 'italic' in seg_formats or 'bold_italic' in seg_formats:
                        merged_fonts['i'] = True
                    if 'underline' in seg_formats:
                        merged_fonts['u'] = 'single'
                    if 'strikethrough' in seg_formats:
                        merged_fonts['strike'] = True
                    apply_body_formatting(run, merged_fonts)

    for segment in segments:
        seg_text = segment['text']
        seg_formats = segment['formats']

        if not seg_text:
            continue

        if 'code' in seg_formats:
            add_text_with_formats(seg_text, seg_formats)
            continue

        last_end = 0
        for match in br_pattern.finditer(seg_text):
            before = seg_text[last_end:match.start()]
            if before:
                add_text_with_formats(before, seg_formats)
            if emitted_any:
                add_soft_break()
            last_end = match.end()

        remaining = seg_text[last_end:]
        if remaining:
            add_text_with_formats(remaining, seg_formats)
    
    logger.debug(f"添加格式化文本到段落: {len(segments)} 个片段, {run_count} 个 run")


def add_formatted_text_to_heading(paragraph, text: str, mode: str = "apply", doc=None):
    """
    将带Markdown格式的文本添加到标题段落（专用于小标题）
    
    支持 <br> 标签转换为 Word 软回车
    
    与 add_formatted_text_to_paragraph 的区别：
    - 不应用正文字体格式
    - 当 mode="apply" 时，支持覆盖标题样式的默认格式
      （未标记的文字显式设为不加粗/不斜体，实现部分加粗效果）
    
    参数:
        paragraph: Word段落对象（已应用标题样式）
        text: 包含Markdown格式标记的文本（支持 <br> 换行）
        mode: 格式处理模式
            - "apply": 应用格式，覆盖样式默认格式（实现部分加粗）
            - "keep": 保留标记原样
            - "remove": 清理标记，让Word标题样式格式自然生效
        doc: Word Document对象（用于获取字符样式，可选）
    """
    from docwen.config.config_manager import config_manager
    
    # 获取代码配置
    code_font = config_manager.get_code_font()
    code_bg_color = config_manager.get_code_background_color()
    
    # 解析格式
    segments = parse_markdown_formatting(text, mode)
    
    # 判断是否需要覆盖样式
    # 只有在 mode="apply" 且文本中存在格式标记时才覆盖
    has_any_format = any(seg['formats'] for seg in segments)
    override_style = (mode == "apply" and has_any_format)
    
    br_pattern = re.compile(r'<br\s*/?>', re.IGNORECASE)
    emitted_any = False

    def add_soft_break():
        run = paragraph.add_run()
        br = OxmlElement('w:br')
        run._r.append(br)

    def add_text_with_formats(seg_text: str, seg_formats: list):
        nonlocal emitted_any
        if not seg_text:
            return
        emoji_parts = _split_text_by_emoji(seg_text)
        for part_text, is_emoji in emoji_parts:
            if not part_text:
                continue
            run = paragraph.add_run(part_text)
            emitted_any = True
            apply_formats_to_run(
                run,
                seg_formats,
                doc=doc,
                code_font=code_font,
                code_bg_color=code_bg_color,
                override_style=override_style
            )
            if is_emoji:
                run.font.name = EMOJI_FONT
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:ascii'), EMOJI_FONT)
                rFonts.set(qn('w:hAnsi'), EMOJI_FONT)
                rFonts.set(qn('w:eastAsia'), EMOJI_FONT)
                rFonts.set(qn('w:cs'), EMOJI_FONT)

    for segment in segments:
        seg_text = segment['text']
        seg_formats = segment['formats']
        
        if not seg_text:
            continue

        if 'code' in seg_formats:
            add_text_with_formats(seg_text, seg_formats)
            continue

        last_end = 0
        for match in br_pattern.finditer(seg_text):
            before = seg_text[last_end:match.start()]
            if before:
                add_text_with_formats(before, seg_formats)
            if emitted_any:
                add_soft_break()
            last_end = match.end()

        remaining = seg_text[last_end:]
        if remaining:
            add_text_with_formats(remaining, seg_formats)
    
    logger.debug(f"添加格式化文本到标题: {len(segments)} 个片段, override_style={override_style}")


