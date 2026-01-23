"""
MD→DOCX公式处理器

将Markdown中的LaTeX公式转换为Word的Office Math对象。

主要功能：
- 解析Markdown中的LaTeX公式（$...$和$$...$$）
- 将LaTeX转换为OMML（Office Math XML）
- 将OMML插入到Word文档段落中
"""

import logging
import re
from typing import List, Dict, Optional
from docx.oxml import parse_xml, register_element_cls
from docx.oxml.ns import nsdecls
from lxml import etree

from docwen.utils.formula_utils import (
    latex_to_mathml,
    mathml_to_omml,
    is_formula_supported,
    parse_latex_from_markdown,
    OMML_NS
)
from docwen.config.config_manager import config_manager

logger = logging.getLogger(__name__)

# Word 命名空间
WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _inject_style_to_omml(omml_tree, style_id: str) -> None:
    """
    向 OMML 元素树中的所有 m:r 和 m:ctrlPr 注入字符样式
    
    Word 公式中的样式需要应用在：
    1. m:r 元素的 w:rPr/w:rStyle 中 - 公式文本 run
    2. m:ctrlPr 元素的 w:rPr/w:rStyle 中 - 控制属性
    
    XML 结构示例:
    <m:r>
      <w:rPr>
        <w:rStyle w:val="InlineFormula"/>
        <w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/>
      </w:rPr>
      <m:t>E=mc²</m:t>
    </m:r>
    
    参数:
        omml_tree: lxml 的 OMML 元素树
        style_id: 字符样式的 styleId（如 "InlineFormula"）
    """
    if not style_id:
        return
    
    m_ns = OMML_NS['m']
    w_ns = WORD_NS
    
    # 注册命名空间前缀，确保生成的 XML 使用正确的前缀
    nsmap = {'m': m_ns, 'w': w_ns}
    
    def ensure_rPr_and_style(parent_elem):
        """
        确保元素包含 w:rPr/w:rStyle，如果不存在则创建
        """
        # 查找现有的 w:rPr
        rPr = None
        for child in parent_elem:
            if child.tag == '{%s}rPr' % w_ns:
                rPr = child
                break
        
        # 如果没有 w:rPr，创建一个
        if rPr is None:
            rPr = etree.Element('{%s}rPr' % w_ns, nsmap={'w': w_ns})
            # w:rPr 应该是第一个子元素
            parent_elem.insert(0, rPr)
        
        # 查找现有的 w:rStyle
        rStyle = None
        for child in rPr:
            if child.tag == '{%s}rStyle' % w_ns:
                rStyle = child
                break
        
        # 如果没有 w:rStyle，创建一个
        if rStyle is None:
            rStyle = etree.Element('{%s}rStyle' % w_ns, nsmap={'w': w_ns})
            # w:rStyle 应该是 w:rPr 的第一个子元素
            rPr.insert(0, rStyle)
        
        # 设置样式 ID
        rStyle.set('{%s}val' % w_ns, style_id)
    
    # 遍历所有 m:r 元素（公式文本 run）
    for m_r in omml_tree.iter('{%s}r' % m_ns):
        ensure_rPr_and_style(m_r)
    
    # 遍历所有 m:ctrlPr 元素（控制属性）
    for ctrlPr in omml_tree.iter('{%s}ctrlPr' % m_ns):
        ensure_rPr_and_style(ctrlPr)
    
    logger.debug(f"已向公式注入样式: {style_id}")


def insert_formula_to_paragraph(paragraph, latex_str: str, is_inline: bool = True, doc=None, apply_style: bool = True):
    """
    将LaTeX公式插入到Word段落中
    
    参数:
        paragraph: python-docx的Paragraph对象
        latex_str: LaTeX公式字符串
        is_inline: 是否为行内公式（目前Office Math都作为行内处理）
        doc: Document对象（用于获取字符样式，可选）
        apply_style: 是否应用公式样式（默认True）
    
    返回:
        bool: 是否插入成功
    """
    if not is_formula_supported():
        logger.warning("公式转换功能不可用")
        return False
    
    try:
        # LaTeX → MathML
        mathml_str = latex_to_mathml(latex_str)
        if not mathml_str:
            logger.error(f"LaTeX转MathML失败: {latex_str}")
            return False
        
        # MathML → OMML
        omml_str = mathml_to_omml(mathml_str)
        if not omml_str:
            logger.error(f"MathML转OMML失败")
            return False
        
        # 解析OMML XML
        omml_tree = etree.fromstring(omml_str.encode('utf-8'))
        
        # 根据配置注入字符样式到公式内部的 m:r 元素
        # 样式应用在 m:r 内部的 w:rPr/w:rStyle 中，不是在 m:oMath 元素上
        if apply_style and is_inline:
            # 获取行内公式样式名称
            inline_style = config_manager.get_inline_formula_style()
            if inline_style:
                # 将样式名称转换为 styleId（移除空格）
                style_id = inline_style.replace(' ', '')
                _inject_style_to_omml(omml_tree, style_id)
        
        # 根据 Open XML 标准，oMath 应该作为段落 (w:p) 的直接子元素
        # 与文本 Run (w:r) 并列，而不是嵌套在 Run 内部
        if is_inline:
            # 行内公式：直接将 oMath 作为段落的子元素插入
            paragraph._element.append(omml_tree)
        else:
            # 块公式：使用 oMathPara 包装
            omath_para = etree.Element(
                '{%s}oMathPara' % OMML_NS['m'],
                nsmap={'m': OMML_NS['m']}
            )
            omath_para.append(omml_tree)
            
            # 插入到段落
            paragraph._element.append(omath_para)
        
        logger.debug(f"成功插入公式: {latex_str[:50]}...")
        return True
    
    except Exception as e:
        logger.error(f"插入公式失败: {e}, LaTeX: {latex_str}", exc_info=True)
        return False


def process_paragraph_formulas(
    paragraph, 
    para_text: str, 
    fonts: dict = None, 
    formatting_mode: str = "apply", 
    doc = None,
    note_ctx = None
) -> bool:
    """
    处理段落文本中的LaTeX公式，并将其插入到paragraph中
    
    这个函数会：
    1. 解析para_text中的所有LaTeX公式
    2. 将文本按公式分割
    3. 依次添加文本run和公式到paragraph
    4. 公式前后的文本会解析Markdown格式标记
    5. 支持处理文本中的脚注/尾注引用
    
    参数:
        paragraph: python-docx的Paragraph对象（已创建但为空）
        para_text: 包含LaTeX公式的段落文本
        fonts: 字体格式信息字典（可选，用于应用基础格式）
        formatting_mode: Markdown格式处理模式（'apply', 'remove', 'keep'）
        doc: Document对象（可选，用于获取字符样式）
        note_ctx: NoteContext对象（可选，用于脚注/尾注处理）
    
    返回:
        bool: 是否成功处理
    """
    from .text_handler import add_formatted_text_to_paragraph
    from .note_handler import process_text_with_note_references
    
    # 内部辅助函数：根据是否有脚注/尾注引用选择处理方式
    def add_text_segment(text_segment):
        """添加文本片段，支持脚注/尾注处理"""
        if note_ctx and note_ctx.has_notes and '[^' in text_segment:
            # 包含脚注/尾注引用，使用专门的处理函数
            process_text_with_note_references(text_segment, paragraph, fonts, formatting_mode, doc, note_ctx)
        else:
            # 普通文本，使用格式化方式添加
            add_formatted_text_to_paragraph(paragraph, text_segment, fonts, formatting_mode, doc=doc)
    
    if not is_formula_supported():
        # 如果不支持公式，使用格式化方式添加文本（支持脚注/尾注）
        add_text_segment(para_text)
        return False
    
    try:
        # 解析公式
        formulas = parse_latex_from_markdown(para_text)
        
        if not formulas:
            # 没有公式，使用格式化方式添加文本（支持脚注/尾注）
            add_text_segment(para_text)
            return True
        
        # 按位置分割文本和公式
        last_pos = 0
        
        for formula in formulas:
            # 添加公式前的文本（解析Markdown格式，支持脚注/尾注）
            if formula['start'] > last_pos:
                text_before = para_text[last_pos:formula['start']]
                if text_before:
                    add_text_segment(text_before)
            
            # 插入公式
            insert_formula_to_paragraph(
                paragraph,
                formula['latex'],
                formula['is_inline']
            )
            
            last_pos = formula['end']
        
        # 添加最后一段文本（解析Markdown格式，支持脚注/尾注）
        if last_pos < len(para_text):
            text_after = para_text[last_pos:]
            if text_after:
                add_text_segment(text_after)
        
        # 注：公式块行距由 Formula Block 段落样式控制，不再在此处强制设置
        # 
        # 【技术说明 - 关于行内公式字符样式】
        # 字符样式可以应用到公式内部的 m:r 元素，而不是 m:oMath 元素本身。
        # 每个 m:r 元素（公式文本 run）都可以包含 w:rPr/w:rStyle 来引用字符样式。
        # 这样做的好处：
        # 1. 公式可以继承字符样式定义的字体、字号等属性
        # 2. 便于在文档中统一管理和识别行内公式
        # 3. 用户可以通过修改 Inline Formula 样式来批量调整公式格式
        
        logger.debug(f"处理段落公式完成: {len(formulas)} 个公式")
        return True
    
    except Exception as e:
        logger.error(f"处理段落公式失败: {e}", exc_info=True)
        # 发生错误时，至少添加纯文本
        if not paragraph.runs:
            paragraph.add_run(para_text)
        return False


def has_latex_formulas(text: str) -> bool:
    """
    检查文本中是否包含LaTeX公式
    
    参数:
        text: 文本字符串
    
    返回:
        bool: 是否包含公式
    """
    # 快速检查是否包含 $ 符号
    if '$' not in text:
        return False
    
    # 更精确的检查
    formulas = parse_latex_from_markdown(text)
    return len(formulas) > 0


def replace_formulas_with_placeholders(text: str) -> tuple:
    """
    将文本中的LaTeX公式替换为占位符
    
    用于在某些处理流程中临时移除公式，避免干扰其他解析逻辑
    
    参数:
        text: 包含LaTeX公式的文本
    
    返回:
        tuple: (替换后的文本, 公式列表)
    """
    formulas = parse_latex_from_markdown(text)
    
    if not formulas:
        return text, []
    
    # 从后往前替换，避免位置偏移
    result = text
    for idx in range(len(formulas) - 1, -1, -1):
        formula = formulas[idx]
        placeholder = f"{{{{FORMULA_{idx}}}}}"
        result = result[:formula['start']] + placeholder + result[formula['end']:]
    
    return result, formulas


def restore_formulas_from_placeholders(text: str, formulas: List[Dict]) -> str:
    """
    将占位符恢复为LaTeX公式
    
    参数:
        text: 包含占位符的文本
        formulas: 公式列表
    
    返回:
        str: 恢复公式后的文本
    """
    result = text
    
    for idx, formula in enumerate(formulas):
        placeholder = f"{{{{FORMULA_{idx}}}}}"
        
        if formula['is_inline']:
            latex_syntax = f"${formula['latex']}$"
        else:
            latex_syntax = f"$$\n{formula['latex']}\n$$"
        
        result = result.replace(placeholder, latex_syntax)
    
    return result


def convert_block_formulas_to_paragraphs(doc, md_lines: List[str]) -> List[str]:
    """
    预处理Markdown行，将块公式（$$...$$）转换为独立段落
    
    这个函数在md_processor处理之前调用，确保块公式被正确处理
    
    参数:
        doc: Document对象（未使用，保留用于兼容）
        md_lines: Markdown文本行列表
    
    返回:
        list: 处理后的Markdown行列表
    """
    result = []
    in_block_formula = False
    formula_lines = []
    
    for line in md_lines:
        # 检测块公式开始
        if line.strip().startswith('$$'):
            if not in_block_formula:
                # 开始块公式
                in_block_formula = True
                formula_lines = [line]
            else:
                # 结束块公式
                formula_lines.append(line)
                
                # 将整个块公式作为一行
                result.append('\n'.join(formula_lines))
                
                in_block_formula = False
                formula_lines = []
        elif in_block_formula:
            # 在块公式中
            formula_lines.append(line)
        else:
            # 普通行
            result.append(line)
    
    # 处理未闭合的块公式
    if in_block_formula and formula_lines:
        result.extend(formula_lines)
    
    return result


# 模块测试代码
if __name__ == "__main__":
    logging.basicConfig(level=logging.DEBUG)
    
    from docx import Document
    
    # 测试
    test_text = """
    这是一段包含行内公式 $E = mc^2$ 的文本。
    
    下面是块公式：
    
    $$
    \\int_{-\\infty}^{\\infty} e^{-x^2} dx = \\sqrt{\\pi}
    $$
    
    还有更多文本和另一个行内公式 $a^2 + b^2 = c^2$。
    """
    
    print("测试文本:")
    print(test_text)
    print("\n" + "="*60 + "\n")
    
    # 测试公式检测
    if has_latex_formulas(test_text):
        print("✓ 检测到LaTeX公式")
        
        # 解析公式
        formulas = parse_latex_from_markdown(test_text)
        print(f"✓ 解析到 {len(formulas)} 个公式:")
        for idx, formula in enumerate(formulas):
            formula_type = "行内" if formula['is_inline'] else "块"
            print(f"  {idx+1}. [{formula_type}] {formula['latex']}")
        
        # 测试插入到Word
        print("\n测试插入到Word文档...")
        doc = Document()
        
        # 简单测试：为每个公式创建一个段落
        for formula in formulas:
            para = doc.add_paragraph()
            success = insert_formula_to_paragraph(
                para,
                formula['latex'],
                formula['is_inline']
            )
            if success:
                print(f"  ✓ 插入成功: {formula['latex'][:30]}...")
            else:
                print(f"  ✗ 插入失败: {formula['latex'][:30]}...")
        
        # 保存测试文档
        try:
            doc.save("test_formula_output.docx")
            print("\n✓ 测试文档已保存: test_formula_output.docx")
        except Exception as e:
            print(f"\n✗ 保存失败: {e}")
    else:
        print("✗ 未检测到LaTeX公式")
