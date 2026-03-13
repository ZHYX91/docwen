from __future__ import annotations

import copy
import datetime
import logging
import re

from docx.shared import Pt

from docwen.utils.heading_utils import convert_to_halfwidth
from docwen.utils.validation_utils import is_value_empty

from ..field_registry import (
    register_placeholder_rules,
    register_special_handler,
    register_special_placeholders,
    register_yaml_processor,
)
from ..processors.placeholder_handler import remove_special_mark, try_remove_element

logger = logging.getLogger(__name__)

ATTACH_NUM_PATTERN = re.compile(
    r"^[一二三四五六七八九十㈠㈡㈢㈣㈤㈥㈦㈧㈨㈩]+、|"
    r"^（[一二三四五六七八九十]+）|"
    r"^\d+[\.．]\s*|"
    r"^[０１２３４５６７８９]+[\.．]\s*|"
    r"^（\d+）\s*|"
    r"^（[０１２３４５６７８９]+）\s*|"
    r"^[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳㉑㉒㉓㉔㉕㉖㉗㉘㉙㉚㉛㉜㉝㉞㉟㊱㊲㊳㊴㊵㊶㊷㊸㊹㊺㊻㊼㊽㊾㊿⓪⓵⓶⓷⓸⓹⓺⓻⓼⓽⓾⓿❶❷❸❹❺❻❼❽❾❿⓫⓬⓭⓮⓯⓰⓱⓲⓳⓴]\s*"
)

ATTACHMENT_DESCRIPTION_PLACEHOLDER = "{{附件说明}}"

GONGWEN_PLACEHOLDER_RULES = {
    "delete_paragraph_if_empty": [
        ["密级和保密期限"],
        ["紧急程度"],
        ["发文字号"],
        ["公开方式"],
        ["主送机关"],
        ["附注"],
        ["抄送机关"],
        ["附件说明"],
        ["份号", "发文字号"],
    ],
    "delete_cell_if_empty": [],
    "delete_row_if_empty": [
        ["抄送机关"],
        ["印发机关", "印发日期"],
    ],
    "delete_table_if_empty": [],
}


def process_gongwen_yaml(data: dict) -> None:
    process_attachment_description(data)
    process_cc_orgs(data)
    process_special_fields(data)


def process_attachment_description(data: dict) -> None:
    if "附件说明" not in data:
        return

    attachments = data["附件说明"]
    if is_value_empty(attachments):
        data["附件说明"] = []
        return

    if not isinstance(attachments, list):
        attachments = [attachments]

    cleaned_attachments = []
    for item in attachments:
        content = str(item).strip() if item is not None else ""
        normalized_content = convert_to_halfwidth(content)
        cleaned = ATTACH_NUM_PATTERN.sub("", normalized_content).strip()
        if cleaned == "" and content != "":
            cleaned = content
        cleaned_attachments.append(cleaned)

    formatted = []
    for i, content in enumerate(cleaned_attachments, 1):
        if len(cleaned_attachments) == 1:
            formatted.append(f"附件：{content}")
        else:
            if i == 1:
                formatted.append(f"附件：{i}. {content}")
            else:
                indent = "\u3000\u3000\u3000" if i < 10 else "\u3000\u3000 "
                formatted.append(f"{indent}{i}. {content}")

    data["附件说明"] = formatted
    logger.debug("处理附件说明完成，共 %d 项", len(formatted))


def process_cc_orgs(data: dict) -> None:
    if "抄送机关" not in data:
        return
    cc_orgs = data["抄送机关"]
    if is_value_empty(cc_orgs):
        data["抄送机关"] = ""
        return
    if not isinstance(cc_orgs, list):
        cc_orgs = [cc_orgs]
    from docwen.utils.text_utils import format_display_value

    valid_orgs = [format_display_value(org).strip() for org in cc_orgs if not is_value_empty(org)]
    data["抄送机关"] = "，".join(valid_orgs)


def process_special_fields(data: dict) -> None:
    if "附注" in data:
        data["附注"] = process_notes(data["附注"])
    if "印发日期" in data:
        data["印发日期"] = format_date(data["印发日期"], suffix="印发")
    if "成文日期" in data:
        data["成文日期"] = format_date(data["成文日期"])


def process_notes(notes: str) -> str:
    if not notes:
        return ""
    if re.match(r"^[（(].*?[)）]$", notes):
        return notes[1:-1]
    return notes


def format_date(date_str, suffix: str = "") -> str:
    if not date_str:
        return ""
    if isinstance(date_str, (datetime.date, datetime.datetime)):
        date_obj = date_str
    else:
        date_formats = ["%Y-%m-%d", "%Y/%m/%d", "%Y年%m月%d日", "%Y.%m.%d", "%Y年%m月%d号"]
        date_obj = None
        for fmt in date_formats:
            try:
                date_obj = datetime.datetime.strptime(str(date_str), fmt)
                break
            except ValueError:
                continue
    if date_obj:
        formatted = f"{date_obj.year}年{date_obj.month}月{date_obj.day}日"
        if suffix:
            formatted += suffix
        return formatted
    return str(date_str)


def process_attachment_description_placeholder(doc, yaml_data: dict) -> None:
    logger.info("开始处理附件说明占位符...")

    attach_desc_para = None
    for para_idx, para in enumerate(doc.paragraphs):
        if ATTACHMENT_DESCRIPTION_PLACEHOLDER in para.text:
            remove_special_mark(para)
            logger.debug("在第 %d 段找到附件说明占位符", para_idx + 1)
            attach_desc_para = para
            break

    if not attach_desc_para:
        logger.info("未找到附件说明占位符")
        return

    attachment_desc = yaml_data.get("附件说明")
    if is_value_empty(attachment_desc):
        logger.info("附件说明为空，删除占位符段落")
        try_remove_element(attach_desc_para._element)
        return

    if not isinstance(attachment_desc, list):
        attachment_desc = [attachment_desc]

    base_style = attach_desc_para.style
    base_rpr = attach_desc_para.runs[0]._element.rPr if attach_desc_para.runs else None

    left_indent = attach_desc_para.paragraph_format.left_indent
    if left_indent:
        char_width = left_indent / 2
    else:
        char_width = Pt(16)
        left_indent = 2 * char_width

    parent = attach_desc_para._element.getparent()
    index = parent.index(attach_desc_para._element)
    is_single = len(attachment_desc) == 1
    hanging_indent_single = int(3 * char_width)
    hanging_indent_multi = int(4.5 * char_width)

    for i, line in enumerate(attachment_desc):
        new_p = doc.add_paragraph(style=base_style)
        new_run = new_p.add_run(str(line))
        if base_rpr is not None:
            new_rpr = new_run._element.get_or_add_rPr()
            for child in base_rpr.getchildren():
                new_rpr.append(copy.deepcopy(child))
        pf = new_p.paragraph_format
        pf.left_indent = left_indent
        pf.first_line_indent = -hanging_indent_single if is_single else -hanging_indent_multi
        new_p._p.getparent().remove(new_p._p)
        parent.insert(index + i, new_p._p)

    parent.remove(attach_desc_para._element)
    logger.info("附件说明占位符处理完成")


register_yaml_processor("gongwen", process_gongwen_yaml)
register_placeholder_rules("gongwen", GONGWEN_PLACEHOLDER_RULES)
register_special_placeholders("gongwen", {ATTACHMENT_DESCRIPTION_PLACEHOLDER})
register_special_handler(ATTACHMENT_DESCRIPTION_PLACEHOLDER, process_attachment_description_placeholder)
