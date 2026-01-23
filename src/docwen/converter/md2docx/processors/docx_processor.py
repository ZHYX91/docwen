"""
DOCX处理器模块

负责文档的主要处理流程，包括：
- 加载文档模板
- 替换占位符
- 处理正文内容（标题、段落、列表、代码块、引用、公式等）
- 保存和后处理

相关模块：
- handlers/table_handler.py: 表格创建与样式处理
- handlers/break_handler.py: 分页/分节/分隔线处理
- style/helper.py: 样式查找辅助函数
- handlers/note_handler.py: 脚注/尾注处理
- processors/placeholder_handler.py: 占位符替换
- processors/xml_processor.py: XML深度处理
"""

import logging
import tempfile
from docx import Document
from docx.shared import Pt, Cm
from docx.text.paragraph import Paragraph
from .placeholder_handler import (
    mark_special_placeholders,
    is_special_marked,
    remove_special_mark,
    process_paragraph_placeholders,
    process_table_cell_placeholders,
    process_attachment_description_placeholder,
    try_remove_element,
    get_body_placeholder_variants,
)
from . import xml_processor
from ..style.injector import ensure_styles

# 样式查找函数（从 style 模块导入）
from ..style.helper import (
    get_heading_style_name,
    get_code_block_style_name,
    get_quote_style_name,
    get_formula_block_style_name,
    get_list_block_style_name,
)

# 表格处理（从 handlers 模块导入）
from ..handlers.table_handler import create_word_table

# 分页/分节/分隔线处理（从 handlers 模块导入）
from ..handlers.break_handler import (
    insert_page_break,
    insert_section_break,
    insert_horizontal_rule,
    append_page_break_to_paragraph,
    append_section_break_to_paragraph,
    append_horizontal_rule_to_paragraph,
)

# 脚注/尾注处理（从 handlers 模块导入）
from ..handlers.note_handler import (
    NoteContext,
    find_note_references,
    is_endnote_id,
    ENDNOTE_PREFIX,
    write_notes_to_docx,
    process_text_with_note_references,
)
from docwen.utils import docx_utils
from ..handlers.formula_handler import process_paragraph_formulas, is_formula_supported
from ..handlers.text_handler import add_formatted_text_to_paragraph, add_formatted_text_to_heading
from docwen.config.config_manager import config_manager
from ..handlers.list_handler import (
    ListFormatManager,
    apply_list_to_paragraph,
    analyze_list_structure,
    group_consecutive_list_items,
    BASE_INDENT,
    INDENT_INCREMENT,
)

# 配置日志
logger = logging.getLogger(__name__)


def load_document(template_path):
    """加载Word文档模板"""
    try:
        doc = Document(template_path)
        logger.info(f"成功加载模板文件: {template_path}")
        return doc
    except Exception as e:
        logger.error(f"加载模板文件失败: {template_path}, 错误: {str(e)}")
        raise

def replace_placeholders(doc, yaml_data, body_data, template_name=None, *, footnotes=None, endnotes=None):
    """
    主处理函数：替换文档中的所有占位符
    
    参数:
        doc: Document对象
        yaml_data: 字典，包含要替换的键值对
        body_data: 列表，包含处理后的正文段落数据
        template_name: 模板名称（可选，用于错误提示）
        footnotes: 脚注字典 {id: content}（可选）
        endnotes: 尾注字典 {id: content}（可选）
        
    返回:
        tuple: (临时文件路径, 警告列表)
               警告列表包含需要通知用户的警告信息
    """
    # 初始化脚注/尾注上下文（作为局部变量传递给其他函数，避免全局变量线程安全问题）
    note_ctx = None
    if footnotes or endnotes:
        note_ctx = NoteContext(footnotes or {}, endnotes or {})
        note_ctx.init_styles_from_doc(doc)
        logger.info(f"脚注/尾注上下文初始化 | 脚注: {len(note_ctx.footnotes)} 个, 尾注: {len(note_ctx.endnotes)} 个")
    
    logger.info("开始处理文档占位符...")
    
    # 第一步：标记特殊占位符段落
    mark_special_placeholders(doc)
    
    # 第二步：处理常规段落占位符
    paragraphs_to_remove = []
    for paragraph in doc.paragraphs:
        # 跳过已标记的特殊占位符
        if not is_special_marked(paragraph):
            logger.debug(f"处理段落: {paragraph.text[:50]}...")
            process_paragraph_placeholders(paragraph, yaml_data, paragraphs_to_remove)
    
    # 第三步：处理表格中的占位符
    rows_to_remove = []
    tables_to_remove = []  # 存储需要删除的表格
    for table in doc.tables:
        row_removal_flags = []  # 标记需要删除的行
        for row_idx, row in enumerate(table.rows):
            # 【新增】收集整行的所有占位符（用于行级删除检查）
            row_placeholder_keys = set()
            for cell in row.cells:
                for para in cell.paragraphs:
                    # 提取占位符键名并添加到整行集合
                    from .placeholder_handler import extract_placeholder_keys
                    cell_keys = extract_placeholder_keys(para.text)
                    row_placeholder_keys.update(cell_keys)
            
            logger.debug(f"表格第{row_idx+1}行的占位符: {row_placeholder_keys}")
            
            # 处理单元格，传入整行占位符集合
            row_should_remove = False  # 标记当前行是否需要删除
            for cell in row.cells:
                for para in cell.paragraphs:
                    # 处理单元格占位符，传入整行占位符集合，返回是否需要删除行
                    if process_table_cell_placeholders(para, yaml_data, row, rows_to_remove, row_placeholder_keys):
                        row_should_remove = True
                        break
                if row_should_remove:
                    break
            row_removal_flags.append(row_should_remove)
        
        # 检查表格是否所有行都需要删除
        # 如果所有行都需要删除，则删除整个表格（避免留下空表格框架）
        if row_removal_flags and all(row_removal_flags):
            logger.info(f"表格中所有行都需要删除，标记删除整个表格")
            tables_to_remove.append(table)
    
    # 第四步：处理正文占位符
    # 收集警告信息列表
    warnings = []
    found_body_placeholder = process_main_content(doc, body_data, yaml_data, template_name, note_ctx=note_ctx)
    if not found_body_placeholder:
        from docwen.i18n import t
        warnings.append(t('messages.warnings.missing_body_placeholder'))

    # 第五步：处理附件说明占位符
    process_attachment_description_placeholder(doc, yaml_data)

    # 第六步：处理图片占位符
    from .placeholder_handler import process_image_placeholders
    process_image_placeholders(doc)

    # 第七步：删除需要移除的元素
    # 先删除表格（如果表格需要删除）
    for table in tables_to_remove:
        try:
            table._element.getparent().remove(table._element)
            logger.info(f"已删除整个表格（只剩一行且需要删除）")
        except Exception as e:
            logger.warning(f"删除表格失败: {str(e)}")
    
    # 删除需要移除的段落和表格行
    remove_marked_elements(paragraphs_to_remove, rows_to_remove)

    # 第八步：深度处理所有XML内容
    temp_path = save_and_process_temp_file(doc, yaml_data, note_ctx=note_ctx)

    return temp_path, warnings

def _iter_body_paragraphs(doc):
    from docx.oxml.ns import qn
    normal = []
    in_table = []
    in_textbox = []
    for p in doc._element.body.iter(qn('w:p')):
        parent = p
        is_in_table = False
        is_in_textbox = False
        while parent is not None:
            parent = parent.getparent()
            if parent is None:
                break
            if parent.tag == qn('w:txbxContent'):
                is_in_textbox = True
                break
            if parent.tag == qn('w:tbl'):
                is_in_table = True
        if is_in_textbox:
            in_textbox.append(p)
        elif is_in_table:
            in_table.append(p)
        else:
            normal.append(p)
    for p in normal + in_table + in_textbox:
        yield Paragraph(p, doc)

def process_main_content(doc, body_data, yaml_data, template_name=None, *, note_ctx=None):
    """
    处理正文占位符
    
    参数:
        doc: Document对象
        body_data: 正文段落数据列表
        yaml_data: YAML元数据字典
        template_name: 模板名称（可选，用于错误提示）
        note_ctx: NoteContext对象（可选，用于脚注/尾注处理）
    
    返回:
        bool: 是否找到正文占位符（如果未找到，正文内容不会被插入）
    """
    # ========================================================================
    # 列表预处理：分组分析并分配 numId
    # 
    # 根据 Markdown 语法，连续的列表项属于同一个列表，非列表内容终止当前列表。
    # 预处理阶段为每个独立的列表组创建一个 Word 列表定义（numId）。
    # ========================================================================
    _preprocess_list_items(doc, body_data)
    
    for paragraph in list(_iter_body_paragraphs(doc)):
        # 获取所有正文占位符变体（支持国际化：{{正文}}、{{body}} 等）
        body_variants = get_body_placeholder_variants()
        
        # 查找段落中是否包含任一正文占位符变体
        found_variant = None
        for variant in body_variants:
            if variant in paragraph.text:
                found_variant = variant
                break
        
        if found_variant:
            # 移除特殊标记
            remove_special_mark(paragraph)
            
            base_style = paragraph.style
            
            if not paragraph.runs:
                paragraph.add_run(found_variant)

            # 提取格式信息
            fonts = docx_utils.extract_format_from_paragraph(paragraph)
            logger.info(f"正文格式提取结果 - 字体: {fonts['eastAsia']}, 字号: {fonts['sz']}")

            # 准备插入新内容
            parent = paragraph._element.getparent()
            index = parent.index(paragraph._element)
            
            # 插入处理后的内容
            # 使用单独的插入位置计数器，因为某些项目（如代码块）可能占用多行
            insert_offset = 0
            
            for i, item in enumerate(body_data):
                item_type = item.get('type')
                level = item.get('level', 0)
                
                # 处理表格
                if item_type == 'table':
                    table_data = item.get('table_data')
                    word_table = create_word_table(doc, table_data, fonts)
                    # 插入表格到正确位置
                    parent.insert(index + insert_offset, word_table._element)
                    insert_offset += 1
                    logger.info(f"插入Word表格: {len(table_data['headers'])}列 x {len(table_data['rows'])}行")
                    continue
                
                # 处理代码块
                if item_type == 'code_block':
                    code_content = item.get('text', '')
                    language = item.get('language', None)
                    in_list_context = item.get('in_list_context', False)
                    list_level = item.get('list_level', 0)
                    
                    # 尝试使用 Code Block 样式（从 style 模块获取）
                    code_block_style_name = get_code_block_style_name(doc, config_manager)
                    
                    # 为代码块的每一行创建段落
                    code_lines = code_content.split('\n')
                    for line_idx, code_line in enumerate(code_lines):
                        if code_block_style_name:
                            # 使用 Code Block 段落样式
                            try:
                                new_p = doc.add_paragraph(style=code_block_style_name)
                                run = new_p.add_run(code_line if code_line else ' ')
                                logger.debug(f"代码块行使用样式: {code_block_style_name}")
                            except KeyError:
                                # 样式不存在，回退到底纹方式
                                new_p = doc.add_paragraph(style=base_style)
                                run = new_p.add_run(code_line if code_line else ' ')
                                _apply_code_block_style(new_p, run, config_manager)
                                logger.debug(f"Code Block 样式不存在，使用底纹兼容")
                        else:
                            # 使用底纹兼容方式
                            new_p = doc.add_paragraph(style=base_style)
                            run = new_p.add_run(code_line if code_line else ' ')
                            _apply_code_block_style(new_p, run, config_manager)
                        
                        # 如果在列表上下文中，设置缩进（与续行内容对齐）
                        if in_list_context:
                            from ..handlers.list_handler import BASE_INDENT, INDENT_INCREMENT
                            from docx.shared import Twips
                            indent_value = BASE_INDENT + INDENT_INCREMENT * list_level
                            new_p.paragraph_format.left_indent = Twips(indent_value)
                            if line_idx == 0:
                                logger.debug(f"代码块在列表上下文中: level={list_level}, indent={indent_value}")
                        
                        # 插入到正确位置（每行递增偏移量）
                        new_p._p.getparent().remove(new_p._p)
                        parent.insert(index + insert_offset, new_p._p)
                        insert_offset += 1
                    
                    logger.debug(f"插入代码块（{len(code_lines)}行，语言: {language or '无'}）")
                    continue
                
                # 处理引用块
                if item_type == 'quote':
                    quote_content = item.get('text', '')
                    quote_level = item.get('level', 1)
                    has_formula = item.get('has_formula', False)
                    
                    # 获取引用样式名称（从 style 模块获取）
                    quote_style_name = get_quote_style_name(doc, quote_level, config_manager)
                    
                    if quote_style_name:
                        try:
                            new_p = doc.add_paragraph(style=quote_style_name)
                            logger.debug(f"引用块使用样式: {quote_style_name}")
                        except KeyError:
                            new_p = doc.add_paragraph(style=base_style)
                            logger.debug(f"引用样式 {quote_style_name} 不存在，使用基础样式")
                    else:
                        new_p = doc.add_paragraph(style=base_style)
                    
                # 处理引用内容（支持公式和格式）
                    if has_formula and is_formula_supported():
                        formatting_mode = config_manager.get_md_to_docx_formatting_mode()
                        process_paragraph_formulas(new_p, quote_content, fonts, formatting_mode, doc, note_ctx=note_ctx)
                    else:
                        formatting_mode = config_manager.get_md_to_docx_formatting_mode()
                        add_formatted_text_to_paragraph(new_p, quote_content, fonts, formatting_mode, doc=doc)
                    
                    # 插入到正确位置
                    new_p._p.getparent().remove(new_p._p)
                    parent.insert(index + insert_offset, new_p._p)
                    insert_offset += 1
                    
                    logger.debug(f"插入引用块（{quote_level}级）: {quote_content[:30]}...")
                    continue
                
                # 处理分隔符/水平线（转换为分页符或分节符）
                if item_type == 'horizontal_rule':
                    hr_type = item.get('hr_type', 'dash')
                    attach_to = item.get('attach_to', 'none')
                    
                    # 检查是否启用分隔符转换
                    if config_manager.is_horizontal_rule_enabled():
                        target = config_manager.get_horizontal_rule_mapping(hr_type)
                        
                        if target == 'ignore':
                            logger.debug(f"分隔符 {hr_type} 配置为忽略，跳过")
                            continue
                        
                        # 如果 attach_to='previous'，将分隔符添加到前一个段落末尾（同一段落）
                        if attach_to == 'previous' and insert_offset > 0:
                            # 获取前一个段落
                            prev_para_elem = parent[index + insert_offset - 1]
                            
                            if target == 'page_break':
                                # 在前一个段落末尾添加分页符
                                append_page_break_to_paragraph(prev_para_elem)
                                logger.debug(f"插入分页符到前一段落末尾（同段，来自 {hr_type} 分隔符）")
                            elif target == 'section_break' or target.startswith('section_'):
                                # 分节符：在前一个段落的 pPr 中添加 sectPr
                                section_type = target.replace('section_', '') if target.startswith('section_') else 'next'
                                append_section_break_to_paragraph(prev_para_elem, doc, section_type)
                                logger.debug(f"插入分节符到前一段落（同段，来自 {hr_type} 分隔符）")
                            elif target.startswith('horizontal_rule_'):
                                # 分隔线：创建新段落应用 Horizontal Rule 样式
                                # attach_to='previous' 情况下，我们仍然创建独立段落
                                # 因为样式段落无法附加到其他段落
                                hr_num = target.split('_')[2]  # 提取数字 1/2/3
                                new_p = insert_horizontal_rule(doc, hr_num)
                                new_p._p.getparent().remove(new_p._p)
                                parent.insert(index + insert_offset, new_p._p)
                                insert_offset += 1
                                logger.debug(f"插入分隔线（Horizontal Rule {hr_num}，来自 {hr_type} 分隔符）")
                            else:
                                logger.warning(f"未知的分隔符转换目标: {target}")
                            # 不增加 insert_offset，因为没有创建新段落
                        else:
                            # attach_to='none' 或没有前一个段落，创建独立段落
                            if target == 'page_break':
                                # 插入分页符
                                new_p = insert_page_break(doc)
                                logger.debug(f"插入分页符（独立，来自 {hr_type} 分隔符）")
                            elif target == 'section_break':
                                # 插入分节符（统一使用下一页类型）
                                new_p = insert_section_break(doc, 'next')
                                logger.debug(f"插入分节符（独立，来自 {hr_type} 分隔符）")
                            elif target.startswith('section_'):
                                # 兼容旧配置：插入分节符
                                section_type = target.replace('section_', '')
                                new_p = insert_section_break(doc, section_type)
                                logger.debug(f"插入分节符（{section_type}，独立，来自 {hr_type} 分隔符）")
                            elif target.startswith('horizontal_rule_'):
                                # 插入分隔线（Horizontal Rule 1/2/3 样式）
                                hr_num = target.split('_')[2]  # 提取数字 1/2/3
                                new_p = insert_horizontal_rule(doc, hr_num)
                                logger.debug(f"插入分隔线（Horizontal Rule {hr_num}，独立，来自 {hr_type} 分隔符）")
                            else:
                                logger.warning(f"未知的分隔符转换目标: {target}")
                                continue
                            
                            # 插入到正确位置
                            new_p._p.getparent().remove(new_p._p)
                            parent.insert(index + insert_offset, new_p._p)
                            insert_offset += 1
                    else:
                        logger.debug("分隔符转换已禁用，跳过")
                    continue
                
                # 处理列表项
                if item_type == 'list_item':
                    list_content = item.get('text', '')
                    list_level = item.get('level', 0)
                    list_type = item.get('list_type', 'ordered')
                    has_formula = item.get('has_formula', False)
                    
                    # 获取或创建列表管理器
                    if not hasattr(doc, '_list_manager'):
                        doc._list_manager = ListFormatManager(doc)
                    
                    # 获取当前列表组的 numId
                    # 通过检查是否是新列表组来决定是否创建新定义
                    if not hasattr(doc, '_current_list_num_id') or item.get('_new_list_group', False):
                        # 分析当前列表组结构
                        # 简化处理：为每个列表项单独分析，或使用预处理的 numId
                        if 'numId' in item:
                            doc._current_list_num_id = item['numId']
                        else:
                            # 创建新的列表定义
                            level_types = {list_level: list_type}
                            doc._current_list_num_id = doc._list_manager.create_list_definition(level_types)
                    
                    num_id = doc._current_list_num_id
                    
                    # 创建列表段落，使用辅助函数查找列表块样式（支持国际化）
                    list_style_name = get_list_block_style_name(doc)
                    if list_style_name:
                        try:
                            new_p = doc.add_paragraph(style=list_style_name)
                            logger.debug(f"列表项使用样式: {list_style_name}")
                        except KeyError:
                            new_p = doc.add_paragraph(style=base_style)
                            logger.debug(f"列表块样式 {list_style_name} 不存在，使用基础样式")
                    else:
                        new_p = doc.add_paragraph(style=base_style)
                        logger.debug("未找到列表块样式，使用基础样式")
                    
                    # 添加内容
                    if has_formula and is_formula_supported():
                        formatting_mode = config_manager.get_md_to_docx_formatting_mode()
                        process_paragraph_formulas(new_p, list_content, fonts, formatting_mode, doc)
                        logger.debug(f"列表项包含公式，已使用公式处理器")
                    else:
                        formatting_mode = config_manager.get_md_to_docx_formatting_mode()
                        add_formatted_text_to_paragraph(new_p, list_content, fonts, formatting_mode, doc=doc)
                    
                    # 应用列表属性
                    if num_id:
                        apply_list_to_paragraph(new_p, num_id, list_level)
                        logger.debug(f"插入{list_type}列表项（{list_level}级）: {list_content[:30]}...")
                    else:
                        logger.warning(f"无法创建列表定义，列表项作为普通段落处理")
                    
                    # 插入到正确位置
                    new_p._p.getparent().remove(new_p._p)
                    parent.insert(index + insert_offset, new_p._p)
                    insert_offset += 1
                    continue
                
                # 处理列表续行内容
                if item_type == 'list_continuation':
                    continuation_content = item.get('text', '')
                    continuation_level = item.get('level', 0)
                    has_formula = item.get('has_formula', False)
                    
                    # 使用辅助函数查找列表块样式（支持国际化），不设置 numPr（无编号）
                    list_style_name = get_list_block_style_name(doc)
                    if list_style_name:
                        try:
                            new_p = doc.add_paragraph(style=list_style_name)
                            logger.debug(f"列表续行使用样式: {list_style_name}")
                        except KeyError:
                            new_p = doc.add_paragraph(style=base_style)
                            logger.debug(f"列表块样式 {list_style_name} 不存在，使用基础样式")
                    else:
                        new_p = doc.add_paragraph(style=base_style)
                        logger.debug("未找到列表块样式，使用基础样式")
                    
                    # 添加内容
                    if has_formula and is_formula_supported():
                        formatting_mode = config_manager.get_md_to_docx_formatting_mode()
                        process_paragraph_formulas(new_p, continuation_content, fonts, formatting_mode, doc)
                        logger.debug(f"列表续行包含公式，已使用公式处理器")
                    else:
                        formatting_mode = config_manager.get_md_to_docx_formatting_mode()
                        add_formatted_text_to_paragraph(new_p, continuation_content, fonts, formatting_mode, doc=doc)
                    
                    # 设置缩进（与对应级别的列表项内容对齐）
                    # 使用列表格式管理器中的缩进配置
                    from ..handlers.list_handler import BASE_INDENT, INDENT_INCREMENT
                    from docx.shared import Twips
                    indent_value = BASE_INDENT + INDENT_INCREMENT * continuation_level
                    new_p.paragraph_format.left_indent = Twips(indent_value)
                    
                    logger.debug(f"插入列表续行内容（{continuation_level}级）: {continuation_content[:30]}...")
                    
                    # 插入到正确位置
                    new_p._p.getparent().remove(new_p._p)
                    parent.insert(index + insert_offset, new_p._p)
                    insert_offset += 1
                    continue
                
                # 创建新段落
                if 1 <= level <= 9:
                    # 使用自定义方法查找并使用标题样式（支持中文/英文/未激活）
                    style_name = get_heading_style_name(doc, level)
                    
                    try:
                        if style_name:
                            new_p = doc.add_paragraph('', style=style_name)
                        else:
                            # 尝试使用默认的 add_heading，这通常需要 'Heading N' 样式存在
                            new_p = doc.add_heading('', level=level)
                    except KeyError as e:
                        # 捕获样式不存在的错误，提供更友好的提示
                        template_title = template_name or doc.core_properties.title or "当前使用的模板文件"
                        error_msg = (
                            f"模板样式缺失：无法找到“标题 {level}”样式。"
                            f"请打开模板文件“{template_title}”，在 Word 中手动点击一次所有级别的内置小标题样式（标题 1 ~ 标题 9）以激活它们，然后保存文件并重试。"
                        )
                        logger.error(error_msg)
                        raise ValueError(error_msg) from e
                    
                    # 检查标题文本是否包含公式
                    title_text = item['text']
                    has_formula = item.get('has_formula', False) or (is_formula_supported() and '$' in title_text)
                    
                    if has_formula and is_formula_supported():
                        # 使用公式处理器处理标题（清空段落后重新填充）
                        # 标题通常不应用正文字体，传入 None
                        heading_formatting_mode = config_manager.get_md_to_docx_heading_formatting_mode()
                        process_paragraph_formulas(new_p, title_text, None, heading_formatting_mode, doc)
                        logger.debug(f"标题包含公式，已使用公式处理器")
                    else:
                        # 获取小标题专用的格式处理模式配置
                        heading_formatting_mode = config_manager.get_md_to_docx_heading_formatting_mode()
                        
                        # 使用标题专用的格式解析函数
                        # 当 mode="apply" 时，支持覆盖标题样式的默认格式（实现部分加粗效果）
                        # 当 mode="remove" 时，清理标记，让Word标题样式格式自然生效
                        # 当 mode="keep" 时，保留标记原样
                        add_formatted_text_to_heading(new_p, title_text, heading_formatting_mode, doc=doc)
                        logger.debug(f"标题已使用格式解析添加，模式: {heading_formatting_mode}")
                else:  # 普通正文段落（包括列表项，保留原始格式）
                    # 检查正文是否包含公式
                    content_text = item['text']
                    has_formula = item.get('has_formula', False) or (is_formula_supported() and '$' in content_text)
                    
                    # 判断是否为纯公式块（只包含 $$...$$ 格式的公式，无其他文字）
                    is_pure_formula_block = has_formula and _is_pure_block_formula(content_text)
                    
                    if is_pure_formula_block:
                        # 纯公式块使用 Formula Block 段落样式
                        formula_block_style_name = get_formula_block_style_name(doc, config_manager)
                        if formula_block_style_name:
                            try:
                                new_p = doc.add_paragraph(style=formula_block_style_name)
                                logger.debug(f"纯公式块使用样式: {formula_block_style_name}")
                            except KeyError:
                                new_p = doc.add_paragraph(style=base_style)
                                logger.debug(f"Formula Block 样式不存在，使用基础样式")
                        else:
                            new_p = doc.add_paragraph(style=base_style)
                    else:
                        new_p = doc.add_paragraph(style=base_style)
                    
                    if has_formula and is_formula_supported():
                        # 使用公式处理器处理正文（支持公式与脚注/尾注共存）
                        formatting_mode = config_manager.get_md_to_docx_formatting_mode()
                        process_paragraph_formulas(new_p, content_text, fonts, formatting_mode, doc, note_ctx=note_ctx)
                        logger.debug(f"正文包含公式，已使用公式处理器")
                    else:
                        # 获取格式处理模式配置
                        formatting_mode = config_manager.get_md_to_docx_formatting_mode()
                        
                        # 检查是否包含脚注/尾注引用
                        if note_ctx and note_ctx.has_notes and '[^' in content_text:
                            # 使用脚注/尾注处理函数
                            process_text_with_note_references(content_text, new_p, fonts, formatting_mode, doc, note_ctx)
                            logger.debug(f"正文包含脚注/尾注引用，已处理")
                        else:
                            # 使用带格式解析的方式添加文本（支持粗体、斜体等Markdown格式）
                            add_formatted_text_to_paragraph(new_p, content_text, fonts, formatting_mode, doc=doc)
                            logger.debug(f"正文已使用格式解析添加，模式: {formatting_mode}")

                
                # 添加正文内容（如果有）
                # ========================================================================
                # 【Word/WPS 样式机制说明 - 混合段落处理】
                #
                # 这里处理"小标题+正文"的组合段落场景，如：
                # Markdown: "### 一、问题。这是问题的详细说明。"
                # 输出 Word: 整段使用 Heading 3 样式，但"这是问题的详细说明。"应显示为正文格式
                #
                # 为什么正文部分使用字体属性而不是字符样式？
                #
                # Word/WPS 的样式优先级机制：
                # 1. 段落样式会"压制"字符样式 - 当段落已应用 Heading 样式时，
                #    run 级别的其他字符样式无法覆盖段落样式的格式
                # 2. 字符样式只能在普通段落中生效
                #
                # 实验验证（2024-12）：
                # - 普通段落中，选中部分字符，应用小标题样式 → 生效
                #   （run.style 变为 "Heading X Char"，即段落样式的关联字符样式）
                # - 小标题段落中，选中部分字符，应用其他样式 → 不生效
                #   （样式被段落 Heading 样式压制，视觉上无变化）
                #
                # 因此，对于组合段落中的正文部分，我们只能通过直接设置字体属性
                # （字体、字号等，即 fonts 参数）来实现与 {{正文}} 占位符格式一致，
                # 而不能使用字符样式（如 "Normal Char"）。
                # ========================================================================
                if item_type == 'heading_with_content':
                    content_text = item['content']
                    has_content_formula = is_formula_supported() and '$' in content_text
                    
                    if has_content_formula and is_formula_supported():
                        # 正文内容包含公式，使用公式处理器
                        # 直接在同一段落中添加公式内容
                        formatting_mode = config_manager.get_md_to_docx_formatting_mode()
                        process_paragraph_formulas(new_p, content_text, fonts, formatting_mode, doc)
                        logger.debug(f"组合段落正文包含公式，已使用公式处理器")
                    else:
                        # 普通正文内容 - 使用格式解析（支持粗体、斜体等Markdown格式）
                        formatting_mode = config_manager.get_md_to_docx_formatting_mode()
                        add_formatted_text_to_paragraph(new_p, content_text, fonts, formatting_mode, doc=doc)
                        logger.debug(f"组合段落正文已使用格式解析添加，模式: {formatting_mode}")
                
                # 插入到正确位置（使用偏移量计数器）
                new_p._p.getparent().remove(new_p._p)
                parent.insert(index + insert_offset, new_p._p)
                insert_offset += 1
            
            # 移除原始占位符段落
            parent.remove(paragraph._element)
            return True  # 找到了正文占位符
    
    # 遍历所有段落后仍未找到任何正文占位符
    logger.warning("模板中未找到正文占位符（如 {{正文}} 或 {{body}}），Markdown 正文内容未被插入")
    return False

def remove_marked_elements(paragraphs_to_remove, rows_to_remove):
    """移除标记为需要删除的段落和表格行"""
    # 移除段落
    for para in paragraphs_to_remove:
        try_remove_element(para._element)
    
    # 移除表格行
    for row in rows_to_remove:
        try:
            parent = row._element.getparent()
            if parent is not None:
                parent.remove(row._element)
        except Exception as e:
            logger.warning(f"移除表格行失败: {str(e)}")


# ============================================================================
# 内部辅助函数
# ============================================================================

def _is_pure_block_formula(text: str) -> bool:
    """
    检查文本是否为纯公式块（只包含 $$...$$ 格式的公式，无其他文字）
    
    参数:
        text: 段落文本
        
    返回:
        bool: 是否为纯公式块
    """
    import re
    
    # 去除首尾空白
    stripped = text.strip()
    
    # 公式块模式：以 $$ 开头，以 $$ 结尾
    # 支持多行公式块和单行公式块
    block_formula_pattern = r'^\$\$[\s\S]*\$\$$'
    
    if re.match(block_formula_pattern, stripped):
        # 确保没有其他非空白内容
        # 移除所有公式块后检查是否只剩空白
        remaining = re.sub(r'\$\$[\s\S]*?\$\$', '', stripped)
        return remaining.strip() == ''
    
    return False


def _apply_code_block_style(paragraph, run, config_mgr):
    """
    应用代码块样式：等宽字体 + 灰色段落底纹（兼容方式）
    
    参数:
        paragraph: Word段落对象
        run: Word Run对象
        config_mgr: 配置管理器
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    # 获取代码配置
    code_font = config_mgr.get_code_font()
    code_bg_color = config_mgr.get_code_background_color()
    
    # 应用等宽字体到Run
    run.font.name = code_font
    # 确保rPr存在
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn('w:eastAsia'), code_font)
    
    # 应用段落底纹（整行灰色背景）
    pPr = paragraph._element.get_or_add_pPr()
    
    # 检查是否已有shd元素
    existing_shd = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
    if existing_shd is not None:
        pPr.remove(existing_shd)
    
    # 创建段落底纹
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), code_bg_color)
    pPr.append(shd)
    
    logger.debug(f"应用代码块样式: 字体={code_font}, 背景色={code_bg_color}")


def save_and_process_temp_file(doc, yaml_data, *, note_ctx=None):
    """
    保存临时文件并进行深度处理
    
    参数:
        doc: Document对象
        yaml_data: YAML元数据字典
        note_ctx: NoteContext对象（可选，用于脚注/尾注写入）
    
    返回:
        str: 临时文件路径
    """
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_docx.name)
    temp_docx.close()
    
    # 先写入脚注/尾注（必须在 xml_processor 之前执行）
    # xml_processor 会处理所有 XML 文件包括 footnotes.xml/endnotes.xml
    # 如果先执行 xml_processor，脚注内容可能被覆盖或丢失
    if note_ctx and note_ctx.has_notes:
        write_notes_to_docx(temp_docx.name, note_ctx)
        logger.info("脚注/尾注已写入文档")
    
    logger.debug("开始深度处理DOCX文件中的占位符")
    xml_processor.process_docx_file(temp_docx.name, yaml_data)
    
    return temp_docx.name


def _preprocess_list_items(doc, body_data):
    """
    列表预处理：对连续列表项分组并分配 numId
    
    根据 Markdown 语法规范：
    - 连续的列表项属于同一个列表
    - 非列表内容（普通段落、标题等）终止当前列表
    - 之后的列表项属于新列表（新的 numId）
    
    参数:
        doc: Document对象
        body_data: 正文段落数据列表（会被原地修改，添加 numId 和 _new_list_group 标记）
    """
    # 获取或创建列表管理器
    if not hasattr(doc, '_list_manager'):
        doc._list_manager = ListFormatManager(doc)
    
    list_manager = doc._list_manager
    
    # 分组连续的列表项
    list_groups = group_consecutive_list_items(body_data)
    
    if not list_groups:
        logger.debug("列表预处理：无列表项需要处理")
        return
    
    logger.info(f"列表预处理：检测到 {len(list_groups)} 个独立列表组")
    
    # 为每个分组创建列表定义
    for group_idx, group in enumerate(list_groups):
        # 分析整个组的结构（检测同级冲突）
        level_types = analyze_list_structure(group)
        
        # 创建列表定义，获取 numId
        num_id = list_manager.create_list_definition(level_types)
        
        logger.debug(f"列表组 {group_idx + 1}: {len(group)} 项, numId={num_id}, 级别类型={level_types}")
        
        # 将 numId 写入每个列表项
        for item_idx, item in enumerate(group):
            item['numId'] = num_id
            # 标记组的第一项，用于重置 _current_list_num_id
            item['_new_list_group'] = (item_idx == 0)
