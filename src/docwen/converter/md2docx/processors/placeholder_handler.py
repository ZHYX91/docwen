"""
占位符处理器模块
实现DOCX文档中占位符的识别、标记和处理功能
使用隐藏文本标记方案替代直接XML操作，提高健壮性

国际化说明：
占位符支持多语言，通过 i18n 模块获取所有语言版本：
- 中文环境模板: {{标题}}、{{正文}}
- 英文环境模板: {{title}}、{{body}}
检测时自动识别所有语言版本的占位符。
"""

import contextlib
import copy
import logging
import re
from pathlib import Path

from docx.shared import Pt, RGBColor

from docwen.config.config_manager import config_manager
from docwen.translation import t_all_locales
from docwen.utils import docx_utils
from docwen.utils.text_utils import format_display_value
from docwen.utils.validation_utils import is_value_empty

# 配置日志
logger = logging.getLogger(__name__)

# 特殊标记格式
SPECIAL_MARKER_PREFIX = "%%SPECIAL_PH:"
SPECIAL_MARKER_SUFFIX = "%%"


def _get_all_placeholder_variants(placeholder_key: str) -> list:
    """
    获取占位符的所有语言变体

    参数:
        placeholder_key: 占位符键名（如 "title", "body"）

    返回:
        list: 所有语言版本的占位符列表，如 ["{{标题}}", "{{title}}"]
    """
    translations = t_all_locales(f"placeholders.{placeholder_key}")
    variants = [f"{{{{{name}}}}}" for name in translations.values()]
    logger.debug("占位符 %s 的所有变体: %s", placeholder_key, variants)
    return variants


def _get_special_placeholders(extra_placeholders: set[str] | None = None) -> dict:
    """
    获取特殊处理的占位符（所有语言版本）

    返回:
        dict: {占位符键名: [所有语言变体列表], ...}
    """
    # 特殊占位符键名列表
    special_keys = ["body"]

    result = {}
    for key in special_keys:
        variants = _get_all_placeholder_variants(key)
        if variants:
            result[key] = variants

    for placeholder in extra_placeholders or set():
        result[placeholder] = [placeholder]

    return result


def _get_special_placeholder_list(extra_placeholders: set[str] | None = None) -> list:
    """
    获取所有特殊占位符的扁平列表（用于检测）

    返回:
        list: 所有特殊占位符的列表
    """
    special = _get_special_placeholders(extra_placeholders)
    result = []
    for variants in special.values():
        result.extend(variants)
    return result


# 占位符处理规则配置
#
# 配置结构说明：
# - 每个规则类型包含若干"组"
# - 每个"组"是一个字段列表（单字段或多字段）
# - 只有当组内所有字段都为空（不存在或值为空）时，才执行对应操作
#
# 支持5种处理方式：
# 1. 默认行为：不在任何规则中的占位符，直接替换为空字符串
# 2. delete_paragraph_if_empty: 删除占位符所在段落（仅在段落中生效）
# 3. delete_cell_if_empty: 清空占位符所在单元格（仅在表格中生效，检查单元格级别）
# 4. delete_row_if_empty: 删除占位符所在表格行（仅在表格中生效，检查整行级别）
# 5. delete_table_if_empty: 删除占位符所在整个表格（保留用于未来扩展）
#
# 示例：
# - ["密级和保密期限"]：单字段组，该字段为空时删除段落
# - ["印发机关", "印发日期"]：多字段组，两个字段都为空，且在同一行时，才删除表格行
PLACEHOLDER_RULES = {
    "delete_paragraph_if_empty": [],
    "delete_cell_if_empty": [],
    "delete_row_if_empty": [],
    "delete_table_if_empty": [],
}


def get_body_placeholder_variants() -> list:
    """
    获取所有有效的正文占位符变体

    供外部模块（如 docx_processor.py）调用，用于查找正文插入点。

    返回:
        list: 所有语言版本的正文占位符列表，如 ["{{body}}", "{{正文}}"]
    """
    return _get_all_placeholder_variants("body")


def _find_placeholder_type(ph_text: str, extra_placeholders: set[str] | None = None) -> str | None:
    """
    根据占位符文本查找其类型名称

    通过遍历所有特殊占位符的变体列表来查找，支持任意语言变体的匹配。

    参数:
        ph_text: 占位符文本，如 "{{正文}}" 或 "{{body}}"

    返回:
        str: 类型名称（如 "正文"），未找到返回 None
    """
    # 获取所有特殊占位符及其变体
    special_placeholders = _get_special_placeholders(extra_placeholders)

    # 遍历查找匹配的类型
    for key, variants in special_placeholders.items():
        if ph_text in variants:
            # 将 key 映射为中文类型名（用于标记）
            # body -> 正文
            type_mapping = {"body": "正文"}
            ph_type = type_mapping.get(key, key)
            if isinstance(ph_type, str) and ph_type.startswith("{{") and ph_type.endswith("}}"):
                ph_type = ph_type[2:-2].strip()
            logger.debug(f"占位符 '{ph_text}' 匹配类型: {ph_type} (key={key})")
            return ph_type

    return None


def mark_special_placeholders(doc, extra_placeholders: set[str] | None = None):
    """
    标记特殊占位符段落 - 使用隐藏文本标记方案
    在包含特殊占位符的段落末尾添加隐藏标记Run
    """
    logger.info("开始标记特殊占位符段落 (使用隐藏文本标记方案)...")

    special_texts = _get_special_placeholder_list(extra_placeholders)
    logger.debug(f"特殊占位符列表: {special_texts}")

    # 遍历所有段落
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if not text:
            continue

        # 检查段落是否包含任何特殊占位符
        for ph_text in special_texts:
            if ph_text in text:
                # 获取占位符类型名称（支持所有语言变体）
                ph_type = _find_placeholder_type(ph_text, extra_placeholders)

                if not ph_type:
                    logger.warning(f"未找到占位符类型: {ph_text}")
                    continue

                # 检查是否已标记
                if is_special_marked(paragraph):
                    logger.debug(f"段落已标记，跳过: {text[:30]}...")
                    continue

                # 添加隐藏标记Run
                marker_text = f"{SPECIAL_MARKER_PREFIX}{ph_type}{SPECIAL_MARKER_SUFFIX}"
                marker_run = paragraph.add_run()
                marker_run.text = marker_text

                # 设置隐藏样式（白色字体，极小字号）
                font = marker_run.font
                font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # 白色
                font.size = Pt(1)  # 极小字号

                logger.debug(f"添加隐藏标记: {marker_text} -> 段落: {text[:30]}...")
                break  # 一个段落只标记一次

    logger.info("特殊占位符标记完成")


def is_special_marked(paragraph):
    """
    检查段落是否被标记为特殊占位符
    通过检查段落末尾的隐藏标记Run

    参数:
        paragraph: 段落对象

    返回:
        bool: 如果是特殊标记段落返回True，否则False
    """
    # 检查段落是否有Run
    if not paragraph.runs:
        return False

    # 检查最后一个Run是否是隐藏标记
    last_run = paragraph.runs[-1]
    last_run_text = last_run.text.strip()

    # 检查是否符合标记格式
    is_marked = last_run_text.startswith(SPECIAL_MARKER_PREFIX) and last_run_text.endswith(SPECIAL_MARKER_SUFFIX)

    if is_marked:
        logger.debug(f"检测到隐藏标记: {last_run_text}")

    return is_marked


def get_special_mark_type(paragraph):
    """
    获取特殊标记的类型

    参数:
        paragraph: 段落对象

    返回:
        str: 标记类型名称，如果不是特殊标记返回None
    """
    if not is_special_marked(paragraph):
        return None

    # 提取标记类型
    last_run_text = paragraph.runs[-1].text.strip()
    mark_type = last_run_text.replace(SPECIAL_MARKER_PREFIX, "").replace(SPECIAL_MARKER_SUFFIX, "")

    logger.debug(f"提取标记类型: {mark_type}")
    return mark_type


def remove_special_mark(paragraph):
    """
    移除特殊标记
    通过删除标记Run实现

    参数:
        paragraph: 段落对象
    """
    if not is_special_marked(paragraph):
        return

    try:
        # 获取标记Run
        last_run = paragraph.runs[-1]

        # 检查是否是标记Run
        last_run_text = last_run.text.strip()
        if last_run_text.startswith(SPECIAL_MARKER_PREFIX) and last_run_text.endswith(SPECIAL_MARKER_SUFFIX):
            # 删除标记Run
            p = last_run._element.getparent()
            p.remove(last_run._element)
            logger.debug("已移除隐藏标记")
    except Exception as e:
        logger.warning(f"移除标记失败: {e!s}")


def extract_placeholder_keys(text: str):
    """
    从文本中提取所有占位符键名

    支持多语言占位符，包括含空格的越南语（如 {{Tiêu đề}}）

    返回: 包含所有占位符键名的列表
    """
    # 使用正则表达式匹配所有 {{key}} 格式的占位符
    # 使用 [^{}]+ 匹配花括号内的任意非花括号字符，支持空格和Unicode
    pattern = r"\{\{([^{}]+)\}\}"
    matches = re.findall(pattern, text)
    # 去除前后空白并去重
    return list({m.strip() for m in matches})


def get_placeholder_rules(key: str):
    """
    查找占位符所属的所有规则和组

    工作原理：
    1. 遍历PLACEHOLDER_RULES中的所有规则类型
    2. 对每个规则类型，遍历其包含的所有组
    3. 如果占位符键在某个组中，将该规则类型和组添加到结果列表
    4. 返回所有匹配的规则列表

    参数:
        key: 占位符键名（如"密级和保密期限"）

    返回:
        list: [(rule_type, group), ...] 规则列表
              每个元组包含：
              - rule_type: 规则类型字符串
                * 'delete_paragraph_if_empty': 删除段落（仅在段落中生效）
                * 'delete_cell_if_empty': 清空单元格（仅在表格中生效）
                * 'delete_row_if_empty': 删除表格行（仅在表格中生效）
                * 'delete_table_if_empty': 删除整个表格
              - group: 包含该键的组（字段列表）

              如果不属于任何规则，返回空列表，表示使用默认行为（替换为空字符串）
    """
    rules = []
    effective_rules = _get_effective_placeholder_rules()

    for rule_type, groups in effective_rules.items():
        for group in groups:
            if key in group:
                rules.append((rule_type, group))
                logger.debug(f"占位符 '{key}' 属于规则 '{rule_type}', 组: {group}")

    # 如果不属于任何规则，返回空列表
    if not rules:
        logger.debug(f"占位符 '{key}' 不属于任何特殊规则，将替换为空字符串")

    return rules


def _get_effective_placeholder_rules() -> dict:
    from ..field_registry import get_merged_placeholder_rules

    merged = get_merged_placeholder_rules()
    result = {rule_type: list(groups) for rule_type, groups in PLACEHOLDER_RULES.items()}
    for rule_type, groups in merged.items():
        result.setdefault(rule_type, [])
        result[rule_type].extend(groups)
    return result


def check_group_all_empty(group: list, yaml_data: dict):
    """
    检查组内所有字段是否都为空

    工作原理：
    1. 遍历组内的每个字段
    2. 如果字段在YAML数据中存在：
       - 获取其值并检查是否为空（使用is_value_empty函数）
       - 如果不为空，立即返回False，表示组不是全空的
    3. 如果字段不存在，视为空值
    4. 只有当组内所有字段都为空时，才返回True

    判断"空"的标准（由is_value_empty函数实现）：
    - None值
    - 空字符串
    - 空列表
    - 仅包含空白字符的字符串

    参数:
        group: 字段名列表，如["印发机关", "印发日期"]
        yaml_data: YAML数据字典

    返回:
        bool: 如果组内所有字段都为空（不存在或值为空）则返回True

    示例:
        组["印发机关", "印发日期"]：
        - 如果两个字段都为空 → True
        - 如果任何一个字段有值 → False
    """
    if not group:
        return False

    # 遍历组内所有字段
    for field in group:
        # 检查字段是否存在
        if field in yaml_data:
            value = yaml_data[field]
            # 如果字段不为空，组就不是全空的
            if not is_value_empty(value):
                logger.debug(f"组中字段 '{field}' 不为空: {value}")
                return False
        else:
            # 字段不存在，视为空
            logger.debug(f"组中字段 '{field}' 不存在，视为空")

    # 所有字段都为空
    logger.debug(f"组 {group} 中所有字段都为空")
    return True


def process_paragraph_placeholders(paragraph, yaml_data, paragraphs_to_remove):
    """
    处理单个段落中的占位符（支持组处理）

    处理流程：
    1. 提取段落中所有占位符键名
    2. 第一遍遍历：检查是否需要删除段落
       - 查找每个占位符所属的规则和组
       - 使用processed_groups避免重复检查同一组
       - 如果是delete_paragraph_if_empty规则且组内全空，标记删除段落
    3. 如果需要删除段落，加入删除列表并返回
    4. 第二遍遍历：正常替换占位符
       - 从YAML数据获取值或使用空字符串
       - 调用replace_placeholder_in_paragraph进行替换

    示例：
    - 段落包含{{密级和保密期限}}，该字段为空 → 删除段落
    - 段落包含{{标题}}，有值 → 替换为实际值
    - 段落包含{{不存在的字段}} → 替换为空字符串

    参数:
        paragraph: 段落对象
        yaml_data: YAML数据字典
        paragraphs_to_remove: 需要删除的段落列表（输出参数）
    """
    text = paragraph.text
    logger.debug(f"处理段落: {text[:50]}...")

    # 步骤1：提取段落中的所有占位符键
    placeholder_keys = extract_placeholder_keys(text)
    if not placeholder_keys:
        logger.debug("段落中无占位符，跳过处理")
        return

    logger.debug(f"找到占位符键: {placeholder_keys}")

    # 用于跟踪已处理的组（避免重复检查同一组）
    # 例如：如果段落包含["印发机关", "印发日期"]，只需检查一次组
    processed_groups = set()
    paragraph_should_remove = False

    # 步骤2：第一遍遍历 - 检查是否需要删除段落
    for key in placeholder_keys:
        # 查找占位符所属的所有规则
        rules = get_placeholder_rules(key)

        # 只处理 delete_paragraph_if_empty 规则
        for rule_type, group in rules:
            if rule_type != "delete_paragraph_if_empty":
                continue  # 跳过非段落规则

            # 跳过已处理的组（避免重复检查）
            group_tuple = tuple(group)
            if group_tuple in processed_groups:
                logger.debug(f"组 {group} 已处理，跳过")
                continue
            processed_groups.add(group_tuple)

            # 【关键】检查组内所有字段是否都在当前段落中
            # 只有组内所有字段都在同一段落，才判断是否删除
            if all(field in placeholder_keys for field in group):
                # 组内所有字段都在当前段落，检查是否都为空
                if check_group_all_empty(group, yaml_data):
                    logger.info(f"组 {group} 所有字段都在段落中且都为空，标记删除段落")
                    paragraph_should_remove = True
                    break  # 找到需要删除的原因，不再继续检查
            else:
                # 组内字段不全在当前段落，不删除
                logger.debug(f"组 {group} 的字段不全在当前段落中，跳过删除检查")

        if paragraph_should_remove:
            break  # 外层循环也要退出

    # 步骤3：如果需要删除段落，添加到移除列表并返回
    if paragraph_should_remove:
        paragraphs_to_remove.append(paragraph)
        logger.debug("段落已加入删除列表")
        return

    # 步骤4：第二遍遍历 - 正常替换占位符
    # 获取列表拼接符配置
    list_separator = config_manager.get_yaml_list_separator()

    for key in placeholder_keys:
        placeholder = f"{{{{{key}}}}}"

        # 获取值并格式化（传入列表拼接符）
        if key in yaml_data:
            value = yaml_data[key]
            display_value = format_display_value(value, list_separator)
            logger.debug(f"字段 '{key}' 的值: {display_value}")
        else:
            # 键不存在，替换为空字符串（默认行为）
            display_value = ""
            logger.debug(f"键 '{key}' 不存在，替换为空字符串")

        # 执行替换
        if placeholder in text:
            replace_placeholder_in_paragraph(paragraph, placeholder, display_value)
            logger.debug(f"成功替换占位符: {placeholder} -> {display_value}")


def process_table_cell_placeholders(paragraph, yaml_data, row, rows_to_remove, row_placeholder_keys):
    """
    处理表格单元格中的占位符（支持组处理）

    处理流程：
    1. 提取单元格段落中所有占位符键名
    2. 第一遍遍历：检查是否需要删除行或清空单元格
       - 查找每个占位符所属的规则和组
       - 使用processed_groups避免重复检查同一组
       - delete_row_if_empty规则：检查组内所有字段是否都在整行中，且都为空时删除整行（优先级高）
       - delete_paragraph_if_empty规则：组内全空时清空单元格内容
    3. 如果需要删除行，加入删除列表并返回True
    4. 如果需要清空单元格，清空所有内容并返回False
    5. 第二遍遍历：正常替换占位符
       - 从YAML数据获取值或使用空字符串
       - 调用replace_placeholder_in_paragraph进行替换

    注意事项：
    - delete_row_if_empty优先级最高，一旦触发立即返回
    - delete_row_if_empty检查组内所有字段是否都在整行中（跨单元格）
    - delete_paragraph_if_empty在表格中表现为清空单元格，而非删除段落
    - 如果表格所有行都被删除，在docx_processor.py中会自动删除整个表格

    示例：
    - 行包含{{抄送机关}}，该字段为空 → 删除整行
    - 行包含{{印发机关}}（单元格1）和{{印发日期}}（单元格2），两者都为空 → 删除整行
    - 行只有{{印发机关}}，印发日期不在该行 → 不删除
    - 单元格包含{{密级和保密期限}}，该字段为空 → 清空单元格
    - 单元格包含{{发文字号}}，有值 → 替换为实际值

    参数:
        paragraph: 表格单元格中的段落对象
        yaml_data: YAML数据字典
        row: 当前表格行对象
        rows_to_remove: 需要删除的行列表（输出参数）
        row_placeholder_keys: 整行的所有占位符键集合（可选，用于行级检查）

    返回:
        bool: 是否需要删除整行（True=删除行，False=不删除行）
    """
    text = paragraph.text
    logger.debug(f"处理表格单元格: {text[:50]}...")

    # 步骤1：提取段落中的所有占位符键
    placeholder_keys = extract_placeholder_keys(text)
    body_keys = {variant[2:-2].strip() for variant in get_body_placeholder_variants()}
    placeholder_keys = [k for k in placeholder_keys if k not in body_keys]
    if not placeholder_keys:
        logger.debug("单元格中无占位符，跳过处理")
        return False

    logger.debug(f"找到占位符键: {placeholder_keys}")

    # 用于跟踪已处理的组（避免重复检查同一组）
    processed_groups = set()
    cell_should_clear = False
    row_should_remove = False

    # 步骤2：第一遍遍历 - 检查是否需要删除行或清空单元格
    for key in placeholder_keys:
        # 查找占位符所属的所有规则
        rules = get_placeholder_rules(key)

        # 只处理表格相关规则：delete_row_if_empty 和 delete_cell_if_empty
        for rule_type, group in rules:
            # 跳过非表格规则
            if rule_type not in ("delete_row_if_empty", "delete_cell_if_empty"):
                continue

            # 跳过已处理的组（避免重复检查）
            group_tuple = tuple(group)
            if group_tuple in processed_groups:
                logger.debug(f"组 {group} 已处理，跳过")
                continue
            processed_groups.add(group_tuple)

            # 处理delete_row_if_empty规则（优先级高）
            if rule_type == "delete_row_if_empty":
                if all(field in row_placeholder_keys for field in group):
                    if check_group_all_empty(group, yaml_data):
                        logger.info(f"组 {group} 所有字段都在整行中且都为空，标记删除行")
                        row_should_remove = True
                        break
                else:
                    logger.debug(f"组 {group} 的字段不全在整行中，跳过删除检查")

            # 处理delete_cell_if_empty规则（优先级中）
            elif rule_type == "delete_cell_if_empty":
                # 【关键】检查组内所有字段是否都在当前单元格中
                # 只有组内所有字段都在同一单元格，才判断是否清空
                if all(field in placeholder_keys for field in group):
                    # 组内所有字段都在当前单元格，检查是否都为空
                    if check_group_all_empty(group, yaml_data):
                        logger.info(f"组 {group} 所有字段都在单元格中且都为空，标记清空单元格")
                        cell_should_clear = True
                else:
                    # 组内字段不全在当前单元格，不清空
                    logger.debug(f"组 {group} 的字段不全在当前单元格中，跳过清空检查")

        if row_should_remove:
            break  # 外层循环也要退出

    # 步骤3：如果行需要删除，添加到移除列表并返回True
    if row_should_remove:
        rows_to_remove.append(row)
        logger.debug("行已加入删除列表")
        return True

    # 步骤4：如果需要清空单元格，清空所有内容
    if cell_should_clear:
        # 获取单元格元素（tc元素）
        cell_element = paragraph._element.getparent()
        # 遍历单元格中的所有段落
        for para_element in cell_element.findall(".//w:p", docx_utils.NAMESPACES):
            # 遍历段落中的所有run
            for run_element in para_element.findall(".//w:r", docx_utils.NAMESPACES):
                # 遍历run中的所有文本元素
                for text_element in run_element.findall(".//w:t", docx_utils.NAMESPACES):
                    text_element.text = ""
        logger.debug("已清空单元格内容")
        return False

    # 步骤5：第二遍遍历 - 正常替换占位符
    # 获取列表拼接符配置
    list_separator = config_manager.get_yaml_list_separator()

    for key in placeholder_keys:
        placeholder = f"{{{{{key}}}}}"

        # 获取值并格式化（传入列表拼接符）
        if key in yaml_data:
            value = yaml_data[key]
            display_value = format_display_value(value, list_separator)
            logger.debug(f"字段 '{key}' 的值: {display_value}")
        else:
            # 键不存在，替换为空字符串（默认行为）
            display_value = ""
            logger.debug(f"键 '{key}' 不存在，替换为空字符串")

        # 执行替换
        if placeholder in text:
            replace_placeholder_in_paragraph(paragraph, placeholder, display_value)
            logger.debug(f"成功替换占位符: {placeholder} -> {display_value}")

    return False


def replace_placeholder_in_paragraph(paragraph, placeholder, value):
    """
    精确替换段落中的占位符（合并相同样式的run后处理）

    参数:
        paragraph: 段落对象
        placeholder: 占位符文本
        value: 要替换的值
    """
    logger.debug(f"开始处理段落占位符: {placeholder}")
    logger.debug(f"段落原始文本: {paragraph.text}")

    # 情况1：整个段落就是占位符（最高优先级）
    if paragraph.text and paragraph.text.strip() == placeholder.strip():
        logger.debug("情况1：整个段落就是占位符")
        paragraph.text = value
        return

    # 步骤2：合并相同样式的连续run
    logger.debug("步骤2：合并相同样式的连续run")
    merge_similar_runs(paragraph)

    # 情况3：在单个run内查找占位符
    logger.debug("情况3：在单个run内查找占位符")
    for i, run in enumerate(paragraph.runs):
        if placeholder in run.text:
            logger.debug(f"在run[{i}]中找到占位符: {run.text}")
            run.text = run.text.replace(placeholder, value)
            return

    # 情况4：处理跨多个run的占位符
    logger.debug("情况4：处理跨多个run的占位符")
    full_text = "".join(run.text for run in paragraph.runs)
    if placeholder in full_text:
        logger.debug(f"在合并文本中找到跨run占位符: {full_text}")

        # 保存第一个run的样式（用于保持格式）
        first_run = paragraph.runs[0] if paragraph.runs else None

        # 清空所有现有run
        for run in paragraph.runs:
            run.text = ""

        # 创建新run并应用原样式
        new_text = full_text.replace(placeholder, value)
        new_run = paragraph.add_run(new_text)

        # 应用原始样式
        if first_run and first_run._element.rPr is not None:
            logger.debug("应用原始样式到新run")
            docx_utils.apply_run_style(new_run, first_run)

        logger.debug(f"跨run替换完成: {new_text}")
        return

    logger.warning(f"在段落中未找到占位符: {placeholder}")


def merge_similar_runs(paragraph):
    """
    合并相同样式的连续run

    参数:
        paragraph: 段落对象

    返回:
        list: 合并后的run列表
    """
    if not paragraph.runs or len(paragraph.runs) < 2:
        logger.debug("无需合并：少于2个run")
        return paragraph.runs

    logger.debug(f"原始run数量: {len(paragraph.runs)}")

    # 创建一个新的run列表用于合并
    merged_runs = []
    current_run = None

    for run in paragraph.runs:
        # 如果是第一个run，直接开始新组
        if current_run is None:
            current_run = {"text": run.text, "rPr": run._element.rPr}
            continue

        # 检查当前run是否与前一个样式相同
        if docx_utils.is_rPr_equal(current_run["rPr"], run._element.rPr):
            logger.debug(f"合并相同样式run: '{current_run['text']}' + '{run.text}'")
            current_run["text"] += run.text
        else:
            # 保存当前组并开始新组
            merged_runs.append(current_run)
            current_run = {"text": run.text, "rPr": run._element.rPr}

    # 添加最后一组
    if current_run is not None:
        merged_runs.append(current_run)

    logger.debug(f"合并后run数量: {len(merged_runs)}")

    # 清空原始run
    for run in paragraph.runs[:]:
        p = run._element.getparent()
        p.remove(run._element)

    # 创建新的合并后的run
    new_runs = []
    for merged in merged_runs:
        new_run = paragraph.add_run(merged["text"])
        if merged["rPr"] is not None:
            docx_utils.apply_run_style(new_run, merged["rPr"])
        new_runs.append(new_run)

    return new_runs


def find_run_at_position(paragraph, position):
    """
    在段落中查找包含指定文本位置的run

    参数:
        paragraph: 段落对象
        position: 文本位置

    返回:
        Run对象 或 None
    """
    current_pos = 0
    for run in paragraph.runs:
        run_length = len(run.text)
        if current_pos <= position < current_pos + run_length:
            return run
        current_pos += run_length
    return paragraph.runs[0] if paragraph.runs else None


def try_remove_element(element):
    """安全移除元素"""
    try:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)
    except Exception as e:
        logger.warning(f"移除元素失败: {e!s}")


def _process_paragraph_images(paragraph, context="document", *, doc=None, col_count: int | None = None):
    """
    处理单个段落中的图片占位符（核心处理逻辑）

    参数:
        paragraph: 段落对象
        context: 上下文标识 ('document' 或 'table')

    返回:
        int: 成功插入的图片数量
    """
    from docx.shared import Emu

    image_placeholder_pattern = re.compile(r"\{\{IMAGE:.*?\}\}")

    def parse_image_placeholder(placeholder: str) -> tuple[str, int | None, int | None]:
        inner = placeholder[len("{{IMAGE:") : -2]
        from docwen.converter.shared.image_placeholder import parse_image_payload

        return parse_image_payload(inner)

    def get_page_usable_width_emu(doc) -> int:
        section = doc.sections[0]
        usable = section.page_width - section.left_margin - section.right_margin
        return int(usable)

    def get_table_cell_usable_width_emu(doc, col_count: int, cell_padding_pt: float = 5.4) -> int:
        if col_count <= 0:
            return get_page_usable_width_emu(doc)
        page_usable = get_page_usable_width_emu(doc)
        cell_width = page_usable // col_count
        cell_padding_emu = int(Pt(cell_padding_pt))
        return max(0, cell_width - (cell_padding_emu * 2))

    def calculate_image_size(
        image_path: str, md_width: int | None, md_height: int | None, usable_width_emu: int | None
    ):
        from PIL import Image

        try:
            with Image.open(image_path) as img:
                original_width_px, original_height_px = img.size
                dpi = img.info.get("dpi")
                dpi_x = None
                if isinstance(dpi, tuple) and dpi and isinstance(dpi[0], (int, float)) and dpi[0] > 0:
                    dpi_x = float(dpi[0])
                if dpi_x is None:
                    dpi_x = 96.0
        except Exception:
            return None, None

        px_to_emu = 914400.0 / dpi_x
        original_width_emu = int(original_width_px * px_to_emu)
        original_height_emu = int(original_height_px * px_to_emu)

        target_width_emu = int(md_width * px_to_emu) if md_width is not None else original_width_emu

        if usable_width_emu is not None and usable_width_emu > 0 and target_width_emu > usable_width_emu:
            target_width_emu = usable_width_emu

        if md_height is not None:
            target_height_emu = int(md_height * px_to_emu)
        else:
            if original_width_emu > 0:
                target_height_emu = int(target_width_emu * (original_height_emu / original_width_emu))
            else:
                target_height_emu = original_height_emu

        return Emu(target_width_emu), Emu(target_height_emu)

    def copy_run_format(src_run, dst_run):
        src_rpr = src_run._element.rPr
        if src_rpr is None:
            return
        dst_rpr = dst_run._element.get_or_add_rPr()
        for child in list(dst_rpr):
            dst_rpr.remove(child)
        for child in src_rpr.getchildren():
            dst_rpr.append(copy.deepcopy(child))

    text = paragraph.text or ""
    placeholders = list(image_placeholder_pattern.finditer(text))
    if not placeholders:
        return 0

    inserted_count = 0
    usable_width_emu = None
    if context == "document" and doc is not None:
        try:
            usable_width_emu = get_page_usable_width_emu(doc)
        except Exception:
            usable_width_emu = None
    if context == "table" and doc is not None and isinstance(col_count, int) and col_count > 0:
        usable_width_emu = get_table_cell_usable_width_emu(doc, col_count)

    image_only = image_placeholder_pattern.sub("", text).strip() == ""

    supported_formats = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".emf", ".wmf"}

    for run in list(paragraph.runs):
        while True:
            match = image_placeholder_pattern.search(run.text or "")
            if not match:
                break

            placeholder_text = match.group(0)
            before = (run.text or "")[: match.start()]
            after = (run.text or "")[match.end() :]

            run.text = before

            image_path, md_width, md_height = parse_image_placeholder(placeholder_text)
            if not Path(image_path).exists():
                logger.warning("图片文件不存在，跳过: %s", image_path)
                if after:
                    new_run = paragraph.add_run(after)
                    copy_run_format(run, new_run)
                    new_run._element.getparent().remove(new_run._element)
                    run._element.addnext(new_run._element)
                    run = new_run
                    continue
                break

            ext = Path(image_path).suffix.lower()
            if ext not in supported_formats:
                logger.warning("不支持的图片格式: %s，跳过: %s", ext, image_path)
                if after:
                    new_run = paragraph.add_run(after)
                    copy_run_format(run, new_run)
                    new_run._element.getparent().remove(new_run._element)
                    run._element.addnext(new_run._element)
                    run = new_run
                    continue
                break

            try:
                width, height = calculate_image_size(
                    image_path,
                    md_width=md_width,
                    md_height=md_height,
                    usable_width_emu=usable_width_emu,
                )

                picture_run = paragraph.add_run()
                picture_run._element.getparent().remove(picture_run._element)
                run._element.addnext(picture_run._element)
                if width is not None and height is not None:
                    picture_run.add_picture(image_path, width=width, height=height)
                else:
                    picture_run.add_picture(image_path)

                inserted_count += 1
                logger.info(
                    "成功插入图片 [%s]: %s | 尺寸: %.2fx%.2f 英寸",
                    context,
                    Path(image_path).name,
                    width.inches if width is not None else 0.0,
                    height.inches if height is not None else 0.0,
                )

                if after:
                    after_run = paragraph.add_run(after)
                    copy_run_format(run, after_run)
                    after_run._element.getparent().remove(after_run._element)
                    picture_run._element.addnext(after_run._element)
                    run = after_run
                else:
                    run = picture_run
                    break

            except Exception as e:
                logger.error("插入图片失败: %s | 错误: %s", image_path, str(e))
                run.text = (run.text or "") + placeholder_text + after
                break

    if inserted_count > 0 and image_only:
        try:
            from ..style.helper import get_image_paragraph_style_name

            style_name = get_image_paragraph_style_name()
            if style_name:
                paragraph.style = style_name
        except Exception:
            pass

        with contextlib.suppress(Exception):
            paragraph.paragraph_format.first_line_indent = None

    return inserted_count


def process_image_placeholders(doc):
    """
    处理文档中的图片占位符 {{IMAGE:路径}}

    处理流程：
    1. 遍历所有文档段落和表格单元格段落
    2. 查找包含 {{IMAGE:路径}} 的段落
    3. 提取图片路径
    4. 检查图片文件是否存在
    5. 在段落中插入图片并设置上下环绕
    6. 删除占位符文本

    参数:
        doc: Document对象

    返回:
        int: 成功插入的图片数量
    """
    logger.info("开始处理图片占位符...")

    total_inserted = 0

    # 1. 处理文档段落中的图片
    logger.debug("处理文档段落中的图片...")
    for paragraph in doc.paragraphs:
        count = _process_paragraph_images(paragraph, "document", doc=doc)
        total_inserted += count

    # 2. 处理表格单元格中的图片
    logger.debug("处理表格单元格中的图片...")
    for table in doc.tables:
        col_count = len(table.columns)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    count = _process_paragraph_images(paragraph, "table", doc=doc, col_count=col_count)
                    total_inserted += count

    logger.info(f"图片占位符处理完成 | 成功插入: {total_inserted} 个（文档+表格）")
    return total_inserted
