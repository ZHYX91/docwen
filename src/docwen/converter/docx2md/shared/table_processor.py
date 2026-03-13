"""
表格处理器模块
处理DOCX文档中的表格，提取内容并转换为段落格式
"""

import logging
from pathlib import Path

from docx.oxml.ns import qn

from docwen.utils.docx_utils import (
    NAMESPACES,
    extract_format_from_paragraph,
)

from .formula_processor import has_formulas_in_paragraph, is_formula_supported, process_paragraph_with_formulas
from .markdown_converter import (
    convert_paragraph_to_markdown_with_styles,  # 带样式检测的段落转换
)

# 配置日志
logger = logging.getLogger(__name__)


def extract_tables_from_document(doc):
    """
    从文档中提取所有表格内容

    参数:
        doc: Document对象

    返回:
        list: 表格内容列表，每个元素为 (位置索引, 段落对象, 格式信息)
    """
    logger.info("开始提取文档中的表格内容")
    table_contents = []

    try:
        # 遍历所有表格
        for table_idx, table in enumerate(doc.tables):
            logger.debug(f"处理表格 {table_idx + 1}, 行数: {len(table.rows)}, 列数: {len(table.columns)}")

            # 提取表格内容
            table_data = extract_table_content(table, table_idx)
            if table_data:
                table_contents.extend(table_data)

        logger.info(f"共提取到 {len(table_contents)} 个表格内容段落")
        return table_contents

    except Exception as e:
        logger.error(f"提取表格内容失败: {e!s}", exc_info=True)
        return []


def extract_table_content(table, table_index):
    """
    提取单个表格的内容

    参数:
        table: Table对象
        table_index: 表格索引

    返回:
        list: 表格内容数据列表
    """
    table_data = []

    try:
        # 获取表格在文档中的位置
        position_info = get_table_position(table, table_index)

        # 遍历表格的所有行和单元格
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                # 获取单元格的XML元素
                tc = cell._tc

                # 检查垂直合并：跳过被合并的单元格
                vMerge = tc.find(".//w:vMerge", namespaces=NAMESPACES)
                if vMerge is not None and vMerge.get(qn("w:val")) == "continue":
                    # 跳过被垂直合并的单元格
                    logger.debug(f"跳过被垂直合并的单元格: 表{table_index + 1}行{row_idx + 1}列{cell_idx + 1}")
                    continue

                # 检查水平合并：跳过被跨越的单元格
                gridSpan = tc.find(".//w:gridSpan", namespaces=NAMESPACES)
                if (
                    gridSpan is not None
                    and cell_idx > 0
                    and row.cells[cell_idx - 1]._tc.find(".//w:gridSpan", namespaces=NAMESPACES) is not None
                ):
                    # 当前单元格是被水平跨越的单元格，跳过
                    logger.debug(f"跳过被水平跨越的单元格: 表{table_index + 1}行{row_idx + 1}列{cell_idx + 1}")
                    continue

                # 处理单元格中的所有段落，合并为一个字符串
                para_texts = []
                first_para = None
                first_para_fonts = None
                all_runs_data = []

                for _para_idx, paragraph in enumerate(cell.paragraphs):
                    # 检查是否包含公式
                    has_formula = is_formula_supported() and has_formulas_in_paragraph(paragraph)

                    if has_formula:
                        # 提取表格内公式时默认不保留格式（简化输出）
                        formula_md = process_paragraph_with_formulas(paragraph, preserve_formatting=False)
                        if formula_md:
                            # 添加公式内容到段落文本列表
                            para_texts.append(formula_md)
                            logger.debug(f"表格单元格包含公式，已转换: {formula_md[:50]}...")

                            if first_para is None:
                                first_para = paragraph
                                first_para_fonts = extract_format_from_paragraph(paragraph)
                            continue
                        else:
                            # 如果公式转换失败，至少保留段落文本
                            logger.warning("表格单元格公式转换失败，尝试提取文本")

                    para_text = paragraph.text.strip()
                    if para_text:
                        # 替换段落内部的换行符为 <br>
                        para_text = para_text.replace("\r\n", "<br>").replace("\n", "<br>").replace("\r", "<br>")
                        para_texts.append(para_text)

                        # 保存第一个段落的格式信息
                        if first_para is None:
                            first_para = paragraph
                            first_para_fonts = extract_format_from_paragraph(paragraph)

                        # 收集所有run的数据
                        for run in paragraph.runs:
                            if run.text.strip():
                                all_runs_data.append(
                                    {
                                        "text": run.text,
                                        "fonts": {},  # 可以扩展为提取run格式
                                    }
                                )

                # 如果单元格有内容，合并所有段落并创建一条记录
                if para_texts:
                    # 用 <br> 连接多个段落
                    cell_text = "<br>".join(para_texts)

                    # 转义管道符，避免破坏Markdown表格结构
                    cell_text = cell_text.replace("|", "\\|")

                    # 构建表格数据（只为整个单元格创建一条记录）
                    cell_data = {
                        "text": cell_text,
                        "fonts": first_para_fonts if first_para_fonts else {},
                        "runs": all_runs_data,
                        "position": position_info,
                        "table_index": table_index,
                        "row_index": row_idx,
                        "cell_index": cell_idx,
                        "para_index": 0,  # 单元格级别，不再区分段落
                        "source_type": "table",
                        "element": first_para._element if first_para else None,
                    }

                    table_data.append(cell_data)
                    logger.debug(
                        f"从表格提取单元格: 表{table_index + 1}行{row_idx + 1}列{cell_idx + 1} - '{cell_text[:50]}...'"
                    )

        return table_data

    except Exception as e:
        logger.error(f"提取表格内容失败: {e!s}")
        return []


def get_table_position(table, table_index):
    """
    获取表格在文档中的位置信息

    参数:
        table: Table对象
        table_index: 表格索引

    返回:
        dict: 位置信息
    """
    try:
        # 获取表格元素
        table_element = table._element

        # 获取表格在文档中的位置
        body = table_element.getparent()
        if body is not None:
            # 获取body的所有子元素
            all_elements = list(body)

            # 找到表格元素的位置
            if table_element in all_elements:
                table_pos = all_elements.index(table_element)

                # 计算表格之前的段落数量
                paragraph_count_before_table = 0
                for i in range(table_pos):
                    elem = all_elements[i]
                    if elem.tag.endswith("}p"):  # 段落元素
                        paragraph_count_before_table += 1

                # 返回表格应该插入的位置（在表格之前的最后一个段落之后）
                # 注意：这里返回的是表格应该被插入的位置，而不是表格本身的位置
                insert_position = paragraph_count_before_table
                logger.debug(
                    f"表格{table_index + 1}位置计算: 表格在元素位置{table_pos}, 之前段落数{paragraph_count_before_table}, 插入位置{insert_position}"
                )

                return {
                    "type": "paragraph",
                    "index": insert_position,  # 在表格原本的位置插入
                    "element": table_element,
                    "absolute_position": table_pos,
                }
            else:
                logger.warning(f"表格{table_index + 1}元素不在body子元素列表中")
                return {"type": "paragraph", "index": -1, "element": None}

        logger.warning(f"表格{table_index + 1}无法获取父元素body")
        return {"type": "paragraph", "index": -1, "element": None}

    except Exception as e:
        logger.warning(f"获取表格位置失败: {e!s}")
        return {"type": "paragraph", "index": -1, "element": None}


def convert_table_to_md(table_data_list):
    """
    将表格数据转换为Markdown格式（用于调试或特定输出）

    注意：此函数处理的是 DOCX 中的 table（文档内表格），
    不要与 xlsx2md 中的 convert_spreadsheet_to_md 混淆（表格文件）

    参数:
        table_data_list: 表格数据列表

    返回:
        str: Markdown格式的表格
    """
    if not table_data_list:
        return ""

    try:
        # 按表格分组
        tables_dict = {}
        for data in table_data_list:
            table_idx = data["table_index"]
            if table_idx not in tables_dict:
                tables_dict[table_idx] = []
            tables_dict[table_idx].append(data)

        # 为每个表格生成Markdown
        markdown_tables = []
        for _table_idx, table_data in tables_dict.items():
            # 确定表格的行列数
            max_row = max(data["row_index"] for data in table_data)
            max_col = max(data["cell_index"] for data in table_data)

            # 创建二维数组存储表格内容
            table_grid = [["" for _ in range(max_col + 1)] for _ in range(max_row + 1)]

            # 填充表格内容
            for data in table_data:
                row = data["row_index"]
                col = data["cell_index"]
                table_grid[row][col] = data["text"]

            # 生成Markdown表格
            markdown_lines = []

            # 表头
            header_line = "| " + " | ".join(table_grid[0]) + " |"
            markdown_lines.append(header_line)

            # 分隔线
            separator_line = "| " + " | ".join(["---"] * (max_col + 1)) + " |"
            markdown_lines.append(separator_line)

            # 数据行
            for row_idx in range(1, max_row + 1):
                row_line = "| " + " | ".join(table_grid[row_idx]) + " |"
                markdown_lines.append(row_line)

            markdown_tables.append("\n".join(markdown_lines))

        return "\n\n".join(markdown_tables)

    except Exception as e:
        logger.error(f"转换表格为Markdown失败: {e!s}")
        return "表格内容提取失败"


# ============================================================================
# 新增功能：支持非公文模式的表格Markdown输出（保留表格格式）
# ============================================================================


def extract_tables_with_structure(doc):
    """
    从文档中提取表格结构（用于非公文模式，保留表格格式）

    参数:
        doc: Document对象

    返回:
        list: 表格信息列表，每个元素包含 {'table': table_obj, 'position': idx}
    """
    logger.info("开始提取文档中的表格结构（非公文模式）")
    tables_info = []

    try:
        for table_idx, table in enumerate(doc.tables):
            # 获取表格位置
            position_info = get_table_position(table, table_idx)

            tables_info.append({"table": table, "table_index": table_idx, "position": position_info.get("index", -1)})

            logger.debug(f"提取表格 {table_idx + 1}, 位置: {position_info.get('index', -1)}")

        logger.info(f"共提取到 {len(tables_info)} 个表格结构")
        return tables_info

    except Exception as e:
        logger.error(f"提取表格结构失败: {e!s}", exc_info=True)
        return []


def find_nested_tables_in_cell(cell):
    """
    检测单元格中是否有嵌套表格

    参数:
        cell: Table Cell对象

    返回:
        list: 嵌套表格列表
    """
    nested_tables = []

    try:
        # 在单元格的XML中查找嵌套的表格元素
        tc_element = cell._tc
        nested_table_elements = tc_element.findall(".//w:tbl", namespaces=NAMESPACES)

        # 排除当前单元格所属的表格（只要直接嵌套的）
        for tbl_elem in nested_table_elements:
            # 检查该表格元素的父元素是否在当前单元格内
            # 简单的方法：如果找到了 w:tbl，就认为是嵌套表格
            # 这里需要更精确的判断，暂时先简单处理
            try:
                from docx.table import Table

                nested_table = Table(tbl_elem, cell._parent)
                nested_tables.append(nested_table)
            except Exception as e:
                logger.warning(f"解析嵌套表格失败: {e}")
                continue

    except Exception as e:
        logger.debug(f"查找嵌套表格失败: {e}")

    return nested_tables


def save_nested_table_as_md(
    nested_table, output_folder, parent_file_name, nested_index, images_info, config_manager, options=None
):
    """
    保存嵌套表格为独立md文件

    参数:
        nested_table: 嵌套的Table对象
        output_folder: 输出文件夹路径
        parent_file_name: 父文件名（不含扩展名）
        nested_index: 嵌套表格索引
        images_info: 图片信息列表
        config_manager: 配置管理器
        options: 选项字典（包含 keep_images, enable_ocr 等）

    返回:
        str: 创建的md文件名
    """
    from docwen.utils.path_utils import generate_output_path

    # 1. 生成md文件路径（使用统一工具）
    dummy_input_path = str(Path(output_folder) / f"{parent_file_name}.docx")

    md_path = generate_output_path(
        input_path=dummy_input_path,
        output_dir=output_folder,
        section=f"nested_table_{nested_index}",
        add_timestamp=True,
        description="fromDocx",
        file_type="md",
    )
    md_filename = Path(md_path).name

    try:
        # 2. 转换嵌套表格为Markdown（递归调用支持图片）
        markdown_content = convert_table_to_md_with_images(
            nested_table,
            -1,  # table_index 不重要
            images_info,
            output_folder,
            f"{parent_file_name}_nested_{nested_index}",
            config_manager,
            options=options,
        )

        # 3. 替换图片标记为实际链接（如果不替换，文件中将只有 {{IMAGE...}}）
        if markdown_content:
            keep_images = options.get("keep_images", True) if options else True
            enable_ocr = options.get("enable_ocr", False) if options else False
            extraction_mode = options.get("extraction_mode", "file") if options else "file"
            ocr_placement_mode = options.get("ocr_placement_mode", "image_md") if options else "image_md"

            logger.debug(f"正在替换嵌套表格 {md_filename} 中的图片标记")
            markdown_content = replace_image_markers_in_table(
                markdown_content,
                images_info,
                keep_images,
                enable_ocr,
                extraction_mode,
                ocr_placement_mode,
                output_folder,
            )

        # 4. 写入文件
        with Path(md_path).open("w", encoding="utf-8") as f:
            f.write(markdown_content)

        logger.info(f"创建嵌套表格md文件: {md_filename} (包含 {len(markdown_content)} 字符)")
        return md_filename, markdown_content

    except Exception as e:
        logger.error(f"创建嵌套表格md文件失败: {md_filename}, 错误: {e}", exc_info=True)
        return f"[嵌套表格{nested_index}]", ""


def convert_table_to_md_with_images(
    table, table_index, images_info, output_folder, parent_file_name, config_manager, options=None
):
    """
    将DOCX表格转为Markdown格式，支持图片、公式和嵌套表格

    参数:
        options: 字典，包含 keep_images, enable_ocr 等
    """
    markdown_lines = []
    nested_table_counter = 1

    # 默认选项
    if options is None:
        options = {
            "keep_images": True,
            "enable_ocr": False,
            "extraction_mode": "file",
            "ocr_placement_mode": "image_md",
        }

    # 获取格式保留配置
    preserve_formatting = config_manager.get_docx_to_md_preserve_formatting()
    preserve_table_header_formatting = config_manager.get_docx_to_md_preserve_table_header_formatting()
    syntax_config = config_manager.get_all_syntax_settings()
    wps_shading_enabled = config_manager.is_wps_shading_enabled()
    word_shading_enabled = config_manager.is_word_shading_enabled()

    try:
        # 遍历表格行
        for row_idx, row in enumerate(table.rows):
            row_cells = []
            is_header_row = row_idx == 0  # 第一行为表头
            last_tc = None

            for cell_idx, cell in enumerate(row.cells):
                if last_tc is not None and getattr(cell, "_tc", None) is last_tc:
                    row_cells.append("")
                    continue
                last_tc = getattr(cell, "_tc", None)

                cell_parts = []  # 使用列表收集单元格的各部分内容

                # 遍历单元格的所有段落
                for para in cell.paragraphs:
                    # 检查段落是否包含公式
                    has_formula = is_formula_supported() and has_formulas_in_paragraph(para)

                    if has_formula:
                        # 提取公式内容，根据行类型和配置决定是否保留格式
                        should_preserve_format = (
                            preserve_table_header_formatting if is_header_row else preserve_formatting
                        )
                        formula_md = process_paragraph_with_formulas(
                            para, preserve_formatting=should_preserve_format, syntax_config=syntax_config
                        )
                        if formula_md:
                            # 移除公式中的换行符以保持表格单行格式
                            formula_md = formula_md.replace("$$\n", "$$").replace("\n$$", "$$")
                            cell_parts.append(formula_md)
                            logger.debug(f"表格单元格[{row_idx},{cell_idx}]包含公式: {formula_md[:50]}...")
                            continue

                    # 处理普通文本
                    para_text_raw = para.text.strip()
                    if para_text_raw:
                        # 根据行类型和配置决定是否保留格式
                        # 表头行：根据 preserve_table_header_formatting 配置
                        # 数据行：根据 preserve_formatting 配置
                        should_preserve_format = (
                            preserve_table_header_formatting if is_header_row else preserve_formatting
                        )

                        if should_preserve_format:
                            # 使用带样式检测的转换函数，支持检测代码/引用相关的字符样式
                            # 包括：行内代码(Inline Code)、代码块关联字符样式(Code Block Char)等
                            para_text, _, _ = convert_paragraph_to_markdown_with_styles(
                                para, config_manager, True, syntax_config, wps_shading_enabled, word_shading_enabled
                            )
                        else:
                            # 不保留格式，只提取纯文本
                            para_text = para_text_raw

                        # 将单元格内的换行符替换为<br>，避免破坏Markdown表格结构
                        para_text = para_text.replace("\r\n", "<br>").replace("\n", "<br>").replace("\r", "<br>")
                        cell_parts.append(para_text)

                    # 检查该段落是否有图片
                    para_images = []
                    for img in images_info:
                        img_para = img.get("paragraph")
                        if img_para is not None and para._element is not None and img_para._element == para._element:
                            para_images.append(img)
                            logger.debug(f"单元格匹配到图片: {img.get('filename')}")

                    for img in para_images:
                        filename = img.get("filename", "")
                        if filename:
                            marker = f"{{{{IMAGE:{filename}}}}}"
                            cell_parts.append(marker)

                # 检测嵌套表格
                nested_tables = find_nested_tables_in_cell(cell)
                if nested_tables:
                    logger.info(
                        f"表{table_index + 1}行{row_idx + 1}列{cell_idx + 1}检测到 {len(nested_tables)} 个嵌套表格"
                    )

                    for nested_table in nested_tables:
                        md_name, _md_content = save_nested_table_as_md(
                            nested_table,
                            output_folder,
                            parent_file_name,
                            nested_table_counter,
                            images_info,
                            config_manager,
                            options=options,
                        )

                        link_settings = config_manager.get_markdown_link_style_settings()
                        md_file_link_style = link_settings.get("md_file_link_style", "wiki_embed")

                        from docwen.utils.markdown_utils import format_md_file_link

                        nested_link = format_md_file_link(md_name, md_file_link_style)
                        cell_parts.append(nested_link)

                        nested_table_counter += 1

                # 智能合并单元格内容
                cell_text = ""
                for i, part in enumerate(cell_parts):
                    if i > 0:
                        # 检查前一部分和当前部分是否都是公式
                        prev_is_formula = cell_parts[i - 1].strip().startswith("$")
                        curr_is_formula = part.strip().startswith("$")

                        # 只在非公式之间添加<br>
                        if not (prev_is_formula and curr_is_formula):
                            cell_text += "<br>"
                    cell_text += part

                # 转义管道符
                cell_text = cell_text.replace("|", "\\|")

                row_cells.append(cell_text)

            # 生成行
            row_line = "| " + " | ".join(row_cells) + " |"
            markdown_lines.append(row_line)

            # 在第一行后添加分隔线
            if row_idx == 0:
                separator_line = "| " + " | ".join(["---"] * len(row_cells)) + " |"
                markdown_lines.append(separator_line)

        return "\n".join(markdown_lines)

    except Exception as e:
        logger.error(f"转换表格为Markdown失败（增强版）: {e}", exc_info=True)
        return "表格转换失败"


def replace_image_markers_in_table(
    table_md,
    images_info,
    keep_images,
    enable_ocr,
    extraction_mode,
    ocr_placement_mode,
    output_folder,
    progress_callback=None,
    cancel_event=None,
):
    """
    替换表格Markdown中的图片标记为实际链接

    参数:
        table_md: 表格Markdown文本
        images_info: 图片信息列表
        keep_images: 是否保留图片
        enable_ocr: 是否启用OCR
        output_folder: 输出文件夹路径
        progress_callback: 进度回调函数
        cancel_event: 取消事件

    返回:
        str: 替换后的Markdown文本
    """

    # 引入图片处理函数（与文档转MD一致）
    try:
        from .image_processor import process_image_with_ocr
    except ImportError:
        logger.warning("无法导入 process_image_with_ocr，图片标记将不被替换")
        return table_md

    # 遍历图片信息，替换标记
    for img_info in images_info:
        filename = img_info.get("filename", "")
        if not filename:
            continue

        marker = f"{{{{IMAGE:{filename}}}}}"

        # 如果Markdown中包含这个标记，替换为实际链接
        if marker in table_md:
            # 处理图片（与文档转MD逻辑一致）
            image_link = process_image_with_ocr(
                img_info,
                keep_images,
                enable_ocr,
                output_folder,
                progress_callback=progress_callback,
                cancel_event=cancel_event,
                extraction_mode=extraction_mode,
                ocr_placement_mode="image_md" if ocr_placement_mode == "main_md" else ocr_placement_mode,
            )

            table_md = table_md.replace(marker, image_link)
            logger.debug(f"替换表格中的图片标记: {marker} -> {image_link}")

    return table_md
