"""
公文格式DOCX转MD转换器

实现公文专用的转换逻辑：
- 三轮公文元素识别
- YAML元数据提取（14个字段）
- 附件内容单独输出
- 支持图片提取和OCR

⚠️ 格式保留说明：
    公文模式**不保留**粗体、斜体、删除线等字符格式，原因如下：
    
    1. 公文格式有严格规范：根据《党政机关公文格式》（GB/T 9704-2012），
       公文正文要求使用统一的3号仿宋字体，不允许使用粗体、斜体等装饰性格式
    2. 转换目标不同：公文模式侧重于结构化元素提取（发文字号、主送机关、
       成文日期等），而非保留视觉格式
    3. 保证输出规范性：丢弃源文档中不规范的格式，确保转换后的Markdown文本干净整洁
    
    如需保留字符格式，请使用简化模式（simple/converter.py）。

⚠️ 国际化说明：
    本模块的状态栏进度信息**不进行国际化**，保持中文硬编码，原因如下：
    
    1. 公文是中国特有概念：根据《党政机关公文格式》（GB/T 9704-2012），
       "公文"专指中国党政机关的正式文件格式，其他国家没有对应概念
    2. 目标用户群体：只有中文用户才会使用"针对公文优化"功能
    3. 代码维护：避免为仅中文用户使用的功能维护多语言翻译
    
    注意：通用步骤（如加载文档、提取图片等）仍使用i18n，
    仅公文专属步骤（如识别公文元素、三轮识别等）保持中文硬编码。

模块结构：
- converter.py: 主转换入口（本文件）
- content_generator.py: 内容生成器
- extractor.py: 公文元素提取
- scorer.py: 元素评分器
"""

import os
import re
import logging
import tempfile
import threading
from typing import Callable, Optional
from docx import Document

# 导入处理模块
from ..shared.content_injector import process_document_with_special_content
from ..shared.image_processor import extract_images_from_docx
from .scorer import create_element_scorer
from .content_generator import generate_main_content, generate_attachment_content
from .extractor import (
    extract_doc_number_and_signers,
    extract_signers_from_text,
    extract_doc_number_and_name
)

# 导入工具函数
from docwen.utils.docx_utils import (
    extract_format_from_paragraph_style, 
    get_effective_run_format
)
from docwen.utils.validation_utils import contains_chinese
from docwen.utils.date_utils import convert_date_format
from docwen.i18n import t
from docwen.utils.text_utils import (
    remove_colon,
    extract_after_colon,
    remove_brackets,
    process_copy_to,
    process_attachment_item
)

logger = logging.getLogger(__name__)


def convert_docx_to_md_gongwen(
    docx_path: str,
    extract_image: bool = True,
    extract_ocr: bool = False,
    progress_callback: Optional[Callable[[str], None]] = None,
    cancel_event: Optional[threading.Event] = None,
    output_folder: Optional[str] = None,
    original_file_path: Optional[str] = None,
    options: Optional[dict] = None
) -> dict:
    """
    公文优化模式的DOCX转MD转换
    
    采用三轮识别策略提取公文元素，生成详细的YAML元数据。
    
    参数:
        docx_path: str - DOCX文件路径（可能是临时副本）
        extract_image: bool - 是否保留图片（由GUI传入，默认True）
        extract_ocr: bool - 是否进行OCR识别（由GUI传入，默认False）
        progress_callback: Callable - 进度回调函数（可选）
        cancel_event: threading.Event - 取消事件（可选）
        output_folder: str - 输出文件夹路径，用于保存图片（可选）
        original_file_path: str - 原始文件路径（用于图片命名，可选）
        options: dict - 转换选项字典，包含序号配置等（可选）
    
    返回:
        dict: 转换结果字典，包含以下键：
            - success: bool - 转换是否成功
            - main_content: str - 主要markdown内容
            - attachment_content: str or None - 附件markdown内容
            - metadata: dict - YAML元数据
            - error: str or None - 错误信息
    """
    # 直接使用GUI传入的参数
    keep_images = extract_image
    enable_ocr = extract_ocr
    
    # 提取序号配置参数
    if options is None:
        options = {}
    
    remove_numbering = options.get('doc_remove_numbering', False)
    add_numbering = options.get('doc_add_numbering', False)
    scheme_name = options.get('doc_numbering_scheme', 'gongwen_standard')
    
    logger.info(f"开始转换DOCX文件: {docx_path}")
    logger.info(f"导出选项 - 提取图片: {keep_images}, OCR识别: {enable_ocr}")
    logger.info(f"序号配置: 清除={remove_numbering}, 添加={add_numbering}, 方案={scheme_name}")
    
    try:
        # 1. 加载DOCX文档
        if progress_callback:
            progress_callback(t('conversion.progress.loading_document'))
        doc = Document(docx_path)
        logger.info(f"成功加载DOCX文档, 包含 {len(doc.paragraphs)} 个段落")
        
        # 2. 处理文本框和表格内容
        if progress_callback:
            progress_callback(t('conversion.progress.processing_textboxes_tables'))
        
        temp_dir = tempfile.gettempdir()
        modified_doc, extracted_path = process_document_with_special_content(
            doc, docx_path, temp_dir, mode='gongwen'
        )
        
        if extracted_path:
            logger.info(f"提取文档已生成: {extracted_path}")
        
        doc = modified_doc
        logger.info(f"处理后的文档包含 {len(doc.paragraphs)} 个段落")
        
        # 3. 提取图片信息
        images_info = []
        if (keep_images or enable_ocr) and output_folder:
            if progress_callback:
                progress_callback(t('conversion.progress.extracting_images', count=''))
            
            try:
                path_for_naming = original_file_path or docx_path
                images_info = extract_images_from_docx(
                    doc, output_folder, path_for_naming,
                    progress_callback=progress_callback
                )
                logger.info(f"图片提取完成，共 {len(images_info)} 张")
            except Exception as e:
                logger.error(f"图片提取失败: {e}", exc_info=True)
                images_info = []
        
        # 4. 三轮公文元素识别（公文专属步骤，不国际化）
        if progress_callback:
            progress_callback("正在识别公文元素...")
        
        yaml_info, skip_para_indices, attachment_content_indices = _identify_elements(
            doc, images_info, progress_callback, cancel_event
        )
        
        if cancel_event and cancel_event.is_set():
            return _cancelled_result(yaml_info)
        
        # 5. 生成主要Markdown内容（公文专属步骤，不国际化）
        if progress_callback:
            progress_callback("正在生成Markdown内容...")
        
        main_content = generate_main_content(
            yaml_info, doc, skip_para_indices, images_info, 
            keep_images, enable_ocr, output_folder, 
            progress_callback, cancel_event,
            remove_numbering, add_numbering, scheme_name
        )
        logger.info("主要Markdown内容生成完成")
        
        # 6. 生成附件Markdown内容
        attachment_content = None
        if attachment_content_indices:
            if cancel_event and cancel_event.is_set():
                return _cancelled_result(yaml_info)
            
            attachment_content = generate_attachment_content(
                doc, attachment_content_indices, docx_path, images_info, 
                keep_images, enable_ocr, output_folder, 
                progress_callback, cancel_event,
                remove_numbering, add_numbering, scheme_name
            )
            logger.info("附件Markdown内容生成完成")
        else:
            logger.info("未检测到附件内容")
        
        return {
            'success': True,
            'main_content': main_content,
            'attachment_content': attachment_content,
            'metadata': yaml_info,
            'error': None
        }
        
    except Exception as e:
        error_msg = f"转换DOCX文件失败: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return {
            'success': False,
            'main_content': None,
            'attachment_content': None,
            'metadata': {},
            'error': error_msg
        }


def _cancelled_result(yaml_info: dict) -> dict:
    """返回取消操作的结果"""
    logger.info("操作被用户取消")
    return {
        'success': False,
        'main_content': None,
        'attachment_content': None,
        'metadata': yaml_info,
        'error': '操作已取消'
    }


def _identify_elements(
    doc, 
    images_info: list, 
    progress_callback: Callable, 
    cancel_event: threading.Event
) -> tuple:
    """
    三轮公文元素识别
    
    参数:
        doc: Document - Word文档对象
        images_info: list - 图片信息列表
        progress_callback: Callable - 进度回调函数
        cancel_event: threading.Event - 取消事件
    
    返回:
        tuple: (yaml_info, skip_para_indices, attachment_content_indices)
    """
    # 创建评分器
    scorer = create_element_scorer()
    logger.info(f"初始化公文元素评分器")
    
    # 收集文档上下文信息
    context = {
        'total_paras': len(doc.paragraphs),
        'para_index': 0,
        'para_non_empty': {},
        'doc': doc
    }
    
    # 初始化YAML信息
    yaml_info = _init_yaml_info()
    
    # 需要跳过的段落索引    
    skip_para_indices = []  
    title_groups = []
    attachment_content_indices = []
    
    # 初始化段落状态管理
    identified_paragraphs = set()
    paragraph_elements = {}
    
    # 第1轮：识别大部分唯一性元素（公文专属步骤，不国际化）
    if progress_callback:
        progress_callback("第1轮：识别唯一性元素...")
    logger.info("开始第1轮识别：大部分唯一性元素")
    
    for idx, para in enumerate(doc.paragraphs):
        if cancel_event and cancel_event.is_set():
            break
        
        para_text = para.text.strip()
        context['para_index'] = idx
        context['para_non_empty'][idx] = bool(para_text)
        
        if not para_text:
            continue
        
        run_dict = _get_run_dict(para, para_text)
        best_element, best_score = scorer.score_round1_unique_elements(run_dict, context)
        
        if best_element and best_score >= 80:
            element_name = scorer.ELEMENT_TYPES.get(best_element, "")
            logger.info(f"第1轮段落 {idx+1} 识别结果: {element_name} (得分: {best_score})")
            
            scorer.update_context(best_element, idx)
            identified_paragraphs.add(idx)
            paragraph_elements[idx] = best_element
    
    # 第2轮：识别发文机关标志和印发机关（公文专属步骤，不国际化）
    if progress_callback:
        progress_callback("第2轮：识别机关标志...")
    logger.info("开始第2轮识别：发文机关标志和印发机关")
    
    for idx, para in enumerate(doc.paragraphs):
        if cancel_event and cancel_event.is_set():
            break
        
        if idx in identified_paragraphs:
            continue
            
        para_text = para.text.strip()
        if not para_text:
            continue
            
        context['para_index'] = idx
        run_dict = _get_run_dict(para, para_text)
        best_element, best_score = scorer.score_round2_unique_elements(run_dict, context)
        
        if best_element and best_score >= 80:
            element_name = scorer.ELEMENT_TYPES.get(best_element, "")
            logger.info(f"第2轮段落 {idx+1} 识别结果: {element_name} (得分: {best_score})")
            
            scorer.update_context(best_element, idx)
            identified_paragraphs.add(idx)
            paragraph_elements[idx] = best_element
    
    # 第3轮：识别非唯一性元素（公文专属步骤，不国际化）
    if progress_callback:
        progress_callback("第3轮：识别非唯一性元素...")
    logger.info("开始第3轮识别：非唯一性元素")
    
    for idx, para in enumerate(doc.paragraphs):
        if cancel_event and cancel_event.is_set():
            break
        
        if idx in identified_paragraphs:
            continue
            
        para_text = para.text.strip()
        if not para_text:
            # 空段落，检查是否有图片
            has_image = any(img['para_index'] == idx for img in images_info)
            if has_image:
                if idx > scorer.last_known_element_index:
                    paragraph_elements[idx] = 'attachment_content'
                    attachment_content_indices.append(idx)
                    skip_para_indices.append(idx)
                    logger.info(f"检测到附件区域的空图片段落 {idx+1}")
                else:
                    paragraph_elements[idx] = 'general_content'
                    logger.info(f"检测到正文区域的空图片段落 {idx+1}")
            continue
            
        context['para_index'] = idx
        run_dict = _get_run_dict(para, para_text)
        best_element, best_score = scorer.score_round3_non_unique_elements(run_dict, context)
        
        if best_element and best_score > 0:
            element_name = scorer.ELEMENT_TYPES.get(best_element, "")
            logger.info(f"第3轮段落 {idx+1} 识别结果: {element_name} (得分: {best_score})")
            
            scorer.update_context(best_element, idx)
            paragraph_elements[idx] = best_element
            
            if best_element == 'attachment_content':
                logger.info(f"检测到附件内容段落 {idx+1}")
                attachment_content_indices.append(idx)
                skip_para_indices.append(idx)
    
    # 提取关键信息到YAML
    _extract_yaml_info(doc, paragraph_elements, yaml_info, skip_para_indices, title_groups)
    
    return yaml_info, skip_para_indices, attachment_content_indices


def _init_yaml_info() -> dict:
    """初始化YAML信息字典"""
    return {
        "aliases": [],
        "标题": "",
        "份号": "",
        "密级和保密期限": "",
        "紧急程度": "",
        "发文字号": "",
        "发文机关标志": "",
        "签发人": [],
        "发文机关署名": "",
        "成文日期": "",
        "印发日期": "",
        "主送机关": "",
        "附注": "",
        "印发机关": "",
        "抄送机关": [],
        "附件说明": [],
        "公开方式": ""
    }


def _get_run_dict(para, para_text: str) -> dict:
    """获取段落的运行字典"""
    para_fonts = extract_format_from_paragraph_style(para) or {}
    effective_fonts = get_effective_run_format(para, para_fonts) or {}
    
    return {
        'text': para_text,
        'fonts': effective_fonts,
        'is_chinese': contains_chinese(para_text)
    }


def _extract_yaml_info(
    doc, 
    paragraph_elements: dict, 
    yaml_info: dict, 
    skip_para_indices: list,
    title_groups: list
):
    """
    从识别的段落元素中提取YAML信息
    
    参数:
        doc: Document - Word文档对象
        paragraph_elements: dict - 段落元素映射
        yaml_info: dict - YAML信息字典（会被修改）
        skip_para_indices: list - 跳过索引列表（会被修改）
        title_groups: list - 标题组列表（会被修改）
    """
    for idx, best_element in paragraph_elements.items():
        para = doc.paragraphs[idx]
        para_text = para.text.strip()
        
        # 处理组合ID（份号+发文字号）
        if best_element == 'combined_id':
            parts = re.split(r'[\s\x20\t]+', para_text, 1)
            if len(parts) >= 2:
                yaml_info['份号'] = parts[0]
                yaml_info['发文字号'] = parts[1]
            else:
                yaml_info['发文字号'] = para_text
            skip_para_indices.append(idx)

        elif best_element == 'title' and not yaml_info['aliases']:
            skip_para_indices.append(idx)
        
        elif best_element == 'copy_id' and not yaml_info['份号']:
            yaml_info['份号'] = para_text
            skip_para_indices.append(idx)
        
        elif best_element == 'security' and not yaml_info['密级和保密期限']:
            yaml_info['密级和保密期限'] = para_text
            skip_para_indices.append(idx)
        
        elif best_element == 'urgency' and not yaml_info['紧急程度']:
            yaml_info['紧急程度'] = para_text
            skip_para_indices.append(idx)
        
        elif best_element == 'doc_number' and not yaml_info['发文字号']:
            yaml_info['发文字号'] = para_text
            skip_para_indices.append(idx)
        
        elif best_element == 'issuing_authority_mark' and not yaml_info['发文机关标志']:
            yaml_info['发文机关标志'] = para_text
            skip_para_indices.append(idx)
        
        elif best_element == 'combined_doc_number_signer':
            doc_num, signers = extract_doc_number_and_signers(para_text)
            if doc_num and not yaml_info['发文字号']:
                yaml_info['发文字号'] = doc_num
            if signers:
                yaml_info['签发人'].extend(signers)
            skip_para_indices.append(idx)
        
        elif best_element == 'signer':
            signer = extract_after_colon(para_text)
            if signer:
                yaml_info['签发人'].append(signer)
            skip_para_indices.append(idx)
        
        elif best_element == 'signer_following':
            signers = extract_signers_from_text(para_text)
            if signers:
                yaml_info['签发人'].extend(signers)
            skip_para_indices.append(idx)
        
        elif best_element == 'combined_doc_number_signer_following':
            doc_num, name = extract_doc_number_and_name(para_text)
            if name:
                yaml_info['签发人'].append(name)
            skip_para_indices.append(idx)
        
        elif best_element == 'issuing_authority_signature' and not yaml_info['发文机关署名']:
            yaml_info['发文机关署名'] = para_text
            skip_para_indices.append(idx)
        
        elif best_element == 'issue_date' and not yaml_info['成文日期']:
            yaml_info['成文日期'] = convert_date_format(para_text)
            skip_para_indices.append(idx)
        
        elif best_element == 'recipient' and not yaml_info['主送机关']:
            yaml_info['主送机关'] = remove_colon(para_text)
            skip_para_indices.append(idx)
        
        elif best_element == 'notes' and not yaml_info['附注']:
            yaml_info['附注'] = remove_brackets(para_text)
            skip_para_indices.append(idx)
        
        elif best_element == 'disclosure' and not yaml_info['公开方式']:
            yaml_info['公开方式'] = extract_after_colon(para_text)
            skip_para_indices.append(idx)
        
        elif best_element == 'copy_to' and not yaml_info['抄送机关']:
            yaml_info['抄送机关'] = process_copy_to(para_text)
            skip_para_indices.append(idx)
        
        elif best_element in ('attachment', 'attachment_following'):
            item = process_attachment_item(para_text)
            yaml_info['附件说明'].append(item)
            skip_para_indices.append(idx)
        
        elif best_element == 'printing_authority' and not yaml_info['印发机关']:
            yaml_info['印发机关'] = para_text
            skip_para_indices.append(idx)
        
        elif best_element == 'printing_date' and not yaml_info['印发日期']:
            date_text = para_text.replace('印发', '').strip()
            yaml_info['印发日期'] = convert_date_format(date_text)
            skip_para_indices.append(idx)
    
    # 处理标题组
    _process_title_groups(doc, paragraph_elements, yaml_info, skip_para_indices, title_groups)


def _process_title_groups(
    doc, 
    paragraph_elements: dict, 
    yaml_info: dict, 
    skip_para_indices: list,
    title_groups: list
):
    """处理标题组"""
    current_title_indices = []
    
    for i in range(len(doc.paragraphs)):
        element = paragraph_elements.get(i)
        if element == 'title':
            if current_title_indices:
                title_groups.append(current_title_indices)
            current_title_indices = [i]
        elif element == 'title_following':
            if current_title_indices:
                current_title_indices.append(i)
        else:
            if current_title_indices:
                title_groups.append(current_title_indices)
                current_title_indices = []
    
    if current_title_indices:
        title_groups.append(current_title_indices)

    # 合并标题文本
    if title_groups:
        first_title_indices = title_groups[0]
        full_title = "".join(
            doc.paragraphs[i].text.replace('\n', '').replace('\r', '').strip() 
            for i in first_title_indices
        )
        
        start_idx, end_idx = first_title_indices[0], first_title_indices[-1]
        logger.info(f"提取并合并标题: 段落 {start_idx+1}-{end_idx+1}, 文本: '{full_title}'")
        
        yaml_info['标题'] = full_title
        yaml_info['aliases'].append(full_title)
        
        for i in first_title_indices:
            if i not in skip_para_indices:
                skip_para_indices.append(i)
    else:
        logger.warning("未检测到任何标题组")
