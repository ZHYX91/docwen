"""
简化模式DOCX转MD转换器

实现普通Word文档的基础转换逻辑。

特性：
- 基于Word样式（Title、Subtitle、Heading 1-6）转换
- 提取Title/Subtitle到YAML元数据
- 不做公文元素识别
- 支持图片提取和OCR
- 支持文本框和表格处理

与公文模式的区别：
- 无三轮元素识别
- YAML只有2个字段（标题、副标题）
- 不处理附件内容
- 标题直接基于Word样式，不做内容分析

模块结构：
- converter.py: 主转换入口（本文件）
- yaml_builder.py: YAML头部生成
- paragraph_handler.py: 段落处理器
"""

import logging
import re
import threading
from collections.abc import Callable
from pathlib import Path

from docx import Document

from docwen.config.config_manager import config_manager
from docwen.translation import t

from ..shared.break_processor import BorderGroupTracker
from ..shared.content_injector import process_document_with_special_content
from ..shared.conversion_context import ConversionContext
from ..shared.image_processor import extract_images_from_docx
from ..shared.list_processor import ListContextManager, ListCounterManager, preprocess_list_ranges
from ..shared.note_processor import NoteExtractor
from ..shared.table_processor import (
    convert_table_to_md_with_images,
    extract_tables_with_structure,
    replace_image_markers_in_table,
)
from .paragraph_handler import ParagraphHandler
from .yaml_builder import build_yaml_header

# 配置日志
logger = logging.getLogger(__name__)


def _is_list_item(line: str) -> bool:
    """
    判断一行内容是否是Markdown列表项

    列表项特征：
    - 无序列表：以 `-`、`*`、`+` 开头（可能有前导空格表示缩进）
    - 有序列表：以 `数字.` 开头（可能有前导空格表示缩进）

    参数:
        line: str - 要检查的行

    返回:
        bool: 是否是列表项

    示例:
        >>> _is_list_item('- 项目1')
        True
        >>> _is_list_item('1. 第一项')
        True
        >>> _is_list_item('普通段落')
        False
    """
    if not line:
        return False

    stripped = line.lstrip()

    # 无序列表：以 - * + 开头，后跟空格
    if stripped and stripped[0] in "-*+" and len(stripped) > 1 and stripped[1] == " ":
        return True

    # 有序列表：以 数字. 开头，后跟空格
    return bool(re.match(r"^\d+\.\s", stripped))


def _smart_join_content(content_lines: list) -> str:
    """
    智能合并内容行

    根据Markdown语法规范，智能选择段落间的分隔符：
    - 连续的列表项之间使用单换行（保持列表连续性）
    - 其他情况使用双换行（Markdown段落分隔）

    参数:
        content_lines: list[str] - 内容行列表，每个元素是一个段落、列表项或其他内容块

    返回:
        str: 合并后的Markdown内容，段落间用适当的换行符分隔

    示例:
        >>> _smart_join_content(['- 项目1', '- 项目2', '正文段落'])
        '- 项目1\\n- 项目2\\n\\n正文段落'

        >>> _smart_join_content(['段落1', '段落2'])
        '段落1\\n\\n段落2'
    """
    if not content_lines:
        return ""

    if len(content_lines) == 1:
        return content_lines[0]

    result_parts = []

    for i, line in enumerate(content_lines):
        result_parts.append(line)

        if i < len(content_lines) - 1:
            current_is_list = _is_list_item(line)
            next_is_list = _is_list_item(content_lines[i + 1])

            # 连续列表项之间使用单换行
            if current_is_list and next_is_list:
                result_parts.append("\n")
            else:
                result_parts.append("\n\n")

    return "".join(result_parts)


def convert_docx_to_md_simple(
    docx_path: str,
    extract_image: bool = True,
    extract_ocr: bool = False,
    progress_callback: Callable[[str], None] | None = None,
    cancel_event: threading.Event | None = None,
    output_folder: str | None = None,
    original_file_path: str | None = None,
    options: dict | None = None,
) -> dict:
    """
    简化模式的DOCX转MD转换

    适用于普通Word文档的基础转换，不做公文元素识别。

    参数:
        docx_path: str - DOCX文件路径（可能是临时副本）
        extract_image: bool - 是否保留图片（由GUI传入，默认True）
        extract_ocr: bool - 是否进行OCR识别（由GUI传入，默认False）
        progress_callback: Callable[[str], None] - 进度回调函数，接收进度消息（可选）
        cancel_event: threading.Event - 取消事件，设置后中止转换（可选）
        output_folder: str - 输出文件夹路径，用于保存图片（可选）
        original_file_path: str - 原始文件路径，用于图片命名和默认标题（可选）
        options: dict - 转换选项字典，包含以下键（可选）：
            - remove_numbering: bool - 是否清除原有序号
            - add_numbering: bool - 是否添加新序号
            - numbering_scheme: str - 序号方案名称

    返回:
        dict: 转换结果字典，包含以下键：
            - success: bool - 转换是否成功
            - main_content: str - 主要markdown内容（含YAML头部）
            - attachment_content: None - 简化模式不处理附件
            - metadata: dict - 提取的YAML元数据（标题、副标题、aliases）
            - error: str or None - 错误信息（如果失败）

    示例:
        >>> result = convert_docx_to_md_simple('document.docx', extract_image=True)
        >>> if result['success']:
        ...     print(result['main_content'])
    """
    # 直接使用GUI传入的参数
    keep_images = extract_image
    enable_ocr = extract_ocr

    # 提取序号配置参数
    if options is None:
        options = {}

    extraction_mode = options.get(
        "to_md_image_extraction_mode",
        config_manager.get_docx_to_md_image_extraction_mode(),
    )
    ocr_placement_mode = options.get(
        "to_md_ocr_placement_mode",
        config_manager.get_docx_to_md_ocr_placement_mode(),
    )

    remove_numbering = options.get("remove_numbering", options.get("doc_remove_numbering", False))
    add_numbering = options.get("add_numbering", options.get("doc_add_numbering", False))
    scheme_name = options.get("numbering_scheme", options.get("doc_numbering_scheme", "gongwen_standard"))

    logger.info(f"开始简化模式转换DOCX文件: {docx_path}")
    logger.info(f"导出选项 - 提取图片: {keep_images}, OCR识别: {enable_ocr}")
    logger.info(f"模式参数 - 图片提取方式: {extraction_mode}, OCR位置: {ocr_placement_mode}")
    logger.info(f"序号配置: 清除={remove_numbering}, 添加={add_numbering}, 方案={scheme_name}")

    try:
        # 1. 加载DOCX文档
        if progress_callback:
            progress_callback(t("conversion.progress.loading_document"))
        doc = Document(docx_path)
        logger.info(f"成功加载DOCX文档, 包含 {len(doc.paragraphs)} 个段落")

        # 2. 处理文本框和表格内容
        if progress_callback:
            progress_callback(t("conversion.progress.processing_textboxes_tables"))

        import tempfile

        temp_dir = tempfile.gettempdir()
        modified_doc, extracted_path = process_document_with_special_content(doc, docx_path, temp_dir, mode="simple")

        if extracted_path:
            logger.info(f"提取文档已生成: {extracted_path}")

        doc = modified_doc
        logger.info(f"处理后的文档包含 {len(doc.paragraphs)} 个段落")

        # 3. 提取图片信息
        images_info = []
        if (keep_images or enable_ocr) and output_folder:
            if progress_callback:
                progress_callback(t("conversion.progress.extracting_images", count=""))

            try:
                path_for_naming = original_file_path or docx_path
                images_info = extract_images_from_docx(
                    doc, output_folder, path_for_naming, progress_callback=progress_callback
                )
                logger.info(f"图片提取完成，共 {len(images_info)} 张")
            except Exception as e:
                logger.error(f"图片提取失败: {e}", exc_info=True)
                images_info = []

        # 4. 提取Title和Subtitle
        if progress_callback:
            progress_callback(t("conversion.progress.extracting_title"))

        metadata, title_indices, subtitle_indices = _extract_title_metadata(
            doc, docx_path, original_file_path, cancel_event
        )

        if cancel_event and cancel_event.is_set():
            return _cancelled_result()

        # 5. 生成Markdown内容
        if progress_callback:
            progress_callback(t("conversion.progress.generating_markdown"))

        # 提取表格结构
        tables_info = extract_tables_with_structure(doc)

        # 获取父文件名
        parent_file_name = Path(original_file_path or docx_path).stem

        skip_indices = title_indices + subtitle_indices
        main_content = _generate_markdown_content_simple(
            metadata=metadata,
            doc=doc,
            docx_path=docx_path,
            tables_info=tables_info,
            skip_indices=skip_indices,
            images_info=images_info,
            keep_images=keep_images,
            enable_ocr=enable_ocr,
            extraction_mode=extraction_mode,
            ocr_placement_mode=ocr_placement_mode,
            output_folder=output_folder,
            parent_file_name=parent_file_name,
            progress_callback=progress_callback,
            cancel_event=cancel_event,
            remove_numbering=remove_numbering,
            add_numbering=add_numbering,
            scheme_name=scheme_name,
        )

        if output_folder and images_info and (not keep_images or extraction_mode == "base64"):
            for img in images_info:
                try:
                    Path(img.get("image_path", "")).unlink(missing_ok=True)
                except Exception:
                    pass

        logger.info("简化模式转换成功")
        return {
            "success": True,
            "main_content": main_content,
            "attachment_content": None,
            "metadata": metadata,
            "error": None,
        }

    except Exception as e:
        error_msg = f"简化模式转换DOCX文件失败: {e!s}"
        logger.error(error_msg, exc_info=True)
        return {"success": False, "main_content": None, "attachment_content": None, "metadata": {}, "error": error_msg}


def _extract_title_metadata(doc, docx_path: str, original_file_path: str | None, cancel_event) -> tuple:
    """
    提取标题和副标题元数据

    扫描文档段落，收集连续的Title和Subtitle样式段落，
    合并为标题和副标题文本。

    参数:
        doc: Document - Word文档对象
        docx_path: str - DOCX文件路径
        original_file_path: str - 原始文件路径（用于默认标题）
        cancel_event: threading.Event - 取消事件

    返回:
        tuple: (metadata, title_indices, subtitle_indices)
            - metadata: dict - 包含 aliases、标题、副标题
            - title_indices: list[int] - Title段落索引列表
            - subtitle_indices: list[int] - Subtitle段落索引列表
    """
    metadata = {"aliases": [], "标题": "", "副标题": ""}
    title_indices = []
    subtitle_indices = []
    has_met_non_title = False
    has_met_non_subtitle = False

    for idx, para in enumerate(doc.paragraphs):
        if cancel_event and cancel_event.is_set():
            break

        style_name = para.style.name if para.style else None

        if not style_name:
            continue

        if style_name == "Title":
            if not has_met_non_title:
                title_indices.append(idx)
                logger.debug(f"检测到Title样式 (段落 {idx + 1})")
        elif style_name == "Subtitle":
            if not has_met_non_subtitle:
                subtitle_indices.append(idx)
                logger.debug(f"检测到Subtitle样式 (段落 {idx + 1})")
        else:
            if title_indices:
                has_met_non_title = True
            if subtitle_indices:
                has_met_non_subtitle = True

            if has_met_non_title and has_met_non_subtitle:
                break

    # 处理标题
    if title_indices:
        title = "".join(doc.paragraphs[i].text.replace("\n", "").replace("\r", "").strip() for i in title_indices)
        metadata["标题"] = title
        metadata["aliases"].append(title)
        logger.info(f"合并标题: 段落 {title_indices[0] + 1}-{title_indices[-1] + 1}, 文本: '{title}'")
    else:
        path_for_title = original_file_path or docx_path
        filename = Path(path_for_title).stem
        metadata["标题"] = filename
        metadata["aliases"].append(filename)
        logger.info(f"未找到Title样式，使用文件名作为标题: '{filename}'")

    # 处理副标题
    if subtitle_indices:
        subtitle = "".join(doc.paragraphs[i].text.replace("\n", "").replace("\r", "").strip() for i in subtitle_indices)
        metadata["副标题"] = subtitle
        logger.info(f"合并副标题: 段落 {subtitle_indices[0] + 1}-{subtitle_indices[-1] + 1}")

    return metadata, title_indices, subtitle_indices


def _cancelled_result() -> dict:
    """返回取消操作的结果"""
    logger.info("操作被用户取消")
    return {"success": False, "main_content": None, "attachment_content": None, "metadata": {}, "error": "操作已取消"}


def _generate_markdown_content_simple(
    metadata: dict,
    doc,
    docx_path: str,
    tables_info: list,
    skip_indices: list,
    images_info: list | None = None,
    keep_images: bool = True,
    enable_ocr: bool = False,
    extraction_mode: str = "file",
    ocr_placement_mode: str = "image_md",
    output_folder: str | None = None,
    parent_file_name: str = "document",
    progress_callback=None,
    cancel_event=None,
    remove_numbering: bool = False,
    add_numbering: bool = False,
    scheme_name: str = "gongwen_standard",
) -> str:
    """
    生成简化模式的Markdown内容

    处理流程：
    1. 生成YAML头部
    2. 预扫描列表范围
    3. 初始化段落处理器
    4. 遍历段落和表格
    5. 添加脚注/尾注定义

    参数:
        metadata: dict - 元数据字典，包含 aliases、标题、副标题
        doc: Document - Word文档对象
        docx_path: str - DOCX文件路径（用于脚注提取）
        tables_info: list - 表格信息列表
        skip_indices: list - 需要跳过的段落索引（Title/Subtitle）
        images_info: list - 图片信息列表（可选）
        keep_images: bool - 是否保留图片（默认True）
        enable_ocr: bool - 是否启用OCR（默认False）
        output_folder: str - 输出文件夹路径（可选）
        parent_file_name: str - 父文件名，用于嵌套表格命名（默认"document"）
        progress_callback: Callable - 进度回调函数（可选）
        cancel_event: threading.Event - 取消事件（可选）
        remove_numbering: bool - 是否清除原有序号（默认False）
        add_numbering: bool - 是否添加新序号（默认False）
        scheme_name: str - 序号方案名称（默认"gongwen_standard"）

    返回:
        str: 完整的Markdown内容，包含YAML头部和正文
    """
    if images_info is None:
        images_info = []

    title = metadata.get("标题", "")
    subtitle = metadata.get("副标题", "")
    logger.info(f"生成简化模式Markdown - 标题: '{title}', 副标题: '{subtitle}'")

    # 提取脚注和尾注
    note_extractor = NoteExtractor(doc, docx_path)
    if note_extractor.has_notes:
        logger.info(
            f"提取到脚注/尾注 | 脚注: {len(note_extractor.footnotes)} 个, 尾注: {len(note_extractor.endnotes)} 个"
        )

    # 获取配置
    preserve_formatting = config_manager.get_docx_to_md_preserve_formatting()
    preserve_heading_formatting = config_manager.get_docx_to_md_preserve_heading_formatting()
    syntax_config = config_manager.get_all_syntax_settings()
    wps_shading_enabled = config_manager.is_wps_shading_enabled()
    word_shading_enabled = config_manager.is_word_shading_enabled()
    list_indent_spaces = config_manager.get_list_indent_spaces()

    # 初始化序号格式化器
    formatter = None
    if add_numbering:
        from docwen.utils.heading_numbering import get_formatter_from_config

        formatter = get_formatter_from_config(config_manager, scheme_name)
        if formatter:
            formatter.reset_counters()
            logger.info(f"已创建序号格式化器，方案：{scheme_name}")
        else:
            logger.warning(f"无法创建序号格式化器（方案：{scheme_name}）")
            add_numbering = False

    # 预扫描列表范围
    list_ranges = preprocess_list_ranges(doc, skip_indices)

    # 创建处理组件
    ctx = ConversionContext(config_manager, list_indent_spaces)
    list_context_manager = ListContextManager(list_ranges)
    list_counter_manager = ListCounterManager()
    border_tracker = BorderGroupTracker()

    # 创建段落处理器
    handler = ParagraphHandler(
        ctx=ctx,
        list_context_manager=list_context_manager,
        list_counter_manager=list_counter_manager,
        border_tracker=border_tracker,
        config_manager=config_manager,
        note_extractor=note_extractor,
        list_indent_spaces=list_indent_spaces,
        preserve_formatting=preserve_formatting,
        preserve_heading_formatting=preserve_heading_formatting,
        syntax_config=syntax_config,
        wps_shading_enabled=wps_shading_enabled,
        word_shading_enabled=word_shading_enabled,
        remove_numbering=remove_numbering,
        add_numbering=add_numbering,
        formatter=formatter,
    )

    # 1. 生成YAML头部
    lines = build_yaml_header(metadata)

    # 2. 处理正文
    content_lines = []

    # 建立表格位置映射
    tables_by_position = {}
    for table_info in tables_info:
        pos = table_info["position"]
        if pos not in tables_by_position:
            tables_by_position[pos] = []
        tables_by_position[pos].append(table_info)

    # 计算图片总数
    total_images = sum(1 for img in images_info if img["para_index"] not in skip_indices)
    current_image_index = 0

    total_paragraphs = len(doc.paragraphs)

    # 遍历所有段落和表格
    for idx in range(total_paragraphs + 1):
        # 处理当前位置的表格
        if idx in tables_by_position:
            for table_info in tables_by_position[idx]:
                table_md = _process_table(
                    table_info,
                    images_info,
                    output_folder,
                    parent_file_name,
                    keep_images,
                    enable_ocr,
                    extraction_mode,
                    ocr_placement_mode,
                    progress_callback,
                    cancel_event,
                    list_context_manager,
                    list_indent_spaces,
                    idx,
                )
                content_lines.append(table_md)
                content_lines.append("")  # 表格后加空行

        # 处理完所有段落后退出
        if idx >= total_paragraphs:
            break

        para = doc.paragraphs[idx]

        # 使用段落处理器处理段落
        result, current_image_index = handler.process_paragraph(
            para=para,
            idx=idx,
            skip_indices=skip_indices,
            images_info=images_info,
            keep_images=keep_images,
            enable_ocr=enable_ocr,
            extraction_mode=extraction_mode,
            ocr_placement_mode=ocr_placement_mode,
            output_folder=output_folder,
            progress_callback=progress_callback,
            cancel_event=cancel_event,
            current_image_index=current_image_index,
            total_images=total_images,
        )
        content_lines.extend(result)

    # 处理器收尾
    final_content = handler.finalize()
    content_lines.extend(final_content)

    # 3. 智能合并正文内容
    if content_lines:
        merged_content = _smart_join_content(content_lines)
        lines.append(merged_content)

    # 4. 添加脚注/尾注定义块
    if note_extractor.has_notes:
        definitions_block = note_extractor.build_definitions_block()
        if definitions_block:
            lines.append("")
            lines.append(definitions_block)
            logger.info("已添加脚注/尾注定义块")

    return "\n".join(lines)


def _process_table(
    table_info: dict,
    images_info: list,
    output_folder: str | None,
    parent_file_name: str,
    keep_images: bool,
    enable_ocr: bool,
    extraction_mode: str,
    ocr_placement_mode: str,
    progress_callback,
    cancel_event,
    list_context_manager: ListContextManager,
    list_indent_spaces: int,
    position: int,
) -> str:
    """
    处理单个表格，返回Markdown内容

    参数:
        table_info: dict - 表格信息字典
        images_info: list - 图片信息列表
        output_folder: str - 输出文件夹路径
        parent_file_name: str - 父文件名
        keep_images: bool - 是否保留图片
        enable_ocr: bool - 是否启用OCR
        progress_callback: Callable - 进度回调
        cancel_event: threading.Event - 取消事件
        list_context_manager: ListContextManager - 列表上下文管理器
        list_indent_spaces: int - 列表缩进空格数
        position: int - 表格位置索引

    返回:
        str: 表格的Markdown内容
    """
    table = table_info["table"]
    table_index = table_info["table_index"]

    # 转换表格为Markdown
    table_md = convert_table_to_md_with_images(
        table,
        table_index,
        images_info,
        output_folder or "",
        parent_file_name,
        config_manager,
        options={
            "keep_images": keep_images,
            "enable_ocr": enable_ocr,
            "extraction_mode": extraction_mode,
            "ocr_placement_mode": ocr_placement_mode,
        },
    )

    # 替换表格中的图片标记
    table_md = replace_image_markers_in_table(
        table_md,
        images_info,
        keep_images,
        enable_ocr,
        extraction_mode,
        ocr_placement_mode,
        output_folder or "",
        progress_callback,
        cancel_event,
    )

    # 检查表格是否在列表上下文中
    table_list_context = list_context_manager.get_context_for_position(position)
    if table_list_context:
        _, t_level = table_list_context
        t_indent = " " * list_indent_spaces * (t_level + 1)
        # 为表格的每一行添加缩进
        indented_table = "\n".join(t_indent + line for line in table_md.split("\n"))
        logger.info(f"插入表格 {table_index + 1}，在列表上下文中，level={t_level}")
        return indented_table
    else:
        logger.info(f"插入表格 {table_index + 1}")
        return table_md
