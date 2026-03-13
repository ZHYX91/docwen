"""docx_spell 单元测试。"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docwen.docx_spell.core import add_comments_for_errors, create_validator_with_options, process_docx
from docwen.docx_spell.spell_checker import TextError
from docwen.docx_spell.utils import plan_run_splits, rebuild_paragraph_with_splits


pytestmark = pytest.mark.unit


@pytest.mark.unit
def test_process_docx_missing_file_raises(tmp_path: Path) -> None:
    missing = tmp_path / "missing.docx"
    with pytest.raises(FileNotFoundError):
        process_docx(str(missing))


@pytest.mark.unit
def test_process_docx_no_rules_just_saves(tmp_path: Path) -> None:
    src = tmp_path / "input.docx"
    Document().save(src)

    out = process_docx(str(src), output_dir=str(tmp_path), proofread_options={})
    assert out is not None
    assert Path(out).exists() is True
    assert Path(out).name != src.name


@pytest.mark.unit
def test_create_validator_with_options_partial_dict_falls_back_to_config_defaults(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import docwen.docx_spell.core as core

    monkeypatch.setattr(
        core.config_manager,
        "get_proofread_engine_config",
        lambda: {
            "enable_symbol_pairing": False,
            "enable_symbol_correction": True,
            "enable_typos_rule": True,
            "enable_sensitive_word": False,
        },
        raising=True,
    )

    validator = create_validator_with_options({"typos_rule": False})

    assert validator.overrides == {
        "enable_symbol_pairing": False,
        "enable_symbol_correction": True,
        "enable_typos_rule": False,
        "enable_sensitive_word": False,
    }


@pytest.mark.unit
def test_rebuild_paragraph_preserves_numPr() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("这是一个列表项，包含错误。")

    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "1")
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)

    errors = [
        TextError(
            start_pos=10,
            end_pos=12,
            error_text="错误",
            suggestion="修正",
            error_type="测试",
            source="unit",
        )
    ]

    split_plan = plan_run_splits(paragraph, errors)
    new_paragraph, _ = rebuild_paragraph_with_splits(paragraph, split_plan, doc)

    assert new_paragraph._element.xpath(".//w:numPr")


@pytest.mark.unit
def test_non_text_paragraph_skips_rebuild(monkeypatch: pytest.MonkeyPatch) -> None:
    class DummyDoc:
        def __init__(self) -> None:
            self.comments = []

        def add_comment(self, run, text: str, author: str, initials: str) -> None:
            self.comments.append((run, text, author, initials))

    def _should_not_be_called(*args, **kwargs):
        raise AssertionError("rebuild_paragraph_with_splits 不应在非文本段落策略下被调用")

    monkeypatch.setattr("docwen.docx_spell.core.rebuild_paragraph_with_splits", _should_not_be_called)

    doc = DummyDoc()
    paragraph = Document().add_paragraph("含图片的段落文本")
    drawing = OxmlElement("w:drawing")
    paragraph.runs[0]._element.append(drawing)

    errors = [
        TextError(
            start_pos=0,
            end_pos=1,
            error_text="含",
            suggestion="含",
            error_type="测试",
            source="unit",
        )
    ]

    assert add_comments_for_errors(paragraph, doc, errors) == 1
    assert len(doc.comments) == 1


@pytest.mark.unit
def test_non_text_paragraph_skips_rebuild_for_math(monkeypatch: pytest.MonkeyPatch) -> None:
    class DummyDoc:
        def __init__(self) -> None:
            self.comments = []

        def add_comment(self, run, text: str, author: str, initials: str) -> None:
            self.comments.append((run, text, author, initials))

    def _should_not_be_called(*args, **kwargs):
        raise AssertionError("rebuild_paragraph_with_splits 不应在非文本段落策略下被调用")

    monkeypatch.setattr("docwen.docx_spell.core.rebuild_paragraph_with_splits", _should_not_be_called)

    doc = DummyDoc()
    paragraph = Document().add_paragraph("含公式的段落文本")
    paragraph.runs[0]._element.append(OxmlElement("w:oMath"))

    errors = [
        TextError(
            start_pos=0,
            end_pos=1,
            error_text="含",
            suggestion="含",
            error_type="测试",
            source="unit",
        )
    ]

    assert add_comments_for_errors(paragraph, doc, errors) == 1
    assert len(doc.comments) == 1


@pytest.mark.unit
def test_non_text_detection_exception_is_conservative(monkeypatch: pytest.MonkeyPatch) -> None:
    import docwen.docx_spell.core as core

    class DummyDoc:
        def __init__(self) -> None:
            self.comments = []

        def add_comment(self, run, text: str, author: str, initials: str) -> None:
            self.comments.append((run, text, author, initials))

    def _should_not_be_called(*args, **kwargs):
        raise AssertionError("rebuild_paragraph_with_splits 不应在检测异常时被调用")

    monkeypatch.setattr("docwen.docx_spell.core.rebuild_paragraph_with_splits", _should_not_be_called)
    monkeypatch.delitem(core.DOCX_NAMESPACES, "m", raising=False)

    doc = DummyDoc()
    paragraph = Document().add_paragraph("普通段落文本")

    errors = [
        TextError(
            start_pos=0,
            end_pos=1,
            error_text="普",
            suggestion="普",
            error_type="测试",
            source="unit",
        )
    ]

    assert add_comments_for_errors(paragraph, doc, errors) == 1
    assert len(doc.comments) == 1


@pytest.mark.unit
def test_non_text_paragraph_skips_rebuild_for_hyperlink(monkeypatch: pytest.MonkeyPatch) -> None:
    class DummyDoc:
        def __init__(self) -> None:
            self.comments = []

        def add_comment(self, run, text: str, author: str, initials: str) -> None:
            self.comments.append((run, text, author, initials))

    def _should_not_be_called(*args, **kwargs):
        raise AssertionError("rebuild_paragraph_with_splits 不应在非文本段落策略下被调用")

    monkeypatch.setattr("docwen.docx_spell.core.rebuild_paragraph_with_splits", _should_not_be_called)

    doc = DummyDoc()
    paragraph = Document().add_paragraph("含超链接的段落文本")
    paragraph._element.append(OxmlElement("w:hyperlink"))

    errors = [
        TextError(
            start_pos=0,
            end_pos=1,
            error_text="含",
            suggestion="含",
            error_type="测试",
            source="unit",
        )
    ]

    assert add_comments_for_errors(paragraph, doc, errors) == 1
    assert len(doc.comments) == 1


@pytest.mark.unit
def test_rebuild_paragraph_preserves_tab_and_line_break() -> None:
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("前\t中\n后")

    text = paragraph.text
    start = text.index("中")
    errors = [
        TextError(
            start_pos=start,
            end_pos=start + 1,
            error_text="中",
            suggestion="中",
            error_type="测试",
            source="unit",
        )
    ]

    split_plan = plan_run_splits(paragraph, errors)
    new_paragraph, _ = rebuild_paragraph_with_splits(paragraph, split_plan, doc)

    assert new_paragraph._element.xpath(".//w:tab")
    assert new_paragraph._element.xpath(".//w:br")


@pytest.mark.unit
def test_composite_scenarios_keep_structures() -> None:
    class DocWithComments:
        def __init__(self, inner) -> None:
            self._inner = inner
            self.comments = []

        def add_paragraph(self, *args, **kwargs):
            return self._inner.add_paragraph(*args, **kwargs)

        def add_comment(self, run, text: str, author: str, initials: str) -> None:
            self.comments.append((run, text, author, initials))

        @property
        def _element(self):
            return self._inner._element

    inner = Document()
    doc = DocWithComments(inner)

    p1 = doc.add_paragraph("这是一个列表项，包含错误。")
    pPr = p1._element.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "1")
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)

    text1 = p1.text
    start1 = text1.index("错误")
    errors1 = [
        TextError(
            start_pos=start1,
            end_pos=start1 + 2,
            error_text="错误",
            suggestion="修正",
            error_type="测试",
            source="unit",
        )
    ]
    assert add_comments_for_errors(p1, doc, errors1) == 1

    p2 = doc.add_paragraph("含超链接段落错误")
    p2._element.append(OxmlElement("w:hyperlink"))
    errors2 = [
        TextError(
            start_pos=0,
            end_pos=1,
            error_text="含",
            suggestion="含",
            error_type="测试",
            source="unit",
        )
    ]
    assert add_comments_for_errors(p2, doc, errors2) == 1

    p3 = doc.add_paragraph()
    p3.add_run("前\t中\n后")
    text3 = p3.text
    start3 = text3.index("中")
    errors3 = [
        TextError(
            start_pos=start3,
            end_pos=start3 + 1,
            error_text="中",
            suggestion="中",
            error_type="测试",
            source="unit",
        )
    ]
    assert add_comments_for_errors(p3, doc, errors3) == 1

    assert len(doc.comments) == 3
    assert inner._element.xpath(".//w:numPr")
    assert inner._element.xpath(".//w:hyperlink")
    assert inner._element.xpath(".//w:tab")
    assert inner._element.xpath(".//w:br")
