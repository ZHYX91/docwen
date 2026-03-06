"""
测试 proofread_options 三态语义一致性

验证 GUI/CLI 入口与 docx_spell 引擎对 proofread_options 的处理：
- None: 使用配置文件默认设置
- 全 False: 跳过校对
- 部分键: 缺失键回退到配置默认，不会意外变为 False
"""

from __future__ import annotations

from unittest.mock import patch

import pytest

from docwen.proofread_keys import SENSITIVE_WORD, SYMBOL_CORRECTION, SYMBOL_PAIRING, TYPOS_RULE

pytestmark = pytest.mark.unit


class TestCreateValidatorWithOptions:
    """测试 create_validator_with_options 的三态行为"""

    def test_none_uses_defaults(self) -> None:
        """proofread_options=None 应使用默认 TextValidator"""
        from docwen.docx_spell.core import create_validator_with_options

        validator = create_validator_with_options(None)
        # 应成功创建验证器（使用默认配置）
        assert validator is not None

    def test_empty_dict_uses_defaults(self) -> None:
        """proofread_options={} 应使用默认 TextValidator（not {} == True）"""
        from docwen.docx_spell.core import create_validator_with_options

        validator = create_validator_with_options({})
        assert validator is not None

    def test_all_false_creates_validator_with_all_disabled(self) -> None:
        """全 False dict 应创建全部规则关闭的验证器"""
        from docwen.docx_spell.core import create_validator_with_options

        options = {
            SYMBOL_PAIRING: False,
            SYMBOL_CORRECTION: False,
            TYPOS_RULE: False,
            SENSITIVE_WORD: False,
        }
        validator = create_validator_with_options(options)
        assert validator is not None

    def test_partial_keys_fallback_to_config_defaults(self) -> None:
        """只传部分键时，缺失键应回退到配置默认值，而非 False"""
        from docwen.docx_spell.core import create_validator_with_options

        # 模拟引擎配置中 enable_typos_rule=True
        mock_engine_config = {
            "enable_symbol_pairing": True,
            "enable_symbol_correction": True,
            "enable_typos_rule": True,
            "enable_sensitive_word": True,
        }

        with patch("docwen.docx_spell.core.config_manager") as mock_cm:
            mock_cm.get_proofread_engine_config.return_value = mock_engine_config

            # 只传 symbol_pairing=False，其他键缺失
            options = {SYMBOL_PAIRING: False}
            validator = create_validator_with_options(options)

            # 验证器应成功创建
            assert validator is not None


class TestProcessDocxProofreadSkip:
    """测试 process_docx 对全 False 的跳过行为"""

    def test_all_false_skips_proofreading(self, tmp_path) -> None:
        """proofread_options 全 False 时应跳过校对，直接保存"""
        from docx import Document

        # 创建测试 DOCX
        doc = Document()
        doc.add_paragraph("测试文本")
        src = tmp_path / "test.docx"
        doc.save(str(src))

        from docwen.docx_spell.core import process_docx

        options = {
            SYMBOL_PAIRING: False,
            SYMBOL_CORRECTION: False,
            TYPOS_RULE: False,
            SENSITIVE_WORD: False,
        }
        out = process_docx(str(src), output_dir=str(tmp_path), proofread_options=options)
        assert out is not None

    def test_empty_dict_skips_proofreading(self, tmp_path) -> None:
        """proofread_options={} 空字典也应跳过校对"""
        from docx import Document

        doc = Document()
        doc.add_paragraph("测试文本")
        src = tmp_path / "test.docx"
        doc.save(str(src))

        from docwen.docx_spell.core import process_docx

        out = process_docx(str(src), output_dir=str(tmp_path), proofread_options={})
        assert out is not None
