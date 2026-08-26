"""Deterministic real-XPS builders and normalized PDF projections for tests."""

from __future__ import annotations

import hashlib
from collections.abc import Iterable
from io import BytesIO
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile

from PIL import Image, ImageDraw, ImageFont

_MINIMAL_XPS_PARTS = {
    "[Content_Types].xml": """<?xml version="1.0" encoding="utf-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Override PartName="/FixedDocumentSequence.fdseq" ContentType="application/vnd.ms-package.xps-fixeddocumentsequence+xml"/>
  <Override PartName="/Documents/1/FixedDocument.fdoc" ContentType="application/vnd.ms-package.xps-fixeddocument+xml"/>
  <Override PartName="/Documents/1/Pages/1.fpage" ContentType="application/vnd.ms-package.xps-fixedpage+xml"/>
  <Override PartName="/Documents/1/Pages/2.fpage" ContentType="application/vnd.ms-package.xps-fixedpage+xml"/>
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
</Types>""",
    "_rels/.rels": """<?xml version="1.0" encoding="utf-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="R1" Type="http://schemas.microsoft.com/xps/2005/06/fixedrepresentation" Target="/FixedDocumentSequence.fdseq"/>
</Relationships>""",
    "FixedDocumentSequence.fdseq": """<FixedDocumentSequence xmlns="http://schemas.microsoft.com/xps/2005/06">
  <DocumentReference Source="Documents/1/FixedDocument.fdoc"/>
</FixedDocumentSequence>""",
    "Documents/1/FixedDocument.fdoc": """<FixedDocument xmlns="http://schemas.microsoft.com/xps/2005/06">
  <PageContent Source="Pages/1.fpage"/>
  <PageContent Source="Pages/2.fpage"/>
</FixedDocument>""",
    "Documents/1/Pages/1.fpage": """<FixedPage xmlns="http://schemas.microsoft.com/xps/2005/06" Width="816" Height="1056" xml:lang="en-US">
  <Path Fill="#FFDC143C" Data="M96,128 L336,128 336,288 96,288 Z"/>
  <Path Fill="#FFFFD700" Data="M420,240 C420,273.14 393.14,300 360,300 326.86,300 300,273.14 300,240 300,206.86 326.86,180 360,180 393.14,180 420,206.86 420,240 Z"/>
</FixedPage>""",
    "Documents/1/Pages/2.fpage": """<FixedPage xmlns="http://schemas.microsoft.com/xps/2005/06" Width="816" Height="1056" xml:lang="en-US">
  <Path Fill="#FF4169E1" Data="M96,128 L336,128 336,288 96,288 Z"/>
  <Path Fill="#FFFFD700" Data="M420,420 C420,453.14 393.14,480 360,480 326.86,480 300,453.14 300,420 300,386.86 326.86,360 360,360 393.14,360 420,386.86 420,420 Z"/>
</FixedPage>""",
}


def create_minimal_xps(path: Path) -> None:
    """Create a two-page, font-free XPS package with distinct vector geometry."""
    with ZipFile(path, "w", ZIP_DEFLATED) as archive:
        for name, body in _MINIMAL_XPS_PARTS.items():
            archive.writestr(name, body)


def create_image_xps(path: Path) -> None:
    """Create a real two-page XPS whose pages contain deterministic PNGs."""
    page_images: dict[str, bytes] = {}
    for page_number, label in enumerate(("DOCWEN XPS PAGE ONE", "DOCWEN XPS PAGE TWO"), start=1):
        image = Image.new("RGB", (640, 320), "white")
        draw = ImageDraw.Draw(image)
        draw.rectangle((24, 24, 616, 296), outline=(20, 70, 140), width=6)
        draw.text((64, 128), label, fill="black", font=ImageFont.load_default(size=16))
        image_path = Path(f"page-{page_number}.png")
        buffer = BytesIO()
        image.save(buffer, format="PNG")
        page_images[f"Resources/Images/{image_path.name}"] = buffer.getvalue()

    parts = dict(_MINIMAL_XPS_PARTS)
    parts["[Content_Types].xml"] = parts["[Content_Types].xml"].replace(
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>',
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>\n'
        '  <Default Extension="png" ContentType="image/png"/>',
    )
    for page_number in (1, 2):
        parts[
            f"Documents/1/Pages/{page_number}.fpage"
        ] = f"""<FixedPage xmlns="http://schemas.microsoft.com/xps/2005/06" Width="816" Height="1056" xml:lang="en-US">
  <Path Data="M88,220 L728,220 728,540 88,540 Z">
    <Path.Fill>
      <ImageBrush ImageSource="/Resources/Images/page-{page_number}.png"
                  Viewbox="0,0,640,320" ViewboxUnits="Absolute"
                  Viewport="88,220,640,320" ViewportUnits="Absolute" TileMode="None"/>
    </Path.Fill>
  </Path>
</FixedPage>"""

    with ZipFile(path, "w", ZIP_DEFLATED) as archive:
        for name, body in parts.items():
            archive.writestr(name, body)
        for name, body in page_images.items():
            archive.writestr(name, body)


def pdf_visual_projection(path: str | Path) -> dict[str, Any]:
    """Project stable page geometry, vector fills, and rendered pixels from a PDF."""
    import fitz

    pages: list[dict[str, Any]] = []
    with fitz.open(str(path)) as document:
        for page in document:
            drawings = [
                {
                    "rect": [round(value, 3) for value in drawing["rect"]],
                    "fill": [round(value, 6) for value in drawing["fill"]],
                }
                for drawing in page.get_drawings()
            ]
            pixmap = page.get_pixmap(alpha=False)
            pages.append(
                {
                    "rect": [round(value, 3) for value in page.rect],
                    "text": page.get_text("text").strip(),
                    "drawings": drawings,
                    "pix_sha256": hashlib.sha256(pixmap.samples).hexdigest(),
                }
            )
    return {
        "pdf_magic": Path(path).read_bytes()[:5].decode("ascii"),
        "page_count": len(pages),
        "pages": pages,
    }


def png_visual_projection(paths: Iterable[str | Path]) -> list[dict[str, Any]]:
    """Project decoded RGBA pixels and basic image properties from PNGs."""
    from PIL import Image

    projection: list[dict[str, Any]] = []
    for page_number, path in enumerate(paths, start=1):
        with Image.open(path) as image:
            rgba = image.convert("RGBA")
            projection.append(
                {
                    "page": page_number,
                    "format": image.format,
                    "mode": image.mode,
                    "size": list(image.size),
                    "alpha_extrema": list(rgba.getchannel("A").getextrema()),
                    "rgba_sha256": hashlib.sha256(rgba.tobytes()).hexdigest(),
                }
            )
    return projection


def raster_visual_projection(paths: Iterable[str | Path]) -> dict[str, Any]:
    """Project decoded RGB pixels across one or more raster containers/frames."""
    from PIL import Image

    containers: list[str] = []
    frames: list[dict[str, Any]] = []
    for path in paths:
        with Image.open(path) as image:
            containers.append(str(image.format))
            for frame_index in range(int(getattr(image, "n_frames", 1))):
                image.seek(frame_index)
                rgb = image.convert("RGB")
                frames.append(
                    {
                        "page": len(frames) + 1,
                        "format": image.format,
                        "mode": image.mode,
                        "size": list(image.size),
                        "rgb_sha256": hashlib.sha256(rgb.tobytes()).hexdigest(),
                    }
                )
    return {
        "container_count": len(containers),
        "container_formats": containers,
        "frame_count": len(frames),
        "frames": frames,
    }


def docx_semantic_projection(path: str | Path) -> dict[str, Any]:
    """Project DOCX structure and decoded media without relying on raw ZIP bytes."""
    from docx import Document
    from lxml import etree

    document = Document(str(path))
    with ZipFile(path) as archive:
        names = sorted(archive.namelist())
        document_xml = etree.fromstring(archive.read("word/document.xml"))
        media: list[dict[str, Any]] = []
        for name in (item for item in names if item.startswith("word/media/")):
            encoded = archive.read(name)
            with Image.open(BytesIO(encoded)) as image:
                rgba = image.convert("RGBA")
                media.append(
                    {
                        "format": image.format,
                        "size": list(image.size),
                        "rgba_sha256": hashlib.sha256(rgba.tobytes()).hexdigest(),
                    }
                )

        word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        drawing_namespace = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        return {
            "zip_valid": archive.testzip() is None,
            "package_part_count": len(names),
            "required_parts_present": all(
                name in names
                for name in (
                    "[Content_Types].xml",
                    "_rels/.rels",
                    "word/document.xml",
                    "word/_rels/document.xml.rels",
                )
            ),
            "paragraph_count": len(document.paragraphs),
            "paragraph_texts": [paragraph.text for paragraph in document.paragraphs],
            "section_count": len(document.sections),
            "sections": [
                {
                    "width_emu": section.page_width,
                    "height_emu": section.page_height,
                    "orientation": int(section.orientation),
                }
                for section in document.sections
            ],
            "table_count": len(document.tables),
            "inline_shape_count": len(document.inline_shapes),
            "page_break_count": len(document_xml.xpath(".//w:br[@w:type='page']", namespaces={"w": word_namespace})),
            "anchored_shape_count": len(document_xml.xpath(".//wp:anchor", namespaces={"wp": drawing_namespace})),
            "drawing_count": len(
                document_xml.xpath(".//wp:inline | .//wp:anchor", namespaces={"wp": drawing_namespace})
            ),
            "media": media,
        }
