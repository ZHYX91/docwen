"""converter 单元测试。"""

from __future__ import annotations

import os
import shutil
from pathlib import Path

import pytest
from docx import Document

from docwen.converter.docx2md.shared.content_injector import save_extracted_document

pytestmark = pytest.mark.unit


def test_save_extracted_document_uniquifies_on_timestamp_collision(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr("docwen.utils.path_utils.generate_timestamp", lambda: "20200101_000000", raising=True)

    base_dir = tmp_path / "base_dir"
    base_dir.mkdir()
    original = base_dir / "input.docx"
    Document().save(original)

    out_dir = tmp_path / "out"
    out_dir.mkdir()

    doc1 = Document()
    doc1.add_paragraph("x")
    p1 = save_extracted_document(doc1, str(original), str(out_dir))

    doc2 = Document()
    doc2.add_paragraph("y")
    p2 = save_extracted_document(doc2, str(original), str(out_dir))

    assert p1 is not None
    assert p2 is not None
    assert Path(p1).exists() is True
    assert Path(p2).exists() is True
    assert p1 != p2
    assert Path(p2).stem.endswith("_001") is True


def test_save_extracted_document_falls_back_when_target_move_fails(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr("docwen.utils.path_utils.generate_timestamp", lambda: "20200101_000000", raising=True)

    base_dir = tmp_path / "base_dir"
    base_dir.mkdir()
    original = base_dir / "input.docx"
    Document().save(original)

    bad_out = tmp_path / "bad_out"
    bad_out.mkdir()

    def _stub_move_file_with_retry(source: str, destination: str, *_args, **_kwargs):
        if os.path.normcase(destination).startswith(os.path.normcase(str(bad_out))):
            return None
        Path(destination).parent.mkdir(parents=True, exist_ok=True)
        shutil.move(source, destination)
        return destination

    monkeypatch.setattr("docwen.utils.workspace_manager.move_file_with_retry", _stub_move_file_with_retry, raising=True)

    doc = Document()
    doc.add_paragraph("x")
    saved = save_extracted_document(doc, str(original), str(bad_out))

    assert saved is not None
    out = Path(saved)
    assert out.exists() is True
    assert os.path.normcase(str(base_dir)) in os.path.normcase(str(out.parent))
