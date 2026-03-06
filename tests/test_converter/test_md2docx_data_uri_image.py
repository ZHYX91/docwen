"""converter 单元测试。"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from docwen.converter.md2docx import core as md2docx_core

pytestmark = pytest.mark.unit


_BASE64_PNG_1X1 = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="


def test_md2docx_embeds_data_uri_image(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    template = tmp_path / "template.docx"
    tpl_doc = Document()
    tpl_doc.add_paragraph("{{正文}}")
    tpl_doc.save(template)

    monkeypatch.setattr(
        "docwen.template.loader.TemplateLoader.get_template_path",
        lambda *_args, **_kwargs: str(template),
        raising=True,
    )
    monkeypatch.setattr(
        "docwen.template.loader.TemplateLoader.load_docx_template",
        lambda *_args, **_kwargs: Document(str(template)),
        raising=True,
    )

    monkeypatch.setattr("docwen.utils.link_processing.config_manager.get_max_embed_depth", lambda: 10, raising=True)
    monkeypatch.setattr(
        "docwen.utils.link_processing.config_manager.get_markdown_embed_image_mode",
        lambda: "embed",
        raising=True,
    )
    monkeypatch.setattr(
        "docwen.utils.link_processing.config_manager.get_wiki_embed_image_mode", lambda: "embed", raising=True
    )
    monkeypatch.setattr("docwen.utils.link_processing.config_manager.get_wiki_link_mode", lambda: "keep", raising=True)
    monkeypatch.setattr(
        "docwen.utils.link_processing.config_manager.get_markdown_link_mode", lambda: "keep", raising=True
    )

    md = f"# Title\n\n![img](data:image/png;base64,{_BASE64_PNG_1X1})\n"
    md_path = tmp_path / "a.md"
    md_path.write_text(md, encoding="utf-8")

    out = tmp_path / "out.docx"
    result = md2docx_core.convert(
        md_path=str(md_path),
        output_path=str(out),
        template_name="dummy",
        progress_callback=None,
        cancel_event=None,
        original_source_path=str(md_path),
        options={},
    )

    assert result is not None
    assert out.exists() is True

    doc = Document(str(out))
    assert len(doc.inline_shapes) >= 1
