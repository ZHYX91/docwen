"""公文字段处理器重构测试。"""

from __future__ import annotations

import copy

import pytest
from docx import Document
from docx.shared import Pt

from docwen.converter.md2docx.field_registry import (
    get_active_special_placeholders,
    get_merged_placeholder_rules,
    run_special_handlers,
)
from docwen.converter.md2docx.fields.gongwen import (
    format_date,
    process_attachment_description,
    process_cc_orgs,
    process_gongwen_yaml,
    process_notes,
)
from docwen.converter.md2docx.processors.placeholder_handler import is_special_marked, mark_special_placeholders

pytestmark = pytest.mark.unit


def test_process_attachment_description_single_attachment() -> None:
    data = {"附件说明": "实施方案"}
    process_attachment_description(data)
    assert data["附件说明"] == ["附件：实施方案"]


def test_process_attachment_description_multiple_attachments() -> None:
    data = {"附件说明": ["实施方案", "经费预算", "人员名单"]}
    process_attachment_description(data)
    assert data["附件说明"][0] == "附件：1. 实施方案"
    assert data["附件说明"][1].startswith("\u3000")


def test_process_attachment_description_strip_existing_numbering() -> None:
    data = {"附件说明": ["1. 实施方案", "2. 经费预算"]}
    process_attachment_description(data)
    assert "1. 1." not in data["附件说明"][0]


def test_process_cc_orgs_join_by_cn_comma() -> None:
    data = {"抄送机关": ["市发改委", "市财政局", "市教育局"]}
    process_cc_orgs(data)
    assert data["抄送机关"] == "市发改委，市财政局，市教育局"


def test_format_date_and_notes() -> None:
    assert format_date("2025-01-06") == "2025年1月6日"
    assert format_date("2025-01-06", suffix="印发") == "2025年1月6日印发"
    assert process_notes("（此件公开发布）") == "此件公开发布"


def test_process_gongwen_yaml_snapshot() -> None:
    data = {
        "标题": "关于做好2025年工作的通知",
        "发文字号": "某发〔2025〕1号",
        "主送机关": "各县（市、区）人民政府",
        "附件说明": ["实施方案", "经费预算"],
        "抄送机关": ["市发改委", "市财政局"],
        "附注": "（此件公开发布）",
        "印发日期": "2025-01-06",
        "成文日期": "2025-01-05",
    }
    copied = copy.deepcopy(data)
    process_gongwen_yaml(copied)
    assert copied["附件说明"] == ["附件：1. 实施方案", "\u3000\u3000\u30002. 经费预算"]
    assert copied["抄送机关"] == "市发改委，市财政局"
    assert copied["附注"] == "此件公开发布"
    assert copied["印发日期"] == "2025年1月6日印发"
    assert copied["成文日期"] == "2025年1月5日"


def test_field_registry_locale_filter_for_rules() -> None:
    zh_rules = get_merged_placeholder_rules("zh_CN")
    en_rules = get_merged_placeholder_rules("en_US")
    assert len(zh_rules.get("delete_paragraph_if_empty", [])) > 0
    assert len(en_rules.get("delete_paragraph_if_empty", [])) == 0


def test_field_registry_locale_filter_for_special_placeholders() -> None:
    zh_special = get_active_special_placeholders("zh_CN")
    en_special = get_active_special_placeholders("en_US")
    assert "{{附件说明}}" in zh_special
    assert "{{附件说明}}" not in en_special


def test_mark_special_placeholders_with_extra_placeholders() -> None:
    doc = Document()
    para = doc.add_paragraph("这里是 {{附件说明}}")
    mark_special_placeholders(doc, {"{{附件说明}}"})
    assert is_special_marked(para)


def test_run_special_handlers_processes_attachment_placeholder() -> None:
    doc = Document()
    para = doc.add_paragraph("{{附件说明}}")
    para.paragraph_format.left_indent = Pt(32)

    yaml_data = {"附件说明": ["实施方案", "经费预算"]}
    process_attachment_description(yaml_data)
    run_special_handlers(doc, yaml_data, current_locale="zh_CN")

    assert "{{附件说明}}" not in "\n".join(p.text for p in doc.paragraphs)
    assert "附件：1. 实施方案" in doc.paragraphs[0].text


def test_end_to_end_special_placeholder_chain(monkeypatch) -> None:
    from docwen.converter.md2docx.processors import docx_processor

    monkeypatch.setattr(
        docx_processor, "save_and_process_temp_file", lambda doc, yaml_data, *, note_ctx=None: "dummy.docx"
    )
    monkeypatch.setattr(
        docx_processor,
        "process_main_content",
        lambda doc, body_data, yaml_data, template_name=None, *, note_ctx=None: True,
    )

    doc = Document()
    doc.add_paragraph("{{标题}}")
    para = doc.add_paragraph("{{附件说明}}")
    para.paragraph_format.left_indent = Pt(32)
    doc.add_paragraph("{{正文}}")

    yaml_data = {"标题": "测试标题", "附件说明": ["实施方案", "经费预算"]}
    process_attachment_description(yaml_data)

    _, warnings = docx_processor.replace_placeholders(doc, yaml_data, body_data=[], template_name=None)
    assert warnings == []
    assert "测试标题" in "\n".join(p.text for p in doc.paragraphs)
    assert "{{附件说明}}" not in "\n".join(p.text for p in doc.paragraphs)

    attachment_paras = [p for p in doc.paragraphs if "附件：" in p.text]
    assert attachment_paras
    assert attachment_paras[0].paragraph_format.first_line_indent is not None
    assert attachment_paras[0].paragraph_format.first_line_indent < 0


def test_end_to_end_special_hanging_indent_single_vs_multi(monkeypatch) -> None:
    from docwen.converter.md2docx.processors import docx_processor

    monkeypatch.setattr(
        docx_processor, "save_and_process_temp_file", lambda doc, yaml_data, *, note_ctx=None: "dummy.docx"
    )
    monkeypatch.setattr(
        docx_processor,
        "process_main_content",
        lambda doc, body_data, yaml_data, template_name=None, *, note_ctx=None: True,
    )

    doc_single = Document()
    para_single = doc_single.add_paragraph("{{附件说明}}")
    para_single.paragraph_format.left_indent = Pt(32)
    yaml_single = {"附件说明": ["实施方案"]}
    process_attachment_description(yaml_single)
    docx_processor.replace_placeholders(doc_single, yaml_single, body_data=[], template_name=None)
    single_paras = [p for p in doc_single.paragraphs if "附件：" in p.text]
    assert single_paras

    doc_multi = Document()
    para_multi = doc_multi.add_paragraph("{{附件说明}}")
    para_multi.paragraph_format.left_indent = Pt(32)
    yaml_multi = {"附件说明": ["实施方案", "经费预算"]}
    process_attachment_description(yaml_multi)
    docx_processor.replace_placeholders(doc_multi, yaml_multi, body_data=[], template_name=None)
    multi_paras = [p for p in doc_multi.paragraphs if "附件：" in p.text]
    assert multi_paras

    single_indent = abs(int(single_paras[0].paragraph_format.first_line_indent or 0))
    multi_indent = abs(int(multi_paras[0].paragraph_format.first_line_indent or 0))
    assert multi_indent > single_indent
