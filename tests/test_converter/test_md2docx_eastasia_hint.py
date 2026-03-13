"""converter 单元测试。"""

from __future__ import annotations

import pytest
from docx import Document
from docx.oxml.ns import qn

from docwen.converter.md2docx.handlers.note_handler import WORD_NS, create_footnote_element
from docwen.converter.md2docx.handlers.text_handler import (
    add_formatted_text_to_heading,
    add_formatted_text_to_paragraph,
)
from docwen.utils.docx_utils import DEFAULT_FONTS

pytestmark = pytest.mark.unit


def test_md2docx_body_run_sets_eastasia_hint_for_chinese_text() -> None:
    doc = Document()
    p = doc.add_paragraph()

    chinese = "这是一段包含“中文双引号”的文本。省略号……以及破折号——测试。"
    english = 'English text with "English quotes" should stay western.'

    add_formatted_text_to_paragraph(p, chinese, base_fonts=DEFAULT_FONTS.copy(), mode="apply", doc=doc)
    add_formatted_text_to_paragraph(p, english, base_fonts=DEFAULT_FONTS.copy(), mode="apply", doc=doc)

    run_cn = next(r for r in p.runs if "中文双引号" in r.text)
    run_en = next(r for r in p.runs if "English text" in r.text)

    hint_cn = run_cn._element.rPr.rFonts.get(qn("w:hint"))
    hint_en = run_en._element.rPr.rFonts.get(qn("w:hint"))

    assert hint_cn == "eastAsia"
    assert hint_en is None


def test_md2docx_body_run_sets_eastasia_hint_for_general_punctuation_only() -> None:
    doc = Document()
    p = doc.add_paragraph()

    punctuation = "“”……——"
    add_formatted_text_to_paragraph(p, punctuation, base_fonts=DEFAULT_FONTS.copy(), mode="apply", doc=doc)

    run = next(r for r in p.runs if r.text == punctuation)
    assert run._element.rPr.rFonts.get(qn("w:hint")) == "eastAsia"


def test_md2docx_heading_run_sets_eastasia_hint_for_chinese_heading() -> None:
    doc = Document()
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]

    heading = "做好“加减乘除”四则运算"
    add_formatted_text_to_heading(p, heading, mode="apply", doc=doc)

    run = next(r for r in p.runs if heading in r.text)
    assert run._element.rPr.rFonts.get(qn("w:hint")) == "eastAsia"


@pytest.mark.parametrize(
    "text",
    [
        "これはテストです「引用」",
        "테스트 “인용” 입니다",
        "。，！？“”",
    ],
)
def test_md2docx_body_run_sets_eastasia_hint_for_east_asian_context(text: str) -> None:
    doc = Document()
    p = doc.add_paragraph()

    add_formatted_text_to_paragraph(p, text, base_fonts=DEFAULT_FONTS.copy(), mode="apply", doc=doc)
    run = next(r for r in p.runs if r.text == text)

    assert run._element.rPr.rFonts.get(qn("w:hint")) == "eastAsia"


def test_md2docx_footnote_run_sets_eastasia_hint_for_chinese_text() -> None:
    footnote = create_footnote_element(1, "脚注内容：“中文引号”。")

    ns = {"w": WORD_NS}
    text_nodes = footnote.xpath(".//w:t", namespaces=ns)
    cn_node = next(n for n in text_nodes if "中文引号" in (n.text or ""))

    run = cn_node.getparent()
    assert run is not None

    rpr = run.find(f"{{{WORD_NS}}}rPr")
    assert rpr is not None

    rfonts = rpr.find(f"{{{WORD_NS}}}rFonts")
    assert rfonts is not None

    assert rfonts.get(f"{{{WORD_NS}}}hint") == "eastAsia"


@pytest.mark.parametrize(
    "content",
    [
        "これはテストです「引用」",
        "테스트 “인용” 입니다",
        "。，！？“”",
    ],
)
def test_md2docx_footnote_run_sets_eastasia_hint_for_east_asian_context(content: str) -> None:
    footnote = create_footnote_element(1, content)

    ns = {"w": WORD_NS}
    text_nodes = footnote.xpath(".//w:t", namespaces=ns)
    node = next(n for n in text_nodes if (n.text or "") == content)

    run = node.getparent()
    assert run is not None

    rpr = run.find(f"{{{WORD_NS}}}rPr")
    assert rpr is not None

    rfonts = rpr.find(f"{{{WORD_NS}}}rFonts")
    assert rfonts is not None

    assert rfonts.get(f"{{{WORD_NS}}}hint") == "eastAsia"
