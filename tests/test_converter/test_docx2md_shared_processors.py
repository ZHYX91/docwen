"""converter 单元测试。"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from lxml import etree
from PIL import Image

from docwen.converter.docx2md.shared import formula_processor as fp
from docwen.converter.docx2md.shared import image_processor as ip

pytestmark = pytest.mark.unit


def test_is_boolean_format_enabled() -> None:
    rpr = etree.Element(f"{{{fp.WORD_NS}}}rPr")
    assert fp._is_boolean_format_enabled(rpr, "b") is False

    etree.SubElement(rpr, f"{{{fp.WORD_NS}}}b")
    assert fp._is_boolean_format_enabled(rpr, "b") is True


def test_apply_format_markers_from_xml_bold_italic() -> None:
    run = etree.Element(f"{{{fp.WORD_NS}}}r")
    rpr = etree.SubElement(run, f"{{{fp.WORD_NS}}}rPr")
    etree.SubElement(rpr, f"{{{fp.WORD_NS}}}b")
    etree.SubElement(rpr, f"{{{fp.WORD_NS}}}i")

    out = fp._apply_format_markers_from_xml(run, "X", preserve_formatting=True, syntax_config={})
    assert out == "___X___"


def test_replace_formulas_in_text() -> None:
    text = "a [[FORMULA_0]] b"
    formulas = [{"latex": "x+y"}]
    out = fp.replace_formulas_in_text(text, formulas)
    assert "x+y" in out


@pytest.mark.slow
def test_extract_images_from_docx_extracts_one(tmp_path: Path) -> None:
    png_path = tmp_path / "img.png"
    Image.new("RGB", (1, 1), (255, 0, 0)).save(png_path)

    doc = Document()
    doc.add_picture(str(png_path))
    docx_path = tmp_path / "in.docx"
    doc.save(docx_path)

    loaded = Document(docx_path)
    out_dir = tmp_path / "images"
    images = ip.extract_images_from_docx(loaded, str(out_dir), str(docx_path))

    assert len(images) == 1
    assert Path(images[0]["image_path"]).exists() is True
