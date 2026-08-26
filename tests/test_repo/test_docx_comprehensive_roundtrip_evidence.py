from __future__ import annotations

import hashlib
import json
import zipfile
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

pytestmark = pytest.mark.contract

PROJECT_ROOT = Path(__file__).resolve().parents[2]
FIXTURE = PROJECT_ROOT / "tests" / "fixtures" / "golden" / "old_system_docx_comprehensive_roundtrip_semantics.json"
SOURCE = PROJECT_ROOT / "tests" / "fixtures" / "golden" / "md_to_docx_old" / "sample_golden.docx"


def _fixture() -> dict:
    return json.loads(FIXTURE.read_text(encoding="utf-8"))


def test_comprehensive_docx_fixture_locks_checked_in_old_artifact_profile() -> None:
    data = _fixture()
    source = data["source"]
    assert data["golden_id"] == "GOLDEN-002"
    assert data["evidence_id"] == "VIS-2026-07-15-028"
    assert SOURCE.is_file()
    assert SOURCE.stat().st_size == source["bytes"] == 39146
    assert hashlib.sha256(SOURCE.read_bytes()).hexdigest() == source["sha256"]

    document = Document(SOURCE)
    assert len(document.paragraphs) == source["paragraph_count"] == 86
    assert sum(bool(paragraph.text.strip()) for paragraph in document.paragraphs) == 82
    assert sum(paragraph.style.name.startswith("Heading ") for paragraph in document.paragraphs) == 34
    assert [[len(table.rows), len(table.columns)] for table in document.tables] == [[4, 2], [4, 2]]


def test_comprehensive_docx_fixture_locks_xml_boundaries_that_exposed_regressions() -> None:
    data = _fixture()["source"]
    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    }
    with zipfile.ZipFile(SOURCE) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))
        names = set(archive.namelist())

    assert len(root.xpath(".//m:oMath", namespaces=namespaces)) == data["omath_count"] == 3
    assert len(root.xpath(".//m:oMathPara", namespaces=namespaces)) == 0
    assert len(root.xpath(".//w:body/w:p[w:pPr/w:pBdr/w:bottom]", namespaces=namespaces)) == 3
    assert len(root.xpath(".//w:footnoteReference", namespaces=namespaces)) == 2
    assert len(root.xpath(".//w:endnoteReference", namespaces=namespaces)) == 2
    assert len(root.xpath(".//w:hyperlink", namespaces=namespaces)) == 1
    assert {"word/footnotes.xml", "word/endnotes.xml"} <= names


def test_comprehensive_docx_fixture_records_red_green_and_current_runtime_contract() -> None:
    data = _fixture()
    pre_fix = data["pre_fix_current_regressions"]
    current = data["post_fix_projects"]["docwen-current"]
    assert pre_fix["horizontal_rule_count"] == 0
    assert pre_fix["expected_horizontal_rule_count_from_both_old_projects"] == 3
    assert pre_fix["first_standalone_formula"] == "$a^{2}+b^{2}=c^{2}$"
    assert pre_fix["classification"] == "two_confirmed_current_only_regressions"

    assert current["success"] is True
    assert current["primary_name"] == "sample_golden.md"
    assert current["primary_in_requested_output_directory"] is True
    assert current["metadata"] == {
        "paragraph_count": 47,
        "heading_count": 34,
        "table_count": 2,
        "image_count": 0,
    }
    assert current["diagnostic_codes"] == ["DOCX2MD-OK", "FINALIZER_DONE"]


def test_comprehensive_docx_fixture_records_normalized_parity_and_reference_improvement() -> None:
    data = _fixture()
    normalized = data["normalized_contract"]
    projects = data["post_fix_projects"]
    assert normalized["frontmatter_dictionaries_equal_across_three_projects"] is True
    assert normalized["old_pyside6_and_current_bodies_equal_after_narrow_syntax_normalization"] is True
    assert normalized["heading_level_sequence_equal_across_three_projects"] is True
    assert normalized["horizontal_rule_count"] == 3
    assert normalized["all_three_formula_values_preserved"] is True
    assert normalized["no_unicode_replacement_characters"] is True
    assert projects["docwen-ref-tk"]["hyperlink_preserved"] is False
    assert projects["docwen-ref-pyside6"]["hyperlink_preserved"] is True
    assert projects["docwen-current"]["hyperlink_preserved"] is True


def test_comprehensive_docx_evidence_updates_actual_golden_inventory_only() -> None:
    golden_files = sorted((PROJECT_ROOT / "tests" / "fixtures" / "golden").glob("*.json"))
    assert len(golden_files) == 85
    assert FIXTURE in golden_files
    for path in golden_files:
        json.loads(path.read_text(encoding="utf-8"))
