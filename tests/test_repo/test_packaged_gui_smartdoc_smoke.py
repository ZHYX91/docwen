"""Contracts for the packaged GUI SmartDoc release smoke."""

from __future__ import annotations

import json
import subprocess
from pathlib import Path

import pytest

pytestmark = pytest.mark.contract


def _write_docx(path: Path, *, include_semantics: bool = True) -> None:
    from docx import Document

    document = Document()
    if include_semantics:
        document.add_heading("DOCWEN PACKAGED GUI SMARTDOC 2026", level=0)
        document.add_paragraph("External document bridge semantic readback.")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Item"
        table.cell(0, 1).text = "Qty"
        table.cell(1, 0).text = "Alpha"
        table.cell(1, 1).text = "7"
    document.save(path)


def test_smartdoc_fixture_uses_document_bridge_format_contracts(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from scripts.release import verify_packaged_gui

    from docwen_core.office_bridge import BridgeResult

    calls: list[tuple[str, str, list[str], list[tuple[str, int]], str | None]] = []

    def fake_convert(
        _input_path: str,
        output_path: str,
        *,
        source_format: str,
        backend_priority: list[str],
        com_candidates: dict[str, object],
        libreoffice_format: str | None,
        **_kwargs: object,
    ) -> BridgeResult:
        candidates = [(candidate.prog_id, candidate.save_format) for candidate in com_candidates.values()]  # type: ignore[attr-defined]
        calls.append((Path(output_path).suffix, source_format, backend_priority, candidates, libreoffice_format))
        Path(output_path).write_bytes(b"external document fixture")
        return BridgeResult(True, output_path=output_path, backend="fixture office")

    monkeypatch.setattr("docwen_core.office_bridge.convert_with_backend_priority", fake_convert)

    cases = verify_packaged_gui._write_smartdoc_smoke_inputs(tmp_path)

    assert [case[0] for case in cases] == ["doc", "rtf", "odt"]
    assert calls == [
        (
            ".doc",
            "docx",
            ["wps_writer", "msoffice_word", "libreoffice"],
            [("Kwps.Application", 0), ("Word.Application", 0)],
            "doc",
        ),
        (
            ".rtf",
            "docx",
            ["wps_writer", "msoffice_word", "libreoffice"],
            [("Kwps.Application", 6), ("Word.Application", 6)],
            "rtf",
        ),
        (
            ".odt",
            "docx",
            ["msoffice_word", "libreoffice"],
            [("Word.Application", 23)],
            "odt",
        ),
    ]


def test_smartdoc_smoke_drives_three_panel_routes(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from scripts.release import verify_packaged_gui

    binary_dir = tmp_path / "dist"
    binary_dir.mkdir()
    binary_name = "DocWen.exe"
    (binary_dir / binary_name).write_text("placeholder", encoding="utf-8")
    monkeypatch.setattr(verify_packaged_gui, "_verify_resource_layout", lambda _path: None)

    def fake_fixtures(work_dir: Path) -> list[tuple[str, Path, str, tuple[str, ...]]]:
        inputs = work_dir / "inputs"
        inputs.mkdir(parents=True)
        cases = []
        for source_format in ("doc", "rtf", "odt"):
            source = inputs / f"packaged-gui-smartdoc-{source_format}.{source_format}"
            source.write_bytes(b"fixture")
            cases.append(
                (
                    source_format,
                    source,
                    f"fixture {source_format}",
                    ("DOCWEN PACKAGED GUI SMARTDOC 2026", "Alpha"),
                )
            )
        return cases

    observed: list[dict[str, str]] = []

    def fake_run_with_env(
        binary_path: Path,
        *args: str,
        cwd: Path,
        env: dict[str, str],
        timeout: int,
    ) -> subprocess.CompletedProcess[str]:
        del args, timeout
        observed.append(dict(env))
        source_format = Path(env["DOCWEN_GUI_TEST_CONVERSION_INPUT"]).suffix.removeprefix(".")
        output_dir = Path(env["DOCWEN_GUI_TEST_CONVERSION_OUTPUT_DIR"])
        output_path = output_dir / f"packaged-gui-smartdoc-{source_format}.docx"
        _write_docx(output_path)
        report_path = Path(env["DOCWEN_GUI_TEST_CONVERSION_REPORT"])
        report_path.write_text(
            json.dumps(
                {
                    "success": True,
                    "status": "completed",
                    "outputPath": str(output_path),
                    "outputExists": True,
                    "error": None,
                }
            ),
            encoding="utf-8",
        )
        log_dir = cwd / "log_home" / "logs"
        log_dir.mkdir(parents=True, exist_ok=True)
        (log_dir / "docwen.log").write_text("ok", encoding="utf-8")
        return subprocess.CompletedProcess([str(binary_path)], 0, stdout="", stderr="")

    monkeypatch.setattr(verify_packaged_gui, "_write_smartdoc_smoke_inputs", fake_fixtures)
    monkeypatch.setattr(verify_packaged_gui, "_run_with_env", fake_run_with_env)
    monkeypatch.setattr(verify_packaged_gui, "_snapshot_relevant_processes", dict)
    monkeypatch.setattr(verify_packaged_gui, "_wait_for_no_new_relevant_processes", lambda _before: None)

    exit_code = verify_packaged_gui.main(
        ["--binary-dir", str(binary_dir), "--binary-name", binary_name, "--smartdoc-smoke"]
    )

    assert exit_code == 0
    assert len(observed) == 3
    assert {env["DOCWEN_GUI_TEST_CONVERSION_TARGET"] for env in observed} == {"docx"}
    assert {env["DOCWEN_GUI_TEST_CONVERSION_SURFACE"] for env in observed} == {"panel"}


def test_smartdoc_report_rejects_staging_name_and_missing_semantics(tmp_path: Path) -> None:
    from scripts.release import verify_packaged_gui

    output_path = tmp_path / "primary_1.docx"
    _write_docx(output_path, include_semantics=False)
    report_path = tmp_path / "report.json"
    report_path.write_text(
        json.dumps(
            {
                "success": True,
                "status": "completed",
                "outputPath": str(output_path),
                "outputExists": True,
                "error": None,
            }
        ),
        encoding="utf-8",
    )

    with pytest.raises(RuntimeError, match="packaged_gui_conversion_docx_name_unexpected"):
        verify_packaged_gui._verify_docx_conversion_smoke_report(
            report_path,
            expected_name="packaged-gui-smartdoc-doc.docx",
            expected_tokens=("DOCWEN PACKAGED GUI SMARTDOC 2026",),
        )

    output_path.rename(tmp_path / "packaged-gui-smartdoc-doc.docx")
    report = json.loads(report_path.read_text(encoding="utf-8"))
    report["outputPath"] = str(tmp_path / "packaged-gui-smartdoc-doc.docx")
    report_path.write_text(json.dumps(report), encoding="utf-8")
    with pytest.raises(RuntimeError, match="packaged_gui_conversion_docx_semantics_missing"):
        verify_packaged_gui._verify_docx_conversion_smoke_report(
            report_path,
            expected_name="packaged-gui-smartdoc-doc.docx",
            expected_tokens=("DOCWEN PACKAGED GUI SMARTDOC 2026",),
        )


def test_packaged_gui_runtime_diagnostics_fail_closed(tmp_path: Path) -> None:
    from scripts.release import verify_packaged_gui

    log_path = tmp_path / "docwen.log"
    log_path.write_text("Failed to load plugin docwen_plugin_document", encoding="utf-8")
    proc = subprocess.CompletedProcess(["DocWen.exe"], 0, stdout="", stderr="")

    with pytest.raises(RuntimeError, match="packaged_gui_runtime_diagnostics_failed"):
        verify_packaged_gui._verify_runtime_diagnostics(proc, log_files=[log_path])


def test_packaged_gui_process_guard_ignores_preexisting_and_rejects_new() -> None:
    from scripts.release import verify_packaged_gui

    verify_packaged_gui._verify_no_new_relevant_processes(
        {11: "wps.exe"},
        {11: "wps.exe"},
    )
    with pytest.raises(RuntimeError, match="packaged_gui_process_residue"):
        verify_packaged_gui._verify_no_new_relevant_processes(
            {11: "wps.exe"},
            {11: "wps.exe", 22: "winword.exe"},
        )
