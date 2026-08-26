"""Legacy document-semantics v1 boundaries under the v3 production route."""

from __future__ import annotations

import base64
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document

from docwen_core.docx_semantics import DocxSemanticImporter
from docwen_core.models.semantic_document import SemanticTable
from docwen_plugin_markdown.document_semantics import analyze_document_semantics
from docwen_plugin_markdown.mistune_extensions import parse_markdown_text
from docwen_plugin_markdown.renderer import MdToDocxRenderer
from docwen_plugin_markdown.yaml_processor import extract_yaml_front_matter
from tests.integration._round_trip_helper import _primary_path, _run, docx_to_md

pytestmark = pytest.mark.integration

_PIXEL_PNG = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="

_CAPTION_REFERENCE = """Figure: Blue square {#fig-blue-square}

![A two-by-two blue square](pixel.png)

As shown in @fig-blue-square, the image is intentionally synthetic.
"""

_ENHANCED_TABLE = """Table: Sales channels {#tbl-sales}

| Region | Sales | < | Total |
|---|---:|---:|---:|
| ^ | Online | Retail | ^ |
| North | 10 | 12 | 22 |
| South | 8 | 9 | 17 |
{header-rows=2 header-cols=1 repeat-header=true}

See @tbl-sales for the synthetic totals.
"""

_INVALID = """Figure: Wrong kind {#fig-wrong-kind}

| A | B |
|---|---|
| 1 | 2 |

: Broken binding {#fig-broken}

This paragraph breaks caption binding.

![Synthetic image](pixel.png)

Figure: First duplicate {#fig-repeat}

![Synthetic image](pixel.png)

Figure: Second duplicate {#fig-repeat}

![Synthetic image](pixel.png)

See @tbl-missing.

| A | < |
|---|---|
| ^ | value |
{header-rows=3 header-cols=0 repeat-header=true}
"""


def _write_source(tmp_path: Path, name: str, text: str, *, image: bool = False) -> Path:
    source_dir = tmp_path / name
    source_dir.mkdir()
    source = source_dir / f"{name}.md"
    source.write_text(text, encoding="utf-8")
    if image:
        (source_dir / "pixel.png").write_bytes(base64.b64decode(_PIXEL_PNG))
    return source


def _projection(markdown: str) -> dict[str, object]:
    _metadata, body = extract_yaml_front_matter(markdown)
    analysis = analyze_document_semantics(parse_markdown_text(body))
    assert not analysis.has_errors, {"diagnostics": analysis.diagnostics, "markdown": markdown}
    assert analysis.oracle_projection["schema"] == "docwen.document_semantics.v1"
    return analysis.oracle_projection


def _document_xml(docx_path: Path) -> str:
    with ZipFile(docx_path) as package:
        return package.read("word/document.xml").decode("utf-8")


def test_legacy_caption_reference_source_is_not_reinterpreted_by_v3(round_trip_runtime, tmp_path: Path) -> None:
    source = _write_source(tmp_path, "caption-reference", _CAPTION_REFERENCE, image=True)
    forward_dir = tmp_path / "forward"
    reverse_dir = tmp_path / "reverse"
    forward_dir.mkdir()
    reverse_dir.mkdir()

    result = _run(
        round_trip_runtime,
        "semantics-caption-forward",
        source,
        source_format="markdown",
        target_format="docx",
        output_dir=forward_dir,
    )
    docx_path = _primary_path(result)
    returned = docx_to_md(
        round_trip_runtime,
        docx_path,
        reverse_dir,
        request_id="semantics-caption-reverse",
    )

    assert "Figure: Blue square {#fig-blue-square}" in returned
    assert "As shown in @fig-blue-square" in returned
    [returned_image] = list(reverse_dir.rglob("*.png"))
    assert returned_image.name in returned
    xml = _document_xml(docx_path)
    assert "SEQ Figure" not in xml
    assert "REF _DW_" not in xml
    assert "w:bookmarkStart" not in xml
    assert 'descr="A two-by-two blue square"' in xml


def test_enhanced_table_round_trip_and_real_ooxml_grid(round_trip_runtime, tmp_path: Path) -> None:
    source = _write_source(tmp_path, "enhanced-table", _ENHANCED_TABLE)
    forward_dir = tmp_path / "forward"
    reverse_dir = tmp_path / "reverse"
    forward_dir.mkdir()
    reverse_dir.mkdir()

    result = _run(
        round_trip_runtime,
        "semantics-table-forward",
        source,
        source_format="markdown",
        target_format="docx",
        output_dir=forward_dir,
    )
    docx_path = _primary_path(result)
    returned = docx_to_md(
        round_trip_runtime,
        docx_path,
        reverse_dir,
        request_id="semantics-table-reverse",
    )

    assert "Sales channels" in returned
    assert "North" in returned
    assert "South" in returned
    assert "See @tbl-sales for the synthetic totals." in returned
    xml = _document_xml(docx_path)
    assert "SEQ Table" not in xml
    assert "w:gridSpan" in xml
    assert 'w:vMerge w:val="restart"' in xml
    assert "w:vMerge" in xml
    assert xml.count("w:tblHeader") == 2
    assert "w:firstColumn" in xml


def test_legacy_citation_is_literal_without_v1_processor_warning(round_trip_runtime, tmp_path: Path) -> None:
    citation = "Existing studies support this result [@smith2025; @wang2024].\n"
    source = _write_source(tmp_path, "citation", citation)
    forward_dir = tmp_path / "forward"
    reverse_dir = tmp_path / "reverse"
    forward_dir.mkdir()
    reverse_dir.mkdir()

    result = _run(
        round_trip_runtime,
        "semantics-citation-forward",
        source,
        source_format="markdown",
        target_format="docx",
        output_dir=forward_dir,
    )
    assert result.success
    warnings = [
        (item.level, item.code, item.message)
        for item in result.diagnostics
        if item.code == "interop.citation.processor_unavailable"
    ]
    assert warnings == []
    returned = docx_to_md(
        round_trip_runtime,
        _primary_path(result),
        reverse_dir,
        request_id="semantics-citation-reverse",
    )
    assert "[@smith2025; @wang2024]" in returned


def test_shorthand_caption_source_form_survives(round_trip_runtime, tmp_path: Path) -> None:
    shorthand = _ENHANCED_TABLE.replace("Table: Sales channels", ": Sales channels", 1)
    source = _write_source(tmp_path, "shorthand", shorthand)
    work_dir = tmp_path / "work"
    reverse_dir = tmp_path / "reverse"
    work_dir.mkdir()
    reverse_dir.mkdir()
    result = _run(
        round_trip_runtime,
        "semantics-shorthand-forward",
        source,
        source_format="markdown",
        target_format="docx",
        output_dir=work_dir,
    )
    returned = docx_to_md(
        round_trip_runtime,
        _primary_path(result),
        reverse_dir,
        request_id="semantics-shorthand-reverse",
    )
    _metadata, returned_body = extract_yaml_front_matter(returned)
    assert returned_body.lstrip().startswith(": Sales channels {#tbl-sales}")
    assert "See @tbl-sales for the synthetic totals." in returned_body


def test_invalid_semantics_fail_closed_with_stable_diagnostics(round_trip_runtime, tmp_path: Path) -> None:
    source = _write_source(tmp_path, "invalid", _INVALID, image=True)
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    result = _run(
        round_trip_runtime,
        "semantics-invalid-forward",
        source,
        source_format="markdown",
        target_format="docx",
        output_dir=output_dir,
    )

    assert not result.success
    assert result.error.error_type == "invalid_document_semantics"
    assert result.artifacts == []
    assert [(item.level, item.code, item.message) for item in result.diagnostics] == [
        (
            "error",
            "interop.table.merge_non_rectangular",
            "Merge markers do not cover a complete rectangle.",
        ),
        ("error", "interop.table.attribute_invalid", "header-rows exceeds the table row count."),
    ]
    assert not list(output_dir.glob("*.docx"))


def test_markdown_targetless_table_caption_survives_nonadjacent_real_docx_import(tmp_path: Path) -> None:
    markdown = """: Unaddressable table

| Column |
|---|
| value |
"""
    analysis = analyze_document_semantics(parse_markdown_text(markdown))
    assert not analysis.has_errors
    document = Document()
    MdToDocxRenderer(document).render(analysis.ast)
    body = document.element.body
    table_element = document.tables[0]._tbl
    caption_element = next(paragraph._p for paragraph in document.paragraphs if "SEQ Table" in paragraph._p.xml)
    assert body.index(caption_element) < body.index(table_element)
    separator = document.add_paragraph("Physical paragraph between caption and table.")._p
    body.remove(separator)
    body.insert(body.index(caption_element) + 1, separator)
    assert body.index(caption_element) + 2 == body.index(table_element)
    output = tmp_path / "markdown-targetless-table.docx"
    document.save(output)

    reopened = Document(output)
    reopened_body = reopened.element.body
    reopened_table = reopened.tables[0]._tbl
    reopened_caption = next(paragraph._p for paragraph in reopened.paragraphs if "SEQ Table" in paragraph._p.xml)
    reopened_separator = next(
        paragraph._p
        for paragraph in reopened.paragraphs
        if paragraph.text == "Physical paragraph between caption and table."
    )
    assert reopened_body.index(reopened_caption) < reopened_body.index(reopened_separator)
    assert reopened_body.index(reopened_separator) < reopened_body.index(reopened_table)
    imported = DocxSemanticImporter().import_document(reopened)
    tables = [block for block in imported.document.blocks if isinstance(block, SemanticTable)]

    assert imported.diagnostics == ()
    assert len(tables) == 1
    assert tables[0].caption is not None
    assert tables[0].caption.target_id is None
    assert tables[0].caption.cached_number == "1"
    assert tables[0].caption.content == "Unaddressable table"
    assert "_DWP_C_" in reopened.element.xml
    assert "_DWP_O_" in reopened.element.xml
