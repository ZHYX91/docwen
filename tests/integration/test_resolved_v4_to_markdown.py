"""Production proof-only resolved-v4 DOCX -> Markdown integration gates."""

from __future__ import annotations

import hashlib
import json
from pathlib import Path
from typing import Any, cast
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from lxml import etree

from docwen_core._docx_semantics_v3_model import CaptionStyleBindingV3, CaptionStyleKeyV3
from docwen_core.docx_numbering_import import AMBIGUOUS_VISIBLE_PREFIX_DIAGNOSTIC
from docwen_core.docx_resolved_numbering import ResolvedNumberingDocxSession
from docwen_core.docx_resolved_numbering_recovery import (
    RESOLVED_V4_SOURCE_SNAPSHOT_MISSING_DIAGNOSTIC,
    ResolvedNumberingV4Recovery,
)
from docwen_core.models.file_ref import FileRef
from docwen_core.models.request import ConversionRequest, OutputPolicy
from docwen_core.models.resolved_numbering import (
    NUMBERING_EXPORT_PLAN_MEDIA_TYPE,
    RESOLVED_DOCUMENT_MEDIA_TYPE,
    HeadingCounterSegment,
    HeadingDefinition,
    HeadingInstance,
    HeadingLevelDefinition,
    HeadingListMaterialization,
    NumberingExportPlanEnvelope,
    NumberingTarget,
    ResolvedDocument,
    ResolvedDocumentEnvelope,
    ResolvedDocumentTarget,
    ResolvedNumberingPlan,
    ResolvedNumberingPort,
)
from docwen_plugin_document.to_markdown.converter import DocxToMarkdownConverter
from docwen_plugin_markdown.to_docx.converter import MdToDocxConverter
from docwen_runtime.config.document_styles import build_document_style_catalog
from tests.support.cancellation import FakeCancellationTokenView
from tests.support.config import FakeConfigView
from tests.support.execution import FakeExecutionContext
from tests.support.logging import FakePluginLogger
from tests.support.progress import FakeProgressSink
from tests.support.workspace import FakeWorkspaceHandle

pytestmark = pytest.mark.integration

_PROJECT_ROOT = Path(__file__).resolve().parents[2]
_FORWARD_FIXTURES = _PROJECT_ROOT / "packages" / "plugins" / "markdown" / "tests" / "fixtures" / "resolved_v4"
_NEUTRAL = _FORWARD_FIXTURES / "resolved-document.rich.json"
_PLAN = _FORWARD_FIXTURES / "numbering-export-plan.rich.json"


def _sha(value: str) -> str:
    return hashlib.sha256(value.encode()).hexdigest()


def _forward_context(tmp_path: Path) -> FakeExecutionContext:
    staging = tmp_path / "forward-staging"
    staging.mkdir()
    refs = (
        FileRef(
            path=str(_NEUTRAL),
            format="markdown",
            category="document",
            size_bytes=_NEUTRAL.stat().st_size,
            input_kind="document",
            input_role="neutral_document",
            logical_path="document.json",
            media_type=RESOLVED_DOCUMENT_MEDIA_TYPE,
        ),
        FileRef(
            path=str(_PLAN),
            format="json",
            category="resource",
            size_bytes=_PLAN.stat().st_size,
            input_kind="resource",
            input_role="numbering_export_plan",
            logical_path="numbering-plan.json",
            media_type=NUMBERING_EXPORT_PLAN_MEDIA_TYPE,
        ),
    )
    request = ConversionRequest(
        request_id="resolved-v4-forward-for-reverse",
        input_refs=list(refs),
        target_format="docx",
        options={"locale": "zh_CN", "heading_merge_mode": "never"},
        output_policy=OutputPolicy(),
    )
    workspace = FakeWorkspaceHandle(str(_NEUTRAL), str(staging), refs)
    styles = build_document_style_catalog(
        {"gui": {"language": {"locale": "zh_CN"}}},
        locales_dir=_PROJECT_ROOT / "i18n" / "locales",
    )
    return FakeExecutionContext(
        request,
        workspace,
        FakeConfigView(),
        FakeProgressSink(),
        FakeCancellationTokenView(),
        FakePluginLogger(),
        document_style_catalog=styles,
    )


def _reverse_context(tmp_path: Path, source: Path, *, request_id: str) -> FakeExecutionContext:
    staging = tmp_path / f"reverse-staging-{request_id}"
    staging.mkdir()
    ref = FileRef(
        path=str(source),
        format="docx",
        category="document",
        size_bytes=source.stat().st_size,
    )
    return FakeExecutionContext(
        ConversionRequest(
            request_id=request_id,
            input_refs=[ref],
            target_format="md",
            options={"to_md_keep_images": True, "remove_numbering": True},
            output_policy=OutputPolicy(),
        ),
        FakeWorkspaceHandle(str(source), str(staging), (ref,)),
        FakeConfigView(),
        FakeProgressSink(),
        FakeCancellationTokenView(),
        FakePluginLogger(),
    )


def _forward_representative(tmp_path: Path) -> Path:
    result = MdToDocxConverter().convert(_forward_context(tmp_path))
    assert result.success, result.error
    return Path(result.artifacts[0].staging_path)


def test_representative_proves_four_kinds_refs_citation_and_preserves_tokens_without_exact_claim(
    tmp_path: Path,
) -> None:
    source = _forward_representative(tmp_path)
    context = _reverse_context(tmp_path, source, request_id="representative")

    result = DocxToMarkdownConverter().convert(context)

    assert result.success, result.error
    primary = Path(next(item.staging_path for item in result.artifacts if item.is_primary))
    markdown = primary.read_text(encoding="utf-8")
    expected_source = json.loads(_NEUTRAL.read_text(encoding="utf-8"))["document"]["authored_markdown"]
    assert len(expected_source.encode()) == 398
    assert markdown != expected_source
    assert RESOLVED_V4_SOURCE_SNAPSHOT_MISSING_DIAGNOSTIC in [item[2] for item in context.progress.diagnostics]
    assert "# Architecture ^h-7f3a" in markdown
    assert "Figure: System overview ^system-overview" in markdown
    assert "Table: Results ^results-main" in markdown
    assert "Equation: ^energy-main" in markdown
    assert "Code: Entry point ^entry-main" in markdown
    assert "Figure: 1" not in markdown
    assert "Table: 1" not in markdown
    assert "Equation: 1" not in markdown
    assert "Code: 1" not in markdown
    assert "Stable: @[[#^h-7f3a]] and @[[#^system-overview|System overview]]." in markdown
    assert "Ordinary: [[#^system-overview]] and ![[Guide#^h-7f3a]]." in markdown
    assert "Citation: @cite-one." in markdown

    reopened = Document(str(source))
    recovery = ResolvedNumberingV4Recovery.load_if_present(source, reopened)
    assert recovery is not None
    assert recovery.caption_signatures == (
        ("figure", "system-overview", "System overview", "1"),
        ("table", "results-main", "Results", "1"),
        ("equation", "energy-main", "", "1"),
        ("code_block", "entry-main", "Entry point", "1"),
    )


@pytest.mark.parametrize("enabled", [True, False])
def test_manual_heading_prefix_survives_enabled_and_disabled_without_legacy_cleanup(
    tmp_path: Path,
    enabled: bool,
) -> None:
    source = _render_manual_heading_package(tmp_path, enabled=enabled)
    context = _reverse_context(tmp_path, source, request_id=f"manual-{enabled}")

    result = DocxToMarkdownConverter().convert(context)

    assert result.success, result.error
    markdown = Path(result.artifacts[0].staging_path).read_text(encoding="utf-8")
    assert "# 2.3 标题 ^head-a" in markdown
    assert "# 标题 ^head-a" not in markdown
    assert "# 1 2.3 标题" not in markdown
    diagnostics = [item[2] for item in context.progress.diagnostics]
    assert RESOLVED_V4_SOURCE_SNAPSHOT_MISSING_DIAGNOSTIC in diagnostics
    assert (AMBIGUOUS_VISIBLE_PREFIX_DIAGNOSTIC in diagnostics) is (not enabled)


def test_reference_cache_tamper_fails_before_artifact_or_staging_publish(tmp_path: Path) -> None:
    source = _forward_representative(tmp_path)
    tampered = tmp_path / "tampered-reference.docx"
    tampered.write_bytes(source.read_bytes())
    _tamper_first_reference_cached_result(tampered)
    context = _reverse_context(tmp_path, tampered, request_id="tampered")

    result = DocxToMarkdownConverter().convert(context)

    assert not result.success
    assert result.artifacts == []
    assert context.workspace.registered_artifacts == []
    assert list(Path(context.workspace.staging_dir).iterdir()) == []


def _caption_bindings(document: Any) -> tuple[CaptionStyleBindingV3, ...]:
    output: list[CaptionStyleBindingV3] = []
    for semantic_key, style_id, name in (
        ("figure_caption", "DWFigureCaption", "Figure Caption V4"),
        ("table_caption", "DWTableCaption", "Table Caption V4"),
        ("equation_caption", "DWEquationCaption", "Equation Caption V4"),
        ("code_block_caption", "DWCodeCaption", "Code Caption V4"),
    ):
        style = document.styles.add_style(style_id, WD_STYLE_TYPE.PARAGRAPH)
        style.name = name
        output.append(CaptionStyleBindingV3(cast(CaptionStyleKeyV3, semantic_key), style.style_id, style.name))
    return tuple(output)


def _manual_port(*, enabled: bool) -> ResolvedNumberingPort:
    source = "# 2.3 标题 ^head-a\n\nTable: 数据\n\n| A |\n|---|\n| B |\n"
    heading_end = source.index("\n")
    table_start = source.index("Table:")
    table_end = len(source) - 1
    heading = ResolvedDocumentTarget(0, heading_end, _sha(source[:heading_end]), "heading", "head-a", 1, "2.3 标题")
    table = ResolvedDocumentTarget(
        table_start,
        table_end,
        _sha(source[table_start:table_end]),
        "table",
        None,
        None,
        "数据",
    )
    definition = HeadingDefinition(
        "main",
        (
            HeadingLevelDefinition(
                1,
                1,
                "arabic_half",
                (HeadingCounterSegment(1, "arabic_half"),),
                "space",
                None,
            ),
        ),
    )
    heading_plan = NumberingTarget(
        0,
        heading_end,
        "heading",
        enabled,
        "head-a",
        "1" if enabled else None,
        HeadingListMaterialization("main", "document", 1) if enabled else None,
    )
    plan = ResolvedNumberingPlan(
        (definition,) if enabled else (),
        (HeadingInstance("document", "main", ()),) if enabled else (),
        (
            heading_plan,
            NumberingTarget(table_start, table_end, "table", False, None, None, None),
        ),
    )
    identity = "a" * 64
    document = ResolvedDocument(source, (heading, table), (), (), (), ())
    return ResolvedNumberingPort(
        ResolvedDocumentEnvelope("manual", _sha(source), identity, document),
        NumberingExportPlanEnvelope("manual", _sha(source), identity, plan),
    )


def _render_manual_heading_package(tmp_path: Path, *, enabled: bool) -> Path:
    document = Document()
    session = ResolvedNumberingDocxSession(
        document,
        _manual_port(enabled=enabled),
        heading_style_ids={1: "Heading1"},
        heading_style_names={"heading_1": "Heading 1"},
        caption_style_bindings=_caption_bindings(document),
    )
    heading = document.add_heading("2.3 标题", level=1)
    caption = document.add_paragraph(style="Table Caption V4")
    table = document.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "A"
    table.cell(1, 0).text = "B"
    targets = session.port.document.targets
    session.bind_heading(heading, source_start=targets[0].source_start, source_end=targets[0].source_end)
    session.bind_caption(
        caption,
        (table._element,),
        source_start=targets[1].source_start,
        source_end=targets[1].source_end,
        kind="table",
    )
    output = tmp_path / f"manual-{enabled}.docx"
    session.write_package(output)
    return output


def _tamper_first_reference_cached_result(path: Path) -> None:
    with ZipFile(path) as package:
        infos = package.infolist()
        members = {item.filename: package.read(item.filename) for item in infos}
    root = etree.fromstring(members["word/document.xml"])
    occurrence = next(
        item
        for item in root.iter(qn("w:sdt"))
        if (tag := item.find(f"{qn('w:sdtPr')}/{qn('w:tag')}")) is not None
        and (tag.get(qn("w:val")) or "").startswith("docwen-ref-occurrence-v1:")
    )
    result_texts = list(occurrence.iter(qn("w:t")))
    assert result_texts and result_texts[0].text == "1"
    result_texts[0].text = "9"
    members["word/document.xml"] = etree.tostring(
        root,
        encoding="UTF-8",
        xml_declaration=True,
        standalone=True,
    )
    rewritten = path.with_suffix(".rewrite.docx")
    with ZipFile(rewritten, "w", compression=ZIP_DEFLATED) as output:
        for info in infos:
            output.writestr(info, members[info.filename])
    rewritten.replace(path)
