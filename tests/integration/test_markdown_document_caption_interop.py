"""Markdown and Document plugin caption interoperability contracts."""

from __future__ import annotations

import base64
import hashlib
import json
from dataclasses import asdict
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from docx.oxml.ns import qn

from docwen_core.docx_resolved_numbering_recovery import ResolvedNumberingV4Recovery
from docwen_core.docx_semantics_v3 import DocxSemanticsV3Recovery
from docwen_core.models.file_ref import FileRef
from docwen_core.models.request import ConversionRequest, OutputPolicy
from docwen_core.models.resolved_numbering import (
    NUMBERING_EXPORT_PLAN_MEDIA_TYPE,
    RESOLVED_DOCUMENT_MEDIA_TYPE,
    CaptionMaterialization,
    NumberingTarget,
    ResolvedDocument,
    ResolvedDocumentTarget,
    ResolvedEmbeddedResource,
    ResolvedNumberingPlan,
    ResolvedReference,
    ResolvedResourceOccurrence,
    canonicalize_numbering_plan,
)
from tests.integration._round_trip_helper import docx_to_md, md_to_docx

pytestmark = [pytest.mark.integration, pytest.mark.pr_gate]


def test_figure_captioned_multi_image_table_round_trips_as_native_table(
    round_trip_runtime: Any,
    tmp_path: Path,
) -> None:
    png = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAIAAAACCAIAAAD91JpzAAAAEklEQVR4nGNUSFjAwMDAxAAGAA0qASTlOPBgAAAAAElFTkSuQmCC"
    )
    (tmp_path / "a.png").write_bytes(png)
    (tmp_path / "b.png").write_bytes(png)
    source_path = tmp_path / "composite.md"
    source = "Figure: Composite ^composite\n\n| A | B |\n|---|---|\n| ![a](a.png) | ![b](b.png) |\n"
    source_path.write_text(source, encoding="utf-8")
    output = md_to_docx(
        round_trip_runtime,
        source_path,
        tmp_path / "docx-output",
        request_id="caption-cross-type",
    )

    reopened = Document(str(output))
    recovery = DocxSemanticsV3Recovery.load(output, reopened)
    body = recovery.logical_body_elements(reopened)
    table = next(item for item in body if item.tag == qn("w:tbl"))
    caption = recovery.caption_for_object(table)
    assert caption is not None
    assert (caption.kind, caption.source_id, caption.title, caption.cached_number) == (
        "figure",
        "composite",
        "Composite",
        "1",
    )
    assert "SEQ Figure" in caption.caption_element.xml
    assert len(list(table.iter(qn("w:drawing")))) == 2
    assert [item.tag for item in body[:2]] == [qn("w:tbl"), qn("w:p")]

    markdown = docx_to_md(
        round_trip_runtime,
        output,
        tmp_path / "markdown-output",
        request_id="caption-cross-type-reverse",
        options={"to_md_keep_images": True, "to_md_enable_ocr": False},
        preserve_numbering=False,
    )
    declaration = markdown.index("Figure: Composite ^composite")
    table_start = markdown.index("| A | B |")
    assert declaration < table_start
    assert markdown.count("![[docx-image_") == 2
    assert "| --- | --- |" in markdown


def test_exact_two_figure_captioned_multi_image_table_round_trips_with_short_target_range(
    round_trip_runtime: Any,
    tmp_path: Path,
) -> None:
    neutral, plan = _write_exact_two_composite_pair(tmp_path)
    refs = (
        FileRef(
            path=str(neutral),
            format="markdown",
            category="document",
            size_bytes=neutral.stat().st_size,
            input_kind="document",
            input_role="neutral_document",
            logical_path="document.json",
            media_type=RESOLVED_DOCUMENT_MEDIA_TYPE,
        ),
        FileRef(
            path=str(plan),
            format="json",
            category="resource",
            size_bytes=plan.stat().st_size,
            input_kind="resource",
            input_role="numbering_export_plan",
            logical_path="numbering-plan.json",
            media_type=NUMBERING_EXPORT_PLAN_MEDIA_TYPE,
        ),
    )
    result = round_trip_runtime.execute(
        ConversionRequest(
            request_id="caption-cross-type-exact-two",
            input_refs=list(refs),
            target_format="docx",
            options={"locale": "en_US", "heading_merge_mode": "never"},
            output_policy=OutputPolicy(output_dir=str(tmp_path / "exact-two")),
        )
    )
    assert result.success, result.error
    [primary] = [artifact for artifact in result.artifacts if artifact.kind == "primary"]
    output = Path(primary.staging_path)

    reopened = Document(str(output))
    recovery = ResolvedNumberingV4Recovery.load_if_present(output, reopened)
    assert recovery is not None
    body = recovery.logical_body_elements(reopened)
    table = next(item for item in body if item.tag == qn("w:tbl"))
    caption = recovery.caption_for_object(table)
    assert caption is not None
    assert (caption.kind, caption.source_id, caption.title, caption.cached_number) == (
        "figure",
        "composite",
        "Composite",
        "1",
    )
    assert "SEQ Figure" in caption.caption_element.xml
    instructions = [item.text or "" for item in reopened.element.iter(qn("w:instrText"))]
    assert any(" REF " in item for item in instructions)
    assert not any("SEQ Table" in item for item in instructions)
    assert len(list(table.iter(qn("w:drawing")))) == 2
    assert [item.tag for item in body[:2]] == [qn("w:tbl"), qn("w:p")]

    markdown = docx_to_md(
        round_trip_runtime,
        output,
        tmp_path / "exact-two-reverse",
        request_id="caption-cross-type-exact-two-reverse",
        options={"to_md_keep_images": True, "to_md_enable_ocr": False},
        preserve_numbering=False,
    )
    authored = json.loads(neutral.read_text(encoding="utf-8"))["document"]["authored_markdown"]
    assert markdown == authored
    assert markdown.index("Figure: Composite\n^composite") < markdown.index("| A | B |")
    assert markdown.count("![[") == 2
    assert "|---|---|" in markdown
    assert "@[[#^composite]]" in markdown


def _write_exact_two_composite_pair(tmp_path: Path) -> tuple[Path, Path]:
    png = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAIAAAACCAIAAAD91JpzAAAAEklEQVR4nGNUSFjAwMDAxAAGAA0qASTlOPBgAAAAAElFTkSuQmCC"
    )
    source = (
        "Figure: Composite\n^composite\n| A | B |\n|---|---|\n| ![[a.png]] | ![[b.png]] |\n\nSee @[[#^composite]].\n"
    )
    declaration = "Figure: Composite"
    target = ResolvedDocumentTarget(
        source_start=0,
        source_end=len(declaration),
        source_slice_sha256=hashlib.sha256(declaration.encode()).hexdigest(),
        kind="figure",
        target_id="composite",
        heading_level=None,
        authored_text="Composite",
    )
    resources: list[ResolvedEmbeddedResource] = []
    occurrences: list[ResolvedResourceOccurrence] = []
    for index, name in enumerate(("a.png", "b.png"), start=1):
        token = f"![[{name}]]"
        start = source.index(token)
        resource_id = f"image-{index}"
        resources.append(
            ResolvedEmbeddedResource(
                resource_id=resource_id,
                role="linked_resource",
                media_type="image/png",
                size_bytes=len(png),
                sha256=hashlib.sha256(png).hexdigest(),
                content=png,
            )
        )
        occurrences.append(
            ResolvedResourceOccurrence(
                source_start=start,
                source_end=start + len(token),
                source_slice_sha256=hashlib.sha256(token.encode()).hexdigest(),
                authored_token=token,
                authored_locator=name,
                resource_id=resource_id,
            )
        )
    materialization = CaptionMaterialization(
        type="simple_seq",
        counter="Figure",
        number_format="arabic_half",
        sequence_action="reset_to_start",
        start_value=1,
        chapter_heading_level=None,
        chapter_heading_style=None,
        chapter_separator=None,
        restart_heading_level=None,
        restart_heading_style=None,
        chapter_cached_number=None,
        sequence_cached_number="1",
        localized_label="Figure",
        label_separator=" ",
    )
    numbering_plan = ResolvedNumberingPlan(
        heading_definitions=(),
        heading_instances=(),
        targets=(
            NumberingTarget(
                source_start=target.source_start,
                source_end=target.source_end,
                kind="figure",
                enabled=True,
                target_id="composite",
                derived_number="1",
                materialization=materialization,
            ),
        ),
    )
    plan_body = json.loads(json.dumps(asdict(numbering_plan)))
    plan_sha256 = hashlib.sha256(canonicalize_numbering_plan(plan_body)).hexdigest()
    reference_token = "@[[#^composite]]"
    reference_start = source.index(reference_token)
    reference = ResolvedReference(
        source_start=reference_start,
        source_end=reference_start + len(reference_token),
        source_slice_sha256=hashlib.sha256(reference_token.encode()).hexdigest(),
        authored_token=reference_token,
        target_source_start=target.source_start,
        target_source_end=target.source_end,
        target_kind="figure",
        target_id="composite",
        cached_number="1",
        alias=None,
    )
    document = ResolvedDocument(source, (target,), (reference,), tuple(occurrences), (), tuple(resources))
    document_body = asdict(document)
    for resource in document_body["resources"]:
        content = resource.pop("content")
        resource["content_base64"] = base64.b64encode(content).decode("ascii")
    source_sha256 = hashlib.sha256(source.encode()).hexdigest()
    neutral_payload = {
        "$schema": "urn:docwen:schema:resolved-document:v1",
        "schema": "docwen.resolved_document.v1",
        "input_id": "caption-cross-type",
        "source_sha256": source_sha256,
        "plan_sha256": plan_sha256,
        "document": document_body,
    }
    plan_payload = {
        "$schema": "urn:docwen:schema:numbering-export-plan:v1",
        "schema": "docwen.numbering_export_plan.v1",
        "input_id": "caption-cross-type",
        "source_sha256": source_sha256,
        "plan_sha256": plan_sha256,
        "plan": plan_body,
    }
    neutral = tmp_path / "composite-neutral.json"
    plan = tmp_path / "composite-plan.json"
    neutral.write_text(json.dumps(neutral_payload, separators=(",", ":")), encoding="utf-8")
    plan.write_text(json.dumps(plan_payload, separators=(",", ":")), encoding="utf-8")
    return neutral, plan
