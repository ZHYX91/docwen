"""Real DOCX round trip driven exclusively by provider-neutral data."""

from __future__ import annotations

from zipfile import ZipFile

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from lxml import etree

from docwen_core.docx_semantics import DocxSemanticImporter, DocxSemanticRenderer
from docwen_core.models.semantic_document import (
    SemanticBibliographyEntry,
    SemanticBibliographyFragment,
    SemanticBibliographyRun,
    SemanticCaption,
    SemanticCitationCluster,
    SemanticCitationItem,
    SemanticDocument,
    SemanticParagraph,
    SemanticReference,
    SemanticTable,
    SemanticTableCell,
    SemanticText,
)

pytestmark = pytest.mark.integration


def _neutral_document() -> SemanticDocument:
    # The table is 4x4 overall.  Its first two rows form a 2x4 repeated
    # column-header region; the final rows supply row_header and data roles.
    table = SemanticTable(
        row_count=4,
        column_count=4,
        repeat_header="always",
        caption=SemanticCaption(
            kind="table",
            target_id="tbl-sales",
            cached_number="7",
            label="Table",
            content="Sales channels",
        ),
        cells=(
            SemanticTableCell(0, 0, "Region", "corner_header", row_span=2),
            SemanticTableCell(0, 1, "Sales", "column_header", column_span=2),
            SemanticTableCell(0, 3, "Total", "column_header"),
            SemanticTableCell(1, 1, "Online", "column_header"),
            SemanticTableCell(1, 2, "Retail", "column_header"),
            SemanticTableCell(1, 3, "Combined", "column_header"),
            SemanticTableCell(2, 0, "North", "row_header", row_span=2),
            SemanticTableCell(2, 1, "10", "data", column_span=2),
            SemanticTableCell(2, 3, "22", "data"),
            SemanticTableCell(3, 1, "8", "data"),
            SemanticTableCell(3, 2, "9", "data"),
            SemanticTableCell(3, 3, "17", "data"),
        ),
    )
    return SemanticDocument(
        blocks=(
            table,
            SemanticParagraph(
                (
                    SemanticText("See table "),
                    SemanticReference("tbl-sales", "7"),
                    SemanticText("."),
                )
            ),
            SemanticParagraph(
                (
                    SemanticText("Citations "),
                    SemanticCitationCluster(
                        cluster_id="cluster-one",
                        items=(SemanticCitationItem("smith2025"),),
                        cached_result="[1]",
                    ),
                    SemanticText(" and "),
                    SemanticCitationCluster(
                        cluster_id="cluster-two",
                        items=(
                            SemanticCitationItem("wang2024"),
                            SemanticCitationItem("smith2025"),
                        ),
                        cached_result="[2, 1]",
                    ),
                    SemanticText("."),
                )
            ),
            # Bibliography heading is deliberately an ordinary independent
            # paragraph, never a field of SemanticBibliographyFragment.
            SemanticParagraph((SemanticText("References"),)),
        ),
        bibliography=SemanticBibliographyFragment(
            entries=(
                SemanticBibliographyEntry(
                    "smith2025",
                    (
                        SemanticBibliographyRun("Smith, A. (2025). ", bold=True),
                        SemanticBibliographyRun(
                            "Neutral documents",
                            italic=True,
                            href="https://example.org/neutral-documents",
                        ),
                        SemanticBibliographyRun("."),
                    ),
                ),
                SemanticBibliographyEntry(
                    "wang2024",
                    (SemanticBibliographyRun("Wang, B. (2024). Structured tables."),),
                ),
            )
        ),
    )


def test_neutral_data_real_docx_round_trip(tmp_path) -> None:
    semantic_document = _neutral_document()
    document = Document()
    document.styles.add_style("Bibliography", WD_STYLE_TYPE.PARAGRAPH)
    document.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
    renderer = DocxSemanticRenderer(document)
    renderer.render_blocks(semantic_document.blocks)
    bibliography_anchor = document.add_paragraph("exclusive bibliography anchor")
    assert semantic_document.bibliography is not None
    renderer.render_bibliography_fragment(
        semantic_document.bibliography,
        placeholder_anchor=bibliography_anchor,
        fallback_style_id="Bibliography",
        hyperlink_style_id="Hyperlink",
    )
    output = tmp_path / "neutral-semantics.docx"
    document.save(output)

    with ZipFile(output) as package:
        xml = package.read("word/document.xml").decode("utf-8")
        custom_xml_roots = [
            etree.fromstring(package.read(name))
            for name in package.namelist()
            if name.startswith("customXml/") and name.endswith(".xml")
        ]
    assert "SEQ Table" in xml
    assert "REF _DW_" in xml
    assert "_DWO_" in xml
    assert "_DWB_BIBLIOGRAPHY" in xml
    assert "_DWC_" in xml
    assert "_DWE_" in xml
    assert "CITATION smith2025" in xml
    assert r"CITATION wang2024 \m smith2025" in xml
    assert "w:gridSpan" in xml
    assert 'w:vMerge w:val="restart"' in xml
    assert "w:cnfStyle" in xml
    assert xml.count("w:tblHeader") == 2
    assert ">7<" in xml
    assert not any(root.xpath("//*[local-name()='Source']") for root in custom_xml_roots)

    reopened = Document(output)
    citation_begins = [
        item
        for item in reopened.element.iter(qn("w:fldChar"))
        if item.get(qn("w:fldCharType")) == "begin" and item.get(qn("w:fldLock")) in {"1", "true", "on"}
    ]
    instructions = [item.text or "" for item in reopened.element.iter(qn("w:instrText"))]
    assert len(citation_begins) == 2
    assert all(item.get(qn("w:dirty")) is None for item in citation_begins)
    assert not any(instruction.strip().upper().startswith("BIBLIOGRAPHY") for instruction in instructions)
    assert all(token not in xml.casefold() for token in ("wenleaf", "csl", "pkwf"))
    hyperlink = next(reopened.element.iter(qn("w:hyperlink")))
    relationship = reopened.part.rels[hyperlink.get(qn("r:id"))]
    assert relationship.reltype == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
    assert relationship.is_external
    assert relationship.target_ref == "https://example.org/neutral-documents"

    imported = DocxSemanticImporter().import_document(reopened)

    assert imported.diagnostics == ()
    assert imported.document == semantic_document


def test_empty_bibliography_elides_only_anchor_after_real_save_reopen(tmp_path) -> None:
    heading = SemanticParagraph((SemanticText("References"),))
    tail = SemanticParagraph((SemanticText("After bibliography"),))
    document = Document()
    renderer = DocxSemanticRenderer(document)
    renderer.render_paragraph(heading)
    anchor = document.add_paragraph("exclusive bibliography anchor")
    renderer.render_paragraph(tail)

    rendered = renderer.render_bibliography_fragment(
        SemanticBibliographyFragment(entries=()),
        placeholder_anchor=anchor,
    )
    output = tmp_path / "empty-neutral-bibliography.docx"
    document.save(output)

    with ZipFile(output) as package:
        xml = package.read("word/document.xml").decode("utf-8")
    imported = DocxSemanticImporter().import_document(Document(output))

    assert rendered == ()
    assert "_DWB_BIBLIOGRAPHY" not in xml
    assert "_DWE_" not in xml
    assert imported.diagnostics == ()
    assert imported.document == SemanticDocument(blocks=(heading, tail))


def test_targetless_table_caption_real_docx_round_trip_uses_explicit_pairing(tmp_path) -> None:
    table = SemanticTable(
        row_count=1,
        column_count=1,
        cells=(SemanticTableCell(0, 0, "value"),),
        caption=SemanticCaption(
            kind="table",
            target_id=None,
            cached_number="7",
            label="Table",
            content="Unaddressable but explicitly paired",
        ),
    )
    intervening_paragraph = SemanticParagraph((SemanticText("Explicit pairing does not use adjacency."),))
    semantic_document = SemanticDocument(blocks=(intervening_paragraph, table))
    document = Document()
    DocxSemanticRenderer(document).render_blocks(semantic_document.blocks)
    body = document.element.body
    intervening_element = document.paragraphs[0]._p
    table_element = document.tables[0]._tbl
    body.remove(intervening_element)
    body.insert(body.index(table_element), intervening_element)
    output = tmp_path / "targetless-neutral-semantics.docx"
    document.save(output)

    with ZipFile(output) as package:
        xml = package.read("word/document.xml").decode("utf-8")
    assert "SEQ Table" in xml
    assert "_DWP_C_" in xml
    assert "_DWP_O_" in xml
    assert "_DW_" not in xml
    assert "_DWO_" not in xml

    imported = DocxSemanticImporter().import_document(Document(output))

    assert imported.diagnostics == ()
    assert imported.document == semantic_document
