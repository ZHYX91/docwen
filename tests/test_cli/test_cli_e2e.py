"""CLI protocol 3 end-to-end product workflows."""

from __future__ import annotations

import json
import os
import sys
from pathlib import Path

import openpyxl
import pytest
from docx import Document
from tests.support.cli import bundle_cli_command
from tests.support.subprocess_runner import run_subprocess

pytestmark = pytest.mark.e2e

_MACOS_PRIMARY_OPERATION_UNAVAILABLE = pytest.mark.skipif(
    sys.platform == "darwin",
    reason="DocWen 0.9 exposes primary document operations only on Windows and Linux",
)


@pytest.fixture(autouse=True)
def _isolate_cli_process_data(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    """Keep source CLI subprocesses away from the real user's data roots."""

    runtime_root = tmp_path.parent / f"{tmp_path.name}-runtime"
    monkeypatch.setenv("DOCWEN_CONFIG_DIR", str(runtime_root / "config_home"))
    monkeypatch.setenv("DOCWEN_LOG_DIR", str(runtime_root / "log_home"))


def _repo_root() -> Path:
    return Path(__file__).resolve().parents[2]


def _write_xlsx(path: Path) -> None:
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.append(["name", "value"])
    sheet.append(["alpha", 1])
    workbook.save(path)
    workbook.close()


def _write_docx(path: Path, text: str) -> None:
    document = Document()
    document.add_paragraph(text)
    document.save(path)


def _cli_env() -> dict[str, str]:
    env = os.environ.copy()
    env["PYTHONUTF8"] = "1"
    env["PYTHONIOENCODING"] = "utf-8"
    return env


def _run(*args: str):
    return run_subprocess(
        [*bundle_cli_command(), *args],
        cwd=_repo_root(),
        env=_cli_env(),
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        check=False,
    )


def _enable_fake_external_office(monkeypatch: pytest.MonkeyPatch) -> None:
    """Keep fake pre-conversion tests independent of the host Office inventory."""
    from docwen_core import office_bridge

    monkeypatch.setattr(office_bridge, "find_soffice_path", lambda: Path("fake-soffice"))


def _payload(proc) -> dict[str, object]:
    assert proc.stdout.strip()
    value = json.loads(proc.stdout)
    assert isinstance(value, dict)
    assert value["protocol_version"] == 3
    assert value["product_version"] == "0.9.0"
    return value


@pytest.mark.pr_gate
@pytest.mark.release_gate
@_MACOS_PRIMARY_OPERATION_UNAVAILABLE
def test_cli_convert_xlsx_to_markdown_document_node(tmp_path: Path) -> None:
    source = tmp_path / "sample.xlsx"
    output_parent = tmp_path / "published"
    _write_xlsx(source)

    proc = _run("convert", str(source), "--to", "md", "--output", str(output_parent), "--json")
    assert proc.returncode == 0, proc.stderr
    assert proc.stderr == ""
    payload = _payload(proc)
    assert payload["success"] is True
    assert payload["command"] == "convert"
    assert payload["data"]["inputs"] == [str(source.resolve())]
    output = Path(payload["data"]["output"])
    assert output.parent.parent == output_parent.resolve()
    assert output.stem == output.parent.name
    assert (output.parent / "docwen-node.json").is_file()
    assert output.is_file()
    assert output.read_text(encoding="utf-8")


@pytest.mark.pr_gate
@pytest.mark.release_gate
@_MACOS_PRIMARY_OPERATION_UNAVAILABLE
@pytest.mark.parametrize(
    ("name", "payload"),
    [
        ("semicolon.csv", "city;value\n北京;1\n上海;2\n".encode()),
        ("utf16.csv", "city,value\n北京,1\n上海,2\n".encode("utf-16")),
    ],
)
def test_cli_delimited_ingress_reaches_markdown_document_node(
    tmp_path: Path,
    name: str,
    payload: bytes,
) -> None:
    source = tmp_path / name
    output_parent = tmp_path / "published"
    source.write_bytes(payload)

    inspected = _run("inspect", str(source), "--json")
    assert inspected.returncode == 0, inspected.stderr
    assert inspected.stderr == ""
    inspection_payload = _payload(inspected)
    assert inspection_payload["success"] is True
    assert inspection_payload["data"]["detected_format"] == "csv"
    assert inspection_payload["data"]["decision"] == "allow"

    converted = _run("convert", str(source), "--to", "md", "--output", str(output_parent), "--json")
    assert converted.returncode == 0, converted.stderr
    assert converted.stderr == ""
    conversion_payload = _payload(converted)
    assert conversion_payload["success"] is True
    output = Path(conversion_payload["data"]["output"])
    assert output.parent.parent == output_parent.resolve()
    assert output.stem == output.parent.name
    assert output.parent.name.startswith(f"{source.stem}_")
    assert output.parent.name.endswith("_fromCsv")
    assert (output.parent / "docwen-node.json").is_file()
    markdown = output.read_text(encoding="utf-8")
    assert "北京" in markdown
    assert "上海" in markdown


def test_cli_convert_reports_missing_input(tmp_path: Path) -> None:
    source = tmp_path / "missing.xlsx"
    proc = _run("convert", str(source), "--to", "md", "--output", str(tmp_path / "published"), "--json")
    assert proc.returncode == 2
    assert proc.stderr == ""
    payload = _payload(proc)
    assert payload["success"] is False
    assert payload["command"] == "convert"
    assert payload["error"]["category"] == "invalid_input"
    assert payload["error"]["code"] == "invalid_input"


def test_cli_validate_is_read_only_unless_report_is_explicit(tmp_path: Path) -> None:
    source = tmp_path / "proofread.md"
    source.write_text("# Title\n\nteh value", encoding="utf-8")
    before = {path.name for path in tmp_path.iterdir()}

    read_only = _run("validate", str(source), "--check", "typo", "--json")
    assert read_only.returncode == 0, read_only.stderr
    read_only_payload = _payload(read_only)
    assert read_only_payload["command"] == "validate"
    assert read_only_payload["data"]["output"] == ""
    assert read_only_payload["data"]["artifacts"] == []
    assert "proofread" in read_only_payload["data"]["details"]
    assert {path.name for path in tmp_path.iterdir()} == before

    report = tmp_path / "explicit-report.json"
    written = _run("validate", str(source), "--check", "typo", "--report", str(report), "--json")
    assert written.returncode == 0, written.stderr
    written_payload = _payload(written)
    assert written_payload["data"]["output"] == str(report.resolve())
    assert report.is_file()
    report_payload = json.loads(report.read_text(encoding="utf-8"))
    assert report_payload["schema"] == "docwen.proofread_report.v2"
    assert report_payload["file"] == source.name
    assert report_payload["source"]["content_sha256"]
    assert report_payload["location_contract"] == {
        "id": "docwen.proofread-text-range",
        "version": 1,
        "coordinate_system": "unicode_code_point",
        "offset_base": 0,
        "line_base": 0,
        "column_base": 0,
        "range_end": "exclusive",
    }
    assert report_payload["checks_enabled"]["typos_rule"] is True
    assert isinstance(report_payload["issues"], list)
    assert isinstance(report_payload["summary"], dict)


def test_cli_validate_docx_writes_annotated_docx_only_to_explicit_report(tmp_path: Path) -> None:
    source = tmp_path / "proofread.docx"
    report = tmp_path / "reviewed.docx"
    _write_docx(source, "teh value")
    source_before = source.read_bytes()

    proc = _run("validate", str(source), "--check", "typo", "--report", str(report), "--json")

    assert proc.returncode == 0, proc.stderr
    assert proc.stderr == ""
    payload = _payload(proc)
    assert payload["command"] == "validate"
    assert payload["data"]["output"] == str(report.resolve())
    assert report.is_file()
    assert source.read_bytes() == source_before
    reviewed = Document(report)
    assert [paragraph.text for paragraph in reviewed.paragraphs] == ["teh value"]


def test_cli_validate_legacy_word_preconverts_then_writes_annotated_docx(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    capsys: pytest.CaptureFixture[str],
) -> None:
    """Exercise public CLI -> Application bridge -> Runtime proofread without real Office."""

    from docwen_application.preconversion import pre_converter
    from docwen_application.preconversion.pre_converter import PreConversionResult
    from docwen_bundle.config_port import ConfigPortAdapter
    from docwen_bundle.runtime_factory import create_runtime_port
    from docwen_cli.main import main
    from docwen_runtime.config import ConfigLoader

    source = tmp_path / "legacy.rtf"
    report = tmp_path / "legacy-reviewed.docx"
    source.write_bytes(b"{\\rtf1\\ansi teh value}")
    source_before = source.read_bytes()
    _enable_fake_external_office(monkeypatch)

    def fake_pre_convert(
        _input_path: str,
        source_format: str,
        *,
        staging_dir: str,
        **_kwargs: object,
    ) -> PreConversionResult:
        converted = Path(staging_dir) / "legacy.docx"
        _write_docx(converted, "teh value")
        return PreConversionResult(str(converted), source_format, "Fake Office")

    monkeypatch.setattr(pre_converter, "pre_convert", fake_pre_convert)

    config_loader: ConfigLoader | None = None

    def shared_loader() -> ConfigLoader:
        nonlocal config_loader
        if config_loader is None:
            config_loader = ConfigLoader(
                base_dir=_repo_root() / "configs",
                user_dir=tmp_path / "config_home",
                runtime_overrides={"logger": {"console_enable": False}},
            )
        return config_loader

    def config_port_factory() -> ConfigPortAdapter:
        return ConfigPortAdapter(shared_loader())

    def runtime_port_factory():
        return create_runtime_port(config_loader=shared_loader())

    exit_code = main(
        ["validate", str(source), "--check", "typo", "--report", str(report), "--json"],
        runtime_port_factory=runtime_port_factory,
        config_port_factory=config_port_factory,
    )

    captured = capsys.readouterr()
    assert exit_code == 0, captured.out or captured.err
    assert captured.err == ""
    payload = json.loads(captured.out)
    assert payload["protocol_version"] == 3
    assert payload["command"] == "validate"
    assert payload["success"] is True
    assert payload["data"]["output"] == str(report.resolve())
    assert source.read_bytes() == source_before
    reviewed = Document(report)
    assert [paragraph.text for paragraph in reviewed.paragraphs] == ["teh value"]


def test_cli_legacy_word_installed_backend_failure_uses_conversion_failed_envelope(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    capsys: pytest.CaptureFixture[str],
) -> None:
    """Preserve installed-backend failure taxonomy across CLI -> Application."""

    from docwen_application.preconversion import pre_converter
    from docwen_application.preconversion.pre_converter import PreConversionFailure
    from docwen_bundle.config_port import ConfigPortAdapter
    from docwen_bundle.runtime_factory import create_runtime_port
    from docwen_cli.main import main
    from docwen_runtime.config import ConfigLoader

    source = tmp_path / "legacy.rtf"
    report = tmp_path / "legacy-reviewed.docx"
    source.write_bytes(b"{\\rtf1\\ansi content}")
    source_before = source.read_bytes()
    _enable_fake_external_office(monkeypatch)

    monkeypatch.setattr(
        pre_converter,
        "pre_convert",
        lambda *_args, **_kwargs: PreConversionFailure(
            message="Microsoft Word export failed",
            error_type="conversion_failed",
            diagnostic_code="OFFICE_BACKEND_FAILED",
            cleanup_message="Private Office workspace cleanup failed: test workspace",
            cleanup_failed=True,
        ),
    )

    config_loader: ConfigLoader | None = None

    def shared_loader() -> ConfigLoader:
        nonlocal config_loader
        if config_loader is None:
            config_loader = ConfigLoader(
                base_dir=_repo_root() / "configs",
                user_dir=tmp_path / "config_home",
                runtime_overrides={"logger": {"console_enable": False}},
            )
        return config_loader

    exit_code = main(
        ["validate", str(source), "--check", "typo", "--report", str(report), "--json"],
        runtime_port_factory=lambda: create_runtime_port(config_loader=shared_loader()),
        config_port_factory=lambda: ConfigPortAdapter(shared_loader()),
    )

    captured = capsys.readouterr()
    assert exit_code == 1
    assert captured.err == ""
    payload = json.loads(captured.out)
    assert payload["protocol_version"] == 3
    assert payload["command"] == "validate"
    assert payload["success"] is False
    assert payload["error"]["category"] == "internal"
    assert payload["error"]["code"] == "conversion_failed"
    assert payload["error"]["details"] == "OFFICE_BACKEND_FAILED"
    assert "Microsoft Word export failed" in payload["error"]["message"]
    assert payload["warnings"] == [
        {
            "level": "warning",
            "code": "OFFICE_CLEANUP_FAILED",
            "message": "Private Office workspace cleanup failed: test workspace",
            "location": "",
        }
    ]
    assert source.read_bytes() == source_before
    assert not report.exists()


def test_cli_legacy_word_preconversion_cancellation_uses_public_cancel_contract(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    capsys: pytest.CaptureFixture[str],
) -> None:
    from docwen_application.preconversion import pre_converter
    from docwen_application.preconversion.pre_converter import PreConversionFailure
    from docwen_bundle.config_port import ConfigPortAdapter
    from docwen_bundle.runtime_factory import create_runtime_port
    from docwen_cli.main import main
    from docwen_runtime.config import ConfigLoader

    source = tmp_path / "legacy.rtf"
    report = tmp_path / "legacy-reviewed.docx"
    source.write_bytes(b"{\\rtf1\\ansi content}")
    _enable_fake_external_office(monkeypatch)

    monkeypatch.setattr(
        pre_converter,
        "pre_convert",
        lambda *_args, **_kwargs: PreConversionFailure(
            message="backend stopped",
            cancelled=True,
            error_type="cancelled",
            diagnostic_code="OFFICE_CONVERSION_CANCELLED",
            cleanup_message="Private Office workspace cleanup failed: test workspace",
            cleanup_failed=True,
        ),
    )

    config_loader: ConfigLoader | None = None

    def shared_loader() -> ConfigLoader:
        nonlocal config_loader
        if config_loader is None:
            config_loader = ConfigLoader(
                base_dir=_repo_root() / "configs",
                user_dir=tmp_path / "config_home",
                runtime_overrides={"logger": {"console_enable": False}},
            )
        return config_loader

    exit_code = main(
        ["validate", str(source), "--check", "typo", "--report", str(report), "--json"],
        runtime_port_factory=lambda: create_runtime_port(config_loader=shared_loader()),
        config_port_factory=lambda: ConfigPortAdapter(shared_loader()),
    )

    captured = capsys.readouterr()
    assert exit_code == 130
    assert captured.err == ""
    payload = json.loads(captured.out)
    assert payload["success"] is False
    assert payload["error"]["category"] == "cancelled"
    assert payload["error"]["code"] == "operation_cancelled"
    assert payload["warnings"] == [
        {
            "level": "warning",
            "code": "OFFICE_CLEANUP_FAILED",
            "message": "Private Office workspace cleanup failed: test workspace",
            "location": "",
        }
    ]
    assert not report.exists()


def test_cli_convert_rejects_existing_output_until_overwrite_is_explicit(tmp_path: Path) -> None:
    source = tmp_path / "sample.md"
    output = tmp_path / "sample.docx"
    source.write_text("# Sample\n", encoding="utf-8")
    output.write_text("keep me", encoding="utf-8")

    rejected = _run("convert", str(source), "--to", "docx", "--output", str(output), "--json")
    assert rejected.returncode == 7
    rejected_payload = _payload(rejected)
    assert rejected_payload["error"]["code"] == "output_exists"
    assert output.read_text(encoding="utf-8") == "keep me"

    accepted = _run(
        "convert",
        str(source),
        "--to",
        "docx",
        "--output",
        str(output),
        "--overwrite",
        "--json",
    )
    assert accepted.returncode == 0, accepted.stderr
    assert _payload(accepted)["data"]["output"] == str(output.resolve())
    assert output.read_bytes() != b"keep me"


def test_cli_unicode_long_output_conflict_is_typed_and_serializable(tmp_path: Path) -> None:
    source = tmp_path / "源文件.md"
    output = tmp_path / (("很长" * 35) + ".docx")
    source.write_text("# 源文件\n", encoding="utf-8")
    output.write_text("保留", encoding="utf-8")

    proc = _run("convert", str(source), "--to", "docx", "--output", str(output), "--json")
    assert proc.returncode == 7
    assert proc.stderr == ""
    payload = _payload(proc)
    assert payload["error"]["category"] == "conflict"
    assert payload["error"]["code"] == "output_exists"
    assert payload["error"]["details"] == {"path": str(output.resolve())}
    assert output.read_text(encoding="utf-8") == "保留"


def test_batch_convert_rejects_same_stem_collision_before_creating_output_dir(tmp_path: Path) -> None:
    first_dir = tmp_path / "first"
    second_dir = tmp_path / "second"
    first_dir.mkdir()
    second_dir.mkdir()
    first = first_dir / "same.xlsx"
    second = second_dir / "same.xlsx"
    _write_xlsx(first)
    _write_xlsx(second)
    output_dir = tmp_path / "outputs"

    proc = _run(
        "batch",
        "convert",
        str(first),
        str(second),
        "--to",
        "md",
        "--output-dir",
        str(output_dir),
        "--json",
    )

    assert proc.returncode == 7
    assert proc.stderr == ""
    payload = _payload(proc)
    assert payload["error"]["category"] == "conflict"
    assert payload["error"]["code"] == "output_collision"
    assert payload["error"]["details"]["collisions"] == [[str(first.resolve()), str(second.resolve())]]
    assert not output_dir.exists()


def test_batch_convert_rejects_existing_deterministic_target(tmp_path: Path) -> None:
    source = tmp_path / "sample.xlsx"
    _write_xlsx(source)
    output_dir = tmp_path / "outputs"
    output_dir.mkdir()
    existing = output_dir / "sample.csv"
    existing.write_text("keep", encoding="utf-8")

    proc = _run(
        "batch",
        "convert",
        str(source),
        "--to",
        "csv",
        "--output-dir",
        str(output_dir),
        "--json",
    )

    assert proc.returncode == 7
    payload = _payload(proc)
    assert payload["error"]["code"] == "output_exists"
    assert payload["error"]["details"]["paths"] == [str(existing.resolve())]
    assert existing.read_text(encoding="utf-8") == "keep"


def test_cli_merge_pdf_uses_exact_output_path(tmp_path: Path) -> None:
    import fitz

    sources = [tmp_path / "first.pdf", tmp_path / "second.pdf"]
    for path, label in zip(sources, ("first", "second"), strict=True):
        document = fitz.open()
        document.new_page(width=240, height=160).insert_text((48, 80), label)
        document.save(path)
        document.close()
    output = tmp_path / "combined.pdf"

    proc = _run("merge", "pdf", *(str(path) for path in sources), "--output", str(output), "--json")
    assert proc.returncode == 0, proc.stderr
    payload = _payload(proc)
    assert payload["command"] == "merge pdf"
    assert payload["data"]["output"] == str(output.resolve())
    merged = fitz.open(output)
    try:
        assert merged.page_count == 2
    finally:
        merged.close()


def test_cli_merge_tables_uses_exact_output_path(tmp_path: Path) -> None:
    sources = [tmp_path / "base.xlsx", tmp_path / "collect.xlsx"]
    for path, rows in zip(
        sources,
        ([["Name", "Score"], ["Alice", 90]], [["Name", "Score"], ["Charlie", 78]]),
        strict=True,
    ):
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        for row in rows:
            sheet.append(row)
        workbook.save(path)
        workbook.close()
    output = tmp_path / "combined.xlsx"

    proc = _run("merge", "tables", *(str(path) for path in sources), "--output", str(output), "--json")
    assert proc.returncode == 0, proc.stderr
    payload = _payload(proc)
    assert payload["command"] == "merge tables"
    assert payload["data"]["output"] == str(output.resolve())
    workbook = openpyxl.load_workbook(output)
    try:
        values = {cell for row in workbook.active.iter_rows(values_only=True) for cell in row if cell is not None}
        assert "Alice,Charlie" in values
    finally:
        workbook.close()


def test_cli_merge_images_uses_exact_output_path(tmp_path: Path) -> None:
    from PIL import Image, ImageSequence

    first = tmp_path / "red.png"
    second = tmp_path / "blue.jpg"
    Image.new("RGB", (32, 24), (255, 0, 0)).save(first)
    Image.new("RGB", (32, 24), (0, 0, 255)).save(second)
    output = tmp_path / "combined.tif"

    proc = _run("merge", "images", str(first), str(second), "--output", str(output), "--json")
    assert proc.returncode == 0, proc.stderr
    payload = _payload(proc)
    assert payload["command"] == "merge images"
    assert payload["data"]["output"] == str(output.resolve())
    image = Image.open(output)
    try:
        assert image.format == "TIFF"
        assert sum(1 for _ in ImageSequence.Iterator(image)) == 2
    finally:
        image.close()
