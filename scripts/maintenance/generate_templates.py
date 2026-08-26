"""Generate localized DOCX templates from the blank template.

The current workspace no longer has the old monolithic package layout or
converter API.  This dev-only tool reads locale metadata from
``i18n/locales/*.toml`` and rewrites the placeholder text in
``scripts/maintenance/空白模板.docx`` while preserving the blank template's
paragraph/run formatting.
"""

from __future__ import annotations

import argparse
import logging
import shutil
from pathlib import Path
from typing import Any

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[2]
SCRIPT_DIR = Path(__file__).resolve().parent
LOCALES_DIR = PROJECT_ROOT / "i18n" / "locales"
TEMPLATES_DIR = PROJECT_ROOT / "templates"
BLANK_TEMPLATE_SRC = SCRIPT_DIR / "空白模板.docx"
DEFAULT_OUTPUT_DIR = PROJECT_ROOT / "samples" / "generated_templates"

logging.basicConfig(level=logging.INFO, format="%(levelname)s - %(name)s - %(message)s")
logger = logging.getLogger(__name__)


def _read_toml(path: Path) -> dict[str, Any]:
    """Read a locale TOML file with the repo's preferred TOML reader."""
    try:
        from docwen_runtime.toml_io import read_toml_file
    except Exception:
        import tomllib

        return tomllib.loads(path.read_text(encoding="utf-8"))
    return read_toml_file(path)


def discover_locale_template_names() -> dict[str, str]:
    """Return ``{locale_code: template_name}`` from locale ``[meta]`` tables."""
    if not LOCALES_DIR.is_dir():
        raise FileNotFoundError(f"Locale directory not found: {LOCALES_DIR}")

    result: dict[str, str] = {}
    for path in sorted(LOCALES_DIR.glob("*.toml"), key=lambda item: item.name.casefold()):
        data = _read_toml(path)
        meta = data.get("meta", {})
        template_name = meta.get("template_name") if isinstance(meta, dict) else None
        if isinstance(template_name, str) and template_name.strip():
            result[path.stem] = template_name.strip()
        else:
            logger.debug("Skipping %s: [meta].template_name is missing", path.name)

    logger.info("Discovered %d locale templates: %s", len(result), ", ".join(result))
    return result


def get_locale_placeholders(locale: str) -> dict[str, str]:
    """Return the localized title/body placeholders for *locale*."""
    locale_path = LOCALES_DIR / f"{locale}.toml"
    if not locale_path.is_file():
        raise FileNotFoundError(f"Locale file not found: {locale_path}")

    data = _read_toml(locale_path)
    placeholders = data.get("placeholders", {})
    if not isinstance(placeholders, dict):
        placeholders = {}
    title = str(placeholders.get("title") or "title").strip()
    body = str(placeholders.get("body") or "body").strip()
    return {"title": title or "title", "body": body or "body"}


def _replace_paragraph_text(paragraph: Any, text: str) -> None:
    """Replace paragraph text while keeping the first run's formatting."""
    if paragraph.runs:
        first_run = paragraph.runs[0]
        first_run.text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _replace_placeholder_text(doc: Any, old_text: str, new_text: str) -> int:
    """Replace placeholder paragraphs/cells and return the replacement count."""
    replacements = 0
    needle = f"{{{{{old_text}}}}}"
    replacement = f"{{{{{new_text}}}}}"

    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == needle:
            _replace_paragraph_text(paragraph, replacement)
            replacements += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip() == needle:
                        _replace_paragraph_text(paragraph, replacement)
                        replacements += 1

    return replacements


def generate_template(locale: str, output_path: Path) -> None:
    """Generate one localized DOCX template."""
    if not BLANK_TEMPLATE_SRC.is_file():
        raise FileNotFoundError(f"Blank template not found: {BLANK_TEMPLATE_SRC}")

    placeholders = get_locale_placeholders(locale)
    doc = Document(str(BLANK_TEMPLATE_SRC))

    title_count = _replace_placeholder_text(doc, "标题", placeholders["title"])
    body_count = _replace_placeholder_text(doc, "正文", placeholders["body"])

    if title_count == 0:
        raise RuntimeError("Blank template does not contain {{标题}}")
    if body_count == 0:
        raise RuntimeError("Blank template does not contain {{正文}}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def install_templates(output_dir: Path) -> int:
    """Copy generated DOCX templates into the runtime templates directory."""
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    installed = 0
    for src in sorted(output_dir.glob("*.docx"), key=lambda item: item.name.casefold()):
        shutil.copy2(src, TEMPLATES_DIR / src.name)
        logger.info("Installed template: %s", src.name)
        installed += 1
    return installed


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    locale_template_names = discover_locale_template_names()
    parser = argparse.ArgumentParser(
        description="Generate localized DOCX templates from the blank template",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument(
        "--locale",
        nargs="+",
        choices=sorted(locale_template_names),
        default=sorted(locale_template_names),
        help="Locales to generate (default: all discovered locales)",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=DEFAULT_OUTPUT_DIR,
        help=f"Output directory (default: {DEFAULT_OUTPUT_DIR})",
    )
    parser.add_argument(
        "--install",
        action="store_true",
        help="Copy generated templates into templates/ after generation",
    )
    parser.set_defaults(locale_template_names=locale_template_names)
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)
    locale_template_names: dict[str, str] = args.locale_template_names
    output_dir: Path = args.output_dir
    results: dict[str, bool] = {}

    logger.info("=" * 60)
    logger.info("Localized DOCX template generation")
    logger.info("Blank template: %s", BLANK_TEMPLATE_SRC)
    logger.info("Output directory: %s", output_dir)
    logger.info("Locales: %s", ", ".join(args.locale))
    logger.info("=" * 60)

    for locale in args.locale:
        template_name = locale_template_names[locale]
        output_path = output_dir / f"{template_name}.docx"
        try:
            generate_template(locale, output_path)
        except Exception as exc:
            results[locale] = False
            logger.error("%s failed: %s", locale, exc, exc_info=True)
            continue

        results[locale] = True
        logger.info("%s generated: %s (%d bytes)", locale, output_path, output_path.stat().st_size)

    success_count = sum(1 for ok in results.values() if ok)
    fail_count = len(results) - success_count

    if args.install and success_count:
        installed = install_templates(output_dir)
        logger.info("Installed %d templates into %s", installed, TEMPLATES_DIR)

    logger.info("Generation summary: success=%d / %d, failed=%d", success_count, len(results), fail_count)
    return 1 if fail_count else 0


if __name__ == "__main__":
    raise SystemExit(main())
