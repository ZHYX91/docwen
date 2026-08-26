from __future__ import annotations

import argparse
import hashlib
import os
import zipfile
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

SUCCESSFUL_WARNING_ACTION = "gongwen"
SUCCESSFUL_WARNING_CODE = "GONGWEN-NEEDS-REVIEW"
SUCCESSFUL_WARNING_MESSAGE = "缺少必需字段：成文日期、发文机关署名"
SUCCESSFUL_WARNING_FIXTURE_SHA256 = "3e1980324a79b5df9ce4ca96017784c530a8e8c9cd6c60e60a583a1408d4efd7"

_FIXED_ZIP_TIMESTAMP = (2024, 1, 1, 0, 0, 0)


def write_successful_warning_fixture(output_path: Path) -> str:
    """Write the deterministic Gongwen DOCX used by packaged warning gates."""

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    target = output_path.resolve()
    target.parent.mkdir(parents=True, exist_ok=True)
    raw_path = target.with_name(f".{target.name}.{os.getpid()}.raw")
    raw_path.unlink(missing_ok=True)
    target.unlink(missing_ok=True)
    try:
        document = Document()
        document.core_properties.author = "DocWen"
        document.core_properties.last_modified_by = "DocWen"
        document.core_properties.created = datetime(2024, 1, 1, tzinfo=UTC)
        document.core_properties.modified = datetime(2024, 1, 1, tzinfo=UTC)
        document.core_properties.revision = 1

        title = document.add_paragraph("关于进一步规范公文处理工作的通知")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _style_run(title.runs[0], font_name="小标宋", size=22)

        paragraphs = (
            "国办发〔2024〕5号",
            "各省、自治区、直辖市人民政府办公厅：",
            "为进一步规范公文处理工作，现就有关事项通知如下：",
            "一、严格公文格式标准。",
            "二、规范公文流转程序。",
            "附件：1. 公文格式标准",
            "          2. 公文流转程序",
            "抄送：省委组织部、省人民政府办公厅",
            "国务院办公厅　　　　2024年1月15日印发",
        )
        for text in paragraphs:
            paragraph = document.add_paragraph(text)
            _style_run(paragraph.runs[0], font_name="仿宋", size=15)

        document.save(str(raw_path))
        _repack_deterministically(raw_path, target)
    finally:
        raw_path.unlink(missing_ok=True)

    digest = hashlib.sha256(target.read_bytes()).hexdigest()
    if SUCCESSFUL_WARNING_FIXTURE_SHA256 and digest != SUCCESSFUL_WARNING_FIXTURE_SHA256:
        target.unlink(missing_ok=True)
        raise RuntimeError(
            f"successful_warning_fixture_hash_mismatch: expected={SUCCESSFUL_WARNING_FIXTURE_SHA256}; actual={digest}"
        )
    return digest


def _style_run(run: Any, *, font_name: str, size: int) -> None:
    from docx.shared import Pt

    font = run.font
    font.name = font_name
    font.size = Pt(size)


def _repack_deterministically(source: Path, target: Path) -> None:
    with (
        zipfile.ZipFile(source, "r") as input_zip,
        zipfile.ZipFile(
            target,
            "w",
            compression=zipfile.ZIP_DEFLATED,
            compresslevel=9,
        ) as output_zip,
    ):
        for source_info in sorted(input_zip.infolist(), key=lambda item: item.filename):
            target_info = zipfile.ZipInfo(source_info.filename, _FIXED_ZIP_TIMESTAMP)
            target_info.compress_type = zipfile.ZIP_DEFLATED
            target_info.create_system = 0
            target_info.external_attr = 0
            target_info.flag_bits = source_info.flag_bits
            output_zip.writestr(target_info, input_zip.read(source_info.filename), compress_type=zipfile.ZIP_DEFLATED)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Write the deterministic packaged successful-warning DOCX.")
    parser.add_argument("output", type=Path)
    args = parser.parse_args(argv)
    digest = write_successful_warning_fixture(args.output)
    print(f"successful_warning_fixture_ok: {args.output.resolve()}; sha256={digest}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
