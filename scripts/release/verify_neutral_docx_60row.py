"""Build and headlessly verify a local noncandidate 60-row DOCX rehearsal artifact."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import subprocess
import sys
from pathlib import Path
from typing import Any, NoReturn
from zipfile import ZipFile

from docx import Document
from lxml import etree

from docwen_core.docx_semantics import DocxSemanticImporter, DocxSemanticRenderer
from docwen_core.models.semantic_document import (
    SemanticBibliographyFragment,
    SemanticCaption,
    SemanticCitationCluster,
    SemanticCitationItem,
    SemanticDocument,
    SemanticParagraph,
    SemanticReference,
    SemanticTable,
    SemanticTableCell,
    SemanticText,
)

ARTIFACT_CLASS = "LOCAL_NONCANDIDATE_REHEARSAL_ARTIFACT"
DOCX_NAME = f"{ARTIFACT_CLASS}-neutral-60-row.docx"
REPORT_NAME = f"{ARTIFACT_CLASS}-neutral-60-row-report.json"
MANUAL_CHECKLIST_NAME = f"{ARTIFACT_CLASS}-neutral-60-row-manual-checklist.md"
REPORT_SCHEMA = "docwen-neutral-docx-60row-headless-report-v1"
EXPECTED_DATA_REGIONS = tuple(f"Region{index:02d}" for index in range(1, 61))


class RehearsalError(RuntimeError):
    """Fail-closed error for an invalid request or failed automatic gate."""


def _fail(code: str, detail: str | None = None) -> NoReturn:
    raise RehearsalError(f"{code}: {detail}" if detail else code)


def _require(condition: bool, gate: str, detail: str | None = None) -> None:
    if not condition:
        _fail(f"automatic_gate_failed.{gate}", detail)


def _source_repo() -> Path:
    return Path(__file__).resolve().parents[2]


def _git(repo: Path, *args: str) -> str:
    completed = subprocess.run(
        ["git", "-C", str(repo), *args],
        check=False,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="strict",
    )
    if completed.returncode != 0:
        _fail("git_command_failed", completed.stderr.strip() or " ".join(args))
    return completed.stdout.rstrip("\n")


def _resolve_source_repo(repo: Path) -> Path:
    try:
        requested = repo.resolve(strict=True)
    except (OSError, RuntimeError) as exc:
        _fail("source_repo_invalid", str(exc))
    if not requested.is_dir():
        _fail("source_repo_not_directory", str(requested))
    discovered = Path(_git(requested, "rev-parse", "--show-toplevel")).resolve(strict=True)
    if discovered != requested:
        _fail("source_repo_must_be_git_top_level", f"requested={requested}; discovered={discovered}")
    return requested


def _prepare_output_root(output_dir: Path, *, source_repo: Path) -> Path:
    if not output_dir.is_absolute():
        _fail("output_dir_must_be_absolute", str(output_dir))
    requested_parent = output_dir.parent
    if not requested_parent.exists():
        _fail("output_parent_missing", str(requested_parent))
    if not requested_parent.is_dir():
        _fail("output_parent_not_directory", str(requested_parent))
    if requested_parent.is_symlink() or requested_parent.is_junction():
        _fail("output_parent_link_or_junction_rejected", str(requested_parent))
    try:
        parent = requested_parent.resolve(strict=True)
    except (OSError, RuntimeError) as exc:
        _fail("output_parent_invalid", str(exc))
    output_root = parent / output_dir.name
    if "NONCANDIDATE" not in output_root.name.upper():
        _fail("output_dir_missing_noncandidate_marker", str(output_root))
    if output_root.exists() or output_root.is_symlink():
        _fail("output_dir_exists", str(output_root))
    source_resolved = source_repo.resolve(strict=True)
    try:
        output_root.relative_to(source_resolved)
    except ValueError:
        pass
    else:
        _fail("output_dir_inside_source_repo", str(output_root))
    return output_root


def build_neutral_document_60row() -> SemanticDocument:
    """Return explicit provider-neutral semantics for a 62x4 table and its surrounding content."""

    cells = [
        SemanticTableCell(0, 0, "Region", "corner_header", row_span=2),
        SemanticTableCell(0, 1, "Sales", "column_header", column_span=2),
        SemanticTableCell(0, 3, "Total", "column_header"),
        SemanticTableCell(1, 1, "Online", "column_header"),
        SemanticTableCell(1, 2, "Retail", "column_header"),
        SemanticTableCell(1, 3, "Combined", "column_header"),
    ]
    for data_index in range(1, 61):
        row = data_index + 1
        suffix = f"{data_index:02d}"
        cells.extend(
            (
                SemanticTableCell(row, 0, f"Region{suffix}", "row_header"),
                SemanticTableCell(row, 1, f"Online{suffix}", "data"),
                SemanticTableCell(row, 2, f"Retail{suffix}", "data"),
                SemanticTableCell(row, 3, f"Total{suffix}", "data"),
            )
        )

    table = SemanticTable(
        row_count=62,
        column_count=4,
        cells=tuple(cells),
        repeat_header="always",
        caption=SemanticCaption(
            kind="table",
            target_id="tbl-neutral-60",
            cached_number="60",
            label="Table",
            content="Provider-neutral 60-row sales table",
        ),
    )
    return SemanticDocument(
        blocks=(
            table,
            SemanticParagraph(
                (
                    SemanticText("See table "),
                    SemanticReference("tbl-neutral-60", "60"),
                    SemanticText(" for all regions."),
                )
            ),
            SemanticParagraph(
                (
                    SemanticText("Citations "),
                    SemanticCitationCluster(
                        cluster_id="cluster-single",
                        items=(SemanticCitationItem("source-one"),),
                        cached_result="[1]",
                    ),
                    SemanticText(" and "),
                    SemanticCitationCluster(
                        cluster_id="cluster-multiple",
                        items=(SemanticCitationItem("source-two"), SemanticCitationItem("source-one")),
                        cached_result="[2, 1]",
                    ),
                    SemanticText("."),
                )
            ),
            SemanticParagraph((SemanticText("References"),)),
            SemanticParagraph((SemanticText("Tail paragraph after empty bibliography."),)),
        )
    )


def _render_and_save(semantic_document: SemanticDocument, output: Path) -> None:
    document = Document()
    renderer = DocxSemanticRenderer(document)
    renderer.render_blocks(semantic_document.blocks[:-1])
    bibliography_anchor = document.add_paragraph("exclusive empty bibliography placeholder anchor")
    tail = semantic_document.blocks[-1]
    if not isinstance(tail, SemanticParagraph):
        _fail("semantic_tail_must_be_paragraph")
    renderer.render_paragraph(tail)
    rendered = renderer.render_bibliography_fragment(
        SemanticBibliographyFragment(entries=()),
        placeholder_anchor=bibliography_anchor,
    )
    _require(rendered == (), "empty_bibliography_render_elides_anchor")
    document.save(output)


def _package_xml_parts(package: ZipFile) -> dict[str, bytes]:
    return {name: package.read(name) for name in package.namelist() if name.endswith((".xml", ".rels"))}


def verify_docx_60row(output: Path, semantic_document: SemanticDocument) -> dict[str, bool]:
    """Verify raw package/XML evidence and an exact python-docx reopen/import round trip."""

    gates: dict[str, bool] = {}
    table = semantic_document.blocks[0]
    if not isinstance(table, SemanticTable):
        _fail("automatic_gate_failed.semantic_first_block_is_table")
    _require((table.row_count, table.column_count) == (62, 4), "semantic_dimensions_62x4")
    gates["semantic_dimensions_62x4"] = True
    _require(table.repeat_header == "always", "semantic_repeat_header_always")
    gates["semantic_repeat_header_always"] = True
    roles = {cell.role for cell in table.cells}
    _require(roles == {"corner_header", "column_header", "row_header", "data"}, "semantic_roles_all_four")
    gates["semantic_roles_all_four"] = True
    model_regions = tuple(cell.text for cell in table.cells if cell.text.startswith("Region") and cell.text != "Region")
    _require(model_regions == EXPECTED_DATA_REGIONS, "semantic_exact_60_data_regions")
    gates["semantic_exact_60_data_regions"] = True
    _require(table.caption is not None and table.caption.cached_number == "60", "semantic_table_caption_60")
    gates["semantic_table_caption_60"] = True

    try:
        with ZipFile(output) as package:
            document_xml = package.read("word/document.xml")
            xml_parts = _package_xml_parts(package)
    except (OSError, KeyError) as exc:
        _fail("docx_package_unreadable", str(exc))

    root = etree.fromstring(document_xml)
    word_namespace = root.nsmap.get("w")
    _require(word_namespace is not None, "wordprocessingml_namespace_present")
    namespaces = {"w": word_namespace}
    rows = root.xpath(".//w:tbl/w:tr", namespaces=namespaces)
    _require(len(rows) == 62, "raw_ooxml_table_rows_62", str(len(rows)))
    gates["raw_ooxml_table_rows_62"] = True
    header_flags = [bool(row.xpath("./w:trPr/w:tblHeader", namespaces=namespaces)) for row in rows]
    _require(header_flags == [True, True, *([False] * 60)], "raw_ooxml_repeat_headers_first_two_only")
    gates["raw_ooxml_repeat_headers_first_two_only"] = True
    grid_spans = root.xpath(".//w:tbl/w:tr/w:tc/w:tcPr/w:gridSpan/@w:val", namespaces=namespaces)
    _require(grid_spans == ["2"], "raw_ooxml_horizontal_merge")
    gates["raw_ooxml_horizontal_merge"] = True
    vertical_restarts = root.xpath(".//w:tbl/w:tr/w:tc/w:tcPr/w:vMerge[@w:val='restart']", namespaces=namespaces)
    vertical_continuations = root.xpath(
        ".//w:tbl/w:tr/w:tc/w:tcPr/w:vMerge[not(@w:val) or @w:val='continue']", namespaces=namespaces
    )
    _require(len(vertical_restarts) == 1 and len(vertical_continuations) == 1, "raw_ooxml_vertical_merge")
    gates["raw_ooxml_vertical_merge"] = True
    conditional_styles = root.xpath(".//w:tbl//w:cnfStyle", namespaces=namespaces)
    _require(len(conditional_styles) >= 2, "raw_ooxml_conditional_styles")
    gates["raw_ooxml_conditional_styles"] = True

    visible_text = tuple(root.xpath(".//w:t/text()", namespaces=namespaces))
    raw_regions = tuple(text for text in visible_text if text.startswith("Region") and text != "Region")
    _require(raw_regions == EXPECTED_DATA_REGIONS, "raw_ooxml_exact_60_visible_regions")
    _require("Region01" in visible_text and "Region60" in visible_text, "raw_ooxml_region_endpoints")
    _require("Region61" not in visible_text, "raw_ooxml_region61_absent")
    gates["raw_ooxml_exact_60_visible_regions"] = True
    gates["raw_ooxml_region_endpoints_and_61_absent"] = True

    instructions = tuple(root.xpath(".//w:instrText/text()", namespaces=namespaces))
    normalized_instructions = tuple(" ".join(instruction.split()) for instruction in instructions)
    _require(any(item.startswith("SEQ Table ") for item in normalized_instructions), "raw_ooxml_seq_field")
    _require(any(item.startswith("REF _DW_") for item in normalized_instructions), "raw_ooxml_ref_field")
    _require("CITATION source-one" in normalized_instructions, "raw_ooxml_single_citation_field")
    _require(
        r"CITATION source-two \m source-one" in normalized_instructions,
        "raw_ooxml_multiple_citation_field",
    )
    gates["raw_ooxml_seq_ref_and_citation_fields"] = True
    locked_begins = root.xpath(
        ".//w:fldChar[@w:fldCharType='begin'][@w:fldLock='true' or @w:fldLock='1' or @w:fldLock='on']",
        namespaces=namespaces,
    )
    _require(len(locked_begins) == 2, "raw_ooxml_two_locked_citations")
    _require(
        all(begin.get(f"{{{word_namespace}}}dirty") is None for begin in locked_begins),
        "raw_ooxml_locked_citations_not_dirty",
    )
    gates["raw_ooxml_two_citations_locked_and_clean"] = True

    all_xml = b"\n".join(xml_parts.values())
    all_xml_lower = all_xml.lower()
    _require(b"_DWB_BIBLIOGRAPHY" not in all_xml, "raw_package_bibliography_boundary_absent")
    _require(b"_DWE_" not in all_xml, "raw_package_bibliography_entries_absent")
    _require(
        not any(item.upper().startswith("BIBLIOGRAPHY") for item in normalized_instructions),
        "raw_ooxml_bibliography_field_absent",
    )
    bibliography_namespace = "http://schemas.openxmlformats.org/officeDocument/2006/bibliography"
    bibliography_sources = sum(
        len(
            etree.fromstring(content).xpath(
                "//*[local-name()='Source' and namespace-uri()=$namespace]",
                namespace=bibliography_namespace,
            )
        )
        for content in xml_parts.values()
    )
    _require(bibliography_sources == 0, "raw_package_bibliography_source_records_absent")
    _require(
        all(token not in all_xml_lower for token in (b"wenleaf", b"csl", b"pkwf")),
        "raw_package_provider_tokens_absent",
    )
    gates["raw_package_empty_bibliography_absent"] = True
    gates["raw_package_provider_tokens_absent"] = True

    reopened = Document(output)
    imported = DocxSemanticImporter().import_document(reopened)
    _require(imported.diagnostics == (), "reopened_import_diagnostics_empty", repr(imported.diagnostics))
    _require(imported.document == semantic_document, "reopened_exact_semantic_round_trip")
    _require(imported.document.bibliography is None, "reopened_bibliography_is_none")
    gates["reopened_import_diagnostics_empty"] = True
    gates["reopened_exact_semantic_round_trip"] = True
    gates["reopened_bibliography_is_none"] = True
    return gates


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while chunk := handle.read(1024 * 1024):
            digest.update(chunk)
    return digest.hexdigest()


def _write_text_exclusive(path: Path, content: str) -> None:
    with path.open("x", encoding="utf-8", newline="\n") as handle:
        handle.write(content)
        handle.flush()
        os.fsync(handle.fileno())


def _manual_checklist() -> str:
    checks = (
        "Pagination and table continuity across pages",
        "Two-row repeated header on every continuation page",
        "Horizontal Sales merge and vertical Region header merge",
        "Caption and cross-reference display the same number after viewer field refresh",
        "Single-item and multi-item citation display",
        "References heading remains, no bibliography entries appear, and the tail paragraph remains",
    )
    rows = "\n".join(f"| {check} | READY_FOR_MANUAL | READY_FOR_MANUAL | READY_FOR_MANUAL |" for check in checks)
    sections = "\n\n".join(
        f"## {viewer}\n\nStatus: READY_FOR_MANUAL" for viewer in ("Microsoft Word", "WPS Writer", "LibreOffice Writer")
    )
    return (
        f"# {ARTIFACT_CLASS}: neutral 60-row DOCX manual checklist\n\n"
        "This checklist records no viewer execution. Each viewer starts at READY_FOR_MANUAL.\n\n"
        "The SEQ/REF source cache starts at 60. A viewer may recalculate the first table and its "
        "reference to 1; acceptance requires the two displayed numbers to match, not preservation "
        "of the cached value.\n\n"
        "| Check | Microsoft Word | WPS Writer | LibreOffice Writer |\n"
        "| --- | --- | --- | --- |\n"
        f"{rows}\n\n"
        f"{sections}\n"
    )


def _source_state(repo: Path) -> dict[str, Any]:
    porcelain = _git(repo, "status", "--porcelain=v1", "--untracked-files=all")
    return {
        "repo": str(repo),
        "head": _git(repo, "rev-parse", "HEAD"),
        "tree": _git(repo, "show", "-s", "--format=%T", "HEAD"),
        "remotes": _git(repo, "remote").splitlines(),
        "status": "CLEAN" if not porcelain else "DIRTY",
        "porcelain": porcelain.splitlines(),
    }


def run_rehearsal(output_dir: Path, *, source_repo: Path | None = None) -> dict[str, Any]:
    """Create a new output root and emit verified noncandidate rehearsal evidence."""

    repo = _resolve_source_repo(source_repo or _source_repo())
    output_root = _prepare_output_root(output_dir, source_repo=repo)
    source_before = _source_state(repo)
    output_root.mkdir()
    docx_path = output_root / DOCX_NAME
    checklist_path = output_root / MANUAL_CHECKLIST_NAME
    report_path = output_root / REPORT_NAME

    semantic_document = build_neutral_document_60row()
    _render_and_save(semantic_document, docx_path)
    gates = verify_docx_60row(docx_path, semantic_document)
    _require(bool(gates) and all(gates.values()), "all_recorded_gates_true")

    _write_text_exclusive(checklist_path, _manual_checklist())
    source_after = _source_state(repo)
    _require(source_after == source_before, "source_snapshot_stable_during_generation")
    gates["source_snapshot_stable_during_generation"] = True
    overall = "SOURCE_HEADLESS_PASS" if source_before["status"] == "CLEAN" else "SOURCE_HEADLESS_VERIFIED_DIRTY_SOURCE"
    report: dict[str, Any] = {
        "schema": REPORT_SCHEMA,
        "artifact_class": ARTIFACT_CLASS,
        "overall": overall,
        "source": source_before,
        "artifacts": {
            "docx": {
                "name": docx_path.name,
                "size_bytes": docx_path.stat().st_size,
                "sha256": _sha256(docx_path),
            },
            "manual_checklist": {
                "name": checklist_path.name,
                "size_bytes": checklist_path.stat().st_size,
                "sha256": _sha256(checklist_path),
            },
        },
        "automatic_gates": gates,
        "evidence_layers": {
            "source_semantics": "HEADLESS_AUTOMATED",
            "raw_docx_package_xml": "HEADLESS_AUTOMATED",
            "python_docx_reopen_import": "HEADLESS_AUTOMATED",
            "viewer_word": "READY_FOR_MANUAL",
            "viewer_wps": "READY_FOR_MANUAL",
            "viewer_libreoffice": "READY_FOR_MANUAL",
        },
        "claim_boundaries": [
            "Local noncandidate rehearsal artifact only.",
            "No Microsoft Word, WPS Writer, or LibreOffice Writer execution was performed.",
            "The cached SEQ/REF result is 60; a viewer may recalculate the first table and reference to 1.",
            "No candidate, release, publication, or final-package claim is made.",
        ],
    }
    _write_text_exclusive(report_path, json.dumps(report, ensure_ascii=False, sort_keys=True, indent=2) + "\n")
    return report


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--output-dir",
        required=True,
        type=Path,
        help="Absolute, nonexistent output directory outside the source repo; its name must contain NONCANDIDATE.",
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    try:
        report = run_rehearsal(args.output_dir)
    except RehearsalError as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 2
    print(json.dumps({"artifact_class": report["artifact_class"], "overall": report["overall"]}, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
