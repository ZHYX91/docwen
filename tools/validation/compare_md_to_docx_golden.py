"""Compare two MD→DOCX conversion outputs — golden baseline vs. new run.

Usage as CLI:
    PYTHONPATH=. .venv/Scripts/python tools/validation/compare_md_to_docx_golden.py \
        --old-docx tests/fixtures/golden/md_to_docx_old/sample_golden.docx \
        --new-docx /tmp/output.docx \
        --json-output /tmp/result.json

Usage as library:
    from tools.validation.compare_md_to_docx_golden import compare_docx_files
    result = compare_docx_files(old_path, new_path)
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

# ── Normalization ──────────────────────────────────────────────────────


def normalize_whitespace(text: str) -> str:
    """Collapse all whitespace sequences to a single space, then strip."""
    return re.sub(r"\s+", " ", text).strip()


# ── Extraction ─────────────────────────────────────────────────────────


def _paragraph_visible_text(paragraph: Any) -> str:
    """Return the complete visible text including inline SDT carriers.

    Fenced-source carriers wrap visible payload runs in one inline SDT, which
    python-docx ``Paragraph.text`` does not traverse.  The recovery readers and
    this golden comparison use the same full-w:t projection.
    """

    from docx.oxml.ns import qn

    return "".join(item.text or "" for item in paragraph._p.iter(f"{qn('w:t')}"))


def _extract_paragraphs(doc: Any) -> list[dict[str, Any]]:
    """Return every paragraph with its style name and normalized text."""
    result: list[dict[str, Any]] = []
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else "None"
        text = normalize_whitespace(_paragraph_visible_text(p))
        result.append({"style": style_name, "text": text})
    return result


def _extract_headings(doc: Any) -> list[dict[str, Any]]:
    """Extract heading paragraphs (style matches ``Heading N`` or ``Title``).

    Returns a list of ``{"level": int, "text": str, "style": str}`` dicts.
    ``Title`` is mapped to level 0.
    """
    result: list[dict[str, Any]] = []
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        text = normalize_whitespace(_paragraph_visible_text(p))
        if not text:
            continue
        m = re.match(r"^Heading (\d+)", style_name)
        if m:
            result.append({"level": int(m.group(1)), "text": text, "style": style_name})
        elif style_name == "Title":
            result.append({"level": 0, "text": text, "style": "Title"})
    return result


def _extract_tables(doc: Any) -> list[list[list[str]]]:
    """Extract all tables as a list of 2D cell-value arrays (normalized)."""
    result: list[list[list[str]]] = []
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [normalize_whitespace(cell.text) for cell in row.cells]
            rows.append(cells)
        result.append(rows)
    return result


# ── Comparison helpers ─────────────────────────────────────────────────


def _heading_key(h: dict[str, Any]) -> tuple[int, str]:
    """Sortable / comparable key for a heading dict."""
    return (h["level"], h["text"])


def _compare_headings(
    old_h: list[dict[str, Any]],
    new_h: list[dict[str, Any]],
    *,
    allowed_new_title: str | None = None,
    allowed_heading_body_merge: tuple[int, str, str] | None = None,
) -> dict[str, Any]:
    """Compare heading sequences and return a result dict."""
    old_keys = [_heading_key(h) for h in old_h]
    new_keys = [_heading_key(h) for h in new_h]

    allowed_title_applied = False
    if (
        allowed_new_title
        and new_keys
        and new_keys[0] == (0, normalize_whitespace(allowed_new_title))
        and (not old_keys or old_keys[0] != new_keys[0])
    ):
        new_keys = new_keys[1:]
        allowed_title_applied = True

    allowed_heading_body_merge_applied = False
    if allowed_heading_body_merge is not None:
        level, old_heading, merged_heading = allowed_heading_body_merge
        old_key = (level, normalize_whitespace(old_heading))
        merged_key = (level, normalize_whitespace(merged_heading))
        if (
            old_keys.count(old_key) == 1
            and merged_key not in old_keys
            and new_keys.count(merged_key) == 1
            and old_key not in new_keys
        ):
            old_keys[old_keys.index(old_key)] = merged_key
            allowed_heading_body_merge_applied = True

    match = old_keys == new_keys
    details: dict[str, Any] = {
        "old_count": len(old_keys),
        "new_count": len(new_keys),
        "allowed_new_title_applied": allowed_title_applied,
        "allowed_heading_body_merge_applied": allowed_heading_body_merge_applied,
    }

    if not match:
        missing = [h for h in old_h if _heading_key(h) not in new_keys]
        extra = [h for h in new_h if _heading_key(h) not in old_keys]
        if missing:
            details["missing_in_new"] = [{"level": h["level"], "text": h["text"]} for h in missing]
        if extra:
            details["extra_in_new"] = [{"level": h["level"], "text": h["text"]} for h in extra]

        # Positional diffs: same entries, different order
        if not missing and not extra and old_keys != new_keys:
            details["order_differs"] = True
            details["old_sequence"] = [{"level": level, "text": t} for level, t in old_keys]
            details["new_sequence"] = [{"level": level, "text": t} for level, t in new_keys]

    return {"match": match, **details}


def _compare_tables(
    old_t: list[list[list[str]]],
    new_t: list[list[list[str]]],
) -> dict[str, Any]:
    """Compare table contents cell by cell."""
    match = old_t == new_t
    details: dict[str, Any] = {
        "old_count": len(old_t),
        "new_count": len(new_t),
    }
    if not match:
        details["table_count_match"] = len(old_t) == len(new_t)
        diffs: list[dict[str, Any]] = []
        max_t = max(len(old_t), len(new_t))
        for ti in range(max_t):
            old_rows = old_t[ti] if ti < len(old_t) else None
            new_rows = new_t[ti] if ti < len(new_t) else None
            if old_rows != new_rows:
                diffs.append(
                    {
                        "table_index": ti,
                        "old": old_rows,
                        "new": new_rows,
                    }
                )
        if diffs:
            details["diffs"] = diffs
    return {"match": match, **details}


def _compare_paragraphs(
    old_p: list[dict[str, Any]],
    new_p: list[dict[str, Any]],
    *,
    allowed_removed_old_paragraph: str | None = None,
    allowed_heading_body_merge: tuple[int, str, str] | None = None,
) -> dict[str, Any]:
    """Compare body paragraphs (non-heading, non-empty) between old and new.

    Compares non-heading text in document order. Blank paragraph differences
    remain ignored, but moving visible content to another part of the document
    is a fidelity failure rather than an acceptable normalization.
    """
    old_texts = [
        p["text"] for p in old_p if p["text"] and not re.match(r"^Heading \d+", p["style"]) and p["style"] != "Title"
    ]
    new_texts = [
        p["text"] for p in new_p if p["text"] and not re.match(r"^Heading \d+", p["style"]) and p["style"] != "Title"
    ]

    allowed_removed_old_paragraph_applied = False
    if allowed_removed_old_paragraph:
        normalized_allowed = normalize_whitespace(allowed_removed_old_paragraph)
        if old_texts.count(normalized_allowed) == 1 and normalized_allowed not in new_texts:
            old_texts.remove(normalized_allowed)
            allowed_removed_old_paragraph_applied = True

    allowed_heading_body_merge_applied = False
    if allowed_heading_body_merge is not None:
        _level, old_heading, merged_heading = allowed_heading_body_merge
        normalized_old_heading = normalize_whitespace(old_heading)
        normalized_merged_heading = normalize_whitespace(merged_heading)
        if normalized_merged_heading.startswith(normalized_old_heading):
            merged_body = normalized_merged_heading[len(normalized_old_heading) :].strip()
            if merged_body and old_texts.count(merged_body) == 1 and merged_body not in new_texts:
                old_texts.remove(merged_body)
                allowed_heading_body_merge_applied = True

    match = old_texts == new_texts
    details: dict[str, Any] = {
        "old_count": len(old_texts),
        "new_count": len(new_texts),
        "allowed_removed_old_paragraph": allowed_removed_old_paragraph,
        "allowed_removed_old_paragraph_applied": allowed_removed_old_paragraph_applied,
        "allowed_heading_body_merge_applied": allowed_heading_body_merge_applied,
    }

    if not match:
        positional_diffs: list[dict[str, Any]] = []
        for index in range(max(len(old_texts), len(new_texts))):
            old_text = old_texts[index] if index < len(old_texts) else None
            new_text = new_texts[index] if index < len(new_texts) else None
            if old_text != new_text:
                positional_diffs.append({"index": index, "old": old_text, "new": new_text})
            if len(positional_diffs) >= 20:
                break
        details["positional_diffs"] = positional_diffs

    return {"match": match, **details}


# ── Main entry point ───────────────────────────────────────────────────


def compare_docx_files(
    old_path: str | Path,
    new_path: str | Path,
    *,
    allowed_new_title: str | None = None,
    allowed_removed_old_paragraph: str | None = None,
    allowed_heading_body_merge: tuple[int, str, str] | None = None,
) -> dict[str, Any]:
    """Compare two DOCX files and return a detailed comparison dict.

    Returns a dict with top-level keys:
    - ``paragraphs_match`` (bool)
    - ``headings_match`` (bool)
    - ``tables_match`` (bool)
    - ``passed`` (bool) — True when all three match
    - ``details`` (dict) — per-category detail including diffs
    - ``old_paragraph_count``, ``new_paragraph_count`` (int)
    - ``old_heading_count``, ``new_heading_count`` (int)
    - ``old_table_count``, ``new_table_count`` (int)
    """
    from docx import Document  # lazy import so CLI --help is fast

    old_doc = Document(str(old_path))
    new_doc = Document(str(new_path))

    old_paras = _extract_paragraphs(old_doc)
    new_paras = _extract_paragraphs(new_doc)
    old_headings = _extract_headings(old_doc)
    new_headings = _extract_headings(new_doc)
    old_tables = _extract_tables(old_doc)
    new_tables = _extract_tables(new_doc)

    para_result = _compare_paragraphs(
        old_paras,
        new_paras,
        allowed_removed_old_paragraph=allowed_removed_old_paragraph,
        allowed_heading_body_merge=allowed_heading_body_merge,
    )
    heading_result = _compare_headings(
        old_headings,
        new_headings,
        allowed_new_title=allowed_new_title,
        allowed_heading_body_merge=allowed_heading_body_merge,
    )
    table_result = _compare_tables(old_tables, new_tables)

    paragraphs_match: bool = para_result["match"]
    headings_match: bool = heading_result["match"]
    tables_match: bool = table_result["match"]
    passed = paragraphs_match and headings_match and tables_match

    return {
        "paragraphs_match": paragraphs_match,
        "headings_match": headings_match,
        "tables_match": tables_match,
        "passed": passed,
        "old_paragraph_count": len(old_paras),
        "new_paragraph_count": len(new_paras),
        "old_heading_count": len(old_headings),
        "new_heading_count": len(new_headings),
        "old_table_count": len(old_tables),
        "new_table_count": len(new_tables),
        "details": {
            "paragraphs": para_result,
            "headings": heading_result,
            "tables": table_result,
        },
    }


# ── CLI adapter ────────────────────────────────────────────────────────


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Compare two MD→DOCX conversion outputs (golden vs. new).",
    )
    parser.add_argument(
        "--old-docx",
        required=True,
        help="Path to the golden (old-system) DOCX file.",
    )
    parser.add_argument(
        "--new-docx",
        required=True,
        help="Path to the new-system DOCX file.",
    )
    parser.add_argument(
        "--json-output",
        default=None,
        help="If given, write the JSON result to this file.",
    )
    parser.add_argument(
        "--allowed-new-title",
        default=None,
        help="Allow one exact leading Title paragraph in the new DOCX.",
    )
    parser.add_argument(
        "--allowed-removed-old-paragraph",
        default=None,
        help="Allow removal of one exact unique visible paragraph from the old DOCX.",
    )
    return parser


def main(argv: list[str] | None = None) -> dict[str, Any]:
    """CLI entry point.  Returns the comparison dict (also prints it)."""
    parser = _build_parser()
    args = parser.parse_args(argv)

    result = compare_docx_files(
        args.old_docx,
        args.new_docx,
        allowed_new_title=args.allowed_new_title,
        allowed_removed_old_paragraph=args.allowed_removed_old_paragraph,
    )

    json_text = json.dumps(result, ensure_ascii=False, indent=2)

    if args.json_output:
        Path(args.json_output).write_text(json_text, encoding="utf-8")
        print(f"Result written to {args.json_output}")

    print(json_text)
    return result


if __name__ == "__main__":
    main()
