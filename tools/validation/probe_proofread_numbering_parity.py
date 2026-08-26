"""Capture FA-11 real-dictionary proofreading and numbering parity evidence.

The public entry point orchestrates three isolated worker processes so each
project uses its own virtual environment and production modules. Reference
repositories remain read-only; copied configuration and all artifacts live
under the requested output directory.
"""

from __future__ import annotations

import argparse
import hashlib
import importlib
import json
import shutil
import subprocess
import sys
from dataclasses import asdict
from pathlib import Path
from typing import Any

PROJECT_KINDS = ("docwen-ref-tk", "docwen-ref-pyside6", "docwen-current")
PROOFREAD_OPTIONS = {
    "symbol_pairing": True,
    "symbol_correction": True,
    "typos_rule": True,
    "sensitive_word": False,
}
CUSTOM_SCHEME_ID = "fa11_malformed_custom"
CUSTOM_SCHEME = {
    "level_1": {"format": "卷{1.roman_upper} "},
    "level_2": {"format": "{1.roman_upper}-{2.letter_upper} "},
    "level_3": {"format": "X{10.arabic_half}/{3.not_a_style} "},
    "level_4": {"format": "{1.roman_upper}.{2.letter_upper}.{4.arabic_half} "},
    "level_6": {"format": "[{6.letter_lower}] "},
}


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def _write_json(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def _copy_reference_configs(project_root: Path, output_dir: Path, typo_fixture: Path) -> Path:
    copied = output_dir / "configs"
    shutil.copytree(project_root / "configs", copied)
    shutil.copyfile(typo_fixture, copied / "proofread_typos.toml")
    return copied


def _patch_reference_config_manager(project_kind: str, config_dir: Path) -> Any:
    from docwen.config.config_manager import ConfigManager

    manager = ConfigManager(str(config_dir))

    import docwen.text_rules.config_loader as rules_loader_module
    import docwen.text_rules.validator as validator_module

    manager_module = importlib.import_module("docwen.config.config_manager")
    manager_module.config_manager = manager
    rules_loader_module.config_manager = manager
    validator_module.config_manager = manager

    loaded_typos = manager.get_typos() if project_kind == "docwen-ref-tk" else manager.proofread_typos.typos
    return manager, loaded_typos


def _reference_worker(args: argparse.Namespace) -> dict[str, Any]:
    project_root = Path(args.project_root).resolve()
    output_dir = Path(args.output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    config_dir = _copy_reference_configs(project_root, output_dir, Path(args.typos).resolve())
    manager, loaded_typos = _patch_reference_config_manager(args.project_kind, config_dir)

    from docwen.docx_spell.core import process_docx
    from docwen.md_spell.core import process_md_file
    from docwen.utils import heading_numbering

    docx_output = output_dir / "official-government-checked.docx"
    process_docx(
        args.docx,
        output_path=str(docx_output),
        proofread_options=PROOFREAD_OPTIONS,
    )

    markdown_report = output_dir / "official-government-proofread-report.json"
    report = process_md_file(
        Path(args.markdown_projection),
        output_path=markdown_report,
        proofread_options=PROOFREAD_OPTIONS,
    )

    if args.project_kind == "docwen-ref-tk":
        original_getter = manager.get_heading_schemes
        manager.get_heading_schemes = lambda: {**original_getter(), CUSTOM_SCHEME_ID: CUSTOM_SCHEME}
    else:
        manager.numbering_add.schemes[CUSTOM_SCHEME_ID] = CUSTOM_SCHEME

    numbering_input = Path(args.numbering).read_text(encoding="utf-8")
    numbering_output_text = heading_numbering.process_md_numbering(
        numbering_input,
        remove_existing=True,
        add_new=True,
        scheme_id=CUSTOM_SCHEME_ID,
    )
    numbering_output = output_dir / "malformed-custom-numbering-numbered.md"
    numbering_output.write_text(numbering_output_text, encoding="utf-8")

    return {
        "project": args.project_kind,
        "dictionary": {
            "config_path": str(config_dir / "proofread_typos.toml"),
            "loaded_typos": loaded_typos,
        },
        "proofread": {
            "docx_output": str(docx_output),
            "docx_sha256": _sha256(docx_output),
            "markdown_report": str(markdown_report),
            "markdown_report_sha256": _sha256(markdown_report),
            "issues": [asdict(issue) for issue in report.issues],
            "summary": report.summary,
        },
        "numbering": {
            "output": str(numbering_output),
            "sha256": _sha256(numbering_output),
            "text": numbering_output_text,
        },
    }


def _current_worker(args: argparse.Namespace) -> dict[str, Any]:
    from docwen_bundle.runtime_factory import create_runtime_port
    from docwen_core.models.file_ref import FileRef
    from docwen_core.models.request import ConversionRequest, OutputPolicy
    from docwen_runtime.config import build_proofread_rules
    from docwen_runtime.config.loader import ConfigLoader
    from docwen_runtime.numbering import NumberingSchemeRegistry

    project_root = Path(args.project_root).resolve()
    output_dir = Path(args.output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    user_dir = output_dir / "user-configs"
    user_typos = user_dir / "proofread" / "typos.toml"
    user_typos.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(Path(args.typos).resolve(), user_typos)

    loader = ConfigLoader(base_dir=project_root / "configs", user_dir=user_dir)
    loaded_rules = build_proofread_rules(loader.config.as_dict())
    runtime = create_runtime_port(config_loader=loader)

    proofread_options = {
        "enable_symbol_pairing": True,
        "enable_symbol_correction": True,
        "enable_typos_rule": True,
        "enable_sensitive_word": False,
    }
    docx_request = ConversionRequest(
        request_id="fa11-current-docx",
        input_refs=[FileRef(path=str(Path(args.docx).resolve()), format="document", category="document")],
        target_format="docx",
        action_name="validate",
        options=proofread_options,
        output_policy=OutputPolicy(output_dir=str(output_dir / "docx-output")),
    )
    docx_result = runtime.execute(docx_request)
    if not docx_result.success:
        raise RuntimeError(f"current DOCX proofread failed: {docx_result.error}")

    markdown_request = ConversionRequest(
        request_id="fa11-current-markdown",
        input_refs=[
            FileRef(
                path=str(Path(args.markdown_projection).resolve()),
                format="markdown",
                category="document",
            )
        ],
        target_format="markdown",
        action_name="validate-md",
        options=proofread_options,
        output_policy=OutputPolicy(output_dir=str(output_dir / "markdown-output")),
    )
    markdown_result = runtime.execute(markdown_request)
    if not markdown_result.success:
        raise RuntimeError(f"current Markdown proofread failed: {markdown_result.error}")

    docx_output = Path(docx_result.artifacts[0].staging_path)
    markdown_report = Path(markdown_result.artifacts[0].staging_path)
    report = json.loads(markdown_report.read_text(encoding="utf-8"))

    snapshot = loader.config.as_dict()
    add_config = snapshot["numbering"]["add"]
    add_config["schemes"][CUSTOM_SCHEME_ID] = CUSTOM_SCHEME
    add_config["settings"]["order"].append(CUSTOM_SCHEME_ID)
    locale = str(snapshot.get("gui", {}).get("language", {}).get("locale") or "zh_CN")
    registry = NumberingSchemeRegistry.from_config_snapshot(
        snapshot,
        locale_path=project_root / "i18n" / "locales" / f"{locale}.toml",
    )

    from docwen_plugin_markdown.common_utils import add_md_numbering, remove_md_numbering

    numbering_input = Path(args.numbering).read_text(encoding="utf-8")
    numbering_output_text = add_md_numbering(
        remove_md_numbering(numbering_input),
        CUSTOM_SCHEME_ID,
        registry,
    )
    numbering_output = output_dir / "malformed-custom-numbering-numbered.md"
    numbering_output.write_text(numbering_output_text, encoding="utf-8")

    return {
        "project": args.project_kind,
        "dictionary": {
            "config_path": str(user_typos),
            "loaded_typos": {key: list(values) for key, values in loaded_rules.typos_map.items()},
        },
        "proofread": {
            "docx_output": str(docx_output),
            "docx_sha256": _sha256(docx_output),
            "markdown_report": str(markdown_report),
            "markdown_report_sha256": _sha256(markdown_report),
            "issues": report["issues"],
            "summary": report["summary"],
        },
        "numbering": {
            "output": str(numbering_output),
            "sha256": _sha256(numbering_output),
            "text": numbering_output_text,
        },
    }


def _worker(args: argparse.Namespace) -> int:
    result = _current_worker(args) if args.project_kind == "docwen-current" else _reference_worker(args)
    _write_json(Path(args.result), result)
    return 0


def _extract_markdown_projection(docx_path: Path, output_path: Path) -> dict[str, int]:
    from docx import Document

    document = Document(docx_path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    output_path.write_text("\n".join(paragraphs) + "\n", encoding="utf-8")
    return {"body_paragraphs": len(document.paragraphs), "nonempty_body_paragraphs": len(paragraphs)}


def _augment_anchor_evidence(projects: dict[str, Any]) -> None:
    from docwen_plugin_proofread.anchor_report import _extract_comments, build_anchor_report_markdown

    for project in projects.values():
        docx_output = Path(project["proofread"]["docx_output"])
        comments = _extract_comments(docx_output)
        anchor_report = docx_output.with_name("official-government-anchor-report.md")
        anchor_report.write_text(build_anchor_report_markdown(docx_output) + "\n", encoding="utf-8")
        project["proofread"]["docx_comments"] = [
            {
                "comment_id": comment.comment_id,
                "author": comment.author,
                "text": comment.text,
                "anchor_paragraph_index": comment.anchor_paragraph_index,
            }
            for comment in comments
        ]
        project["proofread"]["anchor_report"] = str(anchor_report)
        project["proofread"]["anchor_report_sha256"] = _sha256(anchor_report)


def _run_worker(
    *,
    interpreter: Path,
    script: Path,
    project_kind: str,
    project_root: Path,
    docx: Path,
    markdown_projection: Path,
    numbering: Path,
    typos: Path,
    output_dir: Path,
) -> dict[str, Any]:
    result_path = output_dir / "worker-result.json"
    command = [
        str(interpreter),
        str(script),
        "--worker",
        "--project-kind",
        project_kind,
        "--project-root",
        str(project_root),
        "--docx",
        str(docx),
        "--markdown-projection",
        str(markdown_projection),
        "--numbering",
        str(numbering),
        "--typos",
        str(typos),
        "--output-dir",
        str(output_dir),
        "--result",
        str(result_path),
    ]
    completed = subprocess.run(command, cwd=project_root, capture_output=True)
    stdout = completed.stdout.decode("utf-8", errors="replace")
    stderr = completed.stderr.decode("utf-8", errors="replace")
    if completed.returncode != 0:
        raise RuntimeError(
            f"{project_kind} worker failed ({completed.returncode})\nSTDOUT:\n{stdout}\nSTDERR:\n{stderr}"
        )
    return json.loads(result_path.read_text(encoding="utf-8"))


def _orchestrate(args: argparse.Namespace) -> int:
    current_root = Path(args.current_root).resolve()
    tk_root = Path(args.tk_root).resolve()
    pyside_root = Path(args.pyside_root).resolve()
    output_dir = Path(args.output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    script = Path(__file__).resolve()
    docx = Path(args.docx).resolve()
    numbering = Path(args.numbering).resolve()
    current_typos = Path(args.current_typos).resolve()
    legacy_typos = Path(args.legacy_typos).resolve()
    markdown_projection = output_dir / "official-government-body.md"
    corpus_shape = _extract_markdown_projection(docx, markdown_projection)

    specs = (
        ("docwen-ref-tk", tk_root, tk_root / ".venv" / "Scripts" / "python.exe", legacy_typos),
        (
            "docwen-ref-pyside6",
            pyside_root,
            pyside_root / ".venv" / "Scripts" / "python.exe",
            legacy_typos,
        ),
        ("docwen-current", current_root, Path(sys.executable), current_typos),
    )
    projects: dict[str, Any] = {}
    for project_kind, project_root, interpreter, typo_fixture in specs:
        projects[project_kind] = _run_worker(
            interpreter=interpreter,
            script=script,
            project_kind=project_kind,
            project_root=project_root,
            docx=docx,
            markdown_projection=markdown_projection,
            numbering=numbering,
            typos=typo_fixture,
            output_dir=output_dir / project_kind,
        )
    _augment_anchor_evidence(projects)

    result = {
        "probe_id": "FA-11-N1-N2-2026-07-22",
        "inputs": {
            "official_docx": {"path": str(docx), "sha256": _sha256(docx), **corpus_shape},
            "markdown_projection": {
                "path": str(markdown_projection),
                "sha256": _sha256(markdown_projection),
            },
            "numbering_markdown": {"path": str(numbering), "sha256": _sha256(numbering)},
            "current_typos": {"path": str(current_typos), "sha256": _sha256(current_typos)},
            "legacy_typos": {"path": str(legacy_typos), "sha256": _sha256(legacy_typos)},
            "custom_scheme_id": CUSTOM_SCHEME_ID,
            "custom_scheme": CUSTOM_SCHEME,
        },
        "projects": projects,
    }
    _write_json(output_dir / "probe-result.json", result)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


def _parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser()
    parser.add_argument("--worker", action="store_true")
    parser.add_argument("--project-kind", choices=PROJECT_KINDS)
    parser.add_argument("--project-root")
    parser.add_argument("--current-root", default=".")
    parser.add_argument("--tk-root", default="../docwen-ref-tk")
    parser.add_argument("--pyside-root", default="../docwen-ref-pyside6")
    parser.add_argument("--docx", required=True)
    parser.add_argument("--markdown-projection")
    parser.add_argument("--numbering", required=True)
    parser.add_argument("--typos")
    parser.add_argument("--current-typos")
    parser.add_argument("--legacy-typos")
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--result")
    return parser


def main() -> int:
    args = _parser().parse_args()
    return _worker(args) if args.worker else _orchestrate(args)


if __name__ == "__main__":
    raise SystemExit(main())
