"""
docx_spell 批注锚点检查（复现矩阵，一键生成样例 + 校对 + 报告）

特点：
- 仓库内不落样例 DOCX：运行时动态生成，产物写入 .tmp/ 下
- 仅检查/取证，不做任何修复
"""

from __future__ import annotations

import argparse
import base64
import contextlib
import datetime as _dt
import logging
import os
import re
import shutil
import sys
from collections.abc import Iterable
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Inches

from docwen_plugin_proofread.anchor_report import (
    build_anchor_report_markdown,
    extract_comment_texts_from_comments_xml,
    extract_occurrences_from_document_xml,
    read_docx_part,
)

logger = logging.getLogger(__name__)


_PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4nGNgYAAAAAMAAWgmWQ0AAAAASUVORK5CYII="
)


@dataclass(frozen=True)
class Expectation:
    label: str
    error_text: str
    expected_covered: str


@dataclass(frozen=True)
class Case:
    slug: str
    title: str
    source: str
    expected_path: str
    minimal_text: str
    build: callable
    expectations: tuple[Expectation, ...]


def _configure_logging(log_path: Path, verbose: bool) -> None:
    root = logging.getLogger()
    for h in list(root.handlers):
        root.removeHandler(h)
        with contextlib.suppress(Exception):
            h.close()

    root.setLevel(logging.DEBUG if verbose else logging.INFO)
    fmt = logging.Formatter("%(asctime)s %(levelname)s %(name)s: %(message)s")

    log_path.parent.mkdir(parents=True, exist_ok=True)
    fh = logging.FileHandler(log_path, encoding="utf-8")
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(fmt)
    root.addHandler(fh)

    sh = logging.StreamHandler(sys.stderr)
    sh.setLevel(logging.INFO if verbose else logging.WARNING)
    sh.setFormatter(fmt)
    root.addHandler(sh)


def _write_png_1x1(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(_PNG_1X1)


def _all_rules_enabled() -> dict[str, bool]:
    return {
        "enable_symbol_pairing": True,
        "enable_symbol_correction": True,
        "enable_typos_rule": True,
        "enable_sensitive_word": False,
    }


def _run_current_proofread(
    input_docx: Path,
    output_docx: Path,
    proofread_options: dict[str, bool],
) -> Path:
    """Run the canonical current runtime in an output-local config boundary."""
    from docwen_bundle.runtime_factory import create_runtime_port
    from docwen_core.models.file_ref import FileRef
    from docwen_core.models.request import ConversionRequest, OutputPolicy
    from docwen_runtime.config.loader import ConfigLoader

    repo_root = Path(__file__).resolve().parents[1]
    user_dir = output_docx.parent / "user-config"
    proofread_dir = user_dir / "proofread"
    proofread_dir.mkdir(parents=True, exist_ok=True)
    (proofread_dir / "typos.toml").write_text('[entries]\n"材料" = ["才料"]\n', encoding="utf-8")
    log_dir = output_docx.parent / "runtime-logs"
    previous_log_dir = os.environ.get("DOCWEN_LOG_DIR")
    os.environ["DOCWEN_LOG_DIR"] = str(log_dir)
    try:
        loader = ConfigLoader(
            base_dir=repo_root / "configs",
            user_dir=user_dir,
        )
        runtime = create_runtime_port(config_loader=loader)
        request = ConversionRequest(
            request_id=f"anchor-matrix-{output_docx.parent.name}",
            input_refs=[FileRef(path=str(input_docx.resolve()), format="document", category="document")],
            target_format="docx",
            action_name="validate",
            options=proofread_options,
            output_policy=OutputPolicy(output_dir=str(output_docx.parent / "runtime-output")),
        )
        result = runtime.execute(request)
    finally:
        if previous_log_dir is None:
            os.environ.pop("DOCWEN_LOG_DIR", None)
        else:
            os.environ["DOCWEN_LOG_DIR"] = previous_log_dir
    if not result.success or not result.artifacts:
        raise RuntimeError(f"current DOCX proofread failed: {result.error}")
    staging_path = Path(result.artifacts[0].staging_path)
    shutil.copyfile(staging_path, output_docx)
    return output_docx


def _grep_key_log_lines(log_text: str) -> list[str]:
    keep = []
    patterns = [
        r"段落包含非文本节点，跳过重建并降级批注",
        r"检测段落非文本节点时出错，保守降级为非文本段落",
        r"拆分规划完成，共 \d+ 个run",
        r"开始重建段落",
        r"段落重建完成，共标记 \d+ 个错误run",
        r"重建段落后未找到错误run，无法添加批注",
        r"未找到错误 \d+ 对应的run",
        r"未找到包含位置 \d+ 的run，返回第一个run",
        r"找到run: 位置 \d+ 在run内相对位置 \d+",
    ]
    combined = re.compile("|".join(f"(?:{p})" for p in patterns))
    for line in log_text.splitlines():
        if combined.search(line):
            keep.append(line)
    return keep


def _extract_structured_report(docx_path: Path, context_chars: int, redact: bool):
    document_xml = read_docx_part(docx_path, "word/document.xml")
    if document_xml is None:
        raise FileNotFoundError("DOCX 缺少 word/document.xml")

    comments_xml = read_docx_part(docx_path, "word/comments.xml")
    occurrences, diagnostics = extract_occurrences_from_document_xml(document_xml, context_chars, redact)
    comments = extract_comment_texts_from_comments_xml(comments_xml, redact) if comments_xml else {}
    return occurrences, diagnostics, comments


def _build_case_run_split_unclosed_punc() -> Document:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("这是一个")
    p.add_run("测")
    p.add_run("试")
    p.add_run("（")
    p.add_run("未闭合")
    return doc


def _build_case_typo_cross_runs() -> Document:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("本段包含")
    p.add_run("才")
    p.add_run("料")
    p.add_run("二字（应为“材料”）")
    return doc


def _build_case_tab_and_break() -> Document:
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("字段：")
    r1.add_tab()
    p.add_run("０")
    r2 = p.add_run(" 行1")
    r2.add_break()
    p.add_run("行2（未闭合")
    return doc


def _build_case_non_text_degraded(tmp_dir: Path) -> Document:
    doc = Document()
    img_path = tmp_dir / "img" / "1x1.png"
    _write_png_1x1(img_path)
    p = doc.add_paragraph()
    p.add_run("含图片段落：")
    p.add_run().add_picture(str(img_path), width=Inches(0.2))
    p.add_run("才料（")
    return doc


def _cases(tmp_dir: Path) -> list[Case]:
    return [
        Case(
            slug="case-01-unclosed-punc-multi-run",
            title="多 run 未闭合符号（重建路径应可字符级锚定）",
            source="动态生成 DOCX",
            expected_path="重建路径",
            minimal_text="这是一个测试（未闭合",
            build=_build_case_run_split_unclosed_punc,
            expectations=(Expectation(label="未闭合（", error_text="（", expected_covered="（"),),
        ),
        Case(
            slug="case-02-typo-cross-runs",
            title="错别字跨 run（风险：仅标到末段）",
            source="动态生成 DOCX",
            expected_path="重建路径",
            minimal_text="本段包含才料二字（应为“材料”）",
            build=_build_case_typo_cross_runs,
            expectations=(Expectation(label="错别字才料", error_text="才料", expected_covered="才料"),),
        ),
        Case(
            slug="case-03-tab-break-mixed",
            title="含 tab/换行/全角数字/未闭合符号（重建路径边界）",
            source="动态生成 DOCX",
            expected_path="重建路径",
            minimal_text="字段：\\t０ 行1\\n行2（未闭合",
            build=_build_case_tab_and_break,
            expectations=(
                Expectation(label="全角数字０", error_text="０", expected_covered="０"),
                Expectation(label="未闭合（", error_text="（", expected_covered="（"),
            ),
        ),
        Case(
            slug="case-04-non-text-degraded",
            title="含图片段落触发降级路径（定位精度上限=run粒度）",
            source="动态生成 DOCX",
            expected_path="降级路径",
            minimal_text="含图片段落：才料（",
            build=lambda: _build_case_non_text_degraded(tmp_dir),
            expectations=(
                Expectation(label="错别字才料", error_text="才料", expected_covered="才料"),
                Expectation(label="未闭合（", error_text="（", expected_covered="（"),
            ),
        ),
    ]


def _render_case_summary_md(
    case: Case,
    input_docx: Path,
    checked_docx: Path,
    report_md: Path,
    key_log_lines: list[str],
    issues: list[str],
) -> str:
    lines = []
    lines.append(f"# {case.slug} {case.title}")
    lines.append("")
    lines.append(f"- 输入：`{input_docx.name}`")
    lines.append(f"- 输出：`{checked_docx.name}`")
    lines.append(f"- 锚点报告：`{report_md.name}`")
    lines.append(f"- 期望路径：{case.expected_path}")
    lines.append("")
    lines.append("## 最小文本（脱敏样例）")
    lines.append("")
    lines.append("```text")
    lines.append(case.minimal_text)
    lines.append("```")
    lines.append("")
    lines.append("## 日志关键行")
    lines.append("")
    lines.append("```text")
    if key_log_lines:
        lines.extend(key_log_lines)
    else:
        lines.append("(none)")
    lines.append("```")
    lines.append("")
    lines.append("## 自动判定")
    lines.append("")
    if issues:
        for s in issues:
            lines.append(f"- {s}")
    else:
        lines.append("- 未发现与期望不一致的锚点覆盖（仅基于本样例的自动规则）")
    lines.append("")
    return "\n".join(lines)


def _evaluate_expectations(occurrences, comments: dict[str, str], expectations: Iterable[Expectation]) -> list[str]:
    issues = []
    occ_by_cid: dict[str, list] = {}
    for o in occurrences:
        occ_by_cid.setdefault(o.comment_id, []).append(o)

    for exp in expectations:
        matched = False
        for cid, text in comments.items():
            if exp.error_text not in text:
                continue
            matched = True
            occs = occ_by_cid.get(cid, [])
            if not occs:
                issues.append(f"{exp.label}：comments.xml 存在批注，但 document.xml 未找到锚点范围")
                continue
            for o in occs:
                if o.covered_text != exp.expected_covered:
                    issues.append(
                        f"{exp.label}：覆盖不符，期望 `{exp.expected_covered}`，实际 `{o.covered_text}`（段落 {o.paragraph_index}，范围 [{o.start},{o.end})）"
                    )
        if not matched:
            issues.append(f"{exp.label}：未在批注文案中匹配到 error_text=`{exp.error_text}`，无法自动关联")
    return issues


def main(argv: list[str]) -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--out-dir", help="输出目录（默认 .tmp/docx_spell_anchor_matrix/<timestamp>）")
    parser.add_argument("--context-chars", type=int, default=20)
    parser.add_argument("--redact", action="store_true", help="对报告与提取内容做脱敏替换")
    parser.add_argument("--keep", action="store_true", help="保留生成的输入/输出 DOCX 与报告文件")
    parser.add_argument("--verbose", action="store_true")
    args = parser.parse_args(argv)

    ts = _dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    base_out = Path(args.out_dir) if args.out_dir else Path(".tmp") / "docx_spell_anchor_matrix" / ts
    base_out.mkdir(parents=True, exist_ok=True)

    all_rules = _all_rules_enabled()
    logger.info(f"输出目录: {base_out.resolve()}")

    index_lines = []
    index_lines.append("# docx_spell 批注锚点检查（复现矩阵）")
    index_lines.append("")
    index_lines.append(f"- 生成时间：`{_dt.datetime.now().isoformat(timespec='seconds')}`")
    index_lines.append(f"- 输出目录：`{base_out.as_posix()}`")
    index_lines.append("")
    failed_cases: list[str] = []

    for case in _cases(base_out):
        case_dir = base_out / case.slug
        case_dir.mkdir(parents=True, exist_ok=True)

        log_path = case_dir / "run.log"
        _configure_logging(log_path, verbose=args.verbose)

        input_docx = case_dir / "input.docx"
        checked_docx = case_dir / "checked.docx"
        report_md = case_dir / "anchor_report.md"
        summary_md = case_dir / "summary.md"

        logger.info(f"生成样例: {case.slug}")
        doc = case.build()
        doc.save(input_docx)

        logger.info("运行校对并生成带批注 DOCX")
        out_path = _run_current_proofread(input_docx, checked_docx, all_rules)
        if not out_path:
            raise RuntimeError("process_docx 返回空路径")

        logger.info("生成锚点 Markdown 报告")
        md = build_anchor_report_markdown(checked_docx, context_chars=args.context_chars, redact=args.redact)
        report_md.write_text(md, encoding="utf-8")

        occurrences, diagnostics, comments = _extract_structured_report(
            checked_docx, context_chars=args.context_chars, redact=False
        )
        issues = _evaluate_expectations(occurrences, comments, case.expectations)
        if diagnostics.cross_paragraph or diagnostics.start_without_end_ids or diagnostics.end_without_start_ids:
            logger.warning("检测到跨段/未闭合锚点异常，详见锚点报告")
            issues.append("生成的批注锚点包含跨段或未闭合的结构异常")

        log_text = log_path.read_text(encoding="utf-8", errors="replace")
        key_lines = _grep_key_log_lines(log_text)

        summary_md.write_text(
            _render_case_summary_md(case, input_docx, checked_docx, report_md, key_lines, issues),
            encoding="utf-8",
        )

        status = "FAIL" if issues else "PASS"
        index_lines.append(f"- {case.slug}: **{status}** `{(summary_md.relative_to(base_out)).as_posix()}`")
        if issues:
            failed_cases.append(case.slug)

        if not args.keep:
            with contextlib.suppress(Exception):
                input_docx.unlink(missing_ok=True)

    (base_out / "index.md").write_text("\n".join(index_lines) + "\n", encoding="utf-8")
    print(str((base_out / "index.md").resolve()))
    if failed_cases:
        logger.error("锚点矩阵失败: %s", ", ".join(failed_cases))
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
